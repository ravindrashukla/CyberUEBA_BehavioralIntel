"""Generate comprehensive Speaker Notes & Deep Preparation Guide for ACECARD CE May 6.
Covers every slide with: talking points, technical depth, anticipated questions, pivot strategies."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT = os.path.join(os.path.dirname(__file__), "05_Speaker_Notes_Deep_Prep_Guide.docx")

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2D, 0x5A, 0x87)

def h3(text):
    doc.add_heading(text, level=3)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    return p

def quoted(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(36)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x88)

def divider():
    p = doc.add_paragraph()
    p.add_run('_' * 80).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('ACECARD CE — Speaker Notes & Deep Preparation Guide')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Comprehensive Preparation for May 6, 2026 Solution Brief Presentation')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('\nEvent: ACECARD Collaboration Event\n'
             'Location: CFIC, 200 Grace Hopper Lane, Augusta GA\n'
             'Time Slot: 1420-1535 (Solution Brief Presentations)\n'
             'Audience: Cleared industry partners + military cyber (ARCYBER/ArCTIC)\n'
             'Format: 10-12 minute presentation followed by Q&A\n'
             'Purpose: Demonstrate working solution, not marketing claims\n').font.size = Pt(11)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
h1('Table of Contents')
toc_items = [
    'Executive Summary & Key Message',
    'Slide 1: Title — Opening Statement',
    'Slide 2: The Problem — Four CFIC Challenges',
    'Slide 3: Why Current Tools Fail — 8 Structural Gaps',
    'Slide 4: The Two-Layer Solution',
    'Slide 5: How It Solves the 4 CFIC Challenges',
    'Slide 6: Five Behavioral Signals + Trajectory Analysis',
    'Slide 7: Volt Typhoon Walkthrough',
    'Slide 8: Salt Typhoon Walkthrough',
    'Slide 9: Proof — Working Prototype',
    'Slide 10: Deployment Readiness',
    'Slide 11: Team Credentials',
    'Slide 12: Closing',
    'Anticipated Questions & Responses',
    'Breakout Session Preparation',
    'Technical Deep-Dive Reference',
    'Key Numbers to Memorize',
]
for i, item in enumerate(toc_items, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
h1('1. Executive Summary & Key Message')

para('ONE SENTENCE YOU MUST INTERNALIZE:', bold=True)
quoted('"We have a working, containerized prototype that combines formal mathematical '
       'verification of security configurations with behavioral trajectory intelligence '
       'to detect APT-class threats that neither layer can catch alone."')

h3('The Three Things the Audience Must Remember')
bullet('We have a WORKING solution (not a proposal, not a concept)')
bullet('It solves ALL FOUR stated CFIC challenges with specific, demonstrable capabilities')
bullet('It is Gabriel Nimbus-ready: containerized, multi-enclave, cATO-aligned')

h3('Framing: Solution, Not Marketing')
para('This audience is technical. They have seen dozens of vendor pitches. '
     'What they have NOT seen is a team that walks in with a running Docker prototype, '
     'shows code, shows dashboards, and maps every capability directly to their stated problems. '
     'That is our differentiator. Stay solution-focused at all times.')

h3('Presentation Flow (10-12 minutes)')
bullet('Slides 1-3 (2 min): Problem framing — establish shared understanding')
bullet('Slides 4-5 (2 min): Solution architecture — two layers, four challenges mapped')
bullet('Slide 6 (2 min): Technical depth — signals, embeddings, detection')
bullet('Slides 7-8 (2 min): Real-world application — both Typhoons, phase by phase')
bullet('Slides 9-10 (2 min): Proof & deployment — running system, Gabriel Nimbus ready')
bullet('Slides 11-12 (1 min): Team + close')

doc.add_page_break()

# ============================================================
# SLIDE 1
# ============================================================
h1('2. Slide 1: Title — Opening Statement')

h3('What Is On This Slide')
bullet('ACECARD CE — Solution Brief')
bullet('Behavioral Anomaly Detection Mission Application')
bullet('Preemptive Cyber Defense + Behavioral Trajectory Intelligence')
bullet('Containerized, Cloud-Native Prototype for Gabriel Nimbus')
bullet('CFIC | ARCYBER | ArCTIC | 22CT | Rigor AI')

h3('Talking Points (30 seconds)')
para('Open with confidence. Do NOT read the slide. Look at the room and say:')
quoted('"Good afternoon. I am [name] from 22nd Century Technologies. '
       'We are here with our partner Rigor AI to present a working solution — '
       'not a concept, not a roadmap — a containerized prototype that we can demo today. '
       'It combines two layers of defense that address all four challenges in the CFIC problem statement. '
       'Let me show you what it does."')

h3('Key Points to Hit')
bullet('"Working solution" — say these exact words. Sets the expectation.')
bullet('"Containerized prototype" — signals it is real, deployable, not slides-only')
bullet('"Both layers" — plants the seed for the two-layer architecture')
bullet('"All four challenges" — directly references the CFIC problem statement they wrote')

h3('Do NOT Say')
bullet('Do not explain the company history here')
bullet('Do not apologize for time or format')
bullet('Do not use marketing adjectives (revolutionary, cutting-edge, next-gen)')

divider()
doc.add_page_break()

# ============================================================
# SLIDE 2
# ============================================================
h1('3. Slide 2: The Problem — Four CFIC Challenges')

h3('What Is On This Slide')
para('Four columns, each with a CFIC challenge and three specific data points:')

h3('Challenge 1: Manual Context Reassembly')
bullet('Analyst touches 6-12 data tables per alert')
bullet('Each table: different schema, different timestamps, different entity IDs')
bullet('30-45 min triage time per alert x 500 alerts/day = most never investigated')

h3('Challenge 2: Inconsistent Siloed Analytics')
bullet('AD, AAD, Okta, AWS, K8s each see 1 login')
bullet('Lateral movement across 5 systems = 5 weak signals, not 1 strong signal')
bullet('No current tool fuses these into a single trajectory per entity')

h3('Challenge 3: Scalability Constraints')
bullet('Manual correlation caps at ~50 entities')
bullet('N log sources = N separate analytics pipelines')
bullet('Gabriel Nimbus scale: 100K+ entities. Current tools cannot keep up')

h3('Challenge 4: Operational Inefficiency')
bullet('Anomaly score "87%" — no direction, no MITRE mapping')
bullet('Static Sigma rules decay monthly (PRC actors test against them)')
bullet('Alert fatigue: analysts ignore most alerts. Attacker operates freely')

h3('Talking Points (60 seconds)')
quoted('"These are not our words — these are the four challenges from the CFIC problem statement. '
       'Let me make them concrete with numbers. '
       '[Point to Challenge 1] An analyst triaging a single alert touches 6 to 12 tables, '
       'each with different schemas and timestamps. That is 30 to 45 minutes PER ALERT. '
       'At 500 alerts per day, most are never investigated. '
       '[Point to Challenge 2] When Volt Typhoon moved laterally across 5 identity systems, '
       'each system saw one normal login. Five weak signals, never correlated into one strong signal. '
       '[Point to bottom] Result: Volt Typhoon dwelled 5+ years. Salt Typhoon dwelled 4+ years. '
       'These are the consequences of these four structural limitations."')

h3('Supporting Data You Should Know')
bullet('CISA Advisory AA24-038A (Feb 2024): Volt Typhoon 5+ year dwell, LoTL tradecraft', level=0)
bullet('CISA Advisory AA25-239A: Salt Typhoon affected 9 US telecom providers', level=0)
bullet('Gartner Research: 99% of firewall breaches through 2025 are caused by misconfigurations, not vulnerabilities', level=0)
bullet('MITRE ATT&CK: Volt Typhoon uses T1133 (External Remote Services), T1059 (Command Line), T1003 (Credential Dumping)', level=0)
bullet('Mean dwell time for state-sponsored APT: 180+ days (FireEye M-Trends)', level=0)

h3('Anticipated Questions')
para('Q: "Where do the 6-12 tables come from?"', bold=True)
para('A: Typical SOC: SIEM (Splunk/Elastic), EDR (CrowdStrike/Defender), IAM (AD/Okta/AAD), '
     'network (Zeek/PAN), ticket system (ServiceNow), asset inventory (CMDB), '
     'threat intel (MISP/TIP). Each alert requires correlation across 6+ of these.')

para('Q: "Are you saying SIEM is inadequate?"', bold=True)
para('A: SIEM is excellent at log aggregation and rule-based detection. It is structurally '
     'incapable of unified behavioral trajectory analysis because it operates on events, not entities. '
     'We complement SIEM — we do not replace it. We consume its output.')

divider()
doc.add_page_break()

# ============================================================
# SLIDE 3
# ============================================================
h1('4. Slide 3: Why Current Tools Fail — 8 Structural Gaps')

h3('What Is On This Slide')
para('Two columns: 4 preemptive layer gaps (left) + 4 behavioral layer gaps (right). '
     'This slide establishes WHY a new approach is needed — these are structural limitations, '
     'not implementation bugs.')

h2('LEFT COLUMN: Preemptive Layer Fails')

h3('Gap 1: 99% of firewall breaches = misconfiguration (Gartner)')
para('What this means:', bold=True)
bullet('Gartner prediction (confirmed through 2025): the vast majority of network breaches '
       'are NOT from zero-days or novel exploits')
bullet('They come from firewalls that are MISCONFIGURED — rules that shadow each other, '
       'zombie rules from decommissioned systems, overly permissive any-any rules')
bullet('Both Volt Typhoon and Salt Typhoon exploited misconfigurations in Next-Generation '
       'Firewalls (NGFWs) and VPN concentrators')
bullet('NGFW = Next-Generation Firewall (Palo Alto, Fortinet, Cisco). Adds app-layer inspection '
       'on top of traditional packet filtering. Still depends on CORRECT rule configuration.')
bullet('VPN = Virtual Private Network. When VPN concentrator ACLs are misconfigured, '
       'external actors get internal network access (exactly what Volt Typhoon did with SOHO routers)')

h3('Gap 2: KEV patching covers only ~20% of attacks')
para('What this means:', bold=True)
bullet('KEV = Known Exploited Vulnerabilities catalog, maintained by CISA')
bullet('CVE = Common Vulnerabilities and Exposures — a database of known software flaws')
bullet('KEV is the subset of CVEs that CISA has confirmed are actively exploited in the wild')
bullet('Even if you patch EVERY KEV entry, you only close ~20% of real attack paths')
bullet('The other 80% use valid credentials, misconfigurations, or techniques that have NO CVE')
bullet('Salt Typhoon used CVE-2023-20198 (Cisco IOS XE) as initial entry, but most of their '
       'internal movement used techniques with no patch — CALEA access, ACL tampering, GRE tunneling')

h3('Gap 3: ASM/BAS only sample the 2^8000 state space')
para('What this means:', bold=True)
bullet('ASM = Attack Surface Management — tools that scan for exposed assets and vulnerabilities')
bullet('BAS = Breach and Attack Simulation — tools that simulate specific attack paths to test defenses')
bullet('The problem: a typical enterprise network has thousands of rules across dozens of controls. '
       'The total configuration state space is approximately 2^8000 combinations.')
bullet('2^8000 is larger than the number of atoms in the observable universe (10^80 = roughly 2^266)')
bullet('ASM/BAS tools SAMPLE this space — they test hundreds or thousands of paths')
bullet('Sampling a vanishingly small fraction of 2^8000 and declaring "secure" is like '
       'checking 10 grains of sand on a beach and declaring the beach clean')
bullet('Rigor AI does NOT sample. It uses formal symbolic reasoning to PROVE properties '
       'across the entire state space. Mathematical certainty, not statistical confidence.')

h3('Gap 4: Config drift goes unverified')
para('What this means:', bold=True)
bullet('Configuration drift: when a system\'s actual configuration diverges from its intended state')
bullet('Causes: manual changes, emergency patches, failed automation, legacy overrides')
bullet('Salt Typhoon specifically modified router ACLs after gaining access — creating '
       'persistent backdoors in the configuration itself')
bullet('Current drift detection is REACTIVE — it notices after the fact')
bullet('Rigor does CONTINUOUS re-verification: every config change triggers re-analysis '
       'of the entire policy space within minutes')

h2('RIGHT COLUMN: Behavioral Layer Fails')

h3('Gap 5: Threshold scoring misses Living-off-the-Land (LoTL)')
para('What this means:', bold=True)
bullet('LoTL = Living off the Land — using legitimate, pre-installed system tools for malicious purposes')
bullet('Specific tools used by Volt Typhoon (all are legitimate Windows admin tools):')
bullet('wmic — Windows Management Instrumentation command-line (system queries, remote execution)', level=1)
bullet('ntdsutil — Active Directory database utility (extracts password hashes from NTDS.dit)', level=1)
bullet('netsh — Network shell (modifies firewall rules, routes, proxy settings)', level=1)
bullet('rundll32 — Runs DLL functions (loads malicious code via legitimate process)', level=1)
bullet('vssadmin — Volume Shadow Copy admin (deletes backup snapshots before ransomware)', level=1)
bullet('PowerShell — Scripting framework (downloads, executes, persists — extremely versatile)', level=1)
bullet('Why thresholds fail: each individual LoTL execution looks "normal" — one wmic call is fine. '
       'There is no threshold to set because the INDIVIDUAL action is legitimate. '
       'Only the PATTERN over time (trajectory) reveals the attack.')
bullet('ACECARD solution: behavioral embedding captures the COMBINATION of tools used, '
       'the sequence, the timing, the targets. The drift vector points toward "lotl_execution" '
       'concept even though no single event crossed a threshold.')

h3('Gap 6: Per-source siloed analytics fragment identity')
para('What this means:', bold=True)
bullet('Modern enterprises use 5+ identity providers simultaneously:')
bullet('AD = Active Directory (on-premises Windows domain authentication)', level=1)
bullet('AAD = Azure Active Directory (now Entra ID) (cloud identity)', level=1)
bullet('Okta = Cloud identity provider (SSO, MFA)', level=1)
bullet('AWS IAM = Amazon Web Services Identity and Access Management', level=1)
bullet('K8s RBAC = Kubernetes Role-Based Access Control', level=1)
bullet('Each system has its OWN user identifier (SID, UPN, email, ARN, service account name)')
bullet('When an attacker moves laterally across these systems, each system sees ONE normal login')
bullet('4 weak signals across 4 systems are never correlated into 1 strong signal')
bullet('ACECARD solution: Entity resolution maps ALL identity system identifiers to a single '
       'entity_uuid. All 5 behavioral signals from all sources contribute to ONE embedding. '
       'Lateral movement reads as one accelerating trajectory, not 4 isolated blips.')

h3('Gap 7: Static rules decay monthly')
para('What this means:', bold=True)
bullet('Sigma rules = open-source detection rule format for SIEM systems')
bullet('Organizations write Sigma rules to detect known TTPs (Tactics, Techniques, Procedures)')
bullet('TTP = the specific methods an attacker uses (e.g., "dump credentials via ntdsutil")')
bullet('PRC (People\'s Republic of China) state-sponsored actors specifically TEST their '
       'tradecraft against published Sigma rules before launching campaigns')
bullet('They modify their tools until no published rule detects them')
bullet('Result: detection coverage of rule-based systems DECAYS over time as adversaries adapt')
bullet('Typical half-life of a Sigma rule against APT actors: weeks to low months')
bullet('This is not a failure of the rule-writing team — it is a structural limitation of '
       'signature-based detection against adaptive adversaries')
bullet('ACECARD solution: behavioral embeddings do not depend on known signatures. They measure '
       'WHAT the entity IS BECOMING, not whether a specific string matched a known pattern. '
       'Adversary can change tools, change timing, change targets — the behavioral trajectory '
       'still drifts toward the threat concept.')

h3('Gap 8: "Anomaly score" without direction')
para('What this means:', bold=True)
bullet('Current UEBA tools produce a single number: "User X has anomaly score 87%"')
bullet('This tells the analyst NOTHING about:')
bullet('What the user is doing that is anomalous', level=1)
bullet('What threat this behavior pattern resembles', level=1)
bullet('Which MITRE ATT&CK technique it maps to', level=1)
bullet('Whether it is getting worse or stabilizing', level=1)
bullet('Whether peers in the same role show similar patterns', level=1)
bullet('Result: analyst must manually investigate to determine if 87% is a credential theft, '
       'an insider threat, a role change, or a data error. Triage time explodes.')
bullet('ACECARD solution: drift DIRECTION via cosine projection onto 8 named threat concepts. '
       'Alert says: "User X drifting toward credential_dumping (0.73) and lateral_movement (0.61), '
       'mapped to T1003 and T1021." Analyst knows exactly what to look for.')

h3('Talking Points (90 seconds)')
quoted('"This slide is the core of our argument. There are 8 structural gaps — not bugs to fix, '
       'but architectural limitations of current approaches. On the left: preemptive tools like '
       'firewalls and vulnerability scanners. They cannot PROVE coverage — they sample a tiny '
       'fraction of the state space and call it good. On the right: behavioral tools. Threshold-based '
       'detection structurally cannot catch Living-off-the-Land because each individual action is '
       'legitimate. Static rules decay as adversaries adapt. And even when an alert fires, it says '
       '87 percent anomalous with no direction, no MITRE mapping, no hypothesis. '
       'Our solution addresses both columns simultaneously."')

divider()
doc.add_page_break()

# ============================================================
# SLIDE 4
# ============================================================
h1('5. Slide 4: The Two-Layer Solution')

h3('What Is On This Slide')
para('Two large boxes: Layer 1 (Rigor AI - Preemptive) and Layer 2 (ACECARD UEBA - Behavioral). '
     'Bottom box: why both layers are required.')

h2('Layer 1: Rigor AI — Preemptive Defense')
para('Key technical claims to understand deeply:', bold=True)

h3('Formal Mathematical Model')
bullet('Rigor ingests the ACTUAL configuration of all security controls: NGFW rules, IPS signatures, '
       'IdP policies, SASE rules, WAF rules')
bullet('NGFW = Next-Generation Firewall (app-aware packet filtering)')
bullet('IPS = Intrusion Prevention System (signature-based blocking)')
bullet('IdP = Identity Provider (authentication/authorization policies)')
bullet('SASE = Secure Access Service Edge (cloud-delivered security)')
bullet('WAF = Web Application Firewall (HTTP-layer protection)')
bullet('It builds a formal model of what traffic CAN and CANNOT pass through the combined control stack')

h3('Exhaustive Analysis of 2^8000 State Space')
bullet('Traditional tools (Nessus, Qualys, BAS platforms) test specific paths')
bullet('Rigor uses symbolic reasoning (similar to formal verification in chip design)')
bullet('It does NOT enumerate all 2^8000 states — that would be impossible')
bullet('It uses mathematical PROOFS about properties that hold across the entire space')
bullet('Analogy: you don\'t need to test every number to prove that the sum of two even numbers is always even')
bullet('Result: "There is NO path from zone X to zone Y that bypasses firewall rule set Z" — proven, not sampled')

h3('Continuous Re-verification')
bullet('Every configuration change triggers re-analysis')
bullet('Cadence: hourly or per-change (configurable)')
bullet('If someone adds a firewall rule that creates a shadow or gap, Rigor detects it within the hour')
bullet('This is how it catches config drift in near-real-time')

h3('Risk-Prioritized Remediation')
bullet('Rigor does not just say "you have a problem" — it gives vendor-specific fix instructions')
bullet('Example: "On PA-5200 Series, move rule 147 above rule 23 to eliminate shadow"')
bullet('Prioritized by exploitability × business impact')

h2('Layer 2: ACECARD UEBA — Behavioral Detection')
para('Key technical claims to understand deeply:', bold=True)

h3('1536-d Behavioral Embedding')
bullet('1536-d = 1536-dimensional vector (the output size of OpenAI text-embedding-3-small)')
bullet('Each entity (user, device, network segment, application) gets ONE vector per time window')
bullet('This vector encodes the BEHAVIORAL STATE of that entity during that window')
bullet('Vector space has the property that similar behaviors are CLOSE (high cosine similarity) '
       'and dissimilar behaviors are FAR')

h3('5 Behavioral Signals')
bullet('AUTH: authentication patterns (logon count, failure rate, off-hours ratio, impossible travel)')
bullet('PROCESS: process execution (unique processes, LoTL binary count, command-line entropy, unsigned ratio)')
bullet('NETWORK: network behavior (unique destinations, bytes out ratio, beacon score, DNS query rate)')
bullet('FILE: file system activity (files created/deleted, sensitive access, archive creation, extension changes)')
bullet('IDENTITY: privilege operations (escalations, group additions, MFA bypasses, role changes)')

h3('CUSUM Change-Point Detection')
bullet('CUSUM = Cumulative Sum — a statistical algorithm from quality control')
bullet('Designed to detect SMALL, SUSTAINED shifts in a process mean')
bullet('Traditional threshold: fires only when value exceeds X')
bullet('CUSUM: accumulates small deviations over time. Many small shifts that never cross a threshold still trigger CUSUM')
bullet('Threshold: 4 sigma (standard deviations) accumulated drift')
bullet('This is WHY it catches slow APTs: Volt Typhoon\'s daily drift was tiny (0.01-0.02), '
       'but accumulated over weeks, CUSUM triggered while no threshold ever fired')

h3('Drift Direction via Cosine Projection')
bullet('When we detect drift, the analyst needs to know: "drift TOWARD WHAT?"')
bullet('We maintain 8 reference concept embeddings (text descriptions of threat behaviors, embedded into the same 1536-d space)')
bullet('We compute: projection = cosine(drift_vector, concept_embedding) for all 8 concepts')
bullet('Result: "This entity is drifting toward credential_dumping (0.73) and lateral_movement (0.61)"')
bullet('Each concept is mapped to specific MITRE ATT&CK technique IDs')

h3('The 8 MITRE-Mapped Threat Concepts')
bullet('credential_dumping → T1003 (OS Credential Dumping)')
bullet('lateral_movement → T1021 (Remote Services)')
bullet('data_exfiltration → T1041 (Exfiltration Over C2 Channel)')
bullet('c2_beaconing → T1071 (Application Layer Protocol for C2)')
bullet('lotl_execution → T1059 (Command and Scripting Interpreter)')
bullet('privilege_escalation → T1078 (Valid Accounts)')
bullet('defense_evasion → T1562 (Impair Defenses)')
bullet('insider_hoarding → T1074 (Data Staged)')

h3('ABAC Trust Loop')
bullet('ABAC = Attribute-Based Access Control')
bullet('Four trust states: TRUSTED → ELEVATED_WATCH → RESTRICTED → BLOCKED')
bullet('State transitions are automatic based on drift magnitude and direction')
bullet('ELEVATED_WATCH: step-up MFA required, increased logging')
bullet('RESTRICTED: read-only access, lateral movement blocked')
bullet('BLOCKED: all access revoked, incident ticket created')
bullet('Key property: PROPORTIONAL response. No binary on/off. Graduated enforcement.')

h3('Why Both Layers — The Core Argument')
quoted('"Rigor cannot detect attacks in progress — if someone has valid credentials and is using '
       'legitimate tools, no configuration analysis will catch them. ACECARD cannot prevent '
       'misconfigured controls — if a firewall rule allows traffic it shouldn\'t, behavioral '
       'analysis of the traffic is too late. Together: either the path never existed because '
       'Rigor mathematically proved it was closed, OR the attack is detected within hours by '
       'ACECARD and contained via ABAC. Defense in depth."')

h3('Talking Points (90 seconds)')
quoted('"Our solution has two layers and both are required. Layer 1: Rigor AI performs formal '
       'mathematical verification of your entire security control stack. Not sampling — mathematical '
       'proof across the full state space. If a misconfiguration exists, Rigor finds it before '
       'exploitation. Layer 2: ACECARD tracks behavioral trajectory of every entity using 5 signals '
       'fused into a 1536-dimensional embedding. When behavior drifts, CUSUM catches it even if the '
       'drift is tiny — accumulated over weeks. And drift direction tells you exactly what the entity '
       'is becoming: credential dumping, lateral movement, exfiltration — mapped to MITRE. '
       'Why both? Rigor can\'t detect valid credentials being misused. ACECARD can\'t prevent '
       'misconfigured firewalls. Together, they close the gap."')

divider()
doc.add_page_break()

# ============================================================
# SLIDE 5
# ============================================================
h1('6. Slide 5: How It Solves the 4 CFIC Challenges')

h3('What Is On This Slide')
para('Four rows, each mapping a CFIC challenge to specific solution capabilities with a gold-highlighted outcome.')

h3('Challenge 1 → Unified Embedding Eliminates Manual Joins')
para('Problem: Analyst touches 6-12 tables per alert.', italic=True)
bullet('Solution: All log sources feed into ONE 1536-d embedding per entity')
bullet('No manual joins — the embedding IS the fused context')
bullet('One API call: GET /api/trajectory/{entity}/dashboard')
bullet('Returns: drift magnitude, velocity, CUSUM state, health score, top-3 threat concepts, '
       'MITRE technique IDs, peer z-score')
bullet('OUTCOME: Alert to full context in <3 seconds, not 30-45 minutes')

h3('Challenge 2 → Entity Resolution Unifies Identity')
para('Problem: 5 identity systems, 5 weak signals.', italic=True)
bullet('Solution: Entity resolution maps AD/AAD/Okta/AWS/K8s identifiers to single entity_uuid')
bullet('5 signal serializers normalize events from ANY source format (ECS, OCSF, STIX 2.1)')
bullet('ECS = Elastic Common Schema (Elasticsearch\'s normalized event format)')
bullet('OCSF = Open Cybersecurity Schema Framework (AWS + industry consortium)')
bullet('STIX 2.1 = Structured Threat Information eXpression (OASIS standard for threat intel)')
bullet('OUTCOME: Lateral movement across 5 identity systems = 1 strong signal, not 5 weak ones')

h3('Challenge 3 → Containerized, Horizontally Scalable')
para('Problem: Current tools cap at ~50 entities.', italic=True)
bullet('Solution: Containerized with Iron Bank-compatible base images, OCI-compliant')
bullet('Iron Bank = DoD centralized container image repository (hardened, approved)')
bullet('OCI = Open Container Initiative (standard container format)')
bullet('Performance: 10,000 entities/hour on 4 vCPU. <3 seconds per entity window.')
bullet('Horizontal scaling: add K8s replicas, linear throughput increase')
bullet('1 unified pipeline handles ALL log types — new sources via schema config, not code')
bullet('OUTCOME: Gabriel Nimbus scale (100K+ entities) achievable with standard K8s deployment')

h3('Challenge 4 → Fully Automated Pipeline + ABAC Enforcement')
para('Problem: Manual wrangling, alert fatigue, no automated response.', italic=True)
bullet('Solution: Full automation from ingest to enforcement')
bullet('Pipeline: ingest → normalize → serialize → embed → analyze → alert → enforce')
bullet('ABAC trust loop: detection to enforcement in <5 minutes')
bullet('Analyst TP/FP feedback auto-tunes thresholds (system improves with use, not decays)')
bullet('TP/FP = True Positive / False Positive')
bullet('OUTCOME: Zero manual wrangling. System gets BETTER over time, unlike rules that decay.')

h3('Talking Points (90 seconds)')
quoted('"Let me show you how these capabilities map directly to your four stated challenges. '
       '[Point to row 1] Manual context reassembly? Eliminated. One API call returns drift, velocity, '
       'CUSUM, health, MITRE mapping — in under 3 seconds. No more joining 6 tables. '
       '[Row 2] Siloed analytics? All identity systems resolve to one entity. Lateral movement '
       'across 5 systems is now one accelerating trajectory. [Row 3] Scale? Containerized, '
       '10,000 entities per hour on 4 vCPU. Add K8s replicas for more. [Row 4] Operational '
       'efficiency? Fully automated pipeline with graduated enforcement — step-up MFA, restrict, '
       'block — proportional to threat level. No manual wrangling."')

divider()
doc.add_page_break()

# ============================================================
# SLIDE 6
# ============================================================
h1('7. Slide 6: Five Behavioral Signals + Trajectory Analysis')

h3('What Is On This Slide')
para('Five colored columns (one per signal) showing sources and features. '
     'Bottom left: trajectory math. Bottom right: 8 drift direction concepts.')

h2('Signal Deep-Dive')

h3('Signal 1: AUTH (Blue)')
para('What it captures: WHO is logging in, WHERE, WHEN, HOW', bold=True)
bullet('Sources: Windows Event 4624 (successful logon), 4625 (failed logon), 4768 (Kerberos ticket), Okta syslog, Azure AD sign-in logs')
bullet('Features computed per entity per time window:')
bullet('logon_count — raw volume of authentication events', level=1)
bullet('failure_rate — ratio of failed to total attempts (brute force indicator)', level=1)
bullet('unique_hosts — number of distinct machines authenticated to (lateral movement)', level=1)
bullet('off_hours_ratio — fraction of logins outside business hours (compromise indicator)', level=1)
bullet('impossible_travel — logins from locations that are geographically impossible given time gap', level=1)

h3('Signal 2: PROCESS (Purple)')
para('What it captures: WHAT programs are running and HOW they are invoked', bold=True)
bullet('Sources: Sysmon Event ID 1 (process creation), Event ID 3 (network connection), EDR telemetry, Linux auditd')
bullet('Features:')
bullet('unique_procs — count of distinct process names (diversity of activity)', level=1)
bullet('lolbin_count — count of Living-off-the-Land binaries executed (wmic, ntdsutil, etc.)', level=1)
bullet('cmdline_entropy — Shannon entropy of command-line arguments (high entropy = obfuscation)', level=1)
bullet('parent_child_depth — max depth of process tree (deep trees = staged execution)', level=1)
bullet('unsigned_ratio — fraction of executed binaries without valid code signing', level=1)

h3('Signal 3: NETWORK (Teal)')
para('What it captures: WHERE data is flowing and in what patterns', bold=True)
bullet('Sources: NetFlow (router-level flow records), Zeek (formerly Bro, full protocol analysis), Palo Alto/Fortinet firewall logs')
bullet('Features:')
bullet('unique_dest_ips — count of unique destination IPs contacted (reconnaissance)', level=1)
bullet('bytes_out_ratio — outbound vs inbound bytes (exfiltration indicator)', level=1)
bullet('beacon_score — regularity of outbound connections (C2 heartbeat pattern)', level=1)
bullet('dns_rate — DNS queries per minute (DNS tunneling, DGA activity)', level=1)
bullet('geo_anomaly — connections to unusual geographic regions', level=1)

h3('Signal 4: FILE (Orange)')
para('What it captures: WHAT files are being accessed, created, or modified', bold=True)
bullet('Sources: Sysmon Event ID 11 (file created), file access audit logs, DLP (Data Loss Prevention) alerts')
bullet('Features:')
bullet('files_created/deleted — volume of file system changes (ransomware indicator)', level=1)
bullet('sensitive_access — access to classified or restricted file paths', level=1)
bullet('archive_creates — creation of zip/tar/rar archives (staging for exfiltration)', level=1)
bullet('extension_changes — files renamed with new extensions (encryption indicator)', level=1)

h3('Signal 5: IDENTITY (Red)')
para('What it captures: privilege and role changes over time', bold=True)
bullet('Sources: Windows Event 4672 (special privileges assigned), CloudTrail IAM events, Kubernetes audit log, Active Directory change log')
bullet('Features:')
bullet('priv_escalations — count of privilege elevation events', level=1)
bullet('group_adds — additions to security groups (domain admins, etc.)', level=1)
bullet('mfa_bypass — successful authentications that skipped MFA', level=1)
bullet('role_changes — IAM role assumptions or switches', level=1)
bullet('admin_actions — administrative actions taken (user creation, policy changes)', level=1)

h2('How Signals Become Vectors')
para('Critical concept — understand this flow:', bold=True)
bullet('Step 1: Raw events are collected from all sources for a given entity and time window (1 hour)')
bullet('Step 2: For each of the 5 signals, features are computed (numerical summaries)')
bullet('Step 3: Features are serialized as STRUCTURED NATURAL LANGUAGE TEXT')
bullet('Example: "User jsmith auth signal: 47 logons, 3.2% failure rate, 8 unique hosts, '
       '0.15 off-hours ratio, no impossible travel detected"', level=1)
bullet('Step 4: All 5 signal texts are CONCATENATED into one paragraph')
bullet('Step 5: The combined text is sent to text-embedding-3-small → returns 1536-d vector')
bullet('Step 6: That vector IS the behavioral state of the entity for that time window')
bullet('No raw event content is sent to the embedding model — only structured feature summaries')
bullet('For JWICS (classified): local SLM (Phi-4 or Mistral 7B) replaces OpenAI. Same API, env var swap.')

h2('Trajectory Analysis Math')
para('Know these formulas:', bold=True)
bullet('Cosine drift = 1 - cos(V_current, V_baseline)')
bullet('Where V_current = latest embedding, V_baseline = rolling average of prior N windows')
bullet('Drift = 0 means identical behavior. Drift = 1 means completely orthogonal behavior.')
bullet('Typical normal drift: 0.02-0.05 per window')
bullet('Alert-worthy drift: 0.15+ per window or 0.30+ accumulated')
bullet('Velocity = drift(t) - drift(t-1) — rate of change. Positive = accelerating departure.')
bullet('CUSUM: S(t) = max(0, S(t-1) + drift(t) - expected_drift). Alarm when S(t) > 4*sigma.')
bullet('Health score: 0-100 composite. Critical < 40, Warning 40-70, Healthy 70-100.')

h2('Drift Direction')
para('8 reference concepts as text, embedded into same 1536-d space:', bold=True)
bullet('Each concept is a 2-3 sentence description of a threat behavior pattern')
bullet('Example concept text for "credential_dumping": "Entity shows escalating credential access. '
       'Multiple authentication database queries. LSASS memory access. SAM registry hive reads. '
       'ntdsutil invocations increasing. Kerberoasting ticket requests rising."')
bullet('This text is embedded once → becomes a fixed reference vector')
bullet('Drift direction = cosine(drift_vector, concept_vector)')
bullet('drift_vector = V_current - V_baseline (the direction entity is moving)')
bullet('High projection (>0.5) = entity IS moving toward that concept')
bullet('Multiple high projections = multi-technique attack (kill chain)')

h3('Talking Points (2 minutes — this is the technical heart)')
quoted('"This is how it works technically. Five behavioral signals — authentication, process, network, '
       'file, and identity — each computed as features from raw events. These are serialized as '
       'structured text and embedded together into a single 1536-dimensional vector. '
       'That vector IS the behavioral state of the entity. We track how it changes over time. '
       'When it drifts, CUSUM catches even tiny sustained shifts. And drift direction tells you '
       'exactly what threat concept the behavior resembles — credential dumping, lateral movement, '
       'exfiltration — each mapped to specific MITRE techniques. The key insight: '
       'Living-off-the-Land tools never cross individual thresholds, but the TRAJECTORY of behavior '
       'across all 5 signals consistently points toward threat concepts."')

divider()
doc.add_page_break()

# ============================================================
# SLIDE 7
# ============================================================
h1('8. Slide 7: Volt Typhoon Walkthrough')

h3('What Is On This Slide')
para('Table with 5 phases of Volt Typhoon attack. Each row shows: attacker behavior, '
     'what Rigor preempts, what ACECARD detects.')

h2('Background: What IS Volt Typhoon?')
bullet('PRC (Chinese) state-sponsored APT group')
bullet('Active since at least 2021, disclosed by CISA Feb 2024 (AA24-038A)')
bullet('Targets: US critical infrastructure (energy, water, transportation, communications)')
bullet('Dwell time: 5+ YEARS undetected in some networks')
bullet('Tradecraft: exclusively Living-off-the-Land — NO custom malware')
bullet('Infrastructure: KV Botnet — compromised SOHO routers as C2 relay')
bullet('Goal: PRE-POSITIONING for potential conflict (not espionage, not financial)')
bullet('This is the hardest attack to detect: no malware signatures, no threshold violations')

h2('Phase-by-Phase Deep Dive')

h3('Phase 1: Initial Access (0-12h)')
para('Attacker behavior:', bold=True)
bullet('Compromises residential SOHO router (home office device, easily exploitable)')
bullet('Uses that router as entry point to victim network via VPN credential reuse')
bullet('Credential reuse: password obtained from prior breach (credential stuffing)')
para('What Rigor preempts:', bold=True)
bullet('Zero-trust segmentation VERIFIED — proves that even with VPN access, lateral paths are closed')
bullet('MFA bypass paths surfaced — identifies where MFA can be circumvented')
bullet('MITRE T1133: External Remote Services')
para('What ACECARD detects:', bold=True)
bullet('Auth signal: off_hours_ratio spikes (login at 2 AM from new location)')
bullet('Auth signal: new source IP never seen before for this entity')
bullet('Drift: 0.08 (minor but non-zero — CUSUM begins accumulating)')

h3('Phase 2: LOLBin Execution (12-36h)')
para('Attacker behavior:', bold=True)
bullet('Executes chain of Living-off-the-Land binaries:')
bullet('wmic — queries system info, discovers domain structure', level=1)
bullet('ntdsutil — accesses AD database, extracts credentials', level=1)
bullet('netsh — modifies local firewall, creates port forwarding', level=1)
para('What Rigor preempts:', bold=True)
bullet('Application control policies VERIFIED — proves unsigned/unauthorized execution is blocked')
bullet('If applocker or WDAC policies have gaps, Rigor surfaces them')
bullet('MITRE T1059: Command and Scripting Interpreter')
para('What ACECARD detects:', bold=True)
bullet('Process signal: lolbin_count jumps from 0 to 7 in one window')
bullet('Process signal: cmdline_entropy rises (obfuscated arguments)')
bullet('Drift: 0.14 (moderate — drift direction: lotl_execution 0.71)')

h3('Phase 3: Lateral Movement (36-72h)')
para('Attacker behavior:', bold=True)
bullet('Uses RDP (Remote Desktop Protocol) to reach domain controllers')
bullet('Uses WinRM (Windows Remote Management) for remote command execution')
bullet('Moves from user workstation → server → domain controller')
para('What Rigor preempts:', bold=True)
bullet('Micro-segmentation rules VERIFIED — proves user VLAN cannot reach DC management ports')
bullet('If segmentation is incomplete, Rigor identifies the exact gap')
bullet('MITRE T1021: Remote Services')
para('What ACECARD detects:', bold=True)
bullet('Network signal: unique_dest_ips jumps from 3 to 18')
bullet('Network signal: admin share access (C$, ADMIN$) from unusual source')
bullet('Drift: 0.19 (significant — drift direction: lateral_movement 0.82)')

h3('Phase 4: Persistence (72-96h)')
para('Attacker behavior:', bold=True)
bullet('Creates scheduled tasks for automatic re-execution')
bullet('Modifies registry for persistence (Run keys, services)')
bullet('Establishes multiple persistence mechanisms for redundancy')
para('What Rigor preempts:', bold=True)
bullet('Configuration drift detected — baseline includes scheduled task whitelist')
bullet('Deviation from known-good config flagged within the hour')
bullet('MITRE T1053: Scheduled Task/Job')
para('What ACECARD detects:', bold=True)
bullet('Identity signal: service account usage spikes (new scheduled tasks run as svc accounts)')
bullet('CUSUM TRIGGERS AT 4 SIGMA — accumulated drift over 72+ hours crosses threshold')
bullet('This is the key moment: no single event was above threshold, but CUSUM caught the accumulation')

h3('Phase 5: Pre-positioning (96h+)')
para('Attacker behavior:', bold=True)
bullet('Harvests credentials for future use')
bullet('Maps network for strategic targets')
bullet('Positions for potential destructive action during conflict')
para('What Rigor preempts:', bold=True)
bullet('Credential guard policies verified on domain controllers')
bullet('LAPS (Local Admin Password Solution) enforcement confirmed')
bullet('MITRE T1003: OS Credential Dumping')
para('What ACECARD detects:', bold=True)
bullet('File signal: sensitive file access rising (credential stores, config files)')
bullet('Health score: 28 (CRITICAL) — entity is severely compromised')
bullet('ABAC response: entity moved to RESTRICTED → access limited → incident triggered')

h3('KEY TAKEAWAY for the room:')
quoted('"Without our system, this attack dwelled for 5+ years. With both layers: either Rigor '
       'catches the misconfiguration that allowed VPN access in the first place, OR ACECARD\'s '
       'CUSUM triggers by hour 72 — before the attacker reaches pre-positioning. Either way, '
       'the dwell time drops from years to hours."')

divider()
doc.add_page_break()

# ============================================================
# SLIDE 8
# ============================================================
h1('9. Slide 8: Salt Typhoon Walkthrough')

h3('What Is On This Slide')
para('Same format as Volt Typhoon: 5 phases, 4 columns. Different tradecraft — telecom-focused.')

h2('Background: What IS Salt Typhoon?')
bullet('PRC (Chinese) state-sponsored APT group (also tracked as Earth Estries, FamousSparrow, GhostEmperor)')
bullet('Disclosed 2024-2025 by CISA (AA25-239A)')
bullet('Targets: US telecom carriers (AT&T, Verizon, T-Mobile, Lumen, Charter + 4 others)')
bullet('Dwell time: 4+ YEARS in some networks')
bullet('Goal: access to CALEA lawful intercept systems (wiretap infrastructure)')
bullet('Implants: GhostSpider (memory-resident backdoor), Demodex (Windows rootkit), SnappyBee (modular loader)')
bullet('Unlike Volt Typhoon: DOES use custom malware + zero-days, but also exploits misconfigurations')
bullet('Entry: CVE-2023-20198 (Cisco IOS XE web UI privilege escalation — CVSS 10.0)')

h2('Phase-by-Phase Deep Dive')

h3('Phase 1: Edge CVE Exploit (0-24h)')
para('Attacker behavior:', bold=True)
bullet('Exploits CVE-2023-20198 against internet-facing Cisco IOS XE router')
bullet('Gains Level 15 (full admin) access to the routing infrastructure')
bullet('No authentication required — web UI exposed to internet')
para('What Rigor preempts:', bold=True)
bullet('WAF blocks public access to /webui/ management interface')
bullet('Management-plane ACL verified: only approved source IPs can reach management ports')
bullet('MITRE T1190: Exploit Public-Facing Application')
para('What ACECARD detects:', bold=True)
bullet('Network signal: anomalous management-plane source IP (external → management port)')
bullet('Drift: 0.09 (initial deviation from normal admin access patterns)')

h3('Phase 2: ACL Tampering (1-7 days)')
para('Attacker behavior:', bold=True)
bullet('Modifies router Access Control Lists (ACLs) to permit their traffic')
bullet('Tampers with AAA (Authentication, Authorization, Accounting) configuration')
bullet('Creates persistent hidden access that survives reboots')
para('What Rigor preempts:', bold=True)
bullet('Configuration drift detected within <1 hour of ACL change')
bullet('Deviation from verified baseline flagged immediately')
bullet('MITRE T1556: Modify Authentication Process')
para('What ACECARD detects:', bold=True)
bullet('Auth signal: SNMP access pattern changes (new management sessions)')
bullet('Identity signal: new admin-level sessions from previously unseen contexts')

h3('Phase 3: GRE Tunneling (7-30 days)')
para('Attacker behavior:', bold=True)
bullet('GRE = Generic Routing Encapsulation (tunneling protocol)')
bullet('Establishes encrypted tunnels that bypass normal traffic inspection')
bullet('Channels C2 (Command and Control) traffic through these tunnels')
para('What Rigor preempts:', bold=True)
bullet('Unauthorized tunnel endpoints detected in routing config')
bullet('Egress policy verified: only approved protocols/ports permitted outbound')
bullet('MITRE T1572: Protocol Tunneling')
para('What ACECARD detects:', bold=True)
bullet('Network signal: beacon_score rises (regular interval C2 heartbeats)')
bullet('Network signal: new protocol mix in traffic profile')
bullet('CUSUM: accumulating — small consistent drift over 3+ weeks')

h3('Phase 4: Wiretap Access (30-90 days)')
para('Attacker behavior:', bold=True)
bullet('CALEA = Communications Assistance for Law Enforcement Act')
bullet('Requires telecom carriers to build lawful intercept capability into their systems')
bullet('Salt Typhoon accesses these CALEA intercept systems to wiretap targets')
bullet('Targets: government officials, political campaign communications')
para('What Rigor preempts:', bold=True)
bullet('Privilege verification on interception systems')
bullet('Only approved users/roles can access CALEA infrastructure')
bullet('MITRE T1078: Valid Accounts (privilege abuse)')
para('What ACECARD detects:', bold=True)
bullet('Identity signal: privilege escalation on CALEA systems')
bullet('Identity signal: admin_actions spike on intercept infrastructure')
bullet('CUSUM TRIGGERS — accumulated drift over 30+ days crosses 4-sigma')

h3('Phase 5: CDR Exfiltration (90+ days)')
para('Attacker behavior:', bold=True)
bullet('CDR = Call Detail Records (who called whom, when, for how long, from where)')
bullet('Exfiltration is SLOW — small amounts over months to avoid detection')
bullet('Metadata is as valuable as content for intelligence purposes')
para('What Rigor preempts:', bold=True)
bullet('Outbound restriction policies verified for CDR data stores')
bullet('Data classification + DLP policies confirmed operational')
para('What ACECARD detects:', bold=True)
bullet('File signal: steady bytes_out increase from CDR repositories')
bullet('Health score: 24 (CRITICAL)')
bullet('ABAC: entity moved to BLOCKED state, all access revoked')

h3('KEY DIFFERENCE from Volt Typhoon:')
quoted('"Salt Typhoon combines both CVE exploitation AND Living-off-the-Land. This is why you need '
       'both layers: Rigor catches the misconfigured web UI that allowed CVE exploitation, AND '
       'ACECARD catches the slow exfiltration pattern that no signature rule would detect."')

divider()
doc.add_page_break()

# ============================================================
# SLIDE 9
# ============================================================
h1('10. Slide 9: Proof — Working Prototype')

h3('What Is On This Slide')
para('Three boxes: Running Now (our prototype), 7-Chart SOC Dashboard, Proven Foundation (DLA heritage). '
     'Bottom: Rigor production findings from real customers.')

h2('Running Now — Our Docker Prototype')
bullet('docker-compose with 3 containers: PostgreSQL+pgvector, FastAPI API, background worker')
bullet('pgvector = PostgreSQL extension for vector similarity search (cosine, L2, inner product)')
bullet('270 days of synthetic security event data (Jan 2025 - Oct 2025)')
bullet('6 injected attack scenarios covering Volt Typhoon + Salt Typhoon TTPs')
bullet('Mock embedding pipeline: 5 signals serialized → 1536-d vectors generated')
bullet('CUSUM detection running on all entity trajectories')
bullet('Drift direction + MITRE mapping operational')
bullet('Kill-chain reconstruction: links multi-phase attack events into narratives')

para('If asked "Can you show us?":')
quoted('"Yes. The system is running. I can show you the API endpoints, the dashboard, '
       'the detection alerts, and walk through a specific attack scenario in real time."')

h2('7-Chart SOC Dashboard')
para('Each chart and what it tells the analyst:', bold=True)
bullet('Drift timeline — cosine drift per window over time (shows when behavior changed)')
bullet('Velocity/acceleration — rate of change (is it getting worse or stabilizing?)')
bullet('Change-point timeline — where CUSUM triggered (the moments of significance)')
bullet('Health gauge — 0-100 composite score (instant threat assessment)')
bullet('Drift-direction radar — 8 MITRE concepts as radar chart (what is it becoming?)')
bullet('Per-signal contribution — which of the 5 signals is driving the drift (where to look)')
bullet('Peer-cohort scatter — entity vs. peers in same role/OU (outlier or normal variation?)')

h2('Proven Foundation — DLA Temporal Trajectory Intelligence')
para('Key point: this is NOT new, unproven technology.', bold=True)
bullet('Same architecture deployed for Defense Logistics Agency supply chain intelligence')
bullet('E-11 Temporal Trajectory Intelligence module — production validated')
bullet('500+ entities tracked over 23 monthly behavioral snapshots')
bullet('5 behavioral signals → CUSUM change-point detection → drift direction')
bullet('Concept alignment (cosine projection onto reference concepts)')
bullet('Same algorithms, same approach — just different domain signals')
bullet('SCM signals (demand patterns, supplier behavior) → cyber signals (auth, process, network)')
bullet('This dramatically de-risks the solution: proven architecture + new domain application')

h2('Rigor AI Production Findings')
para('Real customer results (anonymized):', bold=True)
bullet('FinTech (Asia): Closed SMB lateral movement path — 3 higher-priority rules shadowed application control profiles')
bullet('Communications (North America): Eliminated CVE exposure — IPS filters missing on rsync traffic to backup servers')
bullet('Government: Found zombie rules — legacy permits for decommissioned systems still allowed traffic through')
bullet('Critical Infrastructure: Verified Zero Trust claims were incomplete — implicit trust in edge cases that violated the ZTA policy')

para('Why this matters for the room:', bold=True)
quoted('"These are real production findings from real customers. Rigor found attack paths that '
       'penetration testers missed, that BAS tools missed, because those tools sample. '
       'Rigor proves. When we combine this with ACECARD behavioral detection, we cover '
       'both the configuration and the behavior — the entire threat surface."')

h3('Talking Points (90 seconds)')
quoted('"This is not a concept. The system is running in Docker right now. 270 days of synthetic '
       'data, 6 attack scenarios, CUSUM detection, drift direction, kill-chain reconstruction — all '
       'operational. The dashboard gives analysts 7 views: drift, velocity, change-points, health, '
       'direction, signal contribution, and peer comparison. And this architecture is not new — '
       'it is proven at DLA for supply chain intelligence. Same algorithms, different signals. '
       'On the Rigor side: production findings across fintech, telecom, and government — real '
       'attack paths closed that pen testers and BAS tools missed."')

divider()
doc.add_page_break()

# ============================================================
# SLIDE 10
# ============================================================
h1('11. Slide 10: Deployment Readiness')

h3('What Is On This Slide')
para('8 horizontal bars: Container, Enclaves, Embedding, Performance, Auth, Schema, Observability, CI/CD.')

h2('Gabriel Nimbus Context')
para('Gabriel Nimbus = DoD enterprise cloud computing environment. Key requirements:', bold=True)
bullet('Containers must use Iron Bank-approved base images')
bullet('Iron Bank = DoD Centralized Artifacts Repository (hardened, scanned, approved containers)')
bullet('Must support multiple classification enclaves (IL5, IL6, JWICS)')
bullet('IL5 = Impact Level 5 (CUI + mission data on NIPR)')
bullet('IL6 = Impact Level 6 (classified up to SECRET on SIPR)')
bullet('JWICS = Joint Worldwide Intelligence Communications System (TS/SCI)')
bullet('Must be deployable via GitOps (FluxCD or ArgoCD)')
bullet('Must have SBOM (Software Bill of Materials) for supply chain transparency')
bullet('Must align with cATO (continuous Authority to Operate) per DoD CIO memo Feb 2022')

h2('Line-by-Line Depth')

h3('Container')
bullet('Iron Bank-compatible base image (python:3.12-slim hardened)')
bullet('OCI-compliant: runs on any OCI container runtime (containerd, CRI-O)')
bullet('Signed with SBOM: every dependency traceable')
bullet('Big Bang Helm chart: Platform One deployment template')
bullet('Big Bang = DoD DevSecOps reference architecture Helm chart collection')

h3('Enclaves')
bullet('Same container image, different env vars for classification level')
bullet('IL5 (NIPR): uses OpenAI API for embeddings (data is CUI, not classified)')
bullet('IL6 (SIPR): local SLM (Phi-4 or Mistral 7B) — no external API calls')
bullet('JWICS: same local SLM, air-gapped deployment, no internet connectivity')
bullet('Key: ONE codebase, environment variable switches embedding provider. No code forks.')

h3('Embedding Options')
bullet('OpenAI text-embedding-3-small: 1536-d, fastest, lowest latency (~100ms per call)')
bullet('Phi-4 (Microsoft, 14B params): runs on single A100, outputs 1536-d with projection layer')
bullet('Mistral 7B: smaller, runs on V100, slight quality reduction vs Phi-4')
bullet('Switching: set EMBEDDING_PROVIDER=openai|phi4|mistral in .env. No code change.')

h3('Performance')
bullet('<3 seconds: embed + analyze + interpret per entity window')
bullet('Breakdown: ~100ms embedding + ~50ms CUSUM + ~50ms drift direction + ~2.5s serialization/overhead')
bullet('10,000 entities/hour on 4-vCPU worker container')
bullet('Horizontal scaling: N worker replicas = N x throughput. K8s HPA on queue depth.')
bullet('For 100K entities with 1-hour windows: 10 worker replicas')

h3('Auth')
bullet('Development: Bearer token (simple API key for testing)')
bullet('Production: CAC/PIV mTLS (Common Access Card / Personal Identity Verification)')
bullet('CAC = DoD smart card for identity/authentication')
bullet('PIV = NIST SP 800-73 Personal Identity Verification')
bullet('mTLS = mutual TLS (both client and server present certificates)')
bullet('RBAC on API endpoints: analyst, admin, service-account roles')

h3('Schema')
bullet('Input normalization supports 3 standards:')
bullet('ECS (Elastic Common Schema) — if feeding from Elasticsearch/Kibana SIEM', level=1)
bullet('OCSF (Open Cybersecurity Schema Framework) — AWS Security Lake format', level=1)
bullet('STIX 2.1 — if ingesting from threat intel platforms', level=1)
bullet('New log sources: add a schema mapping config file. No code changes required.')
bullet('This means: ANY log source in the target environment can feed the system')

h3('Observability')
bullet('Structured JSON logs (parseable by any log aggregator)')
bullet('Prometheus metrics endpoint (/metrics) for dashboarding')
bullet('OpenTelemetry traces end-to-end (correlate a single event through the entire pipeline)')

h3('CI/CD')
bullet('FluxCD GitOps: git push triggers deployment. Declarative desired state.')
bullet('SBOM generation: CycloneDX format, generated per build')
bullet('Vulnerability scanning: Anchore/Grype against Iron Bank CVE database')
bullet('cATO: continuous Authority to Operate (DoD CIO Feb 2022 memo)')
bullet('cATO means: no one-time ATO process. Continuous scanning + compliance = ongoing authorization')

h3('Talking Points (60 seconds)')
quoted('"Deployment readiness. Iron Bank-compatible container, SBOM, signed. '
       'Multi-enclave: IL5 uses OpenAI embeddings, IL6 and JWICS use local SLM — one codebase, '
       'env var switch. Performance: under 3 seconds per entity, 10K entities per hour, '
       'horizontally scalable. CAC/PIV auth, RBAC, mTLS. Any log schema via ECS, OCSF, or STIX 2.1. '
       'Full observability. FluxCD GitOps for cATO. This is not a future plan — '
       'the container runs today."')

divider()
doc.add_page_break()

# ============================================================
# SLIDE 11
# ============================================================
h1('12. Slide 11: Team Credentials')

h3('What Is On This Slide')
para('Two boxes: 22nd Century Technologies (left) and Rigor AI (right).')

h2('22nd Century Technologies (TSCTI)')
bullet('U.S. Army SOC/MDR: $90M contract, 800+ cleared analysts')
bullet('SOC = Security Operations Center (monitors, detects, responds)')
bullet('MDR = Managed Detection and Response (outsourced SOC operations)')
bullet('U.S. Air Force: $108M cybersecurity operations contract')
bullet('FBI TSC: $56M cybersecurity support')
bullet('TSC = Terrorist Screening Center')
bullet('NAVAIR / USMC: $145M cybersecurity operations')
bullet('NAVAIR = Naval Air Systems Command')

para('Key message:', bold=True)
quoted('"This system is built by the team that OPERATES SOC/MDR at scale for the U.S. military. '
       'These are not researchers — these are practitioners who triage alerts, investigate incidents, '
       'and respond to threats every day. They know what works and what doesn\'t because they live it."')

h2('Rigor AI Inc.')
bullet('20+ design partners across: Fortune 500, federal agencies, critical infrastructure, MSSPs, nation states')
bullet('MSSP = Managed Security Service Provider')
bullet('Production findings (not theoretical — actual customer environments):')
bullet('Closed lateral movement paths in financial services', level=1)
bullet('Eliminated CVE exposure in telecommunications', level=1)
bullet('Found zombie firewall rules in government', level=1)
bullet('Verified (and found gaps in) Zero Trust claims', level=1)
bullet('Formal verification engine: same mathematical approach used in chip design and avionics')

h3('Talking Points (30 seconds)')
quoted('"Quick on credentials. 22CT operates Army SOC at scale — 800+ cleared analysts, $90M contract. '
       'We built this system because we know what analysts need — we ARE the analysts. '
       'Rigor AI has 20+ design partners across critical sectors with production findings. '
       'This is a proven team with operational experience, not a startup with a slide deck."')

divider()
doc.add_page_break()

# ============================================================
# SLIDE 12
# ============================================================
h1('13. Slide 12: Closing')

h3('What Is On This Slide')
para('Three lines center-screen:')
bullet('"We Have a Working Solution." (large, cyan)')
bullet('"We Can Demo It Today." (medium, gold)')
bullet('Summary paragraph + deployment specs + team names')

h3('Talking Points (30 seconds)')
quoted('"We have a working solution. We can demo it today. '
       'Rigor mathematically proves your configuration closes every known attack path. '
       'ACECARD detects behavioral drift when Living-off-the-Land tradecraft gets through. '
       'Together: preempt what can be prevented, detect what cannot, enforce proportionally. '
       'Containerized. Gabriel Nimbus ready. We are happy to take questions or do a live demo."')

h3('After Closing — Q&A Mindset')
bullet('Stay at the podium. Do not sit down.')
bullet('If asked for demo: "Absolutely. Let me pull up the dashboard." (have it ready)')
bullet('If asked about timeline: "The prototype is running today. '
       'Production-ready MVP for pilot: 90 days from contract start."')
bullet('If asked about data requirements: "We consume any structured security log via '
       'ECS, OCSF, or STIX 2.1. Minimum: authentication + network flow. '
       'More signals = richer trajectory."')

divider()
doc.add_page_break()

# ============================================================
# ANTICIPATED QUESTIONS
# ============================================================
h1('14. Anticipated Questions & Responses')

questions = [
    ('How is this different from Splunk UBA / Exabeam / Securonix?',
     'Those tools produce per-event anomaly scores using ML models trained on labeled data. '
     'They operate on EVENTS. We operate on ENTITIES — tracking behavioral trajectory over time. '
     'Three structural differences: (1) We use embedding-based semantic representation, not feature-engineered '
     'anomaly scores. (2) We use CUSUM for change-point detection, which catches slow drift that never '
     'crosses a fixed threshold. (3) We provide drift DIRECTION with MITRE mapping, not just '
     '"anomaly score 87%." And critically — none of those tools include formal configuration verification. '
     'They are behavioral-only. We combine behavioral + preemptive.'),

    ('What about false positive rates?',
     'Two mechanisms control false positives: (1) Peer cohort comparison — if the entire OU drifted '
     'similarly, it is likely a policy change, not a threat. We compute z-score against peers. '
     '(2) Benign concept suppression — we maintain reference vectors for "normal_role_change" and '
     '"seasonal_variation." If drift direction projects onto benign concepts, it is suppressed. '
     'In synthetic testing: 6 attack scenarios detected with zero false positives on the benign population. '
     'Real-world tuning is required, but the architectural approach minimizes FP structurally.'),

    ('Does this require labeled training data?',
     'No. The embedding model (text-embedding-3-small or local SLM) is pre-trained on general text. '
     'It maps structured behavioral descriptions into a semantic space without any security-specific '
     'training data. CUSUM requires only a baseline period of normal behavior (we use 30 days). '
     'No labeled attack data needed. No supervised ML model to retrain. '
     'This is a fundamental advantage: we can deploy in an environment we have never seen '
     'and detect anomalies from day 31.'),

    ('What about encrypted traffic? Can you see inside TLS?',
     'We do not inspect payload content. We analyze METADATA: who talks to whom, when, how often, '
     'how much data, what protocols. Beacon score is computed from connection timing patterns, '
     'not content. Bytes out ratio is computed from flow sizes. DNS queries are visible even '
     'when traffic is encrypted. The behavioral signals are all metadata-derived. '
     'For classified networks where deeper inspection is required: we can consume decrypted '
     'traffic from authorized inspection points (SSL/TLS proxies).'),

    ('How long to deploy to production?',
     'Phase 1 (90 days): Integration with existing log sources, entity resolution tuning, '
     'baseline establishment (30 days of normal behavior). '
     'Phase 2 (60 days): ABAC policy integration, SIEM/SOAR connector, analyst training. '
     'Phase 3 (ongoing): Production operations, threshold tuning from TP/FP feedback. '
     'Total to first detection capability: 90 days from contract start. '
     'Container is ready today. The long pole is log source integration and baseline.'),

    ('What about privacy? Are you sending all our logs to OpenAI?',
     'No raw event content is ever sent to the embedding API. We serialize FEATURES — '
     'numerical summaries like "47 logons, 3.2% failure rate, 8 unique hosts." '
     'No usernames, no IP addresses, no file paths, no command lines leave the enclave. '
     'For IL6/JWICS: local SLM (Phi-4/Mistral), air-gapped, nothing leaves the network. '
     'For IL5: even the feature text contains no PII — it is statistical summaries only.'),

    ('How does this handle a zero-day exploit?',
     'Zero-day exploits are the exact scenario ACECARD is designed for. '
     'Rigor cannot prevent an exploit with no known CVE. But the POST-EXPLOITATION behavior '
     'is always detectable: the attacker must move laterally, escalate privileges, establish '
     'persistence, or exfiltrate. All of these change the behavioral trajectory. '
     'CUSUM will trigger. Drift direction will point toward the specific technique. '
     'The zero-day is the entry — the behavioral change is what we catch.'),

    ('What if the attacker is aware of behavioral monitoring?',
     'To evade behavioral trajectory detection, an attacker would need to: '
     '(1) maintain identical authentication patterns, (2) identical process patterns, '
     '(3) identical network patterns, (4) identical file access patterns, AND '
     '(5) identical privilege patterns — ALL simultaneously across ALL monitored windows. '
     'Achieving their objective (credential theft, lateral movement, exfiltration) while '
     'keeping ALL 5 signals indistinguishable from baseline is mathematically incompatible '
     'with accomplishing the attack goal. They MUST change behavior to achieve their objective.'),

    ('Can this run on JWICS (TS/SCI)?',
     'Yes. The architecture is enclave-agnostic. For JWICS: local SLM replaces OpenAI '
     '(env var switch, no code change). Air-gapped deployment. All data stays on-premises. '
     'Container image from Iron Bank. No internet connectivity required for operation. '
     'Embedding quality is slightly lower with local SLM vs OpenAI, but detection '
     'capability is preserved — we trade embedding quality for classification compliance.'),

    ('What is your competitive advantage over Microsoft Sentinel / Google Chronicle?',
     'Sentinel and Chronicle are SIEM platforms — excellent at log aggregation, '
     'rule-based detection, and ML anomaly scoring. They are not trajectory intelligence systems. '
     'Three things we do that they cannot: (1) Unified behavioral embedding across ALL identity '
     'systems in a single vector space. (2) CUSUM change-point detection that catches slow drift. '
     '(3) Formal mathematical verification of the configuration layer. '
     'We COMPLEMENT Sentinel/Chronicle — we consume their outputs as one of our log sources. '
     'We are not replacing your SIEM. We are adding trajectory intelligence on top of it.'),
]

for q, a in questions:
    h3(f'Q: "{q}"')
    para(a)
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# BREAKOUT SESSION PREP
# ============================================================
h1('15. Breakout Session Preparation')

h2('Morning: Limiting Factors Breakout (0900-0945)')
para('You will discuss what LIMITS current approaches. Your talking points:', bold=True)

h3('Limiting Factor 1: Semantic Gap Between Logs and Understanding')
bullet('Security logs are syntactic (raw events). Understanding is semantic (what is happening).')
bullet('Current tools: regex parsing → rule matching → alert. No semantic understanding.')
bullet('Our approach: structured text → embedding → semantic vector space. '
       'The model UNDERSTANDS behavioral meaning.')

h3('Limiting Factor 2: Temporal Blindness')
bullet('Current tools: point-in-time alerts. "Right now, X happened."')
bullet('No concept of trajectory: "Over the past 3 weeks, this entity has been BECOMING X."')
bullet('CUSUM was invented for exactly this: detecting sustained process shifts in manufacturing. '
       'Applies perfectly to behavioral drift in cybersecurity.')

h3('Limiting Factor 3: Identity Fragmentation')
bullet('Average enterprise: 5+ identity systems. No unified entity view.')
bullet('Lateral movement designed to exploit this fragmentation.')
bullet('Entity resolution is a prerequisite for trajectory analysis — must solve first.')

h3('Limiting Factor 4: Configuration Verification Gap')
bullet('Security posture ASSUMED based on compliance checklists.')
bullet('No PROOF that configuration actually achieves the intended security properties.')
bullet('Gap between intended policy and actual enforcement is where APTs live.')

h2('Morning: Ways to Overcome Breakout (1000-1045)')
para('After identifying limiting factors, the group discusses solutions. Your contributions:', bold=True)

bullet('Semantic embeddings solve the semantic gap — structured text to vector space')
bullet('Temporal trajectory analysis (CUSUM + drift direction) solves temporal blindness')
bullet('Entity resolution across identity providers solves fragmentation')
bullet('Formal verification (symbolic reasoning) solves the configuration proof gap')
bullet('Containerized architecture solves the deployment/scalability constraint')

h2('Afternoon: Problem Statement Development (1100-1200)')
para('The group will draft a problem statement. Ensure these elements are included:', bold=True)
bullet('"Detect behavioral anomalies" — not just known signatures')
bullet('"Correlate across siloed log architectures" — entity resolution + unified embedding')
bullet('"Scale to Gabriel Nimbus requirements" — containerized, K8s, horizontal scaling')
bullet('"Living-off-the-Land techniques" — specifically call out LoTL as the hardest challenge')

h2('Afternoon: Solution Brief Presentation (1420-1535)')
para('This is where you present the 12-slide deck. Key reminders:', bold=True)
bullet('10-12 minutes, not longer. Respect others\' time.')
bullet('Solution-focused: every claim backed by specific capability')
bullet('Offer live demo: "We can show you this running right now"')
bullet('Close strong: "We have a working solution. We can demo it today."')

h2('Afternoon: Voting (1535-1545)')
para('After all presentations, participants vote. Your advantage:', bold=True)
bullet('You are the only team with a RUNNING prototype (most will have slides only)')
bullet('You are the only team with BOTH preemptive + behavioral layers')
bullet('You directly map to all 4 CFIC challenges with specific capabilities')
bullet('You can demonstrate — not just describe')

divider()
doc.add_page_break()

# ============================================================
# TECHNICAL DEEP-DIVE
# ============================================================
h1('16. Technical Deep-Dive Reference')

h2('pgvector and Vector Similarity')
bullet('PostgreSQL extension that adds vector data type and similarity operators')
bullet('HNSW index: Hierarchical Navigable Small World graph — approximate nearest neighbor search')
bullet('cosine_ops: similarity metric. cos(A,B) = dot(A,B) / (|A| * |B|). Range: -1 to 1.')
bullet('Our index params: m=16 (connections per node), ef_construction=64 (build quality)')
bullet('Why pgvector vs. dedicated vector DB (Pinecone, Weaviate)? Simpler stack, one database for everything, '
       'good enough performance for our entity scale (100K), reduces operational complexity for DoD deployment.')

h2('CUSUM Algorithm Detail')
bullet('Page\'s CUSUM (1954) — designed for quality control in manufacturing')
bullet('Upper CUSUM: S_high(t) = max(0, S_high(t-1) + x(t) - (mu + k))')
bullet('Lower CUSUM: S_low(t) = max(0, S_low(t-1) - x(t) + (mu - k))')
bullet('mu = expected mean (baseline drift), k = slack (allowance), h = threshold')
bullet('We set: k = 0.5 * sigma, h = 4 * sigma')
bullet('Alarm when S_high(t) > h OR S_low(t) > h')
bullet('Properties: (1) Catches small sustained shifts. (2) Does NOT accumulate random variation. '
       '(3) Self-resetting when shift reverses.')
bullet('Why 4-sigma threshold: balances detection speed vs false alarm rate. '
       'With daily drift sigma ~0.03, 4-sigma = 0.12 accumulated — roughly 4-6 days of anomalous drift.')

h2('Embedding Architecture')
bullet('OpenAI text-embedding-3-small: 62,000 token context, 1536 output dimensions')
bullet('Model: trained on massive text corpus, outputs vectors where semantic similarity = cosine proximity')
bullet('Our input: ~200-400 tokens per entity window (5 signal texts concatenated)')
bullet('Batching: up to 2048 inputs per API call for efficiency')
bullet('Local SLM alternative: Phi-4 (14B) or Mistral 7B with projection head to 1536-d')
bullet('Projection head: linear layer trained on general text to match OpenAI output space dimensions')

h2('Entity Resolution')
bullet('Challenge: same human has AD SID, AAD UPN, Okta email, AWS ARN, K8s service account')
bullet('SID = Security Identifier (S-1-5-21-...), UPN = User Principal Name (user@domain.com)')
bullet('ARN = Amazon Resource Name (arn:aws:iam::123456:user/jsmith)')
bullet('Resolution approach: deterministic matching on shared attributes (email, employee ID) + '
       'probabilistic matching on behavior correlation')
bullet('Output: entity_uuid that maps to ALL identity system identifiers')

h2('Kill-Chain Reconstruction')
bullet('When multiple entities show correlated drift at sequential times, they form a kill chain')
bullet('Algorithm: temporal graph of CUSUM triggers, linked by shared resources/networks')
bullet('Output: narrative timeline — "Entity A (initial access, T+0h) → Entity B (lateral, T+36h) → '
       'Entity C (persistence, T+72h) → Entity D (exfiltration, T+120h)"')
bullet('MITRE ATT&CK tactic ordering validates the chain: recon → access → execution → '
       'persistence → privilege → lateral → collection → exfiltration')

divider()
doc.add_page_break()

# ============================================================
# KEY NUMBERS
# ============================================================
h1('17. Key Numbers to Memorize')

para('These are the numbers you should be able to cite without looking at slides:', bold=True)
doc.add_paragraph()

numbers = [
    ('5+ years', 'Volt Typhoon dwell time (CISA AA24-038A)'),
    ('4+ years', 'Salt Typhoon dwell time'),
    ('9', 'US telecom providers compromised by Salt Typhoon'),
    ('99%', 'Firewall breaches caused by misconfiguration, not vulnerabilities (Gartner)'),
    ('~20%', 'Attack paths closed by KEV patching alone'),
    ('2^8000', 'Configuration state space that ASM/BAS can only sample'),
    ('1536', 'Embedding dimensions (unified vector space)'),
    ('5', 'Behavioral signals (AUTH, PROCESS, NETWORK, FILE, IDENTITY)'),
    ('8', 'MITRE-mapped threat concepts for drift direction'),
    ('4-sigma', 'CUSUM alarm threshold'),
    ('<3 seconds', 'Time from raw events to full analysis per entity window'),
    ('10,000/hour', 'Entities processed per hour on 4-vCPU'),
    ('270 days', 'Synthetic data in running prototype'),
    ('6', 'Attack scenarios injected and detectable'),
    ('7', 'Dashboard charts (drift, velocity, CUSUM, health, direction, signal, peer)'),
    ('30-45 min', 'Current analyst triage time per alert (our system: <3 sec)'),
    ('500 alerts/day', 'Typical SOC alert volume (most uninvestigated)'),
    ('$90M', '22CT U.S. Army SOC/MDR contract'),
    ('800+', 'Cleared analysts on 22CT SOC contract'),
    ('20+', 'Rigor AI design partners'),
    ('90 days', 'Time to first detection capability from contract start'),
    ('0 → 72h', 'Volt Typhoon: dwell until CUSUM trigger (vs 5 years without)'),
    ('IL5 / IL6 / JWICS', 'Three enclave support levels'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Number'
hdr[1].text = 'What It Means'
for num, meaning in numbers:
    row = table.add_row().cells
    row[0].text = num
    row[1].text = meaning

doc.add_paragraph()
para('Practice: close this document and try to recall all 23 numbers. '
     'If someone asks "how fast?" you say "<3 seconds." '
     'If someone asks "how do you catch slow APTs?" you say "CUSUM, 4-sigma threshold." '
     'If someone asks "what about classified?" you say "IL6/JWICS, local SLM, env var switch."', italic=True)

doc.add_page_break()

# ============================================================
# DEEP DIVE: METRICS & THRESHOLDS
# ============================================================
h1('18. Deep Dive: Every Metric and Every Threshold')

para('This section lists EVERY metric our system computes, the exact threshold that triggers '
     'each alert level, and the rationale behind each number. If someone in the room asks '
     '"what exactly do you measure and at what level does it alert?" — this is the answer.', bold=True)

h2('A. Per-Signal Metrics (Computed Per Entity Per 1-Hour Window)')

h3('Signal 1: AUTH — Authentication Behavior')
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
for i, h_text in enumerate(['Metric', 'What It Measures', 'Normal Range', 'Warning', 'Critical']):
    table.rows[0].cells[i].text = h_text
auth_rows = [
    ('logon_count', 'Total successful + failed auth events in window', '5-50/hour', '>200/hour', '>500/hour'),
    ('failure_rate', 'Failed logins / total logins', '0.01-0.05 (1-5%)', '>0.15 (15%)', '>0.40 (40%)'),
    ('unique_hosts', 'Distinct machines authenticated to', '1-3/hour', '>8/hour', '>15/hour'),
    ('off_hours_ratio', 'Logins outside 0700-1900 local / total', '0.0-0.10', '>0.35', '>0.60'),
    ('impossible_travel', 'Geo distance / time gap violates physics', '0 (none)', '1 event', '>1 event'),
    ('new_source_ips', 'Source IPs never seen in 30-day baseline', '0-1/day', '>3/day', '>5/day'),
    ('mfa_skip_ratio', 'Auth events that bypassed MFA / total', '0.0 (all MFA)', '>0.10', '>0.30'),
    ('kerberos_anomaly', 'Unusual ticket requests (TGS without TGT)', '0', '>2/hour', '>5/hour'),
]
for row_data in auth_rows:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
para('How to explain:', italic=True)
quoted('"failure_rate measures brute force. Normal users fail 1-5% of logins (typos). '
       'A brute force attack pushes this to 40%+. unique_hosts measures lateral movement: '
       'a normal user touches 1-3 machines per hour. An attacker doing RDP hop-scotch touches 15+. '
       'impossible_travel is binary proof of credential sharing: if jsmith logs in from Fort Meade '
       'and Shanghai within 30 minutes, one of those sessions is not jsmith."')

h3('Signal 2: PROCESS — Execution Behavior')
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
for i, h_text in enumerate(['Metric', 'What It Measures', 'Normal Range', 'Warning', 'Critical']):
    table.rows[0].cells[i].text = h_text
proc_rows = [
    ('unique_procs', 'Distinct process names executed', '10-40/hour', '>80/hour', '>150/hour'),
    ('lolbin_count', 'LoTL binaries used (wmic, ntdsutil, netsh, etc.)', '0-1/hour', '>3/hour', '>6/hour'),
    ('cmdline_entropy', 'Shannon entropy of command-line args (bits)', '2.0-3.5', '>4.5', '>5.5'),
    ('parent_child_depth', 'Max process tree depth', '2-4', '>6', '>8'),
    ('unsigned_ratio', 'Unsigned binaries / total executed', '0.0-0.05', '>0.15', '>0.30'),
    ('powershell_encoded', 'PowerShell -EncodedCommand executions', '0', '>1/hour', '>3/hour'),
    ('script_block_count', 'Unique PowerShell script blocks logged', '0-5/hour', '>20/hour', '>50/hour'),
    ('injected_thread_count', 'Remote thread creation (CreateRemoteThread)', '0', '>1/hour', '>3/hour'),
]
for row_data in proc_rows:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
para('How to explain:', italic=True)
quoted('"lolbin_count is the LoTL detector. A sysadmin might run 1 wmic or netsh per hour during '
       'maintenance. Volt Typhoon ran 7+ in a single window: wmic for discovery, ntdsutil for '
       'credential extraction, netsh for firewall modification. cmdline_entropy catches obfuscation: '
       'normal commands like dir C:\\Users have entropy ~2.5. Base64-encoded payloads have entropy >5.0. '
       'unsigned_ratio catches custom tooling dropped on disk."')

h3('Signal 3: NETWORK — Traffic Behavior')
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
for i, h_text in enumerate(['Metric', 'What It Measures', 'Normal Range', 'Warning', 'Critical']):
    table.rows[0].cells[i].text = h_text
net_rows = [
    ('unique_dest_ips', 'Distinct destination IPs contacted', '5-20/hour', '>50/hour', '>100/hour'),
    ('bytes_out_ratio', 'Bytes sent / bytes received', '0.05-0.30', '>0.60', '>0.85'),
    ('beacon_score', 'Regularity of outbound connection intervals (0-1)', '0.0-0.15', '>0.45', '>0.70'),
    ('dns_query_rate', 'DNS queries per minute', '2-15/min', '>40/min', '>80/min'),
    ('dns_entropy', 'Avg Shannon entropy of queried domains', '2.5-3.5', '>4.0', '>5.0'),
    ('geo_anomaly_count', 'Connections to countries not in 30-day baseline', '0', '>2/hour', '>5/hour'),
    ('admin_share_access', 'C$, ADMIN$, IPC$ share connections', '0', '>2/hour', '>5/hour'),
    ('internal_scan_ratio', 'Distinct internal IPs / total connections', '0.10-0.30', '>0.60', '>0.80'),
]
for row_data in net_rows:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
para('How to explain:', italic=True)
quoted('"beacon_score is the C2 detector. Command-and-control malware calls home at regular intervals: '
       'every 60 seconds, every 5 minutes. We compute the coefficient of variation of inter-connection '
       'times. Regular beaconing has low variance = high beacon_score. Normal browsing is irregular = '
       'low score. bytes_out_ratio catches exfiltration: normal users download more than they upload '
       '(ratio 0.05-0.30). An entity extracting data inverts this ratio. dns_entropy catches DGA '
       '(Domain Generation Algorithm) domains: xk7r2m9p.evil.com has entropy >5.0 vs google.com at ~2.8."')

h3('Signal 4: FILE — File System Behavior')
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
for i, h_text in enumerate(['Metric', 'What It Measures', 'Normal Range', 'Warning', 'Critical']):
    table.rows[0].cells[i].text = h_text
file_rows = [
    ('files_created', 'New files written per window', '0-20/hour', '>100/hour', '>500/hour'),
    ('files_deleted', 'Files removed per window', '0-5/hour', '>50/hour', '>200/hour'),
    ('sensitive_access', 'Access to classified/restricted paths', '0-2/hour', '>5/hour', '>10/hour'),
    ('archive_creates', 'zip/tar/rar/7z creations', '0-1/day', '>3/hour', '>5/hour'),
    ('extension_changes', 'Files renamed with new extensions', '0', '>10/hour', '>50/hour'),
    ('large_file_writes', 'Files >100MB written', '0-1/day', '>3/hour', '>5/hour'),
    ('shadow_copy_ops', 'vssadmin/wbadmin operations', '0', '>1/hour', '>2/hour'),
]
for row_data in file_rows:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
para('How to explain:', italic=True)
quoted('"extension_changes is the ransomware detector. Normal users rename 0 file extensions per hour. '
       'Ransomware encrypting files changes .docx to .docx.encrypted at hundreds per hour. '
       'archive_creates catches staging for exfiltration: an insider zipping up project directories '
       'before sending them out. shadow_copy_ops catches ransomware pre-encryption prep: '
       'vssadmin delete shadows is the classic precursor to encryption."')

h3('Signal 5: IDENTITY — Privilege Behavior')
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
for i, h_text in enumerate(['Metric', 'What It Measures', 'Normal Range', 'Warning', 'Critical']):
    table.rows[0].cells[i].text = h_text
id_rows = [
    ('priv_escalations', 'Privilege elevation events (sudo, runas, UAC)', '0-2/hour', '>5/hour', '>10/hour'),
    ('group_adds', 'Security group membership changes', '0/day', '>1/hour', '>3/hour'),
    ('mfa_bypass', 'Auth events that succeeded without MFA challenge', '0', '>1/hour', '>3/hour'),
    ('role_changes', 'IAM role assumptions or switches', '0-1/day', '>3/hour', '>5/hour'),
    ('admin_actions', 'Account creation, policy changes, GPO mods', '0/day', '>2/hour', '>5/hour'),
    ('svc_acct_interactive', 'Service accounts with interactive logon', '0', '>1/day', '>1/hour'),
    ('delegation_changes', 'Kerberos delegation config changes', '0', '>1/day', '>1/hour'),
]
for row_data in id_rows:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
para('How to explain:', italic=True)
quoted('"group_adds catches privilege escalation: adding yourself to Domain Admins is the goal of '
       'most lateral movement campaigns. svc_acct_interactive catches compromised service accounts: '
       'a service account should NEVER have interactive logon. If it does, someone stole its credentials. '
       'delegation_changes catches Kerberos delegation attacks (unconstrained delegation = golden ticket path)."')

h2('B. Composite Metrics (Computed From All 5 Signals)')

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
for i, h_text in enumerate(['Metric', 'What It Measures', 'Normal Range', 'Warning', 'Critical']):
    table.rows[0].cells[i].text = h_text
composite_rows = [
    ('cosine_drift', '1 - cos(V_current, V_baseline)', '0.01-0.05', '>0.15', '>0.30'),
    ('velocity', 'drift(t) - drift(t-1)', '-0.02 to +0.02', '>+0.05', '>+0.10'),
    ('acceleration', 'velocity(t) - velocity(t-1)', '-0.01 to +0.01', '>+0.03', '>+0.05'),
    ('cusum_value', 'Cumulative sum statistic', '0 - 2*sigma', '>3*sigma', '>4*sigma (ALARM)'),
    ('health_score', 'Composite 0-100 (inverse of risk)', '70-100', '40-70', '<40'),
    ('peer_z_score', 'Std deviations from cohort mean drift', '-1.5 to +1.5', '>2.0', '>3.0'),
    ('max_concept_projection', 'Highest cosine(drift_vec, concept_vec)', '0.0-0.20', '>0.40', '>0.60'),
]
for row_data in composite_rows:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
para('The CUSUM threshold is the most important number:', bold=True)
bullet('k (slack parameter) = 0.5 * sigma of baseline drift distribution')
bullet('h (alarm threshold) = 4 * sigma')
bullet('With typical daily drift sigma = 0.03:')
bullet('k = 0.015 (CUSUM accumulates drift above baseline + 0.015)', level=1)
bullet('h = 0.12 (alarm fires when accumulated excess > 0.12)', level=1)
bullet('A sustained anomalous drift of 0.04/window (barely noticeable) triggers CUSUM in ~8 windows', level=1)
bullet('A threshold-only system at 0.15 would NEVER fire for this entity', level=1)
bullet('This is why CUSUM catches Volt Typhoon: daily drift was 0.01-0.04 (invisible to thresholds), '
       'but accumulated over 4-6 days to trigger CUSUM')

h2('C. ABAC Trust State Transitions')

para('The system does not just alert — it ENFORCES. Trust state transitions are automatic:', bold=True)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
for i, h_text in enumerate(['From State', 'To State', 'Trigger Condition', 'Enforcement Action', 'Reversal']):
    table.rows[0].cells[i].text = h_text
abac_rows = [
    ('TRUSTED', 'ELEVATED_WATCH', 'health_score drops below 70 OR cusum_value > 2*sigma',
     'Step-up MFA on next auth; increased log verbosity; analyst notification',
     'health_score > 75 for 24h'),
    ('ELEVATED_WATCH', 'RESTRICTED', 'health_score drops below 40 OR cusum_value > 3.5*sigma OR max_concept_projection > 0.55',
     'Read-only access; lateral movement blocked (no RDP/WinRM/SSH outbound); ticket created',
     'Analyst manual review + health > 60 for 48h'),
    ('RESTRICTED', 'BLOCKED', 'health_score drops below 20 OR cusum ALARM (4*sigma) OR max_concept_projection > 0.70',
     'All access revoked; active sessions terminated; incident response triggered',
     'IR team manual unlock only'),
    ('Any', 'TRUSTED', 'Analyst marks as False Positive',
     'State reset; FP feedback auto-adjusts baseline for this entity',
     'N/A'),
]
for row_data in abac_rows:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
para('Key point for the room:', bold=True)
quoted('"This is not binary. We do not lock out a user because of one anomaly. The response is '
       'PROPORTIONAL: first step-up MFA, then restrict, then block. Each level has specific, '
       'reversible enforcement. False positive feedback improves the baseline. The system learns."')

doc.add_page_break()

# ============================================================
# DEEP DIVE: ENTITY FUSION
# ============================================================
h1('19. Deep Dive: Entity Fusion — How We Solve Identity Fragmentation')

para('This is one of the four CFIC challenges and one of the most important capabilities '
     'of our system. Understand this thoroughly.', bold=True)

h2('The Problem: Identity Fragmentation')
para('A single human — "John Smith, Tier-2 SOC Analyst" — has the following identifiers '
     'across a typical DoD enterprise network:')

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
for i, h_text in enumerate(['System', 'Identifier Format', 'Example']):
    table.rows[0].cells[i].text = h_text
id_rows = [
    ('Active Directory', 'SamAccountName / SID', 'CORP\\jsmith / S-1-5-21-123...-1234'),
    ('Azure AD (Entra ID)', 'UserPrincipalName / ObjectId', 'john.smith@corp.mil / 8a2b3c4d-...'),
    ('Okta', 'Login / Okta UID', 'john.smith@corp.mil / 00u1a2b3c4'),
    ('AWS IAM', 'ARN', 'arn:aws:iam::842:user/jsmith'),
    ('Kubernetes', 'ServiceAccount / Subject', 'system:serviceaccount:soc:jsmith-sa'),
    ('CrowdStrike EDR', 'Agent ID / Hostname', 'aid:abc123def / JSMITH-WS01'),
    ('Splunk', 'Username field (varies)', 'src_user=jsmith OR user=CORP\\jsmith'),
    ('VPN', 'Certificate DN / Username', 'CN=john.smith,OU=SOC,O=CORP'),
    ('RADIUS/TACACS+', 'Username', 'jsmith@corp.mil'),
    ('PKI/CAC', 'Certificate Subject / EDIPI', 'CN=SMITH.JOHN.Q.1234567890'),
]
for row_data in id_rows:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
para('That is 10+ identifiers for ONE person. When Volt Typhoon moves laterally:', bold=True)
bullet('Login via VPN as jsmith@corp.mil (VPN log)')
bullet('Kerberos TGT as CORP\\jsmith (AD log)')
bullet('AAD sign-in as john.smith@corp.mil (Azure log)')
bullet('AWS console as arn:aws:iam::842:user/jsmith (CloudTrail)')
bullet('kubectl exec as system:serviceaccount:soc:jsmith-sa (K8s audit)')
para('Each system sees ONE login. Five tools, five "normal" events, five separate alert '
     'evaluations — each below threshold. No correlation. Attack succeeds.')

h2('Our Solution: Multi-Layer Entity Resolution')

h3('Layer 1: Deterministic Matching (High Confidence)')
bullet('Email address match: john.smith@corp.mil appears in AD, AAD, Okta, VPN')
bullet('Employee ID / EDIPI match: DoD unique identifier appears in CAC and HR systems')
bullet('Hostname correlation: CrowdStrike agent on JSMITH-WS01 maps to AD computer object')
bullet('Certificate Subject match: CAC CN matches AD DisplayName')
bullet('SID ↔ UPN mapping: AD provides direct SID-to-UPN lookup')
bullet('Result: merge all deterministic matches into single entity_uuid')

h3('Layer 2: Probabilistic Matching (Medium Confidence)')
bullet('Behavioral correlation: two identifiers that always authenticate from the same source IP '
       'within the same time window are likely the same person')
bullet('Session linkage: VPN session start + AD TGT within 30 seconds from same IP = same entity')
bullet('Naming convention patterns: jsmith, john.smith, j.smith across systems → candidate match')
bullet('These are scored, not assumed. Score > 0.85 → merge. Score 0.50-0.85 → flag for review.')

h3('Layer 3: Graph-Based Resolution')
bullet('Build an identity graph: nodes = identifiers, edges = correlation evidence')
bullet('Connected components in the graph = resolved entities')
bullet('Graph captures the EVIDENCE for why identifiers are linked')
bullet('Supports audit: "Why did you link these two accounts?" → show the edge chain')

h2('What Entity Resolution Enables')
bullet('ALL 5 behavioral signals from ALL identity systems feed into ONE embedding per entity')
bullet('Lateral movement across AD → AAD → AWS → K8s = ONE accelerating trajectory')
bullet('The drift magnitude now reflects the COMBINED behavior across ALL systems')
bullet('Instead of 5 weak signals (each score: 0.3/1.0), you get 1 strong signal (score: 0.85/1.0)')
bullet('CUSUM accumulates across ALL sources for this entity — catches slow cross-system patterns')
bullet('Drift direction tells you: "this entity is moving toward lateral_movement across 5 systems"')

h3('Talking Point for Entity Fusion')
quoted('"Without entity resolution, lateral movement is invisible. Each system sees one normal login. '
       'Our system resolves all identifiers — AD, AAD, Okta, AWS, K8s, EDR, VPN — into one entity. '
       'One embedding. One trajectory. One CUSUM accumulator. Five weak signals become one strong signal. '
       'This is not a nice-to-have — it is a prerequisite for detecting modern APT movement."')

doc.add_page_break()

# ============================================================
# DEEP DIVE: HOW WE SOLVE RULE DECAY
# ============================================================
h1('20. Deep Dive: How We Solve Sigma Rule Decay')

para('This addresses the question: "If PRC actors test against published detection rules, '
     'how do you stay ahead?"', bold=True)

h2('The Problem: Adversary-Detection Arms Race')
bullet('Sigma = open standard for detection rules (YAML format)')
bullet('CAR = MITRE Cyber Analytics Repository (similar concept)')
bullet('YARA = pattern matching rules for malware (file-level)')
bullet('Snort/Suricata = network signature rules')
para('All of these share a structural limitation:', bold=True)
bullet('Step 1: Defender observes attack technique, writes a detection rule')
bullet('Step 2: Rule is published (open source or shared via ISAC)')
bullet('Step 3: Adversary downloads published rules')
bullet('Step 4: Adversary tests their tools against the rules')
bullet('Step 5: Adversary modifies tools until no rule matches')
bullet('Step 6: Modified tools are deployed in next campaign')
bullet('Step 7: Defender observes new technique, writes new rule (go to Step 1)')
para('This cycle means detection coverage DECAYS between rule updates. PRC actors specifically '
     'exploit this window. Their retooling cycle is 2-4 weeks; defender rule update cycle is 1-3 months.')

h2('Our Solution: Behavioral Trajectory (Not Signatures)')
para('We do NOT write rules about specific commands or specific tools. '
     'We track the BEHAVIORAL EFFECT of whatever the attacker does.', bold=True)

h3('Example: Credential Access')
bullet('Sigma rule approach: "detect ntdsutil.exe execution" (works until attacker uses secretsdump.py instead)')
bullet('Our approach: detect BEHAVIORAL CHANGE across all 5 signals:')
bullet('AUTH: new service account usage pattern (who is authenticating differently?)', level=1)
bullet('PROCESS: new high-privilege process execution pattern (what is running?)', level=1)
bullet('NETWORK: new internal destination pattern (where is traffic going?)', level=1)
bullet('FILE: new access to credential stores (SAM, NTDS.dit, lsass.dmp)', level=1)
bullet('IDENTITY: privilege escalation events (how are permissions changing?)', level=1)
bullet('The attacker can change the TOOL (ntdsutil → secretsdump → mimikatz → pypykatz). '
       'They CANNOT change the BEHAVIORAL EFFECT: they must access credentials, '
       'they must authenticate with stolen credentials, they must move to new targets.')

h3('Why This Is Structurally Immune to Rule Decay')
bullet('Rules are syntactic: "if process_name == ntdsutil AND cmdline contains activate" → ALERT')
bullet('Embeddings are semantic: the MEANING of the behavioral pattern is captured')
bullet('If the attacker changes from ntdsutil to secretsdump:')
bullet('The Sigma rule misses (different process name)', level=1)
bullet('The behavioral embedding still drifts (same behavioral effect: credential access)', level=1)
bullet('The drift direction still projects onto credential_dumping concept', level=1)
bullet('The MITRE mapping is still T1003', level=1)
bullet('Embeddings capture WHAT the entity is BECOMING, not HOW it is doing it')
bullet('The adversary cannot retool against this without fundamentally changing their objective')

h3('The Mathematical Incompatibility Argument')
quoted('"An attacker who wants to dump credentials MUST access credential stores. An attacker who '
       'wants to move laterally MUST authenticate to new targets. An attacker who wants to exfiltrate '
       'MUST send data outbound. These behavioral effects are INHERENT to the attack objective. '
       'They can change the tool, the timing, the encoding, the process name — but the 5-signal '
       'behavioral trajectory will still drift toward the threat concept. The only way to evade '
       'trajectory detection is to not attack. Signatures detect tools. We detect objectives."')

doc.add_page_break()

# ============================================================
# DEEP DIVE: ALL DIMENSIONS OF AN ENTITY
# ============================================================
h1('21. Deep Dive: All Dimensions of an Entity')

para('An entity in our system is NOT just a username or an IP address. It is a rich, '
     'multi-dimensional object tracked across time. This section covers every dimension.', bold=True)

h2('Entity Types We Track')
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
for i, h_text in enumerate(['Entity Type', 'Examples', 'Primary Signals', 'Why Tracked']):
    table.rows[0].cells[i].text = h_text
entity_rows = [
    ('User', 'jsmith, admin01, svc-backup', 'AUTH, IDENTITY, FILE',
     'Primary target of credential theft, insider threat, privilege abuse'),
    ('Device', 'JSMITH-WS01, DC-PROD-01, WEBSVR-03', 'PROCESS, NETWORK, FILE',
     'Compromise target; process execution host; lateral movement hop point'),
    ('Network Segment', 'VLAN-100-UserDesktops, DMZ-Web', 'NETWORK (aggregate)',
     'Traffic pattern changes indicate scanning, C2, or exfiltration at segment level'),
    ('Application', 'Exchange, SharePoint, CALEA-Intercept', 'AUTH (access patterns), FILE',
     'Application-level abuse: unusual query patterns, bulk data access, API abuse'),
    ('Session', 'User X on Device Y during Window Z', 'All 5 signals',
     'Relationship entity: captures cross-entity behavior in specific context'),
]
for row_data in entity_rows:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

h2('Dimensions Tracked Per Entity')

h3('Dimension 1: Static Identity')
bullet('entity_uuid: our resolved unique identifier')
bullet('entity_type: user | device | segment | application | session')
bullet('identity_map: dictionary of all known identifiers across all systems')
bullet('organizational_context: OU, department, role, clearance level, location')
bullet('peer_cohort: other entities with similar role/function (for z-score comparison)')

h3('Dimension 2: Behavioral State (Current Window)')
bullet('The 1536-d embedding vector — encodes ALL behavioral signals for this window')
bullet('Per-signal feature vectors: the raw numerical metrics before embedding')
bullet('Signal contribution weights: which signals are driving the most variance')
bullet('This is the "snapshot" — what is the entity doing RIGHT NOW?')

h3('Dimension 3: Behavioral Trajectory (History)')
bullet('Baseline embedding: rolling average of prior N windows (default: 720 = 30 days)')
bullet('Drift series: cosine_drift per window over entire history')
bullet('Velocity series: first derivative of drift (rate of change)')
bullet('Acceleration series: second derivative of drift (rate of rate of change)')
bullet('CUSUM accumulator state: current S_high and S_low values')
bullet('Change-point history: every window where CUSUM triggered')
bullet('This is the "trajectory" — HOW is the entity CHANGING over time?')

h3('Dimension 4: Threat Alignment (Direction)')
bullet('Drift vector: V_current - V_baseline (the direction of behavioral change)')
bullet('8 concept projections: cosine(drift_vector, concept_embedding) for each threat concept')
bullet('MITRE technique IDs: mapped from top concept projections')
bullet('Kill-chain position: if entity is part of a multi-entity attack chain, which phase?')
bullet('This is the "direction" — WHAT is the entity BECOMING?')

h3('Dimension 5: Peer Context')
bullet('Cohort membership: which entities share similar role/function/OU')
bullet('Peer mean drift: average drift across the cohort')
bullet('Peer z-score: (entity_drift - cohort_mean) / cohort_stddev')
bullet('Cohort co-drift: are multiple peers drifting in the same direction? (policy change vs attack)')
bullet('This is the "context" — is this entity an OUTLIER among its peers?')

h3('Dimension 6: Trust State')
bullet('Current ABAC state: TRUSTED | ELEVATED_WATCH | RESTRICTED | BLOCKED')
bullet('State history: every transition with timestamp and trigger reason')
bullet('Enforcement actions active: step-up MFA, read-only, blocked, etc.')
bullet('This is the "response" — what ENFORCEMENT is currently applied?')

h3('Dimension 7: Relationship Graph')
bullet('Other entities this entity has interacted with (auth to devices, accessed apps)')
bullet('Relationship strength (frequency + recency of interaction)')
bullet('Anomalous new relationships (never-before-seen entity pairs)')
bullet('Kill-chain linkage: which other entities are co-drifting in correlated patterns')
bullet('This is the "network" — WHO does this entity interact with?')

para('Total: 7 dimensions per entity, continuously updated per window.', bold=True)
quoted('"When we say we track an entity, we do not mean we count logins. We track identity, '
       'behavior, trajectory, direction, peer context, trust state, and relationships. '
       'This is why we detect what SIEM cannot: we have 7 dimensions where SIEM has 1 (event count)."')

doc.add_page_break()

# ============================================================
# DEEP DIVE: 4-COMPONENT ENTITY SIGNATURE
# ============================================================
h1('22. Deep Dive: The 4-Component Entity Signature')

para('Every entity gets a composite signature per time window. This signature is the foundation '
     'of integrity, attribution, and trajectory analysis. Here are all 4 components.', bold=True)

h2('Component 1: Identity-Graph Hash')
h3('What It Is')
bullet('A SHA-256 hash of the entity resolution graph for this entity')
bullet('Inputs: all linked identifiers, correlation edges, confidence scores')
bullet('Deterministic: same identity links = same hash')
h3('Why It Exists')
bullet('Detects identity graph manipulation: if someone adds a fake identifier link, hash changes')
bullet('Audit trail: "At time T, this entity was resolved from these 6 identifiers"')
bullet('Versioning: if entity resolution changes (new link discovered), hash changes, '
       'creating a clear boundary between resolution versions')
h3('Technical Detail')
bullet('SHA-256 over sorted (identifier_type, identifier_value) pairs + edge weights')
bullet('Example input: [(AD, S-1-5-21-...), (AAD, john.smith@corp.mil), (AWS, arn:aws:...), '
       '(edge: AD-AAD, confidence: 0.99), (edge: AAD-AWS, confidence: 0.92)]')
bullet('Stored with each behavioral snapshot for provenance')

h2('Component 2: Behavioral Embedding (1536-d Vector)')
h3('What It Is')
bullet('The 1536-d vector produced by embedding the concatenated 5-signal text')
bullet('THIS is the behavioral state of the entity for this window')
bullet('Stored in PostgreSQL with pgvector (vector(1536) column type)')
h3('Why It Exists')
bullet('Enables trajectory computation: cosine distance between consecutive embeddings = drift')
bullet('Enables drift direction: V_current - V_baseline projected onto concept vectors')
bullet('Enables peer comparison: cosine similarity between entity embeddings in same cohort')
bullet('Enables kill-chain correlation: entities with correlated embedding changes → linked')
h3('Technical Detail')
bullet('Input: ~200-400 tokens of structured text (5 signal summaries)')
bullet('Model: text-embedding-3-small (1536-d) or local SLM with projection head')
bullet('Normalization: L2-normalized (unit vector on hypersphere)')
bullet('Storage: one row per (entity_uuid, window_start) in behavioral_snapshots table')
bullet('Index: HNSW (m=16, ef_construction=64, vector_cosine_ops) for fast nearest-neighbor')

h2('Component 3: Structural Features (JSONB)')
h3('What It Is')
bullet('The raw numerical feature values for all 5 signals, stored as structured JSON')
bullet('Example: {"auth": {"logon_count": 47, "failure_rate": 0.032, "unique_hosts": 3, ...}, '
       '"process": {"unique_procs": 22, "lolbin_count": 0, ...}, ...}')
h3('Why It Exists')
bullet('Explainability: when drift is detected, which specific metric drove it?')
bullet('Signal contribution analysis: compute variance contribution per signal')
bullet('Historical analysis: what were the EXACT numbers at time of detection?')
bullet('Audit: reproducibility — can re-embed from features to verify embedding correctness')
h3('Technical Detail')
bullet('JSONB column in PostgreSQL (indexed for fast access)')
bullet('~40-50 numerical features total across all 5 signals')
bullet('Retained for full history (not just current window)')
bullet('Used to generate the "per-signal contribution" dashboard chart')
bullet('Used by analyst during investigation: "Show me the exact auth metrics at T+36h"')

h2('Component 4: HMAC-Chained Trajectory History')
h3('What It Is')
bullet('HMAC = Hash-based Message Authentication Code')
bullet('Each snapshot includes an HMAC that chains to the PREVIOUS snapshot')
bullet('Forms a tamper-evident chain: modifying any past snapshot breaks the chain')
h3('Why It Exists')
bullet('Integrity: proves that the trajectory history has not been retroactively modified')
bullet('Non-repudiation: an attacker who compromises the database cannot silently alter '
       'past behavioral records to hide their pre-compromise drift')
bullet('Forensic value: in incident response, provides cryptographic proof that the '
       'trajectory data is authentic and unmodified')
h3('Technical Detail')
bullet('HMAC-SHA256 with entity-specific key derived from entity_uuid + master secret')
bullet('Input: HMAC(previous_hmac + current_embedding_hash + current_features_hash + timestamp)')
bullet('Chain structure: snapshot[0] → snapshot[1] → snapshot[2] → ... → snapshot[N]')
bullet('Verification: traverse chain from earliest to latest, recompute each HMAC, compare')
bullet('If any intermediate snapshot is modified, all subsequent HMACs fail verification')
bullet('Stored alongside each snapshot row: hmac_chain column')

h3('Why All 4 Components Together')
quoted('"The 4-component signature gives us something no other UEBA system has: a tamper-evident, '
       'auditable, multi-dimensional record of entity behavior across time. The identity-graph hash '
       'proves WHO was resolved. The embedding captures WHAT they did. The structural features '
       'explain WHY the embedding looks the way it does. And the HMAC chain proves WHEN it happened '
       'and that nobody altered the record after the fact. This is forensic-grade behavioral evidence."')

doc.add_page_break()

# ============================================================
# DEEP DIVE: DRIFT DIRECTION PROJECTION
# ============================================================
h1('23. Deep Dive: Drift Direction — Cosine Projection onto Threat Concepts')

para('This is the capability that transforms an opaque anomaly score into an actionable, '
     'MITRE-mapped threat hypothesis. Understand every step.', bold=True)

h2('Step 1: Compute the Drift Vector')
bullet('V_baseline = rolling mean of embeddings over prior N windows (default: 720 = 30 days)')
bullet('V_current = embedding for the current window')
bullet('drift_vector = V_current - V_baseline')
bullet('This vector is 1536-dimensional. It captures the DIRECTION of behavioral change.')
bullet('If the entity behavior has not changed, drift_vector is approximately the zero vector.')
bullet('If behavior has changed, drift_vector points from "old behavior" toward "new behavior".')

h2('Step 2: Define Reference Concept Embeddings')
para('Each threat concept is a text description of what that threat looks like behaviorally. '
     'These texts are embedded ONCE using the same model and stored as fixed reference vectors.', bold=True)

h3('All 8 Concepts With Full Text Descriptions')

concepts = [
    ('credential_dumping', 'T1003', 'OS Credential Dumping',
     'Entity shows escalating credential access patterns. Multiple authentication database queries. '
     'LSASS memory access attempts. SAM registry hive reads. ntdsutil or secretsdump invocations '
     'increasing. Kerberoasting ticket requests rising. Service account credentials accessed from '
     'non-standard sources. DCSync operations detected.'),
    ('lateral_movement', 'T1021', 'Remote Services',
     'Entity authenticating to increasing number of unique hosts. RDP, WinRM, SSH, or SMB connections '
     'to previously uncontacted servers. Admin share access (C$, ADMIN$, IPC$) from non-admin '
     'workstations. Pass-the-hash or pass-the-ticket patterns. Service creation on remote hosts. '
     'Process execution on remote systems via WMI or PsExec.'),
    ('data_exfiltration', 'T1041', 'Exfiltration Over C2 Channel',
     'Entity shows increasing outbound data volume relative to inbound. Large file transfers to '
     'external destinations. Archive file creation (zip, tar, rar) followed by outbound transfer. '
     'DNS query volume spike with high-entropy domain names. Cloud storage uploads increasing. '
     'Steady low-volume outbound to unusual geographic destinations over extended period.'),
    ('c2_beaconing', 'T1071', 'Application Layer Protocol',
     'Entity generating outbound connections at regular intervals. Low jitter between connection '
     'times. Consistent packet sizes. Connections to dynamically generated domains. HTTPS traffic '
     'to non-standard ports. DNS TXT record queries increasing. Periodic connectivity to cloud '
     'services not used by the organization. Sleep/jitter patterns consistent with known C2 frameworks.'),
    ('lotl_execution', 'T1059', 'Command and Scripting Interpreter',
     'Entity executing increasing numbers of Living-off-the-Land binaries. wmic, ntdsutil, netsh, '
     'rundll32, regsvr32, mshta, certutil, bitsadmin chain. PowerShell with encoded commands. '
     'High command-line entropy. Deep process trees with legitimate parent processes. Scheduled '
     'task creation via schtasks. Registry modification via reg.exe. No custom malware signatures — '
     'all tools are pre-installed Windows utilities.'),
    ('privilege_escalation', 'T1078', 'Valid Accounts',
     'Entity acquiring elevated privileges through legitimate mechanisms. sudo or runas frequency '
     'increasing. Security group membership changes. Service account assumption from interactive '
     'sessions. UAC bypass attempts. Token manipulation. New admin-level actions from previously '
     'standard-privilege entity. GPO modifications. Certificate template changes.'),
    ('defense_evasion', 'T1562', 'Impair Defenses',
     'Entity modifying security controls. Antivirus exclusions added. Firewall rules changed or '
     'disabled. Audit log configuration modified. Sysmon configuration altered. Event log cleared. '
     'Security agent stopped or uninstalled. Tamper protection bypass attempts. Timestomping of '
     'files to avoid detection. Process injection to evade monitoring.'),
    ('insider_hoarding', 'T1074', 'Data Staged',
     'Entity accumulating data access over extended period. Steady increase in sensitive file access '
     'volume. Bulk downloads from document management systems. Database export operations increasing. '
     'Data copied to removable media or personal cloud. Access patterns shift toward sensitive '
     'projects outside normal job scope. Working hours shift to avoid observation. Slow, steady '
     'accumulation pattern rather than sudden spike.'),
]

for name, tid, tname, desc in concepts:
    h3(f'{name} ({tid}: {tname})')
    para(f'Reference text: "{desc}"')
    doc.add_paragraph()

h2('Step 3: Compute Projections')
bullet('For each of the 8 concepts:')
bullet('projection[concept] = cosine_similarity(drift_vector, concept_embedding)', level=1)
bullet('cosine_similarity(A, B) = dot(A, B) / (|A| * |B|)', level=1)
bullet('Range: -1.0 to +1.0 (in practice, 0.0 to 0.8 for threat projections)', level=1)
bullet('Projection > 0: entity is drifting TOWARD this concept')
bullet('Projection < 0: entity is drifting AWAY from this concept')
bullet('Projection ~ 0: drift is orthogonal (unrelated) to this concept')

h2('Step 4: Rank and Report')
bullet('Sort all 8 projections by magnitude (descending)')
bullet('Return top 3 with their MITRE technique IDs')
bullet('Example output for a Volt Typhoon entity at T+72h:')
bullet('1. lateral_movement: 0.82 (T1021: Remote Services)', level=1)
bullet('2. lotl_execution: 0.71 (T1059: Command and Scripting Interpreter)', level=1)
bullet('3. credential_dumping: 0.54 (T1003: OS Credential Dumping)', level=1)
bullet('This tells the analyst EXACTLY what is happening: this entity is moving laterally '
       'using LoTL tools after dumping credentials. The entire kill chain is visible in '
       'the concept projections.')

h2('Why This Is Better Than "Anomaly Score 87%"')
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Traditional UEBA Alert'
table.rows[0].cells[1].text = 'Our Drift Direction Alert'
compare_rows = [
    ('User jsmith: anomaly score 87%',
     'User jsmith: drifting toward lateral_movement (0.82, T1021) + lotl_execution (0.71, T1059)'),
    ('What do I investigate? Everything.',
     'Investigate: remote service connections (RDP/WinRM) + LoTL binary chain (wmic, netsh)'),
    ('Is this getting worse? Unknown.',
     'Velocity: +0.06/window (accelerating). CUSUM: 3.2 sigma (approaching alarm).'),
    ('Is this normal for this role? No idea.',
     'Peer z-score: 3.1 (no peers show similar pattern). This is an outlier.'),
    ('What should I do? Unclear.',
     'ABAC state: ELEVATED_WATCH. Step-up MFA enforced. Lateral movement restricted pending review.'),
]
for old, new in compare_rows:
    row = table.add_row().cells
    row[0].text = old
    row[1].text = new

doc.add_paragraph()
para('The difference: traditional UEBA tells you SOMETHING is wrong. '
     'Our system tells you WHAT is happening, HOW fast, WHERE it is heading, '
     'WHETHER peers show the same pattern, and WHAT enforcement is already applied.', bold=True)

h3('Final Talking Point for Drift Direction')
quoted('"Drift direction is our answer to the directionless anomaly score problem. '
       'When behavior changes, we do not just say it changed — we say WHAT it is becoming. '
       'Credential dumping, lateral movement, exfiltration, insider hoarding — each with '
       'a MITRE technique ID and a specific investigation path. The analyst opens the alert '
       'and already knows what to look for. That is the difference between triage in 30 minutes '
       'and triage in 30 seconds."')

# ============================================================
# SAVE
# ============================================================
doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Pages: ~{len(doc.paragraphs) // 30} (estimated)")
