"""Build comprehensive deep-dive reference document with visual diagrams.
Includes: data flow visuals, architecture diagrams, callout boxes, comparison tables.
Output: CFIC_Deep_Dive_Reference.docx"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT = os.path.join(os.path.dirname(__file__), "CFIC_Deep_Dive_Reference_v3.docx")

doc = Document()

# Set narrow margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── STYLES ───
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Georgia'
    h.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 1'].paragraph_format.space_before = Pt(24)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 2'].paragraph_format.space_before = Pt(18)
doc.styles['Heading 3'].font.size = Pt(13)
doc.styles['Heading 3'].paragraph_format.space_before = Pt(12)

# ─── HELPERS ───
def set_cell_bg(cell, color_hex):
    """Set background color of a table cell"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, font='Calibri'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_para(text, bold=False, italic=False, size=None, color=None, space_after=None, space_before=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p

def add_bullet(text, level=0, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.bold = bold
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

def callout_box(title, body, color_hex='1B4F72'):
    """Create a visually distinct callout box using a single-cell table"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, 'F0F4F8')
    # Add left border via XML
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:color="{color_hex}"/>'
        f'<w:top w:val="single" w:sz="4" w:color="E0E0E4"/>'
        f'<w:bottom w:val="single" w:sz="4" w:color="E0E0E4"/>'
        f'<w:right w:val="single" w:sz="4" w:color="E0E0E4"/>'
        f'</w:tcBorders>')
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(int(color_hex[:2],16), int(color_hex[2:4],16), int(color_hex[4:],16))

    p2 = cell.add_paragraph()
    run2 = p2.add_run(body)
    run2.font.size = Pt(10.5)
    doc.add_paragraph()  # spacing

def key_insight_box(text):
    """Red-accented insight callout"""
    callout_box('KEY INSIGHT', text, 'C0392B')

def info_box(title, text):
    """Blue-accented information box"""
    callout_box(title, text, '1B4F72')

def visual_flow_table(steps, title=None):
    """Create a horizontal flow diagram using a table with arrows"""
    if title:
        add_para(title, bold=True, size=12, space_before=12)

    # Create table: alternating content and arrow columns
    n_cols = len(steps) * 2 - 1
    table = doc.add_table(rows=2, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (step_num, step_name, step_desc) in enumerate(steps):
        col_idx = i * 2
        # Top row: step number + name
        cell = table.rows[0].cells[col_idx]
        set_cell_bg(cell, '0D1B2A')
        set_cell_text(cell, f'{step_num}. {step_name}', bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

        # Bottom row: description
        cell = table.rows[1].cells[col_idx]
        set_cell_bg(cell, 'F7F8FA')
        set_cell_text(cell, step_desc, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Arrow column (except last)
        if i < len(steps) - 1:
            arrow_idx = col_idx + 1
            cell = table.rows[0].cells[arrow_idx]
            set_cell_text(cell, '→', bold=True, size=16, color=RGBColor(0x1B,0x4F,0x72), align=WD_ALIGN_PARAGRAPH.CENTER)
            cell2 = table.rows[1].cells[arrow_idx]
            cell2.text = ''

    doc.add_paragraph()  # spacing

def architecture_layer(layers):
    """Create a layered architecture visual"""
    table = doc.add_table(rows=len(layers), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    colors = ['4A90D9', '357ABD', '2E6EA6', '1B4F72', '0D1B2A']

    for i, (layer_name, components, purpose) in enumerate(layers):
        color = colors[i] if i < len(colors) else '0D1B2A'

        # Layer label
        cell0 = table.rows[i].cells[0]
        set_cell_bg(cell0, color)
        set_cell_text(cell0, layer_name, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

        # Components
        cell1 = table.rows[i].cells[1]
        set_cell_bg(cell1, 'F7F8FA')
        set_cell_text(cell1, components, size=9.5)

        # Purpose
        cell2 = table.rows[i].cells[2]
        set_cell_bg(cell2, 'EBF5FB')
        set_cell_text(cell2, purpose, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

def comparison_table(title_left, title_right, rows_data, left_color='FDEDEC', right_color='E8F4F8'):
    """Side-by-side comparison visual"""
    table = doc.add_table(rows=len(rows_data)+1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    set_cell_bg(table.rows[0].cells[0], 'C0392B')
    set_cell_text(table.rows[0].cells[0], title_left, bold=True, size=11, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[0].cells[1], '0E6B8A')
    set_cell_text(table.rows[0].cells[1], title_right, bold=True, size=11, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, (left, right) in enumerate(rows_data):
        set_cell_bg(table.rows[i+1].cells[0], left_color)
        set_cell_text(table.rows[i+1].cells[0], left, size=10)
        set_cell_bg(table.rows[i+1].cells[1], right_color)
        set_cell_text(table.rows[i+1].cells[1], right, size=10)

    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Containerized Behavioral Cyber Defense\nfor Gabriel Nimbus')
run.font.size = Pt(26)
run.font.name = 'Georgia'
run.bold = True
run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Deep-Dive Technical Reference')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Preemptive Posture Verification  +  ACECARD UEBA Behavioral Intelligence')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x0E, 0x6B, 0x8A)

doc.add_paragraph()
doc.add_paragraph()

# Visual summary box on title page
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = table.rows[0].cells
set_cell_bg(cells[0], 'FDEDEC')
set_cell_text(cells[0], 'PROBLEM\n\nNation-state APTs dwell for years using LOTL tradecraft that defeats all current tools', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(cells[1], 'EBF5FB')
set_cell_text(cells[1], 'SOLUTION\n\nLayer 1: Formal config verification\nLayer 2: Behavioral trajectory + CUSUM', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(cells[2], 'E8F8F5')
set_cell_text(cells[2], 'RESULT\n\nAttack surface eliminated OR drift detected in hours, not years. 90 days to deploy.', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('22nd Century Technologies Inc.')
run.font.size = Pt(13)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CFIC Collaboration Event  |  Augusta, GA  |  May 6, 2026\nPROPRIETARY // NOT FOR PUBLIC RELEASE')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'Nation-state adversaries — Volt Typhoon and Salt Typhoon — have dwelled inside United States critical '
    'infrastructure for years without detection. These campaigns exploit two structural blind spots:')

# Visual summary of the two gaps
table = doc.add_table(rows=2, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_bg(table.rows[0].cells[0], 'C0392B')
set_cell_text(table.rows[0].cells[0], '  GAP #1', bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
set_cell_bg(table.rows[0].cells[1], 'C0392B')
set_cell_text(table.rows[0].cells[1], '  GAP #2', bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
set_cell_bg(table.rows[1].cells[0], 'FDEDEC')
set_cell_text(table.rows[1].cells[0], 'Defensive configurations are never mathematically PROVEN complete. Pentests sample. ASM enumerates. Nobody proves.', size=10)
set_cell_bg(table.rows[1].cells[1], 'FDEDEC')
set_cell_text(table.rows[1].cells[1], 'Behavioral detection uses fixed thresholds that Living-off-the-Land tradecraft operates below — indefinitely.', size=10)
doc.add_paragraph()

doc.add_paragraph('Our two-layer solution closes both gaps simultaneously:')

# Solution visual
table = doc.add_table(rows=2, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_bg(table.rows[0].cells[0], '1B4F72')
set_cell_text(table.rows[0].cells[0], '  LAYER 1: 22CT PREEMPTIVE', bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
set_cell_bg(table.rows[0].cells[1], '0E6B8A')
set_cell_text(table.rows[0].cells[1], '  LAYER 2: ACECARD UEBA', bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
set_cell_bg(table.rows[1].cells[0], 'EBF5FB')
set_cell_text(table.rows[1].cells[0], 'Formal mathematical verification of every NGFW, IPS, IdP, SASE, WAF configuration.\n\nExhaustively reasons over combinatorial state space.\n\nZero FP/FN within formal model scope.\n\nRe-verified hourly + on every change.\n\nPowered by Rigor AI (innovation partner).', size=10)
set_cell_bg(table.rows[1].cells[1], 'E8F4F8')
set_cell_text(table.rows[1].cells[1], '1536-d behavioral embeddings per entity per hour.\n\nCUSUM change-point catches sub-threshold drift.\n\nDrift DIRECTION mapped to MITRE ATT&CK.\n\nEntity fusion across 10+ identity systems.\n\nK8s-native, Iron Bank, Gabriel Nimbus ready.', size=10)
doc.add_paragraph()

info_box('DEPLOYMENT', 'Containerized, K8s-native, Iron Bank hardened. IL5 (CUI), IL6 (Secret), JWICS (TS/SCI). 90 days from authorization to first detection. DoD ZT RA Pillars 4 & 7 | NIST 800-53 Rev 5 | cATO pathway.')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 2. THE PROBLEM
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('2. The Problem: AI-Enabled Cyber Attacks', level=1)

doc.add_heading('2.1 AI as a Force Multiplier for Adversaries', level=2)

doc.add_paragraph(
    'Artificial intelligence has fundamentally shifted the balance between offense and defense. '
    'Nation-state actors leverage AI across the full attack lifecycle:')

# 3-column visual
table = doc.add_table(rows=2, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['AI-POWERED RECON', 'AI-ASSISTED EVASION', 'AI-EXTENDED DWELL']
colors = ['C0392B', 'B7950B', '0D1B2A']
bodies = [
    '• LLMs generate exploit code from CVEs in minutes\n• AI scans 100K+ endpoints for misconfigs\n• Automated supply-chain mapping\n• AI-driven password spraying adapts to lockout policies',
    '• C2 traffic mutation evades signatures\n• Pre-campaign payload testing vs defender rules\n• Generative AI phishing at scale\n• Polymorphic malware reshapes per-target',
    '• Adaptive beaconing mimics normal traffic\n• AI-driven lateral movement timing\n• Automated credential rotation\n• Self-healing implants re-establish access'
]
for i in range(3):
    set_cell_bg(table.rows[0].cells[i], colors[i])
    set_cell_text(table.rows[0].cells[i], f'  {headers[i]}', bold=True, size=9.5, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(table.rows[1].cells[i], 'F7F8FA')
    set_cell_text(table.rows[1].cells[i], bodies[i], size=9.5)
doc.add_paragraph()

doc.add_heading('2.2 Case Study: Volt Typhoon (MITRE G1017)', level=2)

info_box('THREAT PROFILE',
    'PRC state-sponsored group. Active since mid-2021. Objective: pre-positioning in U.S. critical infrastructure '
    '(power, water, telecom) for future sabotage during geopolitical crisis. NOT espionage — sabotage capability. '
    'Tradecraft: exclusively Living-off-the-Land. C2: KV botnet (compromised SOHO routers).\n\n'
    'Sources: CISA AA24-038A | Microsoft Threat Intelligence | FBI/NSA/CISA Joint Advisory (Dec 2024)')

add_para('Kill Chain — 5 Phases:', bold=True, size=12, space_before=12)

# Kill chain as visual table
table = doc.add_table(rows=6, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
# Header
for i, h in enumerate(['Phase', 'Name', 'Technique', 'Why Invisible']):
    set_cell_bg(table.rows[0].cells[i], '0D1B2A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

vt_phases = [
    ('1', 'Initial Access', 'CVE-2023-20198 (Cisco IOS XE)\nCVE-2024-3400 (PAN-OS)\nFortinet, Ivanti edge devices', 'Targets perimeter devices directly.\nNo phishing, no user interaction.'),
    ('2', 'Credential Access', 'ntdsutil.exe → NTDS.dit dump\nLSASS memory extraction\nwmic.exe, netsh.exe', 'Every tool is a SIGNED Microsoft\nbinary. No malware signature exists.'),
    ('3', 'Lateral Movement', 'RDP, WinRM, SMB\nValid stolen credentials\nC2 via KV botnet (SOHO routers)', 'Standard Windows protocols.\nC2 from residential IPs.\nIndistinguishable from admins.'),
    ('4', 'Persistence', 'Web shells (.aspx/.jspx)\nScheduled tasks\nRegistry modification', 'Common persistence mechanisms.\nNo custom implants.'),
    ('5', 'Defense Evasion', 'Event log clearing\nMRU key deletion\nZERO custom malware deployed', 'LOTL only = no signatures.\nEvidence actively destroyed.'),
]
alt_colors = ['F7F8FA', 'FFFFFF']
for i, (phase, name, tech, why) in enumerate(vt_phases):
    bg = alt_colors[i % 2]
    for j, val in enumerate([phase, name, tech, why]):
        set_cell_bg(table.rows[i+1].cells[j], bg)
        set_cell_text(table.rows[i+1].cells[j], val, size=9, align=WD_ALIGN_PARAGRAPH.CENTER if j==0 else WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

key_insight_box(
    'Every tool used is a legitimate admin binary. Every login uses valid credentials. No malware exists anywhere '
    'in the kill chain. Signature-based detection is STRUCTURALLY blind. Threshold-based anomaly detection cannot '
    'fire because each individual action is below any reasonable threshold. Only the PATTERN OF BEHAVIOR over '
    'time reveals the compromise — and no current tool tracks behavioral patterns across time.')

doc.add_heading('2.3 Case Study: Salt Typhoon', level=2)

info_box('THREAT PROFILE',
    'PRC-affiliated cyber espionage targeting telecommunications. 9+ U.S. telecoms confirmed compromised '
    '(AT&T, Verizon, T-Mobile, Lumen). Objective: CDR metadata + CALEA lawful-intercept access (knows who '
    'law enforcement is wiretapping). Tradecraft: config-resident persistence — lives IN the router config.\n\n'
    'Sources: FBI/CISA advisory (Dec 2024) | Cisco Talos | Trend Micro research')

table = doc.add_table(rows=6, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Phase', 'Name', 'Technique', 'Why Different']):
    set_cell_bg(table.rows[0].cells[i], '0D1B2A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

st_phases = [
    ('1', 'Edge Exploitation', 'CVE-2023-20198 (IOS XE)\nCVE-2018-0171 (Smart Install)\nCVE-2024-21887 (Ivanti)', 'Target IS the network device,\nnot a stepping stone.'),
    ('2', 'Config-Resident\nPersistence', 'Local accounts on routers\nACL whitelist modifications\nGuest Shell enablement', 'Persistence IS the config.\nNo file. No binary. No process.'),
    ('3', 'TACACS+ Harvest', 'On-box PCAP of TCP/49\nJumbledPath ELF binary\nCredential harvesting', 'Captures network engineer\ncredentials at the source.'),
    ('4', 'Collection', 'GRE tunnels → CDR/metadata\nCALEA intercept access\nGhostSpider + Demodex', 'Knows who law enforcement\nis surveilling.'),
    ('5', 'Long-Term Espionage', 'Slow CDR exfiltration\nEncrypted C2\nLOTL on network gear', 'Multi-year dwell.\nDeliberately low volume.'),
]
for i, (phase, name, tech, why) in enumerate(st_phases):
    bg = alt_colors[i % 2]
    for j, val in enumerate([phase, name, tech, why]):
        set_cell_bg(table.rows[i+1].cells[j], bg)
        set_cell_text(table.rows[i+1].cells[j], val, size=9, align=WD_ALIGN_PARAGRAPH.CENTER if j==0 else WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

key_insight_box(
    'Salt Typhoon literally lives INSIDE the router configuration. Their persistence IS an ACL entry, a local '
    'account, a Guest Shell config. There is no file to scan, no process to kill, no binary to signature-match. '
    'Only continuous config verification AND behavioral traffic analysis can detect them.')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 3. WHY CURRENT TOOLS FAIL
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Why Current Security Tools Fail', level=1)

doc.add_paragraph(
    'Each major tool category has STRUCTURAL limitations — not implementation failures — that make them '
    'blind to these specific adversary tradecraft patterns.')

# Visual comparison table
table = doc.add_table(rows=6, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Tool', 'What It Does', 'Why It Fails vs. Volt/Salt Typhoon', 'Verdict']
for i, h in enumerate(headers):
    set_cell_bg(table.rows[0].cells[i], '0D1B2A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

tools = [
    ('SIEM', 'Collects logs, correlates against known rules', 'Cannot detect NOVEL combinations of legitimate tools.\n10K+ alerts/day buries true positives.\nNo behavioral trajectory over time.', 'BLIND'),
    ('EDR/XDR', 'Monitors endpoints for malicious processes', 'LOLBins explicitly ALLOWLISTED.\nCannot run on network gear (Cisco IOS).\nPer-host analysis, no cross-host trajectory.', 'BLIND'),
    ('Commercial\nUEBA', 'Baselines behavior, alerts on threshold breach', 'FIXED thresholds: VT operates below ALL of them.\nNo drift DIRECTION.\nNo identity fusion.\nPer-source silos.', 'BLIND'),
    ('Pentest /\nRed Team', 'Simulates attacks on subset of paths', 'SAMPLING: cannot cover combinatorial state space.\nPoint-in-time (annual) vs continuous adversary.\nKnown TTP bias.', 'BLIND'),
    ('ASM / BAS', 'Enumerates surface, simulates known attacks', 'Cannot PROVE controls block every path.\nStatistical, not mathematical.\nMisses config-resident persistence.', 'BLIND'),
]
for i, (tool, does, fails, verdict) in enumerate(tools):
    bg = alt_colors[i % 2]
    set_cell_bg(table.rows[i+1].cells[0], '1B4F72')
    set_cell_text(table.rows[i+1].cells[0], tool, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[1], bg)
    set_cell_text(table.rows[i+1].cells[1], does, size=9)
    set_cell_bg(table.rows[i+1].cells[2], bg)
    set_cell_text(table.rows[i+1].cells[2], fails, size=9)
    set_cell_bg(table.rows[i+1].cells[3], 'FDEDEC')
    set_cell_text(table.rows[i+1].cells[3], verdict, bold=True, size=9, color=RGBColor(0xC0,0x39,0x2B), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

doc.add_heading('3.1 The Four Structural Gaps', level=2)

table = doc.add_table(rows=5, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Gap', 'Description', 'Exploited By', 'Required Innovation']):
    set_cell_bg(table.rows[0].cells[i], '0D1B2A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

gaps_data = [
    ('#1', 'Configs never PROVEN closed', 'Both (edge device entry)', 'Formal mathematical verification'),
    ('#2', 'Detection has no direction', 'Volt Typhoon (LOTL)', 'Drift direction + MITRE mapping'),
    ('#3', 'Identity fragmented (10+ sys)', 'Both (cross-domain)', 'Entity fusion → one trajectory'),
    ('#4', 'No K8s-native UEBA for Nimbus', 'Capability gap', 'Iron Bank, IL5-JWICS, cATO'),
]
for i, row in enumerate(gaps_data):
    bg = alt_colors[i % 2]
    for j, val in enumerate(row):
        set_cell_bg(table.rows[i+1].cells[j], bg)
        bold = j == 0
        set_cell_text(table.rows[i+1].cells[j], val, size=10, bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER if j==0 else WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

key_insight_box('The gap is structural, not operational. More rules, more alerts, more tools operating on the same '
                'principles will not close it. A fundamentally different approach is required — one that PROVES '
                'configurations are complete (formal verification) and one that detects behavioral DIRECTION '
                'over time (trajectory intelligence with CUSUM).')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 4. LAYER 1: PREEMPTIVE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Layer 1: 22CT Preemptive — Formal Verification', level=1)

doc.add_paragraph(
    'Addresses Gap #1. Instead of testing or sampling, uses formal mathematical methods to PROVE that every '
    'defensive configuration blocks every known attack path — or surfaces the exact gap with a vendor-specific fix.')

doc.add_heading('4.1 What Is Formal Verification?', level=2)

comparison_table(
    'TRADITIONAL SECURITY (Sampling)',
    'FORMAL VERIFICATION (Proving)',
    [
        ('Pentest checks ~100 attack paths', 'Reasons over ALL possible paths (billions)'),
        ('Annual point-in-time assessment', 'Continuous: hourly + on every config change'),
        ('"We tested it and it held"', '"We PROVED it blocks every known TTP"'),
        ('Statistical confidence (90-95%)', 'Mathematical certainty (within model)'),
        ('Misses novel path combinations', 'Finds EVERY gap in the formal model'),
        ('Result: "Probably secure"', 'Result: "Provably secure within scope"'),
    ]
)

doc.add_paragraph(
    'Formal verification exhaustively reasons over ALL possible states of a system to prove that certain '
    'properties hold universally — not just for the cases you tested, but for EVERY possible case. '
    'Borrowed from hardware design (Intel uses it to verify chip correctness) and safety-critical software '
    '(avionics, nuclear, medical devices).')

doc.add_heading('4.2 The Three Intelligences', level=2)

# Three intelligences visual
table = doc.add_table(rows=2, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
intels = [
    ('ATTACK\nINTELLIGENCE', '1B4F72',
     '• MITRE-enriched attack graphs\n• ML/NLP/LLM pipeline\n• Ingests: CVE, CISA, CTI, IOCs\n• Full TTP graph for G1017 + Salt\n• Auto-updates on new advisories'),
    ('DEFENSE\nINTELLIGENCE', '0E6B8A',
     '• Symbolic Model of Computation\n• Every config as formal logic\n• Exhaustive state-space reasoning\n• PROVES every path blocked\n• Or surfaces the exact gap'),
    ('REMEDIATION\nINTELLIGENCE', 'B7950B',
     '• Guardrailed agentic-AI\n• Vendor-specific config fixes\n• Validates no new gaps created\n• SOC playbooks auto-generated\n• ServiceNow/JIRA integration'),
]
for i, (name, color, body) in enumerate(intels):
    set_cell_bg(table.rows[0].cells[i], color)
    set_cell_text(table.rows[0].cells[i], name, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[1].cells[i], 'F7F8FA')
    set_cell_text(table.rows[1].cells[i], body, size=9.5)
doc.add_paragraph()

doc.add_heading('4.3 How It Works: The Process', level=2)

visual_flow_table([
    ('1', 'PARSE', 'Ingest device\nconfig (API/file)'),
    ('2', 'TRANSLATE', 'Rules → formal\nlogic predicates'),
    ('3', 'MODEL', 'Build complete\nstate-space model'),
    ('4', 'VERIFY', 'Prove all attack\npaths blocked'),
    ('5', 'REPORT', 'Surface gaps +\nvendor-specific fix'),
    ('6', 'RE-VERIFY', 'Hourly + on\nevery change'),
], 'Preemptive Verification Pipeline:')

doc.add_paragraph(
    'The combinatorial explosion is why sampling cannot work: a single PAN firewall with 500 rules × 50 zones × '
    '20 applications × 10 user groups × NAT × decryption × threat profiles = BILLIONS of unique packet-flow decisions. '
    'A pentest exercises hundreds. Formal verification covers ALL of them using SAT/SMT solvers that efficiently '
    'prune the search space.')

doc.add_heading('4.4 Zero FP/FN — Scope & Caveats', level=2)

callout_box('ZERO FALSE POSITIVES',
    'If the system reports a gap, that gap genuinely EXISTS. The formal proof is sound — it cannot hallucinate '
    'a vulnerability. This eliminates the single largest operational cost: analyst time on false alarms.', '0E6B8A')
callout_box('ZERO FALSE NEGATIVES',
    'If the system says "all paths blocked," then within the formal model scope, no known attack path is '
    'unmitigated. The proof is complete over the modeled state space.', '0E6B8A')
callout_box('SCOPE LIMITATION',
    'Bounded by the formal model. If a new TTP is discovered that the attack graph does not yet contain, '
    'or if a device has undocumented behavior, the guarantee does not extend. This is why Layer 2 (ACECARD) '
    'exists — to catch what escapes the formal model through behavioral detection.', 'B7950B')

doc.add_heading('4.5 Supported Vendors', level=2)

table = doc.add_table(rows=5, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Vendor', 'Products', 'Coverage']):
    set_cell_bg(table.rows[0].cells[i], '0D1B2A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
vendors = [
    ('Palo Alto Networks', 'PA-Series, Panorama, Prisma', 'NGFW, App-ID, URL, Threat Prevention, Decryption'),
    ('Fortinet', 'FortiGate, FortiManager', 'FW policies, IPS, Web Filter, App Control'),
    ('Cisco', 'ASA, Firepower/FTD, ISE', 'ACLs, IPS/IDS, NAC, Identity-based access'),
    ('Check Point', 'R81+, SmartConsole', 'Security policies, HTTPS inspection, TP'),
]
for i, (v, p, c) in enumerate(vendors):
    bg = alt_colors[i % 2]
    for j, val in enumerate([v, p, c]):
        set_cell_bg(table.rows[i+1].cells[j], bg)
        set_cell_text(table.rows[i+1].cells[j], val, size=9.5)
doc.add_paragraph()

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 5. LAYER 2: ACECARD UEBA
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Layer 2: ACECARD UEBA — Behavioral Trajectory Intelligence', level=1)

doc.add_paragraph(
    'ACECARD (Adaptive Cyber Entity Characterization, Assessment & Risk Detection) closes Gaps #2, #3, and #4. '
    'Transforms raw security telemetry into 1536-d behavioral vectors, tracks how they CHANGE over time.')

doc.add_heading('5.1 Embedding Pipeline — How Raw Events Become Vectors', level=2)

add_para('The 4-Step Pipeline:', bold=True, size=12)

visual_flow_table([
    ('1', 'AGGREGATE', 'Raw events →\nhourly metrics\nper entity'),
    ('2', 'SERIALIZE', 'Metrics →\nstructured\nnatural language'),
    ('3', 'EMBED', 'Text → 1536-d\nvector via\nembedding model'),
    ('4', 'COMPOSE', '5 signal vectors\n→ 1 composite\n(weighted avg)'),
], 'Embedding Pipeline:')

callout_box('WHY TEXT EMBEDDING (not raw numbers)?',
    '• Leverages semantic understanding of pre-trained language models\n'
    '• Captures RELATIONSHIPS between features that numerical vectors miss\n'
    '• Model-agnostic: swap embedding models without retraining anything\n'
    '• Enables direct comparison with threat concepts (also expressed as text)\n'
    '• Same 1536-d space for ALL signals, entities, and concepts = unified comparison', '0E6B8A')

doc.add_heading('5.2 The Five Behavioral Signals', level=2)

# Signals visual table
table = doc.add_table(rows=6, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Signal', 'Data Sources', 'Key Metrics', 'Attack Signature']):
    set_cell_bg(table.rows[0].cells[i], '0E6B8A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

signals = [
    ('AUTH', 'Win 4624/4625/4768\nOkta, AAD, CloudTrail', 'logon_count, failure_rate\nunique_hosts, off_hours_ratio\nimpossible_travel, mfa_skip', 'VT: impossible_travel +\nnew unique_hosts +\noff_hours spike'),
    ('PROCESS', 'Sysmon EID 1/3\nCrowdStrike, Defender\nauditd', 'unique_processes\ncmdline_entropy, lolbin_count\nparent_child_depth', 'VT: lolbin_count 0→7\n(ntdsutil, wmic, netsh)\nfrom quiet workstation'),
    ('NETWORK', 'NetFlow/IPFIX, Zeek\nPAN-OS, FortiGate\nVPC Flow Logs', 'unique_dest_ips\nbytes_out_ratio, beacon_score\ndns_query_rate, geo_anomaly', 'Both: beacon_score↑ (C2)\nunique_dest↑ (lateral)\nbytes_out↑ (exfil)'),
    ('FILE', 'Sysmon EID 11\nWin 4663, DLP\ncloud storage audit', 'files_created/deleted\nsensitive_access\narchive_creates\nshadow_copy_ops', 'VT: shadow_copy spike\n(NTDS.dit dump)\nStaging: archive_creates↑'),
    ('IDENTITY', 'Win 4672/4728\nCloudTrail IAM\nK8s audit, AD change', 'priv_escalations\ngroup_adds, mfa_bypass\nrole_changes\nservice_acct_use', 'Both: new accounts\nACL modifications\nservice account creation'),
]
for i, (sig, sources, metrics, attack) in enumerate(signals):
    bg = alt_colors[i % 2]
    set_cell_bg(table.rows[i+1].cells[0], '0E6B8A')
    set_cell_text(table.rows[i+1].cells[0], sig, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[1], bg)
    set_cell_text(table.rows[i+1].cells[1], sources, size=9)
    set_cell_bg(table.rows[i+1].cells[2], bg)
    set_cell_text(table.rows[i+1].cells[2], metrics, size=9, font='Consolas')
    set_cell_bg(table.rows[i+1].cells[3], 'FEF9E7')
    set_cell_text(table.rows[i+1].cells[3], attack, size=9)
doc.add_paragraph()

doc.add_heading('5.3 CUSUM Change-Point Detection — The Mathematical Core', level=2)

doc.add_paragraph(
    'The Cumulative Sum (CUSUM) algorithm is what makes ACECARD fundamentally different from every commercial UEBA. '
    'It solves the problem that defeats all threshold-based systems.')

# CUSUM formula box
callout_box('THE ALGORITHM — Page\'s CUSUM (1954)',
    'S(t) = max(0, S(t-1) + drift(t) - μ - k)\n\n'
    'S(t) = cumulative sum statistic (starts at 0, resets on negative)\n'
    'drift(t) = cosine distance between current embedding and baseline\n'
    'μ = expected baseline drift (calibrated per entity from normal data)\n'
    'k = slack = 0.5 × σ (half standard deviation of normal drift)\n'
    'ALARM when S(t) > h = 4 × σ (configurable threshold)\n\n'
    'Minimax optimal for i.i.d. shift detection (Moustakides 1986):\n'
    'For any given false alarm rate, no other algorithm detects a shift faster.', '0E6B8A')

add_para('Why CUSUM Catches Volt Typhoon — Worked Example:', bold=True, size=12, space_before=12)

# CUSUM comparison visual
comparison_table(
    'FIXED THRESHOLD (Current UEBA)',
    'CUSUM (Our Approach)',
    [
        ('Threshold set at 0.15 (to avoid false positives)', 'Parameters: μ=0.03, σ=0.03, k=0.015, h=0.12'),
        ('Volt Typhoon drift per window: 0.01–0.06', 'Same drift: 0.01–0.06 per window'),
        ('Maximum single-window deviation: 0.06', 'Excess per window: 0.06 - 0.03 - 0.015 = 0.015'),
        ('0.06 < 0.15 threshold → NO ALERT', 'Cumulative: 0.015 accumulates each window'),
        ('Result: NEVER fires. Ever.', 'Result: alarm at 0.12/0.015 = 8 windows = 8 HOURS'),
        ('Dwell time: YEARS', 'Dwell time: HOURS'),
        ('Every commercial UEBA uses this approach', 'No commercial UEBA uses CUSUM on embeddings'),
    ]
)

key_insight_box(
    'This is not a tuning difference. Fixed thresholds can NEVER catch drift that stays below them, no matter '
    'how you tune them. Lowering the threshold increases false positives exponentially. CUSUM accumulates '
    'evidence mathematically until confidence is achieved — catching slow drift without false alarm explosion.')

doc.add_heading('5.4 Drift Direction & MITRE ATT&CK Mapping', level=2)

doc.add_paragraph(
    'Knowing THAT an entity is drifting is necessary but not sufficient. The SOC analyst needs: "drifting toward WHAT?"')

visual_flow_table([
    ('1', 'DRIFT VECTOR', 'embedding(t)\nminus\nembedding(t-1)'),
    ('2', 'PROJECT', 'Cosine similarity\nvs 8 threat\nconcept vectors'),
    ('3', 'MAP', 'Top concept →\nMITRE technique\nID + score'),
    ('4', 'ALERT', '"Drifting toward\nlateral_movement\n(0.78, T1021)"'),
])

# Concept table
add_para('The 8 Threat Concepts:', bold=True, size=11)
table = doc.add_table(rows=9, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Concept', 'Description', 'MITRE Techniques']):
    set_cell_bg(table.rows[0].cells[i], '0E6B8A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
concepts = [
    ('lateral_movement', 'Many new hosts, admin shares, cross-subnet', 'T1021, T1570'),
    ('credential_theft', 'LSASS, ntdsutil, Kerberoasting', 'T1003, T1558'),
    ('data_exfiltration', 'Large outbound, archives, cloud uploads', 'T1041, T1048'),
    ('privilege_escalation', 'Group adds, role changes, svc account abuse', 'T1078, T1068'),
    ('c2_beacon', 'Periodic external connections, low entropy', 'T1071, T1573'),
    ('lotl_execution', 'LOLBin spike: wmic, certutil, rundll32', 'T1059, T1218'),
    ('insider_hoarding', 'Bulk file access, staging, downloads', 'T1074, T1005'),
    ('reconnaissance', 'Port scanning, DNS enum, dir listing', 'T1046, T1018'),
]
for i, (concept, desc, mitre) in enumerate(concepts):
    bg = alt_colors[i % 2]
    for j, val in enumerate([concept, desc, mitre]):
        set_cell_bg(table.rows[i+1].cells[j], bg)
        set_cell_text(table.rows[i+1].cells[j], val, size=9.5, font='Consolas' if j==0 else 'Calibri')
doc.add_paragraph()

doc.add_heading('5.5 Entity Fusion', level=2)

doc.add_paragraph('One human = 10+ identities across enterprise systems:')

# Identity fragmentation visual
table = doc.add_table(rows=2, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
systems = ['Active\nDirectory', 'Azure AD\n/ Okta', 'AWS IAM\n/ K8s', 'CrowdStrike\n/ VPN', 'TACACS+\n/ PKI']
ids = ['DOMAIN\\jsmith', 'jsmith@co.com', 'arn:aws:...\n/jsmith', 'agent-abc\njsmith@vpn', 'jsmith\n(TCP/49)']
for i in range(5):
    set_cell_bg(table.rows[0].cells[i], '1B4F72')
    set_cell_text(table.rows[0].cells[i], systems[i], bold=True, size=8.5, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[1].cells[i], 'EBF5FB')
    set_cell_text(table.rows[1].cells[i], ids[i], size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, font='Consolas')
doc.add_paragraph()

comparison_table(
    'WITHOUT FUSION',
    'WITH FUSION',
    [
        ('4 weak signals in 4 separate systems', '1 strong trajectory under 1 entity_uuid'),
        ('Each below alert threshold individually', 'Combined drift clearly crosses CUSUM alarm'),
        ('Cross-domain movement = invisible', 'Cross-domain movement = one trajectory'),
        ('Analyst sees nothing', 'Analyst sees HIGH alert with full context'),
    ]
)

doc.add_heading('5.6 ABAC Trust Loop', level=2)

doc.add_paragraph('Detection is only useful if it drives enforcement:')

table = doc.add_table(rows=5, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Trust State', 'Trigger', 'Enforcement']):
    set_cell_bg(table.rows[0].cells[i], '0D1B2A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
states = [
    ('TRUSTED', 'Normal behavior, no drift', 'Full access per role assignment'),
    ('ELEVATED_WATCH', 'CUSUM rising, LOW/MED alert', 'Increased logging + step-up MFA on sensitive'),
    ('RESTRICTED', 'High-confidence drift detected', 'Access limited to essential functions'),
    ('BLOCKED', 'CRITICAL alert confirmed', 'All access revoked pending investigation'),
]
state_colors = ['E8F8F5', 'FEF9E7', 'FDEDEC', 'F5B7B1']
for i, (state, trigger, enforce) in enumerate(states):
    set_cell_bg(table.rows[i+1].cells[0], state_colors[i])
    set_cell_text(table.rows[i+1].cells[0], state, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[1], 'F7F8FA')
    set_cell_text(table.rows[i+1].cells[1], trigger, size=9.5)
    set_cell_bg(table.rows[i+1].cells[2], 'F7F8FA')
    set_cell_text(table.rows[i+1].cells[2], enforce, size=9.5)
doc.add_paragraph()

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 6. HOW THEY STOP ATTACKS TOGETHER
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('6. How They Stop AI-Enabled Attacks Together', level=1)

doc.add_paragraph(
    'Neither layer alone is sufficient. Together they create defense-in-depth where the attack surface '
    'either never exists OR the attack is detected within hours.')

doc.add_heading('6.1 Stopping Volt Typhoon — Phase by Phase', level=2)

table = doc.add_table(rows=6, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Phase', 'Adversary Action', '22CT Preemptive', 'ACECARD Detects']):
    color = ['0D1B2A', '0D1B2A', '1B4F72', '0E6B8A'][i]
    set_cell_bg(table.rows[0].cells[i], color)
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

vt_response = [
    ('Initial\nAccess', 'Edge CVE exploit\n(Fortinet, Cisco, PAN)', 'IPS+WAF PROVEN to\nblock. Compensating\ncontrols verified.', 'Auth: anomalous mgmt\nsource IP. Drift 0.08.'),
    ('Credential\nAccess', 'ntdsutil → NTDS.dit\nLSASS extraction', 'EDR verified. Script-\ncontrol blocks ntdsutil\non non-admin hosts.', 'Process: lolbin 0→7.\nCUSUM alarm T+47min.\nT1003 alignment 0.72.'),
    ('Lateral\nMovement', 'RDP/SMB/WinRM\nto domain controllers', 'Micro-segmentation\nPROVEN. SMB shadow\nrules detected.', 'Network: dest_ips 3→18.\nCohort z=6.17. HIGH.\nT1021 alignment 0.78.'),
    ('Persistence', 'Web shells, local accts,\nACL modification', 'Config Drift catches\nbaseline deviation\nin <1 hour.', 'Identity: svc_acct spike.\nDrift → T1078.\nABAC: step-up MFA.'),
    ('Pre-\nposition', 'Credential harvest,\nC2 via KV botnet', 'Egress filtering\nPROVEN. C2 IPs\nblocked ALL zones.', 'File: sensitive_access↑.\nCRITICAL alert.\nABAC → BLOCKED.'),
]
for i, (phase, attack, preempt, acecard) in enumerate(vt_response):
    bg = alt_colors[i % 2]
    set_cell_bg(table.rows[i+1].cells[0], 'F7F8FA')
    set_cell_text(table.rows[i+1].cells[0], phase, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[1], bg)
    set_cell_text(table.rows[i+1].cells[1], attack, size=9)
    set_cell_bg(table.rows[i+1].cells[2], 'EBF5FB')
    set_cell_text(table.rows[i+1].cells[2], preempt, size=9)
    set_cell_bg(table.rows[i+1].cells[3], 'E8F4F8')
    set_cell_text(table.rows[i+1].cells[3], acecard, size=9)
doc.add_paragraph()

doc.add_heading('6.2 Stopping Salt Typhoon — Phase by Phase', level=2)

table = doc.add_table(rows=6, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Phase', 'Adversary Action', '22CT Preemptive', 'ACECARD Detects']):
    color = ['0D1B2A', '0D1B2A', '1B4F72', '0E6B8A'][i]
    set_cell_bg(table.rows[0].cells[i], color)
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

st_response = [
    ('Edge CVE\nExploit', 'CVE-2023-20198\nIOS XE level-15', 'Mgmt-plane ACL\nPROVEN. /webui\nblocked to public.', 'Network: anomalous\nmgmt source IP.\nDrift 0.09.'),
    ('Account\nCreation', 'Local priv account.\nAAA config changed.', 'Config Drift detects\nbaseline diff in\n<1 hour.', 'Identity: graph hash\nchanges. New SID\nlinked to router.'),
    ('TACACS+\nCapture', 'On-box PCAP TCP/49.\nCred harvesting.', 'Capture restricted\nto authorized roles.\nMgmt VRF enforced.', 'Process: tcpdump\noutside change window.\nCUSUM alarm.'),
    ('JumbledPath\nPivot', 'ELF chains across\ncompromised routers.', 'Outbound to C2\nblocked ALL\ninternal zones.', 'Composite: same uuid\nacross 4 router hops.\nCo-drift detected.'),
    ('GRE Tunnel\n+ Exfil', 'GRE → CDR/metadata.\nSlow exfil (months).', 'GRE peer auth verified.\nCALEA-zone seg\nPROVEN.', 'data_exfiltration 0.44.\nSlow drift: CUSUM\ncatches in hours.'),
]
for i, (phase, attack, preempt, acecard) in enumerate(st_response):
    bg = alt_colors[i % 2]
    set_cell_bg(table.rows[i+1].cells[0], 'F7F8FA')
    set_cell_text(table.rows[i+1].cells[0], phase, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[1], bg)
    set_cell_text(table.rows[i+1].cells[1], attack, size=9)
    set_cell_bg(table.rows[i+1].cells[2], 'EBF5FB')
    set_cell_text(table.rows[i+1].cells[2], preempt, size=9)
    set_cell_bg(table.rows[i+1].cells[3], 'E8F4F8')
    set_cell_text(table.rows[i+1].cells[3], acecard, size=9)
doc.add_paragraph()

doc.add_heading('6.3 Four Integration Points', level=2)

# Integration visual
visual_flow_table([
    ('1', 'COVERAGE→\nSENSITIVITY', 'Preemptive map\ndrives ACECARD\nthresholds'),
    ('2', 'GAPS→\nTRUST FLOOR', 'Unmitigated gap\nlowers trust for\naffected entities'),
    ('3', 'ALERTS→\nTHREAT FEED', 'Confirmed TPs\nfeed Preemptive\nAttack Intel'),
    ('4', 'SHARED\nVOCABULARY', 'MITRE IDs +\nEntity UUIDs =\nSOAR join key'),
], 'Closed-Loop Integration:')

info_box('SYSTEM, NOT BUNDLE',
    'These integration points make 22CT Preemptive + ACECARD a unified system:\n\n'
    '• What Preemptive prevents today reduces ACECARD noise tomorrow\n'
    '• What ACECARD detects today, Preemptive prevents tomorrow\n'
    '• Shared MITRE vocabulary eliminates translation gaps\n'
    '• Entity UUIDs enable one-click investigation across both layers')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 7. ARCHITECTURE & DEPLOYMENT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Architecture & Deployment', level=1)

doc.add_heading('7.1 Solution Architecture — Five Layers', level=2)

add_para('Full-Stack Architecture (Gabriel Nimbus):', bold=True, size=12)

architecture_layer([
    ('L5: PRESENTATION', 'Plotly Dash / React UI  |  7-chart dashboard  |  Alert feed  |  Kill-chain viz  |  CAC/PIV auth', 'Analyst Interface'),
    ('L4: API + INTEGRATION', 'FastAPI (async)  |  REST + WebSocket  |  Istio mTLS  |  SOAR webhooks  |  ServiceNow  |  RBAC', 'External Integration'),
    ('L3: BEHAVIORAL ENGINE', '5 Signal Summarizers  |  Embedding Pipeline  |  CUSUM  |  Drift Direction  |  Entity Fusion  |  ABAC', 'Core Detection'),
    ('L2: DATA PLANE', 'PostgreSQL + pgvector (HNSW)  |  TimescaleDB  |  Redis (cache + pub/sub)  |  S3  |  AES-256 at rest', 'Storage + Search'),
    ('L1: INFRASTRUCTURE', 'K8s (Big Bang / RKE2)  |  Iron Bank  |  Istio mesh  |  Fluentd  |  Prometheus + Grafana  |  FIPS 140-2/3', 'Platform'),
])

doc.add_heading('7.2 Data Flow — End to End', level=2)

add_para('Complete Pipeline: Raw Telemetry → Actionable Alert', bold=True, size=12)

# Detailed data flow as table
table = doc.add_table(rows=7, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Step', 'Name', 'What Happens', 'Infrastructure', 'Output']):
    set_cell_bg(table.rows[0].cells[i], '0D1B2A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

flow_steps = [
    ('1', 'INGEST', 'Raw events normalized\nto ECS/OCSF/STIX 2.1', 'Kafka / Fluentd', 'Normalized\nevents'),
    ('2', 'WINDOW', '1-hour aggregation\nper entity_uuid.\nEntity fusion resolves IDs.', 'TimescaleDB', 'Hourly\nmetrics'),
    ('3', 'SUMMARIZE', '5 signal summarizers\nproduce structured text.\nNo raw content / PII.', 'Python workers', 'Text\ndescriptions'),
    ('4', 'EMBED', 'Text → 1536-d vector.\nLocal model or\nAzure OpenAI GovCloud.', 'pgvector HNSW', '6 vectors\nper entity'),
    ('5', 'ANALYZE', 'Drift + velocity + accel.\nCUSUM 4σ alarm.\nDirection → 8 concepts.', 'Per entity/\nwindow', 'Drift scores\n+ concepts'),
    ('6', 'ALERT +\nENFORCE', 'CRITICAL→INFO tiers.\nMITRE IDs, cohort z.\nABAC trust update.', 'WebSocket\n+ SOAR', 'Actionable\nalerts'),
]
for i, (step, name, what, infra, output) in enumerate(flow_steps):
    bg = alt_colors[i % 2]
    set_cell_bg(table.rows[i+1].cells[0], '1B4F72')
    set_cell_text(table.rows[i+1].cells[0], step, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[1], '0E6B8A')
    set_cell_text(table.rows[i+1].cells[1], name, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[2], bg)
    set_cell_text(table.rows[i+1].cells[2], what, size=9)
    set_cell_bg(table.rows[i+1].cells[3], bg)
    set_cell_text(table.rows[i+1].cells[3], infra, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[4], 'E8F8F5')
    set_cell_text(table.rows[i+1].cells[4], output, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# Performance box
callout_box('PERFORMANCE',
    '• <3 sec per entity (API-backed embedding model)\n'
    '• 10,000 entities/hour per 4-vCPU pod (API-backed)\n'
    '• ~2,000 entities/hour per pod (local CPU model)\n'
    '• pgvector HNSW: <10ms nearest-neighbor search across millions of vectors\n'
    '• Horizontal autoscale: add pods to match entity count', '0E6B8A')

doc.add_heading('7.3 Gabriel Nimbus Deployment Model', level=2)

table = doc.add_table(rows=4, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Classification', 'Embedding Model', 'Key Constraints']):
    set_cell_bg(table.rows[0].cells[i], '0D1B2A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
deploy = [
    ('IL5 (CUI)', 'Azure OpenAI GovCloud\ntext-embedding-3-small', 'FedRAMP High. mTLS. FIPS crypto.\nCan use cloud API within boundary.'),
    ('IL6 (Secret)', 'Local: E5-Mistral-7B or\nnomic-embed-text-v1.5', 'Fully air-gapped. Zero external calls.\nLocal GPU (A100/H100) for embedding.'),
    ('JWICS (TS/SCI)', 'Local: E5-Mistral-7B or\nnomic-embed-text-v1.5', 'Fully air-gapped + isolated network.\nManual STIX import at boundary.'),
]
for i, (cls, model, constraints) in enumerate(deploy):
    bg = alt_colors[i % 2]
    set_cell_bg(table.rows[i+1].cells[0], '1B4F72')
    set_cell_text(table.rows[i+1].cells[0], cls, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[1], bg)
    set_cell_text(table.rows[i+1].cells[1], model, size=9.5, font='Consolas')
    set_cell_bg(table.rows[i+1].cells[2], bg)
    set_cell_text(table.rows[i+1].cells[2], constraints, size=9.5)
doc.add_paragraph()

info_box('AIR-GAPPED OPERATIONS',
    'Single Helm chart for all environments. Same containers, same code, different values file.\n'
    'No API calls to external services. No telemetry. No cloud dependencies.\n'
    'Threat intel updates: STIX 2.1 bundles imported manually at classification boundary.\n'
    'Container updates: Iron Bank pipeline → security scan → transport across air gap.')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 8. COMPLIANCE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('8. DoD Compliance & Accreditation', level=1)

table = doc.add_table(rows=7, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Framework', 'Requirement', 'How We Meet It']):
    set_cell_bg(table.rows[0].cells[i], '0D1B2A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
compliance = [
    ('NIST 800-53 Rev 5', 'SI-4, AU-6, IR-4, SC-7, IA-5, AC-2', 'Continuous monitoring, automated audit analysis, SOAR response, formal boundary verification, auth analytics, identity fusion'),
    ('DoD ZT RA', 'Pillar 4: Data\nPillar 7: Visibility & Analytics', 'Behavioral data access analysis + DLP integration. Comprehensive entity visibility across all systems.'),
    ('RMF / eMASS', 'ATO documentation + continuous monitoring', 'OSCAL-format SSP, auto-generated control statements, continuous monitoring feeds, POA&M auto-gen'),
    ('cATO', 'Continuous authorization (not periodic 3-yr)', 'Continuous control monitoring, real-time risk scoring, automated evidence collection, direct eMASS feeds'),
    ('DISA STIGs', 'Container hardening standards', 'All containers from Iron Bank (pre-scanned). Automated STIG compliance in CI/CD pipeline.'),
    ('FIPS 140-2/3', 'Validated cryptographic modules', 'FIPS-validated TLS libraries. HSM for key storage. Encryption at rest (AES-256-GCM).'),
]
for i, (fw, req, how) in enumerate(compliance):
    bg = alt_colors[i % 2]
    set_cell_bg(table.rows[i+1].cells[0], '1B4F72')
    set_cell_text(table.rows[i+1].cells[0], fw, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[1], bg)
    set_cell_text(table.rows[i+1].cells[1], req, size=9.5)
    set_cell_bg(table.rows[i+1].cells[2], bg)
    set_cell_text(table.rows[i+1].cells[2], how, size=9.5)
doc.add_paragraph()

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 9. TIMELINE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Timeline: 90 Days to First Detection', level=1)

# Timeline visual
table = doc.add_table(rows=5, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Phase', 'Timing', 'Activities', 'Deliverable']):
    set_cell_bg(table.rows[0].cells[i], '0D1B2A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

timeline = [
    ('FOUNDATION', 'Weeks 1–4', '• Deploy K8s containers (Iron Bank)\n• Connect SIEM/EDR/NetFlow connectors\n• Configure entity fusion rules\n• Compute initial baselines (2-4 wks)', 'Data flowing,\nentities resolved,\nbaselines computing'),
    ('DETECTION', 'Weeks 5–8', '• Enable 5-signal embedding pipeline\n• Activate CUSUM + drift direction\n• Configure cohort definitions\n• Dashboard live', 'Alerts firing,\ndrift direction\noperational'),
    ('OPERATIONS', 'Weeks 9–12', '• SOC integration (SOAR, ServiceNow)\n• Threshold tuning from analyst feedback\n• Enable ABAC trust loop\n• Volt/Salt scenario replay', 'Full UEBA,\nABAC enforcing,\nvalidated'),
    ('PREEMPTIVE', 'Week 13+', '• Connect NGFW/IPS/SASE\n• Initial formal verification scan\n• Coverage→sensitivity coupling\n• Detection-to-prevention loop', 'Both layers\noperational,\nclosed-loop'),
]
phase_colors = ['0E6B8A', '1B4F72', '0D1B2A', 'B7950B']
for i, (phase, timing, activities, deliverable) in enumerate(timeline):
    set_cell_bg(table.rows[i+1].cells[0], phase_colors[i])
    set_cell_text(table.rows[i+1].cells[0], phase, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[1], 'F7F8FA')
    set_cell_text(table.rows[i+1].cells[1], timing, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[2], 'F7F8FA')
    set_cell_text(table.rows[i+1].cells[2], activities, size=9.5)
    set_cell_bg(table.rows[i+1].cells[3], 'E8F8F5')
    set_cell_text(table.rows[i+1].cells[3], deliverable, size=9.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

callout_box('RISK MITIGATION',
    '• ATO timeline → mitigated by cATO pathway + existing reciprocity\n'
    '• Data quality from legacy SIEM → mitigated by 2-week connector validation sprint\n'
    '• Baseline period → minimum 2 weeks stable data; longer = better seasonal patterns\n'
    '• Red team validation → planned within 60 days of deployment', 'C0392B')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 10. GLOSSARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Key Technical Terms Glossary', level=1)

terms = [
    ('ABAC', 'Attribute-Based Access Control. Access decisions based on attributes of user, resource, action, environment.'),
    ('ACECARD', 'Adaptive Cyber Entity Characterization, Assessment & Risk Detection. Our UEBA system.'),
    ('Air-Gapped', 'Network isolation with zero connectivity to external networks. Required for IL6/JWICS.'),
    ('Big Bang', 'Platform One\'s pre-configured K8s distribution with built-in security (Istio, Kiali, Jaeger).'),
    ('C2', 'Command and Control. Communication channel between adversary infrastructure and compromised hosts.'),
    ('CAC/PIV', 'Common Access Card / Personal Identity Verification. DoD smart card authentication.'),
    ('CALEA', 'Communications Assistance for Law Enforcement Act. Requires telecoms to enable lawful intercept.'),
    ('cATO', 'Continuous Authority to Operate. Modern DoD accreditation replacing periodic 3-year ATO cycles.'),
    ('CDR', 'Call Detail Records. Telecom metadata: who called whom, when, duration, cell tower.'),
    ('Cohort Z-Score', 'Standard deviations an entity\'s drift deviates from peer group. >3.0 = significant outlier.'),
    ('Cosine Similarity', 'Measure of angular distance between vectors. Range [-1,1]. Used to compare embeddings.'),
    ('CUSUM', 'Cumulative Sum. Sequential detection algorithm (Page 1954) that accumulates small deviations.'),
    ('DCO-IDM', 'Defensive Cyberspace Operations - Internal Defensive Measures. DoD network defense.'),
    ('DISA STIGs', 'Defense Information Systems Agency Security Technical Implementation Guides.'),
    ('Drift Direction', 'Angular direction of behavioral change projected onto threat concept vectors.'),
    ('Drift Vector', 'embedding(t) - embedding(t-1). Direction and magnitude of behavioral change.'),
    ('ECS', 'Elastic Common Schema. Standard field naming for security events.'),
    ('Embedding', 'Vector representation of data in high-dimensional space. Similar things = close vectors.'),
    ('Entity Fusion', 'Resolving multiple identifiers across systems into single canonical entity_uuid.'),
    ('FIPS 140-2/3', 'Federal Information Processing Standard for cryptographic module validation.'),
    ('Gabriel Nimbus', 'ARCYBER\'s classified big data analytics platform on Kubernetes.'),
    ('GRE', 'Generic Routing Encapsulation. Tunneling protocol. Used by Salt Typhoon for data exfiltration.'),
    ('HNSW', 'Hierarchical Navigable Small World. Fast approximate nearest-neighbor algorithm for vectors.'),
    ('i.i.d.', 'Independent and Identically Distributed. Statistical assumption for CUSUM optimality proof.'),
    ('IL5/IL6', 'Impact Level 5 (CUI) / Impact Level 6 (Secret). DoD cloud classification levels.'),
    ('Iron Bank', 'DoD hardened container registry from Platform One. Pre-scanned and approved for DoD.'),
    ('Istio', 'Service mesh providing mTLS, traffic management, and observability for Kubernetes.'),
    ('JWICS', 'Joint Worldwide Intelligence Communications System. TS/SCI classified network.'),
    ('KV Botnet', 'Compromised SOHO routers used by Volt Typhoon as C2 relay infrastructure.'),
    ('LOLBin', 'Living-off-the-Land Binary. Legitimate system tool abused for malicious purposes.'),
    ('LOTL', 'Living-off-the-Land. Using only built-in tools — no custom malware.'),
    ('MITRE ATT&CK', 'Knowledge base of adversary TTPs. Standard taxonomy for cyber threats.'),
    ('mTLS', 'Mutual TLS. Both client and server authenticate via certificates.'),
    ('NIST 800-53', 'Security and Privacy Controls for Information Systems. DoD control standard.'),
    ('NTDS.dit', 'Active Directory database containing all domain user password hashes.'),
    ('OCSF', 'Open Cybersecurity Schema Framework. Vendor-neutral schema for security events.'),
    ('OSCAL', 'Open Security Controls Assessment Language. Machine-readable security documentation.'),
    ('pgvector', 'PostgreSQL extension for vector similarity search. Supports HNSW indexes.'),
    ('RKE2', 'Rancher Kubernetes Engine 2. FIPS-compliant K8s distribution for DoD.'),
    ('RMF', 'Risk Management Framework. DoD/NIST process for authorizing information systems.'),
    ('SAT/SMT', 'Satisfiability / Satisfiability Modulo Theories. Solvers used in formal verification.'),
    ('SOAR', 'Security Orchestration, Automation and Response. Automated security workflows.'),
    ('STIX 2.1', 'Structured Threat Information Expression. Standard for sharing threat intelligence.'),
    ('TACACS+', 'Terminal Access Controller Access-Control System Plus. AAA for network devices (TCP/49).'),
    ('Trajectory', 'Path an entity\'s behavioral embedding traces through vector space over time.'),
    ('UEBA', 'User and Entity Behavior Analytics. Builds baselines and detects deviations.'),
    ('Zero Trust', 'Never trust, always verify. Every access authenticated regardless of network location.'),
]

# Glossary as a two-column table for compactness
half = len(terms) // 2 + len(terms) % 2
table = doc.add_table(rows=half+1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_bg(table.rows[0].cells[0], '0D1B2A')
set_cell_text(table.rows[0].cells[0], 'Term', bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
set_cell_bg(table.rows[0].cells[1], '0D1B2A')
set_cell_text(table.rows[0].cells[1], 'Definition', bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
set_cell_bg(table.rows[0].cells[2], '0D1B2A')
set_cell_text(table.rows[0].cells[2], 'Term', bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
set_cell_bg(table.rows[0].cells[3], '0D1B2A')
set_cell_text(table.rows[0].cells[3], 'Definition', bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))

for i in range(half):
    bg = alt_colors[i % 2]
    # Left column pair
    if i < len(terms):
        set_cell_bg(table.rows[i+1].cells[0], bg)
        set_cell_text(table.rows[i+1].cells[0], terms[i][0], bold=True, size=8.5)
        set_cell_bg(table.rows[i+1].cells[1], bg)
        set_cell_text(table.rows[i+1].cells[1], terms[i][1], size=8.5)
    # Right column pair
    j = i + half
    if j < len(terms):
        set_cell_bg(table.rows[i+1].cells[2], bg)
        set_cell_text(table.rows[i+1].cells[2], terms[j][0], bold=True, size=8.5)
        set_cell_bg(table.rows[i+1].cells[3], bg)
        set_cell_text(table.rows[i+1].cells[3], terms[j][1], size=8.5)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 11. CONCEPTS EXPLAINED SIMPLY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('11. Key Concepts Explained Simply', level=1)

doc.add_paragraph(
    'This section explains core technical concepts in plain language for quick reference.')

# --- LOTL ---
doc.add_heading('11.1 Living-off-the-Land (LOTL)', level=2)

doc.add_paragraph(
    'Living-off-the-Land means the attacker brings NO malware. They use tools that Windows already has installed — '
    'tools Microsoft put there for system administrators.')

comparison_table(
    'Traditional Hacker',
    'Living-off-the-Land (Volt Typhoon)',
    [
        ('Brings their own weapon (malware, virus)', 'Brings NOTHING. Uses built-in Windows tools.'),
        ('Antivirus sees unknown file → BLOCKED', 'Antivirus sees signed Microsoft binary → "Fine"'),
        ('Leaves forensic traces (custom files)', 'Leaves NO foreign artifacts on disk'),
        ('Detectable by signatures', 'UNDETECTABLE by any signature-based tool'),
    ]
)

add_para('The tools they use:', bold=True)
table = doc.add_table(rows=7, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Tool', 'Legitimate Purpose', 'Attacker Uses It For']):
    set_cell_bg(table.rows[0].cells[i], '0D1B2A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
lotl_tools = [
    ('ntdsutil.exe', 'Backup Active Directory database', 'Steal ALL domain passwords'),
    ('wmic.exe', 'Query system info remotely', 'Run commands on other machines'),
    ('netsh.exe', 'Configure network settings', 'Open firewall holes, redirect traffic'),
    ('certutil.exe', 'Manage certificates', 'Download malware from the internet'),
    ('powershell.exe', 'Automation and scripting', 'Run attack scripts'),
    ('rundll32.exe', 'Run DLL files', 'Execute malicious code hidden in DLLs'),
]
for i, (tool, legit, attack) in enumerate(lotl_tools):
    bg = alt_colors[i % 2]
    set_cell_bg(table.rows[i+1].cells[0], bg)
    set_cell_text(table.rows[i+1].cells[0], tool, bold=True, size=9.5, font='Consolas')
    set_cell_bg(table.rows[i+1].cells[1], bg)
    set_cell_text(table.rows[i+1].cells[1], legit, size=9.5)
    set_cell_bg(table.rows[i+1].cells[2], 'FDEDEC')
    set_cell_text(table.rows[i+1].cells[2], attack, size=9.5)
doc.add_paragraph()

key_insight_box(
    'The TOOLS are legitimate. The PERSON running them is not an admin. That\'s why ACECARD tracks lolbin_count '
    'per user. John the accountant running ntdsutil = alarm. John the sysadmin running ntdsutil = normal. '
    'It\'s not about WHAT ran. It\'s about WHO ran it.')

# --- Five Signals Simple ---
doc.add_heading('11.2 The Five Signals — What Each One Asks', level=2)

doc.add_paragraph(
    'All five signals are tied to the same person via their login session. Together they form a complete picture.')

table = doc.add_table(rows=6, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Signal', 'Question It Answers', 'What Changes During Attack']):
    set_cell_bg(table.rows[0].cells[i], '0E6B8A')
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
sig_simple = [
    ('AUTH', 'WHO logged in? From where? When? How often?', 'Login at 2am from new IP. Impossible travel.'),
    ('PROCESS', 'WHAT did they RUN on their machines?', 'Admin tools appear (ntdsutil, wmic, netsh).'),
    ('NETWORK', 'WHERE did they CONNECT TO? How much data?', 'Connections to 18 new servers (was 3).'),
    ('FILE', 'WHAT did they TOUCH? Copy? Delete? Archive?', 'NTDS.dit copied. Sensitive files archived.'),
    ('IDENTITY', 'WHAT POWER did they GAIN? New privileges?', 'Added to Domain Admins. Service acct used.'),
]
for i, (sig, q, attack) in enumerate(sig_simple):
    bg = alt_colors[i % 2]
    set_cell_bg(table.rows[i+1].cells[0], '0E6B8A')
    set_cell_text(table.rows[i+1].cells[0], sig, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(table.rows[i+1].cells[1], bg)
    set_cell_text(table.rows[i+1].cells[1], q, size=10)
    set_cell_bg(table.rows[i+1].cells[2], 'FDEDEC')
    set_cell_text(table.rows[i+1].cells[2], attack, size=10)
doc.add_paragraph()

info_box('HOW SIGNALS CONNECT TO A PERSON',
    'Step 1: John logs into a workstation (AUTH records: "John is on WKST-042")\n'
    'Step 2: Programs run on that workstation during John\'s session → attributed to JOHN (PROCESS)\n'
    'Step 3: Network connections from that workstation → attributed to JOHN (NETWORK)\n\n'
    'The connection is always through LOGIN. No login linking person to machine = processes '
    'only show under the DEVICE entity, not a user entity.')

# --- CUSUM Simple ---
doc.add_heading('11.3 CUSUM — The Bucket Analogy', level=2)

doc.add_paragraph(
    'Imagine a bucket under a dripping faucet.')

callout_box('THE BUCKET',
    'Normal day: A little water drips in (normal behavioral noise). The bucket has a small drain hole — '
    'so it drains out. Bucket never fills up.\n\n'
    'Attack happening: The drip gets slightly faster. Not a flood — just a LITTLE more than the hole can drain. '
    'Each hour, the bucket fills a tiny bit more. After 8 hours — bucket overflows. ALARM.\n\n'
    '• Bucket = the CUSUM statistic (starts empty)\n'
    '• Drip = how much behavior drifted this hour\n'
    '• Drain hole = allowance for normal variation\n'
    '• Overflow line = alarm threshold', '0E6B8A')

comparison_table(
    'Current Tools (No Bucket)',
    'CUSUM (Has a Bucket)',
    [
        ('Looks at each hour independently', 'Keeps a running total across hours'),
        ('Asks: "Is THIS hour bad?"', 'Asks: "Is the SUM of all hours bad?"'),
        ('"Is this drip a flood? No. No. No..."', '"The bucket is slowly filling... filling... ALARM"'),
        ('Catches: big sudden jumps only', 'Catches: slow steady creep'),
        ('Volt Typhoon: NEVER fires', 'Volt Typhoon: fires in 8 HOURS'),
    ]
)

add_para('The Fundamental Difference:', bold=True, size=12, space_before=8)
table = doc.add_table(rows=2, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_bg(table.rows[0].cells[0], 'C0392B')
set_cell_text(table.rows[0].cells[0], '  Fixed Threshold = Alarm on POSITION', bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
set_cell_bg(table.rows[0].cells[1], '0E6B8A')
set_cell_text(table.rows[0].cells[1], '  CUSUM = Alarm on INTENT', bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
set_cell_bg(table.rows[1].cells[0], 'FDEDEC')
set_cell_text(table.rows[1].cells[0], '"How far are you from normal RIGHT NOW?"\n\nAttacker can control position (stay close).\nResult: invisible forever.', size=10)
set_cell_bg(table.rows[1].cells[1], 'E8F4F8')
set_cell_text(table.rows[1].cells[1], '"Are you consistently moving away, EVERY hour?"\n\nAttacker CANNOT hide intent (must keep moving).\nResult: caught in hours.', size=10)
doc.add_paragraph()

# --- Continuous Drift ---
doc.add_heading('11.4 Continuous Drift in One Direction = Problem', level=2)

doc.add_paragraph(
    'Normal users drift randomly — some days left, some days right. It cancels out. '
    'An attacker MUST keep moving in one direction toward their objective.')

comparison_table(
    'Normal User',
    'Attacker (Volt Typhoon)',
    [
        ('Drift: → ← → → ← ← → ←', 'Drift: → → → → → → → →'),
        ('Random direction each day', 'Same direction every hour'),
        ('Cancels out over time', 'Accumulates over time'),
        ('CUSUM stays near zero', 'CUSUM grows → ALARM'),
        ('No consistent pattern', 'Consistent pattern = INTENT'),
    ]
)

key_insight_box(
    'To steal credentials: MUST run credential tools → Process signal drifts toward credential_theft.\n'
    'To move laterally: MUST connect to new hosts → Network signal drifts toward lateral_movement.\n'
    'To exfiltrate data: MUST send data outbound → File + Network drift toward exfiltration.\n\n'
    'They can go SLOW (avoid thresholds). They CANNOT avoid MOVING.\n'
    'CUSUM catches the movement. Drift direction names the objective.')

# --- What the Metrics Mean ---
doc.add_heading('11.5 What the Alert Numbers Mean', level=2)

add_para('unique_dest_ips: 3 → 18', bold=True, size=12)
doc.add_paragraph(
    'How many different machines this person CONNECTED TO (outbound). Not where they logged in from — '
    'where they REACHED OUT TO.')
add_bullet('Normal: John talks to same 3 servers daily (file server, email, SharePoint)')
add_bullet('Attack: John connects to 18 servers — domain controllers, HR, finance. He\'s LOOKING for something.')
add_bullet('Cohort z=6.17: other admins average 4 destinations. John at 18 = extreme outlier among his peers.')

add_para('lolbin_count: 0 → 7', bold=True, size=12, space_before=8)
doc.add_paragraph(
    'How many Living-off-the-Land binaries this person ran (legit admin tools used for attack).')
add_bullet('Normal: John the accountant runs 0 admin tools. Ever.')
add_bullet('Attack: 7 admin tools in one hour (ntdsutil, wmic, netsh...). Same tools an admin would use.')
add_bullet('The tools are legitimate. But John is NOT an admin. WHO ran it is what matters.')

add_para('service_account_usage: 0 → 4', bold=True, size=12, space_before=8)
doc.add_paragraph(
    'Service accounts (svc_backup, svc_sql) are for APPLICATIONS, not people. They often have high privileges '
    'and weak security (no MFA, passwords don\'t expire).')
add_bullet('Normal: John uses his own account. Never touches service accounts.')
add_bullet('Attack: John logs in as svc_backup 4 times. Getting high privilege through a back door.')
add_bullet('Red flag: human using non-human account = trying to get power without normal security checks.')

info_box('WHY ALL THREE TOGETHER MATTER',
    'Each number alone = slightly weird. All three moving in the same direction = attack trajectory.\n\n'
    'unique_dest = reaching into new places (lateral movement)\n'
    'lolbin_count = using attack tools (credential theft)\n'
    'service_account = grabbing power (persistence)\n\n'
    'Three signals, one direction, one entity. CUSUM catches the consistency. Drift direction names it.')

# --- Neither Layer Alone ---
doc.add_heading('11.6 Why Both Layers Are Required', level=2)

doc.add_paragraph(
    'Neither layer alone is sufficient:')

table = doc.add_table(rows=4, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_bg(table.rows[0].cells[0], '0D1B2A')
set_cell_text(table.rows[0].cells[0], '', size=9)
set_cell_bg(table.rows[0].cells[1], '1B4F72')
set_cell_text(table.rows[0].cells[1], 'PREEMPTIVE ALONE', bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(table.rows[0].cells[2], '0E6B8A')
set_cell_text(table.rows[0].cells[2], 'ACECARD ALONE', bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

set_cell_bg(table.rows[1].cells[0], 'E8F8F5')
set_cell_text(table.rows[1].cells[0], 'CAN DO', bold=True, size=10)
set_cell_bg(table.rows[1].cells[1], 'E8F8F5')
set_cell_text(table.rows[1].cells[1], '• Eliminate config gaps (proven)\n• Block known attack paths\n• Zero FP/FN within model', size=9.5)
set_cell_bg(table.rows[1].cells[2], 'E8F8F5')
set_cell_text(table.rows[1].cells[2], '• Detect behavioral evolution\n• Catch LOTL via drift\n• Work without signatures', size=9.5)

set_cell_bg(table.rows[2].cells[0], 'FDEDEC')
set_cell_text(table.rows[2].cells[0], 'CANNOT DO', bold=True, size=10)
set_cell_bg(table.rows[2].cells[1], 'FDEDEC')
set_cell_text(table.rows[2].cells[1], '• Detect valid-credential abuse\n• See LOTL (tools are legit)\n• Track behavioral evolution', size=9.5)
set_cell_bg(table.rows[2].cells[2], 'FDEDEC')
set_cell_text(table.rows[2].cells[2], '• Only detects AFTER entry\n• Cannot prevent config gaps\n• Reactive, not preemptive', size=9.5)

set_cell_bg(table.rows[3].cells[0], 'FEF9E7')
set_cell_text(table.rows[3].cells[0], 'TOGETHER', bold=True, size=10)
set_cell_bg(table.rows[3].cells[1], 'FEF9E7')
set_cell_text(table.rows[3].cells[1], 'Attack surface NEVER EXISTS', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(table.rows[3].cells[2], 'FEF9E7')
set_cell_text(table.rows[3].cells[2], '...or detected in HOURS, not years', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

callout_box('DEFENSE IN DEPTH',
    'Preemptive eliminates the attack surface that exists in configuration — adversary cannot enter through config gaps.\n\n'
    'ACECARD detects behavioral evolution — catches adversaries who bypass via valid credentials and legitimate tools.\n\n'
    'Together: the attack surface either NEVER EXISTS or the attack is detected and contained within hours.\n\n'
    'Preemptive closes the door.  ACECARD watches the room.  Nothing gets through unnoticed.', 'B7950B')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# APPENDIX: 22CT Qualifications
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Appendix: 22CT Qualifications', level=1)

table = doc.add_table(rows=7, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_bg(table.rows[0].cells[0], '0D1B2A')
set_cell_text(table.rows[0].cells[0], 'Qualification', bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
set_cell_bg(table.rows[0].cells[1], '0D1B2A')
set_cell_text(table.rows[0].cells[1], 'Detail', bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))

quals = [
    ('$90M Army SOC/MDR Contract', '800+ cleared analysts. 24x7x365 security operations. Proven incident response at scale.'),
    ('USAF $108M', 'Cybersecurity services for Air Force networks and systems.'),
    ('FBI TSC $56M', 'Terrorist Screening Center technology services.'),
    ('NAVAIR $145M', 'Naval Air Systems Command IT and cybersecurity.'),
    ('Cleared Workforce', 'Secret to TS/SCI. MITRE ATT&CK certified. NIST/RMF trained. Day-1 ready.'),
    ('Innovation Partner: Rigor AI', 'Preemptive layer formal verification technology. 22CT provides DoD integration, compliance, workforce.'),
]
for i, (qual, detail) in enumerate(quals):
    bg = alt_colors[i % 2]
    set_cell_bg(table.rows[i+1].cells[0], 'B7950B')
    set_cell_text(table.rows[i+1].cells[0], qual, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(table.rows[i+1].cells[1], bg)
    set_cell_text(table.rows[i+1].cells[1], detail, size=10)


# ─── SAVE ───
doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Pages: ~25-30 (estimated)")
