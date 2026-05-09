"""Generate all documents for ACECARD CE May 6, 2026 event.
UPDATED: Aligned with non-disclosure approach. No CUSUM, no embedding dimensions,
no formal verification specifics in any external-facing content.
Safe language: "Formal Mathematical Model" and "Digital Twin with time dimension"."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = os.path.dirname(__file__)


def _style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)
        hs.font.name = "Calibri"


def _bold_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def _bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(" -- " + text)
    else:
        p.add_run(text)


# ═══════════════════════════════════════════════════════════════
# DOCUMENT 1: SOLUTION BRIEF (1-2 pages)
# ═══════════════════════════════════════════════════════════════
def build_solution_brief():
    doc = Document()
    _style(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Preemptive Cyber Defense + Behavioral Intelligence")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Solution Brief for CFIC Collaboration Event\n"
        "Containerized, Cloud-Native Prototype for Gabriel Nimbus"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("22nd Century Technologies (TSCTI)")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── PROBLEM STATEMENT ──
    doc.add_heading("Problem Statement", level=1)
    doc.add_paragraph(
        "Two confirmed nation-state cyber campaigns -- Volt Typhoon (5+ year dwell "
        "in U.S. critical infrastructure) and Salt Typhoon (4+ years inside 200+ "
        "telecoms including 9 U.S. carriers) -- have demonstrated that current "
        "security tools are structurally incapable of detecting advanced persistent "
        "threats that use Living-off-the-Land tradecraft and configuration-resident "
        "persistence."
    )
    doc.add_paragraph(
        "No existing SIEM, EDR, XDR, UEBA, pentest, or vulnerability scanner can: "
        "(1) mathematically prove all config paths block an attack, "
        "(2) detect sub-threshold behavioral drift over weeks/months, "
        "(3) map behavioral anomaly to specific MITRE ATT&CK techniques, or "
        "(4) fuse identity across 10+ systems into one behavioral trajectory."
    )

    # ── SOLUTION ARCHITECTURE ──
    doc.add_heading("Solution: Two Complementary Layers", level=1)

    doc.add_heading("Layer 1: Preemptive Verification", level=2)
    doc.add_paragraph(
        "We define a Formal Mathematical Model that reasons exhaustively over "
        "your entire defensive configuration. This is not sampling, not scanning, "
        "not penetration testing -- it is mathematical proof that attack paths are "
        "blocked or identifies exactly where they are not."
    )
    _bullet(doc, "Ingests all threat intelligence continuously (CVE, CISA KEV, MITRE ATT&CK)")
    _bullet(doc, "Models every defensive control formally (NGFW, IPS, WAF, IdP, SASE)")
    _bullet(doc, "Proves whether attack paths exist or are blocked -- zero false positives, zero false negatives")
    _bullet(doc, "Prescribes vendor-specific remediations (Palo Alto, Fortinet, Cisco, Check Point)")
    _bullet(doc, "Re-verifies continuously: hourly + on every configuration change")
    _bullet(doc, "80% of attacks eliminated at network/firewall layer before reaching endpoints")

    doc.add_heading("Layer 2: Behavioral Intelligence (ACECARD UEBA)", level=2)
    doc.add_paragraph(
        "We create a Digital Twin of each entity instance. User interaction is a "
        "living entity with a time dimension -- capturing behavioral evolution "
        "continuously. This detects the 20% of threats that pass through even "
        "perfect configuration: stolen credentials, LOTL tradecraft, zero-days."
    )
    _bullet(doc, "Multi-signal behavioral profiling per entity (user, device, network segment)")
    _bullet(doc, "Advanced change-point detection catches sub-threshold drift over weeks/months")
    _bullet(doc, "Drift direction mapped to MITRE ATT&CK -- not just 'anomalous' but WHAT threat type")
    _bullet(doc, "Cross-system entity fusion: 10+ identity sources stitched into one trajectory")
    _bullet(doc, "Automated trust scoring and adaptive access control (ABAC enforcement)")
    _bullet(doc, "Detects: LOTL, credential abuse, insider threat, slow lateral movement")

    doc.add_heading("Why Both Layers Are Required", level=2)
    doc.add_paragraph(
        "Preemptive eliminates the configuration attack surface (80% of threats). "
        "Behavioral catches what configuration alone cannot stop: stolen valid "
        "credentials, LOTL tradecraft, zero-day post-exploitation. Neither alone "
        "is sufficient. Together: 95-98% coverage. Dwell time reduced from years to hours."
    )

    # ── THREE INTELLIGENCES ──
    doc.add_heading("Three Intelligences (Continuous Closed Loop)", level=1)
    _bullet(doc, "Continuously analyzes all threat feeds, advisories, CVEs, and TTPs. Builds attack graphs automatically.", "Attack Intelligence")
    _bullet(doc, "Formal Mathematical Model identifies every gap in your defensive controls. Exhaustive, not sampled.", "Defense Intelligence")
    _bullet(doc, "Risk-prioritized, vendor-specific remediation playbooks. Zero false positives. Auto-push to ticketing.", "Remediation Intelligence")

    # ── PROOF ──
    doc.add_heading("Proof: Production Findings", level=1)
    _bullet(doc, "Financial services (Asia): closed lateral movement path hidden by shadowed firewall rules")
    _bullet(doc, "Communications (NA): eliminated RCE exposure from missing IPS inspection")
    _bullet(doc, "Critical infrastructure (SE Asia): detected application obfuscation via DMZ misconfiguration")
    _bullet(doc, "EU sovereign government: national-scale attack surface management across 12 agencies")
    _bullet(doc, "20+ design partners across F500, federal, critical infrastructure, MSSPs, and nation states")

    # ── DEPLOYMENT ──
    doc.add_heading("Deployment on Gabriel Nimbus", level=1)
    _bullet(doc, "Containerized: Iron Bank certified, OCI-compliant, Kubernetes-native Helm charts")
    _bullet(doc, "Cloud-native: designed for big data platforms, horizontal pod autoscaling")
    _bullet(doc, "Multi-enclave: IL5 (NIPR) / IL6 (SIPR) / JWICS (air-gapped) -- same codebase, environment config only")
    _bullet(doc, "Zero-trust: Istio service mesh, Vault secrets, CAC/PIV mTLS authentication")
    _bullet(doc, "cATO-aligned: artifacts auto-generated, FluxCD GitOps CI/CD")

    # ── TEAM ──
    doc.add_heading("Team Credentials", level=1)
    doc.add_paragraph(
        "22nd Century Technologies (TSCTI): $90M U.S. Army SOC/MDR contract "
        "(800+ cleared analysts), USAF $108M, FBI TSC $56M, NAVAIR/USMC $145M."
    )
    doc.add_paragraph(
        "Preemptive engine: 20+ design partners. Production findings in "
        "financial services, telecom, and government. "
        "Covers Palo Alto, Fortinet, Cisco, Check Point (N-S and E-W)."
    )

    # ── OFFER ──
    doc.add_heading("Our Offer", level=1)
    doc.add_paragraph(
        "Give us 4 weeks and your data. We will demonstrate verified gap findings "
        "on your configurations (zero false positives) and behavioral detections "
        "on your entity data (with MITRE mapping). No other vendor can make this "
        "offer because no other tool has these capabilities."
    )

    return doc


# ═══════════════════════════════════════════════════════════════
# DOCUMENT 2: BREAKOUT SESSION TALKING POINTS
# ═══════════════════════════════════════════════════════════════
def build_breakout_points():
    doc = Document()
    _style(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Breakout Session Strategy & Talking Points")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CFIC Collaboration Event  |  06 MAY 2026  |  INTERNAL ONLY")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    p = doc.add_paragraph()
    run = p.add_run(
        "REMINDER: Do NOT disclose our proprietary approach. Use safe language:\n"
        "  Preemptive = 'Formal Mathematical Model'\n"
        "  UEBA = 'Digital Twin of the entity instance with time dimension'\n"
        "  We can prove on the whiteboard in person, or demonstrate on their data in 4 weeks."
    )
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    # ── MORNING: LIMITING FACTORS (0925-1020) ──
    doc.add_heading("Morning Session: Current Limitations (0925-1020)", level=1)
    doc.add_paragraph(
        "Goal: Demonstrate deep understanding of WHY current tools fail structurally. "
        "Establish credibility by diagnosing the problem precisely. Do NOT pitch yet -- just diagnose. "
        "Plant seeds so the audience reaches YOUR conclusion independently."
    )

    doc.add_heading("Limitation 1: No Tool Proves Configuration Is Correct", level=2)
    _bullet(doc, "99% of firewall breaches are caused by misconfiguration, not firewall flaws (Gartner)")
    _bullet(doc, "Configuration state space is astronomically large -- pentests and scanners can only sample a vanishing fraction")
    _bullet(doc, "Annual pentests find point-in-time issues; adversaries operate continuously")
    _bullet(doc, "Salt Typhoon lived IN the router configuration itself -- the persistence IS the config")
    _bullet(doc, "No existing tool mathematically proves that all attack paths through your configs are blocked")
    _bullet(doc, "Compensating controls are assumed, never verified across the full policy chain")

    doc.add_heading("Limitation 2: Behavioral Detection Uses Fixed Thresholds", level=2)
    _bullet(doc, "Every commercial UEBA uses static thresholds per data source -- 'alert if >X logins'")
    _bullet(doc, "Volt Typhoon operates below EVERY threshold -- each action is individually 'normal'")
    _bullet(doc, "The threat is in the trajectory over time, not in any single event")
    _bullet(doc, "Current tools say 'anomaly score 87%' -- no hypothesis, no direction, no MITRE mapping")
    _bullet(doc, "Analysts waste 30-45 min per alert reassembling context manually -- most alerts never investigated")
    _bullet(doc, "PRC actors test against published detection rules before campaigns -- signatures decay monthly")

    doc.add_heading("Limitation 3: Identity Is Fragmented Across 10+ Systems", level=2)
    _bullet(doc, "A single user has identities in AD, Azure AD, Okta, AWS IAM, VPN, TACACS+, badge, K8s, email")
    _bullet(doc, "Each system logs independently with different identifiers and schemas")
    _bullet(doc, "Lateral movement across 5 systems reads as 5 separate 'normal' logins -- not 1 attack chain")
    _bullet(doc, "Salt Typhoon used stolen TACACS+ creds on routers + AD creds on servers -- two silos, one attacker")
    _bullet(doc, "No existing product fuses all identity sources into one behavioral trajectory per real-world entity")

    doc.add_heading("Limitation 4: Tools Detect Known-Bad, Not Unknown-Malicious", level=2)
    _bullet(doc, "EDR/XDR: every binary Volt Typhoon uses (wmic, ntdsutil, netsh) is a legitimate admin tool")
    _bullet(doc, "SIEM: correlates known patterns. Novel LOTL combinations have no pre-built rule")
    _bullet(doc, "UEBA: fixed baselines. A 'normal' admin account used by an adversary looks the same")
    _bullet(doc, "Vuln scanners: see what EXISTS, not what is REACHABLE through configuration")
    _bullet(doc, "The gap is structural, not operational. More of the same tools will not solve it")

    # ── MORNING: WAYS TO OVERCOME (1040-1140) ──
    doc.add_heading("Morning Session: Ways to Overcome (1040-1140)", level=1)
    doc.add_paragraph(
        "Goal: Introduce the solution CONCEPTS without revealing proprietary methods. "
        "Speak to WHAT we achieve, not HOW we achieve it technically."
    )

    doc.add_heading("Overcome 1: Mathematical Proof of Defense (Not Sampling)", level=2)
    _bullet(doc, "Build a Formal Mathematical Model of all defensive configurations")
    _bullet(doc, "Reason exhaustively over the complete configuration state space")
    _bullet(doc, "PROVE that attack paths are blocked -- not sample, not scan, PROVE")
    _bullet(doc, "Continuous re-verification: hourly + on every configuration change")
    _bullet(doc, "Result: zero false positives, zero false negatives within model scope")
    _bullet(doc, "If we say there is a gap, there IS a gap. If we say you are covered, you ARE covered.")

    doc.add_heading("Overcome 2: Digital Twin of Every Entity (Not Thresholds)", level=2)
    _bullet(doc, "Create a Digital Twin of each entity instance -- user, device, network segment")
    _bullet(doc, "User interaction is a living entity with a TIME DIMENSION -- captures evolution continuously")
    _bullet(doc, "Detects behavioral drift that accumulates over weeks/months -- below any fixed threshold")
    _bullet(doc, "Maps drift DIRECTION to specific MITRE ATT&CK techniques -- not just 'anomalous'")
    _bullet(doc, "Cross-system entity fusion: one trajectory per real-world entity across all systems")
    _bullet(doc, "Result: catches LOTL, credential abuse, insider threat, slow lateral movement")

    doc.add_heading("Overcome 3: Three Intelligences Working Continuously", level=2)
    _bullet(doc, "Attack Intelligence: continuously ingests threat feeds, builds attack graphs")
    _bullet(doc, "Defense Intelligence: Formal Mathematical Model finds every gap in your controls")
    _bullet(doc, "Remediation Intelligence: vendor-specific fixes, auto-pushed to ticketing, re-verified")
    _bullet(doc, "Closed loop: new threat -> model attack -> verify defense -> prescribe fix -> re-verify")

    doc.add_heading("Overcome 4: Cloud-Native, Containerized, Gabriel Nimbus Ready", level=2)
    _bullet(doc, "Kubernetes-native: Helm charts, Iron Bank containers, horizontal autoscaling")
    _bullet(doc, "Designed for big data platforms -- scales with data volume")
    _bullet(doc, "Multi-enclave: IL5 / IL6 / JWICS -- same codebase, environment config only")
    _bullet(doc, "Zero-trust service mesh, CAC/PIV authentication, cATO-aligned CI/CD")

    # ── AFTERNOON: SOLUTION BRIEF (1305-1405) ──
    doc.add_heading("Afternoon: Problem Statement & Solution Brief (1305-1405)", level=1)
    doc.add_paragraph(
        "Goal: Develop the formal problem statement for the presentation. "
        "Use the deck (22CT_Preemptive_ACECARD_Volt_Salt_Typhoon_v5.pptx)."
    )

    doc.add_heading("Problem Statement (What to Say)", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "Current SOC tooling cannot detect Living-off-the-Land nation-state operations "
        "because: (1) no tool proves configuration is correct -- they only sample, "
        "(2) behavioral detection uses fixed thresholds that LOTL operates below, "
        "(3) identity is fragmented across 10+ systems -- lateral movement is invisible, and "
        "(4) existing tools detect known-bad patterns, not unknown-malicious behavior. "
        "We need a fundamentally different approach: mathematical proof of defense "
        "combined with a behavioral Digital Twin that captures entity evolution over time."
    )
    run.italic = True

    # ── PRESENTATION SLOT (1420-1535) ──
    doc.add_heading("Afternoon: Presentation (1420-1535)", level=1)
    doc.add_paragraph(
        "Use the deck. Lead with problem, show both layers, "
        "prove with Volt/Salt Typhoon scenarios, close with 4-week proof offer."
    )

    doc.add_heading("Suggested Flow (10-15 minutes)", level=2)
    _bullet(doc, "Minutes 1-2: Problem Statement slide -- frame the two campaigns and structural failures")
    _bullet(doc, "Minutes 3-4: Why Current Tools Fail -- the limiting factors (echo what was discussed in morning)")
    _bullet(doc, "Minutes 5-7: Our Approach -- two layers, three intelligences, 80% diagram")
    _bullet(doc, "Minutes 8-10: How we stop Volt Typhoon and Salt Typhoon -- phase by phase")
    _bullet(doc, "Minutes 11-12: Architecture and deployment (Gabriel Nimbus ready)")
    _bullet(doc, "Minutes 13-14: What no existing tool delivers -- the uniqueness slide")
    _bullet(doc, "Minute 15: Close -- 'Give us 4 weeks and your data. We will prove it.'")

    doc.add_heading("Key Phrases to Use", level=2)
    _bullet(doc, "'Formal Mathematical Model' (not: formal verification, SAT solving, symbolic reasoning)")
    _bullet(doc, "'Digital Twin of each entity with time dimension' (not: embeddings, vectors, CUSUM)")
    _bullet(doc, "'Exhaustive reasoning over the complete configuration state space' (not: 2^8000, state explosion)")
    _bullet(doc, "'Sub-threshold drift detection' (not: change-point detection, cumulative sum)")
    _bullet(doc, "'Drift direction mapped to MITRE' (not: cosine projection, threat concepts)")
    _bullet(doc, "'No existing tool, device, or platform has this capability'")
    _bullet(doc, "'We will prove it on your data in 4 weeks'")

    doc.add_heading("Things NOT to Do", level=2)
    _bullet(doc, "Do NOT reveal the mathematical approach (CUSUM, embeddings, formal verification method)")
    _bullet(doc, "Do NOT name specific dimensions, thresholds, or algorithmic details")
    _bullet(doc, "Do NOT compare to specific competitor products by name in the presentation")
    _bullet(doc, "Do NOT oversell -- let the Volt/Salt Typhoon scenarios speak for themselves")
    _bullet(doc, "Do NOT read slides -- speak from understanding of the problem and solution")
    _bullet(doc, "If pressed for technical details: 'We can demonstrate on your data in 4 weeks, or walk through the math on a whiteboard.'")

    return doc


# ═══════════════════════════════════════════════════════════════
# DOCUMENT 3: ONE-PAGE LEAVE-BEHIND
# ═══════════════════════════════════════════════════════════════
def build_leave_behind():
    doc = Document()
    _style(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Preemptive Cyber Defense + Behavioral Intelligence")
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Stopping Nation-State Adversaries Before They Strike  |  "
        "Containerized for Gabriel Nimbus"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    doc.add_heading("The Problem", level=2)
    p = doc.add_paragraph(
        "Volt Typhoon (5+ years, U.S. critical infrastructure) and Salt Typhoon "
        "(4+ years, 200+ telecoms, 9 U.S. carriers) used Living-off-the-Land tradecraft "
        "to dwell undetected inside networks -- because no existing security tool can "
        "mathematically prove configurations block attacks, or detect sub-threshold "
        "behavioral drift that accumulates over months."
    )
    p.paragraph_format.space_after = Pt(4)

    doc.add_heading("The Solution: Two Layers", level=2)

    table = doc.add_table(rows=2, cols=2)
    table.style = "Light Grid Accent 1"
    table.columns[0].width = Inches(3.4)
    table.columns[1].width = Inches(3.4)

    c = table.rows[0].cells[0]
    c.text = "LAYER 1: Preemptive Verification"
    for r in c.paragraphs[0].runs:
        r.bold = True
    c = table.rows[1].cells[0]
    c.text = (
        "Formal Mathematical Model of ALL security controls. "
        "Exhaustively reasons over the complete config state space. "
        "PROVES attack paths are blocked. "
        "Continuous re-verification hourly + on every change. "
        "Zero false positives. Zero false negatives."
    )

    c = table.rows[0].cells[1]
    c.text = "LAYER 2: Behavioral (ACECARD UEBA)"
    for r in c.paragraphs[0].runs:
        r.bold = True
    c = table.rows[1].cells[1]
    c.text = (
        "Digital Twin of each entity instance. "
        "User interaction is a living entity with time dimension. "
        "Detects sub-threshold behavioral drift over weeks/months. "
        "Drift direction mapped to MITRE ATT&CK. "
        "Cross-system entity fusion (10+ identity sources)."
    )

    doc.add_heading("Coverage", level=2)
    p = doc.add_paragraph(
        "80% of attacks stopped at network/firewall layer (Preemptive). "
        "15-18% caught by behavioral detection (ACECARD). "
        "Combined: 95-98% coverage. Dwell time: hours, not years."
    )
    p.paragraph_format.space_after = Pt(4)

    doc.add_heading("How It Stops Volt Typhoon & Salt Typhoon", level=2)
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = "Light Grid Accent 1"
    table2.columns[0].width = Inches(2.0)
    table2.columns[1].width = Inches(4.8)
    table2.rows[0].cells[0].text = "Attack Method"
    table2.rows[0].cells[1].text = "How We Stop It"
    for r in table2.rows[0].cells[0].paragraphs[0].runs:
        r.bold = True
    for r in table2.rows[0].cells[1].paragraphs[0].runs:
        r.bold = True
    table2.rows[1].cells[0].text = "Config exploitation"
    table2.rows[1].cells[1].text = "Mathematical model proves ALL config paths block attack. Every CVE mapped to control verification."
    table2.rows[2].cells[0].text = "LOTL / stolen creds"
    table2.rows[2].cells[1].text = "Digital Twin detects behavioral drift -- same credentials, structurally different trajectory. Direction: MITRE-mapped."

    doc.add_heading("Deployment", level=2)
    p = doc.add_paragraph(
        "Iron Bank containers | Kubernetes Helm | IL5/IL6/JWICS | CAC/PIV mTLS | "
        "Gabriel Nimbus ready | cATO-aligned CI/CD"
    )
    p.paragraph_format.space_after = Pt(4)

    doc.add_heading("Team", level=2)
    p = doc.add_paragraph(
        "22nd Century Technologies: $90M Army SOC/MDR (800+ cleared analysts). "
        "20+ design partners with production findings in F500/federal/CI."
    )
    p.paragraph_format.space_after = Pt(2)

    doc.add_heading("Our Offer", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "Give us 4 weeks and your data. We will prove verified gap findings "
        "and behavioral detections on YOUR environment. No other vendor can "
        "make this offer."
    )
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("22nd Century Technologies  |  PROPRIETARY")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


# ═══════════════════════════════════════════════════════════════
# DOCUMENT 4: QUICK REFERENCE CARD (Internal pocket reference)
# ═══════════════════════════════════════════════════════════════
def build_quick_reference():
    doc = Document()
    _style(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INTERNAL Quick Reference -- Meeting Prep")
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CFIC 06 MAY 2026  |  DO NOT DISTRIBUTE  |  Personal Reference Only")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    # ── SAFE LANGUAGE ──
    doc.add_heading("SAFE LANGUAGE (What to Say Out Loud)", level=1)

    doc.add_heading("Preemptive Layer", level=2)
    _bullet(doc, "'We define a Formal Mathematical Model that reasons exhaustively over your complete configuration'")
    _bullet(doc, "'Not sampling, not scanning -- mathematical proof'")
    _bullet(doc, "'Zero false positives, zero false negatives within model scope'")
    _bullet(doc, "'Re-verified hourly and on every configuration change'")
    _bullet(doc, "'Three intelligences: Attack, Defense, Remediation -- continuous closed loop'")

    doc.add_heading("UEBA Layer", level=2)
    _bullet(doc, "'We create a Digital Twin of each entity instance'")
    _bullet(doc, "'User interaction is a living entity with a time dimension'")
    _bullet(doc, "'Detects sub-threshold behavioral drift that accumulates over weeks and months'")
    _bullet(doc, "'Drift direction mapped to specific MITRE ATT&CK techniques'")
    _bullet(doc, "'Cross-system entity fusion -- 10+ identity sources stitched into one trajectory'")
    _bullet(doc, "'Catches what no threshold-based system can: LOTL, credential abuse, slow lateral movement'")

    doc.add_heading("Closing / Proof", level=2)
    _bullet(doc, "'Give us 4 weeks and your data -- we will prove it'")
    _bullet(doc, "'No other vendor can make this offer because no other tool has these capabilities'")
    _bullet(doc, "'This is not an incremental improvement -- it is a fundamentally new class of defense'")
    _bullet(doc, "'We can walk through the math on a whiteboard right now if you'd like'")

    # ── THREAT FACTS ──
    doc.add_heading("Threat Facts (Cite in Breakout Sessions)", level=1)
    _bullet(doc, "Volt Typhoon: 5+ year dwell in U.S. critical infrastructure (CISA AA24-038A)", "FACT")
    _bullet(doc, "Salt Typhoon: 4+ years, 200+ telecoms, 9 U.S. carriers, CALEA wiretap access", "FACT")
    _bullet(doc, "99% of firewall breaches caused by MISCONFIGURATION, not firewall flaws (Gartner)", "STAT")
    _bullet(doc, "Config state space is exponentially large -- scanners/pentests sample a vanishing fraction", "STAT")
    _bullet(doc, "PRC actors test against published Sigma/YARA rules before campaigns", "FACT")
    _bullet(doc, "Both campaigns used only legitimate admin tools -- zero malware signatures", "FACT")
    _bullet(doc, "Average SOC triage: 30-45 min per alert at 500 alerts/day -- most never investigated", "STAT")

    # ── TEAM NUMBERS ──
    doc.add_heading("Team Numbers (If Asked)", level=1)
    _bullet(doc, "22CT: U.S. Army SOC/MDR $90M, 800+ cleared analysts")
    _bullet(doc, "22CT: USAF $108M cybersecurity operations")
    _bullet(doc, "22CT: FBI TSC $56M")
    _bullet(doc, "22CT: NAVAIR/USMC $145M")
    _bullet(doc, "20+ design partners -- F500, federal, CI, MSSPs, nation states")
    _bullet(doc, "IL5/IL6/JWICS deployment capability")
    _bullet(doc, "GSA MAS, CIO-SP3, SEWP V contract vehicles")

    # ── Q&A RESPONSES ──
    doc.add_heading("Anticipated Q&A (Safe Answers)", level=1)

    qa = [
        ("How does the formal model actually work?",
         "It reasons exhaustively over the complete configuration state space to prove "
         "whether attack paths exist. The mathematical approach is proprietary, but we "
         "can walk through it on a whiteboard or demonstrate on your data in 4 weeks."),
        ("How is this different from Exabeam/Securonix/Sentinel UEBA?",
         "Those use fixed threshold-based anomaly scoring. We create a Digital Twin of "
         "each entity with a time dimension. The Twin captures behavioral evolution "
         "continuously and detects drift that never crosses any individual threshold. "
         "They say 'anomaly score 87%.' We say 'drifting toward lateral_movement at "
         "MITRE T1021 -- here's the behavioral evidence.'"),
        ("Can this run on JWICS / air-gapped?",
         "Yes. Same containerized application. Embedding and AI inference runs locally "
         "with no external API call. Environment variable switches between cloud and "
         "air-gapped modes. Same Helm charts, same containers."),
        ("What makes you say no other tool can do this?",
         "Name any SIEM, EDR, UEBA, scanner, or pentest tool. Ask: can it mathematically "
         "prove all config paths block an attack? Can it detect behavioral drift that "
         "never crosses a threshold? Can it map that drift to a specific MITRE technique? "
         "Can it fuse identity across 10+ systems into one trajectory? The answer is no "
         "to all four for every product on the market."),
        ("What about false positives?",
         "The Preemptive layer has zero false positives by mathematical construction -- "
         "if we say there is a gap, there IS a gap. The Behavioral layer uses multiple "
         "validation layers including peer cohort comparison and feedback learning to "
         "minimize false positives. In practice, very low FP rate because we detect "
         "structural drift, not individual anomalous events."),
        ("What log sources do you need?",
         "Any source that produces standard formatted events -- Windows EventLog, Sysmon, "
         "EDR feeds, NetFlow, identity systems (AD, Okta, Azure AD), cloud audit logs. "
         "We normalize all sources into a unified pipeline."),
        ("How long to show value?",
         "Preemptive: verified gap findings within 2 weeks of config ingestion. "
         "Behavioral: meaningful drift detection after 30-day baseline establishment. "
         "Full operational capability within 90 days."),
        ("How does this relate to what AWS/Microsoft/Google are doing with agents?",
         "Those platforms provide runtime plumbing -- orchestration, identity, guardrails. "
         "They do NOT provide detection intelligence. Our engines are the domain intelligence "
         "substrate that those agent platforms need to act on. Agents without domain logic "
         "are just expensive chatbots."),
        ("How do you stop UNKNOWN attacks -- zero-days, novel TTPs with no signature?",
         "Two reasons. First, our Preemptive layer proves the DEFENSE is correct -- it does "
         "not need to recognize the specific attack. If the formal model proves all paths "
         "through your config are blocked, then ANY attack -- known or unknown -- cannot "
         "traverse those paths. Unknown CVE? If the compensating control (segmentation, "
         "inspection) is proven in place, the unknown exploit still cannot reach its target. "
         "Second, our Digital Twin detects behavioral CONSEQUENCE, not attack signatures. "
         "An unknown attack still requires the adversary to DO things: move laterally, access "
         "unusual systems, use tool combinations the real user never does. The Twin captures "
         "that the trajectory is shifting -- regardless of whether we have ever seen this "
         "specific attack before. We do not detect attacks. We prove the defense is complete, "
         "and we detect behavioral change. Signatures are irrelevant to our approach."),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        run = p.add_run("Q: " + q)
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)
        p = doc.add_paragraph("A: " + a)

    # ── AGENDA NOTES ──
    doc.add_heading("Agenda Strategy", level=1)
    doc.add_paragraph(
        "0740-0800: Arrive early. Network. Identify key stakeholders."
    )
    doc.add_paragraph(
        "0815-0900: Listen to problem set overview. Take notes on exact language "
        "they use to describe challenges. Mirror this language in your presentation."
    )
    doc.add_paragraph(
        "0925-1020: LIMITING FACTORS breakout. Be the person who diagnoses "
        "clearly. Use the threat facts above. Do not pitch yet -- just establish "
        "that you deeply understand the problem."
    )
    doc.add_paragraph(
        "1040-1140: WAYS TO OVERCOME breakout. Now introduce the concepts "
        "(Formal Mathematical Model, Digital Twin). Let others react. "
        "Gauge appetite and adjust afternoon presentation accordingly."
    )
    doc.add_paragraph(
        "1300-1405: SOLUTION BRIEF development. Work with your team to "
        "refine the problem statement using morning language. Prepare the deck flow."
    )
    doc.add_paragraph(
        "1420-1535: PRESENTATION. Deliver the deck. Close with 4-week proof offer."
    )
    doc.add_paragraph(
        "1535-1545: VOTING. Your solution is unique -- no one else in the room "
        "can offer mathematical proof + behavioral Digital Twin. Let that speak."
    )

    return doc


# ═══════════════════════════════════════════════════════════════
# GENERATE ALL
# ═══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    docs = [
        ("01_Solution_Brief.docx", build_solution_brief),
        ("02_Breakout_Talking_Points.docx", build_breakout_points),
        ("03_One_Page_Leave_Behind.docx", build_leave_behind),
        ("04_Quick_Reference_Card.docx", build_quick_reference),
    ]
    for filename, builder in docs:
        d = builder()
        path = os.path.join(OUT_DIR, filename)
        try:
            d.save(path)
        except PermissionError:
            alt = filename.replace('.docx', '_v2.docx')
            path = os.path.join(OUT_DIR, alt)
            d.save(path)
            filename = alt + " (original locked)"
        size = os.path.getsize(path) / 1024
        print(f"  {filename}: {size:.1f} KB")
    print(f"\nAll {len(docs)} documents saved to: {OUT_DIR}")
