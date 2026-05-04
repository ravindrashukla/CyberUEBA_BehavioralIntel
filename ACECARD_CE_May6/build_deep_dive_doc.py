"""Build comprehensive deep-dive reference document for CFIC/ACECARD UEBA.
Covers: problem, gaps, preemptive layer, ACECARD UEBA, architecture, deployment.
Output: CFIC_Deep_Dive_Reference.docx"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT = os.path.join(os.path.dirname(__file__), "CFIC_Deep_Dive_Reference.docx")

doc = Document()

# ─── STYLES ───
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Georgia'
    h.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)

doc.styles['Heading 1'].font.size = Pt(24)
doc.styles['Heading 2'].font.size = Pt(18)
doc.styles['Heading 3'].font.size = Pt(14)
doc.styles['Heading 4'].font.size = Pt(12)

def add_para(text, bold=False, italic=False, size=None, color=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.runs[0].bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    return table


# ════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Containerized Behavioral Cyber Defense\nfor Gabriel Nimbus')
run.font.size = Pt(28)
run.font.name = 'Georgia'
run.bold = True
run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Deep-Dive Technical Reference')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Preemptive Posture Verification + ACECARD UEBA Behavioral Intelligence')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x0E, 0x6B, 0x8A)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('22nd Century Technologies Inc.')
run.font.size = Pt(14)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CFIC Collaboration Event | Augusta, GA | May 6, 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROPRIETARY // NOT FOR PUBLIC RELEASE')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. The Problem: AI-Enabled Cyber Attacks',
    '   2.1 AI as a Force Multiplier for Adversaries',
    '   2.2 Case Study: Volt Typhoon (MITRE G1017)',
    '   2.3 Case Study: Salt Typhoon',
    '   2.4 The Structural Challenge',
    '3. Why Current Security Tools Fail',
    '   3.1 SIEM (Security Information & Event Management)',
    '   3.2 EDR/XDR (Endpoint Detection & Response)',
    '   3.3 Commercial UEBA (User & Entity Behavior Analytics)',
    '   3.4 Penetration Testing & Red Teaming',
    '   3.5 ASM & BAS (Attack Surface / Breach Simulation)',
    '   3.6 The Four Structural Gaps',
    '4. Layer 1: 22CT Preemptive — Formal Verification',
    '   4.1 What Is Formal Verification?',
    '   4.2 The Three Intelligences',
    '   4.3 How It Works: Symbolic Model of Computation',
    '   4.4 Zero False Positives/Negatives (Within Model Scope)',
    '   4.5 Supported Vendors & Coverage',
    '5. Layer 2: ACECARD UEBA — Behavioral Trajectory Intelligence',
    '   5.1 What Are Behavioral Embeddings?',
    '   5.2 The Five Behavioral Signals',
    '   5.3 Text Serialization & Embedding',
    '   5.4 CUSUM Change-Point Detection',
    '   5.5 Drift Direction & MITRE ATT&CK Mapping',
    '   5.6 Entity Fusion',
    '   5.7 ABAC Trust Loop',
    '   5.8 Cohort Analysis',
    '6. How They Stop AI-Enabled Attacks Together',
    '   6.1 Stopping Volt Typhoon — Phase by Phase',
    '   6.2 Stopping Salt Typhoon — Phase by Phase',
    '   6.3 Four Integration Points',
    '   6.4 Coverage Matrix',
    '7. Architecture & Deployment',
    '   7.1 Five-Layer Architecture',
    '   7.2 Data Flow Pipeline',
    '   7.3 Gabriel Nimbus Deployment',
    '   7.4 Air-Gapped Operations',
    '8. DoD Compliance & Accreditation',
    '9. Timeline: 90 Days to First Detection',
    '10. Key Technical Terms Glossary',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith('   '):
        p.paragraph_format.left_indent = Cm(1.0)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'Nation-state adversaries — specifically Volt Typhoon and Salt Typhoon — have demonstrated '
    'the ability to dwell inside United States critical infrastructure for years without detection. '
    'These campaigns exploit two structural blind spots in current cyber defense: (1) defensive configurations '
    'are never mathematically proven complete, and (2) behavioral detection uses fixed thresholds that '
    'Living-off-the-Land tradecraft operates below indefinitely.')

doc.add_paragraph(
    'This document describes a two-layer defense system designed to close both gaps simultaneously:')

add_bullet('Layer 1 — 22CT Preemptive: Formal mathematical verification of every firewall, IPS, IdP, SASE, and WAF '
           'configuration. Exhaustively reasons over the combinatorial config state space to PROVE that every known '
           'attack path is blocked. Zero false positives, zero false negatives within the scope of the formal model. '
           'Re-verified hourly and on every config change. Powered by Rigor AI (innovation partner).')
add_bullet('Layer 2 — ACECARD UEBA: Behavioral trajectory intelligence using 1536-dimensional embeddings, '
           'CUSUM change-point detection, and drift direction analysis mapped to MITRE ATT&CK techniques. '
           'Detects not WHAT is anomalous but what an entity is BECOMING — catching slow APT drift that '
           'never crosses any fixed threshold.')

doc.add_paragraph(
    'Together, these layers create a closed-loop system where the attack surface either never exists '
    '(Preemptive eliminates it) or the attack is detected and contained within hours (ACECARD catches '
    'behavioral drift). The system is containerized, Kubernetes-native, Iron Bank hardened, and designed '
    'for deployment into Gabriel Nimbus at IL5 (CUI), IL6 (Secret), and JWICS (TS/SCI) classification levels.')

doc.add_paragraph(
    'Timeline: 90 days from authorization to first detection capability. Aligned with DoD Zero Trust '
    'Reference Architecture (Pillars 4: Data, 7: Visibility & Analytics), NIST 800-53 Rev 5, and cATO pathway.')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 2. THE PROBLEM
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('2. The Problem: AI-Enabled Cyber Attacks', level=1)

doc.add_heading('2.1 AI as a Force Multiplier for Adversaries', level=2)

doc.add_paragraph(
    'Artificial intelligence has fundamentally shifted the balance between offense and defense in cybersecurity. '
    'Nation-state actors now leverage AI/ML capabilities across the full attack lifecycle, creating campaigns '
    'that are faster to plan, harder to detect, and more adaptive than anything in the pre-AI era.')

doc.add_heading('AI-Powered Reconnaissance', level=3)
doc.add_paragraph(
    'Large Language Models (LLMs) can analyze CVE descriptions and generate working exploit code in minutes — '
    'a task that previously required days of skilled human effort. AI systems scan hundreds of thousands of '
    'endpoints for misconfiguration patterns, identifying the weakest entry points across an entire enterprise. '
    'Automated supply-chain mapping tools use AI to identify dependencies, third-party integrations, and weak '
    'links that can be exploited for initial access.')
add_bullet('LLMs generate exploit code from CVE descriptions in minutes')
add_bullet('AI scans 100K+ endpoints to identify misconfiguration patterns')
add_bullet('Automated supply-chain mapping identifies weakest entry points')
add_bullet('AI-driven password spraying adapts to lockout policies in real-time')

doc.add_heading('AI-Assisted Evasion', level=3)
doc.add_paragraph(
    'Adversary AI can mutate command-and-control (C2) traffic patterns to evade signature-based detection. '
    'Machine learning models pre-test payloads against defender rules before launching campaigns, ensuring '
    'that the actual attack bypasses known detection signatures. Generative AI produces convincing phishing '
    'content at scale — personalized, contextually appropriate, and linguistically natural.')
add_bullet('C2 traffic mutation to evade signature detection')
add_bullet('Pre-campaign payload testing against defender rule sets')
add_bullet('Generative AI phishing at scale (personalized, contextual)')
add_bullet('Polymorphic malware that reshapes itself per-target')

doc.add_heading('AI-Extended Dwell Time', level=3)
doc.add_paragraph(
    'Adaptive beaconing uses AI to vary callback intervals based on network traffic patterns, making C2 '
    'communication indistinguishable from normal background noise. AI decides when to move laterally based '
    'on monitoring defender activity — pausing when SOC teams are active and advancing when they are not. '
    'Automated credential rotation maintains persistence even when individual accounts are disabled.')
add_bullet('Adaptive beaconing that mimics normal traffic patterns')
add_bullet('AI-driven lateral movement timing based on defender activity')
add_bullet('Automated credential rotation for persistent access')
add_bullet('Self-healing implants that re-establish access after partial remediation')

doc.add_heading('2.2 Case Study: Volt Typhoon (MITRE G1017)', level=2)

doc.add_paragraph(
    'Volt Typhoon is a People\'s Republic of China (PRC) state-sponsored cyber group that has been active '
    'since at least mid-2021. Their primary objective is pre-positioning within U.S. critical infrastructure '
    '— power grids, water systems, transportation, telecommunications — for potential future disruption '
    'during a geopolitical crisis (e.g., a Taiwan contingency). They are not stealing data; they are '
    'establishing the ability to cause physical destruction on command.')

add_para('Key Characteristics:', bold=True)
add_bullet('Active since: at least mid-2021 (likely earlier)')
add_bullet('Targets: U.S. critical infrastructure (power, water, telecom, transportation)')
add_bullet('Objective: Pre-positioning for future sabotage, not espionage')
add_bullet('Tradecraft: Almost exclusively Living-off-the-Land (LOTL)')
add_bullet('C2: KV botnet (compromised SOHO routers — legitimate Netgear/Cisco devices)')
add_bullet('Sources: CISA AA24-038A, Microsoft Threat Intelligence, FBI/NSA joint advisory')

doc.add_heading('Volt Typhoon Kill Chain', level=3)

add_para('Phase 1 — Initial Access', bold=True)
doc.add_paragraph(
    'Volt Typhoon exploits known vulnerabilities in internet-facing edge devices: Fortinet FortiOS, '
    'Ivanti Connect Secure, Cisco IOS XE (CVE-2023-20198), and Palo Alto Networks PAN-OS (CVE-2024-3400). '
    'These are perimeter devices — firewalls, VPN concentrators, and routers — that sit at the boundary '
    'between the internet and internal networks. A single unpatched edge device provides direct access '
    'to the internal network. No phishing, no user interaction required.')

add_para('Phase 2 — Credential Access', bold=True)
doc.add_paragraph(
    'Once inside, Volt Typhoon uses exclusively legitimate Windows administration tools to harvest credentials. '
    'ntdsutil.exe (a built-in Active Directory utility) is used to dump the NTDS.dit database — the file that '
    'contains all domain password hashes. LSASS memory extraction captures credentials of currently logged-in users. '
    'The critical point: every tool used is a signed, legitimate Microsoft binary. There is no malware to detect.')
add_bullet('ntdsutil.exe — legitimate AD backup utility, used to dump credential database')
add_bullet('LSASS extraction — dumps memory of the Local Security Authority Subsystem Service')
add_bullet('wmic.exe — Windows Management Instrumentation command line (legitimate admin tool)')
add_bullet('netsh.exe — network configuration utility (legitimate)')

add_para('Phase 3 — Lateral Movement', bold=True)
doc.add_paragraph(
    'Movement across the network uses standard Windows protocols: RDP (Remote Desktop Protocol), '
    'WinRM (Windows Remote Management), and SMB (Server Message Block). All connections use valid, '
    'stolen credentials. The attacker routes external C2 traffic through the KV botnet — a network of '
    'compromised small-office/home-office (SOHO) routers (Netgear, Cisco, ASUS). This makes C2 traffic '
    'appear to originate from legitimate residential IP addresses.')

add_para('Phase 4 — Persistence', bold=True)
doc.add_paragraph(
    'Persistence mechanisms include: web shells (.aspx and .jspx files on IIS/Tomcat servers), '
    'scheduled tasks (Windows Task Scheduler), and registry modifications. These ensure access survives '
    'reboots and routine maintenance. The multi-year pre-positioning means these persistence mechanisms '
    'must remain undetected indefinitely.')

add_para('Phase 5 — Defense Evasion', bold=True)
doc.add_paragraph(
    'Volt Typhoon actively clears evidence: Windows event logs are wiped, Most Recently Used (MRU) registry '
    'keys are deleted, and no custom malware is deployed. Because they use only built-in admin tools '
    '(Living-off-the-Land), there are no malware signatures to detect, no suspicious binaries to flag, '
    'and no anomalous process names to alert on. Every action they take is indistinguishable from '
    'a legitimate system administrator.')

add_para('Why Traditional Detection Fails Against Volt Typhoon:', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
doc.add_paragraph(
    'Every single tool used is a legitimate admin binary. Every login uses valid, stolen credentials. '
    'No malware exists anywhere in the kill chain. Signature-based detection is structurally blind. '
    'Threshold-based anomaly detection cannot fire because each individual action is below any '
    'reasonable alert threshold. Only the PATTERN OF BEHAVIOR over time reveals the compromise.')

doc.add_heading('2.3 Case Study: Salt Typhoon', level=2)

doc.add_paragraph(
    'Salt Typhoon is a separate PRC-affiliated cyber espionage campaign targeting telecommunications providers. '
    'As of late 2024, at least 9 major U.S. telecom carriers and dozens of international providers have been '
    'confirmed compromised. The operation targets Call Detail Records (CDRs), metadata, and — critically — '
    'access to CALEA (Communications Assistance for Law Enforcement Act) lawful-intercept systems, '
    'meaning the adversary can see who U.S. law enforcement is wiretapping.')

add_para('Key Characteristics:', bold=True)
add_bullet('Targets: 9+ U.S. telecoms, dozens globally (AT&T, Verizon, T-Mobile, Lumen confirmed)')
add_bullet('Objective: Espionage — CDR metadata, CALEA lawful-intercept access')
add_bullet('Tradecraft: Config-resident persistence (lives IN the router configuration)')
add_bullet('Key tools: JumbledPath (connection-chaining ELF), GhostSpider, Demodex')
add_bullet('Dwell time: Multi-year')
add_bullet('Sources: FBI/CISA advisory (Dec 2024), Cisco Talos, Trend Micro research')

doc.add_heading('Salt Typhoon Kill Chain', level=3)

add_para('Phase 1 — Edge Exploitation', bold=True)
doc.add_paragraph(
    'Salt Typhoon targets network infrastructure devices directly: Cisco IOS XE (CVE-2023-20198 for initial access, '
    'CVE-2018-0171 Smart Install abuse), and Ivanti Connect Secure (CVE-2024-21887). They targeted over 1,000 '
    'Cisco devices globally. The key difference from Volt Typhoon: Salt Typhoon\'s target IS the network device '
    'itself, not just a stepping stone to internal hosts.')

add_para('Phase 2 — Persistence (Config-Resident)', bold=True)
doc.add_paragraph(
    'Salt Typhoon creates local privileged accounts directly on router operating systems. They modify Access Control '
    'Lists (ACLs) to whitelist their IP addresses. They enable Cisco Guest Shell (a Linux container running inside '
    'IOS XE) for persistent command execution. Critically, their persistence mechanism IS the router configuration '
    'itself — not a file on disk, not a binary, not a process. This means traditional file-based scanning '
    'or process monitoring cannot detect their presence.')

add_para('Phase 3 — TACACS+ Credential Harvest', bold=True)
doc.add_paragraph(
    'TACACS+ (Terminal Access Controller Access-Control System Plus) runs on TCP port 49 and is used for '
    'AAA (Authentication, Authorization, Accounting) on network devices. Salt Typhoon performs on-box packet '
    'captures of TACACS+ traffic, harvesting network engineer credentials as they authenticate to routers. '
    'The tool "JumbledPath" — a custom ELF binary — chains connections across multiple compromised routers '
    'to obscure the attacker\'s origin.')

add_para('Phase 4 — Collection (GRE Tunnels + CALEA)', bold=True)
doc.add_paragraph(
    'Generic Routing Encapsulation (GRE) tunnels are established to siphon Call Detail Records and metadata. '
    'The adversary targets CALEA lawful-intercept systems — the infrastructure that U.S. law enforcement uses '
    'for court-ordered wiretaps. Access to CALEA means the adversary knows which targets are under surveillance. '
    'Additional implants (GhostSpider, Demodex) are deployed for long-term collection on specific high-value targets.')

add_para('Phase 5 — Long-Term Espionage', bold=True)
doc.add_paragraph(
    'Salt Typhoon maintains slow, persistent exfiltration of CDR/metadata over months and years. '
    'C2 communications are encrypted and routed through legitimate infrastructure. On network gear, they use '
    'Guest Shell (LOTL) rather than deploying custom malware. The exfiltration rate is deliberately low to '
    'avoid triggering volume-based alerts.')

add_para('Why This Is Different From Traditional Attacks:', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
doc.add_paragraph(
    'Salt Typhoon literally lives INSIDE the router configuration. Their persistence IS an ACL entry, '
    'a local account, a Guest Shell configuration. There is no file to scan, no process to kill, no '
    'binary to signature-match. The only way to detect them is to continuously verify that the router '
    'configuration matches an approved baseline — and to detect the behavioral drift of traffic patterns '
    'that reveals credential harvesting and data exfiltration.')

doc.add_heading('2.4 The Structural Challenge', level=2)

doc.add_paragraph(
    'Both campaigns exploit the same fundamental architectural weakness: the gap between "we deployed defenses" '
    'and "we can PROVE those defenses work against these specific adversaries." Current security operates on '
    'assumption and sampling. Penetration tests sample a subset of attack paths. Vulnerability scanners check '
    'known CVEs. SIEM correlates known patterns. But nobody PROVES that every configuration is correct, '
    'and nobody tracks behavioral DIRECTION over time.')

doc.add_paragraph(
    'This is not a tuning problem. More rules, more alerts, more tools operating on the same principles '
    'will not close this gap. A fundamentally different approach is required.')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 3. WHY CURRENT TOOLS FAIL
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Why Current Security Tools Fail', level=1)

doc.add_paragraph(
    'This section examines each major category of security tooling and explains — structurally, not operationally — '
    'why it cannot detect Volt Typhoon or Salt Typhoon tradecraft. These are not implementation failures; they are '
    'architectural limitations inherent to each approach.')

doc.add_heading('3.1 SIEM (Security Information & Event Management)', level=2)

add_para('What It Does:', bold=True)
doc.add_paragraph(
    'SIEM systems (Splunk, QRadar, Sentinel, Elastic SIEM) collect logs from across the enterprise, normalize '
    'them, and correlate events against predefined rules and known attack patterns. They are the central nervous '
    'system of most SOC operations.')

add_para('Why It Fails Against These Adversaries:', bold=True)
add_bullet('Rule-based correlation only catches KNOWN patterns. Novel combinations of legitimate tools are invisible.')
add_bullet('Alert fatigue: 10,000+ events/day. Analysts investigate <5% of alerts. True positives buried in noise.')
add_bullet('No behavioral trajectory. SIEM sees individual events, not patterns of behavioral evolution over weeks.')
add_bullet('Volt Typhoon uses only legitimate admin commands — they look identical to normal admin activity in logs.')
add_bullet('Salt Typhoon operates on network gear that often has limited/no SIEM integration.')

doc.add_heading('3.2 EDR/XDR (Endpoint Detection & Response)', level=2)

add_para('What It Does:', bold=True)
doc.add_paragraph(
    'EDR/XDR systems (CrowdStrike Falcon, Microsoft Defender for Endpoint, SentinelOne, Carbon Black) '
    'monitor endpoint processes, file system changes, and network connections. They use behavioral analysis '
    'and ML to detect malicious activity on individual hosts.')

add_para('Why It Fails:', bold=True)
add_bullet('Living-off-the-Land binaries (LOLBins) are explicitly ALLOWLISTED. ntdsutil, wmic, netsh are legitimate.')
add_bullet('Cannot reason about cross-host behavioral evolution — each endpoint is analyzed independently.')
add_bullet('Blind on network gear. EDR agents cannot run on Cisco IOS XE, Fortinet FortiOS, or PAN-OS.')
add_bullet('Salt Typhoon operates entirely on routers/switches — no endpoint to monitor.')
add_bullet('Even on endpoints, valid credential usage from ntdsutil is indistinguishable from legitimate backup operations.')

doc.add_heading('3.3 Commercial UEBA (User & Entity Behavior Analytics)', level=2)

add_para('What It Does:', bold=True)
doc.add_paragraph(
    'UEBA systems (Exabeam, Securonix, Microsoft Sentinel UEBA, Gurucul) build behavioral baselines for users '
    'and entities, then alert when behavior deviates beyond a threshold. They assign "risk scores" based on '
    'deviation magnitude.')

add_para('Why It Fails:', bold=True)
add_bullet('FIXED THRESHOLDS: "Alert when risk score > 85." Volt Typhoon operates at scores of 20-40 — forever.')
add_bullet('No drift DIRECTION: "User is 87% anomalous" tells you nothing about WHAT they\'re becoming.')
add_bullet('Per-source silos: each data source analyzed independently. No cross-system trajectory.')
add_bullet('No identity fusion: AD shows one thing, Okta another, VPN another. 4 weak signals, no alert.')
add_bullet('Cannot explain WHY something is anomalous — just that it IS. Analysts waste time investigating benign drift.')
add_bullet('Living-off-the-Land tradecraft produces minimal deviation because the TOOLS are legitimate — only the PATTERN over time reveals intent.')

add_para('The Core Problem with Threshold-Based Detection:', bold=True)
doc.add_paragraph(
    'Consider an attacker who drifts 0.03 units per day from their behavioral baseline. A typical threshold is '
    'set at 0.15 (to avoid false positives from normal variation). The attacker NEVER triggers an alert — ever. '
    'They can operate for years. This is exactly what Volt Typhoon does. The per-window deviation is tiny. '
    'Only the CUMULATIVE drift over time reveals the compromise. No commercial UEBA uses cumulative detection.')

doc.add_heading('3.4 Penetration Testing & Red Teaming', level=2)

add_para('What It Does:', bold=True)
doc.add_paragraph(
    'Pentesters and red teams simulate adversary behavior against the organization\'s defenses. They attempt '
    'to exploit vulnerabilities, move laterally, and achieve objectives (data exfiltration, domain admin, etc.).')

add_para('Why It Fails:', bold=True)
add_bullet('SAMPLING problem: a pentest exercises a subset of attack paths. It cannot cover the combinatorially explosive configuration state space.')
add_bullet('Point-in-time: a pentest runs for 2-4 weeks, then stops. The adversary operates 365/24/7.')
add_bullet('Known TTP bias: red teams exercise known techniques. Novel chains and combinations are missed.')
add_bullet('Annual cadence: most organizations pentest 1-2x/year. Configurations drift daily.')
add_bullet('Volt Typhoon\'s tradecraft is so subtle that most red teams cannot replicate it — they use noisier, faster techniques.')

doc.add_heading('3.5 ASM & BAS (Attack Surface Management / Breach & Attack Simulation)', level=2)

add_para('What It Does:', bold=True)
doc.add_paragraph(
    'ASM tools (Mandiant ASM, Cortex Xpanse, CyCognito) enumerate the external attack surface. BAS tools '
    '(SafeBreach, AttackIQ, Cymulate) simulate known attacks against production defenses to validate that controls work.')

add_para('Why It Fails:', bold=True)
add_bullet('ASM enumerates — it cannot PROVE controls block every path. Knowing you have an exposed port is not the same as proving your firewall blocks all exploitation of that port.')
add_bullet('BAS uses statistical confidence, not mathematical certainty. "90% of our tests passed" still means 10% are exploitable.')
add_bullet('Neither can reason about config-resident persistence (Salt Typhoon living in router ACLs).')
add_bullet('Neither provides behavioral trajectory analysis — they validate point-in-time control effectiveness.')
add_bullet('The combinatorial explosion of multi-vendor, multi-zone configurations exceeds what sampling can cover.')

doc.add_heading('3.6 The Four Structural Gaps', level=2)

doc.add_paragraph(
    'Synthesizing the above, four structural gaps emerge that no current tool addresses:')

add_table(
    ['Gap #', 'Description', 'Exploited By', 'Required Innovation'],
    [
        ['1', 'Configurations never PROVEN closed', 'Both (edge device entry)', 'Formal mathematical verification'],
        ['2', 'Behavioral detection has no direction', 'Volt Typhoon (LOTL)', 'Drift direction + MITRE mapping'],
        ['3', 'Identity fragmented across 10+ systems', 'Both (cross-domain movement)', 'Entity fusion into one trajectory'],
        ['4', 'No cloud-native UEBA for Gabriel Nimbus', 'N/A (capability gap)', 'K8s-native, Iron Bank, IL5-JWICS'],
    ]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 4. LAYER 1: PREEMPTIVE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Layer 1: 22CT Preemptive — Formal Verification', level=1)

doc.add_paragraph(
    'The 22CT Preemptive layer addresses Gap #1: configurations are never mathematically proven complete. '
    'Instead of testing or sampling, Preemptive uses formal mathematical methods to PROVE that every defensive '
    'configuration blocks every known attack path — or surfaces the exact gap with a vendor-specific fix.')

doc.add_heading('4.1 What Is Formal Verification?', level=2)

doc.add_paragraph(
    'Formal verification is a mathematical technique borrowed from hardware design and safety-critical software. '
    'It exhaustively reasons over ALL possible states of a system to prove that certain properties hold universally — '
    'not just for the cases you tested, but for EVERY possible case.')

add_para('Analogy:', bold=True)
doc.add_paragraph(
    'A penetration test is like checking if specific doors are locked. Formal verification is like '
    'mathematically proving that the building has no openings of any kind — every wall, window, door, vent, '
    'and pipe is verified as sealed. Not by checking each one, but by reasoning about the complete structure.')

add_para('Key Properties:', bold=True)
add_bullet('Completeness: reasons over the ENTIRE state space, not a sample')
add_bullet('Soundness: if it says "secure," it IS secure within the model')
add_bullet('Exhaustiveness: no path is missed, regardless of how obscure')
add_bullet('Continuous: re-verifies hourly and on every configuration change')

add_para('Why This Matters for Cybersecurity:', bold=True)
doc.add_paragraph(
    'A modern enterprise firewall (e.g., Palo Alto PA-5400) can have millions of possible rule combinations. '
    'With multi-zone architectures, NAT, application-ID, URL filtering, and threat prevention profiles, '
    'the configuration state space is combinatorially explosive. Human review catches the obvious issues. '
    'Pentests sample a fraction. Formal verification covers ALL of them.')

doc.add_heading('4.2 The Three Intelligences', level=2)

add_para('Attack Intelligence', bold=True)
doc.add_paragraph(
    'A continuously updated graph of adversary TTPs (Tactics, Techniques, and Procedures) built by ML, NLP, '
    'and LLM pipelines. Ingests: CVE databases, CISA advisories, Cyber Threat Intelligence feeds, IOCs '
    '(Indicators of Compromise), and vendor PSIRTs (Product Security Incident Response Teams). '
    'Maintains the complete TTP graph for Volt Typhoon (G1017) and Salt Typhoon, updated automatically '
    'on every new advisory or intelligence publication.')
add_bullet('Input: CVE, CISA, MITRE ATT&CK, CTI feeds, IOCs, PSIRTs')
add_bullet('Output: Complete attack graph — all known paths from initial access to objective')
add_bullet('Update cadence: automatic on every new advisory')

add_para('Defense Intelligence', bold=True)
doc.add_paragraph(
    'The Symbolic Model of Computation — the core innovation. This component models every defensive '
    'configuration as formal logic: firewall rules become logical predicates, IPS signatures become constraints, '
    'identity provider policies become access control assertions. The formal engine then reasons exhaustively '
    'over this model, proving that every attack path (from Attack Intelligence) is blocked by at least one '
    'defensive control — or identifying the exact gap where coverage fails.')
add_bullet('Models: NGFW rules, IPS signatures, IdP policies, SASE configs, WAF rules as formal logic')
add_bullet('Reasons: over combinatorially explosive state space (millions of rule combinations)')
add_bullet('Proves: every attack path is blocked — or surfaces the specific gap')
add_bullet('Scope: proofs are valid within the formal model (does not cover zero-day in the model itself)')

add_para('Remediation Intelligence', bold=True)
doc.add_paragraph(
    'When Defense Intelligence surfaces a gap, Remediation Intelligence generates vendor-specific, risk-prioritized '
    'configuration fixes. These are not generic recommendations ("patch your firewall") but specific rule changes '
    '("add rule X at position Y in zone Z on device D"). Each proposed fix is validated to ensure it does not '
    'introduce new gaps or disrupt business traffic. Output formats include SOC playbooks, ServiceNow/JIRA tickets, '
    'and direct device API calls.')
add_bullet('Generates: vendor-specific config fixes (PAN, Fortinet, Cisco, Check Point syntax)')
add_bullet('Validates: fix introduces no new gaps and no business disruption')
add_bullet('Delivers: SOC playbooks, ServiceNow/JIRA integration, direct API push')
add_bullet('Guardrails: agentic-AI reasoners with human-in-the-loop for critical changes')

doc.add_heading('4.3 How It Works: Symbolic Model of Computation', level=2)

doc.add_paragraph(
    'The Symbolic Model of Computation is the mathematical engine at the heart of Defense Intelligence. '
    'It works by translating device configurations into formal logic, then using theorem-proving techniques '
    'to verify security properties.')

add_para('Process:', bold=True)
add_bullet('1. PARSE: Ingest device configuration (running-config, API export, NGFW policy)')
add_bullet('2. TRANSLATE: Convert rules into formal logic predicates (e.g., "if source ∈ zone_A AND dest_port = 443 AND app = ssl THEN action = allow")')
add_bullet('3. MODEL: Build complete state-space model of all possible packet flows')
add_bullet('4. VERIFY: For each attack path from Attack Intelligence, prove it is blocked')
add_bullet('5. REPORT: Surface any unblocked path with: exact path, responsible device, proposed fix')
add_bullet('6. RE-VERIFY: Trigger on config change event or hourly schedule')

add_para('Combinatorial Explosion — Why This Requires Formal Methods:', bold=True)
doc.add_paragraph(
    'Consider a single Palo Alto firewall with: 500 rules × 50 zones × 20 applications × 10 user groups × '
    'NAT policies × decryption policies × threat profiles. The number of unique packet-flow decisions is '
    'on the order of billions. A pentest might exercise hundreds of paths. Formal verification reasons over '
    'ALL of them in minutes, using SAT/SMT solvers that efficiently prune the search space without missing any case.')

doc.add_heading('4.4 Zero False Positives / Zero False Negatives (Within Model Scope)', level=2)

doc.add_paragraph(
    'This claim requires careful qualification:')

add_para('Zero False Positives:', bold=True)
doc.add_paragraph(
    'If the system reports a gap, that gap genuinely exists in the configuration. The formal proof is sound — '
    'it cannot "hallucinate" a vulnerability that is not there. This eliminates the single largest operational '
    'cost of traditional security tools: analyst time spent investigating false alarms.')

add_para('Zero False Negatives:', bold=True)
doc.add_paragraph(
    'If the system reports "all paths blocked," then within the scope of the formal model, no known attack path '
    'is unmitigated. The proof is complete over the modeled state space.')

add_para('Scope Limitation:', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
doc.add_paragraph(
    'The guarantee is bounded by the formal model. If a new TTP is discovered that the attack graph does not yet '
    'contain, or if a device has undocumented behavior not captured in the model, the guarantee does not extend '
    'to those cases. This is why the system updates automatically on every new advisory — to keep the attack graph '
    'current. It is also why Layer 2 (ACECARD) exists: to catch what escapes the formal model through behavioral '
    'detection.')

doc.add_heading('4.5 Supported Vendors & Coverage', level=2)

add_table(
    ['Vendor', 'Product', 'Coverage'],
    [
        ['Palo Alto Networks', 'PA-Series, Panorama, Prisma Access', 'NGFW rules, App-ID, URL, Threat Prevention, Decryption'],
        ['Fortinet', 'FortiGate, FortiManager', 'Firewall policies, IPS, Web Filter, Application Control'],
        ['Cisco', 'ASA, Firepower/FTD, ISE', 'ACLs, IPS/IDS, NAC policies, Identity-based access'],
        ['Check Point', 'R81+, SmartConsole', 'Security policies, HTTPS inspection, Threat Prevention'],
    ]
)

doc.add_paragraph(
    'Additional integrations: CrowdStrike (EDR verification), Okta/Azure AD (IdP policy), '
    'Zscaler (SASE), Cloudflare/Imperva (WAF). Coverage expands continuously.')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 5. LAYER 2: ACECARD UEBA
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Layer 2: ACECARD UEBA — Behavioral Trajectory Intelligence', level=1)

doc.add_paragraph(
    'ACECARD (Adaptive Cyber Entity Characterization, Assessment & Risk Detection) is a behavioral intelligence '
    'platform that closes Gaps #2, #3, and #4. It transforms raw security telemetry into 1536-dimensional '
    'behavioral vectors, then tracks how those vectors CHANGE over time — detecting not what is anomalous '
    'but what an entity is BECOMING.')

doc.add_heading('5.1 What Are Behavioral Embeddings?', level=2)

doc.add_paragraph(
    'An embedding is a numerical representation of something complex in a high-dimensional vector space. '
    'Just as word embeddings (Word2Vec, BERT) represent words as vectors where semantic similarity = geometric '
    'proximity, behavioral embeddings represent an entity\'s behavior as a vector where behavioral similarity = '
    'geometric proximity.')

add_para('Key Properties of Behavioral Embeddings:', bold=True)
add_bullet('Dimensionality: 1536 dimensions (same as OpenAI text-embedding-3-small)')
add_bullet('Semantic: similar behaviors → similar vectors (close in cosine distance)')
add_bullet('Compositional: individual signal embeddings combine into a composite behavioral fingerprint')
add_bullet('Temporal: each time window produces a new embedding, enabling trajectory tracking')
add_bullet('Comparable: entities with similar roles cluster together; drift AWAY from cluster = anomaly')

add_para('Why 1536 Dimensions?', bold=True)
doc.add_paragraph(
    'This matches the dimensionality of the embedding model (text-embedding-3-small or equivalent local model). '
    'Higher dimensions capture more nuance in behavioral patterns. The unified dimensionality means all signals, '
    'all entities, and all threat concepts live in the SAME vector space — enabling direct cosine comparison '
    'between a user\'s drift vector and a threat concept like "lateral_movement."')

add_para('Why Text Embedding (Not Direct Numerical Features)?', bold=True)
doc.add_paragraph(
    'Traditional UEBA feeds raw metrics (logon_count, bytes_transferred) into statistical models. ACECARD '
    'instead serializes these metrics as structured natural language text, then embeds that text. This approach:')
add_bullet('Leverages the semantic understanding of pre-trained language models')
add_bullet('Captures relationships between features that numerical vectors miss')
add_bullet('Makes the system model-agnostic — swap embedding models without retraining')
add_bullet('Enables direct comparison with threat concepts also expressed as text')

doc.add_heading('5.2 The Five Behavioral Signals', level=2)

doc.add_paragraph(
    'ACECARD decomposes entity behavior into five orthogonal signals. Each captures a different attack dimension. '
    'Together, they form a complete behavioral fingerprint that reveals any class of cyber attack.')

add_para('Signal 1: AUTH (Authentication & Access)', bold=True)
doc.add_paragraph(
    'Captures how an entity authenticates and what it accesses. Source data: Windows Event IDs 4624 (logon), '
    '4625 (failed logon), 4768 (Kerberos TGT), Okta authentication logs, Azure AD sign-in logs, AWS CloudTrail.')
add_bullet('Metrics: logon_count, failure_rate, unique_hosts_accessed, off_hours_ratio, impossible_travel_flag, mfa_skip_ratio, geo_diversity, first_seen_hosts')
add_bullet('Attack signature: Volt Typhoon credential access shows impossible_travel + new unique_hosts + off_hours access patterns')

add_para('Signal 2: PROCESS (Process & Execution)', bold=True)
doc.add_paragraph(
    'Captures what software runs on an entity. Source data: Sysmon Event ID 1 (process creation), EID 3 '
    '(network connection), CrowdStrike process telemetry, Microsoft Defender, Linux auditd.')
add_bullet('Metrics: unique_processes, cmdline_entropy, lolbin_count (Living-off-the-Land binary usage), parent_child_depth, unsigned_ratio, encoded_command_count')
add_bullet('Attack signature: Volt Typhoon shows sudden lolbin_count increase (ntdsutil, wmic, netsh) from previously quiet workstation')

add_para('Signal 3: NETWORK (Network Activity)', bold=True)
doc.add_paragraph(
    'Captures network communication patterns. Source data: NetFlow/IPFIX, Zeek (formerly Bro) logs, '
    'PAN-OS Traffic logs, FortiGate flow logs, AWS VPC Flow Logs.')
add_bullet('Metrics: unique_dest_ips, bytes_out_ratio, beacon_score, dns_query_rate, geo_anomaly_score, admin_share_access_count, new_external_connections')
add_bullet('Attack signature: both Typhoons show beacon_score increase (C2) + new unique_dest_ips (lateral) + bytes_out_ratio shift (exfil)')

add_para('Signal 4: FILE (File System Activity)', bold=True)
doc.add_paragraph(
    'Captures file operations. Source data: Sysmon Event ID 11 (file creation), Windows Security 4663 '
    '(object access), DLP systems, cloud storage audit logs.')
add_bullet('Metrics: files_created, files_deleted, sensitive_access_count, archive_creates, extension_changes, shadow_copy_operations')
add_bullet('Attack signature: NTDS.dit dump shows shadow_copy_operations spike; data staging shows archive_creates + sensitive_access increase')

add_para('Signal 5: IDENTITY (Identity & Privilege)', bold=True)
doc.add_paragraph(
    'Captures privilege and identity changes. Source data: Windows 4672 (special privileges assigned), '
    '4728 (member added to security group), CloudTrail IAM events, Kubernetes audit logs, AD change logs.')
add_bullet('Metrics: privilege_escalations, group_membership_adds, mfa_bypass_events, role_changes, service_account_usage, administrative_actions')
add_bullet('Attack signature: persistence via new local accounts, ACL modifications, service account creation')

doc.add_heading('5.3 Text Serialization & Embedding', level=2)

doc.add_paragraph(
    'The embedding pipeline converts raw metrics into behavioral vectors through a three-step process:')

add_para('Step 1: Aggregate', bold=True)
doc.add_paragraph(
    'Raw events within a 1-hour window are aggregated into metrics per entity. Example: 47 authentication events '
    'for user U in hour H → logon_count=47, failure_rate=0.04, unique_hosts=3, off_hours_ratio=0.0, etc.')

add_para('Step 2: Serialize to Text', bold=True)
doc.add_paragraph(
    'Metrics are formatted as structured natural language:')
add_para(
    '"AUTH signal for entity user:jsmith in window 2026-03-15T14:00Z: '
    'logon_count=47 (baseline=42, z=0.8), failure_rate=0.04 (baseline=0.02, z=1.1), '
    'unique_hosts=3 (baseline=2, z=0.5), off_hours_ratio=0.0 (baseline=0.0, z=0.0). '
    'Context: role=domain_admin, department=IT, typical_hours=08-18."',
    italic=True)

add_para('Step 3: Embed', bold=True)
doc.add_paragraph(
    'The text is passed through an embedding model to produce a 1536-dimensional vector. '
    'Default (air-gapped): local embedding model such as E5-Mistral-7B-Instruct or nomic-embed-text-v1.5. '
    'IL5 option: Azure OpenAI text-embedding-3-small via GovCloud endpoint. '
    'The result is a single vector that captures the SHAPE of this signal\'s behavior in this time window.')

add_para('Step 4: Compose', bold=True)
doc.add_paragraph(
    'The 5 individual signal embeddings are combined into a single composite embedding using weighted averaging '
    '(not concatenation — that would produce 7680 dimensions and break comparability). Default weights are equal '
    '(0.2 each) but can be tuned per entity type. The composite embedding represents the entity\'s complete '
    'behavioral state in that time window.')

doc.add_heading('5.4 CUSUM Change-Point Detection', level=2)

doc.add_paragraph(
    'The Cumulative Sum (CUSUM) algorithm is the mathematical core of ACECARD\'s detection capability. '
    'It solves the problem that defeats all threshold-based systems: detecting slow, sustained drift that '
    'NEVER crosses any fixed threshold.')

add_para('Background: Page\'s CUSUM (1954)', bold=True)
doc.add_paragraph(
    'Developed by E.S. Page in 1954 for industrial quality control, CUSUM detects the point at which a '
    'process shifts from one statistical regime to another. It was later proven by Moustakides (1986) to be '
    'minimax optimal for detecting mean shifts in i.i.d. (independent and identically distributed) sequences — '
    'meaning: for a given false alarm rate, no other algorithm detects a shift faster.')

add_para('The Algorithm:', bold=True)
doc.add_paragraph('S(t) = max(0, S(t-1) + drift(t) - μ - k)', style='Normal')
p = doc.add_paragraph()
run = p.add_run('S(t) = max(0, S(t-1) + drift(t) - μ - k)')
run.font.name = 'Consolas'
run.font.size = Pt(12)
run.bold = True

doc.add_paragraph('Where:')
add_bullet('S(t) = cumulative sum statistic at time t (starts at 0, resets to 0 on negative)')
add_bullet('drift(t) = cosine distance between current embedding and baseline embedding')
add_bullet('μ (mu) = expected baseline drift (mean drift under normal conditions, calibrated per entity)')
add_bullet('k = slack parameter = 0.5 × σ (half standard deviation of normal drift)')
add_bullet('ALARM when S(t) exceeds threshold h = 4 × σ')

add_para('Why CUSUM Catches Volt Typhoon:', bold=True)
doc.add_paragraph('Consider Volt Typhoon\'s operational tempo:')
add_bullet('Baseline drift: μ = 0.03, σ = 0.03 (normal daily variation)')
add_bullet('APT drift per window: 0.06 (consistently above normal, but below any fixed threshold)')
add_bullet('Fixed threshold (traditional UEBA) = 0.15 → NEVER fires')
add_bullet('CUSUM parameters: k = 0.015, h = 0.12')
add_bullet('Excess per window: 0.06 - 0.03 - 0.015 = 0.015 (accumulated, never reset)')
add_bullet('Windows to alarm: 0.12 / 0.015 = 8 windows = 8 HOURS')

add_para('Result:', bold=True, color=RGBColor(0x0E, 0x6B, 0x8A))
doc.add_paragraph(
    'Where traditional UEBA allows years of undetected dwell, CUSUM detects the same activity in 8 hours. '
    'This is not a tuning difference — it is a structural algorithmic advantage. Fixed thresholds can NEVER '
    'catch drift that stays below them. CUSUM accumulates evidence until confidence is achieved.')

add_para('Optimality Guarantee:', bold=True)
doc.add_paragraph(
    'Moustakides (1986) proved that Page\'s CUSUM is minimax optimal: for any given average run length between '
    'false alarms (ARL0), no sequential detection procedure has a shorter detection delay. Caveat: this optimality '
    'holds for i.i.d. observations. Real behavioral drift may have serial correlation, so the theoretical '
    'optimality is a bound, not an exact guarantee. In practice, CUSUM still dramatically outperforms fixed thresholds.')

doc.add_heading('5.5 Drift Direction & MITRE ATT&CK Mapping', level=2)

doc.add_paragraph(
    'Knowing THAT an entity is drifting is necessary but not sufficient. The SOC analyst needs to know: '
    '"drifting TOWARD WHAT?" ACECARD answers this with drift direction analysis.')

add_para('How It Works:', bold=True)
add_bullet('1. Compute drift vector: V(drift) = embedding(t) - embedding(t-1)')
add_bullet('2. Compare drift vector against 8 pre-defined threat concept embeddings using cosine similarity')
add_bullet('3. Output: ranked list of concepts with alignment scores and mapped MITRE technique IDs')

add_para('The 8 Threat Concepts:', bold=True)
doc.add_paragraph(
    'Each concept is a text description of a threat behavior pattern, embedded into the same 1536-d space:')

add_table(
    ['Concept', 'Description', 'Primary MITRE Techniques'],
    [
        ['lateral_movement', 'User accessing many new hosts, admin shares, cross-subnet traffic', 'T1021 (Remote Services), T1570'],
        ['credential_theft', 'LSASS access, ntdsutil, Kerberoasting patterns', 'T1003 (Credential Dumping), T1558'],
        ['data_exfiltration', 'Large outbound transfers, archive creation, cloud upload spikes', 'T1041 (Exfiltration over C2), T1048'],
        ['privilege_escalation', 'New group memberships, role changes, service account abuse', 'T1078 (Valid Accounts), T1068'],
        ['c2_beacon', 'Regular interval external connections, low-entropy payloads', 'T1071 (Application Layer Protocol)'],
        ['lotl_execution', 'Living-off-the-Land binary usage spike (wmic, certutil, rundll32)', 'T1059 (Command Scripting), T1218'],
        ['insider_hoarding', 'Bulk file access, staging, unusual download patterns', 'T1074 (Data Staged), T1005'],
        ['reconnaissance', 'Port scanning, DNS enumeration, directory listing', 'T1046 (Network Service Scanning)'],
    ]
)

add_para('Example Output:', bold=True)
add_para(
    '"Entity user:jsmith drift direction: lateral_movement (0.78, T1021), credential_theft (0.45, T1003), '
    'c2_beacon (0.22, T1071). Highest alignment: lateral_movement with 0.78 cosine similarity."',
    italic=True)

doc.add_paragraph(
    'This transforms an opaque "anomaly score" into an actionable, explainable alert that tells the analyst '
    'exactly what behavior pattern the entity is evolving toward, with a specific MITRE technique ID for '
    'playbook lookup and response automation.')

doc.add_heading('5.6 Entity Fusion', level=2)

doc.add_paragraph(
    'In most enterprises, a single human appears as 10+ distinct identities across different systems:')

add_bullet('Active Directory: DOMAIN\\jsmith')
add_bullet('Azure AD: jsmith@company.com')
add_bullet('Okta: john.smith@company.com')
add_bullet('AWS IAM: arn:aws:iam::123456:user/jsmith')
add_bullet('Kubernetes: system:serviceaccount:namespace:jsmith-sa')
add_bullet('CrowdStrike: agent-id-abc123 (on their workstation)')
add_bullet('VPN: jsmith@vpn.company.com')
add_bullet('Splunk: splunk-user-jsmith')
add_bullet('PKI: CN=John Smith, OU=Engineering')
add_bullet('TACACS+: jsmith (network device authentication)')

doc.add_paragraph(
    'Without fusion, an adversary moving across these systems generates 4-5 weak, unrelated signals — '
    'none strong enough to alert. WITH fusion, those same signals combine into a single strong trajectory '
    'showing clear lateral movement across systems.')

add_para('How Fusion Works:', bold=True)
add_bullet('Graph-based identity resolution: connect identifiers that share attributes (email, name, device, IP)')
add_bullet('One entity_uuid per resolved entity, regardless of how many identifiers they have')
add_bullet('All 5 behavioral signals aggregated under the unified entity_uuid')
add_bullet('Cross-system behavior becomes a single trajectory in the same vector space')
add_bullet('Lateral movement across identity boundaries now reads as ONE drift trajectory, not 4 weak signals')

doc.add_heading('5.7 ABAC Trust Loop', level=2)

doc.add_paragraph(
    'ABAC (Attribute-Based Access Control) is the enforcement mechanism that makes detection actionable. '
    'Instead of just generating alerts for human review, ACECARD can dynamically adjust access permissions '
    'based on behavioral trust state.')

add_para('Trust States:', bold=True)
add_bullet('TRUSTED: normal behavior, full access per role')
add_bullet('ELEVATED_WATCH: CUSUM rising, alert generated. Increased logging, step-up MFA on sensitive resources.')
add_bullet('RESTRICTED: high-confidence drift detected. Access limited to essential functions. Analyst notification.')
add_bullet('BLOCKED: CRITICAL alert confirmed. All access revoked pending investigation.')

add_para('Enforcement Points:', bold=True)
add_bullet('Identity Provider (conditional access policies)')
add_bullet('NGFW (dynamic address groups)')
add_bullet('SASE (user risk-based policies)')
add_bullet('Application layer (feature flags, read-only mode)')

doc.add_heading('5.8 Cohort Analysis', level=2)

doc.add_paragraph(
    'Cohort analysis detects anomalies by comparing an entity to its PEERS rather than just its own baseline. '
    'Entities are grouped by role (all Domain Admins, all network engineers, all finance users) and their '
    'behavioral embeddings are compared.')

add_para('Cohort Z-Score:', bold=True)
doc.add_paragraph(
    'For each entity, compute: z = (drift_magnitude - cohort_mean) / cohort_std. A z-score > 3.0 indicates '
    'the entity is drifting significantly more than its peers — even if the absolute drift is small. '
    'This catches Volt Typhoon targeting a single admin account among 50 similar accounts.')

add_para('Co-Drifting Detection:', bold=True)
doc.add_paragraph(
    'If multiple entities in different cohorts drift in the SAME direction simultaneously, this signals '
    'a coordinated campaign (botnet, coordinated insider, or multi-stage APT). Co-drift is detected by '
    'computing pairwise cosine similarity of drift vectors across entities.')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 6. HOW THEY STOP ATTACKS TOGETHER
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('6. How They Stop AI-Enabled Attacks Together', level=1)

doc.add_paragraph(
    'Neither layer alone is sufficient. Preemptive eliminates the attack surface that exists in configuration — '
    'but cannot detect adversaries who bypass controls through valid credentials and legitimate tools. ACECARD '
    'detects behavioral evolution — but only AFTER the adversary is inside. Together, they create defense-in-depth '
    'where the attack surface either never exists OR the attack is detected within hours.')

doc.add_heading('6.1 Stopping Volt Typhoon — Phase by Phase', level=2)

add_table(
    ['Kill-Chain Phase', '22CT Preemptive Response', 'ACECARD Response'],
    [
        ['Initial Access (Edge CVE)', 'PROVES IPS/WAF blocks exploitation of CVE-2023-20198, CVE-2024-3400. If compensating control needed, generates and validates fix. Attack surface ELIMINATED.', 'If bypass occurs: Auth signal detects anomalous management-plane source IP. Drift = 0.08 from baseline.'],
        ['Credential Access (ntdsutil)', 'PROVES EDR script-control policy blocks ntdsutil execution on non-admin workstations. Validates allow/block lists are complete.', 'Process signal: lolbin_count jumps 0→7 in one window. CUSUM alarm fires at T+47 minutes. drift_direction = credential_theft (0.72, T1003).'],
        ['Lateral Movement (RDP/SMB)', 'PROVES micro-segmentation blocks cross-zone SMB. Detects SMB shadow rules. Validates no allow-any exists.', 'Network signal: unique_dest_ips jumps 3→18. Cohort z-score = 6.17 (>3.0 threshold). CRITICAL alert. Trust state → RESTRICTED.'],
        ['Persistence (web shells)', 'Config Drift detects baseline deviation within 1 hour of web shell deployment or scheduled task creation.', 'Identity signal: service_account_usage spikes. drift_direction = privilege_escalation (0.68, T1078). ABAC: step-up MFA enforced.'],
        ['Pre-positioning (C2)', 'PROVES egress filtering blocks known C2 infrastructure IPs across ALL zones. Validates no bypass path exists.', 'File signal: sensitive_access rises. data_exfiltration concept = 0.44. Combined with Network beacon_score increase. CRITICAL → ABAC BLOCKED.'],
    ]
)

doc.add_heading('6.2 Stopping Salt Typhoon — Phase by Phase', level=2)

add_table(
    ['Kill-Chain Phase', '22CT Preemptive Response', 'ACECARD Response'],
    [
        ['Edge CVE Exploit (IOS XE)', 'PROVES management-plane ACLs restrict /webui access to authorized admin IPs only. WAF verified blocking public access.', 'Network signal: anomalous management-plane source IP detected. drift = 0.09. Entity: router management interface.'],
        ['Account Creation (config)', 'Config Drift re-verifies every hour. Detects new local account creation, ACL whitelist additions, Guest Shell enablement. Alerts within 1 hour.', 'Identity signal: graph hash changes (new SID linked to router). drift_direction = privilege_escalation. Trust floor lowered for segment.'],
        ['TACACS+ Capture (TCP/49)', 'PROVES packet capture permissions restricted to authorized roles via management VRF segmentation. Unauthorized capture blocked.', 'Process signal: tcpdump/packet-capture outside change management window. CUSUM alarm on router entity.'],
        ['JumbledPath Pivot', 'PROVES outbound connections to JumbledPath C2 infrastructure blocked across all internal zones.', 'Composite signal: same entity_uuid appears across 4 router hops (entity fusion). Co-drift detected across compromised devices.'],
        ['GRE Tunnel + Exfiltration', 'PROVES GRE peer authentication required. CALEA-zone segmentation verified. Unauthorized GRE tunnel establishment blocked.', 'data_exfiltration concept = 0.44 on router entity. Slow drift caught by CUSUM (monthly accumulation). CRITICAL alert.'],
    ]
)

doc.add_heading('6.3 Four Integration Points', level=2)

doc.add_paragraph(
    'The two layers are not simply stacked — they actively feed each other through four integration points:')

add_para('Integration 1: Coverage Map → Sensitivity', bold=True)
doc.add_paragraph(
    'Preemptive publishes a MITRE coverage map: which techniques are config-blocked (with proof). '
    'ACECARD uses this to RAISE detection thresholds on covered techniques (less noise from things that '
    'can\'t happen) and LOWER thresholds on uncovered techniques (more sensitivity where gaps exist). '
    'Result: alert volume drops 30-50% on mitigated TTPs without sacrificing detection of unmitigated ones.')

add_para('Integration 2: Gaps → Trust Floor', bold=True)
doc.add_paragraph(
    'When Preemptive finds an unmitigated control gap, it signals ACECARD to lower the trust floor for '
    'all entities in the affected network segment. Those entities move to ELEVATED_WATCH until the gap is closed. '
    'Step-up MFA and tighter logging are applied automatically through the ABAC integration.')

add_para('Integration 3: Confirmed Alerts → Threat Feed', bold=True)
doc.add_paragraph(
    'When a SOC analyst confirms an ACECARD alert as a true positive, the confirmed finding — including '
    'MITRE technique IDs and entity context — feeds into Preemptive\'s Attack Intelligence. Field-observed '
    'TTPs drive preemptive policy changes. This closes the detection-to-prevention loop: what ACECARD detects '
    'today, Preemptive prevents tomorrow.')

add_para('Integration 4: Shared Vocabulary (MITRE IDs + Entity UUIDs)', bold=True)
doc.add_paragraph(
    'Both layers use MITRE ATT&CK technique IDs as their common language. ACECARD outputs "drift toward T1021" '
    'and Preemptive outputs "T1021 blocked in zones A, B; unmitigated in zone C." Entity UUIDs are the join key '
    'for SOAR playbooks. A single alert contains: what happened (MITRE ID), which entity (UUID), how confident '
    '(CUSUM statistic + concept alignment score), and is it config-covered (Preemptive proof).')

doc.add_heading('6.4 Coverage Matrix', level=2)

add_table(
    ['Adversary Capability', 'Owner', 'Why This Layer'],
    [
        ['Edge-device CVE exploitation', 'PREEMPTIVE', 'Formal proof: patch enforced or compensating control verified'],
        ['Misconfigured / shadowed FW rules', 'PREEMPTIVE', 'Math model finds shadows across multi-vendor config'],
        ['Router ACL tampering / config drift', 'PREEMPTIVE', 'Hourly re-verification against approved baseline'],
        ['Compromised valid accounts', 'ACECARD', 'Policy-allowed; only behavioral drift reveals abuse'],
        ['Living-off-the-Land binaries', 'ACECARD', 'lotl_execution concept + process signal structural shift'],
        ['Cross-domain lateral movement', 'ACECARD', 'Entity fusion joins all identities into one trajectory'],
        ['Slow APT dwell (months/years)', 'ACECARD', 'CUSUM accumulates sub-threshold drift; alarm in hours'],
        ['C2 to known malicious infra', 'BOTH', 'Preemptive blocks known IPs; ACECARD catches unknown via beacon_score'],
        ['Zero-day / unmapped TTPs', 'ACECARD', 'Behavioral anomaly detected even with no existing MITRE label'],
        ['Insider threat / data hoarding', 'ACECARD', 'insider_hoarding concept (T1074) + cohort z-score outlier'],
    ]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 7. ARCHITECTURE & DEPLOYMENT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Architecture & Deployment', level=1)

doc.add_heading('7.1 Five-Layer Architecture', level=2)

add_table(
    ['Layer', 'Name', 'Components', 'Purpose'],
    [
        ['L5', 'Presentation', 'Plotly Dash / React UI, 7-chart dashboard, alert feed, kill-chain visualization, ABAC admin console, CAC/PIV authentication', 'Analyst interface'],
        ['L4', 'API + Integration', 'FastAPI (async), REST + WebSocket, Istio mTLS, SOAR webhooks, ServiceNow/JIRA, RBAC', 'External integration'],
        ['L3', 'Behavioral Engine', '5 Signal Summarizers, Embedding Pipeline, Trajectory Analysis, CUSUM, Drift Direction, Entity Fusion, ABAC', 'Core detection'],
        ['L2', 'Data Plane', 'PostgreSQL + pgvector (HNSW indexes), TimescaleDB (time-series), Redis (cache + pub/sub), S3 (object store), AES-256 encryption at rest', 'Storage + search'],
        ['L1', 'Infrastructure', 'Kubernetes (Big Bang / RKE2), Iron Bank containers, Istio service mesh, Fluentd logging, Prometheus + Grafana monitoring, FIPS 140-2/140-3', 'Platform'],
    ]
)

doc.add_heading('7.2 Data Flow Pipeline', level=2)

doc.add_paragraph(
    'The complete data flow from raw telemetry to actionable alert, entirely within the enclave:')

add_para('Step 1: INGEST', bold=True)
doc.add_paragraph(
    'Raw security events arrive via Kafka or Fluentd from SIEM, EDR, NetFlow, IdP, and K8s audit logs. '
    'Events are normalized to ECS (Elastic Common Schema), OCSF (Open Cybersecurity Schema Framework), '
    'or STIX 2.1 (Structured Threat Information Expression) format. No raw event content is stored — '
    'only normalized, structured fields.')

add_para('Step 2: WINDOW', bold=True)
doc.add_paragraph(
    'Events are aggregated into 1-hour windows per entity_uuid. Entity fusion (graph-based identity resolution) '
    'maps all identifiers to a single canonical entity. Example: 47 authentication events for user U in hour H '
    'become a single set of metrics: logon_count=47, failure_rate=0.04, etc.')

add_para('Step 3: SUMMARIZE', bold=True)
doc.add_paragraph(
    'Five signal summarizers produce structured text descriptions of each signal for the time window. '
    'The summarized text contains metrics and context only — no raw event payloads, no PII, no credentials. '
    'Example: "AUTH signal: logon_count=47, unique_hosts=3, off_hours=0.0. Role: domain_admin."')

add_para('Step 4: EMBED', bold=True)
doc.add_paragraph(
    'Concatenated signal text is passed through the embedding model to produce 1536-d vectors. '
    'Default: local embedding model (E5-Mistral-7B-Instruct or nomic-embed-text-v1.5) running on GPU. '
    'IL5 option: Azure OpenAI GovCloud text-embedding-3-small endpoint. '
    'Vectors are stored in PostgreSQL with pgvector extension using HNSW (Hierarchical Navigable Small World) '
    'indexes for efficient nearest-neighbor search.')

add_para('Step 5: ANALYZE', bold=True)
doc.add_paragraph(
    'Trajectory analysis computes: cosine drift from previous window, velocity (rate of drift change), '
    'acceleration (rate of velocity change). CUSUM runs continuously, accumulating excess drift. '
    'Drift direction projects the drift vector onto 8 threat concepts via cosine similarity. '
    'Cohort z-scores computed against peer entities.')

add_para('Step 6: ALERT + ENFORCE', bold=True)
doc.add_paragraph(
    'Alert generation: CUSUM alarm fires → alert created with MITRE IDs, concept alignment scores, '
    'cohort z-score, confidence level. Alert tiers: CRITICAL (immediate ABAC enforcement), HIGH (analyst '
    'queue), MEDIUM (enrichment), LOW/INFO (logging only). ABAC trust state updated. WebSocket push to '
    'SOC dashboard. SOAR webhook for automated response playbooks.')

add_para('Performance:', bold=True)
add_bullet('<3 seconds per entity (API-backed embedding model)')
add_bullet('10,000 entities/hour per 4-vCPU pod (API-backed)')
add_bullet('~2,000 entities/hour per pod (local CPU embedding model)')
add_bullet('Horizontal autoscale: add pods to match entity count')
add_bullet('pgvector HNSW: <10ms nearest-neighbor search across millions of vectors')

doc.add_heading('7.3 Gabriel Nimbus Deployment', level=2)

doc.add_paragraph(
    'Gabriel Nimbus is ARCYBER\'s classified big data analytics platform. ACECARD + Preemptive are designed '
    'for native deployment within this environment:')

add_para('Deployment Model:', bold=True)
add_bullet('Containerized: all components as OCI containers from Iron Bank (DoD hardened container registry)')
add_bullet('Orchestration: Kubernetes via Big Bang (Platform One\'s K8s distribution) or RKE2 (Rancher)')
add_bullet('Service mesh: Istio with mTLS (mutual TLS) for all inter-service communication')
add_bullet('Deployment artifact: single Helm chart with environment-specific values files')
add_bullet('Air-gapped: identical containers, local embedding model, zero external API calls')

add_para('Classification Levels:', bold=True)
add_bullet('IL5 (Impact Level 5): Controlled Unclassified Information (CUI). Azure GovCloud option for embedding API.')
add_bullet('IL6 (Impact Level 6): Secret. Fully air-gapped. Local embedding model only.')
add_bullet('JWICS (Joint Worldwide Intelligence Communications System): TS/SCI. Fully air-gapped. Isolated network.')

doc.add_heading('7.4 Air-Gapped Operations', level=2)

doc.add_paragraph(
    'For IL6/JWICS deployments, the system operates with zero external dependencies:')

add_bullet('Embedding model: E5-Mistral-7B-Instruct or nomic-embed-text-v1.5 running locally on GPU (NVIDIA A100/H100)')
add_bullet('Threat intelligence: manual update via STIX 2.1 bundles imported at classification boundary')
add_bullet('Container updates: Iron Bank pipeline, security scanning, manual transport across air gap')
add_bullet('No API calls to external services. No telemetry. No cloud dependencies.')
add_bullet('FIPS 140-2 Level 1 (software) / FIPS 140-3 (hardware HSM for key storage)')

add_para('Embedding Model Performance (Air-Gapped):', bold=True)
add_bullet('E5-Mistral-7B on A100 GPU: ~500 embeddings/second')
add_bullet('nomic-embed-text on 4-vCPU: ~50 embeddings/second (CPU only)')
add_bullet('Batched processing: entities queued and processed in batches of 100')
add_bullet('Throughput: 2,000-10,000 entities/hour depending on hardware')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 8. COMPLIANCE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('8. DoD Compliance & Accreditation', level=1)

doc.add_paragraph(
    'The system is designed from the ground up for DoD compliance, not retrofitted after development:')

add_para('NIST 800-53 Rev 5 Controls:', bold=True)
add_bullet('SI-4 (System Monitoring): continuous behavioral monitoring with automated alerting')
add_bullet('AU-6 (Audit Record Review): automated analysis of audit records via embedding pipeline')
add_bullet('IR-4 (Incident Handling): SOAR integration for automated response playbooks')
add_bullet('SC-7 (Boundary Protection): formal verification of boundary device configurations')
add_bullet('IA-5 (Authenticator Management): behavioral analysis of authentication patterns')
add_bullet('AC-2 (Account Management): identity fusion detects unauthorized account usage')

add_para('DoD Zero Trust Reference Architecture:', bold=True)
add_bullet('Pillar 4 (Data): behavioral analysis of data access patterns, DLP integration, encryption at rest')
add_bullet('Pillar 7 (Visibility & Analytics): comprehensive behavioral visibility across all entity types')

add_para('RMF (Risk Management Framework) / eMASS:', bold=True)
add_bullet('System Security Plan (SSP) in OSCAL format (machine-readable)')
add_bullet('Control Implementation Statements auto-generated from system configuration')
add_bullet('Continuous monitoring feeds directly into eMASS for real-time risk posture')
add_bullet('POA&M (Plan of Action & Milestones) auto-generation for identified gaps')

add_para('cATO (Continuous Authority to Operate):', bold=True)
doc.add_paragraph(
    'The cATO pathway enables continuous authorization rather than periodic 3-year ATO cycles. '
    'ACECARD supports this through: continuous control monitoring, automated evidence collection, '
    'real-time risk scoring, and direct eMASS data feeds. Preemptive\'s formal verification provides '
    'continuous proof of control effectiveness — the strongest evidence possible for cATO.')

add_para('Additional Compliance:', bold=True)
add_bullet('DISA STIGs: automated STIG compliance scanning of all containers')
add_bullet('Iron Bank: all containers sourced from DoD hardened container registry')
add_bullet('FIPS 140-2/140-3: cryptographic modules validated for classified use')
add_bullet('DCO-IDM (Defensive Cyberspace Operations - Internal Defensive Measures): alert format compatible')
add_bullet('CAC/PIV: Common Access Card / Personal Identity Verification for user authentication')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 9. TIMELINE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Timeline: 90 Days to First Detection', level=1)

doc.add_paragraph(
    'The deployment timeline assumes: AO (Authorizing Official) approved pilot authorization OR existing '
    'cATO reciprocity that allows container deployment into Gabriel Nimbus without a new full ATO.')

add_para('Phase 1: Foundation (Weeks 1-4)', bold=True)
add_bullet('Deploy K8s containers (from Iron Bank) into Gabriel Nimbus environment')
add_bullet('Establish data connectors: SIEM, EDR, NetFlow, IdP, K8s audit via ECS/OCSF')
add_bullet('Configure entity fusion rules: map AD, AAD, Okta, AWS, CrowdStrike identifiers')
add_bullet('Initial baseline computation: 2-4 weeks of data establishes normal behavioral patterns')
add_bullet('DELIVERABLE: Data flowing, entities resolved, baselines computing')

add_para('Phase 2: Detection (Weeks 5-8)', bold=True)
add_bullet('Enable 5-signal embedding pipeline with local embedding model')
add_bullet('Activate trajectory analysis: cosine drift, velocity, acceleration')
add_bullet('Enable CUSUM (4-sigma threshold, calibrated from Phase 1 baselines)')
add_bullet('Activate drift direction analysis onto 8 MITRE threat concepts')
add_bullet('Configure cohort definitions from AD/LDAP group membership')
add_bullet('DELIVERABLE: Alerts firing, drift direction operational, dashboard live')

add_para('Phase 3: Operations (Weeks 9-12)', bold=True)
add_bullet('SOC integration: SOAR webhooks, ServiceNow/JIRA ticket creation')
add_bullet('Threshold tuning based on analyst TP/FP feedback from first detection period')
add_bullet('Enable ABAC trust loop: trust state updates, conditional access enforcement')
add_bullet('Replay Volt Typhoon and Salt Typhoon scenarios (synthetic data injection) for validation')
add_bullet('Red team exercise: validate detection of known TTPs')
add_bullet('DELIVERABLE: Full UEBA capability, ABAC enforcing, validated against known adversary TTPs')

add_para('Phase 4: Preemptive Integration (Week 13+)', bold=True)
add_bullet('Connect 22CT Preemptive to edge defense estate (NGFW, IPS, SASE, WAF)')
add_bullet('Initial formal verification scan of all defensive configurations')
add_bullet('Establish coverage map → sensitivity coupling with ACECARD')
add_bullet('Enable trust-floor coupling: Preemptive gaps → ACECARD elevated watch')
add_bullet('Activate detection-to-prevention loop: confirmed alerts → Preemptive threat feed')
add_bullet('DELIVERABLE: Both layers operational, closed-loop defense, continuous re-verification')

add_para('Risk Factors:', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
add_bullet('ATO timeline: mitigated by cATO pathway + existing reciprocity agreements')
add_bullet('Data quality from legacy SIEM: mitigated by 2-week connector validation sprint (Phase 1)')
add_bullet('Baseline period: minimum 2 weeks of stable data needed; longer is better for seasonal patterns')
add_bullet('Red team validation: planned within 60 days of deployment to confirm detection capability')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 10. GLOSSARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Key Technical Terms Glossary', level=1)

terms = [
    ('ABAC', 'Attribute-Based Access Control. Access decisions based on attributes of the user, resource, action, and environment — more granular than RBAC (Role-Based).'),
    ('ACECARD', 'Adaptive Cyber Entity Characterization, Assessment & Risk Detection. The UEBA system described in this document.'),
    ('Air-Gapped', 'Network isolation with no connectivity to external networks. Required for IL6/JWICS classified environments.'),
    ('ARL0', 'Average Run Length to false alarm. Key metric for sequential detection: how long between false positives under normal conditions.'),
    ('ASM', 'Attack Surface Management. Tools that enumerate an organization\'s external-facing assets and vulnerabilities.'),
    ('ATO', 'Authority to Operate. DoD accreditation that a system may process classified data at a given impact level.'),
    ('BAS', 'Breach and Attack Simulation. Tools that automatically simulate attacks against production defenses.'),
    ('Beacon Score', 'Statistical measure of periodic external communication — high scores indicate potential C2 callback patterns.'),
    ('Big Bang', 'Platform One\'s pre-configured Kubernetes distribution with built-in security (Istio, Kiali, Jaeger, etc.).'),
    ('C2', 'Command and Control. Communication channel between adversary infrastructure and compromised hosts.'),
    ('CAC/PIV', 'Common Access Card / Personal Identity Verification. DoD smart card for identity authentication.'),
    ('CALEA', 'Communications Assistance for Law Enforcement Act. Requires telecoms to enable lawful-intercept capability.'),
    ('cATO', 'Continuous Authority to Operate. Modern DoD accreditation approach that replaces periodic 3-year ATO with continuous monitoring.'),
    ('CDR', 'Call Detail Records. Metadata about telecommunications: who called whom, when, duration, cell tower.'),
    ('Cohort Z-Score', 'How many standard deviations an entity\'s drift deviates from its peer group average.'),
    ('Cosine Similarity', 'Measure of angular distance between two vectors. Range [-1, 1]. Used to compare embeddings.'),
    ('CUSUM', 'Cumulative Sum. Sequential detection algorithm (Page 1954) that accumulates small deviations to detect regime change.'),
    ('DCO-IDM', 'Defensive Cyberspace Operations - Internal Defensive Measures. DoD network defense operations.'),
    ('DISA STIGs', 'Defense Information Systems Agency Security Technical Implementation Guides. Hardening standards for DoD systems.'),
    ('Drift Direction', 'The angular direction of behavioral change projected onto known threat concept vectors.'),
    ('Drift Vector', 'embedding(t) - embedding(t-1). The direction and magnitude of behavioral change between time windows.'),
    ('ECS', 'Elastic Common Schema. Standard field naming for security events (Elastic/OpenSearch ecosystem).'),
    ('EDR', 'Endpoint Detection and Response. Security tools that monitor endpoint processes and file activity.'),
    ('eMASS', 'Enterprise Mission Assurance Support Service. DoD system for tracking RMF accreditation status.'),
    ('Embedding', 'A vector representation of data in high-dimensional space where similarity = proximity.'),
    ('Entity Fusion', 'Resolving multiple identifiers across systems into a single canonical entity_uuid.'),
    ('FIPS 140-2/3', 'Federal Information Processing Standard for cryptographic module validation.'),
    ('Gabriel Nimbus', 'ARCYBER\'s classified big data analytics platform on Kubernetes.'),
    ('GRE', 'Generic Routing Encapsulation. Tunneling protocol used by Salt Typhoon for data exfiltration.'),
    ('HNSW', 'Hierarchical Navigable Small World. Algorithm for approximate nearest-neighbor search in vector databases.'),
    ('i.i.d.', 'Independent and Identically Distributed. Statistical assumption for CUSUM optimality proof.'),
    ('IL5/IL6', 'Impact Level 5 (CUI) / Impact Level 6 (Secret). DoD cloud classification levels.'),
    ('Iron Bank', 'DoD hardened container registry (Platform One). All containers scanned and approved for DoD use.'),
    ('Istio', 'Service mesh providing mTLS, traffic management, and observability for Kubernetes.'),
    ('JWICS', 'Joint Worldwide Intelligence Communications System. TS/SCI classified network.'),
    ('Kill Chain', 'Sequential phases of a cyber attack from initial access to objective completion.'),
    ('KV Botnet', 'Network of compromised SOHO routers used by Volt Typhoon as C2 relay infrastructure.'),
    ('LOLBin', 'Living-off-the-Land Binary. Legitimate system tool abused for malicious purposes (ntdsutil, wmic, etc.).'),
    ('LOTL', 'Living-off-the-Land. Adversary tradecraft using only built-in tools — no custom malware deployed.'),
    ('MITRE ATT&CK', 'Knowledge base of adversary tactics, techniques, and procedures. Standard taxonomy for cyber threats.'),
    ('mTLS', 'Mutual TLS. Both client and server authenticate via certificates — standard for zero trust service mesh.'),
    ('NIST 800-53', 'Security and Privacy Controls for Information Systems. The standard DoD uses for control selection.'),
    ('NTDS.dit', 'Active Directory database file containing all domain user credentials (password hashes).'),
    ('OCSF', 'Open Cybersecurity Schema Framework. Vendor-neutral schema for security events (AWS/Splunk-led).'),
    ('OSCAL', 'Open Security Controls Assessment Language. Machine-readable format for security documentation.'),
    ('pgvector', 'PostgreSQL extension for vector similarity search. Supports HNSW and IVFFlat indexes.'),
    ('POA&M', 'Plan of Action & Milestones. Document tracking known security weaknesses and remediation plans.'),
    ('RKE2', 'Rancher Kubernetes Engine 2. FIPS-compliant K8s distribution approved for DoD use.'),
    ('RMF', 'Risk Management Framework. DoD/NIST process for authorizing information systems.'),
    ('SAT/SMT', 'Satisfiability / Satisfiability Modulo Theories. Mathematical solvers used in formal verification.'),
    ('SIEM', 'Security Information and Event Management. Central log collection and correlation platform.'),
    ('SOAR', 'Security Orchestration, Automation and Response. Platforms that automate security workflows.'),
    ('SSP', 'System Security Plan. Document describing system architecture, controls, and compliance posture.'),
    ('STIX 2.1', 'Structured Threat Information Expression. Standard format for sharing cyber threat intelligence.'),
    ('TACACS+', 'Terminal Access Controller Access-Control System Plus. AAA protocol for network devices (TCP/49).'),
    ('Trajectory', 'The path an entity\'s behavioral embedding traces through vector space over time.'),
    ('UEBA', 'User and Entity Behavior Analytics. Technology that builds behavioral baselines and detects deviations.'),
    ('XDR', 'Extended Detection and Response. Cross-domain EDR that correlates across endpoint, network, and cloud.'),
    ('Zero Trust', 'Security model: never trust, always verify. Every access request authenticated and authorized regardless of network location.'),
]

for term, definition in terms:
    p = doc.add_paragraph()
    run = p.add_run(f'{term}: ')
    run.bold = True
    p.add_run(definition)
    p.paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════════════════════════════
# FINAL: 22CT Qualifications (back matter)
# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Appendix: 22nd Century Technologies Qualifications', level=1)

doc.add_paragraph(
    '22nd Century Technologies Inc. (TSCTI) is a prime federal IT services contractor with deep '
    'cybersecurity experience across DoD, IC, and civilian agencies.')

add_para('Key Contracts:', bold=True)
add_bullet('$90M Army SOC/MDR: 800+ cleared analysts, 24x7x365 security operations, proven incident response at scale')
add_bullet('USAF $108M: cybersecurity services for Air Force networks')
add_bullet('FBI TSC $56M: Terrorist Screening Center technology services')
add_bullet('NAVAIR $145M: Naval Air Systems Command IT and cybersecurity')

add_para('Cleared Workforce:', bold=True)
add_bullet('Secret to TS/SCI cleared personnel')
add_bullet('MITRE ATT&CK certified')
add_bullet('NIST 800-53 / DoD RMF trained')
add_bullet('Day-1 ready for classified environment operations')

add_para('Innovation Partner — Rigor AI:', bold=True)
doc.add_paragraph(
    'The Preemptive layer\'s formal verification technology is powered by Rigor AI, 22CT\'s innovation partner. '
    'Rigor AI provides the Symbolic Model of Computation engine, the three Intelligences framework, and the '
    'formal verification algorithms. 22CT provides the DoD integration, compliance expertise, cleared workforce, '
    'and Gabriel Nimbus deployment capability.')

# ─── SAVE ───
doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Sections: 10 + Glossary + Appendix")
