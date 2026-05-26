"""Generate comprehensive Word document expanding every slide of the
Rigor AI + ACECARD UEBA combined deck (32 slides).

Business-friendly language with all acronyms expanded."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_key_term(doc, term, definition):
    p = doc.add_paragraph()
    run = p.add_run(term + ": ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x55, 0x99)
    p.add_run(definition)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(" -- " + text)
    else:
        p.add_run(text)


def add_impact_box(doc, title, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    p.add_run(text)


def add_insight_box(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("KEY INSIGHT: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x77, 0x00)
    p.add_run(text)


def add_analogy(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("ANALOGY: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x55, 0x44, 0x99)
    p.add_run(text)


def slide_heading(doc, num, title):
    doc.add_heading(f"Slide {num}: {title}", level=1)


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        hs.font.name = "Calibri"

    # =========================================================
    # TITLE PAGE
    # =========================================================
    for _ in range(3):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Rigor AI + ACECARD UEBA")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Combined Cyber Defense Architecture")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "\nStopping Nation-State Threats (Volt Typhoon & Salt Typhoon)\n"
        "Through Mathematical Verification + Behavioral Intelligence"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "\n\nA Complete Business-Friendly Explanation of Every Concept\n"
        "All Acronyms Expanded -- No Prior Technical Knowledge Required"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x55, 0x99)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\nMay 2026")

    doc.add_page_break()

    # =========================================================
    # TABLE OF CONTENTS
    # =========================================================
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "Slide 1: Title and Overview",
        "Slide 2: The Threat -- Volt Typhoon and Salt Typhoon",
        "Slide 3: Why Current Defenses Fail (8 Structural Gaps)",
        "Slide 4: The Two-Layer Answer (Rigor + ACECARD)",
        "Slides 5-6: Rigor AI -- Approach and Operational Flow",
        "Slides 7-8: How Rigor Stops Volt/Salt Typhoon Phase-by-Phase",
        "Slides 9-10: TTP-by-TTP Coverage Tables",
        "Slide 11: Nine Rigor Use Cases",
        "Slide 12: Where Rigor Needs Help (Honest Assessment)",
        "Slide 13: Rigor Proof Points",
        "Slide 14: Why Current UEBA Fails",
        "Slide 15: Three ACECARD Innovations",
        "Slide 16: The Math of Behavioral Drift",
        "Slide 17: Five Behavioral Signals",
        "Slide 18: Pipeline and Trajectory Analysis",
        "Slides 19-20: How ACECARD Detects Volt/Salt Typhoon Signal-by-Signal",
        "Slide 21: Drift Direction with 8 MITRE Concepts",
        "Slide 22: Cohort Comparison and Alert Tiers",
        "Slide 23: ABAC Trust Loop",
        "Slides 24-25: Combined Architecture and Data Flow",
        "Slides 26-27: Attack Walkthroughs (Both Layers)",
        "Slide 28: Combined Coverage Matrix",
        "Slide 29: Complementary Tools",
        "Slide 30: Reference Architecture",
        "Slide 31: Verdict",
        "Slide 32: Next Steps",
        "Appendix A: Complete Acronym Glossary",
        "Appendix B: MITRE ATT&CK Framework Primer",
    ]
    for t in toc:
        p = doc.add_paragraph(t)
        if t.startswith("Slide") or t.startswith("Appendix"):
            p.runs[0].bold = True

    doc.add_page_break()

    # =========================================================
    # SLIDE 1
    # =========================================================
    slide_heading(doc, 1, "Title and Overview")

    doc.add_heading("What This Deck Is About", level=2)
    doc.add_paragraph(
        "This presentation describes a two-layer cyber defense system designed to "
        "stop the most sophisticated nation-state hackers currently targeting U.S. "
        "critical infrastructure and military networks."
    )

    doc.add_heading("The Two Components", level=2)
    add_key_term(
        doc,
        "Rigor AI (Layer 1 -- Preemptive/Preventive)",
        "A formal mathematical model that exhaustively reasons over the ENTIRE "
        "configuration state space of your network's security controls. It does NOT "
        "sample or simulate -- it mathematically PROVES whether a given attack path "
        "can or cannot succeed given your current firewall rules, access controls, and "
        "network segmentation. Think of it as a mathematician who checks EVERY "
        "possible combination of locked and unlocked doors in your building, not just "
        "the ones a tester happened to try.",
    )
    add_key_term(
        doc,
        "ACECARD UEBA (Layer 2 -- Detective/Behavioral)",
        "ACECARD stands for 'Anomalous Cyber Entity Characterization And Rapid "
        "Detection.' UEBA stands for 'User and Entity Behavior Analytics.' This layer "
        "tracks the behavioral profile of every user, device, and network segment as "
        "a mathematical fingerprint (a vector of 1,536 numbers). When behavior slowly "
        "drifts toward patterns matching known attack techniques, it detects the "
        "drift -- even when every individual action looks legitimate.",
    )

    doc.add_heading("Key Terms on This Slide", level=2)
    add_key_term(
        doc,
        "Gabriel Nimbus",
        "The U.S. Department of Defense's cloud computing environment. Think of it as "
        "the military's version of Amazon Web Services (AWS) or Microsoft Azure.",
    )
    add_key_term(
        doc,
        "CFIC",
        "Cyber Futures Innovation Center -- the Army organization responsible for "
        "developing next-generation cyber capabilities.",
    )
    add_key_term(
        doc,
        "ARCYBER",
        "U.S. Army Cyber Command -- the Army organization responsible for conducting "
        "cyber operations.",
    )
    add_key_term(
        doc,
        "ArCTIC",
        "Army Cyber Technology and Innovation Center.",
    )
    add_key_term(
        doc,
        "DLA Entity Digital Model (E-11)",
        "A proven AI system built for the Defense Logistics Agency that tracks "
        "supply chain entities (parts, suppliers) as behavioral vectors. The ACECARD "
        "system applies this same proven architecture to cybersecurity.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 2
    # =========================================================
    slide_heading(doc, 2, "The Threat -- Volt Typhoon and Salt Typhoon")

    doc.add_heading("Who Are These Groups?", level=2)
    doc.add_paragraph(
        "These are two separate hacking groups operated by the People's Republic of "
        "China (PRC). They are not criminals seeking money -- they are government "
        "intelligence operatives with virtually unlimited resources, operating under "
        "state direction."
    )

    doc.add_heading("Volt Typhoon", level=2)
    add_bullet(doc, "Chinese state-sponsored hacking group (also tracked as 'Bronze Silhouette' and 'Vanguard Panda')")
    add_bullet(doc, "Targets: U.S. critical infrastructure -- power plants, water treatment, telecom, transportation, military logistics")
    add_bullet(doc, "Goal: Pre-position inside these systems for potential disruption during a future conflict (e.g., over Taiwan)")
    add_bullet(doc, "Dwell time: Over 5 YEARS inside some targets before discovery")
    add_bullet(doc, "Signature technique: Living-off-the-Land (LOtL) -- uses only legitimate tools already installed on victim systems")
    add_bullet(doc, "C2 method: Routes commands through thousands of compromised home routers (KV Botnet)")
    add_bullet(doc, "Reference: CISA Advisory AA24-038A (February 2024)")

    doc.add_heading("Salt Typhoon", level=2)
    add_bullet(doc, "Separate Chinese state-sponsored group (also tracked as 'Earth Estries' and 'FamousSparrow')")
    add_bullet(doc, "Targets: Telecommunications carriers -- AT&T, Verizon, T-Mobile, Lumen Technologies")
    add_bullet(doc, "Goal: Access to call records, text messages, and law enforcement wiretap (CALEA) systems")
    add_bullet(doc, "Technique: Exploits network infrastructure devices (routers, switches) rather than endpoints")
    add_bullet(doc, "Custom malware: GhostSpider (backdoor), Demodex (rootkit), SnappyBee (RAT)")
    add_bullet(doc, "Pivot capability: Once inside a carrier, access all their enterprise customers")
    add_bullet(doc, "Reference: CISA Advisory AA25-239A")

    doc.add_heading("Key Terms", level=2)
    add_key_term(
        doc, "APT (Advanced Persistent Threat)",
        "A prolonged, targeted attack by a well-resourced adversary (usually a nation-state). "
        "'Advanced' = sophisticated tools. 'Persistent' = they stay for months or years. "
        "'Threat' = specific objectives (espionage, pre-positioning for disruption).",
    )
    add_key_term(
        doc, "TTPs (Tactics, Techniques, and Procedures)",
        "The three-level framework describing HOW an attacker operates. Tactics = the strategic "
        "goal (e.g., 'gain initial access'). Techniques = the method (e.g., 'exploit VPN flaw'). "
        "Procedures = the exact steps (e.g., 'send crafted HTTP request to /api/endpoint'). The "
        "MITRE ATT&CK framework catalogs hundreds of observed TTPs.",
    )
    add_key_term(
        doc, "CISA",
        "Cybersecurity and Infrastructure Security Agency -- the U.S. government agency "
        "responsible for protecting critical infrastructure from cyber threats. They publish "
        "advisories (AA = Activity Alert) when major threats are discovered.",
    )
    add_key_term(
        doc, "LOtL (Living off the Land)",
        "Attack technique where hackers use only tools already installed on victim systems "
        "(PowerShell, WMI, certutil, netsh). No malware to detect -- every tool used is "
        "a legitimate Microsoft administration utility.",
    )
    add_key_term(
        doc, "C2 (Command and Control)",
        "The communication channel between the attacker and their compromised systems. "
        "Used to send commands and receive stolen data. KV Botnet uses thousands of "
        "compromised home routers as relay points to hide this traffic.",
    )
    add_key_term(
        doc, "CALEA (Communications Assistance for Law Enforcement Act)",
        "U.S. law requiring telecom carriers to build wiretap capabilities into their "
        "networks for FBI and law enforcement use. Salt Typhoon gained access to these "
        "systems -- a foreign government accessed America's own surveillance infrastructure.",
    )
    add_key_term(
        doc, "Dwell Time",
        "How long an attacker remains inside a network before being discovered. "
        "Volt Typhoon's 5+ year dwell time is extraordinarily long and indicates "
        "their operations were essentially invisible to existing security tools.",
    )

    add_impact_box(
        doc, "WHY THIS MATTERS",
        "These are not hypothetical threats. They have already penetrated U.S. critical "
        "infrastructure and telecom carriers. The FBI confirmed Chinese hackers had access "
        "to wiretap systems used by U.S. law enforcement. Traditional security tools "
        "failed to detect them for years.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 3
    # =========================================================
    slide_heading(doc, 3, "Why Current Defenses Fail Against These Adversaries")

    doc.add_paragraph(
        "This slide explains WHY existing cybersecurity tools -- which organizations "
        "have spent billions of dollars deploying -- fundamentally cannot stop these "
        "threats. There are TWO structural gap categories: one in preemption (preventing "
        "attacks) and one in detection (finding attacks in progress). Traditional tooling "
        "cannot close either gap."
    )

    doc.add_heading("PREEMPTIVE LAYER FAILS", level=2)
    doc.add_paragraph(
        "These gaps explain why attacks succeed at getting in and moving through "
        "your network even when preventive tools are deployed."
    )

    add_key_term(
        doc, "99% of Firewall Breaches Are Misconfiguration (Gartner)",
        "Gartner predicted that through 2023, 99% of firewall breaches would be caused by "
        "firewall misconfigurations, not flaws in the firewall itself. Both Typhoon groups "
        "exploit unpatched, misconfigured NGFWs and VPNs. The firewall works fine -- it "
        "was just set up wrong. Source: Gartner Technology Insight for Network Security "
        "Policy Management, February 2019.",
    )
    add_key_term(
        doc, "KEV Patching Covers Only ~20% of Attacks",
        "The CISA KEV (Known Exploited Vulnerabilities) catalog -- which federal agencies "
        "are mandated to patch -- covers only about 20% of attack vectors. The other 80% "
        "(most Salt Typhoon paths) have no CVE patch because they exploit configuration "
        "weaknesses, not software bugs.",
    )
    add_key_term(
        doc, "ASM/BAS Only Sample the 2^8000 State Space",
        "ASM (Attack Surface Management) and BAS (Breach and Attack Simulation) tools "
        "test your defenses by running simulated attacks. But a modern enterprise network "
        "has a configuration state space of approximately 2^8000 possible states. This is "
        "an astronomically large number -- larger than the number of atoms in the observable "
        "universe (which is only about 2^266). These tools sample perhaps thousands or "
        "millions of paths. They CANNOT PROVE protection. Coverage is statistical, not "
        "deterministic. An attacker only needs to find ONE path these tools missed.",
    )
    add_key_term(
        doc, "Configuration Drift Goes Unverified",
        "Salt Typhoon modifies router ACLs to create persistent access. These configuration "
        "changes create new attack paths that did not exist when the last security assessment "
        "was performed. Drift detection in traditional tools (if it exists at all) operates "
        "on syntax (did the config text change?), not semantics (did the security posture "
        "change?).",
    )

    doc.add_heading("BEHAVIORAL LAYER FAILS", level=2)
    doc.add_paragraph(
        "These gaps explain why attacks remain invisible to monitoring and detection "
        "tools even after the attacker is inside your network."
    )

    add_key_term(
        doc, "Gap 5: Threshold Scoring Misses LoTL (Living-off-the-Land)",
        "SIEM and UEBA tools fire alerts only when a metric crosses a fixed threshold "
        "line. Living-off-the-Land NEVER crosses thresholds -- Volt Typhoon's process "
        "tree was 'normal' at EVERY step. The tools they use -- ntdsutil, rundll32, "
        "vssadmin, netsh, wmic -- are legitimate Windows administration tools that IT "
        "teams run every day. Antivirus trusts them because they are signed by Microsoft. "
        "No threshold catches 'a legitimate employee running a legitimate Microsoft tool.' "
        "The individual action is always 'normal.' Only the PATTERN of actions over time "
        "reveals the attack.",
    )

    doc.add_heading("What are these tools?", level=3)
    add_key_term(
        doc, "ntdsutil",
        "A built-in Windows tool for maintaining Active Directory (the master database "
        "of all users and passwords in an organization). IT administrators use it for "
        "database maintenance. Attackers use it to COPY the entire password database. "
        "It is a standard Microsoft utility -- antivirus will never flag it.",
    )
    add_key_term(
        doc, "rundll32",
        "A built-in Windows utility that executes code from DLL (Dynamic Link Library) "
        "files. Legitimate use: running specific Windows functions. Attacker use: loading "
        "malicious code through a trusted Microsoft-signed executable so antivirus "
        "does not intervene.",
    )
    add_key_term(
        doc, "vssadmin",
        "Volume Shadow Copy administrator -- manages backup snapshots of hard drives. "
        "Legitimate use: creating and managing backups. Attacker use: accessing locked "
        "password databases through snapshots, or deleting all backups before ransomware.",
    )
    add_key_term(
        doc, "netsh",
        "Network Shell -- a Windows command-line tool for configuring network settings. "
        "Legitimate use: adjusting firewall rules, network interfaces. Attacker use: "
        "creating hidden port-forwarding tunnels that route traffic from inside the "
        "network to the outside, bypassing all firewall rules.",
    )
    add_key_term(
        doc, "wmic",
        "Windows Management Instrumentation Command-line -- a tool for querying and "
        "managing Windows computers remotely. Legitimate use: IT teams use it to check "
        "system status, install software, run diagnostics on remote computers. Attacker "
        "use: executing commands on other computers across the network without installing "
        "any software.",
    )

    add_key_term(
        doc, "Gap 6: Per-Source Siloed Analytics Fragment Identity",
        "When an attacker uses stolen credentials, they touch multiple identity systems: "
        "AD (Active Directory -- on-premises Windows identity), AAD (Azure Active "
        "Directory -- Microsoft's cloud identity), AWS IAM (Amazon's cloud access "
        "management), Okta (cloud single-sign-on provider), and K8s (Kubernetes -- "
        "container orchestration platform). Each system sees ONE login that looks "
        "'normal.' Lateral movement that involves all five systems reads as 4 weak "
        "signals instead of 1 strong one -- because no current tool FUSES these "
        "identities into a single behavioral trajectory per real-world entity.",
    )

    doc.add_heading("What are these identity systems?", level=3)
    add_key_term(
        doc, "AD (Active Directory)",
        "Microsoft's on-premises directory service. The master database of every user "
        "account, computer, and permission in a Windows network. If an organization uses "
        "Windows, they almost certainly use AD. It is the FIRST thing attackers target "
        "because compromising it means owning every account.",
    )
    add_key_term(
        doc, "AAD (Azure Active Directory, now called Microsoft Entra ID)",
        "Microsoft's cloud-based identity service. The cloud version of Active Directory. "
        "Manages sign-ins for Microsoft 365 (Outlook, Teams, SharePoint), Azure cloud "
        "resources, and thousands of third-party applications. Most large organizations "
        "use BOTH AD (on-prem) and AAD (cloud) -- the attacker who compromises one can "
        "often pivot to the other.",
    )
    add_key_term(
        doc, "AWS IAM (Amazon Web Services Identity and Access Management)",
        "Amazon's system for controlling who can access what in their cloud platform. "
        "Manages API keys, roles, and permissions for all AWS services (servers, databases, "
        "storage). An attacker with stolen AWS credentials can provision resources, access "
        "databases, or exfiltrate data.",
    )
    add_key_term(
        doc, "Okta",
        "A cloud-based single-sign-on (SSO) provider. Employees log into Okta once, then "
        "access all their applications (Salesforce, Slack, AWS, etc.) without re-entering "
        "passwords. Compromising an Okta account = access to everything that account is "
        "connected to. In 2023, Okta itself was breached by attackers.",
    )
    add_key_term(
        doc, "K8s (Kubernetes)",
        "An open-source container orchestration platform. 'Containers' are lightweight "
        "packages of software that run applications. Kubernetes manages thousands of "
        "containers across many servers. Its audit log records who deployed what, who "
        "accessed which service, and who changed configurations. Attackers target K8s "
        "to deploy cryptominers, access secrets, or pivot to other systems.",
    )

    add_key_term(
        doc, "Gap 7: Static Rules Decay Monthly",
        "Detection rules (called Sigma rules) are written by security researchers when "
        "new attack techniques are discovered. These rules are then published openly. "
        "PRC (People's Republic of China) threat actors actively TEST their tools "
        "against published Sigma rules BEFORE launching campaigns. They modify their "
        "techniques until the published rules no longer detect them. This means detection "
        "coverage decays by the month -- rules that worked in January may be useless by "
        "March because the attacker has adapted.",
    )

    doc.add_heading("What are Sigma rules?", level=3)
    add_key_term(
        doc, "Sigma Rules",
        "An open standard for writing detection rules that can be used across different "
        "SIEM platforms (Splunk, Sentinel, QRadar, Elastic). A Sigma rule describes a "
        "specific pattern to look for in logs -- for example, 'alert when ntdsutil.exe "
        "runs on a non-domain-controller.' The problem: these rules are PUBLIC. Attackers "
        "read the same rules defenders use, then modify their behavior to avoid triggering "
        "them. It is an arms race where the attacker always has the advantage because "
        "they can test against the rules before attacking.",
    )
    add_key_term(
        doc, "SIEM (Security Information and Event Management)",
        "A centralized system that collects security logs from across the entire "
        "organization (servers, workstations, firewalls, applications) and searches "
        "them for suspicious patterns. Products include Splunk, Microsoft Sentinel, "
        "IBM QRadar, Elastic Security. The SIEM is only as good as its RULES -- if "
        "no rule matches the attack, the SIEM sees nothing.",
    )

    add_key_term(
        doc, "Gap 8: 'Anomaly Score' Without Direction",
        "When commercial UEBA does flag something, the analyst gets a message like: "
        "'User jsmith has an anomaly score of 87%.' That tells the analyst NOTHING "
        "actionable. 87% anomalous HOW? Is it credential theft? Lateral movement? "
        "Data exfiltration? Or did the user just start a new project? Without a "
        "HYPOTHESIS -- without direction -- the analyst must manually investigate "
        "every alert. There is no MITRE ATT&CK mapping, no concept alignment, no "
        "indication of what KIND of threat this resembles. Triage time explodes. "
        "Alert fatigue sets in. Analysts start ignoring alerts. And the attacker "
        "operates undetected.",
    )

    doc.add_heading("Additional Terms on This Slide", level=2)

    add_key_term(
        doc, "NGFW (Next-Generation Firewall)",
        "Modern firewalls that inspect network traffic at the application level, not "
        "just at the IP address and port level. Traditional firewalls ask: 'Is traffic "
        "going to port 443?' NGFWs ask: 'Is this traffic actually HTTPS, or is someone "
        "using port 443 to hide something else?' Products include Palo Alto Networks, "
        "Fortinet FortiGate, Cisco Firepower, Check Point. Both Typhoon groups exploit "
        "misconfigured NGFWs as their primary entry point.",
    )
    add_key_term(
        doc, "VPN (Virtual Private Network)",
        "An encrypted tunnel connecting a remote user (working from home, traveling) to "
        "the corporate network. The VPN gateway is an 'edge device' -- it sits at the "
        "boundary between the internet and the internal network. Volt Typhoon exploits "
        "unpatched VPN gateways from Fortinet, Ivanti, and Palo Alto as their primary "
        "initial access method.",
    )
    add_key_term(
        doc, "ASM (Attack Surface Management)",
        "Tools that continuously map your organization's exposed assets -- what an "
        "attacker can see from the outside. This includes: internet-facing servers, "
        "open ports, expired SSL certificates, forgotten test environments, cloud "
        "misconfigurations. Examples: Mandiant ASM, Qualys EASM, CrowdStrike Falcon "
        "Surface. They scan and discover, but they SAMPLE -- they cannot test every "
        "possible attack path.",
    )
    add_key_term(
        doc, "BAS (Breach and Attack Simulation)",
        "Tools that simulate real attacks against your defenses to find weaknesses "
        "BEFORE real attackers do. Like a burglar alarm company that sends testers to "
        "try breaking in. Products include SafeBreach, AttackIQ, Cymulate. They run "
        "specific attack scenarios and report what succeeded. The problem: they can "
        "only test attack paths that someone has PROGRAMMED. They sample perhaps "
        "thousands of paths out of 2^8000 possible paths.",
    )
    add_key_term(
        doc, "2^8000 State Space",
        "The total number of possible configurations of all your security controls "
        "(every firewall rule can be on/off, every access policy can be allow/deny, "
        "every network segment can be connected/isolated, every authentication setting "
        "can be enforced/relaxed). 2^8000 means 2 multiplied by itself 8,000 times. "
        "For scale: 2^256 is already more than the number of atoms in the observable "
        "universe. 2^8000 is incomprehensibly, astronomically larger -- a number with "
        "over 2,400 digits. No computer can enumerate these combinations. No simulation "
        "tool can 'try them all.' Only mathematical proof (formal verification) can "
        "reason over the entire space without enumeration.",
    )
    add_key_term(
        doc, "ACL (Access Control List)",
        "A set of rules on a router or firewall that determines what traffic is "
        "allowed or blocked. Think of it as the guest list at a building entrance -- "
        "if your name is on the list, you get in; if not, you are turned away. Salt "
        "Typhoon modifies router ACLs to add themselves to the 'allowed' list, creating "
        "persistent hidden access.",
    )
    add_key_term(
        doc, "CVE (Common Vulnerabilities and Exposures)",
        "A standardized ID number assigned to a known security flaw in a specific product. "
        "Example: CVE-2024-3400 is a specific flaw in Palo Alto GlobalProtect VPN. When "
        "a flaw is discovered, the vendor releases a patch (software update). The flaw "
        "is publicly documented. Attackers scan the internet for organizations that have "
        "NOT applied the patch yet. The window between disclosure and patching is when "
        "organizations are most vulnerable.",
    )
    add_key_term(
        doc, "KEV (Known Exploited Vulnerabilities)",
        "CISA maintains a catalog of CVEs that are ACTIVELY being exploited by real "
        "attackers right now. Being on the KEV list means this is not theoretical -- "
        "someone is currently using this flaw to break into organizations. Federal "
        "agencies are legally required (BOD 22-01) to patch KEV items within specific "
        "timelines (typically 14-21 days). However, KEV covers only ~20% of attack "
        "vectors -- the rest exploit configuration weaknesses that have no CVE.",
    )
    add_key_term(
        doc, "MITRE ATT&CK Mapping",
        "The MITRE ATT&CK framework is a public knowledge base cataloging hundreds of "
        "real-world adversary techniques. Each technique has a T-number (e.g., T1003 = "
        "OS Credential Dumping, T1021 = Remote Services). When an alert includes MITRE "
        "mapping, the analyst immediately knows: what type of attack this is, what stage "
        "of the kill chain they are in, what to investigate next, and what the attacker's "
        "likely next move will be. WITHOUT mapping, the analyst must figure all of this "
        "out from scratch for every single alert.",
    )

    doc.add_heading("The Bottom Line", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "Result: Volt Typhoon dwelled 5+ years. Salt Typhoon dwelled 4+ years. "
        "Neither was caught by config-only or behavior-only defenses. Both structural "
        "gaps must be closed simultaneously -- which is why Rigor (preemptive) + "
        "ACECARD (behavioral) exist as a combined architecture."
    )
    run.bold = True

    add_insight_box(
        doc,
        "The fundamental problem is NOT that we lack security tools. It is that our "
        "tools are structurally incapable of detecting these specific threat patterns. "
        "More money spent on the same TYPES of tools will not help. Two fundamentally "
        "different approaches are required: mathematical proof (Rigor) for the "
        "preemptive layer, and behavioral trajectory intelligence (ACECARD) for the "
        "detection layer.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 4
    # =========================================================
    slide_heading(doc, 4, "The Two-Layer Answer -- Rigor + ACECARD")

    doc.add_paragraph(
        "This slide introduces the solution: a two-layer architecture where each layer "
        "addresses the gaps the other cannot. Together, they provide defense-in-depth "
        "that covers both PREVENTION (stopping attacks before they start) and DETECTION "
        "(catching attacks that evade prevention)."
    )

    doc.add_heading("Layer 1: Rigor AI -- Preemptive/Preventive", level=2)
    doc.add_paragraph(
        "Rigor operates BEFORE an attack happens. It analyzes your complete security "
        "configuration and mathematically proves which attack paths exist. It does not "
        "run simulations or sample scenarios -- it uses formal mathematical reasoning "
        "to exhaustively evaluate the ENTIRE 2^8000 state space."
    )
    add_analogy(
        doc,
        "Imagine you have a building with 8,000 doors, and each door can be locked or "
        "unlocked. A simulation tool might try 10,000 combinations and tell you it could "
        "not get in. But there are 2^8000 possible combinations. Rigor is like a "
        "mathematician who can PROVE that no combination allows entry -- without trying "
        "them one by one. This is the difference between testing and proving.",
    )

    add_bullet(doc, "Analyzes firewall rules, access control lists, network segmentation, identity policies")
    add_bullet(doc, "Identifies ALL possible attack paths -- not just the ones a tester happens to find")
    add_bullet(doc, "Provides mathematically guaranteed coverage (unlike sampling-based tools)")
    add_bullet(doc, "Operates on a Symbolic Model of Computation -- reasoning over logical rules, not running code")
    add_bullet(doc, "Proactive: finds vulnerabilities before attackers do")

    doc.add_heading("Layer 2: ACECARD UEBA -- Detective/Behavioral", level=2)
    doc.add_paragraph(
        "ACECARD operates DURING an attack. Even perfect prevention cannot guarantee zero "
        "breaches (zero-day vulnerabilities, stolen credentials, insider threats). ACECARD "
        "detects when an entity's behavior STRUCTURALLY CHANGES -- indicating compromise -- "
        "even when every individual action appears legitimate."
    )
    add_analogy(
        doc,
        "A security guard checks IDs at the door (prevention = Rigor). But what if someone "
        "has a valid ID? You also need cameras watching behavior patterns inside the building "
        "(detection = ACECARD). A person with valid credentials who starts visiting unusual "
        "areas at unusual times has a changed behavioral pattern -- even though each individual "
        "action is 'allowed.'",
    )

    add_bullet(doc, "Tracks behavioral fingerprint (1,536-dimensional vector) per entity per time window")
    add_bullet(doc, "Detects SLOW drift over days/weeks/months (CUSUM algorithm)")
    add_bullet(doc, "Explains WHAT the entity is drifting toward (drift direction with MITRE ATT&CK mapping)")
    add_bullet(doc, "Compares against peer cohort (same-role employees) to distinguish anomaly from normal variation")
    add_bullet(doc, "Automatically adjusts access (ABAC trust loop) -- proportional response without human delay")

    doc.add_heading("Why Both Layers Are Necessary", level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Attack Type", "Rigor (Prevention)", "ACECARD (Detection)"]):
        table.rows[0].cells[i].text = h
    data = [
        ("Misconfigured firewall allows path", "CATCHES (proves path exists)", "May not see (no behavior yet)"),
        ("Attacker exploits zero-day", "Cannot prevent (unknown flaw)", "CATCHES (behavioral drift)"),
        ("Insider with valid credentials", "Cannot prevent (authorized user)", "CATCHES (behavioral change)"),
        ("Slow APT over 6 months", "Cannot prevent (incremental)", "CATCHES (CUSUM accumulation)"),
        ("Attacker uses LOtL tools", "Cannot prevent (legitimate tools)", "CATCHES (unusual tool patterns)"),
    ]
    for i, (attack, rigor, acecard) in enumerate(data):
        row = table.rows[i + 1]
        row.cells[0].text = attack
        row.cells[1].text = rigor
        row.cells[2].text = acecard

    add_insight_box(
        doc,
        "Neither layer alone is sufficient. Prevention without detection leaves you blind "
        "to attacks that bypass your controls. Detection without prevention means you are "
        "always reacting after the fact. The two layers together close the entire gap.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDES 5-6
    # =========================================================
    slide_heading(doc, "5-6", "Rigor AI -- Approach and Operational Flow")

    doc.add_heading("What Rigor Actually Does (Non-Technical Explanation)", level=2)
    doc.add_paragraph(
        "Rigor AI takes your entire network configuration -- every firewall rule, every "
        "access control policy, every network segmentation boundary, every identity "
        "permission -- and builds a mathematical model of it. Then it uses formal logic "
        "(the same kind of math used to prove theorems) to determine ALL possible "
        "attack paths through that configuration."
    )

    doc.add_heading("How It Differs from Traditional Tools", level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Aspect", "Traditional (ASM/BAS/Pen Test)", "Rigor AI"]):
        table.rows[0].cells[i].text = h
    data = [
        ("Method", "Sample & test paths", "Mathematical proof over all paths"),
        ("Coverage", "Thousands to millions of tests", "Complete (2^8000 state space)"),
        ("Guarantee", "No path found = 'probably safe'", "No path exists = PROVEN safe"),
        ("Speed", "Hours to run simulations", "Seconds (symbolic reasoning)"),
        ("False negatives", "Many (missed paths)", "None (mathematically complete)"),
    ]
    for i, (aspect, trad, rigor) in enumerate(data):
        row = table.rows[i + 1]
        row.cells[0].text = aspect
        row.cells[1].text = trad
        row.cells[2].text = rigor

    doc.add_heading("Key Technical Concepts", level=2)
    add_key_term(
        doc, "Symbolic Model of Computation",
        "Instead of running actual traffic or simulations, Rigor represents your network as "
        "a set of logical symbols and rules. It then applies mathematical transformations to "
        "determine what is reachable from what. Like solving an equation vs. trying random "
        "numbers until one works.",
    )
    add_key_term(
        doc, "Formal Verification",
        "A field of computer science that uses mathematical proof to guarantee properties of "
        "a system. Originally developed for hardware chip design (where bugs cost billions), "
        "now applied to cybersecurity. Intel uses formal verification to prove their chips "
        "work correctly; Rigor uses it to prove your network is correctly configured.",
    )
    add_key_term(
        doc, "State Space Exhaustion",
        "Instead of sampling paths, Rigor reasons over ALL possible states simultaneously. "
        "The 2^8000 state space is not enumerated one by one (that would take longer than "
        "the age of the universe). Instead, symbolic methods represent entire classes of "
        "states as single logical expressions, making exhaustive reasoning tractable.",
    )
    add_key_term(
        doc, "NGFW (Next-Generation Firewall)",
        "Modern firewalls that inspect traffic at the application level, not just IP "
        "addresses and ports. Products from Palo Alto, Fortinet, Check Point. Rigor "
        "analyzes their complete rule sets.",
    )
    add_key_term(
        doc, "IdP (Identity Provider)",
        "The system that verifies who you are when you log in. Examples: Okta, Azure AD, "
        "Ping Identity. Manages usernames, passwords, MFA (multi-factor authentication), "
        "and single sign-on.",
    )
    add_key_term(
        doc, "SASE (Secure Access Service Edge)",
        "A cloud-delivered network security model that combines networking and security "
        "functions. Products from Zscaler, Netskope, Palo Alto Prisma. Rigor can analyze "
        "the policies configured in these platforms.",
    )
    add_key_term(
        doc, "WAF (Web Application Firewall)",
        "A specialized firewall that protects web applications by filtering HTTP traffic. "
        "Products from Cloudflare, Akamai, AWS. Rigor verifies that WAF rules are "
        "correctly configured to block known attack patterns.",
    )

    doc.add_heading("Operational Flow", level=2)
    doc.add_paragraph("Rigor operates in a continuous cycle:")
    add_bullet(doc, "Ingest current configurations from all security controls (firewalls, IdPs, cloud security groups, network segmentation policies)")
    add_bullet(doc, "Build symbolic model representing the complete security posture")
    add_bullet(doc, "Run formal analysis to identify ALL reachable attack paths")
    add_bullet(doc, "Report findings: which paths exist, what attacker capabilities are needed, which controls to modify")
    add_bullet(doc, "Validate remediation: after changes, re-verify that the path is truly closed")
    add_bullet(doc, "Continuous monitoring: re-analyze whenever configurations change (drift detection at the config level)")

    doc.add_page_break()

    # =========================================================
    # SLIDES 7-8
    # =========================================================
    slide_heading(doc, "7-8", "How Rigor Stops Volt Typhoon and Salt Typhoon Phase-by-Phase")

    doc.add_heading("Volt Typhoon -- Rigor's Response at Each Phase", level=2)

    doc.add_heading("Phase 1: Initial Access via Edge Device Exploitation", level=3)
    doc.add_paragraph(
        "Attack: Volt Typhoon exploits unpatched VPN gateways and firewalls "
        "(Fortinet, Ivanti, Palo Alto, Versa) to gain initial foothold."
    )
    doc.add_paragraph(
        "Rigor's Response: Verifies that even if an edge device is compromised, the "
        "attacker cannot reach critical internal systems. Proves network segmentation "
        "is correctly configured so that a compromised DMZ device cannot pivot inward. "
        "Identifies any misconfigured rules that would allow such traversal."
    )

    doc.add_heading("Phase 2: Living-off-the-Land Execution", level=3)
    doc.add_paragraph(
        "Attack: Uses built-in Windows tools (wmic, ntdsutil, netsh, rundll32) to "
        "execute commands and extract credentials."
    )
    doc.add_paragraph(
        "Rigor's Response: Verifies that application-control policies and network "
        "micro-segmentation prevent unauthorized lateral use of administrative tools. "
        "Proves that even with valid credentials on one workstation, the security "
        "architecture prevents reach to domain controllers and file servers."
    )

    doc.add_heading("Phase 3: Lateral Movement", level=3)
    doc.add_paragraph(
        "Attack: Moves between systems using RDP, SMB, WinRM with stolen credentials."
    )
    doc.add_paragraph(
        "Rigor's Response: Maps ALL possible lateral movement paths given the current "
        "firewall rules, VLAN segmentation, and access policies. Identifies where a "
        "stolen credential + network access combination allows reaching high-value targets. "
        "Reports: 'With Domain Admin credentials from workstation in VLAN 10, the attacker "
        "can reach the SQL server in VLAN 50 via path X -- recommend closing ports Y and Z.'"
    )

    doc.add_heading("Phase 4: C2 via KV Botnet", level=3)
    doc.add_paragraph(
        "Attack: Commands routed through compromised residential routers to avoid "
        "detection by IP reputation systems."
    )
    doc.add_paragraph(
        "Rigor's Response: Verifies that outbound connection policies are restrictive "
        "enough that only whitelisted destinations are reachable from critical systems. "
        "Proves that even if a C2 channel is established from a workstation, it cannot "
        "exfiltrate data from segmented high-value assets."
    )

    doc.add_heading("Salt Typhoon -- Rigor's Response at Each Phase", level=2)

    doc.add_heading("Phase 1: Network Infrastructure Exploitation", level=3)
    doc.add_paragraph(
        "Attack: Exploits Cisco IOS XE, Citrix NetScaler vulnerabilities to compromise "
        "core routing infrastructure."
    )
    doc.add_paragraph(
        "Rigor's Response: Analyzes the management plane access controls for all network "
        "devices. Proves whether management interfaces are reachable from unauthorized "
        "sources. Identifies any path from the internet or user VLANs to router management "
        "ports (SSH, SNMP, NETCONF). Reports: 'Router R1 management port is reachable from "
        "VLAN 20 via rule exception #47 -- this should not exist.'"
    )

    doc.add_heading("Phase 2: ACL/AAA Modification", level=3)
    doc.add_paragraph(
        "Attack: After compromising a router, modifies Access Control Lists and "
        "Authentication/Authorization/Accounting configs to create persistent hidden access."
    )
    doc.add_paragraph(
        "Rigor's Response: Continuously ingests router configurations. When ACLs change, "
        "immediately re-verifies the security posture. Detects: 'New rule #52 allows "
        "traffic from 10.0.0.0/8 to reach management VLAN -- this opens 47 previously "
        "closed attack paths.' Configuration drift = formal re-analysis."
    )

    doc.add_heading("Phase 3: GRE Tunneling and Exfiltration", level=3)
    doc.add_paragraph(
        "Attack: Creates GRE tunnels to exfiltrate data (call records, metadata) to "
        "external attacker infrastructure."
    )
    doc.add_paragraph(
        "Rigor's Response: Verifies that all configured GRE/IPIP/VXLAN tunnel endpoints "
        "are authorized. Detects new tunnel configurations that create unauthorized paths "
        "out of the network. Proves: 'New GRE tunnel from router R3 to IP 203.0.113.5 "
        "creates an exfiltration path bypassing all DLP controls.'"
    )

    doc.add_page_break()

    # =========================================================
    # SLIDES 9-10
    # =========================================================
    slide_heading(doc, "9-10", "TTP-by-TTP Coverage Tables")

    doc.add_paragraph(
        "These slides provide a detailed mapping of specific attacker techniques "
        "(from the MITRE ATT&CK framework) to how Rigor addresses each one. This is "
        "the evidence that Rigor's formal verification approach is not just theoretical "
        "-- it maps directly to real-world attack behaviors."
    )

    doc.add_heading("Key Terms for Understanding These Tables", level=2)
    add_key_term(
        doc, "MITRE ATT&CK",
        "A publicly available knowledge base of adversary tactics and techniques based on "
        "real-world observations. Maintained by MITRE Corporation (a U.S. federally funded "
        "research organization). Categories include: Initial Access, Execution, Persistence, "
        "Privilege Escalation, Defense Evasion, Credential Access, Discovery, Lateral "
        "Movement, Collection, Exfiltration, Command & Control.",
    )
    add_key_term(
        doc, "T-Number (e.g., T1190, T1003)",
        "Each MITRE ATT&CK technique has a unique identifier. T1190 = 'Exploit Public-"
        "Facing Application.' T1003 = 'OS Credential Dumping.' These are the standard "
        "language security teams worldwide use to describe attack methods.",
    )
    add_key_term(
        doc, "CVE (Common Vulnerabilities and Exposures)",
        "A standardized ID for a known security flaw. Example: CVE-2024-3400 is a specific "
        "flaw in Palo Alto GlobalProtect. Manufacturers release patches; attackers scan "
        "for systems that have not applied the patch.",
    )
    add_key_term(
        doc, "KEV (Known Exploited Vulnerabilities)",
        "CISA maintains a catalog of CVEs that are ACTIVELY being used by attackers. Being "
        "on the KEV list means this is not theoretical -- real attackers are using this flaw "
        "right now. Federal agencies are required to patch KEV items within specific timelines.",
    )

    doc.add_heading("Example Coverage Entries (Volt Typhoon)", level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["TTP", "What the Attacker Does", "How Rigor Prevents It"]):
        table.rows[0].cells[i].text = h
    data = [
        ("T1190: Exploit Public App", "Exploits CVEs in VPN/firewall", "Proves patch state + reachability = no valid path"),
        ("T1078: Valid Accounts", "Uses stolen credentials", "Verifies least-privilege + segmentation limits blast radius"),
        ("T1059: Command Scripting", "Runs wmic, PowerShell, netsh", "Verifies app-control policies block unauthorized use"),
        ("T1021: Remote Services", "RDP/SMB/WinRM lateral movement", "Proves micro-segmentation blocks unauthorized paths"),
        ("T1071: Application Protocol", "C2 over HTTPS/DNS", "Verifies outbound policies restrict unauthorized destinations"),
        ("T1003: Credential Dumping", "LSASS dump, ntdsutil", "Verifies credential-guard policies + DC isolation"),
    ]
    for i, (ttp, attack, rigor) in enumerate(data):
        row = table.rows[i + 1]
        row.cells[0].text = ttp
        row.cells[1].text = attack
        row.cells[2].text = rigor

    doc.add_page_break()

    # =========================================================
    # SLIDE 11
    # =========================================================
    slide_heading(doc, 11, "Nine Rigor Use Cases")

    doc.add_paragraph(
        "Beyond stopping Volt/Salt Typhoon specifically, Rigor provides nine "
        "categories of operational value:"
    )

    use_cases = [
        ("1. Attack Path Analysis",
         "Find ALL paths an attacker could take from a given starting point to a target. "
         "Not just 'can they get there?' but 'how many ways can they get there, and which "
         "is shortest?'"),
        ("2. Compliance Verification",
         "Mathematically prove your network meets a compliance standard (NIST 800-53, "
         "CMMC, PCI-DSS, HIPAA). Not 'we checked the boxes' but 'it is mathematically "
         "impossible to violate this control.'"),
        ("3. Change Impact Analysis",
         "Before making a configuration change, predict all security implications. "
         "'If we open this port for the new application, how many new attack paths does "
         "it create?'"),
        ("4. Micro-Segmentation Validation",
         "Verify that your network segmentation is working as intended. Prove that VLAN X "
         "truly cannot reach VLAN Y, even through indirect paths via other VLANs."),
        ("5. Firewall Rule Optimization",
         "Identify redundant, conflicting, or overly permissive firewall rules. "
         "Some organizations have 50,000+ rules accumulated over years -- which ones "
         "actually matter?"),
        ("6. Zero Trust Verification",
         "Verify that your Zero Trust architecture truly enforces 'never trust, always "
         "verify.' Prove that no implicit trust relationships exist in the configuration."),
        ("7. Cloud Security Group Audit",
         "For cloud environments (AWS, Azure, GCP), verify that security groups, network "
         "ACLs, and IAM policies combine to prevent unauthorized access paths."),
        ("8. Merger/Acquisition Risk Assessment",
         "When two networks must be connected (M&A, partnership), identify all attack paths "
         "that the connection would create BEFORE connecting them."),
        ("9. Incident Response Prioritization",
         "During an active breach, instantly determine which systems the attacker can reach "
         "from their known position. Prioritize containment actions based on mathematical "
         "reachability, not guesswork."),
    ]
    for title, desc in use_cases:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =========================================================
    # SLIDE 12
    # =========================================================
    slide_heading(doc, 12, "Where Rigor Needs Help -- Honest Assessment")

    doc.add_paragraph(
        "This slide is critically important because it demonstrates intellectual honesty. "
        "No single tool solves everything. Rigor has specific limitations that explain "
        "why ACECARD is needed as a complement."
    )

    doc.add_heading("What Rigor Cannot Do", level=2)
    limitations = [
        ("Cannot detect attacks in progress",
         "Rigor analyzes CONFIGURATIONS, not live traffic. If an attacker is already inside "
         "using valid credentials, Rigor does not see them. It can tell you the door was "
         "unlockable, but not that someone walked through it."),
        ("Cannot handle zero-day exploits proactively",
         "If a vulnerability is unknown (zero-day), Rigor cannot model it in advance. It "
         "can verify whether the BLAST RADIUS is limited (can the attacker reach anything "
         "useful even if they exploit the zero-day?), but it cannot prevent the initial "
         "exploitation."),
        ("Cannot detect behavioral change",
         "A legitimate user who becomes a malicious insider does not change any configuration. "
         "Their credentials are valid, their access is authorized. Only behavioral monitoring "
         "can detect the CHANGE in how they use their access."),
        ("Cannot detect credential theft/reuse",
         "If an attacker steals valid credentials and uses them from expected locations, "
         "the configuration still looks correct. The ACCESS is authorized -- only the "
         "BEHAVIOR reveals that a different person is using the account."),
        ("Cannot adapt in real-time",
         "Rigor re-analyzes when configurations change, but it does not watch behavior "
         "continuously. It cannot say 'this user is acting strange right now' -- that "
         "requires behavioral telemetry and continuous monitoring."),
    ]
    for title, desc in limitations:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    add_insight_box(
        doc,
        "This is exactly why ACECARD exists. Every limitation Rigor has is a strength "
        "of ACECARD. Rigor prevents what can be prevented. ACECARD detects what cannot "
        "be prevented. Together = complete coverage.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 13
    # =========================================================
    slide_heading(doc, 13, "Rigor Proof Points -- Real Customer Findings")

    doc.add_paragraph(
        "This slide provides concrete evidence from real Rigor deployments across "
        "financial services, telecom, government, and critical infrastructure customers."
    )

    doc.add_heading("Selected Customer Findings", level=2)

    doc.add_heading("FinTech | Asia -- Closed SMB Lateral Movement Path", level=3)
    doc.add_paragraph(
        "SMB application-control profiles were completely shadowed by 3 higher-priority "
        "firewall rules -- this allowed unprotected SMBv1 traffic between branches. "
        "Rigor's formal model exposed the multi-rule shadow that manual review missed. "
        "This is exactly the kind of path Volt Typhoon exploits for lateral movement."
    )

    doc.add_heading("Communications | North America -- Eliminated Known CVE Exposure", level=3)
    doc.add_paragraph(
        "IPS (Intrusion Prevention System) filters were missing on rsync traffic to backup "
        "servers. This exposed the enterprise to lateral malware movement via an unmonitored "
        "protocol. Rigor identified the gap by formally verifying that ALL traffic to ALL "
        "server segments had appropriate IPS coverage -- something manual audits had missed."
    )

    doc.add_heading("Government / Critical Infrastructure -- Additional Findings", level=3)
    add_bullet(doc, "Discovered firewall rules that were correct individually but created exploitable paths in combination (multi-rule interaction effects)")
    add_bullet(doc, "Identified 'zombie rules' -- legacy permit rules for decommissioned systems that still allowed unauthorized traffic")
    add_bullet(doc, "Found cloud security group misconfigurations that exposed internal services to the internet via indirect routing")
    add_bullet(doc, "Verified Zero Trust claims that turned out to be incomplete (implicit trust relationships existed in edge cases)")

    doc.add_heading("What 'Proof' Means in This Context", level=2)
    doc.add_paragraph(
        "When Rigor says 'no attack path exists,' this is a mathematical proof -- the same "
        "type of certainty as '2 + 2 = 4.' It is not a statistical claim ('probably safe') "
        "or an empirical claim ('we tested and nothing happened'). It is a logical certainty "
        "given the model. The only caveat: if the model does not perfectly represent reality "
        "(e.g., a typo in importing firewall rules), the proof applies to the model, not to "
        "the real system. This is why continuous configuration ingestion is critical."
    )

    doc.add_heading("Industry Precedent for Formal Methods", level=2)
    add_bullet(doc, "Intel/AMD/ARM use formal verification to prove chip correctness (30+ year track record)")
    add_bullet(doc, "AWS uses formal methods (model checking) to verify S3, IAM, VPC services")
    add_bullet(doc, "Microsoft uses DAFNY/VCC to prove security properties of Windows")
    add_bullet(doc, "FAA requires formal methods for safety-critical avionics software")
    add_bullet(doc, "NASA uses formal verification for autonomous systems and Mars rover code")

    doc.add_page_break()

    # =========================================================
    # SLIDE 14
    # =========================================================
    slide_heading(doc, 14, "Why Current UEBA Misses Volt Typhoon")

    doc.add_paragraph(
        "Both Volt Typhoon and Salt Typhoon relied on Living-off-the-Land tradecraft. "
        "Both walked past commercial UEBA tooling (Exabeam, Securonix, Microsoft Sentinel "
        "UEBA) for YEARS. Volt Typhoon's dwell time exceeded 5 years in U.S. critical "
        "infrastructure. This slide explains the four structural gaps that made this possible."
    )

    doc.add_heading("The Four Structural Gaps in Commercial UEBA", level=2)
    doc.add_paragraph(
        "Each gap is exactly what Volt Typhoon and Salt Typhoon exploit."
    )

    failures = [
        ("Gap 1: Threshold-Based Scoring",
         "Commercial UEBA alerts fire only when a metric crosses a fixed threshold line. "
         "Living-off-the-Land NEVER crosses thresholds -- every binary used (ntdsutil, "
         "rundll32, vssadmin, netsh, wmic) is a legitimate Windows tool used by IT teams "
         "every day. There is no threshold that catches 'a legitimate tool running normally.'"),
        ("Gap 2: No Behavioral Trajectory",
         "Current products compare 'last 24 hours' to 'last 30 days' using static sliding "
         "windows. They cannot track slow drift that accumulates over months. Volt Typhoon "
         "changes behavior by 1% per week -- no 30-day window comparison catches this. But "
         "after 6 months, behavior has changed 26%. Only a trajectory (time-series shape) "
         "reveals this."),
        ("Gap 3: Anomaly Without Direction",
         "When commercial UEBA does flag something, it says 'this is anomalous' -- weird, "
         "different, unusual. It does NOT say 'this entity is drifting toward "
         "lateral_movement (T1021) with velocity 0.03/hour.' Without direction, every alert "
         "requires full manual investigation. Alert fatigue causes analysts to ignore 90%+ "
         "of alerts."),
        ("Gap 4: No Cross-Entity Identity Fusion",
         "Volt Typhoon uses stolen credentials across Windows AD, Azure AD, VPN systems, "
         "and cloud platforms. Each system sees a 'normal' login. No current commercial "
         "product fuses all identity signals into a single trajectory per real-world entity. "
         "Without this fusion, the attack appears as isolated normal events in separate "
         "tools -- invisible in aggregate."),
    ]
    for title, desc in failures:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =========================================================
    # SLIDE 15
    # =========================================================
    slide_heading(doc, 15, "Three ACECARD Innovations")

    doc.add_paragraph(
        "These are the three specific technical innovations that differentiate ACECARD "
        "from existing UEBA products."
    )

    doc.add_heading("Innovation 1: Behavioral Embedding in Unified Semantic Space", level=2)
    doc.add_paragraph(
        "Every entity's behavior is converted into a 1,536-dimensional mathematical vector "
        "(a list of 1,536 numbers). This is called an 'embedding.' All entity types (users, "
        "devices, network segments, applications) exist in the SAME mathematical space."
    )
    add_analogy(
        doc,
        "Think of GPS coordinates. Every location on Earth has a latitude and longitude "
        "(2 numbers). You can calculate the distance between any two locations. Similarly, "
        "ACECARD gives every entity a 'behavioral coordinate' with 1,536 dimensions (instead "
        "of 2). You can calculate the 'distance' between any two behavioral states. When an "
        "entity's coordinate moves, that is drift.",
    )
    add_bullet(doc, "1,536 dimensions captures nuances that 2 or 10 dimensions cannot")
    add_bullet(doc, "Same space for all entity types means you can compare users to devices to network segments")
    add_bullet(doc, "Semantic embedding means the system understands MEANING, not just numbers")

    doc.add_heading("Innovation 2: CUSUM Change-Point Detection for Slow Drift", level=2)
    doc.add_paragraph(
        "CUSUM (Cumulative Sum) is a statistical algorithm that detects SUSTAINED small "
        "changes that accumulate over time. Traditional alerts fire when a value exceeds a "
        "threshold. CUSUM fires when the CUMULATIVE deviation from baseline exceeds a "
        "threshold -- even if no single measurement is alarming."
    )
    add_analogy(
        doc,
        "Imagine your water bill. If it suddenly doubles, you notice immediately (threshold "
        "alert). But what if it increases by $2 every month? After a year, you are paying "
        "$24 more -- but no single month looked alarming. CUSUM catches the sustained upward "
        "trend. This is EXACTLY how APT attackers operate: tiny changes, consistently applied, "
        "accumulating over months.",
    )
    add_bullet(doc, "Catches Volt Typhoon's multi-year slow creep")
    add_bullet(doc, "Catches insider threat escalation over months")
    add_bullet(doc, "Catches Salt Typhoon's gradual data exfiltration")
    add_bullet(doc, "Threshold: 4 standard deviations (4-sigma) to avoid false positives")

    doc.add_heading("Innovation 3: Drift Direction with MITRE ATT&CK Mapping", level=2)
    doc.add_paragraph(
        "When drift is detected, ACECARD does not just say 'something changed.' It tells "
        "you WHAT the entity is drifting toward by projecting the drift vector onto 8 "
        "reference threat concepts, each mapped to specific MITRE ATT&CK techniques."
    )
    add_analogy(
        doc,
        "If someone tells you 'a storm is coming,' that is anomaly detection. If they tell "
        "you 'a Category 4 hurricane is heading northeast at 15 mph toward Miami,' that is "
        "drift direction. The first triggers generic panic. The second enables specific, "
        "proportional action.",
    )
    add_bullet(doc, "8 threat concepts: credential_dumping, lateral_movement, data_exfiltration, c2_beaconing, lotl_execution, privilege_escalation, defense_evasion, insider_data_hoarding")
    add_bullet(doc, "Each concept has a reference text that is also embedded in the same 1,536-d space")
    add_bullet(doc, "Projection score = cosine similarity between drift vector and concept embedding")
    add_bullet(doc, "Output: 'This entity is drifting toward lateral_movement (0.78) and credential_dumping (0.45)'")
    add_bullet(doc, "Each concept maps to MITRE ATT&CK technique IDs for standardized reporting")

    doc.add_page_break()

    # =========================================================
    # SLIDE 16
    # =========================================================
    slide_heading(doc, 16, "The Math of Behavioral Drift")

    doc.add_paragraph(
        "This slide explains the mathematics behind drift detection. Do not be "
        "intimidated by the formulas -- each one has a simple plain-English meaning."
    )

    doc.add_heading("Core Formulas Explained", level=2)

    add_key_term(
        doc, "Cosine Drift: drift(t) = 1 - cosine(V_t, V_baseline)",
        "Plain English: 'How different is this entity's behavior RIGHT NOW compared to "
        "their established normal?' A value of 0.0 means identical to baseline. A value "
        "of 0.05 means noticeably different. A value of 0.15+ means highly anomalous. "
        "Cosine similarity measures the ANGLE between two vectors -- if they point in the "
        "same direction, similarity = 1 (drift = 0). If perpendicular, similarity = 0 "
        "(drift = 1).",
    )
    add_key_term(
        doc, "Velocity: velocity(t) = drift(t) - drift(t-1)",
        "Plain English: 'Is the entity drifting FASTER or SLOWER?' Positive velocity means "
        "they are diverging from baseline at an increasing rate. Negative velocity means "
        "they are returning to normal. A sudden velocity spike indicates an abrupt change.",
    )
    add_key_term(
        doc, "Acceleration: accel(t) = velocity(t) - velocity(t-1)",
        "Plain English: 'Is the RATE of drift itself changing?' Acceleration = the rate of "
        "change of the rate of change. A spike in acceleration means the entity just "
        "experienced a sudden regime shift -- their behavior is not just changing, it is "
        "changing faster and faster.",
    )
    add_key_term(
        doc, "CUSUM: S_t = max(0, S_{t-1} + (x_t - mean) - k); alarm when S_t > h",
        "Plain English: 'Has the CUMULATIVE total of small deviations exceeded the alarm "
        "level?' Each time period, we add the deviation above the mean (minus a small "
        "allowance k). If the running sum exceeds threshold h (= 4 standard deviations), "
        "we flag a change-point. Key insight: many small deviations that individually "
        "look harmless will accumulate until CUSUM triggers.",
    )
    add_key_term(
        doc, "Health Score: health = 100 - weighted_drift - weighted_velocity - weighted_acceleration - weighted_cusum_count",
        "Plain English: 'On a 0-100 scale, how healthy is this entity?' A score of 100 "
        "means perfectly normal. Below 70 with a recent change-point = medium alert. "
        "Below 40 = critical alert requiring immediate investigation.",
    )
    add_key_term(
        doc, "Drift Direction: projection = cosine(drift_vector, concept_embedding)",
        "Plain English: 'Which threat pattern does this behavioral change most resemble?' "
        "We compute the drift_vector (current - baseline), then measure how closely it "
        "aligns with each of 8 known threat patterns. The highest-scoring concept tells "
        "the analyst what kind of attack this looks like.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 17
    # =========================================================
    slide_heading(doc, 17, "Five Behavioral Signals")

    doc.add_paragraph(
        "ACECARD monitors five distinct categories of behavior for every entity. Each "
        "signal captures a different dimension of 'what this entity does.' Together, they "
        "form a complete behavioral portrait."
    )

    doc.add_heading("Signal 1: AUTH (Authentication Behavior)", level=2)
    doc.add_paragraph(
        "What it measures: How the entity logs in -- when, from where, how often, "
        "success vs. failure rate, what protocols used."
    )
    add_bullet(doc, "Data sources: Windows Event 4624 (successful login), 4625 (failed login), 4768/4769 (Kerberos tickets), Okta, Azure AD")
    add_bullet(doc, "Features: logon_count, failure_rate, unique_hosts_accessed, kerberos_ticket_types, off_hours_ratio, impossible_travel_flag, auth_protocol_mix")
    add_bullet(doc, "What drift means: New login sources, unusual failure patterns, off-hours access, new protocols = possible credential theft/reuse")

    doc.add_heading("Signal 2: PROCESS (Process/Execution Behavior)", level=2)
    doc.add_paragraph(
        "What it measures: What programs the entity runs -- which executables, "
        "what command-line arguments, parent-child process relationships."
    )
    add_bullet(doc, "Data sources: Sysmon Event ID 1 (process creation), EDR (CrowdStrike, Microsoft Defender), Linux auditd")
    add_bullet(doc, "Features: unique_processes, cmdline_entropy, parent_child_depth, lolbin_count, unsigned_binary_ratio, new_process_rate, scripting_engine_calls")
    add_bullet(doc, "What drift means: New tools appearing, LOLBin usage (wmic, certutil), deep process chains, encoded command lines = possible execution/LOtL attack")

    add_key_term(
        doc, "LOLBin (Living-off-the-Land Binary)",
        "A legitimate system tool that attackers abuse for malicious purposes. Examples: "
        "wmic.exe, certutil.exe, mshta.exe, rundll32.exe, regsvr32.exe. These are "
        "Microsoft-signed tools that every Windows computer has -- antivirus trusts them.",
    )

    doc.add_heading("Signal 3: NETWORK (Network Behavior)", level=2)
    doc.add_paragraph(
        "What it measures: Network communication patterns -- who the entity talks to, "
        "how much data, what ports, any beaconing patterns."
    )
    add_bullet(doc, "Data sources: NetFlow, Zeek (formerly Bro), Palo Alto/Fortinet traffic logs")
    add_bullet(doc, "Features: unique_destination_IPs, bytes_out, bytes_in_out_ratio, port_diversity, dns_query_rate, beacon_score, geo_anomaly_count, encrypted_traffic_percentage")
    add_bullet(doc, "What drift means: New destinations, high outbound ratio, periodic beaconing, unusual DNS = possible C2 communication or data exfiltration")

    add_key_term(
        doc, "Beacon Score",
        "A measure of how periodic (regular-interval) an entity's outbound connections "
        "are. Legitimate traffic is bursty and irregular. C2 malware often 'phones home' "
        "at predictable intervals (every 30 seconds, every 5 minutes). High beacon score "
        "= suspicious regularity.",
    )

    doc.add_heading("Signal 4: FILE (File Access Behavior)", level=2)
    doc.add_paragraph(
        "What it measures: How the entity interacts with files -- creating, deleting, "
        "renaming, accessing sensitive documents, creating archives."
    )
    add_bullet(doc, "Data sources: Sysmon Event ID 11 (file creation), Windows file audit (Event 4663), DLP systems")
    add_bullet(doc, "Features: files_created, files_deleted, files_renamed, sensitive_access_count, archive_creates, bulk_operations_flag, extension_change_rate")
    add_bullet(doc, "What drift means: Bulk sensitive file access, archive creation, extension changes (.docx to .zip) = possible staging for exfiltration or ransomware")

    add_key_term(
        doc, "DLP (Data Loss Prevention)",
        "Software that monitors and blocks sensitive data from leaving the organization. "
        "It watches for patterns like credit card numbers, classified markings, or PHI "
        "(Protected Health Information) in emails, file transfers, and cloud uploads.",
    )

    doc.add_heading("Signal 5: IDENTITY (Identity/Privilege Behavior)", level=2)
    doc.add_paragraph(
        "What it measures: Changes to the entity's identity and privileges -- "
        "escalations, group membership changes, MFA events, administrative actions."
    )
    add_bullet(doc, "Data sources: Windows Event 4672 (special privileges), CloudTrail IAM events, Kubernetes audit logs, Active Directory group changes")
    add_bullet(doc, "Features: privilege_escalations, group_adds, mfa_bypass_attempts, role_changes, service_account_use, admin_actions, token_refreshes")
    add_bullet(doc, "What drift means: New admin actions, group additions, MFA bypasses, service account usage by a non-service entity = possible privilege escalation")

    add_key_term(
        doc, "MFA (Multi-Factor Authentication)",
        "Requiring more than one proof of identity to log in. Factor 1 = something you know "
        "(password). Factor 2 = something you have (phone, hardware key). Factor 3 = "
        "something you are (fingerprint, face). MFA bypass attempts are highly suspicious.",
    )

    doc.add_heading("How the 5 Signals Become One Vector", level=2)
    doc.add_paragraph(
        "Each signal's metrics are serialized as structured text (e.g., 'User jsmith auth "
        "window 2025-05-02T14:00Z: logon_count=47, failure_rate=0.02...'). All 5 signal "
        "texts are concatenated and embedded via OpenAI's text-embedding-3-small model into "
        "a single 1,536-dimensional vector. This fused vector IS the entity's behavioral "
        "fingerprint for that time window."
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 18
    # =========================================================
    slide_heading(doc, 18, "Pipeline and Trajectory Analysis")

    doc.add_heading("The 6-Step Pipeline", level=2)

    steps = [
        ("Step 1: INGEST",
         "Raw events arrive from data sources in standardized formats (ECS, OCSF, or "
         "STIX 2.1). These are the raw logs that security tools already collect."),
        ("Step 2: WINDOW",
         "Events are grouped into 1-hour time windows per entity. All events for user "
         "'jsmith' between 14:00 and 15:00 become one window."),
        ("Step 3: SERIALIZE",
         "The 5 signal serializers convert raw events into structured text descriptions "
         "of the entity's behavior in that window. Only aggregate METRICS are included "
         "-- no raw event content, no PII (Personally Identifiable Information)."),
        ("Step 4: HASH",
         "Entity identifiers (usernames, hostnames) are hashed using SHA-256 before "
         "embedding. This ensures no PII is sent to the embedding provider."),
        ("Step 5: EMBED",
         "The concatenated signal text is sent to OpenAI's text-embedding-3-small model "
         "(or a local alternative for classified networks). Output: one 1,536-d vector."),
        ("Step 6: STORE",
         "The vector is stored in pgvector (PostgreSQL with vector extensions) alongside "
         "entity_id, window_start timestamp, and cutoff date. HNSW indexes enable fast "
         "similarity searches across thousands of entities."),
    ]
    for title, desc in steps:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    doc.add_heading("Key Terms", level=2)
    add_key_term(
        doc, "ECS (Elastic Common Schema)",
        "A standardized way to format security events. Developed by Elastic (makers of "
        "Elasticsearch/Splunk alternative). Widely deployed in DoD SOCs. Ensures that "
        "events from different sources (Windows, Linux, firewalls) use the same field names.",
    )
    add_key_term(
        doc, "OCSF (Open Cybersecurity Schema Framework)",
        "An AWS-led schema standard for cybersecurity events. Similar goal to ECS but "
        "newer and cloud-native. Supported by CrowdStrike, Splunk, IBM.",
    )
    add_key_term(
        doc, "STIX 2.1 (Structured Threat Information Expression)",
        "A standard for sharing threat intelligence (indicators, attack patterns, "
        "observed data). Used for replay scenarios and CTI enrichment.",
    )
    add_key_term(
        doc, "pgvector",
        "An extension for PostgreSQL (a popular database) that adds support for storing "
        "and searching high-dimensional vectors. Enables finding 'which entities have "
        "similar behavioral patterns' in milliseconds.",
    )
    add_key_term(
        doc, "HNSW (Hierarchical Navigable Small World)",
        "A fast algorithm for searching through millions of vectors to find the most "
        "similar ones. Instead of comparing against every stored vector (too slow), "
        "HNSW builds a graph structure that enables logarithmic-time search.",
    )

    doc.add_heading("Baseline Computation", level=2)
    doc.add_paragraph(
        "The baseline is a 30-day rolling average of stored embeddings. It represents "
        "'what this entity normally does.' Drift is measured as deviation from this "
        "baseline. The baseline updates every 24 hours and can be frozen at a specific "
        "date for investigation purposes."
    )

    doc.add_page_break()

    # =========================================================
    # SLIDES 19-20
    # =========================================================
    slide_heading(doc, "19-20", "How ACECARD Detects Volt/Salt Typhoon Signal-by-Signal")

    doc.add_heading("Volt Typhoon Detection Timeline", level=2)

    vt_phases = [
        ("Phase 1: Initial Access (0-12 hours)",
         "Auth Signal: off_hours_ratio spikes (login at 3 AM from new IP). Network Signal: "
         "VPN gateway shows connection from unusual residential IP (compromised SOHO router). "
         "ACECARD detects auth embedding drift.",),
        ("Phase 2: LOLBin Execution (12-36 hours)",
         "Process Signal: lolbin_count jumps from 0 to 7 (wmic, ntdsutil, netsh, rundll32, "
         "certutil, vssadmin, mshta). Command-line entropy rises (encoded commands). "
         "ACECARD detects process embedding drift toward 'lotl_execution' concept.",),
        ("Phase 3: Lateral Movement (36-72 hours)",
         "Network Signal: unique_destination_IPs jumps from 3 to 18. Auth Signal: Kerberos "
         "ticket requests to new servers. Identity Signal: admin shares accessed. ACECARD "
         "detects multi-signal drift toward 'lateral_movement' concept.",),
        ("Phase 4: Persistence (72-96 hours)",
         "Identity Signal: service_account_use spikes, scheduled tasks created. Process "
         "Signal: registry modification tools used. ACECARD CUSUM triggers -- cumulative "
         "drift across 4 phases exceeds 4-sigma threshold.",),
        ("Phase 5: Pre-positioning (96+ hours)",
         "File Signal: sensitive_access_count rises steadily. Network Signal: low-volume "
         "periodic outbound (beacon_score increases). ACECARD health score drops below 40 "
         "-- CRITICAL alert with drift direction: credential_dumping (0.72), "
         "lateral_movement (0.65), c2_beaconing (0.48).",),
    ]
    for title, desc in vt_phases:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    doc.add_heading("Salt Typhoon Detection Timeline", level=2)
    st_phases = [
        ("Phase 1: Edge Exploitation (0-24 hours)",
         "Network Signal (on router entity): management plane access from anomalous IP. "
         "Auth Signal: new admin session on network device. ACECARD detects device entity "
         "drift toward 'compromised_endpoint.'",),
        ("Phase 2: Credential Harvest (1-7 days)",
         "Auth Signal: SNMP community string access patterns change. Identity Signal: "
         "config pull operations from unusual source. Process Signal: show running-config "
         "equivalent commands executed repeatedly.",),
        ("Phase 3: Tunneling (7-30 days)",
         "Network Signal: beacon_score rises slowly. New protocol mix (GRE where none "
         "existed). Traffic pattern periodicity detected. CUSUM begins accumulating drift. "
         "Still below threshold -- but tracking.",),
        ("Phase 4: Wiretap Access (30-90 days)",
         "Identity Signal: privilege_escalation events on CALEA systems. Admin_actions "
         "spike on systems that rarely see admin activity. CUSUM crosses 4-sigma threshold. "
         "Drift direction: privilege_escalation (0.71), data_exfiltration (0.55).",),
        ("Phase 5: Data Collection (90+ days)",
         "File Signal: steady bytes_out increase over weeks. Network Signal: consistent "
         "small data volumes outbound at regular intervals. ACECARD health score: 28 "
         "(CRITICAL). Kill-chain reconstruction links all 5 phases into coherent narrative.",),
    ]
    for title, desc in st_phases:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    add_insight_box(
        doc,
        "Note that ACECARD detects the attack at DIFFERENT PHASES depending on speed. "
        "Fast attacks (like brute force) trigger within hours. Slow APTs (like Volt "
        "Typhoon's 5-year campaign) trigger via CUSUM accumulation over weeks -- but "
        "still FAR earlier than traditional tools that never detected them at all.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 21
    # =========================================================
    slide_heading(doc, 21, "Drift Direction with 8 MITRE Threat Concepts")

    doc.add_paragraph(
        "This slide details the 8 threat concepts that ACECARD uses to EXPLAIN what "
        "an entity is drifting toward. Each concept is a prose description of a threat "
        "behavior pattern, embedded into the same 1,536-d space as entity behaviors."
    )

    doc.add_heading("The 8 Threat Concepts", level=2)

    concepts = [
        ("1. credential_dumping (T1003: Credential Access)",
         "Reference text: 'Entity behavior resembling credential harvesting: high auth "
         "failures, LSASS access, ticket anomalies, ntdsutil usage.'",
         "What it catches: Mimikatz, LSASS memory dumps, SAM database extraction, DCSync attacks, "
         "Kerberoasting, AS-REP roasting."),
        ("2. lateral_movement (T1021: Lateral Movement)",
         "Reference text: 'Entity behavior resembling lateral traversal: new host "
         "connections, admin share access, remote execution, RDP/SMB/WinRM spikes.'",
         "What it catches: Pass-the-hash, PsExec, WinRM PowerShell remoting, RDP hopping, "
         "administrative share access (C$, ADMIN$)."),
        ("3. data_exfiltration (T1041: Exfiltration)",
         "Reference text: 'Entity behavior resembling data exfiltration: high bytes_out "
         "ratio, bulk file access, archive creation, DNS tunneling indicators.'",
         "What it catches: Large file uploads, staging data in archives (.zip, .rar), DNS "
         "tunneling (encoding data in DNS queries), cloud storage exfil (Dropbox, Google Drive)."),
        ("4. c2_beaconing (T1071: Command & Control)",
         "Reference text: 'Entity behavior resembling C2 communication: periodic "
         "connections, high beacon score, DGA domains, encoded payloads.'",
         "What it catches: Cobalt Strike beacons, Sliver C2, KV Botnet relays, DNS-over-HTTPS "
         "C2 channels, Domain Generation Algorithms (DGA)."),
        ("5. lotl_execution (T1059: Execution)",
         "Reference text: 'Entity behavior resembling living-off-the-land: high LOLBin "
         "count, scripting engines (PowerShell, WScript), encoded command lines.'",
         "What it catches: Volt Typhoon's primary TTP. PowerShell Empire, WMI abuse, certutil "
         "for download, mshta for script execution, rundll32 for DLL loading."),
        ("6. privilege_escalation (T1078: Privilege Escalation)",
         "Reference text: 'Entity behavior resembling privilege escalation: privilege "
         "events, new admin actions, role changes, UAC bypass, token manipulation.'",
         "What it catches: Token impersonation, UAC bypass, sudo abuse, group policy modification, "
         "service account creation, adding users to privileged groups."),
        ("7. defense_evasion (T1562: Defense Evasion)",
         "Reference text: 'Entity behavior resembling defense evasion: log gaps, unsigned "
         "processes, timestamp anomalies, AV/EDR tampering.'",
         "What it catches: Log clearing (wevtutil cl), disabling Windows Defender, timestomping "
         "(changing file timestamps), process hollowing, DLL side-loading."),
        ("8. insider_data_hoarding (T1074: Collection)",
         "Reference text: 'Entity behavior resembling insider staging: bulk sensitive file "
         "access, local staging to USB/temp directories, print volume spikes.'",
         "What it catches: Departing employee downloading all their files, systematic access "
         "to sensitive documents outside role, USB mass storage events, excessive printing."),
    ]
    for title, ref_text, catches in concepts:
        doc.add_heading(title, level=3)
        doc.add_paragraph(ref_text)
        doc.add_paragraph(catches)

    doc.add_heading("How Projection Works", level=2)
    doc.add_paragraph(
        "The drift_vector (V_current - V_baseline) represents HOW the entity has changed. "
        "Each concept has its own embedding vector. The projection score is the cosine "
        "similarity between the drift_vector and the concept_embedding. Scores range from "
        "-1 to +1. A score of 0.7+ indicates strong alignment with that threat concept. "
        "The top-3 concepts are reported with each alert, giving analysts immediate context "
        "about what kind of attack this resembles."
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 22
    # =========================================================
    slide_heading(doc, 22, "Cohort Comparison and Alert Tiers")

    doc.add_heading("Peer Cohort Comparison", level=2)
    doc.add_paragraph(
        "Context matters. A system administrator connecting to 30 servers is normal. A "
        "finance analyst doing the same is alarming. ACECARD compares each entity's "
        "drift against their PEER COHORT -- entities with the same role, organizational "
        "unit, or security group."
    )

    add_key_term(
        doc, "Cohort",
        "A group of similar entities used for comparison. Defined by role (e.g., all "
        "financial analysts), organizational unit (e.g., all users in Marketing), or "
        "security group (e.g., all Domain Admins). Minimum cohort size: 5 entities.",
    )
    add_key_term(
        doc, "Z-Score",
        "A statistical measure of how many standard deviations an entity's drift is "
        "from the cohort average. z = (entity_drift - cohort_mean) / cohort_std. "
        "A z-score > 2.0 means the entity is drifting MORE than 95% of their peers -- "
        "flagged as an outlier.",
    )

    doc.add_heading("Alert Severity Tiers", level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Severity", "Condition", "SOC Action", "Response Time"]):
        table.rows[0].cells[i].text = h
    tiers = [
        ("CRITICAL", "Health < 40", "Immediate investigation", "< 5 minutes"),
        ("HIGH", "Health < 40 OR velocity z > 3", "Tier-2 analyst triage", "< 15 minutes"),
        ("MEDIUM", "Health < 70 AND recent change-point", "Tier-1 review", "< 1 hour"),
        ("LOW", "Health 70-85 AND cohort z > 2", "Watchlist + auto-enrich", "Next shift"),
        ("INFO", "Any CUSUM trigger (4-sigma)", "Log for trend analysis", "Batch review"),
    ]
    for i, (sev, cond, action, time) in enumerate(tiers):
        row = table.rows[i + 1]
        row.cells[0].text = sev
        row.cells[1].text = cond
        row.cells[2].text = action
        row.cells[3].text = time

    add_key_term(
        doc, "SOC (Security Operations Center)",
        "The team of analysts who monitor security alerts 24/7. Typically organized in tiers: "
        "Tier-1 handles initial triage, Tier-2 handles complex investigations, Tier-3 handles "
        "advanced threat hunting and incident response.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 23
    # =========================================================
    slide_heading(doc, 23, "ABAC Trust Loop")

    doc.add_heading("What is ABAC?", level=2)
    add_key_term(
        doc, "ABAC (Attribute-Based Access Control)",
        "An access control model where permissions are determined by ATTRIBUTES of the user, "
        "resource, and environment -- rather than fixed roles (RBAC). In ACECARD, one of "
        "those attributes is the entity's behavioral trust state. A user whose behavior is "
        "drifting gets automatically restricted without any human action needed.",
    )

    doc.add_heading("The Self-Reinforcing Loop", level=2)
    doc.add_paragraph("ACECARD's feedback loop works in 6 steps:")

    steps_fb = [
        ("1. ALERT", "CUSUM/health/velocity triggers an alert to the SOC dashboard."),
        ("2. TRIAGE", "Analyst investigates using the 7-chart dashboard -- sees drift timeline, signal decomposition, peer comparison, threat concept alignment."),
        ("3. VERDICT", "Analyst marks the alert as TP (True Positive = real threat) or FP (False Positive = benign). Adds context notes."),
        ("4. LEARN", "Labeled data feeds back into the threshold tuning model. TP verdicts strengthen the signal weights. FP verdicts raise thresholds for that entity/role combination to reduce future false alarms."),
        ("5. TAG", "Entity's trust state is updated in the ABAC engine based on the verdict and current health score."),
        ("6. ENFORCE", "ABAC automatically adjusts the entity's access: step-up MFA, restrict to read-only, or terminate session -- proportional to risk level."),
    ]
    for title, desc in steps_fb:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    doc.add_heading("Four Trust States", level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Trust State", "Condition", "Access Level", "Enforcement"]):
        table.rows[0].cells[i].text = h
    states = [
        ("TRUSTED", "Health >= 70, no recent alerts", "Full normal access", "Standard MFA"),
        ("ELEVATED_WATCH", "Health 40-70 or cohort z > 2", "Normal + monitoring", "Step-up MFA on sensitive resources"),
        ("RESTRICTED", "Health < 40 or TP alert", "Read-only access", "No lateral movement, no downloads"),
        ("BLOCKED", "Active incident confirmed", "No access", "Account disabled, session terminated"),
    ]
    for i, (state, cond, access, enforce) in enumerate(states):
        row = table.rows[i + 1]
        row.cells[0].text = state
        row.cells[1].text = cond
        row.cells[2].text = access
        row.cells[3].text = enforce

    add_insight_box(
        doc,
        "The trust state is CONTINUOUS, not binary. Most systems are either 'blocked' or "
        "'not blocked.' ACECARD provides PROPORTIONAL response -- escalating restrictions "
        "as confidence grows. This means earlier intervention with less disruption. A user "
        "in ELEVATED_WATCH still works normally but gets an extra MFA prompt for sensitive "
        "systems. If behavior worsens, they move to RESTRICTED. Only confirmed incidents "
        "cause BLOCKED.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDES 24-25
    # =========================================================
    slide_heading(doc, "24-25", "Combined Architecture and Data Flow")

    doc.add_paragraph(
        "This section shows how Rigor and ACECARD work together as an integrated system, "
        "not as two separate tools."
    )

    doc.add_heading("Data Flow: Prevention Layer (Rigor)", level=2)
    add_bullet(doc, "Configuration sources: Firewall configs (Palo Alto, Fortinet, Check Point), Cloud security groups (AWS, Azure, GCP), IdP policies (Okta, Azure AD), Network segmentation (VLANs, micro-seg), SD-WAN policies")
    add_bullet(doc, "Processing: Symbolic model construction -> Formal verification -> Attack path enumeration")
    add_bullet(doc, "Outputs: Proven attack paths, Compliance verification results, Change impact predictions, Remediation recommendations")
    add_bullet(doc, "Handoff to ACECARD: Rigor's findings inform ACECARD's alert context -- 'this entity is drifting AND sits on a proven attack path'")

    doc.add_heading("Data Flow: Detection Layer (ACECARD)", level=2)
    add_bullet(doc, "Event sources: SIEM/log forwarders (Splunk, Sentinel, ELK), EDR (CrowdStrike, Defender), Network sensors (Zeek, NetFlow), Cloud audit (CloudTrail, Azure Activity Log)")
    add_bullet(doc, "Processing: Schema normalization (ECS/OCSF) -> Signal extraction (5 signals) -> Embedding (1536-d) -> Trajectory analysis -> Drift direction")
    add_bullet(doc, "Outputs: Alerts (Sigma-compatible), Trust tags (ABAC), Dashboard (7 charts), Kill-chain narratives")
    add_bullet(doc, "Handoff to Rigor: ACECARD detections trigger Rigor re-analysis -- 'this entity drifted; what paths exist FROM their current position?'")

    doc.add_heading("Integration Points", level=2)
    add_bullet(doc, "Rigor to ACECARD: Attack path context enriches alerts ('entity has access to 3 proven paths to domain controller')")
    add_bullet(doc, "ACECARD to Rigor: Behavioral alerts trigger path re-analysis from the drifting entity's position")
    add_bullet(doc, "Shared: Common entity identifiers allow correlation between configuration risk and behavioral risk")
    add_bullet(doc, "Enforcement: ABAC tags from ACECARD can trigger Rigor re-evaluation ('if we restrict this entity, which paths close?')")

    doc.add_heading("Deployment Architecture", level=2)
    add_key_term(
        doc, "Iron Bank",
        "The DoD's repository of hardened, approved container images. All containers "
        "must pass security scanning and vulnerability assessment before being added.",
    )
    add_key_term(
        doc, "Big Bang",
        "A Helm chart (Kubernetes deployment package) that bundles DoD-approved tools "
        "into a single deployable stack.",
    )
    add_key_term(
        doc, "IL5 / IL6 / JWICS",
        "Impact Levels define the sensitivity of data that can be processed. "
        "IL5 = NIPR (unclassified but sensitive), IL6 = SIPR (Secret), JWICS = "
        "Joint Worldwide Intelligence Communications System (Top Secret/SCI). "
        "The system deploys across ALL levels with appropriate isolation.",
    )
    add_key_term(
        doc, "cATO (Continuous Authority to Operate)",
        "Instead of a one-time security assessment (traditional ATO that takes 12-18 months), "
        "cATO provides ongoing authorization through continuous monitoring. DoD CIO policy "
        "from February 2022. Enables faster deployment of new capabilities.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDES 26-27
    # =========================================================
    slide_heading(doc, "26-27", "Attack Walkthroughs -- Both Layers Working Together")

    doc.add_heading("Walkthrough: Volt Typhoon with Rigor + ACECARD", level=2)

    doc.add_heading("Before the Attack (Rigor)", level=3)
    doc.add_paragraph(
        "Rigor continuously analyzes the network configuration. It identifies: "
        "'Workstation VLAN can reach Domain Controller via SMB (port 445) -- this path "
        "should be restricted to IT admin VLAN only.' Report sent to security team. "
        "If remediated: attack path is CLOSED before Volt Typhoon arrives."
    )

    doc.add_heading("If Not Remediated -- During the Attack (ACECARD)", level=3)
    doc.add_paragraph(
        "Volt Typhoon compromises an edge device and begins LOtL operations. "
        "ACECARD detects:"
    )
    add_bullet(doc, "Hour 6: Auth signal drift -- new source IP, off-hours login")
    add_bullet(doc, "Hour 18: Process signal drift -- LOLBin count spikes (lolbin_count: 0 -> 5)")
    add_bullet(doc, "Hour 30: Network signal drift -- new destination IPs (unique_dest: 3 -> 12)")
    add_bullet(doc, "Hour 42: CUSUM triggers -- cumulative drift exceeds 4-sigma")
    add_bullet(doc, "Hour 42: Drift direction report: lateral_movement (0.78), credential_dumping (0.65)")
    add_bullet(doc, "Hour 42: ABAC trust state: TRUSTED -> RESTRICTED")
    add_bullet(doc, "Hour 42: Automatic enforcement: lateral movement ports CLOSED for this entity")

    doc.add_heading("Combined Outcome", level=3)
    doc.add_paragraph(
        "Without either tool: Volt Typhoon operates undetected for 5+ years. "
        "With Rigor alone: The specific path may be closed, but if the attacker finds "
        "another (via zero-day), they proceed undetected. "
        "With ACECARD alone: Detected in ~42 hours, but the path existed. "
        "With BOTH: Either the path never existed (Rigor closed it) OR the attack is "
        "detected and contained within hours (ACECARD + ABAC). Defense in depth."
    )

    doc.add_heading("Walkthrough: Salt Typhoon with Rigor + ACECARD", level=2)

    doc.add_heading("Before the Attack (Rigor)", level=3)
    doc.add_paragraph(
        "Rigor analyzes router management plane access. Identifies: 'Cisco router R7 "
        "management interface (SSH port 22) is reachable from VLAN 40 (contractor VLAN) "
        "due to permissive ACL rule #23. This should be restricted to VLAN 99 (NOC only).' "
        "Report sent. If remediated: Salt Typhoon cannot reach the router management plane."
    )

    doc.add_heading("If Not Remediated -- During the Attack (ACECARD)", level=3)
    doc.add_paragraph("Salt Typhoon exploits the router and begins infrastructure compromise:")
    add_bullet(doc, "Day 2: Network entity (router R7) shows anomalous management session from non-NOC IP")
    add_bullet(doc, "Day 5: Auth signal drift on router entity -- SNMP access pattern change")
    add_bullet(doc, "Day 14: Network signal -- new GRE tunnel endpoint detected (beacon_score rises)")
    add_bullet(doc, "Day 21: CUSUM triggers on router entity -- sustained low drift accumulates")
    add_bullet(doc, "Day 21: Drift direction: compromised_endpoint (0.69), c2_beaconing (0.58)")
    add_bullet(doc, "Day 21: Alert to SOC: 'Router R7 behavioral drift -- investigate immediately'")
    add_bullet(doc, "Day 21: Rigor triggered for re-analysis FROM router R7's position -- identifies all assets reachable")

    doc.add_page_break()

    # =========================================================
    # SLIDE 28
    # =========================================================
    slide_heading(doc, 28, "Combined Coverage Matrix")

    doc.add_paragraph(
        "This slide provides a comprehensive matrix showing which attack techniques "
        "are covered by which layer (Rigor, ACECARD, or both)."
    )

    table = doc.add_table(rows=13, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["MITRE Tactic", "Example Techniques", "Rigor Coverage", "ACECARD Coverage"]):
        table.rows[0].cells[i].text = h
    matrix = [
        ("Initial Access", "Exploit public app, Phishing", "Proves paths are closed", "Detects auth drift post-exploit"),
        ("Execution", "LOLBins, Scripting, WMI", "Verifies app-control", "Detects process signal shift"),
        ("Persistence", "Scheduled tasks, Registry", "Verifies config integrity", "Detects identity drift"),
        ("Privilege Escalation", "Token manipulation, UAC bypass", "Verifies least-privilege", "Detects privilege signal drift"),
        ("Defense Evasion", "Log clearing, AV disable", "Verifies audit integrity", "Detects log gaps as signal"),
        ("Credential Access", "Dumping, Kerberoasting", "Verifies credential isolation", "Detects auth pattern change"),
        ("Discovery", "Network scanning, AD recon", "Verifies segmentation", "Detects network signal change"),
        ("Lateral Movement", "RDP, SMB, WinRM, PsExec", "Proves/blocks paths", "Detects multi-entity drift"),
        ("Collection", "File staging, Screen capture", "N/A", "Detects file signal drift"),
        ("Exfiltration", "DNS tunnel, Cloud upload", "Verifies outbound restrictions", "Detects bytes_out drift"),
        ("Command & Control", "Beaconing, DGA, HTTPS", "Verifies egress policies", "Detects beacon_score drift"),
        ("Impact", "Ransomware, Wiper", "Verifies backup isolation", "Detects rapid file changes"),
    ]
    for i, (tactic, tech, rigor, acecard) in enumerate(matrix):
        row = table.rows[i + 1]
        row.cells[0].text = tactic
        row.cells[1].text = tech
        row.cells[2].text = rigor
        row.cells[3].text = acecard

    add_insight_box(
        doc,
        "Every MITRE ATT&CK tactic has coverage from at least one layer, and most have "
        "coverage from BOTH layers. This is true defense-in-depth -- if one layer misses "
        "(Rigor cannot detect ongoing credential abuse; ACECARD cannot prevent a "
        "misconfigured firewall), the other layer catches it.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 29
    # =========================================================
    slide_heading(doc, 29, "Complementary Tools")

    doc.add_paragraph(
        "Rigor + ACECARD does not REPLACE your existing security stack -- it adds the "
        "two capabilities that are structurally missing. Here is how it integrates with "
        "tools you likely already have."
    )

    doc.add_heading("Integration with Existing Tools", level=2)
    tools = [
        ("SIEM (Splunk, Sentinel, QRadar)",
         "ACECARD consumes events FROM your SIEM. It adds behavioral intelligence that "
         "SIEM rules cannot provide. Alerts feed BACK into the SIEM as enriched events."),
        ("EDR (CrowdStrike, Defender, SentinelOne)",
         "EDR provides the raw process, file, and network telemetry that feeds ACECARD's "
         "signal serializers. EDR detects known malware; ACECARD detects behavioral drift "
         "from legitimate-tool abuse."),
        ("SOAR (Splunk SOAR, Palo Alto XSOAR, Swimlane)",
         "ACECARD alerts trigger automated playbooks in your SOAR platform. Example: "
         "CRITICAL alert -> isolate host + notify analyst + collect forensic artifacts."),
        ("IAM/IdP (Okta, Azure AD, Ping Identity)",
         "ACECARD's ABAC trust tags integrate with your identity provider. Step-up MFA "
         "enforcement happens at the IdP level. RESTRICTED state = conditional access policy."),
        ("NDR (Darktrace, Vectra, ExtraHop)",
         "Network Detection and Response tools provide additional network telemetry. "
         "ACECARD differs by tracking TRAJECTORY (change over time) rather than point-in-time anomaly."),
        ("Vulnerability Scanner (Qualys, Tenable, Rapid7)",
         "Rigor uses vulnerability scan results to inform its attack path analysis. "
         "A vulnerable system on a proven attack path = CRITICAL finding."),
        ("Firewall/NGFW (Palo Alto, Fortinet, Check Point)",
         "Rigor directly ingests firewall rule sets for formal verification. "
         "Configuration changes trigger immediate re-analysis."),
    ]
    for tool, desc in tools:
        doc.add_heading(tool, level=3)
        doc.add_paragraph(desc)

    add_key_term(
        doc, "SOAR (Security Orchestration, Automation, and Response)",
        "Tools that automate incident response workflows. When an alert fires, SOAR "
        "automatically runs a 'playbook' of response actions (isolate host, block IP, "
        "notify team, collect logs) without waiting for a human.",
    )
    add_key_term(
        doc, "NDR (Network Detection and Response)",
        "Security tools that analyze network traffic for threats. Unlike firewalls (which "
        "block/allow), NDR passively monitors traffic to detect anomalies.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 30
    # =========================================================
    slide_heading(doc, 30, "Reference Architecture")

    doc.add_paragraph(
        "This slide shows the complete deployment architecture -- how Rigor + ACECARD "
        "is deployed in a DoD-compliant environment."
    )

    doc.add_heading("Components", level=2)
    add_bullet(doc, "Container runtime: Kubernetes (K8s) with Big Bang baseline, Iron Bank images", "Infrastructure")
    add_bullet(doc, "Database: PostgreSQL with pgvector extension for embedding storage", "Storage")
    add_bullet(doc, "Message queue: Kafka or RabbitMQ for event streaming", "Messaging")
    add_bullet(doc, "API: FastAPI (Python) REST endpoints under /api/trajectory/*", "Compute")
    add_bullet(doc, "Embedding: OpenAI API (IL5) or local SLM (IL6/JWICS)", "AI")
    add_bullet(doc, "Visualization: Plotly.js 7-chart dashboard", "UI")
    add_bullet(doc, "Auth: Bearer token (dev), CAC/PIV via mTLS (production)", "Security")
    add_bullet(doc, "Observability: Prometheus metrics, OpenTelemetry traces, structured JSON logs", "Monitoring")
    add_bullet(doc, "CI/CD: FluxCD GitOps, SBOM generation, vulnerability scanning (Anchore/Grype)", "Pipeline")

    doc.add_heading("Performance", level=2)
    add_bullet(doc, "Embed + analyze + interpret: < 3 seconds per entity window")
    add_bullet(doc, "Throughput: 10,000 entities/hour on 4-vCPU node")
    add_bullet(doc, "Horizontal scaling via Kubernetes replicas")
    add_bullet(doc, "Dashboard chart load: < 500ms (pre-computed Plotly JSON)")

    doc.add_heading("Multi-Enclave Deployment", level=2)
    doc.add_paragraph(
        "The system deploys identically across classification levels. At IL5 (NIPR), "
        "OpenAI's embedding API is used. At IL6 (SIPR) and on JWICS (TS/SCI), a local "
        "Small Language Model (Phi-4, Mistral, or sentence-transformers) provides the "
        "same 1,536-d embeddings without requiring any external API call. The embedding "
        "provider is swapped via environment variable -- no code change needed."
    )

    add_key_term(
        doc, "SLM (Small Language Model)",
        "A compact AI model that can run on local hardware without cloud connectivity. "
        "Examples: Microsoft Phi-4, Mistral 7B. Used for air-gapped or classified "
        "environments where external API calls are not permitted.",
    )
    add_key_term(
        doc, "SBOM (Software Bill of Materials)",
        "A complete list of all software components, libraries, and dependencies in an "
        "application. Required for DoD software supply chain security. Like a food "
        "ingredient label, but for software.",
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 31
    # =========================================================
    slide_heading(doc, 31, "Verdict")

    doc.add_heading("The Core Argument", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "Nation-state adversaries (Volt Typhoon, Salt Typhoon) have demonstrated that "
        "current defense approaches -- signature-based detection, rule-based alerts, "
        "sampling-based assessment -- are structurally incapable of stopping them."
    )
    run.bold = True

    doc.add_paragraph(
        "The evidence is clear: Volt Typhoon operated inside U.S. critical infrastructure "
        "for 5+ years undetected. Salt Typhoon compromised major telecom carriers and "
        "accessed law enforcement wiretap systems. Billions of dollars in existing security "
        "tools failed to detect them."
    )

    doc.add_heading("The Solution Requires Two Fundamentally Different Approaches", level=2)

    doc.add_heading("1. Mathematical Prevention (Rigor AI)", level=3)
    add_bullet(doc, "Exhaustively verifies the 2^8000 configuration state space -- no sampling")
    add_bullet(doc, "PROVES attack paths exist or do not exist -- mathematical certainty")
    add_bullet(doc, "Closes paths BEFORE attackers arrive")
    add_bullet(doc, "Limitation: Cannot detect attacks in progress using valid credentials")

    doc.add_heading("2. Behavioral Detection (ACECARD UEBA)", level=3)
    add_bullet(doc, "Tracks behavioral trajectory of every entity via 1,536-d embeddings")
    add_bullet(doc, "CUSUM catches slow drift that accumulates over months (APT pattern)")
    add_bullet(doc, "Drift direction explains WHAT is happening (not just 'something is weird')")
    add_bullet(doc, "ABAC trust loop provides automatic, proportional enforcement")
    add_bullet(doc, "Limitation: Cannot prevent attacks -- only detects and responds")

    doc.add_heading("Together", level=3)
    p = doc.add_paragraph()
    run = p.add_run(
        "Rigor prevents what can be prevented. ACECARD detects what cannot be prevented. "
        "Together they cover the complete attack lifecycle with no structural gaps."
    )
    run.bold = True
    run.font.size = Pt(12)

    doc.add_heading("Proven Foundation", level=2)
    doc.add_paragraph(
        "This is not theoretical. The ACECARD trajectory engine is directly adapted from "
        "the E-11 Temporal Trajectory Intelligence module built for DLA (Defense Logistics "
        "Agency), which has been running in production with 500+ entities, 23 monthly "
        "snapshots, 5 behavioral signals, and CUSUM change-point detection. Same "
        "architecture, same algorithms, different signals -- applied to cyber defense."
    )

    doc.add_page_break()

    # =========================================================
    # SLIDE 32
    # =========================================================
    slide_heading(doc, 32, "Next Steps -- Build Your Dome")

    doc.add_paragraph(
        "Preempt the next Typhoon. Detect what slips past."
    )

    doc.add_heading("Who Delivers This", level=2)
    doc.add_paragraph(
        "Rigor AI is engaged with 20+ design partners across Fortune 500, federal "
        "agencies, critical infrastructure, MSSPs, and nation states."
    )
    doc.add_paragraph(
        "ACECARD UEBA is operationally tested by 22nd Century Technologies (TSCTI) -- "
        "the company running:"
    )
    add_bullet(doc, "$90M U.S. Army SOC/MDR contract with 800+ cleared analysts")
    add_bullet(doc, "USAF $108M cybersecurity operations")
    add_bullet(doc, "FBI TSC $56M operations")
    add_bullet(doc, "NAVAIR/USMC $145M cybersecurity")

    doc.add_heading("Four-Step Engagement", level=2)
    add_bullet(doc, "Step 1 -- Pilot Rigor AI: Formal policy analysis on your edge defense estate against Volt + Salt Typhoon TTPs. Immediate finding report.")
    add_bullet(doc, "Step 2 -- Pilot ACECARD UEBA: Replay Volt + Salt Typhoon scenarios on your telemetry. Tune thresholds and cohorts to your environment.")
    add_bullet(doc, "Step 3 -- Integrate: Rigor's verified findings raise ACECARD's trust floor on entities traversing closed gaps. Both layers inform each other.")
    add_bullet(doc, "Step 4 -- Operationalize: SaaS / on-prem / IL5 / IL6 / JWICS / cATO. Continuous re-verification. Measurable risk reduction.")

    doc.add_heading("Timeline", level=2)
    add_bullet(doc, "Demo environment (synthetic Volt/Salt Typhoon scenarios, full pipeline): Available NOW")
    add_bullet(doc, "Pilot (single enclave, real data, limited entity count): 4-6 weeks from authorization")
    add_bullet(doc, "Production (full entity coverage, ABAC enforcement, multi-enclave): 3-4 months")

    doc.add_heading("What We Need from You", level=2)
    add_bullet(doc, "Access to existing log sources (SIEM forwarding rules or API keys)")
    add_bullet(doc, "Active Directory / IAM export for cohort definition (role, OU, group membership)")
    add_bullet(doc, "Firewall/NGFW configuration exports for Rigor formal analysis")
    add_bullet(doc, "SOC analyst engagement for feedback loop (TP/FP verdicts)")
    add_bullet(doc, "Kubernetes cluster access (or Gabriel Nimbus environment) for deployment")

    doc.add_page_break()

    # =========================================================
    # APPENDIX A: COMPLETE ACRONYM GLOSSARY
    # =========================================================
    doc.add_heading("Appendix A: Complete Acronym Glossary", level=1)

    glossary = [
        ("ABAC", "Attribute-Based Access Control -- access decisions based on entity attributes (role, location, behavioral trust state)"),
        ("ACL", "Access Control List -- rules on a router/firewall determining what traffic is allowed or blocked"),
        ("AD", "Active Directory -- Microsoft's directory service managing all user accounts and permissions in a Windows network"),
        ("APT", "Advanced Persistent Threat -- prolonged targeted attack by a well-resourced adversary (nation-state)"),
        ("ARCYBER", "U.S. Army Cyber Command"),
        ("ArCTIC", "Army Cyber Technology and Innovation Center"),
        ("ASM", "Attack Surface Management -- tools that map exposed assets and attack vectors"),
        ("BAS", "Breach and Attack Simulation -- tools that simulate attacks against defenses (sampling-based)"),
        ("C2 / C&C", "Command and Control -- communication channel between attacker and compromised systems"),
        ("CALEA", "Communications Assistance for Law Enforcement Act -- requires telecoms to build wiretap capability"),
        ("cATO", "Continuous Authority to Operate -- ongoing authorization through continuous monitoring (vs. one-time ATO)"),
        ("CDR", "Call Detail Records -- metadata for phone calls (who, when, duration, location)"),
        ("CFIC", "Cyber Futures Innovation Center"),
        ("CISA", "Cybersecurity and Infrastructure Security Agency (U.S.)"),
        ("CMMC", "Cybersecurity Maturity Model Certification -- DoD supply chain security standard"),
        ("CUSUM", "Cumulative Sum -- statistical algorithm detecting sustained small changes over time"),
        ("CVE", "Common Vulnerabilities and Exposures -- standardized ID for known security flaws"),
        ("DGA", "Domain Generation Algorithm -- malware technique creating random domain names for C2"),
        ("DLL", "Dynamic Link Library -- shared code library on Windows"),
        ("DLP", "Data Loss Prevention -- software blocking sensitive data from leaving the organization"),
        ("DMZ", "Demilitarized Zone -- network segment between internal network and internet"),
        ("DNS", "Domain Name System -- translates website names to IP addresses"),
        ("DoD", "Department of Defense"),
        ("ECS", "Elastic Common Schema -- standardized event format for security logs"),
        ("EDR", "Endpoint Detection and Response -- security software monitoring computers for suspicious activity"),
        ("GRE", "Generic Routing Encapsulation -- protocol creating network tunnels"),
        ("HNSW", "Hierarchical Navigable Small World -- fast algorithm for searching vectors"),
        ("HTTPS", "Hypertext Transfer Protocol Secure -- encrypted web communication"),
        ("IAM", "Identity and Access Management -- systems managing who can access what"),
        ("IdP", "Identity Provider -- system verifying identity (Okta, Azure AD, Ping)"),
        ("IL5/IL6", "Impact Level 5/6 -- DoD data sensitivity classifications (CUI / Secret)"),
        ("IOC", "Indicator of Compromise -- evidence of breach (IP address, file hash, domain)"),
        ("JWICS", "Joint Worldwide Intelligence Communications System -- Top Secret/SCI network"),
        ("K8s", "Kubernetes -- container orchestration platform"),
        ("KEV", "Known Exploited Vulnerabilities -- CISA catalog of actively exploited flaws"),
        ("KV Botnet", "Volt Typhoon's network of compromised SOHO routers used for C2"),
        ("LOLBin", "Living-off-the-Land Binary -- legitimate system tool abused by attackers"),
        ("LOtL", "Living off the Land -- attack technique using only built-in tools"),
        ("LSASS", "Local Security Authority Subsystem Service -- Windows process storing credentials in memory"),
        ("MFA", "Multi-Factor Authentication -- requiring multiple proofs of identity"),
        ("MITRE ATT&CK", "Knowledge base of adversary tactics and techniques based on real-world observations"),
        ("mTLS", "Mutual TLS -- both client and server present certificates for authentication"),
        ("NDR", "Network Detection and Response -- passive network traffic analysis for threats"),
        ("NGFW", "Next-Generation Firewall -- inspects traffic at application level"),
        ("NIST", "National Institute of Standards and Technology"),
        ("OCSF", "Open Cybersecurity Schema Framework -- AWS-led event schema standard"),
        ("OU", "Organizational Unit -- hierarchical grouping in Active Directory"),
        ("PII", "Personally Identifiable Information -- data that can identify an individual"),
        ("RAT", "Remote Access Trojan -- malware giving attacker remote control"),
        ("RBAC", "Role-Based Access Control -- access decisions based on assigned roles"),
        ("RDP", "Remote Desktop Protocol -- full graphical remote access to computers"),
        ("SASE", "Secure Access Service Edge -- cloud-delivered network security model"),
        ("SBOM", "Software Bill of Materials -- inventory of all software components"),
        ("SIEM", "Security Information and Event Management -- centralized log collection and analysis"),
        ("SLM", "Small Language Model -- compact AI model for local/air-gapped deployment"),
        ("SMB", "Server Message Block -- Windows file sharing protocol"),
        ("SNMP", "Simple Network Management Protocol -- standard for monitoring network devices"),
        ("SOC", "Security Operations Center -- team monitoring security alerts 24/7"),
        ("SOAR", "Security Orchestration, Automation, and Response -- automated incident response"),
        ("SOHO", "Small Office/Home Office -- refers to residential routers and small business equipment"),
        ("SSL/TLS", "Secure Sockets Layer / Transport Layer Security -- encryption for internet communication"),
        ("STIX", "Structured Threat Information Expression -- standard for sharing threat intelligence"),
        ("TTP", "Tactics, Techniques, and Procedures -- three-level description of attacker behavior"),
        ("UEBA", "User and Entity Behavior Analytics -- security analytics detecting behavioral deviations"),
        ("VLAN", "Virtual Local Area Network -- logical network segmentation"),
        ("VPN", "Virtual Private Network -- encrypted connection between remote user and corporate network"),
        ("WAF", "Web Application Firewall -- protects web applications by filtering HTTP traffic"),
        ("WinRM", "Windows Remote Management -- allows running commands on remote computers"),
        ("Zero-Day", "Vulnerability unknown to vendor -- no patch available, zero days to prepare"),
    ]
    for term, defn in glossary:
        add_key_term(doc, term, defn)

    doc.add_page_break()

    # =========================================================
    # APPENDIX B: MITRE ATT&CK PRIMER
    # =========================================================
    doc.add_heading("Appendix B: MITRE ATT&CK Framework Primer", level=1)

    doc.add_paragraph(
        "The MITRE ATT&CK framework is the standard language used by security teams "
        "worldwide to describe how attackers operate. It organizes attack behaviors into "
        "a matrix of Tactics (columns) and Techniques (rows). Understanding this framework "
        "helps you interpret ACECARD's drift direction reports."
    )

    doc.add_heading("The 14 Tactics (Attack Phases)", level=2)
    tactics = [
        ("Reconnaissance", "Attacker gathers information about the target before attacking"),
        ("Resource Development", "Attacker acquires infrastructure and tools for the operation"),
        ("Initial Access", "Attacker gets their first foothold inside the network"),
        ("Execution", "Attacker runs malicious code or commands"),
        ("Persistence", "Attacker maintains access even after reboots or password changes"),
        ("Privilege Escalation", "Attacker gains higher-level permissions"),
        ("Defense Evasion", "Attacker avoids detection by security tools"),
        ("Credential Access", "Attacker steals usernames and passwords"),
        ("Discovery", "Attacker maps the network to find valuable targets"),
        ("Lateral Movement", "Attacker moves between systems within the network"),
        ("Collection", "Attacker gathers data of interest for theft"),
        ("Command & Control", "Attacker establishes communication with compromised systems"),
        ("Exfiltration", "Attacker steals the collected data"),
        ("Impact", "Attacker disrupts, destroys, or manipulates systems/data"),
    ]
    for tactic, desc in tactics:
        add_key_term(doc, tactic, desc)

    doc.add_heading("How ACECARD Uses ATT&CK", level=2)
    doc.add_paragraph(
        "Each of ACECARD's 8 drift direction concepts maps to one or more ATT&CK "
        "techniques. When an alert fires, the report includes the specific T-numbers "
        "(technique IDs) so analysts can immediately look up the full ATT&CK documentation "
        "for investigation procedures, known adversaries who use that technique, and "
        "recommended mitigations."
    )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("--- End of Document ---")
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


if __name__ == "__main__":
    doc = build_document()
    output_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Rigor_ACECARD_Business_Guide.docx")
    doc.save(output_path)
    print(f"Document saved: {output_path}")
    print(f"Size: {os.path.getsize(output_path) / 1024:.1f} KB")
