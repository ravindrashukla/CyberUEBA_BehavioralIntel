# -*- coding: utf-8 -*-
"""Apply 5 rounds of review refinements to the US Army / Pentagon cyber whitepaper."""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

P = "WP DLA/Presentation/V-Intelligence_Layered_Cybersecurity_Whitepaper_US_Army_Pentagon.docx"
d = Document(P)

def acc(p): return "".join(t.text or "" for t in p._p.iter(qn('w:t')))
def find(prefix):
    for p in d.paragraphs:
        if acc(p).strip().startswith(prefix): return p
def find_contains(s):
    for p in d.paragraphs:
        if s in acc(p): return p
def settext(p, new):
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]: r.text = ""
    else:
        p.add_run(new)
def ins_after(anchor, text, style):
    np = OxmlElement('w:p'); anchor._p.addnext(np)
    pp = Paragraph(np, anchor._parent); pp.style = d.styles[style]
    if text: pp.add_run(text)
    return pp
def ins_before(anchor, text, style):
    np = OxmlElement('w:p'); anchor._p.addprevious(np)
    pp = Paragraph(np, anchor._parent); pp.style = d.styles[style]
    if text: pp.add_run(text)
    return pp

BODY, H1, H2, BUL = "Normal", "Heading 1", "Heading 2", "List Bullet"
log = []

# ROUND 1: structural
h13 = find("1.3 Two Structural Gaps")
a = ins_after(h13, "The adversary exploits two structural gaps the prior model leaves open.", BODY)
a = ins_after(a, "First, a preventive-proof gap. Prevention tools sample the configuration and attack surface; they never prove it is closed. A path that was not sampled, or a rule that drifted after the last test, stays invisible until it is used. The adversary needs only the one path no one proved was shut.", BODY)
a = ins_after(a, "Second, a behavioral-detection gap. Detection tools threshold on magnitude and watch for spikes. A nation-state operator using valid credentials and living-off-the-land techniques produces no spike, fragments one identity across many systems, and stays below every alert for years.", BODY)
a = ins_after(a, "The two halves of the architecture close these gaps in turn: preventive proof at the edge, and behavioral watch inside.", BODY)
log.append("R1: filled 1.3")

h82 = find("8.2 How to Defend")
ins_after(h82, "Each layer answers a specific class of AI-enabled adversary move, in the shared MITRE vocabulary: preemptive proof at the network edge (Layer 1), behavioral detection of valid-account and living-off-the-land abuse (Layer 2), runtime protection of applications and APIs (Layer 3), governance of data and the agents that reach it (Layer 4), and assured code before deployment (Layer 5). The defensive AI at each layer is itself constrained and auditable, as Section 8.4 describes.", BODY)
log.append("R1: filled 8.2")

p9 = find("No single product closes this")
settext(p9, acc(p9).replace(
    "Detection tools watch for spikes that slow, credential-borne intrusions are designed never to produce.",
    "Detection tools watch for spikes; slow, credential-borne intrusions are designed never to produce one."))
p10 = find("This paper presents a layered answer")
settext(p10, acc(p10).replace(
    "The proprietary internals of our methods are deliberately summarized at the capability level: the proprietary internals (representations, formulas, dimensions, and scoring logic) are not disclosed here, while architectural structure is described only where it aids understanding.",
    "The proprietary internals of our methods (representations, formulas, dimensions, and scoring logic) are deliberately summarized at the capability level and are not disclosed here, while architectural structure is described only where it aids understanding."))
log.append("R1: fixed exec-summary garble + doubled clause")

# ROUND 2: Layer 2 proof + Layer 1 bridge
ill = find("Illustrative result. In an internal blind")
ins_after(ill, "In that evaluation, all four nation-state-style attackers were surfaced (4 of 4) at an 8.1 percent false-positive operating point, where classical point-anomaly methods (Isolation Forest, One-Class SVM, and Local Outlier Factor) caught none of the four and a z-score method caught one. The four surfaced at ranks 1, 2, 7, and 24 on a risk-ordered list of 250 entities. These are controlled, synthetic-data proof-of-concept results, indicative rather than a guarantee of field performance.", BODY)
log.append("R2: added Layer 2 proof numbers")
pb = find("The approach is operational across financial services")
if pb:
    settext(pb, "Beyond the design, the approach is operational across government, critical infrastructure, financial services, and technology, and has surfaced lurking, exploitable misconfigurations that human review and existing tools had missed. Representative findings:")
log.append("R2: bridged Layer 1 proof points")

# ROUND 3: Army framing
p8 = find("The cybersecurity contest has changed shape")
settext(p8, acc(p8) + " For the U.S. Army and the Pentagon this is not abstract: these campaigns pre-position in the commercial critical infrastructure that mobilization and power projection depend on, and they target the enterprise and tactical networks, and the IL5, IL6, and JWICS enclaves, that Army Cyber Command (ARCYBER) defends.")
pc = find("The adversary has changed the game on two axes")
settext(pc, acc(pc).replace(
    "critical-infrastructure defenders move from years of undetected access to immediate identification",
    "Army and Pentagon cyber defenders, and the critical infrastructure they protect, move from years of undetected access to rapid identification"))
log.append("R3: Army framing in exec summary + conclusion")

# ROUND 4: accuracy / naming
for tb in d.tables:
    for r in tb.rows:
        for c in r.cells:
            for p in c.paragraphs:
                t = acc(p)
                if t.strip() == "What it proves or catches":
                    settext(p, "Mission objective")
                if "22nd Century's behavioral-twin detection approach" in t:
                    settext(p, t.replace(
                        "22nd Century's behavioral-twin detection approach (the Layer 2 product).",
                        "V-Intelligence's behavioral-twin detection approach (the Layer 2 capability)."))
pm = find("Mythos is the V-Intelligence code and supply-chain assurance capability")
if pm:
    settext(pm, acc(pm).replace(
        "Mythos is the V-Intelligence code and supply-chain assurance capability.",
        "Mythos is the V-Intelligence code and supply-chain assurance capability, a defensive capability distinct from any similarly named offensive tooling."))
p80 = find_contains("on the order of 80 percent of attacks when their posture is provably correct")
if p80:
    settext(p80, acc(p80).replace(
        "Next-generation firewalls alone can prevent on the order of 80 percent of attacks when their posture is provably correct",
        "Industry analysis indicates next-generation firewalls can prevent a large share of attacks when their posture is provably correct"))
log.append("R4: Table1 header, BEI->V-Intelligence, Mythos disambig, 80% hedge")

# ROUND 5: Validation Boundaries + Pilot Path
concl = find("11. Conclusion")
vb = ins_before(concl, "11. Technical Evidence and Validation Boundaries", H1)
b = ins_after(vb, "This paper distinguishes what has been validated from what has not, so the evidence can be weighed honestly.", BODY)
for t in [
    "Layer 1 (preemptive): the customer findings are representative engagements across government, critical infrastructure, financial services, and technology; they demonstrate the method, not a universal guarantee.",
    "Layer 2 (behavioral): the 4-of-4 result at an 8.1 percent operating point is a controlled, synthetic-data proof of concept on 250 entities, not yet field-validated on Army telemetry.",
    "Layer 5 (code): the vulnerability-discovery findings are from internal multi-week runs, reported internally and to be independently validated.",
    "Layers 3 and 4 (application, API, data, and agentic): presented as architecture and early capability; specific runtime results are roadmap.",
    "All layers map to MITRE ATT&CK and ATLAS but are not yet third-party audited."]:
    b = ins_after(b, t, BUL)
pp = ins_before(concl, "12. Recommended Pilot and Validation Path", H1)
b = ins_after(pp, "The recommended next step is a bounded pilot on real Army telemetry, advisory and human-in-the-loop, measured side by side against current tooling.", BODY)
for t in [
    "Layer 1: an agentless, configuration-only assessment of one enclave's firewall and edge estate, with first actionable findings in under fifteen days.",
    "Layer 2: a blind behavioral evaluation on a bounded Army enclave dataset, scored at a fixed false-positive operating point against current tooling.",
    "Success metrics: detection at a fixed false-positive rate, time to first finding, false-positive economics, and analyst-rated explainability.",
    "Environment: an IL5 or IL6 enclave, extensible to JWICS.",
    "Acquisition path: an Other Transaction Authority (OTA) prototype maturing toward a program of record (PoR) as Technology Readiness Level rises."]:
    b = ins_after(b, t, BUL)
settext(concl, "13. Conclusion")
log.append("R5: added Validation Boundaries + Pilot Path; renumbered Conclusion to 13")

d.save(P)
for l in log: print(" ", l)
print("saved.")
