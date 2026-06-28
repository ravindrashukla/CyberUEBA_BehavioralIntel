# -*- coding: utf-8 -*-
"""Revise the whitepaper Layer 2 to the enhanced two-layer framing:
production threat-profile detector (4/4 @ 0 FP) + discovery digital twin (8.1%); pgvector DB."""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
P = "WP DLA/Presentation/V-Intelligence_Layered_Cybersecurity_Whitepaper_US_Army_Pentagon.docx"
d = Document(P)
def acc(p): return "".join(t.text or "" for t in p._p.iter(qn('w:t')))
def find_pref(pref, after=None):
    started = after is None
    for p in d.paragraphs:
        t = acc(p).strip()
        if after and t.startswith(after): started = True
        if started and t.startswith(pref): return p
def settext(p, new):
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]: r.text = ""
    else: p.add_run(new)
def ins_after(anchor, text, style="Normal", bold=False):
    np = OxmlElement('w:p'); anchor._p.addnext(np)
    pp = Paragraph(np, anchor._parent); pp.style = d.styles[style]
    if text:
        r = pp.add_run(text); r.bold = bold
    return pp

# 1) Insert two-layer framing after the Layer 2 "The V-Intelligence Approach" heading
h = find_pref("The V-Intelligence Approach", after="4. Layer 2")
a = ins_after(h, "V-Intelligence runs two complementary detectors at this layer. The production engine is a precision multi-front threat-profile detector that matches each entity against named known-bad technique fingerprints, scored relative to its role-group peers and fired only on a corroborated match. Alongside it, a semantic behavioral digital twin provides discovery and direction: it reads behavior as meaning, tracks where an entity is heading, and names that direction in MITRE ATT&CK terms as an unsupervised net. Behavioral snapshots are stored as vectors in a PostgreSQL pgvector database; the representation, scoring, and detection methods are proprietary.")
a = ins_after(a, "The precision layer: the multi-front threat-profile detector (production).", bold=True)
a = ins_after(a, "Named technique fingerprints. Each profile is a measurable signature of a known attack: mass collection, insider collection, reconnaissance fan-out, living-off-the-land process use, C2 beaconing, DGA DNS, and data exfiltration, among others.", "List Bullet")
a = ins_after(a, "Cohort-relative and corroborated. A profile fires only when an entity is extreme against its own role-group peers and several corroborating features agree, which keeps precision high and false positives at zero.", "List Bullet")
a = ins_after(a, "Verified result: 4 of 4 nation-state-style attackers caught at 0 false positives (precision 100 percent): the insider by mass collection plus rare destinations, the slow APT by a 386-day C2 beacon plus 160 DGA domains, the living-off-the-land operator by anomalous process use, and the telecom-pivot campaign by reconnaissance fan-out.", "List Bullet")
a = ins_after(a, "The discovery layer: the behavioral digital twin.", bold=True)
print("inserted two-layer framing + precision detector")

# 2) Reframe the 8.1% paragraph as the discovery net (complementary, not the headline)
p131 = find_pref("In that evaluation, on 250 entities over 485 days")
if p131:
    settext(p131, "As an unsupervised discovery net, the digital twin also surfaces all four attackers by rank (Salt #1, Insider #2, Slow APT #7, Volt #24) at an 8.1 percent operating point, where classical point-anomaly methods (Isolation Forest, One-Class SVM, and Local Outlier Factor) catch none of the four and a z-score method catches one; on its own it cleanly separates two of the four. Its value is not a lower false-positive rate than the precision layer, but breadth and direction: it needs no prior fingerprint and it names where an entity is heading in MITRE terms. These are controlled, synthetic-data proof-of-concept results on 250 entities over 485 days (about 14 million events across five log sources), indicative rather than a guarantee of field performance.")
    print("reframed 8.1% paragraph as discovery net")

# 3) Validation Boundaries line for Layer 2
vb = find_pref("Layer 2 (behavioral): the 4-of-4 result")
if vb:
    settext(vb, "Layer 2 (behavioral): the production threat-profile detector reaches 4 of 4 at 0 false positives, and the discovery twin surfaces 4 of 4 by rank at an 8.1 percent operating point; both are controlled, synthetic-data proof-of-concept results on 250 entities, not yet field-validated on Army telemetry.")
    print("updated validation boundary")

d.save(P)
print("saved.")
