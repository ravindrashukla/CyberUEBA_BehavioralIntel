# -*- coding: utf-8 -*-
"""Apply approved trims to DLA FINAL as tracked changes:
 §5: delete the 'Core operators' paragraph; delete the two competitor-name parentheticals.
 §7: trim the SCRM activity list; condense the regulation list.
Nothing else touched. §6 forecasting bullets are KEPT per instruction.
"""
import copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

F="WP DLA/22nd Century DLA White BEI EDM Paper 062126 FINAL.docx"
AUTHOR="V-Intelligence"; DATE="2026-06-21T14:00:00Z"
d=Document(F)
_id=[80000]
def nid(): _id[0]+=1; return str(_id[0])
def _run(text,rpr=None):
    r=OxmlElement('w:r')
    if rpr is not None: r.append(copy.deepcopy(rpr))
    t=OxmlElement('w:t'); t.set(qn('xml:space'),'preserve'); t.text=text; r.append(t); return r
def _wrap(tag,run):
    e=OxmlElement(tag); e.set(qn('w:id'),nid()); e.set(qn('w:author'),AUTHOR); e.set(qn('w:date'),DATE); e.append(run); return e
def accp(p): return "".join(t.text or "" for t in p._p.iter(qn('w:t')))

def tracked_replace(OLD,NEW,label):
    for p in d.paragraphs:
        if OLD not in accp(p): continue
        runs=p._p.findall(qn('w:r'))
        for r in runs:
            t=r.find(qn('w:t'))
            if t is not None and t.text and OLD in t.text:
                full=t.text; i=full.index(OLD); before,after=full[:i],full[i+len(OLD):]
                rpr=r.find(qn('w:rPr'))
                t.text=before; t.set(qn('xml:space'),'preserve')
                r.addnext(_run(after,rpr))
                if NEW: r.addnext(_wrap('w:ins',_run(NEW,rpr)))
                delrun=_run(OLD,rpr)
                for tt in list(delrun.iter(qn('w:t'))):
                    dt=OxmlElement('w:delText'); dt.set(qn('xml:space'),'preserve'); dt.text=tt.text; tt.getparent().replace(tt,dt)
                r.addnext(_wrap('w:del',delrun))
                print(f"[{label}] applied (single-run)"); return True
        full="".join("".join(t.text or "" for t in r.findall(qn('w:t'))) for r in runs)
        if OLD in full and runs:
            for t in runs[0].findall(qn('w:t')): t.text=full; t.set(qn('xml:space'),'preserve')
            for r in runs[1:]:
                for t in r.findall(qn('w:t')): t.text=""
            t0=runs[0].find(qn('w:t')); i=t0.text.index(OLD)
            before,after=t0.text[:i],t0.text[i+len(OLD):]
            rpr=runs[0].find(qn('w:rPr'))
            t0.text=before; t0.set(qn('xml:space'),'preserve')
            runs[0].addnext(_run(after,rpr))
            if NEW: runs[0].addnext(_wrap('w:ins',_run(NEW,rpr)))
            delrun=_run(OLD,rpr)
            for tt in list(delrun.iter(qn('w:t'))):
                dt=OxmlElement('w:delText'); dt.set(qn('xml:space'),'preserve'); dt.text=tt.text; tt.getparent().replace(tt,dt)
            runs[0].addnext(_wrap('w:del',delrun))
            print(f"[{label}] applied (merged-runs)"); return True
    print(f"[{label}] NOT FOUND"); return False

def del_para_tracked(startswith,label):
    for p in d.paragraphs:
        if accp(p).strip().startswith(startswith):
            pel=p._p; runs=pel.findall(qn('w:r'))
            if runs:
                delw=OxmlElement('w:del'); delw.set(qn('w:id'),nid()); delw.set(qn('w:author'),AUTHOR); delw.set(qn('w:date'),DATE)
                runs[0].addprevious(delw)
                for r in runs:
                    for t in r.findall(qn('w:t')):
                        dt=OxmlElement('w:delText'); dt.set(qn('xml:space'),'preserve'); dt.text=t.text; r.replace(t,dt)
                    delw.append(r)
            ppr=pel.find(qn('w:pPr'))
            if ppr is None: ppr=OxmlElement('w:pPr'); pel.insert(0,ppr)
            rpr=ppr.find(qn('w:rPr'))
            if rpr is None: rpr=OxmlElement('w:rPr'); ppr.append(rpr)
            dm=OxmlElement('w:del'); dm.set(qn('w:id'),nid()); dm.set(qn('w:author'),AUTHOR); dm.set(qn('w:date'),DATE); rpr.append(dm)
            print(f"[{label}] paragraph deleted (tracked)"); return True
    print(f"[{label}] NOT FOUND"); return False

# ---- Section 5 ----
del_para_tracked("Core operators include behavioral composites","S5 Core operators paragraph")
tracked_replace(" (for example Kinaxis, o9, and Blue Yonder)","","S5 competitor names 1")
tracked_replace(" (for example Interos, Everstream, and Resilinc)","","S5 competitor names 2")

# ---- Section 7 ----
OLD1=("supplier qualification, exclusion-list screening, contractual flow-downs, OEM and authorized-reseller sourcing, "
      "component authenticity, chain-of-custody documentation, serial-level traceability, anti-counterfeit inspection, "
      "cybersecurity validation, incident reporting, and continuous updates to the SCRM plan")
NEW1=("supplier qualification, component authenticity, anti-counterfeit inspection, and cybersecurity validation, with "
      "continuous updates to the SCRM plan")
tracked_replace(OLD1,NEW1,"S7 SCRM activity list")
OLD2="NIST SP 800-53, NIST SP 800-161, DFARS, NDAA, FAR-based sourcing controls, and related federal risk-management expectations"
NEW2="NIST SP 800-161 and related federal requirements"
tracked_replace(OLD2,NEW2,"S7 regulation list")

# Track Changes ON
st=d.settings.element
if st.find(qn('w:trackChanges')) is None:
    tc=OxmlElement('w:trackChanges')
    ref=st.find(qn('w:defaultTabStop'))
    if ref is not None: ref.addnext(tc)
    elif len(st): st.insert(0,tc)
    else: st.append(tc)
d.save(F)
d2=Document(F)
print(f"\ntracked insertions={len(d2.element.body.findall('.//'+qn('w:ins')))} deletions={len(d2.element.body.findall('.//'+qn('w:del')))} | trackChanges on={d2.settings.element.find(qn('w:trackChanges')) is not None}")
