"""Consistency fixes on DLA Final.docx:
  #1 add self-attention support sentence in 'Mathematical and Model Design Principles'
  #2 broaden the glossary BEI definition
  #3 (inspect only) figure caption doubling structure
"""
from docx import Document
from docx.oxml.ns import qn

PATH = "WP DLA/DLA Final.docx"
d = Document(PATH)

def accp(p):
    return "".join(t.text or "" for t in p._p.iter(qn('w:t')))

def replace_span(old, new, label):
    """Replace old->new even if it spans runs, by editing the w:t that holds old,
    else merging consecutive w:t in the paragraph that contains it."""
    for p in d.paragraphs:
        if old in accp(p):
            for t in p._p.iter(qn('w:t')):
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new)
                    print(f"  [{label}] single-run OK")
                    return True
            # spans runs: collapse all w:t text into first, blank rest
            ts = [t for t in p._p.iter(qn('w:t'))]
            joined = "".join(t.text or "" for t in ts)
            if old in joined:
                ts[0].text = joined.replace(old, new)
                ts[0].set(qn('xml:space'), 'preserve')
                for t in ts[1:]:
                    t.text = ""
                print(f"  [{label}] multi-run OK")
                return True
    print(f"  [{label}] NOT FOUND")
    return False

# also handle table cells (glossary is a table)
def replace_in_tables(old, new, label):
    for tb in d.tables:
        for row in tb.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    if old in accp(p):
                        for t in p._p.iter(qn('w:t')):
                            if t.text and old in t.text:
                                t.text = t.text.replace(old, new)
                                print(f"  [{label}] table OK")
                                return True
                        ts = [t for t in p._p.iter(qn('w:t'))]
                        joined = "".join(t.text or "" for t in ts)
                        if old in joined:
                            ts[0].text = joined.replace(old, new)
                            ts[0].set(qn('xml:space'), 'preserve')
                            for t in ts[1:]:
                                t.text = ""
                            print(f"  [{label}] table multi-run OK")
                            return True
    print(f"  [{label}] NOT FOUND in tables")
    return False

# ---- #1 self-attention support ----
OLD1 = "predictive risk models. The contribution is the composition"
NEW1 = ("predictive risk models. The embedding step is what carries the language-model foundation into this design: "
        "each entity's serialized state is encoded by the same class of transformer, self-attention-based models that "
        "give large language models their grasp of context, so the resulting vector reflects how an entity's signals "
        "relate to one another rather than being read one field at a time. The contribution is the composition")
replace_span(OLD1, NEW1, "1 self-attention")

# ---- #2 glossary BEI ----
OLD2 = "The cross-domain early-warning framework that analyzes entities"
NEW2 = ("The behavioral-intelligence approach, applied across demand, supplier-risk, and network early-warning, that "
        "analyzes entities")
replace_in_tables(OLD2, NEW2, "2 glossary BEI")

d.save(PATH)
print("saved #1 and #2")

# ---- #3 inspect figure caption doubling ----
print("\n=== Figure caption structure (Figure 1 paragraph) ===")
d2 = Document(PATH)
for p in d2.paragraphs:
    if "Metrics to Mission Decision Intelligence" in accp(p):
        for el in p._p.iter():
            tag = el.tag.split('}')[1]
            if tag in ('docPr', 'cNvPr'):
                print(f"  <{tag}> name={el.get('name')!r} descr={el.get('descr')!r}")
            if tag == 't' and el.text and "Figure 1" in el.text:
                anc = [a.tag.split('}')[1] for a in el.iterancestors()]
                box = [a for a in anc if a in ('txbxContent', 'wsp', 'pic')]
                print(f"  <w:t in {box}> {el.text!r}")
        break
