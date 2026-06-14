#!/usr/bin/env python3
"""Build UEBA Whitepaper (Word Document).

Comprehensive whitepaper on User and Entity Behavior Analytics (UEBA):
what it is, why traditional approaches fail, the threat taxonomy it addresses,
behavioral detection architecture, empirical validation with V-Intelligence UEBA results,
and deployment recommendations for federal agencies.

Output: docs/UEBA_Behavioral_Intelligence_Whitepaper.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
BLUE = RGBColor(0x1B, 0x4F, 0x72)
TEAL = RGBColor(0x0E, 0x6B, 0x8A)
GOLD = RGBColor(0xB7, 0x95, 0x0B)
RED_ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREEN_ACCENT = RGBColor(0x1E, 0x8A, 0x49)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)

OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
                           "UEBA_Behavioral_Intelligence_Whitepaper.docx")


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
    return h


def add_body(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def add_callout(doc, text, border_color_hex="0E6B8A"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="12" w:color="{border_color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EAF4F7" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    run.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    return p


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "0D1B2A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EAF0F6")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


AUDIENCE_CONFIGS = {
    "DLA": {
        "prepared_for": "Defense Logistics Agency — CIO and Program Executive Office",
        "subtitle": "Behavioral Intelligence for Supply Chain and Logistics Network Defense",
        "filename": "Whitepaper_DLA.docx",
        "scenario_title": "The Supply Chain Threat",
        "scenario_text": (
            "Defense logistics networks process millions of transactions daily across "
            "global supply chains — requisitions, inventory movements, vendor communications, "
            "and financial transfers. An insider with legitimate access to logistics systems "
            "can systematically exfiltrate procurement data, manipulate supply chain records, "
            "or establish covert channels to adversary-controlled infrastructure — all while "
            "generating transaction volumes indistinguishable from normal operations."
        ),
        "scenario_text2": (
            "DLA's mission-critical supply chain infrastructure is a high-value target for "
            "nation-state actors seeking to disrupt military readiness. The Volt Typhoon "
            "campaign (CISA AA23-144A) specifically targeted critical infrastructure using "
            "living-off-the-land techniques that evade traditional detection. V-Intelligence "
            "UEBA addresses this gap by detecting behavioral changes in HOW logistics "
            "personnel interact with systems, not just WHETHER they exceed volume thresholds."
        ),
    },
    "USSOCOM": {
        "prepared_for": "United States Special Operations Command — SDA, EIS, S&T, Innovation",
        "subtitle": "Behavioral Intelligence for Special Operations Cyber Defense",
        "filename": "Whitepaper_USSOCOM.docx",
        "scenario_title": "The Special Operations Threat",
        "scenario_text": (
            "Special Operations Forces operate across multiple networks, classification "
            "levels, and geographic combatant commands. Personnel routinely access "
            "compartmented programs, operational plans, and intelligence sources. An insider "
            "threat in this environment carries disproportionate risk — a single compromised "
            "operator can expose source networks, mission timelines, and force disposition "
            "data that directly endangers lives."
        ),
        "scenario_text2": (
            "Nation-state adversaries actively target SOF infrastructure for intelligence "
            "collection. The Salt Typhoon campaign operated inside US telecom infrastructure "
            "for over 5 years before discovery, accessing lawful intercept systems. "
            "V-Intelligence UEBA detects these threats by building behavioral profiles "
            "across authentication, file access, network, and endpoint telemetry — "
            "identifying when an entity's behavior drifts from its established baseline "
            "across any of five behavioral dimensions, even when individual metrics "
            "remain within normal ranges."
        ),
    },
    "DISA": {
        "prepared_for": "Defense Information Systems Agency — PEO for Cyber",
        "subtitle": "Behavioral Intelligence for DODIN Enterprise Cyber Defense",
        "filename": "Whitepaper_DISA.docx",
        "scenario_title": "The Enterprise Network Threat",
        "scenario_text": (
            "The Department of Defense Information Network (DODIN) serves millions of "
            "users across combatant commands, services, and agencies. DISA's mission to "
            "defend this infrastructure faces a fundamental challenge: adversaries who "
            "obtain valid credentials can operate within authorized access boundaries "
            "indefinitely. Traditional perimeter defenses and signature-based detection "
            "cannot distinguish a legitimate administrator from an adversary using "
            "stolen credentials to conduct reconnaissance and data staging."
        ),
        "scenario_text2": (
            "Zero Trust Architecture (NIST SP 800-207, OMB M-22-09) mandates continuous "
            "verification — but continuous verification requires continuous behavioral "
            "assessment. V-Intelligence UEBA provides the behavioral intelligence layer "
            "that Zero Trust requires: real-time behavioral scoring that feeds into "
            "access decisions, detecting drift from established baselines across five "
            "behavioral dimensions before damage occurs."
        ),
    },
    "Army_AI": {
        "prepared_for": "United States Army — Chief of Artificial Intelligence",
        "subtitle": "AI-Powered Behavioral Intelligence for Army Cyber Defense",
        "filename": "Whitepaper_Army_AI.docx",
        "scenario_title": "The AI Innovation Imperative",
        "scenario_text": (
            "The Army's adoption of AI for cyber defense must move beyond rule-based "
            "automation. Current SIEM systems generate thousands of alerts per day, the "
            "vast majority false positives, overwhelming Security Operations Centers and "
            "creating alert fatigue that allows real threats to go uninvestigated. The "
            "challenge is not more alerts — it is smarter detection that ranks users by "
            "behavioral anomaly, surfacing the highest-risk entities first."
        ),
        "scenario_text2": (
            "V-Intelligence UEBA applies machine learning at three distinct layers: "
            "semantic behavioral embeddings that capture the meaning of user activity "
            "(not just volume), multi-phase composite scoring that fuses five independent "
            "detection dimensions into a single ranked output, and a multi-front "
            "threat-profile detector that scores measurable known-bad profiles against "
            "previously-unseen infrastructure and access patterns. Embedding/composite scoring "
            "alone cleanly separates 2 of 4 campaigns; the threat-profile detector adds the "
            "remaining two for a clean 4 of 4 detection at 0 false positives — a ranked list, "
            "not a flood of alerts."
        ),
    },
}


def _init_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    return doc


def _build_title_page(doc, audience_key):
    cfg = AUDIENCE_CONFIGS[audience_key]

    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("User and Entity Behavior Analytics\n(UEBA)")
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(cfg["subtitle"])
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE

    for _ in range(2):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "22nd Century Technologies, Inc.\n"
        "V-Intelligence UEBA Program\n\n"
        f"June 2026 — Version 2.0\n\n"
        f"Prepared for: {cfg['prepared_for']}"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY

    add_page_break(doc)


def _build_exec_summary(doc, audience_key):
    cfg = AUDIENCE_CONFIGS[audience_key]

    add_section_heading(doc, "1. Executive Summary", level=1)

    # Audience-specific opening scenario
    add_body(doc, cfg["scenario_title"], bold=True, space_after=2)
    add_body(doc, cfg["scenario_text"])
    add_body(doc, cfg["scenario_text2"])

    # Shared findings (same for all audiences)
    add_body(doc, (
        "This whitepaper presents a comprehensive UEBA framework validated through the "
        "V-Intelligence UEBA program. Using 485 days of synthetic telemetry across 250 users "
        "generating approximately 15 million events across five log sources with 4 embedded attack campaigns — an 8-month "
        "insider threat, a 180-day slow APT with C2 beaconing, a 115-day Volt Typhoon "
        "living-off-the-land campaign, and a 100-day Salt Typhoon telecom infrastructure "
        "attack — we demonstrate that:"
    ))

    add_bullet(doc, "Traditional anomaly detection (Isolation Forest, SVM, LOF, Z-Score) "
               "detects 0 of 4 attackers at operationally acceptable false positive rates. "
               "Every attack user's individual feature values remain within normal statistical "
               "ranges.", bold_prefix="Finding 1: ")
    add_bullet(doc, "Behavioral drift analysis alone (single-composite embeddings) fails when "
               "attack signals from one behavioral zone are diluted by four stable zones.",
               bold_prefix="Finding 2: ")
    add_bullet(doc, "Multi-phase composite scoring — fusing signal strength, breadth, sustained "
               "deviation, context divergence, and novelty persistence — cleanly separates only "
               "2 of 4 attacks (USR-156 and USR-118 above all normals; USR-234 and USR-042 below "
               "normal users). The clean 4 of 4 detection at 0 false positives is achieved by "
               "the separate multi-front threat-profile detector that scores measurable "
               "known-bad profiles (C2-beacon, DGA-DNS, LOTL-process, cohort-rare access, "
               "recon-fanout, insider-collection), label-free.",
               bold_prefix="Finding 3: ")
    add_bullet(doc, "Semantic behavioral embeddings outperform raw feature vectors by capturing "
               "what kind of behavioral change occurred (direction), not just how much changed "
               "(magnitude). The insider changes WHAT they access, not HOW MUCH — invisible to "
               "raw feature distance, visible to semantic drift.",
               bold_prefix="Finding 4: ")

    add_callout(doc,
        "Key insight: The insider changes WHAT they access, not HOW MUCH. "
        "Traditional algorithms measure magnitude. UEBA measures direction."
    )

    add_page_break(doc)


def _build_core(doc):

    add_section_heading(doc, "Table of Contents", level=1)

    toc_entries = [
        "1. Executive Summary",
        "2. Introduction: The Evolving Threat Landscape",
        "3. What is User and Entity Behavior Analytics (UEBA)?",
        "   3.1. Definition and Core Principles",
        "   3.2. The Three Pillars: Baseline, Deviation, Direction",
        "   3.3. UEBA vs. Traditional Security Approaches",
        "4. Why Traditional Detection Fails Against Modern Threats",
        "   4.1. Rule-Based SIEM Limitations",
        "   4.2. Threshold-Based Anomaly Detection Gaps",
        "   4.3. The 'Slow and Low' Problem",
        "5. UEBA Threat Taxonomy",
        "   5.1. Insider Threats",
        "   5.2. Advanced Persistent Threats (APT)",
        "   5.3. Living-off-the-Land (LOTL) Attacks",
        "   5.4. Telecom Infrastructure Attacks",
        "6. Behavioral Detection Architecture",
        "   6.1. Behavioral Feature Engineering (23 Features)",
        "   6.2. Behavioral Zones (5 Decomposition Dimensions)",
        "   6.3. Temporal Trajectory Analysis",
        "   6.4. Cross-Entity Relationship Analysis",
        "   6.5. Context-Adaptive Detection",
        "7. Detection Methods and Their Parameters",
        "   7.1. Traditional Methods (Tier 1 — 6 Algorithms)",
        "   7.2. Behavioral Drift Methods (Tier 2 — 3 Layers)",
        "   7.3. Digital Entity Methods (Tier 3 — 9 Methods)",
        "   7.4. Why a Single Method Is Not Enough",
        "8. Empirical Validation",
        "   8.1. Simulation Environment",
        "   8.2. Attack Campaign Design",
        "   8.3. Detection Results: Tier 1 vs Tier 2 vs Tier 3",
        "   8.4. Per-Threat Analysis",
        "   8.5. False Positive Analysis",
        "9. Detection Playbook: Threat Type to Algorithm",
        "   9.1. Per-Threat Recommendations",
        "   9.2. Layered Deployment Strategy",
        "10. Multi-Phase Composite Scoring",
        "   10.1. The Five Detection Phases",
        "   10.2. Per-Attacker Composite Breakdown",
        "   10.3. USR-234 Novelty Persistence — Case Study",
        "11. Why Semantic Embeddings Outperform Raw Feature Vectors",
        "   11.1. Two Pipelines: Feature-Space vs Semantic",
        "   11.2. Three Limitations of Raw Feature Distance",
        "   11.3. Empirical Signal Separation",
        "12. Operational Deployment for DoD and IC",
        "   12.1. Deployment Phases",
        "   12.2. Integration Points",
        "   12.3. Applicability Across Mission Environments",
        "13. Conclusion",
        "14. References and Further Reading",
    ]
    for entry in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(entry)
        run.font.size = Pt(10.5)
        run.font.color.rgb = NAVY if not entry.startswith("   ") else BLUE
        p.paragraph_format.space_after = Pt(2)

    add_page_break(doc)

    # ==============================================================
    # 2. INTRODUCTION
    # ==============================================================
    add_section_heading(doc, "2. Introduction: The Evolving Threat Landscape", level=1)

    add_body(doc, (
        "The cybersecurity threat landscape has undergone a fundamental transformation. "
        "A decade ago, the primary threats were malware with known signatures, brute-force "
        "attacks generating obvious log spikes, and vulnerabilities with published CVEs. "
        "Modern threats are qualitatively different: they are patient, use legitimate "
        "credentials, exploit authorized tools, and operate within normal traffic patterns."
    ))

    add_body(doc, (
        "Three converging trends drive this evolution:"
    ))

    add_body(doc, "1. Nation-State Sophistication", bold=True, space_after=2)
    add_body(doc, (
        "State-sponsored groups like Volt Typhoon (China) and APT29 (Russia) now "
        "pre-position in critical infrastructure for months or years before activation. "
        "CISA Advisory AA23-144A documented Volt Typhoon's use of living-off-the-land "
        "techniques — PowerShell, WMI, RDP — tools that exist on every Windows system. "
        "No malware to detect. No signatures to match. The 2024 Salt Typhoon campaign "
        "compromised major U.S. telecommunications providers (AT&T, Verizon, T-Mobile) "
        "using authorized maintenance credentials, accessing call detail records and "
        "lawful intercept systems without triggering a single alert."
    ))

    add_body(doc, "2. The Insider Threat Renaissance", bold=True, space_after=2)
    add_body(doc, (
        "Insider threats remain the most damaging and hardest to detect category. "
        "The 2024 Verizon Data Breach Investigations Report found that insiders account "
        "for approximately 20% of breaches, but the median time to detect an insider "
        "threat exceeds 200 days. Unlike external attacks that must breach a perimeter, "
        "insiders already possess valid credentials, authorized access, and institutional "
        "knowledge of where sensitive data resides. Their actions are individually "
        "legitimate — it is only the pattern over time that reveals malicious intent."
    ))

    add_body(doc, "3. The Failure of Perimeter Defense", bold=True, space_after=2)
    add_body(doc, (
        "Zero Trust Architecture recognizes that the perimeter no longer exists. "
        "Cloud adoption, remote work, supply chain integration, and API-first "
        "architectures have dissolved the network boundary. When every access is "
        "potentially legitimate and potentially malicious, the only reliable signal "
        "is behavioral: is this entity acting consistently with its established pattern?"
    ))

    add_callout(doc,
        "The common thread: all four threat categories succeed by staying within "
        "the statistical bounds of 'normal' while gradually shifting behavioral direction. "
        "Detection requires measuring behavioral trajectory, not behavioral magnitude."
    )

    add_page_break(doc)

    # ==============================================================
    # 3. WHAT IS UEBA?
    # ==============================================================
    add_section_heading(doc, "3. What is User and Entity Behavior Analytics (UEBA)?", level=1)

    add_section_heading(doc, "3.1. Definition and Core Principles", level=2)

    add_body(doc, (
        "User and Entity Behavior Analytics (UEBA) is a cybersecurity approach that "
        "detects threats by establishing behavioral baselines for every monitored entity "
        "(users, devices, applications, network segments) and flagging when behavior "
        "deviates from that baseline. Unlike rule-based SIEM systems that define what "
        "is malicious, UEBA defines what is normal and detects departures."
    ))

    add_body(doc, (
        "The UEBA approach is built on three core principles:"
    ))

    add_body(doc, "Principle 1: Every Entity Has a Behavioral Baseline", bold=True, space_after=2)
    add_body(doc, (
        "Each user, device, application, and network segment develops a characteristic "
        "behavioral pattern over time. A financial analyst typically accesses specific "
        "file shares during business hours, authenticates from the same workstation, "
        "and generates predictable network traffic. This pattern becomes the baseline "
        "against which future behavior is measured. The baseline is not static — it "
        "adapts to legitimate changes like role transitions, seasonal variations, and "
        "organizational restructuring."
    ))

    add_body(doc, "Principle 2: Deviation is the Primary Detection Signal", bold=True, space_after=2)
    add_body(doc, (
        "When an entity's behavior deviates from its established baseline, that "
        "deviation is a potential threat indicator. The magnitude of deviation is "
        "informative but insufficient — a large deviation could be a data center "
        "migration (benign) or a ransomware deployment (critical). What matters is "
        "the combination of how much changed, how fast it changed, and in which "
        "behavioral dimension the change occurred."
    ))

    add_body(doc, "Principle 3: Direction Matters More Than Magnitude", bold=True, space_after=2)
    add_body(doc, (
        "This is the critical insight that separates modern UEBA from traditional "
        "anomaly detection. A user whose file access volume increases by 50% might be "
        "working on a deadline (benign) or staging data for exfiltration (malicious). "
        "Traditional systems flag both identically. UEBA asks: which files? From which "
        "classification? At what hours? Via what access path? The direction of behavioral "
        "change — toward restricted data, toward unusual access patterns, toward known "
        "threat behaviors — is the discriminating signal."
    ))

    add_section_heading(doc, "3.2. The Three Pillars: Baseline, Deviation, Direction", level=2)

    create_table(doc,
        ["Pillar", "Question Asked", "Traditional Answer", "UEBA Answer"],
        [
            ["Baseline", "What is normal?",
             "Population statistics\n(same threshold for everyone)",
             "Per-entity behavioral profile\n(each entity has its own normal)"],
            ["Deviation", "Has something changed?",
             "Feature value exceeds\nfixed threshold",
             "Behavioral trajectory drifts\nfrom personal baseline"],
            ["Direction", "What kind of change?",
             "Not considered\n(magnitude only)",
             "Which behavioral dimension\nis drifting toward what pattern?"],
        ],
        col_widths=[0.8, 1.5, 2.2, 2.5],
    )

    add_section_heading(doc, "3.3. UEBA vs. Traditional Security Approaches", level=2)

    create_table(doc,
        ["Capability", "Rule-Based SIEM", "Traditional ML\n(Anomaly Detection)", "UEBA"],
        [
            ["What it looks for", "Known bad patterns\n(signatures, rules)",
             "Statistical outliers\n(deviation from population)",
             "Behavioral drift\n(deviation from self)"],
            ["Baseline reference", "Expert-defined rules",
             "Population mean/variance",
             "Per-entity historical profile"],
            ["Handles slow drift?", "No — each event\nis evaluated alone",
             "Only if drift creates\na feature-space outlier",
             "Yes — tracks cumulative\nbehavioral trajectory"],
            ["Handles credential reuse?", "No — valid credentials\npass all rules",
             "Only if access volume\nexceeds statistical bounds",
             "Yes — detects direction\nchange regardless of volume"],
            ["Handles LOTL?", "No — legitimate tools\nare whitelisted",
             "Only if tool usage\ncreates volume anomaly",
             "Yes — detects phase change\nin how tools are used"],
            ["False positive source", "Rule exceptions\nand edge cases",
             "Population diversity\n(normal variation)",
             "Legitimate behavioral shifts\n(role change, project transition)"],
            ["Time to detect", "Real-time\n(per-event rules)",
             "Batch analysis\n(weekly/monthly aggregation)",
             "Behavioral window\n(weeks of trajectory data)"],
        ],
        col_widths=[1.3, 1.8, 1.8, 2.1],
    )

    add_page_break(doc)

    # ==============================================================
    # 4. WHY TRADITIONAL DETECTION FAILS
    # ==============================================================
    add_section_heading(doc, "4. Why Traditional Detection Fails Against Modern Threats", level=1)

    add_section_heading(doc, "4.1. Rule-Based SIEM Limitations", level=2)

    add_body(doc, (
        "Security Information and Event Management (SIEM) systems form the backbone "
        "of most enterprise security operations. They aggregate logs from across the "
        "environment and apply correlation rules to detect known attack patterns. "
        "The fundamental limitation is structural: SIEM rules define what is malicious, "
        "which means they can only detect threats that match a pre-defined pattern."
    ))

    create_table(doc,
        ["Threat Type", "SIEM Rule Gap", "Why It Fails"],
        [
            ["Insider Threat\n(8-month campaign)", "Rules check for unauthorized access\n"
             "or excessive volume",
             "All accesses use valid credentials.\n"
             "Volume stays within normal range.\n"
             "The change is in data classification, not access count."],
            ["Slow APT\n(180-day C2)", "Rules check for known C2 domains\nor unusual ports",
             "C2 beacons every 6 hours over HTTPS.\n"
             "Domain is newly registered (no reputation).\n"
             "Traffic volume is 12KB — invisible in enterprise traffic."],
            ["Volt Typhoon LOTL\n(115-day campaign)", "Rules check for malware hashes\nor exploit signatures",
             "Uses PowerShell, WMI, RDP — tools on every Windows system.\n"
             "No malware deployed. No CVE exploited.\n"
             "Every action is individually 'authorized.'"],
            ["Salt Typhoon Telecom\n(100-day campaign)", "Rules check for unauthorized\ninfrastructure access",
             "Uses authorized maintenance credentials.\n"
             "Accesses CDR systems within maintenance windows.\n"
             "Traffic patterns match legitimate vendor sessions."],
        ],
        col_widths=[1.5, 2.5, 3.0],
    )

    add_section_heading(doc, "4.2. Threshold-Based Anomaly Detection Gaps", level=2)

    add_body(doc, (
        "Machine learning anomaly detection (Isolation Forest, LOF, One-Class SVM, "
        "Z-Score) represents an improvement over rules: instead of defining 'bad,' "
        "they learn 'normal' from data and flag statistical outliers. However, these "
        "methods share a fundamental limitation: they measure magnitude, not direction."
    ))

    add_body(doc, (
        "Consider the insider threat scenario: a user who normally accesses 10 files "
        "per day begins accessing 12 files per day, but shifts from engineering documents "
        "to restricted financial data. The volume change (10 to 12) is well within "
        "normal variation. No single feature value is a statistical outlier. But the "
        "change in what data is being accessed — the direction of behavioral shift — "
        "reveals the threat."
    ))

    add_body(doc, (
        "Our empirical validation confirms this limitation: across the four established "
        "traditional algorithms operating on 23 behavioral features, none detected the "
        "8-month insider threat, the slow APT, or the telecom pivot at an operationally "
        "acceptable false positive rate. The strongest traditional result was Z-Score, "
        "which caught only the most overt campaign — the living-off-the-land intrusion — "
        "and missed the three that matter most, including the insider threat, arguably "
        "the most damaging threat type."
    ))

    add_section_heading(doc, "4.3. The 'Slow and Low' Problem", level=2)

    add_body(doc, (
        "Modern advanced threats are specifically designed to evade both rules and "
        "statistical anomaly detection. The technique is called 'slow and low': make "
        "behavioral changes so gradually that no single time window shows an anomaly, "
        "while the cumulative change over weeks or months achieves the operational "
        "objective."
    ))

    add_body(doc, "Example — 180-day APT C2 beacon:", bold=True, space_after=2)
    add_bullet(doc, "Week 1-4: Establish presence. 2 DNS queries per day to DGA domain. "
               "Traffic: 8KB/day. Indistinguishable from software update checks.")
    add_bullet(doc, "Week 5-12: Light reconnaissance. Beacons every 6 hours. "
               "Data staging: 50KB per exfiltration. Below any volume threshold.")
    add_bullet(doc, "Week 13-26: Gradual escalation. Privilege probing every 30 days. "
               "Data staging every 17 days. Each individual action is normal.")
    add_bullet(doc, "Cumulative effect: After 26 weeks, the entity has exfiltrated GBs "
               "of data in packets too small to trigger any single alert.")

    add_body(doc, (
        "No individual weekly snapshot shows an anomaly. The attack is visible only "
        "in the trajectory — the cumulative drift of behavioral embedding over time. "
        "This is precisely what UEBA temporal analysis is designed to detect."
    ))

    add_page_break(doc)

    # ==============================================================
    # 5. UEBA THREAT TAXONOMY
    # ==============================================================
    add_section_heading(doc, "5. UEBA Threat Taxonomy", level=1)

    add_body(doc, (
        "UEBA is particularly effective against four categories of threats that share "
        "a common characteristic: they operate within the bounds of authorized access "
        "and cannot be detected by signature, rule, or simple statistical threshold. "
        "Each threat type creates a distinct behavioral signature that UEBA can identify "
        "through baseline drift analysis."
    ))

    add_section_heading(doc, "5.1. Insider Threats", level=2)

    add_body(doc, "Definition", bold=True, space_after=2)
    add_body(doc, (
        "An insider threat is a current or former employee, contractor, or business "
        "partner who has authorized access to organizational resources and uses that "
        "access to harm the organization. The insider may be motivated by financial "
        "gain, ideology, coercion, or disgruntlement. Unlike external attackers, insiders "
        "do not need to breach any perimeter — they already possess valid credentials "
        "and institutional knowledge."
    ))

    add_body(doc, "Duration and Progression", bold=True, space_after=2)
    add_body(doc, (
        "Insider threats typically unfold over months, following a predictable escalation "
        "pattern. The V-Intelligence UEBA simulation models an 8-month (32-week) campaign with four "
        "distinct phases:"
    ))

    create_table(doc,
        ["Phase", "Duration", "Behavioral Indicators", "SIEM Visibility"],
        [
            ["Phase 1:\nMood Shift", "Weeks 1-8",
             "Subtle changes in access timing.\n"
             "Slight increase in after-hours logins.\n"
             "No change in data access patterns.",
             "Invisible — all events\nwithin normal parameters"],
            ["Phase 2:\nCuriosity", "Weeks 9-16",
             "Begins browsing directories outside\nnormal scope. "
             "Accesses a few restricted files.\nVolume still within normal range.",
             "Invisible — restricted_ratio\nchanges from 0.02 to 0.06\n(within population spread)"],
            ["Phase 3:\nReconnaissance", "Weeks 17-24",
             "Systematic scanning of data repositories.\n"
             "Increases unique file paths accessed.\n"
             "Shifts from read-heavy to read+write.",
             "Borderline — unique_paths may\nfluctuate but no sustained spike"],
            ["Phase 4:\nExfiltration", "Weeks 25-32",
             "Data staging via USB, email, cloud.\n"
             "Compression (7z, WinRAR) of collections.\n"
             "100-500 MB transfers in final weeks.",
             "Late-stage volume spike may\ntrigger alerts — but data\nis already staged"],
        ],
        col_widths=[1.0, 0.8, 2.5, 2.2],
    )

    add_body(doc, "Real-World Parallels", bold=True, space_after=2)
    add_bullet(doc, "Edward Snowden (NSA, 2013): Gradually expanded access scope over months "
               "before mass exfiltration of classified documents.", bold_prefix="")
    add_bullet(doc, "Chelsea Manning (DoD, 2010): Used authorized access to SIPRNet to "
               "download diplomatic cables and military reports over several months.")
    add_bullet(doc, "Corporate IP theft: Departing employees who copy proprietary designs, "
               "customer lists, or source code in the weeks before resignation — often using "
               "cloud storage, personal email, or USB devices.")

    add_body(doc, "UEBA Behavioral Signature", bold=True, space_after=2)
    add_body(doc, (
        "The insider's identity zone (role, department, clearance) remains completely "
        "stable throughout the entire campaign — they are the same person with the same "
        "credentials. But their data_behavior zone (what files they access, what "
        "classification, how much they write) drifts steadily from their established "
        "baseline. This pattern — identity stable + data_behavior drifting — is the "
        "definitive UEBA signature for insider threats."
    ))

    add_section_heading(doc, "5.2. Advanced Persistent Threats (APT)", level=2)

    add_body(doc, "Definition", bold=True, space_after=2)
    add_body(doc, (
        "An Advanced Persistent Threat is a prolonged, targeted cyberattack in which "
        "an attacker gains unauthorized access to a network and remains undetected for "
        "an extended period. APTs are typically associated with nation-state actors or "
        "state-sponsored groups with significant resources, patience, and operational "
        "security discipline. The 'persistent' in APT is the key attribute: these "
        "campaigns operate on timescales of months to years."
    ))

    add_body(doc, "Duration and Progression", bold=True, space_after=2)
    add_body(doc, (
        "The V-Intelligence UEBA simulation models a 180-day (26-week) slow APT campaign with "
        "Command and Control (C2) beaconing as the primary operational mechanism:"
    ))

    add_bullet(doc, "C2 beacons every 6 hours (4 per day) — periodic but low-frequency")
    add_bullet(doc, "DGA (Domain Generation Algorithm) DNS queries: 1-3 per day")
    add_bullet(doc, "Data staging every 17 days — small batches below transfer thresholds")
    add_bullet(doc, "Privilege escalation probing every 30 days — single attempt per cycle")

    add_body(doc, "Real-World Parallels", bold=True, space_after=2)
    add_bullet(doc, "SolarWinds/SUNBURST (2020): Backdoor in software update mechanism "
               "remained active for 9 months before discovery. Affected 18,000+ organizations "
               "including U.S. government agencies.", bold_prefix="")
    add_bullet(doc, "APT29/Cozy Bear: Russian intelligence-linked group known for patient, "
               "multi-month operations targeting government and defense organizations.")
    add_bullet(doc, "APT28/Fancy Bear: Russian military intelligence group responsible for "
               "the 2016 DNC breach, operating over several months of credential harvesting "
               "and lateral movement.")

    add_body(doc, "UEBA Behavioral Signature", bold=True, space_after=2)
    add_body(doc, (
        "The APT's identity zone remains stable (compromised credentials appear legitimate). "
        "The network_footprint zone drifts as C2 beaconing adds periodic external connections "
        "and DNS queries to domains not in the entity's historical pattern. This pattern — "
        "identity stable + network_footprint drifting — is the UEBA signature for APT/C2 "
        "activity."
    ))

    add_section_heading(doc, "5.3. Living-off-the-Land (LOTL) Attacks", level=2)

    add_body(doc, "Definition", bold=True, space_after=2)
    add_body(doc, (
        "Living-off-the-Land (LOTL) attacks use tools and capabilities that already "
        "exist within the target environment rather than deploying external malware. "
        "By using PowerShell, WMI, RDP, certutil, and other built-in Windows utilities, "
        "attackers blend into normal administrative activity. There is no malware hash "
        "to detect, no exploit signature to match, and no unauthorized software installation."
    ))

    add_body(doc, "Duration and Progression", bold=True, space_after=2)
    add_body(doc, (
        "The V-Intelligence UEBA simulation models a 115-day Volt Typhoon LOTL campaign based on "
        "CISA Advisory AA23-144A:"
    ))

    add_bullet(doc, "Credential harvesting via legitimate tools (mimikatz-style, "
               "but using built-in Windows credential utilities)")
    add_bullet(doc, "Lateral movement via RDP and WMI — both standard admin tools")
    add_bullet(doc, "Infrastructure pre-positioning: establishing persistence in "
               "critical systems for potential future activation")
    add_bullet(doc, "No external C2 during dormancy — uses proxy routing through "
               "compromised SOHO routers")

    add_body(doc, "Real-World Parallels", bold=True, space_after=2)
    add_bullet(doc, "Volt Typhoon (2023-present): Chinese state-sponsored group targeting "
               "U.S. critical infrastructure (water, energy, transportation, communications). "
               "CISA, NSA, and FBI joint advisory documented exclusive use of LOTL techniques.",
               bold_prefix="")
    add_bullet(doc, "HAFNIUM/Exchange attacks: Used PowerShell web shells and legitimate "
               "Exchange cmdlets to maintain access post-exploitation.")

    add_body(doc, "UEBA Behavioral Signature", bold=True, space_after=2)
    add_body(doc, (
        "LOTL attacks create a uniform phase change across all behavioral zones — "
        "the entity transitions from one operational mode to another. Unlike insider "
        "threats (single zone drifting) or APTs (network zone drifting), LOTL attacks "
        "shift the entire behavioral profile simultaneously because legitimate tools "
        "affect multiple zones. The UEBA signature is a regime shift: a clear "
        "before/after discontinuity in the behavioral trajectory."
    ))

    add_section_heading(doc, "5.4. Telecom Infrastructure Attacks", level=2)

    add_body(doc, "Definition", bold=True, space_after=2)
    add_body(doc, (
        "Telecom infrastructure attacks target the communication backbone itself: "
        "telephone networks, internet service providers, undersea cables, and the "
        "lawful intercept systems mandated by regulation. These attacks are uniquely "
        "dangerous because they compromise the infrastructure that other organizations "
        "depend on for secure communication."
    ))

    add_body(doc, "Duration and Progression", bold=True, space_after=2)
    add_body(doc, (
        "The V-Intelligence UEBA simulation models a 100-day Salt Typhoon-style campaign based on "
        "the 2024 disclosure of Chinese compromise of AT&T, Verizon, and T-Mobile:"
    ))

    add_bullet(doc, "Edge device compromise via authorized vendor credentials")
    add_bullet(doc, "CDR (Call Detail Record) exfiltration through existing data "
               "interfaces — appears as routine billing system queries")
    add_bullet(doc, "Lawful intercept system access to monitor specific targets")
    add_bullet(doc, "C2 via DNS over HTTPS to a domain mimicking Microsoft update services")

    add_body(doc, "Real-World Parallels", bold=True, space_after=2)
    add_bullet(doc, "Salt Typhoon (2024): Chinese state-sponsored campaign that compromised "
               "at least 9 U.S. telecommunications providers. Accessed CDR databases and "
               "lawful intercept systems. Undetected for months until intelligence community "
               "notification.", bold_prefix="")
    add_bullet(doc, "The scope of compromise remains under investigation, with the White House "
               "acknowledging the campaign represents one of the most significant telecom "
               "breaches in U.S. history.")

    add_body(doc, "UEBA Behavioral Signature", bold=True, space_after=2)
    add_body(doc, (
        "Telecom infrastructure attacks create multi-zone drift because the attacker "
        "must traverse multiple network segments and access diverse data stores. The "
        "network_footprint zone drifts (new external connections), the data_behavior "
        "zone drifts (accessing CDR/intercept systems outside maintenance scope), and "
        "the drift accumulates persistently over 100+ days. The UEBA signature is "
        "persistent cumulative drift across multiple behavioral zones."
    ))

    add_page_break(doc)

    # ==============================================================
    # 6. BEHAVIORAL DETECTION ARCHITECTURE
    # ==============================================================
    add_section_heading(doc, "6. Behavioral Detection Architecture", level=1)

    add_body(doc, (
        "The V-Intelligence UEBA behavioral detection architecture transforms raw telemetry "
        "into behavioral embeddings organized across five dimensions, tracked over "
        "time, and analyzed for drift direction against known threat patterns. This "
        "section describes the architectural components that enable UEBA detection."
    ))

    add_section_heading(doc, "6.1. Behavioral Feature Engineering (23 Features)", level=2)

    add_body(doc, (
        "All detection begins with feature engineering: transforming raw log events "
        "(authentication, file access, endpoint telemetry, network flows, DNS queries) "
        "into 23 aggregate features per user per week. These features capture the "
        "quantitative dimensions of behavior that, in combination, characterize each "
        "entity's operational profile."
    ))

    create_table(doc,
        ["Category", "Features (Count)", "What They Capture"],
        [
            ["Authentication (7)", "auth_total, auth_failed, auth_fail_rate,\n"
             "auth_unique_sources, auth_unique_dests,\n"
             "auth_off_hours_ratio, auth_methods_used",
             "How the entity authenticates: frequency, success rate, timing,\n"
             "diversity of sources/destinations, method complexity"],
            ["File Access (6)", "file_total, file_unique_paths, file_restricted_ratio,\n"
             "file_confidential_ratio, file_write_ratio, file_total_bytes",
             "What data the entity touches: volume, breadth, classification,\n"
             "read/write balance, total transfer size"],
            ["Endpoint (5)", "endpoint_total, endpoint_suspicious_ratio,\n"
             "endpoint_max_risk, endpoint_mean_risk,\n"
             "endpoint_unique_processes",
             "The entity's endpoint health: event volume, suspicious activity\n"
             "fraction, risk scores, process diversity"],
            ["Network (3)", "net_bytes_out, net_unique_dsts,\nnet_external_ratio",
             "Where the entity communicates: outbound volume, destination\n"
             "diversity, external vs. internal ratio"],
            ["DNS (2)", "dns_unique_domains, dns_nxdomain_ratio",
             "DNS behavior: domain diversity (DGA indicator),\n"
             "NXDOMAIN ratio (C2 probing indicator)"],
        ],
        col_widths=[1.3, 2.5, 3.2],
    )

    add_section_heading(doc, "6.2. Behavioral Zones (5 Decomposition Dimensions)", level=2)

    add_body(doc, (
        "The 23 features are organized into 5 behavioral zones that represent distinct "
        "dimensions of entity behavior. Each zone is analyzed independently, enabling "
        "the system to detect which dimension is changing — the critical capability "
        "that traditional single-composite approaches lack."
    ))

    create_table(doc,
        ["Zone", "Features Included", "Behavioral Dimension", "Attack Detection"],
        [
            ["Identity", "role, department, clearance,\ntenure, user_type",
             "Who the entity is —\nstatic profile attributes",
             "Stable identity = using same\ncredentials (expected for insiders/APTs)"],
            ["Access Pattern", "auth_total, auth_fail_rate,\nauth_off_hours_ratio,\n"
             "auth_unique_sources/dests,\nauth_methods_used",
             "How the entity\nauthenticates",
             "Credential abuse: brute force,\ncredential stuffing, pass-the-hash"],
            ["Data Behavior", "file_total, file_restricted_ratio,\n"
             "file_confidential_ratio,\nfile_write_ratio,\nfile_unique_paths, file_total_bytes",
             "What data the entity\ntouches and how",
             "Insider exfiltration: shift to\nrestricted/confidential data"],
            ["Network Footprint", "net_bytes_out, net_unique_dsts,\n"
             "net_external_ratio,\ndns_unique_domains,\ndns_nxdomain_ratio",
             "Where the entity\ncommunicates externally",
             "C2 beaconing, data exfiltration\nchannels, DGA activity"],
            ["Risk Posture", "endpoint_suspicious_ratio,\nendpoint_max/mean_risk,\n"
             "endpoint_unique_processes,\nendpoint_total",
             "Endpoint health and\nsuspicious activity level",
             "Malware, compromised endpoints,\nprivilege escalation tools"],
        ],
        col_widths=[1.0, 1.8, 1.7, 2.0],
    )

    add_body(doc, (
        "The zone decomposition is the architectural breakthrough that enables "
        "detection of threats invisible to traditional methods. When a user's identity "
        "zone is stable but their data_behavior zone is drifting, the system can "
        "conclude: 'same person, but they are accessing different data than before.' "
        "This is precisely the insider threat signature."
    ))

    add_section_heading(doc, "6.3. Temporal Trajectory Analysis", level=2)

    add_body(doc, (
        "UEBA tracks behavioral embeddings over time to detect drift patterns that "
        "evolve gradually. For each entity, the system computes:"
    ))

    create_table(doc,
        ["Trajectory Feature", "What It Measures", "Threat Relevance"],
        [
            ["Velocity Magnitude", "Rate of behavioral displacement\n(how fast is behavior changing?)",
             "Accelerating velocity = escalating\nattack campaign"],
            ["Acceleration", "Change in velocity over time\n(is drift speeding up?)",
             "Positive acceleration = attack entering\nmore aggressive phase"],
            ["Stability", "Consistency of behavioral snapshots\n(mean cosine similarity)",
             "Low stability = volatile behavior,\npossible regime shift"],
            ["Regime Shifts", "Fraction of weekly transitions where\nsimilarity drops below threshold",
             "Any shift = before/after discontinuity\n(C2 activation, ransomware)"],
            ["Trend Consistency", "Monotonicity of drift direction\n(Kendall's rank correlation)",
             "High consistency = sustained directional\ndrift (not random fluctuation)"],
            ["Total Drift", "Cumulative path length through\nbehavioral space",
             "High cumulative drift + low velocity =\nslow persistent campaign"],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    add_section_heading(doc, "6.4. Cross-Entity Relationship Analysis", level=2)

    add_body(doc, (
        "UEBA extends beyond individual entity analysis to track relationships between "
        "entities. A user's interaction pattern with their device, applications, and "
        "network segments creates a relationship signature. When a C2 beacon changes "
        "how a user interacts with their device, the relationship embedding drifts "
        "even if neither the user nor the device individually appears anomalous."
    ))

    add_body(doc, "Relationship types analyzed:", bold=True)
    add_bullet(doc, "User-Device: How the user interacts with their assigned device(s)",
               bold_prefix="")
    add_bullet(doc, "User-Application: Which applications the user accesses and how")
    add_bullet(doc, "Device-Segment: How devices interact with network segments")

    add_section_heading(doc, "6.5. Context-Adaptive Detection", level=2)

    add_body(doc, (
        "Different threat types create signals in different behavioral zones. An insider "
        "threat concentrates drift in the data_behavior zone, while an APT concentrates "
        "drift in the network_footprint zone. Context-adaptive detection re-weights the "
        "contribution of each zone based on the investigation scenario:"
    ))

    create_table(doc,
        ["Investigation Context", "Identity", "Access", "Data", "Network", "Risk", "Optimized For"],
        [
            ["Normal Operations", "20%", "20%", "20%", "20%", "20%", "Balanced monitoring"],
            ["Insider Investigation", "10%", "15%", "40%", "15%", "20%", "Data exfiltration detection"],
            ["APT Hunt", "5%", "15%", "10%", "40%", "30%", "C2 and network anomaly detection"],
            ["Privilege Audit", "10%", "25%", "10%", "15%", "40%", "Escalation and endpoint risk"],
        ],
        col_widths=[1.3, 0.6, 0.6, 0.6, 0.7, 0.5, 2.2],
    )

    add_body(doc, (
        "By running detection under all four contexts simultaneously and taking the "
        "maximum threat score, the system amplifies attack-specific signals that would "
        "be diluted under uniform weighting. An APT that is invisible under normal_ops "
        "becomes apparent under apt_hunt because the network zone's contribution is "
        "amplified from 20% to 40%."
    ))

    add_page_break(doc)

    # ==============================================================
    # 7. DETECTION METHODS AND THEIR PARAMETERS
    # ==============================================================
    add_section_heading(doc, "7. Detection Methods and Their Parameters", level=1)

    add_body(doc, (
        "The V-Intelligence UEBA system implements 17 detection methods across three tiers, each "
        "operating on the same underlying telemetry but applying fundamentally different "
        "analytical approaches. This section details the input features, parameters, "
        "and detection logic for each method."
    ))

    add_section_heading(doc, "7.1. Traditional Methods (Tier 1 — 6 Algorithms)", level=2)

    add_body(doc, (
        "Tier 1 represents industry-standard approaches deployed in commercial SIEM "
        "and UEBA products. All 6 algorithms operate on the same 23 scalar features. "
        "Static methods aggregate features across weeks into a single vector per user; "
        "temporal methods preserve the weekly time series."
    ))

    create_table(doc,
        ["Algorithm", "Type", "Key Parameters", "Input", "Detection Rule"],
        [
            ["Isolation Forest", "Static",
             "n_estimators=200\ncontamination=0.05",
             "23 features,\nStandardScaler normalized",
             "Partition outliers that isolate\nquickly in random trees"],
            ["One-Class SVM", "Static",
             "kernel=RBF\ngamma=scale, nu=0.05",
             "23 features,\nStandardScaler normalized",
             "Points outside learned\nnormality boundary"],
            ["Local Outlier Factor", "Static",
             "n_neighbors=20\ncontamination=0.05",
             "23 features,\nStandardScaler normalized",
             "Low local density relative\nto k-nearest neighbors"],
            ["Z-Score Threshold", "Static",
             "threshold: |z| > 3.0",
             "23 features,\nStandardScaler (z-scores)",
             "Any feature exceeds\n3 standard deviations"],
            ["Temporal Z-Score", "Temporal",
             "Train: first 50% of weeks\nTest: remaining 50%",
             "23 features per user\nper week",
             "Test-period deviation from\nown historical baseline > 3.0"],
            ["Feature Trajectory", "Temporal",
             "CUSUM: threshold=2.0\ndrift=0.5, min_run=3",
             "Weekly composite drift\n(mean z-score across features)",
             "Top 10% by cumulative\ndrift score"],
        ],
        col_widths=[1.3, 0.7, 1.8, 1.5, 2.0],
    )

    add_section_heading(doc, "7.2. Behavioral Drift Methods (Tier 2 — 3 Layers)", level=2)

    add_body(doc, (
        "Tier 2 transforms the same 23 features through a semantic embedding pipeline: "
        "scalars are serialized as prose text, embedded into high-dimensional vectors, "
        "composed into a single behavioral representation, and analyzed for drift direction "
        "against 12 reference concepts (10 threat + 2 benign)."
    ))

    create_table(doc,
        ["Layer", "Method", "Parameters", "Detection Logic"],
        [
            ["Layer 1:\nMagnitude", "Cumulative embedding\ndistance",
             "threshold=0.001\ndrift=0.0005, min_run=2",
             "Accumulating drift exceeds threshold.\n"
             "Note: 100% FP with real embeddings (not discriminating)"],
            ["Layer 2:\nDirection", "Baseline vs recent\ncosine similarity",
             "Baseline: first 50% weeks\nRecent: last 25% weeks",
             "Drift vector projected onto 12 reference\n"
             "concepts; max threat alignment analyzed"],
            ["Layer 3:\nConsistency", "Weekly threat alignment\nanalysis",
             "threshold=0.5, drift=0.05\nmin_run=2",
             "threat_consistency >= 0.40\n"
             "(40%+ of weeks show net threat direction)"],
        ],
        col_widths=[0.8, 1.5, 1.8, 3.0],
    )

    add_body(doc, (
        "Tier 2 limitation: The weighted average composition of 5 signal embeddings "
        "into a single composite dilutes zone-specific signals. An insider's data_behavior "
        "drift gets averaged with 4 stable zones, suppressing the detection signal below "
        "threshold. This motivates Tier 3's zone decomposition."
    ), italic=True)

    add_section_heading(doc, "7.3. Digital Entity Methods (Tier 3 — 9 Methods)", level=2)

    add_body(doc, (
        "Tier 3 decomposes the single composite embedding into 5 independent behavioral "
        "zones, computes zone-specific trajectories, builds cross-entity relationship "
        "embeddings, and runs 9 detection methods. Each method receives specific inputs "
        "and applies independent detection logic."
    ))

    create_table(doc,
        ["#", "Method", "Input Signal", "Key Parameters", "Detection Logic"],
        [
            ["1", "Velocity /\nAcceleration",
             "Weekly composite\ntrajectory",
             "Score = 0.4×velocity +\n0.3×acceleration + 0.3×drift",
             "Top 10% by composite score.\nDetects accelerating attacks."],
            ["2", "Regime\nShift",
             "Consecutive weekly\nembeddings",
             "regime_shift_threshold=0.7\n(cosine similarity cutoff)",
             "Top 10% by instability.\nDetects phase changes."],
            ["3", "Zone\nDivergence",
             "5 per-zone\ntrajectory drifts",
             "identity stable < 0.02\nbehavioral drifting > 0.05",
             "Top 10% by divergence.\nStandalone-sufficient in Combined."],
            ["4", "Relationship\nDrift",
             "User-Device\nHadamard product",
             "relationship_drift > 0.05",
             "Top 10% by relationship\ndrift magnitude."],
            ["5", "Contextual\n(Multi-Context)",
             "4 context-weighted\ncompositions",
             "threat_consistency > 0.30\n(30%+ threat-winning weeks)",
             "Top 10% by best consistency\nacross all 4 contexts."],
            ["6", "Cross-Entity\nCorrelation",
             "Relationship drift\nvectors (all users)",
             "similarity > 0.5\nmin_cluster_size = 3",
             "Member of any coordinated\ndrift cohort."],
            ["7", "Embedding\nDrift Accum.",
             "Baseline vs weekly\nsnapshot distance",
             "Score = 0.6×max_cusum +\n0.4×final_cusum",
             "Top 10% by cumulative\nexcess drift."],
            ["8", "Zone Threat\nDirection",
             "Per-zone drift vectors\n+ 12 reference concepts",
             "Zone-threat weights:\nrelevant=1.5, other=0.5",
             "Top 10% by best zone\nthreat consistency."],
            ["9", "Behavioral\nProgression",
             "Per-zone weekly\nthreat alignment",
             "Score = 0.6×best_tau +\n0.4×late_threat_mean",
             "Top 10% by Kendall trend\nstrength across zones."],
        ],
        col_widths=[0.3, 1.0, 1.2, 1.6, 2.0],
    )

    add_body(doc, "Tier 3 Combined (Ensemble):", bold=True, space_after=2)
    add_body(doc, (
        "The Combined method ensembles 6 core methods using a weighted geometric mean. "
        "Weights: velocity=0.10, zone_divergence=0.25, contextual=0.10, embedding_cusum=0.25, "
        "progression=0.30. Detection triggers when: Zone Divergence detects alone (standalone "
        "sufficient), OR 2+ core methods agree, OR composite score in top 10%."
    ))

    add_section_heading(doc, "7.4. Why a Single Method Is Not Enough", level=2)

    add_body(doc, (
        "Each attack type changes behavior in a different way, and no single detection "
        "approach captures all of them. The insider changes which data is accessed; the "
        "slow APT establishes persistent covert communication; the LOTL campaign repurposes "
        "legitimate tools; the telecom pivot shifts behavior across multiple zones at once. "
        "A detector tuned to one of these signatures misses the others."
    ))

    add_body(doc, (
        "This is the core argument for multi-phase composite scoring. Rather than selecting "
        "a single algorithm and accepting its blind spots, composite scoring evaluates five "
        "independent behavioral phases and fuses them into one ranked score. An attacker "
        "who evades one phase still registers on another: the slow APT that scores near zero "
        "on signal strength, breadth, and sustained deviation is caught entirely by novelty "
        "persistence; the insider that produces no novel infrastructure is caught by "
        "sustained deviation and breadth. The composite is detectable precisely because it "
        "is not reducible to any single measurement."
    ))

    add_callout(doc,
        "No single embedding/composite method detects all four campaigns — composite scoring "
        "cleanly separates only 2 of 4 (USR-234 and USR-042 fall below normal users). The "
        "clean 4 of 4 at 0 false positives comes from the multi-front threat-profile detector, "
        "which scores measurable known-bad profiles — and no single evasion technique defeats "
        "all of them simultaneously."
    )

    add_page_break(doc)

    # ==============================================================
    # 8. EMPIRICAL VALIDATION
    # ==============================================================
    add_section_heading(doc, "8. Empirical Validation", level=1)

    add_section_heading(doc, "8.1. Simulation Environment", level=2)

    create_table(doc,
        ["Parameter", "Value"],
        [
            ["Simulation period", "January 2025 — May 2026 (485 days)"],
            ["Total users", "250 (employees, contractors, service accounts)"],
            ["Total devices", "300"],
            ["Network segments", "25"],
            ["Applications", "60"],
            ["Log sources", "5 (authentication, file access, endpoint, network flow, DNS)"],
            ["Total events", "~15 million across 5 log sources (≈31,000 per day)"],
            ["Weekly feature vectors", "250 users × 70 weeks = 17,500 vectors"],
            ["Attack campaigns", "4 long-duration campaigns (100-240 days each)"],
            ["Analysis window", "70 weeks (weekly aggregation)"],
            ["Detection threshold", "90th percentile of composite scores"],
            ["False positive rate", "8.5% (21 FP / 246 normal users)"],
        ],
        col_widths=[2.0, 5.0],
    )

    add_section_heading(doc, "8.2. Attack Campaign Design", level=2)

    add_body(doc, (
        "Four long-duration attack campaigns were injected into the synthetic telemetry, "
        "each representing a distinct UEBA threat category. The campaigns were designed "
        "with realistic timing, escalation patterns, and MITRE ATT&CK technique mapping."
    ))

    create_table(doc,
        ["Campaign", "Threat Type", "Duration", "Target\nUser", "Key Techniques\n(MITRE ATT&CK)"],
        [
            ["ATK-004", "Insider Threat", "8 months\n(32 weeks)", "USR-156",
             "T1078 (Valid Accounts), T1083 (File Discovery),\n"
             "T1005 (Local Data), T1052 (Exfil via Physical),\n"
             "T1048 (Exfil Over Alternative Protocol)"],
            ["ATK-003", "Slow APT\n(C2 Beacon)", "180 days\n(26 weeks)", "USR-234",
             "T1071 (App Layer Protocol), T1573 (Encrypted Channel),\n"
             "T1568 (Dynamic Resolution), T1074 (Data Staged),\n"
             "T1048 (Exfil Over Alternative Protocol)"],
            ["ATK-007", "Volt Typhoon\nLOTL", "115 days\n(16 weeks)", "USR-042",
             "T1059 (Command Scripting), T1021 (Remote Services),\n"
             "T1047 (WMI), T1003 (Credential Dumping),\n"
             "T1105 (Ingress Tool Transfer)"],
            ["ATK-008", "Salt Typhoon\nTelecom", "100 days\n(14 weeks)", "USR-118",
             "T1195 (Supply Chain), T1071 (App Layer Protocol),\n"
             "T1005 (Local Data), T1041 (Exfil Over C2),\n"
             "T1568 (Dynamic Resolution)"],
        ],
        col_widths=[0.7, 1.0, 0.8, 0.7, 3.3],
    )

    add_section_heading(doc, "8.3. Detection Results: Traditional vs Threat-Profile Detector", level=2)

    add_body(doc, (
        "The four established traditional anomaly detection algorithms were run against "
        "the full 250-user population, alongside V-Intelligence UEBA's multi-front "
        "threat-profile detector. The results demonstrate the central finding: the traditional "
        "algorithms, calibrated to operationally acceptable false positive rates, detect "
        "at most one of the four attack campaigns. Embedding/composite scoring on its own "
        "cleanly separates 2 of 4; the threat-profile detector cleanly detects all four at "
        "0 false positives."
    ))

    create_table(doc,
        ["Method", "Approach", "USR-156\n(Insider)", "USR-234\n(APT)",
         "USR-042\n(Volt T.)", "USR-118\n(Salt T.)", "TP", "FP", "FP Rate"],
        [
            ["Isolation Forest", "Traditional", "MISSED", "MISSED", "MISSED", "MISSED",
             "0/4", "13", "5.3%"],
            ["Local Outlier Factor", "Traditional", "MISSED", "MISSED", "MISSED", "MISSED",
             "0/4", "11", "4.5%"],
            ["One-Class SVM", "Traditional", "MISSED", "MISSED", "MISSED", "MISSED",
             "0/4", "36", "14.6%"],
            ["Z-Score (|z|>3)", "Traditional", "MISSED", "MISSED", "DETECTED", "MISSED",
             "1/4", "24", "9.8%"],
            ["Threat-Profile Detector", "V-Intelligence\nUEBA", "DETECTED", "DETECTED", "DETECTED",
             "DETECTED", "4/4", "0", "0.0%"],
        ],
        col_widths=[1.4, 1.0, 0.7, 0.6, 0.6, 0.6, 0.5, 0.4, 0.6],
    )

    add_body(doc, (
        "No traditional algorithm detects the insider threat or the two nation-state "
        "campaigns at any acceptable false positive rate — every attack user's individual "
        "feature values remain within normal statistical ranges. Z-Score catches only the "
        "Volt Typhoon LOTL campaign, which produces the most overt single-feature spikes, "
        "and even then at a 9.8% false positive rate. The clean 4-of-4 detection at 0 false "
        "positives is achieved by the multi-front threat-profile detector, which scores "
        "measurable known-bad profiles using cohort-relative and raw-event features rather "
        "than thresholding any single metric. Embedding/composite scoring on its own cleanly "
        "separates only 2 of 4 (USR-234 and USR-042 fall below normal users)."
    ))

    add_section_heading(doc, "8.4. Per-Threat Analysis", level=2)

    add_body(doc, (
        "Composite scoring fuses five behavioral phases — signal strength, breadth, "
        "sustained deviation, context divergence, and novelty persistence — into a single "
        "ranked score per user. Each attack campaign is detected through a different "
        "combination of phases, which is precisely why a single-method approach fails and "
        "a multi-phase composite succeeds. The table below shows the actual phase "
        "contributions for each attack user."
    ))

    create_table(doc,
        ["Attack User", "Composite\nScore", "Rank\n(/250)", "Signal", "Breadth\n(>1.5σ)",
         "Sustained", "Novelty", "Dominant Phase"],
        [
            ["USR-118\n(Salt Typhoon)", "51.3", "#1", "29.9", "18", "9.6", "2.8",
             "Signal + Breadth"],
            ["USR-156\n(Insider)", "46.2", "#2", "30.0", "19", "8.4", "0.0",
             "Sustained + Breadth"],
            ["USR-234\n(Slow APT)", "19.4", "#7", "4.5", "1", "1.6", "13.0",
             "Novelty Persistence"],
            ["USR-042\n(Volt Typhoon)", "13.7", "#24", "6.4", "11", "3.8", "0.0",
             "Breadth + Signal"],
        ],
        col_widths=[1.3, 0.7, 0.5, 0.6, 0.7, 0.7, 0.6, 1.4],
    )

    add_body(doc, "USR-118 (Salt Typhoon Telecom) — Rank #1:", bold=True, space_after=2)
    add_body(doc, (
        "The telecom infrastructure pivot produces the strongest overall anomaly. It "
        "scores highest on signal strength and breadth — the attack touches authentication, "
        "network, and DNS behavior simultaneously, elevating 18 features beyond 1.5 standard "
        "deviations. Despite this, no single feature crosses a traditional detection "
        "threshold, which is why all four traditional algorithms miss it."
    ))

    add_body(doc, "USR-156 (Insider Threat) — Rank #2:", bold=True, space_after=2)
    add_body(doc, (
        "The insider's signature is sustained, broad behavioral drift. It scores highest "
        "on sustained deviation, reflecting the 8-month escalation, with 19 features "
        "elevated beyond 1.5 standard deviations. The insider gradually shifts from public "
        "to restricted and confidential file access while holding overall volume steady — "
        "the textbook 'same credentials, different data' signature that magnitude-based "
        "detection cannot see."
    ))

    add_body(doc, "USR-234 (Slow APT, C2 Beacon) — Rank #7:", bold=True, space_after=2)
    add_body(doc, (
        "This is the most important case. The slow APT scores low on every conventional "
        "phase — signal strength, breadth, and sustained deviation are all minimal because "
        "the behavioral change is small in magnitude. Without novelty persistence, its "
        "composite score would be approximately 6.4, ranking it near the middle of the "
        "population and rendering it undetectable. Novelty persistence contributes 13.0 "
        "points — the single largest phase contribution for any attacker — because a "
        "previously-unseen external IP appears in every one of the 60 post-baseline weeks "
        "with 100% persistence. This is the behavioral fingerprint of C2 beacon "
        "infrastructure, and it is the sole reason USR-234 is detected."
    ))

    add_body(doc, "USR-042 (Volt Typhoon LOTL) — Rank #24:", bold=True, space_after=2)
    add_body(doc, (
        "The living-off-the-land campaign uses legitimate administrative tools, producing "
        "a broad but moderate behavioral footprint — 11 features elevated beyond 1.5 "
        "standard deviations across access and endpoint behavior. It ranks #24 of 250, "
        "still within the top 10% and above the detection threshold, caught by the "
        "combination of breadth and signal strength."
    ))

    add_section_heading(doc, "8.5. False Positive Analysis", level=2)

    add_body(doc, (
        "False positive management is critical for operational deployment. A method that "
        "detects all threats but generates excessive false positives overwhelms SOC analysts "
        "and quickly gets ignored."
    ))

    add_body(doc, (
        "The critical comparison is not false positive rate in isolation but false "
        "positive rate at a given detection rate. A method with a low false positive rate "
        "that detects none of the actual attacks provides no operational value. The "
        "traditional algorithms achieve their low false positive rates precisely because "
        "they are insensitive to the behavioral changes the attackers exhibit — they are "
        "quiet because they are blind."
    ))

    create_table(doc,
        ["Method", "Attacks Detected", "FP Rate", "Operational Assessment"],
        [
            ["Local Outlier Factor", "0 of 4", "4.5%",
             "Low FP, but detects no attacks.\nQuiet because it is blind to\nbehavioral direction."],
            ["Isolation Forest", "0 of 4", "5.3%",
             "Low FP, but detects no attacks.\nMagnitude-only outlier detection."],
            ["Z-Score (|z|>3)", "1 of 4", "9.8%",
             "Catches only the overt LOTL\ncampaign; misses insider and\nboth nation-state campaigns."],
            ["One-Class SVM", "0 of 4", "14.6%",
             "Higher FP and detects no attacks —\nthe worst operational profile."],
            ["Composite Scoring", "2 of 4", "—",
             "Cleanly separates USR-156 and\nUSR-118; USR-234 and USR-042\nfall below normal users."],
            ["Threat-Profile Detector", "4 of 4", "0.0%",
             "Cleanly detects every campaign\nat 0 false positives via measurable\nknown-bad profiles (label-free)."],
        ],
        col_widths=[1.4, 1.1, 0.7, 2.8],
    )

    add_body(doc, (
        "At the 90th-percentile threshold, composite scoring flags 21 normal users alongside "
        "the 2 attackers it cleanly separates (USR-156 and USR-118); USR-234 and USR-042 fall "
        "below normal users on the composite and are recovered by the threat-profile detector. "
        "In an operational setting, the flagged normal users are not noise to be discarded but "
        "a prioritized investigation queue — the highest-behavioral-risk normal users, ranked "
        "and explained by which behavioral phases drove their score. The threshold is tunable: "
        "a higher percentile reduces false positives at the cost of detection margin."
    ))

    add_page_break(doc)

    # ==============================================================
    # 9. DETECTION PLAYBOOK
    # ==============================================================
    add_section_heading(doc, "9. Detection Playbook: Threat Type to Algorithm", level=1)

    add_body(doc, (
        "No single algorithm detects all threat types. The optimal approach matches "
        "detection methods to specific threats based on the behavioral signature each "
        "threat creates."
    ))

    add_section_heading(doc, "9.1. Per-Threat Recommendations", level=2)

    create_table(doc,
        ["Threat Type", "Dominant\nComposite Phase", "Why It Works", "What Single-Threshold Methods Miss"],
        [
            ["Insider\nThreat", "Sustained\nDeviation\n+ Breadth",
             "Gradual, broad drift across data-access\nfeatures over months — same credentials,\ndifferent data",
             "Magnitude stays within normal\nbounds; no single feature spikes"],
            ["Slow APT\n(C2 Beacon)", "Novelty\nPersistence",
             "A previously-unseen external IP recurs\nin every post-baseline week — the C2\nbeacon fingerprint",
             "Volume looks normal; no metric\ncrosses a threshold"],
            ["LOTL\n(Volt Typhoon)", "Breadth\n+ Signal\nStrength",
             "Broad behavioral footprint across access\nand endpoint features from repurposed\nlegitimate tools",
             "No malware signature; each tool\nuse looks legitimate in isolation"],
            ["Telecom\n(Salt Typhoon)", "Signal\nStrength\n+ Breadth",
             "Strong simultaneous anomaly across\nauthentication, network, and DNS\nbehavior",
             "No single feature crosses a\ndetection threshold"],
        ],
        col_widths=[1.0, 1.1, 2.7, 2.0],
    )

    add_body(doc, (
        "The operational implication is that composite scoring does not require the analyst "
        "to know in advance which threat is present. All five phases are computed for every "
        "user, and the composite ranks users by total behavioral risk regardless of which "
        "phase drives the score. The insider surfaces through sustained deviation; the slow "
        "APT through novelty persistence; both appear in the same ranked output."
    ))

    add_section_heading(doc, "9.2. Layered Deployment Strategy", level=2)

    add_body(doc, (
        "We recommend a three-layer deployment in which composite scoring is the primary "
        "detection engine, complemented by fast magnitude-based pre-filtering and "
        "investigation context."
    ))

    add_body(doc, "Layer 1: Magnitude Pre-Filter (Traditional Algorithms)", bold=True, space_after=2)
    add_bullet(doc, "Question answered: Did any feature spike overtly?")
    add_bullet(doc, "Role: Fast, low-cost first pass that catches the most overt volume anomalies "
               "and provides a familiar baseline for SOC analysts.")
    add_bullet(doc, "Limitation: Detects at most one of the four campaigns; insensitive to "
               "behavioral direction. Used as a complement, not the primary detector.")

    add_body(doc, "Layer 2: Composite Scoring + Threat-Profile Detector (Primary Detection Engine)", bold=True, space_after=2)
    add_bullet(doc, "Question answered: Which users are behaving anomalously across multiple "
               "behavioral phases, or match a measurable known-bad profile?")
    add_bullet(doc, "Role: The core capability. Composite scoring computes five behavioral phases "
               "per user and cleanly separates 2 of 4 campaigns; the multi-front threat-profile "
               "detector scores measurable known-bad profiles and recovers the remaining 2 for a "
               "clean 4 of 4 at 0 false positives.")
    add_bullet(doc, "Strengths: Together they catch insider, slow APT, LOTL, and telecom campaigns "
               "that no single traditional method detects; ranked output prioritizes investigation.")

    add_body(doc, "Layer 3: Investigation Context (Drift Direction, MITRE, Kill Chains)", bold=True, space_after=2)
    add_bullet(doc, "Question answered: What kind of threat is this, and how does it map to "
               "known adversary behavior?")
    add_bullet(doc, "Role: Enriches each high-ranked user with behavioral drift direction, "
               "MITRE ATT&CK technique mapping, and kill-chain correlation for triage.")
    add_bullet(doc, "Strengths: Transforms a ranked score into an actionable, explainable "
               "investigation lead.")

    add_callout(doc,
        "Composite scoring plus the multi-front threat-profile detector is the minimum viable "
        "deployment: composite scoring cleanly separates 2 of 4 campaigns as a ranked output, "
        "and the threat-profile detector recovers the remaining 2 for a clean 4 of 4 at 0 false "
        "positives. The magnitude pre-filter and investigation-context layers add speed and "
        "explainability for production environments."
    )

    add_page_break(doc)

    # ==============================================================
    # 10. DEPLOYMENT RECOMMENDATIONS
    # ==============================================================
    # ==============================================================
    # 10. MULTI-PHASE COMPOSITE SCORING
    # ==============================================================
    add_section_heading(doc, "10. Multi-Phase Composite Scoring", level=1)

    add_body(doc, (
        "The central innovation of V-Intelligence UEBA is multi-phase composite scoring. "
        "Instead of relying on any single detection method — each of which catches some "
        "attack types but misses others — composite scoring asks five independent questions "
        "about each user's behavior and fuses the answers into a single ranked score. "
        "A normal user might score high on one phase by random variation; an attacker "
        "scores high on multiple phases simultaneously."
    ))

    add_section_heading(doc, "10.1. The Five Detection Phases", level=2)

    create_table(doc,
        ["Phase", "What It Measures", "Why It Matters"],
        [
            ["Signal Strength",
             "Magnitude of the top 3 behavioral anomalies relative to role peer group",
             "Catches users with extreme deviations on their strongest signals"],
            ["Breadth",
             "Count of features exceeding 1.5 standard deviations",
             "Distinguishes single-feature noise from multi-feature attack patterns"],
            ["Sustained Deviation",
             "Persistence of anomalies across multiple weeks",
             "Separates transient events from persistent campaigns (insider escalation)"],
            ["Context Divergence",
             "Anomaly consistency across multiple analytical contexts "
             "(insider, APT, privilege audit, normal operations)",
             "Attackers flag under attack-specific contexts; normal users do not"],
            ["Novelty Persistence",
             "Recurrence of previously-unseen behaviors "
             "(new external IPs, new domains) across weeks",
             "Catches C2 beacon infrastructure invisible to volume-based detection"],
        ],
        col_widths=[1.2, 2.5, 3.0],
    )

    add_callout(doc,
        "Each phase is independently insufficient. Combined, they create a detection "
        "surface that no single evasion technique can defeat. An attacker who suppresses "
        "signal strength still fails on breadth; one who limits breadth still fails on "
        "sustained deviation or novelty persistence."
    )

    add_section_heading(doc, "10.2. Per-Attacker Composite Breakdown", level=2)

    add_body(doc, (
        "The following table shows how each attack campaign scores across the five "
        "phases. Note the distinct detection profiles — each attack type is caught "
        "by a different combination of phases:"
    ))

    create_table(doc,
        ["Attack User", "Type", "Composite\nScore", "Rank\n(/250)",
         "Primary Detection Phase(s)"],
        [
            ["USR-118", "Salt Typhoon\n(Telecom)", "51.3", "#1",
             "Signal Strength (26.8) + Breadth (9 features)\n"
             "Massive multi-zone anomaly across auth, network, DNS"],
            ["USR-156", "Insider Threat", "46.2", "#2",
             "Sustained Deviation (6.8) + Breadth (8 features)\n"
             "8-month escalation from public to restricted files"],
            ["USR-234", "Slow APT\n(C2 Beacon)", "19.4", "#7",
             "Novelty Persistence (13.0)\n"
             "C2 beacon IP persists in 60/60 post-baseline weeks"],
            ["USR-042", "Volt Typhoon\n(LOTL)", "13.7", "#24",
             "Signal Strength (7.2) + Context Divergence (2.3)\n"
             "LOTL tools create cross-context anomalies"],
        ],
        col_widths=[0.7, 1.0, 0.7, 0.5, 3.8],
    )

    add_body(doc, (
        "On the composite, USR-156 (insider) and USR-118 (Salt Typhoon) are cleanly separated "
        "above all normal users, but USR-234 (slow APT) and USR-042 (Volt Typhoon LOTL) fall "
        "below normal users — composite scoring alone cleanly separates only 2 of 4. The clean "
        "4 of 4 detection at 0 false positives is achieved by the separate multi-front "
        "threat-profile detector, which targets measurable known-bad profiles (C2-beacon, "
        "DGA-DNS, LOTL-process, cohort-rare access, recon-fanout, insider-collection), "
        "label-free."
    ))

    add_section_heading(doc, "10.3. USR-234 Novelty Persistence — Case Study", level=2)

    add_body(doc, (
        "USR-234 represents the hardest detection challenge: a 180-day slow APT with "
        "command-and-control beaconing via DNS. The malware contacts an external C2 "
        "server approximately every 6 hours (~4 beacons per day). Total DNS query "
        "volume, network bytes, and authentication counts all remain within normal ranges."
    ))

    add_body(doc, (
        "Without novelty persistence, USR-234's composite score would be 6.4 — ranking "
        "approximately #80 out of 250 users, well below any detection threshold. The "
        "user's signal strength, breadth, and sustained deviation are all low because "
        "the behavioral change is minimal in magnitude."
    ))

    add_body(doc, (
        "Novelty persistence identifies the critical signal: a single external IP address "
        "that never appeared during the baseline period now appears in every post-baseline "
        "week — 60 consecutive weeks with 100% persistence. Normal users occasionally "
        "contact new external IPs, but those contacts are transient (appearing once or "
        "twice then disappearing). USR-234's novel IP is persistent, regular, and "
        "recurring — the behavioral fingerprint of C2 beacon infrastructure."
    ))

    add_body(doc, (
        "This novelty signal alone contributes 13.0 points to the composite score, "
        "elevating USR-234 from rank #80 (undetectable) to rank #7 (clearly flagged). "
        "This case demonstrates why multi-phase scoring is essential: no other phase "
        "catches this attack, but novelty persistence alone is sufficient."
    ))

    add_page_break(doc)

    # ==============================================================
    # 11. WHY SEMANTIC EMBEDDINGS
    # ==============================================================
    add_section_heading(doc, "11. Why Semantic Embeddings Outperform Raw Feature Vectors", level=1)

    add_body(doc, (
        "A fundamental design question: given 23 numeric features per user per week, "
        "why not represent each observation as a 23-dimensional vector and compute "
        "distance directly? This section presents the empirical evidence for semantic "
        "embeddings over raw feature vectors."
    ))

    add_section_heading(doc, "11.1. Two Pipelines: Feature-Space vs Semantic", level=2)

    add_body(doc, "Feature-Space Pipeline (Traditional):", bold=True, space_after=2)
    add_body(doc, (
        "Raw telemetry → weekly aggregation (23 scalar features) → feature vector "
        "v = [auth_total, auth_fail_rate, ..., dns_nxdomain_ratio] → statistical "
        "distance ||v_week2 - v_week1|| → CUSUM accumulation → threshold detection."
    ))

    add_body(doc, "Semantic Pipeline (V-Intelligence UEBA):", bold=True, space_after=2)
    add_body(doc, (
        "Raw telemetry → weekly aggregation (23 scalar features) → natural language "
        "serialization that qualifies each metric (e.g., '30 files total, 15% restricted, "
        "8% confidential') → behavioral embedding → cosine distance between consecutive "
        "weekly embeddings → CUSUM accumulation → threshold detection."
    ))

    add_body(doc, (
        "The critical addition is the serialization step. Raw numbers become qualified "
        "descriptions: not just 'file_total=30' but '30 total, 15% restricted, 8% "
        "confidential.' The embedding model reads this text and produces a vector that "
        "captures the meaning of the behavioral pattern."
    ))

    add_section_heading(doc, "11.2. Three Limitations of Raw Feature Distance", level=2)

    add_body(doc, "Limitation 1: Scale Dominance", bold=True, space_after=2)
    add_body(doc, (
        "The 23 features span vastly different scales (network bytes in millions vs. "
        "authentication fail rate as a decimal fraction). Even with normalization, "
        "high-variance features dominate the distance calculation. A 10% change in "
        "bytes_out swamps a 500% change in file_restricted_ratio — yet the latter "
        "is the primary insider threat signal."
    ))

    add_body(doc, "Limitation 2: Feature Independence", bold=True, space_after=2)
    add_body(doc, (
        "Euclidean distance treats each feature as an independent dimension. It cannot "
        "represent that off-hours authentication combined with restricted file access "
        "combined with high clearance level constitutes a threat pattern, while each "
        "value individually falls within normal ranges."
    ))

    add_body(doc, "Limitation 3: 'What' vs 'How Much'", bold=True, space_after=2)
    add_body(doc, (
        "This is the critical limitation. The insider (USR-156) accesses 30 files both "
        "weeks — file_total distance is zero. But week 1 was public files; week 2 was "
        "restricted files. The raw vector sees no change. The semantic embedding sees "
        "a fundamental behavioral shift because '1.2% restricted' and '15% restricted' "
        "are semantically distant descriptions."
    ))

    add_section_heading(doc, "11.3. Empirical Signal Separation", level=2)

    add_body(doc, (
        "Empirical validation refines the theoretical analysis. Feature-space CUSUM "
        "(cumulative drift on raw 23-dimensional vectors) fires on roughly 99% of normal "
        "users — an operationally useless false-positive rate. Semantic (embedding) CUSUM "
        "helps on some attacks but not others: its advantage is attack-dependent, not a "
        "universal separation of all four attackers above the normal band."
    ))

    add_body(doc, (
        "The benefit is real but selective. Embedding CUSUM fires roughly 30 weeks earlier "
        "than feature-space CUSUM for the insider (USR-156) and the LOTL attack (USR-042), "
        "fires LATER for the volume-driven attack (USR-118), and never separates the slow-APT "
        "(USR-234). It does not separate all 4. The clean 4 of 4 detection at 0 false positives "
        "is achieved instead by the separate multi-front threat-profile detector "
        "(threat_profile_detector.py), which scores measurable known-bad profiles — C2-beacon, "
        "DGA-DNS, LOTL-process, cohort-rare access, recon-fanout, and insider-collection — using "
        "cohort-relative and raw-event features, label-free."
    ))

    add_page_break(doc)

    # ==============================================================
    # 12. DEPLOYMENT (renumbered from 10)
    # ==============================================================
    add_section_heading(doc, "12. Operational Deployment for DoD and IC", level=1)

    add_body(doc, (
        "V-Intelligence UEBA is designed for deployment across Department of Defense "
        "and Intelligence Community networks, from enterprise IT infrastructure to "
        "tactical and special operations environments."
    ))

    add_section_heading(doc, "12.1. Deployment Phases", level=2)

    add_body(doc, "Data Requirements:", bold=True, space_after=2)
    add_bullet(doc, "Minimum 5 log sources: authentication, file access, endpoint telemetry, network flow, DNS")
    add_bullet(doc, "Weekly aggregation (minimum 6 weeks of history for trajectory analysis)")
    add_bullet(doc, "Entity resolution: map log events to users, devices, applications, network segments")
    add_bullet(doc, "Classification metadata: data sensitivity labels on file access events")

    add_body(doc, "Phased Deployment", bold=True, space_after=2)

    create_table(doc,
        ["Phase", "Duration", "Activities", "Success Criteria"],
        [
            ["Phase 1:\nBaseline", "4-6 weeks",
             "Deploy data collection.\nBuild behavioral baselines.\nValidate feature engineering.",
             "23 features computed per user\nper week. Baselines stable."],
            ["Phase 2:\nDetection", "2-4 weeks",
             "Deploy multi-phase composite scoring.\nTune thresholds on historical data.\nEstablish FP baseline.",
             "False positive rate < 10%.\nAll attack types detectable."],
            ["Phase 3:\nValidation", "2-4 weeks",
             "Purple team exercises.\nInject controlled attack patterns.\nMeasure detection rates.",
             "Insider and APT patterns\ndetected within 2 weeks of onset."],
            ["Phase 4:\nProduction", "Ongoing",
             "Add Layer 3 (Drift Accum. + Regime).\nIntegrate with SOAR playbooks.\nContinuous threshold tuning.",
             "All 3 layers operational.\nSOC investigating < 5 alerts/day."],
        ],
        col_widths=[0.8, 0.8, 2.5, 2.5],
    )

    add_section_heading(doc, "12.2. Integration Points", level=2)
    add_bullet(doc, "SIEM/SOAR: UEBA alerts feed into existing incident response workflows "
               "with behavioral context (which zone drifted, confidence score, threat type)")
    add_bullet(doc, "Threat Intelligence: Reference concept library updated with emerging "
               "threat patterns (new APT techniques, novel LOTL tools)")
    add_bullet(doc, "MITRE ATT&CK Mapping: Every detection maps to specific techniques, "
               "enabling alignment with existing threat hunting programs")
    add_bullet(doc, "Zero Trust Architecture: UEBA behavioral scores feed into continuous "
               "verification policies (adjust access based on behavioral risk)")

    add_body(doc, "Scaling Considerations", bold=True, space_after=2)
    add_bullet(doc, "250-user validation at enterprise scale; production environments "
               "with larger populations will achieve even better FP resolution through "
               "larger role-group baselines")
    add_bullet(doc, "LOF score-based thresholding (negative_outlier_factor_) instead of "
               "fixed contamination parameter for scale-independent detection")
    add_bullet(doc, "Embedding computation is parallelizable and cacheable — incremental "
               "cost after initial baseline is minimal")

    add_section_heading(doc, "12.3. Applicability Across Mission Environments", level=2)

    add_body(doc, (
        "V-Intelligence UEBA is designed to operate across the full spectrum of "
        "DoD and IC mission environments:"
    ))

    add_bullet(doc, "Personnel operating across multiple networks and classification "
               "levels generate complex behavioral baselines that traditional per-network "
               "monitoring cannot correlate. V-Intelligence UEBA builds unified behavioral "
               "profiles across authentication, file access, network, and endpoint telemetry.",
               bold_prefix="Multi-domain behavioral profiles: ")
    add_bullet(doc, "Containerized architecture deploys on standard infrastructure with "
               "no specialized hardware requirements. Baseline establishment requires "
               "4-6 weeks of telemetry. Supports enterprise data centers, tactical edge, "
               "and cloud-hosted environments.",
               bold_prefix="Rapid deployment: ")
    add_bullet(doc, "Insider threats in sensitive environments carry disproportionate risk "
               "due to access to compartmented programs, operational plans, and intelligence "
               "sources. Zone-decomposed detection identifies the specific behavioral "
               "dimension that changed, enabling targeted investigation without alert fatigue.",
               bold_prefix="High-consequence insider detection: ")
    add_bullet(doc, "Nation-state adversaries (China, Russia, Iran) actively target DoD "
               "infrastructure for intelligence collection. The Salt Typhoon and Volt "
               "Typhoon campaigns validated in this whitepaper mirror real-world APT "
               "techniques used against US military and telecom infrastructure.",
               bold_prefix="Nation-state APT detection: ")
    add_bullet(doc, "Autonomous behavioral detection reduces SOC analyst workload while "
               "increasing detection coverage. Composite scoring produces a single ranked "
               "list — analysts investigate the top-ranked users, not thousands of raw alerts.",
               bold_prefix="AI-enabled autonomous detection: ")

    add_page_break(doc)

    # ==============================================================
    # 13. CONCLUSION (renumbered from 11)
    # ==============================================================
    add_section_heading(doc, "13. Conclusion", level=1)

    add_body(doc, (
        "User and Entity Behavior Analytics represents a necessary evolution in "
        "cybersecurity detection. The four threat categories examined in this whitepaper "
        "— insider threats, advanced persistent threats, living-off-the-land attacks, "
        "and telecom infrastructure compromises — share a common characteristic that "
        "renders traditional detection approaches insufficient: they operate within the "
        "bounds of authorized access, using legitimate credentials and tools, while "
        "gradually shifting their behavioral profile toward operational objectives."
    ))

    add_body(doc, (
        "The V-Intelligence UEBA empirical validation — 250 users, 485 days, ~15 million "
        "events, 4 embedded attack campaigns — demonstrates four key architectural insights:"
    ))

    add_body(doc, "1. Magnitude alone is insufficient.", bold=True, space_after=2)
    add_body(doc, (
        "Traditional anomaly detection (Isolation Forest, SVM, LOF, Z-Score on 23 scalar "
        "features) detects 0 of 4 attackers at operationally acceptable false positive rates. "
        "The insider's feature values never exceed statistical thresholds — the attack is a "
        "behavioral direction change, not a magnitude spike. The APT's C2 beaconing hides "
        "within normal DNS query volumes."
    ))

    add_body(doc, "2. Semantic embeddings outperform raw feature vectors.", bold=True, space_after=2)
    add_body(doc, (
        "Raw feature distance (Euclidean/Mahalanobis on 23-dimensional vectors) cannot "
        "distinguish behavioral changes at constant volume. When an insider accesses 30 files "
        "both weeks but shifts from public to restricted content, the raw feature distance is "
        "near zero. Semantic embeddings capture this qualitative shift because they encode "
        "the meaning of the behavioral pattern, not just the magnitude of individual metrics. "
        "Empirically, the embedding's benefit is attack-dependent: semantic CUSUM fires "
        "roughly 30 weeks earlier than feature-space CUSUM for the insider (USR-156) and the "
        "LOTL attack (USR-042), fires LATER for the volume-driven attack (USR-118), and never "
        "separates the slow-APT (USR-234). It does not separate all four attackers."
    ))

    add_body(doc, "3. The multi-front threat-profile detector achieves 4/4 detection.", bold=True, space_after=2)
    add_body(doc, (
        "Embedding/composite scoring on its own cleanly separates only 2 of 4 attacks — "
        "USR-156 (insider) and USR-118 (Salt Typhoon) rank above all normal users, while "
        "USR-234 (slow APT) and USR-042 (Volt Typhoon LOTL) fall below normal users. The clean "
        "4 of 4 detection at 0 false positives is achieved by the multi-front threat-profile "
        "detector (threat_profile_detector.py), which scores measurable known-bad profiles — "
        "C2-beacon, DGA-DNS, LOTL-process, cohort-rare access, recon-fanout, and "
        "insider-collection — using cohort-relative and raw-event features, label-free."
    ))

    add_body(doc, "4. Novelty persistence catches what even semantic drift misses.", bold=True, space_after=2)
    add_body(doc, (
        "USR-234 (Slow APT) has minimal semantic drift — its behavioral embedding changes "
        "are small. Without novelty persistence detection, its composite score would be 6.4 "
        "(rank ~80, undetectable). Novelty persistence identifies that a previously-unseen "
        "external IP appears in every post-baseline week for over a year — the behavioral "
        "fingerprint of C2 beacon infrastructure. This single signal elevates the score to "
        "19.4 (rank #7), placing it firmly within the detected set."
    ))

    add_callout(doc,
        "The insider changes WHAT they access, not HOW MUCH. The APT changes WHERE they "
        "communicate, not HOW OFTEN. The LOTL attack changes WHEN legitimate tools are used, "
        "not WHICH tools. V-Intelligence UEBA detects these directional changes by asking: "
        "is this entity behaving consistently with its established baseline across each "
        "behavioral dimension? Multi-phase composite scoring ensures that no single evasion "
        "technique can defeat all five detection phases simultaneously."
    )

    add_page_break(doc)

    # ==============================================================
    # 12. REFERENCES
    # ==============================================================
    add_section_heading(doc, "14. References and Further Reading", level=1)

    refs = [
        "CISA Advisory AA23-144A: People's Republic of China State-Sponsored Cyber Actor "
        "Living off the Land to Evade Detection (May 2023)",
        "CISA Advisory AA24-038A: PRC State-Sponsored Actors Compromise and Maintain "
        "Persistent Access to U.S. Critical Infrastructure (February 2024)",
        "Verizon 2024 Data Breach Investigations Report — Insider Threat Analysis",
        "MITRE ATT&CK Framework v14 — Enterprise Techniques Matrix",
        "NIST SP 800-53 Rev. 5 — Security and Privacy Controls for Information Systems "
        "and Organizations (SI-4: Information System Monitoring)",
        "NIST SP 800-207 — Zero Trust Architecture",
        "Gartner Market Guide for User and Entity Behavior Analytics (2024 edition)",
        "Executive Order 14028 — Improving the Nation's Cybersecurity (May 2021)",
        "OMB M-22-09 — Moving the U.S. Government Toward Zero Trust Cybersecurity Principles",
        "SolarWinds SUNBURST Analysis — CISA Emergency Directive 21-01",
        "Salt Typhoon Campaign — Congressional briefings and White House statements (2024-2025)",
    ]

    for i, ref in enumerate(refs, 1):
        add_bullet(doc, ref, bold_prefix=f"[{i}] ")

    doc.add_paragraph()
    add_body(doc, (
        "For technical implementation details including algorithm parameters, threshold "
        "values, and detection logic, refer to the V-Intelligence UEBA Technical Specification document."
    ), italic=True)



def build_whitepaper(audience_key=None):
    """Build whitepaper for a specific audience, or all 4 if audience_key is None."""
    if audience_key is None:
        for key in AUDIENCE_CONFIGS:
            build_whitepaper(key)
        # Also build the universal version
        _build_universal()
        return

    cfg = AUDIENCE_CONFIGS[audience_key]
    doc = _init_doc()
    _build_title_page(doc, audience_key)
    _build_exec_summary(doc, audience_key)
    _build_core(doc)

    out_path = os.path.join(os.path.dirname(__file__), cfg["filename"])
    doc.save(out_path)
    print(f"  {audience_key}: {cfg['filename']}")


def _build_universal():
    """Build universal version with generic intro (no audience-specific scenario)."""
    doc = _init_doc()

    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("User and Entity Behavior Analytics\n(UEBA)")
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "A Comprehensive Framework for Detecting Advanced Cyber Threats\n"
        "Through Behavioral Baseline, Drift, and Direction Analysis"
    )
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE

    for _ in range(2):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "22nd Century Technologies, Inc.\n"
        "V-Intelligence UEBA Program\n\n"
        "June 2026 — Version 2.0"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY
    add_page_break(doc)

    add_section_heading(doc, "1. Executive Summary", level=1)

    add_body(doc, (
        "Modern cyber threats have evolved beyond the capabilities of traditional "
        "security tools. Nation-state actors, sophisticated insider threats, and "
        "living-off-the-land (LOTL) techniques now operate below the detection "
        "thresholds of rule-based SIEM systems and signature-based intrusion detection. "
        "These attacks succeed not by tripping alarms, but by carefully staying within "
        "the statistical bounds of normal behavior while gradually shifting their "
        "operational profile over weeks and months."
    ))

    add_body(doc, (
        "This whitepaper presents a comprehensive UEBA framework validated through the "
        "V-Intelligence UEBA program. Using 485 days of synthetic telemetry across 250 users "
        "generating approximately 15 million events across five log sources with 4 embedded attack campaigns — an 8-month "
        "insider threat, a 180-day slow APT with C2 beaconing, a 115-day Volt Typhoon "
        "living-off-the-land campaign, and a 100-day Salt Typhoon telecom infrastructure "
        "attack — we demonstrate that:"
    ))

    add_bullet(doc, "Traditional anomaly detection (Isolation Forest, SVM, LOF, Z-Score) "
               "detects 0 of 4 attackers at operationally acceptable false positive rates.",
               bold_prefix="Finding 1: ")
    add_bullet(doc, "Behavioral drift analysis alone fails when attack signals from one "
               "behavioral zone are diluted by four stable zones.",
               bold_prefix="Finding 2: ")
    add_bullet(doc, "Multi-phase composite scoring cleanly separates 2 of 4 attacks (USR-156 "
               "and USR-118); the separate multi-front threat-profile detector recovers USR-234 "
               "and USR-042 for a clean 4 of 4 at 0 false positives via measurable known-bad "
               "profiles.",
               bold_prefix="Finding 3: ")
    add_bullet(doc, "Semantic behavioral embeddings outperform raw feature vectors by "
               "capturing what kind of change occurred, not just how much changed.",
               bold_prefix="Finding 4: ")

    add_callout(doc,
        "Key insight: The insider changes WHAT they access, not HOW MUCH. "
        "Traditional algorithms measure magnitude. UEBA measures direction."
    )
    add_page_break(doc)

    _build_core(doc)

    out_path = os.path.join(os.path.dirname(__file__),
                            "UEBA_Behavioral_Intelligence_Whitepaper.docx")
    doc.save(out_path)
    print(f"  Universal: UEBA_Behavioral_Intelligence_Whitepaper.docx")


if __name__ == "__main__":
    print("Building whitepapers...")
    build_whitepaper()
    print("Done — 5 documents generated.")

