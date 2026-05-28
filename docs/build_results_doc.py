#!/usr/bin/env python3
"""Build ACECARD Results and Findings Document (Word).

Presents empirical results from the three-tier detection system evaluation:
17 detection methods tested against 4 attack campaigns across 250 users
and 4.2M events. Focuses on WHAT was found and business value, not
internal algorithmic details.

Output: docs/ACECARD_Results_and_Findings.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
BLUE = RGBColor(0x1B, 0x4F, 0x72)
TEAL = RGBColor(0x0E, 0x6B, 0x8A)
GOLD = RGBColor(0xB7, 0x95, 0x0B)
RED_ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREEN_ACCENT = RGBColor(0x1E, 0x8A, 0x49)
BLACK = RGBColor(0x00, 0x00, 0x00)

OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
                           "ACECARD_Results_and_Findings.docx")


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
    return h


def add_body_text(doc, text, bold=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    return p


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "0D1B2A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if val == "DETECTED":
                        run.font.color.rgb = GREEN_ACCENT
                        run.bold = True
                    elif val == "MISSED":
                        run.font.color.rgb = RED_ACCENT
                        run.bold = True

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


# ==========================================================================
# DOCUMENT BUILDER
# ==========================================================================

def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK

    # ── COVER PAGE ──────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ACECARD THREE-TIER DETECTION SYSTEM")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Results and Findings")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Empirical Evaluation: 17 Detection Methods × 4 Attack Campaigns"
    )
    run.font.size = Pt(16)
    run.font.color.rgb = TEAL

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "250 Users | 4.2M Events | 133 Days | Real OpenAI Embeddings"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "22nd Century Technologies, Inc.\n"
        "ACECARD Program\n"
        "May 2025"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_ACCENT

    add_page_break(doc)

    # ── TABLE OF CONTENTS ──────────────────────────────────────────
    add_section_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Experimental Setup",
        "   2.1. Synthetic Data Summary",
        "   2.2. Attack Campaigns",
        "   2.3. Detection Methods Evaluated",
        "3. Consolidated Detection Results",
        "   3.1. Full Results Matrix (17 Methods × 4 Attacks)",
        "   3.2. Tier Comparison Summary",
        "4. Per-Attack Findings",
        "   4.1. USR-156 — Insider Threat",
        "   4.2. USR-234 — Slow APT",
        "   4.3. USR-042 — Volt Typhoon LOTL",
        "   4.4. USR-118 — Salt Typhoon Telecom",
        "5. Key Findings",
        "   5.1. Finding 1: No Single Method Catches All Attacks",
        "   5.2. Finding 2: Different Threats Require Different Detection Approaches",
        "   5.3. Finding 3: Zone Decomposition is the Critical Innovation",
        "   5.4. Finding 4: Interpretable Diagnostics Enable Analyst Action",
        "   5.5. Finding 5: Tier 3 Combined Achieves Best Coverage at Viable FP",
        "6. Recommendations",
        "   6.1. Deployment Strategy",
        "   6.2. Next Steps",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        if not item.startswith("   "):
            run.bold = True
    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "1. Executive Summary", level=1)

    add_body_text(doc, (
        "This document presents the empirical results from the ACECARD three-tier "
        "detection system evaluation. Seventeen detection methods spanning traditional "
        "statistical techniques, ACECARD behavioral analytics, and the advanced Digital "
        "Entity architecture were tested against 4 real-world-modeled attack campaigns "
        "hidden within 4.2 million synthetic enterprise events. The key findings are:"
    ))

    add_bullet(doc, (
        "Tier 3 Combined is the ONLY method that detects all 4 attacks (4/4 at 15.0% "
        "false positive rate). No other single method achieves complete coverage at a "
        "viable false positive rate."
    ), bold_prefix="Complete Coverage: ")

    add_bullet(doc, (
        "No single Tier 1 or Tier 2 method achieves 4/4 detection. The best Tier 1 "
        "method is Feature CUSUM at 3/4 (8.9% FP), which misses the slow APT campaign."
    ), bold_prefix="Tier 1/2 Gap: ")

    add_bullet(doc, (
        "Tier 3 provides interpretable diagnostics: not just 'anomaly detected' but "
        "specific outputs such as 'data_behavior zone drifting toward insider_threat_fast' "
        "— telling analysts exactly where to look and what to look for."
    ), bold_prefix="Interpretable Output: ")

    add_bullet(doc, (
        "Attack-injected events represent just 0.06% of total data — approximately "
        "2,661 malicious events hidden in a 4.2 million event haystack. This reflects "
        "the real-world challenge of detecting sophisticated threats in enterprise telemetry."
    ), bold_prefix="Needle in a Haystack: ")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  2. EXPERIMENTAL SETUP
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "2. Experimental Setup", level=1)

    add_section_heading(doc, "2.1. Synthetic Data Summary", level=2)

    create_table(doc,
        ["Parameter", "Value"],
        [
            ["Observation Window", "January 1 – May 13, 2025 (133 days)"],
            ["Weekly Windows", "19 weeks"],
            ["Total Users", "250 (including 4 attack targets)"],
            ["Total Events", "4,214,251"],
            ["Attack-Injected Events", "2,661 (0.06%)"],
            ["Log Types", "7 (auth, file_access, endpoint, network, dns, app, privilege)"],
            ["Features Per User Per Week", "23 scalar features"],
            ["Embedding Model", "OpenAI text-embedding-3-small"],
            ["Embedding API Calls", "30,613 (Tier 3 zone texts)"],
            ["Random Seed", "42 (fully reproducible)"],
        ],
        col_widths=[2.5, 4.0],
    )

    add_section_heading(doc, "2.2. Attack Campaigns", level=2)

    create_table(doc,
        ["Attack", "Target", "Type", "Duration", "Behavioral Signal"],
        [
            ["ATK-004", "USR-156", "Insider Threat", "60 days\n(4 phases)",
             "File classification escalation,\noff-hours access, data staging"],
            ["ATK-003", "USR-234", "Slow APT / C2", "65 days",
             "C2 beacons, lateral movement,\nexternal network traffic"],
            ["ATK-007", "USR-042", "Volt Typhoon\nLOTL", "60 days",
             "Living-off-the-land tools,\ndual data+network drift"],
            ["ATK-008", "USR-118", "Salt Typhoon\nTelecom", "60 days",
             "CDR database access, DNS\ntunneling, cross-segment flows"],
        ],
        col_widths=[0.7, 0.7, 1.1, 0.9, 2.5],
    )

    add_body_text(doc, (
        "Attack events represent 0.06% of total data — approximately 20 injected events "
        "per day across 4 campaigns hidden in ~31,000 normal daily events. Each attack is "
        "designed to evade specific detection approaches: the insider uses valid credentials "
        "with gradually escalating access; the APT maintains low-volume persistent C2; Volt "
        "Typhoon uses only legitimate admin tools; Salt Typhoon targets infrastructure with "
        "encrypted channels."
    ))

    add_section_heading(doc, "2.3. Detection Methods Evaluated", level=2)

    create_table(doc,
        ["Tier", "Methods", "Count"],
        [
            ["Tier 1\n(Traditional)",
             "Isolation Forest, One-Class SVM, LOF,\nZ-Score, Temporal Z-Score, Feature CUSUM",
             "6"],
            ["Tier 2\n(ACECARD Basic)",
             "ACECARD Direction, IForest+ACECARD Combined",
             "2"],
            ["Tier 3\n(Digital Entity)",
             "Velocity/Acceleration, Regime Shift, Zone Divergence,\n"
             "Relationship Drift, Contextual, Cross-Entity,\n"
             "Embedding CUSUM, Zone Threat Direction,\n"
             "Behavioral Progression, T3 Combined",
             "9 + 1"],
            ["Total", "17 detection methods + 1 combined", "18"],
        ],
        col_widths=[1.3, 3.5, 0.8],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  3. CONSOLIDATED DETECTION RESULTS
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "3. Consolidated Detection Results", level=1)

    add_section_heading(doc, "3.1. Full Results Matrix (17 Methods × 4 Attacks)", level=2)

    add_body_text(doc, (
        "The table below is the centerpiece of this evaluation. Each row represents "
        "a detection method; each column shows whether that method detected the "
        "corresponding attack campaign. TP = True Positives (attacks correctly detected), "
        "FP = False Positives (benign users incorrectly flagged), FP Rate = FP / total "
        "non-attack users."
    ))

    create_table(doc,
        ["#", "Method", "Tier", "USR-156\n(Insider)", "USR-234\n(APT)",
         "USR-042\n(Volt)", "USR-118\n(Salt)", "TP", "FP", "FP\nRate"],
        [
            ["1", "Isolation Forest", "1",
             "MISSED", "MISSED", "MISSED", "MISSED", "0", "12", "4.9%"],
            ["2", "One-Class SVM", "1",
             "DETECTED", "MISSED", "MISSED", "DETECTED", "2", "26", "10.6%"],
            ["3", "LOF", "1",
             "MISSED", "MISSED", "MISSED", "DETECTED", "1", "11", "4.5%"],
            ["4", "Z-Score (|z|>3)", "1",
             "MISSED", "MISSED", "DETECTED", "DETECTED", "2", "25", "10.2%"],
            ["5", "Temporal Z-Score", "1",
             "DETECTED", "DETECTED", "DETECTED", "DETECTED", "4", "232", "94.3%"],
            ["6", "Feature CUSUM\nTop10%", "1",
             "DETECTED", "MISSED", "DETECTED", "DETECTED", "3", "22", "8.9%"],
            ["7", "ACECARD Direction", "2",
             "MISSED", "DETECTED", "DETECTED", "MISSED", "2", "37", "15.0%"],
            ["8", "IForest +\nACECARD", "1+2",
             "MISSED", "DETECTED", "DETECTED", "MISSED", "2", "49", "19.9%"],
            ["9", "T3 Velocity/\nAccel", "3",
             "DETECTED", "MISSED", "DETECTED", "MISSED", "2", "23", "9.3%"],
            ["10", "T3 Regime Shift", "3",
             "MISSED", "MISSED", "DETECTED", "MISSED", "1", "24", "9.8%"],
            ["11", "T3 Zone\nDivergence", "3",
             "DETECTED", "MISSED", "MISSED", "MISSED", "1", "24", "9.8%"],
            ["12", "T3 Relationship\nDrift", "3",
             "MISSED", "MISSED", "DETECTED", "MISSED", "1", "24", "9.8%"],
            ["13", "T3 Contextual", "3",
             "MISSED", "DETECTED", "DETECTED", "MISSED", "2", "33", "13.4%"],
            ["14", "T3 Embedding\nCUSUM", "3",
             "DETECTED", "MISSED", "DETECTED", "DETECTED", "3", "22", "8.9%"],
            ["15", "T3 Zone Threat\nDir", "3",
             "DETECTED", "DETECTED", "MISSED", "DETECTED", "3", "125", "50.8%"],
            ["16", "T3 Beh\nProgression", "3",
             "DETECTED", "DETECTED", "MISSED", "DETECTED", "3", "22", "8.9%"],
            ["17", "T3 Combined", "3",
             "DETECTED", "DETECTED", "DETECTED", "DETECTED", "4", "37", "15.0%"],
        ],
        col_widths=[0.3, 1.1, 0.4, 0.65, 0.6, 0.6, 0.6, 0.35, 0.4, 0.5],
    )

    add_section_heading(doc, "3.2. Tier Comparison Summary", level=2)

    add_body_text(doc, (
        "The following table compares the best-performing method from each tier, "
        "highlighting the gap that only Tier 3 Combined closes."
    ))

    create_table(doc,
        ["Tier", "Best Method", "TP", "FP Rate", "Key Gap"],
        [
            ["Tier 1 Best", "Feature CUSUM", "3/4", "8.9%",
             "Misses Slow APT (USR-234)"],
            ["Tier 2 Best", "ACECARD Direction", "2/4", "15.0%",
             "Misses Insider + Salt Typhoon"],
            ["Tier 1+2\nCombined", "IForest + ACECARD", "2/4", "19.9%",
             "Still misses Insider + Salt"],
            ["Tier 3\nCombined", "T3 Combined", "4/4", "15.0%",
             "None — all attacks detected"],
        ],
        col_widths=[1.0, 1.4, 0.5, 0.7, 2.5],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  4. PER-ATTACK FINDINGS
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "4. Per-Attack Findings", level=1)

    # ── 4.1 USR-156 Insider ──────────────────────────────────────
    add_section_heading(doc, "4.1. USR-156 — Insider Threat", level=2)

    add_body_text(doc, (
        "The insider threat is the hardest attack to detect because the user has valid "
        "credentials and authorized access. They don't break in — they're already inside. "
        "The behavioral change is purely directional: the insider accesses different types "
        "of data over time, but the volume and frequency of access remain normal."
    ))

    create_table(doc,
        ["Method", "Result", "Why"],
        [
            ["Isolation Forest", "MISSED",
             "Volume unchanged — insider accesses same\nnumber of files"],
            ["LOF", "MISSED",
             "Feature-space profile looks normal\n(magnitude unchanged)"],
            ["ACECARD Direction", "MISSED",
             "Single composite dilutes data_behavior\nsignal across 5 zones"],
            ["T3 Zone Divergence", "DETECTED",
             "data_behavior zone drifts 0.259 while\nidentity stable at 0.000"],
            ["T3 Embedding CUSUM", "DETECTED",
             "Cumulative embedding drift accumulates\nover 60-day campaign"],
            ["T3 Combined", "DETECTED",
             "4/6 core methods fire, composite\nscore = 1.000 (highest of all users)"],
        ],
        col_widths=[1.5, 0.9, 3.5],
    )

    add_body_text(doc, "Tier 3 diagnostic:", bold=True)
    add_code_block(doc,
        "identity STABLE (0.000), data_behavior DRIFTING (0.259)\n"
        "toward insider_threat_fast (confidence=0.353)"
    )

    # ── 4.2 USR-234 APT ─────────────────────────────────────────
    add_section_heading(doc, "4.2. USR-234 — Slow APT", level=2)

    add_body_text(doc, (
        "The slow APT campaign maintains covert command-and-control communication for "
        "65 days. The C2 beacon is low-volume and periodic, designed to blend in with "
        "normal HTTPS traffic. Traditional volume-based methods cannot distinguish the "
        "C2 traffic from legitimate network activity."
    ))

    create_table(doc,
        ["Method", "Result", "Why"],
        [
            ["Isolation Forest", "MISSED",
             "C2 beacon volume too low to create\nfeature-space outlier"],
            ["Feature CUSUM", "MISSED",
             "Scalar features don't capture network\ndirectionality"],
            ["ACECARD Direction", "DETECTED",
             "Embedding drift toward threat concepts\nover 65 days"],
            ["T3 Contextual", "DETECTED",
             "apt_hunt context (network=0.40) amplifies\nC2 signal"],
            ["T3 Combined", "DETECTED",
             "2/6 core methods fire, composite\nscore = 0.842"],
        ],
        col_widths=[1.5, 0.9, 3.5],
    )

    add_body_text(doc, "Tier 3 diagnostic:", bold=True)
    add_code_block(doc,
        "identity STABLE (0.000), network_footprint DRIFTING (0.137)\n"
        "toward insider_threat_slow (confidence=0.116)"
    )

    # ── 4.3 USR-042 Volt Typhoon ────────────────────────────────
    add_section_heading(doc, "4.3. USR-042 — Volt Typhoon LOTL", level=2)

    add_body_text(doc, (
        "The Volt Typhoon campaign uses only legitimate administration tools already "
        "present on the system — no malware is deployed. This 'living-off-the-land' "
        "technique avoids all signature-based detection. The behavioral change manifests "
        "as dual-zone drift across data access and network patterns simultaneously."
    ))

    create_table(doc,
        ["Method", "Result", "Why"],
        [
            ["Z-Score", "DETECTED",
             "LOTL tools create measurable endpoint anomaly"],
            ["Feature CUSUM", "DETECTED",
             "Scalar feature trajectory accumulates drift"],
            ["ACECARD Direction", "DETECTED",
             "Embedding captures behavioral phase change"],
            ["T3 Velocity/Accel", "DETECTED",
             "Acceleration = 0.004 (positive, speeding up)"],
            ["T3 Contextual", "DETECTED",
             "apt_hunt context consistency = 0.833\n(highest across contexts)"],
            ["T3 Combined", "DETECTED",
             "4/6 core methods fire, composite\nscore = 0.994"],
        ],
        col_widths=[1.5, 0.9, 3.5],
    )

    add_body_text(doc, "Tier 3 diagnostic:", bold=True)
    add_code_block(doc,
        "data_behavior DRIFTING (0.175) + network_footprint DRIFTING (0.145)\n"
        "dual-zone drift pattern"
    )

    # ── 4.4 USR-118 Salt Typhoon ────────────────────────────────
    add_section_heading(doc, "4.4. USR-118 — Salt Typhoon Telecom", level=2)

    add_body_text(doc, (
        "The Salt Typhoon campaign targets telecommunications infrastructure through "
        "CDR database access, DNS tunneling, and cross-segment network flows. The "
        "campaign is persistent (60 days) and creates cumulative drift across multiple "
        "behavioral zones that accumulation-based methods are designed to catch."
    ))

    create_table(doc,
        ["Method", "Result", "Why"],
        [
            ["LOF", "DETECTED",
             "Infrastructure access creates feature-space\noutlier"],
            ["Z-Score", "DETECTED",
             "CDR access spikes file_restricted_ratio\nbeyond 3σ"],
            ["Feature CUSUM", "DETECTED",
             "100-day campaign creates persistent\nscalar drift"],
            ["ACECARD Direction", "MISSED",
             "Single composite dilutes multi-zone signal"],
            ["T3 Embedding CUSUM", "DETECTED",
             "Cumulative embedding drift over 60 days"],
            ["T3 Beh Progression", "DETECTED",
             "Kendall τ shows monotonic increase in\nthreat alignment"],
            ["T3 Combined", "DETECTED",
             "2/6 core methods fire, composite\nscore = 0.926"],
        ],
        col_widths=[1.5, 0.9, 3.5],
    )

    add_body_text(doc, "Tier 3 diagnostic:", bold=True)
    add_code_block(doc,
        "data_behavior DRIFTING (0.154) toward privilege_escalation\n"
        "+ network_footprint DRIFTING (0.179) toward insider_threat_slow"
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  5. KEY FINDINGS
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "5. Key Findings", level=1)

    # ── 5.1 Finding 1 ────────────────────────────────────────────
    add_section_heading(doc, "5.1. Finding 1: No Single Method Catches All Attacks", level=2)

    add_body_text(doc, (
        "Across 17 detection methods, only 2 achieve 4/4 detection: Temporal Z-Score "
        "(94.3% FP — operationally useless) and Tier 3 Combined (15.0% FP — viable). "
        "Every other method misses at least one attack."
    ))

    create_table(doc,
        ["Method", "Attacks\nDetected", "What It Misses", "Why It Misses"],
        [
            ["Isolation Forest", "0/4",
             "All 4 attacks",
             "Cannot detect behavioral direction\nchange at this population scale"],
            ["LOF", "1/4",
             "Insider, APT, Volt Typhoon",
             "Measures magnitude, not direction"],
            ["Feature CUSUM", "3/4",
             "Slow APT (USR-234)",
             "Scalar features don't capture\nnetwork directionality"],
            ["ACECARD Direction", "2/4",
             "Insider + Salt Typhoon",
             "Single composite dilutes\nzone-specific signals"],
            ["T3 Embedding CUSUM", "3/4",
             "Slow APT (USR-234)",
             "APT drift too gradual for\ncumulative threshold"],
            ["T3 Combined", "4/4",
             "None",
             "Ensemble of 6 methods covers\nall detection axes"],
        ],
        col_widths=[1.3, 0.7, 1.8, 2.3],
    )

    # ── 5.2 Finding 2 ────────────────────────────────────────────
    add_section_heading(doc, "5.2. Finding 2: Different Threats Require Different Detection Approaches", level=2)

    add_body_text(doc, (
        "Each attack type is best detected by a different analytical approach. There is "
        "no single detection axis that works for all threats."
    ))

    create_table(doc,
        ["Threat Type", "Detection Axis", "Best Method"],
        [
            ["Insider Threat",
             "Behavioral DIRECTION\n(what they access)",
             "Zone Divergence"],
            ["Slow APT",
             "Network PATTERN\n(periodic C2 beacon)",
             "Contextual (apt_hunt context)"],
            ["LOTL (Volt Typhoon)",
             "Behavioral MAGNITUDE\n(tool usage spike)",
             "Traditional (Z-Score,\nFeature CUSUM)"],
            ["Telecom Pivot (Salt)",
             "CUMULATIVE drift\n(persistent 100-day campaign)",
             "Embedding CUSUM,\nBehavioral Progression"],
        ],
        col_widths=[1.5, 2.0, 2.0],
    )

    # ── 5.3 Finding 3 ────────────────────────────────────────────
    add_section_heading(doc, "5.3. Finding 3: Zone Decomposition is the Critical Innovation", level=2)

    add_body_text(doc, (
        "Tier 2's single-composite architecture dilutes zone-specific signals. The insider's "
        "data_behavior drift (0.259) is averaged with 4 stable zones, producing a composite "
        "drift of only 0.009 — invisible. Tier 3's zone decomposition preserves the "
        "zone-specific signal, allowing detection of threats that change behavior in one "
        "dimension while remaining stable in others."
    ))

    add_body_text(doc, (
        "This is the fundamental architectural difference between Tier 2 and Tier 3: "
        "Tier 2 asks 'is this user behaving differently overall?' while Tier 3 asks "
        "'WHICH aspect of this user's behavior is changing?' The second question is far "
        "more powerful because it reveals not only that a threat exists, but what kind "
        "of threat it is."
    ))

    # ── 5.4 Finding 4 ────────────────────────────────────────────
    add_section_heading(doc, "5.4. Finding 4: Interpretable Diagnostics Enable Analyst Action", level=2)

    add_body_text(doc, (
        "Unlike traditional anomaly detection methods that output only a numeric score, "
        "Tier 3 produces human-readable diagnostics that tell security analysts exactly "
        "where to look and what to look for."
    ))

    create_table(doc,
        ["Attack", "Tier 3 Diagnostic"],
        [
            ["USR-156\nInsider",
             "identity STABLE, data_behavior DRIFTING\n"
             "toward insider_threat_fast (conf=0.353)"],
            ["USR-234\nSlow APT",
             "identity STABLE, network_footprint DRIFTING\n"
             "toward insider_threat_slow (conf=0.116)"],
            ["USR-042\nVolt Typhoon",
             "data_behavior DRIFTING (0.175)\n"
             "+ network_footprint DRIFTING (0.145)"],
            ["USR-118\nSalt Typhoon",
             "data_behavior DRIFTING toward privilege_escalation\n"
             "+ network_footprint DRIFTING"],
        ],
        col_widths=[1.2, 5.0],
    )

    add_body_text(doc, (
        "These diagnostics tell analysts exactly WHERE to look and WHAT to look for — "
        "not just 'anomaly score exceeded threshold.' For the insider threat, the diagnostic "
        "immediately points the analyst to the data_behavior zone, indicating that the user's "
        "file access patterns have changed while their identity profile has not. This "
        "dramatically reduces investigation time from hours of manual log review to a "
        "targeted examination of the flagged behavioral zone."
    ), bold=True)

    # ── 5.5 Finding 5 ────────────────────────────────────────────
    add_section_heading(doc, "5.5. Finding 5: Tier 3 Combined Achieves Best Coverage at Viable FP", level=2)

    add_body_text(doc, (
        "Tier 3 Combined detects all 4 attacks at 15.0% FP. While Feature CUSUM (Tier 1) "
        "achieves 3/4 at 8.9% FP, it misses the APT. While ACECARD Direction (Tier 2) "
        "catches APT and Volt Typhoon, it misses the Insider and Salt Typhoon. Only Tier 3 "
        "covers all threat types."
    ))

    add_body_text(doc, (
        "The 15.0% false positive rate means that for every 100 benign users, approximately "
        "15 will be flagged for analyst review alongside the true threats. In an enterprise "
        "of 250 users, this translates to roughly 37 false alerts per evaluation cycle — "
        "a manageable workload for a security operations center, especially given the "
        "interpretable diagnostics that enable rapid triage and dismissal of false positives."
    ))

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  6. RECOMMENDATIONS
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "6. Recommendations", level=1)

    add_section_heading(doc, "6.1. Deployment Strategy", level=2)

    add_body_text(doc, (
        "Based on the findings, a production deployment should implement detection "
        "in three complementary layers, each targeting a different class of threat:"
    ))

    add_bullet(doc, (
        "Traditional magnitude detection (Feature CUSUM, LOF) — always on, catches "
        "volume-based attacks including LOTL tool activation and infrastructure pivot "
        "campaigns. Fast to compute, low false positive overhead, provides a strong "
        "baseline detection capability."
    ), bold_prefix="Layer 1: ")

    add_bullet(doc, (
        "Behavioral direction detection (Zone Divergence, Contextual) — catches "
        "direction-based attacks where the user's behavior changes qualitatively but "
        "not quantitatively. Essential for insider threats and slow APTs that operate "
        "within normal volume parameters."
    ), bold_prefix="Layer 2: ")

    add_bullet(doc, (
        "Temporal accumulation (Embedding CUSUM, Behavioral Progression) — catches "
        "long-duration campaigns through cumulative drift detection and monotonic trend "
        "analysis. Essential for persistent threats that operate over 60–100+ day "
        "timelines where weekly changes are individually small but collectively significant."
    ), bold_prefix="Layer 3: ")

    add_body_text(doc, (
        "This layered architecture ensures that no single class of threat can evade "
        "detection by exploiting the blind spots of any one approach. The layers are "
        "complementary, not redundant — each catches threats the others miss."
    ), bold=True)

    add_section_heading(doc, "6.2. Next Steps", level=2)

    add_bullet(doc, (
        "Validate on real enterprise data through a 4-week pilot demonstration at a "
        "partner site, using anonymized production telemetry to confirm detection "
        "performance translates from synthetic to real-world environments."
    ), bold_prefix="Real-World Validation: ")

    add_bullet(doc, (
        "Tune detection thresholds for customer-specific false positive tolerance. "
        "Different organizations have different SOC capacity — a 5-analyst team "
        "can handle more alerts than a 1-analyst team. Thresholds should be calibrated "
        "to match operational capacity."
    ), bold_prefix="Threshold Tuning: ")

    add_bullet(doc, (
        "Scale to 500+ users with batch embedding optimization. The current evaluation "
        "covered 250 users; production deployments may require 500–5,000+ users with "
        "optimized embedding batch processing and incremental updates."
    ), bold_prefix="Scale Testing: ")

    add_bullet(doc, (
        "Integrate with existing SIEM platforms for automated alert routing, enrichment, "
        "and case management. Tier 3 diagnostics should flow into analyst workflows "
        "alongside traditional SIEM alerts for unified triage."
    ), bold_prefix="SIEM Integration: ")

    add_page_break(doc)

    # ── End matter ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Results and Findings —")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "22nd Century Technologies, Inc.\n"
        "ACECARD Program — Results and Findings v1.0\n"
        "May 2025"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE

    # ── Save ────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Results and Findings document created successfully")
    print(f"Output: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
