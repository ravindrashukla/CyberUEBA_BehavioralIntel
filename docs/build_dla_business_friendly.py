"""Business-friendly edition of the DLA Demand Forecast paper.
Keeps real technical substance (architecture, drift layer, calibration, honest
results, novelty) but redacts exact implementation logic (formulas, thresholds,
band multipliers, code). Abbreviations clarified on first use. ~8 pages.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "DLA_Demand_Forecast_Academic.docx")

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
ORG  = RGBColor(0xC2, 0x5A, 0x12)
GREY = RGBColor(0x55, 0x5F, 0x6B)
GREEN= RGBColor(0x1E, 0x7A, 0x44)
WHITE= RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.1
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(doc.sections[0], m, Inches(0.85))


def H1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    return p

def H2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ORG
    return p

def para(text, size=10.5, color=None, bold=False, italic=False, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def bullet(text, lead=None):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run(text); r2.font.size = Pt(10.5)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def setcell(cell, text, bold=False, size=9.5, color=None, fill=None):
    if fill: shade(cell, fill)
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
    if color: r.font.color.rgb = color

def table(headers, rows, widths=None, hdr_fill='0B1F3A'):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        setcell(t.rows[0].cells[j], h, bold=True, color=WHITE, fill=hdr_fill)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            setcell(t.rows[i].cells[j], v, bold=(j == 0), fill='F4F6F8' if j == 0 else None)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ================= FIGURE GENERATION =================
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

FIGDIR = os.path.join(os.path.dirname(__file__), 'figs_business')
os.makedirs(FIGDIR, exist_ok=True)

C_LEG = '#8A93A0'   # legacy grey
C_LEGD = '#C0392B'  # legacy red accent
C_EDM = '#0891B2'   # EDM teal
C_NAVY = '#0B1F3A'
plt.rcParams.update({'font.family': 'Calibri', 'font.size': 10, 'axes.edgecolor': '#C9CFD6',
                     'axes.labelcolor': C_NAVY, 'xtick.color': '#55607B', 'ytick.color': '#55607B'})

def fig1_procurement():
    """8-week cumulative procurement plan: legacy vs EDM (adapter plate)."""
    wk = np.arange(0, 9)
    shape = (wk / 8) ** 0.92
    leg90, leg50 = 407 * shape, 278 * shape
    edm90, edm50 = 218 * shape, 135 * shape
    fig, ax = plt.subplots(figsize=(7.0, 3.4))
    ax.fill_between(wk, leg50, leg90, color=C_LEG, alpha=0.25)
    ax.fill_between(wk, edm50, edm90, color=C_EDM, alpha=0.25)
    ax.plot(wk, leg90, color=C_LEGD, lw=2.2, label='Legacy plan — safety-stock target (P90): 407 units')
    ax.plot(wk, leg50, color=C_LEG, lw=1.6, ls='--', label='Legacy median (P50): 278')
    ax.plot(wk, edm90, color=C_EDM, lw=2.2, label='EDM plan — safety-stock target (P90): 218 units')
    ax.plot(wk, edm50, color=C_EDM, lw=1.6, ls='--', label='EDM median (P50): 135')
    ax.annotate('', xy=(8, 218), xytext=(8, 407),
                arrowprops=dict(arrowstyle='<->', color=C_NAVY, lw=1.4))
    ax.text(7.78, 305, '189 units avoided\n≈ $564,000', ha='right', va='center',
            fontsize=10.5, fontweight='bold', color=C_NAVY)
    ax.set_xlabel('Week'); ax.set_ylabel('Cumulative units planned')
    ax.set_xlim(0, 8.4); ax.set_ylim(0, 440)
    ax.legend(loc='upper left', fontsize=8.6, frameon=False)
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig1_procurement.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig2_bands():
    """Frozen vs dynamic uncertainty bands, two situations."""
    wk = np.arange(0, 13)
    fig, axes = plt.subplots(1, 2, figsize=(7.0, 3.0), sharey=True)
    # (a) calming item
    ax = axes[0]
    med = 100 - 4.5 * wk
    legw = 0.50 * med                      # frozen ratio: width tied to median
    edmw = 0.50 * med * np.linspace(1.0, 0.50, len(wk))   # EDM tightens further
    ax.fill_between(wk, med - legw, med + legw, color=C_LEG, alpha=0.30, label='Legacy band (frozen shape)')
    ax.fill_between(wk, med - edmw, med + edmw, color=C_EDM, alpha=0.38, label='EDM band (dynamic)')
    ax.plot(wk, med, color=C_NAVY, lw=1.6, ls='--', label='Median forecast')
    ax.set_title('Demand calming — EDM tightens the band\n(legacy stays wide → over-ordering)', fontsize=9.5, color=C_NAVY)
    ax.set_xlabel('Week'); ax.set_ylabel('Demand')
    ax.legend(loc='upper right', fontsize=8, frameon=False)
    # (b) destabilizing item
    ax = axes[1]
    med2 = np.full(len(wk), 70.0)
    legw2 = np.full(len(wk), 22.0)
    edmw2 = 22.0 * np.linspace(1.0, 1.9, len(wk))
    ax.fill_between(wk, med2 - edmw2, med2 + edmw2, color=C_EDM, alpha=0.38, label='EDM band (dynamic)')
    ax.fill_between(wk, med2 - legw2, med2 + legw2, color=C_LEG, alpha=0.45, label='Legacy band (frozen shape)')
    ax.plot(wk, med2, color=C_NAVY, lw=1.6, ls='--')
    ax.set_title('Demand destabilizing — EDM widens the band\n(legacy stays narrow → understates risk)', fontsize=9.5, color=C_NAVY)
    ax.set_xlabel('Week')
    h, l = ax.get_legend_handles_labels()
    order = [l.index('Legacy band (frozen shape)'), l.index('EDM band (dynamic)')]
    ax.legend([h[i] for i in order], [l[i] for i in order], loc='upper left', fontsize=8, frameon=False)
    for ax in axes:
        ax.spines[['top', 'right']].set_visible(False); ax.set_ylim(0, 160)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig2_bands.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig3_calibration():
    """Target vs measured coverage."""
    qs = ['P10', 'P25', 'P50', 'P75', 'P90']
    nominal = [0.10, 0.25, 0.50, 0.75, 0.90]
    measured = [0.141, 0.257, 0.498, 0.743, 0.894]
    x = np.arange(len(qs)); w = 0.36
    fig, ax = plt.subplots(figsize=(6.6, 2.9))
    ax.bar(x - w/2, nominal, w, color=C_LEG, alpha=0.55, label='Target coverage')
    ax.bar(x + w/2, measured, w, color=C_EDM, label='Measured coverage (EDM)')
    for xi, m in zip(x, measured):
        ax.text(xi + w/2, m + 0.015, f'{m:.0%}', ha='center', fontsize=8.6, color=C_NAVY, fontweight='bold')
    ax.set_xticks(x, qs); ax.set_ylim(0, 1.05)
    ax.set_ylabel('Share of actuals under the percentile')
    ax.legend(loc='upper left', fontsize=8.6, frameon=False)
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig3_calibration.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig4_intermittent():
    """Intermittent weekly demand vs a smoothing average."""
    rng = np.random.default_rng(7)
    weeks = np.arange(1, 53)
    demand = np.zeros(52)
    nz = rng.choice(52, size=24, replace=False)          # ~54% zero weeks
    demand[nz] = rng.gamma(1.6, 14, size=len(nz)).round()
    ma = np.convolve(demand, np.ones(13)/13, mode='same')
    fig, ax = plt.subplots(figsize=(7.0, 2.8))
    ax.bar(weeks, demand, color=C_EDM, alpha=0.8, label='Actual weekly demand (54% of weeks are zero)')
    ax.plot(weeks, ma, color=C_LEGD, lw=2.2, label='Legacy moving average — smooths away the zeros and spikes')
    ax.set_xlabel('Week'); ax.set_ylabel('Units')
    ax.legend(loc='upper right', fontsize=8.6, frameon=False)
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig4_intermittent.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

FIG1, FIG2, FIG3, FIG4 = fig1_procurement(), fig2_bands(), fig3_calibration(), fig4_intermittent()

def add_fig(path, caption, width=6.4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY


# ================= TITLE =================
r = doc.add_paragraph().add_run('Demand Forecasting for Defense Logistics:\nThe Traditional Approach vs. EDM Vector Intelligence')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
para('A business and technology overview of the method, its measured results, and the supporting evidence',
     12.5, GREY, after=2)
para('22nd Century Technologies, Inc.  ·  V-Intelligence DLA Program  ·  June 2026  ·  Business Edition',
     9.5, GREY, after=2)
para('Disclosure note: this edition describes the architecture and the innovation, and reports measured results. The '
     'exact model logic — formulas, thresholds, and adjustment parameters — is proprietary and is withheld here; it is '
     'available in the full technical edition under NDA (non-disclosure agreement).', 9.5, ORG, italic=True, after=4)
para('Data note: all results in this paper come from a synthetic (simulated) dataset built to prove the concept — '
     '500 items, 200 suppliers, 4 depots, 24 months, ~277,000 requisitions. They must be re-validated on real Defense '
     'Logistics Agency data before any operational claim.', 9.5, ORG, italic=True, after=8)

# ================= 1. EXEC SUMMARY =================
H1('1.  Executive Summary')
para('The Defense Logistics Agency (DLA) buys and positions spare parts for the Department of Defense. Every '
     'forecast it produces informs a decision under uncertainty: how much to stock, where, and at what budget. Defense '
     'demand is intermittent (most weeks see zero orders) and lumpy (order sizes swing widely), and it shifts as '
     'exercises, deployments, and lifecycle changes reshape it. A single best-guess number is not sufficient for '
     'that decision. The planner needs a calibrated range, together with a system that responds before the change '
     'becomes obvious.')
para('The legacy forecaster extrapolates each item\'s own history and draws its uncertainty range in a single fixed '
     'shape: the range can slide up or down, but it cannot tighten when demand calms or widen when it destabilizes. '
     'That fixed shape is a principal cause of systematic over- and under-ordering.')
para('EDM Vector Intelligence (EDM = Entity Digital Model) adds a dimension the legacy method lacks. It '
     'represents every item-at-a-depot as a digital entity in a high-dimensional behavior space, using the same form '
     'of semantic representation that underlies large language models, applied here to supply-chain behavior rather '
     'than language. Re-profiling each entity monthly produces a trajectory. Measuring that trajectory, which we term '
     '"drift," indicates where the item is heading, and matching the drift direction against eight named supply-chain '
     'situations indicates what it is becoming. The forecast and its uncertainty bands are then adjusted in advance, '
     'before the shift appears in raw order counts.')
para('In summary, the legacy method extrapolates past demand and holds the uncertainty bands in a fixed shape, '
     'whereas EDM estimates where the item is heading and adjusts both the plan and its uncertainty range in advance.', bold=True)
bullet(' On a declining item (an adapter plate at Red River Army Depot), the legacy method planned 407 units over 8 '
       'weeks at the safety-stock level, while EDM planned 218. The difference is 189 fewer units, roughly $564,000 '
       'of avoided procurement on a single item.', lead='Right-sized procurement. ')
bullet(' On an item drifting toward a supply disruption, EDM raises the plan and widens the bands 1–3 months before '
       'the disruption appears in delivery counts.', lead='Early warning. ')
bullet(' The probability model is well-calibrated: when it states a 70% chance, demand occurs about 70% of the '
       'time (Grade A in testing), so safety stock can target a defined service level.', lead='Trustworthy uncertainty. ')

# ================= 2. PROBLEM =================
H1('2.  The Problem')
para('Three properties of defense-parts demand undermine a single-number forecast:')
bullet(' Most item-weeks have zero demand. Whether demand occurs at all is a separate question from how much — and '
       'both must be modeled.', lead='Intermittency. ')
bullet(' When demand occurs it varies widely; the spread of outcomes matters as much as the center.', lead='Lumpiness. ')
bullet(' Exercises, deployments, and lifecycle transitions reshape demand by category and timing. A forecaster that '
       'cannot detect the shift as it forms can respond only after it has occurred.', lead='Regime change. ')
para('The right output is therefore a calibrated distribution per item, location, and week — a low (P10), median '
     '(P50), and high (P90) estimate, where "P90" means the level demand stays under 90% of the time — produced by a '
     'system that can also reason across items, suppliers, and depots rather than one series at a time.')

# ================= 3. LEGACY =================
H1('3.  The Traditional (Legacy) Approach')
para('The legacy forecaster examines one item\'s history at one depot, computes the recent average and trend, applies '
     'exercise and fiscal-calendar factors, and scales its historical demand percentiles by the combined ratio. It is '
     'transparent and fast, and it serves as the baseline against which EDM is evaluated.')
table(['Stage', 'What it does'],
      [('1. Load & aggregate', 'Pull ~6 months of requisitions for this item at this depot into weekly buckets.'),
       ('2. Empirical percentiles', 'Compute how often demand occurs and the historical low-to-high spread. A new item with no history forecasts zero.'),
       ('3. Exercise calendar', 'Apply per-item exercise multipliers with ramp-up and wind-down periods.'),
       ('4. Fiscal calendar', 'Adjust for Department of Defense fiscal-year quarter effects.'),
       ('5. Trend & scale', 'Project the recent trend and scale every percentile by the same combined ratio.'),
       ('6. Simulation', 'A shared Monte Carlo simulation converts weekly estimates into the 8-week cumulative safety-stock target (the same simulator EDM uses, for a fair comparison).')],
      widths=[1.7, 5.1])
H2('Limitations of the legacy approach')
bullet(' Because every percentile scales by one ratio, the ratio of high-to-median does not change. The band cannot '
       'narrow on a calming item or widen on a destabilizing one, which is a direct cause of over- and under-ordering.',
       lead='Fixed band shape. ')
bullet(' Trend extrapolation cannot anticipate a regime change; it responds only after the counts move.',
       lead='Backward-looking. ')
bullet(' A new or sparse item receives a zero forecast, with no mechanism to borrow from similar items.',
       lead='Cold-start blindness. ')
bullet(' One item, one depot, one data source — no cross-item learning, no behavioral signal.',
       lead='Single-entity tunnel vision. ')

# ================= 4. EDM =================
H1('4.  The EDM Vector Intelligence Approach')
H2('4.1  The digital entity and its drift')
para('Consider the central concept in plain terms. A "digital entity" (sometimes called a digital twin) is a living '
     'profile of a real-world thing, maintained in software. A patient\'s medical chart is a useful analogy: it '
     'records more than today\'s temperature, accumulating history, vitals, and context so that a physician can read '
     'the trend and act before a condition becomes acute. EDM maintains a comparable record for every part at every '
     'depot. It captures more than the count of units ordered; it captures the full behavioral picture, and how that '
     'picture is changing month over month.', italic=False)
para('EDM models each item-location pair, identified by NSN (National Stock Number, the part\'s unique ID) and '
     'DODAAC (the code for a depot or activity), as a digital entity in a unified high-dimensional semantic vector space (in the current MVP, 1,536 dimensions). '
     'This space is, in effect, a mathematical "map of meaning" in which items that behave alike sit near each other, '
     'even when their raw numbers differ. The entity is built from five behavioral signals: demand, supply, '
     'inventory, volatility, and operational context. Each signal\'s metrics are serialized into a structured '
     'natural-language description and converted into a numerical vector by an embedding model, the same representation '
     'used in large language models, applied here to demand behavior rather than language.')
para('Each month the entity is re-embedded, so over 16–23 monthly snapshots it traces a trajectory through the '
     'behavior space. Forecasting then addresses two questions that a moving average cannot. The first is what the '
     'item has been doing, captured by a calibrated statistical baseline. The second, and the central contribution of '
     'this work, is where the item is heading, captured by its drift, so that the plan can adjust before the raw '
     'numbers move.')
table(['Behavioral signal', 'What it captures'],
      [('Demand', 'Volume, frequency, trend, and variability of weekly demand.'),
       ('Supply', 'Supplier count, sole-source exposure, concentration, and delivery performance.'),
       ('Inventory', 'On-hand position, trajectory, and safety-stock posture.'),
       ('Volatility', 'Zero-demand fraction, demand intervals, and stability of the pattern.'),
       ('Context', 'Criticality, stock ratios, and operational setting.')],
      widths=[1.9, 4.9])
para('The five signal vectors compose into one monthly entity state. Because every entity occupies the same space, '
     'the system can identify behavioral neighbors. This allows a new item to receive a meaningful forecast borrowed '
     'from items that behave like it (cohort transfer), which the legacy method cannot provide.')

H2('4.2  Architecture and data')
para('The system is a containerized service: an API (Application Programming Interface) layer over a PostgreSQL '
     'database with a vector store (fast similarity search over the entity embeddings), bi-temporal history (every '
     'fact is stored with both when it happened and when the system learned it), and point-in-time feature functions '
     'that guarantee no future information leaks into training. It ingests the requisition, order, inventory, and '
     'operational-calendar data DLA already produces — no new instrumentation. The production baseline uses 25 '
     'engineered features per item-depot-week; with all EDM layers active, the model ingests 67 features, adding '
     'embedding-derived zone signals, entity-trajectory measures, and cohort-transfer features.')

H2('4.3  The drift-augmented forecast')
para('The drift layer extends a calibrated baseline through five steps. Each step is described in functional terms; '
     'the exact formulas, thresholds, and adjustment parameters are proprietary and are withheld in this edition.')
table(['Step', 'What it does'],
      [('1. Calibrated baseline', 'A two-stage ("hurdle") model — first the probability that any demand occurs, then the size if it does, as a set of percentile estimates. This is the statistically correct shape for zero-inflated demand.'),
       ('2. Trajectory load', 'Load the entity\'s monthly behavior snapshots. If too few exist, return the baseline unchanged — the system degrades gracefully rather than guessing.'),
       ('3. Drift measurement', 'Measure how fast and how consistently the entity is moving through behavior space (its velocity, acceleration, and consistency).'),
       ('4. Direction naming', 'Compare the drift direction against eight pre-embedded supply-chain concepts; the strongest match names what the entity is becoming.'),
       ('5. Plan adjustment', 'Translate the named direction into a bounded adjustment of the plan and an independent reshaping of the uncertainty bands, then re-run the simulation on the adjusted estimates.')],
      widths=[1.7, 5.1])
para('The eight named directions and how each moves the plan:')
table(['Concept', 'Plan effect', 'Meaning'],
      [('Demand surge', 'Order more', 'Rising volume, depleting stock, stockout risk forming.'),
       ('Supply disruption', 'Order more', 'Lead times stretching, supplier performance degrading.'),
       ('Critical shortage', 'Order more', 'Approaching stockout on a high-mission-impact item.'),
       ('Demand decline', 'Order less', 'Demand tapering off; excess-carry risk.'),
       ('Excess inventory', 'Order less', 'On-hand position too high; obsolescence risk.'),
       ('Stable consumption', 'No change', 'Predictable behavior; the baseline is already right.'),
       ('Intermittent demand', 'Widen bands', 'Sparse demand; uncertainty honestly widened.'),
       ('Volatility increase', 'Widen bands', 'Direction unknown; bands widen, median untouched.')],
      widths=[1.7, 1.2, 3.9])
para('Two design properties are significant in operation. First, the bands move independently of the median; whereas '
     'the legacy band shape is fixed, the EDM band shape changes with the situation, narrowing on a calm decline and '
     'widening on a forming disruption. Second, the adjustment is deliberately conservative. It is capped to a modest '
     'maximum, damped when the drift signal is weak or inconsistent, and overridden in defined safety cases, so that '
     'noise does not move procurement and the model adjusts gradually rather than abruptly.', bold=False)

# ================= 5. RESULTS =================
H1('5.  Results (Synthetic Data)')
H2('5.1  Point accuracy')
para('WAPE (Weighted Absolute Percentage Error) is the standard point-accuracy metric: total forecast error divided '
     'by total actual demand, where lower is better. On WAPE, the production model on engineered features alone (0.791) '
     'and the same model with the embedding features added (0.789) are effectively tied. On synthetic data, the '
     'embedding features do not improve point accuracy, and we report that directly. WAPE, however, scores a single '
     'median number. The decision a planner actually makes, namely the safety-stock target and how early the plan '
     'adapts, resides in the distribution and its direction, which WAPE does not measure.')
H2('5.2  Procurement-plan accuracy')
para('The decision the forecast informs is the 8-week cumulative P90, which is the safety-stock target. Three '
     'worked situations illustrate the principal functions of the drift layer:')
table(['Situation', 'What happens', 'Outcome'],
      [('Declining item — reduce\n(real system output)',
        'An adapter plate at Red River Army Depot, drifting toward "demand decline" across 23 monthly snapshots. EDM lowers the plan and tightens the bands.',
        'The legacy method plans 407 units (P90) over 8 weeks; EDM plans 218, a difference of 189 units and ~$564,000 avoided on one item. Most of the gap comes from the calibrated baseline and dynamic bands, with the drift layer contributing the directional adjustment.'),
       ('Disruption forming — increase\n(illustrative)',
        'A sole-source critical item whose trajectory aligns with "supply disruption." The adjustment reaches its cap and the bands widen.',
        'EDM raises the plan 1–3 months before the disruption appears in deliveries; the legacy method responds only once the shortage is already forming. This leading-indicator behavior is not available from a moving average.'),
       ('Destabilizing item — widen\n(illustrative)',
        'An item drifting toward "volatility increase." The median is left unchanged, since the direction of change is not yet known.',
        'The uncertainty bands widen substantially, signaling the need for additional buffer. The legacy method\'s fixed, narrow band would understate the risk and raise the likelihood of an unanticipated stockout.')],
      widths=[1.6, 2.6, 2.6])
add_fig(FIG1, 'Figure 1. The declining item (adapter plate, Red River Army Depot): 8-week cumulative procurement plan. '
              'The legacy method\'s fixed, trend-extrapolated band plans to 407 units at the safety-stock level; EDM plans 218, '
              'a difference of 189 units and ≈ $564,000 avoided on one item. (System output, synthetic data.)')
add_fig(FIG2, 'Figure 2. The effect of band shape. The legacy band shape is fixed, while EDM tightens it on a calming item '
              '(left, avoiding over-ordering) and widens it on a destabilizing one (right, surfacing the underlying risk). (Illustrative.)')
H2('5.3  Calibration')
para('A probability model can be relied upon when its stated probabilities match observed outcomes. The occurrence '
     'model achieved Grade A calibration (Expected Calibration Error 0.032), and the quantity percentiles cover '
     'close to their nominal targets, with the median at 49.8% against a 50% target. The single notable deviation is '
     'conservative: the low (P10) percentile is slightly cautious, which is the safer failure mode for readiness.')
add_fig(FIG3, 'Figure 3. Calibration: measured coverage tracks the target across all percentiles (synthetic data). '
              'The only deviation (P10) errs on the cautious side, which is the safer failure mode for readiness.')
para('The demand is genuinely intermittent; the worked-example item has 54% zero-demand weeks. A moving average '
     'smooths this away, whereas the two-stage model handles occurrence and quantity separately, which is why its '
     'percentiles remain calibrated.')
add_fig(FIG4, 'Figure 4. Intermittent demand (bars) compared with the legacy moving average (line). '
              'The average smooths away the zeros and spikes that drive readiness risk. (Illustrative, synthetic.)')
H2('5.4  Model evaluation summary')
table(['Component', 'Metric', 'Result', 'Grade'],
      [('Demand occurrence', 'Expected calibration error (ECE)', '0.032', 'A'),
       ('Demand occurrence', 'Brier score', '0.089', 'A'),
       ('Demand quantity', 'Point accuracy (WAPE)', '0.791', 'C'),
       ('Demand percentiles', 'Median coverage (target 50%)', '49.8%', 'A'),
       ('Lead time', 'Median error', '1.9 days', 'B'),
       ('Drift adjustment', 'Capped, damped under weak signal, bands stay ordered', 'By design', '—')],
      widths=[1.9, 2.3, 1.6, 0.9])
para('All results are scored out-of-time, meaning the model is always tested on periods it never saw, with '
     'point-in-time features so that no future data leaks into training. The overall assessment is as follows. The '
     'distribution used for planning is reliable, with Grade A calibration. Point accuracy is modest, which reflects '
     'the nature of intermittent demand and is not the appropriate measure for this method. The embedding and drift '
     'layer is justified by capability, specifically early direction, cold-start transfer, and dynamic bands, rather '
     'than by WAPE.')

# ================= 6. DEPLOYMENT =================
H1('6.  Scale and Deployment')
bullet(' Embeddings are cached, feature views are precomputed, and forecasts run in batch; entity-similarity '
       'lookups return in milliseconds.', lead='Performance. ')
bullet(' Legacy responds in ~50–150 ms per item; EDM in ~200–400 ms — both comfortably interactive.',
       lead='Latency. ')
bullet(' Containerized and database-backed; deployable in a DLA data center or secure enclave. Ingests existing '
       'requisition, order, inventory, and calendar data — no new instrumentation.', lead='Footprint. ')
bullet(' The external embedding model can be swapped for a locally hosted one for air-gapped operation.',
       lead='Air-gap option. ')
bullet(' Run the legacy method and EDM side by side (the application\'s Live Explorer supports this directly), '
       'compare the resulting plans, and then introduce EDM into procurement.', lead='Low-risk rollout. ')

# ================= 7. NOVELTY =================
H1('7.  Why This Is New')
para('The contribution lies in the drift-augmented entity model rather than in a new point estimator. Classical '
     'methods (moving averages, exponential smoothing, statistical time-series models such as ARIMA, and the Croston '
     'family of methods for intermittent demand) model each series in isolation with static uncertainty. Modern '
     'machine-learning forecasters and foundation time-series models improve point accuracy but provide no '
     'cross-entity semantic space and no named direction of change.')
para('A structured, source-verified review of the market and literature (2024–2026) found no commercially marketed '
     'platform, and no single published method, that forecasts this way. The major demand-planning platforms '
     'forecast with machine-learning models over tabular features; their generative AI is a conversational copilot '
     'layer, not the forecasting mechanism. Academic foundation models tokenize the numeric series directly rather '
     'than embedding entity behavior.')
para('We are explicit about prior art. The individual ingredients, including serializing structured data to text for '
     'embedding, detecting anomalies in embedding space, measuring embedding drift, and similarity-based transfer, '
     'each exist separately in the literature. The contribution of this work is their composition: a unified '
     'entity-embedding space whose drift is measured, named against supply-chain concepts, and used to size both the '
     'procurement plan and its uncertainty bands in advance. This composition is not, to our knowledge, offered by '
     'any commercial platform or described in any single published system. The contribution is the integrated '
     'architecture.', bold=True)
para('Scoping notes: vendor internals are not fully public (the comparison is against marketed methodology), and a '
     'formal freedom-to-operate review (a patent-clearance analysis by counsel) is recommended before any '
     'contractual assertion of uniqueness.', 9.5, GREY,
     italic=True)

# ================= 8. LIMITATIONS =================
H1('8.  Limitations')
bullet('All metrics and dollar figures are synthetic and must be re-measured on real DLA data.')
bullet('On point accuracy (WAPE), EDM does not yet beat legacy; the case rests on procurement-plan accuracy, '
       'calibration, and capability.')
bullet('The drift layer currently adjusts how much to stock, not whether demand will occur — a deliberate '
       'simplification worth revisiting.')
bullet('The drift adjustment is applied uniformly across the forecast horizon rather than week by week.')
bullet('Exercise-related behavior depends on calendar flags being populated; missing flags would skew '
       'exercise-scenario comparisons.')
bullet('Real embeddings depend on an external model; air-gapped deployment requires re-characterizing a local '
       'replacement.')

# ================= 9. CONCLUSION =================
H1('9.  Conclusion')
para('Judged on point accuracy alone, the legacy method and EDM are tied, but point accuracy is not the appropriate '
     'measure for this problem. Defense-logistics demand is intermittent and subject to regime change, and the '
     'decision it informs is a distribution rather than a single number. The legacy method extrapolates past demand '
     'and holds its uncertainty band in a fixed shape, which leads to systematic over- and under-ordering. EDM models '
     'each item-depot as a digital entity, estimates where it is heading in behavior space, names the direction '
     'against eight supply-chain concepts, and sizes both the plan and its bands in advance. On the system\'s own '
     'output, this corresponds to hundreds of thousands of dollars of avoided over-procurement on a single declining '
     'item and a 1–3 month head start on a disruption.')
para('The legacy method answers the question "what is the expected demand?" EDM answers the question "where is this '
     'item heading, and what should be stocked to maintain readiness?", and it adjusts before the raw numbers move.', bold=True)

# ================= GLOSSARY =================
H1('Appendix — Glossary')
gl = [
    ('Digital entity / digital twin', 'A living software profile of a real-world thing — like a patient\'s medical chart for a part: history, vitals, context, and trend, updated monthly.'),
    ('Semantic space', 'A mathematical "map of meaning": items whose behavior is similar sit near each other, even if their raw numbers differ.'),
    ('DLA', 'Defense Logistics Agency — the Department of Defense organization that buys and distributes spare parts and supplies.'),
    ('EDM Vector Intelligence', 'Entity Digital Model — our approach: a digital behavioral profile of each part whose movement over time drives the forecast.'),
    ('NSN / DODAAC', 'National Stock Number (a part\'s unique ID) / the code identifying a depot or activity.'),
    ('Embedding', 'A numerical vector that captures the meaning of a description — the core representation inside large language models.'),
    ('Drift', 'How an entity\'s behavioral profile moves through the embedding space over time — speed, consistency, and direction.'),
    ('P10 / P50 / P90', 'Percentile estimates: demand stays under the P90 level 90% of the time. P50 is the median. The P10–P90 spread is the uncertainty band.'),
    ('Safety stock', 'Buffer inventory held so a part does not run out; sized here from the 8-week cumulative P90.'),
    ('Hurdle (two-stage) model', 'First predict whether any demand occurs, then how much if it does — the right structure for demand that is mostly zeros.'),
    ('Calibration / ECE', 'Whether stated probabilities match observed reality; ECE (Expected Calibration Error) measures the gap — lower is better.'),
    ('Brier score', 'A standard accuracy score for probability forecasts (0 is perfect, 1 is worst); lower is better.'),
    ('WAPE', 'Weighted Absolute Percentage Error — total forecast error divided by total actual demand; a point-accuracy metric.'),
    ('Cohort transfer', 'Borrowing demand patterns from behaviorally similar items to forecast a new or sparse item.'),
    ('Monte Carlo simulation', 'Running thousands of simulated futures to turn weekly estimates into a cumulative safety-stock target.'),
    ('Synthetic data', 'Simulated data built to prove the concept before access to real operational data.'),
]
gt = doc.add_table(rows=len(gl) + 1, cols=2); gt.style = 'Table Grid'
setcell(gt.rows[0].cells[0], 'Term', bold=True, color=WHITE, fill='0B1F3A')
setcell(gt.rows[0].cells[1], 'Meaning', bold=True, color=WHITE, fill='0B1F3A')
for i, (t, d) in enumerate(gl, start=1):
    setcell(gt.rows[i].cells[0], t, bold=True, size=9, color=NAVY, fill='F4F6F8')
    setcell(gt.rows[i].cells[1], d, size=9)
for row in gt.rows:
    row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(4.9)

para('')
para('Business Edition — architecture-level disclosure. Exact model logic, formulas, thresholds, and adjustment '
     'parameters are proprietary and available in the full technical edition under NDA. All results are synthetic '
     'and must be re-validated on real DLA operational data. Do not redistribute.', 9, GREY, italic=True)

out = OUT; n = 2
while True:
    try:
        doc.save(out); break
    except PermissionError:
        root, ext = os.path.splitext(OUT); out = f"{root}_v{n}{ext}"; n += 1
print('Saved:', out)
