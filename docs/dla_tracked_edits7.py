"""Seventh pass: fix Section-4 figure issues in the current REV file.
  - Restore the 'Our Layered Analytical Architecture' figure (from previous_final) into
    the Layered Architecture subsection, as a TRACKED insertion with caption.
  - Renumber all figure captions and in-text 'Figure N' references to a single clean
    sequence (direct corrections, not tracked, since this is a mechanical numbering fix):
      1 From Metrics to Mission | 2 Architecture | 3 Layered (restored) | 4 Demand |
      5 Legacy | 6 Supplier Zones | 7 Network | 8 Evidence | 9 Pilot
"""
import os, copy
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PREV = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_previous_final.docx"
REV = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_final_REV_v2.docx"
AUTHOR = "ReviewPanel"; DATE = "2026-06-19T18:00:00Z"
IMG = os.path.join("docs", "_fig_layered.png")

def acc(p): return "".join(x.text or "" for x in p._p.iter(qn('w:t')))

# --- extract the Layered Analytical Architecture image from previous_final ---
prev = Document(PREV)
idxs = [i for i, p in enumerate(prev.paragraphs) if 'Layered Analytical Architecture' in acc(p)]
blip = None; width_in = 6.5
for i in idxs:
    for j in range(max(0, i - 2), min(len(prev.paragraphs), i + 3)):
        bl = prev.paragraphs[j]._p.findall('.//' + qn('a:blip'))
        if bl:
            blip = bl[0]
            ext = prev.paragraphs[j]._p.findall('.//' + qn('wp:extent'))
            if ext:
                width_in = min(int(ext[0].get('cx')) / 914400, 6.5)
            break
    if blip is not None:
        break
if blip is None:
    raise SystemExit("Layered figure image not found in previous_final")
open(IMG, 'wb').write(prev.part.related_parts[blip.get(qn('r:embed'))].blob)
print("extracted Layered figure:", os.path.getsize(IMG), "bytes, width_in=%.2f" % width_in)

# --- open current ---
d = Document(REV)
_id = [17000]
def nid():
    _id[0] += 1; return str(_id[0])
def insmark():
    e = OxmlElement('w:ins'); e.set(qn('w:id'), nid()); e.set(qn('w:author'), AUTHOR); e.set(qn('w:date'), DATE); return e
def ins_wrap(run):
    e = insmark(); e.append(run); return e

# --- capture an existing caption's style to clone ---
cap_src = next(p for p in d.paragraphs if 'One BEI / EDM Architecture' in acc(p) and acc(p).strip().startswith('Figure'))
src_pPr = copy.deepcopy(cap_src._p.find(qn('w:pPr')))
src_rPr = cap_src.runs[0]._r.find(qn('w:rPr'))
src_rPr = copy.deepcopy(src_rPr) if src_rPr is not None else None

# --- renumber: list of (unique_context, old, new) ; direct edits ---
OPS = [
    # captions
    ("Figure 1 - One BEI", "Figure 1", "Figure 2"),
    ("Figure 2 - Demand Forecasting", "Figure 2", "Figure 4"),
    ("Figure 3 - Legacy Approach", "Figure 3", "Figure 5"),
    ("Figure 4 - Our Supplier Risk Zones", "Figure 4", "Figure 6"),
    ("Figure 5 - Logistics Network", "Figure 5", "Figure 7"),
    ("Figure 6 - Technical Evidence", "Figure 6", "Figure 8"),
    ("Figure 7 - Bounded Pilot", "Figure 7", "Figure 9"),
    # in-text references
    ("as shown in Figure 1, represents", "Figure 1", "Figure 2"),
    ("as shown in Figure 1 below", "Figure 1", "Figure 3"),
    ("As shown in Figure 3, the legacy", "Figure 3", "Figure 5"),
    ("shown in Figure 2, changes the unit", "Figure 2", "Figure 4"),
    ("as shown in Figure 2. Five signal", "Figure 2", "Figure 4"),
    ("as shown in Figure 4, is better treated", "Figure 4", "Figure 6"),
    ("detailed in Figure 4, below", "Figure 4", "Figure 6"),
    ("Figure 4 shows both. The top row", "Figure 4", "Figure 6"),
]
def renumber(context, old, new):
    for p in d.paragraphs:
        if context in acc(p):
            for t in p._p.iter(qn('w:t')):
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new, 1); return True
    return False
applied = sum(1 for c, o, n in OPS if renumber(c, o, n))
print("renumber ops applied:", applied, "of", len(OPS))

# --- restore Layered figure (tracked) after the layered-architecture sentence ---
anchor = next(p for p in d.paragraphs if 'Our approach consists of a layered analytical architecture' in acc(p))
tmp_p = d.add_paragraph(); tmp_r = tmp_p.add_run(); tmp_r.add_picture(IMG, width=Inches(width_in))
drawing = copy.deepcopy(tmp_r._r.find(qn('w:drawing')))
dp = drawing.find('.//' + qn('wp:docPr'))
if dp is not None:
    dp.set('id', '9003'); dp.set('name', 'Figure 3 Layered Architecture')
# figure paragraph (centered, tracked)
figp = OxmlElement('w:p'); pPr = OxmlElement('w:pPr')
jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)
rPr = OxmlElement('w:rPr'); rPr.append(insmark()); pPr.append(rPr); figp.append(pPr)
runel = OxmlElement('w:r'); runel.append(drawing); figp.append(ins_wrap(runel))
# caption paragraph (cloned style, tracked)
capp = OxmlElement('w:p')
capp.append(copy.deepcopy(src_pPr))
crPr = capp.find(qn('w:pPr')).find(qn('w:rPr'))
if crPr is None:
    crPr = OxmlElement('w:rPr'); capp.find(qn('w:pPr')).append(crPr)
crPr.append(insmark())
crun = OxmlElement('w:r')
if src_rPr is not None: crun.append(copy.deepcopy(src_rPr))
ct = OxmlElement('w:t'); ct.set(qn('xml:space'), 'preserve'); ct.text = "Figure 3 - Our Layered Analytical Architecture"
crun.append(ct)
capp.append(ins_wrap(crun))
anchor._p.addnext(capp); anchor._p.addnext(figp)
tmp_p._p.getparent().remove(tmp_p._p)

o = REV; k = 3
while True:
    try:
        d.save(o); break
    except PermissionError:
        root, ext = os.path.splitext(REV); o = f"{root.replace('_v2','')}_v{k}{ext}"; k += 1
print("Saved:", o)
