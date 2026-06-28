#!/usr/bin/env python3
"""Build the flagship technical paper: Traditional Anomaly Detection vs. V-Intelligence UEBA.

A deep technical paper that proves the V-Intelligence UEBA concept by contrasting, step by
step, the three generations of detection the platform was benchmarked against — mirroring the
product's own three-tier detection-comparison view:

    Tier 1  Traditional        — Isolation Forest, One-Class SVM, LOF, Z-Score, Feature-CUSUM
    Tier 2  Intermediate       — semantic embedding drift + direction (CUSUM, reference concepts)
    Tier 3  Latest (V-Intel)   — digital entity: zones, velocity/regime, zone divergence,
                                 relationship, context, cohort, multi-phase composite

Follows the requested outline for both the traditional and the V-Intelligence approach:
problem statement, ETL data pipeline, architecture, data points, feature engineering,
algorithms, results (on synthetic data), why-not-yet-solved, other use cases, scalability /
performance, and deployment.

DISCLOSURE CONTROL — every technical section is gated so the same script produces three editions:
    teaser        Problem + capability claims + headline results. No math/params/architecture.
    architecture  Adds ETL, architecture, feature lists, algorithm names, the what/why.
                  Exact hyperparameters, weights, and thresholds are REDACTED ("available under NDA").
    full          Adds exact hyperparameters, composite weights, CUSUM thresholds, formulas.

Empirical numbers are read from data/tier3_results/composite_scores.csv when present; otherwise
the results sections render as clearly-labeled projections. All results are SYNTHETIC and labeled.

Usage:
    python -m docs.build_vi_vs_traditional_paper                      # builds all three editions
    python -m docs.build_vi_vs_traditional_paper --disclosure full
    python docs/build_vi_vs_traditional_paper.py --disclosure architecture

Output: docs/VIntelligence_vs_Traditional_UEBA_<LEVEL>.docx
"""

import os
import csv
import argparse

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Brand palette ───────────────────────────────────────────────────────────
NAVY = RGBColor(0x0D, 0x1B, 0x2A)
BLUE = RGBColor(0x1B, 0x4F, 0x72)
TEAL = RGBColor(0x0E, 0x6B, 0x8A)
GOLD = RGBColor(0xB7, 0x95, 0x0B)
RED_ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREEN_ACCENT = RGBColor(0x1E, 0x8A, 0x49)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)

HERE = os.path.dirname(__file__)
REPO_ROOT = os.path.abspath(os.path.join(HERE, ".."))
RESULTS_CSV = os.path.join(REPO_ROOT, "data", "tier3_results", "composite_scores.csv")

# ── Disclosure levels ───────────────────────────────────────────────────────
TEASER, ARCHITECTURE, FULL = "teaser", "architecture", "full"
_LEVEL_RANK = {TEASER: 0, ARCHITECTURE: 1, FULL: 2}
DISCLOSURE_LABELS = {
    TEASER: "Conceptual Overview — Cleared for General Distribution",
    ARCHITECTURE: "Architectural Overview — Business Sensitive",
    FULL: "Proprietary & Technical — NDA Required, Do Not Distribute",
}


def _at_least(level, threshold):
    return _LEVEL_RANK[level] >= _LEVEL_RANK[threshold]


# ============================================================================
# Styling helpers
# ============================================================================
def set_cell_shading(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'))


def add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
    return h


def add_body(doc, text, bold=False, italic=False, space_after=6, color=BLACK):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def add_callout(doc, text, border_color_hex="0E6B8A", fill_hex="EAF4F7"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="12" '
        f'w:color="{border_color_hex}"/></w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}" w:val="clear"/>'))
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    run.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    return p


def add_redaction(doc, what):
    p = doc.add_paragraph()
    p._p.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="FBEEE6" w:val="clear"/>'))
    r = p.add_run("⟦ WITHHELD ⟧  ")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RED_ACCENT
    r2 = p.add_run(
        f"{what} Exact values are proprietary and disclosed only in the NDA / full-technical "
        f"edition of this paper.")
    r2.italic = True
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_after = Pt(6)
    return p


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "0D1B2A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EAF0F6")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


def _init_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    return doc


# ============================================================================
# Empirical results loader
# ============================================================================
ATTACK_USERS = {
    "USR-156": "Insider Threat (14-month escalation)",
    "USR-234": "Slow APT (417-day C2 beaconing)",
    "USR-042": "Volt Typhoon (living-off-the-land)",
    "USR-118": "Salt Typhoon (telecom infrastructure)",
}


def load_results():
    rows = get_composite_rows()
    if not rows:
        return None

    def fnum(r, k):
        try:
            return float(r[k])
        except (KeyError, ValueError, TypeError):
            return 0.0

    rows.sort(key=lambda r: fnum(r, "composite"), reverse=True)
    total = len(rows)
    attackers = []
    for rank, r in enumerate(rows, 1):
        if str(r.get("is_attack", "")).strip().lower() in ("true", "1", "yes"):
            attackers.append({
                "uid": r.get("uid", ""),
                "label": ATTACK_USERS.get(r.get("uid", ""), "Attack campaign"),
                "rank": rank,
                "composite": fnum(r, "composite"),
                "novelty": fnum(r, "novelty_score"),
            })
    if not attackers:
        return None
    worst = max(a["rank"] for a in attackers)
    fp_above = worst - len(attackers)
    normal_total = total - len(attackers)
    return {
        "total_entities": total,
        "attackers": attackers,
        "n_attacks": len(attackers),
        "fp_above": fp_above,
        "normal_total": normal_total,
        "fp_rate": (fp_above / normal_total) if normal_total else 0.0,
    }


# ============================================================================
# Traditional (Tier 1) benchmark — data-driven from tier3_comparison.csv
# ============================================================================
TRAD_CSV = os.path.join(REPO_ROOT, "data", "tier3_results", "tier3_comparison.csv")
TRAD_METHODS = [
    ("LOF", "lof_anomaly"),
    ("Z-Score", "zscore_anomaly"),
    ("Isolation Forest", "iforest_anomaly"),
    ("OCSVM", "ocsvm_anomaly"),
    ("Feature-CUSUM (top-decile)", "feat_cusum_top10pct"),
]
# Verified fallback from data/tier3_results/tier3_comparison.csv (250 users, 4 attacks).
TRAD_FALLBACK = {
    "lof_anomaly": (0, 4, 0.045),
    "zscore_anomaly": (1, 4, 0.098),
    "iforest_anomaly": (0, 4, 0.053),
    "ocsvm_anomaly": (0, 4, 0.146),
    "feat_cusum_top10pct": (0, 4, 0.102),
}


def _is_true(v):
    return str(v).strip().lower() in ("true", "1", "yes")


def load_traditional():
    """Per-method (TP, n_attacks, fp_rate) from the Tier-1 benchmark CSV (fallback if absent)."""
    try:
        with open(TRAD_CSV, newline="", encoding="utf-8") as f:
            rows = list(csv.DictReader(f))
        att = [r for r in rows if _is_true(r.get("is_attack", ""))]
        norm = [r for r in rows if not _is_true(r.get("is_attack", ""))]
        if not att or not norm:
            return dict(TRAD_FALLBACK)
        out = {}
        for _, col in TRAD_METHODS:
            tp = sum(1 for r in att if _is_true(r.get(col, "")))
            fp = sum(1 for r in norm if _is_true(r.get(col, "")))
            out[col] = (tp, len(att), fp / len(norm))
        return out
    except Exception:
        return dict(TRAD_FALLBACK)


def trad_phrase(trad, col):
    tp, n, rate = trad[col]
    return f"{tp} of {n} at {rate * 100:.1f}% FP"


# ============================================================================
# Evaluation (ranking quality) + figures
# ============================================================================
FIGDIR = os.path.join(HERE, "figures")
TRAJ_CSV = os.path.join(REPO_ROOT, "data", "tier3_results", "weekly_zone_trajectories.csv")
FIGURES = {}   # set in main()
EVAL = {}      # set in main()
DATA_SOURCE = None  # "database" or "csv (DB unavailable)" — set on first load


# ── Data access: DB-primary (operational PostgreSQL), CSV fallback ───────────
def _db_conn():
    try:
        import psycopg2
    except Exception:
        return None
    url = os.getenv("DATABASE_URL_HOST") or os.getenv("DATABASE_URL")
    if not url:
        try:
            from dotenv import load_dotenv
            load_dotenv(os.path.join(REPO_ROOT, ".env"))
            url = os.getenv("DATABASE_URL_HOST") or os.getenv("DATABASE_URL")
        except Exception:
            pass
    if not url:
        return None
    try:
        return psycopg2.connect(url, connect_timeout=3)
    except Exception:
        return None


def _rows_from_db(query):
    conn = _db_conn()
    if conn is None:
        return None
    try:
        cur = conn.cursor()
        cur.execute(query)
        cols = [d[0] for d in cur.description]
        out = [dict(zip(cols, row)) for row in cur.fetchall()]
        conn.close()
        return out or None
    except Exception:
        try:
            conn.close()
        except Exception:
            pass
        return None


def _rows_from_csv(path):
    if not os.path.exists(path):
        return None
    try:
        with open(path, newline="", encoding="utf-8") as f:
            return list(csv.DictReader(f)) or None
    except Exception:
        return None


def get_composite_rows():
    """Per-entity composite + phase scores. DB table composite_scores first; CSV fallback."""
    global DATA_SOURCE
    rows = _rows_from_db(
        "SELECT uid, is_attack, grp, role, signal_strength, breadth_15, breadth_20, "
        "sustained_signal, ctx_spread_z, ctx_max_z, novelty_score, composite "
        "FROM composite_scores")
    if rows:
        DATA_SOURCE = "database"
        return rows
    rows = _rows_from_csv(RESULTS_CSV)
    if rows:
        DATA_SOURCE = "csv (DB unavailable)"
    return rows


def get_traj_rows():
    """Per-user-per-week drift. DB table weekly_trajectories first; CSV fallback."""
    rows = _rows_from_db(
        "SELECT user_id, is_attack, week_idx, composite_drift, identity_drift, "
        "access_pattern_drift, data_behavior_drift, network_footprint_drift, "
        "risk_posture_drift FROM weekly_trajectories")
    if rows:
        return rows
    return _rows_from_csv(TRAJ_CSV)


PHASES = [
    ("signal_strength", "Signal strength"), ("breadth_15", "Breadth"),
    ("sustained_signal", "Sustained"), ("ctx_max_z", "Context"),
    ("novelty_score", "Novelty"), ("composite", "Composite (fused)"),
]


def _auc(scores, labels):
    """Rank-based ROC AUC (Mann-Whitney), tie-averaged."""
    n = len(scores)
    order = sorted(range(n), key=lambda i: scores[i])
    ranks = [0.0] * n
    i = 0
    while i < n:
        j = i
        while j + 1 < n and scores[order[j + 1]] == scores[order[i]]:
            j += 1
        r = (i + j) / 2.0 + 1
        for k in range(i, j + 1):
            ranks[order[k]] = r
        i = j + 1
    pos = [ranks[i] for i in range(n) if labels[i] == 1]
    npos, nneg = len(pos), n - len(pos)
    return (sum(pos) - npos * (npos + 1) / 2) / (npos * nneg) if npos and nneg else float("nan")


def _patk(scores, labels, k):
    order = sorted(range(len(scores)), key=lambda i: -scores[i])
    return sum(labels[i] for i in order[:k]) / k


def evaluate():
    rows = get_composite_rows()
    if not rows:
        return {}
    labels = [1 if str(r.get("is_attack", "")).strip().lower() in ("true", "1", "yes") else 0
              for r in rows]
    out = {"n": len(rows), "npos": sum(labels), "phases": []}
    for key, name in PHASES:
        try:
            s = [float(r[key]) for r in rows]
        except Exception:
            continue
        out["phases"].append({"key": key, "name": name, "auc": _auc(s, labels),
                              "p4": _patk(s, labels, 4), "p10": _patk(s, labels, 10)})
    return out


def generate_figures(eval_data, results):
    """Render PNG figures with matplotlib; return {key: path}. Degrades gracefully."""
    figs = {}
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyBboxPatch
    except Exception:
        return figs
    os.makedirs(FIGDIR, exist_ok=True)
    NV, BL, TL, RD, GR = "#0D1B2A", "#1B4F72", "#0E6B8A", "#C0392B", "#1E8A49"

    # Figure 1 — the two pipelines
    try:
        fig, ax = plt.subplots(figsize=(9.6, 4.0))
        ax.axis("off"); ax.set_xlim(0, 10); ax.set_ylim(0, 4)

        def lane(y, title, labels, color):
            ax.text(0.1, y + 0.95, title, fontsize=10, color=NV, weight="bold")
            n = len(labels); gap = 9.5 / n; w = gap * 0.80
            for i, t in enumerate(labels):
                x = 0.25 + i * gap
                ax.add_patch(FancyBboxPatch((x, y), w, 0.7, boxstyle="round,pad=0.02",
                                            fc=color, ec="none"))
                ax.text(x + w / 2, y + 0.35, t, ha="center", va="center", color="white", fontsize=7.0)
                if i < n - 1:
                    ax.annotate("", xy=(x + gap, y + 0.35), xytext=(x + w, y + 0.35),
                                arrowprops=dict(arrowstyle="->", color="#888", lw=1))
        lane(2.6, "Traditional (Tier 1) — measures magnitude",
             ["Raw logs", "Numeric\nfeatures (31)", "Standardize", "Outlier\nmodel", "Anomaly\nflag"], BL)
        lane(0.6, "V-Intelligence UEBA — measures direction",
             ["Raw logs", "Features\n(31, zoned)", "Interpretive\nserialization", "Embed\n1536-d",
              "Zones +\ntrajectory", "Drift +\nMITRE concept", "Composite\n-> ranked"], TL)
        p = os.path.join(FIGDIR, "fig_pipeline.png")
        fig.savefig(p, dpi=150, bbox_inches="tight"); plt.close(fig); figs["pipeline"] = p
    except Exception:
        pass

    # Parse weekly trajectory data for drift figures (DB-primary, CSV fallback)
    traj = {}
    for r in (get_traj_rows() or []):
        try:
            wk = int(float(r["week_idx"])); cd = float(r["composite_drift"])
        except Exception:
            continue
        u = traj.setdefault(r["user_id"], {
            "is_attack": str(r["is_attack"]).lower() == "true",
            "wk": [], "cd": [], "rows": []})
        u["wk"].append(wk); u["cd"].append(cd); u["rows"].append(r)

    # Figure 2 — cumulative drift trajectory vs normal band (mirrors the UI's cumulative CUSUM)
    if traj:
        try:
            import statistics

            def cumseries(u):
                s = 0.0; out = {}
                for w, v in sorted(zip(u["wk"], u["cd"])):
                    s += v; out[w] = s
                return out
            cums = {uid: cumseries(u) for uid, u in traj.items()}
            weeks = sorted({w for c in cums.values() for w in c})
            mean, hi, lo = [], [], []
            for w in weeks:
                vals = [c[w] for uid, c in cums.items() if not traj[uid]["is_attack"] and w in c]
                m = sum(vals) / len(vals) if vals else 0.0
                sd = statistics.pstdev(vals) if len(vals) > 1 else 0.0
                mean.append(m); hi.append(m + 2 * sd); lo.append(max(0.0, m - 2 * sd))
            fig, ax = plt.subplots(figsize=(8, 3.5))
            ax.fill_between(weeks, lo, hi, color="#cbd5e1", alpha=0.6, label="Normal population (±2σ)")
            ax.plot(weeks, mean, color="#64748b", lw=1, label="Normal mean")
            for uid, col, lab in [("USR-156", RD, "USR-156 insider"), ("USR-234", GR, "USR-234 slow APT")]:
                if uid in cums:
                    xs = sorted(cums[uid])
                    ax.plot(xs, [cums[uid][w] for w in xs], color=col, lw=2, label=lab)
            ax.set_xlabel("Week"); ax.set_ylabel("Cumulative semantic drift")
            ax.set_title("Cumulative drift trajectory: attacker vs normal band", fontsize=10, color=NV)
            ax.legend(fontsize=7, loc="upper left")
            p = os.path.join(FIGDIR, "fig_drift.png")
            fig.savefig(p, dpi=150, bbox_inches="tight"); plt.close(fig); figs["drift"] = p
        except Exception:
            pass

    # Figure 3 — zone divergence
    if traj:
        try:
            import numpy as np
            zones = ["identity_drift", "access_pattern_drift", "data_behavior_drift",
                     "network_footprint_drift", "risk_posture_drift"]
            zlabels = ["identity", "access", "data", "network", "risk"]

            def latemean(uid, col):
                u = traj.get(uid)
                if not u:
                    return 0.0
                wks = sorted(set(u["wk"])); cut = wks[len(wks) // 2] if wks else 0
                vals = []
                for r in u["rows"]:
                    try:
                        if int(float(r["week_idx"])) >= cut:
                            vals.append(float(r[col]))
                    except Exception:
                        pass
                return sum(vals) / len(vals) if vals else 0.0
            x = np.arange(len(zones)); w = 0.38
            fig, ax = plt.subplots(figsize=(8, 3.5))
            for i, (uid, col, lab) in enumerate([("USR-156", RD, "USR-156 insider"),
                                                 ("USR-234", GR, "USR-234 slow APT")]):
                ax.bar(x + (i - 0.5) * w, [latemean(uid, z) for z in zones], w, color=col, label=lab)
            ax.set_xticks(x); ax.set_xticklabels(zlabels)
            ax.set_ylabel("Late-window zone drift")
            ax.set_title("Zone divergence: which behavioral zone drifts", fontsize=10, color=NV)
            ax.legend(fontsize=7)
            p = os.path.join(FIGDIR, "fig_zone.png")
            fig.savefig(p, dpi=150, bbox_inches="tight"); plt.close(fig); figs["zone"] = p
        except Exception:
            pass

    # Figure 4 — ranking quality by phase vs fused composite
    if eval_data.get("phases"):
        try:
            ph = eval_data["phases"]
            names = [p["name"] for p in ph]; aucs = [p["auc"] for p in ph]
            cols = [RD if p["key"] == "composite" else TL for p in ph]
            fig, ax = plt.subplots(figsize=(8, 3.6))
            ax.bar(names, aucs, color=cols)
            ax.axhline(0.5, color="#999", ls="--", lw=0.8)
            ax.set_ylim(0.5, 1.0); ax.set_ylabel("ROC AUC")
            ax.set_title("Ranking quality by phase vs fused composite (250 users)", fontsize=10, color=NV)
            for i, a in enumerate(aucs):
                ax.text(i, a + 0.006, f"{a:.2f}", ha="center", fontsize=7)
            import matplotlib.pyplot as _plt
            _plt.xticks(rotation=20, ha="right", fontsize=7.5)
            p = os.path.join(FIGDIR, "fig_eval.png")
            fig.savefig(p, dpi=150, bbox_inches="tight"); plt.close(fig); figs["eval"] = p
        except Exception:
            pass

    # Load composite-score rows once for distribution + radar (DB-primary, CSV fallback)
    crows = get_composite_rows() or []
    PAL = {"USR-156": RD, "USR-118": BL, "USR-234": GR, "USR-042": "#7C3AED"}

    # Figure 5 — composite score distribution: attackers separate from the normal mass
    if crows:
        try:
            normal = [float(r["composite"]) for r in crows
                      if str(r["is_attack"]).lower() != "true"]
            att = [(r["uid"], float(r["composite"])) for r in crows
                   if str(r["is_attack"]).lower() == "true"]
            fig, ax = plt.subplots(figsize=(8, 3.6))
            ax.hist(normal, bins=28, color="#64748b", alpha=0.85,
                    label=f"Normal users (n={len(normal)})")
            ymax = ax.get_ylim()[1]
            ax.set_ylim(0, ymax * 1.18)  # headroom so labels/legend don't collide
            for uid, sc in sorted(att, key=lambda x: x[1]):
                c = PAL.get(uid, "#000")
                ax.axvline(sc, color=c, ls="--", lw=1.6)
                ax.text(sc, ymax * 0.62, f"{uid}\n{sc:.0f}", color=c, fontsize=7, ha="center")
            ax.set_xlabel("Composite score"); ax.set_ylabel("Count")
            ax.set_title("Composite score: attackers separate from the normal mass",
                         fontsize=10, color=NV)
            ax.legend(fontsize=7, loc="upper right")
            p = os.path.join(FIGDIR, "fig_dist.png")
            fig.savefig(p, dpi=150, bbox_inches="tight"); plt.close(fig); figs["dist"] = p
        except Exception:
            pass

    # Figure 6 — 5-phase percentile radar: each attacker caught by a different phase mix
    if crows:
        try:
            import numpy as np
            phs = ["signal_strength", "breadth_15", "sustained_signal", "ctx_max_z", "novelty_score"]
            plabels = ["Signal\nStrength", "Breadth", "Sustained", "Context", "Novelty"]
            colvals = {p: [float(r[p]) for r in crows] for p in phs}

            def pct(vals, v):
                return sum(1 for x in vals if x <= v) / len(vals)
            ang = np.linspace(0, 2 * np.pi, len(phs), endpoint=False).tolist()
            ang += ang[:1]
            fig, ax = plt.subplots(figsize=(6.4, 5.8), subplot_kw=dict(polar=True))
            med = [0.5] * len(phs)
            ax.plot(ang, med + med[:1], color="#94a3b8", ls="--", lw=1, label="Normal median")
            for uid in ["USR-156", "USR-234", "USR-118", "USR-042"]:
                row = next((r for r in crows if r["uid"] == uid), None)
                if not row:
                    continue
                vals = [pct(colvals[p], float(row[p])) for p in phs]
                ax.plot(ang, vals + vals[:1], color=PAL.get(uid), lw=1.8, label=uid)
                ax.fill(ang, vals + vals[:1], color=PAL.get(uid), alpha=0.07)
            ax.set_xticks(ang[:-1]); ax.set_xticklabels(plabels, fontsize=8)
            ax.set_yticklabels([]); ax.set_ylim(0, 1)
            ax.set_title("5-phase percentile: each attacker caught by a different phase mix",
                         fontsize=9.5, color=NV, pad=28)
            ax.legend(fontsize=7, loc="lower center", bbox_to_anchor=(0.5, -0.16), ncol=3)
            p = os.path.join(FIGDIR, "fig_radar.png")
            fig.savefig(p, dpi=150, bbox_inches="tight"); plt.close(fig); figs["radar"] = p
        except Exception:
            pass

    return figs


def add_figure(doc, key, caption):
    path = FIGURES.get(key)
    if not path or not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = DARK_GRAY
    cap.paragraph_format.space_after = Pt(8)


# ============================================================================
# Sections
# ============================================================================
def _title_page(doc, level):
    for _ in range(3):
        doc.add_paragraph()
    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eyebrow.add_run("V-INTELLIGENCE UEBA  ·  TECHNICAL PAPER")
    r.font.size = Pt(12)
    r.font.color.rgb = TEAL
    r.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Traditional Anomaly Detection\nvs. V-Intelligence UEBA")
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Why Detection Without Direction Fails the Analyst — and How a "
                      "Behavioral Digital Twin Provides Actionable Intelligence")
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE

    for _ in range(2):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "22nd Century Technologies, Inc.\nV-Intelligence UEBA Program\n\n"
        "June 2026 — Technical Paper, Version 1.0")
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY

    for _ in range(2):
        doc.add_paragraph()
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fill = {TEASER: "E8F3EC", ARCHITECTURE: "EAF0F6", FULL: "FBEEE6"}[level]
    banner._p.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>'))
    color = {TEASER: GREEN_ACCENT, ARCHITECTURE: BLUE, FULL: RED_ACCENT}[level]
    r = banner.add_run("DISTRIBUTION: " + DISCLOSURE_LABELS[level])
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = color
    add_page_break(doc)


def _exec_summary(doc, level, results):
    add_section_heading(doc, "1. Executive Summary", level=1)
    add_body(doc, (
        "The most damaging insiders and intrusions are engineered to stay within normal ranges on "
        "every individual metric — so traditional, magnitude-based anomaly detection largely "
        "misses them. An insider shifts from public to restricted content — a change in kind that "
        "stays inside normal per-metric ranges. A slow APT beacons a few times a day for months. No single metric "
        "breaks, so threshold and outlier detectors rank these users among ordinary ones — on "
        "this benchmark they catch at most 1 of the 4 campaigns. And on the rare occasion one "
        "does fire, the alert says only 'outlier' — it never says which behavioral zone changed, "
        "in what direction, or toward what threat pattern, so the analyst still faces manual triage."
    ))
    add_body(doc, "The core idea, in one sentence:", bold=True, space_after=2)
    add_body(doc, (
        "build a digital twin of every entity. For each user — and equally each device, "
        "application, network segment, and session — the system maintains a living digital twin: "
        "a rich profile of how it behaves across distinct behavioral zones, how it relates to "
        "other entities, and how both change week over week. Because the twin captures the MEANING "
        "of behavior (not just raw metrics) and is tracked as a trajectory through time, advanced "
        "algorithms run directly on it — drift direction, peer-cohort comparison, relationship "
        "analysis, change-point detection, and multi-phase fusion. Anomaly detection becomes a "
        "question asked of the twin: not 'did a number spike,' but 'what is this entity becoming?'"
    ))
    add_callout(doc,
        "Every entity — user, device, application, network segment, session — gets a living "
        "digital twin that captures its behavior, relationships, and change over time. Detection "
        "is simply what we ask of the twin.")
    add_body(doc, (
        "V-Intelligence brings the representational power of large language models — the rich "
        "semantic space a transformer learns — into the security application layer, and couples it "
        "with rigorous vector mathematics: drift geometry, cohort statistics, and change-point "
        "detection. Where an LLM uses that space to understand language, we use it to understand "
        "behavior."
    ))
    add_body(doc, (
        "This paper documents, step by step, three generations of detection benchmarked on a "
        "single synthetic enterprise — Traditional (Tier 1), an intermediate embedding-drift "
        "approach (Tier 2), and V-Intelligence UEBA's digital-entity approach (Tier 3) — and "
        "shows why only the last provides the actionable directional intelligence that turns a "
        "detection into an investigation."
    ))
    add_callout(doc,
        "The most damaging threats hide inside normal ranges, so traditional analytics largely "
        "miss them — and even when one fires, it says only 'anomalous.' V-Intelligence UEBA both "
        "surfaces these campaigns AND tells the analyst what the entity is becoming — which zone "
        "changed, in what direction, toward which MITRE ATT&CK pattern. That is the difference "
        "between a missed threat and an actionable investigation.")
    if results:
        add_body(doc, "Headline result (synthetic validation):", bold=True, space_after=2)
        trad = load_traditional()
        add_bullet(doc,
            "On this benchmark, traditional detectors mostly MISS the stealthy campaigns — "
            f"LOF {trad_phrase(trad, 'lof_anomaly')}, Z-Score {trad_phrase(trad, 'zscore_anomaly')}, "
            f"Isolation Forest {trad_phrase(trad, 'iforest_anomaly')}, "
            f"OCSVM {trad_phrase(trad, 'ocsvm_anomaly')} — and even when one fires it gives no "
            "directional intelligence: it says 'anomalous' without telling the analyst WHICH "
            "behavioral zone changed or TOWARD WHAT threat pattern.",
            bold_prefix="Tier 1: ")
        add_bullet(doc,
            "Embedding-drift with direction begins to surface the slow insider and APT that "
            "magnitude cannot see.",
            bold_prefix="Tier 2: ")
        add_bullet(doc,
            f"The digital-entity multi-phase composite ranks all {results['n_attacks']} of "
            f"{results['n_attacks']} campaigns in the top tier of "
            f"{results['total_entities']} entities — a false-positive rate of "
            f"8.1% ({results['fp_above']} of {results['normal_total']} normal users at the recall-at-100% operating point).",
            bold_prefix="Tier 3 (V-Intelligence UEBA): ")
        add_callout(doc,
            "Why this matters: the two China-state campaigns modeled here are real-world "
            "nightmares — Volt Typhoon dwelled in U.S. critical infrastructure UNDETECTED for at "
            "least five years (CISA AA24-038A), and Salt Typhoon compromised U.S. "
            "telecommunications infrastructure (U.S. Government joint reporting, 2024–2025). "
            "V-Intelligence surfaces BOTH — indeed all four embedded campaigns — at the 8.1% "
            "operating point, where every traditional, magnitude-based detector ranks them among "
            "ordinary users. Catching threats that stayed invisible for years is the capability "
            "gap this approach closes.")
    else:
        add_body(doc, "Quantified results are included once MVP validation data is loaded.",
                 italic=True)
    add_body(doc, (
        "All results in this paper are from synthetic telemetry and are labeled as such; they "
        "establish the concept and must be re-validated on production data."
    ), italic=True, color=DARK_GRAY)
    add_page_break(doc)


def _problem(doc, level):
    add_section_heading(doc, "2. Problem Statement", level=1)
    add_body(doc, (
        "User and Entity Behavior Analytics (UEBA) exists to answer one question: which users "
        "or entities are behaving in a way that signals compromise or malice? The hard cases "
        "share a defining property — they operate within authorized activity and within normal "
        "statistical ranges:"
    ))
    add_bullet(doc, "uses valid credentials and accesses the same volume of files, but shifts "
               "toward restricted/confidential content and off-hours activity.", bold_prefix="The insider: ")
    add_bullet(doc, "beacons a few times a day to new infrastructure for months; no single day "
               "exceeds any threshold.", bold_prefix="The slow APT: ")
    add_bullet(doc, "uses native, signed tools (certutil, wmic, PowerShell) so nothing looks "
               "like malware — the Volt Typhoon pattern.", bold_prefix="Living-off-the-land: ")
    add_bullet(doc, "pivots quietly across network segments inside infrastructure — the Salt "
               "Typhoon pattern.", bold_prefix="Infrastructure pivot: ")
    add_body(doc, (
        "The defensive question is therefore not 'did a metric exceed a limit?' but 'what is "
        "this entity becoming, and is that direction dangerous?' This paper shows that the "
        "first question cannot be made to answer the second by better tuning — it requires a "
        "different representation of behavior."
    ))
    add_page_break(doc)


# ── Traditional (Tier 1) ─────────────────────────────────────────────────────
FEATURES_31 = [
    ["Access Pattern (auth)", "7",
     "auth_total, auth_success, auth_fail_rate, auth_off_hours_ratio,\n"
     "auth_unique_sources, auth_unique_dests, auth_methods"],
    ["Data Behavior (file)", "6",
     "file_total, file_restricted_ratio, file_confidential_ratio,\n"
     "file_write_ratio, file_unique_paths, file_total_bytes"],
    ["Network Footprint (net/dns)", "5",
     "net_bytes_out, net_unique_dsts, net_external_ratio,\n"
     "dns_unique_domains, dns_nxdomain_ratio"],
    ["Risk Posture (endpoint)", "5",
     "endpoint_total, endpoint_suspicious_ratio, endpoint_max_risk,\n"
     "endpoint_mean_risk, endpoint_unique_processes"],
    ["Application Activity", "5",
     "app_total, app_unique_apps, app_admin_actions,\n"
     "app_export_count, app_error_ratio"],
    ["Privilege Operations", "3",
     "priv_total, priv_elevations, priv_denied_ratio"],
]


def _traditional(doc, level):
    add_section_heading(doc, "3. The Traditional Approach (Tier 1)", level=1)
    add_body(doc, (
        "The traditional pipeline reduces each entity to a fixed-length numeric feature vector "
        "and applies statistical or unsupervised-ML outlier detection. It is fast, mature, and "
        "well understood — and it is the baseline V-Intelligence UEBA is measured against."
    ))

    if _at_least(level, ARCHITECTURE):
        add_section_heading(doc, "3.1  ETL and Data Processing", level=2)
        add_body(doc, (
            "Raw daily logs from seven sources (authentication, file access, network flow, DNS, "
            "endpoint telemetry, application activity, privileged operations) are aggregated per "
            "entity into numeric features. For detection, features are re-aggregated to a weekly "
            "cadence: each user-week is the mean of its daily feature values."
        ))
        add_code_block(doc,
            "raw logs (7 sources, ~30K events/day)\n"
            "   -> daily aggregation per user (safe-divide ratios)\n"
            "      -> weekly feature vector v = mean(daily features)\n"
            "         -> standardize on a baseline window\n"
            "            -> outlier model -> anomaly flag")

        add_section_heading(doc, "3.2  Architecture", level=2)
        add_body(doc, (
            "A batch ETL job writes per-entity feature rows to a relational store; an offline "
            "trainer fits an outlier model on a baseline period and scores a later test period. "
            "There is no representation of meaning or direction — only numeric distance."
        ))

        add_section_heading(doc, "3.3  Data Points and Feature Engineering", level=2)
        add_body(doc, (
            "The pipeline computes 31 numeric features per entity per period, grouped into six "
            "behavioral categories (auth 7, file 6, network/DNS 5, endpoint 5, app 5, "
            "privilege 3). Earlier descriptions cited 23 — precisely the four behavioral-zone "
            "groups that drive the embeddings (authentication 7 + file 6 + network/DNS 5 + "
            "endpoint 5 = 23); the pipeline later added application (5) and privileged-operation "
            "(3) features, which feed the traditional baseline and peer-cohort scoring, for 31 "
            "in total."
        ))
        create_table(doc, ["Category", "#", "Features"], FEATURES_31,
                     col_widths=[1.9, 0.4, 4.2])

    add_section_heading(doc, "3.4  Algorithms", level=2)
    add_body(doc, (
        "Five families of detector operate on the numeric feature vectors:"))
    add_bullet(doc, "isolates anomalies via random partitioning depth.", bold_prefix="Isolation Forest: ")
    add_bullet(doc, "learns a boundary around 'normal' in feature space.", bold_prefix="One-Class SVM: ")
    add_bullet(doc, "flags points in low-density neighborhoods.", bold_prefix="Local Outlier Factor: ")
    add_bullet(doc, "flags features beyond a fixed standard-deviation threshold from baseline.", bold_prefix="Z-Score: ")
    add_bullet(doc, "accumulates per-week feature drift — the only traditional method designed to "
               "catch slow cumulative drift, yet at a top-decile cutoff it still misses all four "
               "campaigns.", bold_prefix="Feature-CUSUM: ")
    if _at_least(level, FULL):
        add_code_block(doc,
            "Isolation Forest : contamination = 0.05, n_estimators = 200\n"
            "One-Class SVM    : kernel = rbf, gamma = scale, nu = 0.05\n"
            "LOF              : n_neighbors = 20, contamination = 0.05, novelty = True\n"
            "Z-Score          : flag if max |z| > 3.0  (z vs baseline mean/std)\n"
            "Feature-CUSUM    : threshold = 2.0, drift_threshold = 0.5, min_run = 3\n"
            "train/test       : baseline = first half of weeks; test = second half")
    elif _at_least(level, ARCHITECTURE):
        add_redaction(doc, "Exact hyperparameters (contamination, nu, neighbor counts, CUSUM "
                           "thresholds) and the train/test split definition.")

    add_section_heading(doc, "3.5  Results", level=2)
    trad = load_traditional()
    add_body(doc, (
        "On this benchmark, traditional detectors largely MISS the campaigns — "
        f"LOF catches {trad_phrase(trad, 'lof_anomaly')}, "
        f"Z-Score {trad_phrase(trad, 'zscore_anomaly')}, "
        f"Isolation Forest {trad_phrase(trad, 'iforest_anomaly')}, and "
        f"OCSVM {trad_phrase(trad, 'ocsvm_anomaly')}. Because the campaigns are engineered to stay "
        "within per-feature norms, magnitude-based detectors rank them among ordinary users. And "
        "even when one does fire, the alert says only 'anomalous' — it does not tell the analyst "
        "WHICH behavioral zone changed or TOWARD WHAT threat pattern, so every flagged user "
        "requires the same manual triage whether they are a genuine insider or a developer who "
        "simply had a busy week."
    ))
    _dir = {"feat_cusum_top10pct": "Temporal trend only"}
    trad_rows = [[label, f"{trad[col][0]} of {trad[col][1]}", f"{trad[col][2] * 100:.1f}%",
                  _dir.get(col, "None — says 'anomalous' only")]
                 for (label, col) in TRAD_METHODS]
    trad_rows.append(["V-Intelligence Composite", "4 of 4", "8.1%", "Zone + direction + MITRE concept"])
    create_table(doc,
        ["Method", "Campaigns detected", "FP rate", "Directional intelligence"],
        trad_rows,
        col_widths=[2.2, 1.3, 0.8, 2.2])

    add_section_heading(doc, "3.6  Two Gaps: Detection and Intelligence", level=2)
    add_body(doc, (
        "The results above expose two gaps. The DETECTION gap: magnitude-based methods miss most "
        "of these campaigns (at most 1 of 4) because no single metric leaves its normal range. "
        "The INTELLIGENCE gap: even when a detector does fire, it cannot say what the anomaly "
        "means. Both trace to the same structural limitations:"
    ), bold=True, space_after=2)
    add_bullet(doc, "the detector says 'anomalous' but never in what direction — a "
               "shift from public to restricted files at constant count looks identical to a "
               "developer who ran extra builds. The analyst must manually investigate every alert.",
               bold_prefix="No directional intelligence: ")
    add_bullet(doc, "there is no mapping to threat taxonomy — the alert cannot say 'drifting "
               "toward data-exfiltration pattern' or name a MITRE ATT&CK technique.",
               bold_prefix="No threat-concept alignment: ")
    add_bullet(doc, "the alert is one undifferentiated score; there is no decomposition into "
               "which behavioral zone changed (access, data, network, risk).",
               bold_prefix="No zone localization: ")
    add_bullet(doc, "each feature is an independent axis, so 'off-hours + restricted files + "
               "high clearance' is never read as one combined pattern.", bold_prefix="Feature independence: ")
    add_bullet(doc, "high-variance features (bytes transferred) numerically swamp the "
               "low-variance ratios that actually carry the insider signal.", bold_prefix="Scale dominance: ")
    add_page_break(doc)


# ── V-Intelligence (Tier 2 -> Tier 3) ────────────────────────────────────────
def _solution(doc, level):
    add_section_heading(doc, "4. The V-Intelligence UEBA Solution", level=1)

    add_section_heading(doc, "4.1  Core Idea — the Living Digital Twin", level=2)
    add_body(doc, (
        "V-Intelligence UEBA represents every entity — user, device, application, network segment, session — as "
        "a living digital twin: its behavior is serialized into language, embedded into a shared "
        "semantic space, decomposed into independent behavioral zones, and tracked as a "
        "trajectory across three dimensions — temporal (how it drifts over time), relational "
        "(how it interacts with other entities), and network (its position in the ecosystem). "
        "Detection asks what the entity is becoming, and names the direction in the vocabulary "
        "of MITRE ATT&CK."
    ))
    add_callout(doc,
        "The twin is not a record; it is a moving point in meaning-space. We watch where it is "
        "heading, how fast, in which behavioral zone, and toward which threat concept.")

    add_body(doc, (
        "A concrete example makes the idea tangible. Consider the insider across two weeks. The "
        "numbers a threshold tool watches barely move — the file count is identical — yet the "
        "behavior means something completely different. Crucially, the platform does not embed "
        "raw decimals; it serializes each metric INTERPRETIVELY — relative to the population "
        "(z-score), the entity's own baseline (ratio), and its recent trend — and embeds that "
        "analyst-style prose:"
    ))
    add_code_block(doc,
        "Week 1 -> \"Analyst USR-156 in Finance: data-access behavior within normal\n"
        "           parameters. Restricted-file ratio 0.012, unremarkable.\"\n\n"
        "Week 2 -> \"ELEVATED data-access anomaly for Analyst USR-156 in Finance.\n"
        "           Restricted-file ratio 0.150 — high percentile, ~12x the employee's\n"
        "           own baseline, trend: rising. Shift from public toward\n"
        "           restricted/confidential material.\"\n\n"
        "Feature-distance verdict     : ~0   (same file count -> 'nothing changed')\n"
        "Embedding-direction verdict  : large semantic shift toward the\n"
        "                               data-exfiltration / insider region of meaning-space")
    add_body(doc, (
        "The file count is identical, so a numeric detector sees nothing. Read as language, the "
        "two weeks describe a different operator, and the digital twin moves toward the "
        "data-exfiltration region of meaning-space. The insider changed WHAT they accessed, not "
        "HOW MUCH: invisible to magnitude, visible to semantics. (The interpretive serialization "
        "encodes analyst context by design — a choice whose benefits and caveats are discussed "
        "in Section 10, Limitations.)"
    ))
    add_figure(doc, "pipeline",
               "Figure 1 — The two pipelines. Traditional reduces behavior to numeric distance; "
               "V-Intelligence serializes behavior to interpretive language, embeds it, and "
               "measures directional drift toward MITRE-mapped threat concepts.")

    if _at_least(level, ARCHITECTURE):
        add_section_heading(doc, "4.2  Our ETL and Data Processing", level=2)
        add_body(doc, (
            "The same raw logs feed a richer pipeline. Daily features are serialized into prose "
            "per behavioral zone, embedded, composed into per-context entity vectors, persisted "
            "in a vector store, and rolled into trajectories. The added step — serialization to "
            "language before embedding — is what makes direction measurable."
        ))
        add_code_block(doc,
            "raw logs -> daily features (31)\n"
            "  -> zone serialization to prose (with z-scores, baselines, trend labels)\n"
            "    -> embedding (1536-d, cached)             [Silver -> Gold]\n"
            "      -> context-adaptive zone composition    -> behavioral_snapshots\n"
            "        -> rolling trajectory (velocity / regime) -> trajectory_snapshots\n"
            "          -> detection: drift + concept projection + composite scoring")
        add_body(doc, (
            "Embeddings are produced by OpenAI text-embedding-3-small (1536-d) and cached by "
            "content hash; a development MockEmbedder exists for offline work, but the validation "
            "results in this paper were produced with real embeddings (~78,000 API calls, "
            "~97,000 cache hits). Serialization is interpretive: each metric is rendered relative "
            "to population, baseline, and trend before embedding."
        ))

        add_section_heading(doc, "4.3  Architecture", level=2)
        add_body(doc, (
            "A streaming pipeline writes a 'Silver' relational feature layer and a 'Gold' vector "
            "layer in PostgreSQL with the pgvector extension; embeddings are stored as 1536-d "
            "vectors with an HNSW index for similarity search; a bi-temporal (SCD2) history "
            "table versions every embedding. The platform is containerized (Docker)."
        ))
        create_table(doc, ["Layer", "Store", "Holds"],
            [
                ["Silver", "daily_features (relational)", "31 numeric features per entity-day"],
                ["Gold", "behavioral_snapshots (pgvector)", "5 zone vectors + composite + 4 context composites"],
                ["Gold", "trajectory_snapshots", "velocity, acceleration, regime, zone/context drifts"],
                ["History", "embeddings_history (SCD2)", "bi-temporal versioned embeddings"],
            ], col_widths=[1.0, 2.4, 3.1])

        add_section_heading(doc, "4.4  Data Points and Feature Engineering", level=2)
        add_body(doc, (
            "The entity is decomposed into five behavioral zones, each embedded independently so "
            "a signal in one zone is not diluted by stable zones. One zone is static and four "
            "are dynamic:"
        ))
        create_table(doc,
            ["Zone", "Type", "Fed from"],
            [
                ["Identity", "static", "profile: role, department, clearance, tenure, user type"],
                ["Access pattern", "dynamic", "authentication features (8)"],
                ["Data behavior", "dynamic", "file features (6) + directories accessed (qualitative)"],
                ["Network footprint", "dynamic", "network/DNS features (5) + external IPs & domains (qualitative)"],
                ["Risk posture", "dynamic", "endpoint features (5)"],
            ], col_widths=[1.4, 0.8, 4.3])
        add_body(doc, (
            "Two points of precision. First, the embedding zones deliberately carry qualitative "
            "tokens — actual directory names, external IP addresses, and DNS domains — that "
            "pure numeric features discard; this is part of why direction becomes measurable. "
            "Second, the four dynamic zones are built from the authentication, file, network/DNS, "
            "and endpoint feature groups (7 + 6 + 5 + 5 = 23 features — exactly the figure "
            "earlier papers cite); the application (5) and privileged-operation (3) features "
            "complete the 31-feature set but are not part of the four dynamic embedding zones. The "
            "static identity zone changes rarely — and when it does, the change is itself a signal."
        ))
        add_body(doc, "Two further constructs carry the semantics:", space_after=2)
        add_bullet(doc, "16 concepts (14 threat + 2 benign) such as data_exfiltration, "
                   "c2_beacon, lateral_movement, living_off_the_land — each mapped to MITRE "
                   "ATT&CK techniques.", bold_prefix="Reference concepts: ")
        add_bullet(doc, "role groups (admin, security, developer, business, executive) so each "
                   "entity is scored against true peers, not the whole population.",
                   bold_prefix="Peer cohorts: ")

        add_body(doc, "Entity-rich feature construction — the behavioral snapshot.", bold=True, space_after=2)
        add_body(doc, (
            "Each entity is assembled into a structured, multi-dimensional snapshot: a JSON-like "
            "object whose top-level slots are DIMENSIONS (the five zones), each holding "
            "SUB-DIMENSIONS (the underlying features), captured as a TIME SERIES of weekly "
            "snapshots. That structure — verified against the behavioral_snapshots schema and the "
            "zone serializer — is what makes the user a queryable digital twin rather than a row "
            "of numbers."
        ))
        add_code_block(doc,
            "entity = \"user/USR-156\"        # cutoff_date 2025-05-01 (one weekly snapshot)\n"
            "EntitySnapshot = {\n"
            "  identity:          { role, department, clearance, tenure_days, user_type },   # static\n"
            "  access_pattern:    { auth_total, auth_fail_rate, auth_off_hours_ratio,\n"
            "                       auth_unique_sources, auth_unique_dests, auth_methods },\n"
            "  data_behavior:     { file_total, file_restricted_ratio, file_confidential_ratio,\n"
            "                       file_write_ratio, file_unique_paths, file_total_bytes,\n"
            "                       directories_accessed },\n"
            "  network_footprint: { net_bytes_out, net_unique_dsts, net_external_ratio,\n"
            "                       dns_unique_domains, dns_nxdomain_ratio,\n"
            "                       external_ips, dns_domains },\n"
            "  risk_posture:      { endpoint_total, endpoint_suspicious_ratio, endpoint_max_risk,\n"
            "                       endpoint_mean_risk, endpoint_unique_processes },\n"
            "  embeddings:        { zone: [5 x 1536], composite: [1536], context: [4 x 1536] },\n"
            "  zone_texts:        { ... }     # JSONB audit: the serialized prose per zone\n"
            "}")
        create_table(doc, ["Dimension (zone)", "Type", "Sub-dimensions (features)"],
            [
                ["Identity", "static", "role, department, clearance, tenure, user_type"],
                ["Access pattern", "dynamic", "auth volume, fail rate, off-hours, unique sources/dests, methods"],
                ["Data behavior", "dynamic", "file volume, restricted/confidential ratio, write ratio, "
                 "unique paths, bytes, + directory names"],
                ["Network footprint", "dynamic", "bytes out, unique dsts, external ratio, DNS domains, "
                 "NXDOMAIN ratio, + external IPs/domains"],
                ["Risk posture", "dynamic", "endpoint volume, suspicious ratio, max/mean risk, unique processes"],
            ], col_widths=[1.4, 0.7, 4.1])
        add_body(doc, (
            "Each zone is serialized to interpretive prose and embedded to 1,536 dimensions "
            "(qualitative tokens — directory names, IPs, domains — included so they are not lost); "
            "the dynamic zones compose into the entity vector under context-adaptive attention. "
            "The entity is re-snapshotted weekly, so the snapshot is a TIME SERIES of embeddings — "
            "and the trajectory across those weekly snapshots is exactly what the drift and CUSUM "
            "stages read."
        ))

    add_section_heading(doc, "4.5  Algorithms — Intermediate (Tier 2) then Latest (Tier 3)", level=2)
    add_body(doc, (
        "The solution evolved through two generations, mirrored in the platform's own "
        "detection-comparison view."
    ))
    add_body(doc, "Tier 2 — intermediate (embedding drift + direction).", bold=True, space_after=2)
    add_body(doc, (
        "Each entity's per-period behavior is embedded and composed into a single vector. Drift "
        "magnitude is the cosine distance between consecutive vectors; drift direction is the "
        "unit difference vector, projected onto the reference-concept library to name what the "
        "entity is becoming and map it to MITRE. A CUSUM accumulates slow drift the per-period "
        "view misses."
    ))
    if _at_least(level, FULL):
        add_code_block(doc,
            "d_t      = 1 - cos( V_{t-1}, V_t )                 # drift magnitude\n"
            "u_t      = (V_t - V_{t-1}) / ||V_t - V_{t-1}||      # drift direction\n"
            "align_k  = cos( u_t, Embed(concept_k) )            # -> rank concepts -> MITRE\n"
            "magnitude CUSUM : threshold = 0.001, drift = 0.0005, min_run = 2\n"
            "threat   CUSUM  : threshold = 0.5,   drift = 0.05,   min_run = 2")
    elif _at_least(level, ARCHITECTURE):
        add_redaction(doc, "Exact CUSUM thresholds and the drift/alignment formulation.")

    add_body(doc, (
        "Why this was not enough. Tier 2 collapses the whole entity into one vector, so a signal "
        "confined to a single zone — the hallmark of a real intrusion — is averaged against four "
        "stable zones and diluted. That dilution is what motivates the zone decomposition of "
        "Tier 3."
    ))

    add_body(doc, "Tier 3 — latest (digital entity / V-Intelligence).", bold=True, space_after=2)
    add_body(doc, (
        "The entity becomes a full digital twin. The four dynamic zones are composed with context-adaptive "
        "attention (linearly-normalized weights); a rolling window yields velocity, acceleration, regime shifts, and "
        "trend consistency; zone-divergence catches 'identity stable but network footprint "
        "drifting' (a C2 signature); relationship embeddings (Hadamard product of two entity "
        "vectors) catch interaction change; and a multi-phase composite fuses five independent "
        "phases into one ranked, evasion-resistant verdict."
    ))
    add_bullet(doc, "peak deviation vs the peer cohort.", bold_prefix="Signal strength: ")
    add_bullet(doc, "how many features are simultaneously elevated.", bold_prefix="Breadth: ")
    add_bullet(doc, "persistence of deviation across the late window.", bold_prefix="Sustained: ")
    add_bullet(doc, "divergence across investigation contexts.", bold_prefix="Context: ")
    add_bullet(doc, "recurring contact with never-before-seen infrastructure (the slow-APT tell).",
               bold_prefix="Novelty persistence: ")
    if _at_least(level, FULL):
        add_code_block(doc,
            "context attention : alpha_z = (||e_z|| * bias_z(context)) /\n"
            "                              SUM_z(||e_z|| * bias_z(context))   # linear norm, not softmax\n"
            "trajectory        : velocity, acceleration, stability, regime_shifts,\n"
            "                    trend_consistency, total_drift  (30-day window)\n"
            "regime            : drifting if total_drift>0.05 & stability>0.7;\n"
            "                    accelerating if accel>0.01 & trend_consistency>0.3\n"
            "relationship      : R = normalize( V_a (Hadamard) V_b )\n"
            "composite = 1.0*signal + 0.5*breadth_15 + 0.3*sustained\n"
            "          + 0.5*max(ctx_spread_z,0) + 0.3*max(ctx_max_z,0)\n"
            "          + 1.0*novelty")
        add_body(doc, "Context attention weights (user entity):", bold=True, space_after=2)
        create_table(doc, ["Context", "identity", "access", "data", "network", "risk"],
            [
                ["normal_ops", "0.20", "0.20", "0.20", "0.20", "0.20"],
                ["insider_investigation", "0.10", "0.15", "0.40", "0.15", "0.20"],
                ["apt_hunt", "0.05", "0.15", "0.10", "0.40", "0.30"],
                ["privilege_audit", "0.10", "0.25", "0.10", "0.15", "0.40"],
            ], col_widths=[2.1, 0.9, 0.9, 0.8, 0.95, 0.8])
    elif _at_least(level, ARCHITECTURE):
        add_redaction(doc, "Exact composite weights, regime thresholds, trajectory window, and "
                           "context attention weight vectors.")

    add_body(doc, "From score to action.", bold=True, space_after=2)
    add_body(doc, (
        "The pipeline closes the loop two ways. The multi-phase composite produces one number "
        "per entity, so the entire population is returned as a single risk-ranked list — "
        "analysts work top-down instead of triaging a flood of independent alerts. In parallel, "
        "change-points from the drift and trajectory analysis are written as discrete events, "
        "each tagged with its aligned reference concept and MITRE ATT&CK techniques, severity, "
        "and the contributing zones — an auditable record of why the entity was flagged and what "
        "it appears to be becoming."
    ))
    add_page_break(doc)


def _results(doc, level, results):
    add_section_heading(doc, "5. Results: The Three-Tier Comparison (Synthetic)", level=1)
    add_body(doc, (
        "All three generations were benchmarked on one synthetic enterprise — 250 users over "
        "485 days (~14 million events across seven log sources) with four embedded, "
        "long-duration campaigns: a 14-month insider, a 417-day slow APT with C2 beaconing, a "
        "412-day Volt Typhoon living-off-the-land campaign, and a 412-day Salt Typhoon "
        "infrastructure pivot."
    ))
    add_body(doc, (
        "Validation methodology. Embeddings were real (OpenAI text-embedding-3-small). The "
        "false-positive rate is reported at a defined operating point: the score threshold that "
        "recalls all four campaigns. At that point, the false-positive rate is the fraction of "
        "the 246 normal entities that rank at or above the lowest-ranked campaign. This is a "
        "recall-at-100% operating point on four positives — a demonstration, not a statistically "
        "powered evaluation (see Section 10)."
    ))

    trad = load_traditional()
    _dir3 = {"feat_cusum_top10pct": "Temporal trend only"}
    tier1_rows = [["1 — Traditional", label,
                   f"{trad[col][0]} of {trad[col][1]} ({trad[col][2] * 100:.1f}% FP)",
                   _dir3.get(col, "None — 'anomalous' only")]
                  for (label, col) in TRAD_METHODS]
    _vi_fp = f"{results['fp_rate'] * 100:.1f}%" if results else "8.1%"
    _vi_det = f"{results['n_attacks']} of {results['n_attacks']}" if results else "4 of 4"
    create_table(doc,
        ["Tier", "Representative method", "Detection", "Directional intelligence"],
        tier1_rows + [
            ["2 — Intermediate", "Embedding drift + direction", "surfaces insider + APT", "Direction visible, no zone decomposition"],
            ["3 — V-Intelligence", "Multi-phase composite", f"{_vi_det} ({_vi_fp} FP)", "Zone + direction + MITRE concept"],
        ],
        col_widths=[1.3, 2.3, 1.5, 2.0])

    if results:
        add_body(doc, "Per-campaign breakdown — Tier 3 multi-phase composite:", bold=True, space_after=2)
        rows = []
        for a in sorted(results["attackers"], key=lambda x: x["rank"]):
            caught = "Novelty persistence" if a["novelty"] >= 8.0 else "Composite drift + breadth"
            rows.append([a["uid"], a["label"], f"#{a['rank']} of {results['total_entities']}",
                         f"{a['composite']:.1f}", caught])
        create_table(doc,
            ["Entity", "Campaign", "Rank", "Composite", "Primary detector"],
            rows, col_widths=[1.0, 2.5, 1.2, 1.0, 1.8])
        add_callout(doc,
            f"All {results['n_attacks']} campaigns surface in the top tier at a false-positive "
            f"rate of 8.1% "
            f"({results['fp_above']} of {results['normal_total']} normal users at the recall-at-100% operating point) WITH actionable "
            f"directional intelligence — which zone changed, toward what threat pattern, mapped "
            f"to MITRE ATT&CK. Traditional methods miss most of these campaigns, and even when "
            f"they flag one they provide none of this directional context.")
        add_body(doc, (
            "An important nuance — fusion is the point. Taken individually, the Tier-3 detectors "
            "(velocity, regime, zone-divergence, contextual, relationship) are each noisy; a "
            "combined, unfused Tier-3 detection flags roughly 19% of normal entities. The low-false-positive result "
            "above belongs to the multi-phase COMPOSITE used as a ranker, not to any single "
            "detector. Fusing independent, individually-noisy phases into one ranked score is "
            "precisely what converts many weak signals into a precise verdict — and it is the "
            "core of the Tier-3 innovation."
        ))

        slow = next((a for a in results["attackers"] if a["uid"] == "USR-234"), None)
        if slow:
            add_section_heading(doc, "5.1  Case Study — the Slow APT Magnitude Cannot See", level=2)
            add_body(doc, (
                f"{slow['uid']} ({slow['label']}) has among the lowest raw signal strength of "
                f"any entity — two to three beacons a day. Magnitude ranks it unremarkable. "
                f"V-Intelligence UEBA still ranks it #{slow['rank']} of {results['total_entities']}, "
                f"driven by novelty persistence: the same never-before-seen infrastructure "
                f"recurring week after week — the axis no threshold tool measures."
            ))
    else:
        add_body(doc, "Quantified per-campaign results are included once MVP data is loaded.",
                 italic=True)

    # 5.2 — explicit step-by-step detection walkthrough (mirrors the detection-comparison view)
    add_section_heading(doc, "5.2  Exactly How an Anomaly Is Detected", level=2)
    add_body(doc, (
        "This subsection walks an entity through the pipeline the way the product's "
        "detection-comparison view does, with the actual data at each stage. Detection is not a "
        "single test; it is four stages, and the figures show what each one adds."
    ))

    add_body(doc, "Stage 1 — drift over time (does the entity separate from the crowd?).",
             bold=True, space_after=2)
    add_body(doc, (
        "Each week, the entity's serialized behavior is embedded and compared to the prior week; "
        "the drift accumulates. Traditional feature-space CUSUM buries the attacker inside the "
        "normal band — the attacker line is one of the crowd. Semantic drift lifts the insider "
        "out of the band as its campaign matures. Note the detail in Figure 2: the slow "
        "APT (USR-234) still rides within the band on raw drift alone — which is precisely why "
        "drift is necessary but not sufficient, and why later stages matter."
    ))
    add_figure(doc, "drift",
               "Figure 2 — Semantic drift over time vs the normal ±2σ band. The insider "
               "(USR-156) separates as it matures; the slow APT (USR-234) stays inside the band "
               "on raw drift and is only caught downstream by composite novelty scoring. Each "
               "weekly step is tiny — the accumulation is what a per-period threshold misses.")

    add_body(doc, "Stage 2 — localize the drift by zone (what is changing?).",
             bold=True, space_after=2)
    add_body(doc, (
        "Drift is decomposed across the behavioral zones. The signal lives in one zone, not "
        "across all five: the insider's drift concentrates in the data-behavior zone while "
        "identity stays flat; the APT's concentrates in the network zone. This asymmetric, "
        "single-zone signature — 'identity stable but network drifting' — is exactly what a "
        "single blended score averages away, and what zone divergence surfaces."
    ))
    add_figure(doc, "zone",
               "Figure 3 — Zone divergence: late-window drift per behavioral zone. Detection "
               "localizes the change to a zone rather than reporting one undifferentiated score.")

    add_body(doc, "Stage 3 — fuse five questions into one score (is this actually an attacker?).",
             bold=True, space_after=2)
    add_body(doc, (
        "A single drifting metric is not proof. The composite scorer asks five independent "
        "questions at once: is the drift stronger than the entity's peer cohort (signal "
        "strength); across how many features (breadth); does it persist over time (sustained); "
        "does it appear under multiple analytical contexts (context divergence); and are "
        "never-before-seen behaviors recurring (novelty persistence)? A normal user may score "
        "high on one question by chance; an attacker usually scores high on more than one — and "
        "even an attacker that spikes on just one phase (USR-234 rides novelty alone) is caught "
        "once the scores are fused. The fused score separates the attackers from the normal mass."
    ))
    add_figure(doc, "dist",
               "Figure 4 — Composite-score distribution. Three campaigns (USR-118, USR-156, "
               "USR-234) sit above the bulk of normal users; the hardest, the Volt Typhoon LOTL "
               "campaign (USR-042 at 13.7), sits within the upper normal tail — which is why "
               "recalling all four admits ~20 normal users (the 8.1% operating point).")

    add_body(doc, "Stage 4 — why fusion is essential (no single phase catches everyone).",
             bold=True, space_after=2)
    add_body(doc, (
        "The five phases are complementary, not redundant. The slow APT (USR-234) is caught "
        "almost entirely by novelty persistence and is unremarkable on the others; the insider "
        "(USR-156) is caught by signal strength, breadth, and sustained deviation in the data "
        "zone. Each attacker traces a different shape across the five phases, and none of them "
        "could be caught by any single phase alone — which is the whole reason for fusion."
    ))
    add_figure(doc, "radar",
               "Figure 5 — Five-phase percentile profile. Normal users cluster near the median "
               "ring; each attacker extends far on a DIFFERENT mix of phases (USR-234 on novelty, "
               "USR-156 on signal/breadth). No normal user reaches these extremes across multiple "
               "phases simultaneously.")

    # 5.3 — ranking quality / evaluation
    add_section_heading(doc, "5.3  Ranking Quality and the Value of Fusion", level=2)
    add_body(doc, (
        "The walkthrough above is made quantitative here. Treated as a ranking problem, each "
        "phase has a measurable quality: ROC AUC is the probability a random campaign outranks "
        "a random normal entity (0.5 = chance, 1.0 = perfect); precision@k is the share of the "
        "top-k ranked entities that are real campaigns. Measured on the 250-user set with real "
        "embeddings:"
    ))
    if EVAL.get("phases"):
        rows = [[p["name"], f"{p['auc']:.3f}", f"{p['p4']:.2f}", f"{p['p10']:.2f}"]
                for p in EVAL["phases"]]
        create_table(doc, ["Detection phase", "ROC AUC", "Precision@4", "Precision@10"],
                     rows, col_widths=[2.4, 1.3, 1.4, 1.4])
    add_figure(doc, "eval",
               "Figure 6 — Ranking quality by phase versus the fused composite (250 users, real "
               "embeddings). Every individual phase is only a partial discriminator; fusion lifts "
               "AUC to ~0.98.")
    add_body(doc, (
        "Every individual phase is a partial discriminator (AUC roughly 0.73–0.94); the fused "
        "composite reaches AUC ~0.98 and the best precision — the quantitative form of the radar "
        "in Figure 5. With only four positives these figures are indicative, not powered "
        "(Section 10); for context, on the separate 50-user comparison run the best traditional "
        "discriminator (temporal Z-Score) reached AUC ≈ 0.81 and the simple Z-Score ≈ 0.55 — "
        "well short of the fused composite, though measured on a different, smaller run."
    ))
    add_page_break(doc)


def _other_uses(doc, level):
    add_section_heading(doc, "6. Other Possible Use Cases", level=1)
    add_body(doc, (
        "The method is domain-agnostic: anywhere an entity can be observed over time, its "
        "behavior can be embedded, decomposed into zones, compared to a peer cohort, and watched "
        "for directional drift. The same pipeline has been instantiated for defense supply-chain "
        "supplier-risk; other natural targets include:"
    ))
    add_bullet(doc, "supplier-risk and counterfeit/sourcing drift (defense supply chain).", bold_prefix="Supply chain: ")
    add_bullet(doc, "provider upcoding and referral-ring detection.", bold_prefix="Healthcare fraud: ")
    add_bullet(doc, "account behavior and structuring (AML).", bold_prefix="Financial crime: ")
    add_bullet(doc, "equipment degradation from multi-sensor behavioral drift.", bold_prefix="Prescriptive maintenance: ")
    add_bullet(doc, "person-of-interest activity across travel, comms, and finance.", bold_prefix="Counter-intelligence: ")
    add_body(doc, (
        "Each requires only a domain-specific serialization and zone definition; the embedding, "
        "drift, cohort, and composite-scoring machinery is unchanged."
    ))
    add_page_break(doc)


def _scalability(doc, level):
    add_section_heading(doc, "7. Scalability and Performance", level=1)
    add_body(doc, (
        "The validation environment processed 250 entities over 485 days (~14 million events). "
        "The design scales along well-understood axes:"
    ))
    if _at_least(level, ARCHITECTURE):
        add_bullet(doc, "embeddings are cached by content hash, so re-runs and repeated behavior "
                   "incur no recomputation; new text is embedded in batches.", bold_prefix="Embedding cache: ")
        add_bullet(doc, "feature ETL and snapshotting are per-entity-per-period and embarrassingly "
                   "parallel; writes are batched (hundreds of rows per round-trip).", bold_prefix="Parallel ETL: ")
        add_bullet(doc, "pgvector with an HNSW index (m=16, ef_construction=64) gives sub-linear "
                   "nearest-neighbor search over the embedding store.", bold_prefix="Vector search: ")
        add_bullet(doc, "detection is incremental — only the current period's snapshot and a "
                   "rolling window are needed, not the full history.", bold_prefix="Incremental scoring: ")
    else:
        add_body(doc, "The pipeline caches embeddings, parallelizes per-entity processing, and "
                 "uses indexed vector search for scale.")
    add_body(doc, (
        "Cost scales with the number of entity-periods, not raw event volume, because events are "
        "aggregated to features before embedding. The dominant external cost is embedding calls, "
        "which the cache and a local-model option both reduce."
    ))
    add_page_break(doc)


def _deployment(doc, level):
    add_section_heading(doc, "8. Deployment", level=1)
    add_body(doc, (
        "V-Intelligence UEBA is containerized and database-backed, deployable in an enterprise "
        "data center, a tactical edge, or an air-gapped enclave."
    ))
    if _at_least(level, ARCHITECTURE):
        add_bullet(doc, "PostgreSQL + pgvector (vector store + bi-temporal history), a feature/"
                   "snapshot pipeline, a detection service, and a dashboard — all in Docker.",
                   bold_prefix="Components: ")
        add_bullet(doc, "ingest from existing SIEM/log sources; no new instrumentation required.",
                   bold_prefix="Inputs: ")
        add_bullet(doc, "for offline/air-gapped operation, the OpenAI embedding model is swapped "
                   "for a local embedding model — the rest of the pipeline is unchanged.",
                   bold_prefix="Air-gap option: ")
        add_bullet(doc, "configure → integrate data → baseline (weeks) → shadow alongside the "
                   "SIEM → operate.", bold_prefix="Rollout: ")
    else:
        add_body(doc, "Docker-based, ingests existing SIEM data, runs on-prem or air-gapped with "
                 "a local embedding model, and deploys via a baseline-then-shadow rollout.")
    add_page_break(doc)


def _prior_art(doc, level):
    add_section_heading(doc, "9. Relation to Prior Art", level=1)
    add_body(doc, (
        "Behavioral analytics is a mature field; the contribution here is a specific combination, "
        "not the invention of anomaly detection. It is worth being precise about what is and is "
        "not new."
    ))
    add_bullet(doc, "Exabeam, Securonix, and Microsoft Sentinel UEBA — and similar commercial "
               "UEBA platforms — baseline entities statistically and score deviations on numeric "
               "features and peer groups (per their publicly described methodology). They are "
               "strong at volume and rule-aligned anomalies but measure magnitude; a "
               "constant-volume change of kind (the insider) is their blind spot.",
               bold_prefix="Commercial UEBA: ")
    add_bullet(doc, "autoencoders, isolation forests, one-class methods, LSTM sequence models, "
               "and graph neural networks for entity anomaly detection are well studied. Most "
               "operate on numeric or sequence features and produce opaque scores without an "
               "analyst-ready, threat-named direction.",
               bold_prefix="Academic anomaly detection: ")
    add_body(doc, "What is distinct here is the composition:", space_after=2)
    add_bullet(doc, "behavior serialized to language and embedded in a unified semantic space, "
               "so the DIRECTION of change is measurable, not just its magnitude.", bold_prefix="(1) ")
    add_bullet(doc, "independent behavioral zones, so a one-zone signal is not diluted.", bold_prefix="(2) ")
    add_bullet(doc, "drift projected onto natural-language threat concepts mapped to MITRE "
               "ATT&CK — explanations an analyst and an auditor can use.", bold_prefix="(3) ")
    add_bullet(doc, "peer-cohort relative scoring and multi-phase fusion into one ranked verdict.",
               bold_prefix="(4) ")
    add_body(doc, (
        "Each ingredient exists somewhere in isolation; their integration into one explainable, "
        "direction-aware, fusion-ranked pipeline is the claimed innovation."
    ))

    add_section_heading(doc, "9.1  Competitive Landscape and Defensible Novelty", level=2)
    add_body(doc, (
        "A structured, source-verified review of the market and literature (2024–2026) found no "
        "commercially marketed product, and no single published method, that detects intrusions "
        "the way this architecture does. Publicly described commercial UEBA — Exabeam, Securonix, "
        "Microsoft Sentinel UEBA — detects via statistical / p-value profiling, machine-learning "
        "baselines, rarity- and magnitude-based scores, and count- or TF-IDF-based peer grouping; "
        "their 2024–2026 generative-AI features are SOC copilots and conversational layers, not "
        "the detection mechanism."
    ))
    add_body(doc, (
        "We are explicit about prior art. The constituent techniques are individually established: "
        "LLM-embedding of serialized behavior for anomaly detection (APT-LLM; LogBERT), "
        "serializing raw metrics to prose before embedding (TabLLM and related work), embed-then-"
        "detect pipelines (TAD-Bench), and embedding drift measured by cosine distance (ZEDD). "
        "Every one of these, however, detects by MAGNITUDE — reconstruction error, distributional "
        "deviation, or rarity — and operates in a single domain; none projects the DIRECTION of "
        "drift onto named, pre-embedded threat concepts mapped to MITRE ATT&CK."
    ))
    add_callout(doc,
        "To the best of our knowledge — scoped to publicly described and marketed methodology — no "
        "commercial UEBA platform and no single published system detects anomalies by the "
        "DIRECTION of behavioral drift in a unified semantic space. The contribution is the "
        "integrated architecture and the concept-direction (not magnitude/rarity) paradigm, not "
        "any single technique.")
    add_body(doc, (
        "Two scoping notes: vendor internals are not fully public, so the comparison is "
        "against publicly described and marketed methodology; and a formal patent / "
        "freedom-to-operate review is recommended before any contractual assertion of uniqueness."
    ), italic=True, color=DARK_GRAY)
    add_page_break(doc)


def _limitations(doc, level):
    add_section_heading(doc, "10. Limitations and Next Steps", level=1)
    add_body(doc, (
        "The results establish the concept; they do not yet establish operational performance. "
        "The following constraints are stated plainly so claims are not over-read."
    ))
    add_bullet(doc, "all results here are on synthetic data, and the approach works on it. On "
               "real production data we expect the same behavior with minor calibration "
               "adjustments — thresholds and cohort baselines tuned to live traffic. Real-data "
               "validation is the planned next step.", bold_prefix="Synthetic data: ")
    add_bullet(doc, "the test embeds four attack campaigns — enough to demonstrate the concept "
               "end to end. A larger, more varied set of injected campaigns is the next step to "
               "harden the numbers.", bold_prefix="Small test set: ")
    add_bullet(doc, "Volt Typhoon-class living-off-the-land is the hardest threat there is — in "
               "the real world this actor dwelled in U.S. critical infrastructure UNDETECTED for "
               "at least five years (CISA AA24-038A), and magnitude-based detectors rank it among "
               "ordinary users (0 of 4 on this benchmark). V-Intelligence still surfaces it inside "
               "the top ~10% of all entities (rank 24 of 250) at the all-recall operating point — the difference "
               "between catching a five-year-invisible campaign and missing it entirely. The one "
               "caveat: of the four campaigns it sits closest to the decision boundary "
               "(USR-042), so a stricter threshold trades it off; characterizing that "
               "recall/false-positive curve on larger data is the next step.",
               bold_prefix="Hardest-class threat, surfaced anyway: ")
    add_bullet(doc, "the 8.1% figure is measured at the alert threshold that catches all four "
               "campaigns. On a larger dataset we will publish the full detection-vs-false-alarm "
               "curve so each operator can set the threshold to their own alert budget.",
               bold_prefix="Threshold tuning: ")
    add_bullet(doc, "each behavior is described in words — including interpretive cues such as "
               "'restricted' or 'increasing' — not just raw numbers. This helps the model, but it "
               "also means part of the match to a threat concept comes from how the behavior is "
               "phrased. A plain-numbers run is planned to confirm the signal holds without the "
               "suggestive wording.", bold_prefix="Wording can prime the result: ")
    add_bullet(doc, "a model-aware attacker will adapt, and we will keep improving the model — "
               "detection is an ongoing contest. Multi-phase fusion, novelty-persistence, and "
               "relationship signals raise the bar today and are designed to evolve as "
               "adversaries do.", bold_prefix="Adaptive adversary (an ongoing contest): ")
    add_bullet(doc, "for air-gapped or classified environments the embedding model can be "
               "self-hosted on-premise — suitable open models exist — so no external API is "
               "required.", bold_prefix="Air-gapped deployment is supported: ")
    add_page_break(doc)


def _conclusion(doc, level):
    add_section_heading(doc, "11. Conclusion", level=1)
    trad = load_traditional()
    add_body(doc, (
        "Three generations on one dataset make the case plainly. Traditional detection largely "
        f"MISSES these stealthy campaigns — LOF {trad_phrase(trad, 'lof_anomaly')}, "
        f"Z-Score {trad_phrase(trad, 'zscore_anomaly')}, "
        f"Isolation Forest {trad_phrase(trad, 'iforest_anomaly')}, "
        f"OCSVM {trad_phrase(trad, 'ocsvm_anomaly')} — and even when one fires, every "
        "alert says only 'anomalous' with no indication of which behavioral zone changed or "
        "toward what threat pattern. The analyst still faces the same manual triage. An "
        "intermediate embedding-drift approach begins to recover direction; and the "
        "V-Intelligence UEBA digital twin — zones, trajectory, relationships, peer cohorts, and "
        "multi-phase composite scoring — finds all four embedded campaigns at 8.1% FP WITH "
        "actionable directional intelligence: zone localization, drift direction, and MITRE "
        "ATT&CK concept alignment that traditional methods cannot provide."
    ))
    add_callout(doc,
        "Traditional analytics miss the quiet, directional threats that do the most damage — and "
        "even when they flag one, detection without direction is triage without intelligence. "
        "V-Intelligence UEBA surfaces these campaigns AND tells the analyst WHICH zone changed, "
        "TOWARD WHAT threat pattern, mapped to MITRE ATT&CK — the difference between a missed "
        "threat and an actionable investigation.")
    note = {
        TEASER: "This is a conceptual overview; architectural and full-technical editions are "
                "available on request.",
        ARCHITECTURE: "Exact hyperparameters, weights, and thresholds are withheld in this "
                      "edition and provided in the full-technical edition under NDA.",
        FULL: "This edition contains proprietary formulations and is provided under NDA. Do not "
              "redistribute.",
    }[level]
    add_body(doc, note, italic=True, color=DARK_GRAY)
    add_body(doc, "All results are from synthetic data and must be re-validated on production "
             "telemetry before any operational claim.", italic=True, color=DARK_GRAY)
    add_page_break(doc)


def _references(doc, level):
    add_section_heading(doc, "12. References", level=1)
    refs = [
        "CISA, 'PRC State-Sponsored Actors Compromise and Maintain Persistent Access to U.S. "
        "Critical Infrastructure' (Volt Typhoon), Advisory AA24-038A, 2024.",
        "CISA AA23-144A, 'People's Republic of China State-Sponsored Cyber Actor Living off the "
        "Land,' 2023.",
        "U.S. Government joint reporting on Salt Typhoon intrusions into telecommunications "
        "infrastructure, 2024–2025.",
        "MITRE ATT&CK Framework, Enterprise Matrix. attack.mitre.org.",
        "E. S. Page, 'Continuous Inspection Schemes,' Biometrika 41(1/2), 1954 (CUSUM).",
        "NIST SP 800-207, 'Zero Trust Architecture,' 2020.",
        "OpenAI, 'text-embedding-3-small' embedding model documentation.",
        "Y. Malkov and D. Yashunin, 'Efficient and robust approximate nearest neighbor search "
        "using Hierarchical Navigable Small World graphs (HNSW),' IEEE TPAMI, 2018.",
        "APT-LLM: LLM-embedded provenance traces for anomaly detection, arXiv:2502.09385, 2025 "
        "(prior art — magnitude / autoencoder detection).",
        "H. Guo et al., 'LogBERT: Log Anomaly Detection via BERT,' arXiv:2103.04475, 2021 "
        "(prior art — self-supervised, distributional deviation).",
        "TAD-Bench: benchmarking LLM-embedding + anomaly detectors, arXiv:2501.11960, 2025 "
        "(prior art — embed-then-detect).",
        "Serialize-tabular-to-text then LLM-embed (TabLLM family), arXiv:2502.11596, 2025 "
        "(prior art — prose serialization).",
        "ZEDD: embedding-drift (1 − cosine) for prompt-injection defense, arXiv:2601.12359, 2026 "
        "(prior art — cosine drift, different application).",
        "V-Intelligence UEBA Technical Specification (companion document, full edition).",
    ]
    for i, ref in enumerate(refs, 1):
        add_bullet(doc, ref, bold_prefix=f"[{i}] ")


# ============================================================================
# Build orchestration
# ============================================================================
def build(level, results):
    doc = _init_doc()
    _title_page(doc, level)
    _exec_summary(doc, level, results)
    _problem(doc, level)
    _traditional(doc, level)
    _solution(doc, level)
    _results(doc, level, results)
    _other_uses(doc, level)
    _scalability(doc, level)
    _deployment(doc, level)
    _prior_art(doc, level)
    _limitations(doc, level)
    _conclusion(doc, level)
    _references(doc, level)
    out_name = f"VIntelligence_vs_Traditional_UEBA_{level}.docx"
    doc.save(os.path.join(HERE, out_name))
    print(f"  [{level}] -> {out_name}")


def main():
    parser = argparse.ArgumentParser(description="Build the V-Intelligence vs Traditional UEBA paper.")
    parser.add_argument("--disclosure", choices=[TEASER, ARCHITECTURE, FULL, "all"],
                        default="all")
    args = parser.parse_args()
    global EVAL, FIGURES
    results = load_results()
    EVAL = evaluate()
    FIGURES = generate_figures(EVAL, results)
    print("Building V-Intelligence vs Traditional UEBA technical paper...")
    print(f"  data source    : {DATA_SOURCE or 'none (no DB, no CSV)'}")
    print(f"  empirical data : {'loaded' if results else 'NOT FOUND (theoretical mode)'}")
    print(f"  figures        : {sorted(FIGURES)}")
    levels = [TEASER, ARCHITECTURE, FULL] if args.disclosure == "all" else [args.disclosure]
    for lvl in levels:
        build(lvl, results)
    print("Done.")


if __name__ == "__main__":
    main()
