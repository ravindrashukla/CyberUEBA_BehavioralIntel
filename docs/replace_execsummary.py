"""Replace the Executive Summary in DLA Final.docx with the merged version.
Keeps the 'Executive Summary' Heading 1, preserves Figure 1, bold sub-headers,
List Bullet callouts. Clean replacement (no track changes)."""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = "WP DLA/DLA Final.docx"
d = Document(PATH)

def accp(p):
    return "".join(t.text or "" for t in p._p.iter(qn('w:t')))

hi = next(i for i, p in enumerate(d.paragraphs) if accp(p).strip() == "Executive Summary")
body = []
j = hi + 1
while j < len(d.paragraphs):
    p = d.paragraphs[j]
    st = p.style.name if p.style else ""
    if st.startswith("Heading") and accp(p).strip():
        break
    body.append(p)
    j += 1
anchor = body[0]

fig_run = None
for p in body:
    for r in p._p.findall(qn('w:r')):
        if r.findall('.//' + qn('w:drawing')):
            fig_run = r
            break
    if fig_run is not None:
        break
print("figure run found:", fig_run is not None, "| body paras:", len(body))

def para(style='Normal'):
    return anchor.insert_paragraph_before('', style=style)

def body_p(t):
    p = para(); p.add_run(t); return p

def header(t):
    p = para(); r = p.add_run(t); r.bold = True; return p

def bullet(label, rest):
    p = para('List Bullet'); rl = p.add_run(label); rl.bold = True
    if rest:
        p.add_run(rest)
    return p

def note(t):
    p = para(); r = p.add_run(t); r.italic = True; return p

def figure():
    p = para(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fig_run is not None:
        p._p.append(copy.deepcopy(fig_run))
    return p

CH1 = ("Today, critical operational intelligence lives in silos. Information about a part, supplier, depot, route, or "
       "user is scattered across disparate tables and structures, each holding only a fragment. Because legacy "
       "mathematical models confine analysis to individual attributes, one field at a time, decisions are consistently "
       "made on isolated slices of data rather than the complete, comprehensive reality of the entity. We rarely see "
       "the full picture.")
CH2 = ("For the Defense Logistics Agency, that gap has a mission cost. DLA operates in a contested logistics "
       "environment where adversarial exploitation and disruption can emerge through demand volatility, supplier "
       "fragility, sub-tier dependency, counterfeit exposure, sourcing shifts, route vulnerability, cyber-enabled "
       "disruption, or correlated movement across the defense industrial base. In this environment, the mission risk "
       "is not simply a late shipment, an inaccurate forecast, or a delayed supplier alert. The mission risk is the "
       "loss of decision time, the interval between the first observable behavioral signal and the point at which "
       "readiness, mission assurance, or operational resiliency is affected. Current planning, forecasting, and "
       "supplier-monitoring methods remain necessary, but they can leave DLA reacting after thresholds break rather "
       "than seeing where demand, supplier behavior, inventory posture, or network exposure is moving. The "
       "requirement is a stronger decision-intelligence layer that preserves decision time and supports decision "
       "advantage before operational consequences are fully visible.")
IN1 = ("Our innovation, Entity Digital Model (EDM) Vector Intelligence and Behavioral Entity Intelligence (BEI), "
       "collectively EDM / BEI, solves this fragmentation by representing every entity, and its network of "
       "connections, as a single, holistic behavioral digital twin.")
IN2 = ("The breakthrough is taking the core mathematical engine behind Large Language Models, the Transformer "
       "architecture's self-attention mechanism, out of the text domain and applying it to operational "
       "decision-making. Just as self-attention weighs the words in a sentence against one another to derive context, "
       "EDM / BEI weighs every element of an entity's operational data against every other element. Defining the "
       "entity as a digital twin is what captures that mechanism. Each item-location, supplier, depot, route, "
       "relationship, and network becomes a rich, high-dimensional behavioral representation within a shared digital "
       "space; its movement is measured over time, compared against its own history and peer cohorts, aligned to "
       "known supply-chain concepts, and translated into explainable planning or risk signals. Figure 1 illustrates "
       "this shift from backward-looking metrics to forward-looking, mission-relevant decision intelligence.")
IN3 = ("That digital foundation is the launching pad. By looking at the whole entity rather than individual "
       "attributes, EDM / BEI unlocks a powerful family of algorithms, including forecasting, risk classification, "
       "survival, anomaly, and network analyses, all running over the same unified representation. It right-sizes "
       "procurement where conventional methods over-order and catches demand they miss entirely, and it provides "
       "capabilities legacy methods cannot, such as calibrated uncertainty, early warning from behavioral drift, and "
       "forecasts for items with little or no history.")
POC = ("To validate the versatility of this foundation, we tested EDM / BEI across three distinct mission domains "
       "using synthetic data. The same behavioral signals carried across all three, which is strong evidence that "
       "this is a generalized capability, not a narrowly tuned point solution. Early results are highly encouraging:")
MV = ("For the mission, this means tighter, better-justified procurement, earlier warning that protects readiness, "
      "and faster, explainable decisions, all delivered from one reusable foundation rather than a new tool built for "
      "each problem. The detailed evidence is presented in Section 2, “Proof of Concept: Decision-Value "
      "Evidence.”")
PILOT = ("The recommended next step is a bounded pilot that runs current DLA baselines and EDM / BEI outputs side by "
         "side across selected NSNs, suppliers, depots, and mission-relevant scenarios. Use Case 1, Demand "
         "Forecasting and Inventory Planning, is the prioritized starting case, because it provides the clearest path "
         "to validate forecast quality, uncertainty management, procurement-plan accuracy, early warning, and planner "
         "usability against existing DLA data. From there, the same operating model extends to supplier-risk and "
         "logistics-network early-warning use cases, in support of broader mission assurance, operational resiliency, "
         "supply-chain risk mitigation, and decision advantage.")
ROAD = ("Critically, the pilot builds the data foundation and the behavioral digital-twin representation once. On that "
        "same foundation, the platform is designed to address a mapped roadmap of 17 identified use cases across four "
        "horizons (demand, supply and lead-time risk, network and resilience, and portfolio intelligence), each "
        "reusing the data, embeddings, and features already built, so every wave is faster than the one before. "
        "Beyond the demand pilot, these include cold-start forecasting, stockout prediction, lead-time estimation, "
        "supplier-capability discovery, route-bottleneck identification, cascading-failure and combined-shock "
        "analysis, drift early-warning alerts, item-health watchlists, and semantic substitute search. The path "
        "forward is to map the attributes that matter for each new scenario, simulate the corresponding digital "
        "twins, and demonstrate how widely the foundation scales.")
NOTE = ("Note: This initial assessment was conducted using synthetic data. To ensure mission readiness, every result "
        "and algorithm will be rigorously re-validated on operational Defense Logistics Agency (DLA) data prior to "
        "production use.")

header("The Challenge: Fragmented Data, and the Loss of Decision Time")
body_p(CH1)
body_p(CH2)
header("The Innovation: Bringing Self-Attention to the Decision Domain")
body_p(IN1)
body_p(IN2)
figure()
body_p(IN3)
header("What Makes EDM / BEI Unique")
bullet("Unified entity space:", " every entity type, from parts and suppliers to depots and users, is represented in "
       "one shared space, so cross-entity relationships become measurable rather than implied.")
bullet("Direction over threshold:", " it measures where behavior is heading, surfacing slow, sub-threshold change "
       "that magnitude-based methods miss by design.")
bullet("One foundation, many domains:", " the same representation and behavioral signals carry across logistics and "
       "cyber, making this a reusable platform rather than a single-purpose tool.")
header("Proof of Concept: Generalizable Impact Across Domains")
body_p(POC)
bullet("Demand Forecasting:", " EDM Vector Intelligence reduced the eight-week procurement buy by 40 to 54 percent on "
       "worked items while fully covering realized demand (for example, right-sizing a 379-unit adapter plate buy to "
       "225 units, preventing roughly $460,000 in over-procurement). It also produced forecasts for items that legacy "
       "methods report as zero, catching demand they miss entirely.")
bullet("Supplier Risk:", " Flags higher-risk suppliers and estimates time-to-failure, with explainable behavioral "
       "drivers.")
bullet("Network Early Warning (Cyber):", " Detected all four slow, stealthy campaigns (4 of 4) at an 8.1 percent "
       "false-positive point, where traditional threshold methods caught zero to one. These campaigns were modeled on "
       "the tradecraft of living-off-the-land, state-sponsored actors such as Volt Typhoon and Salt Typhoon, which "
       "operate within normal activity thresholds, the same stealth that let them pre-position in U.S. critical "
       "infrastructure and compromise telecommunications networks while evading signature- and threshold-based "
       "detection. Surfacing this class of slow, sub-threshold movement, before any threshold is breached, is "
       "precisely the gap EDM / BEI closes.")
body_p(MV)
header("The Recommended Next Step and Path Forward")
body_p(PILOT)
body_p(ROAD)
note(NOTE)

for p in body:
    p._p.getparent().remove(p._p)

d.save(PATH)
print("Executive Summary replaced in DLA Final.docx")
