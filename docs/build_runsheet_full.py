# -*- coding: utf-8 -*-
"""Brandon Pugh session PLAYBOOK (INTERNAL) - full detail: how to run the session.
Every block: goal, time, show, verbatim SAY track, likely pushback, transition.
Output overwrites Brandon_Pugh_Working_Session_Run-Sheet_INTERNAL.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6); GRAY=RGBColor(0x55,0x5F,0x6B)
RED=RGBColor(0xA6,0x1B,0x1B); GREEN=RGBColor(0x1F,0x7A,0x33); AMBER=RGBColor(0xB5,0x5A,0x00)
SAYBG=RGBColor(0x10,0x40,0x20)
d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(10.5)

def H1(t):
    p=d.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
    return p
def H2(t):
    p=d.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=BLUE
    return p
def P(t,italic=False,bold=False,color=None,size=None):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold
    if color: r.font.color.rgb=color
    if size: r.font.size=Pt(size)
    return p
def B(label,rest="",sub=False):
    p=d.add_paragraph(style='List Bullet 2' if sub else 'List Bullet')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def NUM(label,rest=""):
    p=d.add_paragraph(style='List Number')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def callout(text,color=RED,size=9.5):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2)
    r=p.add_run(text); r.bold=True; r.font.color.rgb=color; r.font.size=Pt(size); return p
def meta(label,val):
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(label+"  "); r.bold=True; r.font.color.rgb=GRAY; r.font.size=Pt(9.5)
    r2=p.add_run(val); r2.font.size=Pt(9.5); return p
def say(lines):
    """Verbatim talking track block, shaded-feel via indent + label."""
    lp=d.add_paragraph(); lp.paragraph_format.space_after=Pt(1)
    lr=lp.add_run("SAY (talking track):"); lr.bold=True; lr.font.color.rgb=GREEN; lr.font.size=Pt(9.5)
    for ln in lines:
        p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.25); p.paragraph_format.space_after=Pt(3)
        r=p.add_run("“"+ln+"”"); r.italic=True; r.font.size=Pt(10.5)
def table(headers,rows,widths=None,fs=8.5,hdr_fs=8.5):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(hdr_fs)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""
            for j,part in enumerate(v.split("||")):
                para=cells[i].paragraphs[0] if j==0 else cells[i].add_paragraph()
                para.add_run(part).font.size=Pt(fs)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    d.add_paragraph(); return t
def rule():
    p=d.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    r=p.add_run("_"*68); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC); r.font.size=Pt(8)

# ============ TITLE ============
tp=d.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("How to Run the Session - Playbook"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
sp=d.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Brandon Pugh Technical Deep-Dive"); r.font.size=Pt(13); r.font.color.rgb=BLUE
sp2=d.add_paragraph(); sp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp2.add_run("22nd Century  |  V-Intelligence Four-Layer Critical-Infrastructure Defense"); r.font.size=Pt(10); r.italic=True; r.font.color.rgb=GRAY
wn=d.add_paragraph(); wn.alignment=WD_ALIGN_PARAGRAPH.CENTER
rw=wn.add_run("INTERNAL playbook - not the leave-behind. SAY tracks are scripts to adapt, not read robotically."); rw.bold=True; rw.font.color.rgb=RED; rw.font.size=Pt(9.5)

# ============ 1 FRAME ============
H1("1. Frame (read before you walk in)")
B("Audience: ","Brandon Pugh, cybersecurity lead at R Street, advisor on Pentagon / critical-infrastructure cyber. Technical enough to go deep; his lens is policy and national defense. Rewards candor; discounts anything tuned or overclaimed.")
B("Format: ","2.5 to 3 hour technical working session. He has the layered whitepaper.")
B("The one story: ","an integrated four-layer defense-in-depth architecture under one MITRE vocabulary, covering the full attack lifecycle from code to network edge to behavior to API/data runtime. The behavioral layer (V-Intelligence) is demoed live; the others are presented as architecture and capability.")
B("Three objectives: ","(1) establish the four-layer architecture closes the Volt/Salt gap end to end; (2) prove the behavioral layer live and honestly; (3) secure a bounded Phase-1 pilot plus a policy / thought-leadership relationship.")
B("Golden rule: ","national problem, then architecture, then proof. Show what runs (Layer 2). Label what is integrated or in development, out loud. Candor is the advantage, not a weakness.")
callout("Internal guardrail, never blur in the room: only Layer 2 is 22CT-built and demoable live today. Layers 1, 3, 4 are "
 "architecture and integrated/in-development capability, NOT live demos. The four layers share a MITRE vocabulary and a "
 "vision; they are not one turnkey running platform today, and full cross-layer fusion is the roadmap.")

# ============ 2 ARCHITECTURE ============
H1("2. The architecture at a glance")
table(["Layer","Secures","Incumbent today","Core shortcoming","22CT wedge","Maturity"],
 [["1 - Preemptive Network Assurance","Firewall/NGFW, IdP, IPS, SASE, WAF config","Legacy SIEM; breach-sim & policy tools; vuln mgmt / pen test","Reactive or samples ~2^8000 config space; vuln mgmt covers ~20%; can't prove segmentation","Formal math proof over full state space; residual-capacity threat blocking","Integrated"],
  ["2 - Behavioral Entity Intelligence (UEBA)","Valid-account, living-off-the-land behavior inside","Traditional UEBA; point-anomaly ML (iForest, OC-SVM, LOF, z-score)","Threshold-based; magnitude not meaning; per-source silos; '87% anomalous', no direction","Meaning over magnitude: 1536-d embeddings, CUSUM trajectory, MITRE-direction, multi-front composite","LIVE demo (22CT-built)"],
  ["3 - Application, API & Data Runtime","APIs, microservices, data, AI-agent paths","Traditional WAFs; standalone API-sec/ASPM; gateways; bot tools","Miss business-logic abuse; signatures miss authenticated abuse (95% of API attacks); agentic AI ungoverned","Unified discover+comply+protect on behavioral/identity core; runtime-native mitigation; agentic-AI governance","Integrated"],
  ["4 - Code & Supply-Chain","Code, deps, containers, IaC, CI/CD","SAST/DAST; SCA; container/cloud & secrets scanners","Known-CVE matching; can't find novel zero-days; reactive to disclosure; alert fatigue","AI-driven novel zero-day discovery pre-deploy; multi-gate CI/CD","In development"]],
 widths=[1.05,0.95,1.2,1.55,1.55,0.7],fs=7.5,hdr_fs=8)

# ============ 3 ROLES & PRE-FLIGHT ============
H1("3. Roles and pre-flight")
table(["Role","Owner","Job"],
 [["Lead / narrator","You (Ravindra)","Architecture, Layers 1/3/4 framing, the ask, the honesty line"],
  ["Demo driver","(assign)","Layer 2 live demo; files open and tested in advance"],
  ["Scribe","(assign)","His questions and pilot-scope asks to a parking lot"]],
 widths=[1.3,1.7,3.9],fs=9.5)
P("Demo machine (Layer 2), open and tested BEFORE he arrives:",bold=True)
B("",("simulator/attacks/volt_typhoon.py | data/generated/ (USR-042 stream) | "
   "data/tier3_results/weekly_zone_trajectories.csv + composite_scores.csv | "
   "detection/ (cusum.py, drift_direction.py, reference_concepts.py, composite_scorer.py) | "
   "comparison/run_comparison.py | UI up on host 8001."))
B("Dry-run once: ","derive the 8.1% live from composite_scores.csv (sort by composite; count normals above Volt's 13.70). If it stalls, you have the briefing-book backup slide.")

# ============ 4 AGENDA ============
H1("4. Agenda (target 3h incl. buffer)")
table(["#","Block","Time","Driver"],
 [["0","Open: his goals, set the arc","5 min","Lead"],
  ["1","Problem + 4-layer architecture","25 min","Lead"],
  ["2","Layer 1 - Preemptive Network Assurance","20 min","Lead"],
  ["3","Layer 2 - Behavioral Entity Intelligence (LIVE)","45 min","Demo"],
  ["4","Layer 3 - Application, API & Data Runtime","20 min","Lead"],
  ["5","Layer 4 - Code & Supply-Chain Assurance","15 min","Lead"],
  ["6","How the layers work together + federal fit","20 min","Both"],
  ["7","Roadmap + the ask","15 min","Lead"],
  ["-","Buffer for his tangents","15 min","-"]],
 widths=[0.4,4.0,0.8,0.9],fs=9)

# ============ 5 BLOCK BY BLOCK (FULL) ============
H1("5. Block-by-block - how to run each segment")

# --- BLOCK 0 ---
H2("Block 0 - Open (5 min, Lead)")
meta("Goal:","Get his agenda, set the arc, set the honesty frame in the first two minutes.")
meta("Show:","Nothing. Talk.")
say(["Brandon, thanks. Before I drive, what do you most want out of the next three hours, and what's your hard stop?",
 "Here's the arc I'd propose. I'll frame the national problem, then walk our four-layer architecture, then we slow down "
 "and I show you the behavioral layer running live against four nation-state campaigns. I'll be explicit the whole way "
 "about what's built and runs today versus what's architecture and what's roadmap. If I ever blur that, call me on it."])
meta("Transition:","“Let me start with why the architecture looks the way it does.”")

# --- BLOCK 1 ---
H2("Block 1 - Problem + 4-layer architecture (25 min, Lead)")
meta("Goal:","Land the two structural gaps (prevention and detection both fail) and that four layers under one vocabulary close them.")
meta("Show:","The defense-in-depth diagram (the 4-layer figure).")
say(["Two events reset the threat model. Volt Typhoon, CISA AA24-038A: pre-positioning inside U.S. critical "
 "infrastructure, legitimate Windows tooling only, no malware, more than five years of dwell. Salt Typhoon, "
 "AA25-239A: telecom espionage at carrier scale, slow, sub-threshold, valid accounts. The common technique is "
 "MITRE T1078, valid accounts. They log in. They don't break in.",
 "That breaks defense two ways at once. Prevention fails because the tools sample the configuration space instead "
 "of proving it. Detection fails because the tools threshold on magnitude, and a nation-state living off the land "
 "never crosses a threshold.",
 "So we built four layers under one MITRE vocabulary across the whole lifecycle. Layer 4 keeps the exploitable "
 "code out. Layer 1 proves the network edge is actually closed. Layer 2, the one I'll show you live, catches the "
 "valid-account behavior once they're inside, which is the gap Layer 1 can't see. Layer 3 catches the API and data "
 "exfiltration and the agent abuse at runtime. Each layer catches what the others miss.",
 "One thing up front: the behavioral layer is ours and you'll see it run. The other layers are architecture and "
 "integrated capability. Full cross-layer fusion is what we'd build together, and I'll flag every seam."])
meta("Likely pushback:","“Isn't this just defense-in-depth rebranded?”  ->  “The novelty isn't having layers; "
 "it's one MITRE vocabulary across all four plus a behavioral/identity core shared between Layer 2 and Layer 3, so the "
 "layers reinforce instead of sit in silos.”")
meta("Transition:","“Start at the edge. Layer 1.”")

# --- BLOCK 2 ---
H2("Block 2 - Layer 1: Preemptive Network Assurance (20 min, Lead)")
meta("Goal:","Prevention done as proof, not sampling. Cite evidence. Do NOT call it a live demo.")
meta("Show:","Layer 1 evidence slide(s).")
say(["Layer 1's job is to prove every known attack path through the network controls is closed before traffic flows.",
 "Today's incumbents can't. Legacy SIEM is reactive, it fires after. Breach-simulation and policy tools only sample "
 "the configuration space, which is astronomically large, so they're resource-bound and never complete. Vulnerability "
 "management covers roughly 20 percent of attacks and perfect patching isn't practical. And none of them can prove "
 "segmentation is correct, so a shadowed rule or a cross-VLAN path gets exploited before anything alarms. The "
 "industry numbers are stark: 99 percent of firewall breaches are misconfiguration, and more than 60 percent of "
 "incidents through 2029 trace to misconfigured controls.",
 "Our wedge is a formal mathematical model of the entire configuration, exhaustive and deterministic and "
 "verifiable, with MITRE-enriched attack models queried against it. It uses the residual capacity in the controls "
 "to add threat-blocking rules from live intel, and it catches drift within the hour.",
 "Concretely, on real assessments this found a credential-theft path from a shadowed rule, SMB lateral movement "
 "from a three-rule group shadow, a multi-vendor update-traffic path with no inspection, and it proved about a "
 "quarter of internet-facing rules had weak or no inspection. It compressed rule sets 45 percent and produced "
 "verified attack-surface compliance for a national agency."])
meta("Maturity line:","“This layer is integrated capability. I'm showing you the evidence, not running it live today.”")
meta("Likely pushback:","“Isn't this firewall policy analysis?”  ->  “Policy tools lint rules heuristically. "
 "This proves reachability and segmentation across the full state space and against MITRE attacks. Proof, not linting.”")
meta("Transition:","“Layer 1 proves the door is shut. But Volt and Salt already have a valid key. That's Layer 2, and this one I run live.”")

# --- BLOCK 3 ---
H2("Block 3 - Layer 2: Behavioral Entity Intelligence - LIVE (45 min, Demo driver)")
meta("Goal:","Prove the mechanism live and honestly. This is the centerpiece. Slow down.")
meta("Show:","Live: the attack script, the trajectory CSVs, the composite scores, the live 8.1% derivation.")
P("Run it in this order:",bold=True)
NUM("The setup. ","“Four nation-state-style attackers are hidden among 246 normal users, 250 entities over 485 days, "
 "about 14 million events across five log sources. Full ground truth. This is a blind, gradeable test.”")
NUM("The digital-twin scorecard. ","“We model each entity as a behavioral digital twin. Of the eight properties that "
 "defines, five are built today, multi-signal, semantic embedding, trajectory, runtime explainability, partial deploy; "
 "and three, identity fusion, a versioned store, and the graph, are partial or next. I'm telling you that up front.”")
NUM("Drift / CUSUM. ","Show a healthy flat trajectory against an attacker's accumulating drift. “CUSUM is "
 "minimax-optimal change detection. It turns years of dwell into days. It's the broad net; the composite decides who matters.”")
NUM("Direction / MITRE. ","“We don't say '87 percent anomalous.' We say the entity is drifting toward "
 "insider_threat_slow with alignment 0.68, and we name the MITRE technique. Direction, not just magnitude. "
 "Fourteen threat concepts and two benign anchors.”")
NUM("Composite. ","“Five behavioral fronts, signal strength, breadth, sustained, context divergence, and novelty "
 "persistence, fuse into one ranked, explainable verdict.”")
NUM("The proof, and get ahead of it. ","Walk Salt at #1 (51.27) and Volt at #24 (13.70). “Be straight with you: our "
 "top normal user outscores two of our four attackers on raw score. That's exactly why we deliver a ranked list, not a "
 "threshold. The honest claim is recall at the top of the ranking.”")
NUM("The live 8.1%. ","“Sort everyone by composite. To catch the stealthiest, Volt at 13.70, the line sits there; "
 "20 of 246 normals score above it. So 8.1 percent false positive, and all four caught. On the same data, Isolation "
 "Forest, One-Class SVM and LOF each catch zero of four; a z-score method catches one.”")
meta("Maturity line:","“This is the one layer that runs live and is 22CT-built. Identity fusion, tamper-evident forensics, "
 "and automated response are the next build phase, not running today.”")
meta("Likely pushback:","“How is this different from Exabeam or Securonix?”  ->  “Meaning versus counts, direction "
 "versus magnitude, identity fusion, and a measured zero-of-four versus four-of-four on identical data, delivered as a "
 "ranked, MITRE-mapped, explainable verdict.”")
meta("Transition:","“Layer 2 catches the human and machine behavior. But the data and the APIs are where Salt actually "
 "collects and exfiltrates. That's Layer 3.”")

# --- BLOCK 4 ---
H2("Block 4 - Layer 3: Application, API & Data Runtime (20 min, Lead)")
meta("Goal:","Cover the runtime exfil altitude and the new agentic-AI blind spot. Architecture, not a 22CT demo.")
meta("Show:","Layer 3 architecture slide.")
say(["Layer 3 is where Salt-style data collection and exfiltration actually happen: the APIs, the microservices, the "
 "data, and now the AI-agent pathways.",
 "Incumbents struggle here. Traditional WAFs don't understand API nuance and miss business-logic abuse. Signature "
 "detection misses authenticated abuse, and 95 percent of API attacks come from authenticated sources. Stitching "
 "together separate discovery, WAF, bot, and testing tools creates blind spots and overhead. And the new blind spot "
 "is agentic AI: agents talk over APIs, roughly 99 percent of AI-related vulnerabilities are API-related, and "
 "over-permissioned or rogue agent connectors are ungoverned.",
 "Our wedge is one platform that discovers, governs, and protects, on the same behavioral and identity core as Layer "
 "2 applied at runtime, with native real-time mitigation, block, rate-limit, header injection, deception, inline or "
 "passive, and explicit agentic-AI governance including real-time PII and credential redaction. No app instrumentation."])
meta("Maturity line:","“Integrated capability. Architecture story here, not a live 22CT demo.”")
meta("Likely pushback:","“How is this different from API-security point tools?”  ->  “A unified behavioral/identity "
 "core extended natively to agentic-AI governance with no-instrumentation runtime mitigation, sharing the behavioral "
 "philosophy of Layer 2, which is what makes the cross-layer story coherent.”")
meta("Transition:","“All of that assumes the code itself isn't already backdoored. Layer 4.”")

# --- BLOCK 5 ---
H2("Block 5 - Layer 4: Code & Supply-Chain Assurance (15 min, Lead)")
meta("Goal:","Close the lifecycle at the source. Honest about maturity: in development.")
meta("Show:","Layer 4 slide.")
say(["Layer 4's job is to find and remove the exploitable vulnerability before deploy, the code, dependency, "
 "container, or IaC flaw that grants initial access in the first place.",
 "Incumbents, SAST, DAST, software composition analysis, container and secrets scanners, are largely matching "
 "known CVEs. They can't find novel zero-days, they're reactive to disclosure, which is exactly the window Volt and "
 "Salt exploit, and they drown developers in false positives. Supply-chain injection slips past dependency scanners "
 "that only check for known-bad versions.",
 "Our wedge is AI-driven novel zero-day discovery across code, dependencies, containers, and IaC, pre-deploy not "
 "post-disclosure, gated through a multi-stage CI/CD pipeline so vulnerable code can't reach production."])
meta("Maturity line:","“This layer is in development. I'm presenting the approach, not a running product.”")
callout("FLAG before the meeting: the '2,000+ zero-days in 7 weeks' figure must be sourced before you state it as fact, "
 "or present it as a target. The deck now labels Mythos 'in development.' Brandon will check the number.",color=AMBER)
meta("Likely pushback:","“Isn't this SAST/SCA?”  ->  “Those match known CVEs. This targets novel zero-day discovery "
 "before disclosure.”")
meta("Transition:","“Now let me put the four together, because no single box does this.”")

# --- BLOCK 6 ---
H2("Block 6 - How the layers work together + federal fit (20 min, Both)")
meta("Goal:","Show one Typhoon intrusion crossing all four layers, then land federal deployability.")
meta("Show:","The 4-layer diagram again, trace the attack path across it.")
say(["One vocabulary, MITRE ATT&CK, ties all four layers together. Watch one Typhoon intrusion cross the stack.",
 "Layer 4 removes the exploitable code that would have granted initial access. Layer 1 proves the misconfigured "
 "edge path Volt enters through is closed. Layer 2 catches the valid-account, living-off-the-land behavior once "
 "they're inside, the gap Layer 1 structurally can't see. And Layer 3 catches the API and data exfiltration and the "
 "agent abuse at runtime, which is Salt's collection altitude.",
 "On federal fit, and here I'm speaking to what runs, Layer 2: it's containerized, it runs on logs agencies already "
 "collect, it deploys across IL5, IL6, and JWICS, it holds no PII in the vector, and the offline embedding mode runs "
 "today with a local semantic model as a Phase-1 build for air-gapped enclaves.",
 "The honest integration line: no single box does this. We've built an integrated four-layer architecture under one "
 "vocabulary. The behavioral layer is ours and runs live; the rest is architecture and integrated capability; full "
 "cross-layer fusion is the roadmap."])
meta("Transition:","“Which brings me to what we'd actually do together.”")

# --- BLOCK 7 ---
H2("Block 7 - Roadmap + the ask (15 min, Lead)")
meta("Goal:","Convert to a bounded Phase-1 pilot and a policy relationship. Be specific and modest.")
meta("Show:","The phased-build slide.")
say(["The path to operational for Layer 2 is phased. Phase 0, the detection core, is now. Phase 1 is the pilot: "
 "real-data ingestion, basic identity resolution, a versioned twin store, and a local semantic model. Phase 2 is "
 "streaming and scale. Phase 3 is the graph and full identity fusion. Phase 4 is evidence, forensics, and response.",
 "The ask is a bounded Phase-1 pilot. Connectors for one enclave's identity and endpoint logs, identity resolution "
 "for a defined population, the validated core run on real telemetry, and detection, explainability, and "
 "false-positive economics measured side-by-side against current tooling, advisory and human-in-the-loop. Identity "
 "fusion is the first thing we co-design.",
 "And separately, R Street thought leadership on the preemptive-plus-behavioral approach and the national "
 "valid-account, living-off-the-land gap, plus a follow-on engineering session."])
P("Closing line (say it slowly):",bold=True)
P("“What you saw run today is the behavioral detection core, on the logs you already collect, catching four "
 "nation-state campaigns that the tools agencies deploy score as normal. It sits inside a four-layer architecture, "
 "preemptive, behavioral, API and data, and code, under one MITRE vocabulary. The behavioral layer is ours and runs "
 "live; the rest is architecture and integrated capability; full fusion is what we'd build. The honest ask is a "
 "bounded pilot on real data.”",italic=True)

# ============ 6 NUMBERS ============
H1("6. Numbers to keep straight (he will check)")
B("Dataset: ","250 entities, 485 days (70 weeks), ~14M events, 5 log sources.")
B("FP rate: ","8.1% composite (about 20/246 normals above Volt's 13.70), 4/4 caught. Use 8.1, never 8.5.")
B("Ranks: ","Salt USR-118 #1 (51.27); Insider USR-156 #2 (46.24); Slow APT USR-234 #7 (19.44); Volt USR-042 #24 (13.70).")
B("Baselines: ","classical anomaly methods 0/4 (FP 4.5-14.6%); z-score 1/4 (9.8%).")
B("Separation is by RANK, not a clean gap: ","top normal user (USR-081, 21.5) outscores Slow APT (19.4) and Volt (13.7). Always present a ranked list.")
B("Operating point ","was chosen knowing ground truth in the PoC: a reporting benchmark to re-validate, not a deployable threshold.")
B("Context facts: ","~99% of AI-related vulnerabilities are API-related; 95% of API attacks come from authenticated sources.")

# ============ 7 TOUGH Q ============
H1("7. Tough-question quick card")
H2("Cross-cutting")
B("'One product or a bundle?' ","Integrated four-layer architecture under one MITRE vocabulary. Behavioral layer runs live; the rest is architecture and integrated/in-development; full fusion is roadmap.")
B("'What can you show me today?' ","The behavioral layer, live, end to end. The others as architecture and evidence.")
B("'How integrated is it today?' ","Shared vocabulary and vision now; turnkey cross-layer fusion is roadmap. Don't overclaim.")
H2("Layer-specific")
B("L1 'vs policy-analysis tools?' ","Proof over the full state space + MITRE, not rule-linting.")
B("L2 'vs traditional UEBA?' ","Meaning vs counts; 0/4 vs 4/4 on identical data; ranked, explainable, MITRE-mapped.")
B("L2 '8% FP at scale?' ","Ranked list, not threshold alerts; analysts work top-down; tune on real data.")
B("L2 'synthetic, real telemetry?' ","Blind test to prove the mechanism; next step is the bounded pilot.")
B("L3 'vs API-security point tools?' ","Unified behavioral/identity core + native agentic-AI governance, no instrumentation.")
B("L4 'vs SAST/SCA?' ","Novel zero-day discovery pre-disclosure, not known-CVE matching. (Substantiate the 2,000 figure.)")
B("'Data sovereignty / air-gap?' ","Offline embedding today; local semantic model is Phase 1; no PII in embeddings.")
B("'Is identity fusion / forensics / response built?' ","Detection core is built and demoed; those three are the next phase. Not presented as running today.")

# ============ 8 FIX LIST ============
H1("8. Pre-meeting fix list (status)")
B("8.5% to 8.1% (slide 14): ","DONE.",);
B("HMAC forensics (slide 11): ","relabeled '(roadmap)'. DONE.")
B("ABAC / proportional response (slide 11): ","relabeled '(roadmap)'. DONE.")
B("Known-bad 0-FP (slide 20): ","relabeled 'roadmap, not in current build'. DONE.")
B("Mythos / Layer-4 zero-days (slides 4,6): ","labeled '(in development)'. OPEN: source the 2,000 figure or cut it.")
B("'Integrated platform' language: ","present as architecture + integrated/in-development; cross-layer fusion is roadmap.")
B("Baselines: ","one canonical set (0/4 classical, 1/4 z-score); retire the others.")

fp=d.add_paragraph(); fr=fp.add_run("Internal preparation playbook. SAY tracks are scripts to adapt. Build-vs-integrated-vs-roadmap "
 "distinctions and the reconciled number set are load-bearing. Not for distribution.")
fr.italic=True; fr.font.size=Pt(9); fr.font.color.rgb=GRAY

out="WP DLA/Presentation/Brandon_Pugh_Working_Session_Run-Sheet_INTERNAL.docx"
os.makedirs("WP DLA/Presentation",exist_ok=True)
d.save(out)
print("saved:",out,"| paragraphs:",len(d.paragraphs),"| tables:",len(d.tables))
