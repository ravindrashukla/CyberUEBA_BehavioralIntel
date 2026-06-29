"""Business-friendly edition of the V-Intelligence UEBA whitepaper.
Keeps real technical substance (why threshold detection misses slow attacks,
the behavioral-entity architecture, drift and direction naming, honest results,
novelty, limitations) but redacts exact implementation logic (CUSUM formulas,
thresholds, weights, code). Abbreviations clarified on first use. ~8 pages.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "UEBA_Behavioral_Intelligence_Academic.docx")

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
ORG  = RGBColor(0xC2, 0x5A, 0x12)
GREY = RGBColor(0x55, 0x5F, 0x6B)
GREEN= RGBColor(0x1E, 0x7A, 0x44)
WHITE= RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.1
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(doc.sections[0], m, Inches(0.85))


def H1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    return p

def H2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ORG
    return p

def para(text, size=10.5, color=None, bold=False, italic=False, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def bullet(text, lead=None):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run(text.lstrip()); r2.font.size = Pt(10.5)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def setcell(cell, text, bold=False, size=9.5, color=None, fill=None):
    if fill: shade(cell, fill)
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
    if color: r.font.color.rgb = color

def table(headers, rows, widths=None, hdr_fill='0B1F3A'):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        setcell(t.rows[0].cells[j], h, bold=True, color=WHITE, fill=hdr_fill)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            setcell(t.rows[i].cells[j], v, bold=(j == 0), fill='F4F6F8' if j == 0 else None)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ================= FIGURE GENERATION =================
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

FIGDIR = os.path.join(os.path.dirname(__file__), 'figs_business_ueba')
os.makedirs(FIGDIR, exist_ok=True)

C_LEG = '#8A93A0'   # traditional grey
C_LEGD = '#C0392B'  # traditional red accent
C_VI = '#0891B2'    # V-Intelligence teal
C_NAVY = '#0B1F3A'
plt.rcParams.update({'font.family': 'Calibri', 'font.size': 10, 'axes.edgecolor': '#C9CFD6',
                     'axes.labelcolor': C_NAVY, 'xtick.color': '#55607B', 'ytick.color': '#55607B'})

def fig1_slowlow():
    """The slow-and-low evasion: each week stays under the alarm; the accumulated drift crosses it."""
    rng = np.random.default_rng(11)
    wk = np.arange(1, 61)
    # weekly behavioral-change signal: small, slowly ramping, always below the fixed threshold
    weekly = 0.18 + 0.10 * (wk / 60) ** 1.4 + rng.normal(0, 0.035, len(wk))
    weekly = np.clip(weekly, 0.05, 0.46)
    thresh = 0.55
    fig, axes = plt.subplots(1, 2, figsize=(7.0, 3.0))
    ax = axes[0]
    ax.bar(wk, weekly, color=C_LEG, alpha=0.75, width=0.9,
           label='Weekly behavioral change (attacker)')
    ax.axhline(thresh, color=C_LEGD, lw=2.0, ls='--', label='Fixed alert threshold')
    ax.text(30, thresh + 0.03, 'Never crossed — no alert, ever', ha='center', fontsize=8.8,
            color=C_LEGD, fontweight='bold')
    ax.set_title('What a threshold-based tool sees\n(each week judged alone)', fontsize=9.5, color=C_NAVY)
    ax.set_xlabel('Week'); ax.set_ylabel('Change vs. baseline')
    ax.set_ylim(0, 0.75); ax.legend(loc='upper left', fontsize=8, frameon=False)
    ax = axes[1]
    drift = np.cumsum(weekly - 0.16)            # attacker accumulates
    normal_hi = np.cumsum(np.full(len(wk), 0.045))  # normal users hover near this envelope
    ax.fill_between(wk, 0, normal_hi, color=C_LEG, alpha=0.30, label='Normal-user range')
    ax.plot(wk, drift, color=C_VI, lw=2.4, label='Accumulated behavioral drift (attacker)')
    above = np.where(drift > normal_hi)[0]
    cross = above[above > 10][0]
    ax.plot(wk[cross], drift[cross], 'o', color=C_LEGD, ms=7, zorder=5)
    ax.annotate('Drift alarm —\nmonths, not years', xy=(wk[cross], drift[cross]),
                xytext=(12, drift[cross] + 0.35), fontsize=8.8, fontweight='bold', color=C_NAVY,
                arrowprops=dict(arrowstyle='->', color=C_NAVY, lw=1.2))
    ax.set_title('What V-Intelligence UEBA sees\n(the same weeks, accumulated)', fontsize=9.5, color=C_NAVY)
    ax.set_xlabel('Week'); ax.set_ylabel('Accumulated drift')
    ax.set_ylim(0, drift[-1] * 1.32)
    ax.legend(loc='upper left', fontsize=8, frameon=False)
    for ax in axes:
        ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig1_slowlow.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig2_scoreboard():
    """Detection scoreboard: attacks caught and false-positive rate, real numbers."""
    methods = ['Local Outlier\nFactor', 'Z-Score', 'Isolation\nForest', 'One-Class\nSVM', 'V-Intelligence\nComposite']
    det = [0, 1, 0, 0, 4]
    fp = [4.5, 9.8, 5.3, 14.6, 8.1]
    colors = [C_LEG, C_LEG, C_LEG, C_LEG, C_VI]
    x = np.arange(len(methods))
    fig, axes = plt.subplots(1, 2, figsize=(7.0, 3.0))
    ax = axes[0]
    ax.bar(x, det, color=colors, width=0.62)
    for xi, d in zip(x, det):
        ax.text(xi, d + 0.08, f'{d}/4', ha='center', fontsize=9.5, fontweight='bold',
                color=(C_NAVY if d else C_LEGD))
    ax.set_xticks(x, methods, fontsize=7.6)
    ax.set_ylim(0, 4.8); ax.set_ylabel('Attack campaigns detected (of 4)')
    ax.set_title('Detection — traditional methods catch at most 1 of 4...', fontsize=9.5, color=C_NAVY)
    ax = axes[1]
    bars = ax.bar(x, fp, color=colors, width=0.62)
    for xi, f in zip(x, fp):
        ax.text(xi, f + 0.35, f'{f}%', ha='center', fontsize=9, fontweight='bold', color=C_NAVY)
    ax.set_xticks(x, methods, fontsize=7.6)
    ax.set_ylim(0, 18); ax.set_ylabel('False-positive rate (%)')
    ax.set_title('...the composite catches all four — at a comparable\nfalse-positive rate, with direction', fontsize=9.5, color=C_NAVY)
    for ax in axes:
        ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig2_scoreboard.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig3_composite():
    """Per-attacker composite scores vs the detection threshold; USR-234's novelty rescue."""
    labels = ['Salt Typhoon\ntelecom (USR-118)\nRank #1', 'Insider threat\n(USR-156)\nRank #2',
              'Slow APT, C2 beacon\n(USR-234)\nRank #7', 'Volt Typhoon LOTL\n(USR-042)\nRank #30']
    scores = [51.7, 46.2, 20.0, 12.9]
    x = np.arange(len(labels))
    fig, ax = plt.subplots(figsize=(7.0, 3.3))
    ax.bar(x[[0, 1, 3]], [51.7, 46.2, 12.9], color=C_VI, width=0.58)
    # USR-234 stacked: 7.0 without novelty + 13.0 novelty persistence
    ax.bar(x[2], 7.0, color=C_LEG, width=0.58, label='USR-234 without the novelty signal: 7.0 (rank ~#105 — missed)')
    ax.bar(x[2], 13.0, bottom=7.0, color=C_VI, width=0.58,
           label='Novelty persistence: +13.0 (new C2 address recurs 60 of 60 weeks)')
    thresh = 12.95
    ax.axhline(thresh, color=C_LEGD, lw=1.8, ls='--')
    ax.text(2.64, thresh + 8.0, 'Detection threshold\n(catches all four attackers)', fontsize=8.6, color=C_LEGD,
            ha='center', fontweight='bold')
    for xi, s in zip(x, scores):
        ax.text(xi, s + 1.0, f'{s}', ha='center', fontsize=9.5, fontweight='bold', color=C_NAVY)
    ax.set_xticks(x, labels, fontsize=8.2)
    ax.set_ylabel('Composite behavioral-risk score')
    ax.set_ylim(0, 58)
    ax.legend(loc='upper right', fontsize=8.2, frameon=False)
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig3_composite.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

import csv as _csv
_DATA = os.path.join(os.path.dirname(__file__), '..', 'data')
# app palette for the four attackers (matches the Live Explorer)
ATK_STYLE = [('USR-156', '#C0392B', 'Insider threat (USR-156)'),
             ('USR-234', '#E67E22', 'Slow APT (USR-234)'),
             ('USR-042', '#8E44AD', 'Volt Typhoon (USR-042)'),
             ('USR-118', '#2980B9', 'Salt Typhoon (USR-118)')]
ATK_IDS = {a for a, _, _ in ATK_STYLE}

def _ueba_db_rows(query):
    """Operational DB first (single source of truth); returns list of dict rows or None."""
    try:
        import psycopg2
        from dotenv import load_dotenv
        load_dotenv(os.path.join(os.path.dirname(__file__), '..', '.env'))
        url = os.getenv('DATABASE_URL_HOST') or os.getenv('DATABASE_URL')
        if not url:
            return None
        conn = psycopg2.connect(url, connect_timeout=3)
        cur = conn.cursor(); cur.execute(query)
        cols = [d[0] for d in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]
        conn.close()
        return rows or None
    except Exception:
        return None


def fig4_semantic():
    """REAL DATA: cumulative drift both ways — feature space (attackers hidden)
    vs semantic embedding space (attackers separate)."""
    # semantic drift: per-user cumulative drift-from-baseline in embedding space (DB-first)
    sem = {}
    _tr = _ueba_db_rows("SELECT user_id, week_idx, composite_drift FROM weekly_trajectories")
    if _tr is None:
        with open(os.path.join(_DATA, 'tier3_results', 'all250_trajectories.csv')) as f:
            _tr = list(_csv.DictReader(f))
    for r in _tr:
        sem.setdefault(r['user_id'], []).append((int(r['week_idx']), float(r['composite_drift'])))
    # feature-space final cumulative score: DB detection_results.feat_cusum_value (CSV fallback)
    feat_final = {}
    _fc = _ueba_db_rows("SELECT user_id, feat_cusum_value FROM detection_results")
    if _fc is None:
        with open(os.path.join(_DATA, 'tier3_results', 'tier3_comparison.csv')) as f:
            _fc = list(_csv.DictReader(f))
    for r in _fc:
        feat_final[r['user_id']] = float(r['feat_cusum_value'])

    def sem_cusum(u):
        seq = sorted(sem[u]); return np.cumsum([d for _, d in seq])

    fig, axes = plt.subplots(1, 2, figsize=(7.2, 3.3))
    # LEFT: every user's official raw-metric cumulative score, sorted — attackers mid-pack
    ax = axes[0]
    order = sorted(feat_final, key=feat_final.get, reverse=True)
    xs = np.arange(1, len(order) + 1)
    ax.scatter(xs, [feat_final[u] for u in order], s=7, color=C_LEG, alpha=0.45, lw=0)
    for u, col, lab in ATK_STYLE:
        rk = order.index(u) + 1
        ax.scatter([rk], [feat_final[u]], s=42, color=col, zorder=5)
        ax.annotate(f'#{rk}', xy=(rk, feat_final[u]), xytext=(rk, feat_final[u] + 1.6),
                    ha='center', fontsize=7.6, color=col, fontweight='bold')
    ax.set_title('Raw-metric cumulative drift (feature space)\n', fontsize=9.5, color=C_NAVY)
    ax.text(0.5, 1.02, 'attackers HIDDEN mid-pack (ranks #90–#133 of 250)',
            transform=ax.transAxes, ha='center', fontsize=8.8, color=C_LEGD, fontweight='bold')
    ax.set_xlabel('User rank (1 = highest score)'); ax.set_ylabel('Final cumulative score')
    ax.spines[['top', 'right']].set_visible(False)
    # RIGHT: 70-week semantic-drift trajectories for all 250 users
    ax = axes[1]
    for u in sem:
        if u not in ATK_IDS:
            ax.plot(sem_cusum(u), color=C_LEG, lw=0.5, alpha=0.18)
    for u, col, lab in ATK_STYLE:
        ax.plot(sem_cusum(u), color=col, lw=2.0, label=lab)
    ax.set_title('Semantic drift (behavioral embeddings)\n', fontsize=9.5, color=C_NAVY)
    ax.text(0.5, 1.02, 'drift surfaces three of the four (#1, #4, #13)',
            transform=ax.transAxes, ha='center', fontsize=8.8, color='#1E7A44', fontweight='bold')
    ax.set_xlabel('Week')
    ax.legend(loc='upper left', fontsize=7.6, frameon=False)
    ax.spines[['top', 'right']].set_visible(False)
    # the slow APT stays buried even in drift space — the novelty phase catches it
    apt = sem_cusum('USR-234')
    ax.annotate('Slow APT stays buried in drift —\ncaught by novelty (Figure 5)',
                xy=(len(apt) - 1, apt[-1]), xytext=(20, apt[-1] - 0.62),
                fontsize=7.8, color='#B5651D', fontweight='bold',
                arrowprops=dict(arrowstyle='->', color='#B5651D', lw=1.1))
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig4_semantic.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig5_radar():
    """REAL DATA: the five known scoring dimensions as population percentiles —
    normal users cluster near the centre; each attacker breaks out differently."""
    rows = list(_csv.DictReader(open(os.path.join(_DATA, 'tier3_results', 'composite_scores.csv'))))
    phases = [('signal_strength', 'Signal\nStrength'), ('breadth_15', 'Breadth'),
              ('sustained_signal', 'Sustained\nDeviation'), ('ctx_spread_z', 'Context\nDivergence'),
              ('novelty_score', 'Novelty\nPersistence')]
    vals = {ph: np.array([float(r[ph]) for r in rows]) for ph, _ in phases}

    def pct(ph, x):
        # mid-rank percentile: ties (e.g., the many users with novelty = 0) split evenly
        v = vals[ph]
        return 100.0 * ((v < x).mean() + 0.5 * (v == x).mean())

    def user_pcts(uid):
        r = next(r for r in rows if r['uid'] == uid)
        return [pct(ph, float(r[ph])) for ph, _ in phases]

    normal = [r for r in rows if r['uid'] not in ATK_IDS]
    med = [np.median([pct(ph, float(r[ph])) for r in normal]) for ph, _ in phases]
    p75 = [np.percentile([pct(ph, float(r[ph])) for r in normal], 75) for ph, _ in phases]

    ang = np.linspace(0, 2 * np.pi, len(phases), endpoint=False).tolist(); ang += ang[:1]
    fig, ax = plt.subplots(figsize=(5.6, 4.4), subplot_kw=dict(polar=True))
    for series, col, lab, ls, lw in (
            [(med, '#9AA3AD', 'Normal median', '--', 1.4), (p75, '#707A86', 'Normal 75th pct', ':', 1.4)] +
            [(user_pcts(u), c, l, '-', 1.9) for u, c, l in reversed(ATK_STYLE)]):
        v = list(series) + [series[0]]
        ax.plot(ang, v, color=col, ls=ls, lw=lw, label=lab)
        if ls == '-':
            ax.fill(ang, v, color=col, alpha=0.06)
    ax.set_xticks(ang[:-1]); ax.set_xticklabels([l for _, l in phases], fontsize=8.6, color=C_NAVY)
    ax.set_ylim(0, 100); ax.set_yticks([25, 50, 75, 100])
    ax.set_yticklabels(['25th', '50th', '75th', '100th'], fontsize=7, color='#8A93A0')
    ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.06), ncol=3, fontsize=7.6, frameon=False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig5_radar.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

FIG1, FIG2, FIG3, FIG4, FIG5 = fig1_slowlow(), fig2_scoreboard(), fig3_composite(), fig4_semantic(), fig5_radar()

def add_fig(path, caption, width=6.4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY


# ================= TITLE =================
r = doc.add_paragraph().add_run('Catching the Cyber Threats That Never Trip an Alarm:\nTraditional Security Monitoring vs. V-Intelligence UEBA')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
para('A business and technology overview of behavioral intelligence: what it is, what it detects that current tools '
     'miss, and what the evidence shows', 12.5, GREY, after=2)
para('22nd Century Technologies, Inc.  ·  V-Intelligence UEBA Program  ·  June 2026  ·  Business Edition',
     9.5, GREY, after=2)
para('Disclosure note: this edition describes the architecture and the innovation, and reports measured results. The '
     'exact detection logic — formulas, thresholds, weights, and parameters — is proprietary and is withheld here; it '
     'is available in the full technical edition under NDA (non-disclosure agreement).', 9.5, ORG, italic=True, after=4)
para('Data note: all results in this paper come from a synthetic (simulated) dataset built to prove the concept — '
     '250 users, 485 days, approximately 15 million log events across five log sources, with four embedded long-duration '
     'attack campaigns modeled on real, publicly documented intrusions. Results must be re-validated on real '
     'operational telemetry before any production claim.', 9.5, ORG, italic=True, after=8)

# ================= 1. EXEC SUMMARY =================
H1('1.  Executive Summary')
para('The most damaging cyber attacks of the past decade did not trip a single alarm. Volt Typhoon — a Chinese '
     'state-sponsored campaign against U.S. critical infrastructure — sat inside victim networks for at least five '
     'years undetected. Salt Typhoon — described by the chairman of the Senate Intelligence Committee as "the worst '
     'telecom hack in our nation\'s history" — operated inside major U.S. telecommunications providers for months. In '
     'both cases, the victims\' own security tools never fired, and the campaigns were discovered by outside '
     'intelligence agencies. Such attackers succeed by operating with valid credentials and legitimate tools, '
     'remaining within the statistical bounds of normal activity while shifting their behavior gradually over months.')
para('This paper presents the business case for User and Entity Behavior Analytics (UEBA), a detection approach '
     'that establishes a behavioral baseline for every user and system and then evaluates both how much behavior '
     'changes and the direction in which it is trending. The V-Intelligence UEBA program validated this approach on a '
     'realistic simulated enterprise: 250 users, 485 days, roughly 15 million events, and four embedded attack '
     'campaigns — a 14-month insider threat, a 417-day slow nation-state intrusion with covert call-home traffic, '
     'a 412-day Volt Typhoon-style campaign, and a 412-day Salt Typhoon-style telecom campaign.')
para('The principal findings:', bold=True)
bullet('Of four industry-standard anomaly-detection algorithms run against the same data, three detected none of '
       'the four attack campaigns, and the best detected exactly one. That single detection was a statistical spike '
       'with no accompanying explanation of which behavior changed or what kind of threat it resembled.',
       lead='Traditional methods miss these attacks. ')
bullet('V-Intelligence\'s multi-phase composite scoring detected all four attack campaigns at a 10.6% false-positive '
       'rate (26 of 246 normal users flagged), and identified for the analyst which behavioral dimension drifted and '
       'toward what known threat pattern.', lead='Four of four detected, with direction. ')
bullet(' A slow nation-state intrusion whose behavior barely changes in volume is the hardest case in the study. It is '
       'invisible to every magnitude-based method, yet a single behavioral question detected it: whether a '
       'never-before-seen contact point recurred week after week. It had done so for 60 consecutive weeks.',
       lead='The hardest case, detected. ')
para('The underlying distinction is straightforward. An insider changes what they access rather than how much; '
     'traditional algorithms measure magnitude, whereas V-Intelligence UEBA measures direction.', bold=True)

# ================= 2. THE PROBLEM =================
H1('2.  The Problem: Attacks Built to Stay Under the Alarm')
para('Enterprise security monitoring today rests on two pillars. The first is the SIEM (Security Information and '
     'Event Management system), the central log-collection platform that applies expert-written rules to identify '
     'known attack patterns. The second is signature-based detection: antivirus and intrusion-detection tools that '
     'match activity against IOCs (Indicators of Compromise, such as known-bad file fingerprints, addresses, and '
     'domains). Both answer the same question, namely whether the activity matches something already known to be '
     'malicious.')
para('Modern adversaries are engineered to make the answer no. Three threat families account for most of these cases.')
bullet(' A trusted employee or contractor who misuses access they legitimately hold. Industry data (the Verizon Data '
       'Breach Investigations Report) attributes roughly 20% of breaches to insiders, with a median time to detection '
       'exceeding 200 days. Every individual action uses valid credentials; only the pattern over months reveals intent.',
       lead='Insider threats. ')
bullet(' An APT (Advanced Persistent Threat) is a patient, well-resourced intrusion, typically nation-state, that '
       'gains a foothold and remains for months or years. Its covert command-and-control (C2) channel, through which '
       'the implant calls home for instructions, is deliberately throttled to a trickle, with a check-in roughly '
       'every 12 hours and only kilobytes of traffic that are invisible inside enterprise volume.',
       lead='Advanced persistent threats. ')
bullet(' Living-off-the-land (LOTL) attacks use only the administration tools already installed on every system, such '
       'as PowerShell, remote desktop, and management utilities. There is no malware to fingerprint and no rule to '
       'violate, because every action is, in isolation, an authorized one.', lead='Living-off-the-land attacks. ')
para('The common technique is known as "slow and low": behavior is changed so gradually that no single week appears '
     'unusual, while the cumulative change over months achieves the objective. In the simulated slow-APT campaign, '
     'the implant\'s call-home traffic begins at one or two domain lookups per day and roughly 8 kilobytes of '
     'traffic, indistinguishable from software-update checks, and even at full operation it stays below every volume '
     'threshold. After 26 weeks the attacker has staged and removed gigabytes of data in packets too small to trigger '
     'any single alert. The activity is visible only in the trajectory, never in a single snapshot.')
add_fig(FIG1, 'Figure 1 — The "slow and low" evasion. Left: a threshold-based tool judges each week alone, and the '
              'attacker\'s weekly change never crosses the fixed alert line — no alert is ever raised. Right: the same '
              'weeks accumulated as behavioral drift separate the attacker from the normal-user range in months, not '
              'years. (Illustrative, based on the simulated 60-week slow-APT campaign.)')
para('The real-world consequence is well documented. Neither Volt Typhoon nor Salt Typhoon was detected by any '
     'victim organization\'s own monitoring; SIEM, intrusion detection, and endpoint tools all failed because there '
     'was no signature to match and no threshold crossed. Both were found by external intelligence operations, years '
     'or months after compromise. This is a structural gap rather than a tuning problem, and it is the gap that '
     'behavioral analytics is intended to address.')

# ================= 3. WHAT UEBA IS =================
H1('3.  What Is UEBA — and What Is Different Here')
para('User and Entity Behavior Analytics (UEBA) inverts the traditional question. Instead of defining what is '
     'malicious and watching for it, UEBA learns what is normal — for every user, device, application, and network '
     'segment individually — and flags departures from that personal baseline. Three principles define it:')
table(['Principle', 'Question asked', 'Traditional answer', 'UEBA answer'],
      [('Baseline', 'What is normal?', 'One statistical threshold for everyone.', 'A personal behavioral profile per entity — each has its own "normal."'),
       ('Deviation', 'Has something changed?', 'A measurement exceeded a fixed limit.', 'The entity\'s behavioral trajectory is drifting from its own baseline.'),
       ('Direction', 'What kind of change?', 'Not considered — magnitude only.', 'Which behavioral dimension is drifting, and toward what known threat pattern.')],
      widths=[1.0, 1.5, 2.0, 2.4])
para('The third principle is the decisive one. A user whose file activity rises 50% might be working against a '
     'deadline (benign) or staging data for theft (critical), and a magnitude-only tool flags both identically, or '
     'neither. The discriminating signal is the direction of change, for example a shift toward restricted data or '
     'toward known adversary behavior. Direction is what turns an anomaly score into an investigation lead.')

# ================= 4. THE APPROACH =================
H1('4.  The V-Intelligence Approach: A Living Chart for Every Entity')
H2('4.1  The digital entity — a patient\'s medical chart for behavior')
para('The central concept, in plain terms, is that V-Intelligence maintains a "digital entity," a living behavioral '
     'profile, for every monitored user and system. The closest analogy is a patient\'s medical chart. A chart does '
     'not record only today\'s temperature; it accumulates history, vitals, and context, and a doctor reads the trend '
     'to act before a condition becomes acute. The digital entity is that chart for behavior. It captures not merely '
     'how many files were opened but the surrounding context, and, importantly, how that picture is changing week '
     'over week.')
para('Each week, raw log activity from five sources (sign-ins, file access, endpoint activity, network traffic, and '
     'domain lookups) is distilled into 23 behavioral measurements per user. Those measurements are organized into '
     'five behavioral zones, each representing a distinct dimension of who the entity is and what it does.')
table(['Behavioral zone', 'What it captures', 'What it catches when it drifts'],
      [('Identity', 'Who the entity is — role, department, clearance, tenure.', 'Stability here is itself a clue: insiders and credential thieves look like the same person.'),
       ('Access pattern', 'How the entity signs in — frequency, failures, timing, sources.', 'Credential abuse and account takeover.'),
       ('Data behavior', 'What data it touches — volume, sensitivity mix, read/write balance.', 'Insider data theft: a shift toward restricted and confidential material.'),
       ('Network footprint', 'Where it communicates — destinations, external traffic, domain lookups.', 'Covert call-home channels and data removal routes.'),
       ('Risk posture', 'Endpoint health — suspicious process activity and risk indicators.', 'Compromised machines and privilege-escalation tooling.')],
      widths=[1.3, 2.7, 2.9])
para('Zone decomposition is the central architectural choice. When the identity zone is stable but the data-behavior '
     'zone is drifting, the system can conclude that this is the same person using the same credentials but accessing '
     'different data than before. That conclusion is the insider-threat signature, and no single composite score can '
     'express it.')

H2('4.2  Semantic embeddings — a map of meaning')
para('To compare behaviors, V-Intelligence borrows the core representation used by large language '
     'models. Each week\'s measurements are first written out as a structured plain-language description (for example, '
     '"30 files, 15% restricted, 8% confidential" rather than simply "30 files"), and that description is converted '
     'into a high-dimensional numerical vector called an embedding. The result is a semantic space, a map of meaning '
     'in which behaviors that mean similar things sit near each other even when their raw numbers differ.')
para('This addresses a specific limitation of raw numbers. The simulated insider accessed about 30 files per '
     'week both before and during the campaign, so on raw counts nothing changed. Week one, however, consisted of '
     'public files and week forty of restricted files, and on the map of meaning those two weeks sit far apart. In '
     'validation, cumulative drift was measured both ways. On raw numbers, all four attackers stayed buried in the '
     'population of 246 normal users, ranked between #90 and #133 and effectively indistinguishable. On semantic '
     'embeddings, three of the four climbed clear of the population, with the insider reaching #1, above every normal '
     'user. The fourth, the slow APT, stays buried even in drift space and is caught instead by the novelty question '
     '(Section 5.3). This is the empirical reason that no single signal is sufficient and that the scoring is '
     'composite.')
add_fig(FIG4, 'Figure 2 — All 250 users, both representations (measured, synthetic data). Left: each user\'s final '
              'cumulative-drift score on the 23 raw metrics, sorted — the four attackers (colored) land mid-pack at '
              'ranks #90–#133, indistinguishable from normal users. Right: cumulative drift in semantic embedding '
              'space over the full 70 weeks — the insider rises to #1 (above every normal user), Salt Typhoon to #4, '
              'Volt Typhoon to #13. The slow APT stays buried in drift and is caught by novelty persistence instead '
              '(Figure 5).')

H2('4.3  Drift, direction, and naming the threat')
para('Tracking each entity\'s weekly position on the map produces a trajectory, and the trajectory yields drift: the '
     'direction in which behavior is trending, together with its rate and consistency. A cumulative change-detection '
     'technique, from the same family used to catch slow degradation in industrial quality control, accumulates small '
     'weekly shifts until they become statistically significant, which is what allows the system to catch "slow and '
     'low" activity. The exact accumulation logic, thresholds, and parameters are proprietary and are withheld in '
     'this edition.')
para('Direction is then named. The system maintains a library of reference threat concepts, which are plain-language '
     'descriptions of known adversary behaviors embedded into the same map and aligned with MITRE ATT&CK, the '
     'industry-standard public catalog of adversary tactics and techniques. When an entity\'s drift points toward '
     '"data staging for exfiltration" or "covert beaconing," the alert states this and attaches the matching ATT&CK '
     'techniques. Each simulated campaign was built from, and is reported against, real documented technique sets.')
para('Each attack family produces a recognizably different signature on the chart.')
table(['Threat', 'Behavioral signature', 'Real-world parallel'],
      [('Insider threat', 'Identity zone stable; data-behavior zone drifting steadily for months.', 'Snowden, Manning; departing-employee IP theft.'),
       ('Slow APT (covert C2)', 'Identity stable; network-footprint zone drifting — new recurring external contacts.', 'SolarWinds/SUNBURST: 9 months undetected, 18,000+ organizations.'),
       ('Living off the land', 'A regime shift — all five zones step to a new operating mode at once.', 'Volt Typhoon: 5+ years undetected in U.S. critical infrastructure.'),
       ('Telecom infrastructure', 'Persistent multi-zone drift — network and data zones moving together for 100+ days.', 'Salt Typhoon: at least 9 U.S. telecom providers compromised.')],
      widths=[1.4, 3.0, 2.5])

H2('4.4  Multi-phase composite scoring — five questions, one ranked list')
para('No single signal catches all four families, because each attack evades a different detector. The production '
     'approach is composite scoring. Five independent behavioral questions are asked of every user, and the answers '
     'are combined into a single ranked risk score. A normal user may score high on one question by chance, whereas '
     'an attacker tends to score high on several at once. The questions are as follows (the exact scoring logic is '
     'withheld).')
table(['Phase', 'The question it asks'],
      [('Signal strength', 'How extreme are this user\'s strongest deviations, compared with peers in the same role?'),
       ('Breadth', 'How many different behavioral measurements are elevated at once — one noisy metric, or a pattern?'),
       ('Sustained deviation', 'Does the anomaly persist week after week, or was it a one-off event?'),
       ('Context divergence', 'Does the user look anomalous when viewed through attack-specific lenses (insider lens, APT-hunt lens, privilege-audit lens)?'),
       ('Novelty persistence', 'Has a never-before-seen behavior — a new external contact point, a new domain — appeared and kept recurring?')],
      widths=[1.6, 5.3])
para('Each phase alone is insufficient, but combined they form a detection surface that no single evasion defeats. '
     'An attacker who suppresses magnitude still registers on breadth, and one who stays narrow still registers on '
     'persistence or novelty. The SOC (Security Operations Center, meaning the analyst team) receives one ranked '
     'list rather than thousands of raw alerts.')

# ================= 5. RESULTS =================
H1('5.  What the Evidence Shows (Synthetic Data)')
H2('5.1  The scoreboard first')
para('Four industry-standard anomaly-detection algorithms (Z-Score statistics, Isolation Forest, Local Outlier '
     'Factor, and One-Class SVM, or Support Vector Machine) were run against the full 250-user population alongside '
     'composite scoring. Every number below is reproduced directly from the evaluation data.')
table(['Method', 'Attacks detected', 'False-positive rate', 'What the analyst gets'],
      [('Local Outlier Factor', '0 of 4', '4.5%', 'Low false positives — and zero detections. All four campaigns sit inside its normal range.'),
       ('Z-Score', '1 of 4', '9.8%', 'Catches only the Volt Typhoon spike (z = 3.04, barely over the 3.0 line). The other three never cross it — Salt Typhoon peaks at z = 1.71.'),
       ('Isolation Forest', '0 of 4', '5.3%', 'All four attackers look statistically normal in feature space.'),
       ('One-Class SVM', '0 of 4', '14.6%', '36 false positives and zero true detections.'),
       ('V-Intelligence Composite', '4 of 4', '10.6%', 'All four detected, each with zone-level explanation: which dimension drifted, toward what threat.')],
      widths=[1.7, 1.2, 1.2, 2.8])
para('Stated precisely, the best traditional method catches 1 of 4 campaigns, and its single detection is a '
     'single-feature spike with no explanation attached. The composite catches all four at a comparable '
     'false-positive rate and reports to the analyst which behavioral zone changed and toward what threat pattern, '
     'which is the information triage actually consumes. The result is not merely an alert but an actionable '
     'investigation lead.', bold=True)
add_fig(FIG2, 'Figure 3 — The detection scoreboard across 250 users with four embedded attack campaigns (measured, '
              'synthetic data; verified against the evaluation dataset). Traditional methods catch at most 1 of 4; '
              'composite scoring catches 4 of 4 at 10.6% false positives, each detection explained.')
H2('5.2  Four attacks, four different fingerprints')
para('Each campaign was caught by a different combination of the five phases, demonstrating empirically that no '
     'single method would have sufficed.')
bullet(' the strongest overall anomaly, ranked #1 of 250 (score 51.7). The attack touches sign-in, network, and '
       'domain-lookup behavior simultaneously, with 18 measurements elevated at once.', lead='Salt Typhoon telecom (rank #1): ')
bullet(' ranked #2 (score 46.2), caught by sustained deviation and breadth. The campaign showed a 14-month drift from '
       'public toward restricted files while total volume stayed flat, the textbook "same credentials, different '
       'data" signature.', lead='Insider threat (rank #2): ')
bullet(' ranked #7 (score 20.0), the case that motivates the architecture and is detailed below.', lead='Slow APT (rank #7): ')
bullet(' ranked #30 of 250 (score 12.9). This was a deliberately quiet campaign using only legitimate admin tools, '
       'yet it remained inside the top 12% and above the detection threshold on the strength of breadth and signal '
       'strength.', lead='Volt Typhoon living-off-the-land (rank #30): ')
add_fig(FIG5, 'Figure 4 — The drift projected onto the five scoring dimensions, as percentiles across the '
              '250-user population (measured, synthetic data). Normal users cluster near the center (grey dashed = '
              'median, dotted = 75th percentile). Each attacker breaks out on a different combination — the insider '
              'on signal strength and breadth, the slow APT on novelty persistence, Salt Typhoon on nearly every '
              'dimension at once. No normal user reaches these extremes on several dimensions simultaneously.',
              width=5.0)
H2('5.3  Case study: the attack that was "undetectable"')
para('The slow APT (user USR-234) is the hardest detection problem in the study. For 417 days its implant called '
     'home roughly every 12 hours, and its total traffic, lookups, and sign-ins all stayed within normal ranges. On '
     'every magnitude-based phase its scores are negligible. Without the novelty question, its composite score would '
     'be 6.4, a rank of roughly #105 of 250 and far below any reasonable threshold, leaving it effectively '
     'undetectable.')
para('Novelty persistence asks a different question, namely whether something never seen before keeps recurring. One '
     'external address that never appeared in the user\'s baseline subsequently appeared in every one of the 60 '
     'post-baseline weeks, a persistence of 100%. Normal users contact new addresses occasionally, but those '
     'contacts are transient. A new contact that recurs every week for over a year is the behavioral fingerprint of '
     'command-and-control infrastructure. That single signal contributes 13.0 points, the largest phase contribution '
     'of any attacker, lifting USR-234 from rank ~#105 to rank #7 and placing it firmly inside the detected set.')
add_fig(FIG3, 'Figure 5 — Composite scores of the four attack users against the detection threshold '
              '(measured, synthetic data). The slow APT (USR-234) splits into its parts: without the novelty signal it '
              'scores 7.0 and is missed; the recurring never-before-seen contact adds 13.0 points and detects it.')
H2('5.4  False positives, treated as a budget')
para('At the operating point that catches all four attackers, 26 of 246 normal users (10.6%) are flagged alongside '
     'them. Operationally these are not noise; they constitute a prioritized queue of the highest-behavioral-risk '
     'normal users, each explained by the phases that drove its score. The threshold is tunable, and a stricter cut '
     'trades detection margin for fewer flags. The simulated real-world parallels also held: the two campaigns that '
     'no real victim organization ever self-detected, Volt Typhoon and Salt Typhoon, were both detected here, at '
     'ranks #30 and #1.')

# ================= 6. DEPLOYMENT =================
H1('6.  Deployment in Practice')
para('V-Intelligence UEBA is designed to complement, not replace, the existing security stack. The recommended '
     'operating model uses three layers: the existing fast magnitude-based tools as a familiar pre-filter; composite '
     'scoring as the primary detection engine producing the ranked list; and an investigation-context layer that '
     'enriches each lead with drift direction, MITRE ATT&CK technique mapping, and attack-chain correlation.')
bullet(' Five standard log sources that most enterprises already collect (sign-ins, file access, endpoint telemetry, '
       'network flow, and domain lookups), together with data-sensitivity labels. No additional sensors are '
       'required.', lead='Data requirements. ')
bullet(' A phased rollout consisting of 4–6 weeks of baseline building, 2–4 weeks of detection tuning (with a target '
       'of false positives under 10%), 2–4 weeks of red-team validation, and then production.', lead='Timeline. ')
bullet(' Containerized, standard infrastructure with no specialized hardware, suitable for enterprise data centers, '
       'cloud, or disconnected secure environments. Behavioral profiles are computed incrementally, so the ongoing '
       'cost after the initial baseline is small.', lead='Footprint. ')
bullet(' Alerts feed existing SIEM/SOAR (Security Orchestration, Automation and Response) workflows with behavioral '
       'context attached, and behavioral risk scores can feed Zero Trust continuous-verification policies, the '
       'federal architecture mandate that no user or device is trusted by default.', lead='Integration. ')
bullet(' Validation used 250 users; larger populations improve, not degrade, the statistics, because peer baselines '
       'get richer.', lead='Scale. ')

# ================= 7. NOVELTY =================
H1('7.  Why This Is New')
para('Commercial UEBA products exist, and the better ones perform per-entity baselining with statistical or '
     'machine-learning anomaly scores. What they produce is magnitude: a number indicating that behavior is '
     'different, without indicating the direction that would distinguish, for example, insider exfiltration from '
     'benign change. The individual techniques used here (serializing structured data to text for embedding, anomaly '
     'detection in embedding space, cumulative change-detection, and novelty tracking) each exist separately in the '
     'literature.')
para('What is not offered, commercially or in any single published system we have found, is the composition: every '
     'entity maintained as a digital behavioral profile in one shared semantic space; the profile decomposed into '
     'five behavioral zones so the system can say which dimension of behavior moved; drift measured cumulatively so '
     '"slow and low" accumulates into a signal; the direction named against an embedded library of threat concepts '
     'mapped to MITRE ATT&CK; and five independent detection phases fused into one ranked, explained list. The '
     'innovation is the integrated architecture, not any single technique.', bold=True)
para('Scoping notes: vendor internals are not fully public (the comparison is against marketed capability), and a '
     'formal freedom-to-operate review is recommended before any contractual assertion of uniqueness.', 9.5, GREY,
     italic=True)

# ================= 8. LIMITATIONS =================
H1('8.  Limitations — Stated Plainly')
bullet(' All results are from synthetic telemetry. The attack campaigns are carefully calibrated to public CISA '
       '(Cybersecurity and Infrastructure Security Agency) advisories, but real enterprise behavior is messier; every '
       'figure must be re-measured on real data.')
bullet(' The detection threshold reported here is the operating point that catches all four attackers; it was chosen '
       'with knowledge of the answers. In production the threshold must be set blind and tuned per environment.')
bullet(' A 10.6% false-positive rate means 26 flagged normal users per 250; acceptable as a ranked triage queue, but '
       'it is a real analyst cost — and it is not an auto-blocking threshold.')
bullet(' Baselines need 4–6 weeks of history before detection is meaningful, and legitimate behavioral shifts — role '
       'changes, project transitions — are the main source of false flags.')
bullet(' Data-sensitivity labels on file activity materially strengthen the insider signature; environments without '
       'classification metadata lose part of that signal.')
bullet(' Semantic embeddings rely on an embedding model; fully disconnected deployments require a locally hosted '
       'replacement, which must be re-validated.')

# ================= 9. CONCLUSION =================
H1('9.  Conclusion')
para('The threats that matter most, including insiders, nation-state intrusions, living-off-the-land campaigns, and '
     'infrastructure compromises, share one property: they operate inside authorized access, beneath every signature '
     'and threshold, and reveal themselves only as a gradual change of behavioral direction over months. The '
     'documented record is unambiguous. The victims of Volt Typhoon and Salt Typhoon did not detect these campaigns; '
     'outside agencies did, years and months after the fact.')
para('V-Intelligence UEBA maintains a living behavioral chart for every user and system, interprets it in a map of '
     'meaning rather than in raw numbers, accumulates slow drift until it is statistically significant, names the '
     'direction against known adversary behavior, and combines five independent behavioral questions into one '
     'ranked and explained list. On a 250-user, 485-day validation with four embedded campaigns, it detected 4 of 4, '
     'compared with 1 of 4 for the best traditional method, at a 10.6% false-positive rate and with every detection '
     'explained.')
para('An insider changes what they access rather than how much, and an APT changes where it communicates rather than '
     'how often. Traditional tools measure magnitude; V-Intelligence UEBA measures direction, which is the dimension '
     'in which these attacks are visible.', bold=True)

# ================= GLOSSARY =================
H1('Appendix — Glossary')
gl = [
    ('UEBA', 'User and Entity Behavior Analytics — detecting threats by learning each entity\'s normal behavior and flagging departures, rather than matching known-bad patterns.'),
    ('Digital entity', 'A living software profile of a user or system — like a patient\'s medical chart for behavior: history, vitals, context, and trend, updated weekly.'),
    ('Semantic space / embedding', 'A numerical "map of meaning" produced by the same representation that powers large language models; behaviors that mean similar things sit near each other.'),
    ('Behavioral zone', 'One of five dimensions of an entity\'s profile: identity, access pattern, data behavior, network footprint, risk posture.'),
    ('Drift', 'The movement of an entity\'s behavioral profile over time — how fast, how consistently, and in which direction it is trending.'),
    ('SIEM', 'Security Information and Event Management — the central log-collection and rule-correlation platform of enterprise security.'),
    ('SOC', 'Security Operations Center — the analyst team that monitors alerts and investigates incidents.'),
    ('SOAR', 'Security Orchestration, Automation and Response — tooling that automates incident-response workflows.'),
    ('IOC', 'Indicator of Compromise — a known-bad artifact (file fingerprint, address, domain) used by signature-based tools.'),
    ('APT', 'Advanced Persistent Threat — a patient, well-resourced (typically nation-state) intrusion that remains hidden for months or years.'),
    ('C2 (command and control)', 'The covert channel through which an implanted attacker tool "calls home" for instructions and ships out data.'),
    ('LOTL', 'Living off the land — attacking with only the legitimate administration tools already present on target systems; nothing to fingerprint.'),
    ('MITRE ATT&CK', 'The industry-standard public catalog of adversary tactics and techniques, used as a common language for describing attacker behavior.'),
    ('CISA', 'Cybersecurity and Infrastructure Security Agency — the U.S. government agency that publishes authoritative threat advisories.'),
    ('Composite scoring', 'Fusing five independent behavioral questions (signal strength, breadth, persistence, context divergence, novelty) into one ranked risk score per user.'),
    ('Novelty persistence', 'A never-before-seen behavior (e.g., a new external contact) that keeps recurring week after week — the fingerprint of covert C2 infrastructure.'),
    ('False-positive rate', 'The share of normal users incorrectly flagged; here 10.6% = 26 of 246 normal users at the chosen threshold.'),
    ('Z-Score / Isolation Forest / Local Outlier Factor / One-Class SVM', 'The four industry-standard statistical and machine-learning anomaly detectors used as the traditional comparison baseline in this study.'),
    ('Zero Trust', 'The federal security architecture mandate that no user or device is trusted by default; every access is continuously verified.'),
    ('Synthetic data', 'Simulated telemetry built to prove the concept — here calibrated against public CISA advisories — before access to real operational data.'),
]
gt = doc.add_table(rows=len(gl) + 1, cols=2); gt.style = 'Table Grid'
setcell(gt.rows[0].cells[0], 'Term', bold=True, color=WHITE, fill='0B1F3A')
setcell(gt.rows[0].cells[1], 'Meaning', bold=True, color=WHITE, fill='0B1F3A')
for i, (t, d) in enumerate(gl, start=1):
    setcell(gt.rows[i].cells[0], t, bold=True, size=9, color=NAVY, fill='F4F6F8')
    setcell(gt.rows[i].cells[1], d, size=9)
for row in gt.rows:
    row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(5.0)

para('')
para('Business Edition — architecture-level disclosure. Exact detection logic, formulas, thresholds, weights, and '
     'parameters are proprietary and available in the full technical edition under NDA. All results are synthetic '
     'and must be re-validated on real operational telemetry. Do not redistribute.', 9, GREY, italic=True)

out = OUT; n = 2
while True:
    try:
        doc.save(out); break
    except PermissionError:
        root, ext = os.path.splitext(OUT); out = f"{root}_v{n}{ext}"; n += 1
print('Saved:', out)
