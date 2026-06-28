"""Business-friendly edition of the Supplier Risk SCM paper.
Keeps real technical substance (the problem, why traditional supplier-risk
scoring falls short, the cohort/drift architecture, honest measured results,
novelty) but redacts exact implementation logic (formulas, thresholds, zone
weights, accumulator parameters, fusion weights, code). Abbreviations
clarified on first use. ~8 pages.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "Supplier_Risk_SCM_Academic.docx")

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
ORG  = RGBColor(0xC2, 0x5A, 0x12)
GREY = RGBColor(0x55, 0x5F, 0x6B)
GREEN= RGBColor(0x1E, 0x7A, 0x44)
WHITE= RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.1
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(doc.sections[0], m, Inches(0.85))


def H1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    return p

def H2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ORG
    return p

def para(text, size=10.5, color=None, bold=False, italic=False, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def bullet(text, lead=None):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run(text); r2.font.size = Pt(10.5)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def setcell(cell, text, bold=False, size=9.5, color=None, fill=None):
    if fill: shade(cell, fill)
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
    if color: r.font.color.rgb = color

def table(headers, rows, widths=None, hdr_fill='0B1F3A'):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        setcell(t.rows[0].cells[j], h, bold=True, color=WHITE, fill=hdr_fill)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            setcell(t.rows[i].cells[j], v, bold=(j == 0), fill='F4F6F8' if j == 0 else None)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ================= FIGURE GENERATION =================
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

FIGDIR = os.path.join(os.path.dirname(__file__), 'figs_business_scm')
os.makedirs(FIGDIR, exist_ok=True)

C_LEG = '#8A93A0'   # traditional grey
C_LEGD = '#C0392B'  # traditional red accent
C_BEI = '#0891B2'   # BEI teal
C_NAVY = '#0B1F3A'
plt.rcParams.update({'font.family': 'Calibri', 'font.size': 10, 'axes.edgecolor': '#C9CFD6',
                     'axes.labelcolor': C_NAVY, 'xtick.color': '#55607B', 'ytick.color': '#55607B'})

def fig1_campaigns():
    """Cybersecurity proof: 4 advanced campaigns, traditional vs BEI detection."""
    camps = ['Insider\nthreat', 'Slow command-\nand-control', 'Nation-state\ncampaign #1', 'Nation-state\ncampaign #2']
    trad = [1, 0, 0, 0]
    bei = [1, 1, 1, 1]
    x = np.arange(len(camps)); w = 0.36
    fig, ax = plt.subplots(figsize=(7.0, 3.0))
    ax.bar(x - w/2, trad, w, color=C_LEG, label='Best traditional method (statistical threshold) — 1 of 4, at a 9.8% false-positive rate')
    ax.bar(x + w/2, bei, w, color=C_BEI, label='Behavioral Entity Intelligence — 4 of 4, each with a named behavioral direction')
    for xi, (tv, bv) in enumerate(zip(trad, bei)):
        ax.text(xi - w/2, tv + 0.04, 'Caught' if tv else 'Missed', ha='center', fontsize=8.8,
                color=(C_NAVY if tv else C_LEGD), fontweight='bold')
        ax.text(xi + w/2, bv + 0.04, 'Caught', ha='center', fontsize=8.8, color=C_NAVY, fontweight='bold')
    ax.set_xticks(x, camps)
    ax.set_yticks([0, 1]); ax.set_yticklabels(['Missed', 'Detected'])
    ax.set_ylim(0, 1.35)
    ax.legend(loc='upper left', fontsize=8.4, frameon=False)
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig1_campaigns.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig2_cohort():
    """Peer-cohort drift: one supplier separates while the audit threshold is never crossed."""
    rng = np.random.default_rng(11)
    months = np.arange(0, 25)
    fig, ax = plt.subplots(figsize=(7.0, 3.2))
    base = 10 + 0.04 * months
    lo, hi = base - 3.2, base + 3.2
    ax.fill_between(months, lo, hi, color=C_BEI, alpha=0.18, label='Peer-cohort normal range (similar suppliers)')
    for s in range(6):
        line = base + rng.normal(0, 0.8, len(months)).cumsum() * 0.25
        line = np.clip(line, lo + 0.4, hi - 0.4)
        ax.plot(months, line, color=C_BEI, lw=1.0, alpha=0.55)
    drift = 9.5 + np.concatenate([np.zeros(6), 0.95 * np.arange(1, 20) ** 1.05]) + rng.normal(0, 0.35, len(months))
    ax.plot(months, drift, color=C_LEGD, lw=2.4, label='Drifting supplier — separates from its cohort')
    ax.axhline(40, color='#55607B', lw=1.4, ls='--')
    ax.text(0.3, 41.2, 'Audit / contract compliance threshold — never crossed', fontsize=8.8, color='#55607B')
    ax.annotate('Flagged by cohort-relative drift\nwhile every absolute metric is in bounds',
                xy=(14, drift[14]), xytext=(3.0, 27),
                arrowprops=dict(arrowstyle='->', color=C_NAVY, lw=1.2),
                fontsize=9, color=C_NAVY, fontweight='bold')
    ax.set_xlabel('Month'); ax.set_ylabel('High-risk-country sourcing share (%)')
    ax.set_xlim(0, 24); ax.set_ylim(0, 58)
    ax.legend(loc='upper left', fontsize=8.6, frameon=True, framealpha=0.95, edgecolor='none')
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig2_cohort.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig3_model():
    """Measured model results: discrimination and survival calibration."""
    fig, axes = plt.subplots(1, 2, figsize=(7.0, 3.0))
    # (a) discrimination
    ax = axes[0]
    labels = ['Risk classifier\n(12 base signals)', 'Risk classifier\n(+ behavioral\nfeatures)', 'Survival model\n(time-to-failure\nranking)']
    vals = [0.7709, 0.7780, 0.7935]
    bars = ax.bar(np.arange(3), vals, 0.55, color=[C_BEI, C_BEI, '#0E7490'])
    for xi, v in enumerate(vals):
        ax.text(xi, v + 0.012, f'{v:.2f}', ha='center', fontsize=9.2, color=C_NAVY, fontweight='bold')
    ax.axhline(0.5, color=C_LEGD, lw=1.4, ls='--')
    ax.text(2.45, 0.515, 'Chance = 0.50', fontsize=8.4, color=C_LEGD, ha='right')
    ax.set_ylim(0, 1.0); ax.set_xticks(np.arange(3), labels, fontsize=8.2)
    ax.set_ylabel('Discrimination (1.0 = perfect)')
    ax.set_title('How well risk is ranked\n(AUC / C-index, held-out test)', fontsize=9.5, color=C_NAVY)
    # (b) survival calibration
    ax = axes[1]
    hz = ['30 days', '60 days', '90 days']
    ece = [0.084, 0.090, 0.129]
    grades = ['Grade B', 'Grade B', 'Grade C']
    bars = ax.bar(np.arange(3), ece, 0.5, color=[C_BEI, C_BEI, C_LEG])
    for xi, (v, g) in enumerate(zip(ece, grades)):
        ax.text(xi, v + 0.005, f'{v:.3f}\n{g}', ha='center', fontsize=8.8, color=C_NAVY, fontweight='bold')
    ax.set_ylim(0, 0.20); ax.set_xticks(np.arange(3), hz)
    ax.set_ylabel('Calibration error (lower = better)')
    ax.set_title('How honest the stated probabilities are\n(survival-curve calibration by horizon)', fontsize=9.5, color=C_NAVY)
    for ax in axes:
        ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig3_model.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig4_slowdrift():
    """Slow drift: every month under the alert level, cumulative trend caught."""
    rng = np.random.default_rng(5)
    months = np.arange(1, 13)
    step = np.clip(rng.normal(0.6, 0.18, 12), 0.2, 0.95)
    cum = step.cumsum()
    alert = 1.0
    fig, ax = plt.subplots(figsize=(7.0, 2.9))
    ax.bar(months, step, 0.55, color=C_LEG, label='Monthly behavioral change — every month below the per-period alert level')
    ax.axhline(alert, color=C_LEGD, lw=1.5, ls='--')
    ax.text(12.35, alert + 0.12, 'Per-period alert level\n(traditional monitoring)', fontsize=8.4, color=C_LEGD, ha='right')
    ax2 = ax.twinx()
    ax2.plot(months, cum, color=C_BEI, lw=2.4, marker='o', ms=4,
             label='Accumulated drift — the sustained trend is flagged while still forming')
    ax2.set_ylabel('Accumulated drift', color=C_BEI)
    ax2.tick_params(axis='y', colors=C_BEI)
    ax2.spines[['top']].set_visible(False)
    ax.set_xlabel('Month'); ax.set_ylabel('Monthly change')
    ax.set_ylim(0, 2.0); ax2.set_ylim(0, 8.5)
    h1, l1 = ax.get_legend_handles_labels(); h2, l2 = ax2.get_legend_handles_labels()
    ax.legend(h1 + h2, l1 + l2, loc='upper left', fontsize=8.2, frameon=False)
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig4_slowdrift.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

FIG1, FIG2, FIG3, FIG4 = fig1_campaigns(), fig2_cohort(), fig3_model(), fig4_slowdrift()

def add_fig(path, caption, width=6.4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY


# ================= TITLE =================
r = doc.add_paragraph().add_run('Supplier Risk for Defense Supply Chains:\nTraditional Monitoring vs. Behavioral Entity Intelligence')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
para('A business and technology overview of the innovation, what it detects, and what the evidence shows',
     12.5, GREY, after=2)
para('22nd Century Technologies, Inc.  ·  V-Intelligence Supply Chain Risk Program  ·  June 2026  ·  Business Edition',
     9.5, GREY, after=2)
para('Disclosure note: this edition describes the architecture and the innovation, and reports measured results. The '
     'exact model logic — formulas, detection thresholds, zone weights, accumulator settings, and fusion parameters — '
     'is proprietary and is withheld here; it is available in the full technical edition under NDA (non-disclosure '
     'agreement).', 9.5, ORG, italic=True, after=4)
para('Data note: all supply-chain results in this paper come from a synthetic (simulated) dataset built to prove the '
     'concept — 500 stock numbers, 200 suppliers, 4 depots, 24 months of transactions — and must be re-validated on '
     'real Defense Logistics Agency data before any operational claim. The cybersecurity results cited as the proven '
     'foundation come from the same framework\'s evaluation in its original security domain.', 9.5, ORG, italic=True, after=8)

# ================= 1. EXEC SUMMARY =================
H1('1.  Executive Summary')
para('The Defense Logistics Agency (DLA), the Department of Defense organization that buys and distributes spare '
     'parts and supplies, depends on thousands of suppliers. Among the most damaging supplier risks are those that '
     'no single performance metric reveals. Sourcing can shift quietly to factories in an adversary nation, '
     'counterfeit parts can enter through layers of intermediaries, and nominally independent suppliers can share '
     'one fragile sub-tier source. Such risks typically build over months and surface only after mission-critical '
     'systems are affected.')
para('Traditional supplier monitoring checks individual metrics, such as on-time delivery, quality score, and '
     'pricing, each against its own threshold. A supplier can pass every check while substantially changing what it '
     'sources and where it sources from. The reported numbers stay within bounds even as the underlying behavior '
     'changes.')
para('Behavioral Entity Intelligence (BEI) maintains a living behavioral profile, a "digital entity" comparable to '
     'a patient\'s medical chart, for every supplier, part, and depot, and it tracks the direction in which each '
     'profile is moving. Each supplier is scored against its own history and against a peer cohort of comparable '
     'suppliers, in the way a physician reads results against patients of similar age and condition.')
para('The distinction is straightforward: traditional monitoring measures how much each metric moved, whereas BEI '
     'characterizes what the supplier is becoming and identifies the drift while it is still forming.', bold=True)
bullet('In a cybersecurity evaluation, the same framework detected all four of the advanced, long-duration attack '
       'campaigns it was challenged with. The best traditional method caught one of the four, at a 9.8% '
       'false-positive rate, and offered no explanation of what had changed.', lead='Established pattern. ')
bullet('On a realistic synthetic supply-chain dataset, the supplier-risk model ranks failing suppliers well above '
       'chance (discrimination of roughly 0.77 to 0.79 on held-out tests) and states failure probabilities at the '
       '30-, 60-, and 90-day horizons with measured calibration.', lead='Measured supply-chain results. ')
bullet('Every risk score carries a plain-language reason, identifying which behavioral dimension drifted and in '
       'what direction, so that a contracting officer can act on it and defend the resulting decision.',
       lead='Explainable by construction. ')

# ================= 2. PROBLEM =================
H1('2.  The Problem: Risks No Single Metric Shows')
para('Today\'s supplier monitoring has four structural weaknesses that adversaries and market dynamics routinely exploit:')
bullet('A supplier can pass every individual check, with delivery, quality, and price all within bounds, while '
       'shifting component sourcing and substituting materials. No system monitors the combined trajectory.',
       lead='Threshold blindness. ')
bullet('Suppliers and parts are each monitored in isolation. Risk often resides in a specific combination, such as '
       'one supplier\'s quality for one particular part degrading while both appear healthy on average.',
       lead='Single-entity tunnel vision. ')
bullet('Annual risk assessments are point-in-time snapshots. A supplier rated low-risk in January can be in '
       'financial distress by March and shipping substituted components by July, all before the next scheduled '
       'review.', lead='Static risk scores. ')
bullet('Prime suppliers pass audits, but their sub-tier sourcing is effectively invisible. The risk may reside two '
       'hops away, at an unknown third-tier factory in an adversary nation.',
       lead='Sub-tier opacity. ')
para('These are not hypothetical. A 2012 Senate Armed Services Committee investigation identified roughly one '
     'million suspected counterfeit electronic parts in the Department of Defense supply chain, most entering '
     'through multi-tier distribution. Components banned under Section 889 of the National Defense Authorization Act '
     '(the prohibition on specified Chinese telecommunications equipment) have surfaced in federal systems through '
     'intermediaries that obscured their origin. The SolarWinds compromise (2020) reached roughly 18,000 '
     'organizations through a supply chain that looked legitimate at every inspection point, and COVID-19 revealed '
     '"independent" suppliers sharing one hidden sub-tier source that failed together.')
para('Supply-chain risk management (SCRM) regulations, including the counterfeit-parts rules under DFARS (the '
     'Defense Federal Acquisition Regulation Supplement), Section 889, Executive Order 14017, and the National '
     'Institute of Standards and Technology (NIST) SP 800-161 framework, all call for continuous monitoring. '
     'Current systems struggle to deliver it, and that gap is the subject of this work.')

# ================= 3. TRADITIONAL =================
H1('3.  The Traditional Approach (and Why It Falls Short)')
para('Conventional supplier monitoring is built on a simple loop: define a metric, set a threshold, alert on a '
     'breach, and re-assess each supplier on a periodic, often annual, cycle. The approach is transparent and '
     'inexpensive to operate, and it is the baseline against which BEI must be measured.')
table(['Stage', 'What it does'],
      [('1. Define metrics', 'On-time delivery rate, quality/rejection score, price compliance — each tracked separately.'),
       ('2. Set thresholds', 'Each metric gets its own red line, usually written into the contract.'),
       ('3. Periodic review', 'Each reporting period is judged on its own: is this month\'s number above the line?'),
       ('4. Annual assessment', 'A point-in-time risk rating per supplier, refreshed on a yearly cycle.'),
       ('5. Audit on exception', 'Investigation happens after a threshold breach or a field failure — that is, after the damage.')],
      widths=[1.7, 5.2])
H2('Why it is not enough')
para('Four gaps follow from this loop. First, magnitude metrics measure how much a value moved, never what kind of '
     'change occurred, yet the most damaging changes are changes of kind: the same volume, price, and on-time rate, '
     'but with components now sourced elsewhere. Second, slow drift of a small amount each month never crosses a '
     'monthly threshold, and nothing accumulates it. Third, every supplier is judged against the same population '
     'average, which is frequently the wrong comparison group. Fourth, risks that reside in a relationship, such as '
     'one supplier\'s quality for one part, or in the network, such as a shared hidden sub-tier source, fall '
     'structurally outside a per-supplier, per-metric view. Section 4 addresses each gap in turn.')

# ================= 4. BEI APPROACH =================
H1('4.  The Behavioral Entity Intelligence Approach')
H2('4.1  The digital entity and its five behavioral zones')
para('A "digital entity," sometimes called a digital twin, is a living profile of a real-world thing maintained '
     'in software. It corresponds to the patient\'s medical chart from the Executive Summary, recording not only '
     'the current reading but the accumulated history, vitals, and context whose trend allows a physician to act '
     'before a condition becomes acute. BEI maintains such a chart for every supplier, every National Stock Number '
     '(NSN, a part\'s unique identifier), every depot, and every manufacturer.')
para('Each supplier\'s chart is decomposed into five behavioral zones, independent dimensions tracked separately '
     'so that a shift in one, such as sourcing geography, is not masked by stability in another, such as strong '
     'delivery performance:')
table(['Behavioral zone', 'What it captures'],
      [('Identity', 'Registration status, certifications, contract status, tier level.'),
       ('Performance trajectory', 'On-time delivery trend, lead-time consistency, quality scores over time.'),
       ('Geographic / sourcing risk', 'Country-of-origin mix, manufacturing-site diversity, adversary-nation exposure.'),
       ('Network position', 'How many depots served and items supplied; customer concentration; sole-source contracts.'),
       ('Risk assessment', 'Financial-health indicators, insurance/bond status, expiring-contract exposure.')],
      widths=[1.9, 5.0])
para('Each period, the raw metrics of each zone are written out as a short structured description and converted '
     'into a numerical vector by an embedding model, the representational machinery inside large language models, '
     'repurposed here for supplier behavior. The result is a position in a semantic space, a mathematical "map of '
     'meaning" in which suppliers that behave alike sit near each other. The risk signal is most often a change in '
     'the kind of behavior rather than its size. A shift in where components originate barely moves the raw '
     'numbers, but it moves the supplier\'s position on the map.')

H2('4.2  Drift — direction, not just magnitude')
para('Because the profile is re-computed every period, each supplier traces a trajectory across the map. Drift '
     'magnitude answers whether something changed and by how much. Drift direction answers the question traditional '
     'monitoring cannot ask: what this supplier is turning into. The direction is then named. A library of risk '
     'concepts, including adversary-nation sourcing, counterfeit substitution, demand diversion, and financial '
     'distress, is embedded into the same map, and the drift is matched against it. The output is a plain-language '
     'reason, such as "drifting toward adversary-nation sourcing," that maps directly to the relevant regulation, '
     'such as Section 889 or the DFARS counterfeit-parts clause. An analyst can act on it, and a contracting '
     'officer can defend the resulting decision.')
para('Slow drift receives specific treatment. The most damaging shifts move by a small amount each month, below '
     'any single-period alert level. BEI maintains a running tally that accumulates the small monthly excesses and '
     'fires on the sustained trend, using a classical cumulative-sum change detector whose settings are withheld in '
     'this edition. In the technical edition\'s worked example, a twelve-month adversary-nation sourcing migration '
     'that every monthly threshold would pass is flagged within roughly the first half-year, while it is still '
     'forming.')
add_fig(FIG4, 'Figure 1 — Why slow drift defeats thresholds: every month\'s change (bars) stays below the alert '
              'level, so traditional monitoring never fires; the accumulated trajectory (line) reveals the sustained '
              'shift while it is still forming. (Illustrative.)')

H2('4.3  Peer cohorts — the right comparison group')
para('A behavioral value is anomalous only relative to comparable suppliers. A dozen international sources is normal '
     'for a semiconductor fabricator and alarming for a domestic fastener maker, and judging both against the '
     'population average would obscure both signals. BEI segments the supplier base into peer cohorts along the '
     'dimensions that define what is normal, namely what a supplier provides, where it is located, how critical it '
     'is, and its tier and size, and it standardizes each supplier\'s behavior against the normal members of its '
     'own cohort.')
para('Two baselines inform a single verdict: each supplier is scored against its own history (self-baseline drift) '
     'and against its cohort (peer-relative drift). A supplier flagged on both, drifting from its past and away '
     'from its peers, is a high-confidence risk.', bold=True)
add_fig(FIG2, 'Figure 2 — Peer-cohort drift: comparable suppliers stay inside the cohort\'s normal band (teal); one '
              'supplier separates month after month while never crossing the audit threshold (dashed). '
              'Cohort-relative scoring surfaces it; threshold monitoring never would. (Illustrative.)')
para('This cohort mechanism is not new for BEI. It is the same role-group scoring used in the operational '
     'cybersecurity deployment, in which each user is judged against others in the same role. Here suppliers '
     'substitute for users, and cohorts substitute for roles.')

H2('4.4  Relationships and the network')
para('Some risks reside between entities rather than within any one of them. BEI builds behavioral profiles for '
     'specific relationships, including a supplier\'s performance for a given part, a supplier\'s country-of-origin '
     'mix, and a part\'s bill of materials (the SBOM, a product\'s "ingredients list"), and it tracks their drift '
     'independently. A supplier and a part can each appear healthy on average while the quality of that specific '
     'combination steadily degrades, a pattern of targeted substitution visible only at the relationship level.')
para('Above the relationships sits the network view, a map of which entities depend on which. It traces supply '
     'paths through multiple tiers to find adversary-nation connections two or three hops behind a compliant prime, '
     'infers hidden common dependencies when two nominally independent suppliers move in lockstep, and propagates a '
     'hypothetical supplier failure through the graph to identify every part, depot, and mission affected.')

H2('4.5  The predictive risk layer')
para('An implemented predictive layer converts the behavioral evidence into a calibrated, time-aware risk score '
     'using two models. The first, a gradient-boosted classifier (XGBoost, a widely used machine-learning method), '
     'estimates the probability of a significant delivery failure within 90 days, constrained so that the score can '
     'move only in the expected direction: additional quality incidents can only raise risk, and a better on-time '
     'rate can only lower it. The second, a survival model that functions as an actuarial life table for suppliers, '
     'estimates the probability that a supplier is still performing at 30, 60, and 90 days, indicating when, not '
     'merely whether, to expect difficulty. Every score carries a factor-attribution readout (SHAP, short for SHapley Additive exPlanations, a standard '
     'explanation method) showing which factors raised or lowered it.')
para('Operational guardrails surround the models. A debarred supplier is set to maximum risk regardless of the '
     'model, and if the debarment check itself fails, the system fails safe to manual review. A recent change of '
     'CAGE code (Commercial and Government Entity code, a supplier\'s federal identifier) is surfaced as an '
     'ownership-change warning. Scores also incorporate destination-specific performance when a depot is specified. '
     'Exact feature definitions, label rules, and blend settings are withheld in this edition.')

# ================= 5. RESULTS =================
H1('5.  Results')
H2('5.1  The proven foundation: cybersecurity')
para('The framework is not a research proposal. In its original domain of insider-threat and intrusion detection, '
     'it was challenged with four advanced, long-duration attack campaigns deliberately designed to evade '
     'detection: a patient insider threat, a slow command-and-control intrusion, and two nation-state campaigns '
     'modeled on real adversary tradecraft. The best traditional statistical method caught one of the four, at a '
     '9.8% false-positive rate, meaning nearly one in ten normal entities was incorrectly flagged, and it provided '
     'no explanation. BEI detected all four, ranked each attacker among the highest-risk entities, and named the '
     'behavioral dimension that changed. The attackers never crossed a threshold; their behavioral direction '
     'changed. This is the same failure mode that defeats traditional supplier monitoring.')
add_fig(FIG1, 'Figure 3 — The cybersecurity proof point: four advanced campaigns, deliberately built to stay under '
              'every per-metric threshold. Traditional statistical detection caught one; Behavioral Entity '
              'Intelligence caught all four with a named behavioral direction for each. (Measured results, '
              'cybersecurity evaluation, synthetic data.)')
para('The same pattern has since been demonstrated in a second domain, defense demand forecasting, where reading '
     'the direction of each item\'s behavioral drift produced calibrated forecasts that a legacy moving average '
     'cannot match. Supplier risk is the third application of this pattern. In each domain, the damaging actor '
     'operates within authorized bounds, using valid credentials in the security case and compliant contracts in '
     'the supply case, while gradually shifting direction toward an objective.')
H2('5.2  The supply-chain proof of concept (synthetic data)')
para('The framework has been fully instantiated for the supply-chain domain and exercised on a realistic synthetic '
     'dataset: 500 National Stock Numbers, 200 suppliers, 4 depots (3 Indo-Pacific and 1 in the continental United '
     'States), and 24 months of transactions, with demand decomposed into six driver components. Demonstrated '
     'capabilities include regime detection (items flagged while transitioning from stable to depleting, before '
     'stockout), supplier health grades A through F that reflect behavioral direction rather than current-period '
     'metrics, four operating relationship types, and natural-language decision support, such as answering "Which '
     'Indo-Pacific parts are at risk of stockout in the next 90 days?"')
H2('5.3  Measured model performance')
para('The supplier-risk model was trained and audited on 1,174 supplier-evaluation windows (rolling 90-day '
     'snapshots, with features always computed strictly before the outcome window so that no future information '
     'leaks in), with a stratified 20% held-out test set of 235 windows and a failure base rate of about 20%. A '
     'window counts as a failure if the supplier\'s delivery performance collapses in the following 90 days, '
     'defined by an on-time floor and repeated severely late deliveries; the exact rule is withheld.')
table(['Model', 'Metric (plain meaning)', 'Result', 'Grade'],
      [('Risk classifier (12 base signals)', 'AUC (area under the curve) — how well failing suppliers are ranked above healthy ones (1.0 = perfect, 0.5 = chance)', '0.7709', 'B'),
       ('Risk classifier (+ behavioral features)', 'AUC', '0.7780', 'B'),
       ('Risk classifier', 'Average precision — usefulness of the top of the ranked list', '0.4918', '—'),
       ('Risk classifier', 'Calibration — do stated probabilities match reality (Brier score / expected calibration error, ECE; lower = better)', 'Brier 0.1588 / ECE 0.1019', 'C'),
       ('Survival model', 'C-index (concordance index) — how well time-to-failure is ranked', '0.7935', 'B'),
       ('Survival model', 'Calibration error at 30 / 60 / 90 days', '0.084 / 0.090 / 0.129', 'B / B / C')],
      widths=[2.0, 3.0, 1.4, 0.7])
add_fig(FIG3, 'Figure 4 — Measured supplier-risk performance on the held-out test set (synthetic data). Left: all '
              'three discrimination scores sit well above chance; adding behavioral features nudges the classifier '
              'from 0.77 to 0.78. Right: survival-curve calibration is Grade B at the 30- and 60-day horizons, '
              'Grade C at 90 days. (Measured results, synthetic data.)')
para('Read plainly, the results show discrimination that is well above chance but not perfect (Grade B), classifier '
     'calibration at Grade C that represents the main area for improvement, and survival horizons most useful for '
     'action (30 and 60 days) at Grade B. The most influential factors, namely quality incidents, recent order '
     'count, on-time rate, worst delay, and count of critical parts supplied, form an operationally sensible '
     'ordering. The behavioral features add a modest measured lift, and their distinctive value, namely directional '
     'explanation, cohort context, and slow-drift accumulation, is capability the base signals do not provide at '
     'all.')

# ================= 6. DEPLOYMENT =================
H1('6.  Data and Deployment')
bullet('BEI consumes data the enterprise already produces, including procurement transactions, delivery records, '
       'quality inspections, country-of-origin declarations, bills of materials, financial filings, and '
       'Government-Industry Data Exchange Program (GIDEP) counterfeit-and-defect alerts.',
       lead='No new instrumentation. ')
bullet('Suppliers appear under different identifiers across systems. BEI unifies them on the CAGE code and the UEI '
       '(Unique Entity Identifier) so that history follows the supplier through name and ownership changes.',
       lead='Entity resolution. ')
bullet('A supplier missing one data feed is scored on the remaining zones, with a data-coverage confidence '
       'indicator on every score. Every feature traces back to its source record, so a score that informs a '
       'contract decision can withstand challenge or audit.', lead='Graceful degradation and provenance. ')
bullet('The rollout proceeds in four phases: data integration (2 to 4 weeks), baseline building (4 to 6 weeks), '
       'detection deployment (2 to 4 weeks), and then production monitoring, with initial value targeted within 8 '
       'weeks.', lead='Phased rollout. ')
bullet('Scores feed existing workflows, including DLA\'s supply-chain risk program, contracting-officer award '
       'screens (so that a drifting supplier is flagged before award rather than after delivery), and priority '
       'inspection routing for parts from elevated-risk suppliers.', lead='Augments rather than replaces. ')

# ================= 7. NOVELTY =================
H1('7.  Why This Is New')
para('Supplier scorecards, anomaly detection, and financial-risk ratings all exist; the contribution here is not '
     'any single technique but their composition. Commercial supplier-risk platforms score current-period metrics '
     'and external risk feeds. None represents each supplier as a behavioral entity in a shared semantic space, '
     'measures the direction of its drift, names that direction against supply-chain risk concepts, scores it '
     'against the appropriate peer cohort, and accumulates sub-threshold change over months into a single '
     'explained, ranked verdict.')
para('We are explicit about prior art. The individual ingredients, including text serialization for embedding, '
     'anomaly detection in embedding space, embedding-drift measurement, cohort-relative scoring, cumulative-sum '
     'change detection, and gradient-boosted risk models, each exist separately in the literature. What no '
     'commercial or published system offers is their integration into a pipeline in which each stage removes a '
     'blind spot of the others: direction in place of magnitude, zones in place of monoliths, cohorts in place of '
     'population baselines, accumulation in place of instantaneous thresholds, and fusion in place of any single '
     'detector. The innovation is the integrated architecture, which has already been demonstrated in an adjacent '
     'domain.', bold=True)
para('Scoping notes: vendor internals are not fully public (the comparison is against marketed methodology), and a '
     'formal freedom-to-operate review is recommended before any contractual assertion of uniqueness.', 9.5, GREY,
     italic=True)

# ================= 8. LIMITATIONS =================
H1('8.  Limitations')
bullet('All supply-chain metrics are derived from synthetic data and must be re-measured on real DLA data before '
       'any production claim. That re-measurement is the explicit next step.')
bullet('Classifier calibration is currently Grade C, while the survival model\'s near horizons are Grade B. '
       'Improving it is the top model-quality priority.')
bullet('On the synthetic set, the behavioral features add only a modest lift to ranking accuracy (0.77 to 0.78). '
       'The case for them rests on directional explanation, cohort context, and slow-drift detection, which are '
       'capabilities the base signals lack entirely.')
bullet('Several relationship types (supplier-country, part-SBOM, and supplier-supplier) are designed extensions of '
       'the four that have been built and demonstrated. They reuse established machinery but are not yet exercised.')
bullet('Behavioral profiling is only as reliable as the underlying data infrastructure. Entity resolution, feed '
       'latency, and sparse data all require the engineering treatment described in Section 6.')
bullet('The cybersecurity result transfers by structural analogy. Supply-chain detection efficacy against real '
       'adversarial behavior remains to be demonstrated on operational data.')

# ================= 9. CONCLUSION =================
H1('9.  Conclusion')
para('Traditional supplier monitoring rates a supplier low-risk as long as every metric stays below its own '
     'threshold, regardless of what is changing beneath the surface. The most damaging risks exploit precisely '
     'that condition: the supplier changes what it sources rather than how much it delivers. Delivery remains '
     'strong and quality remains acceptable, yet the component origins have shifted, the sub-tier factory has '
     'changed, and the bill of materials has drifted.')
para('Behavioral Entity Intelligence tracks direction instead. It maintains a living behavioral chart per '
     'supplier, decomposed into independent zones, scored against its own history and its true peer cohort, '
     'accumulated over months, analyzed across relationships and the network, and translated into a calibrated, '
     'explainable, time-aware risk score. The pattern already caught every advanced campaign that traditional '
     'methods missed in cybersecurity, and the supply-chain instantiation now has measured results on a realistic '
     'synthetic dataset.')
para('Traditional monitoring asks whether a metric is within bounds. BEI asks what a supplier is becoming and '
     'whether it is unlike its peers, and it flags the answer while the drift is still forming. 22nd Century '
     'Technologies is prepared to conduct a 4-week proof of concept on DLA production data, against actual '
     'transaction history for a selected commodity group.', bold=True)

# ================= GLOSSARY =================
H1('Appendix — Glossary')
gl = [
    ('Behavioral Entity Intelligence (BEI)', 'Our approach: a living behavioral profile per supplier/part/depot whose direction of change — not any single metric — drives risk detection.'),
    ('Digital entity / digital twin', 'A living software profile of a real-world thing — like a patient\'s medical chart for a supplier: history, vitals, context, and trend, updated continuously.'),
    ('Behavioral zone', 'One independent dimension of an entity\'s profile (e.g. geographic sourcing), tracked separately so a shift in one is not masked by stability in others.'),
    ('Embedding', 'A numerical vector that captures the meaning of a description — the core representation inside large language models.'),
    ('Semantic space', 'A mathematical "map of meaning": suppliers whose behavior is similar sit near each other, even if their raw numbers differ.'),
    ('Drift', 'How an entity\'s behavioral profile moves through the semantic space over time — its speed, consistency, and direction.'),
    ('Peer cohort', 'The group of genuinely comparable suppliers (same commodity, geography, criticality, tier) against which a supplier\'s behavior is judged.'),
    ('Cumulative-sum change detection', 'A running tally that adds up many small period-over-period changes and flags the sustained trend that no single period would reveal.'),
    ('Regime', 'A named behavioral state (e.g. Stable, Drifting, Critical) assigned from the trajectory\'s speed and direction.'),
    ('NSN', 'National Stock Number — a part\'s unique identifier in the federal supply system.'),
    ('CAGE code / UEI', 'Commercial and Government Entity code / Unique Entity Identifier — the identifiers that resolve one supplier across federal systems.'),
    ('SCRM', 'Supply-Chain Risk Management — the discipline (and regulatory regime) of managing risks in what is sourced, from whom, and through whom.'),
    ('SBOM', 'Software/hardware Bill of Materials — the "ingredients list" of components inside a delivered product.'),
    ('Section 889 / DFARS', 'The statutory ban on specified Chinese telecom components / the Defense Federal Acquisition Regulation Supplement, including counterfeit-parts rules.'),
    ('GIDEP', 'Government-Industry Data Exchange Program — the federal alert network for counterfeit and defective parts.'),
    ('XGBoost', 'A widely used gradient-boosted machine-learning classifier; here it estimates the probability of supplier failure within 90 days.'),
    ('Survival model', 'A model of time-to-event — like an actuarial life table for suppliers: the chance a supplier is still performing at day 30, 60, 90.'),
    ('AUC / C-index', 'AUC (Area Under the Curve) and C-index: discrimination scores — how well the model ranks failing suppliers above healthy ones (1.0 = perfect, 0.5 = chance). C-index is the survival-model equivalent.'),
    ('Average precision', 'A discrimination score weighted toward the top of the ranked list — how useful the highest-risk rankings are.'),
    ('Calibration / ECE / Brier', 'Whether stated probabilities match observed reality; ECE (Expected Calibration Error) and the Brier score measure the gap — lower is better.'),
    ('SHAP', 'A standard explanation method showing which factors pushed an individual score up or down.'),
    ('False-positive rate', 'The share of normal entities incorrectly flagged — the cost of detection paid in analyst review time.'),
    ('Network graph analysis', 'Treating the supply chain as a map of who depends on whom, to find multi-tier adversary connections, hidden shared dependencies, and failure cascades.'),
    ('Sole source', 'A part with only one qualified supplier — a built-in single point of failure.'),
    ('Synthetic data', 'Simulated data built to prove the concept before access to real operational data.'),
]
gt = doc.add_table(rows=len(gl) + 1, cols=2); gt.style = 'Table Grid'
setcell(gt.rows[0].cells[0], 'Term', bold=True, color=WHITE, fill='0B1F3A')
setcell(gt.rows[0].cells[1], 'Meaning', bold=True, color=WHITE, fill='0B1F3A')
for i, (t, d) in enumerate(gl, start=1):
    setcell(gt.rows[i].cells[0], t, bold=True, size=9, color=NAVY, fill='F4F6F8')
    setcell(gt.rows[i].cells[1], d, size=9)
for row in gt.rows:
    row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(5.0)

para('')
para('Business Edition — architecture-level disclosure. Exact model logic, formulas, detection thresholds, zone '
     'weights, accumulator settings, and fusion parameters are proprietary and available in the full technical '
     'edition under NDA. All supply-chain results are synthetic and must be re-validated on real DLA operational '
     'data. Do not redistribute.', 9, GREY, italic=True)

out = OUT; n = 2
while True:
    try:
        doc.save(out); break
    except PermissionError:
        root, ext = os.path.splitext(OUT); out = f"{root}_v{n}{ext}"; n += 1
print('Saved:', out)
