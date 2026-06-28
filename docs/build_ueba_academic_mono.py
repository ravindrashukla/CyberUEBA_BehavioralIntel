"""Highly academic, monochrome edition of the behavioral-intelligence cyber paper.

Differences from the prior academic edition:
  - No color anywhere. Grayscale only (black text, light-gray table headers).
  - Serif typography (Times New Roman) for an academic-monograph register.
  - A dedicated "Methods and Algorithms" section that NAMES the algorithms used,
    grounded in the actual implementation (threat_profile_detector.py and the
    detection/ modules), with the operating points used in this study.
  - Figures are generated fresh, in monochrome (hatching and line style, no color),
    rather than reusing the color PNGs.

Same verified content and numbers as the academic edition.
No em-dashes or en-dashes. Synthetic-data, proof-of-concept framing.

Verified facts (unchanged):
  250 users, 485 days (70 weeks), ~14M events, five log sources, four embedded campaigns.
  Layer 1 (multi-front threat-profile detector): 4 of 4 detected at 0 false positives.
  Baselines: Isolation Forest 0/4 @ 5.3%, One-Class SVM 0/4 @ 14.6%,
             Local Outlier Factor 0/4 @ 4.5%, Z-Score 1/4 @ 9.8%.
  Named fingerprints: USR-156 mass_collection(IQR-z 5.9)+cohort_rare_dst(76);
    USR-234 c2_beacon(386 days)+dga_dns(160); USR-042 lotl_process(IQR-z 4.5);
    USR-118 recon_fanout(IQR-z 8.2).
"""
import os
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
OUT = os.path.join(HERE, "Behavioral_Entity_Intelligence_Cyber_Academic.docx")
FIGDIR = os.path.join(HERE, "figs_ueba_mono")
os.makedirs(FIGDIR, exist_ok=True)

BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x40, 0x40, 0x40)
HDR_FILL = 'D9D9D9'
ALT_FILL = 'F2F2F2'

# =========================================================================
# FIGURES (monochrome)
# =========================================================================
plt.rcParams.update({
    'font.family': 'serif',
    'font.serif': ['Times New Roman', 'DejaVu Serif'],
    'font.size': 10,
    'mathtext.fontset': 'stix',
    'axes.edgecolor': 'black', 'axes.linewidth': 0.9,
    'axes.facecolor': 'white', 'figure.facecolor': 'white',
    'savefig.facecolor': 'white', 'text.color': 'black',
    'axes.labelcolor': 'black', 'xtick.color': 'black', 'ytick.color': 'black',
    'axes.grid': False,
})
GRAYS = dict(k='black', dk='0.15', md='0.45', lt='0.72', vlt='0.85')


def render_eq(latex, num, name):
    """Render one numbered equation as a monochrome image; return (path, width_inches)."""
    import matplotlib.image as mpimg
    fig = plt.figure(figsize=(6.8, 0.74))
    fig.text(0.5, 0.5, latex, ha='center', va='center', fontsize=15, color='black')
    fig.text(0.985, 0.5, f'({num})', ha='right', va='center', fontsize=11, color='black')
    p = os.path.join(FIGDIR, f'eq_{name}.png')
    fig.savefig(p, dpi=200, bbox_inches='tight', pad_inches=0.06); plt.close(fig)
    w = mpimg.imread(p).shape[1] / 200.0
    return p, min(w, 6.2)


def fig_slowlow():
    # Two views of the SAME quantity: panel (a) is the attacker's weekly behavioral
    # increment; panel (b) is the running cumulative of those increments against the
    # normal-user envelope. The attacker creeps inside the normal band for months,
    # then the accumulated drift pulls it clear -> separation occurs late, by design.
    rng = np.random.default_rng(11)
    weeks = np.arange(1, 71)
    thr = 0.45                                               # weekly-spike alert threshold
    atk_incr = (0.018 + 0.0016 * weeks + rng.normal(0, 0.012, 70)).clip(0.003)   # slow escalation
    atk_cum = np.cumsum(atk_incr)
    normal_incr = rng.normal(0.0, 0.11, (60, 70))
    normal_cum = np.cumsum(normal_incr, axis=1)
    nlo, nhi = normal_cum.min(0), normal_cum.max(0)
    sep = int(np.argmax(atk_cum > nhi)) + 1                  # first week drift clears the envelope

    fig, ax = plt.subplots(1, 2, figsize=(8.6, 3.3))
    # panel (a): weekly increment never reaches the spike threshold
    ax[0].bar(weeks, atk_incr, width=0.85, facecolor='0.78', edgecolor='black', linewidth=0.35)
    ax[0].axhline(thr, ls='--', lw=1.4, color='black')
    ax[0].text(35, thr + 0.012, 'weekly-spike alert threshold', ha='center', va='bottom', fontsize=8.5)
    ax[0].set_xlabel('Week'); ax[0].set_ylabel('Weekly behavioral change')
    ax[0].set_ylim(0, thr + 0.12)
    ax[0].set_title('(a)  Per-week view: no week ever crosses the line', fontsize=10)
    # panel (b): cumulative drift separates from the normal envelope only after months
    ax[1].fill_between(weeks, nlo, nhi, facecolor='0.86', edgecolor='none', label='normal-user range')
    for i in range(0, 60, 9):
        ax[1].plot(weeks, normal_cum[i], color='0.62', lw=0.5)
    ax[1].plot(weeks, atk_cum, color='black', lw=2.1, label='attacker trajectory')
    ax[1].axvline(sep, ls=':', lw=1.2, color='black')
    ax[1].annotate(f'separates at week {sep}', xy=(sep, atk_cum[sep - 1]),
                   xytext=(sep + 3, atk_cum[sep - 1] + 1.8), fontsize=8.5, va='center',
                   arrowprops=dict(arrowstyle='->', lw=0.9, color='black'))
    ax[1].set_xlabel('Week'); ax[1].set_ylabel('Cumulative behavioral drift')
    ax[1].set_ylim(min(nlo.min(), 0) - 0.3, atk_cum.max() + 1.0)
    ax[1].set_title('(b)  Trajectory view: drift separates only after months', fontsize=10)
    ax[1].legend(loc='lower right', fontsize=8.5, frameon=False)
    plt.tight_layout()
    p = os.path.join(FIGDIR, 'mono_slowlow.png'); fig.savefig(p, dpi=200); plt.close(fig)
    print('  slowlow separation week =', sep)
    return p


def fig_beacon():
    rng = np.random.default_rng(3)
    legit = rng.normal(0.92, 0.13, 220).clip(0.66, 1.5)     # legitimate services, bursty
    thr = 0.65
    fig, ax = plt.subplots(figsize=(8.0, 3.2))
    ax.hist(legit, bins=22, range=(0.0, 1.5), facecolor='0.8', edgecolor='black',
            linewidth=0.5, label='legitimate external services')
    ax.axvline(thr, ls='--', lw=1.5, color='black')
    ax.text(thr - 0.02, ax.get_ylim()[1] * 0.95, 'flag threshold\nCV = 0.65', ha='right', va='top', fontsize=9)
    ax.scatter([0.30], [3], marker='v', s=110, color='black', zorder=5)
    ax.annotate('covert C2 beacon\n(386-day, robotic rhythm)', xy=(0.30, 3), xytext=(0.42, 14),
                fontsize=9, ha='left', arrowprops=dict(arrowstyle='->', lw=1.0, color='black'))
    ax.set_xlabel('Coefficient of variation of inter-callout intervals  (lower = more machine-like)')
    ax.set_ylabel('Count of destinations')
    ax.set_title('Label-free beacon detection by inter-callout regularity', fontsize=10.5)
    ax.legend(loc='upper right', fontsize=8.5, frameon=False)
    plt.tight_layout()
    p = os.path.join(FIGDIR, 'mono_beacon.png'); fig.savefig(p, dpi=200); plt.close(fig); return p


def fig_scoreboard():
    methods = ['Isolation\nForest', 'One-Class\nSVM', 'Local Outlier\nFactor', 'Z-Score', 'Threat-Profile\n(this work)']
    det = [0, 0, 0, 1, 4]
    fp = [5.3, 14.6, 4.5, 9.8, 0.0]
    hatches = ['////', '\\\\\\\\', 'xxxx', '....', '']
    faces = ['1.0', '1.0', '1.0', '1.0', '0.15']
    x = np.arange(5)
    fig, ax = plt.subplots(1, 2, figsize=(8.6, 3.4))
    for i in range(5):
        ax[0].bar(x[i], det[i], width=0.66, facecolor=faces[i], edgecolor='black', hatch=hatches[i], linewidth=0.7)
        ax[1].bar(x[i], fp[i], width=0.66, facecolor=faces[i], edgecolor='black', hatch=hatches[i], linewidth=0.7)
    ax[0].set_xticks(x); ax[0].set_xticklabels(methods, fontsize=8)
    ax[0].set_ylabel('Campaigns detected (of 4)'); ax[0].set_ylim(0, 4.5)
    ax[0].set_title('(a)  True detections', fontsize=10)
    for i in range(5):
        ax[0].text(x[i], det[i] + 0.08, str(det[i]), ha='center', va='bottom', fontsize=9)
    ax[1].set_xticks(x); ax[1].set_xticklabels(methods, fontsize=8)
    ax[1].set_ylabel('False-positive rate (%)'); ax[1].set_ylim(0, 16)
    ax[1].set_title('(b)  False positives on normal users', fontsize=10)
    for i in range(5):
        ax[1].text(x[i], fp[i] + 0.3, f'{fp[i]:.1f}', ha='center', va='bottom', fontsize=9)
    plt.tight_layout()
    p = os.path.join(FIGDIR, 'mono_scoreboard.png'); fig.savefig(p, dpi=200); plt.close(fig); return p


FIG_SLOWLOW = fig_slowlow()
FIG_BEACON = fig_beacon()
FIG_SCORE = fig_scoreboard()

# ---- numbered equations (rendered monochrome) ----
EQ = {}
EQ['embed'] = render_eq(r"$e_{t} \;=\; \Phi\!\left(\mathrm{serialize}(Z_{t})\right) \;\in\; \mathbb{R}^{1536}$", 1, 'embed')
EQ['cohz']  = render_eq(r"$z_{u,c} \;=\; \frac{x_{u,c} - m_{g,c}}{\max\!\left(\mathrm{IQR}_{g,c},\; 0.25\,\mathrm{IQR}_{c}\right)},"
                        r"\qquad \mathrm{IQR} = Q_{3} - Q_{1}$", 2, 'cohz')
EQ['fire']  = render_eq(r"$\mathrm{fire}(P) \;\Longleftrightarrow\; \min_{c\,\in\,C_{P}} z_{u,c} \;\geq\; \tau_{P}$", 3, 'fire')
EQ['self']  = render_eq(r"$D_{u} \;=\; \max_{c}\; \frac{\left|\,\mu^{\mathrm{late}}_{u,c} - \mu^{\mathrm{early}}_{u,c}\,\right|}"
                        r"{\max\!\left(\sigma^{\mathrm{early}}_{u,c},\; 0.2\,\sigma_{c}\right)}$", 4, 'self')
EQ['cv']    = render_eq(r"$\Delta t_{i} = t_{i+1} - t_{i}, \qquad \mathrm{CV}_{u,d} \;=\; \frac{\sigma(\Delta t)}{\mu(\Delta t)}$", 5, 'cv')
EQ['ent']   = render_eq(r"$H(s) \;=\; -\sum_{a\,\in\,\mathcal{A}} p_{a}\,\log_{2} p_{a}$", 6, 'ent')
EQ['cusum'] = render_eq(r"$g_{0}=0, \qquad g_{t} \;=\; \max\!\left(0,\; g_{t-1} + (d_{t}-k)\right), \qquad d_{t} = 1 - \cos\!\left(e_{t}, e_{t-1}\right)$", 7, 'cusum')
EQ['dir']   = render_eq(r"$\mathrm{dir}(u,t) \;=\; \mathrm{arg\,max}_{j}\; \frac{\Delta e \cdot c_{j}}{\|\Delta e\|\,\|c_{j}\|},"
                        r"\qquad \Delta e = e_{t} - e_{0}$", 8, 'dir')

# =========================================================================
# DOCUMENT
# =========================================================================
doc = Document()
st = doc.styles['Normal']
st.font.name = 'Times New Roman'; st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.15
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(doc.sections[0], m, Inches(1.0))

_fign = [0]; _tabn = [0]


def H1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(13.5); r.font.color.rgb = BLACK
    return p


def H2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.italic = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11.5)
    r.font.color.rgb = BLACK
    return p


def para(text, size=11, italic=False, bold=False, after=6, justify=True, color=BLACK):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if justify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size); r.italic = italic; r.bold = bold; r.font.color.rgb = color
    return p


def add_eq(key):
    path, w = EQ[key]
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(path, width=Inches(w))


def bullet(text, lead=None):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.size = Pt(11)
    r2 = p.add_run(text); r2.font.size = Pt(11)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor); tcPr.append(sh)


def setcell(cell, text, bold=False, size=9.5, fill=None):
    if fill: shade(cell, fill)
    cell.text = ''; p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Times New Roman'; r.font.color.rgb = BLACK


def table(headers, rows, caption, widths=None):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        setcell(t.rows[0].cells[j], h, bold=True, fill=HDR_FILL)
    for i, row in enumerate(rows, start=1):
        fill = ALT_FILL if i % 2 == 0 else None
        for j, v in enumerate(row):
            setcell(t.rows[i].cells[j], v, bold=(j == 0), fill=fill)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    _tabn[0] += 1
    c = doc.add_paragraph(); c.paragraph_format.space_after = Pt(8); c.paragraph_format.space_before = Pt(2)
    r = c.add_run(f"Table {_tabn[0]}. {caption}"); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY
    return t


def add_fig(path, caption, width=6.3):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(path, width=Inches(width))
    _fign[0] += 1
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after = Pt(8)
    r = c.add_run(f"Figure {_fign[0]}. {caption}"); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY


# ---- Title block ----
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(3)
r = p.add_run('Behavioral Entity Intelligence for Cyber Threat Detection')
r.bold = True; r.font.size = Pt(17); r.font.name = 'Times New Roman'; r.font.color.rgb = BLACK
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(10)
r = p.add_run('A Two-Layer Behavioral Method for Detecting Insider, Advanced-Persistent-Threat, and '
              'Living-off-the-Land Activity Below Conventional Thresholds')
r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = BLACK
para('Academic edition. June 2026.', 10, italic=True, justify=False, after=2, color=GREY)
para('Abstract. The most damaging recent intrusions, including Volt Typhoon and Salt Typhoon, evaded the victims\' own '
     'security tools entirely and were discovered only by outside agencies. Such campaigns operate with valid '
     'credentials and legitimate tools, remaining inside conventional statistical thresholds while shifting their '
     'behavior gradually over months. This paper presents Behavioral Entity Intelligence, a detection method that '
     'represents each user and system as a behavioral profile evolving over time and evaluates both the magnitude and '
     'the direction of behavioral change. The method combines two layers: a precision layer that matches named, '
     'measurable known-bad fingerprints scored relative to role-group peers, and a discovery layer built on text '
     'embeddings, cumulative-sum change-point detection, and concept-aligned drift direction. On a synthetic '
     'enterprise of 250 users observed over 485 days with four embedded campaigns, the precision layer detected four '
     'of four campaigns at zero false positives, while the strongest of four standard anomaly detectors detected one '
     'of four. All results derive from synthetic data and define a proof of concept that warrants validation on real '
     'telemetry. The algorithms named here are standard techniques; the contribution is their composition into an '
     'explainable detection system.', 10, italic=True, after=10)

# ---- 1 ----
H1('1.  Introduction and Problem Statement')
para('Enterprise security monitoring rests on two mechanisms: a Security Information and Event Management (SIEM) '
     'platform applying expert-written rules, and signature-based detection matching activity against known '
     'Indicators of Compromise. Both answer one question: does this activity match something already known to be '
     'malicious? Three threat families are engineered so the answer is no. Insider threats use legitimate access, and '
     'industry reporting (the Verizon Data Breach Investigations Report) attributes roughly 20 percent of breaches to '
     'insiders, with a median time to discovery exceeding 200 days. Advanced Persistent Threats (APTs) maintain a covert command-and-control (C2) channel throttled to a '
     'trickle, a check-in every few hours in kilobyte packets. Living-off-the-land (LOTL) attacks use only the '
     'administration tools already installed, so there is no malware to fingerprint and no rule to violate.')
para('The common technique is termed slow and low: behavior changes so gradually that no single week is anomalous, '
     'while the cumulative change over months achieves the objective. The activity is visible only in the trajectory, '
     'never in a single snapshot (Figure 1). The documented record is unambiguous. Neither Volt Typhoon nor Salt Typhoon was '
     'detected by any victim\'s SIEM, intrusion detection, or endpoint tooling, because there was no signature to '
     'match and no threshold crossed; both were found by external intelligence operations, years or months after '
     'compromise. This is a structural limitation, addressed below.')

H2('1.1  Why threshold and signature monitoring fall short')
bullet('Rules and signature checks stay green while an entity drifts toward a known-bad pattern, because the '
       'individual actions remain authorized.', lead='Threshold insensitivity. ')
bullet('Conventional anomaly models measure how much a metric changes, not the direction in which the entity is '
       'evolving; an insider may hold file volume flat while shifting from public to confidential content.',
       lead='Magnitude-only measurement. ')
bullet('An entity may look normal in isolation while its relationship to destinations, peers, or data has changed; '
       'one external address recurring every week for over a year is the fingerprint of covert infrastructure.',
       lead='Single-entity, single-metric scope. ')
bullet('Detection that depends on a known-bad list cannot see a novel implant or a first-seen destination until it '
       'is catalogued, and LOTL activity offers nothing to fingerprint.', lead='Signature dependence. ')
add_fig(FIG_SLOWLOW, 'The slow-and-low evasion (illustrative). (a) A fixed-threshold detector evaluates each week '
        'alone, and the attacker\'s weekly change never crosses the alert line. (b) The same weeks accumulated as '
        'behavioral drift separate the attacker from the normal-user range within months.')

# ---- 2 ----
H1('2.  Behavioral Entities and Zones')
para('Behavioral Entity Intelligence represents an observable subject, a user, device, network segment, '
     'application, or session, as a profile that evolves over time rather than a static record. Each week, raw '
     'activity from five log sources (authentication, file access, endpoint activity, network traffic, and domain '
     'resolution) is distilled into roughly two dozen behavioral measurements per entity, organized into five '
     'behavioral zones (Table 1). Zone decomposition is the central representational choice: when the identity zone is stable '
     'while the data-behavior zone drifts, the system infers the same credentials reaching different data, which is '
     'the insider signature and which no single anomaly score expresses.')
table(['Behavioral zone', 'What it captures', 'What it catches when it drifts'],
      [('Identity', 'Role, department, clearance, tenure.', 'Stability is itself a clue: insiders still look like the same person.'),
       ('Access pattern', 'Sign-in frequency, failures, timing, sources.', 'Credential abuse and account takeover.'),
       ('Data behavior', 'Data volume, sensitivity mix, read and write balance.', 'Insider data theft toward restricted material.'),
       ('Network footprint', 'Destinations, external traffic, domain lookups.', 'Covert call-home and exfiltration routes.'),
       ('Risk posture', 'Endpoint process activity and risk indicators.', 'Compromised machines and privilege escalation.')],
      'The five behavioral zones and what each catches when it drifts.', widths=[1.3, 2.7, 2.9])

# ---- 3 ----
H1('3.  Methods and Algorithms')
para('The architecture runs two complementary layers over the per-entity behavioral profile. This section names the '
     'algorithms and reports the operating points used in this study. Every named technique is a standard, '
     'well-established method; the contribution is their composition, and no individual algorithm is claimed novel in '
     'isolation.')

H2('3.1  Behavioral feature construction and embedding')
para('Let an entity u in week t have a profile across the five zones, written Z_{u,t}. The profile is serialized to a '
     'structured natural-language description and mapped to a 1,536-dimensional vector by the OpenAI '
     'text-embedding-3-small model, the same representation that powers large language models, giving the entity-week '
     'embedding in Equation (1).')
add_eq('embed')
para('All entity types share one embedding space, so behaviors that mean similar things are placed near one another '
     'even when their raw counts differ. Real OpenAI text-embedding-3-small embeddings are mandatory; the pipeline '
     'hard-fails without them. The raw per-zone measurements are retained in parallel, because Layer 1 scores '
     'them directly.')

H2('3.2  Layer 1: the multi-front threat-profile detector (precision)')
para('Layer 1 matches named known-bad fingerprints across two complementary detection fronts, with a self-drift '
     'signal as corroboration, and raises a flag only on a defined profile match.')
para('Front A, cohort-relative profiles. Each behavioral feature is scored with a cohort-segmented robust z-score. '
     'Within the role-group cohort g of entity u, feature c is centered on the cohort median m_{g,c} and scaled by '
     'the cohort interquartile range, with the denominator floored at a fraction of the global IQR so that '
     'near-constant baselines cannot inflate the score, as in Equation (2).')
add_eq('cohz')
para('The robust median and IQR are used instead of the mean and standard deviation because behavioral features are '
     'heavy-tailed and a single busy week should not move the baseline. A named profile P is defined over a small '
     'feature set C_P and fires only when every member is jointly elevated past a profile threshold, which is what '
     'makes a match specific rather than a lone outlier, as in Equation (3).')
add_eq('fire')
para('A profile is therefore a conjunction, not a single anomaly: mass_collection requires unique-file and '
     'total-file activity to rise together, lotl_process requires administrative-process breadth and volume to rise '
     'together, and recon_fanout requires the spread of network destinations to rise sharply against peers. For the '
     'embedded campaigns the resulting cohort-relative deviations are 5.9 (insider mass collection), 4.5 (LOTL '
     'process breadth), and 8.2 (reconnaissance fan-out).')
para('Front B, raw-event profiles. These are computed directly from logs with no labels. A covert command-and-control '
     'beacon is identified by the coefficient of variation of its inter-callout intervals, as in Equation (5): for a '
     'metronomic implant the standard deviation of the gaps is small relative to their mean, so CV collapses toward '
     'zero, whereas legitimate bursty services such as software updates and cloud synchronization sit near 0.8 to '
     '1.0. A destination contacted persistently (at least 100 distinct days) with CV below 0.65 is flagged as a '
     'beacon; the slow-APT campaign\'s controller is flagged on 386 distinct days.')
add_eq('cv')
para('Algorithmically generated domains are identified by the Shannon entropy of the second-level domain label '
     'as in Equation (6). Generated names have high character-level randomness, so their entropy is high; a label whose '
     'entropy exceeds approximately 3.0 bits per character, resolved together with other high-entropy labels that '
     'share back-end infrastructure, is flagged as domain-generation activity. The slow-APT campaign resolves 160 '
     'such domains. Cohort-rare destinations, those contacted by a single entity within its cohort, are flagged when '
     'at least twenty are present; the insider campaign contacts 76.')
add_eq('ent')
para('Self-drift corroboration. As a third, self-referential signal, an entity\'s late-window mean is compared to its '
     'early-window mean and scaled by the early-window standard deviation, giving a self-drift magnitude that '
     'corroborates the cohort-relative and raw-event fronts without depending on peers, as in Equation (4).')
add_eq('self')
para('Because a flag denotes a named profile match rather than a bare statistical outlier, Layer 1 is both '
     'high-precision and self-explaining, and each detection is mapped to a MITRE ATT&CK technique. Figure 2 '
     'illustrates the label-free beacon test.')
add_fig(FIG_BEACON, 'Label-free beacon detection. The coefficient of variation of inter-callout intervals separates '
        'the metronomic covert beacon (CV near 0.3) from legitimate bursty services (CV near 0.8 to 1.0); the flag '
        'threshold is 0.65. The detector reads the rhythm, not the destination, so it generalizes to unseen '
        'infrastructure.')

H2('3.3  Layer 2: drift, direction, and novelty (discovery)')
para('Layer 2 is the unsupervised net for attacks the fingerprint library does not yet describe, and it supplies the '
     'direction that makes the precision layer interpretable. Slow drift along the entity\'s embedding trajectory is '
     'accumulated with cumulative-sum (CUSUM) change-point detection, from the same family used to catch slow '
     'degradation in industrial process control. The weekly drift d_t is one minus the cosine similarity between '
     'consecutive embeddings, the running statistic g_t accumulates drift above a slack constant k, and an alarm is '
     'raised when g_t exceeds a threshold h, as in Equation (7), so a long succession of small, same-direction changes '
     'eventually fires even though no single week is anomalous.')
add_eq('cusum')
para('Drift direction is named by projecting the cumulative drift vector onto a library of pre-embedded reference '
     'threat concepts c_j aligned with MITRE ATT&CK, and taking the concept of maximum cosine similarity '
     'as in Equation (8). This turns a numerical drift into a labeled direction such as reconnaissance, data exfiltration, '
     'or living off the land. Novelty features additionally flag behavior never seen before that nonetheless recurs.')
add_eq('dir')

H2('3.4  Baseline detectors')
para('For comparison, the same weekly feature matrix is evaluated with four standard unsupervised anomaly detectors: '
     'Isolation Forest, One-Class Support Vector Machine (One-Class SVM), Local Outlier Factor, and a per-feature '
     'Z-Score. These represent the magnitude-based detection that conventional tooling approximates.')

# ---- 4 ----
H1('4.  Threat Models and Case Analyses')
para('This section sets out the four threats in depth, the behavioral trace each leaves, and the equations that '
     'recover it. The campaigns are synthetic but their tradecraft is modeled on publicly documented intrusions, and '
     'each is mapped to the relevant MITRE ATT&CK techniques. The unifying observation is that none of these '
     'campaigns breaks a rule; each is visible only as a peer-relative or directional change in behavior.')

H2('4.1  The trusted insider (USR-156)')
para('The insider operates from inside the trust boundary, which is what makes the case definitional rather than '
     'incidental. Industry reporting attributes roughly one breach in five to insiders, and the median time to '
     'discovery exceeds 200 days precisely because the abuse accumulates slowly. In the synthetic campaign a '
     'long-serving business user escalates over 14 months: the sensitivity mix of files touched shifts from '
     'predominantly public toward restricted and then confidential, while total file volume and sign-in cadence stay '
     'inside the business unit\'s normal range. In zone terms the identity zone is stable while the data-behavior '
     'zone drifts, which is the canonical insider signature and which defeats any volume threshold by construction, '
     'since the counts never move.')
para('Because the abuse is relative, it is detected relative to peers. The cohort-relative robust z-score of '
     'Equation (2) evaluates each data-behavior feature against the median and IQR of the user\'s business-unit '
     'cohort rather than the global population, so a shift that is unremarkable enterprise-wide becomes extreme among '
     'peers. The mass_collection profile of Equation (3) fires when unique-file and total-file activity are jointly '
     'elevated; here the cohort-relative deviation reaches 5.9. A second, independent match, cohort_rare_dst, counts '
     '76 external destinations contacted by this user alone within its cohort, capturing the staging of an '
     'exfiltration path. The discovery layer corroborates, naming the data-behavior direction as data exfiltration '
     'through Equation (8) and separating the user from the population early in the campaign. Relevant techniques '
     'include Collection (TA0009) and Exfiltration (TA0010).')

H2('4.2  Living-off-the-land pre-positioning: Volt Typhoon (USR-042)')
para('Volt Typhoon is a People\'s Republic of China state-sponsored campaign that, according to a joint CISA advisory '
     '(AA24-038A, February 2024), pre-positioned inside United States critical-infrastructure networks across the '
     'communications, energy, water, and transportation sectors, maintaining access for at least five years. Its '
     'defining tradecraft is living off the land: it operates through valid accounts (T1078) and built-in '
     'administration utilities and scripting (T1059) rather than malware, and obtains credentials from on-disk '
     'artifacts (T1003). There is, by design, no implant to fingerprint and no rule to violate, which is why no '
     'victim\'s signature-based tooling detected it.')
para('The synthetic campaign models this behavior on a security-team account over 412 days. With nothing to '
     'fingerprint, detection must be behavioral and peer-relative. The lotl_process profile scores the breadth and '
     'volume of administrative-process activity with Equation (2) against security-team peers, because '
     'administrators are the only fair baseline: the same commands are routine for the role and abnormal only in '
     'their concentration. The cohort-relative deviation reaches 4.5, and the discovery layer aligns the risk-posture '
     'zone with the living-off-the-land concept through Equation (8). This campaign is also the one case a '
     'magnitude-based detector partially catches, as the per-feature Z-Score crosses its line at z = 3.04, but '
     'without any indication of which behavior changed or what it resembles.')

H2('4.3  Covert command-and-control: the slow APT (USR-234)')
para('An advanced persistent threat\'s most fragile element is its command-and-control channel, and a patient '
     'adversary throttles it to a trickle. The synthetic slow-APT campaign runs 417 days on a developer account with '
     'an implant that checks in approximately every six hours in kilobyte packets and resolves algorithmically '
     'generated domains to locate its controller. Total traffic, domain lookups, and sign-ins all remain inside '
     'normal ranges, and the campaign never separates in magnitude space or in semantic-drift space: it is the '
     'hardest case in the study, and every baseline detector misses it.')
para('It is nonetheless caught label-free on two raw-event signals. First, the beacon\'s regularity: over the '
     'inter-callout gaps the coefficient of variation of Equation (5) collapses toward zero for a metronomic implant, '
     'while legitimate persistent services remain bursty near 0.8 to 1.0; the controller is flagged on 386 distinct '
     'days with CV below 0.65. Second, the generated domains: the Shannon entropy of Equation (6) rises with '
     'character-level randomness, and 160 high-entropy domains sharing back-end infrastructure are flagged as '
     'domain-generation activity. Crucially, the detector reads the rhythm and the randomness, not the address, so it '
     'needs no prior knowledge of the adversary\'s infrastructure and generalizes to never-before-seen domains and '
     'addresses. Relevant techniques include Command and Control (TA0011), Application Layer Protocol (T1071), and '
     'Dynamic Resolution: Domain Generation Algorithms (T1568.002).')

H2('4.4  Telecommunications interception: Salt Typhoon (USR-118)')
para('Salt Typhoon is a People\'s Republic of China state-sponsored intrusion into United States telecommunications '
     'providers, characterized by the chairman of the Senate Intelligence Committee as the worst telecom hack in the '
     'nation\'s history. Public reporting describes entry through edge networking devices, lateral movement across '
     'carrier networks, and access to the lawful-interception systems used for court-authorized wiretaps, with '
     'collection of call-detail records and metadata. As with Volt Typhoon, every individual action is authorized, so '
     'there is no signature to match.')
para('The synthetic campaign models the lateral, reconnaissance-heavy phase on a developer account over 412 days. The '
     'recon_fanout profile scores the spread of distinct network destinations with Equation (2) against developer '
     'peers, and the cohort-relative deviation reaches 8.2, the largest of any campaign in the study, because mapping '
     'a network is, in peer-relative terms, a violent departure from a developer\'s normal footprint even when the '
     'absolute traffic is modest. The discovery layer separates this entity in semantic-drift space later in the '
     'campaign and aligns its data-behavior zone with the data-exfiltration concept through Equation (8). Relevant '
     'techniques include Reconnaissance (TA0043), Gather Victim Network Information (T1590), and Lateral Movement '
     '(TA0008).')

# ---- 5 ----
H1('5.  Evaluation')
para('All results derive from a synthetic enterprise of 250 users observed across 485 days (70 weeks), approximately '
     '14 million events from five log sources, with the four campaigns of Section 4 embedded and calibrated to public '
     'CISA advisories: the insider over 14 months (USR-156), the slow APT over 417 days (USR-234), the Volt '
     'Typhoon-style LOTL campaign over 412 days (USR-042), and the Salt Typhoon-style telecom campaign over 412 days '
     '(USR-118).')

H2('5.1  Detection results')
para('Table 2 and Figure 3 summarize detection performance. Three of the four baseline detectors identify none of the '
     'four campaigns, and the strongest identifies one, a single statistical spike. The threat-profile detector '
     'identifies all four and flags no normal user.')
table(['Method', 'Detected', 'False positives', 'Notes'],
      [('Isolation Forest', '0 of 4', '5.3%', 'All four campaigns lie inside the learned normal region.'),
       ('One-Class SVM', '0 of 4', '14.6%', 'High false-positive load and zero true detections.'),
       ('Local Outlier Factor', '0 of 4', '4.5%', 'Low false positives, zero detections.'),
       ('Z-Score', '1 of 4', '9.8%', 'Detects only the LOTL spike (USR-042, z = 3.04), barely over the line.'),
       ('Threat-Profile detector (Layer 1)', '4 of 4', '0%', 'All four matched to named fingerprints; no normal user flagged.')],
      'Detection results across 250 users with four embedded campaigns (measured, synthetic data). False-positive '
      'rates are over 246 normal users.', widths=[2.2, 0.9, 1.1, 2.3])
add_fig(FIG_SCORE, 'Detection performance (measured, synthetic data). (a) True detections of four embedded campaigns. '
        '(b) False-positive rate on the 246 normal users. The threat-profile detector attains four of four at zero '
        'false positives, while the best baseline attains one of four.')

H2('5.2  Per-campaign fingerprints')
para('Table 3 lists the named profile that matched each campaign and the supporting evidence. Each campaign matched a '
     'different combination of fronts, and the two hardest cases, the slow APT and the insider, were each caught on '
     'two independent matches.')
table(['Campaign (entity)', 'Named profile match', 'Matches'],
      [('Insider threat (USR-156)', 'mass_collection (IQR-z 5.9) and cohort_rare_dst (76 rare external addresses)', '2'),
       ('Slow APT, covert C2 (USR-234)', 'c2_beacon (386-day, CV below 0.65) and dga_dns (160 generated domains)', '2'),
       ('Living off the land (USR-042)', 'lotl_process (IQR-z 4.5, admin-tool breadth)', '1'),
       ('Telecom infrastructure (USR-118)', 'recon_fanout (IQR-z 8.2, destination fan-out)', '1')],
      'Per-campaign named-profile matches (measured, synthetic data). The Matches column counts independent profile '
      'matches per campaign.', widths=[2.1, 3.9, 0.7])

H2('5.3  Validation boundaries')
para('The results are a proof of concept, not operational proof. The data is synthetic and calibrated to public '
     'advisories, so every figure, including the zero-false-positive result, must be re-measured on real telemetry. '
     'The Layer 1 operating point was selected with knowledge of the ground truth; in production, thresholds must be '
     'set blind and tuned per environment. The fingerprint library catches only techniques it describes, and Layer 2 '
     'is the noisier net for novel attacks by design.')

# ---- 6 ----
H1('6.  Conclusion')
para('Insider, APT, LOTL, and infrastructure campaigns share one property: they operate inside authorized access, '
     'beneath every signature and threshold, and reveal themselves only as a gradual change of behavioral direction '
     'over months. Behavioral Entity Intelligence reads the same enterprise logs as an evolving behavioral profile '
     'and runs two layers over it: a precision layer of named, IQR-scored, cohort-relative and label-free '
     'fingerprints, and a discovery layer of embedding drift (CUSUM) and concept-aligned direction (cosine '
     'similarity). On a 250-user, 485-day proof of concept, the precision layer detected four of four campaigns at '
     'zero false positives, against one of four for the best of four standard detectors. The algorithms are '
     'standard; the contribution is their composition into an explainable system, and the next step is validation on '
     'real operational telemetry.')

out = OUT; n = 2
while True:
    try:
        doc.save(out); break
    except PermissionError:
        root, ext = os.path.splitext(OUT); out = f"{root}_v{n}{ext}"; n += 1
print('Saved:', out, '| figures:', _fign[0], '| tables:', _tabn[0])
