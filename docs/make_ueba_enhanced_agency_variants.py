"""Post-process the enhanced UEBA business edition into agency-tailored variants.

Inserts a single 'Prepared for: <agency>' line after the program byline — exactly
the tailoring used for the prior _v2_<AGENCY> editions — without regenerating any
figures, so content stays byte-identical apart from that one line.
"""
import os, copy
from docx import Document
from docx.shared import Pt, RGBColor

HERE = os.path.dirname(__file__)
SRC = os.path.join(HERE, "UEBA_Behavioral_Intelligence_Enhanced_BusinessEdition.docx")
NAVY = RGBColor(0x0B, 0x1F, 0x3A)

AGENCIES = {
    "DLA":     "Prepared for: Defense Logistics Agency - Office of the CIO and PEO",
    "DISA":    "Prepared for: Defense Information Systems Agency - PEO for Cyber",
    "USSOCOM": "Prepared for: United States Special Operations Command - SDA, EIS, S&T, Innovation",
    "Army":    "Prepared for: United States Army - Office of the Chief of Artificial Intelligence",
}
BYLINE_KEY = "Business Edition (Enhanced)"


def build(agency, line):
    doc = Document(SRC)
    # locate the program byline paragraph
    anchor = next(p for p in doc.paragraphs if BYLINE_KEY in p.text)
    # new paragraph cloned from the byline's XML, then retargeted
    new_p = copy.deepcopy(anchor._p)
    anchor._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, anchor._parent)
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    para.paragraph_format.space_after = Pt(20)
    r = para.add_run(line)
    r.bold = True; r.font.size = Pt(12); r.font.name = 'Calibri'; r.font.color.rgb = NAVY
    out = os.path.join(HERE, f"UEBA_Behavioral_Intelligence_Enhanced_BusinessEdition_{agency}.docx")
    o = out; n = 2
    while True:
        try:
            doc.save(o); break
        except PermissionError:
            root, ext = os.path.splitext(out); o = f"{root}_v{n}{ext}"; n += 1
    print("Saved:", o)


if __name__ == "__main__":
    if not os.path.exists(SRC):
        raise SystemExit("Build the base edition first: python docs/build_ueba_enhanced_business.py")
    for ag, line in AGENCIES.items():
        build(ag, line)
