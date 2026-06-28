"""De-AI pass + digital-twin-platform paragraph (tracked, REV_v2).
  1) Add a 'digital twin as a platform / breadth' paragraph (user's point).
  2) Consolidate the repeated 25/29/42/67 feature-count and soften 'isolates it'.
  3) Reword the Conclusion's verbatim mission-value quartet into a natural sentence.
Honest disclosure / caveats are left untouched.
"""
import os, copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REV = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_final_REV_v2.docx"
AUTHOR = "ReviewPanel"; DATE = "2026-06-20T10:00:00Z"
d = Document(REV)
_id = [29000]
def nid():
    _id[0] += 1; return str(_id[0])
def acc(p): return "".join(x.text or "" for x in p._p.iter(qn('w:t')))
def _run(text, rpr=None):
    r = OxmlElement('w:r')
    if rpr is not None: r.append(copy.deepcopy(rpr))
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t); return r
def _wrap(tag, run):
    e = OxmlElement(tag); e.set(qn('w:id'), nid()); e.set(qn('w:author'), AUTHOR); e.set(qn('w:date'), DATE)
    e.append(run); return e

def tracked_replace(old, new, label):
    for p in d.paragraphs:
        if old in acc(p):
            run = next((r for r in p.runs if r.text and old in r.text), None)
            if run is None:
                runs = p.runs; fulltext = "".join(r.text for r in runs)
                if old not in fulltext: continue
                for t in runs[0]._r.findall(qn('w:t')): t.text = fulltext
                for r in runs[1:]:
                    for t in r._r.findall(qn('w:t')): t.text = ""
                run = runs[0]
            full = run.text; i = full.index(old)
            before, after = full[:i], full[i + len(old):]
            rpr = run._r.find(qn('w:rPr'))
            for t in run._r.findall(qn('w:t')): t.text = before
            delrun = _run(old, rpr)
            for t in list(delrun.iter(qn('w:t'))):
                dt = OxmlElement('w:delText'); dt.set(qn('xml:space'), 'preserve'); dt.text = t.text
                t.getparent().replace(t, dt)
            run._r.addnext(_run(after, rpr))
            run._r.addnext(_wrap('w:ins', _run(new, rpr)))
            run._r.addnext(_wrap('w:del', delrun))
            print(f"  [{label}] applied"); return True
    print(f"  [{label}] OLD NOT FOUND"); return False

# 1) digital-twin platform paragraph (after the 'Core operators' paragraph)
anchor = next(p for p in d.paragraphs if acc(p).startswith('Core operators include behavioral composites'))
PLAT = ("Because every entity is reduced to one consistent twin, the same foundation supports far more than the three "
        "use cases shown here. From a single data foundation the twin can carry additional features and feed several "
        "model families, including classification, survival, anomaly, graph, and forecasting models, and it extends to "
        "further decision scenarios as data allows, such as allocation, disposal and obsolescence timing, counterfeit "
        "screening, and readiness prioritization. The three use cases in this paper are the proof points; the "
        "digital-twin representation is the reusable layer beneath them, and broader scenario coverage would be "
        "validated as the framework extends to richer DLA data.")
np_ = OxmlElement('w:p'); pPr = OxmlElement('w:pPr'); rPr = OxmlElement('w:rPr')
ins = OxmlElement('w:ins'); ins.set(qn('w:id'), nid()); ins.set(qn('w:author'), AUTHOR); ins.set(qn('w:date'), DATE)
rPr.append(ins); pPr.append(rPr); np_.append(pPr); np_.append(_wrap('w:ins', _run(PLAT)))
anchor._p.addnext(np_)
print("  [1 digital-twin platform paragraph] inserted")

# 2) feature-count consolidation + soften 'isolates it'
tracked_replace(
    "The validated configuration used 29. Of the 67 features, 42 (the 23 zone-derived, 13 drift, and 6 cohort "
    "families) are generated from the embedding representation rather than from raw tabular fields; that is the "
    "embedding's specific contribution, and the tabular-only baseline (V1, 25 features) versus the embedding-enriched "
    "variant (V2, 29 features) isolates it.",
    "The measured configuration used 29 of these features; the full 67-feature architecture activates as more data "
    "becomes available. Of the engineered features, 42 (the zone, drift, and cohort families) are derived from the "
    "embedding rather than from raw tabular fields, which is the embedding's distinct contribution; a tabular-only "
    "variant (V1, 25 features) and the embedding-enriched variant (V2, 29 features) are compared to test that "
    "contribution directly.",
    "2 feature-count")

# 3) conclusion quartet -> natural sentence
tracked_replace(
    "That capability directly supports mission assurance, operational resiliency, supply-chain risk mitigation, and "
    "decision advantage in a contested logistics environment.",
    "That is the decision advantage DLA needs in a contested logistics environment: stronger readiness, earlier risk "
    "mitigation, and faster, better-informed action.",
    "3 conclusion quartet")

o = REV; k = 3
while True:
    try:
        d.save(o); break
    except PermissionError:
        root, ext = os.path.splitext(REV); o = f"{root.replace('_v2','')}_v{k}{ext}"; k += 1
print("saved:", o)
