# -*- coding: utf-8 -*-
"""Accept ALL tracked changes in the DLA 062126 FINAL and save a clean copy.
Handles: insertions/moveTo (unwrap), deletions/moveFrom (remove), paragraph-mark
deletions (merge with next paragraph), formatting-change records (drop), and turns
trackChanges off. Original file is never modified; output is a new file.
"""
from docx import Document
from docx.oxml.ns import qn

SRC="WP DLA/22nd Century DLA White BEI EDM Paper 062126 FINAL.docx"
OUT="WP DLA/22nd Century DLA White BEI EDM Paper 062126 FINAL CLEAN.docx"

doc=Document(SRC)
el=doc.element

def count(root):
    n=lambda tag: len(root.findall('.//'+qn(tag)))
    return {"ins":n('w:ins'),"del":n('w:del'),"delText":n('w:delText'),
            "moveFrom":n('w:moveFrom'),"moveTo":n('w:moveTo'),
            "pPrChange":n('w:pPrChange'),"rPrChange":n('w:rPrChange')}

before=count(el)
paras_before=len(doc.paragraphs)

# 1) insertions and moveTo -> unwrap (promote children, keep content)
for tag in ('w:ins','w:moveTo'):
    for node in el.findall('.//'+qn(tag)):
        parent=node.getparent()
        if parent is None: continue
        pos=parent.index(node)
        for child in list(node):
            parent.insert(pos,child); pos+=1
        parent.remove(node)

# 2) content deletions and moveFrom -> remove entirely (skip paragraph-mark dels in w:rPr)
for tag in ('w:del','w:moveFrom'):
    for node in el.findall('.//'+qn(tag)):
        parent=node.getparent()
        if parent is None: continue
        if parent.tag==qn('w:rPr'):  # paragraph-mark deletion, handled in step 3
            continue
        parent.remove(node)

# 3) paragraph-mark deletions: <w:pPr><w:rPr><w:del/></w:rPr></w:pPr> -> merge para with next
merged=0
for p in list(el.findall('.//'+qn('w:p'))):
    pPr=p.find(qn('w:pPr'))
    if pPr is None: continue
    rPr=pPr.find(qn('w:rPr'))
    if rPr is None: continue
    delmark=rPr.find(qn('w:del'))
    if delmark is None: continue
    rPr.remove(delmark)
    nxt=p.getnext()
    if nxt is not None and nxt.tag==qn('w:p'):
        nxt_pPr=nxt.find(qn('w:pPr'))
        insert_at=0 if nxt_pPr is None else list(nxt).index(nxt_pPr)+1
        content=[c for c in list(p) if c.tag!=qn('w:pPr')]
        for c in reversed(content):
            nxt.insert(insert_at,c)
        p.getparent().remove(p); merged+=1
    # if no mergeable next paragraph, leave the (now empty) paragraph in place

# 4) formatting-change records -> drop (keep current formatting)
for tag in ('w:pPrChange','w:rPrChange','w:tblPrChange','w:trPrChange','w:tcPrChange','w:sectPrChange'):
    for node in el.findall('.//'+qn(tag)):
        par=node.getparent()
        if par is not None: par.remove(node)

# 5) turn trackChanges OFF
st=doc.settings.element
tc=st.find(qn('w:trackChanges'))
if tc is not None: st.remove(tc)

doc.save(OUT)

# ---- verify ----
chk=Document(OUT)
after=count(chk.element)
def acc(p): return "".join(t.text or "" for t in p._p.iter(qn('w:t')))
NAMES=["Kinaxis","Blue Yonder","Interos","Everstream","Resilinc","o9"]
present=[nm for nm in NAMES if any(nm in acc(p) for p in chk.paragraphs)]
print("BEFORE:",before,"| paragraphs:",paras_before)
print("AFTER :",after,"| paragraphs:",len(chk.paragraphs),"| paragraphs merged:",merged)
print("competitor names still present in clean copy:",present or "NONE")
print("trackChanges setting present:",chk.settings.element.find(qn('w:trackChanges')) is not None)
print("saved:",OUT)
