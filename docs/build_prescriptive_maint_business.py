"""Business-friendly edition of the Prescriptive Maintenance paper.
Keeps real technical substance (the maintenance ladder, the entity/zone model,
zone divergence, regimes, supply-chain integration, honest benchmark evidence,
novelty) but redacts exact implementation logic (drift mathematics, zone
composition, regime criteria, thresholds, code). Abbreviations clarified on
first use. ~8 pages.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "Prescriptive_Maintenance_Academic.docx")

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
ORG  = RGBColor(0xC2, 0x5A, 0x12)
GREY = RGBColor(0x55, 0x5F, 0x6B)
GREEN= RGBColor(0x1E, 0x7A, 0x44)
WHITE= RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.1
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(doc.sections[0], m, Inches(0.85))


def H1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    return p

def H2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ORG
    return p

def para(text, size=10.5, color=None, bold=False, italic=False, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def bullet(text, lead=None):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run(text); r2.font.size = Pt(10.5)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def setcell(cell, text, bold=False, size=9.5, color=None, fill=None):
    if fill: shade(cell, fill)
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
    if color: r.font.color.rgb = color

def table(headers, rows, widths=None, hdr_fill='0B1F3A'):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        setcell(t.rows[0].cells[j], h, bold=True, color=WHITE, fill=hdr_fill)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            setcell(t.rows[i].cells[j], v, bold=(j == 0), fill='F4F6F8' if j == 0 else None)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ================= FIGURE GENERATION =================
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

FIGDIR = os.path.join(os.path.dirname(__file__), 'figs_business_maint')
os.makedirs(FIGDIR, exist_ok=True)

C_LEG = '#8A93A0'   # traditional grey
C_LEGD = '#C0392B'  # traditional red accent
C_BEI = '#0891B2'   # BEI teal
C_NAVY = '#0B1F3A'
plt.rcParams.update({'font.family': 'Calibri', 'font.size': 10, 'axes.edgecolor': '#C9CFD6',
                     'axes.labelcolor': C_NAVY, 'xtick.color': '#55607B', 'ytick.color': '#55607B'})

def fig1_relationship():
    """Why thresholds miss it: each sensor in spec, but their relationship drifts."""
    wk = np.arange(0, 13)
    rng = np.random.default_rng(3)
    temp = 165 + 1.0 * wk + rng.normal(0, 1.6, len(wk))          # threshold 200 F
    vib = 0.28 + 0.008 * wk ** 1.25 + rng.normal(0, 0.005, len(wk))  # threshold 0.5 g
    fig, axes = plt.subplots(1, 2, figsize=(7.0, 3.0))
    # (a) the sensors a CBM system watches
    ax = axes[0]
    ax.plot(wk, temp, color=C_LEG, lw=2.0, label='Temperature (°F)')
    ax.axhline(200, color=C_LEGD, lw=1.4, ls='--')
    ax.text(0.2, 202.0, 'Temperature alarm 200°F — never crossed', fontsize=8.0, color=C_LEGD)
    ax2 = ax.twinx()
    ax2.plot(wk, vib, color='#55607B', lw=2.0, ls='-.', label='Vibration (g)')
    ax2.axhline(0.5, color=C_LEGD, lw=1.4, ls=':')
    ax2.text(12.2, 0.488, 'Vibration alarm 0.5g —\nnever crossed', fontsize=8.0, color=C_LEGD,
             ha='right', va='top')
    ax2.set_ylim(0.2, 0.64); ax2.set_ylabel('Vibration (g)', color='#55607B')
    ax.set_ylim(150, 215); ax.set_xlabel('Week'); ax.set_ylabel('Temperature (°F)', color=C_LEG)
    ax.set_title('What traditional monitoring sees:\neach sensor stays within its limit', fontsize=9.5, color=C_NAVY)
    ax.spines[['top']].set_visible(False); ax2.spines[['top']].set_visible(False)
    # (b) the relationship between them
    ax = axes[1]
    rel = (vib / 0.28) / (temp / 165.0)
    base = np.ones(len(wk))
    ax.fill_between(wk, 0.96, 1.04, color=C_LEG, alpha=0.30)
    ax.text(11.8, 0.975, 'normal band for this unit', fontsize=8.2, color='#55607B', ha='right')
    ax.plot(wk, base, color=C_LEG, lw=1.2, ls='--')
    ax.plot(wk, rel, color=C_BEI, lw=2.4, marker='o', ms=3.5,
            label='Vibration-to-temperature relationship')
    ymax = rel.max() + 0.10
    ax.text(0.4, ymax - 0.02, 'steady behavioral drift —\nthe failure precursor,\nvisible weeks early',
            fontsize=8.6, color=C_NAVY, fontweight='bold', va='top')
    ax.annotate('', xy=(9.2, np.interp(9.2, wk, rel) + 0.01), xytext=(5.0, ymax - 0.05),
                arrowprops=dict(arrowstyle='->', color=C_NAVY, lw=1.2))
    ax.set_xlabel('Week'); ax.set_ylabel('Relationship index (baseline = 1.0)')
    ax.set_ylim(0.9, ymax)
    ax.set_title('What BEI sees: the relationship\nbetween the same two sensors', fontsize=9.5, color=C_NAVY)
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig1_relationship.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig2_zones():
    """Zone divergence: all-zones aging vs one-zone failure signature."""
    wk = np.arange(0, 13)
    rng = np.random.default_rng(11)
    zones = ['Vibration', 'Thermal', 'Power', 'Fluid chemistry', 'Duty cycle']
    fig, axes = plt.subplots(1, 2, figsize=(7.0, 3.0), sharey=True)
    # (a) all zones drift together: normal aging
    ax = axes[0]
    for z in zones:
        d = 0.012 * wk + rng.normal(0, 0.006, len(wk))
        ax.plot(wk, d, color=C_LEG, lw=1.6, alpha=0.85)
    ax.text(6.0, 0.30, 'all five zones move together\n→ normal aging\n→ scheduled overhaul, no early parts buy',
            fontsize=8.6, color='#55607B', ha='center')
    ax.set_title('Pattern A — uniform drift:\nexpected wear', fontsize=9.5, color=C_NAVY)
    ax.set_xlabel('Week'); ax.set_ylabel('Drift from this unit\'s baseline')
    # (b) one zone diverges: specific component degrading
    ax = axes[1]
    for z in zones[1:]:
        d = 0.010 * wk + rng.normal(0, 0.006, len(wk))
        ax.plot(wk, d, color=C_LEG, lw=1.6, alpha=0.85)
    dvib = 0.004 * wk ** 1.85 + rng.normal(0, 0.008, len(wk))
    ax.plot(wk, dvib, color=C_BEI, lw=2.6, marker='o', ms=3.5, label='Vibration zone')
    ax.annotate('vibration zone diverges, others stable\n→ bearing wear diagnosed\n→ order the bearing kit now',
                xy=(11, dvib[11]), xytext=(1.0, 0.46),
                fontsize=8.6, color=C_NAVY, fontweight='bold',
                arrowprops=dict(arrowstyle='->', color=C_NAVY, lw=1.2))
    ax.set_title('Pattern B — one zone diverges:\nspecific failure forming', fontsize=9.5, color=C_NAVY)
    ax.set_xlabel('Week')
    for ax in axes:
        ax.spines[['top', 'right']].set_visible(False); ax.set_ylim(-0.04, 0.62)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig2_zones.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig3_regimes():
    """Regime ladder and the lead-time gap vs a threshold alarm."""
    wk = np.linspace(0, 20, 400)
    drift = 0.018 * wk + 0.0032 * np.maximum(wk - 8, 0) ** 2.2
    fig, ax = plt.subplots(figsize=(7.0, 3.2))
    bands = [(0, 6, '#EDF1F5', 'STABLE'), (6, 12, '#DCE8EF', 'DRIFTING'),
             (12, 17, '#C7E3EC', 'ACCELERATING'), (17, 20, '#F6DCD7', 'REGIME SHIFT')]
    for x0, x1, c, lab in bands:
        ax.axvspan(x0, x1, color=c)
        ax.text((x0 + x1) / 2, 1.06, lab, ha='center', fontsize=8.4, color=C_NAVY, fontweight='bold')
    ax.plot(wk, drift, color=C_BEI, lw=2.6, label='Behavioral drift from baseline (BEI)')
    ax.axhline(0.85, color=C_LEGD, lw=1.6, ls='--')
    ax.text(0.4, 0.875, 'Traditional sensor alarm threshold', fontsize=8.4, color=C_LEGD)
    x_alarm = wk[np.argmax(drift >= 0.85)]
    ax.plot([x_alarm], [0.85], 'o', color=C_LEGD, ms=7)
    ax.text(14.6, 0.91, 'threshold alarm —\nfailure already forming', fontsize=8.4, color=C_LEGD, ha='center')
    ax.plot([7.0], [np.interp(7.0, wk, drift)], 'o', color=C_BEI, ms=7)
    ax.text(7.0, 0.32, 'BEI flags DRIFTING regime —\ninspection scheduled,\nparts availability verified',
            fontsize=8.4, color=C_NAVY, ha='center', fontweight='bold')
    ax.annotate('', xy=(x_alarm, 0.70), xytext=(7.0, 0.70),
                arrowprops=dict(arrowstyle='<->', color=C_NAVY, lw=1.4))
    ax.text((7.0 + x_alarm) / 2, 0.74, 'weeks of lead time → parts pre-positioned', ha='center',
            fontsize=8.8, color=C_NAVY, fontweight='bold')
    ax.set_xlabel('Week'); ax.set_ylabel('Drift from baseline')
    ax.set_xlim(0, 20); ax.set_ylim(0, 1.13)
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig3_regimes.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig4_benchmarks():
    """Documented predictive-maintenance benchmark ranges (industry, not BEI-measured)."""
    rows = [('Unplanned failures', 40, 60, '% reduction'),
            ('Total maintenance cost', 25, 40, '% reduction'),
            ('Safety-stock value', 15, 25, '% reduction'),
            ('Operational Readiness rate', 3, 8, 'percentage-point gain'),
            ('Parts availability lead time', 2, 4, 'weeks earlier')]
    fig, ax = plt.subplots(figsize=(7.0, 2.9))
    y = np.arange(len(rows))[::-1]
    for yi, (lab, lo, hi, unit) in zip(y, rows):
        ax.barh(yi, hi - lo, left=lo, height=0.52, color=C_BEI, alpha=0.85)
        ax.plot([lo, hi], [yi, yi], color=C_NAVY, lw=0, marker='|', ms=14, mew=2)
        sep = '' if unit.startswith('%') else ' '
        ax.text(hi + 1.2, yi, f'{lo}–{hi}{sep}{unit}', va='center', fontsize=8.8,
                color=C_NAVY, fontweight='bold')
    ax.set_yticks(y, [r[0] for r in rows])
    ax.set_xlim(0, 80); ax.set_xlabel('Documented improvement range')
    ax.set_title('Predictive / condition-based maintenance program outcomes — defense and industry benchmarks\n'
                 '(the opportunity a behavioral approach targets; not BEI-measured results)',
                 fontsize=9.5, color=C_NAVY)
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig4_benchmarks.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

FIG1, FIG2, FIG3, FIG4 = fig1_relationship(), fig2_zones(), fig3_regimes(), fig4_benchmarks()

def add_fig(path, caption, width=6.4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY


# ================= TITLE =================
r = doc.add_paragraph().add_run('Prescriptive Maintenance for Defense Platforms:\nThe Traditional Approach vs. Behavioral Entity Intelligence')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
para('A business and technology overview of the innovation, its intended outcomes, and the supporting evidence base',
     12.5, GREY, after=2)
para('22nd Century Technologies, Inc.  ·  V-Intelligence Program  ·  Prepared for DLA | US Army | USSOCOM  ·  June 2026  ·  Business Edition',
     9.5, GREY, after=2)
para('Disclosure note: this edition describes the architecture and the innovation at a conceptual level. The exact '
     'model logic — the drift mathematics, zone composition, regime-classification criteria, thresholds, and '
     'adjustment parameters — is proprietary and is withheld here; it is available in the full technical edition '
     'under NDA (non-disclosure agreement).', 9.5, ORG, italic=True, after=4)
para('Evidence note: this is a concept paper. The behavioral framework it applies is proven in an operational '
     'cybersecurity deployment and applied in defense demand forecasting; no maintenance-specific proof of concept '
     'has yet been conducted. All improvement ranges quoted are documented benchmarks from predictive- and condition-based-'
     'maintenance programs across defense and industry — they describe the opportunity, not measured BEI results.',
     9.5, ORG, italic=True, after=8)

# ================= 1. EXEC SUMMARY =================
H1('1.  Executive Summary')
para('When a generator fails at a Forward Operating Base (a deployed military outpost), when a Stryker armored '
     'vehicle\'s drivetrain seizes during a tactical movement, or when an F-35 engine needs unplanned work during '
     'sustained operations, the cost extends well beyond the repair bill to mission capability, operational tempo, '
     'and the safety of the people who depend on that equipment. Unplanned maintenance typically costs 3–10 times the '
     'equivalent planned maintenance, and according to Government Accountability Office (GAO) analyses, roughly 40% '
     'of maintenance delays are parts-related: the mechanic knows what is broken but the part is not in stock at the '
     'right location.')
para('Current best practice, condition-based maintenance (CBM), watches each sensor against a fixed limit: '
     'temperature must stay below 200°F, vibration below 0.5g, oil pressure above 40 PSI (pounds per square inch). '
     'Equipment, however, usually fails because the relationship between readings changes rather than because any '
     'single reading crosses a line. A turbine can run comfortably below every threshold while its '
     'vibration-to-temperature relationship drifts steadily for weeks, and no alarm fires. The failure then appears '
     'to arrive without warning, although the warning was present throughout in a place threshold monitoring is not '
     'designed to look.')
para('Behavioral Entity Intelligence (BEI) assigns every piece of equipment, identified by its individual serial '
     'number rather than its model alone, a digital entity: a behavioral profile analogous to a patient\'s medical '
     'chart. A chart records more than today\'s temperature; it accumulates history and context, and a clinician '
     'reads its trend to act before a condition becomes acute. BEI reads how each unit behaves across five behavioral '
     'zones, tracks how that behavior moves over time, and classifies each unit into one of four health regimes. It '
     'then closes the loop with the supply chain: a forming failure informs the Defense Logistics Agency (DLA) of '
     'what part will be needed, when, and where to stage it, well before the work order exists.')
para('In summary, where traditional maintenance asks whether any reading is out of specification, BEI asks whether a '
     'specific unit\'s behavior is changing direction, and converts the answer into an early maintenance action and '
     'parts shipment.', bold=True)
bullet(' Relationship drift between sensors is the failure precursor, and BEI is designed to surface it weeks before '
       'any individual threshold trips.', lead='Early warning. ')
bullet(' Which zone drifts, and which stay stable, identifies the specific failing component and therefore the '
       'specific part to order.', lead='Specific diagnosis. ')
bullet(' Predicted failures become demand signals to DLA, converting emergency orders into planned ones. Documented '
       'predictive-maintenance programs report 40–60% fewer unplanned failures and 25–40% lower maintenance cost.',
       lead='Supply chain integration. ')

# ================= 2. THE MAINTENANCE LADDER =================
H1('2.  The Problem: Climbing the Maintenance Ladder')
para('Maintenance philosophy has climbed a ladder over decades, each rung answering a better question:')
table(['Rung', 'Question it answers', 'In plain terms'],
      [('Reactive', '"What broke?"', 'Run equipment until it fails, then fix it. Cheapest to set up and most expensive to operate, given emergency repairs, secondary damage, and lost missions.'),
       ('Preventive', '"What is due?"', 'Replace parts on a fixed schedule (hours, miles, calendar) whether they need it or not. Avoids some failures, wastes healthy parts, still misses early failures.'),
       ('Predictive (CBM)', '"What is out of spec?"', 'Monitor actual condition with sensors and act when a reading crosses a limit. The current military best practice, though it watches readings one at a time.'),
       ('Prescriptive (BEI)', '"What is changing, what do I do, and what do I ship?"', 'Read behavioral direction, diagnose the specific failing component, schedule the fix, and pre-position the parts — all before failure.')],
      widths=[1.4, 1.9, 3.6])
para('This paper addresses the step from the third rung to the fourth: from observing that a reading has crossed a '
     'line to determining that a unit is heading toward a specific failure, identifying the corresponding action, '
     'and setting the required part in motion.')
H2('Why condition-based maintenance misses failures')
bullet(' Each sensor is checked against its own limit, independently. Failures caused by a change in how readings '
       'relate to each other never trip an alarm.', lead='Individual thresholds. ')
bullet(' In a turbine engine, the relationship between rotation speed, exhaust gas temperature, and fuel flow '
       'predicts compressor degradation; in a ground vehicle, transmission temperature versus vibration versus speed '
       'predicts gearbox failure. No single value is abnormal — only the pattern is.', lead='Correlation blindness. ')
bullet(' Manufacturer thresholds describe a generic unit in reference conditions. A generator that has run 18 '
       'months at high altitude develops different thermal characteristics than the same model at sea level; a '
       'population-level limit cannot see unit-specific degradation.', lead='Static baselines. ')
bullet(' A bearing whose vibration grows by 0.002g per week will not cross a threshold for months — yet that steady '
       'climb is a reliable failure predictor. Point-in-time checks cannot see a trajectory.', lead='Trend blindness. ')
add_fig(FIG1, 'Figure 1 — Why threshold monitoring misses the failure. Left: both sensors stay safely inside their '
              'alarm limits for twelve weeks. Right: the relationship between the same two sensors drifts steadily '
              'out of this unit\'s normal band — the behavioral precursor BEI is built to detect. (Illustrative.)')

# ================= 3. THE BEI APPROACH =================
H1('3.  The Behavioral Entity Intelligence Approach')
H2('3.1  Every unit becomes a digital entity')
para('A "digital entity" (a form of digital twin) is a software profile of one specific physical unit. Two '
     'F-35 engines of identical manufacture develop different profiles because each accumulates its own operating '
     'history, environmental exposure, maintenance record, and duty cycle. The profile resides in a semantic space, '
     'a mathematical "map of meaning" of the same kind that underlies large language models, in which '
     'units that behave alike sit near each other even when their raw sensor numbers differ. A language model uses '
     'that space to interpret sentences; BEI uses it to interpret how equipment is behaving.')
para('Each entity\'s behavior is decomposed into five zones, each an independent lens on the same machine watching a '
     'different dimension of health:')
table(['Behavioral zone', 'What it captures (in plain terms)'],
      [('Vibration profile', 'The mechanical health of rotating parts — bearing wear, shaft misalignment, blade or gear damage, read from how the machine shakes.'),
       ('Thermal signature', 'How heat moves through the machine — coolant restrictions, blocked exhaust paths, failing insulation, read from temperatures and heating/cooling rates.'),
       ('Power / performance', 'How efficiently fuel becomes output — compressor degradation, fouled injectors, winding wear, read from consumption and output ratios.'),
       ('Fluid chemistry', 'What the oil and coolant reveal — metal particles in the oil point to the specific component shedding them.'),
       ('Operational duty cycle', 'How hard the unit is being worked — load patterns, start/stop frequency, environmental exposure. Hard-run units age differently.')],
      widths=[1.8, 5.1])
H2('3.2  Zone divergence — the key diagnostic signal')
para('The most informative signal is not how far any one zone has drifted but the pattern across zones, much as a '
     'clinician reads one vital sign moving while the others hold steady and infers which organ to examine. When all '
     'five zones drift together, the unit is aging normally; the appropriate response is the scheduled overhaul, and '
     'no part needs to be bought early. When one zone diverges while the rest stay stable, a specific component is '
     'degrading, and the drifting zone identifies it. Vibration drifting alone points to bearing wear and a bearing '
     'kit; thermal drifting alone points to the fluid system, such as a coolant pump or heat exchanger; power '
     'drifting alone points to fuel injectors or electrical windings. Vibration and thermal drifting together is the '
     'pattern of greatest concern, indicating a cascading failure in which mechanical wear is producing secondary '
     'heat damage; the appropriate response is to inspect within two weeks and order the full kit immediately.')
add_fig(FIG2, 'Figure 2 — Zone divergence. Left: all five zones drift together — normal aging, handled by the '
              'scheduled overhaul. Right: the vibration zone diverges while the others hold — bearing wear, with the '
              'specific replacement part identified weeks before failure. (Illustrative.)')
H2('3.3  Trajectory, regimes, and remaining useful life')
para('BEI tracks each entity\'s movement over time: how fast its behavior is drifting from baseline (velocity), '
     'whether that rate is itself increasing (acceleration, the strongest predictor of imminent failure, since '
     'wear creates secondary stress that compounds), and whether the drift is steady (progressive wear) or '
     'oscillating (an intermittent fault such as a loose connection or sticking valve). The underlying mathematics is '
     'proprietary and is withheld in this edition. The operationally relevant output is the classification: every '
     'unit is assigned to one of four regimes, and each regime drives both a maintenance response and a supply-chain '
     'action.')
table(['Regime', 'What it means', 'Maintenance response', 'Supply chain action'],
      [('STABLE', 'Behavior inside the unit\'s normal envelope.', 'Normal operations; maintain at the next scheduled interval.', 'No action; standard inventory suffices.'),
       ('DRIFTING', 'Steady movement away from baseline.', 'Schedule an inspection; monitor weekly.', 'Verify parts availability; flag for procurement.'),
       ('ACCELERATING', 'Degradation is speeding up.', 'Maintain within two weeks; restrict operating tempo if possible.', 'Order parts; pre-position at the nearest depot; expedite if lead time exceeds the window.'),
       ('REGIME SHIFT', 'A fundamental break in behavior.', 'Immediate inspection; take the unit out of service ("ground" or "deadline" it) until checked.', 'Emergency parts request; pull stock from other locations if needed.')],
      widths=[1.2, 1.8, 2.0, 1.9])
para('From the trajectory, BEI estimates remaining useful life (RUL), the time until the unit needs intervention, '
     'with explicit uncertainty: a steadily drifting unit receives a tight estimate and an oscillating one a wider '
     'band. Unlike a manufacturer\'s generic cycle count ("this engine type lasts N flight hours"), the estimate is '
     'grounded in the unit\'s own trajectory and validated against fleet peers, namely other units whose behavioral '
     'path was similar and the outcomes those units experienced.')
add_fig(FIG3, 'Figure 3 — The four health regimes and the lead-time gap. A traditional threshold alarm fires only '
              'when a reading finally crosses its limit; BEI flags the DRIFTING regime weeks earlier, when the '
              'inspection can be scheduled and the parts pre-positioned. (Illustrative; the traditional alarm '
              'moment is projected onto the same timeline for comparison.)')
H2('3.4  Beyond single units: fleet, environment, and supplier intelligence')
para('Because every entity lives in the same behavioral space, BEI can reason across the fleet, not just within one '
     'machine:')
bullet(' Units with similar profiles cluster into cohorts. If 10% of a cohort has already suffered a specific '
       'failure, the remainder are at elevated risk, which supports fleet-wide preventive action drawn from peer '
       'evidence. For example, 15 UH-60 helicopter engines across 3 installations developing similar vibration drift '
       'points to a common component batch, triggering inspection of every affected unit before the first failure.', lead='Fleet cohorts. ')
bullet(' The same model degrades differently in desert heat, maritime salt air, or high altitude, and under convoy '
       'duty versus intermittent patrols. BEI learns each unit\'s context instead of assuming a reference condition.',
       lead='Environment. ')
bullet(' If parts from one supplier are consistently associated with faster behavioral drift across the fleet, BEI '
       'surfaces a supplier-quality signal that feeds procurement and vendor evaluation, making quality erosion '
       'visible well before failure-rate statistics move.', lead='Supplier quality. ')
H2('3.5  The supply chain link — what makes it prescriptive')
para('Predicting failure without ensuring parts availability achieves little, since the maintainer knows the '
     'equipment will fail yet still cannot prevent it. BEI\'s behavioral predictions therefore generate three demand '
     'signals directly into DLA operations:')
table(['Signal', 'Where it comes from', 'What DLA does with it'],
      [('WHAT part', 'The zone-divergence pattern maps to specific components (thermal drift → coolant pump; vibration drift → bearing kit).', 'Identify the National Stock Number (NSN — the part\'s unique catalog ID) and check inventory across depots.'),
       ('WHEN needed', 'The regime and remaining-useful-life estimate set the window (DRIFTING = longer horizon; ACCELERATING = near-term).', 'Align procurement lead time to the predicted need date; order now if lead time exceeds the window.'),
       ('WHERE staged', 'The unit\'s current and projected location from Army logistics systems.', 'Pre-position at the depot or supply point closest to the unit; pull from excess stock elsewhere if needed.')],
      widths=[1.2, 2.9, 2.8])
para('The loop then closes: completed maintenance resets the unit\'s behavioral baseline, and BEI monitors the '
     'post-repair profile to confirm the repair was effective. Each cycle further trains the system and converts '
     'unplanned failure demand into planned maintenance demand, which is the largest single lever on emergency '
     'procurement cost and expedited shipping.')

# ================= 4. EVIDENCE =================
H1('4.  What the Evidence Shows')
para('Two points should be distinguished carefully. First, the behavioral framework itself is not hypothetical: the '
     'same entity-profile, zone-divergence, and trajectory machinery is proven in an operational cybersecurity '
     'deployment, where it isolates threat signatures, and is applied in defense demand forecasting. Second, it has '
     'not yet been applied to maintenance data. No maintenance proof of concept has been conducted, and this paper '
     'claims no BEI-measured maintenance results. What the evidence base establishes is the scale of the '
     'opportunity: documented outcomes from predictive- and condition-based-maintenance programs across defense and '
     'industry, as reported in Department of Defense (DoD), Department of Energy, and commercial reliability '
     'studies.')
table(['Impact area', 'Today (threshold-based practice)', 'Documented benchmark range*'],
      [('Unplanned failures', 'Detected at the point of occurrence; no advance warning; emergency maintenance.', '40–60% reduction across a monitored fleet'),
       ('Maintenance cost', 'Unplanned work costs 3–10x the planned equivalent (expedited parts, overtime, secondary damage).', '25–40% reduction in total maintenance cost per platform'),
       ('Mission readiness', 'Readiness rates capped by unplanned downtime.', '3–8 percentage-point improvement in Operational Readiness rate'),
       ('Parts inventory', 'High safety stock held against demand uncertainty.', '15–25% reduction in safety-stock value without service degradation'),
       ('Supply responsiveness', 'Supply chain reacts to work orders after failure.', '2–4 week improvement in parts availability lead time')],
      widths=[1.5, 2.9, 2.5])
para('*Benchmark ranges from predictive- and condition-based-maintenance programs across defense and industry; not '
     'BEI-measured results.', 9, GREY, italic=True)
add_fig(FIG4, 'Figure 4 — The documented opportunity. Outcome ranges reported by predictive- and condition-based-'
              'maintenance programs in defense and industry. These are the benchmarks a behavioral approach targets — '
              'not BEI-measured results; BEI-specific performance would be established in a proof of concept.')
para('The per-platform stakes give those percentages concrete meaning. A single unplanned aircraft-engine event runs '
     '$150K–$500K (an engine swap means 72–96 hours of grounded aircraft); naval propulsion $200K–$1M+; rotary wing '
     '(UH-60, CH-47) $100K–$400K; ground vehicles $25K–$100K; communications and radar $30K–$150K; generators '
     '$10K–$40K. Illustrative annual savings derived from published DoD cost structures range from $5K–$20K per '
     'generator to $100K–$500K per naval propulsion plant; these figures indicate scale and are not projected BEI '
     'results.')

# ================= 5. DEPLOYMENT =================
H1('5.  Deployment Requirements')
bullet(' BEI uses data the military already collects: platform sensor feeds, maintenance logs (GCSS-Army — the '
       'Army\'s logistics system — and its Air Force equivalents), oil-analysis lab results, usage and deployment '
       'records, weather and location data, and DLA parts history. Initial deployment requires no new sensors.',
       lead='No new instrumentation. ')
bullet(' Behavioral profiling runs at the platform or base level ("edge"), so regime-shift alerts work even with '
       'limited connectivity; fleet-wide cohort analysis and demand-signal generation run at the enterprise level.',
       lead='Edge plus enterprise. ')
bullet(' Connects to existing systems — GCSS-Army for maintenance and supply, DLA\'s Enterprise Business System for '
       'procurement — through standard DoD data interfaces. Nothing is replaced.', lead='Integrates, not replaces. ')
bullet(' Four phases: connect data and register entities (2–4 weeks); accumulate operating history to build each '
       'unit\'s baseline (3–6 months — the longest phase); deploy and validate detection against known maintenance '
       'events (2–4 weeks); integrate demand signals with DLA (4–8 weeks). End to end: roughly 6–9 months.',
       lead='Timeline. ')
bullet(' A 4-week proof of concept on a selected platform and installation would demonstrate behavioral profiling, '
       'zone-divergence detection, regime classification, and demand-signal generation on existing sensor and '
       'maintenance data — establishing the BEI-specific numbers this paper deliberately does not claim.',
       lead='First step. ')

# ================= 6. NOVELTY =================
H1('6.  Why This Is New')
para('Predictive maintenance is a mature industry, and the individual ingredients here, including sensor analytics, '
     'anomaly detection, remaining-useful-life models, and oil analysis, already exist. The contribution is their '
     'composition: representing every serial-numbered unit as a behavioral entity in a shared semantic space; '
     'reading health from relationship drift and zone divergence rather than from individual thresholds; using four '
     'named regimes to drive graduated action; deriving fleet-cohort and supplier-quality inference from the same '
     'space; and connecting behavioral predictions directly to DLA demand signals (what part, when, and where) so '
     'that the supply chain moves before the work order exists. Conventional CBM reports that a reading is out of '
     'specification; this architecture identifies for the maintainer what to fix, for the supply chain what to ship, '
     'and for the commander how long until the equipment is returned to service.', bold=False)
para('The framework was not invented for this paper. The same zone-divergence logic operationally isolates threat '
     'signatures in cybersecurity, and the same entity-drift machinery drives the demand-forecasting work. '
     'Maintenance is the third application of a single proven behavioral architecture applied across multiple '
     'mission domains, which is itself part of the claim.', bold=True)

# ================= 7. LIMITATIONS =================
H1('7.  Limitations')
bullet('This is a concept paper applying a proven framework by analogy; no maintenance-specific proof of concept '
       'has been conducted, and no BEI-measured maintenance performance exists yet.')
bullet('All improvement ranges are industry benchmarks for predictive/condition-based maintenance generally, not '
       'measurements of this system. Actual results will vary by platform, fleet size, environment, and current '
       'maintenance maturity.')
bullet('The behavioral signatures, drift magnitudes, lead times, and regime thresholds described would all need '
       'to be established and calibrated per platform type on real sensor data.')
bullet('Baseline establishment needs 3–6 months of operating history per unit — value arrives gradually, not on '
       'day one.')
bullet('Detection quality depends on the data already being collected: sparse sensor coverage or incomplete '
       'maintenance logs on a given platform would limit what the zones can see.')

# ================= 8. CONCLUSION =================
H1('8.  Conclusion')
para('Equipment changes how it operates before it stops operating. A turbine approaching failure does not suddenly '
     'lose thrust; the relationships among its temperature, vibration, fuel flow, and output shift gradually over '
     'weeks. Threshold-based maintenance, examining each reading in isolation, structurally cannot detect that shift '
     'and registers the failure only once it has occurred. Behavioral Entity Intelligence monitors the '
     'relationships and their direction for each individual unit, against that unit\'s own baseline, diagnoses the '
     'specific component from the pattern of which zone diverges, classifies the urgency into four regimes, and '
     'converts the prediction into a parts shipment already moving toward the appropriate depot.')
para('Traditional maintenance answers whether anything is out of specification. Prescriptive maintenance answers '
     'what the unit is becoming, what to fix, what to ship, and how much time remains, and it does so weeks earlier.',
     bold=True)

# ================= GLOSSARY =================
H1('Appendix — Glossary')
gl = [
    ('Digital entity / digital twin', 'A living software profile of one specific physical unit — like a patient\'s medical chart for a machine: history, vitals, context, and trend.'),
    ('BEI', 'Behavioral Entity Intelligence — our approach: reading how equipment behaves and where that behavior is heading, rather than checking readings against fixed limits.'),
    ('Semantic space', 'A mathematical "map of meaning": units whose behavior is similar sit near each other, even if their raw sensor numbers differ. The representation that powers large language models.'),
    ('CBM', 'Condition-based maintenance — today\'s best practice: monitor sensors and act when a reading crosses a manufacturer threshold.'),
    ('Prescriptive maintenance', 'The rung above predictive: not just when failure will occur, but what action to take now and what parts to stage where.'),
    ('Behavioral zone', 'One of five lenses on a unit\'s health: vibration, thermal, power/performance, fluid chemistry, and duty cycle.'),
    ('Zone divergence', 'The diagnostic pattern of which zones drift and which stay stable — it names the specific failing component.'),
    ('Drift / velocity / acceleration', 'How a unit\'s behavior is moving away from its own baseline; how fast; and whether that rate is increasing (the strongest predictor of imminent failure).'),
    ('Regime', 'A unit\'s health classification — STABLE, DRIFTING, ACCELERATING, or REGIME SHIFT — each driving a defined maintenance and supply response.'),
    ('RUL', 'Remaining useful life — the estimated time until a unit needs intervention, stated with honest uncertainty bounds.'),
    ('DLA', 'Defense Logistics Agency — the DoD organization that buys and distributes spare parts and supplies.'),
    ('NSN', 'National Stock Number — a part\'s unique catalog ID across the defense supply system.'),
    ('GCSS-Army / EBS', 'The Army\'s logistics and maintenance system of record / DLA\'s Enterprise Business System for parts procurement.'),
    ('OR rate', 'Operational Readiness rate — the share of a fleet mission-capable at a given time.'),
    ('FOB', 'Forward Operating Base — a deployed military outpost.'),
    ('GAO', 'Government Accountability Office — the U.S. Congress\'s audit and evaluation agency.'),
    ('Cohort', 'A group of units with similar behavioral profiles; what happens to some members is an early warning for the rest.'),
]
gt = doc.add_table(rows=len(gl) + 1, cols=2); gt.style = 'Table Grid'
setcell(gt.rows[0].cells[0], 'Term', bold=True, color=WHITE, fill='0B1F3A')
setcell(gt.rows[0].cells[1], 'Meaning', bold=True, color=WHITE, fill='0B1F3A')
for i, (t, d) in enumerate(gl, start=1):
    setcell(gt.rows[i].cells[0], t, bold=True, size=9, color=NAVY, fill='F4F6F8')
    setcell(gt.rows[i].cells[1], d, size=9)
for row in gt.rows:
    row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(5.0)

para('')
para('Business Edition — architecture-level disclosure. Exact model logic, drift mathematics, zone composition, '
     'regime criteria, and thresholds are proprietary and available in the full technical edition under NDA. This is '
     'a concept paper: all improvement ranges are industry benchmarks, and BEI-specific maintenance performance '
     'would be established through a proof of concept on real platform data. Do not redistribute.', 9, GREY, italic=True)

out = OUT; n = 2
while True:
    try:
        doc.save(out); break
    except PermissionError:
        root, ext = os.path.splitext(OUT); out = f"{root}_v{n}{ext}"; n += 1
print('Saved:', out)
