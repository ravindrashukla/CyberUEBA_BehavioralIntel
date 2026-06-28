"""Fourth pass: restore the 'Nation's Logistics Combat Support Agency / warfighter
readiness' framing in the executive summary by tracked-replacing the current generic
para 3. The LLM/transformer sentence in para 4 is left untouched.
"""
import os, copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REV = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_final_REV.docx"
AUTHOR = "ReviewPanel"; DATE = "2026-06-19T15:00:00Z"
d = Document(REV)
_id = [11000]
def nid():
    _id[0] += 1; return str(_id[0])

def _run(text, rpr=None):
    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t)
    return r

def _ins_wrap(run):
    ins = OxmlElement('w:ins'); ins.set(qn('w:id'), nid()); ins.set(qn('w:author'), AUTHOR); ins.set(qn('w:date'), DATE)
    ins.append(run); return ins

def del_run_inplace(r_elem):
    for t in list(r_elem.iter(qn('w:t'))):
        dt = OxmlElement('w:delText'); dt.set(qn('xml:space'), 'preserve'); dt.text = t.text
        t.getparent().replace(t, dt)
    delw = OxmlElement('w:del'); delw.set(qn('w:id'), nid()); delw.set(qn('w:author'), AUTHOR); delw.set(qn('w:date'), DATE)
    r_elem.addprevious(delw); delw.append(r_elem)

NEW = ("As the Nation's Logistics Combat Support Agency, DLA must sustain warfighter readiness across a global supply "
    "chain that is increasingly nonlinear, relationship-dependent, and exposed to disruption. Current planning, "
    "forecasting, and supplier-monitoring methods remain necessary, but they can leave DLA reacting after thresholds "
    "break rather than seeing where demand, supplier behavior, inventory posture, or network exposure is moving. The "
    "requirement is a stronger decision-intelligence layer that improves supply-chain risk mitigation and supports "
    "decision advantage before operational consequences are fully visible.")

# locate current para 3 (unique phrase from the current generic version)
p3 = next(p for p in d.paragraphs if 'rather than reveal where demand' in p.text)
rpr = p3.runs[0]._r.find(qn('w:rPr')) if p3.runs else None
# mark existing runs deleted (snapshot list first)
for r in list(p3.runs):
    del_run_inplace(r._r)
# insert the new text as a tracked insertion at end of paragraph
p3._p.append(_ins_wrap(_run(NEW, rpr=rpr)))

o = REV; k = 2
while True:
    try:
        d.save(o); break
    except PermissionError:
        root, ext = os.path.splitext(REV); o = f"{root}_v{k}{ext}"; k += 1
print("Saved:", o)

# re-extract accepted text
d2 = Document(o)
from docx.table import Table
out = []
for child in d2.element.body.iterchildren():
    if child.tag == qn('w:p'):
        t = "".join(x.text or "" for x in child.iter(qn('w:t'))).strip()
        if t: out.append(t)
    elif child.tag == qn('w:tbl'):
        out.append("[TABLE]")
        for row in Table(child, d2).rows:
            out.append(" | ".join(c.text.strip().replace("\n", " ") for c in row.cells))
        out.append("[/TABLE]")
open("WP DLA/_DLA_paper_REV_accepted.txt", "w", encoding="utf-8").write("\n".join(out))
print("accepted-text lines:", len(out))
