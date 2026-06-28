"""Academic edition of the V-Intelligence UEBA whitepaper, structured to mirror the
DLA Decision-Intelligence academic paper: problem -> technical foundation -> use cases
-> technical evidence and validation boundaries -> recommended pilot -> conclusion.

Two-layer detection thesis (verified against CyberUEBA_BehavioralIntel_Enhanced):
  Layer 1 (precision)  Multi-front threat-profile detector: named known-bad fingerprints
     (cohort-relative + raw-event + self-drift). 4/4 injected attacks at 0 false positives.
  Layer 2 (discovery)  Behavioral embeddings + drift + novelty: direction naming against
     MITRE ATT&CK and the unsupervised net for attacks not yet profiled.

Reuses the data-grounded figures already generated under docs/figs_ueba_enhanced/.
No em-dashes. Numbered figures and tables. Synthetic-data, proof-of-concept framing.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "Behavioral_Intelligence_for_Cyber_Defense_Academic.docx")
FIG = os.path.join(os.path.dirname(__file__), "figs_ueba_enhanced")

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
ORG  = RGBColor(0xC2, 0x5A, 0x12)
GREY = RGBColor(0x55, 0x5F, 0x6B)
WHITE= RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.12
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(doc.sections[0], m, Inches(0.9))

_fign = [0]; _tabn = [0]

def H1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    p.style = doc.styles['Heading 1']
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    return p

def H2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    p.style = doc.styles['Heading 2']
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ORG
    return p

def para(text, size=10.5, color=None, bold=False, italic=False, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def bullet(text, lead=None):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run(text.lstrip()); r2.font.size = Pt(10.5)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor); tcPr.append(sh)

def setcell(cell, text, bold=False, size=9.5, color=None, fill=None):
    if fill: shade(cell, fill)
    cell.text = ''; p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
    if color: r.font.color.rgb = color

def table(headers, rows, caption, widths=None, hdr_fill='0B1F3A'):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        setcell(t.rows[0].cells[j], h, bold=True, color=WHITE, fill=hdr_fill)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            setcell(t.rows[i].cells[j], v, bold=(j == 0), fill='F4F6F8' if j == 0 else None)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    _tabn[0] += 1
    c = doc.add_paragraph(); c.paragraph_format.space_after = Pt(8)
    r = c.add_run(f"Table {_tabn[0]} - {caption}"); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY
    return t

def add_fig(name, caption, width=6.4):
    path = os.path.join(FIG, name)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(path, width=Inches(width))
    _fign[0] += 1
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after = Pt(8)
    r = c.add_run(f"Figure {_fign[0]} - {caption}"); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY


# ===================== TITLE =====================
r = doc.add_paragraph().add_run('Behavioral Entity Intelligence for Cyber Threat Detection')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
para('Detecting insider, advanced-persistent-threat, living-off-the-land, and infrastructure-compromise activity '
     'that stays within conventional thresholds.', 12.5, GREY, after=2)
para('22nd Century Technologies, Inc.  ·  V-Intelligence UEBA Program  ·  June 2026  ·  Academic Edition',
     9.5, GREY, after=2)
para('Disclosure note: this edition describes the architecture, the detection method, and measured results. The exact '
     'detection logic (named-profile definitions, formulas, thresholds, weights, and parameters) is proprietary and is '
     'withheld here; it is available in the full technical edition under a non-disclosure agreement.', 9.5, ORG,
     italic=True, after=4)
para('Data note: all results in this paper come from a synthetic (simulated) enterprise dataset built to prove the '
     'concept: 250 users, 485 days (70 weeks), approximately 14 million log events across five log sources, with four '
     'embedded long-duration attack campaigns modeled on real, publicly documented intrusions. Results must be '
     're-validated on real operational telemetry before any production claim.', 9.5, ORG, italic=True, after=8)

# ===================== 1. EXECUTIVE SUMMARY =====================
H1('1.  Executive Summary')
para('The most damaging cyber intrusions of the past decade did not trip a single alarm. Volt Typhoon, a state-'
     'sponsored campaign against United States critical infrastructure, remained inside victim networks for at least '
     'five years undetected. Salt Typhoon, described by the chairman of the Senate Intelligence Committee as the worst '
     'telecom hack in the nation\'s history, operated inside major United States telecommunications providers for '
     'months. In both cases the victims\' own security tools never fired, and the campaigns were discovered by outside '
     'intelligence agencies. Such adversaries succeed by operating with valid credentials and legitimate tools, '
     'staying within the statistical bounds of normal activity while shifting their behavior gradually over months.')
para('This paper presents Behavioral Entity Intelligence (BEI) applied to User and Entity Behavior Analytics (UEBA): '
     'a detection approach that learns a behavioral baseline for every user and system and then evaluates both how '
     'much behavior changes and the direction in which it is trending. The detection engine has two layers, and that '
     'is the central contribution.', bold=False)
bullet('a library of measurable known-bad behavioral fingerprints scored relative to an entity\'s role-group peers '
       '(cohort-relative profiles) and read directly from raw events (covert call-home beacons, algorithmically '
       'generated domains, rare external destinations). Each flag is a named profile match mapped to a MITRE ATT&CK '
       'technique, so it is precise and explainable. On the validation set it matched all four attack campaigns at '
       'zero false positives.', lead='Layer 1, the multi-front threat-profile detector (precision). ')
bullet('a behavioral map of meaning built from the representation that powers large language models, which measures '
       'drift and novelty and names the direction of change against MITRE ATT&CK. It is the unsupervised net for '
       'attacks not yet in the fingerprint library, and it supplies the direction that makes the precision layer\'s '
       'beacon and domain fingerprints interpretable.', lead='Layer 2, behavioral embeddings, drift, and novelty (discovery). ')
para('The proof-of-concept results, stated plainly:', bold=True)
bullet('Of four industry-standard anomaly-detection algorithms run against the same data, three detected none of the '
       'four campaigns and the best detected exactly one, a statistical spike with no explanation of which behavior '
       'changed or what threat it resembled.', lead='Traditional methods miss these attacks. ')
bullet('The precision layer matched all four campaigns to named attack fingerprints and flagged no normal user. Every '
       'detection carries the specific profile it matched and the ATT&CK technique it represents.',
       lead='Four of four detected, zero false positives, each one named. ')
para('An insider changes what they access, not how much; a slow intrusion changes where it communicates, not how '
     'often. Traditional algorithms measure magnitude; BEI measures direction, peer-relative deviation, and named-'
     'profile match, which are the dimensions in which these attacks are actually visible. These results are from '
     'synthetic data and define a proof of concept, not a fielded capability.', bold=True)

# ===================== 2. PROOF OF CONCEPT =====================
H1('2.  Proof of Concept: Detection-Value Evidence')
para('The current proof of concept demonstrates that BEI can detect slow, hidden, directional cyber threats that '
     'stay within conventional thresholds. The results should be read as proof of concept, not as operational proof '
     'on a defended network. They show that the architecture can match long-duration campaigns to named, measurable '
     'behavioral fingerprints, separate attackers from a large normal population in a semantic representation, and '
     'attach an explainable direction to each detection, warranting validation through a bounded pilot on real '
     'operational telemetry.')
para('Across a 250-user, 485-day synthetic enterprise with four embedded campaigns, the precision layer matched 4 of '
     '4 attacks to named fingerprints at 0 false positives, while the strongest traditional detector caught 1 of 4 '
     'with no explanation. The discovery layer cleanly separated the insider and the telecom-infrastructure campaigns '
     'on its own and named the behavioral direction of all four.')

# ===================== 3. THE PROBLEM =====================
H1('3.  The Cyber Detection Problem')
para('Enterprise security monitoring rests on two pillars. The first is the Security Information and Event Management '
     '(SIEM) system, the central log platform that applies expert-written rules to identify known attack patterns. '
     'The second is signature-based detection: antivirus and intrusion-detection tools that match activity against '
     'Indicators of Compromise (IOCs), the known-bad file fingerprints, addresses, and domains. Both answer the same '
     'question: does this activity match something already known to be malicious? Modern adversaries are engineered '
     'to make the answer no. Three threat families account for most of these cases.')
bullet('A trusted employee or contractor who misuses access they legitimately hold. Industry data (the Verizon Data '
       'Breach Investigations Report) attributes roughly 20 percent of breaches to insiders, with a median time to '
       'detection exceeding 200 days. Every individual action uses valid credentials; only the pattern over months '
       'reveals intent.', lead='Insider threats. ')
bullet('An Advanced Persistent Threat (APT) is a patient, well-resourced intrusion, typically nation-state, that '
       'gains a foothold and remains for months or years. Its covert command-and-control (C2) channel, through which '
       'the implant calls home for instructions, is throttled to a trickle: a check-in every several hours, only '
       'kilobytes of traffic, invisible inside enterprise volume.', lead='Advanced persistent threats. ')
bullet('Living-off-the-land (LOTL) attacks use only the administration tools already installed on every system, such '
       'as command shells, remote desktop, and management utilities. There is no malware to fingerprint and no rule '
       'to violate, because every action is, in isolation, an authorized one.', lead='Living-off-the-land attacks. ')
para('The common technique is known as slow and low: behavior changes so gradually that no single week looks unusual, '
     'while the cumulative change over months achieves the objective. The activity is visible only in the trajectory, '
     'never in a single snapshot.')
add_fig('fig_slowlow.png', 'The slow-and-low evasion. Left: a threshold-based tool judges each week alone, and the '
        'attacker\'s weekly change never crosses the fixed alert line, so no alert is ever raised. Right: the same '
        'weeks accumulated as behavioral drift separate the attacker from the normal-user range in months, not years. '
        '(Illustrative, based on the simulated slow-APT campaign.)')
para('The real-world consequence is documented. Neither Volt Typhoon nor Salt Typhoon was caught by any victim\'s own '
     'monitoring; SIEM, intrusion detection, and endpoint tools all failed because there was no signature to match '
     'and no threshold crossed. Both were found by external intelligence operations, years or months after '
     'compromise. This is a structural gap, not a tuning problem, and it is the gap behavioral analytics is built to '
     'close.')

H2('3.1  Why Traditional Security Monitoring Falls Short')
bullet('rules and signature checks can stay green while an entity drifts toward a known-bad pattern, because the '
       'individual actions remain authorized. (e.g., a user keeps normal sign-in volume while shifting from public to '
       'restricted files week over week.)', lead='Threshold insensitivity. ')
bullet('conventional anomaly models measure how much a metric changes, not what direction the entity is becoming. '
       '(e.g., a user accesses about thirty files a week before and during a campaign, but the content shifts from '
       'public to confidential.)', lead='Magnitude-only measurement. ')
bullet('a user, device, or account may look normal on its own while its relationship to destinations, peers, or data '
       'has changed. (e.g., one external address recurs every week for over a year, the fingerprint of covert '
       'infrastructure.)', lead='Single-entity, single-metric scope. ')
bullet('detection that depends on a known-bad list cannot see a novel implant, a new domain, or a first-seen '
       'destination until it has already been catalogued. (e.g., living-off-the-land activity uses only built-in '
       'tools, so there is nothing to fingerprint.)', lead='Signature dependence. ')
bullet('backward-looking methods react after a threshold finally breaks, by which time a patient adversary has '
       'already staged or moved what it needed. (e.g., gigabytes leave over months in packets too small to trip any '
       'single alert.)', lead='Delayed response. ')
para('These limitations do not mean current tools lack value; they indicate where current tools should be '
     'complemented. The proposition advanced here is that a security operations team can retain its existing fast, '
     'familiar detectors while adding a behavioral layer that matches named known-bad fingerprints, measures drift, '
     'and names its direction before the consequence is observable.')

# ===================== 4. TECHNICAL FOUNDATION =====================
H1('4.  Technical Foundation: Behavioral Entities, Drift, and Detection Signals')
para('Behavioral Entity Intelligence represents an observable subject as a profile that evolves over time rather '
     'than as a static record. An entity may be a user, device, network segment, application, or session. The '
     'closest analogy is a patient\'s medical chart: it does not record only today\'s temperature; it accumulates '
     'history, vitals, and context, and a clinician reads the trend to act before a condition becomes acute. Each '
     'week, raw activity from five log sources (sign-ins, file access, endpoint activity, network traffic, and domain '
     'lookups) is distilled into roughly two dozen behavioral measurements per user, organized into five behavioral '
     'zones.')

H2('4.1  Key Behavioral Zones')
table(['Behavioral zone', 'What it captures', 'What it catches when it drifts'],
      [('Identity', 'Who the entity is: role, department, clearance, tenure.',
        'Stability here is itself a clue: insiders and credential thieves still look like the same person.'),
       ('Access pattern', 'How the entity signs in: frequency, failures, timing, sources.', 'Credential abuse and account takeover.'),
       ('Data behavior', 'What data it touches: volume, sensitivity mix, read and write balance.', 'Insider data theft, a shift toward restricted and confidential material.'),
       ('Network footprint', 'Where it communicates: destinations, external traffic, domain lookups.', 'Covert call-home channels and data-removal routes.'),
       ('Risk posture', 'Endpoint health: suspicious process activity and risk indicators.', 'Compromised machines and privilege-escalation tooling.')],
      'Five behavioral zones and what each catches when it drifts.', widths=[1.3, 2.7, 2.9])
para('Zone decomposition is the central representational choice. When the identity zone is stable but the data-'
     'behavior zone is drifting, the system can conclude that this is the same person using the same credentials but '
     'accessing different data than before, which is the insider-threat signature and which no single anomaly score '
     'can express.')

H2('4.2  The Two-Layer Detection Architecture')
para('No single signal catches every family, because each attack evades a different detector. The architecture '
     'therefore runs two complementary layers over the same behavioral profile.')
bullet('a library of measurable known-bad fingerprints. Each fingerprint is a precise, named description of how a '
       'specific attack technique shows up in behavior, and a flag is raised only when an entity matches one. It '
       'fuses three fronts: cohort-relative profiles (how an entity behaves versus its role-group peers, where a '
       'named profile fires only when the right combination of measurements is jointly elevated); raw-event profiles '
       'read directly from logs with no labels (the covert beacon found by its robotic rhythm, generated domains '
       'found by their randomness and shared infrastructure, rare destinations contacted by one entity alone in its '
       'cohort); and self-drift as supporting corroboration. Because a flag means a named profile match rather than a '
       'statistical outlier, the result is both high-precision and self-explaining.',
       lead='Layer 1, the multi-front threat-profile detector (precision). ')
bullet('each week\'s measurements are written as a structured plain-language description and converted into a high-'
       'dimensional numerical vector called an embedding, the same representation that powers large language models. '
       'The result is a semantic space, a map of meaning in which behaviors that mean similar things sit near each '
       'other even when their raw numbers differ. Tracking each entity\'s weekly position yields drift (direction and '
       'rate of change) and novelty (a never-before-seen behavior that keeps recurring); drift is named against a '
       'library of reference threat concepts aligned with MITRE ATT&CK. This layer is the unsupervised net for '
       'attacks the fingerprint library does not yet contain, and it supplies the direction the precision layer '
       'reports.', lead='Layer 2, behavioral embeddings, drift, and novelty (discovery). ')

H2('4.3  Mathematical and Model Design Principles')
para('The architecture follows a consistent pattern. Behavioral measurements are scored relative to the entity\'s '
     'own history and its role-group cohort, using a robust peer-relative deviation that ignores normal spread, so a '
     'named profile fires only on genuine cohort-relative extremes. Raw-event fingerprints are computed label-free: '
     'a covert beacon is identified by the machine-like regularity of its call-out rhythm (a low coefficient of '
     'variation that legitimate, bursty services do not have), and algorithmically generated domains by high name-'
     'randomness and shared back-end infrastructure. Slow drift is accumulated with a cumulative change-detection '
     'method from the same family used to catch slow degradation in industrial quality control, and drift direction '
     'is projected by cosine similarity against pre-embedded threat concepts. The contribution is the composition of '
     'these established techniques into an explainable detection system; no individual method is claimed to be novel '
     'in isolation. Exact thresholds, weights, and parameters are proprietary and are withheld in this edition.')
add_fig('fig_scoreboard.png', 'The detection scoreboard across 250 users with four embedded campaigns (measured, '
        'synthetic data). Traditional methods catch at most 1 of 4; the discovery layer separates 2 of 4 on its own; '
        'the multi-front threat-profile detector catches 4 of 4 at zero false positives, each detection named.')

# ===================== 5-7. USE CASES =====================
H1('5.  Use Case 1: Insider Threat Detection')
para('A trusted business user (entity USR-156) ran a fourteen-month campaign, drifting from public toward restricted '
     'and confidential files while total file volume stayed flat, the textbook signature of the same credentials '
     'used to reach different data. The precision layer matched two named fingerprints: a mass-collection profile, '
     'where the user\'s unique-file and total-file activity were jointly elevated far above business-unit peers '
     '(cohort-relative deviation 5.9), and a cohort-rare-destination profile, where 76 external addresses were '
     'contacted by this user alone within its cohort. The discovery layer independently lifted this user clear of the '
     'entire population in semantic-drift space by week four, the earliest separation of any campaign.')

H1('6.  Use Case 2: Advanced Persistent Threat and Covert Command-and-Control')
para('A developer account (entity USR-234) hosted a 417-day slow nation-state intrusion whose implant called home '
     'roughly every several hours in kilobyte-sized packets, with total traffic, lookups, and sign-ins all inside '
     'normal ranges. This is the hardest detection problem in the study: on every magnitude-based method its scores '
     'are negligible, and it never separates from the population even in semantic-drift space. The precision layer '
     'matched it on two independent raw-event fingerprints, with no labels and no thresholds crossed.')
bullet('among all of this user\'s external contacts, one destination was contacted persistently, on 386 distinct '
       'days, with a robotic, evenly spaced rhythm. Legitimate persistent services such as software updates and cloud '
       'synchronization are bursty and irregular; a beacon is metronomic. The detector flags the regularity, not the '
       'destination, so it needs no prior knowledge of the attacker\'s infrastructure.', lead='Covert C2 beacon. ')
bullet('the same implant resolved 160 algorithmically generated domains, high-randomness names sharing back-end '
       'infrastructure, the signature of domain-generation malware locating its controller.', lead='Generated domains. ')
add_fig('fig_beacon.png', 'Label-free beacon detection (illustrative, modeled on the simulated slow-APT campaign). '
        'Left: the attacker\'s call-outs are evenly spaced; a legitimate service is bursty. Right: the rhythm-'
        'irregularity measure (coefficient of variation of the gaps), where below the flag line means machine-like, '
        'the fingerprint of a covert beacon. The detector reads the rhythm, not the address, so it generalizes to '
        'infrastructure it has never seen.')

H1('7.  Use Case 3: Living-off-the-Land and Critical-Infrastructure Compromise')
para('Two campaigns model the intrusions that no victim organization self-detected. A security-team account '
     '(entity USR-042) ran a Volt Typhoon-style living-off-the-land campaign using only legitimate administration '
     'tools; the precision layer matched a living-off-the-land process profile, where the breadth and volume of '
     'administrative process activity were jointly elevated against security-team peers (cohort-relative deviation '
     '4.5). A second developer account (entity USR-118) ran a Salt Typhoon-style telecom-infrastructure campaign; '
     'the precision layer matched a reconnaissance fan-out profile, where the spread of network destinations was '
     'elevated far above developer peers (cohort-relative deviation 8.2). Both campaigns were matched to named, '
     'measurable fingerprints, and the discovery layer additionally surfaced the telecom campaign in semantic-drift '
     'space.')
add_fig('fig_fingerprint.png', 'The four named-profile matches (measured, synthetic data). Left: cohort-relative '
        'profiles fire only when an entity\'s peer-relative deviation clears the flag line. Right: raw-event profiles '
        'are read straight from the logs, including the slow APT\'s persistent beacon and generated domains and the '
        'insider\'s rare-destination fan-out.')

# ===================== 8. EVIDENCE & BOUNDARIES =====================
H1('8.  Technical Evidence and Validation Boundaries')
para('This section distinguishes proof of concept from operational proof. Every number below is reproduced directly '
     'from the evaluation data on the synthetic enterprise; none has been validated on real telemetry.')
table(['Attack (entity)', 'Named profile match', 'Fronts'],
      [('Insider threat (USR-156)', 'mass_collection (peer-z 5.9) and cohort_rare_dst (76 rare external addresses)', '2'),
       ('Slow APT, covert C2 (USR-234)', 'c2_beacon (386-day persistent, robotic rhythm) and dga_dns (160 generated domains)', '2'),
       ('Living off the land (USR-042)', 'lotl_process (peer-z 4.5, unusual admin-tool breadth)', '1'),
       ('Telecom infrastructure (USR-118)', 'recon_fanout (peer-z 8.2, network-destination fan-out)', '1')],
      'The four named-profile matches (measured, synthetic data). Each campaign matched a different combination of '
      'fronts, with no normal user flagged.', widths=[2.1, 3.9, 0.7])
table(['Method', 'Attacks detected', 'False positives', 'What the analyst receives'],
      [('Local Outlier Factor', '0 of 4', '4.5%', 'Low false positives and zero detections; all four campaigns sit inside its normal range.'),
       ('Isolation Forest', '0 of 4', '5.3%', 'All four attackers look statistically normal in feature space.'),
       ('One-Class SVM', '0 of 4', '14.6%', 'Many false positives and zero true detections.'),
       ('Z-Score', '1 of 4', '9.8%', 'Catches only one spike (the living-off-the-land campaign), barely over the line; the other three never cross it.'),
       ('Discovery layer (alone)', '2 of 4', 'n/a', 'Cleanly separates the insider and telecom campaigns and explains the direction of all four.'),
       ('Multi-front Threat-Profile', '4 of 4', '0', 'All four matched to a named attack fingerprint and ATT&CK technique; no normal user flagged.')],
      'The detection scoreboard across 250 users with four embedded campaigns (measured, synthetic data).',
      widths=[1.8, 1.1, 1.0, 2.8])
add_fig('fig_separation.png', 'Signal separation, both lenses, over the full 70-week campaign (measured on the '
        'synthetic enterprise). The raw-magnitude lens separates the high-volume attacks late; the semantic meaning '
        'lens catches the subtle insider far sooner but trails on the volume attack. Each lens wins on different '
        'attacks, and neither catches the slow APT, which the threat-profile beacon does.')

H2('8.1  Detection Evidence Summary')
table(['Technical point', 'Source detail', 'Decision interpretation'],
      [('Synthetic scope', '250 users, 485 days (70 weeks), approximately 14 million events, five log sources.', 'Large enough for concept testing, but requires real-telemetry remeasurement.'),
       ('Entity representation', 'Per-week behavioral profile in five zones; semantic embedding per week.', 'Detection becomes trajectory-aware rather than snapshot-based.'),
       ('Precision result', 'Multi-front threat-profile detector matched 4 of 4 campaigns at 0 false positives.', 'Named, explainable detection that no normal user tripped.'),
       ('Discovery result', 'Semantic drift cleanly separated 2 of 4 on its own and named the direction of all four.', 'Unsupervised net for novel attacks plus interpretable direction.'),
       ('Traditional comparison', 'Best conventional detector caught 1 of 4; three caught none.', 'Magnitude-only methods miss directional, sub-threshold behavior.'),
       ('Boundary, threshold', 'The operating point was chosen with knowledge of the ground truth.', 'In production the threshold must be set blind and tuned per environment.'),
       ('Boundary, novelty', 'A fingerprint library only catches techniques it describes.', 'The discovery layer is the net for novel attacks and is noisier by design.'),
       ('Boundary, data', 'All results are synthetic, calibrated to public CISA advisories.', 'Every figure, including the zero-false-positive result, must be re-measured on real data.')],
      'Detection technical evidence and validation boundaries (synthetic data).', widths=[1.5, 3.0, 2.5])

# ===================== 9. PILOT =====================
H1('9.  Recommended Pilot Validation Path')
para('The recommended next step is a bounded pilot that runs the existing security stack and BEI outputs side by '
     'side on real operational telemetry, in advisory mode, without feeding automated blocking until results are '
     'validated by analysts. The pilot validates how early, how reliably, and how explainably the behavioral layer '
     'detects long-duration activity that the current stack does not surface.')
table(['Pilot element', 'Recommended scope', 'Success measures'],
      [('Baseline build', '4 to 6 weeks of the five standard log sources; per-entity profiles and role cohorts.', 'Stable baselines and cohorts established.'),
       ('Precision detection', 'Named-profile matching on the live environment, thresholds tuned blind.', 'Named matches with the false-positive rate within an agreed budget.'),
       ('Discovery detection', 'Semantic drift and novelty as the unsupervised net for unprofiled activity.', 'Direction naming and novel-behavior surfacing reviewed by analysts.'),
       ('Red-team validation', 'Replay of known-bad activity or historical incidents on this environment\'s data.', 'Known-bad activity matched, named, and explained.'),
       ('Integration', 'Ranked, explained leads into the existing SIEM and SOAR workflow.', 'Analyst-accepted leads in steady state, with ATT&CK technique attached.')],
      'Recommended pilot scope and success measures.', widths=[1.5, 3.1, 2.4])
para('The behavioral layer is designed to complement, not replace, the existing stack. It uses five standard log '
     'sources most enterprises already collect, runs in containerized standard infrastructure suitable for '
     'enterprise, cloud, or disconnected secure environments, and feeds existing Security Operations Center workflows '
     'with the matched profile and ATT&CK technique attached.')

# ===================== 10. CONCLUSION =====================
H1('10.  Conclusion')
para('The threats that matter most, including insiders, nation-state intrusions, living-off-the-land campaigns, and '
     'infrastructure compromises, share one property: they operate inside authorized access, beneath every signature '
     'and threshold, and reveal themselves only as a gradual change of behavioral direction over months. The '
     'documented record is unambiguous: the victims of Volt Typhoon and Salt Typhoon did not detect these campaigns; '
     'outside agencies did, years and months after the fact.')
para('Behavioral Entity Intelligence reads the same enterprise logs as a living behavioral profile for every entity '
     'and runs two complementary layers over it: a precision layer of named, measurable known-bad fingerprints '
     'scored against role-group peers, and a discovery layer of semantic drift and novelty that names the direction '
     'against known adversary behavior. On a 250-user, 485-day proof of concept with four embedded campaigns, the '
     'precision layer matched 4 of 4 attacks to named attack fingerprints at zero false positives, against 1 of 4 '
     'with no explanation for the best traditional method. These results are from synthetic data and justify '
     'validation on real operational telemetry through a bounded pilot.', bold=True)
para('An insider changes what they access, not how much; an APT changes where it communicates, not how often. '
     'Traditional tools measure magnitude. BEI measures direction, peer-relative deviation, and named-profile match, '
     'the dimensions in which these attacks are actually visible.', bold=True)

out = OUT; n = 2
while True:
    try:
        doc.save(out); break
    except PermissionError:
        root, ext = os.path.splitext(OUT); out = f"{root}_v{n}{ext}"; n += 1
print('Saved:', out, '| figures:', _fign[0], '| tables:', _tabn[0])
