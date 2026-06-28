# -*- coding: utf-8 -*-
"""Digital Twin Working Session - Run-of-Show for the Army cyber audience
(Esquibel / DiGrezio / Herrmann). Digital twin is the centerpiece; three lenses:
prove the mechanism (operator), show it's fieldable (acquisition), frame why it
matters (policy). Real data points throughout.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6); GRAY=RGBColor(0x55,0x5F,0x6B)
RED=RGBColor(0xA6,0x1B,0x1B); GREEN=RGBColor(0x1F,0x7A,0x33); AMBER=RGBColor(0xB5,0x5A,0x00)
d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(10.5)

def H1(t):
    p=d.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
    return p
def H2(t):
    p=d.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=BLUE
    return p
def P(t,italic=False,bold=False,color=None,size=None):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold
    if color: r.font.color.rgb=color
    if size: r.font.size=Pt(size)
    return p
def B(label,rest="",sub=False):
    p=d.add_paragraph(style='List Bullet 2' if sub else 'List Bullet')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def NUM(label,rest=""):
    p=d.add_paragraph(style='List Number')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def meta(label,val):
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(label+"  "); r.bold=True; r.font.color.rgb=GRAY; r.font.size=Pt(9.5)
    r2=p.add_run(val); r2.font.size=Pt(9.5); return p
def say(lines):
    lp=d.add_paragraph(); lp.paragraph_format.space_after=Pt(1)
    lr=lp.add_run("SAY:"); lr.bold=True; lr.font.color.rgb=GREEN; lr.font.size=Pt(9.5)
    for ln in lines:
        p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.25); p.paragraph_format.space_after=Pt(3)
        r=p.add_run("“"+ln+"”"); r.italic=True; r.font.size=Pt(10.5)
def callout(text,color=RED,size=9.5):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2)
    r=p.add_run(text); r.bold=True; r.font.color.rgb=color; r.font.size=Pt(size); return p
def table(headers,rows,widths=None,fs=8.5,hdr_fs=8.5):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(hdr_fs)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""
            for j,part in enumerate(v.split("||")):
                para=cells[i].paragraphs[0] if j==0 else cells[i].add_paragraph()
                para.add_run(part).font.size=Pt(fs)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    d.add_paragraph(); return t

# ===== TITLE =====
tp=d.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("Digital Twin Working Session - Run-of-Show"); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=NAVY
sp=d.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Army Cyber Enterprise audience  |  Esquibel (ARCYBER) - DiGrezio (Acquisition) - Herrmann (Policy)"); r.font.size=Pt(11.5); r.font.color.rgb=BLUE
sp2=d.add_paragraph(); sp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp2.add_run("22nd Century  |  V-Intelligence Behavioral Entity Intelligence"); r.font.size=Pt(10); r.italic=True; r.font.color.rgb=GRAY
wn=d.add_paragraph(); wn.alignment=WD_ALIGN_PARAGRAPH.CENTER
rw=wn.add_run("INTERNAL playbook. Digital twin is the centerpiece - they asked for it."); rw.bold=True; rw.font.color.rgb=RED; rw.font.size=Pt(9.5)

# ===== 1 WHO'S IN THE ROOM =====
H1("1. Who is in the room (and why it changes the pitch)")
P("Brandon Pugh became Principal Cyber Advisor to the Secretary of the Army (June 2025). This is the Army cyber "
  "enterprise, not a think tank. The three attendees map to three lenses you must satisfy in one session.",bold=False)
table(["Attendee","Role (confidence)","Lens","What they will judge","How to pitch"],
 [["Dr. Judy Esquibel","Senior Cyber Advisor, ARCYBER; PhD; Jack Voltaic critical-infra lead (confirmed)","Technical operator (your ally)","Detection fidelity vs real hunt ops; LOTL / valid-account; FP economics; baseline drift","Go deep on the twin mechanics and the threat model; do not oversimplify"],
  ["LTC Micah DiGrezio","PM Cyber & Space, Army acquisition (likely)","Acquisition / fielding","Can it be bought, integrated, sustained; telemetry reqs; footprint; acquisition pathway","Frame as a fieldable program: data, integration, sustainment, OTA/PoR"],
  ["Kellsie Herrmann","Army cyber/tech-policy, DC; Georgetown SFS (likely - confirm title)","Policy / strategy","Why it matters; authorities; readiness; strategic significance","Narrative + significance; keep the math light; give her what she can carry upward"]],
 widths=[1.0,1.65,0.95,1.5,1.4],fs=7.5,hdr_fs=8)
callout("Confirm exact titles from the meeting host before the session, especially Herrmann's. Group center of gravity = "
 "operational/Army. Esquibel is your high-fidelity technical seat and best ally; win her and the room follows.")

# ===== 2 THE ONE STORY =====
H1("2. The session in one line")
P("The digital twin is the centerpiece. Run its anatomy as the spine and satisfy three lenses: prove the mechanism "
  "(Esquibel), show it is fieldable (DiGrezio), and frame why it matters (Herrmann). The twin lives inside our "
  "four-layer architecture, but today it is the main event.",bold=False)
callout("Golden rule unchanged: show what runs (the behavioral twin / Layer 2). Label identity fusion, forensics, and "
 "automated response as roadmap, out loud. With ARCYBER in the room, candor is mandatory.",color=RED)

# ===== 3 AGENDA =====
H1("3. Agenda (digital-twin-centered, ~3h incl. buffer)")
table(["#","Block","Primary seat","Time"],
 [["0","Open: their goals, set the arc","All","5 min"],
  ["1","The problem + why a digital twin (Volt/Salt, valid-account, behavior-as-meaning)","Herrmann","15 min"],
  ["2","The digital twin anatomy - 7 parts, deep","Esquibel","45 min"],
  ["3","LIVE demo: the four attackers + the 8.1% derivation","Esquibel","45 min"],
  ["4","Fielding the twin: data, footprint, integration, sustainment, acquisition","DiGrezio","30 min"],
  ["5","Why it matters + authorities + readiness + the ask","Herrmann","20 min"],
  ["-","Buffer for hard questions","All","20 min"]],
 widths=[0.4,3.9,0.95,0.7],fs=9)

# ===== 4 BLOCK 1 =====
H1("4. Block 1 - The problem + why a digital twin (15 min)")
meta("Seat:","Herrmann's 'why it matters' - but Esquibel will nod along if you get the threat model right.")
say(["Two campaigns reset the Army's threat model. Volt Typhoon: pre-positioned in U.S. critical infrastructure, "
 "legitimate Windows tooling only, no malware, five-plus years of dwell. Salt Typhoon: telecom espionage at carrier "
 "scale, slow, sub-threshold. Both use MITRE T1078, valid accounts. They log in; they don't break in.",
 "Threshold tools never fire, because nothing crosses a line. So we stopped asking 'did a metric spike' and started "
 "asking 'what is this entity becoming.' We build a behavioral digital twin of every user, device, and segment, track "
 "where it is heading over time, and name that direction in MITRE terms. That is the whole idea."])
meta("Esquibel hook:","Tie it to Jack Voltaic / critical-infrastructure resilience - the valid-account, "
 "living-off-the-land gap is exactly what infra hunt teams struggle to see.")
meta("Transition:","“Let me open up the twin and show you exactly what it is made of.”")

# ===== 5 BLOCK 2 - THE ANATOMY =====
H1("5. Block 2 - The digital twin anatomy (45 min, deep - Esquibel)")
P("Walk the seven parts. Every number here is real, from the code and result files.",italic=True)

H2("Part 1 - Five signals (the senses)")
B("","For a user, five signals, each from a different log source, each turned into PROSE then embedded into 1536-d. "
   "Weights: auth 0.25, privilege 0.20, data_access 0.20, network 0.20, communication 0.15.")
B("Key decision: ","raw metrics become natural language before embedding. Meaning over magnitude. Two users who both "
   "touched '47 files' embed differently if the context differs (restricted vs public, off-hours vs business hours).")

H2("Part 2 - Composition (weighted average, not concatenation)")
B("","The five signal vectors fuse into one 1536-d composite by weighted average, then L2-normalize. Averaging keeps "
   "the twin in the SAME space as the reference concepts, which is what lets us measure direction later.")

H2("Part 3 - Behavioral zones")
B("","The twin decomposes into five zones tracked independently: identity, access_pattern, data_behavior, "
   "network_footprint, risk_posture. zone_divergence (some zones moving while others stay frozen) is the "
   "living-off-the-land tell.")

H2("Part 4 - Trajectory + CUSUM (time)")
B("","Each week a fresh composite; the sequence is the trajectory (70 weeks). drift_magnitude = 1 - cosine(week t, "
   "week t+1). CUSUM accumulates excess drift and fires on sustained low-level change no single week would trip.")
B("Live example - Volt's access_pattern_drift, weeks 64-69: ","0.068, 0.021, 0.017, 0.021, 0.232, 0.212. Flat near "
   "0.02 for weeks, then a 10x jump in the final two weeks. The ramp is visible in the trajectory while his composite "
   "drift only creeps 0.013 to 0.033.")

H2("Part 5 - Drift direction / MITRE projection (meaning)")
B("","drift_vector = normalize(v_new - v_old), projected (cosine) onto 14 threat + 2 benign concepts, each a prose "
   "description embedded in the same space. Alignment >= 0.15 to a threat concept flags it and emits the MITRE techniques.")
B("Output is not '87% anomalous' but: ","'network_footprint DRIFTING (0.082) toward c2_beacon [T1071, T1573].' "
   "Zone-specific filtering stops cross-domain false alignments.")
B("Concept-to-attacker map: ","living_off_the_land (T1218/T1053/T1059.001) -> Volt; telecom_infrastructure_pivot "
   "(T1557/T1040/T1005) -> Salt; insider_threat_slow (T1078/T1083/T1005) -> Insider/Slow APT; c2_beacon "
   "(T1071/T1573/T1568) -> Slow APT.")

H2("Part 6 - Composite scorer (5 phases -> one ranked number)")
B("","Features become group-relative z-scores (each user vs the normals in their own role group: admin, security, "
   "developer, business, executive), then fuse by this exact formula:")
P("composite = signal_strength x1.0  +  breadth(z>1.5) x0.5  +  sustained x0.3  +  ctx_spread x0.5  +  ctx_max x0.3  +  novelty x1.0",
  italic=True)

H2("Part 7 - Novelty persistence (the C2 catch)")
B("","Tracks external IPs never seen in the first 10 baseline weeks. A persistent novel IP yields "
   "novelty = min(persistence/5, 10), plus a bonus if it appears in >50% of weeks. This is the single most important "
   "stealth component - see Slow APT below.")
meta("Build/roadmap line:","“All seven parts run live. Identity fusion across systems, the versioned twin store, "
 "the behavioral graph, tamper-evident forensics, and automated response are roadmap, not running today.”")

# ===== 6 BLOCK 3 - THE PROOF =====
H1("6. Block 3 - LIVE demo: the four attackers + 8.1% (45 min, Esquibel)")
P("Four attackers hidden among 246 normals; 250 entities, 485 days, ~14M events, 5 log sources; full ground truth. "
  "Every composite below re-derives live from the formula in Part 6.",italic=True)
table(["Attacker","UID","Rank","Composite","Caught by"],
 [["Salt Typhoon","USR-118","#1","51.27","Breadth(18) + context spread(8.65). QA Engineer in the Legal dept (cross-domain). Broad, sustained, multi-front."],
  ["Insider","USR-156","#2","46.24","Highest signal_strength of all (30.02), breadth 19. Classic broad insider."],
  ["Slow APT","USR-234","#7","19.44","Novelty = 13.0 ALONE. signal_strength 4.47, breadth = 1: invisible to every threshold method. A novel C2 IP every week for 60 weeks lifts it to rank 7."],
  ["Volt Typhoon","USR-042","#24","13.70","Breadth(11) of small elevations, none extreme (breadth z>2.0 = only 3), novelty 0. The living-off-the-land signature: many small fronts, nothing dramatic."]],
 widths=[1.0,0.8,0.5,0.8,2.9],fs=8.5)
B("The live 8.1% derivation: ","sort all 250 by composite. Volt at 13.70 is the stealthiest scorer, so the line must "
  "reach 13.70 to catch all four. 20 of 246 normals sit above that line = 8.1% false positive, 4/4 caught. On the "
  "same data, Isolation Forest, One-Class SVM and LOF each catch 0/4; a z-score method catches 1/4.")
B("Say the hard truth first: ","“Our top normal user (composite 21.5) outscores two of our four attackers. That "
  "is exactly why we deliver a RANKED LIST, not a threshold. The honest claim is recall at the top of the ranking.”")
meta("Esquibel will pressure-test:","FP at Army scale, LOTL fidelity, baseline drift / poisoning. Answers in Section 8.")

# ===== 7 BLOCK 4 - FIELDING =====
H1("7. Block 4 - Fielding the twin (30 min, DiGrezio)")
P("This block is NEW for this audience and it is what wins your acquisition seat. Speak programs, not algorithms.",italic=True)
B("Data / telemetry requirements: ","runs on logs the Army already collects: sign-in, network, endpoint, file, and "
  "identity. Five sources. No new sensors required; the twin is built from existing telemetry.")
B("Deployment footprint: ","containerized; deployable across IL5, IL6, and JWICS; offline embedding mode runs today; "
  "a local semantic model replaces any commercial model in air-gapped enclaves (Phase 1); no PII held in the vector.")
B("Integration: ","sits on the existing SIEM / log pipeline and returns a risk-ranked, MITRE-mapped list to analysts. "
  "It augments hunt operations; it does not rip and replace.")
B("Maturity / TRL honesty: ","the detection core (the seven parts) is built and demoed today. Identity fusion, "
  "versioned store, graph, forensics, and response are the next build phase. Do not present them as fielded.")
B("Acquisition pathway: ","bounded Phase-1 pilot on one enclave's real telemetry -> OTA / prototype -> program of "
  "record. Pilot measures detection, explainability, and FP economics side-by-side against current tooling, advisory "
  "and human-in-the-loop.")
B("Sustainment: ","model and concept updates are data, not re-instrumentation; the 14 threat concepts are versioned "
  "text and extensible as ARCYBER adds new threat patterns.")
meta("DiGrezio will ask:","“What does it take to field this in an Army enclave?” Answer: connectors for one "
 "enclave's identity + endpoint logs, identity resolution for a defined population, run the validated core on real "
 "telemetry. That is the pilot, and it is the ask.")

# ===== 8 BLOCK 5 - WHY IT MATTERS + ASK =====
H1("8. Block 5 - Why it matters + the ask (20 min, Herrmann)")
B("Strategic significance: ","valid-account, living-off-the-land intrusion is the nation-state technique of record "
  "against U.S. critical infrastructure (Volt, Salt). The tools agencies field today score these campaigns as normal. "
  "A behavioral twin closes that gap on telemetry already collected.")
B("Authorities / privacy: ","no PII in the behavioral vector; the twin models behavior, not identity content. This "
  "eases the authorities and privacy posture for monitoring a defined population.")
B("Readiness / workforce: ","output is a ranked, explainable, MITRE-mapped list, which raises analyst leverage rather "
  "than adding alert noise. It fits hunt workflows ARCYBER already runs.")
P("The ask (say it slowly): “What you saw run today is a behavioral digital twin, on the logs you already "
  "collect, catching four nation-state campaigns your current tools score as normal. The detection core is ours and "
  "runs live; identity fusion, forensics, and response are what we would build with you. The honest ask is a bounded "
  "pilot on real Army telemetry, advisory and human-in-the-loop, with detection and false-positive economics measured "
  "side-by-side against current tooling.”",italic=True)

# ===== 9 PER-PERSON HARD Q =====
H1("9. Per-person hard questions")
H2("Esquibel (operator)")
B("'8% FP at Army scale?' ","Ranked list, not threshold alerts; analysts work top-down; the operating point is a PoC "
  "reporting benchmark to re-validate and tune on real telemetry.")
B("'How real is LOTL detection?' ","Volt is the worked example: many small elevations, no novelty, caught by breadth "
  "across fronts. We show the per-front z-scores live.")
B("'Baseline drift / data poisoning?' ","Baseline is the entity's own first-10-weeks plus role-group normals; "
  "hardening the detector (adversarial ML, MITRE ATLAS) is on the roadmap and we say so.")
B("'How does this fit ARCYBER hunt?' ","It pre-ranks entities for hunt teams and names the MITRE direction, so hunters "
  "start at the top of an explainable list rather than triaging raw alerts.")
H2("DiGrezio (acquisition)")
B("'Telemetry / data requirements?' ","Five existing log sources; no new sensors.")
B("'Footprint and accreditation?' ","Containerized; IL5/IL6/JWICS; offline / air-gap mode; no PII in the vector.")
B("'Integration with our stack?' ","Augments the existing SIEM/log pipeline; outputs a ranked MITRE-mapped list.")
B("'Maturity / what is actually built?' ","Detection core built and demoed; identity fusion, forensics, response = "
  "next phase. Honest TRL.")
B("'Acquisition path?' ","Bounded pilot -> OTA/prototype -> PoR.")
H2("Herrmann (policy)")
B("'Why does this matter now?' ","Volt/Salt prove valid-account intrusion is the dominant critical-infra technique and "
  "current tools miss it.")
B("'Privacy / authorities?' ","No PII in the behavioral representation; models behavior, not content.")
B("'Workforce impact?' ","Raises analyst leverage with explainable ranking; does not add noise.")

# ===== 10 NUMBERS =====
H1("10. Numbers to keep straight (Esquibel will check)")
B("Dataset: ","250 entities, 485 days (70 weeks), ~14M events, 5 log sources.")
B("FP: ","8.1% composite (~20/246 normals above Volt's 13.70), 4/4 caught. Use 8.1, never 8.5.")
B("Ranks: ","Salt USR-118 #1 (51.27); Insider USR-156 #2 (46.24); Slow APT USR-234 #7 (19.44); Volt USR-042 #24 (13.70).")
B("Baselines: ","Isolation Forest / OC-SVM / LOF each 0/4; z-score 1/4.")
B("Separation is by RANK, not a clean gap: ","top normal (21.5) outscores Slow APT (19.4) and Volt (13.7). Present a ranked list.")
B("Novelty math: ","Slow APT = novel IP persisting 60 weeks -> min(60/5,10)=10, +3 for 100% frequency = 13.0, which alone makes rank 7.")

fp=d.add_paragraph(); fr=fp.add_run("Internal preparation playbook. Built-vs-roadmap distinctions and the reconciled "
 "number set are load-bearing with an ARCYBER audience. Not for distribution.")
fr.italic=True; fr.font.size=Pt(9); fr.font.color.rgb=GRAY

out="WP DLA/Presentation/Digital_Twin_Working_Session_RunOfShow_INTERNAL.docx"
os.makedirs("WP DLA/Presentation",exist_ok=True)
d.save(out)
print("saved:",out,"| paragraphs:",len(d.paragraphs),"| tables:",len(d.tables))
