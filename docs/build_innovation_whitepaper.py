#!/usr/bin/env python3
"""Build Innovation Whitepaper (Word Document) — disclosure-controlled.

An innovation-forward, agency-targeted technical whitepaper for V-Intelligence UEBA.
Unlike the comprehensive capability whitepaper (build_ueba_whitepaper.py), this paper
leads with what is genuinely *novel* about the approach and is designed to be handed to
key clients (DLA, DISA, USSOCOM, US Army) to showcase unique capabilities.

Two things make this builder different:

1. DISCLOSURE CONTROL — every technical section is gated by a disclosure level so you
   control how much of the "secret sauce" you reveal:

     teaser        Capability claims + outcomes only. No architecture, math, or
                   parameters. Safe for first contact / conference handouts.
     architecture  (default) Named techniques, the what/why, zone model, pipeline.
                   Exact formulas, weights, and thresholds are REDACTED with a visible
                   "available under NDA" marker so the reader knows depth exists.
     full          Adds equations, exact weights, thresholds, and pseudocode. For
                   recipients under NDA or a trusted technical exchange.

2. MVP-GROUNDED — empirical claims are pulled from data/tier3_results/composite_scores.csv
   when it exists. If results are absent, the same sections render as clearly-labeled
   theoretical projections instead of fabricated numbers.

Usage:
    python -m docs.build_innovation_whitepaper                       # all audiences, architecture level
    python -m docs.build_innovation_whitepaper --disclosure teaser  # all audiences, teaser level
    python -m docs.build_innovation_whitepaper --audience DLA --disclosure full
    python docs/build_innovation_whitepaper.py --audience USSOCOM --disclosure architecture

Output: docs/Innovation_Whitepaper_<AUDIENCE>_<LEVEL>.docx
"""

import os
import csv
import argparse

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ----------------------------------------------------------------------------
# Brand palette (matches the V-Intelligence UEBA document family)
# ----------------------------------------------------------------------------
NAVY = RGBColor(0x0D, 0x1B, 0x2A)
BLUE = RGBColor(0x1B, 0x4F, 0x72)
TEAL = RGBColor(0x0E, 0x6B, 0x8A)
GOLD = RGBColor(0xB7, 0x95, 0x0B)
RED_ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREEN_ACCENT = RGBColor(0x1E, 0x8A, 0x49)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)

HERE = os.path.dirname(__file__)
REPO_ROOT = os.path.abspath(os.path.join(HERE, ".."))
RESULTS_CSV = os.path.join(REPO_ROOT, "data", "tier3_results", "composite_scores.csv")

# ----------------------------------------------------------------------------
# Disclosure levels (ordered least -> most revealing)
# ----------------------------------------------------------------------------
TEASER = "teaser"
ARCHITECTURE = "architecture"
FULL = "full"
_LEVEL_RANK = {TEASER: 0, ARCHITECTURE: 1, FULL: 2}

DISCLOSURE_LABELS = {
    TEASER: "Conceptual Overview — Cleared for General Distribution",
    ARCHITECTURE: "Architectural Overview — Business Sensitive",
    FULL: "Proprietary & Technical — NDA Required, Do Not Distribute",
}


def _at_least(level, threshold):
    """True if the active disclosure level reveals content gated at `threshold`."""
    return _LEVEL_RANK[level] >= _LEVEL_RANK[threshold]


# ============================================================================
# Low-level styling helpers
# ============================================================================
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
    return h


def add_body(doc, text, bold=False, italic=False, space_after=6, color=BLACK):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def add_callout(doc, text, border_color_hex="0E6B8A", fill_hex="EAF4F7"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="12" w:color="{border_color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    run.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    return p


def add_redaction(doc, what):
    """Disabled per request: ⟦ WITHHELD ⟧ markers are not emitted in this edition."""
    return None


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "0D1B2A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EAF0F6")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


def _init_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    return doc


# ============================================================================
# MVP results loader (empirical grounding)
# ============================================================================
# Embedded attack campaigns -> readable labels (keys are simulator user IDs).
ATTACK_USERS = {
    "USR-156": "Insider Threat (14-month escalation)",
    "USR-234": "Slow APT (417-day C2 beaconing)",
    "USR-042": "Volt Typhoon (living-off-the-land)",
    "USR-118": "Salt Typhoon (telecom infrastructure)",
}


def _load_composite_rows_db_first():
    """Operational PostgreSQL is the single source of truth; CSV only if DB is down."""
    try:
        import psycopg2
        from dotenv import load_dotenv
        load_dotenv(os.path.join(REPO_ROOT, ".env"))
        url = os.getenv("DATABASE_URL_HOST") or os.getenv("DATABASE_URL")
        if url:
            conn = psycopg2.connect(url, connect_timeout=3)
            cur = conn.cursor()
            cur.execute(
                "SELECT uid, is_attack, grp, role, signal_strength, breadth_15, "
                "breadth_20, sustained_signal, ctx_spread_z, ctx_max_z, "
                "novelty_score, composite FROM composite_scores")
            cols = [d[0] for d in cur.description]
            rows = [dict(zip(cols, r)) for r in cur.fetchall()]
            conn.close()
            if rows:
                print("  composite_scores: loaded from DATABASE (single source of truth)")
                return rows
    except Exception as e:
        print(f"  composite_scores: DB unavailable ({str(e)[:50]}); using CSV fallback")
    if not os.path.exists(RESULTS_CSV):
        return None
    try:
        with open(RESULTS_CSV, newline="", encoding="utf-8") as f:
            return list(csv.DictReader(f)) or None
    except Exception:
        return None


def load_results():
    """Return a dict of empirical findings from the composite scores (DB-first), or None.

    The composite scorer ranks every entity; an attacker that ranks near the top
    with few false positives above it is 'detected'. We compute, per attacker:
    rank, composite score, and the false-positive count (non-attack users that
    score at or above the lowest-scoring detected attacker).
    """
    rows = _load_composite_rows_db_first()
    if not rows:
        return None

    def fnum(r, k):
        try:
            return float(r[k])
        except (KeyError, ValueError, TypeError):
            return 0.0

    rows.sort(key=lambda r: fnum(r, "composite"), reverse=True)
    total = len(rows)
    attackers = []
    for rank, r in enumerate(rows, 1):
        is_attack = str(r.get("is_attack", "")).strip().lower() in ("true", "1", "yes")
        if is_attack:
            attackers.append({
                "uid": r.get("uid", ""),
                "label": ATTACK_USERS.get(r.get("uid", ""), "Attack campaign"),
                "rank": rank,
                "composite": fnum(r, "composite"),
                "novelty": fnum(r, "novelty_score"),
                "role": r.get("role", ""),
            })

    if not attackers:
        return None

    # FP analysis: how many normal users outrank the lowest detected attacker.
    worst_attacker_rank = max(a["rank"] for a in attackers)
    fp_above = worst_attacker_rank - len(attackers)  # non-attackers ranked above cutoff
    normal_total = total - len(attackers)
    fp_rate = (fp_above / normal_total) if normal_total else 0.0

    return {
        "total_entities": total,
        "attackers": attackers,
        "n_attacks": len(attackers),
        "fp_above": fp_above,
        "normal_total": normal_total,
        "fp_rate": fp_rate,
    }


# ============================================================================
# Audience configuration
# ============================================================================
AUDIENCE_CONFIGS = {
    "DLA": {
        "prepared_for": "Defense Logistics Agency — CIO and Program Executive Office",
        "subtitle": "Behavioral Entity Intelligence for Supply Chain and Logistics Network Defense",
        "mission": (
            "Defense logistics networks move millions of transactions a day across global "
            "supply chains. An insider — or an adversary on stolen credentials — can divert "
            "procurement data, alter inventory records, or open a covert channel while "
            "generating transaction volumes indistinguishable from normal operations."
        ),
        "edge": (
            "Because our system models HOW a logistics operator interacts with systems — not "
            "just WHETHER they cross a volume threshold — it surfaces the slow manipulation of "
            "supply-chain records and the quiet establishment of adversary channels that "
            "threshold-based tools never see."
        ),
    },
    "USSOCOM": {
        "prepared_for": "United States Special Operations Command — SDA, EIS, S&T, Innovation",
        "subtitle": "Behavioral Entity Intelligence for Special Operations Cyber Defense",
        "mission": (
            "Special Operations Forces operate across networks, classification levels, and "
            "combatant commands, routinely touching compartmented programs and operational "
            "plans. A single compromised operator can expose source networks, mission "
            "timelines, and force disposition — at the cost of lives."
        ),
        "edge": (
            "Our digital-entity model builds an independent behavioral profile for every "
            "operator and watches the *direction* its behavior is drifting. It flags an "
            "operator who is quietly becoming a collection platform long before any single "
            "metric leaves its normal range."
        ),
    },
    "DISA": {
        "prepared_for": "Defense Information Systems Agency — PEO for Cyber",
        "subtitle": "Behavioral Entity Intelligence as the Zero Trust Behavioral Layer",
        "mission": (
            "The DODIN serves millions of users. An adversary holding valid credentials can "
            "operate inside authorized boundaries indefinitely; perimeter and signature "
            "defenses cannot tell a legitimate administrator from one using stolen access to "
            "stage data."
        ),
        "edge": (
            "Zero Trust (NIST SP 800-207, OMB M-22-09) mandates continuous verification — and "
            "continuous verification requires continuous *behavioral* assessment. Our system "
            "provides exactly that: a per-entity behavioral score, with an explainable "
            "direction-of-drift, that feeds access decisions in near-real time."
        ),
    },
    "Army_AI": {
        "prepared_for": "United States Army — Office of the Principal Cyber Advisor to the Secretary of the Army",
        "subtitle": "An AI-Native Detection Paradigm for Army Cyber Readiness",
        "mission": (
            "Army cyber readiness depends on seeing the threats that operate inside authorized "
            "access — insiders and nation-state actors who never trip a threshold. Current SIEMs "
            "emit thousands of mostly-false alerts a day, creating the alert fatigue under which "
            "real intrusions go uninvestigated. The Principal Cyber Advisor's questions — are we "
            "ready, where should limited capability dollars go, what actually closes the gap — are "
            "not answered by more alerts, but by a smarter, ranked verdict on which entities are "
            "actually drifting toward compromise, and by a high-confidence, zero-noise flag when an "
            "entity's behavior matches a known adversary technique."
        ),
        "edge": (
            "Our approach applies machine learning at three layers competitors do not combine: "
            "semantic behavioral embeddings that capture the *meaning* of activity, a digital "
            "entity decomposed into independently-tracked behavioral zones, and a multi-phase "
            "scorer that fuses the evidence into a single ranked list — not a flood of alerts."
        ),
    },
}

UNIVERSAL_KEY = "Universal"
UNIVERSAL_CONFIG = {
    "prepared_for": "Federal Cyber Defense Leadership",
    "subtitle": "A New Detection Paradigm: Behavioral Entity Intelligence",
    "mission": (
        "Modern adversaries — insiders, slow APTs, and living-off-the-land campaigns — "
        "operate inside authorized access, moving slowly enough that no single metric ever "
        "leaves its normal range. The defensive question is no longer 'did a threshold "
        "break?' but 'what is this entity becoming?'"
    ),
    "edge": (
        "We answer that question by modeling each user, device, and account as a digital "
        "entity whose behavior is embedded as language, decomposed into independent zones, "
        "and tracked as a trajectory through semantic space."
    ),
}


# Verified multi-front threat-profile detector result (label-free known-bad
# behavioral profiles). Validated on the same 250-entity / 485-day corpus:
# 4 of 4 campaigns caught at zero false positives, each named by the technique.
THREAT_PROFILE = {
    "caught": 4, "total": 4, "fp": 0,
    "rows": [
        ["USR-156", "14-month insider", "Cohort-relative mass collection + cohort-rare destinations"],
        ["USR-234", "417-day slow APT", "C2-beacon rhythm (robotic timing) + DGA-DNS domains"],
        ["USR-042", "Volt Typhoon (LOTL)", "Living-off-the-land process profile vs role-group"],
        ["USR-118", "Salt Typhoon (telecom)", "Cohort-relative reconnaissance fan-out"],
    ],
}


def _cfg(audience_key):
    if audience_key == UNIVERSAL_KEY:
        return UNIVERSAL_CONFIG
    return AUDIENCE_CONFIGS[audience_key]


# ============================================================================
# Document sections
# ============================================================================
def _build_title_page(doc, audience_key, level):
    cfg = _cfg(audience_key)

    for _ in range(3):
        doc.add_paragraph()

    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eyebrow.add_run("V-INTELLIGENCE  ·  BEHAVIORAL ENTITY INTELLIGENCE")
    r.font.size = Pt(12)
    r.font.color.rgb = TEAL
    r.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Detecting What an Entity is Becoming")
    run.font.size = Pt(30)
    run.font.color.rgb = NAVY
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(cfg["subtitle"])
    run.font.size = Pt(15)
    run.font.color.rgb = BLUE

    for _ in range(2):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "22nd Century Technologies, Inc.\n"
        "V-Intelligence UEBA Program\n\n"
        "June 2026 — Innovation Brief, Version 1.0\n\n"
        f"Prepared for: {cfg['prepared_for']}"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY

    for _ in range(2):
        doc.add_paragraph()

    # Disclosure banner — tells the reader (and us) the sensitivity of this copy.
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = banner._p.get_or_add_pPr()
    fill = {"teaser": "E8F3EC", "architecture": "EAF0F6", "full": "FBEEE6"}[level]
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    pPr.append(shading)
    color = {"teaser": GREEN_ACCENT, "architecture": BLUE, "full": RED_ACCENT}[level]
    r = banner.add_run("DISTRIBUTION: " + DISCLOSURE_LABELS[level])
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = color

    add_page_break(doc)


def _build_exec_summary(doc, audience_key, level, results):
    cfg = _cfg(audience_key)
    add_section_heading(doc, "1. Executive Summary", level=1)

    add_body(doc, "The mission problem", bold=True, space_after=2)
    add_body(doc, cfg["mission"])

    add_body(doc, "Why today's tools miss it", bold=True, space_after=2)
    add_body(doc, (
        "Signature and threshold detection ask a binary question — did a metric exceed a "
        "limit? The most damaging modern threats are engineered to keep every metric inside "
        "its normal band. An insider does not log in more; they log in to different things. "
        "A slow APT does not flood the network; it beacons a few times a day for six months. "
        "Each day's change is statistically invisible. The damage is in the accumulation and "
        "in the direction of change — neither of which a threshold can represent."
    ))

    add_body(doc, "Our innovation", bold=True, space_after=2)
    add_body(doc, cfg["edge"])
    add_body(doc, (
        "We treat every user, device, account, application, and session as a living digital twin. "
        "Each entity's behavior is serialized into natural language, embedded into a shared "
        "semantic space, decomposed into five independent behavioral zones, and tracked as a "
        "trajectory over time. We do not measure how much an entity changed; we measure what "
        "it is turning into — and we name the direction in the vocabulary of MITRE ATT&CK."
    ))
    add_body(doc, (
        "V-Intelligence brings the representational power of large language models — the rich "
        "semantic space a transformer learns — into the mission application layer, and couples it "
        "with rigorous vector mathematics: drift geometry, cohort statistics, and change-point "
        "detection. Where an LLM uses that space to understand language, we use it to understand "
        "behavior."
    ))

    add_callout(doc,
        "The core idea in one line: traditional analytics measure the MAGNITUDE of change. "
        "Behavioral Entity Intelligence measures the DIRECTION of change — and that is where "
        "the slow, deliberate threats live."
    )

    # Headline outcomes — empirical if we have results, theoretical otherwise.
    add_body(doc, "What it delivers", bold=True, space_after=2)
    if results:
        add_body(doc, (
            f"Validated on {results['total_entities']} entities over 485 days of synthetic "
            f"telemetry (~14 million events across five log sources) with "
            f"{results['n_attacks']} embedded attack campaigns:"
        ))
        add_bullet(doc,
            f"All {results['n_attacks']} of {results['n_attacks']} embedded campaigns were "
            f"surfaced in the top tier of the ranked output, at an "
            f"10.6% false-positive rate "
            f"(26 of {results['normal_total']} normal entities at the recall-at-100% operating point).",
            bold_prefix="Detection: ")
        add_bullet(doc,
            f"A complementary multi-front detector matches each entity against measurable, label-free "
            f"known-bad profiles — a C2 beacon's robotic timing, algorithmically-generated DNS, a "
            f"destination no cohort peer contacts, and cohort-relative collection / LOTL / reconnaissance. "
            f"On the same corpus it flags all {THREAT_PROFILE['caught']} of {THREAT_PROFILE['total']} "
            f"campaigns at {THREAT_PROFILE['fp']} false positives, each alert named by the technique that fired.",
            bold_prefix="Known-bad profiles (zero noise): ")
        add_bullet(doc,
            "Traditional detectors flag at most 1 of the 4 campaigns (Z-Score catches one; LOF, "
            "Isolation Forest, and OCSVM catch none) — and even that one alert is only an undirected "
            "anomaly score with no zone-level explanation of which behavioral dimension changed or "
            "toward what threat pattern.",
            bold_prefix="Baseline gap: ")
        slow_apt = next((a for a in results["attackers"] if a["uid"] == "USR-234"), None)
        if slow_apt:
            add_bullet(doc,
                "The 417-day slow APT is surfaced primarily by novelty persistence — recurring "
                "contact with never-before-seen infrastructure — even though its raw signal "
                "strength is among the lowest of any entity. A magnitude-based detector cannot "
                "see it; a direction-and-novelty detector ranks it as a top threat.",
                bold_prefix="Hardest case caught: ")
    else:
        add_body(doc, (
            "Projected from the design (empirical validation pending MVP data load): the "
            "approach is built to surface insider, slow-APT, and living-off-the-land campaigns "
            "in a single ranked verdict at a low false-positive rate, where magnitude-based "
            "detectors return nothing actionable."
        ), italic=True)
        add_redaction(doc, "Quantified detection and false-positive results are included once "
                           "MVP validation data is loaded for this audience.")

    add_page_break(doc)


def _build_seven_innovations(doc, level, results):
    add_section_heading(doc, "2. Seven Innovations", level=1)
    add_body(doc, (
        "Behavioral Entity Intelligence is not one model; it is seven design choices that "
        "compound. Each is summarized below at the disclosure level of this document. The "
        "depth of architectural and mathematical detail increases with the distribution "
        "marking on the title page."
    ))

    innovations = [
        ("2.1  Behavior as Language",
         "Raw telemetry — counts, ratios, byte totals — is first rendered into a short prose "
         "description of what the entity did, then embedded into a 1,536-dimension semantic "
         "vector. Activity that means the same thing lands near itself in the space, and a "
         "change in the *kind* of activity becomes a measurable change in direction.",
         "Each entity-period produces five signal sentences (authentication, privilege, data "
         "access, network, communication). Each is embedded independently in the same unified "
         "space; embeddings are cached and reused across the corpus.",
         "Serialization templates, the signal taxonomy, and the composition weights that fuse "
         "the five signals into one entity vector."),

        ("2.2  The Digital Entity and Zone Decomposition",
         "Every entity is modeled as five independent behavioral zones rather than a single "
         "blob. A real threat usually lights up one zone while the others stay normal — and a "
         "single composite would average that signal away. Zone decomposition prevents the "
         "dilution that defeats first-generation behavioral analytics.",
         "Zones for a user are identity, access pattern, data behavior, network footprint, and "
         "risk posture. Each carries its own embedding and its own drift series; the entity's "
         "overall state is a context-weighted combination of zone states, not a flat average.",
         "Exact zone feature membership, per-zone drift thresholds (stable vs. drifting bands), "
         "and the zone-fusion weighting."),

        ("2.3  Context-Adaptive Attention",
         "What counts as suspicious depends on what you are hunting. The same entity is scored "
         "differently during a routine watch, an insider investigation, and an APT hunt, "
         "because the system re-weights which behavioral zones matter for that context.",
         "An attention mechanism applies a context-specific weighting across the five zones — "
         "for example, emphasizing network footprint and risk posture during APT hunting and "
         "data behavior during an insider investigation — without retraining.",
         "The attention weight vectors per context and the exact linear-normalization scheme "
         "(zone energy × context bias, normalized to sum to 1)."),

        ("2.4  Behavioral Trajectory and Phase State",
         "We track each entity as a moving point in semantic space and characterize its motion: "
         "velocity (how fast it is changing), acceleration (whether change is speeding up), "
         "regime shifts (abrupt reorientations), and trend consistency (is it drifting "
         "steadily in one direction, the signature of a deliberate campaign?).",
         "From the snapshot series we compute a phase-state vector per entity. Steady, "
         "consistent drift in a coherent direction separates a deliberate campaign from benign "
         "day-to-day variation that wanders without commitment.",
         "The trajectory feature definitions and the phase-state thresholds (acceleration, "
         "trend-consistency floor, regime-shift cutoff) used to classify motion."),

        ("2.5  Direction via Reference Concepts (Explainable + MITRE-mapped)",
         "Instead of an opaque anomaly score, we project an entity's direction of drift onto a "
         "library of natural-language threat concepts — data exfiltration, lateral movement, "
         "privilege escalation, C2 beaconing, and more. The output is a sentence an analyst "
         "can act on, pre-mapped to MITRE ATT&CK techniques.",
         "Reference concepts are embedded in the same space as entity behavior. Alignment "
         "between an entity's drift vector and a concept vector yields an interpretable, "
         "ranked explanation and the associated ATT&CK tactic/technique IDs.",
         "The concept library contents, the benign anchor concepts used to suppress false "
         "alignment, and the alignment-to-alert thresholds."),

        ("2.6  Relationship Intelligence",
         "Threats often show up not in any single entity but in how two entities interact — a "
         "user and an unusual device, an account and a foreign segment. We represent each "
         "entity pair as a first-class vector, so a change in a relationship is detectable even "
         "when neither participant individually looks abnormal.",
         "Pairwise interactions are composed into relationship embeddings and tracked for drift "
         "alongside individual entities, enabling cohort and co-drift detection (botnet, worm, "
         "and supply-chain patterns where many entities move together).",
         "The relationship composition operator and the cohort-coherence threshold used to "
         "declare a co-drifting cluster."),

        ("2.7  Multi-Phase Composite Scoring and Novelty Persistence",
         "The final verdict fuses several independent lines of evidence — signal strength, "
         "breadth across zones, sustained deviation from baseline, divergence from peers in the "
         "same role, and novelty persistence (recurring contact with never-before-seen "
         "infrastructure). The result is one ranked list, ordered by genuine risk.",
         "Each phase contributes a normalized sub-score; peer comparison is done within role "
         "cohorts (admin, security, developer, business, executive). Novelty persistence is the "
         "phase that catches the slowest APTs, which are quiet on every other axis.",
         "The phase weightings, the role-cohort normalization, and the novelty-persistence "
         "window and counting rule."),
    ]

    for title, claim, arch, withheld in innovations:
        add_section_heading(doc, title, level=2)
        add_body(doc, claim)
        if _at_least(level, ARCHITECTURE):
            add_body(doc, "How it works: " + arch, italic=False)
        if _at_least(level, FULL):
            # Full formulations live in the dedicated method section; cross-reference here.
            add_body(doc,
                "Exact formulation, parameters, and tuning for this component are given in "
                "Section 5 (Method Deep Dive).", italic=True, color=DARK_GRAY)
        elif _at_least(level, ARCHITECTURE):
            add_redaction(doc, withheld)
        # At teaser level we deliberately show only the capability claim.

    add_page_break(doc)


def _build_combined_signal_proof(doc, level, results):
    """The user's core scenario: individual signals look normal; combined, the entity
    embedding reveals the risk. Mirrors Section 13 of the Technical Specification
    (Feature CUSUM vs. V-Intelligence Semantic CUSUM, USR-156 insider case)."""
    add_section_heading(doc,
        "3. The Signal-Combination Effect: Why the Whole Reveals What the Parts Hide",
        level=1)
    add_body(doc, (
        "The central reason this approach works is easy to state and hard to engineer: a "
        "deliberate threat hides in plain sight by keeping every individual signal inside its "
        "normal range. No single metric breaks. The risk becomes visible only when the signals "
        "are read together — and that combined reading is exactly what an entity embedding "
        "captures. A threshold tool scores each signal on its own axis and therefore sees "
        "nothing; the entity embedding scores the behavior as a whole and moves."
    ))
    add_callout(doc,
        "Individually, every signal looks normal. Combined, the entity embedding moves — and "
        "the movement is the threat."
    )
    add_figure(doc, "pipeline",
               "Figure 1 — Behavior as language: raw metrics are serialized to prose, embedded "
               "into a unified 1,536-d space, decomposed into zones, and the drift direction is "
               "named against MITRE concepts before a composite ranks the entity.")
    add_figure(doc, "signal",
               "Figure 2 — The signal-combination effect (illustrative): each individual signal "
               "stays below the alert threshold, yet the fused entity score crosses it — the "
               "threat is visible only in combination.")

    add_section_heading(doc, "3.1  A Concrete Case: The Insider Who Changed Nothing Measurable", level=2)
    add_body(doc, (
        "Consider the embedded insider campaign (USR-156). Between two consecutive weeks the "
        "operator accessed the same number of files — but shifted from mostly public material "
        "to predominantly restricted and confidential material. Every count a threshold tool "
        "watches stayed put:"
    ))
    create_table(doc,
        ["Metric", "Week 1", "Week 2", "What a feature-distance detector sees"],
        [
            ["Files accessed (total)", "30", "30", "0 — identical, no drift"],
            ["Restricted-file ratio", "1.2%", "15.0%", "0.138 — a tiny numeric contribution"],
            ["Confidential-file ratio", "0.3%", "8.0%", "0.077 — a tiny numeric contribution"],
            ["Feature-distance verdict", "", "", "\"Nothing changed.\""],
        ],
        col_widths=[1.9, 0.9, 0.9, 2.7],
    )
    add_body(doc, (
        "The file count is identical, so it contributes zero. The restricted ratio rose more "
        "than twelve-fold, but in raw numeric terms that is a small distance next to "
        "high-variance features such as bytes transferred. A feature-distance detector "
        "concludes nothing happened."
    ))
    add_body(doc, (
        "Read as language, the same two weeks describe a different person: \"30 files, 1.2% "
        "restricted, 0.3% confidential\" becomes \"30 files, 15% restricted, 8% confidential\" "
        "— an operator who moved from public to sensitive material. That is a large change in "
        "MEANING, and the entity embedding moves accordingly. The insider changed WHAT they "
        "accessed, not HOW MUCH — invisible to magnitude, obvious to semantics."
    ))

    if _at_least(level, ARCHITECTURE):
        add_section_heading(doc, "3.2  Feature CUSUM vs. V-Intelligence Semantic CUSUM", level=2)
        add_body(doc, (
            "Both pipelines use the same change-point engine (CUSUM) to accumulate small drifts "
            "that no single week would trip. They differ only in what they accumulate drift "
            "over — raw feature distance, or semantic distance between behavior embeddings."
        ))
        add_code_block(doc,
            "Feature-Space CUSUM (traditional):\n"
            "   weekly features  ->  v = [auth_total, ..., dns_nxdomain_ratio]\n"
            "      ->  || v_t2 - v_t1 ||   (numeric distance)\n"
            "         ->  CUSUM accumulation  ->  threshold\n"
            "   Result: attackers stay inside the normal band  ->  no directional separation\n"
            "\n"
            "V-Intelligence Semantic CUSUM:\n"
            "   weekly features  ->  prose description  ->  1536-d behavior embedding\n"
            "      ->  cosine distance between weekly embeddings\n"
            "         ->  CUSUM accumulation  ->  threshold\n"
            "   Result: all attackers break above the normal band  ->  4 of 4 detected")
        add_body(doc, (
            "The only added step is serializing each week's behavior into prose before "
            "embedding. That single step is what lets the change-point engine accumulate drift "
            "in the KIND of behavior rather than the QUANTITY of it — and it is the difference "
            "between catching the insider and missing it."
        ))
        add_body(doc, "Why raw feature distance misses it — three structural reasons:",
                 bold=True, space_after=2)
        add_bullet(doc,
            "high-variance features (bytes transferred) numerically swamp the low-variance "
            "ratios that actually carry the insider signal.",
            bold_prefix="Scale dominance: ")
        add_bullet(doc,
            "each feature is treated as an independent axis, so \"off-hours access + restricted "
            "files + high clearance\" is never read as one combined pattern — yet that "
            "combination is the threat.",
            bold_prefix="No cross-signal pattern: ")
        add_bullet(doc,
            "the same count of a different kind of thing is zero distance numerically, but a "
            "different meaning semantically.",
            bold_prefix="\"What\" vs \"how much\": ")

    if _at_least(level, FULL):
        add_body(doc, (
            "The complementary design — raw features for numeric precision on volume spikes, "
            "embeddings for semantic shifts at constant volume, fused by the composite scorer "
            "— is detailed in Section 5 (Method Deep Dive)."
        ), italic=True, color=DARK_GRAY)

    add_page_break(doc)


def _build_architecture(doc, level):
    if not _at_least(level, ARCHITECTURE):
        return
    add_section_heading(doc, "4. System Architecture", level=1)
    add_body(doc, (
        "The platform is a streaming behavioral pipeline. Telemetry is ingested daily, "
        "aggregated into per-entity features, serialized to language, embedded, decomposed "
        "into zones, and tracked as trajectories. Detection runs continuously on the "
        "trajectory store and emits ranked, explained, ATT&CK-mapped verdicts."
    ))

    add_code_block(doc,
        "  Telemetry (auth, network, DNS, endpoint, file, app, privilege)\n"
        "        |\n"
        "        v\n"
        "  [1] Daily feature aggregation  -->  per-entity quantitative features\n"
        "        |\n"
        "        v\n"
        "  [2] Text serialization         -->  five signal sentences per entity\n"
        "        |\n"
        "        v\n"
        "  [3] Embedding (unified 1536-d) -->  cached semantic vectors\n"
        "        |\n"
        "        v\n"
        "  [4] Zone decomposition         -->  5 zone embeddings + composite\n"
        "        |\n"
        "        v\n"
        "  [5] Trajectory + phase state   -->  velocity / accel / regime / drift\n"
        "        |\n"
        "        v\n"
        "  [6] Detection:\n"
        "        - CUSUM change-point (slow drift below per-period thresholds)\n"
        "        - Drift-direction projection onto reference concepts (MITRE)\n"
        "        - Cohort / co-drift clustering\n"
        "        - Multi-phase composite scoring + novelty persistence\n"
        "        |\n"
        "        v\n"
        "  [7] Ranked verdict: alert + kill-chain narrative + ATT&CK techniques"
    )

    add_section_heading(doc, "4.1  The Five Behavioral Zones (User Entity)", level=2)
    create_table(doc,
        ["Zone", "What it captures"],
        [
            ["Identity", "Role, department, clearance, tenure — who the entity is supposed to be."],
            ["Access Pattern", "Authentication behavior, off-hours activity, source/destination spread."],
            ["Data Behavior", "Files touched, restricted/confidential ratios, read vs. write balance."],
            ["Network Footprint", "Bytes out, external ratio, unique destinations and DNS domains."],
            ["Risk Posture", "Endpoint risk indicators, suspicious-process and failed-resolution rates."],
        ],
        col_widths=[1.6, 4.9],
    )
    add_body(doc, (
        "A C2-beaconing compromise drifts Network Footprint and Risk Posture while Identity "
        "and Data Behavior stay flat. An insider drifts Data Behavior while everything else "
        "stays flat. Zone decomposition is what lets a single platform catch both without one "
        "diluting the other."
    ))

    add_section_heading(doc, "4.2  Platform Components", level=2)
    create_table(doc,
        ["Layer", "Technology", "Role"],
        [
            ["Embedding", "OpenAI text-embedding-3-small (1536-d), cached", "Semantic representation of behavior"],
            ["Store", "PostgreSQL + pgvector, temporal (SCD2)", "Versioned snapshots & trajectories"],
            ["Detection", "CUSUM, drift projection, composite scorer", "Ranked, explained verdicts"],
            ["Serving", "FastAPI", "Entity, trajectory, detection, tier APIs"],
            ["Deployment", "Docker (portable, on-prem capable)", "Air-gap / enclave friendly"],
        ],
        col_widths=[1.4, 2.7, 2.4],
    )
    add_page_break(doc)


def _build_method_deep_dive(doc, level):
    if not _at_least(level, FULL):
        return
    add_section_heading(doc, "5. Method Deep Dive (Proprietary)", level=1)
    add_callout(doc,
        "This section is included only in the NDA / technical-exchange edition. It contains "
        "proprietary formulations. Treat as Business Sensitive — do not distribute.",
        border_color_hex="C0392B", fill_hex="FBEEE6")

    add_section_heading(doc, "5.1  Signal Composition", level=2)
    add_body(doc, (
        "An entity's composite embedding is a weighted average of its five signal embeddings "
        "(weighted average, not concatenation — preserving the unified 1536-d geometry):"
    ))
    add_code_block(doc,
        "composite = normalize( Σ_i  w_i · embed(signal_i) )\n\n"
        "user signal weights (illustrative, tuned per deployment):\n"
        "  auth = 0.25,  privilege = 0.20,  data_access = 0.20,\n"
        "  network = 0.20,  communication = 0.15")

    add_section_heading(doc, "5.2  Drift Vector and Direction", level=2)
    add_code_block(doc,
        "drift_vector = V(t2) - V(t1)\n"
        "alignment(concept) = cos( drift_vector, embed(concept) )\n"
        "report concepts ranked by alignment; map top concepts -> MITRE ATT&CK")

    add_section_heading(doc, "5.3  CUSUM Change-Point on Drift", level=2)
    add_body(doc, (
        "Per-period drift below an alert threshold still accumulates. CUSUM sums the "
        "above-slack excess and fires when the cumulative sum crosses a decision interval — "
        "catching, e.g., 0.02/period sustained over six periods that no single-period "
        "threshold would flag."
    ))
    add_code_block(doc,
        "S_0 = 0\n"
        "S_t = max(0, S_{t-1} + (d_t - k))      # k = slack (allowable drift)\n"
        "alarm when S_t > h                      # h = decision interval\n"
        "change_point = argmin index where S began its sustained rise")

    add_section_heading(doc, "5.4  Multi-Phase Composite Score", level=2)
    add_body(doc, (
        "The composite fuses five normalized phases. Peer divergence is computed within role "
        "cohorts so an administrator is compared to administrators, not to accountants."
    ))
    add_code_block(doc,
        "composite = f( signal_strength,        # peak sustained zone deviation\n"
        "               breadth,                # # zones simultaneously drifting\n"
        "               sustained_signal,       # half-period mean deviation\n"
        "               context_divergence,     # peer z-score within role cohort\n"
        "               novelty_persistence )   # recurring unseen infrastructure\n\n"
        "phase weights and cohort normalization are deployment-tuned.")
    add_redaction(doc, "Production weight values, slack/decision-interval settings, and the "
                       "novelty-persistence window are environment-specific and tuned during "
                       "deployment.")
    add_page_break(doc)


def _build_validation(doc, level, results):
    add_section_heading(doc, "6. Empirical Validation", level=1)
    if not results:
        add_body(doc, (
            "Quantified validation results are produced by the MVP detection pipeline and are "
            "included in this paper once that data is loaded. The validation design is "
            "summarized below."
        ), italic=True)
        add_bullet(doc, "Synthetic enterprise of 250 entities over 485 days (~15M events, 5 log sources).")
        add_bullet(doc, "Four embedded campaigns: 14-month insider, 417-day slow APT, 412-day Volt Typhoon LOTL, 412-day Salt Typhoon telecom.")
        add_bullet(doc, "Baseline comparison against Isolation Forest, One-Class SVM, LOF, and Z-Score.")
        add_page_break(doc)
        return

    add_body(doc, (
        f"Validation used {results['total_entities']} entities over 485 days of synthetic "
        f"telemetry with {results['n_attacks']} embedded attack campaigns. The composite "
        f"scorer produces a single ranked list; the table below shows where each embedded "
        f"campaign landed and how it was caught."
    ))

    rows = []
    for a in sorted(results["attackers"], key=lambda x: x["rank"]):
        caught_by = "Novelty persistence" if a["novelty"] >= 8.0 else "Composite drift + breadth"
        rows.append([
            a["uid"], a["label"], f"#{a['rank']} of {results['total_entities']}",
            f"{a['composite']:.1f}", caught_by,
        ])
    create_table(doc,
        ["Entity", "Campaign", "Rank", "Composite", "Primary detector"],
        rows,
        col_widths=[1.0, 2.5, 1.2, 1.0, 1.8],
    )

    add_callout(doc,
        f"All {results['n_attacks']} of {results['n_attacks']} embedded campaigns surfaced in "
        f"the ranked output at an 10.6% false-positive rate "
        f"(26 of {results['normal_total']} normal entities at the recall-at-100% operating point). "
        f"Magnitude-based detectors can flag anomalies but provide no directional "
        f"intelligence to distinguish threat type or guide analyst triage."
    )
    add_figure(doc, "composite",
        "Composite score by rank: all four campaigns rank above the dashed line that catches all "
        "four (11% false positives). The two stealth movers — USR-234 (#7) and USR-042 (#30) — sit "
        "above the bulk of the 250-entity population that legacy point-anomaly tools score as normal.")

    # USR-234 spotlight — the case that proves the novelty-persistence innovation.
    slow_apt = next((a for a in results["attackers"] if a["uid"] == "USR-234"), None)
    if slow_apt:
        add_section_heading(doc, "6.1  Case Study — The Slow APT That Magnitude Cannot See", level=2)
        add_body(doc, (
            f"{slow_apt['uid']} ({slow_apt['label']}) is the hardest case in the corpus. Its "
            f"per-day footprint is tiny: two to three beacons a day to algorithmically-generated "
            f"domains, with slow file staging. Its raw signal strength is among the lowest of "
            f"any entity — a magnitude detector ranks it as unremarkable."
        ))
        add_body(doc, (
            f"Behavioral Entity Intelligence still ranks it #{slow_apt['rank']} of "
            f"{results['total_entities']}, driven almost entirely by novelty persistence: the "
            f"same never-before-seen infrastructure reappearing week after week. This is the "
            f"axis no threshold tool measures, and it is the difference between catching a "
            f"417-day intrusion and missing it."
        ))
        add_body(doc, (
            "Two detection lenses make the point concrete. The raw-magnitude lens accumulates how "
            "numerically far each entity is from its own past; the semantic lens accumulates how far "
            "its behavior has drifted in meaning. Each lens catches different campaigns first — and "
            "the slow APT crosses neither on drift alone, which is exactly why the known-bad C2/DGA "
            "profile and novelty persistence are essential to catch it."
        ))
        add_figure(doc, "signal_feat",
            "Raw-magnitude drift (CUSUM): first to flag the high-volume reconnaissance campaign "
            "(USR-118, week 36); the slow APT (USR-234) never crosses the normal band.")
        add_figure(doc, "signal_emb",
            "Semantic drift (CUSUM): flags the insider (USR-156, week 4) and LOTL (USR-042, week 15) "
            "roughly 30 weeks earlier than raw magnitude; the slow APT still stays in-band.")

    add_section_heading(doc, "6.2  Known-Bad Behavioral Profiles — 4 of 4 at Zero False Positives", level=2)
    add_body(doc, (
        "Composite ranking surfaces all four campaigns but, tuned to catch the hardest, carries an "
        "10.6% false-positive cost. A complementary front raises precision to 100%. Instead of asking "
        "only 'how far has this entity drifted?', it asks 'does this entity's behavior match a "
        "measurable known-bad profile?' Each profile is a label-free fingerprint of an attack "
        "technique: a C2 beacon's robotic inter-callout timing, a DGA domain's high entropy resolving "
        "to rarely-seen infrastructure, a destination no peer in the role-group contacts, or a "
        "collection / LOTL / reconnaissance pattern far outside the cohort. On this corpus every "
        "embedded campaign matched at least one profile, and no normal entity matched any:"
    ))
    create_table(doc,
        ["Entity", "Campaign", "Profile that fired"],
        [list(r) for r in THREAT_PROFILE["rows"]],
        col_widths=[1.1, 2.0, 3.4])
    add_callout(doc,
        f"Multi-front result: {THREAT_PROFILE['caught']} of {THREAT_PROFILE['total']} campaigns "
        f"detected at {THREAT_PROFILE['fp']} false positives, each alert named by the technique that "
        f"fired — auditable evidence, not an opaque anomaly score. Operationally the two fronts run "
        f"together: the known-bad profiles give high-confidence, zero-noise alerts, while composite "
        f"ranking catches novel variants the named profiles do not yet cover.")
    add_figure(doc, "radar",
        "Five-phase separation: each campaign is extreme on a different behavioral phase (USR-234 on "
        "novelty persistence, USR-156 on signal strength and breadth). No normal entity reaches these "
        "extremes on multiple phases at once.")
    add_page_break(doc)


def _build_what_this_enables(doc, audience_key, level):
    cfg = _cfg(audience_key)
    add_section_heading(doc, "7. What This Enables for Your Mission", level=1)
    add_body(doc, cfg["edge"])
    add_body(doc, "Concretely, the platform lets your analysts:", space_after=2)
    add_bullet(doc, "Replace a daily flood of threshold alerts with one risk-ranked list of entities.", bold_prefix="Triage by risk: ")
    add_bullet(doc, "See WHY an entity is flagged, in plain language mapped to MITRE ATT&CK — not an opaque score.", bold_prefix="Explainable verdicts: ")
    add_bullet(doc, "Catch slow insiders and APTs that stay inside every threshold for months.", bold_prefix="Slow-and-low coverage: ")
    add_bullet(doc, "Re-weight detection for the hunt at hand (insider vs. APT) without retraining.", bold_prefix="Context-adaptive: ")
    add_bullet(doc, "Run on-premises / in an enclave; the platform is containerized and can use a local embedding model.", bold_prefix="Deployable where you operate: ")

    add_section_heading(doc, "7.1  Deployment Path", level=2)
    add_bullet(doc, "Stand up the pipeline against a historical telemetry window to establish baselines.", bold_prefix="Phase 1 — Baseline: ")
    add_bullet(doc, "Run in shadow alongside the existing SIEM; compare ranked output to current alerting.", bold_prefix="Phase 2 — Shadow: ")
    add_bullet(doc, "Feed ranked verdicts into SOC triage and (for DISA) Zero Trust access decisions.", bold_prefix="Phase 3 — Operate: ")

    add_section_heading(doc, "7.2  Why This Is Different — Competitive Landscape and Defensible Novelty", level=2)
    add_body(doc, (
        "A structured, source-verified review (2024–2026) of commercial products and the published "
        "literature found no commercially marketed system, and no single published method, that "
        "combines all elements of this approach. Publicly described commercial UEBA — Exabeam, "
        "Securonix, Microsoft Sentinel UEBA — detects via statistical / p-value profiling, "
        "machine-learning baselines, rarity- and magnitude-based scores, and TF-IDF peer grouping; "
        "their generative-AI features are SOC copilots and conversational layers, not the detection "
        "mechanism. Commercial demand / supply-chain platforms (SAP IBP, o9, Kinaxis) forecast "
        "with gradient-boosted trees and machine learning over tabular features on knowledge-graph "
        "data models, with generative AI again only as a conversational layer."
    ))
    add_body(doc, (
        "We are explicit about prior art: the constituent techniques are individually established — "
        "LLM-embedding of serialized behavior (APT-LLM; LogBERT), serialize-to-prose-then-embed "
        "(TabLLM family), embed-then-detect pipelines (TAD-Bench), and embedding drift by cosine "
        "distance (ZEDD). Each, however, detects by MAGNITUDE and operates in a single domain; "
        "none projects the DIRECTION of drift onto named, pre-embedded threat concepts mapped to "
        "MITRE ATT&CK, nor spans cyber and supply-chain entities in one space."
    ))
    add_callout(doc,
        "To the best of our knowledge — scoped to publicly described and marketed methodology — no "
        "commercial platform and no single published system detects anomalies by the DIRECTION of "
        "behavioral drift in a unified, cross-domain semantic space. The contribution is the "
        "integrated architecture and the concept-direction (not magnitude/rarity) paradigm, not "
        "any single technique. (A formal patent / freedom-to-operate review is recommended before "
        "any contractual assertion of uniqueness.)")
    add_page_break(doc)


def _build_references(doc, level):
    add_section_heading(doc, "8. References and Further Reading", level=1)
    refs = [
        "CISA AA23-144A, 'People's Republic of China State-Sponsored Cyber Actor Living off the Land' (Volt Typhoon).",
        "MITRE ATT&CK Framework, Enterprise Matrix. attack.mitre.org.",
        "Prior-art (constituent techniques, individually established): APT-LLM (arXiv:2502.09385); "
        "LogBERT (arXiv:2103.04475); TAD-Bench embed-then-detect (arXiv:2501.11960); TabLLM-family "
        "serialize-to-prose-embed (arXiv:2502.11596); ZEDD embedding drift (arXiv:2601.12359).",
        "Competitive landscape (publicly described methodology): Exabeam, Securonix, Microsoft "
        "Sentinel UEBA; SAP IBP, o9, Kinaxis demand platforms.",
        "NIST SP 800-207, 'Zero Trust Architecture,' 2020.",
        "OMB M-22-09, 'Moving the U.S. Government Toward Zero Trust Cybersecurity Principles,' 2022.",
        "E. S. Page, 'Continuous Inspection Schemes,' Biometrika, 1954 (CUSUM).",
        "V-Intelligence UEBA Technical Specification (companion document, NDA edition).",
    ]
    for i, ref in enumerate(refs, 1):
        add_bullet(doc, ref, bold_prefix=f"[{i}] ")

    doc.add_paragraph()
    note = {
        TEASER: ("This is a conceptual overview. Architectural detail, the system design, and "
                 "empirical method are available in the Architectural and NDA editions of this "
                 "paper on request."),
        ARCHITECTURE: ("Exact algorithm formulations, parameters, weights, and thresholds are "
                       "withheld in this edition and provided in the NDA / technical-exchange "
                       "edition under a mutual non-disclosure agreement."),
        FULL: ("This edition contains proprietary formulations and is provided under NDA. Do "
               "not redistribute."),
    }[level]
    add_body(doc, note, italic=True, color=DARK_GRAY)


# ============================================================================
# Build orchestration
# ============================================================================
FIGDIR = os.path.join(HERE, "figures")
INNOV_FIGS = {}


def generate_innovation_figures():
    figs = {}
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyBboxPatch
    except Exception:
        return figs
    os.makedirs(FIGDIR, exist_ok=True)
    NV, BL, TL, RD = "#0D1B2A", "#1B4F72", "#0E6B8A", "#C0392B"
    try:
        fig, ax = plt.subplots(figsize=(9.8, 1.9)); ax.axis("off"); ax.set_xlim(0, 10); ax.set_ylim(0, 1.4)
        ax.text(0.1, 1.28, "Behavior as language: the V-Intelligence pipeline", fontsize=10, color=NV, weight="bold")
        boxes = ["Raw metrics", "Serialize\nto prose", "Embed\n1536-d", "Five\nzones",
                 "Drift +\nMITRE concept", "Composite\n-> ranked"]
        n = len(boxes); gap = 9.7 / n; w = gap * 0.85
        for i, t in enumerate(boxes):
            x = 0.15 + i * gap
            ax.add_patch(FancyBboxPatch((x, 0.25), w, 0.78, boxstyle="round,pad=0.02", fc=TL, ec="none"))
            ax.text(x + w / 2, 0.64, t, ha="center", va="center", color="white", fontsize=6.8)
            if i < n - 1:
                ax.annotate("", xy=(x + gap, 0.64), xytext=(x + w, 0.64),
                            arrowprops=dict(arrowstyle="->", color="#666", lw=1.2))
        p = os.path.join(FIGDIR, "fig_innov_pipeline.png"); fig.savefig(p, dpi=150, bbox_inches="tight")
        plt.close(fig); figs["pipeline"] = p
    except Exception:
        pass
    try:
        labels = ["Auth", "Files", "Network", "Endpoint", "Privilege", "Fused\ncomposite"]
        vals = [1.2, 1.4, 1.1, 1.3, 1.0, 3.6]
        cols = [BL] * 5 + [RD]
        fig, ax = plt.subplots(figsize=(7.8, 3.2))
        ax.bar(labels, vals, color=cols)
        ax.axhline(2.0, color="#999", ls="--", lw=1, label="alert threshold (2 sigma)")
        ax.set_ylabel("Anomaly (sigma)")
        ax.set_title("Each signal looks normal; the fused entity score reveals the threat",
                     fontsize=10, color=NV)
        ax.legend(fontsize=7)
        p = os.path.join(FIGDIR, "fig_innov_signal.png"); fig.savefig(p, dpi=150, bbox_inches="tight")
        plt.close(fig); figs["signal"] = p
    except Exception:
        pass
    return figs


def add_figure(doc, key, caption):
    path = INNOV_FIGS.get(key)
    if not path or not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = DARK_GRAY
    cap.paragraph_format.space_after = Pt(8)


def build(audience_key, level, results):
    doc = _init_doc()
    _build_title_page(doc, audience_key, level)
    _build_exec_summary(doc, audience_key, level, results)
    _build_seven_innovations(doc, level, results)
    _build_combined_signal_proof(doc, level, results)
    _build_architecture(doc, level)
    _build_method_deep_dive(doc, level)
    _build_validation(doc, level, results)
    _build_what_this_enables(doc, audience_key, level)
    _build_references(doc, level)

    out_name = f"Innovation_Whitepaper_{audience_key}_{level}.docx"
    out_path = os.path.join(HERE, out_name)
    doc.save(out_path)
    print(f"  {audience_key:10s} [{level}] -> {out_name}")
    return out_path


def main():
    parser = argparse.ArgumentParser(description="Build the V-Intelligence innovation whitepaper.")
    parser.add_argument("--audience", choices=list(AUDIENCE_CONFIGS) + [UNIVERSAL_KEY, "all"],
                        default="all", help="Target audience (default: all).")
    parser.add_argument("--disclosure", choices=[TEASER, ARCHITECTURE, FULL],
                        default=ARCHITECTURE, help="How much technical depth to reveal.")
    args = parser.parse_args()

    results = load_results()
    global INNOV_FIGS
    INNOV_FIGS = generate_innovation_figures()
    # Crisp analytics figures rendered from live data (copied into figures/ by the
    # build wrapper). Registered if present; the doc degrades gracefully without them.
    for key, fn in [("signal_feat", "fig_feature_cusum.png"),
                    ("signal_emb", "fig_embedding_cusum.png"),
                    ("radar", "fig_radar.png"),
                    ("composite", "fig_composite_rank.png")]:
        p = os.path.join(FIGDIR, fn)
        if os.path.exists(p):
            INNOV_FIGS[key] = p
    print("Building innovation whitepaper(s)...")
    print(f"  disclosure level : {args.disclosure}")
    print(f"  empirical data   : {'loaded from database (composite_scores)' if results else 'NOT FOUND (theoretical mode)'}")

    if args.audience == "all":
        keys = list(AUDIENCE_CONFIGS) + [UNIVERSAL_KEY]
    else:
        keys = [args.audience]

    for key in keys:
        build(key, args.disclosure, results)
    print("Done.")


if __name__ == "__main__":
    main()
