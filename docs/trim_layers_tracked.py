# -*- coding: utf-8 -*-
"""Tracked trim of cyber whitepaper sections 5,6,7 (Layers 3,4,5). Section 8 untouched.
Removes (as tracked deletions): 'The Threat at This Layer', 'Active Devices and Components',
'How It Works Today', 'Federal and Critical-Infrastructure Fit'. Trims 'Existing Applications
and Tools' to the first representative bullet. Sets Track Changes ON.
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC="WP DLA/Presentation/V-Intelligence_Layered_Cybersecurity_Whitepaper_v2.docx"
AUTHOR="V-Intelligence"; DATE="2026-06-21T13:00:00Z"
d=Document(SRC)
def accp(p): return "".join(t.text or "" for t in p._p.iter(qn('w:t')))
def style(p): return p.style.name if p.style else ""

_id=[70000]
def nid():
    _id[0]+=1; return str(_id[0])
def del_para_tracked(pel):
    runs=pel.findall(qn('w:r'))
    if runs:
        delw=OxmlElement('w:del'); delw.set(qn('w:id'),nid()); delw.set(qn('w:author'),AUTHOR); delw.set(qn('w:date'),DATE)
        runs[0].addprevious(delw)
        for r in runs:
            for t in r.findall(qn('w:t')):
                dt=OxmlElement('w:delText'); dt.set(qn('xml:space'),'preserve'); dt.text=t.text; r.replace(t,dt)
            delw.append(r)
    ppr=pel.find(qn('w:pPr'))
    if ppr is None: ppr=OxmlElement('w:pPr'); pel.insert(0,ppr)
    rpr=ppr.find(qn('w:rPr'))
    if rpr is None: rpr=OxmlElement('w:rPr'); ppr.append(rpr)
    dm=OxmlElement('w:del'); dm.set(qn('w:id'),nid()); dm.set(qn('w:author'),AUTHOR); dm.set(qn('w:date'),DATE); rpr.append(dm)

REMOVE={"The Threat at This Layer","Active Devices and Components","How It Works Today",
        "Federal and Critical-Infrastructure Fit"}
TARGET_SECS={"5","6","7"}

cur_sec=None; cur_h2=None
to_delete=[]; tools={}
for p in d.paragraphs:
    s=accp(p).strip(); st=style(p)
    if st.startswith("Heading 1") and s:
        cur_sec=s[0] if s[:1].isdigit() else None; cur_h2=None; continue
    if st.startswith("Heading 2") and s:
        cur_h2=s
        if cur_sec in TARGET_SECS and s in REMOVE:
            to_delete.append(p._p)   # delete the H2 heading itself
        continue
    if cur_sec in TARGET_SECS:
        if cur_h2 in REMOVE:
            to_delete.append(p._p)
        elif cur_h2=="Existing Applications and Tools" and st=="List Bullet":
            tools.setdefault(cur_sec,[]).append(p._p)

# trim tools: keep first bullet, delete the rest
trimmed=0
for sec,bl in tools.items():
    for pel in bl[1:]:
        to_delete.append(pel); trimmed+=1

for pel in to_delete:
    del_para_tracked(pel)

# Track Changes ON in settings (insert in a tolerated position)
stg=d.settings.element
if stg.find(qn('w:trackChanges')) is None:
    tc=OxmlElement('w:trackChanges')
    # insert after w:defaultTabStop if present else after first child
    ref=stg.find(qn('w:defaultTabStop')) or (list(stg)[0] if len(stg) else None)
    if ref is not None: ref.addnext(tc)
    else: stg.append(tc)

d.save(SRC)
print(f"tracked-deleted paragraphs: {len(to_delete)} (incl. {trimmed} trimmed tool bullets) | sections 5,6,7; section 8 untouched")
# verify
d2=Document(SRC)
print("tracked deletions now:", len(d2.element.body.findall('.//'+qn('w:del'))))
print("trackChanges setting on:", d2.settings.element.find(qn('w:trackChanges')) is not None)
