#!/usr/bin/env python3
"""Build Behavioral Entity Intelligence (BEI) Concept Paper.

A comprehensive framework document describing the universal concept of
behavioral entity profiling with temporal, relational, and network dimensions.
Covers 12+ domain applications from cybersecurity to supply chain to
counter-intelligence.

Output: docs/Behavioral_Entity_Intelligence_Framework.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
BLUE = RGBColor(0x1B, 0x4F, 0x72)
TEAL = RGBColor(0x0E, 0x6B, 0x8A)
GOLD = RGBColor(0xB7, 0x95, 0x0B)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x8A, 0x49)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)

OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
                           "Behavioral_Entity_Intelligence_Framework.docx")


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
    return h


def body(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def callout(doc, text, border_color="0E6B8A"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="12" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EAF4F7" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    run.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return p


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "0D1B2A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EAF0F6")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def page_break(doc):
    doc.add_page_break()


# ======================================================================
# DOCUMENT
# ======================================================================

def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # ── TITLE PAGE ──
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Behavioral Entity Intelligence")
    run.font.size = Pt(34)
    run.font.color.rgb = NAVY
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "A Universal Framework for Anomaly Detection Through\n"
        "Entity Profiles, Temporal Dynamics, Relationship Analysis,\n"
        "and Network Intelligence"
    )
    run.font.size = Pt(15)
    run.font.color.rgb = BLUE

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "22nd Century Technologies, Inc.\n\n"
        "June 2026\n\n"
        "Concept Paper — Version 1.0"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY

    page_break(doc)

    # ── TABLE OF CONTENTS ──
    heading(doc, "Table of Contents", level=1)
    toc = [
        "1. Abstract",
        "2. The Problem: Why Traditional Analytics Fail",
        "3. The Nine Dimensions of Behavioral Entity Intelligence",
        "   3.1. Identity — Who or What Is the Entity?",
        "   3.2. Behavior — What Does the Entity Do?",
        "   3.3. Time — Temporal Evolution and Drift",
        "   3.4. Space — Geographic and Topological Position",
        "   3.5. Relationship — Pairwise Entity Dynamics",
        "   3.6. Network — Graph-Level Ecosystem Intelligence",
        "   3.7. Sequence — Order of Operations",
        "   3.8. Periodicity — Rhythm and Cadence",
        "   3.9. Context — Environmental Meaning",
        "4. Three-Layer Architecture",
        "   4.1. Layer 1: Entity Behavioral Profiles",
        "   4.2. Layer 2: Relationship Behavioral Profiles",
        "   4.3. Layer 3: Network Intelligence",
        "5. Multi-Phase Composite Scoring",
        "6. Domain Applications",
        "   6.1. Cybersecurity — Insider Threat and APT Detection",
        "   6.2. Defense Supply Chain — DLA, SCRM, SBOM",
        "   6.3. Prescriptive Maintenance — Equipment Reliability",
        "   6.4. Healthcare Fraud — Medical Billing and Claims",
        "   6.5. Insurance Fraud — Unemployment, Property, Casualty",
        "   6.6. Clinical Trials — Patient Safety and Site Integrity",
        "   6.7. Counter-Intelligence — Adversary Activity Detection",
        "   6.8. Anti-Money Laundering — Financial Crime Detection",
        "   6.9. Counter-Terrorism — Network Disruption",
        "   6.10. Counter-Narcotics — Trafficking Pattern Detection",
        "   6.11. Energy and Critical Infrastructure",
        "   6.12. Workforce Analytics — Retention and Insider Risk",
        "7. Data Requirements by Domain",
        "8. Validated Results: Cybersecurity Proof of Concept",
        "9. Implementation Architecture",
        "10. Conclusion",
    ]
    for entry in toc:
        p = doc.add_paragraph()
        run = p.add_run(entry)
        run.font.size = Pt(10.5)
        run.font.color.rgb = NAVY if not entry.startswith("   ") else BLUE
        p.paragraph_format.space_after = Pt(2)
    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 1. ABSTRACT
    # ══════════════════════════════════════════════════════════════
    heading(doc, "1. Abstract", level=1)

    body(doc, (
        "Every organization monitors entities — users, suppliers, patients, equipment, "
        "accounts, personnel. Traditional analytics detect anomalies by asking: is any "
        "single metric abnormal? This approach systematically fails against sophisticated "
        "threats that operate within normal statistical ranges while gradually shifting "
        "behavioral direction."
    ))

    body(doc, (
        "This paper introduces Behavioral Entity Intelligence (BEI) — a universal "
        "analytical framework that defines any observable subject as a multi-dimensional "
        "behavioral entity, tracks its evolution across nine dimensions (identity, "
        "behavior, time, space, relationship, network, sequence, periodicity, and "
        "context), and detects anomalies through multi-phase composite scoring. The "
        "framework operates at three layers: individual entity profiles, pairwise "
        "relationship dynamics, and graph-level network intelligence."
    ))

    body(doc, (
        "The core insight is domain-agnostic: the entity changes WHAT it does, not "
        "HOW MUCH. The insider accesses the same number of files but shifts from public "
        "to restricted content. The supplier delivers the same volume but shifts component "
        "sourcing from domestic to adversary-nation origins. The patient submits the same "
        "number of claims but shifts procedure codes upward. Traditional threshold-based "
        "systems see identical volumes and report no anomaly. BEI detects the directional "
        "change because it measures behavioral meaning, not just magnitude."
    ))

    body(doc, (
        "Validated through a cybersecurity proof of concept — 250 users, 485 days, "
        "4.2 million daily events, 4 embedded attack campaigns — BEI's multi-front "
        "threat-profile detector achieves 4 of 4 attack detection at 0% false positive rate, "
        "where traditional methods detect 0 of 4 at operationally acceptable thresholds. "
        "Composite scoring alone cleanly separates 2 of 4; the threat-profile detector closes the gap."
    ))

    callout(doc,
        "The entity changes WHAT it does, not HOW MUCH. Traditional analytics measure "
        "magnitude. Behavioral Entity Intelligence measures direction."
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 2. THE PROBLEM
    # ══════════════════════════════════════════════════════════════
    heading(doc, "2. The Problem: Why Traditional Analytics Fail", level=1)

    body(doc, (
        "Across every domain — cybersecurity, supply chain, healthcare, finance, "
        "intelligence — the most damaging threats share a common characteristic: "
        "they operate within the bounds of authorized activity. The insider uses "
        "valid credentials. The fraudulent provider submits properly formatted claims. "
        "The money launderer keeps individual transactions below reporting thresholds. "
        "The compromised supplier delivers on time and on spec — while gradually "
        "substituting components."
    ))

    body(doc, "Traditional analytics fail for three structural reasons:", bold=True)

    body(doc, "1. Threshold Blindness", bold=True, space_after=2)
    body(doc, (
        "Rule-based systems define anomaly thresholds on individual metrics: "
        "transactions above $10,000, login attempts above 50, delivery delays "
        "above 5 days. Sophisticated actors stay below every individual threshold "
        "while achieving their objectives through directional change across "
        "multiple dimensions simultaneously."
    ))

    body(doc, "2. Magnitude-Only Measurement", bold=True, space_after=2)
    body(doc, (
        "Statistical anomaly detection (z-scores, isolation forests, clustering) "
        "measures HOW MUCH a metric deviates from the mean. It cannot distinguish "
        "between a benign spike (seasonal demand increase) and a malicious shift "
        "(demand diverted to an unauthorized recipient). Both produce the same "
        "statistical magnitude. The DIRECTION of change — what kind of deviation — "
        "is invisible to magnitude-only methods."
    ))

    body(doc, "3. Single-Entity Tunnel Vision", bold=True, space_after=2)
    body(doc, (
        "Traditional systems analyze entities in isolation. A supplier looks normal. "
        "A part looks normal. A location looks normal. But the RELATIONSHIP between "
        "them has changed — the supplier is now sourcing this specific part from a "
        "different country of origin. No single-entity analysis detects this. Only "
        "relationship-level behavioral intelligence reveals the shift."
    ))

    callout(doc,
        "The gap is structural, not parametric. Better thresholds, more features, "
        "and larger training sets do not solve it. The solution requires a fundamentally "
        "different analytical approach: measure behavioral direction, not just magnitude, "
        "across entity relationships over time."
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 3. THE NINE DIMENSIONS
    # ══════════════════════════════════════════════════════════════
    heading(doc, "3. The Nine Dimensions of Behavioral Entity Intelligence", level=1)

    body(doc, (
        "BEI characterizes every observable entity through nine complementary "
        "dimensions. The first six capture distinct aspects of entity state and "
        "interaction. The last three — sequence, periodicity, and context — provide "
        "analytical lenses that transform raw observations into actionable intelligence."
    ))

    table(doc,
        ["#", "Dimension", "Question Answered", "What It Captures"],
        [
            ["1", "Identity", "WHO or WHAT is the entity?",
             "Static/slow-changing attributes: role, type, classification, "
             "certification, clearance, geographic registration"],
            ["2", "Behavior", "WHAT does the entity DO?",
             "Actions decomposed into zones: access patterns, transaction types, "
             "operational activities, resource consumption"],
            ["3", "Time", "HOW does behavior evolve?",
             "Temporal baseline, drift velocity, acceleration, regime shifts, "
             "trend consistency across observation windows"],
            ["4", "Space", "WHERE does behavior occur?",
             "Geographic location, network topology, logical position, "
             "jurisdiction, facility, deployment zone"],
            ["5", "Relationship", "HOW does the entity connect to others?",
             "Pairwise behavioral signatures between entities. "
             "Drift in relationship even when individuals look stable"],
            ["6", "Network", "WHAT is the entity's ecosystem position?",
             "Graph centrality, community membership, influence propagation, "
             "multi-hop risk chains, cohort correlation"],
            ["7", "Sequence", "In what ORDER do events occur?",
             "Process flow compliance, kill chain progression, "
             "escalation patterns, stage skipping"],
            ["8", "Periodicity", "WHAT is the entity's rhythm?",
             "Regular cadences (daily, weekly, seasonal), "
             "deviation from established rhythm, new periodic patterns"],
            ["9", "Context", "WHAT does the behavior MEAN in this setting?",
             "Same behavior interpreted differently under different "
             "analytical contexts (investigation type, operational tempo, season)"],
        ],
        col_widths=[0.3, 0.9, 1.8, 3.5],
    )

    heading(doc, "3.1. Identity — Who or What Is the Entity?", level=2)
    body(doc, (
        "The identity dimension captures the relatively static attributes of an entity: "
        "a user's role, department, and clearance level; a supplier's geographic "
        "registration, certification status, and contract vehicle; a patient's "
        "demographics, insurance tier, and provider network; an equipment unit's "
        "model, age, and maintenance history. Identity changes infrequently — when "
        "it does change, the change itself is a significant signal (role change, "
        "supplier re-registration, provider network switch)."
    ))

    heading(doc, "3.2. Behavior — What Does the Entity Do?", level=2)
    body(doc, (
        "Behavior is decomposed into independent zones — distinct dimensions of "
        "activity that can drift independently. This decomposition is critical: "
        "a single composite behavioral score averages zone-specific signals, "
        "diluting the attack indicator. Zone decomposition preserves it."
    ))

    body(doc, "Examples of behavioral zone decomposition by domain:", bold=True)
    table(doc,
        ["Domain", "Entity", "Behavioral Zones"],
        [
            ["Cybersecurity", "User",
             "Identity zone, Access patterns, Data behavior, "
             "Network footprint, Risk posture"],
            ["Supply Chain", "Supplier",
             "Delivery performance, Pricing behavior, Quality metrics, "
             "Sourcing geography, Capacity utilization"],
            ["Healthcare", "Provider",
             "Procedure mix, Diagnosis coding, Referral network, "
             "Patient demographics, Charge patterns"],
            ["Finance", "Account",
             "Transaction geography, Merchant categories, Transfer patterns, "
             "Login behavior, Balance trajectory"],
            ["Maintenance", "Equipment",
             "Vibration profile, Thermal signature, Power consumption, "
             "Fluid chemistry, Error code patterns"],
            ["Intelligence", "Person of Interest",
             "Travel patterns, Communication frequency, Financial activity, "
             "Social connections, Access to sensitive information"],
        ],
        col_widths=[1.0, 1.0, 4.5],
    )

    heading(doc, "3.3. Time — Temporal Evolution and Drift", level=2)
    body(doc, (
        "The time dimension captures how behavior evolves through periodic snapshots "
        "(daily, weekly, monthly). Key temporal features include:"
    ))
    bullet(doc, "The magnitude and direction of behavioral change between consecutive periods",
           bold_prefix="Drift velocity: ")
    bullet(doc, "Whether the rate of change is itself increasing — escalation detection",
           bold_prefix="Acceleration: ")
    bullet(doc, "Sudden discontinuities in behavioral patterns (cosine similarity between "
           "consecutive snapshots drops below threshold)",
           bold_prefix="Regime shift: ")
    bullet(doc, "Whether anomalies persist across multiple periods or are transient noise",
           bold_prefix="Sustained deviation: ")
    bullet(doc, "Cumulative drift over long periods that individual-period thresholds miss",
           bold_prefix="Trend consistency: ")

    heading(doc, "3.4. Space — Geographic and Topological Position", level=2)
    body(doc, (
        "The spatial dimension captures where behavior occurs — geographic coordinates, "
        "network zones, facility identifiers, jurisdictions, or deployment locations. "
        "Spatial context transforms interpretation: a financial transaction from New York "
        "to London is routine; the same transaction routed through a sanctioned "
        "jurisdiction is a red flag. A login from the user's home office is normal; "
        "from a country they've never visited requires investigation."
    ))

    heading(doc, "3.5. Relationship — Pairwise Entity Dynamics", level=2)
    body(doc, (
        "The relationship dimension captures how pairs of entities behave TOGETHER. "
        "The mathematical foundation is the Hadamard product: the element-wise product "
        "of two entity behavioral vectors, producing a relationship-specific behavioral "
        "signature that captures interaction patterns invisible to individual entity analysis."
    ))

    body(doc, "The critical insight:", bold=True, space_after=2)
    body(doc, (
        "Entity A looks normal. Entity B looks normal. But the A×B relationship has "
        "drifted. In cybersecurity: a user's profile is stable and a device's profile "
        "is stable, but the user is now using a different device — the User×Device "
        "relationship changed. In supply chain: a supplier's performance is stable and "
        "an NSN's demand is stable, but the supplier is now sourcing THIS specific part "
        "from a different manufacturer — the Supplier×NSN×Manufacturer relationship "
        "reveals the risk."
    ))

    table(doc,
        ["Domain", "Relationship", "What It Captures", "Anomaly Signal"],
        [
            ["Cyber", "User × Device",
             "Which devices does this user normally use?",
             "User on unfamiliar device = credential theft"],
            ["Supply Chain", "Supplier × NSN",
             "How does this supplier perform for this specific part?",
             "Quality drift for one part while others stable = targeted substitution"],
            ["Supply Chain", "Supplier × Country",
             "Where does this supplier source components?",
             "Gradual shift to adversary-nation sourcing = SCRM risk"],
            ["Supply Chain", "NSN × SBOM",
             "What components make up this assembly?",
             "Bill of materials composition changes without notification"],
            ["Healthcare", "Provider × Diagnosis",
             "What diagnoses does this provider code?",
             "Diagnosis mix shifts toward higher-reimbursement codes = upcoding"],
            ["Finance", "Account × Geography",
             "Where does this account transact?",
             "Counterparty geography shifts toward sanctioned jurisdictions"],
            ["Intelligence", "Person × Person",
             "Who does this person communicate with?",
             "New persistent contact with known intelligence officer"],
            ["Maintenance", "Equipment × Component",
             "Which components fail on this unit?",
             "New failure mode appears = design flaw or counterfeit part"],
        ],
        col_widths=[0.8, 1.3, 2.2, 2.2],
    )

    heading(doc, "3.6. Network — Graph-Level Ecosystem Intelligence", level=2)
    body(doc, (
        "The network dimension extends pairwise relationships into a full graph. "
        "This enables multi-hop analysis that no individual or pairwise view can provide."
    ))

    bullet(doc, "A Tier-1 supplier passes all compliance checks. "
           "But trace the graph: their Tier-2 subcomponent supplier sources from a "
           "Tier-3 manufacturer with Chinese military connections. The individual entity "
           "is clean. The 3-hop supply chain path reveals the risk.",
           bold_prefix="Multi-hop risk propagation: ")
    bullet(doc, "Two 'independent' suppliers develop correlated delivery "
           "failures. Graph analysis reveals they share a common sub-tier component "
           "source — a single point of failure masked by apparent supplier diversity.",
           bold_prefix="Hidden common dependencies: ")
    bullet(doc, "When one entity's behavior changes, how does it affect "
           "connected entities? A supplier failure propagates through the network to "
           "affect mission readiness at specific locations. Graph-level analysis "
           "predicts the cascade before it materializes.",
           bold_prefix="Influence propagation: ")
    bullet(doc, "Are entities that should be independent actually "
           "behaving as a coordinated group? In healthcare: multiple 'independent' "
           "providers with correlated billing patterns = kickback ring. In finance: "
           "multiple accounts with synchronized transaction timing = structuring.",
           bold_prefix="Cohort correlation: ")

    heading(doc, "3.7. Sequence — Order of Operations", level=2)
    body(doc, (
        "The sequence dimension captures not just WHAT events occur but in what ORDER. "
        "Many anomalies are invisible to frequency-based analysis but obvious in "
        "sequential analysis:"
    ))
    bullet(doc, "Order → Ship → Inspect → Accept is normal. "
           "Order → Accept (skipping inspection) is a sequence anomaly — potential "
           "collusion between buyer and supplier to bypass quality controls.",
           bold_prefix="Supply chain: ")
    bullet(doc, "Reconnaissance → Initial Access → Execution → "
           "Persistence → Exfiltration follows the MITRE ATT&CK kill chain. "
           "Detecting progression through stages is more actionable than detecting "
           "individual events.",
           bold_prefix="Cybersecurity: ")
    bullet(doc, "Diagnosis → Treatment → Follow-up is normal. "
           "Treatment → Diagnosis (retroactive coding to justify treatment already "
           "delivered) is a billing fraud sequence signature.",
           bold_prefix="Healthcare: ")
    bullet(doc, "Placement (cash entry) → Layering "
           "(transfer chain) → Integration (legitimate-appearing investment) follows "
           "a known stage progression. Each individual transaction looks legitimate; "
           "the sequence reveals the scheme.",
           bold_prefix="Money laundering: ")

    heading(doc, "3.8. Periodicity — Rhythm and Cadence", level=2)
    body(doc, (
        "Entities exhibit regular rhythms — daily work patterns, weekly delivery "
        "schedules, monthly billing cycles, seasonal demand curves. Deviation from "
        "established periodicity is a powerful anomaly signal:"
    ))
    bullet(doc, "C2 beacon malware contacts its server every 6 hours "
           "with clockwork regularity. Normal DNS queries are aperiodic. "
           "The RHYTHM is the detection signal, not the volume.",
           bold_prefix="C2 beaconing: ")
    bullet(doc, "An equipment unit's vibration frequency shifts from "
           "its established oscillation pattern — bearing degradation changes the "
           "mechanical rhythm before any absolute threshold is exceeded.",
           bold_prefix="Equipment degradation: ")
    bullet(doc, "A drug trafficking network ships on a 3-week cycle "
           "aligned with production harvests. Detecting the cadence and predicting "
           "the next shipment window enables interdiction planning.",
           bold_prefix="Trafficking patterns: ")

    heading(doc, "3.9. Context — Environmental Meaning", level=2)
    body(doc, (
        "The same behavior carries different meaning in different contexts. "
        "Context-adaptive analysis applies multiple analytical lenses to the same "
        "behavioral data and identifies entities that score anomalous under "
        "specific threat scenarios:"
    ))
    bullet(doc, "Under an 'insider investigation' lens, data access "
           "patterns are weighted heavily. Under an 'APT hunt' lens, network "
           "footprint dominates. A user who scores high under the insider lens but "
           "normal under APT hunt is likely an insider, not an external compromise.",
           bold_prefix="Multi-context scoring: ")
    bullet(doc, "Demand spike during deployment preparation is expected; "
           "during peacetime drawdown is anomalous. The same volume change has "
           "opposite interpretations depending on operational context.",
           bold_prefix="Operational tempo: ")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 4. THREE-LAYER ARCHITECTURE
    # ══════════════════════════════════════════════════════════════
    heading(doc, "4. Three-Layer Architecture", level=1)

    body(doc, (
        "BEI organizes analysis into three layers of increasing sophistication. "
        "Each layer builds on the one below it, and anomalies can be detected at "
        "any layer independently."
    ))

    heading(doc, "4.1. Layer 1: Entity Behavioral Profiles", level=2)
    body(doc, (
        "Each entity is represented as a multi-zone behavioral profile captured at "
        "regular intervals. For each observation period, the entity's behavioral "
        "state is serialized into a structured description and converted into a "
        "high-dimensional representation that captures the MEANING of the behavior, "
        "not just the magnitude of individual metrics. Drift between consecutive "
        "periods is measured as directional change in this semantic space."
    ))

    heading(doc, "4.2. Layer 2: Relationship Behavioral Profiles", level=2)
    body(doc, (
        "For each meaningful pair of entities, a relationship profile is computed "
        "as the interaction-specific behavioral signature. This captures HOW two "
        "entities behave together — patterns that are invisible when analyzing "
        "either entity in isolation. Relationship drift is tracked over time "
        "independently of individual entity drift."
    ))

    body(doc, (
        "A critical design principle: the relationship profile is not simply the "
        "concatenation or average of two entity profiles. It is a distinct "
        "representation that captures the interaction pattern itself. Two entities "
        "with individually stable profiles can have a dramatically drifting "
        "relationship — this is the primary detection signal for supply chain "
        "substitution, credential sharing, and referral fraud."
    ))

    heading(doc, "4.3. Layer 3: Network Intelligence", level=2)
    body(doc, (
        "The network layer analyzes the full graph of entity relationships, enabling "
        "multi-hop risk propagation, hidden dependency discovery, influence cascade "
        "prediction, and coordinated behavior detection. Network-level anomalies are "
        "often the highest-value findings because they reveal systemic risks invisible "
        "to entity-level or relationship-level analysis."
    ))

    table(doc,
        ["Layer", "Analyzes", "Detection Signal", "Example"],
        [
            ["1. Entity", "Individual behavioral profiles",
             "Single entity drifts from its baseline",
             "User's data access pattern shifts from public to restricted files"],
            ["2. Relationship", "Pairwise interaction dynamics",
             "Relationship drifts while individuals look stable",
             "Supplier's quality for ONE specific part degrades while others stable"],
            ["3. Network", "Graph-level ecosystem patterns",
             "Multi-hop risk chains, coordinated behavior",
             "Tier-3 manufacturer in adversary nation connected through 2 intermediaries"],
        ],
        col_widths=[1.0, 1.5, 2.0, 2.0],
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 5. MULTI-PHASE COMPOSITE SCORING
    # ══════════════════════════════════════════════════════════════
    heading(doc, "5. Multi-Phase Composite Scoring", level=1)

    body(doc, (
        "At every layer, anomaly detection uses multi-phase composite scoring — "
        "five independent detection phases fused into a single ranked score. "
        "A normal entity might score high on one phase by random variation; "
        "a genuinely anomalous entity scores high on multiple phases simultaneously."
    ))

    table(doc,
        ["Phase", "What It Measures", "Why It Matters"],
        [
            ["Signal Strength",
             "Magnitude of the strongest anomaly signals "
             "relative to peer group",
             "Catches entities with extreme deviations on their most anomalous dimensions"],
            ["Breadth",
             "Count of dimensions exceeding a significance threshold",
             "Distinguishes single-dimension noise from multi-dimension anomaly patterns"],
            ["Sustained Deviation",
             "Persistence of anomalies across multiple observation periods",
             "Separates transient events from persistent campaigns or degradation"],
            ["Context Divergence",
             "Consistency of anomaly across multiple analytical contexts",
             "Anomalous entities flag under threat-specific contexts; normal entities do not"],
            ["Novelty Persistence",
             "Recurrence of previously-unseen behaviors across periods",
             "Catches new patterns (new contacts, new routes, new components) that persist"],
        ],
        col_widths=[1.2, 2.8, 2.5],
    )

    callout(doc,
        "Each phase is independently insufficient. Combined, they create a detection "
        "surface that no single evasion technique can defeat. An actor who suppresses "
        "signal strength still fails on breadth; one who limits breadth still fails on "
        "sustained deviation or novelty persistence."
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 6. DOMAIN APPLICATIONS
    # ══════════════════════════════════════════════════════════════
    heading(doc, "6. Domain Applications", level=1)

    body(doc, (
        "The following sections demonstrate how the BEI framework applies across "
        "twelve distinct mission domains. For each domain, we define the primary "
        "entities, behavioral zones, key relationships, and the specific anomaly "
        "patterns that BEI detects where traditional analytics fail."
    ))

    # ── 6.1 Cybersecurity ──
    heading(doc, "6.1. Cybersecurity — Insider Threat and APT Detection", level=2)
    table(doc,
        ["Component", "Definition"],
        [
            ["Primary Entities", "User, Device, Application, Network Segment"],
            ["Behavioral Zones", "Identity, Access patterns, Data behavior, Network footprint, Risk posture"],
            ["Key Relationships", "User×Device, User×Application, Device×Segment"],
            ["Time Granularity", "Weekly behavioral snapshots"],
            ["Data Sources", "Authentication logs, file access, network flow, DNS, endpoint telemetry"],
            ["Anomaly Detected", "Insider shifting from public to restricted file access at constant volume;\n"
             "APT C2 beacon: persistent novel IP appearing every week for 60 weeks;\n"
             "LOTL: legitimate admin tools used in anomalous temporal/contextual patterns"],
            ["Traditional Gap", "All 4 attack users remain within normal statistical ranges on all 23 individual features"],
            ["BEI Result", "4/4 detected at 0% FP via the multi-front threat-profile detector "
             "(composite scoring alone separates 2/4)"],
        ],
        col_widths=[1.5, 5.0],
    )

    # ── 6.2 Supply Chain ──
    heading(doc, "6.2. Defense Supply Chain — DLA, SCRM, SBOM", level=2)
    table(doc,
        ["Component", "Definition"],
        [
            ["Primary Entities", "Supplier, NSN (Part), Location (Depot/Base), Manufacturer"],
            ["Behavioral Zones", "Delivery performance, Pricing trends, Quality metrics,\n"
             "Sourcing geography, Capacity utilization, Contract compliance"],
            ["Key Relationships", "Supplier×NSN (part-specific performance),\n"
             "Supplier×Country (sourcing origin tracking),\n"
             "NSN×Location (demand patterns per site),\n"
             "NSN×SBOM (bill of materials composition),\n"
             "Manufacturer×Supplier (who makes vs who distributes)"],
            ["Transactional Entity", "Demand signal for NSN X at Location Y — itself an entity with\n"
             "behavioral profile: requestor mix, urgency code distribution,\n"
             "quantity patterns, seasonal profile, lead time variation"],
            ["Time Granularity", "Weekly or monthly snapshots"],
            ["Data Sources", "Procurement transactions, delivery records, quality inspections,\n"
             "GIDEP alerts, country-of-origin declarations, SBOM registries"],
            ["Anomaly Detected", "Supplier gradually shifts component sourcing from domestic to\n"
             "adversary-nation origins while maintaining delivery and price metrics;\n"
             "NSN bill of materials composition drifts without notification;\n"
             "Demand pattern changes WHO is requesting at WHAT urgency (diversion);\n"
             "Two 'independent' suppliers develop correlated failures (shared sub-tier)"],
            ["Traditional Gap", "Individual supplier metrics (on-time %, quality score) all pass thresholds"],
            ["BEI Result", "Relationship-level drift (Supplier×Country, NSN×SBOM) reveals risk invisible\n"
             "to entity-level monitoring"],
        ],
        col_widths=[1.5, 5.0],
    )

    # ── 6.3 Maintenance ──
    heading(doc, "6.3. Prescriptive Maintenance — Equipment Reliability", level=2)
    table(doc,
        ["Component", "Definition"],
        [
            ["Primary Entities", "Equipment unit (specific serial number), Component, Facility"],
            ["Behavioral Zones", "Vibration profile, Thermal signature, Power consumption,\n"
             "Fluid chemistry, Error code patterns, Operational duty cycle"],
            ["Key Relationships", "Equipment×Component (which components fail on this unit?),\n"
             "Equipment×Facility (does environment affect degradation?),\n"
             "Component×Supplier (does supplier origin predict failure rate?)"],
            ["Time Granularity", "Daily or weekly sensor aggregates"],
            ["Data Sources", "IoT sensors, maintenance logs, CMMS records, fluid analysis,\n"
             "operational tempo data, environmental conditions"],
            ["Anomaly Detected", "RPM-to-temperature-to-fuel-flow RELATIONSHIP shifts 8% over 6 weeks —\n"
             "compressor degradation invisible to individual sensor thresholds;\n"
             "New vibration frequency appears at specific RPM range — bearing wear;\n"
             "Zone divergence: thermal zone drifting while vibration stable = fluid issue"],
            ["Mission Impact", "Equipment behavioral drift predicts WHAT parts will be needed WHEN —\n"
             "feeding demand forecasting and supply chain pre-positioning;\n"
             "Prevents mission-degrading failures through behavioral early warning"],
            ["Traditional Gap", "Every individual sensor reading remains within spec; "
             "the PATTERN change is invisible to single-threshold monitoring"],
        ],
        col_widths=[1.5, 5.0],
    )

    # ── 6.4 Healthcare ──
    heading(doc, "6.4. Healthcare Fraud — Medical Billing and Claims", level=2)
    table(doc,
        ["Component", "Definition"],
        [
            ["Primary Entities", "Provider (physician/facility), Patient/Beneficiary, Procedure, Diagnosis"],
            ["Behavioral Zones", "Procedure mix, Diagnosis coding distribution, Referral network,\n"
             "Patient demographics, Charge amounts, Visit frequency"],
            ["Key Relationships", "Provider×Diagnosis (coding pattern per provider),\n"
             "Provider×Provider (referral network dynamics),\n"
             "Patient×Provider (beneficiary loyalty patterns),\n"
             "Procedure×Diagnosis (clinical appropriateness)"],
            ["Time Granularity", "Monthly claim aggregates"],
            ["Data Sources", "CMS claims data, EHR records, provider enrollment, NPI registry"],
            ["Anomaly Detected", "Provider gradually upcodes from 99213→99214→99215 over 18 months —\n"
             "each individual claim is valid, the DRIFT is the fraud signal;\n"
             "Provider×Provider referral relationship suddenly intensifies = kickback;\n"
             "Patient cluster with suspiciously similar demographics at one provider = phantom billing"],
            ["Traditional Gap", "Each claim passes individual edit checks; fraud is in the behavioral trend"],
        ],
        col_widths=[1.5, 5.0],
    )

    # ── 6.5 Insurance ──
    heading(doc, "6.5. Insurance Fraud — Unemployment, Property, Casualty", level=2)
    table(doc,
        ["Component", "Definition"],
        [
            ["Primary Entities", "Claimant, Employer, Repair vendor, Adjuster"],
            ["Behavioral Zones", "Claim frequency, Loss type distribution, Documentation timing,\n"
             "Vendor selection, Geographic patterns, Certification compliance"],
            ["Key Relationships", "Claimant×Vendor (do specific claimants always use the same repair shop?),\n"
             "Claimant×Claimant (do claimants in a ring have correlated claim timing?),\n"
             "Adjuster×Vendor (does an adjuster consistently route to one vendor?)"],
            ["Anomaly Detected", "Unemployment claimant maintains perfect certification compliance (volume normal)\n"
             "but job search quality degrades — directional change invisible to thresholds;\n"
             "Claimant's vendor selection drifts toward a specific shop network = staged accidents;\n"
             "Adjuster×Vendor relationship intensifies = collusion"],
        ],
        col_widths=[1.5, 5.0],
    )

    # ── 6.6 Clinical Trials ──
    heading(doc, "6.6. Clinical Trials — Patient Safety and Site Integrity", level=2)
    table(doc,
        ["Component", "Definition"],
        [
            ["Primary Entities", "Patient, Clinical site, Investigator, Study drug"],
            ["Behavioral Zones", "Vital sign trajectories, Adverse event frequency, Lab results,\n"
             "Medication compliance, Visit timing, Protocol deviations"],
            ["Key Relationships", "Site×Patient cohort (are patients at one site suspiciously similar?),\n"
             "Investigator×Protocol (does this investigator deviate from protocol?),\n"
             "Patient×Drug response (does response profile match expected pharmacology?)"],
            ["Anomaly Detected", "Clinical site's enrollment pattern drifts — patients suspiciously\n"
             "similar in demographics and response = potential data fabrication;\n"
             "Patient response trajectory diverges from cohort baseline = early safety signal;\n"
             "Site×Patient relationship shows impossibly uniform visit timing = protocol fraud"],
        ],
        col_widths=[1.5, 5.0],
    )

    # ── 6.7 Counter-Intelligence ──
    heading(doc, "6.7. Counter-Intelligence — Adversary Activity Detection", level=2)
    table(doc,
        ["Component", "Definition"],
        [
            ["Primary Entities", "Person of interest, Diplomatic personnel, Intelligence officer,\n"
             "Facility, Communication endpoint"],
            ["Behavioral Zones", "Travel patterns, Communication frequency/timing, Financial transactions,\n"
             "Social contact network, Access to classified information, Cover activities"],
            ["Key Relationships", "Person×Person (who meets whom, communication patterns),\n"
             "Person×Location (dead drop sites, meeting points, travel overlap),\n"
             "Person×Organization (affiliation changes, cover employment),\n"
             "Person×Financial entity (unusual funding sources)"],
            ["Network Analysis", "Handler → Agent → Sub-agent chains;\n"
             "Communication network topology changes (new nodes, dropped contacts);\n"
             "Travel pattern correlation between seemingly unconnected individuals"],
            ["Sequence Analysis", "Recruitment → Assessment → Development → Handling follows\n"
             "known intelligence tradecraft progression;\n"
             "Travel → Meeting → Communication → Financial activity sequence"],
            ["Anomaly Detected", "Diplomat's travel pattern begins overlapping with known intelligence\n"
             "officer's movements — coordination signal invisible to individual monitoring;\n"
             "Person×Financial relationship shows new persistent funding source;\n"
             "Communication periodicity matches known covert signaling cadences"],
        ],
        col_widths=[1.5, 5.0],
    )

    # ── 6.8 AML ──
    heading(doc, "6.8. Anti-Money Laundering — Financial Crime Detection", level=2)
    table(doc,
        ["Component", "Definition"],
        [
            ["Primary Entities", "Account, Beneficial owner, Shell company, Correspondent bank, Jurisdiction"],
            ["Behavioral Zones", "Transaction velocity, Counterparty geography, Amount distribution,\n"
             "Timing patterns, Currency mix, Account lifecycle stage"],
            ["Key Relationships", "Account×Account (transfer chains, round-tripping),\n"
             "Company×Company (ownership layers, shared directors),\n"
             "Account×Geography (jurisdiction risk scoring),\n"
             "Beneficial owner×Shell company (control networks)"],
            ["Sequence Analysis", "Placement (cash entry) → Layering (transfer chain) →\n"
             "Integration (legitimate-appearing investment);\n"
             "Each stage has distinct behavioral signatures"],
            ["Anomaly Detected", "Account maintains normal total volume but gradually shifts counterparty\n"
             "geography toward sanctioned jurisdictions over 6 months — each individual\n"
             "transaction is below reporting thresholds;\n"
             "Company×Company ownership network reveals circular structures;\n"
             "Account×Timing periodicity matches structuring patterns"],
        ],
        col_widths=[1.5, 5.0],
    )

    # ── 6.9 CT ──
    heading(doc, "6.9. Counter-Terrorism — Network Disruption", level=2)
    table(doc,
        ["Component", "Definition"],
        [
            ["Primary Entities", "Suspect, Cell member, Facilitator, Financial supporter, Location"],
            ["Behavioral Zones", "Travel patterns, Communication activity, Financial flows,\n"
             "Procurement patterns, Online activity, Social network changes"],
            ["Key Relationships", "Person×Person (communication network, meeting patterns),\n"
             "Person×Location (training sites, reconnaissance targets),\n"
             "Person×Financial entity (funding chains, hawala networks),\n"
             "Cell×Cell (inter-cell coordination signals)"],
            ["Sequence Analysis", "Radicalization → Planning → Preparation → Execution stages;\n"
             "Procurement of specific materials in specific sequence;\n"
             "Communication pattern changes as operational date approaches"],
            ["Anomaly Detected", "Cell member's communication network suddenly contracts (operational security\n"
             "tightening before action); Financial supporter's transfer pattern shifts to\n"
             "new recipients in target geography; Procurement sequence matches known\n"
             "attack preparation patterns"],
        ],
        col_widths=[1.5, 5.0],
    )

    # ── 6.10 CN ──
    heading(doc, "6.10. Counter-Narcotics — Trafficking Pattern Detection", level=2)
    table(doc,
        ["Component", "Definition"],
        [
            ["Primary Entities", "Smuggler, Courier, Distributor, Financial facilitator,\n"
             "Transport vehicle, Route corridor"],
            ["Behavioral Zones", "Travel routes, Shipment timing, Financial flows,\n"
             "Communication patterns, Vehicle usage, Storage facility activity"],
            ["Key Relationships", "Person×Route (trafficking corridor usage),\n"
             "Person×Person (distribution network),\n"
             "Vehicle×Route×Timing (transport patterns),\n"
             "Financial entity×Geography (proceeds laundering paths)"],
            ["Periodicity Analysis", "Trafficking follows harvest/production cycles (3-week, seasonal);\n"
             "Detecting the cadence enables interdiction timing prediction;\n"
             "Rhythm disruption indicates law enforcement pressure or route change"],
            ["Anomaly Detected", "New travel route appears that correlates with known trafficking corridor\n"
             "timing; Courier×Route relationship suddenly changes (route burned);\n"
             "Financial flow periodicity shifts (new laundering mechanism);\n"
             "Distribution network topology changes (new nodes = expansion)"],
        ],
        col_widths=[1.5, 5.0],
    )

    # ── 6.11 Energy ──
    heading(doc, "6.11. Energy and Critical Infrastructure", level=2)
    table(doc,
        ["Component", "Definition"],
        [
            ["Primary Entities", "Substation, Smart meter, Generation facility, SCADA controller,\n"
             "Grid operator"],
            ["Behavioral Zones", "Load profile shape, Voltage stability, Harmonic distortion,\n"
             "Consumption curve, Communication patterns (for SCADA/ICS)"],
            ["Key Relationships", "Meter×Grid segment (consumption vs generation balance),\n"
             "Controller×Protocol (SCADA command patterns),\n"
             "Facility×Weather (expected vs actual load correlation)"],
            ["Anomaly Detected", "Smart meter's consumption curve shape changes (same total kWh, different\n"
             "time-of-use profile) = energy theft or equipment degradation;\n"
             "SCADA controller's communication pattern subtly shifts = potential cyber attack\n"
             "(Stuxnet-style behavioral modification);\n"
             "Substation load-to-weather correlation drifts = equipment degradation"],
        ],
        col_widths=[1.5, 5.0],
    )

    # ── 6.12 Workforce ──
    heading(doc, "6.12. Workforce Analytics — Retention and Insider Risk", level=2)
    table(doc,
        ["Component", "Definition"],
        [
            ["Primary Entities", "Employee, Team, Project, Manager"],
            ["Behavioral Zones", "Productivity output, Collaboration network breadth,\n"
             "Tool/system usage, Schedule patterns, Communication frequency"],
            ["Key Relationships", "Employee×Team (engagement with team activities),\n"
             "Employee×Project (contribution patterns),\n"
             "Employee×Manager (communication frequency and sentiment trajectory)"],
            ["Anomaly Detected", "Employee's collaboration network shrinks 40% over 3 months while\n"
             "productivity stays flat = flight risk or disengagement;\n"
             "Employee×Manager communication frequency drops = management relationship issue;\n"
             "Employee's system access pattern shifts to include HR and job posting systems"],
        ],
        col_widths=[1.5, 5.0],
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 7. DATA REQUIREMENTS
    # ══════════════════════════════════════════════════════════════
    heading(doc, "7. Data Requirements by Domain", level=1)

    body(doc, (
        "Each domain requires specific data sources to populate the nine BEI dimensions. "
        "The framework is designed to work with existing organizational data — no new "
        "instrumentation is required in most cases."
    ))

    table(doc,
        ["Domain", "Minimum Data Sources", "Observation\nCadence", "Baseline\nPeriod"],
        [
            ["Cybersecurity", "Auth logs, file access, network flow,\nDNS, endpoint telemetry",
             "Weekly", "6-8 weeks"],
            ["Supply Chain", "Procurement transactions, delivery records,\nquality inspections, COO declarations",
             "Weekly/\nMonthly", "6 months"],
            ["Maintenance", "IoT sensors, maintenance logs, CMMS,\nfluid analysis, operational data",
             "Daily/\nWeekly", "3-6 months"],
            ["Healthcare", "Claims data (CMS/private), EHR records,\nNPI registry, referral data",
             "Monthly", "12 months"],
            ["Insurance", "Claims submissions, adjuster records,\nvendor invoices, certification logs",
             "Monthly", "12 months"],
            ["Clinical Trials", "EDC data, lab results, AE reports,\nvisit logs, drug accountability",
             "Per-visit", "Run-in\nperiod"],
            ["Counter-Intel", "Travel records, SIGINT, FININT,\nopen source, access logs",
             "Daily/\nWeekly", "6-12\nmonths"],
            ["AML", "Transaction records, KYC/CDD data,\nownership registries, SWIFT messages",
             "Daily", "12 months"],
            ["Counter-Terror", "Travel, SIGINT, FININT, HUMINT,\nprocurement records, social media",
             "Daily", "Ongoing"],
            ["Counter-Narcotics", "Travel, shipment, financial,\ncommunication, surveillance data",
             "Daily/\nWeekly", "6 months"],
            ["Energy", "SCADA telemetry, smart meter data,\nweather, grid operations logs",
             "Hourly/\nDaily", "3-6 months"],
            ["Workforce", "HR systems, collaboration tools,\ncode repos, calendar, email metadata",
             "Weekly", "3-6 months"],
        ],
        col_widths=[1.0, 2.5, 0.8, 0.8],
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 8. VALIDATED RESULTS
    # ══════════════════════════════════════════════════════════════
    heading(doc, "8. Validated Results: Cybersecurity Proof of Concept", level=1)

    body(doc, (
        "The BEI framework has been validated through a comprehensive cybersecurity "
        "proof of concept — the V-Intelligence UEBA program. This validation "
        "demonstrates that the framework produces measurable, superior results "
        "compared to traditional analytical approaches."
    ))

    table(doc,
        ["Parameter", "Value"],
        [
            ["Simulation period", "485 days (January 2025 — May 2026)"],
            ["User population", "250 users (employees, contractors, service accounts)"],
            ["Daily event volume", "4.2 million events across 5 log sources"],
            ["Attack campaigns", "4 long-duration (100-240 days each)"],
            ["Attack types", "Insider threat, Slow APT (C2), Volt Typhoon LOTL, Salt Typhoon telecom"],
            ["Traditional detection", "0 of 4 at operationally acceptable FP rates"],
            ["BEI composite scoring", "2 of 4 cleanly separated above the normal population"],
            ["BEI threat-profile detector", "4 of 4 detected at 0% false positive rate"],
            ["Detection method", "Multi-front threat-profile matching: C2-beacon, DGA-DNS,\n"
             "LOTL-process, cohort-rare access, recon-fanout, insider-collection"],
        ],
        col_widths=[1.8, 4.8],
    )

    body(doc, "Per-attacker results:", bold=True)
    table(doc,
        ["Attack User", "Type", "Duration", "Composite\nRank (/250)", "Caught By", "Primary Detection Signal"],
        [
            ["USR-118", "Salt Typhoon\n(Telecom)", "100 days", "#1 (above norm)", "Composite",
             "Signal strength + breadth across\nmultiple behavioral zones"],
            ["USR-156", "Insider\nThreat", "8 months", "#2 (above norm)", "Composite",
             "Sustained deviation over 8-month\nescalation; data behavior zone drift"],
            ["USR-234", "Slow APT\n(C2 Beacon)", "180 days", "below norm", "Threat-profile",
             "C2-beacon front: novel IP recurs with\nbeacon cadence in post-baseline weeks"],
            ["USR-042", "Volt Typhoon\n(LOTL)", "115 days", "below norm", "Threat-profile",
             "LOTL-process front: legitimate tools in\nan attacker-shaped usage pattern"],
        ],
        col_widths=[0.7, 0.8, 0.7, 0.7, 0.5, 3.1],
    )

    callout(doc,
        "Traditional anomaly detection (Isolation Forest, SVM, LOF, Z-Score) scores "
        "all four attack users as NORMAL — no individual feature exceeds statistical "
        "thresholds. BEI's composite scoring lifts 2 of 4 (USR-156, USR-118) clearly above the "
        "normal population; the multi-front threat-profile detector then catches all four at 0% FP "
        "by matching attacker-shaped behavior invisible to magnitude-only methods."
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 9. IMPLEMENTATION ARCHITECTURE
    # ══════════════════════════════════════════════════════════════
    heading(doc, "9. Implementation Architecture", level=1)

    body(doc, (
        "BEI is implemented as a containerized, database-backed platform deployable "
        "on standard infrastructure. The architecture is domain-agnostic — the same "
        "platform supports cybersecurity, supply chain, healthcare, and other domains "
        "by configuring entity definitions, behavioral zones, and relationship mappings."
    ))

    body(doc, "Core components:", bold=True)
    bullet(doc, "PostgreSQL with pgvector extension for "
           "high-dimensional behavioral vector storage and similarity search",
           bold_prefix="Vector database: ")
    bullet(doc, "RESTful API layer for data ingestion, "
           "entity profile management, and anomaly query",
           bold_prefix="API layer: ")
    bullet(doc, "Automated pipeline for periodic behavioral "
           "snapshot computation, drift analysis, and composite scoring",
           bold_prefix="Processing pipeline: ")
    bullet(doc, "Interactive dashboard for entity exploration, "
           "anomaly investigation, and drill-down analysis",
           bold_prefix="Dashboard: ")
    bullet(doc, "Containerized deployment (Docker) for "
           "enterprise data centers, tactical edge, and cloud environments",
           bold_prefix="Deployment: ")

    body(doc, "Deployment phases:", bold=True)
    table(doc,
        ["Phase", "Duration", "Activities"],
        [
            ["1. Configuration", "1-2 weeks",
             "Define entities, behavioral zones, relationships,\n"
             "and data source mappings for the target domain"],
            ["2. Data Integration", "2-4 weeks",
             "Connect to organizational data sources;\n"
             "validate feature engineering pipeline"],
            ["3. Baseline", "4-6 weeks",
             "Build behavioral baselines for all entities;\n"
             "establish normal behavioral ranges"],
            ["4. Detection", "2-4 weeks",
             "Deploy composite scoring;\n"
             "tune thresholds on historical data; validate FP rate"],
            ["5. Production", "Ongoing",
             "Continuous behavioral monitoring;\n"
             "integration with existing investigation workflows"],
        ],
        col_widths=[1.0, 0.8, 4.7],
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 10. CONCLUSION
    # ══════════════════════════════════════════════════════════════
    heading(doc, "10. Conclusion", level=1)

    body(doc, (
        "Behavioral Entity Intelligence represents a fundamental advance in anomaly "
        "detection. By defining any observable subject as a multi-dimensional behavioral "
        "entity — tracked across nine dimensions, analyzed at three layers, and scored "
        "through five composite phases — BEI detects threats that are structurally "
        "invisible to traditional threshold-based and magnitude-only analytics."
    ))

    body(doc, (
        "The framework is validated through a cybersecurity proof of concept where the "
        "multi-front threat-profile detector achieves 4 of 4 attack detection at 0% false "
        "positive rate (composite scoring alone separates 2 of 4) while traditional methods "
        "detect 0 of 4. More importantly, the framework is domain-agnostic — the same "
        "analytical architecture that detects insider threats in cybersecurity detects "
        "supplier substitution fraud in supply chains, upcoding in healthcare billing, "
        "C2 beaconing in APT campaigns, and degradation patterns in equipment maintenance."
    ))

    body(doc, (
        "The core insight is universal: sophisticated threats change WHAT they do, not "
        "HOW MUCH. They stay within normal volume ranges while shifting behavioral "
        "direction. Traditional systems measure magnitude. BEI measures direction. "
        "This directional intelligence — applied across entity profiles, relationship "
        "dynamics, and network graphs — is the analytical foundation required to "
        "detect the threats that matter most."
    ))

    callout(doc,
        "Behavioral Entity Intelligence: Nine dimensions. Three layers. Five scoring "
        "phases. One universal framework — from cybersecurity to supply chain to "
        "healthcare to counter-intelligence. The entity changes WHAT it does, not "
        "HOW MUCH. BEI sees the difference."
    )

    # ── Save ──
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"BEI Concept Paper: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
