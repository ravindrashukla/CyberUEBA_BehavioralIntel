#!/usr/bin/env python3
"""Regenerate Figure 3 with subplot titles that match each panel (comment #18)
and swap it into the doc in place. Saves a new file; original untouched."""
import os, glob
from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(__file__)
SRC = os.path.join(HERE, "UEBA_Behavioral_Intelligence_Business_Friendly_v5.docx")
OUT = os.path.join(HERE, "UEBA_Behavioral_Intelligence_Business_Friendly_v6.docx")
PNG = os.path.join(HERE, "_fig3_fixed.png")

doc = Document(SRC)

# locate the "Figure 3 —" caption, then the nearest preceding image paragraph
cap_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("Figure 3"):
        cap_idx = i
        break
if cap_idx is None:
    raise SystemExit("Figure 3 caption not found")

rId = None
extent = None
for j in range(cap_idx, max(cap_idx - 6, -1), -1):
    p = doc.paragraphs[j]
    blips = p._p.findall(".//" + qn("a:blip"))
    if blips:
        rId = blips[0].get(qn("r:embed"))
        ext = p._p.findall(".//" + qn("wp:extent"))
        if ext:
            cx = int(ext[0].get("cx")); cy = int(ext[0].get("cy"))
            extent = (cx, cy)
        break
if rId is None:
    raise SystemExit("Figure 3 image not found")
aspect = (extent[0] / extent[1]) if extent else 2.6
print(f"Figure 3 rId={rId} aspect={aspect:.2f}")

# ── regenerate the figure ──
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

methods = ["Local Outlier\nFactor", "Z-Score", "Isolation\nForest",
           "One-Class\nSVM", "V-Intelligence\nComposite"]
detect = [0, 1, 0, 0, 4]
fp = [4.5, 9.8, 5.3, 14.6, 8.1]
GREY = "#8A95A1"
TEAL = "#0E6B8A"
RED = "#C0392B"
NAVY = "#0B1F3A"
colors = [GREY, GREY, GREY, GREY, TEAL]

h = 4.0
w = h * aspect
fig, (axL, axR) = plt.subplots(1, 2, figsize=(w, h), dpi=200)

# left: detection count
barsL = axL.bar(range(5), detect, color=colors, width=0.62, edgecolor="white")
axL.set_ylim(0, 4.6)
axL.set_ylabel("Attack campaigns detected (of 4)", fontsize=9, color=NAVY)
axL.set_title("How many of the 4 campaigns each method catches",
              fontsize=9.5, color=NAVY, fontweight="bold", pad=10)
for i, v in enumerate(detect):
    axL.text(i, v + 0.12, f"{v}/4", ha="center", fontsize=9.5, fontweight="bold",
             color=(RED if v == 0 else (TEAL if v == 4 else NAVY)))
axL.set_xticks(range(5)); axL.set_xticklabels(methods, fontsize=7.5)
axL.spines[["top", "right"]].set_visible(False)
axL.tick_params(axis="y", labelsize=8)

# right: false-positive rate
axR.bar(range(5), fp, color=colors, width=0.62, edgecolor="white")
axR.set_ylim(0, 16.5)
axR.set_ylabel("False-positive rate (%)", fontsize=9, color=NAVY)
axR.set_title("False-positive rate at that detection level",
              fontsize=9.5, color=NAVY, fontweight="bold", pad=10)
for i, v in enumerate(fp):
    axR.text(i, v + 0.4, f"{v:.1f}%", ha="center", fontsize=9, fontweight="bold", color=NAVY)
axR.set_xticks(range(5)); axR.set_xticklabels(methods, fontsize=7.5)
axR.spines[["top", "right"]].set_visible(False)
axR.tick_params(axis="y", labelsize=8)

fig.suptitle("Detection scoreboard — 250 users, 4 embedded campaigns",
             fontsize=10.5, color=NAVY, fontweight="bold", y=1.02)
plt.tight_layout()
fig.savefig(PNG, bbox_inches="tight", facecolor="white")
plt.close(fig)
print("figure regenerated")

# ── swap the image blob in place ──
part = doc.part.related_parts[rId]
with open(PNG, "rb") as f:
    part._blob = f.read()

doc.save(OUT)
print(f"Saved: {os.path.basename(OUT)}")
os.remove(PNG)
