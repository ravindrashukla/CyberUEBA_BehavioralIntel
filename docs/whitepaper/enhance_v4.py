#!/usr/bin/env python3
"""v4: correct the §4.4 explainer placement (it must follow the phase table),
add a 'how the score is built' flow diagram, and resolve the remaining
comments (#7 activities, #8 slop, #9 Volt breadth, #20 forward-reference).
Builds on v3; original untouched. Output: ..._v4.docx
"""
import os, glob
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(__file__)
V3 = os.path.join(HERE, "UEBA_Behavioral_Intelligence_Business_Friendly_v3.docx")
OUT = os.path.join(HERE, "UEBA_Behavioral_Intelligence_Business_Friendly_v5.docx")
DIAG = os.path.join(HERE, "_score_flow.png")
RUST = RGBColor(0xC2, 0x5A, 0x12)
NAVY = "#0B1F3A"
RUSTHEX = "#C25A12"

# ── 1. build the flow diagram ──
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

fig, ax = plt.subplots(figsize=(11, 3.0), dpi=200)
ax.set_xlim(0, 100); ax.set_ylim(0, 30); ax.axis("off")

def box(x, y, w, h, title, sub, fc, ec):
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.3,rounding_size=1.2",
                 linewidth=1.4, edgecolor=ec, facecolor=fc))
    ax.text(x + w/2, y + h*0.62, title, ha="center", va="center",
            fontsize=8.5, fontweight="bold", color=NAVY)
    if sub:
        ax.text(x + w/2, y + h*0.27, sub, ha="center", va="center",
                fontsize=6.6, color="#333333")

def arrow(x1, y1, x2, y2):
    ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>",
                 mutation_scale=12, linewidth=1.3, color="#777777"))

box(1, 9, 15, 12, "5 Behavioral Zones", "identity · access · data\nnetwork · risk", "#EAF0F6", NAVY)
box(20, 9, 14, 12, "Per-Zone Embedding\n→ Drift", "each zone embedded,\ndrifts on its own", "#EAF0F6", NAVY)
box(38, 9, 15, 12, "Peer-relative\nz-scores", "vs. the user's\nrole cohort", "#EAF0F6", NAVY)
box(57, 16, 17, 11, "Four deviation phases", "signal · breadth\nsustained · context", "#FDF1E7", RUSTHEX)
box(57, 2, 17, 11, "Novelty persistence", "new contact that\nkeeps recurring?", "#FDF1E7", RUSTHEX)
box(79, 9, 18, 12, "Composite Score", "one ranked number\nper user", "#E7F0EA", "#1E8A49")

arrow(16, 15, 20, 15)
arrow(34, 15, 38, 15)
arrow(53, 15, 57, 21.5)   # to phases
arrow(53, 15, 57, 7.5)    # to novelty
arrow(74, 21.5, 79, 16)   # phases to score
arrow(74, 7.5, 79, 14)    # novelty to score
ax.text(50, 28.2, "How a behavioral score is assembled", ha="center",
        fontsize=10, fontweight="bold", color=NAVY)
plt.tight_layout()
fig.savefig(DIAG, bbox_inches="tight", facecolor="white")
plt.close(fig)
print("diagram built")

# ── 2. load v3 and edit ──
doc = Document(V3)


def set_text(p, text):
    f = p.runs[0]
    size, bold, italic = f.font.size, f.font.bold, f.font.italic
    color = f.font.color.rgb if (f.font.color and f.font.color.type is not None) else None
    for r in list(p.runs):
        r.text = ""
    r0 = p.runs[0]; r0.text = text
    if size: r0.font.size = size
    r0.font.bold = bold; r0.font.italic = italic
    if color is not None: r0.font.color.rgb = color


# 2a. remove the 3 mis-placed explainer paragraphs from §4.3
remove_keys = ["How a score becomes a number.",
               "weekly drift is summarized into a few simple traits",
               "in the score is one standard deviation of peer-relative deviation"]
for p in list(doc.paragraphs):
    if any(k in p.text for k in remove_keys):
        p._p.getparent().remove(p._p)

# 2b. simple phrase fixes (within-run, preserve format)
PAIRS = [
    ("What the evidence shows, stated honestly:", "What the evidence shows:"),   # #8 slop
    ("where behaviors that mean similar things sit near each other, even when their raw numbers differ.",
     "where activities that play similar roles — a sign-in, an app launch, a data "
     "download, an API call — sit near each other when they mean similar things, "
     "even when their raw numbers differ."),                                      # #7
    ("the case that justifies the architecture, detailed below.",
     "despite looking ordinary on every volume measure — the hardest case in the "
     "study, examined in Section 5.3."),                                          # #20
    ("a deliberately quiet campaign using only legitimate admin tools",
     "a deliberately quiet campaign using only legitimate admin tools, so its "
     "signal is spread thin across many features at once rather than concentrated "
     "in one spike — caught by breadth"),                                         # #9
]
for p in doc.paragraphs:
    for old, new in PAIRS:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
        joined = "".join(r.text for r in p.runs)
        if old in joined and not any(old in r.text for r in p.runs):
            set_text(p, joined.replace(old, new))

# 2c. re-insert explainer + diagram in §4.4, right after the phase table
#     anchor = the paragraph that follows the phase table
anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith("Each phase alone is insufficient"):
        anchor = p
        break
if anchor is None:
    raise SystemExit("anchor 'Each phase alone is insufficient' not found")

def ins_para(anchor_p, text, bold=False, color=None, size=10.5, sb=4, sa=6, italic=False):
    np = anchor_p.insert_paragraph_before()
    r = np.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    np.paragraph_format.space_before = Pt(sb)
    np.paragraph_format.space_after = Pt(sa)
    return np

ins_para(anchor, "How a score becomes a number.", bold=True, color=RUST, size=11, sb=8, sa=3)
ins_para(anchor,
    "Each zone's weekly drift is summarized into a few simple traits — how large, "
    "how broad, how persistent, how it trends, and how late it escalates. Every trait "
    "is then expressed as a z-score: the number of standard deviations it sits above "
    "the norm for the user's role peers, so a developer is judged against developers, "
    "not executives. The five phases in the table above are summaries of those "
    "z-scores — the strongest few, how many are elevated, how persistent they are, and "
    "how the picture changes under attack-specific lenses — plus novelty persistence, "
    "which simply asks whether a never-before-seen contact keeps recurring.")
ins_para(anchor,
    "So a “point” in the score is one standard deviation of peer-relative "
    "deviation, and the five phases are combined with fixed weights (withheld here) "
    "into a single ranked number. Normal users sit near the center of that "
    "peer-relative space; only a genuine attacker reaches the edges on several phases "
    "at once.")
# diagram
pic_para = anchor.insert_paragraph_before()
pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
pic_para.add_run().add_picture(DIAG, width=Inches(6.2))
cap = anchor.insert_paragraph_before()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cap.add_run("Schematic — how five behavioral zones become one ranked score. "
                 "Four phases read the deviation structure; novelty persistence is a "
                 "separate recurrence signal.")
cr.italic = True; cr.font.size = Pt(8.5); cr.font.color.rgb = RGBColor(0x55,0x55,0x55)

doc.save(OUT)
print(f"Saved: {os.path.basename(OUT)}")
if os.path.exists(DIAG):
    os.remove(DIAG)
