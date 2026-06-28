#!/usr/bin/env python3
"""Enhance the KS-commented business whitepaper: de-slop the language Kunal
flagged, smooth clunky prose, fix specific clarity gaps, and insert a crisp
'how the score is built' explainer. Preserves images/tables/formatting.
Saves a NEW file; original KS copy untouched.
"""
import glob, os, copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

HERE = os.path.dirname(__file__)
SRC = glob.glob(os.path.join(HERE, "*_KS.docx"))[0]
OUT = os.path.join(HERE, "UEBA_Behavioral_Intelligence_Business_Friendly_v3.docx")
RUST = RGBColor(0xC2, 0x5A, 0x12)

doc = Document(SRC)

# ── simple within-run phrase replacements (preserve formatting) ──
PAIRS = [
    ("The honest scoreboard first", "The scoreboard"),                       # #14
    ("Limitations — Stated Plainly", "Limitations"),                    # #24
    ("The value claim, stated precisely:", "The value claim:"),              # #16
    ("(measured, synthetic data; verified against the evaluation dataset)",
     "(synthetic dataset, verified against the evaluation run)"),            # #19
    ("(measured, synthetic data)", "(synthetic dataset)"),                   # #19
    ("the insider changes WHAT they access, not HOW MUCH",
     "the insider changes which data they access, not how much"),           # #27
    ("The insider changes WHAT they access, not HOW MUCH.",
     "The insider changes which data they access, not how much."),          # #27
    ("The APT changes WHERE it communicates, not HOW OFTEN.",
     "The APT changes where it communicates, not how often."),              # #27
    ("disconnected secure environments", "air-gapped secure environments"), # #25
    ("fully disconnected deployments", "air-gapped deployments"),           # #25
    ("The product is not just a better alarm; it is a better lead.",
     "The product is not just a better alarm — it gives the analyst a "
     "place to start: the named threat, the zone that changed, and the "
     "evidence behind the score."),                                          # #17
    ("did not detect them; outside agencies did, years and months late.",
     "did not detect them; outside agencies did, years and months late — "
     "long after the adversaries had mapped infrastructure, harvested "
     "credentials, and pre-positioned for disruption."),                     # #26
    ("We report the result without spin — every number below is reproduced "
     "directly from the evaluation data.",
     "Every number below comes directly from the evaluation data."),         # #15
]


def replace_phrases(p):
    changed = False
    for old, new in PAIRS:
        # within a single run
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
                changed = True
        # spanning runs: fall back to whole-paragraph set on first run
        if old in p.text and not changed and old not in "".join(r.text for r in p.runs):
            pass
    # cross-run fallback: if the joined text still contains an 'old', rewrite paragraph text
    joined = "".join(r.text for r in p.runs)
    for old, new in PAIRS:
        if old in joined and not any(old in r.text for r in p.runs):
            fmt = p.runs[0] if p.runs else None
            newjoined = joined.replace(old, new)
            for r in list(p.runs):
                r.text = ""
            if p.runs:
                p.runs[0].text = newjoined
            changed = True
    return changed


for p in doc.paragraphs:
    replace_phrases(p)


# ── whole-paragraph rewrites (multi-run, supply clean prose) ──
def set_text(p, text):
    if not p.runs:
        p.add_run(text)
        return
    f = p.runs[0]
    name, size, bold, italic = f.font.name, f.font.size, f.font.bold, f.font.italic
    color = f.font.color.rgb if (f.font.color and f.font.color.type is not None) else None
    for r in list(p.runs):
        r.text = ""
    r0 = p.runs[0]
    r0.text = text
    if size: r0.font.size = size
    r0.font.bold = bold
    r0.font.italic = italic
    if color is not None:
        r0.font.color.rgb = color


REWRITES = {
    "Tracking each entity's weekly position on the map":
        "Tracking an entity's position week over week produces a trajectory — "
        "and the trajectory reveals drift: the direction its behavior is trending, "
        "how fast, and how steadily. A cumulative change-detection method — the "
        "same family used to catch slow equipment degradation in industrial quality "
        "control — adds up these small weekly shifts until they become impossible "
        "to dismiss. This is what catches “slow and low.” The exact "
        "accumulation logic, thresholds, and parameters are proprietary and withheld "
        "in this edition.",
    "Direction is then named. The system maintains a library":
        "The system then names the direction. It holds a library of known attack "
        "patterns — short, plain-language descriptions of how real adversaries "
        "behave, each mapped to MITRE ATT&CK, the public catalog of attacker "
        "techniques. When an entity's drift points toward one of them — data "
        "staging for exfiltration, say, or covert beaconing — the alert names "
        "that pattern and attaches the matching ATT&CK techniques. Every simulated "
        "campaign in this study was built from real, documented technique sets and "
        "is scored against them.",
    "Stability here is itself a clue":
        "A stable identity zone is itself the clue: an insider and a credential "
        "thief both keep using the same legitimate account, so the identity zone "
        "barely moves while the data or network zones drift — and that "
        "split is the signature.",
}
for p in doc.paragraphs:
    for key, new in list(REWRITES.items()):
        if key in p.text:
            set_text(p, new)
            del REWRITES[key]
            break


# ── table cell fix: Baseline 'per entity' confusion (#2/#3) ──
for t in doc.tables:
    for row in t.rows:
        c0 = row.cells[0].text.strip().lower()
        if c0 == "baseline":
            # last cell is the 'UEBA answer'
            target = row.cells[-1]
            for para in target.paragraphs:
                if para.runs:
                    set_text(para, "Each entity's own baseline, calibrated to its "
                             "role cohort — a developer's “normal” "
                             "differs from an executive's.")
                    break
        if c0 == "deviation":
            target = row.cells[-1]
            for para in target.paragraphs:
                if "drifting from its own baseline" in para.text and para.runs:
                    set_text(para, "The entity's behavioral trajectory is drifting "
                             "from its own baseline and from its role-peer norm.")
                    break


# ── insert 'How the score is built' explainer after the §4.4 phase table (#21/#22) ──
anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith("Each attack family writes"):
        anchor = p
        break
if anchor is not None:
    def new_para_before(anchor_p, text, bold=False, color=None, size=10.5,
                        sb=4, sa=6):
        from docx.text.paragraph import Paragraph
        new_p = copy.deepcopy(anchor_p._p)
        for r in list(new_p.findall(qn("w:r"))):
            new_p.remove(r)
        anchor_p._p.addprevious(new_p)
        para = Paragraph(new_p, anchor_p._parent)
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
        para.paragraph_format.space_before = Pt(sb)
        para.paragraph_format.space_after = Pt(sa)
        return para

    new_para_before(anchor, "How a score becomes a number.", bold=True, color=RUST,
                    size=11, sb=6, sa=3)
    new_para_before(anchor,
        "Each zone’s weekly drift is summarized into a few simple traits — "
        "how large, how broad, how persistent, how it trends, and how late it "
        "escalates. Every trait is then expressed as a z-score: the number of "
        "standard deviations it sits above the norm for the user’s role peers, "
        "so a developer is judged against developers, not executives. The five "
        "phases in the table above are summaries of those z-scores — the "
        "strongest few, how many are elevated, how persistent they are, and how the "
        "picture changes under attack-specific lenses — plus novelty "
        "persistence, which simply asks whether a never-before-seen contact keeps "
        "recurring.")
    new_para_before(anchor,
        "So a “point” in the score is one standard deviation of "
        "peer-relative deviation, and the five phases are combined with fixed "
        "weights (withheld here) into a single ranked number. The axes in the next "
        "figure show each user’s position as a percentile within that "
        "peer-relative space — which is why normal users cluster at the center "
        "and only a genuine attacker reaches the edges on several phases at once.")

doc.save(OUT)
print(f"Saved: {os.path.basename(OUT)}")
