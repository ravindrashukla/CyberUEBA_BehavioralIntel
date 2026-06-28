#!/usr/bin/env python3
"""Integrate a crisp 'verification' block into the business whitepaper as
subsections 5.5 and 5.6 (no renumbering), and emit audience-tailored versions.

Preserves the original. Outputs:
  UEBA_Behavioral_Intelligence_Business_Friendly_v2.docx          (integrated)
  UEBA_..._Business_Friendly_v2_DLA.docx  / _USSOCOM / _DISA / _Army  (fronts)
"""
import os
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

HERE = os.path.dirname(__file__)
SRC = os.path.join(HERE, "UEBA_Behavioral_Intelligence_Business_Friendly.docx")

RUST = RGBColor(0xC2, 0x5A, 0x12)   # subsection heading color (matches doc)
NAVY = RGBColor(0x0B, 0x1F, 0x3A)
HDR_FILL = "0B1F3A"


def sub_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RUST
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p


def body(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)
    return p


def make_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HDR_FILL}"/>'))
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if ri % 2 == 1:
                c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEF2F7"/>'))
    return t


def build_v2():
    doc = Document(SRC)

    # find anchor = the "6.  Deployment in Practice" heading
    anchor = None
    for p in doc.paragraphs:
        tx = p.text.strip()
        if tx.startswith("6.") and "Deployment" in tx:
            anchor = p
            break
    if anchor is None:
        raise SystemExit("Anchor '6. Deployment' not found")

    # build new content at end of doc, collect body elements in order
    new_elems = []

    def take(el):
        new_elems.append(el)

    # 5.5
    take(sub_heading(doc, "5.5  Were the attackers actually hard to find?")._p)
    take(body(doc,
        "A detection result is only as strong as the difficulty of the problem it solved. Before "
        "trusting the scoreboard above, the fair question is whether the simulated attackers were "
        "genuinely evasive — or merely obvious. Three measured checks confirm they were hard to "
        "find. The hardest case, the slow APT (USR-234), is used throughout.")._p)
    take(body(doc,
        "Statistical twins. Averaged across all weeks, the slow APT is indistinguishable from an "
        "ordinary user — and on two measures it is actually quieter:")._p)
    take(make_table(doc,
        ["Weekly measure", "Slow APT", "Normal user", "Verdict"],
        [
            ["File accesses", "70.9", "68.7", "~ identical"],
            ["Restricted-file ratio", "0.046", "0.060", "attacker LOWER"],
            ["Network bytes sent", "782M", "711M", "within variance"],
            ["Failed-DNS ratio", "0.039", "0.044", "attacker LOWER"],
        ])._tbl)
    take(body(doc,
        "A whisper in a stadium. In the raw logs, the attacker's command-and-control beacon moves "
        "120-487 bytes across a whole day - roughly 380,000 times smaller than the single largest "
        "ordinary network flow on the same day (184 MB). No volume threshold can isolate the "
        "whisper without drowning in normal traffic.")._p)
    take(body(doc,
        "Ordinary among peers. Measured against its 62 role peers (the developer cohort), the "
        "attacker sits inside the normal middle-50% range on file access, restricted-file ratio, "
        "and network volume. By volume, it is an unremarkable developer. Only the direction of its "
        "behavior over time - and the recurrence of one novel address in 60 of 60 weeks - sets it "
        "apart.")._p)

    # 5.6
    take(sub_heading(doc, "5.6  How invisible weekly changes add up")._p)
    take(body(doc,
        "If the attacker never spikes, how is it caught? By accumulation. Each week the system "
        "measures how far behavior moved, subtracts a small allowance for normal week-to-week "
        "wobble (measured at 0.0138 on this data), and adds the remainder to a running total that "
        "resets to zero whenever a week looks normal. A sustained attacker drifts a little further "
        "in the same direction every week, so the total climbs and never resets; a normal user's "
        "keeps resetting. This running total is called CUSUM - the mathematical answer to 'slow "
        "and low.'")._p)
    take(make_table(doc,
        ["Entity", "Type", "Cumulative drift", "vs. normal user"],
        [
            ["USR-156", "Insider threat", "0.779", "29x"],
            ["USR-042", "Volt Typhoon (LOTL)", "0.474", "18x"],
            ["USR-234", "Slow APT (C2 beacon)", "0.162", "6x"],
            ["USR-118", "Salt Typhoon (telecom)", "0.131", "5x"],
            ["USR-001", "Normal user", "0.026", "baseline"],
        ])._tbl)
    take(body(doc,
        "Every attacker accumulates between 5 and 29 times the normal user's drift - even though, "
        "in any single week, none of them crossed an alarm threshold. Accumulation converts many "
        "invisible weekly changes into one unmistakable signal. (Every figure in 5.5-5.6 is "
        "reproducible live in the platform's interactive 'Proof of Realism' view, which lets a "
        "reviewer pick any two users and watch volume statistics overlap while cumulative-drift "
        "curves diverge.)")._p)

    # move new elements (currently at end) to just before the anchor, in order
    for el in new_elems:
        anchor._p.addprevious(el)

    out = os.path.join(HERE, "UEBA_Behavioral_Intelligence_Business_Friendly_v2.docx")
    doc.save(out)
    print(f"Saved: {os.path.basename(out)}")
    return out


def build_audience(v2_path):
    audiences = {
        "DLA": "Defense Logistics Agency - Office of the CIO and PEO",
        "USSOCOM": "United States Special Operations Command - SDA, EIS, S&T, Innovation",
        "DISA": "Defense Information Systems Agency - PEO for Cyber",
        "Army": "United States Army - Office of the Chief of Artificial Intelligence",
    }
    for key, line in audiences.items():
        doc = Document(v2_path)
        # find the metadata paragraph that contains 'Business Edition' and add a prepared-for line after it
        target = None
        for p in doc.paragraphs:
            if "Business Edition" in p.text:
                target = p
                break
        if target is None:
            target = doc.paragraphs[2]
        newp = copy.deepcopy(target._p)
        # clear runs in the copy and set the prepared-for text
        for r in list(newp.findall(qn("w:r"))):
            newp.remove(r)
        target._p.addnext(newp)
        from docx.text.paragraph import Paragraph
        para = Paragraph(newp, target._parent)
        run = para.add_run(f"Prepared for: {line}")
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = NAVY
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        out = os.path.join(HERE, f"UEBA_Behavioral_Intelligence_Business_Friendly_v2_{key}.docx")
        doc.save(out)
        print(f"Saved: {os.path.basename(out)}")


if __name__ == "__main__":
    v2 = build_v2()
    build_audience(v2)
    print("Done.")
