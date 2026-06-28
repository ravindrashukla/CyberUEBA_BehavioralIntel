# -*- coding: utf-8 -*-
"""Strip every partner / Rigor / Cequence reference from existing docs, in place.
Keeps 'Mythos' (22nd Century's own Layer-4/code capability, not a partner).
Plain text replacement (not tracked). Robust to runs split mid-phrase.
"""
import os
from docx import Document

REPLACE = [
 ("V-Intelligence, Rigor AI, and Mythos", "V-Intelligence and Mythos"),
 ("Rigor AI: ", ""),
 ("(Rigor, API, Mythos)", "(Preemptive, API, Code)"),
 ("Layer 1 preemptive (Rigor)", "Layer 1 preemptive"),
 ("Layer 1 Rigor (preemptive)", "Layer 1 (preemptive)"),
 ("are separate products and are not in this codebase",
  "are part of the 22nd Century four-layer platform and are not in this behavioral codebase"),
 ("(SEPARATE products)", "(22nd Century platform)"),
 ("some are partner or separate capabilities", "they are part of the 22nd Century four-layer platform"),
 ("partner or separate capabilities", "part of the 22nd Century four-layer platform"),
 ("(Rigor)", ""),
 ("Rigor AI", "the Preemptive layer"),
 ("Cequence Security", "the API/Data layer"),
 ("Cequence", "the API/Data layer"),
 ("Rigor", ""),
]

def para_replace(p):
    full="".join(r.text for r in p.runs)
    new=full; changed=False
    for old,rep in REPLACE:
        if old in new: new=new.replace(old,rep); changed=True
    if not changed: return 0          # only touch paragraphs with a real partner hit
    new=new.replace("  "," ")          # tidy spacing only on changed paragraphs
    if new!=full and p.runs:
        p.runs[0].text=new
        for r in p.runs[1:]: r.text=""
        return 1
    return 0

def strip_doc(path):
    d=Document(path)
    n=0
    for p in d.paragraphs: n+=para_replace(p)
    for tb in d.tables:
        for row in tb.rows:
            for c in row.cells:
                for p in c.paragraphs: n+=para_replace(p)
    if n:
        d.save(path)
    return n

TARGETS=[
 "Brandon_Pugh_Tough_Questions_Cheat_Sheet_INTERNAL.docx",
 "Cyber_DeepDive_Briefing_Book_Brandon_Pugh_INTERNAL.docx",
 "V-Intelligence_Layered_Cybersecurity_Whitepaper_v2.docx",
 "V-Intelligence_Layered_Cybersecurity_Whitepaper.docx",
]
base="WP DLA/Presentation"
for fn in TARGETS:
    path=os.path.join(base,fn)
    if not os.path.exists(path): print("  MISSING",fn); continue
    try:
        n=strip_doc(path); print(f"  {n:2d} edits  {fn}")
    except PermissionError:
        print(f"  LOCKED    {fn}  <-- close in Word, re-run")
    except Exception as e:
        print(f"  ERROR     {fn}: {e}")
