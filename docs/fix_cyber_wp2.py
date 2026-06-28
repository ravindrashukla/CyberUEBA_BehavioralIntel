# -*- coding: utf-8 -*-
"""Apply review-round refinements to the US Army / Pentagon cyber whitepaper."""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
P = "WP DLA/Presentation/V-Intelligence_Layered_Cybersecurity_Whitepaper_US_Army_Pentagon.docx"
d = Document(P)
def acc(p): return "".join(t.text or "" for t in p._p.iter(qn('w:t')))
def find_contains(s):
    for p in d.paragraphs:
        if s in acc(p): return p
def settext(p, new):
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]: r.text = ""
    else: p.add_run(new)
def ins_after(anchor, text, style):
    np = OxmlElement('w:p'); anchor._p.addnext(np)
    pp = Paragraph(np, anchor._parent); pp.style = d.styles[style]
    if text: pp.add_run(text)
    return pp
log=[]

# (1)+(2) Layer 2 proof: add dataset specifics + label ranks
p125 = find_contains("all four nation-state-style attackers were surfaced (4 of 4)")
if p125:
    settext(p125, "In that evaluation, on 250 entities over 485 days (about 14 million events across five log sources), all four nation-state-style attackers were surfaced (4 of 4) at an 8.1 percent false-positive operating point, where classical point-anomaly methods (Isolation Forest, One-Class SVM, and Local Outlier Factor) caught none of the four and a z-score method caught one. The four surfaced at ranks Salt #1, Insider #2, Slow APT #7, and Volt #24 on a risk-ordered list, with the stealthiest living-off-the-land campaign still well inside the operating band. These are controlled, synthetic-data proof-of-concept results, indicative rather than a guarantee of field performance.")
    log.append("L2 proof: dataset specifics + rank labels")

# (5) trim redundant caveat in the 'Illustrative result' paragraph
pill = find_contains("Illustrative result. In an internal blind")
if pill:
    t = acc(pill)
    for tail in [", indicative rather than a guarantee of field performance",
                 " These are controlled, synthetic-data results, indicative rather than a guarantee of field performance."]:
        t = t.replace(tail, "")
    settext(pill, t)
    log.append("trimmed duplicate caveat in Illustrative result")

# (3) serial form IL5 and IL6 and JWICS -> IL5, IL6, and JWICS (everywhere)
for p in d.paragraphs:
    t = acc(p)
    if "IL5 and IL6 and JWICS" in t:
        settext(p, t.replace("IL5 and IL6 and JWICS", "IL5, IL6, and JWICS"))
        log.append("fixed IL5/IL6/JWICS serial form")

# (4)+(7) glossary subheadings -> Heading 2; serial comma
GLOSS = {"Network and Infrastructure", "Detection and Analytics", "Application, API and Data",
         "Code and Supply Chain", "Federal and Acquisition"}
for p in d.paragraphs:
    t = acc(p).strip()
    if t in GLOSS and not (p.style.name or "").startswith("Heading"):
        if t == "Application, API and Data":
            settext(p, "Application, API, and Data")
        p.style = d.styles["Heading 2"]
        log.append(f"glossary subheading -> H2: {t[:30]}")

# (6) Army hooks on Layer 3/4/5 Federal Fit
secnum = 0; in_target = False
army_line = "Deployable in IL5 and IL6 enclaves, extensible to JWICS, and relevant to ARCYBER operations and the Army software factory and data-governance mission."
fed_fits = []
cur_layer = None
for p in d.paragraphs:
    s = p.style.name or ""; t = acc(p).strip()
    if s.startswith("Heading 1") and (t.startswith("5. Layer 3") or t.startswith("6. Layer 4") or t.startswith("7. Layer 5")):
        cur_layer = t[:9]
    elif s.startswith("Heading 1"):
        cur_layer = None
    if cur_layer and s.startswith("Heading 2") and t.startswith("Federal and Critical"):
        fed_fits.append(p)
for fp in fed_fits:
    army = army_line.replace("IL5 and IL6", "IL5, IL6,")  # keep serial style
    ins_after(fp, army, "List Bullet")
log.append(f"Army hooks added to {len(fed_fits)} Federal-Fit sections (L3/4/5)")

d.save(P)
for l in log: print(" ", l)
print("saved.")
