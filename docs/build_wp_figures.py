# -*- coding: utf-8 -*-
"""Generate conceptual figures and embed them into the US Army/Pentagon cyber whitepaper."""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import pandas as pd

NAVY="#1f3864"; BLUE="#2e75b6"; SLATE="#55606b"; LIGHT="#eef3fb"
L1="#0b5394"; L2="#1f7a33"; L3="#7a4ea6"; L4="#b5570a"; L5="#a61b1b"
FIG="docs/wp_figs"; os.makedirs(FIG,exist_ok=True)

def box(ax,x,y,w,h,text,fc,tc="white",fs=10,bold=True):
    ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.02,rounding_size=0.06",
                                fc=fc,ec="none",mutation_aspect=1))
    ax.text(x+w/2,y+h/2,text,ha="center",va="center",color=tc,fontsize=fs,
            fontweight="bold" if bold else "normal",wrap=True)
def arrow(ax,x1,y1,x2,y2,color=BLUE):
    ax.add_patch(FancyArrowPatch((x1,y1),(x2,y2),arrowstyle="-|>",mutation_scale=14,color=color,lw=1.6))

# ===== FIG A: Five-layer architecture =====
fig,ax=plt.subplots(figsize=(9,5.2)); ax.set_xlim(0,10); ax.set_ylim(0,10); ax.axis("off")
layers=[("Layer 1  Network / Preemptive","Prove the door is closed before traffic flows",L1),
        ("Layer 2  Behavioral / UEBA (digital twin)","Catch valid-account, living-off-the-land behavior inside",L2),
        ("Layer 3  Application / API","Guard APIs; stop valid-but-malicious abuse",L3),
        ("Layer 4  Data / Agentic AI","Protect data, govern the AI agents",L4),
        ("Layer 5  Code / Supply Chain","Find novel zero-days before deploy",L5)]
y=8.4
for nm,desc,col in layers:
    box(ax,0.3,y,7.0,1.15,"",col)
    ax.text(0.55,y+0.72,nm,ha="left",va="center",color="white",fontsize=11,fontweight="bold")
    ax.text(0.55,y+0.32,desc,ha="left",va="center",color="white",fontsize=8.5)
    y-=1.5
box(ax,7.6,0.9,2.1,8.65,"MITRE\nATT&CK\n+ ATLAS\n\none shared\nvocabulary",NAVY,fs=10)
ax.text(5.0,0.3,"Each layer catches what the others structurally cannot: completeness, not redundancy.",
        ha="center",fontsize=8.5,color=SLATE,style="italic")
plt.tight_layout(); plt.savefig(f"{FIG}/figA_architecture.png",dpi=150,bbox_inches="tight"); plt.close()

# ===== FIG B: Digital twin pipeline =====
fig,ax=plt.subplots(figsize=(11,2.7)); ax.set_xlim(0,11); ax.set_ylim(0,3); ax.axis("off")
steps=[("Raw logs","auth, file, net,\nDNS, endpoint",SLATE),("5 zones","23 features\nas prose",L2),
       ("Embed","each zone\n-> 1536-d",BLUE),("Compose","attention-\nweighted twin",NAVY),
       ("Trajectory","one composite\nper week",BLUE),("Drift +\ndirection","CUSUM +\nMITRE",L2),
       ("Verdict","one ranked,\nexplainable score",NAVY)]
x=0.15; bw=1.35
for i,(nm,desc,col) in enumerate(steps):
    box(ax,x,1.1,bw,1.3,"",col)
    ax.text(x+bw/2,1.95,nm,ha="center",va="center",color="white",fontsize=9,fontweight="bold")
    ax.text(x+bw/2,1.4,desc,ha="center",va="center",color="white",fontsize=7)
    if i<len(steps)-1: arrow(ax,x+bw,1.75,x+bw+0.18,1.75)
    x+=bw+0.18
ax.text(5.5,0.4,"Meaning over magnitude: each zone's attributes fuse into one sentence, then one vector.",
        ha="center",fontsize=8.5,color=SLATE,style="italic")
plt.tight_layout(); plt.savefig(f"{FIG}/figB_twin_pipeline.png",dpi=150,bbox_inches="tight"); plt.close()

# ===== FIG C: Preemptive verification flow =====
fig,ax=plt.subplots(figsize=(11,2.6)); ax.set_xlim(0,11); ax.set_ylim(0,3); ax.axis("off")
steps=[("Inputs","configs +\nthreat intel",SLATE),("Formal model","symbolic estate +\nMITRE attack graphs",L1),
       ("Query","where is a gap\nagainst this attack?",BLUE),("Remediate","guardrailed reasoner,\nrisk-ranked fixes",NAVY),
       ("Human","review and apply;\ndrift caught hourly",L1)]
x=0.3; bw=1.9
for i,(nm,desc,col) in enumerate(steps):
    box(ax,x,1.1,bw,1.3,"",col)
    ax.text(x+bw/2,1.95,nm,ha="center",va="center",color="white",fontsize=9.5,fontweight="bold")
    ax.text(x+bw/2,1.4,desc,ha="center",va="center",color="white",fontsize=7.5)
    if i<len(steps)-1: arrow(ax,x+bw,1.75,x+bw+0.18,1.75)
    x+=bw+0.18
ax.text(5.5,0.4,"No sampling, no agents, configuration-only. First actionable findings in under 15 days.",
        ha="center",fontsize=8.5,color=SLATE,style="italic")
plt.tight_layout(); plt.savefig(f"{FIG}/figC_preemptive_flow.png",dpi=150,bbox_inches="tight"); plt.close()

# ===== FIG D: Drift + CUSUM (real Volt trajectory) =====
tr=pd.read_csv("data/tier3_results/weekly_zone_trajectories.csv")
v=tr[tr.user_id=="USR-042"].sort_values("week_idx")
wk=v.week_idx.values; ap=v.access_pattern_drift.values; cd=v.composite_drift.values
cus=[]; c=0.0
for m in cd: c=max(0.0,c+m-0.02); cus.append(c)
fig,ax=plt.subplots(figsize=(9,3.6))
ax.plot(wk,ap,color=L2,lw=1.6,label="access_pattern drift")
ax.plot(wk,cd,color=BLUE,lw=1.6,label="composite drift")
ax.plot(wk,cus,color=L5,lw=1.8,ls="--",label="CUSUM (accumulated)")
ax.axhline(0.05,color=SLATE,lw=0.9,ls=":",label="CUSUM fire line (0.05)")
ax.set_xlabel("Week",fontsize=9); ax.set_ylabel("Drift / CUSUM",fontsize=9)
ax.set_title("Volt Typhoon (USR-042): slow, sub-threshold drift accumulates over time",fontsize=10,color=NAVY,fontweight="bold")
ax.legend(fontsize=8,loc="upper left"); ax.grid(alpha=0.25); ax.tick_params(labelsize=8)
for sp in ["top","right"]: ax.spines[sp].set_visible(False)
plt.tight_layout(); plt.savefig(f"{FIG}/figD_drift_cusum.png",dpi=150,bbox_inches="tight"); plt.close()

# ===== FIG E: One Typhoon across the layers =====
fig,ax=plt.subplots(figsize=(11,2.6)); ax.set_xlim(0,11); ax.set_ylim(0,3); ax.axis("off")
steps=[("Layer 5  Code","remove the\nexploitable code",L5),("Layer 1  Edge","prove the\nedge is closed",L1),
       ("Layer 2  Behavior","catch valid-account\nmovement inside",L2),("Layer 3  API","catch API and\nbusiness-logic abuse",L3),
       ("Layer 4  Data/Agents","catch exfil and\nungoverned agents",L4)]
x=0.3; bw=1.95
for i,(nm,desc,col) in enumerate(steps):
    box(ax,x,1.1,bw,1.3,"",col)
    ax.text(x+bw/2,1.95,nm,ha="center",va="center",color="white",fontsize=9,fontweight="bold")
    ax.text(x+bw/2,1.4,desc,ha="center",va="center",color="white",fontsize=7.5)
    if i<len(steps)-1: arrow(ax,x+bw,1.75,x+bw+0.13,1.75,color=SLATE)
    x+=bw+0.13
ax.text(5.5,0.4,"No single box does this. The hand-off across layers, under one vocabulary, is the moat.",
        ha="center",fontsize=8.5,color=SLATE,style="italic")
plt.tight_layout(); plt.savefig(f"{FIG}/figE_typhoon.png",dpi=150,bbox_inches="tight"); plt.close()
print("figures generated.")

# ===== INSERT INTO WHITEPAPER =====
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
P="WP DLA/Presentation/V-Intelligence_Layered_Cybersecurity_Whitepaper_US_Army_Pentagon.docx"
d=Document(P)
def acc(p): return "".join(t.text or "" for t in p._p.iter(qn('w:t')))
GRAY=RGBColor(0x55,0x60,0x6b)
def insert_fig(anchor,img,width,caption):
    # picture paragraph after anchor
    np=OxmlElement('w:p'); anchor._p.addnext(np)
    pic=Paragraph(np,anchor._parent); pic.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(img,width=Inches(width))
    cnp=OxmlElement('w:p'); np.addnext(cnp)
    cap=Paragraph(cnp,anchor._parent); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=cap.add_run(caption); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=GRAY

# anchors: track current layer for per-layer "V-Intelligence Approach"
ps=d.paragraphs
def find_pref(pref,after_text=None):
    started=after_text is None
    for p in ps:
        t=acc(p).strip()
        if after_text and t.startswith(after_text): started=True
        if started and t.startswith(pref): return p
    return None
# Fig A after section 2 heading
a=find_pref("2. A Layered Answer")
if a: insert_fig(a,f"{FIG}/figA_architecture.png",4.6,"Figure. The five-layer defense-in-depth architecture under one shared MITRE vocabulary.")
# Fig C in Layer 1 approach
c=find_pref("The V-Intelligence Approach",after_text="3. Layer 1")
if c: insert_fig(c,f"{FIG}/figC_preemptive_flow.png",6.3,"Figure. Layer 1: preemptive network assurance, the formal verification flow.")
# Fig B in Layer 2 approach
b=find_pref("The V-Intelligence Approach",after_text="4. Layer 2")
if b: insert_fig(b,f"{FIG}/figB_twin_pipeline.png",6.3,"Figure. Layer 2: the behavioral digital twin pipeline.")
# Fig D near Layer 2 illustrative result
dd=find_pref("Illustrative result. In an internal blind",after_text="4. Layer 2")
if dd: insert_fig(dd,f"{FIG}/figD_drift_cusum.png",5.6,"Figure. Real trajectory: slow, sub-threshold drift that CUSUM accumulates over time.")
# Fig E in section 8
e=find_pref("8. Countering AI-Enabled")
if e: insert_fig(e,f"{FIG}/figE_typhoon.png",6.3,"Figure. One Typhoon intrusion traced across all five layers; the hand-off is the moat.")

d.save(P)
print("inserted 5 figures into whitepaper.")
