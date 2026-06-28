"""Refocus Use Case 3 in DLA Final.docx Executive Summary onto supply chain,
with cybersecurity UEBA used only as the proving analogy. Also de-emphasize
'cyber' in the 'One foundation' unique bullet. Clean edit."""
from docx import Document
from docx.oxml.ns import qn

PATH = "WP DLA/DLA Final.docx"
d = Document(PATH)

def accp(p):
    return "".join(t.text or "" for t in p._p.iter(qn('w:t')))

LABEL = "Logistics Network Early Warning:"
REST = (" BEI surfaces slow, sub-threshold, directional change across the logistics network, including supplier "
        "behavior, sourcing shifts, route exposure, and correlated movement across the industrial base, giving "
        "warning before disruption reaches a readiness threshold. The underlying detection principle was validated "
        "most rigorously in a cybersecurity proof of concept, drawn here purely as the analogy that proves it works: "
        "BEI detected all four slow, stealthy campaigns (4 of 4) at an 8.1 percent false-positive point, where "
        "traditional threshold methods caught zero to one, the same sub-threshold, living-off-the-land pattern seen "
        "in state-sponsored actors such as Volt Typhoon and Salt Typhoon. That ability to see slow, directional "
        "movement before any threshold breaks is exactly what early warning across a contested supply chain "
        "requires.")

uc3_done = False
for p in d.paragraphs:
    if accp(p).startswith("Network Early Warning (Cyber):"):
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
        rl = p.add_run(LABEL); rl.bold = True
        p.add_run(REST)
        uc3_done = True
        break

uniq_done = False
OLD = "across logistics and cyber"
NEW = "across very different mission domains"
for p in d.paragraphs:
    if accp(p).startswith("One foundation, many domains:") and OLD in accp(p):
        for t in p._p.iter(qn('w:t')):
            if t.text and OLD in t.text:
                t.text = t.text.replace(OLD, NEW)
                uniq_done = True
                break
        break

d.save(PATH)
print("Use Case 3 refocused:", uc3_done, "| unique bullet de-emphasized:", uniq_done)
