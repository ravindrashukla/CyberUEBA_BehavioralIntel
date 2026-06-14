#!/usr/bin/env python3
"""Build V-Intelligence UEBA Technical Specification Document (Word).

Comprehensive technical specification covering Tier 1 (Traditional)
and Tier 2 (V-Intelligence UEBA Basic) implementations — data pipeline, algorithms,
parameters, detection logic, and empirical results.

Output: docs/V_Intelligence_UEBA_Technical_Specification.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
BLUE = RGBColor(0x1B, 0x4F, 0x72)
TEAL = RGBColor(0x0E, 0x6B, 0x8A)
GOLD = RGBColor(0xB7, 0x95, 0x0B)
RED_ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREEN_ACCENT = RGBColor(0x1E, 0x8A, 0x49)
BLACK = RGBColor(0x00, 0x00, 0x00)

OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
                           "V_Intelligence_UEBA_Technical_Specification.docx")


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
    return h


def add_body_text(doc, text, bold=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    return p


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "0D1B2A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if val == "DETECTED":
                        run.font.color.rgb = GREEN_ACCENT
                        run.bold = True
                    elif val == "MISSED":
                        run.font.color.rgb = RED_ACCENT
                        run.bold = True

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


# ==========================================================================
# DOCUMENT BUILDER
# ==========================================================================

def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK

    # ── COVER PAGE ──────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("V-INTELLIGENCE UEBA TECHNICAL SPECIFICATION")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tier 1 (Traditional) | Tier 2 (V-Intelligence UEBA Basic) | Tier 3 (Digital Entity)")
    run.font.size = Pt(16)
    run.font.color.rgb = TEAL

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Data Pipeline, Algorithms, Parameters,\n"
        "Detection Logic, and Empirical Results\n"
        "17 Methods | 4 Attack Campaigns"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("22nd Century Technologies, Inc.\nV-Intelligence UEBA Program\nMay 2025")
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_ACCENT

    add_page_break(doc)

    # ── TABLE OF CONTENTS ──────────────────────────────────────────
    add_section_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Document Overview",
        "2. Data Generation Pipeline",
        "   2.1. Synthetic Telemetry Parameters",
        "   2.2. Entity Population",
        "   2.3. Log Types and Daily Event Volumes",
        "   2.4. Attack Scenario Injection",
        "3. Feature Engineering (Shared by All Tiers)",
        "   3.1. Weekly Aggregation Window",
        "   3.2. Feature Definitions (23 Features)",
        "   3.3. Feature Matrix Dimensions",
        "4. Tier 1: Traditional Detection",
        "   4.1. Static Anomaly Detection (4 algorithms)",
        "   4.2. Temporal Anomaly Detection (2 algorithms)",
        "   4.3. Algorithm Parameters",
        "   4.4. Tier 1 Results",
        "5. Tier 2: V-Intelligence UEBA Basic",
        "   5.1. Text Serialization (Scalar → Prose)",
        "   5.2. Embedding (Prose → 1536-d Vector)",
        "   5.3. Signal Composition (5 Signals → 1 Composite)",
        "   5.4. Three-Layer Detection",
        "   5.5. Reference Concepts (12 Concepts)",
        "   5.6. Detection Criteria",
        "   5.7. Tier 2 Results",
        "6. Tier 3: Digital Entity Detection",
        "   6.1. Hierarchical Zone Embeddings (5 Zones)",
        "   6.2. Context-Adaptive Composition (4 Scenarios)",
        "   6.3. Temporal Trajectory (Velocity/Acceleration/Regime)",
        "   6.4. Relationship Embeddings (Hadamard Products)",
        "   6.5. Tier 3 Detection Methods (9 Methods)",
        "   6.6. Tier 3 Results",
        "7. Optimal Detection Configuration",
        "8. Consolidated Results Table (17 Methods × 4 Attacks)",
        "9. UEBA Framework and Threat Taxonomy",
        "   9.1. What is UEBA?",
        "   9.2. Why SIEM Rules Fail Against Modern Threats",
        "   9.3. UEBA Threat Taxonomy (4 Classes)",
        "10. Detection Playbook: Threat Type → Algorithm",
        "   10.1. Per-Threat Recommendations",
        "   10.2. Layered Deployment Strategy",
        "11. Source File Reference",
        "12. Multi-Phase Composite Scoring — Deep Dive",
        "   12.1. The Five Phases",
        "   12.2. Composite Formula",
        "   12.3. Per-Attacker Composite Breakdown",
        "   12.4. USR-234 Novelty Persistence — Detailed Analysis",
        "13. Why Embeddings, Not Raw Feature Vectors?",
        "   13.1. Two Pipelines: Feature-Space vs Semantic",
        "   13.2. The Raw Vector Approach (Detailed)",
        "   13.3. Three Fundamental Limitations of Raw Vectors",
        "   13.4. What Embeddings Gain",
        "   13.5. What Embeddings Lose",
        "   13.6. The Design Decision: Both, Not Either",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        if not item.startswith("   "):
            run.bold = True
    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  1. DOCUMENT OVERVIEW
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "1. Document Overview", level=1)

    add_body_text(doc, (
        "This document is the authoritative technical specification for the V-Intelligence UEBA "
        "three-tier detection system: Tier 1 (Traditional Detection), Tier 2 (V-Intelligence UEBA "
        "Basic), and Tier 3 (Digital Entity). It records every parameter, algorithm "
        "configuration, data pipeline step, and detection criterion used to produce "
        "the empirical results across 17 detection methods against 4 attack campaigns. "
        "All results were generated using real OpenAI text-embedding-3-small embeddings "
        "(1536-d) against synthetic telemetry data."
    ))

    add_body_text(doc, (
        "V-Intelligence UEBA is a User and Entity Behavior Analytics (UEBA) system. UEBA detects "
        "threats by establishing behavioral baselines for every user and device, then "
        "flagging when behavior drifts from that baseline. Unlike rule-based SIEM systems "
        "that match known attack signatures, UEBA asks: 'Is this entity behaving "
        "differently than it used to?' This approach is essential because modern attackers "
        "use valid credentials, legitimate tools, and authorized access — they don't "
        "trigger signatures. The only detectable signal is behavioral change over time."
    ))

    add_body_text(doc, (
        "Key result: Embedding/composite scoring cleanly separates only 2 of the 4 attack "
        "campaigns — USR-156 (insider) and USR-118 (Salt Typhoon telecom) rank above all "
        "normal users, while USR-234 (slow APT, ranked #7) and USR-042 (Volt Typhoon LOTL, "
        "ranked #24) fall BELOW normal users on the composite. The clean 4/4 detection at "
        "0 false positives is achieved instead by the multi-front threat-profile detector "
        "(threat_profile_detector.py), which scores measurable known-bad profiles — C2-beacon, "
        "DGA-DNS, LOTL-process, cohort-rare access, recon-fanout, and insider-collection — using "
        "cohort-relative and raw-event features in a label-free manner. The critical finding is "
        "that different threat types require different detection algorithms — no single "
        "traditional or embedding method catches all 4 attacks at viable false positive rates."
    ))

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  2. DATA GENERATION PIPELINE
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "2. Data Generation Pipeline", level=1)

    add_section_heading(doc, "2.1. Synthetic Telemetry Parameters", level=2)

    create_table(doc,
        ["Parameter", "Value"],
        [
            ["Observation Window", "January 1, 2025 – May 13, 2025 (485 days)"],
            ["Weekly Windows", "70 weeks"],
            ["Total Users", "250 (including 4 attack targets)"],
            ["User Types", "~80% employee, ~15% contractor, ~5% privileged"],
            ["Departments", "15"],
            ["Roles", "31"],
            ["Attack Injection Rate", "0.06% of total events (~2,661 attack events in 4,214,251 total)"],
            ["Random Seed", "42 (reproducible)"],
        ],
        col_widths=[2.5, 4.0],
    )

    add_section_heading(doc, "2.2. Entity Population", level=2)

    create_table(doc,
        ["Entity", "Count", "ID Format", "Key Attributes"],
        [
            ["Users", "250", "USR-NNN", "department, role, clearance, user_type"],
            ["Devices", "~400", "DEV-NNN", "device_type, os_type, IP, owner_user_id"],
            ["Network Segments", "~10", "SEG-NN", "zone_type, trust_level, subnet"],
            ["Applications", "~20", "APP-NN", "app_type, data_classification"],
        ],
        col_widths=[1.5, 0.8, 1.0, 3.2],
    )

    add_section_heading(doc, "2.3. Log Types and Daily Event Volumes", level=2)

    create_table(doc,
        ["Log Type", "Source File Pattern", "Key Columns", "Approx. Events/Day"],
        [
            ["Authentication", "data/generated/auth/YYYY-MM-DD.csv",
             "user_id, success, source_ip, dest_system, auth_method, timestamp", "~5,000"],
            ["File Access", "data/generated/file_access/YYYY-MM-DD.csv",
             "user_id, file_path, operation, data_classification, file_size_bytes", "~8,000"],
            ["Network Flows", "data/generated/network/YYYY-MM-DD.csv",
             "device_id, src_ip, dst_ip, bytes_in, bytes_out, protocol", "~15,000"],
            ["DNS Queries", "data/generated/dns/YYYY-MM-DD.csv",
             "device_id, query_name/query_domain, response_code", "~10,000"],
            ["Endpoint Telemetry", "data/generated/endpoint/YYYY-MM-DD.csv",
             "device_id, user_id, process_name, event_type, risk_score", "~6,000"],
            ["Email (via endpoint)", "Simulated within endpoint/auth logs",
             "Attachment counts, external recipients", "Embedded"],
            ["Web Proxy (via DNS)", "Derived from DNS + network flows",
             "Blocked domains, unique domains visited", "Derived"],
        ],
        col_widths=[1.3, 2.0, 2.5, 1.0],
    )

    add_section_heading(doc, "2.4. Attack Scenario Injection", level=2)

    create_table(doc,
        ["Attack ID", "Target", "Type", "Start Date", "Duration", "Description"],
        [
            ["ATK-004", "USR-156", "Insider Threat", "2025-03-10", "60 days (2-month escalation)",
             "Gradual escalation: expanding file access scope, increasing off-hours "
             "activity, accessing restricted/confidential documents outside role, "
             "subtle privilege accumulation. Volume stays normal."],
            ["ATK-003", "USR-234", "Slow APT / C2", "2025-03-10", "65 days",
             "C2 beaconing via DNS to newly-registered domains at regular intervals, "
             "increasing unique DNS destinations, elevated network external ratio. "
             "Creates measurable network footprint."],
            ["ATK-007", "USR-042", "Volt Typhoon\nLOTL", "2025-03-15", "60 days",
             "Living-off-the-land: LOLBin abuse (PowerShell, wmic, certutil), "
             "WMI lateral movement, no malware dropped, slow data staging via "
             "legitimate admin tools. Minimal endpoint signatures."],
            ["ATK-008", "USR-118", "Salt Typhoon\nTelecom", "2025-03-15", "60 days",
             "Telecom infrastructure targeting: router config exfiltration, "
             "call metadata harvesting, DNS tunneling for C2, encrypted "
             "exfiltration channels. Strong network footprint."],
        ],
        col_widths=[0.7, 0.7, 1.0, 0.9, 0.7, 2.5],
    )

    add_body_text(doc, (
        "Attack design principle: Four campaigns test different detection axes. "
        "USR-156 (insider) changes behavioral DIRECTION without changing VOLUME. "
        "USR-234 (slow APT) creates a subtle network footprint via low-frequency "
        "C2 beaconing. USR-042 (Volt Typhoon) uses only legitimate tools — no "
        "malware, no signatures, pure LOTL. USR-118 (Salt Typhoon) targets "
        "infrastructure with DNS tunneling and encrypted exfiltration. Together "
        "these test direction vs. magnitude, endpoint vs. network, and tool-based "
        "vs. behavior-based detection."
    ), bold=True)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  3. FEATURE ENGINEERING
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "3. Feature Engineering (Shared by All Tiers)", level=1)

    add_section_heading(doc, "3.1. Weekly Aggregation Window", level=2)
    add_body_text(doc, (
        "All features are computed per user per 7-day window. The pipeline iterates "
        "through the observation period one week at a time, loading CSV files for "
        "each day within the window. For each user in each window, 23 scalar features "
        "are computed from the 5 log types."
    ))

    add_body_text(doc, (
        "Source: comparison/run_comparison.py → engineer_weekly_features() (line 72). "
        "User-to-device mapping is used to link network and DNS logs (which are keyed "
        "by device_id) to users (keyed by user_id)."
    ))

    add_section_heading(doc, "3.2. Feature Definitions (23 Features)", level=2)

    create_table(doc,
        ["#", "Feature Name", "Log Source", "Computation"],
        [
            ["1", "auth_total", "Authentication", "Count of all auth events for user in week"],
            ["2", "auth_failed", "Authentication", "Count where success == False"],
            ["3", "auth_fail_rate", "Authentication", "auth_failed / max(auth_total, 1)"],
            ["4", "auth_unique_sources", "Authentication", "source_ip.nunique()"],
            ["5", "auth_unique_dests", "Authentication", "dest_system.nunique()"],
            ["6", "auth_off_hours_ratio", "Authentication", "Fraction with hour < 8 or hour >= 18"],
            ["7", "auth_methods_used", "Authentication", "auth_method.nunique()"],
            ["8", "file_total", "File Access", "Count of all file access events"],
            ["9", "file_unique_paths", "File Access", "file_path.nunique()"],
            ["10", "file_restricted_ratio", "File Access", "Fraction with data_classification == 'restricted'"],
            ["11", "file_confidential_ratio", "File Access", "Fraction with data_classification == 'confidential'"],
            ["12", "file_write_ratio", "File Access", "Fraction with operation == 'write'"],
            ["13", "file_total_bytes", "File Access", "Sum of file_size_bytes"],
            ["14", "endpoint_total", "Endpoint", "Count of all endpoint telemetry events"],
            ["15", "endpoint_suspicious_ratio", "Endpoint", "Fraction with event_type == 'suspicious'"],
            ["16", "endpoint_max_risk", "Endpoint", "Max of risk_score column"],
            ["17", "endpoint_mean_risk", "Endpoint", "Mean of risk_score column"],
            ["18", "endpoint_unique_processes", "Endpoint", "process_name.nunique()"],
            ["19", "net_bytes_out", "Network", "Sum of bytes_out for user's devices"],
            ["20", "net_unique_dsts", "Network", "dst_ip.nunique() for user's devices"],
            ["21", "net_external_ratio", "Network", "Fraction where dst_ip is not RFC1918"],
            ["22", "dns_unique_domains", "DNS", "query_domain.nunique() for user's devices"],
            ["23", "dns_nxdomain_ratio", "DNS", "Fraction with response_code == 'NXDOMAIN'"],
        ],
        col_widths=[0.3, 1.8, 1.0, 3.4],
    )

    add_section_heading(doc, "3.3. Feature Matrix Dimensions", level=2)

    create_table(doc,
        ["Dimension", "Value"],
        [
            ["Users", "250"],
            ["Weeks", "70"],
            ["Features per user per week", "23"],
            ["Total feature vectors", "17,500 (250 × 70)"],
            ["Feature matrix shape (aggregated)", "250 × 23 (for static Tier 1)"],
            ["Feature matrix shape (temporal)", "17,500 × 23 (for temporal methods)"],
            ["Output file", "data/comparison_results/weekly_features.csv"],
        ],
        col_widths=[2.5, 4.0],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  4. TIER 1: TRADITIONAL DETECTION
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "4. Tier 1: Traditional Detection", level=1)

    add_body_text(doc, (
        "Tier 1 implements 6 detection algorithms representing industry-standard "
        "approaches deployed in commercial SIEM/UEBA products. These operate on the "
        "23 scalar features directly — no embedding or semantic transformation."
    ))

    add_section_heading(doc, "4.1. Static Anomaly Detection (4 Algorithms)", level=2)

    add_body_text(doc, (
        "Static methods aggregate each user's 70 weekly feature vectors into a single "
        "mean vector (250 × 23 matrix), apply StandardScaler, then run anomaly detection. "
        "Source: run_traditional_detection() (line 200)."
    ))

    add_body_text(doc, "Algorithm 1: Isolation Forest", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Rationale"],
        [
            ["contamination", "0.05", "Expect ~5% anomaly rate in enterprise environment"],
            ["n_estimators", "200", "More trees = stable isolation score"],
            ["random_state", "42", "Reproducibility"],
            ["Input", "StandardScaler(mean-aggregated features)", "250 × 23 matrix"],
            ["Decision", "predict == -1 → anomaly", "sklearn convention"],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    add_body_text(doc, "Algorithm 2: One-Class SVM", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Rationale"],
        [
            ["kernel", "rbf", "Radial basis function for non-linear boundaries"],
            ["gamma", "scale", "1 / (n_features × X.var()), sklearn default"],
            ["nu", "0.05", "Upper bound on fraction of outliers"],
            ["Input", "StandardScaler(mean-aggregated features)", "250 × 23 matrix"],
            ["Decision", "predict == -1 → anomaly", ""],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    add_body_text(doc, "Algorithm 3: Local Outlier Factor (LOF)", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Rationale"],
        [
            ["n_neighbors", "20", "Local density neighborhood (40% of population)"],
            ["contamination", "0.05", "Expected anomaly fraction"],
            ["Input", "StandardScaler(mean-aggregated features)", "250 × 23 matrix"],
            ["Decision", "predict == -1 → anomaly", ""],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    add_body_text(doc, "Algorithm 4: Z-Score Threshold", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Rationale"],
        [
            ["threshold", "|z| > 3.0", "Standard 3-sigma rule"],
            ["Input", "StandardScaler values (already z-scores)", "250 × 23 matrix"],
            ["Decision", "max(|z|) across all 23 features > 3.0 → anomaly", ""],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    add_section_heading(doc, "4.2. Temporal Anomaly Detection (2 Algorithms)", level=2)

    add_body_text(doc, (
        "Temporal methods preserve week-over-week time series. They train on the "
        "first half of the observation period and test on the second half. "
        "Source: run_temporal_traditional() (line 253) and run_feature_cusum() (line 307)."
    ))

    add_body_text(doc, "Algorithm 5: Temporal Z-Score", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Rationale"],
        [
            ["Training period", "Weeks 0–8 (first half)", "Establish per-user baseline"],
            ["Testing period", "Weeks 9–18 (second half)", "Detect deviations"],
            ["Baseline", "Per-user mean and std per feature from training weeks", ""],
            ["Test metric", "z-score of each test week vs. training baseline", ""],
            ["Detection", "max(|z|) across features and test weeks > 2.0 → anomaly", ""],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    add_body_text(doc, "Algorithm 6: Feature CUSUM", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Rationale"],
        [
            ["Baseline", "Per-user mean/std from first half of weeks", ""],
            ["Weekly drift metric", "Mean |z-score| across all 23 features", "Composite drift per week"],
            ["CUSUM threshold", "2.0", "Cumulative sum trigger level"],
            ["CUSUM drift_threshold", "0.5", "Baseline drift allowance per week"],
            ["min_run_length", "3", "Consecutive weeks above threshold to confirm"],
            ["Top-10% variant", "Users with CUSUM value ≥ 90th percentile", "Reduces FP rate"],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    add_section_heading(doc, "4.3. Algorithm Parameters Summary", level=2)

    create_table(doc,
        ["Algorithm", "Type", "Key Parameters", "Preprocessing"],
        [
            ["Isolation Forest", "Static", "contamination=0.05, n_estimators=200",
             "Mean aggregation → StandardScaler"],
            ["One-Class SVM", "Static", "kernel=rbf, gamma=scale, nu=0.05",
             "Mean aggregation → StandardScaler"],
            ["LOF", "Static", "n_neighbors=20, contamination=0.05",
             "Mean aggregation → StandardScaler"],
            ["Z-Score", "Static", "|z| > 3.0 threshold",
             "Mean aggregation → StandardScaler"],
            ["Temporal Z-Score", "Temporal", "train/test split at week 9, |z| > 2.0",
             "Per-user baseline from training weeks"],
            ["Feature CUSUM", "Temporal", "threshold=2.0, drift=0.5, run=3",
             "Per-user z-score trajectory"],
        ],
        col_widths=[1.5, 0.8, 2.5, 1.8],
    )

    add_section_heading(doc, "4.4. Tier 1 Results", level=2)

    create_table(doc,
        ["Method", "USR-156\n(Insider)", "USR-234\n(Slow APT)", "USR-042\n(Volt Typh)",
         "USR-118\n(Salt Typh)", "True\nPos", "False\nPos", "FP Rate"],
        [
            ["Isolation Forest", "MISSED", "MISSED", "MISSED", "MISSED", "0", "12", "4.9%"],
            ["One-Class SVM", "DETECTED", "MISSED", "MISSED", "DETECTED", "2", "26", "10.6%"],
            ["LOF", "MISSED", "MISSED", "MISSED", "DETECTED", "1", "11", "4.5%"],
            ["Z-Score (|z|>3)", "MISSED", "MISSED", "DETECTED", "DETECTED", "2", "25", "10.2%"],
            ["Temporal Z-Score", "DETECTED", "DETECTED", "DETECTED", "DETECTED", "4", "232", "94.3%"],
            ["Feature CUSUM Top10%", "DETECTED", "MISSED", "DETECTED", "DETECTED", "3", "22", "8.9%"],
        ],
        col_widths=[1.3, 0.8, 0.8, 0.8, 0.8, 0.5, 0.5, 0.6],
    )

    add_body_text(doc, (
        "Tier 1 finding: the traditional SIEM static detectors (Isolation Forest, One-Class "
        "SVM, LOF) detect approximately 0 of 4 attacks. Feature CUSUM appears to detect more, "
        "but the underlying Feature CUSUM signal fires on roughly 99% of normal users — an "
        "operationally useless false-positive rate — so its apparent coverage is not usable in "
        "practice. USR-234 (slow APT) is invisible to all Tier 1 methods except Temporal "
        "Z-Score (which has 94.3% FP, rendering it useless). No Tier 1 method detects all 4 at "
        "viable false positive rates."
    ), bold=True)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  5. TIER 2: V-INTELLIGENCE UEBA BASIC
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "5. Tier 2: V-Intelligence UEBA Basic", level=1)

    add_body_text(doc, (
        "Tier 2 uses the same 23 features as Tier 1 but transforms them through "
        "a semantic embedding pipeline: scalars → prose text → 1536-d vectors → "
        "weighted composition → drift direction analysis against reference concepts. "
        "Source: run_acecard_detection() (line 446)."
    ))

    add_section_heading(doc, "5.1. Text Serialization (Scalar → Prose)", level=2)

    add_body_text(doc, (
        "Each user's 23 weekly features are grouped into 5 behavioral signals and "
        "serialized as structured prose. The text format is designed to carry semantic "
        "meaning that the embedding model can encode — 'restricted file access ratio "
        "0.45' means something different from '0.01' in 1536-d semantic space."
    ))

    create_table(doc,
        ["Signal", "Features Used", "Text Template"],
        [
            ["auth", "auth_total, auth_fail_rate,\nauth_off_hours_ratio,\n"
             "auth_unique_sources,\nauth_unique_dests",
             "\"User {uid} auth: {auth_total} events, {auth_fail_rate} fail rate, "
             "{auth_off_hours_ratio} off-hours, {auth_unique_sources} sources, "
             "{auth_unique_dests} destinations\""],
            ["privilege", "auth_methods_used,\nendpoint_max_risk",
             "\"User {uid} privilege: {auth_methods_used} methods, "
             "{endpoint_max_risk} max risk\""],
            ["data_access", "file_total, file_restricted_ratio,\n"
             "file_confidential_ratio,\nfile_unique_paths, file_total_bytes",
             "\"User {uid} files: {file_total} accesses, {file_restricted_ratio} "
             "restricted, {file_confidential_ratio} confidential, {file_unique_paths} "
             "paths, {file_total_bytes} bytes\""],
            ["network", "net_bytes_out,\nnet_unique_dsts",
             "\"User {uid} network: {net_bytes_out} bytes out, "
             "{net_unique_dsts} destinations\""],
            ["communication", "endpoint_total,\nendpoint_suspicious_ratio,\n"
             "endpoint_unique_processes",
             "\"User {uid} endpoint: {endpoint_total} events, "
             "{endpoint_suspicious_ratio} suspicious, "
             "{endpoint_unique_processes} processes\""],
        ],
        col_widths=[1.0, 1.8, 3.8],
    )

    add_section_heading(doc, "5.2. Embedding (Prose → 1536-d Vector)", level=2)

    create_table(doc,
        ["Parameter", "Value"],
        [
            ["Embedding Model", "OpenAI text-embedding-3-small"],
            ["Dimension", "1536"],
            ["Embedding Calls", "5 signals × 250 users × 70 weeks = 87,500 calls"],
            ["+ 12 reference concepts", "23,762 total embeddings"],
            ["Caching", "SHA-256 hash of text → .npy file in data/embedding_cache/"],
            ["Cache hit rate (rerun)", "100% (0 API calls, 23,762 cache hits)"],
            ["Vector normalization", "L2-normalized (unit vectors) by OpenAI API"],
        ],
        col_widths=[2.5, 4.0],
    )

    add_section_heading(doc, "5.3. Signal Composition (5 Signals → 1 Composite)", level=2)

    add_body_text(doc, (
        "The 5 signal vectors are combined via weighted average into a single composite "
        "vector per user per week. Source: embeddings/composer.py → compose()."
    ))

    create_table(doc,
        ["Signal", "Weight", "Rationale"],
        [
            ["auth (authentication)", "0.25", "Primary identity signal — most discriminating"],
            ["privilege", "0.20", "Escalation detection"],
            ["data_access", "0.20", "Core insider threat signal"],
            ["network", "0.20", "C2 and exfiltration detection"],
            ["communication (endpoint)", "0.15", "Supporting behavioral context"],
        ],
        col_widths=[2.0, 0.8, 3.7],
    )

    add_body_text(doc, "Composition formula:", bold=True)
    add_code_block(doc,
        "composite = Σ(weight_i × signal_i) / Σ(weight_i)\n"
        "composite = composite / ‖composite‖₂        # L2 normalize to unit vector\n"
        "\n"
        "Result: one 1536-d unit vector per user per week\n"
        "Total composites: 250 users × 70 weeks = 17,500 vectors"
    )

    add_body_text(doc, (
        "Missing signal handling: If fewer than 5 signals are available, the formula "
        "normalizes by the sum of available weights, ensuring the composite remains a "
        "valid unit vector."
    ))

    add_section_heading(doc, "5.4. Three-Layer Detection", level=2)

    add_body_text(doc, "LAYER 1: Magnitude CUSUM", bold=True)
    add_body_text(doc, (
        "Computes cosine distance between consecutive weekly composites, then applies "
        "CUSUM change-point detection on the drift magnitude series."
    ))
    create_table(doc,
        ["Parameter", "Value"],
        [
            ["Drift metric", "1.0 − cosine_similarity(composite[week_i], composite[week_i+1])"],
            ["CUSUM threshold", "0.001"],
            ["CUSUM drift_threshold", "0.0005"],
            ["min_run_length", "2"],
            ["Observation", "With real embeddings, ALL users trigger magnitude CUSUM because "
             "even tiny drift (0.002–0.007) accumulates above 0.001 over 17 weeks. "
             "Magnitude alone does not discriminate. This layer is recorded but not used "
             "for detection."],
        ],
        col_widths=[2.0, 4.5],
    )

    add_body_text(doc, "LAYER 2: Overall Drift Direction", bold=True)
    add_body_text(doc, (
        "Compares the baseline behavioral center to the recent behavioral center "
        "by computing the drift vector and projecting it onto 12 reference concepts."
    ))
    create_table(doc,
        ["Parameter", "Value"],
        [
            ["Baseline vectors", "First half of weekly composites (weeks 0–34)"],
            ["Recent vectors", "Last quarter of composites (last ~17 weeks)"],
            ["Baseline centroid", "Mean of baseline vectors, L2-normalized"],
            ["Recent centroid", "Mean of recent vectors, L2-normalized"],
            ["Drift vector", "drift_vector(baseline_centroid, recent_centroid) → unit vector"],
            ["Projection", "cosine_similarity(drift_vector, concept_vector) for all 12 concepts"],
            ["Metrics recorded", "overall_threat_align, overall_benign_align, primary_threat"],
        ],
        col_widths=[2.0, 4.5],
    )

    add_body_text(doc, "LAYER 3: Weekly Direction Consistency (Primary Detection)", bold=True)
    add_body_text(doc, (
        "For each consecutive weekly pair, computes the drift vector and measures "
        "alignment with threat vs. benign concepts. The fraction of weeks where "
        "threat alignment exceeds benign alignment is the 'threat consistency' metric."
    ))
    create_table(doc,
        ["Parameter", "Value"],
        [
            ["Per-week drift vector", "drift_vector(composite[i], composite[i+1])"],
            ["Per-week max_threat", "max(cosine_sim(drift_vec, threat_concept)) across 10 threat concepts"],
            ["Per-week max_benign", "max(cosine_sim(drift_vec, benign_concept)) across 2 benign concepts"],
            ["Threat-winning week", "Week where max_threat − max_benign > 0.05"],
            ["threat_consistency", "threat_winning_weeks / total_weeks"],
            ["Detection threshold", "threat_consistency ≥ 0.40"],
            ["CUSUM on net threat", "CUSUM on (max_threat − max_benign) series, threshold=0.5, drift=0.05, run=2"],
        ],
        col_widths=[2.0, 4.5],
    )

    add_section_heading(doc, "5.5. Reference Concepts (12 Concepts)", level=2)

    add_body_text(doc, (
        "12 reference concept descriptions are embedded into the same 1536-d space. "
        "Drift vectors are projected onto these to classify the nature of behavioral change. "
        "Source: detection/reference_concepts.py."
    ))

    create_table(doc,
        ["#", "Concept Name", "Category", "Severity", "MITRE Techniques"],
        [
            ["1", "compromised_endpoint", "Threat", "Critical", "T1059, T1547, T1562, T1071"],
            ["2", "data_exfiltration", "Threat", "Critical", "T1005, T1074, T1048, T1567"],
            ["3", "privilege_escalation", "Threat", "High", "T1078, T1068, T1134, T1548"],
            ["4", "lateral_movement", "Threat", "High", "T1021, T1570, T1550, T1072"],
            ["5", "insider_threat_slow", "Threat", "High", "T1078, T1083, T1005, T1052"],
            ["6", "insider_threat_fast", "Threat", "Critical", "T1005, T1052, T1048, T1567"],
            ["7", "credential_stuffing", "Threat", "High", "T1110, T1078"],
            ["8", "c2_beacon", "Threat", "Critical", "T1071, T1573, T1568, T1102"],
            ["9", "reconnaissance", "Threat", "Medium", "T1046, T1018, T1087, T1135"],
            ["10", "supply_chain_compromise", "Threat", "Critical", "T1195, T1059, T1071"],
            ["11", "normal_role_change", "Benign", "Low", "(none)"],
            ["12", "seasonal_variation", "Benign", "Low", "(none)"],
        ],
        col_widths=[0.3, 1.8, 0.7, 0.7, 2.0],
    )

    add_section_heading(doc, "5.6. Detection Criteria", level=2)

    create_table(doc,
        ["Detection Method", "Criterion", "Used for Final Result?"],
        [
            ["Magnitude CUSUM", "cusum_detect(drift_magnitudes) triggers", "No — 100% FP with real embeddings"],
            ["V-Intelligence UEBA Direction", "threat_consistency ≥ 0.40", "Yes — primary detection"],
            ["V-Intelligence UEBA Top-10%", "threat_consistency ≥ 90th percentile", "Yes — alternate threshold"],
            ["IForest + V-Intelligence UEBA", "iforest_anomaly OR acecard_direction_detected", "Yes — combined ensemble"],
        ],
        col_widths=[1.8, 2.5, 2.2],
    )

    add_section_heading(doc, "5.7. Tier 2 Results", level=2)

    create_table(doc,
        ["Method", "USR-156\n(Insider)", "USR-234\n(Slow APT)", "USR-042\n(Volt Typh)",
         "USR-118\n(Salt Typh)", "True\nPos", "False\nPos", "FP Rate"],
        [
            ["V-Intelligence UEBA Direction", "MISSED", "DETECTED", "DETECTED", "MISSED", "2", "37", "15.0%"],
            ["IForest + V-Intelligence UEBA", "MISSED", "DETECTED", "DETECTED", "MISSED", "2", "49", "19.9%"],
        ],
        col_widths=[1.3, 0.8, 0.8, 0.8, 0.8, 0.5, 0.5, 0.6],
    )

    add_body_text(doc, (
        "Tier 2 finding: V-Intelligence UEBA Direction detects USR-234 (slow APT) and USR-042 "
        "(Volt Typhoon) — the two attacks with clear embedding drift signatures. However, "
        "it misses USR-156 (insider) and USR-118 (Salt Typhoon) because the single-composite "
        "architecture dilutes zone-specific signals. The IForest + V-Intelligence UEBA combined ensemble "
        "achieves the same detections at higher FP (19.9%) due to IForest's false positives "
        "adding to V-Intelligence UEBA's."
    ), bold=True)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  6. TIER 3: DIGITAL ENTITY DETECTION
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "6. Tier 3: Digital Entity Detection", level=1)

    add_body_text(doc, (
        "Tier 3 decomposes the single composite embedding from Tier 2 into 5 hierarchical "
        "behavioral zones, computes zone-specific trajectories, builds cross-entity "
        "relationship embeddings, and runs 9 detection methods. This addresses Tier 2's "
        "core limitation: a single 1536-d composite averages out zone-specific signals "
        "that reveal attack type. Source: comparison/run_tier3.py."
    ))

    add_section_heading(doc, "6.1. Hierarchical Zone Embeddings (5 Zones)", level=2)

    add_body_text(doc, (
        "Each user's 23 features are partitioned into 5 behavioral zones. Each zone "
        "is independently serialized as prose text and embedded into a 1536-d vector. "
        "Source: models/hierarchical_zones.py → serialize_zone(), build_zone_embeddings()."
    ))

    create_table(doc,
        ["Zone", "Features", "Detects"],
        [
            ["identity", "role, department, clearance, tenure, user_type",
             "Role change (benign)"],
            ["access_pattern", "auth_total, auth_fail_rate, auth_off_hours_ratio,\n"
             "auth_unique_sources, auth_unique_dests, auth_methods_used",
             "Credential abuse"],
            ["data_behavior", "file_total, file_restricted_ratio,\n"
             "file_confidential_ratio, file_write_ratio,\n"
             "file_unique_paths, file_total_bytes",
             "Insider exfiltration"],
            ["network_footprint", "net_bytes_out, net_unique_dsts,\n"
             "net_external_ratio, dns_unique_domains,\n"
             "dns_nxdomain_ratio",
             "C2 / APT"],
            ["risk_posture", "endpoint_suspicious_ratio,\n"
             "endpoint_max_risk, endpoint_mean_risk,\n"
             "endpoint_unique_processes, endpoint_total",
             "Endpoint compromise"],
        ],
        col_widths=[1.2, 2.5, 2.8],
    )

    add_body_text(doc, (
        "Zone serialization example (data_behavior zone):"
    ), bold=True)
    add_code_block(doc,
        "\"User USR-156 data_behavior zone: file_total=342, \"\n"
        "\"file_restricted_ratio=0.18, file_confidential_ratio=0.12, \"\n"
        "\"file_write_ratio=0.35, file_unique_paths=87, file_total_bytes=4521890\"\n"
        "\n"
        "→ embed(text) → 1536-d vector representing data access behavior"
    )

    add_body_text(doc, (
        "Embedding volume: 5 zones × 250 users × 70 weeks = 87,500 zone embeddings. "
        "With 12 reference concepts = 23,762 total. All cached via SHA-256 hash."
    ))

    add_section_heading(doc, "6.2. Context-Adaptive Composition (4 Scenarios)", level=2)

    add_body_text(doc, (
        "The 5 zone embeddings are composed into a single composite via softmax-weighted "
        "attention. Weights change based on investigation context, amplifying different "
        "behavioral dimensions for different threat hunts. "
        "Source: models/hierarchical_zones.py → compose_zones()."
    ))

    create_table(doc,
        ["Context", "identity", "access", "data", "network", "risk", "Use Case"],
        [
            ["normal_ops", "0.20", "0.20", "0.20", "0.20", "0.20", "Default monitoring"],
            ["insider_investigation", "0.10", "0.15", "0.40", "0.15", "0.20",
             "Data exfil hunt"],
            ["apt_hunt", "0.05", "0.10", "0.10", "0.40", "0.30",
             "C2/network hunt"],
            ["privilege_audit", "0.10", "0.25", "0.10", "0.15", "0.40",
             "Escalation audit"],
        ],
        col_widths=[1.3, 0.6, 0.6, 0.6, 0.7, 0.5, 1.3],
    )

    add_body_text(doc, "Composition formula:", bold=True)
    add_code_block(doc,
        "softmax_weights = softmax(context_weights)  # normalize to sum=1\n"
        "composite = Σ(softmax_weight_i × zone_embedding_i)\n"
        "composite = composite / ‖composite‖₂         # L2 normalize"
    )

    add_section_heading(doc, "6.3. Temporal Trajectory (Velocity/Acceleration/Regime)", level=2)

    add_body_text(doc, (
        "For each entity, the weekly composite embedding snapshots are analyzed to extract "
        "6 trajectory features. Per-zone trajectories are also computed to detect "
        "zone-specific drift. Source: models/temporal_trajectory.py."
    ))

    create_table(doc,
        ["Feature", "Computation", "Interpretation"],
        [
            ["velocity_magnitude", "‖last_composite − first_composite‖₂",
             "Total behavioral displacement"],
            ["acceleration", "Slope of consecutive cosine distance series",
             "Speeding up or slowing down"],
            ["stability", "Mean cosine_sim(composite[i], composite[i+1])",
             "Low = volatile; high = stable"],
            ["regime_shifts", "Fraction of pairs with cosine_sim < 0.7",
             "Fundamental phase changes"],
            ["trend_consistency", "Kendall τ of cosine distance from baseline",
             "Monotonic drift pattern"],
            ["total_drift", "Sum of all consecutive cosine distances",
             "Cumulative path length"],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    add_body_text(doc, "Per-zone trajectory analysis:", bold=True)
    add_body_text(doc, (
        "The same 6 trajectory features are computed independently for each of the 5 "
        "zone embedding series. Zone-specific drift reveals which behavioral dimension "
        "is changing: identity stable + data_behavior drifting = insider threat; "
        "identity stable + network_footprint drifting = APT/C2."
    ))

    add_section_heading(doc, "6.4. Relationship Embeddings (Hadamard Products)", level=2)

    add_body_text(doc, (
        "Cross-entity relationships are captured via Hadamard (element-wise) products "
        "of composite embeddings. The resulting vector encodes the interaction pattern "
        "between two entities. Source: models/relationship_embeddings.py."
    ))

    add_body_text(doc, "Hadamard formula:", bold=True)
    add_code_block(doc,
        "relationship_vector = (user_composite ⊙ device_composite) / \n"
        "                      ‖user_composite ⊙ device_composite‖₂\n"
        "\n"
        "Relationship types:\n"
        "  UserDevice  = user_composite ⊙ device_composite\n"
        "  UserApp     = user_composite ⊙ app_composite\n"
        "  DeviceSegment = device_composite ⊙ segment_composite"
    )

    add_body_text(doc, (
        "When a C2 beacon changes how a user interacts with their device, the "
        "relationship vector drifts even if neither entity individually changes. "
        "Drift in relationship embeddings is measured as cosine distance between "
        "baseline and recent relationship vectors."
    ))

    add_section_heading(doc, "6.5. Tier 3 Detection Methods (9 Methods)", level=2)

    add_body_text(doc, (
        "Each method receives specific inputs from the zone/trajectory/relationship "
        "pipeline and applies independent detection logic. All methods use rank-normalized "
        "scoring and flag the top 10% by anomaly score (90th percentile threshold)."
    ))

    add_body_text(doc, "Method 1: Velocity/Acceleration", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Input", "Weekly composite embeddings (1536-d)", "Attention-weighted 5-zone composition"],
            ["velocity_magnitude", "Avg L2 norm of consecutive diffs", "Rate of behavioral displacement"],
            ["acceleration", "Mean change in velocity magnitude", "Whether drift is speeding up"],
            ["total_drift", "1 − cosine_sim(first, last)", "Overall start-to-end displacement"],
            ["Score formula", "0.4×velocity + 0.3×acceleration + 0.3×total_drift", "Rank-normalized components"],
            ["Detection", "Top 10% by composite score", "90th percentile threshold"],
            ["acceleration_threshold", "0.01", "Minimum acceleration for 'accelerating' regime"],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    add_body_text(doc, "Method 2: Regime Shift", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Input", "Weekly composite embeddings (consecutive pairs)", "Same as velocity input"],
            ["regime_shift_threshold", "0.7", "Cosine sim below this = phase change"],
            ["regime_shifts", "Fraction of pairs below 0.7", "0 = no shifts, 1 = every pair shifts"],
            ["stability", "Mean consecutive cosine sim", "1.0 = stable, <0.7 = volatile"],
            ["Score formula", "1 − rank_normalize(stability)", "Lower stability = higher anomaly"],
            ["Detection", "Top 10% by regime score", "90th percentile threshold"],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    add_body_text(doc, "Method 3: Zone Divergence", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Input", "Per-zone trajectory drift (5 zones)", "identity, access, data, network, risk drift"],
            ["zone_stable_threshold", "0.02", "Identity drift below this = stable profile"],
            ["zone_drifting_threshold", "0.05", "Behavioral zone above this = drifting"],
            ["divergence_score", "max(access,data,net,risk) − identity", "Measures independent zone motion"],
            ["Detection", "Top 10% by divergence score", "Standalone-sufficient in T3 Combined"],
            ["Key insight", "Identity stable + data drifting = insider", "Zone decomposition reveals attack type"],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    add_body_text(doc, "Method 4: Relationship Drift", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Input", "UserDevice Hadamard vectors per week", "user_emb ⊙ device_emb, L2-normalized"],
            ["relationship_drift_threshold", "0.05", "Minimum drift to flag"],
            ["Drift metric", "1 − cosine_sim(rel_baseline, rel_recent)", "Measures interaction pattern change"],
            ["Detection", "Top 10% by relationship drift", "90th percentile threshold"],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    add_body_text(doc, "Method 5: Contextual (Multi-Context)", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Input", "4 context-reweighted composites per user/week", "normal_ops, insider, apt_hunt, priv_audit"],
            ["Per-week threat alignment", "cosine_sim(drift, threat_concept)", "Against 10 threat + 2 benign concepts"],
            ["Threat week criterion", "max_threat − max_benign > 0.05", "Week where drift is threat-aligned"],
            ["threat_consistency", "threat_weeks / total_weeks", "Fraction of threat-winning weeks"],
            ["contextual_threat_threshold", "0.30", "Minimum consistency for detection"],
            ["Detection", "Top 10% by best consistency across 4 contexts", "Best-of-4 contexts reported"],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    add_body_text(doc, "Method 6: Cross-Entity Correlation", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Input", "User relationship drift vectors", "All UserDevice Hadamard drifts"],
            ["cohort_similarity", "0.5", "Cosine sim threshold for clustering"],
            ["cohort_min_size", "3", "Minimum users in a correlated cohort"],
            ["Detection", "Member of any detected cohort", "Coordinated drift = coordinated attack"],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    add_body_text(doc, "Method 7: Embedding CUSUM", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Input", "Weekly composites vs 4-week baseline", "Baseline: mean of first 4 weeks, L2-normalized"],
            ["Distance metric", "1 − cosine_sim(baseline, week)", "Per-week deviation from baseline"],
            ["CUSUM formula", "max(0, cusum + dist − baseline_mean_dist)", "Accumulates excess drift"],
            ["Score formula", "0.6×max_cusum + 0.4×final_cusum", "Rank-normalized components"],
            ["Detection", "Top 10% by composite CUSUM score", "90th percentile threshold"],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    add_body_text(doc, "Method 8: Per-Zone Threat Direction", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Input", "Per-zone snapshots (≥5 weeks) + 12 reference concepts", "Zone-specific drift vectors"],
            ["Zone-threat map", "data→[exfil, insider], net→[c2, lateral],\n"
             "access→[cred_stuff, recon], risk→[compromised, priv_esc]", "Relevant threats per zone"],
            ["Relevant weight", "1.5", "Amplifies expected zone-threat alignment"],
            ["Other weight", "0.5", "De-emphasizes cross-zone threats"],
            ["Threat week criterion", "max_threat − max_benign > 0.02", "Lower threshold than Contextual"],
            ["Detection", "Top 10% by best zone consistency", "High recall but 50.8% FP"],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    add_body_text(doc, "Method 9: Behavioral Progression", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Input", "Per-zone weekly threat alignment scores (≥6 weeks)", "Same zone-threat map as Method 8"],
            ["Trend metric", "Kendall's τ rank correlation over time", "Detects monotonic increase in threat"],
            ["best_tau", "Highest τ across all zones", "τ > 0 = escalating threat alignment"],
            ["late_threat_mean", "Mean threat alignment in final 25% of weeks", "Late-stage threat signal"],
            ["Score formula", "0.6×best_tau + 0.4×late_threat_mean", "Rank-normalized components"],
            ["Detection", "Top 10% by composite score", "90th percentile threshold"],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    add_body_text(doc, "Tier 3 Combined Detection (Ensemble):", bold=True)
    add_body_text(doc, (
        "The Combined method ensembles all 6 core discriminating methods using a weighted "
        "geometric mean, with Zone Divergence as a standalone-sufficient override."
    ))
    create_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Core methods (6)", "Velocity, Zone Div, Contextual,\n"
             "CUSUM, Progression, Regime Shift", "Excludes Relationship, Cross-Entity, Zone Threat Dir"],
            ["Geometric weights", "velocity=0.10, zone_div=0.25,\n"
             "contextual=0.10, cusum=0.25, progression=0.30", "Zone Div + CUSUM weighted highest"],
            ["Final score", "0.6×geometric_composite + 0.4×core_count", "Rank-normalized"],
            ["Detection rule", "≥2 core methods agree\n"
             "OR zone_divergence alone\n"
             "OR composite in top 10%", "Zone Divergence standalone-sufficient"],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    add_section_heading(doc, "6.5.1. Tier 3 Configuration Summary", level=2)

    create_table(doc,
        ["Parameter", "Value", "Used By"],
        [
            ["acceleration_threshold", "0.01", "Velocity/Acceleration, Regime classification"],
            ["trend_consistency_min", "0.5", "Velocity/Acceleration regime classification"],
            ["regime_shift_threshold", "0.7", "Regime Shift (cosine sim cutoff)"],
            ["zone_stable_threshold", "0.02", "Zone Divergence (identity drift)"],
            ["zone_drifting_threshold", "0.05", "Zone Divergence (behavioral drift)"],
            ["relationship_drift_threshold", "0.05", "Relationship Drift"],
            ["contextual_threat_threshold", "0.30", "Contextual Detection"],
            ["threat_consistency_threshold", "0.40", "Contextual + Progression"],
            ["cohort_similarity", "0.5", "Cross-Entity Correlation"],
            ["cohort_min_size", "3", "Cross-Entity Correlation"],
            ["detection_percentile", "90th (top 10%)", "All methods (rank-based threshold)"],
        ],
        col_widths=[2.2, 1.5, 2.8],
    )

    add_section_heading(doc, "6.6. Tier 3 Results", level=2)

    create_table(doc,
        ["Method", "USR-156\n(Insider)", "USR-234\n(Slow APT)", "USR-042\n(Volt Typh)",
         "USR-118\n(Salt Typh)", "True\nPos", "False\nPos", "FP Rate"],
        [
            ["T3 Velocity/Accel", "DETECTED", "MISSED", "DETECTED", "MISSED", "2", "23", "9.3%"],
            ["T3 Regime Shift", "MISSED", "MISSED", "DETECTED", "MISSED", "1", "24", "9.8%"],
            ["T3 Zone Divergence", "DETECTED", "MISSED", "MISSED", "MISSED", "1", "24", "9.8%"],
            ["T3 Relationship", "MISSED", "MISSED", "DETECTED", "MISSED", "1", "24", "9.8%"],
            ["T3 Contextual", "MISSED", "DETECTED", "DETECTED", "MISSED", "2", "33", "13.4%"],
            ["T3 Cross-Entity", "—", "—", "—", "—", "—", "30", "—"],
            ["T3 Embed CUSUM", "DETECTED", "MISSED", "DETECTED", "DETECTED", "3", "22", "8.9%"],
            ["T3 Zone Threat Dir", "DETECTED", "DETECTED", "MISSED", "DETECTED", "3", "125", "50.8%"],
            ["T3 Beh Progression", "DETECTED", "DETECTED", "MISSED", "DETECTED", "3", "22", "8.9%"],
            ["T3 Combined", "DETECTED", "DETECTED", "DETECTED", "DETECTED", "4", "37", "15.0%"],
        ],
        col_widths=[1.3, 0.8, 0.8, 0.8, 0.8, 0.5, 0.5, 0.6],
    )

    add_body_text(doc, (
        "Tier 3 key finding: embedding/composite scoring cleanly separates only 2 of 4 — "
        "USR-156 (insider) and USR-118 (Salt Typhoon) rank above all normal users, while "
        "USR-234 (slow APT, #7) and USR-042 (Volt Typhoon LOTL, #24) fall BELOW normal users "
        "on the composite. The clean 4/4 at 0 false positives is achieved by the separate "
        "multi-front threat-profile detector (threat_profile_detector.py), which scores "
        "measurable known-bad profiles — C2-beacon, DGA-DNS, LOTL-process, cohort-rare access, "
        "recon-fanout, and insider-collection — using cohort-relative and raw-event features, "
        "label-free. Zone Divergence uniquely detects the insider threat via zone-specific "
        "decomposition."
    ), bold=True)

    add_body_text(doc, "Zone drift detail for attack users:", bold=True)

    create_table(doc,
        ["User", "Attack", "identity\ndrift", "data_behavior\ndrift",
         "network\ndrift", "risk\ndrift", "Detection Signature"],
        [
            ["USR-156", "Insider", "0.0000", "0.2586", "0.0121", "0.0050",
             "data_behavior ↑ toward insider_threat_fast"],
            ["USR-234", "Slow APT", "0.0000", "0.0062", "0.1367", "0.0058",
             "network_footprint ↑ identity stable"],
            ["USR-042", "Volt Typhoon", "0.0000", "0.1745", "0.1445", "0.0054",
             "dual data + network drift"],
            ["USR-118", "Salt Typhoon", "0.0000", "0.1540", "0.1786", "0.0036",
             "dual data + network drift"],
        ],
        col_widths=[0.7, 0.8, 0.7, 0.9, 0.7, 0.5, 2.0],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  7. OPTIMAL ENSEMBLE
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "7. Optimal Detection Configuration", level=1)

    add_body_text(doc, (
        "With 250 users and amplified attack signals, the multi-front threat-profile detector "
        "(threat_profile_detector.py) is the optimal detection configuration — it is the method "
        "that cleanly detects all 4 attacks at 0 false positives by scoring measurable "
        "known-bad profiles (C2-beacon, DGA-DNS, LOTL-process, cohort-rare access, recon-fanout, "
        "insider-collection). Embedding/composite scoring on its own cleanly separates only 2 of "
        "4 (USR-156 and USR-118 above all normals; USR-234 and USR-042 below normal users). Note "
        "that the Feature CUSUM signal fires on roughly 99% of normal users, so its apparent "
        "coverage is not operationally usable."
    ))

    create_table(doc,
        ["Method", "USR-156\n(Insider)", "USR-234\n(Slow APT)", "USR-042\n(Volt Typh)",
         "USR-118\n(Salt Typh)", "False\nPos", "FP Rate"],
        [
            ["Feature CUSUM (Tier 1)", "DETECTED", "MISSED", "DETECTED", "DETECTED", "22", "8.9%"],
            ["V-Intelligence UEBA Direction (Tier 2)", "MISSED", "DETECTED", "DETECTED", "MISSED", "37", "15.0%"],
            ["T3 Combined (Tier 3)", "DETECTED", "DETECTED", "DETECTED", "DETECTED", "37", "15.0%"],
            ["Feat CUSUM + V-Intelligence UEBA Dir", "DETECTED", "DETECTED", "DETECTED", "DETECTED", "~55", "~22%"],
        ],
        col_widths=[1.5, 0.8, 0.8, 0.8, 0.8, 0.5, 0.6],
    )

    add_body_text(doc, (
        "Across all 17 methods evaluated, embedding/composite scoring cleanly separates only "
        "2 of 4 attack campaigns. The clean 4/4 detection at 0 false positives is achieved by "
        "the multi-front threat-profile detector (threat_profile_detector.py), which targets "
        "measurable known-bad profiles using cohort-relative and raw-event features, label-free."
    ), bold=True)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  8. CONSOLIDATED RESULTS TABLE
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "8. Consolidated Results Table (17 Methods × 4 Attacks)", level=1)

    create_table(doc,
        ["#", "Method", "Tier", "USR-156\n(Insider)", "USR-234\n(APT)",
         "USR-042\n(Volt T)", "USR-118\n(Salt T)", "True\nPos", "False\nPos", "FP\nRate"],
        [
            ["1", "Isolation Forest", "1", "MISSED", "MISSED", "MISSED", "MISSED",
             "0", "12", "4.9%"],
            ["2", "One-Class SVM", "1", "DETECTED", "MISSED", "MISSED", "DETECTED",
             "2", "26", "10.6%"],
            ["3", "LOF", "1", "MISSED", "MISSED", "MISSED", "DETECTED",
             "1", "11", "4.5%"],
            ["4", "Z-Score (|z|>3)", "1", "MISSED", "MISSED", "DETECTED", "DETECTED",
             "2", "25", "10.2%"],
            ["5", "Temporal Z-Score", "1", "DETECTED", "DETECTED", "DETECTED", "DETECTED",
             "4", "232", "94.3%"],
            ["6", "Feature CUSUM", "1", "DETECTED", "MISSED", "DETECTED", "DETECTED",
             "3", "22", "8.9%"],
            ["7", "V-Intelligence UEBA Direction", "2", "MISSED", "DETECTED", "DETECTED", "MISSED",
             "2", "37", "15.0%"],
            ["8", "IForest+V-Intelligence UEBA", "1+2", "MISSED", "DETECTED", "DETECTED", "MISSED",
             "2", "49", "19.9%"],
            ["9", "T3 Velocity", "3", "DETECTED", "MISSED", "DETECTED", "MISSED",
             "2", "23", "9.3%"],
            ["10", "T3 Regime Shift", "3", "MISSED", "MISSED", "DETECTED", "MISSED",
             "1", "24", "9.8%"],
            ["11", "T3 Zone Diverge", "3", "DETECTED", "MISSED", "MISSED", "MISSED",
             "1", "24", "9.8%"],
            ["12", "T3 Relationship", "3", "MISSED", "MISSED", "DETECTED", "MISSED",
             "1", "24", "9.8%"],
            ["13", "T3 Contextual", "3", "MISSED", "DETECTED", "DETECTED", "MISSED",
             "2", "33", "13.4%"],
            ["14", "T3 Embed CUSUM", "3", "DETECTED", "MISSED", "DETECTED", "DETECTED",
             "3", "22", "8.9%"],
            ["15", "T3 Zone Threat", "3", "DETECTED", "DETECTED", "MISSED", "DETECTED",
             "3", "125", "50.8%"],
            ["16", "T3 Beh Progress", "3", "DETECTED", "DETECTED", "MISSED", "DETECTED",
             "3", "22", "8.9%"],
            ["17", "T3 Combined", "3", "DETECTED", "DETECTED", "DETECTED", "DETECTED",
             "4", "37", "15.0%"],
        ],
        col_widths=[0.3, 1.2, 0.4, 0.6, 0.6, 0.6, 0.6, 0.4, 0.4, 0.5],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  9. SOURCE FILE REFERENCE
    # ══════════════════════════════════════════════════════════════
    # ══════════════════════════════════════════════════════════════
    #  9. UEBA FRAMEWORK AND THREAT TAXONOMY
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "9. UEBA Framework and Threat Taxonomy", level=1)

    add_section_heading(doc, "9.1. What is UEBA?", level=2)

    add_body_text(doc, (
        "User and Entity Behavior Analytics (UEBA) detects threats by establishing "
        "behavioral baselines for every user and device, then flagging when behavior "
        "drifts from that baseline. Unlike traditional SIEM (Security Information and "
        "Event Management) systems that match known attack signatures — '3 failed logins "
        "within 5 minutes = alert' — UEBA asks: 'Is this person behaving differently "
        "than they used to?'"
    ))

    add_body_text(doc, (
        "UEBA operates on three principles: (1) Every entity has a behavioral baseline "
        "learned from historical activity. (2) Behavioral deviation from baseline is the "
        "primary detection signal. (3) The deviation must be analyzed for direction, not "
        "just magnitude — what kind of behavioral change, not just how much change."
    ))

    add_body_text(doc, (
        "V-Intelligence UEBA implements UEBA through semantic embeddings: behavioral telemetry is "
        "serialized to natural language, embedded into 1536-dimensional vector space, "
        "and tracked over time. Drift in this space represents behavioral change. The "
        "direction of drift — toward known threat concepts or away from baseline patterns — "
        "determines whether the change is benign (role change, new project) or malicious "
        "(data exfiltration, C2 communication)."
    ))

    add_section_heading(doc, "9.2. Why SIEM Rules Fail Against Modern Threats", level=2)

    add_body_text(doc, (
        "Modern attackers specifically design their campaigns to evade rule-based detection. "
        "The 4 attack campaigns in this study demonstrate the gap:"
    ))

    create_table(doc,
        ["Threat Type", "SIEM Rule Gap", "UEBA Signal"],
        [
            ["Insider Threat\n(USR-156)",
             "All access uses valid credentials and\nauthorized permissions. No rule violation.",
             "Data access pattern drifts from\nbaseline over months."],
            ["Slow APT / C2\n(USR-234)",
             "C2 traffic is small, periodic, encrypted\n— looks like normal HTTPS.",
             "Network communication pattern drifts\n(periodic beacon signature)."],
            ["Nation-State LOTL\n(USR-042)",
             "No malware deployed. Uses PowerShell,\nWMI, RDP — tools admins use daily.",
             "Endpoint/process behavior drifts when\nlegitimate tools are repurposed."],
            ["Telecom Pivot\n(USR-118)",
             "Uses authorized maintenance credentials\nto access network infrastructure.",
             "Network footprint drifts as infrastructure\naccess patterns change."],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    add_body_text(doc, (
        "The common thread: the attacker's identity stays the same, but their behavior "
        "changes. Rule-based systems cannot detect what they're not looking for. UEBA "
        "detects the behavioral change itself, regardless of the specific attack technique."
    ), bold=True)

    add_section_heading(doc, "9.3. UEBA Threat Taxonomy (4 Classes)", level=2)

    add_body_text(doc, (
        "All 4 attack campaigns in this study are core UEBA use cases. Each produces "
        "a distinct type of behavioral drift that requires a different detection approach."
    ))

    add_section_heading(doc, "9.3.1. Insider Threat (USR-156)", level=3)
    add_body_text(doc, (
        "A trusted employee with legitimate access who gradually escalates privileges "
        "and data access over months. They don't break in — they're already inside. The "
        "attacker slowly moves from normal files to restricted/confidential documents, "
        "increases off-hours activity, and stages data for exfiltration. Each individual "
        "action looks routine; the threat is only visible in the cumulative behavioral "
        "direction over time."
    ))
    add_bullet(doc, "Campaign duration: ", bold_prefix="")
    add_body_text(doc, "8 months (32+ weeks). Behavioral escalation across 4 phases.")
    add_bullet(doc, "Real-world parallels: ", bold_prefix="")
    add_body_text(doc, (
        "Edward Snowden (NSA), Chelsea Manning (DOD), corporate IP theft cases where "
        "employees copy data before resignation."
    ))
    add_bullet(doc, "UEBA signature: ", bold_prefix="")
    add_body_text(doc, (
        "Identity zone stable (same role, same department, same clearance). Data_behavior "
        "zone drifting (accessing different data than historical pattern). The user looks "
        "the same but acts differently in the data dimension."
    ))

    add_section_heading(doc, "9.3.2. Slow APT — Advanced Persistent Threat (USR-234)", level=3)
    add_body_text(doc, (
        "A sophisticated, nation-state-level attack maintaining covert Command and Control "
        "(C2) communication for months. 'Advanced' = custom tools and techniques. "
        "'Persistent' = maintains access for weeks to months, not a smash-and-grab. "
        "'Threat' = the goal is espionage, data theft, or infrastructure sabotage. The C2 "
        "beacon sends small periodic signals to an external server, slowly staging and "
        "exfiltrating data."
    ))
    add_bullet(doc, "Campaign duration: ", bold_prefix="")
    add_body_text(doc, "180 days (26 weeks). Slow C2 establishment then data staging.")
    add_bullet(doc, "Real-world parallels: ", bold_prefix="")
    add_body_text(doc, (
        "SolarWinds/SUNBURST (Russia SVR, 9+ months), APT29/Cozy Bear, APT28/Fancy Bear, "
        "numerous government and defense contractor breaches."
    ))
    add_bullet(doc, "UEBA signature: ", bold_prefix="")
    add_body_text(doc, (
        "Identity zone stable. Network_footprint zone drifting — C2 beaconing creates a "
        "persistent network signature with periodic connections to external infrastructure, "
        "DNS anomalies, and gradually increasing outbound data volume."
    ))

    add_section_heading(doc, "9.3.3. Nation-State LOTL — Living-off-the-Land (USR-042)", level=3)
    add_body_text(doc, (
        "An attacker who avoids deploying malware entirely. Instead, they use legitimate "
        "tools already installed on the system — PowerShell, WMI, certutil, scheduled tasks, "
        "remote desktop — to move laterally and maintain access. 'Living off the land' means "
        "blending in with normal admin activity. No malware signatures to detect, no "
        "suspicious executables — just authorized tools used in unauthorized ways."
    ))
    add_bullet(doc, "Campaign duration: ", bold_prefix="")
    add_body_text(doc, "115 days (16+ weeks). Volt Typhoon (China) campaign pattern.")
    add_bullet(doc, "Real-world parallels: ", bold_prefix="")
    add_body_text(doc, (
        "Volt Typhoon (China, 2023-present) targeting US critical infrastructure — energy, "
        "water, telecom. CISA advisory AA23-144A. Also: Hafnium (Exchange servers), APT41."
    ))
    add_bullet(doc, "UEBA signature: ", bold_prefix="")
    add_body_text(doc, (
        "Uniform change across all behavioral zones — endpoint, network, and access patterns "
        "all shift when LOTL tools activate. Creates a strong before/after behavioral phase "
        "change rather than zone-specific drift."
    ))

    add_section_heading(doc, "9.3.4. Telecom Infrastructure Pivot (USR-118)", level=3)
    add_body_text(doc, (
        "An attacker who compromises telecommunications infrastructure — routers, switches, "
        "call processing systems, lawful intercept systems — to gain broad visibility into "
        "communications. Unlike typical data breaches targeting files, this attack pivots "
        "through network devices to intercept voice calls, text messages, and metadata at "
        "scale."
    ))
    add_bullet(doc, "Campaign duration: ", bold_prefix="")
    add_body_text(doc, "100 days (14+ weeks). Salt Typhoon (China) campaign pattern.")
    add_bullet(doc, "Real-world parallels: ", bold_prefix="")
    add_body_text(doc, (
        "Salt Typhoon (China, 2024) compromised AT&T, Verizon, T-Mobile, and Lumen "
        "Technologies. Accessed lawful intercept systems targeting senior US government officials."
    ))
    add_bullet(doc, "UEBA signature: ", bold_prefix="")
    add_body_text(doc, (
        "Broad multi-zone behavioral change — network footprint, data access, and endpoint "
        "patterns all shift during infrastructure pivot. Creates persistent cumulative drift "
        "and a clear phase change detectable by accumulation-based methods."
    ))

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  10. DETECTION PLAYBOOK
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "10. Detection Playbook: Threat Type → Algorithm", level=1)

    add_body_text(doc, (
        "No single detection algorithm catches all 4 threat types at viable false positive "
        "rates. The critical finding is that the threat type determines which algorithm to "
        "deploy. This section maps each threat class to its recommended detection method."
    ))

    add_section_heading(doc, "10.1. Per-Threat Recommendations", level=2)

    create_table(doc,
        ["Threat Type", "Recommended\nMethod", "FP\nRate", "Why It Works", "What Fails"],
        [
            ["Insider Threat\n(USR-156)",
             "Zone\nDivergence", "6.5%",
             "Only method that decomposes behavior\ninto zones. Detects 'identity stable +\ndata drifting.'",
             "LOF, IForest, Z-Score\n(measure magnitude,\nnot direction)"],
            ["Slow APT\n(USR-234)",
             "LOF + Zone\nDivergence", "0%*/\n6.5%",
             "LOF catches outlier network footprint.\nZone Div catches 'identity stable +\nnetwork drifting.'",
             "IForest (not sensitive\nto slow drift)"],
            ["Nation-State LOTL\n(USR-042)",
             "LOF +\nRegime Shift", "0%*/\n6.5%",
             "LOF catches endpoint/network anomaly.\nRegime Shift detects before/after\nphase break.",
             "Zone Divergence\n(uniform change,\nnot zone-specific)"],
            ["Telecom Pivot\n(USR-118)",
             "LOF +\nEmbed CUSUM", "0%*/\n6.5%",
             "LOF catches network footprint.\nEmbedding CUSUM catches persistent\n100-day cumulative drift.",
             "Zone Divergence\n(broad multi-zone\nchange)"],
        ],
        col_widths=[1.2, 0.9, 0.5, 2.2, 1.3],
    )

    add_body_text(doc, (
        "*FP rates reflect contamination=0.05 on 250-user synthetic dataset. "
        "Score-based thresholding planned for production deployment."
    ))

    add_section_heading(doc, "10.2. Layered Deployment Strategy", level=2)

    add_body_text(doc, (
        "Based on the per-threat analysis, a production UEBA deployment should layer "
        "detection methods by threat model:"
    ))

    add_bullet(doc, "Layer 1 — Traditional Magnitude Detection (LOF): ", bold_prefix="")
    add_body_text(doc, (
        "Catches attacks that change HOW MUCH a user does. Effective against LOTL, "
        "telecom pivot, slow APT. Low FP, fast computation. Deploy as always-on baseline."
    ))
    add_bullet(doc, "Layer 2 — Behavioral Direction Detection (Zone Divergence): ", bold_prefix="")
    add_body_text(doc, (
        "Catches attacks that change WHAT KIND of activity occurs. Essential for insider "
        "threats and slow APTs where magnitude is unchanged. Decomposes behavior into 5 "
        "zones and detects zone-specific drift."
    ))
    add_bullet(doc, "Layer 3 — Temporal Accumulation (Embedding CUSUM, Regime Shift): ", bold_prefix="")
    add_body_text(doc, (
        "Catches long-duration campaigns through cumulative drift and phase change detection. "
        "Essential for 100+ day campaigns where weekly drift is tiny but persistent."
    ))

    add_body_text(doc, (
        "The minimum viable deployment is the multi-front threat-profile detector "
        "(threat_profile_detector.py), the method that cleanly detects all 4 threats at 0 "
        "false positives via measurable known-bad profiles. Embedding/composite scoring on "
        "its own cleanly separates only 2 of 4. The Feature CUSUM signal fires on roughly 99% "
        "of normal users and is therefore not operationally usable."
    ), bold=True)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  11. SOURCE FILE REFERENCE
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "11. Source File Reference", level=1)

    create_table(doc,
        ["File", "Purpose", "Key Functions"],
        [
            ["comparison/run_comparison.py", "Tier 1+2 comparison pipeline",
             "engineer_weekly_features(), run_traditional_detection(), "
             "run_temporal_traditional(), run_feature_cusum(), "
             "run_acecard_detection(), build_comparison_report()"],
            ["comparison/run_tier3.py", "Tier 3 Digital Entity pipeline",
             "build_entity_zoo(), run_tier3_detection(), "
             "print_three_tier_comparison()"],
            ["models/cyber_entity.py", "Digital Entity dataclasses",
             "CyberEntity, PhaseState, Tier3Config"],
            ["models/hierarchical_zones.py", "Zone embeddings + composition",
             "serialize_zone(), build_zone_embeddings(), "
             "compose_zones(), ZONE_FEATURES, CONTEXT_WEIGHTS"],
            ["models/temporal_trajectory.py", "Velocity/acceleration/regime",
             "compute_trajectory_features(), detect_regime(), "
             "build_phase_state()"],
            ["models/relationship_embeddings.py", "Hadamard cross-entity",
             "hadamard(), compute_user_device_vector()"],
            ["embeddings/embedder.py", "OpenAI embedding client",
             "Embedder.embed_text(), embed_batch(), SHA-256 disk caching"],
            ["embeddings/composer.py", "Signal composition + drift math",
             "compose(), cosine_similarity(), drift_vector(), drift_magnitude()"],
            ["detection/cusum.py", "CUSUM change-point detection",
             "cusum_detect(), cusum_scan_entity(), batch_cusum_scan()"],
            ["detection/reference_concepts.py", "12 reference concept definitions",
             "ConceptLibrary, embed_concepts(), all_threat/benign_vectors()"],
            ["detection/drift_direction.py", "Drift direction analysis",
             "compute_drift_vector(), concept_alignment(), analyze_entity_drift()"],
            ["simulator/config.py", "Simulation parameters",
             "SIM_START, SIM_END, ATTACK_SCENARIOS (4 campaigns)"],
            ["data/tier3_results/", "Tier 3 output directory",
             "tier3_comparison.csv (250 users × 17 methods)"],
            ["data/embedding_cache/", "Embedding cache (SHA-256 → .npy)",
             "~47,500 cached embeddings, 100% hit rate on rerun"],
        ],
        col_widths=[2.0, 1.5, 3.0],
    )

    add_page_break(doc)

    # ── Section 12: Composite Scoring Deep Dive ────────────────
    add_section_heading(doc, "12. Multi-Phase Composite Scoring — Deep Dive", level=1)

    add_body_text(doc,
        "Traditional anomaly detection asks a single question: 'Is this value abnormal?' "
        "Composite scoring asks five questions simultaneously, then fuses the answers into "
        "a single ranked score per user. A normal user might score high on one phase by "
        "chance — an attacker scores high on multiple phases simultaneously.")

    add_section_heading(doc, "12.1. The Five Phases", level=2)

    create_table(doc,
        ["Phase", "What It Measures", "Formula", "Weight"],
        [
            ["Signal Strength",
             "How extreme are the top 3 z-scores vs the user's role peer group? "
             "Captures the magnitude of the strongest anomalies.",
             "sum(top 3 z-scores)", "1.0"],
            ["Breadth",
             "How many features exceed 1.5 standard deviations? A normal user might "
             "spike on one feature; attackers spike across many features simultaneously.",
             "count(features > 1.5σ)", "0.5"],
            ["Sustained Deviation",
             "Does the anomaly persist across weeks, or is it a one-time blip? "
             "Takes the top 2 zone-level sustained z-scores to measure persistence.",
             "sum(top 2 zone sustained z-scores)", "0.3"],
            ["Context Divergence",
             "Does the user look anomalous under multiple analytical contexts — "
             "insider investigation, APT hunt, privilege audit, normal ops? "
             "Measures spread and maximum across context-specific drift scores.",
             "max(ctx_spread_z, 0) + max(ctx_max_z, 0)", "0.5 + 0.3"],
            ["Novelty Persistence",
             "Are there novel behaviors (new external IPs, newly-registered domains) "
             "that keep recurring week after week? Persistent novel IPs are a strong "
             "indicator of C2 beacon infrastructure.",
             "min(persistence/5, 10) + weeks_frac×3", "1.0"],
        ],
        col_widths=[1.2, 2.5, 1.5, 0.6],
    )

    add_section_heading(doc, "12.2. Composite Formula", level=2)

    add_code_block(doc,
        "composite = signal_strength × 1.0\n"
        "          + breadth_15 × 0.5\n"
        "          + sustained_signal × 0.3\n"
        "          + max(ctx_spread_z, 0) × 0.5\n"
        "          + max(ctx_max_z, 0) × 0.3\n"
        "          + novelty_score × 1.0")

    add_body_text(doc,
        "Z-scores are computed per role peer group (developer, analyst, IT, executive, etc.) "
        "with attack users excluded from the baseline. This ensures that a developer's "
        "network activity is compared to other developers, not to executives with "
        "fundamentally different access patterns.")

    add_section_heading(doc, "12.3. Per-Attacker Composite Breakdown", level=2)

    create_table(doc,
        ["User", "Attack Type", "Composite", "Signal", "Breadth", "Sustained",
         "Ctx Spread", "Ctx Max", "Novelty", "Rank"],
        [
            ["USR-118", "Salt Typhoon", "51.3", "26.8", "9", "5.3",
             "5.7", "4.2", "0.0", "#1"],
            ["USR-156", "Insider Threat", "46.2", "23.1", "8", "6.8",
             "4.9", "3.4", "0.0", "#2"],
            ["USR-234", "Slow APT (C2)", "19.4", "4.5", "1", "1.6",
             "1.5", "0.8", "13.0", "#7"],
            ["USR-042", "Volt Typhoon", "13.7", "7.2", "3", "2.1",
             "2.3", "1.1", "0.0", "#24"],
        ],
        col_widths=[0.6, 1.0, 0.6, 0.5, 0.5, 0.6, 0.6, 0.5, 0.5, 0.4],
    )

    add_body_text(doc,
        "Key observations from the breakdown:")
    add_bullet(doc, "USR-118 (Salt Typhoon): ", bold_prefix="Rank #1. ")
    add_body_text(doc,
        "Dominates signal strength (26.8) and breadth (9 features above 1.5σ). "
        "The telecom infrastructure pivot creates anomalies across multiple behavioral "
        "zones simultaneously — authentication, network, and DNS patterns all shift. "
        "Despite this, no single feature exceeds the z>3 threshold, making it invisible "
        "to traditional single-threshold detection.", space_after=4)
    add_bullet(doc, "USR-156 (Insider): ", bold_prefix="Rank #2. ")
    add_body_text(doc,
        "Highest sustained deviation (6.8) — the 8-month escalation creates persistent "
        "anomalies that accumulate over time. The insider gradually shifts from public to "
        "restricted/confidential files while maintaining normal volume, creating sustained "
        "drift in the data_behavior zone.", space_after=4)
    add_bullet(doc, "USR-234 (Slow APT): ", bold_prefix="Rank #7. ")
    add_body_text(doc,
        "Without novelty persistence, this user's composite would be only 6.4 — barely "
        "noticeable. The novelty score of 13.0 is the sole reason for detection. "
        "See Section 12.4 for detailed analysis.", space_after=4)
    add_bullet(doc, "USR-042 (Volt Typhoon): ", bold_prefix="Rank #24. ")
    add_body_text(doc,
        "Living-off-the-land attacks use legitimate admin tools (PowerShell, wmic, certutil), "
        "creating moderate anomalies across context divergence. Detected at rank #24, "
        "still within the p90 threshold.", space_after=4)

    add_section_heading(doc, "12.4. USR-234 Novelty Persistence — Detailed Analysis", level=2)

    add_body_text(doc,
        "USR-234 simulates a Slow APT with command-and-control (C2) beaconing via DNS. "
        "The malware contacts an external C2 server approximately every 6 hours (~4 beacons/day) "
        "over 180 days. This low-and-slow pattern is invisible to volume-based detection: "
        "total DNS queries, network bytes, and authentication counts all remain within normal ranges.")

    add_body_text(doc, "Novelty persistence metrics for USR-234:", bold=True)

    create_table(doc,
        ["Metric", "Value", "Interpretation"],
        [
            ["persistent_novel_ips", "1",
             "One external IP that never appeared during the baseline period"],
            ["novel_ip_max_persistence", "60",
             "This IP appeared in 60 out of 60 post-baseline weeks (100%)"],
            ["novel_ip_weeks_frac", "1.00",
             "The novel IP is present in every single post-baseline week"],
        ],
        col_widths=[1.8, 0.6, 4.0],
    )

    add_body_text(doc, "Novelty score calculation:", bold=True)
    add_code_block(doc,
        "persistence = 60  (> threshold of 10)\n"
        "  → base_score = min(60 / 5.0, 10.0) = 10.0\n"
        "\n"
        "weeks_frac = 1.0  (> threshold of 0.5)\n"
        "  → bonus = 1.0 × 3.0 = 3.0\n"
        "\n"
        "novelty_score = 10.0 + 3.0 = 13.0")

    add_body_text(doc,
        "Without novelty persistence, USR-234's composite score would be 6.4 "
        "(signal=4.5 + breadth=0.5 + sustained=0.5 + context=1.0). "
        "At 6.4, this user would rank approximately #80 out of 250 — well below any "
        "detection threshold. Novelty persistence elevates the score to 19.4 (rank #7), "
        "placing it firmly within the detected set.")

    add_body_text(doc,
        "The detection insight: normal users occasionally contact new external IPs, "
        "but those contacts are transient — they appear once or twice then disappear. "
        "USR-234's novel IP appears every single week for over a year. This persistent, "
        "regular contact with previously-unseen infrastructure is the behavioral signature "
        "of C2 beaconing.")

    add_page_break(doc)

    # ── Section 13: Embedding vs Raw Feature Vector ────────────
    add_section_heading(doc, "13. Why Embeddings, Not Raw Feature Vectors?", level=1)

    add_body_text(doc,
        "A natural question arises: if we have 23 numeric features per user per week, "
        "why not represent each observation as a 23-dimensional vector and compute "
        "Euclidean distance directly? This section analyzes why semantic embeddings "
        "outperform raw feature vectors for behavioral anomaly detection.")

    add_section_heading(doc, "13.1. Two Pipelines: Feature-Space vs Semantic", level=2)

    add_body_text(doc, "Feature-Space CUSUM Pipeline (Traditional):", bold=True)
    add_code_block(doc,
        "Raw telemetry (auth, file, network, DNS, endpoint logs)\n"
        "  → Weekly aggregation (23 scalar features per user per week)\n"
        "    → Feature vector: v = [auth_total, auth_fail_rate, ..., dns_nxdomain_ratio]\n"
        "      → Statistical distance: ||v_week2 - v_week1|| (Euclidean/Mahalanobis)\n"
        "        → CUSUM accumulation of weekly distances\n"
        "          → Threshold detection")

    add_body_text(doc,
        "This pipeline operates entirely on numeric values. Each user-week is a point "
        "in 23-dimensional space. Drift is measured as the geometric distance between "
        "consecutive points. The CUSUM accumulates small drifts to catch slow changes "
        "that individual-week thresholds miss.")

    add_body_text(doc,
        "Result: All four attack users remain within the normal 5-95th percentile band. "
        "The method detects 0 of 4 attackers.", bold=True)

    add_body_text(doc, "V-Intelligence UEBA Semantic CUSUM Pipeline:", bold=True)
    add_code_block(doc,
        "Raw telemetry (auth, file, network, DNS, endpoint logs)\n"
        "  → Weekly aggregation (23 scalar features per user per week)\n"
        "    → Text serialization: features converted to natural language description\n"
        "      e.g., 'User USR-234, ML Engineer. Authentication: 45 total,\n"
        "            2.1% fail rate, 8.3% off-hours. File access: 30 total,\n"
        "            15% restricted, 8% confidential. Network: 2.1MB out,\n"
        "            28 unique destinations, 18% external...'\n"
        "      → Embedding: text → OpenAI text-embedding-3-small → 1536-d vector\n"
        "        → Cosine distance between consecutive weekly embeddings\n"
        "          → CUSUM accumulation of semantic drift\n"
        "            → Threshold detection")

    add_body_text(doc,
        "The critical addition is the text serialization step. Raw numbers are converted "
        "into prose that qualifies each metric: not just 'file_total=30' but "
        "'30 total, 15% restricted, 8% confidential.' The embedding model reads this "
        "text and produces a vector that captures the meaning of the behavioral pattern, "
        "not just the magnitude of individual metrics.")

    add_body_text(doc,
        "Result: the semantic pipeline helps, but its benefit is attack-dependent rather than "
        "universal. Embedding CUSUM fires roughly 30 weeks earlier for the insider (USR-156) "
        "and the LOTL attack (USR-042), fires LATER than feature-space CUSUM for the "
        "volume-driven attack (USR-118), and never separates the slow-APT (USR-234). It is not "
        "the case that all four attackers break out above the normal band on semantic CUSUM.",
        bold=True)

    add_section_heading(doc, "13.2. The Raw Vector Approach (Detailed)", level=2)

    add_body_text(doc,
        "Given a feature vector v = [auth_total, auth_fail_rate, auth_off_hours_ratio, "
        "file_total, file_restricted_ratio, ..., dns_nxdomain_ratio], we can compute "
        "drift as ||v_week2 - v_week1|| (Euclidean distance) or cosine distance. "
        "This is computationally simpler and requires no external API calls.")

    add_body_text(doc,
        "This is precisely what the Feature-Space CUSUM (Tier 1) computes. "
        "Empirical result: all four attack users remain within the normal band. "
        "The method detects 0 of 4 attackers.")

    add_section_heading(doc, "13.3. Three Fundamental Limitations of Raw Vectors", level=2)

    add_body_text(doc, "Limitation 1: Scale Dominance", bold=True)
    add_body_text(doc,
        "The 23 features span vastly different scales. net_bytes_out may be 2,100,000 "
        "while auth_fail_rate is 0.021. Even with z-score normalization, features with "
        "higher variance dominate the distance calculation. A 10% change in bytes_out "
        "swamps a 500% change in file_restricted_ratio — yet the latter is the primary "
        "insider threat signal for USR-156.")

    add_body_text(doc, "Limitation 2: Feature Independence Assumption", bold=True)
    add_body_text(doc,
        "Euclidean distance treats each feature as an independent dimension. It cannot "
        "represent that (auth_off_hours + file_restricted + high_clearance) together "
        "constitute a threat pattern, while each value individually falls within normal "
        "ranges. Embeddings capture these cross-feature semantic relationships because "
        "the serialized text reads: 'off-hours access to restricted files by a cleared "
        "employee' — the language model understands the combined meaning.")

    add_body_text(doc, "Limitation 3: The 'What vs How Much' Problem", bold=True)
    add_body_text(doc,
        "This is the critical limitation. Consider USR-156 (Insider Threat):")

    create_table(doc,
        ["Metric", "Week 1", "Week 2", "Raw Vector Distance"],
        [
            ["file_total", "30", "30", "0 (identical)"],
            ["file_restricted_ratio", "1.2%", "15.0%", "0.138 (small)"],
            ["file_confidential_ratio", "0.3%", "8.0%", "0.077 (small)"],
            ["Overall file_total distance", "", "", "0 — No drift detected"],
        ],
        col_widths=[1.8, 1.0, 1.0, 1.8],
    )

    add_body_text(doc,
        "The insider accesses the same number of files (30 both weeks), so file_total "
        "contributes zero distance. The restricted_ratio changes from 1.2% to 15% — "
        "a 12.5x increase — but in absolute terms, 0.138 is a small Euclidean contribution "
        "compared to other features. The raw vector says 'nothing changed.'")

    add_body_text(doc,
        "The embedding, however, is computed from text that says: 'File access: 30 total, "
        "1.2% restricted, 0.3% confidential' vs 'File access: 30 total, 15% restricted, "
        "8% confidential.' The language model understands that the user shifted from "
        "accessing mostly public files to predominantly restricted/confidential files — "
        "a semantically large change that produces significant embedding drift.")

    add_section_heading(doc, "13.4. What Embeddings Gain", level=2)

    create_table(doc,
        ["Capability", "Raw Vector", "Semantic Embedding"],
        [
            ["Cross-feature pattern recognition",
             "No — each dimension independent",
             "Yes — language model encodes correlations"],
            ["Qualitative feature encoding",
             "No — file paths, domain names, IP addresses cannot be represented",
             "Yes — text serialization captures qualitative attributes"],
            ["Scale-invariant comparison",
             "Requires explicit normalization",
             "Naturally scale-invariant (semantic similarity)"],
            ["'What' vs 'How Much' distinction",
             "No — same count = zero distance",
             "Yes — different types at same count = different meaning"],
            ["Computational cost",
             "O(1) — trivial distance calculation",
             "Requires API call to embedding model"],
            ["Numeric precision",
             "Exact — all values preserved",
             "Approximate — small numeric differences may be lost"],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    add_section_heading(doc, "13.5. What Embeddings Lose", level=2)

    add_body_text(doc,
        "Embeddings are not strictly superior. Two forms of information loss occur:")

    add_bullet(doc,
        "Precision loss: The embedding compresses 23 specific numbers into a 1536-dimensional "
        "semantic vector. Exact feature values cannot be extracted back. The direction of "
        "change is preserved, but precise magnitudes are smoothed.",
        bold_prefix="1. ")
    add_bullet(doc,
        "Small numeric differences: If the serialized text contains 'auth_fail_rate: 0.021' "
        "vs 'auth_fail_rate: 0.023', the language model may not distinguish these as "
        "meaningfully different. Very small numeric changes can be lost in the embedding.",
        bold_prefix="2. ")

    add_section_heading(doc, "13.6. The Design Decision: Both, Not Either", level=2)

    add_body_text(doc,
        "V-Intelligence UEBA uses both approaches in a complementary three-tier architecture:")

    create_table(doc,
        ["Tier", "Method", "Catches", "Misses"],
        [
            ["Tier 1 (Traditional)", "Raw feature vectors + statistical algorithms",
             "Volume-based attacks with obvious feature spikes",
             "Attacks at constant volume with behavioral changes"],
            ["Tier 2 (V-Intelligence UEBA Basic)", "Semantic embeddings + drift detection",
             "Behavioral changes invisible to raw features",
             "Slow APT with minimal semantic signature"],
            ["Tier 3 (Composite Scoring)", "Zone-decomposed z-scores + novelty persistence",
             "2 of 4 cleanly separated (USR-156, USR-118); clean 4/4 at 0 FP comes from the "
             "separate multi-front threat-profile detector",
             "USR-234 and USR-042 fall below normal users on the composite"],
        ],
        col_widths=[1.5, 2.0, 1.5, 1.5],
    )

    add_body_text(doc,
        "The tradeoff is explicit: raw features provide numeric precision for catching "
        "volume-based anomalies; embeddings provide semantic understanding for catching "
        "behavioral shifts at constant volume. The composite scoring tier fuses both "
        "perspectives — using group-relative z-scores from raw features with novelty "
        "persistence from qualitative analysis — but on its own cleanly separates only 2 of 4 "
        "attacks. The clean 4/4 detection at 0 false positives is provided by the separate "
        "multi-front threat-profile detector that targets measurable known-bad profiles.")

    add_page_break(doc)

    # ── End matter ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Technical Specification —")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "22nd Century Technologies, Inc.\n"
        "V-Intelligence UEBA Program — Three-Tier Technical Specification\n"
        "Document Version 4.0 — June 2026"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE

    # ── Save ────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Technical specification created successfully")
    print(f"Output: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
