#!/usr/bin/env python3
"""Build V-Intelligence UEBA Implementation Specification (Word).

Internal proprietary engineering blueprint with every formula, threshold,
weight, and edge case needed to recreate the V-Intelligence UEBA system.

Output: docs/V_Intelligence_UEBA_Implementation_Specification.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
BLUE = RGBColor(0x1B, 0x4F, 0x72)
TEAL = RGBColor(0x0E, 0x6B, 0x8A)
GOLD = RGBColor(0xB7, 0x95, 0x0B)
RED_ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREEN_ACCENT = RGBColor(0x1E, 0x8A, 0x49)
BLACK = RGBColor(0x00, 0x00, 0x00)

OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
                           "V_Intelligence_UEBA_Implementation_Specification.docx")


# ══════════════════════════════════════════════════════════════════════════════
#  Helper functions (identical pattern to build_tech_spec.py)
# ══════════════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
    return h


def add_body_text(doc, text, bold=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    return p


def add_callout(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEAL
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(18)
    return p


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "0D1B2A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK

    # ── TITLE PAGE ─────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("V-INTELLIGENCE UEBA")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Implementation Specification")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = BLUE

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Internal Engineering Reference — PROPRIETARY")
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("22nd Century Technologies, Inc.\nJune 2026 — Version 1.0")
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CLASSIFICATION: PROPRIETARY — NOT FOR EXTERNAL DISTRIBUTION")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_ACCENT

    add_page_break(doc)

    # ── TABLE OF CONTENTS ──────────────────────────────────────────────────
    add_section_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. System Architecture Overview",
        "2. Data Pipeline",
        "   2.1. Log Sources",
        "   2.2. Daily Feature Extraction",
        "   2.3. Weekly Aggregation (23 Features)",
        "3. Embedding Infrastructure",
        "   3.1. OpenAI Embedder",
        "   3.2. Two-Level Cache",
        "   3.3. Real Embeddings Mandatory (No Mock/Fallback)",
        "4. Text Serialization (Scalar to Prose)",
        "   4.1. Basic Zone Serialization",
        "   4.2. Interpretive Serialization",
        "   4.3. Population Context Functions",
        "   4.4. Critical Anomaly Detection",
        "5. Hierarchical Zone Architecture",
        "   5.1. User Zones (5 Zones)",
        "   5.2. Device Zones (5 Zones)",
        "   5.3. Zone Serialization Templates",
        "6. Signal Composition Weights",
        "   6.1. Entity Type Weights (5 Types)",
        "   6.2. Composition Algorithm",
        "7. Context-Adaptive Composition",
        "   7.1. Context Weight Tables",
        "   7.2. Linear-Normalized Attention",
        "   7.3. Zone Composition Logic",
        "8. Composite Scoring Formula",
        "   8.1. Five-Phase Scoring",
        "   8.2. Role Groups (31 Roles)",
        "   8.3. Z-Score Baseline Computation",
        "9. Novelty Detection",
        "   9.1. Baseline Period",
        "   9.2. Domain Entropy",
        "   9.3. Role-Expected Directories",
        "   9.4. Sensitive Directories",
        "   9.5. Novel IP Persistence",
        "   9.6. Annotation Tags",
        "10. Reference Concepts Library",
        "   10.1. Threat Concepts (14)",
        "   10.2. Benign Concepts (2)",
        "   10.3. Zone-Relevant Concept Mappings",
        "11. Temporal Trajectory",
        "   11.1. Velocity Vector",
        "   11.2. Six Trajectory Features",
        "   11.3. Regime Detection Rules",
        "12. Relationship Embeddings",
        "   12.1. Hadamard Product",
        "   12.2. Relationship Types",
        "   12.3. Relationship Drift",
        "13. Database Schema",
        "   13.1. Entity Tables",
        "   13.2. Log Tables",
        "   13.3. Behavioral Embedding Tables",
        "   13.4. Trajectory and Detection Tables",
        "   13.5. HNSW Indexes",
        "14. Configuration Reference",
        "   14.1. Tier3Config Defaults",
        "   14.2. PhaseState Defaults",
        "   14.3. Environment Variables",
        "   14.4. Drift Detection Parameters",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        if not item.startswith("   "):
            run.bold = True
    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  1. SYSTEM ARCHITECTURE OVERVIEW
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "1. System Architecture Overview", level=1)

    add_body_text(doc, (
        "V-Intelligence UEBA is a three-tier behavioral analytics system that detects "
        "insider threats, advanced persistent threats (APTs), living-off-the-land (LOTL) "
        "attacks, and telecom infrastructure pivots by tracking behavioral drift in "
        "1536-dimensional semantic embedding space. This document provides the complete "
        "engineering blueprint with every formula, threshold, weight, and edge case "
        "needed to recreate the system from scratch."
    ))

    add_section_heading(doc, "1.1. Infrastructure Components", level=2)

    create_table(doc,
        ["Component", "Technology", "Port", "Container"],
        [
            ["Database", "PostgreSQL 16 + pgvector extension", "5437 (external)", "db"],
            ["API Server", "FastAPI (Python)", "8000", "api"],
            ["Dashboard", "Streamlit", "8501", "(host)"],
            ["Background Worker", "Python scheduler service", "(internal)", "worker"],
            ["Embedding Model", "OpenAI text-embedding-3-small", "API", "(external)"],
        ],
        col_widths=[1.5, 2.5, 1.2, 1.3],
    )

    add_section_heading(doc, "1.2. Data Flow Pipeline", level=2)

    add_code_block(doc,
        "Raw Telemetry (auth, file_access, endpoint, network, dns)\n"
        "  --> Daily Feature Extraction (34 daily features per user)\n"
        "    --> Weekly Aggregation (23 features per user per week)\n"
        "      --> Text Serialization (features --> prose per zone)\n"
        "        --> Embedding (prose --> 1536-d vectors per zone)\n"
        "          --> Context-Adaptive Composition (5 zones --> 1 composite)\n"
        "            --> Temporal Snapshots (weekly composite series)\n"
        "              --> Detection (17 methods across 3 tiers)"
    )

    add_section_heading(doc, "1.3. Docker Deployment", level=2)

    add_body_text(doc, (
        "The system deploys via docker-compose with three containers: db (PostgreSQL+pgvector), "
        "api (FastAPI), and worker (background scheduler). Code changes are live via volume "
        "mount; dependency changes require --build. The DATABASE_URL environment variable "
        "controls connectivity: within Docker the host is 'db', externally 'localhost:5437'. "
        "Never hardcode localhost for database connections."
    ))

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  2. DATA PIPELINE
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "2. Data Pipeline", level=1)

    add_section_heading(doc, "2.1. Log Sources", level=2)

    create_table(doc,
        ["Log Source", "Key Columns", "Daily Volume"],
        [
            ["auth", "user_id, success, source_ip, dest_system, auth_method, timestamp", "~5,000 events"],
            ["file_access", "user_id, file_path, operation, data_classification, file_size_bytes", "~8,000 events"],
            ["endpoint", "device_id, user_id, process_name, event_type, risk_score", "~6,000 events"],
            ["network", "device_id, src_ip, dst_ip, bytes_in, bytes_out, protocol", "~15,000 events"],
            ["dns", "device_id, query_name/query_domain, response_code", "~10,000 events"],
        ],
        col_widths=[1.2, 3.5, 1.3],
    )

    add_section_heading(doc, "2.2. Daily Feature Extraction", level=2)

    add_body_text(doc, (
        "Daily features are extracted from each log source per user per day. The daily_features "
        "table stores 34 columns capturing granular per-day activity. These daily features are "
        "then aggregated into weekly windows of 7 days each."
    ))

    add_section_heading(doc, "2.3. Weekly Aggregation (23 Features)", level=2)

    add_body_text(doc, (
        "All detection tiers operate on weekly-aggregated features. The pipeline iterates "
        "through the observation period one week at a time, loading CSV files for each day "
        "within the window. User-to-device mapping links network and DNS logs (keyed by "
        "device_id) to users (keyed by user_id). Source: comparison/run_comparison.py, "
        "engineer_weekly_features()."
    ))

    create_table(doc,
        ["#", "Feature Name", "Log Source", "Computation"],
        [
            ["1", "auth_total", "auth", "Count of all auth events for user in week"],
            ["2", "auth_fail_rate", "auth", "auth_failed / max(auth_total, 1)"],
            ["3", "auth_off_hours_ratio", "auth", "Fraction with hour < 8 or hour >= 18"],
            ["4", "auth_unique_sources", "auth", "source_ip.nunique()"],
            ["5", "auth_unique_dests", "auth", "dest_system.nunique()"],
            ["6", "auth_methods_used", "auth", "auth_method.nunique()"],
            ["7", "file_total", "file_access", "Count of all file access events"],
            ["8", "file_restricted_ratio", "file_access", "Fraction with data_classification == 'restricted'"],
            ["9", "file_confidential_ratio", "file_access", "Fraction with data_classification == 'confidential'"],
            ["10", "file_write_ratio", "file_access", "Fraction with operation == 'write'"],
            ["11", "file_unique_paths", "file_access", "file_path.nunique()"],
            ["12", "file_total_bytes", "file_access", "Sum of file_size_bytes"],
            ["13", "net_bytes_out", "network", "Sum of bytes_out for user's devices"],
            ["14", "net_unique_dsts", "network", "dst_ip.nunique() for user's devices"],
            ["15", "net_external_ratio", "network", "Fraction where dst_ip is not RFC1918"],
            ["16", "dns_unique_domains", "dns", "query_domain.nunique() for user's devices"],
            ["17", "dns_nxdomain_ratio", "dns", "Fraction with response_code == 'NXDOMAIN'"],
            ["18", "endpoint_total", "endpoint", "Count of all endpoint telemetry events"],
            ["19", "endpoint_suspicious_ratio", "endpoint", "Fraction with event_type == 'suspicious'"],
            ["20", "endpoint_max_risk", "endpoint", "Max of risk_score column"],
            ["21", "endpoint_mean_risk", "endpoint", "Mean of risk_score column"],
            ["22", "endpoint_unique_processes", "endpoint", "process_name.nunique()"],
            ["23", "endpoint_total", "endpoint", "Count of all endpoint events (duplicate retained for compatibility)"],
        ],
        col_widths=[0.3, 1.8, 1.0, 3.4],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  3. EMBEDDING INFRASTRUCTURE
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "3. Embedding Infrastructure", level=1)

    add_body_text(doc, (
        "All embeddings are 1536-dimensional vectors produced by OpenAI text-embedding-3-small. "
        "All entity types share the same embedding space, enabling cross-entity comparison "
        "via cosine similarity. Source: embeddings/embedder.py."
    ))

    add_section_heading(doc, "3.1. OpenAI Embedder", level=2)

    create_table(doc,
        ["Parameter", "Value", "Source"],
        [
            ["Model", "text-embedding-3-small", "EMBEDDING_MODEL constant"],
            ["Dimensions", "1536", "EMBEDDING_DIM constant"],
            ["Batch size", "MAX_BATCH_SIZE = 2000", "OpenAI supports large batches for short texts"],
            ["Rate limiting", "RATE_LIMIT_PAUSE = 0.2 seconds between batches", "Prevents 429 errors"],
            ["Retry attempts", "3", "Exponential backoff: 2^(attempt+1) seconds"],
            ["Retry backoff", "2s, 4s, 8s", "2^1, 2^2, 2^3 seconds"],
            ["Client timeout", "60.0 seconds", "OpenAI(timeout=60.0)"],
            ["Lazy initialization", "Client created on first use", "self._client = None at init"],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    add_section_heading(doc, "3.2. Two-Level Cache", level=2)

    add_body_text(doc, "Cache Architecture:", bold=True)
    add_code_block(doc,
        "Level 1: In-memory dict  (self._mem_cache: dict[str, np.ndarray])\n"
        "Level 2: Sharded disk    (data/embedding_cache/{xx}/{hash}.npy)\n"
        "\n"
        "Cache key: SHA-256 hash of text (hashlib.sha256(text.encode('utf-8')).hexdigest())\n"
        "Shard:     First 2 hex characters of hash\n"
        "Path:      data/embedding_cache/ab/abcdef0123456789...npy\n"
        "\n"
        "Consolidated cache: data/embedding_cache/_consolidated.pkl\n"
        "  - Loaded at startup when preload=True\n"
        "  - Contains all embeddings in a single pickle dict\n"
        "  - Falls back to scanning shard directories if .pkl not found"
    )

    add_body_text(doc, "Cache lookup order:", bold=True)
    add_bullet(doc, "Check in-memory dict (self._mem_cache) by SHA-256 key")
    add_bullet(doc, "Check sharded disk path (cache_dir/shard/hash.npy)")
    add_bullet(doc, "Check flat disk path (cache_dir/hash.npy) for legacy entries")
    add_bullet(doc, "On cache miss: call OpenAI API, store in both memory and disk")

    add_body_text(doc, "Cache validation:", bold=True)
    add_bullet(doc, "Loaded vectors must have shape (1536,) exactly")
    add_bullet(doc, "Invalid shapes or corrupted files are deleted (path.unlink())")
    add_bullet(doc, "Preload reports memory usage: count * 1536 * 4 bytes / 1MB")

    add_section_heading(doc, "3.3. Real Embeddings Mandatory (No Mock/Fallback)", level=2)

    add_body_text(doc, (
        "Real OpenAI text-embedding-3-small (1536-d) embeddings are mandatory. Mock and "
        "fallback embeddings have been removed from the codebase: no part of the pipeline "
        "substitutes synthetic vectors. Every pipeline stage hard-fails without OPENAI_API_KEY; "
        "the test suite skips embedding-dependent tests when no key is present rather than "
        "falling back to a substitute embedder."
    ))

    add_callout(doc, "IMPLEMENTATION NOTE: Hash-based or random substitute vectors produce "
        "orthogonal vectors for different inputs (~0.0 cosine similarity), whereas real "
        "embeddings produce semantically meaningful similarity (~0.3-0.7 for related texts). "
        "Because of this fundamental difference, substitute embeddings cannot validate "
        "detection logic and are not used.")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  4. TEXT SERIALIZATION
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "4. Text Serialization (Scalar to Prose)", level=1)

    add_body_text(doc, (
        "The critical transformation that makes behavioral embedding detection work: "
        "raw numeric features are converted to natural language text before embedding. "
        "This enables the language model to capture semantic meaning ('15% restricted "
        "file access' means something different from '0.1% restricted file access') "
        "that raw numeric vectors cannot represent. Source: models/hierarchical_zones.py."
    ))

    add_section_heading(doc, "4.1. Basic Zone Serialization", level=2)

    add_body_text(doc, "Each zone has a fixed text template. Examples:", bold=True)

    add_body_text(doc, "Identity zone:", bold=True)
    add_code_block(doc,
        '"User {uid} identity: role={role}, department={department}, "\n'
        '"clearance={clearance}, tenure_days={tenure_days}, type={user_type}"'
    )

    add_body_text(doc, "Access pattern zone:", bold=True)
    add_code_block(doc,
        '"User {uid} access: auth_events={auth_total:.0f}, "\n'
        '"fail_rate={auth_fail_rate:.4f}, off_hours={auth_off_hours_ratio:.4f}, "\n'
        '"unique_sources={auth_unique_sources:.0f}, "\n'
        '"unique_dests={auth_unique_dests:.0f}, methods={auth_methods_used:.0f}"'
    )

    add_body_text(doc, "Data behavior zone:", bold=True)
    add_code_block(doc,
        '"User {uid} data: file_accesses={file_total:.0f}, "\n'
        '"restricted_ratio={file_restricted_ratio:.4f}, "\n'
        '"confidential_ratio={file_confidential_ratio:.4f}, "\n'
        '"write_ratio={file_write_ratio:.4f}, "\n'
        '"unique_paths={file_unique_paths:.0f}, bytes={file_total_bytes:.0f}"\n'
        '+ optional: ". Directories accessed: {qual_file_dirs}"'
    )

    add_body_text(doc, "Network footprint zone:", bold=True)
    add_code_block(doc,
        '"User {uid} network: bytes_out={net_bytes_out:.0f}, "\n'
        '"unique_dsts={net_unique_dsts:.0f}, external_ratio={net_external_ratio:.4f}, "\n'
        '"dns_domains={dns_unique_domains:.0f}, nxdomain_ratio={dns_nxdomain_ratio:.4f}"\n'
        '+ optional: ". External destinations: {qual_net_ext_ips}"\n'
        '+ optional: ". DNS domains queried: {qual_dns_domains}"'
    )

    add_body_text(doc, "Risk posture zone:", bold=True)
    add_code_block(doc,
        '"User {uid} risk: endpoint_events={endpoint_total:.0f}, "\n'
        '"suspicious_ratio={endpoint_suspicious_ratio:.4f}, "\n'
        '"max_risk={endpoint_max_risk:.4f}, mean_risk={endpoint_mean_risk:.4f}, "\n'
        '"unique_processes={endpoint_unique_processes:.0f}"'
    )

    add_section_heading(doc, "4.2. Interpretive Serialization", level=2)

    add_body_text(doc, (
        "When a BehavioralContext is provided, the serializer produces semantically rich "
        "text that encodes population percentiles, user baselines, and trend labels. This "
        "is the key innovation: the same numeric value produces dramatically different "
        "text (and therefore different embeddings) depending on whether it is normal or "
        "anomalous relative to the population and the user's own history."
    ))

    add_body_text(doc, (
        "The interpretive serializer uses a tiered approach based on the peak z-score "
        "of the zone's features:"
    ))

    add_bullet(doc, "Peak z > 2.5: CRITICAL ANOMALY text with threat language and investigation recommendations", bold_prefix="Critical tier: ")
    add_bullet(doc, "Peak z between 1.5 and 2.5: Moderately elevated text with specific feature assessments", bold_prefix="Elevated tier: ")
    add_bullet(doc, "Peak z < 1.5: Normal parameters text with 'unremarkable' language", bold_prefix="Normal tier: ")

    add_section_heading(doc, "4.3. Population Context Functions", level=2)

    add_body_text(doc, "_percentile_label(z) -- Maps z-score to human-readable label:", bold=True)
    create_table(doc,
        ["Z-Score Range", "Label"],
        [
            ["z > 3.0", "extreme"],
            ["z > 2.0", "very high"],
            ["z > 1.0", "elevated"],
            ["z > 0.5", "above average"],
            ["z > -0.5", "normal"],
            ["z > -1.0", "below average"],
            ["z > -2.0", "low"],
            ["z <= -2.0", "very low"],
        ],
        col_widths=[2.0, 2.0],
    )

    add_body_text(doc, "_trend_label(values) -- Compares current to recent 3 values:", bold=True)
    create_table(doc,
        ["Condition", "Label"],
        [
            ["< 3 values available", "insufficient data"],
            ["All 3 recent values strictly increasing", "increasing"],
            ["All 3 recent values strictly decreasing", "decreasing"],
            ["Otherwise", "fluctuating"],
        ],
        col_widths=[3.0, 2.0],
    )

    add_body_text(doc, "_baseline_ratio(val, baseline) -- Labels relative to user's own baseline:", bold=True)
    create_table(doc,
        ["Ratio (val/baseline)", "Label"],
        [
            ["baseline < 1e-6 and val < 1e-6", "no baseline"],
            ["baseline < 1e-6 and val >= 1e-6", "new activity"],
            ["ratio > 3.0", "{ratio:.1f}x baseline (major escalation)"],
            ["ratio > 1.5", "{ratio:.1f}x baseline (elevated)"],
            ["ratio > 1.1", "{ratio:.1f}x baseline (slightly above)"],
            ["ratio > 0.9", "at baseline"],
            ["ratio > 0.5", "{ratio:.1f}x baseline (reduced)"],
            ["ratio <= 0.5", "{ratio:.1f}x baseline (significantly reduced)"],
        ],
        col_widths=[2.5, 3.5],
    )

    add_section_heading(doc, "4.4. Critical Anomaly Detection", level=2)

    add_body_text(doc, (
        "When the peak z-score of any feature within a zone exceeds 2.5, the serializer "
        "switches to CRITICAL ANOMALY mode. This produces text with explicit threat language "
        "that the embedding model encodes as semantically close to threat reference concepts."
    ))

    add_body_text(doc, "Critical anomaly z-score threshold: 2.5", bold=True)

    add_body_text(doc, "Example critical anomaly text for access_pattern zone:", bold=True)
    add_code_block(doc,
        '"CRITICAL ANOMALY in authentication for {role} {uid} in {dept}. "\n'
        '"Authentication failure rate at EXTREME level: {fail_rate:.3f} "\n'
        '"({z:.1f} standard deviations above population mean, "\n'
        '"{baseline_ratio}). This pattern is strongly consistent with "\n'
        '"credential stuffing, brute force attack, or compromised credential reuse. "\n'
        '"Immediate investigation recommended."'
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  5. HIERARCHICAL ZONE ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "5. Hierarchical Zone Architecture", level=1)

    add_body_text(doc, (
        "Instead of one composite 1536-d embedding per entity (Tier 2), Tier 3 builds "
        "5 independent zone embeddings per entity, then composes via context-adaptive "
        "attention. This enables zone-specific drift detection: 'identity stable but "
        "network_footprint drifting' = C2 pattern. Source: models/hierarchical_zones.py."
    ))

    add_section_heading(doc, "5.1. User Zones (5 Zones)", level=2)

    create_table(doc,
        ["Zone", "Fields", "Purpose"],
        [
            ["identity", "role, department, clearance, tenure_days, user_type",
             "Static profile; should NOT drift unless legitimate role change"],
            ["access_pattern", "auth_total, auth_fail_rate, auth_off_hours_ratio, auth_unique_sources, auth_unique_dests, auth_methods_used",
             "Authentication behavior; catches credential abuse, lateral movement"],
            ["data_behavior", "file_total, file_restricted_ratio, file_confidential_ratio, file_write_ratio, file_unique_paths, file_total_bytes",
             "Data access patterns; catches insider threat, data exfiltration"],
            ["network_footprint", "net_bytes_out, net_unique_dsts, net_external_ratio, dns_unique_domains",
             "Network communication; catches C2 beacons, APT, DNS tunneling"],
            ["risk_posture", "endpoint_suspicious_ratio, endpoint_max_risk, endpoint_mean_risk, endpoint_unique_processes, endpoint_total, dns_nxdomain_ratio",
             "Endpoint risk profile; catches compromised endpoints, LOTL tools"],
        ],
        col_widths=[1.2, 2.8, 2.5],
    )

    add_body_text(doc, "USER_ZONE_ORDER = [\"identity\", \"access_pattern\", \"data_behavior\", \"network_footprint\", \"risk_posture\"]", bold=True)

    add_section_heading(doc, "5.2. Device Zones (5 Zones)", level=2)

    create_table(doc,
        ["Zone", "Fields"],
        [
            ["identity", "device_type, os, segment_id"],
            ["process_behavior", "endpoint_total, endpoint_suspicious_ratio, endpoint_unique_processes"],
            ["traffic_pattern", "net_bytes_out, net_unique_dsts, net_external_ratio"],
            ["resource_usage", "endpoint_max_risk, endpoint_mean_risk"],
            ["auth_profile", "auth_total, auth_fail_rate, auth_off_hours_ratio"],
        ],
        col_widths=[1.5, 5.0],
    )

    add_section_heading(doc, "5.3. Zone Composition Behavior", level=2)

    add_body_text(doc, (
        "The compose_zones() function excludes static zones (identity) from the composite "
        "by default because they never drift and dilute the behavioral signal. The identity "
        "zone is still used separately as a stability reference in zone divergence detection."
    ))

    add_code_block(doc,
        "STATIC_ZONES = {\"identity\"}\n"
        "if exclude_static:\n"
        "    active_zones = {k: v for k, v in zone_embeddings.items()\n"
        "                    if k not in STATIC_ZONES}\n"
        "else:\n"
        "    active_zones = zone_embeddings\n"
        "\n"
        "# Fallback: if all zones are static, use all of them\n"
        "if not active_zones:\n"
        "    active_zones = zone_embeddings"
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  6. SIGNAL COMPOSITION WEIGHTS
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "6. Signal Composition Weights", level=1)

    add_body_text(doc, (
        "Signal composition is the Tier 2 mechanism for combining 5 behavioral signal "
        "embeddings into a single composite vector per entity. Source: embeddings/composer.py."
    ))

    add_section_heading(doc, "6.1. Entity Type Weights", level=2)

    create_table(doc,
        ["Entity Type", "Signal 1", "Weight", "Signal 2", "Weight", "Signal 3", "Weight", "Signal 4", "Weight", "Signal 5", "Weight"],
        [
            ["user", "auth", "0.25", "privilege", "0.20", "data_access", "0.20", "network", "0.20", "communication", "0.15"],
            ["device", "process", "0.25", "traffic", "0.25", "resource", "0.20", "auth", "0.15", "config", "0.15"],
            ["segment", "volume", "0.20", "connections", "0.25", "protocols", "0.20", "threats", "0.20", "exposure", "0.15"],
            ["app", "access", "0.25", "queries", "0.20", "errors", "0.20", "performance", "0.15", "config", "0.20"],
            ["session", "activity", "0.20", "risk_accum", "0.25", "data_movement", "0.20", "lateral", "0.20", "temporal", "0.15"],
        ],
        col_widths=[0.7, 0.7, 0.4, 0.7, 0.4, 0.8, 0.4, 0.7, 0.4, 0.8, 0.4],
    )

    add_section_heading(doc, "6.2. Composition Algorithm", level=2)

    add_code_block(doc,
        "def compose(signal_vectors, entity_type):\n"
        "    weights = SIGNAL_WEIGHTS[entity_type]\n"
        "    composite = zeros(1536, float64)\n"
        "    total_weight = 0.0\n"
        "\n"
        "    for signal_name, vector in signal_vectors.items():\n"
        "        w = weights.get(signal_name, 0.0)\n"
        "        if w > 0:\n"
        "            composite += w * vector.astype(float64)\n"
        "            total_weight += w\n"
        "\n"
        "    # Normalize by actual weight sum (handles missing signals)\n"
        "    composite = composite / total_weight\n"
        "\n"
        "    # L2-normalize to unit vector\n"
        "    norm = linalg.norm(composite)\n"
        "    if norm > 0:\n"
        "        composite = composite / norm\n"
        "\n"
        "    return composite.astype(float32)"
    )

    add_callout(doc, "EDGE CASE: If fewer than 5 signals are available, the formula "
        "normalizes by the sum of available weights (total_weight), ensuring the "
        "composite remains a valid unit vector regardless of missing signals.")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  7. CONTEXT-ADAPTIVE COMPOSITION
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "7. Context-Adaptive Composition", level=1)

    add_body_text(doc, (
        "The Tier 3 zone composition uses context-adaptive weights to bias attention toward "
        "zones relevant to the investigation scenario. Four predefined contexts change how "
        "zones are weighted during composition. Source: models/hierarchical_zones.py, "
        "CONTEXT_WEIGHTS dict and compose_zones() function."
    ))

    add_section_heading(doc, "7.1. Context Weight Tables", level=2)

    create_table(doc,
        ["Context", "identity", "access_pattern", "data_behavior", "network_footprint", "risk_posture", "Use Case"],
        [
            ["normal_ops", "0.20", "0.20", "0.20", "0.20", "0.20", "Default monitoring"],
            ["insider_investigation", "0.10", "0.15", "0.40", "0.15", "0.20", "Data exfiltration hunt"],
            ["apt_hunt", "0.05", "0.15", "0.10", "0.40", "0.30", "C2/network threat hunt"],
            ["privilege_audit", "0.10", "0.25", "0.10", "0.15", "0.40", "Escalation audit"],
        ],
        col_widths=[1.2, 0.6, 0.8, 0.8, 0.9, 0.7, 1.3],
    )

    add_section_heading(doc, "7.2. Linear-Normalized Attention", level=2)

    add_body_text(doc, (
        "V-Intelligence UEBA uses direct linear normalization (NOT softmax) for attention weights. "
        "Softmax would compress weight differentiation (intended 0.40 becomes ~0.23); linear "
        "normalization preserves the intended weight ratios."
    ))

    add_body_text(doc, "Attention formula:", bold=True)
    add_code_block(doc,
        "def softmax_attention(zone_vecs, context_weights):\n"
        "    raw = {}\n"
        "    for name, vec in zone_vecs.items():\n"
        "        energy = float(np.linalg.norm(vec))  # L2 norm of zone vector\n"
        "        bias = context_weights.get(name, 0.2)\n"
        "        raw[name] = energy * bias\n"
        "\n"
        "    total = sum(raw.values())\n"
        "\n"
        "    if total == 0:  # Edge case: all zero vectors\n"
        "        uniform = 1.0 / len(zone_vecs)\n"
        "        return {k: uniform for k in zone_vecs}\n"
        "\n"
        "    return {k: v / total for k, v in raw.items()}\n"
        "\n"
        "# Formula: alpha_i = (||zone_i|| * context_weight[zone_i]) / sum(all)"
    )

    add_section_heading(doc, "7.3. Zone Composition Logic", level=2)

    add_code_block(doc,
        "def compose_zones(zone_embeddings, context='normal_ops',\n"
        "                  entity_type='user', exclude_static=True):\n"
        "    # 1. Remove static zones (identity) from composition\n"
        "    # 2. Get context weights for entity type\n"
        "    # 3. Compute attention weights (linear normalization)\n"
        "    # 4. Weighted sum: composite = sum(alpha_i * zone_i)\n"
        "    # 5. L2-normalize composite to unit vector\n"
        "    # 6. Return as float32 array"
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  8. COMPOSITE SCORING FORMULA
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "8. Composite Scoring Formula", level=1)

    add_body_text(doc, (
        "The multi-phase composite scoring system fuses five independent anomaly signals "
        "into a single ranked score per user. A normal user might score high on one phase "
        "by chance; an attacker scores high on multiple phases simultaneously. "
        "Source: detection/composite_scorer.py."
    ))

    add_section_heading(doc, "8.1. Five-Phase Scoring", level=2)

    add_body_text(doc, "EXACT composite formula:", bold=True)
    add_code_block(doc,
        "composite = signal_strength * 1.0\n"
        "          + breadth_15 * 0.5\n"
        "          + sustained_signal * 0.3\n"
        "          + max(ctx_spread_z, 0) * 0.5\n"
        "          + max(ctx_max_z, 0) * 0.3\n"
        "          + novelty_score * 1.0"
    )

    add_body_text(doc, "Phase 1: Signal Strength (weight: 1.0)", bold=True)
    add_code_block(doc,
        "z_vals = sorted([row[c] for c in z_columns], reverse=True)\n"
        "signal_strength = sum(z_vals[:3])  # Sum of top-3 z-scores"
    )

    add_body_text(doc, "Phase 2: Signal Breadth (weight: 0.5)", bold=True)
    add_code_block(doc,
        "breadth_15 = sum(1 for v in z_vals if v > 1.5)  # Used in formula\n"
        "breadth_20 = sum(1 for v in z_vals if v > 2.0)  # Recorded, not used"
    )

    add_body_text(doc, "Phase 3: Sustained Zone Deviation (weight: 0.3)", bold=True)
    add_code_block(doc,
        "DRIFT_ZONES = ['access_pattern_drift', 'data_behavior_drift',\n"
        "               'network_footprint_drift', 'risk_posture_drift']\n"
        "\n"
        "zone_sustained_zs = []  # z-scored sustained deviation per zone\n"
        "for z in DRIFT_ZONES:\n"
        "    zn = z.replace('_drift', '')\n"
        "    zcol = f'z_{zn}_sustained'\n"
        "    zone_sustained_zs.append(row[zcol])\n"
        "sustained_signal = sum(sorted(zone_sustained_zs, reverse=True)[:2])  # Top 2"
    )

    add_body_text(doc, "Phase 4: Context Divergence (weights: 0.5 + 0.3)", bold=True)
    add_code_block(doc,
        "ctx_spread_z = row['z_ctx_spread']  # Z-scored context spread\n"
        "ctx_max_z = row['z_ctx_max']        # Z-scored context max\n"
        "\n"
        "# Both floored at 0 (negative values ignored)\n"
        "contribution = max(ctx_spread_z, 0) * 0.5 + max(ctx_max_z, 0) * 0.3\n"
        "\n"
        "# Context columns used:\n"
        "CONTEXT_COLS = ['composite_drift_normal_ops',\n"
        "                'composite_drift_insider_investigation',\n"
        "                'composite_drift_apt_hunt',\n"
        "                'composite_drift_privilege_audit']\n"
        "ctx_max = max(late_stage_means across 4 contexts)\n"
        "ctx_spread = max(late_stage_means) - min(late_stage_means)"
    )

    add_body_text(doc, "Phase 5: Novelty Persistence (weight: 1.0)", bold=True)
    add_code_block(doc,
        "novelty_score = 0.0\n"
        "\n"
        "# Component 1: IP persistence score\n"
        "persistence = novel_ip_max_persistence  # Max weeks any novel IP appears\n"
        "if persistence > 10:\n"
        "    novelty_score = min(persistence / 5.0, 10.0)\n"
        "\n"
        "# Component 2: Novel fraction bonus\n"
        "novel_frac = novel_ip_weeks_frac  # Fraction of post-baseline weeks with novel IPs\n"
        "if novel_frac > 0.5:\n"
        "    novelty_score += novel_frac * 3.0\n"
        "\n"
        "# Max possible novelty_score = 10.0 + 1.0 * 3.0 = 13.0"
    )

    add_section_heading(doc, "8.2. Role Groups (31 Roles)", level=2)

    add_body_text(doc, (
        "Z-scores are computed per role peer group. Attack users are EXCLUDED from the "
        "baseline computation. Standard deviation uses ddof=1."
    ))

    create_table(doc,
        ["Group", "Roles", "Count"],
        [
            ["admin", "IT Admin, SysAdmin, DBA, Network Engineer, Cloud Architect, DevOps Engineer, SRE", "7"],
            ["security", "SOC Operator, Security Analyst, CISO", "3"],
            ["developer", "Software Engineer, Senior Engineer, Staff Engineer, ML Engineer, Data Scientist, QA Engineer, Test Lead", "7"],
            ["business", "Accountant, Analyst, Account Manager, Financial Analyst, Sales Rep, HR Manager, HR Specialist, Recruiter, General Counsel", "9"],
            ["executive", "CEO, CFO, COO, CTO, VP Sales", "5"],
        ],
        col_widths=[1.0, 4.0, 0.5],
    )

    add_body_text(doc, "Total: 31 roles across 5 role groups.", bold=True)

    add_section_heading(doc, "8.3. Z-Score Baseline Computation", level=2)

    add_code_block(doc,
        "def compute_group_zscores(user_feats):\n"
        "    for col in feature_columns:\n"
        "        zcol = f'z_{col}'\n"
        "        for grp in groups:\n"
        "            # CRITICAL: only normal (non-attack) users in baseline\n"
        "            normal_mask = (df.grp == grp) & (~df.is_attack)\n"
        "            mu = df.loc[normal_mask, col].mean()\n"
        "            sigma = df.loc[normal_mask, col].std()  # ddof=1 default\n"
        "            if sigma > 1e-10:\n"
        "                df.loc[grp_mask, zcol] = (df.loc[grp_mask, col] - mu) / sigma\n"
        "            else:\n"
        "                df.loc[grp_mask, zcol] = 0.0  # No variation in group"
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  9. NOVELTY DETECTION
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "9. Novelty Detection", level=1)

    add_body_text(doc, (
        "Novelty detection enriches qualitative features (file directories, external IPs, "
        "DNS domains) with semantic annotations that help the embedding model distinguish "
        "benign variation from suspicious behavioral shifts. Source: detection/novelty_features.py."
    ))

    add_section_heading(doc, "9.1. Baseline Period", level=2)

    add_body_text(doc, "BASELINE_WEEKS = 10", bold=True)
    add_body_text(doc, (
        "The first 10 weeks of observation establish the baseline of known directories, "
        "IPs, and domains for each user. Weeks 0-9 are baseline; weeks 10+ are monitored "
        "for novel items."
    ))

    add_section_heading(doc, "9.2. Domain Entropy", level=2)

    add_body_text(doc, "Shannon entropy on character frequency of domain name (before first dot):", bold=True)
    add_code_block(doc,
        "def _domain_entropy(domain):\n"
        "    name = domain.split('.')[0]  # Take subdomain/label only\n"
        "    freq = {c: count for c in name.lower()}\n"
        "    length = len(name)\n"
        "    entropy = -sum((count/length) * log2(count/length)\n"
        "                   for count in freq.values())\n"
        "    return entropy\n"
        "\n"
        "# Threshold: entropy > 3.5 triggers HIGH-ENTROPY tag"
    )

    add_section_heading(doc, "9.3. Internal Domain Detection", level=2)

    add_code_block(doc,
        "def _is_internal_domain(domain):\n"
        "    parts = domain.lower().strip().split('.')\n"
        "    return parts[-1] in {'corp', 'internal', 'local', 'lan', 'intranet'}"
    )

    add_section_heading(doc, "9.4. Role-Expected Directories", level=2)

    add_body_text(doc, "ROLE_EXPECTED_DIRS -- full dict for all 5 role groups:", bold=True)
    create_table(doc,
        ["Role Group", "Expected Directories"],
        [
            ["developer", "/engineering, /shared, /docs, Documents/Project, Documents/Code, Documents/Product, Documents/Tech"],
            ["admin", "/infrastructure, /it-ops, /shared, /engineering, Documents/Network, Documents/Infrastructure"],
            ["security", "/security, /siem, /shared, /incident, Documents/Security, Documents/Incident"],
            ["business", "/finance, /hr, /sales, /shared, /legal, Documents/Budget, Documents/RFP, Documents/Vendor, Documents/Customer, Documents/Q4, Documents/Compliance, Documents/Merger, Documents/Patent"],
            ["executive", "/executive, /finance, /strategy, /shared, Documents/Board, Documents/Strategy"],
        ],
        col_widths=[1.0, 5.5],
    )

    add_section_heading(doc, "9.5. Sensitive Directories", level=2)

    add_body_text(doc, "SENSITIVE_DIRS -- directories mapped to allowed role groups:", bold=True)
    create_table(doc,
        ["Directory Pattern", "Allowed Groups"],
        [
            ["/executive", "executive"],
            ["/finance", "business, executive"],
            ["/hr", "business, executive"],
            ["/security/audit", "security"],
            ["/security/policies", "security, admin"],
            ["/legal", "business, executive"],
        ],
        col_widths=[2.0, 4.5],
    )

    add_section_heading(doc, "9.6. Novel IP Persistence Metrics", level=2)

    add_code_block(doc,
        "# Per-user novelty metrics from compute_novelty_metrics():\n"
        "\n"
        "persistent_novel_ips:    count of novel IPs appearing in >= 5 post-baseline weeks\n"
        "novel_ip_max_persistence: max weeks any single novel IP appears\n"
        "novel_ip_weeks_frac:     fraction of post-baseline weeks with ANY novel IPs\n"
        "\n"
        "# IP is 'novel' if it never appeared during baseline (weeks 0-9)\n"
        "# IP is 'persistent' if it appears in >= 5 post-baseline weeks\n"
        "\n"
        "# Novel IPs are extracted by splitting on ';' and stripping [NOVEL] tags"
    )

    add_section_heading(doc, "9.7. Annotation Tags", level=2)

    create_table(doc,
        ["Tag", "Condition", "Applied To"],
        [
            ["[NOVEL]", "Item not seen in baseline weeks 0-9", "Directories, IPs, Domains"],
            ["[ATYPICAL-{GROUP}]", "Directory is in SENSITIVE_DIRS and role group not in allowed set", "Directories only"],
            ["[HIGH-ENTROPY=X.X]", "Domain entropy > 3.5 AND domain is not internal", "DNS domains only"],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  10. REFERENCE CONCEPTS LIBRARY
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "10. Reference Concepts Library", level=1)

    add_body_text(doc, (
        "16 reference concept descriptions are embedded into the same 1536-d space as entity "
        "behaviors. Drift vectors are projected onto these concepts to classify the nature "
        "of behavioral change: 'what is this entity BECOMING?' "
        "Source: detection/reference_concepts.py."
    ))

    add_section_heading(doc, "10.1. Threat Concepts (14)", level=2)

    create_table(doc,
        ["#", "Name", "Category", "Severity", "MITRE Techniques"],
        [
            ["1", "compromised_endpoint", "Threat", "Critical", "T1059, T1547, T1562, T1071"],
            ["2", "data_exfiltration", "Threat", "Critical", "T1005, T1074, T1048, T1567"],
            ["3", "privilege_escalation", "Threat", "High", "T1078, T1068, T1134, T1548"],
            ["4", "lateral_movement", "Threat", "High", "T1021, T1570, T1550, T1072"],
            ["5", "insider_threat_slow", "Threat", "High", "T1078, T1083, T1005, T1052"],
            ["6", "insider_threat_fast", "Threat", "Critical", "T1005, T1052, T1048, T1567"],
            ["7", "credential_stuffing", "Threat", "High", "T1110, T1078"],
            ["8", "c2_beacon", "Threat", "Critical", "T1071, T1573, T1568, T1102"],
            ["9", "reconnaissance", "Threat", "Medium", "T1046, T1018, T1087, T1135"],
            ["10", "supply_chain_compromise", "Threat", "Critical", "T1195, T1059, T1071"],
            ["11", "living_off_the_land", "Threat", "Critical", "T1218, T1053, T1059.001, T1036"],
            ["12", "telecom_infrastructure_pivot", "Threat", "Critical", "T1557, T1040, T1005, T1556"],
            ["13", "dns_tunneling_exfil", "Threat", "Critical", "T1071.004, T1048.003, T1041"],
            ["14", "credential_harvesting_slow", "Threat", "High", "T1003, T1558, T1110.002, T1555"],
        ],
        col_widths=[0.3, 2.2, 0.6, 0.6, 2.5],
    )

    add_body_text(doc, "Full concept descriptions (embedded as text):", bold=True)

    _concept_descs = [
        ("compromised_endpoint", "A device showing signs of compromise: unusual process execution including encoded PowerShell commands, connection to known C2 infrastructure, registry modifications for persistence, and defense evasion through disabling security tools. Network beaconing at regular intervals to external IPs not seen before."),
        ("data_exfiltration", "An entity systematically collecting and transferring sensitive data outside the organization. Large file copies from restricted shares to staging directories, followed by compression and upload to external cloud storage or email to personal accounts. Data volume increasing exponentially over days."),
        ("privilege_escalation", "A user or process progressively gaining higher access levels. Initial standard access followed by requests for elevated permissions, exploitation of misconfigurations to access admin functions, or token manipulation. Access scope expanding from department-level to organization-wide."),
        ("lateral_movement", "An entity moving systematically across the network from system to system. Authentication to devices outside normal scope, SMB/WinRM/RDP connections to multiple hosts in sequence, credential reuse across systems, and accessing administrative shares on remote hosts."),
        ("insider_threat_slow", "A trusted user whose behavior is gradually shifting toward data theft. Slowly expanding file access scope, increasing off-hours activity, accessing documents outside their role, subtle privilege accumulation. Changes are individually innocuous but directionally concerning when viewed over months."),
        ("insider_threat_fast", "A user rapidly collecting and exfiltrating sensitive data, likely triggered by imminent termination or grievance. Sudden spike in file downloads, bulk access to restricted documents, USB device usage, large email attachments to personal addresses, all within days."),
        ("credential_stuffing", "Mass authentication attempts against multiple user accounts from external IPs. High failure rate across many accounts in a short window, automated patterns with consistent timing, originating from known proxy/VPN/botnet infrastructure."),
        ("c2_beacon", "A device establishing covert command-and-control communication. Regular periodic connections to external infrastructure with consistent intervals (with jitter), small packet sizes, DNS-based or HTTP-based tunneling, and connections to newly-registered domains or cloud services not previously seen."),
        ("reconnaissance", "Systematic discovery and mapping of network resources, user accounts, and system configurations. Port scanning, LDAP enumeration, share discovery, account enumeration, and DNS zone transfer attempts. Methodical progression through network segments."),
        ("supply_chain_compromise", "A trusted application exhibiting new suspicious behaviors after an update. Unexpected network connections to unfamiliar domains, new process spawning patterns, access to data stores not previously touched, and subtle changes in API call patterns."),
        ("living_off_the_land", "An entity using only built-in operating system tools for malicious purposes with no malware deployed. Execution of certutil for downloads, wmic for remote process execution, netsh for port forwarding, schtasks for persistence, and PowerShell with obfuscated commands."),
        ("telecom_infrastructure_pivot", "An entity targeting telecommunications infrastructure for surveillance. Access to call detail record databases, CDR metadata harvesting, lawful intercept system manipulation, router and switch configuration access, and cross-segment network flows to restricted management zones."),
        ("dns_tunneling_exfil", "An entity using DNS queries as a covert data exfiltration channel. High-entropy subdomain labels encoding stolen data, elevated TXT and MX record queries to attacker-controlled domains, persistent low-throughput data channel disguised as legitimate DNS resolution."),
        ("credential_harvesting_slow", "An entity systematically collecting credentials through native tools over an extended campaign. LSASS memory access, SAM database dumps via reg.exe, Kerberoasting service ticket requests, and gradual credential accumulation distinct from brute-force attacks."),
    ]
    for cname, cdesc in _concept_descs:
        add_bullet(doc, cdesc, bold_prefix=f"{cname}: ")

    add_section_heading(doc, "10.2. Benign Concepts (2)", level=2)

    create_table(doc,
        ["#", "Name", "Category", "Severity", "Description"],
        [
            ["15", "normal_role_change", "Benign", "Low",
             "A user whose behavior is shifting due to legitimate organizational change. New application access corresponding to role transfer, different work hours matching new team timezone, new peer group communication patterns."],
            ["16", "seasonal_variation", "Benign", "Low",
             "Entity behavior fluctuating with business cycles. End-of-quarter spikes in finance system usage, annual audit-driven access pattern changes, holiday-period reduced activity, or project-deadline-driven temporary increases in off-hours work."],
        ],
        col_widths=[0.3, 1.5, 0.6, 0.6, 3.5],
    )

    add_section_heading(doc, "10.3. Zone-Relevant Concept Mappings", level=2)

    add_body_text(doc, (
        "Each zone is compared only against semantically relevant concepts. This prevents "
        "cross-domain false alignments (e.g., data_behavior zone spuriously aligning with "
        "c2_beacon, a network concept). Source: detection/drift_direction.py, "
        "ZONE_RELEVANT_CONCEPTS dict."
    ))

    create_table(doc,
        ["Zone", "Relevant Concepts"],
        [
            ["identity", "(empty -- no concepts mapped)"],
            ["access_pattern", "credential_stuffing, privilege_escalation, lateral_movement, insider_threat_slow, insider_threat_fast, credential_harvesting_slow, living_off_the_land, normal_role_change, seasonal_variation"],
            ["data_behavior", "data_exfiltration, insider_threat_slow, insider_threat_fast, privilege_escalation, telecom_infrastructure_pivot, normal_role_change, seasonal_variation"],
            ["network_footprint", "c2_beacon, lateral_movement, reconnaissance, compromised_endpoint, dns_tunneling_exfil, insider_threat_slow, living_off_the_land, telecom_infrastructure_pivot, seasonal_variation"],
            ["risk_posture", "compromised_endpoint, privilege_escalation, supply_chain_compromise, credential_harvesting_slow, living_off_the_land, dns_tunneling_exfil, normal_role_change"],
        ],
        col_widths=[1.3, 5.2],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  11. TEMPORAL TRAJECTORY
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "11. Temporal Trajectory", level=1)

    add_body_text(doc, (
        "Temporal trajectory analysis extracts 6 features from a sequence of weekly "
        "embedding snapshots, capturing how an entity moves through embedding space over "
        "time. Full 1536-d velocity vectors (not just scalar drift magnitude) are computed. "
        "Source: models/temporal_trajectory.py."
    ))

    add_section_heading(doc, "11.1. Velocity Vector", level=2)

    add_code_block(doc,
        "def compute_velocity_vector(snapshots):\n"
        "    # velocity = normalized(snapshots[-1] - snapshots[0])\n"
        "    # Returns 1536-d unit vector pointing in direction of behavioral change\n"
        "    # Returns zero vector if < 2 snapshots or ||diff|| < 1e-8\n"
        "\n"
        "    velocity = snapshots[-1].astype(float64) - snapshots[0].astype(float64)\n"
        "    norm = np.linalg.norm(velocity)\n"
        "    if norm < 1e-8:\n"
        "        return zeros_like(velocity, float32)\n"
        "    return (velocity / norm).astype(float32)"
    )

    add_section_heading(doc, "11.2. Six Trajectory Features", level=2)

    create_table(doc,
        ["Feature", "Formula", "Interpretation"],
        [
            ["velocity_magnitude", "mean(||snapshots[i+1] - snapshots[i]||) for all consecutive pairs", "Average rate of behavioral displacement per step"],
            ["acceleration", "mean(||diff[i+1]|| - ||diff[i]||) for consecutive diffs", "Whether drift is speeding up (+) or slowing down (-)"],
            ["stability", "mean(cosine_similarity(snapshots[i], snapshots[i+1])) for consecutive pairs", "1.0 = stable, < 0.7 = volatile"],
            ["regime_shifts", "count(cosine_sim < 0.7) / max(n-1, 1)", "Fraction of consecutive pairs with phase change"],
            ["trend_consistency", "mean pairwise cosine of consecutive difference vectors; skipped if < 2 diffs", "1.0 = drift in consistent direction; 0.0 = random walk"],
            ["total_drift", "1.0 - cosine_similarity(snapshots[0], snapshots[-1])", "Overall start-to-end behavioral displacement"],
        ],
        col_widths=[1.2, 2.8, 2.5],
    )

    add_callout(doc, "EDGE CASE: If fewer than 2 snapshots, returns zero-filled features "
        "with stability=1.0, trend_consistency=1.0, current_regime='stable'.")

    add_section_heading(doc, "11.3. Regime Detection Rules", level=2)

    add_body_text(doc, "Exact conditions (evaluated in order; first match wins):", bold=True)

    create_table(doc,
        ["Priority", "Condition", "Regime Label"],
        [
            ["1", "regime_shifts > 0", "regime_shift"],
            ["2", "acceleration > 0.01 AND trend_consistency > 0.3", "accelerating"],
            ["3", "total_drift > 0.05 AND stability > 0.7", "drifting"],
            ["4", "else (default)", "stable"],
        ],
        col_widths=[0.7, 3.3, 1.5],
    )

    add_body_text(doc, "Per-zone trajectory analysis:", bold=True)
    add_body_text(doc, (
        "The same 6 trajectory features and regime detection are computed independently "
        "for each of the 5 zone embedding series via compute_zone_trajectories(). This "
        "enables detection signatures like 'identity zone stable but network_footprint "
        "zone in regime_shift' = APT/C2 pattern."
    ))

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  12. RELATIONSHIP EMBEDDINGS
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "12. Relationship Embeddings", level=1)

    add_body_text(doc, (
        "Cross-entity relationships are captured via Hadamard (element-wise) products of "
        "composite embeddings. The resulting vector encodes the interaction pattern between "
        "two entities. When a C2 beacon changes how a user interacts with their device, "
        "the relationship vector drifts even if neither entity individually changes. "
        "Source: models/relationship_embeddings.py."
    ))

    add_section_heading(doc, "12.1. Hadamard Product", level=2)

    add_code_block(doc,
        "def hadamard(a, b):\n"
        "    product = a.astype(float64) * b.astype(float64)  # Element-wise multiply\n"
        "    norm = np.linalg.norm(product)\n"
        "    if norm < 1e-8:\n"
        "        return zeros_like(product, float32)  # Edge case: zero vector\n"
        "    return (product / norm).astype(float32)   # L2-normalized\n"
        "\n"
        "# Formula: relationship = normalize(a odot b)\n"
        "# where odot is element-wise (Hadamard) product"
    )

    add_callout(doc, "EDGE CASE: Returns zero vector if ||a * b|| < 1e-8. This occurs "
        "when the entity embeddings are orthogonal (no shared semantic dimensions).")

    add_section_heading(doc, "12.2. Relationship Types", level=2)

    create_table(doc,
        ["Relationship", "Formula", "Function"],
        [
            ["UserDevice", "user_emb * device_emb, L2-normalized", "compute_user_device_vector()"],
            ["UserApp", "user_emb * app_emb, L2-normalized", "compute_user_app_vector()"],
            ["DeviceSegment", "device_emb * segment_emb, L2-normalized", "compute_device_segment_vector()"],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    add_section_heading(doc, "12.3. Relationship Drift", level=2)

    add_code_block(doc,
        "def relationship_drift(rel_old, rel_new):\n"
        "    # Cosine distance between two relationship snapshots\n"
        "    # Range [0, 2]: 0 = identical, 1 = orthogonal, 2 = opposite\n"
        "    dot = np.dot(rel_old, rel_new)\n"
        "    na = np.linalg.norm(rel_old)\n"
        "    nb = np.linalg.norm(rel_new)\n"
        "    if na < 1e-8 or nb < 1e-8:\n"
        "        return 0.0\n"
        "    return float(1.0 - dot / (na * nb))"
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  13. DATABASE SCHEMA
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "13. Database Schema", level=1)

    add_body_text(doc, (
        "PostgreSQL 16 with pgvector extension for vector(1536) columns. HNSW indexes on "
        "composite vectors for similarity search. Schema defined in db/schema.sql."
    ))

    add_section_heading(doc, "13.1. Entity Tables", level=2)

    create_table(doc,
        ["Table", "Primary Key", "Key Columns"],
        [
            ["users", "user_id (UUID)", "username, user_type, department, role, clearance_level, manager_id, hire_date, is_active"],
            ["devices", "device_id (UUID)", "hostname, device_type, os_type, os_version, ip_address, mac_address, segment_id, owner_user_id, criticality"],
            ["network_segments", "segment_id (UUID)", "segment_name, subnet_cidr (CIDR), vlan_id, zone_type, trust_level, gateway_ip, connected_segments (TEXT[])"],
            ["applications", "app_id (UUID)", "app_name, app_type, owner_team, data_classification, hosting_segment, criticality"],
            ["sessions", "session_id (UUID)", "user_id, device_id, start_time, end_time, src_ip, auth_method, risk_score, is_active"],
        ],
        col_widths=[1.3, 1.2, 4.0],
    )

    add_section_heading(doc, "13.2. Log Tables (High-Volume)", level=2)

    create_table(doc,
        ["Table", "Key Columns", "Indexed On"],
        [
            ["auth_logs", "timestamp, user_id, device_id, source_ip, dest_system, success, auth_method, failure_reason, geo_location", "(user_id, timestamp), (device_id, timestamp)"],
            ["network_flows", "timestamp, src_ip, dst_ip, src_port, dst_port, protocol, bytes_in, bytes_out, duration_ms, tcp_flags", "(device_id, timestamp), (segment_id, timestamp)"],
            ["dns_queries", "timestamp, client_ip, query_domain, query_type, response_code, response_ip, device_id", "(device_id, timestamp)"],
            ["endpoint_telemetry", "timestamp, device_id, process_name, parent_process, command_line, user_id, file_access_path, event_type", "(device_id, timestamp), (user_id, timestamp)"],
            ["file_access_logs", "timestamp, user_id, file_path, operation, bytes, sensitivity_level, device_id, app_id", "(user_id, timestamp), (device_id, timestamp)"],
            ["app_activity_logs", "timestamp, user_id, app_id, action, resource, status_code, response_time_ms, data_volume_bytes", "(user_id, timestamp), (app_id, timestamp)"],
            ["privilege_operations", "timestamp, user_id, operation, target_resource, previous_level, new_level, approved_by, justification", "(user_id, timestamp)"],
        ],
        col_widths=[1.3, 3.2, 2.0],
    )

    add_section_heading(doc, "13.3. Behavioral Embedding Tables", level=2)

    add_body_text(doc, "Each entity type has a dedicated behavioral embedding table:", bold=True)

    create_table(doc,
        ["Table", "Entity", "Signal Columns (all vector(1536))"],
        [
            ["user_behavioral_embeddings", "users", "beh_authentication, beh_privilege, beh_data_access, beh_network, beh_communication, beh_composite"],
            ["device_behavioral_embeddings", "devices", "beh_process, beh_network_traffic, beh_resource, beh_authentication, beh_configuration, beh_composite"],
            ["segment_behavioral_embeddings", "segments", "beh_traffic_volume, beh_connections, beh_protocols, beh_threat_indicators, beh_service_exposure, beh_composite"],
            ["app_behavioral_embeddings", "apps", "beh_access_patterns, beh_query_behavior, beh_error_rates, beh_performance, beh_configuration, beh_composite"],
            ["session_behavioral_embeddings", "sessions", "beh_activity, beh_risk_accum, beh_data_movement, beh_lateral, beh_temporal, beh_composite"],
        ],
        col_widths=[2.2, 0.8, 3.5],
    )

    add_section_heading(doc, "13.4. Trajectory and Detection Tables", level=2)

    add_body_text(doc, "behavioral_snapshots (weekly embedding snapshots):", bold=True)
    create_table(doc,
        ["Column", "Type", "Description"],
        [
            ["snapshot_id", "BIGSERIAL PK", "Auto-incrementing primary key"],
            ["entity_type", "TEXT NOT NULL", "user, device, segment, app, session"],
            ["entity_id", "UUID NOT NULL", "References entity table"],
            ["cutoff_date", "DATE NOT NULL", "Week ending date"],
            ["beh_signal_1 .. beh_signal_5", "vector(1536)", "5 individual signal embeddings"],
            ["beh_composite", "vector(1536)", "Weighted composite embedding"],
            ["computed_at", "TIMESTAMPTZ", "When snapshot was computed"],
        ],
        col_widths=[2.0, 1.5, 3.0],
    )

    add_body_text(doc, "trajectory_events (detected behavioral events):", bold=True)
    create_table(doc,
        ["Column", "Type", "Description"],
        [
            ["event_id", "BIGSERIAL PK", "Auto-incrementing primary key"],
            ["entity_type", "TEXT NOT NULL", "Entity type"],
            ["entity_id", "UUID NOT NULL", "Entity reference"],
            ["event_date", "TIMESTAMPTZ NOT NULL", "When the event occurred"],
            ["event_type", "TEXT NOT NULL", "behavioral_shift, anomaly, concept_alignment, cusum_alert"],
            ["severity", "TEXT NOT NULL", "critical, high, medium, low, info"],
            ["magnitude", "NUMERIC", "Drift magnitude (cosine distance)"],
            ["drift_concept", "TEXT", "Reference concept name"],
            ["concept_alignment", "NUMERIC", "Cosine similarity with concept vector"],
            ["contributing_signals", "JSONB", "Per-signal breakdown"],
            ["mitre_techniques", "TEXT[]", "Associated MITRE ATT&CK technique IDs"],
            ["kill_chain_id", "UUID FK", "References kill_chain_sequences"],
        ],
        col_widths=[1.5, 1.5, 3.5],
    )

    add_body_text(doc, "kill_chain_sequences (correlated attack narratives):", bold=True)
    create_table(doc,
        ["Column", "Type", "Description"],
        [
            ["chain_id", "UUID PK", "Unique identifier"],
            ["chain_name", "TEXT NOT NULL", "Human-readable chain name"],
            ["start_time", "TIMESTAMPTZ NOT NULL", "When chain was first detected"],
            ["end_time", "TIMESTAMPTZ", "When chain was resolved or closed"],
            ["status", "TEXT NOT NULL", "active, resolved, false_positive"],
            ["confidence", "NUMERIC", "Overall confidence score"],
            ["involved_users", "TEXT[]", "Array of user IDs"],
            ["involved_devices", "TEXT[]", "Array of device IDs"],
            ["involved_segments", "TEXT[]", "Array of segment IDs"],
            ["attack_narrative", "TEXT", "AI-generated attack narrative"],
            ["mitre_tactics", "TEXT[]", "Mapped MITRE tactics"],
        ],
        col_widths=[1.5, 1.5, 3.5],
    )

    add_section_heading(doc, "13.5. HNSW Indexes", level=2)

    add_body_text(doc, (
        "HNSW (Hierarchical Navigable Small World) indexes on beh_composite columns "
        "for approximate nearest-neighbor search using cosine distance."
    ))

    add_code_block(doc,
        "CREATE INDEX idx_user_beh_composite_hnsw\n"
        "    ON user_behavioral_embeddings\n"
        "    USING hnsw (beh_composite vector_cosine_ops)\n"
        "    WITH (m = 16, ef_construction = 64);\n"
        "\n"
        "-- Same pattern for: device, segment, app, session, behavioral_snapshots\n"
        "-- Parameters: m=16 (max connections per node), ef_construction=64 (build quality)"
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  14. CONFIGURATION REFERENCE
    # ══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "14. Configuration Reference", level=1)

    add_section_heading(doc, "14.1. Tier3Config Defaults", level=2)

    add_body_text(doc, (
        "All Tier 3 detection thresholds are defined in the Tier3Config dataclass. "
        "Source: models/cyber_entity.py."
    ))

    create_table(doc,
        ["Parameter", "Default Value", "Used By"],
        [
            ["acceleration_threshold", "0.01", "Regime detection: minimum acceleration for 'accelerating' regime"],
            ["trend_consistency_min", "0.5", "Regime detection: minimum trend consistency for 'accelerating'"],
            ["regime_shift_threshold", "0.7", "Regime detection: cosine similarity below this = phase change"],
            ["zone_stable_threshold", "0.02", "Zone divergence: identity drift below this = stable profile"],
            ["zone_drifting_threshold", "0.05", "Zone divergence: behavioral zone above this = drifting"],
            ["relationship_drift_threshold", "0.05", "Relationship drift: minimum drift to flag"],
            ["contextual_threat_threshold", "0.30", "Contextual detection: minimum threat consistency"],
            ["cohort_similarity", "0.5", "Cross-entity: cosine similarity threshold for clustering"],
            ["cohort_min_size", "3", "Cross-entity: minimum users in a correlated cohort"],
            ["threat_consistency_threshold", "0.40", "Contextual + progression: minimum threat-winning week fraction"],
        ],
        col_widths=[2.2, 1.0, 3.3],
    )

    add_section_heading(doc, "14.2. PhaseState Defaults", level=2)

    add_body_text(doc, (
        "Default values for a newly-initialized PhaseState (no temporal data yet). "
        "Source: models/cyber_entity.py."
    ))

    create_table(doc,
        ["Field", "Default Value", "Type"],
        [
            ["velocity_vector", "np.zeros(1536, dtype=np.float32)", "ndarray"],
            ["velocity_magnitude", "0.0", "float"],
            ["acceleration", "0.0", "float"],
            ["stability", "1.0", "float"],
            ["regime_shifts", "0.0", "float"],
            ["trend_consistency", "1.0", "float"],
            ["total_drift", "0.0", "float"],
            ["current_regime", "\"stable\"", "str"],
        ],
        col_widths=[1.8, 2.5, 1.0],
    )

    add_section_heading(doc, "14.3. Environment Variables", level=2)

    create_table(doc,
        ["Variable", "Default Value", "Description"],
        [
            ["DATABASE_URL", "postgresql://cyber_ueba:password@db:5432/cyber_ueba", "PostgreSQL connection string"],
            ["DB_HOST", "db", "Database hostname (Docker: db, external: localhost)"],
            ["DB_PORT", "5432", "Database port (internal Docker; external maps to 5437)"],
            ["DB_NAME", "cyber_ueba", "Database name"],
            ["DB_USER", "cyber_ueba", "Database username"],
            ["DB_PASSWORD", "password", "Database password"],
            ["OPENAI_API_KEY", "(required)", "OpenAI API key for text-embedding-3-small"],
            ["EMBEDDING_MODEL", "text-embedding-3-small", "OpenAI embedding model name"],
            ["EMBEDDING_DIM", "1536", "Embedding vector dimensionality"],
            ["APP_ENV", "development", "Application environment"],
            ["LOG_LEVEL", "INFO", "Logging verbosity"],
            ["SNAPSHOT_INTERVAL_DAYS", "30", "Days between behavioral snapshots"],
            ["ALERT_DRIFT_THRESHOLD", "0.15", "Minimum drift magnitude to generate alert"],
            ["CUSUM_THRESHOLD", "0.05", "CUSUM accumulation trigger threshold"],
            ["CUSUM_DRIFT_ALLOWANCE", "0.02", "CUSUM drift allowance (baseline drift per step)"],
        ],
        col_widths=[2.0, 2.5, 2.0],
    )

    add_section_heading(doc, "14.4. Drift Detection Parameters", level=2)

    add_body_text(doc, "Production drift detection (services/scheduler.py):", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["ALERT_DRIFT_THRESHOLD", "0.15", "Cosine distance threshold for generating an alert"],
            ["CUSUM_THRESHOLD", "0.05", "CUSUM accumulation level to trigger change-point detection"],
            ["CUSUM_DRIFT_ALLOWANCE", "0.02", "Expected baseline drift per snapshot interval (subtracted from CUSUM)"],
        ],
        col_widths=[2.2, 0.6, 3.7],
    )

    add_body_text(doc, "Tier 2 comparison pipeline (comparison/run_comparison.py):", bold=True)
    create_table(doc,
        ["Parameter", "Value", "Context"],
        [
            ["CUSUM threshold (magnitude)", "0.001", "On cosine distance between consecutive weekly composites"],
            ["CUSUM drift_threshold (magnitude)", "0.0005", "Baseline drift allowance per week"],
            ["min_run_length (magnitude)", "2", "Consecutive weeks above threshold"],
            ["CUSUM threshold (net threat)", "0.5", "On (max_threat - max_benign) series"],
            ["CUSUM drift_threshold (net threat)", "0.05", "Baseline net threat allowance"],
            ["threat_consistency threshold", "0.40", "Fraction of threat-winning weeks for detection"],
            ["threat_winning margin", "0.05", "max_threat - max_benign > 0.05 for a week to count"],
        ],
        col_widths=[2.5, 0.6, 3.4],
    )

    add_section_heading(doc, "14.5. Simulation Parameters", level=2)

    add_body_text(doc, "Source: simulator/config.py.", bold=True)
    create_table(doc,
        ["Parameter", "Value"],
        [
            ["SIM_START", "2025-01-01"],
            ["SIM_END", "2026-05-01"],
            ["SIM_DAYS", "486 (computed)"],
            ["N_USERS", "250"],
            ["N_DEVICES", "400"],
            ["N_SEGMENTS", "25"],
            ["N_APPLICATIONS", "60"],
            ["N_ROLES", "30"],
            ["USER_TYPE_DIST", "employee: 80%, contractor: 15%, service_account: 5%"],
            ["AUTH_EVENTS_PER_USER_DAY", "8"],
            ["NETWORK_FLOWS_PER_DEVICE_DAY", "40"],
            ["DNS_QUERIES_PER_DEVICE_DAY", "10"],
            ["ENDPOINT_EVENTS_PER_DEVICE_DAY", "8"],
            ["FILE_ACCESS_PER_USER_DAY", "10"],
            ["APP_EVENTS_PER_USER_DAY", "5"],
            ["WORK_HOURS", "(8, 18)"],
            ["WORK_DAYS", "[0, 1, 2, 3, 4] (Mon-Fri)"],
            ["Random Seed", "42"],
        ],
        col_widths=[2.5, 4.0],
    )

    add_section_heading(doc, "14.6. Attack Scenario Configuration", level=2)

    create_table(doc,
        ["Attack ID", "Type", "Target", "Start Date", "Duration/Details"],
        [
            ["ATK-001", "brute_force", "50 users", "2026-03-15", "4 hours, credential stuffing from external botnet"],
            ["ATK-002", "credential_theft_lateral", "USR-087", "2026-02-01", "5 days, phishing to credential theft to lateral movement"],
            ["ATK-003", "apt_slow", "USR-234", "2025-03-10", "417 days, C2 beacon every 360 min, DGA prob 0.6, staging every 5 days"],
            ["ATK-004", "insider_threat", "USR-156", "2025-03-10", "14 months escalation, gradual access expansion"],
            ["ATK-005", "ransomware", "DEV-342", "2026-04-10", "6 hours, deployment via compromised RDP"],
            ["ATK-006", "supply_chain", "APP-017", "2025-10-01", "90 days dormancy, then backdoor activation"],
            ["ATK-007", "volt_typhoon", "USR-042", "2025-03-15", "412 days, LOTL + credential harvest, proxy 203.0.113.88"],
            ["ATK-008", "salt_typhoon", "USR-118", "2025-03-15", "412 days, CDR exfil, C2 domain cdn-check.microsoftupdate-service.com"],
        ],
        col_widths=[0.7, 1.2, 0.7, 0.9, 3.0],
    )

    add_page_break(doc)

    # ── END MATTER ─────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Implementation Specification —")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROPRIETARY — 22nd Century Technologies, Inc.")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_ACCENT

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "This document contains trade secrets and proprietary information.\n"
        "Do not distribute outside the organization.\n"
        "Unauthorized disclosure may result in legal action."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "22nd Century Technologies, Inc.\n"
        "V-Intelligence UEBA Program\n"
        "Implementation Specification — Version 1.0\n"
        "June 2026"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY

    # ── SAVE ───────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Implementation specification created successfully")
    print(f"Output: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
