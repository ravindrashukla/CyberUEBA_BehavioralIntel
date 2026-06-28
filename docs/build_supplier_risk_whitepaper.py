#!/usr/bin/env python3
"""Build Supplier Risk SCM Whitepaper (Word Document).

Whitepaper on using Behavioral Entity Intelligence (BEI) for Supplier Risk
Management in Defense Supply Chains.  Targets DLA leadership.

Includes a Technical Deep Dive (Section 8) with the explicit mathematics and
algorithms — embedding, zone composition + context-adaptive attention, drift
direction/magnitude, reference-concept projection, CUSUM change-point, peer-cohort
z-scores, multi-phase composite scoring, and Hadamard relationship embeddings —
so the real innovation is explicit. (For an audience-tunable disclosure level, see
build_innovation_whitepaper.py.)

Output: docs/Supplier_Risk_SCM_Whitepaper.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
BLUE = RGBColor(0x1B, 0x4F, 0x72)
TEAL = RGBColor(0x0E, 0x6B, 0x8A)
GOLD = RGBColor(0xB7, 0x95, 0x0B)
RED_ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREEN_ACCENT = RGBColor(0x1E, 0x8A, 0x49)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)

OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
                           "Supplier_Risk_SCM_Whitepaper.docx")


# ── Helper functions (same pattern as build_tech_spec.py / build_ueba_whitepaper.py) ──

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
    return h


def add_body_text(doc, text, bold=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def add_callout(doc, text, border_color_hex="0E6B8A"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="12" w:color="{border_color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EAF4F7" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    run.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    return p


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "0D1B2A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EAF0F6")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


# ==========================================================================
# DOCUMENT BUILDER
# ==========================================================================

# ==========================================================================
# FIGURE ENGINE (rich visuals from the real synthetic-MVP supplier-risk results)
# ==========================================================================
import os as _os
FIGDIR = _os.path.join(_os.path.dirname(__file__), "figures")
SUP_FIGS = {}

# Verified-from-app synthetic MVP supplier-risk metrics (model scorecard)
_SUP_DISC = [("XGBoost AUC (V1)", 0.7709), ("XGBoost AUC (V2)", 0.7780), ("RSF C-index", 0.7935)]
_SUP_CALIB = [("30 d", 0.0841), ("60 d", 0.0904), ("90 d", 0.1294)]
_SUP_SHAP = [("quality_incidents_6m", 0.1653), ("order_count_90d", 0.1436),
             ("ontime_rate_90d", 0.0957), ("worst_delay_90d", 0.0887),
             ("critical_nsn_count", 0.0887)]


def generate_supplier_figures():
    figs = {}
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyBboxPatch
        import numpy as np
    except Exception:
        return figs
    _os.makedirs(FIGDIR, exist_ok=True)
    NV, BL, TL, RD, GR = "#0D1B2A", "#1B4F72", "#0E6B8A", "#C0392B", "#1E8A49"

    def _flow(name, title, boxes, color, figkey):
        fig, ax = plt.subplots(figsize=(9.8, 1.9)); ax.axis("off"); ax.set_xlim(0, 10); ax.set_ylim(0, 1.4)
        ax.text(0.1, 1.28, title, fontsize=10, color=NV, weight="bold")
        n = len(boxes); gap = 9.7 / n; w = gap * 0.85
        for i, t in enumerate(boxes):
            x = 0.15 + i * gap
            ax.add_patch(FancyBboxPatch((x, 0.25), w, 0.78, boxstyle="round,pad=0.02", fc=color, ec="none"))
            ax.text(x + w / 2, 0.64, t, ha="center", va="center", color="white", fontsize=6.8)
            if i < n - 1:
                ax.annotate("", xy=(x + gap, 0.64), xytext=(x + w, 0.64),
                            arrowprops=dict(arrowstyle="->", color="#666", lw=1.2))
        p = _os.path.join(FIGDIR, name); fig.savefig(p, dpi=150, bbox_inches="tight"); plt.close(fig)
        figs[figkey] = p

    # 1 — BEI supplier pipeline
    _flow("fig_sup_pipeline.png", "Behavioral Entity Intelligence — supplier-risk pipeline",
          ["Supplier data\n(orders, quality,\nfinancial)", "5 behavioral\nzones", "1536-d\nembedding",
           "Cohort z-score\n+ drift (CUSUM)", "XGBoost P(fail≤90d)\n+ RSF survival",
           "Ranked risk\n+ survival curve"], TL, "pipeline")

    # 2 — five supplier zones
    try:
        zones = [("Identity", "CAGE/UEI, certs,\ntier"), ("Performance", "on-time, lead time,\nquality"),
                 ("Geographic", "country-of-origin,\nadversary exposure"),
                 ("Network", "depots, NSNs,\nconcentration"), ("Risk", "financial health,\nbond, contract")]
        fig, ax = plt.subplots(figsize=(9.6, 1.9)); ax.axis("off"); ax.set_xlim(0, 10); ax.set_ylim(0, 1.4)
        ax.text(0.1, 1.30, "Supplier digital entity — five behavioral zones", fontsize=10, color=NV, weight="bold")
        n = len(zones); gap = 9.7 / n; w = gap * 0.9
        for i, (z, sub) in enumerate(zones):
            x = 0.15 + i * gap
            ax.add_patch(FancyBboxPatch((x, 0.2), w, 0.86, boxstyle="round,pad=0.02", fc=BL, ec="none"))
            ax.text(x + w / 2, 0.86, z, ha="center", va="center", color="white", fontsize=8, weight="bold")
            ax.text(x + w / 2, 0.5, sub, ha="center", va="center", color="white", fontsize=6)
        p = _os.path.join(FIGDIR, "fig_sup_zones.png"); fig.savefig(p, dpi=150, bbox_inches="tight")
        plt.close(fig); figs["zones"] = p
    except Exception:
        pass

    # 3 — discrimination (AUC / C-index)
    try:
        names = [d[0] for d in _SUP_DISC]; vals = [d[1] for d in _SUP_DISC]
        fig, ax = plt.subplots(figsize=(7.4, 3.2))
        ax.bar(names, vals, color=[TL, TL, GR])
        ax.axhline(0.5, color="#999", ls="--", lw=0.9, label="chance (0.5)")
        ax.set_ylim(0.5, 1.0); ax.set_ylabel("Discrimination")
        for i, v in enumerate(vals):
            ax.text(i, v + 0.006, f"{v:.3f}", ha="center", fontsize=8)
        ax.set_title("Supplier-risk discrimination (synthetic): all > 0.77", fontsize=10, color=NV)
        ax.legend(fontsize=7)
        p = _os.path.join(FIGDIR, "fig_sup_disc.png"); fig.savefig(p, dpi=150, bbox_inches="tight")
        plt.close(fig); figs["disc"] = p
    except Exception:
        pass

    # 4 — RSF survival calibration by horizon
    try:
        names = [c[0] for c in _SUP_CALIB]; vals = [c[1] for c in _SUP_CALIB]
        fig, ax = plt.subplots(figsize=(7.0, 3.0))
        ax.bar(names, vals, color=TL)
        for i, v in enumerate(vals):
            ax.text(i, v + 0.003, f"{v:.3f}", ha="center", fontsize=8)
        ax.set_ylabel("Expected Calibration Error (lower better)")
        ax.set_title("Survival-curve calibration by horizon (synthetic)", fontsize=10, color=NV)
        p = _os.path.join(FIGDIR, "fig_sup_calib.png"); fig.savefig(p, dpi=150, bbox_inches="tight")
        plt.close(fig); figs["calib"] = p
    except Exception:
        pass

    # 5 — top SHAP feature importance
    try:
        names = [s[0] for s in _SUP_SHAP][::-1]; vals = [s[1] for s in _SUP_SHAP][::-1]
        fig, ax = plt.subplots(figsize=(7.6, 3.0))
        ax.barh(names, vals, color=TL)
        ax.set_xlabel("mean |SHAP| (feature importance)")
        ax.set_title("What drives supplier risk — top features (synthetic)", fontsize=10, color=NV)
        p = _os.path.join(FIGDIR, "fig_sup_shap.png"); fig.savefig(p, dpi=150, bbox_inches="tight")
        plt.close(fig); figs["shap"] = p
    except Exception:
        pass

    # 6 — peer-cohort self-vs-cohort drift (illustrative)
    try:
        t = np.arange(0, 12)
        cohort = np.full(12, 0.10)
        fig, ax = plt.subplots(figsize=(7.8, 3.2))
        ax.fill_between(t, cohort - 0.04, cohort + 0.04, color="#cbd5e1", alpha=0.7, label="Cohort normal band")
        ax.plot(t, cohort, color="#64748b", lw=1, label="Cohort baseline")
        supplier = np.concatenate([np.linspace(0.10, 0.12, 6), np.linspace(0.13, 0.30, 6)])
        ax.plot(t, supplier, color=RD, lw=2, label="This supplier (drifting from cohort)")
        ax.set_xlabel("Month"); ax.set_ylabel("Behavioral drift")
        ax.set_title("Peer-cohort drift: supplier separating from its cohort (illustrative)",
                     fontsize=10, color=NV)
        ax.legend(fontsize=7, loc="upper left")
        p = _os.path.join(FIGDIR, "fig_sup_cohort.png"); fig.savefig(p, dpi=150, bbox_inches="tight")
        plt.close(fig); figs["cohort"] = p
    except Exception:
        pass

    return figs


def add_figure(doc, key, caption):
    path = SUP_FIGS.get(key)
    if not path or not _os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = DARK_GRAY
    cap.paragraph_format.space_after = Pt(8)


def build():
    global SUP_FIGS
    SUP_FIGS = generate_supplier_figures()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ══════════════════════════════════════════════════════════════
    #  TITLE PAGE
    # ══════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Behavioral Entity Intelligence\nfor Supplier Risk Management")
    run.font.size = Pt(30)
    run.font.color.rgb = NAVY
    run.bold = True

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Detecting Hidden Supply Chain Risks Through\n"
        "Entity Profiling, Relationship Dynamics, and Network Analysis"
    )
    run.font.size = Pt(15)
    run.font.color.rgb = BLUE

    for _ in range(4):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "22nd Century Technologies, Inc.\n\n"
        "June 2026 — Version 1.0\n\n"
        "Prepared for: Defense Logistics Agency"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "Table of Contents", level=1)

    toc_entries = [
        "1. Executive Summary",
        "2. The Supplier Risk Challenge",
        "   2.1. Why Traditional Supplier Monitoring Fails",
        "   2.2. Real-World Supply Chain Compromises",
        "   2.3. SCRM Regulatory Landscape",
        "3. Behavioral Entity Intelligence for Supply Chain",
        "   3.1. The Entity Model",
        "   3.2. Supplier Behavioral Zones",
        "   3.3. NSN Item Behavioral Zones",
        "   3.4. The Time Dimension",
        "   3.5. Supplier Segmentation and Peer-Cohort Drift",
        "4. Relationship-Level Risk Detection",
        "   4.1. Pairwise Relationships",
        "   4.2. Network-Level Risk",
        "   4.3. The What vs How Much Principle",
        "5. Detection Capabilities",
        "   5.1. Counterfeit Part Detection",
        "   5.2. Adversary-Nation Sourcing (SCRM)",
        "   5.3. Supplier Concentration Risk",
        "   5.4. Demand Diversion and Fraud",
        "   5.5. Contract and Financial Risk",
        "6. Operational Deployment",
        "   6.1. Data Sources and Collection",
        "   6.2. Deployment Phases",
        "   6.3. Integration Points",
        "7. Proven Foundation: From Cybersecurity to Supply Chain",
        "   7.1. What Has Been Proven in Cybersecurity",
        "   7.2. Why the Proof Transfers",
        "   7.3. Supply Chain Proof of Concept",
        "   7.4. Illustrative Detection Scenarios",
        "8. Technical Deep Dive: The Mathematics of Behavioral Drift",
        "   8.1. From Metrics to Meaning: Serialization and Embedding",
        "   8.2. Zone Composition and Context-Adaptive Attention",
        "   8.3. Drift: Measuring Direction, Not Just Magnitude",
        "   8.4. Naming the Direction: Reference-Concept Projection",
        "   8.5. Accumulating Slow Drift: CUSUM Change-Point",
        "   8.6. Peer-Cohort Scoring",
        "   8.7. Multi-Phase Composite Score",
        "   8.8. Relationship and Co-Drift Detection",
        "   8.9. Why the Composition Is the Innovation",
        "   8.10. The Predictive Risk Layer: XGBoost + Survival Model (Implemented MVP)",
        "9. Business Value",
        "10. Conclusion",
    ]
    for entry in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(entry)
        run.font.size = Pt(10.5)
        run.font.color.rgb = NAVY if not entry.startswith("   ") else BLUE
        p.paragraph_format.space_after = Pt(2)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "1. Executive Summary", level=1)

    add_body_text(doc, (
        "Defense supply chains face risks that are invisible to traditional monitoring. "
        "Suppliers gradually shift component sourcing to adversary nations. Counterfeit "
        "parts enter the supply chain through multi-tier intermediaries. Coordinated "
        "failures reveal hidden common dependencies that no single-supplier audit would "
        "uncover. These risks accumulate slowly, often over months, and become visible "
        "only after mission-critical systems are affected."
    ))

    add_body_text(doc, (
        "Traditional supplier monitoring checks individual metrics — on-time delivery "
        "rates, quality scores, pricing compliance — each against its own threshold. A "
        "supplier can pass every metric while systemic risks grow beneath the surface. "
        "The on-time rate stays within target. The quality score remains acceptable. The "
        "pricing stays within contract bounds. But the supplier has quietly shifted a "
        "substantial share of component sourcing from domestic manufacturers to facilities "
        "in an adversary nation. No single metric captures this."
    ))

    add_body_text(doc, (
        "Behavioral Entity Intelligence (BEI) detects these risks by tracking behavioral "
        "direction across three analytical layers: entity profiles that capture each "
        "supplier's multi-dimensional behavioral trajectory over time, pairwise "
        "relationships that reveal how specific supplier-part-location combinations are "
        "changing, and network graphs that expose hidden dependencies, single points of "
        "failure, and multi-hop adversary-nation connections."
    ))

    add_body_text(doc, (
        "This approach is not theoretical. The BEI framework has been proven in an "
        "operational cybersecurity deployment, where it detected all four advanced, "
        "long-duration attack campaigns — an insider threat, a slow command-and-control "
        "campaign, and two nation-state intrusions. Traditional threshold-based detection "
        "caught at most one of the four; the rest stayed within normal statistical ranges "
        "on every individual metric while shifting behavioral direction. BEI caught the "
        "directional change."
    ))

    add_body_text(doc, (
        "The same concept has since proven out in a second, very different domain — defense "
        "demand forecasting. There, modeling each item-at-a-location as a digital entity and "
        "reading the DIRECTION of its demand drift produced calibrated, drift-aware forecasts "
        "that right-size procurement where a legacy moving average cannot. Two domains, one "
        "pattern: model the entity as a digital twin, track where its behavior is heading, and "
        "act on the direction — not the magnitude of any single metric. Supplier risk is the "
        "natural third application of that proven pattern: a supplier rarely trips a single "
        "threshold, but its behavior drifts — sourcing shifts toward adversary nations, quality "
        "slips relative to its peer cohort, delivery patterns change — and the same approach "
        "surfaces the anomalous supplier before any one metric crosses a red line."
    ))

    add_body_text(doc, (
        "The framework has also been instantiated for the supply chain domain and "
        "exercised on a realistic synthetic dataset — 500 National Stock Numbers, 200 "
        "suppliers, 4 depots spanning the Indo-Pacific and CONUS, and 24 months of "
        "transaction history — demonstrating that the same behavioral profiling, regime "
        "detection, supplier risk grading, and relationship analysis operate directly on "
        "supply chain data. This whitepaper describes how that proven framework applies "
        "to defense supplier risk — and what it makes possible."
    ))

    add_callout(doc,
        "Key insight: The supplier changes WHAT it sources, not HOW MUCH it delivers. "
        "Traditional metrics measure magnitude; BEI measures behavioral direction. The same "
        "principle that exposed hidden cyber threats and sharpened defense demand forecasts now "
        "exposes hidden supply chain risk — the third domain for one proven pattern."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  2. THE SUPPLIER RISK CHALLENGE
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "2. The Supplier Risk Challenge", level=1)

    add_section_heading(doc, "2.1. Why Traditional Supplier Monitoring Fails", level=2)

    add_body_text(doc, (
        "Current supplier monitoring systems operate on a simple model: define a metric, "
        "set a threshold, and alert when the threshold is breached. This approach has four "
        "structural weaknesses that adversaries and market dynamics routinely exploit."
    ))

    add_body_text(doc, "Threshold Blindness", bold=True, space_after=2)
    add_body_text(doc, (
        "A supplier passes every individual threshold — on-time delivery, quality score, "
        "and pricing all within bounds — while simultaneously shifting component sourcing, "
        "reducing manufacturing oversight, and substituting materials. Each metric is "
        "within bounds. The composite behavioral trajectory tells a different story, but "
        "no system is watching the trajectory."
    ))

    add_body_text(doc, "Single-Entity Tunnel Vision", bold=True, space_after=2)
    add_body_text(doc, (
        "Traditional monitoring evaluates suppliers in isolation and parts in isolation. "
        "Supplier A looks fine. Part X looks fine. But the relationship between Supplier A "
        "and Part X has changed: the failure rate for this specific supplier-part "
        "combination has risen sharply while the supplier's overall quality score remains "
        "stable. The risk lives in the relationship, not in either entity."
    ))

    add_body_text(doc, "Static Risk Scores", bold=True, space_after=2)
    add_body_text(doc, (
        "Annual risk assessments produce a point-in-time snapshot. Between assessments, "
        "supplier behavior can drift significantly. A supplier assessed as low-risk in "
        "January may experience financial distress in March, begin substituting components "
        "in May, and deliver counterfeit parts in July — all before the next annual "
        "review. Continuous behavioral monitoring closes this gap."
    ))

    add_body_text(doc, "Tier-2/3 Opacity", bold=True, space_after=2)
    add_body_text(doc, (
        "Prime suppliers pass audits, but their sub-tier sourcing is effectively invisible. "
        "A Tier-1 supplier in good standing may source critical components from a Tier-2 "
        "distributor who sources from a Tier-3 manufacturer in an adversary nation. The "
        "prime supplier's metrics look clean because the risk is two hops away."
    ))

    add_section_heading(doc, "2.2. Real-World Supply Chain Compromises", level=2)

    add_body_text(doc, (
        "The following incidents illustrate the types of supply chain risks that "
        "traditional monitoring consistently fails to detect."
    ))

    add_body_text(doc, "Counterfeit Electronic Components in DoD Systems", bold=True, space_after=2)
    add_body_text(doc, (
        "The Senate Armed Services Committee investigation (2012) identified approximately "
        "1 million suspected counterfeit electronic parts in the Department of Defense "
        "supply chain. These counterfeits entered through multi-tier distribution networks "
        "where components were relabeled, remarked, or entirely fabricated. Many passed "
        "initial quality inspections because the counterfeits were functionally adequate "
        "for standard testing but failed under operational stress conditions."
    ))

    add_body_text(doc, "Section 889 Violations", bold=True, space_after=2)
    add_body_text(doc, (
        "Section 889 of the National Defense Authorization Act for Fiscal Year 2019 "
        "prohibits federal agencies from procuring telecommunications equipment from "
        "specified Chinese companies (Huawei, ZTE, and others). Despite the prohibition, "
        "components from these manufacturers have been found embedded in systems sold to "
        "federal agencies through multi-tier supply chains where the ultimate origin was "
        "obscured by intermediary distributors."
    ))

    add_body_text(doc, "SolarWinds Supply Chain Attack", bold=True, space_after=2)
    add_body_text(doc, (
        "The SolarWinds compromise (2020) demonstrated that software supply chains carry "
        "the same risks as physical supply chains. A compromised software build system "
        "injected malicious code into a legitimate software update, affecting approximately "
        "18,000 organizations including multiple federal agencies. The attack succeeded "
        "because the SBOM (Software Bill of Materials) was compromised at the build stage "
        "— the software appeared legitimate at every inspection point."
    ))

    add_body_text(doc, "COVID-19 Supply Chain Disruptions", bold=True, space_after=2)
    add_body_text(doc, (
        "The pandemic exposed hidden single points of failure across defense supply "
        "chains. Suppliers that appeared independent shared common sub-tier sources for "
        "raw materials, specialized components, or manufacturing capacity. When one "
        "facility shut down, multiple supposedly independent supply lines failed "
        "simultaneously, revealing concentration risks invisible to supplier-by-supplier "
        "monitoring."
    ))

    add_section_heading(doc, "2.3. SCRM Regulatory Landscape", level=2)

    add_body_text(doc, (
        "Defense supply chain risk management is governed by an expanding set of "
        "regulations, each addressing a specific dimension of supply chain risk. BEI "
        "provides the continuous monitoring capability that these regulations require "
        "but current systems struggle to deliver."
    ))

    create_table(doc,
        ["Regulation / Directive", "Scope", "BEI Relevance"],
        [
            ["DFARS 252.246-7008", "Counterfeit electronic parts\ndetection and avoidance",
             "BEI tracks supplier-part-manufacturer\nrelationship drift to detect sourcing\n"
             "changes indicative of counterfeit risk"],
            ["Section 889, NDAA FY2019", "Prohibition on Chinese telecom\ncomponents in federal systems",
             "BEI monitors country-of-origin\nrelationships continuously, not just\nat contract award"],
            ["Executive Order 14017", "Supply chain resilience\nacross critical sectors",
             "BEI identifies hidden concentration\nrisks and single points of failure\n"
             "through network analysis"],
            ["NIST SP 800-161 Rev 1", "Cybersecurity Supply Chain\nRisk Management (C-SCRM)",
             "BEI provides the continuous\nassessment and monitoring\n"
             "layer that C-SCRM frameworks require"],
            ["DLA SCRM Program", "DLA-specific supply chain\nrisk management processes",
             "BEI integrates with existing DLA\nSCRM workflows, adding behavioral\n"
             "intelligence to current risk scoring"],
        ],
        col_widths=[1.8, 2.2, 2.8],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  3. BEHAVIORAL ENTITY INTELLIGENCE FOR SUPPLY CHAIN
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "3. Behavioral Entity Intelligence for Supply Chain", level=1)

    add_section_heading(doc, "3.1. The Entity Model", level=2)
    add_figure(doc, "pipeline",
               "Figure 1 — The BEI supplier-risk pipeline: supplier data → five behavioral zones → "
               "1,536-d embedding → cohort-relative z-scores and drift → a calibrated XGBoost + "
               "survival model → a ranked risk score with a time-dependent survival curve.")

    add_body_text(doc, (
        "BEI models the defense supply chain as a network of interacting entities, each "
        "with a multi-dimensional behavioral profile that evolves over time. Four primary "
        "entity types capture the key actors and objects in the supply chain."
    ))

    create_table(doc,
        ["Entity Type", "Examples", "Behavioral Profile"],
        [
            ["Supplier", "Prime contractors, distributors,\nsmall businesses, foreign vendors",
             "Delivery performance, pricing behavior,\ngeographic sourcing patterns,\n"
             "financial health indicators"],
            ["NSN (Part)", "Consumables, repairables,\ncritical components, assemblies",
             "Demand patterns, failure rates,\nsupplier concentration,\n"
             "inventory position trajectories"],
            ["Location\n(Depot/Base)", "DLA depots, military installations,\n"
             "forward operating locations",
             "Demand composition, requisition urgency\npatterns, consumption velocity,\n"
             "seasonal profiles"],
            ["Manufacturer", "Original equipment manufacturers,\ncomponent fabricators,\n"
             "authorized distributors",
             "Production consistency, quality trends,\ncountry-of-origin stability,\n"
             "certification status"],
        ],
        col_widths=[1.2, 2.2, 3.0],
    )

    add_body_text(doc, (
        "Each entity maintains a behavioral profile composed of multiple independent "
        "dimensions — behavioral zones — that are tracked over time. The zones decompose "
        "entity activity into orthogonal dimensions so that a change in one dimension "
        "(e.g., geographic sourcing) is detected even when other dimensions (e.g., "
        "delivery timeliness) remain stable."
    ))

    add_section_heading(doc, "3.2. Supplier Behavioral Zones", level=2)
    add_figure(doc, "zones",
               "Figure 2 — The supplier digital entity, decomposed into five independently-tracked "
               "behavioral zones so a shift in one (e.g. geographic sourcing) is not masked by "
               "stability in the others (e.g. delivery performance).")

    add_body_text(doc, (
        "Each supplier's behavioral profile is decomposed into five zones that capture "
        "independent dimensions of supplier activity. This decomposition ensures that "
        "a change in sourcing geography is detected even if delivery performance "
        "remains exemplary."
    ))

    create_table(doc,
        ["Zone", "What It Captures", "Example Indicators"],
        [
            ["Identity", "Registration, certifications,\ncontract status, tier level",
             "CAGE code stability, SAM.gov status,\n"
             "ISO certification currency,\nsmall business classification"],
            ["Performance\nTrajectory", "Delivery timeliness trends,\nlead time consistency,\n"
             "quality scores over time",
             "On-time delivery rate trajectory,\naverage days early/late trend,\n"
             "quality rejection rate direction"],
            ["Location /\nGeographic Risk", "Sourcing geography,\ngeographic concentration,\n"
             "geopolitical risk exposure",
             "Country-of-origin distribution,\nmanufacturing site diversity,\n"
             "adversary-nation exposure percentage"],
            ["Network\nPosition", "How many depots served,\nhow many items supplied,\n"
             "supplier ecosystem role",
             "Customer concentration index,\nproduct breadth score,\n"
             "sole-source contract count"],
            ["Risk\nAssessment", "Composite risk indicators,\nfinancial health,\n"
             "contract urgency",
             "Altman Z-score trajectory,\ninsurance/bond status,\n"
             "expiring contract percentage"],
        ],
        col_widths=[1.2, 2.0, 3.2],
    )

    add_callout(doc,
        "Why zones matter: A supplier's on-time delivery stays strong (Performance zone "
        "stable) while sourcing shifts from domestic to adversary-nation origins "
        "(Geographic Risk zone drifting). Without zone decomposition, the stable "
        "performance masks the geographic risk."
    )

    add_section_heading(doc, "3.3. NSN Item Behavioral Zones", level=2)

    add_body_text(doc, (
        "Each National Stock Number (NSN) has its own behavioral profile decomposed into "
        "five zones. Tracking items independently from suppliers enables detection of risks "
        "that emerge from the item's lifecycle rather than from any single supplier."
    ))

    create_table(doc,
        ["Zone", "What It Captures", "Example Indicators"],
        [
            ["Identity", "Nomenclature, criticality,\nplatform, supply class",
             "Federal Supply Class, criticality code,\nweapon system association,\n"
             "shelf life classification"],
            ["Demand\nProfile", "Demand patterns, seasonality,\nexercise sensitivity",
             "Monthly demand trajectory,\nseasonal demand coefficients,\n"
             "exercise-driven demand spikes"],
            ["Supply\nNetwork", "Supplier count, sole source\nflag, supplier concentration",
             "Active supplier count trend,\nHerfindahl concentration index,\n"
             "new entrant / exit frequency"],
            ["Inventory\nPosition", "On-hand levels, safety stock\nratios, consumption velocity",
             "Days of supply trajectory,\nsafety stock adequacy ratio,\n"
             "stockout frequency trend"],
            ["Financial\nProfile", "Unit pricing, shelf life,\nsubstitutability",
             "Unit price trajectory,\nprice variance across suppliers,\n"
             "authorized substitute count"],
        ],
        col_widths=[1.2, 2.0, 3.2],
    )

    add_section_heading(doc, "3.4. The Time Dimension", level=2)

    add_body_text(doc, (
        "The critical distinction between BEI and traditional supplier monitoring is "
        "the treatment of time. Traditional systems evaluate each reporting period "
        "independently: is this month's delivery rate above threshold? BEI evaluates "
        "behavioral trajectory: is this month's delivery rate consistent with the "
        "supplier's established pattern, and if not, in which direction is it moving?"
    ))

    add_bullet(doc, "Weekly or monthly behavioral snapshots capture entity state over "
               "time, building a trajectory of behavioral evolution.",
               bold_prefix="Continuous Monitoring: ")
    add_bullet(doc, "Gradual changes invisible to point-in-time audits become visible "
               "when the system tracks cumulative behavioral displacement over weeks "
               "and months.",
               bold_prefix="Drift Detection: ")
    add_bullet(doc, "Each entity is classified into a behavioral regime — Stable, "
               "Drifting, or Critical — based on the velocity, acceleration, and "
               "direction of its behavioral trajectory.",
               bold_prefix="Regime Detection: ")

    add_body_text(doc, (
        "A supplier classified as Drifting in the Geographic Risk zone and Stable in all "
        "other zones is exhibiting the behavioral signature of gradual adversary-nation "
        "sourcing shift — precisely the pattern that annual audits miss."
    ))

    # ── 3.5 Supplier Segmentation and Peer-Cohort Drift ──────────
    add_section_heading(doc, "3.5. Supplier Segmentation and Peer-Cohort Drift", level=2)
    add_figure(doc, "cohort",
               "Figure 3 — Peer-cohort drift (illustrative). Each supplier is scored against the "
               "normal band of its cohort (commodity, geography, criticality); the one separating "
               "from its cohort is surfaced even while every absolute metric stays in bounds.")

    add_body_text(doc, (
        "Tracking a supplier against its own history answers one question — is this supplier "
        "behaving the way it used to? A second, equally important question is: is this "
        "supplier behaving the way its peers do? A behavioral value that is unremarkable for "
        "one kind of supplier can be a strong risk signal for another. A semiconductor "
        "fabricator drawing on a dozen international sources is normal; a domestic fastener "
        "manufacturer suddenly doing the same is anomalous. The signal is meaningful only "
        "relative to comparable suppliers."
    ))

    add_body_text(doc, (
        "BEI therefore segments the supplier population into peer cohorts and evaluates each "
        "supplier's behavioral drift against the baseline of its cohort — not against the "
        "population as a whole. Cohorts are defined along the dimensions that determine what "
        "normal behavior should look like:"
    ))

    add_bullet(doc, "Federal Supply Class / commodity group — electronics, fasteners, "
               "textiles, chemicals, subsistence, aviation spares. What a supplier provides "
               "shapes its normal sourcing geography, lead times, and price volatility.",
               bold_prefix="What they supply: ")
    add_bullet(doc, "Domestic vs. foreign, region, and proximity to adversary-nation "
               "manufacturing. Geography sets the baseline for country-of-origin mix and "
               "logistics behavior.",
               bold_prefix="Where they are located: ")
    add_bullet(doc, "Sole-source status, weapon-system association, and mission-essential "
               "designation. Critical suppliers warrant a tighter cohort baseline and a "
               "lower drift tolerance.",
               bold_prefix="Criticality: ")
    add_bullet(doc, "Prime, distributor, or sub-tier manufacturer, and small-business / "
               "large-business classification, which shape expected breadth and concentration.",
               bold_prefix="Tier and size: ")

    add_body_text(doc, (
        "For each cohort, BEI establishes a cohort behavioral baseline from the suppliers in "
        "that cohort that are behaving normally, then measures how far each individual "
        "supplier has drifted relative to that baseline. A supplier whose drift is "
        "indistinguishable from its cohort is left alone; a supplier whose drift separates "
        "it from its cohort — even while every absolute metric stays within contract bounds "
        "— is surfaced for review."
    ))

    add_callout(doc,
        "Two baselines, one verdict: BEI scores each supplier against its OWN history "
        "(self-baseline drift) and against its COHORT (peer-relative drift). A supplier "
        "flagged on both — drifting from its past and away from its peers — is a "
        "high-confidence risk."
    )

    add_body_text(doc, (
        "This peer-cohort method is not new or unproven for BEI. It is exactly the mechanism "
        "the framework already uses in its operational cybersecurity deployment, where users "
        "are segmented into role groups and each user's behavioral drift is scored relative "
        "to the normal baseline of its group. Suppliers simply substitute for users, and "
        "supply-chain cohorts substitute for workforce roles."
    ))

    create_table(doc,
        ["Element", "Cybersecurity (proven)", "Supplier Risk (applied)"],
        [
            ["Population", "Workforce of users", "Vendor base of suppliers"],
            ["Peer cohort", "Role group: admin, security, developer,\nbusiness, executive",
             "Cohort: commodity / FSC, geography,\ncriticality, tier"],
            ["Cohort baseline", "Normal behavior of users in the role group",
             "Normal behavior of suppliers in the cohort"],
            ["Anomaly signal", "User drifts away from its role group",
             "Supplier drifts away from its cohort"],
            ["Why it works", "An action normal for an admin is alarming\nfor an accountant",
             "A sourcing pattern normal for one commodity\nis alarming for another"],
        ],
        col_widths=[1.3, 2.5, 2.7],
    )

    add_body_text(doc, (
        "Peer-cohort drift also strengthens the relationship and concentration analysis of "
        "Section 4: when several suppliers in the same cohort drift together in the same "
        "direction, the correlated movement points to a shared sub-tier dependency or a "
        "common market or geopolitical shock affecting that segment — a cohort-level risk "
        "no single-supplier review would surface."
    ))

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  4. RELATIONSHIP-LEVEL RISK DETECTION
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "4. Relationship-Level Risk Detection", level=1)

    add_body_text(doc, (
        "Entity-level monitoring — tracking each supplier or part independently — misses "
        "risks that live in the relationships between entities. BEI extends behavioral "
        "analysis to pairwise relationships and network-level structures."
    ))

    add_section_heading(doc, "4.1. Pairwise Relationships", level=2)

    add_body_text(doc, (
        "A pairwise relationship captures how two specific entities interact. The "
        "behavioral profile of the relationship can change even when both individual "
        "entity profiles appear normal."
    ))

    create_table(doc,
        ["Relationship", "What It Captures", "Risk Detected"],
        [
            ["Supplier x NSN", "How does THIS supplier perform\nfor THIS specific part?",
             "Quality drift for one part while\nstable for others = targeted\nsubstitution"],
            ["Supplier x\nCountry of Origin", "Where is THIS supplier sourcing\ncomponents?",
             "Gradual shift to adversary-nation\norigins = SCRM risk"],
            ["NSN x Location", "Demand patterns for THIS part\nat THIS depot",
             "Unusual demand composition at\nspecific locations = diversion"],
            ["NSN x Bill of\nMaterials", "Component composition\nchanges without notification",
             "SBOM drift = counterfeit or\nunauthorized substitution risk"],
            ["Supplier x\nSupplier", "Behavioral correlation between\n\"independent\" suppliers",
             "Hidden common sub-tier source\n= concentration risk"],
        ],
        col_widths=[1.3, 2.2, 2.8],
    )

    add_body_text(doc, (
        "Example: Supplier A's overall quality rate looks healthy. Part X's quality rate "
        "across all suppliers looks healthy. Both entities pass at the aggregate level. "
        "But Supplier A's quality specifically for Part X has been steadily degrading over "
        "months — a targeted degradation invisible at the entity level, visible only when "
        "the specific supplier-part relationship is tracked over time."
    ))

    add_body_text(doc, (
        "The Supplier-NSN and NSN-Location relationships are already built and demonstrated "
        "in the supply chain proof of concept (Section 7.3). The Supplier-Country, "
        "NSN-Bill-of-Materials, and Supplier-Supplier relationships use the identical "
        "relationship machinery and are added by introducing the corresponding entity "
        "profiles — country-of-origin, SBOM component, and supplier peer entities — making "
        "them direct, low-risk extensions rather than new research."
    ))

    add_section_heading(doc, "4.2. Network-Level Risk", level=2)

    add_body_text(doc, (
        "Beyond pairwise relationships, BEI analyzes the supply chain as a network graph "
        "to detect structural risks that no entity-level or pairwise analysis can reveal."
    ))

    add_body_text(doc, "Multi-Hop Risk Propagation", bold=True, space_after=2)
    add_body_text(doc, (
        "A Tier-1 supplier is fully compliant with all SCRM requirements. Its Tier-2 "
        "sub-supplier sources from an unknown Tier-3 manufacturer in an adversary nation. "
        "Entity-level monitoring of the Tier-1 supplier shows all green. Network analysis "
        "traces the supply path through multiple hops and identifies the adversary-nation "
        "connection."
    ))

    add_body_text(doc, "Hidden Single Points of Failure", bold=True, space_after=2)
    add_body_text(doc, (
        "Two suppliers appear independent: different CAGE codes, different locations, "
        "different contract vehicles. But both source a critical sub-component from the "
        "same Tier-2 manufacturer. If that manufacturer fails, both supply lines fail "
        "simultaneously. BEI detects this through behavioral correlation analysis — when "
        "two supposedly independent suppliers exhibit correlated behavioral changes, the "
        "system infers a hidden common dependency."
    ))

    add_body_text(doc, "Cascade Prediction", bold=True, space_after=2)
    add_body_text(doc, (
        "When one supplier fails, which depots are affected? Which missions are degraded? "
        "Network analysis propagates the impact of a supplier disruption through the "
        "supply graph, identifying every NSN, depot, and mission that depends on the "
        "failed supplier either directly or through sub-tier dependencies."
    ))

    add_section_heading(doc, "4.3. The What vs How Much Principle", level=2)

    add_body_text(doc, (
        "The foundational principle of BEI for supply chain risk is the distinction "
        "between what changed and how much changed."
    ))

    add_body_text(doc, (
        "Consider a supplier that delivers the same volume, at the same price, with the "
        "same on-time rate — but quietly shifts component sourcing from domestic origins "
        "to Chinese origins. Traditional metrics are all green: volume unchanged, price "
        "unchanged, delivery unchanged. Every threshold-based monitor shows a healthy "
        "supplier."
    ))

    add_body_text(doc, (
        "BEI relationship analysis detects the change: the Supplier x Country of Origin "
        "behavioral relationship has drifted. The supplier's Performance zone is stable. "
        "The Geographic Risk zone is changing direction. The system flags this as a "
        "SCRM risk — not because any magnitude threshold was crossed, but because the "
        "behavioral direction changed."
    ))

    add_callout(doc,
        "This is the supply chain equivalent of the insider threat: behavior changes "
        "WHAT, not HOW MUCH. The supplier looks the same by every traditional measure "
        "while fundamentally altering its sourcing strategy."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  5. DETECTION CAPABILITIES
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "5. Detection Capabilities", level=1)

    add_body_text(doc, (
        "BEI's multi-layer analysis — entity profiles, pairwise relationships, and "
        "network graphs — enables detection of five categories of supply chain risk that "
        "traditional monitoring consistently misses."
    ))

    add_section_heading(doc, "5.1. Counterfeit Part Detection", level=2)

    add_body_text(doc, (
        "Counterfeit parts enter the defense supply chain through sophisticated "
        "distribution networks where components are relabeled, remarked, or entirely "
        "fabricated. BEI detects counterfeit risk through three behavioral signatures."
    ))

    add_bullet(doc, "NSN failure rate diverges between locations receiving from different "
               "suppliers. When the same part fails markedly more often from one supplier "
               "than another, the NSN x Supplier relationship drift signals a sourcing "
               "quality difference that aggregate part-level statistics conceal.",
               bold_prefix="Failure Rate Divergence: ")
    add_bullet(doc, "The NSN x Manufacturer relationship changes — a different "
               "manufacturer of origin appears on inspection records while the label "
               "remains the same. BEI tracks this relationship continuously, not just "
               "at contract award.",
               bold_prefix="Manufacturer Relationship Drift: ")
    add_bullet(doc, "The component composition of a delivered assembly changes without "
               "engineering notification. BEI monitors SBOM behavioral profiles for "
               "drift in component identity, source, or specification.",
               bold_prefix="SBOM Composition Drift: ")

    add_section_heading(doc, "5.2. Adversary-Nation Sourcing (SCRM)", level=2)

    add_body_text(doc, (
        "Adversary-nation sourcing risk is the gradual shift of component origins to "
        "countries designated as supply chain risks. This shift typically occurs over "
        "months, making it invisible to annual audits."
    ))

    add_bullet(doc, "The Supplier x Country of Origin relationship drifts over "
               "months as sourcing geography shifts incrementally. Each monthly snapshot "
               "looks similar to the previous one; the cumulative trajectory reveals a "
               "fundamental change.",
               bold_prefix="Supplier x Country Drift: ")
    add_bullet(doc, "Network analysis traces supply paths through intermediaries "
               "to identify adversary-nation connections that are two or three hops "
               "removed from the prime supplier. A compliant Tier-1 supplier may "
               "unknowingly source from an adversary-nation Tier-3 manufacturer.",
               bold_prefix="Multi-Hop Network Analysis: ")
    add_bullet(doc, "BEI provides continuous behavioral monitoring of country-of-origin "
               "relationships, extending Section 889 compliance beyond the point-in-time "
               "assessment at contract award to ongoing operational monitoring.",
               bold_prefix="Section 889 Compliance: ")

    add_section_heading(doc, "5.3. Supplier Concentration Risk", level=2)

    add_body_text(doc, (
        "Concentration risk arises when the supply chain depends on fewer independent "
        "sources than it appears to have. BEI detects hidden concentration through "
        "behavioral correlation and network centrality analysis."
    ))

    add_bullet(doc, "When two supposedly independent suppliers exhibit correlated "
               "behavioral changes — simultaneous delivery delays, coordinated price "
               "increases, or synchronized quality shifts — BEI infers a hidden common "
               "dependency at the sub-tier level.",
               bold_prefix="Supplier Behavioral Correlation: ")
    add_bullet(doc, "Graph analysis identifies entities whose removal would "
               "disproportionately impact the supply network. A supplier that appears "
               "minor by contract value may be the sole source for a component used "
               "in 30 weapon systems.",
               bold_prefix="Network Centrality Analysis: ")
    add_bullet(doc, "If Supplier X fails, the system propagates the impact "
               "through the supply graph to identify every affected NSN, depot, and "
               "mission — including indirect impacts through sub-tier dependencies.",
               bold_prefix="Impact Propagation Modeling: ")

    add_section_heading(doc, "5.4. Demand Diversion and Fraud", level=2)

    add_body_text(doc, (
        "Demand diversion occurs when parts ordered for one purpose are redirected to "
        "unauthorized uses. BEI detects diversion through relationship-level anomalies."
    ))

    add_bullet(doc, "The same part is requisitioned by a different set of requestors "
               "with elevated urgency codes. The total volume looks normal; the "
               "requestor composition has changed.",
               bold_prefix="Requestor x NSN x Urgency Drift: ")
    add_bullet(doc, "Demand volume at a specific location remains within normal "
               "bounds, but the mix of items being demanded changes. BEI detects the "
               "composition shift through NSN x Location relationship analysis.",
               bold_prefix="Location Demand Composition: ")
    add_bullet(doc, "Orders that progress through acceptance without standard "
               "inspection steps create a behavioral sequence anomaly — the expected "
               "workflow pattern has changed.",
               bold_prefix="Sequence Anomaly Detection: ")

    add_section_heading(doc, "5.5. Contract and Financial Risk", level=2)

    add_body_text(doc, (
        "Supplier financial distress frequently precedes quality degradation and delivery "
        "failures. BEI correlates behavioral drift with financial health indicators to "
        "provide early warning."
    ))

    add_bullet(doc, "When a supplier's Performance Trajectory zone begins drifting "
               "concurrently with deterioration in its Risk Assessment zone, BEI "
               "flags the correlation as a financial distress indicator.",
               bold_prefix="Behavioral-Financial Correlation: ")
    add_bullet(doc, "Contracts approaching expiration combined with degrading "
               "delivery performance signal a supplier that may be disinvesting in "
               "the relationship — a flight risk for critical supply lines.",
               bold_prefix="Contract Expiry Risk: ")
    add_bullet(doc, "A supplier maintains stable pricing while failure rates "
               "increase. The cost per unit is unchanged, but the cost per "
               "functioning unit is rising. BEI detects this through the divergence "
               "between the financial and performance behavioral zones.",
               bold_prefix="Price-Quality Divergence: ")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  6. OPERATIONAL DEPLOYMENT
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "6. Operational Deployment", level=1)

    add_section_heading(doc, "6.1. Data Sources and Collection", level=2)

    add_body_text(doc, (
        "BEI integrates data from existing DLA and DoD systems. No new data collection "
        "is required — the system operates on transaction records, quality reports, and "
        "supplier profiles already maintained across the enterprise. The categories of "
        "data the framework consumes are summarized below; the subsections that follow "
        "describe the specific systems of record they come from, how each behavioral zone "
        "is populated, and how suppliers are resolved across systems."
    ))

    create_table(doc,
        ["Data Source", "Description", "Cadence", "Integration"],
        [
            ["Procurement\nTransactions", "Purchase orders, delivery orders,\ncontract actions",
             "Daily", "Batch extract from\ncontracting systems"],
            ["Delivery\nRecords", "Receipt confirmations, delivery\ndates, quantity variances",
             "Daily", "Automated feed from\nwarehouse management"],
            ["Quality\nInspections", "Inspection results, rejection rates,\ndefect classifications",
             "Per event", "Quality system\nintegration"],
            ["Country-of-Origin\nDeclarations", "Manufacturer origin certifications,\ntrade compliance records",
             "Per shipment", "Customs and trade\ncompliance systems"],
            ["SBOM\nRegistries", "Software and hardware bills of\nmaterials, component lists",
             "Per delivery", "Engineering data\nmanagement systems"],
            ["Financial\nFilings", "Supplier financial health indicators,\nbond and insurance status",
             "Quarterly", "SAM.gov, D&B,\ncommercial sources"],
            ["GIDEP\nAlerts", "Government-Industry Data Exchange\nProgram notifications",
             "As issued", "GIDEP automated\nalert feed"],
        ],
        col_widths=[1.2, 2.2, 0.8, 1.8],
    )

    # ── How each zone is fed ──
    add_body_text(doc, "How Each Behavioral Zone Is Populated", bold=True, space_after=2)
    add_body_text(doc, (
        "Each supplier behavioral zone (Section 3.2) is computed from a specific set of "
        "source records. The mapping below shows where the inputs for each zone originate "
        "and how they are collected, making explicit that every behavioral signal traces "
        "back to an authoritative system the enterprise already maintains."
    ))

    create_table(doc,
        ["Behavioral Zone", "Example Inputs", "Primary Source System(s)", "Collection / Cadence"],
        [
            ["Identity",
             "CAGE/UEI status, certifications,\ntier, small-business class",
             "SAM.gov, CAGE registry,\nJoint Certification Program",
             "Registry sync / API,\ndaily"],
            ["Performance\nTrajectory",
             "On-time delivery, lead time,\nquality rejection rate",
             "PIEE / WAWF (receipt &\nacceptance), SPRS, PDREP",
             "Batch extract,\ndaily / per-event"],
            ["Geographic /\nSourcing Risk",
             "Country-of-origin mix,\nadversary-nation exposure",
             "Country-of-origin declarations,\ncustoms/CBP, contract COO clauses",
             "Per-shipment feed"],
            ["Network\nPosition",
             "Depots served, NSNs supplied,\ncustomer concentration",
             "DLA EBS, DIBBS,\nFPDS-NG contract actions",
             "Batch extract,\ndaily"],
            ["Risk /\nFinancial",
             "Financial health, bond/insurance,\nexclusions, debarment",
             "Dun & Bradstreet, SAM.gov\nexclusions, SPRS risk",
             "Quarterly (financial),\ndaily (exclusions)"],
        ],
        col_widths=[1.2, 1.9, 1.9, 1.5],
    )

    # ── Systems of record ──
    add_body_text(doc, "Authoritative Systems of Record", bold=True, space_after=2)
    add_body_text(doc, (
        "BEI is fed from the established DoD and DLA systems that already govern supplier "
        "registration, performance, quality, and procurement. The exact set is tailored "
        "per deployment to the data the program office can provide; the systems below are "
        "the typical sources."
    ))

    create_table(doc,
        ["System of Record", "What It Provides", "Steward"],
        [
            ["SAM.gov",
             "Registration, Unique Entity ID (UEI),\nCAGE linkage, exclusions / debarment",
             "GSA"],
            ["SPRS (Supplier Performance\nRisk System)",
             "Supplier performance and risk scores,\nNIST SP 800-171 cybersecurity scores",
             "DoD (Navy)"],
            ["PDREP (Product Data Reporting\n& Evaluation Program)",
             "Quality deficiency reports, supplier\naudits, test and inspection results",
             "DoD (Navy)"],
            ["GIDEP",
             "Counterfeit and nonconformance\nalerts across government and industry",
             "DoD / Industry"],
            ["PIEE / WAWF, FPDS-NG",
             "Receiving, acceptance, invoicing;\nfederal contract action history",
             "DoD / GSA"],
            ["DLA EBS, DIBBS",
             "DLA procurement, solicitation, award,\nand inventory transactions",
             "DLA"],
            ["Dun & Bradstreet\n(commercial)",
             "Independent supplier financial-health\nindicators and corporate linkage",
             "Commercial"],
        ],
        col_widths=[1.9, 3.2, 1.4],
    )

    add_callout(doc,
        "No new instrumentation. BEI consumes data the enterprise already produces. The "
        "CAGE code (with UEI) is the universal join key that resolves a single supplier "
        "across SAM.gov, SPRS, PDREP, EBS, and commercial financial data."
    )

    # ── Entity resolution & data quality ──
    add_body_text(doc, "Entity Resolution and Data Quality", bold=True, space_after=2)
    add_body_text(doc, (
        "Behavioral profiling is only as trustworthy as the data plumbing beneath it. Four "
        "engineering concerns are addressed during integration:"
    ))
    add_bullet(doc, "A single supplier appears under different identifiers across systems. "
               "BEI keys on the CAGE code and UEI to unify records, and reconciles name "
               "changes, address changes, and parent/subsidiary structures so that a "
               "supplier's history follows it across re-registrations.",
               bold_prefix="Entity resolution: ")
    add_bullet(doc, "Sources update on different clocks — daily transactions, per-shipment "
               "origin declarations, quarterly financials. Behavioral snapshots are computed "
               "on a fixed cadence (weekly or monthly) and tolerate slower-moving inputs "
               "by carrying the last known value forward with an age flag.",
               bold_prefix="Latency and cadence: ")
    add_bullet(doc, "A supplier missing one data category is not dropped. Zones degrade "
               "gracefully: a supplier with no commercial financial feed is still scored on "
               "the other four zones, and a data-coverage confidence indicator accompanies "
               "every risk score so analysts know how complete the picture is.",
               bold_prefix="Missing and sparse data: ")
    add_bullet(doc, "Every behavioral feature traces back to the source record that produced "
               "it. This provenance is essential when a behavioral risk score informs a "
               "contract-award decision and must withstand supplier challenge or audit.",
               bold_prefix="Provenance and auditability: ")

    add_section_heading(doc, "6.2. Deployment Phases", level=2)

    add_body_text(doc, (
        "BEI deployment follows a four-phase approach designed to deliver initial value "
        "within 8 weeks while building toward full production capability."
    ))

    create_table(doc,
        ["Phase", "Duration", "Activities", "Deliverables"],
        [
            ["Phase 1:\nData Integration", "2-4 weeks",
             "Connect to procurement, delivery,\nquality, and supplier databases.\n"
             "Entity resolution across systems.\nData quality assessment.",
             "Integrated data pipeline\nEntity resolution map\nData quality report"],
            ["Phase 2:\nBaseline\nEstablishment", "4-6 weeks",
             "Ingest 4-6 weeks of historical data.\n"
             "Build behavioral profiles for all\nentities. Establish zone baselines.\n"
             "Calibrate drift detection thresholds.",
             "Entity behavioral profiles\nBaseline reference points\nCalibration report"],
            ["Phase 3:\nDetection\nDeployment", "2-4 weeks",
             "Deploy detection capabilities.\n"
             "Threshold tuning against known\nhistorical incidents. SOC analyst\n"
             "training on alert triage.",
             "Detection rules deployed\nAlert triage procedures\nAnalyst training complete"],
            ["Phase 4:\nProduction\nMonitoring", "Ongoing",
             "Continuous behavioral monitoring.\n"
             "Integration with acquisition and\nSOC workflows. Quarterly model\n"
             "recalibration.",
             "Operational dashboard\nWeekly risk reports\nAcquisition decision support"],
        ],
        col_widths=[1.0, 0.8, 2.5, 2.2],
    )

    add_section_heading(doc, "6.3. Integration Points", level=2)

    add_body_text(doc, (
        "BEI is designed to integrate with DLA's existing systems and workflows, "
        "augmenting current capabilities rather than replacing them."
    ))

    add_bullet(doc, "BEI's behavioral risk scores feed into DLA's existing SCRM "
               "risk assessments, providing continuous monitoring between annual reviews.",
               bold_prefix="DLA SCRM Program: ")
    add_bullet(doc, "GIDEP alerts are correlated with BEI behavioral detections. "
               "A GIDEP alert about a supplier coinciding with behavioral drift in "
               "that supplier's profile elevates the risk score and triggers "
               "accelerated investigation.",
               bold_prefix="GIDEP Integration: ")
    add_bullet(doc, "BEI cross-references entity profiles against CAGE and SAM.gov "
               "registrations. Changes in supplier registration status (new CAGE codes, "
               "SAM.gov exclusions, changed addresses) trigger behavioral profile updates.",
               bold_prefix="CAGE/SAM.gov Databases: ")
    add_bullet(doc, "Behavioral risk scores are presented to contracting officers at "
               "the point of contract award decision. A supplier with a drifting "
               "Geographic Risk zone is flagged before the award, not after delivery.",
               bold_prefix="Contracting Officer Workflow: ")
    add_bullet(doc, "Parts from suppliers with elevated behavioral risk scores are "
               "routed to priority inspection queues. Limited inspection resources are "
               "allocated based on behavioral intelligence rather than random sampling.",
               bold_prefix="Quality Assurance Routing: ")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  7. VALIDATED RESULTS
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "7. Proven Foundation: From Cybersecurity to Supply Chain", level=1)

    add_section_heading(doc, "7.1. What Has Been Proven in Cybersecurity", level=2)

    add_body_text(doc, (
        "The behavioral detection concept described in this whitepaper is not a research "
        "proposal. It has been built, deployed, and proven in an operational cybersecurity "
        "setting. In that deployment, the BEI framework analyzed a large population of "
        "users over an extended period and was challenged with four advanced, "
        "long-duration attack campaigns deliberately designed to evade detection: a "
        "patient insider threat, a slow command-and-control intrusion, and two "
        "nation-state campaigns modeled on real-world adversary tradecraft."
    ))

    add_body_text(doc, (
        "Traditional detection methods — the statistical and threshold-based techniques "
        "that are the supply chain equivalent of on-time-rate and quality-score "
        "monitoring — catch at most one of the four campaigns (Z-Score detects 1 of 4 at "
        "9.8% FP) and provide no directional intelligence to distinguish threat types or guide "
        "triage. The BEI framework detected all four while adding zone-level explanation, "
        "ranking each attacker among the highest-risk entities and telling analysts WHICH "
        "behavioral dimension changed. The attackers were caught not because any single "
        "value crossed a threshold, but because their behavioral direction changed — "
        "exactly the failure mode that defeats traditional supplier monitoring."
    ))

    add_callout(doc,
        "The proof point is concrete: in cybersecurity, BEI matched traditional detection "
        "rates while adding zone-level directional intelligence that traditional methods "
        "cannot provide. The actionable triage came from behavioral direction, not metric "
        "magnitude — the same signal that exposes hidden supplier risk."
    )

    add_section_heading(doc, "7.2. Why the Proof Transfers", level=2)

    add_body_text(doc, (
        "The cybersecurity and supply chain problems are structurally identical. In both, "
        "the most damaging actors operate within authorized bounds — valid credentials in "
        "one case, compliant contracts in the other — while gradually shifting behavioral "
        "direction toward an objective. The BEI framework is domain-agnostic: the same "
        "analytical architecture that profiles a user profiles a supplier; the same "
        "relationship analysis that links a user to a device links a supplier to a part "
        "and a country of origin; the same temporal trajectory analysis that detects an "
        "escalating insider detects a supplier drifting toward distress."
    ))

    create_table(doc,
        ["Framework Element", "Proven In Cybersecurity", "Applied To Supplier Risk"],
        [
            ["Entity profile",
             "User, device, application behavioral baselines",
             "Supplier, NSN, location, manufacturer baselines"],
            ["Behavioral zones",
             "Identity, access, data, network, risk",
             "Identity, performance, geographic, network, risk"],
            ["Peer-cohort scoring",
             "Group-relative drift vs. role group\n(admin, security, developer, ...)",
             "Group-relative drift vs. supplier cohort\n(commodity, geography, criticality)"],
            ["Relationship analysis",
             "User-device, user-application interactions",
             "Supplier-part, supplier-country, part-SBOM interactions"],
            ["Temporal trajectory",
             "Detecting escalating insider over months",
             "Detecting supplier drift toward distress or substitution"],
            ["Directional detection",
             "Caught all 4 advanced campaigns missed by traditional methods",
             "Designed to catch sourcing shifts, counterfeits, hidden concentration"],
        ],
        col_widths=[1.4, 2.5, 2.6],
    )

    add_body_text(doc, (
        "What changes between domains is the vocabulary of the entities and the data "
        "sources that feed them. What stays constant is the detection principle and the "
        "analytical machinery — both proven."
    ))

    add_section_heading(doc, "7.3. Supply Chain Proof of Concept", level=2)

    add_body_text(doc, (
        "The BEI framework has been instantiated for the supply chain domain and exercised "
        "on a realistic synthetic dataset. This proof of concept demonstrates that the "
        "framework's capabilities — entity profiling, zone decomposition, relationship "
        "analysis, temporal trajectory, and regime detection — operate on supply chain "
        "data exactly as they operate on cybersecurity data."
    ))

    create_table(doc,
        ["Parameter", "Value"],
        [
            ["National Stock Numbers (NSNs)", "500"],
            ["Suppliers", "200"],
            ["Depots", "4 (3 Indo-Pacific, 1 CONUS)"],
            ["Transaction History", "24 months"],
            ["Demand Components Modeled", "6 (baseline, exercise, seasonal,\n"
             "maintenance, combat, failure)"],
            ["Supplier Risk Grades", "A through F (composite behavioral +\n"
             "financial + performance)"],
        ],
        col_widths=[2.5, 4.0],
    )

    add_body_text(doc, "Demonstrated Capabilities", bold=True, space_after=2)

    add_bullet(doc, "The framework identifies items transitioning from STABLE to DEPLETING "
               "behavioral regimes before stockout occurs, providing actionable lead time "
               "for procurement intervention.",
               bold_prefix="Phase Space Regime Detection: ")
    add_bullet(doc, "Each supplier receives a composite health grade (A through F) "
               "incorporating behavioral trajectory, financial health indicators, and "
               "delivery performance trends. The grade reflects behavioral direction, "
               "not just current-period metrics.",
               bold_prefix="Supplier Risk Scoring: ")
    add_bullet(doc, "Demand for each NSN is decomposed into six independent components, "
               "enabling the system to distinguish exercise-driven demand spikes from "
               "baseline consumption changes and combat-related surges.",
               bold_prefix="Demand Decomposition: ")
    add_bullet(doc, "Four relationship types are built and operating: Supply "
               "(Supplier-NSN), Demand (NSN-Location), Transit (NSN-Route), and "
               "Fulfillment (Supplier-NSN-Location). The country-of-origin, "
               "bill-of-materials, and supplier-to-supplier relationships described in "
               "Section 4 are natural extensions of the same relationship machinery, "
               "added by introducing the corresponding entity profiles.",
               bold_prefix="Relationship Analysis: ")
    add_bullet(doc, "Analysts query the system in natural language — 'Which Indo-Pacific "
               "parts are at risk of stockout in the next 90 days?' — and receive "
               "answers contextualized with behavioral intelligence, not just inventory "
               "snapshots.",
               bold_prefix="LLM-Powered Decision Support: ")

    add_body_text(doc, "Measured Supplier-Risk Results (Synthetic MVP)", bold=True, space_after=2)
    add_body_text(doc, (
        "Beyond demonstrating the capabilities, the supply chain MVP implements a calibrated, "
        "explainable supplier-risk model and measures it. The model was trained and audited on "
        "1,174 supplier-evaluation windows (walk-forward 90-day snapshots), with a stratified "
        "20% held-out test set (235 windows) and a failure base rate of approximately 20%. "
        "A supplier-window is labeled a failure if, in the following 90 days, its on-time rate "
        "falls below 0.5 or it incurs two or more severely late (>15-day) deliveries. The "
        "architecture and features are detailed in Section 8.10; the measured results are below."
    ))

    create_table(doc,
        ["Model", "Metric", "Value", "Grade"],
        [
            ["XGBoost classifier\n(P(failure ≤ 90d))", "AUC (V1, 12 features)", "0.7709", "B"],
            ["XGBoost classifier", "AUC (V2, +behavioral features)", "0.7780", "B"],
            ["XGBoost classifier", "Average Precision", "0.4918", "—"],
            ["XGBoost classifier", "Brier / ECE (calibration)", "0.1588 / 0.1019", "C"],
            ["Random Survival Forest", "C-index (held-out)", "0.7935", "B"],
            ["Random Survival Forest", "Survival ECE @ 30 / 60 / 90d", "0.084 / 0.090 / 0.129", "B / B / C"],
        ],
        col_widths=[1.9, 2.3, 1.4, 0.9],
    )

    add_body_text(doc, (
        "The model's most influential features (by SHAP) are quality incidents, recent order "
        "count, on-time rate, worst delay, and critical-NSN count — an explainable, "
        "operationally sensible ordering. Monotone constraints guarantee the score moves in the "
        "intuitive direction (more quality incidents and sole-source exposure raise risk; higher "
        "on-time rate lowers it)."
    ))
    add_figure(doc, "disc",
               "Figure 4 — Supplier-risk discrimination (synthetic): XGBoost AUC 0.77 (V1) / 0.78 "
               "(V2) and survival C-index 0.79 — all well above chance.")
    add_figure(doc, "calib",
               "Figure 5 — Survival-curve calibration by horizon (synthetic): Expected Calibration "
               "Error 0.084 / 0.090 / 0.129 at 30 / 60 / 90 days — Grade B at the near horizons.")
    add_figure(doc, "shap",
               "Figure 6 — What drives supplier risk (top SHAP features, synthetic): quality "
               "incidents, recent order count, on-time rate, worst delay, critical-NSN count — an "
               "operationally sensible, explainable ordering.")

    add_callout(doc,
        "The cybersecurity deployment proved detection efficacy quantitatively. The supply "
        "chain MVP now adds a measured, calibrated supplier-risk model — AUC 0.77, survival "
        "C-index 0.79 — built and audited on a realistic synthetic dataset. ALL supply-chain "
        "numbers are from synthetic data and must be re-validated on real DLA data before any "
        "production claim; the next step is that real-data re-measurement."
    )

    add_section_heading(doc, "7.4. Illustrative Detection Scenarios", level=2)

    add_body_text(doc, (
        "The following scenarios illustrate what BEI is designed to surface when applied "
        "to defense supplier data. They describe the behavioral signal and the mechanism "
        "by which BEI exposes it — the same class of directional signal proven detectable "
        "in the cybersecurity deployment and demonstrated operationally in the supply "
        "chain proof of concept."
    ))

    create_table(doc,
        ["Scenario", "Behavioral Signal", "Traditional\nVisibility", "How BEI Surfaces It"],
        [
            ["Supplier sourcing shift\nto adversary nation",
             "Geographic Risk zone drifting\nwhile Performance zone stable",
             "Invisible until\nnext annual audit",
             "Continuous zone-level\ndrift detection flags the\ngeographic shift as it forms"],
            ["Hidden supplier\nconcentration",
             "Two suppliers show correlated\nbehavioral changes (delivery,\npricing, quality)",
             "Invisible —\nsuppliers appear\nindependent",
             "Network cohort analysis\nflags the correlated\nbehavior as a shared\ndependency"],
            ["NSN transitioning\nto stockout",
             "Inventory Position zone\naccelerating toward depletion\nregime",
             "Visible only when\nstockout occurs",
             "Regime detection flags\nthe STABLE-to-DEPLETING\ntransition before depletion"],
            ["Counterfeit part\nentry",
             "NSN failure rate diverges\nbetween supplier sources",
             "Detected only after\nfield failure",
             "Relationship analysis\nflags the supplier-part\nfailure divergence"],
            ["Supplier financial\ndistress",
             "Performance and Risk\nzones drift concurrently",
             "Detected at next\nquarterly review",
             "Multi-zone trajectory\nflags the concurrent drift\nas it develops"],
        ],
        col_widths=[1.4, 2.0, 1.4, 1.9],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  8. TECHNICAL DEEP DIVE
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "8. Technical Deep Dive: The Mathematics of Behavioral Drift", level=1)

    add_body_text(doc, (
        "The preceding sections describe BEI conceptually. This section makes the innovation "
        "explicit — the representations, transforms, and detection algorithms that turn raw "
        "supply-chain records into a ranked, explained risk verdict. The notation is kept "
        "deliberately concrete so a technical evaluator can reason about, and reproduce, the "
        "method. The central claim is stated up front and demonstrated through the section: "
        "no single equation is the innovation; the innovation is the COMPOSITION of a "
        "semantic representation in which the direction of behavioral change is measurable, "
        "decomposed into independent zones, scored against the right peer cohort, accumulated "
        "over time, and fused across mutually-reinforcing detection phases."
    ))

    add_callout(doc,
        "Traditional analytics live in raw feature space and measure magnitude — how far a "
        "metric moved. BEI lives in a semantic embedding space and measures direction — what "
        "the supplier is becoming. Everything below follows from that single change of space."
    )

    # ── 8.1 ──
    add_section_heading(doc, "8.1. From Metrics to Meaning: Serialization and Embedding", level=2)
    add_body_text(doc, (
        "For each supplier i and period t, the raw features of each behavioral zone z are "
        "first serialized into a short natural-language description, then mapped to a "
        "high-dimensional semantic vector by an embedding model:"
    ))
    add_code_block(doc,
        "e[i,z,t] = Embed( serialize( features(i, z, t) ) )   ∈  R^1536\n\n"
        "example serialization (Geographic / Sourcing zone):\n"
        "  \"Supplier 7731, electronics, prime. Country-of-origin mix:\n"
        "   72% domestic, 18% allied, 10% high-risk foreign. 4 manufacturing\n"
        "   sites across 3 countries. 1 new high-risk site this period.\"")
    add_body_text(doc, (
        "Why embed text rather than compute distance on the raw numbers? Because the risk "
        "signal is often a change in the KIND of value, not its size. A shift from 1% to 12% "
        "high-risk-foreign sourcing is a small Euclidean distance among features dominated by "
        "large-magnitude quantities (dollars, units), yet it is an enormous change in MEANING. "
        "The embedding places semantically similar behavior near itself, so a change in kind "
        "becomes a measurable change in geometry."
    ))

    # ── 8.2 ──
    add_section_heading(doc, "8.2. Zone Composition and Context-Adaptive Attention", level=2)
    add_body_text(doc, (
        "Each supplier's overall state at period t is a composition of its five zone vectors. "
        "The baseline composition is a weighted average, renormalized to the unit sphere:"
    ))
    add_code_block(doc,
        "V[i,t] = normalize( Σ_z  w_z · e[i,z,t] )\n\n"
        "illustrative supplier zone weights (tuned per deployment):\n"
        "  performance = 0.25,  geographic = 0.25,  network = 0.20,\n"
        "  risk = 0.20,  identity = 0.10")
    add_body_text(doc, (
        "Innovation — zone independence. Because zones are composed rather than collapsed into "
        "a single average of raw metrics, a strong signal in one zone (geographic drift) is "
        "not washed out by four stable zones. A single monolithic score would average the "
        "signal away; zone decomposition preserves it."
    ))
    add_body_text(doc, (
        "Context-adaptive attention takes this further. Rather than fixed weights, the "
        "contribution of each zone is set by a context-biased linear normalization of zone "
        "salience (routine monitoring, SCRM hunt, financial-distress review):"
    ))
    add_code_block(doc,
        "α_z = (||e[i,z,t]|| · b_z(context)) / Σ_z (||e[i,z,t]|| · b_z(context))   # linear norm\n"
        "V[i,t] = normalize( Σ_z  α_z · e[i,z,t] )\n\n"
        "under a SCRM-hunt context, b_z raises attention on the geographic\n"
        "and risk zones; under a distress review, on risk and performance.")
    add_body_text(doc, (
        "The same supplier is therefore scored differently depending on what is being hunted — "
        "without retraining the model."
    ))

    # ── 8.3 ──
    add_section_heading(doc, "8.3. Drift: Measuring Direction, Not Just Magnitude", level=2)
    add_body_text(doc, (
        "Behavioral change between consecutive periods is captured two ways. The magnitude of "
        "change is the cosine distance between snapshots; the direction of change is the unit "
        "vector of their difference:"
    ))
    add_code_block(doc,
        "drift magnitude:  d[i,t] = 1 − cos( V[i,t−1], V[i,t] )      ∈ [0, 2]\n"
        "drift direction:  u[i,t] = ( V[i,t] − V[i,t−1] ) / || V[i,t] − V[i,t−1] ||")
    add_body_text(doc, (
        "Magnitude answers \"did something change, and how much?\" Direction answers \"what is "
        "this supplier turning into?\" The second question is the one traditional monitoring "
        "cannot ask, because raw-feature distance has magnitude but no interpretable direction. "
        "The next subsection gives the direction a name."
    ))

    # ── 8.4 ──
    add_section_heading(doc, "8.4. Naming the Direction: Reference-Concept Projection", level=2)
    add_body_text(doc, (
        "A library of risk concepts {c_k} — for example adversary-nation sourcing, counterfeit "
        "substitution, demand diversion, financial distress — is embedded into the SAME "
        "semantic space as supplier behavior. The drift direction is projected onto each "
        "concept by cosine alignment:"
    ))
    add_code_block(doc,
        "align_k = cos( u[i,t], Embed(c_k) )\n"
        "primary risk = argmax_k align_k   (with benign anchor concepts to suppress\n"
        "                                   false alignment from routine variation)")
    add_body_text(doc, (
        "Innovation — explainability by construction. The output is not an opaque anomaly "
        "score but a ranked, plain-language reason (\"drifting toward adversary-nation "
        "sourcing\") that maps directly to the relevant SCRM authority (e.g., Section 889, "
        "DFARS 252.246-7008). An analyst can act on it, and a contracting officer can defend "
        "the decision."
    ))

    # ── 8.5 ──
    add_section_heading(doc, "8.5. Accumulating Slow Drift: CUSUM Change-Point", level=2)
    add_body_text(doc, (
        "The most damaging supply-chain shifts are slow: a supplier moves sourcing a little "
        "each month, so no single period crosses an alert threshold. A one-sided cumulative-sum "
        "(CUSUM) accumulates the small excesses and fires on the sustained trend:"
    ))
    add_code_block(doc,
        "S_0 = 0\n"
        "S_t = max( 0,  S_{t−1} + ( d[i,t] − k ) )        # k = baseline drift allowance\n"
        "alarm when S_t ≥ h for ≥ r consecutive periods   # h = decision threshold\n"
        "change-point = index where the triggering run began\n\n"
        "operational defaults from the proven deployment: k = 0.02, h = 0.05, r = 2")
    add_body_text(doc, (
        "Worked example. A supplier drifts d = 0.03 per month toward foreign sourcing — well "
        "below any single-period alert level. Each month S grows by 0.03 − 0.02 = 0.01, so S "
        "crosses h = 0.05 after roughly five months and flags a change-point at the start of "
        "the drift. A twelve-month adversary-nation migration that every monthly threshold "
        "would pass is caught while it is still forming."
    ))

    # ── 8.6 ──
    add_section_heading(doc, "8.6. Peer-Cohort Scoring", level=2)
    add_body_text(doc, (
        "A behavioral value is anomalous only relative to comparable suppliers. Each supplier "
        "is assigned to a cohort g by stable identity attributes — commodity / Federal Supply "
        "Class, geography, criticality, tier — and every feature is standardized against the "
        "NORMAL members of that cohort:"
    ))
    add_code_block(doc,
        "cohort:   g(i) = ( commodity, geography, criticality, tier )\n"
        "μ[g,f] , σ[g,f]  estimated over normal suppliers in cohort g\n"
        "z[i,f] = ( x[i,f] − μ[g(i),f] ) / σ[g(i),f]")
    add_body_text(doc, (
        "Innovation — the right reference group. A multinational sourcing footprint is normal "
        "for a semiconductor cohort and alarming for a domestic-fastener cohort; scoring "
        "against the population would hide both signals, scoring against the cohort exposes "
        "them. This is exactly the role-group standardization proven in the cybersecurity "
        "deployment (users scored against administrators, analysts, executives), with supplier "
        "cohorts substituting for workforce roles."
    ))

    # ── 8.7 ──
    add_section_heading(doc, "8.7. Multi-Phase Composite Score", level=2)
    add_body_text(doc, (
        "The final verdict fuses five independent phases computed from the cohort z-vector and "
        "the drift series. Each phase targets a different evasion strategy:"
    ))
    add_code_block(doc,
        "signal_strength = Σ  top-3 of { z[i,f] }            # peak deviation\n"
        "breadth         = | { f : z[i,f] > 1.5 } |          # how many features elevated\n"
        "sustained       = Σ  top-2 of { late-window mean z[i,zone] }   # persistence\n"
        "divergence      = z of context-specific composite spread       # cohort/context\n"
        "novelty         = recurrence of never-before-seen entities\n"
        "                  (new country-of-origin, new sub-tier manufacturer)\n\n"
        "composite = β1·signal + β2·breadth + β3·sustained\n"
        "          + β4·divergence + β5·novelty           # rank suppliers by composite")
    add_body_text(doc, (
        "Innovation — evasion resistance through independence. A supplier that suppresses "
        "magnitude still fails breadth; one that limits breadth still fails sustained "
        "deviation or novelty persistence. No single masking strategy defeats the fused score. "
        "Novelty persistence in particular is the phase that catches the slowest campaigns, "
        "which are quiet on every other axis — the supply-chain analogue of the slow APT whose "
        "only tell is the same never-before-seen infrastructure reappearing period after period."
    ))

    # ── 8.8 ──
    add_section_heading(doc, "8.8. Relationship and Co-Drift Detection", level=2)
    add_body_text(doc, (
        "Risk that lives between entities is captured by composing a relationship vector as the "
        "element-wise (Hadamard) product of two entity vectors, renormalized:"
    ))
    add_code_block(doc,
        "R[a×b, t] = normalize( V[a,t] ⊙ V[b,t] )\n\n"
        "drift in R[a×b] while V[a] and V[b] are individually stable reveals\n"
        "targeted change: Supplier×NSN (one part degrading), Supplier×Country\n"
        "(sourcing shift), NSN×SBOM (composition change).")
    add_body_text(doc, (
        "Hidden concentration is detected as co-drift: cluster suppliers whose drift directions "
        "are mutually aligned above a coherence threshold τ — cos(u[a,t], u[b,t]) ≥ τ — and a "
        "cluster of \"independent\" suppliers moving together points to a shared sub-tier "
        "dependency or a common geopolitical shock no single-supplier review would surface."
    ))

    # ── 8.9 ──
    add_section_heading(doc, "8.9. Why the Composition Is the Innovation", level=2)
    add_body_text(doc, (
        "Each component above exists, in some form, elsewhere. The defensible innovation is "
        "their integration into a single pipeline in which each stage compensates for the "
        "blind spot of the others:"
    ))
    create_table(doc,
        ["Stage", "What it adds", "Blind spot it removes"],
        [
            ["Semantic embedding", "Direction of change is measurable",
             "Raw-feature distance has magnitude but no meaning"],
            ["Zone decomposition", "One drifting zone is not diluted",
             "Monolithic scores average the signal away"],
            ["Context attention", "Scoring adapts to the hunt",
             "Fixed weights miss context-specific risk"],
            ["Reference concepts", "Plain-language, authority-mapped reason",
             "Opaque anomaly scores are not actionable"],
            ["CUSUM", "Slow, sub-threshold drift accumulates",
             "Per-period thresholds miss gradual shifts"],
            ["Peer cohort", "Right reference group",
             "Population baselines hide cohort-specific anomalies"],
            ["Multi-phase fusion", "Evasion resistance",
             "Any single detector can be gamed"],
        ],
        col_widths=[1.5, 2.5, 2.5],
    )
    add_body_text(doc, (
        "End to end: a supplier's monthly records are serialized and embedded; zone vectors are "
        "composed under the active context; drift magnitude and direction are computed; the "
        "direction is named against the concept library and mapped to SCRM authority; CUSUM "
        "accumulates the slow trend; cohort z-scores place the behavior against true peers; and "
        "the five-phase composite produces a single ranked, explained verdict. That integrated "
        "path — not any one formula — is what detects the supply-chain threats that magnitude-"
        "based monitoring is structurally unable to see."
    ))

    add_callout(doc,
        "The moat is the pipeline, not a single equation. Direction over magnitude, zones over "
        "monoliths, cohorts over population baselines, accumulation over instantaneous "
        "thresholds, and fusion over any single detector — composed into one explainable, "
        "ranked verdict."
    )

    # ── 8.10 ──
    add_section_heading(doc, "8.10. The Predictive Risk Layer: XGBoost + Survival Model (Implemented MVP)", level=2)
    add_body_text(doc, (
        "The behavioral layer of Sections 8.1–8.9 produces features and embeddings that "
        "describe how a supplier is behaving and in what direction it is drifting. The "
        "implemented MVP turns that behavioral evidence into an actionable, calibrated, "
        "explainable risk score with a time dimension, using a dual predictive model. The "
        "behavioral embeddings and cohort/drift features plug in as inputs; the predictive "
        "layer adds probability calibration, a survival curve, and SHAP explanations."
    ))

    add_body_text(doc, "Two models, two jobs.", bold=True, space_after=2)
    add_bullet(doc,
        "Gradient-boosted classifier estimating P(significant delivery failure within 90 days). "
        "It provides the headline risk score and, via SHAP, the per-feature explanation. "
        "Monotone constraints force the score in the intuitive direction — higher on-time rate "
        "lowers risk; sole-source status, quality incidents, and worst delay raise it — so the "
        "model can never produce a counter-intuitive, indefensible ranking.",
        bold_prefix="Model 1 — XGBoost classifier: ")
    add_bullet(doc,
        "A survival model estimating the full curve P(supplier still OK at day T), reported at "
        "30, 60, and 90 days. This time-dependent risk feeds inventory and sourcing decisions "
        "(when, not just whether, to expect trouble).",
        bold_prefix="Model 2 — Random Survival Forest: ")

    add_body_text(doc, "Features (V1 = 12 leakage-free signals; V2 adds the behavioral layer):", bold=True, space_after=2)
    add_code_block(doc,
        "delivery / performance : ontime_rate_90d, ontime_trend_slope,\n"
        "                         avg_lead_time_ratio, delivery_variance_90d,\n"
        "                         worst_delay_90d, order_count_90d, order_volume_trend\n"
        "quality                : quality_incidents_6m\n"
        "structural / exposure  : is_sole_source, geographic_risk,\n"
        "                         critical_nsn_count, sole_source_nsn_count\n"
        "V2 behavioral add-ons  : supplier zone-embedding features (delivery, portfolio,\n"
        "                         capacity, stability, network) + drift velocity")
    add_body_text(doc, (
        "Labels are built walk-forward: every 90 days each supplier is snapshotted, features are "
        "computed from the prior 90-day window, and the label is read from the next 90 days "
        "(failure = future on-time < 0.5 or two-plus deliveries more than 15 days late). The "
        "survival target is the number of days to the first severely late delivery. All "
        "metadata is read point-in-time from a bi-temporal store so training never sees the "
        "future."
    ))

    add_body_text(doc, "Operational guardrails (where the model defers to hard rules):", bold=True, space_after=2)
    add_bullet(doc, "A debarred supplier short-circuits to maximum risk with a directive to "
               "verify status on SAM.gov; if the debarment check itself fails, the model "
               "fails closed (risk = max, manual review) rather than guessing.",
               bold_prefix="Debarment pre-filter: ")
    add_bullet(doc, "A CAGE-code change in the last 12 months (from bi-temporal history) is "
               "surfaced as an ownership/restructuring/re-registration warning — an enrichment "
               "signal, not a silent model feature.",
               bold_prefix="CAGE-code change: ")
    add_bullet(doc, "When a destination depot is specified, the global score is blended with "
               "destination-specific performance (0.6 global / 0.4 destination) so risk "
               "reflects how the supplier performs to that location.",
               bold_prefix="Destination-aware blend: ")

    add_body_text(doc, (
        "Measured performance on the synthetic MVP — AUC 0.7709, RSF C-index 0.7935, calibration "
        "ECE in the 0.08–0.13 range — is reported in Section 7.3. The relationship between the "
        "two layers is the point: the behavioral pipeline supplies direction and novelty that "
        "raw performance features miss, while the predictive layer supplies calibration, a "
        "survival horizon, and defensible explanations. Neither alone is sufficient; together "
        "they produce a ranked, time-aware, explainable supplier-risk verdict."
    ))

    add_callout(doc,
        "Behavioral layer answers \"what is this supplier becoming, and is it unlike its peers?\" "
        "Predictive layer answers \"what is the calibrated probability and timing of failure, and "
        "why?\" The MVP implements both — measured, explainable, and guardrailed — on synthetic "
        "data pending real-data re-validation."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  9. BUSINESS VALUE
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "9. Business Value", level=1)

    add_body_text(doc, (
        "BEI delivers measurable value across six dimensions of supply chain risk "
        "management, transforming reactive monitoring into proactive risk intelligence."
    ))

    add_body_text(doc, "Reduced Counterfeit Exposure", bold=True, space_after=2)
    add_body_text(doc, (
        "Continuous behavioral monitoring of supplier-part-manufacturer relationships "
        "detects counterfeit risk signals months before parts fail in the field. The "
        "Government Accountability Office estimates average counterfeit remediation "
        "costs at $1.8 million per incident, including recall, inspection, replacement, "
        "and mission impact assessment. BEI's early detection capability reduces both "
        "the frequency and severity of counterfeit incidents."
    ))

    add_body_text(doc, "Earlier Detection of SCRM Risks", bold=True, space_after=2)
    add_body_text(doc, (
        "Traditional SCRM assessments operate on annual review cycles. BEI monitors "
        "supplier behavior continuously, surfacing adversary-nation sourcing shifts as "
        "the behavioral change forms rather than at the next scheduled review. For "
        "slow-moving supply chain compromises that unfold over quarters, the difference "
        "between continuous detection and periodic review is the difference between "
        "prevention and remediation."
    ))

    add_body_text(doc, "Reduced Mission Impact from Supply Disruptions", bold=True, space_after=2)
    add_body_text(doc, (
        "Cascade prediction identifies the mission impact of supplier failures before "
        "they occur. When a supplier enters the Drifting regime, the system immediately "
        "identifies every depot, NSN, and mission that would be affected, enabling "
        "proactive sourcing decisions rather than reactive scrambling."
    ))

    add_body_text(doc, "Compliance Automation", bold=True, space_after=2)
    add_body_text(doc, (
        "Continuous monitoring of Section 889 and DFARS compliance replaces periodic "
        "manual audits. The system continuously tracks country-of-origin relationships "
        "and flags compliance deviations as they emerge, reducing both compliance risk "
        "and the labor cost of manual reviews."
    ))

    add_body_text(doc, "Acquisition Intelligence", bold=True, space_after=2)
    add_body_text(doc, (
        "Behavioral risk scores inform contract award decisions at the point of "
        "decision. A contracting officer evaluating two suppliers can see not just "
        "current performance metrics but behavioral trajectory: is this supplier's "
        "quality improving or degrading? Is its geographic sourcing stable or shifting? "
        "This intelligence transforms acquisition from a backward-looking assessment to "
        "a forward-looking prediction."
    ))

    add_body_text(doc, "The Scale of the Problem BEI Addresses", bold=True, space_after=2)
    add_body_text(doc, (
        "The following figures describe the documented cost of the supply chain risks "
        "that BEI is designed to detect earlier. They establish the scale of the problem, "
        "drawn from external sources; the value of earlier behavioral detection is "
        "proportional to these stakes."
    ))

    create_table(doc,
        ["Risk Category", "Documented Cost of the Problem", "Source", "Where BEI Helps"],
        [
            ["Counterfeit\nRemediation", "~$1.8M average per incident", "GAO reports",
             "Behavioral monitoring of\nsupplier-part-manufacturer\nrelationships surfaces\ncounterfeit risk signals"],
            ["Stockout /\nMission Impact", "Mission-dependent, can reach\ntens of millions per event",
             "DLA readiness\nassessments",
             "Regime detection flags\ndepletion trajectories\nbefore stockout"],
            ["SCRM Compliance\nViolation", "Significant remediation and\nreporting cost per finding",
             "DFARS enforcement\nhistory",
             "Continuous country-of-origin\nmonitoring replaces\nperiodic audits"],
            ["Supplier Failure\nCascade", "Scope-dependent, can be\nsubstantial across a portfolio",
             "COVID-19 supply\nchain analyses",
             "Network analysis exposes\nhidden concentration before\nit cascades"],
        ],
        col_widths=[1.3, 1.7, 1.3, 2.0],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  9. CONCLUSION
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "10. Conclusion", level=1)

    add_body_text(doc, (
        "Traditional supplier monitoring measures individual metrics at points in time. "
        "Each metric has a threshold. Each threshold is evaluated independently. A "
        "supplier that stays below every threshold is classified as low-risk, regardless "
        "of what is changing beneath the surface."
    ))

    add_body_text(doc, (
        "Behavioral Entity Intelligence measures behavioral direction across entity "
        "relationships continuously. It asks not 'Is this metric within bounds?' but "
        "'Is this entity behaving the same way it has historically?' and 'In which "
        "dimension is the behavior changing?' The decomposition into independent "
        "behavioral zones, the tracking of pairwise relationships, and the analysis "
        "of network structure together create a detection capability that addresses "
        "the four structural weaknesses of traditional monitoring: threshold blindness, "
        "single-entity tunnel vision, static risk scores, and sub-tier opacity."
    ))

    add_body_text(doc, (
        "The central insight applies to supply chains as it does to cybersecurity: "
        "the supplier changes WHAT it sources, not HOW MUCH it delivers. The delivery "
        "rate stays strong. The quality score stays acceptable. But the component origins "
        "shifted. The sub-tier manufacturing site changed. The SBOM composition drifted. "
        "These changes are invisible to traditional metrics and visible to behavioral "
        "intelligence — as proven when the same framework caught every advanced threat "
        "that traditional methods missed in the cybersecurity domain."
    ))

    add_callout(doc,
        "Next step: 22nd Century Technologies is prepared to conduct a 4-week proof "
        "of concept on DLA production data, demonstrating BEI detection capabilities "
        "against actual supply chain transaction history across a selected commodity "
        "group."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  END MATTER
    # ══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("22nd Century Technologies, Inc.")
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("V-Intelligence UEBA Program — Supply Chain Risk Management")
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("June 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
