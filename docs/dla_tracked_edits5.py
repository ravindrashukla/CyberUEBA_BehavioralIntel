"""Fifth pass on the DLA REV file:
  F) Add the 'U.S. military superiority' mission sentence into the Combat Support
     Agency paragraph (extends the existing tracked insertion).
  G) Bring Figure 1 'From Metrics to Mission Decision Intelligence' (extracted from
     previous_final) into the executive summary, after the Combat Support Agency
     paragraph, as a tracked insertion with caption.
"""
import os, copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC_PREV = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_previous_final.docx"
REV = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_final_REV_v2.docx"
AUTHOR = "ReviewPanel"; DATE = "2026-06-19T16:00:00Z"
IMG = os.path.join("docs", "_fig1_metrics.png")

# --- extract Figure 1 PNG from previous_final ---
prev = Document(SRC_PREV)
cap_i = next(i for i, p in enumerate(prev.paragraphs)
             if 'From Metrics to Mission' in "".join(x.text or "" for x in p._p.iter(qn('w:t'))))
blip = None
for j in range(cap_i, cap_i - 6, -1):
    bl = prev.paragraphs[j]._p.findall('.//' + qn('a:blip'))
    if bl:
        blip = bl[0]; break
rId = blip.get(qn('r:embed'))
open(IMG, 'wb').write(prev.part.related_parts[rId].blob)
print("extracted Figure 1 ->", IMG, os.path.getsize(IMG), "bytes")

# --- open current ---
d = Document(REV)
_id = [13000]
def nid():
    _id[0] += 1; return str(_id[0])
def insmark():
    e = OxmlElement('w:ins'); e.set(qn('w:id'), nid()); e.set(qn('w:author'), AUTHOR); e.set(qn('w:date'), DATE); return e
def _run(text, italic=False, sz=None, color=None):
    r = OxmlElement('w:r')
    if italic or sz or color:
        rPr = OxmlElement('w:rPr')
        if italic: rPr.append(OxmlElement('w:i'))
        if sz:
            s = OxmlElement('w:sz'); s.set(qn('w:val'), str(sz)); rPr.append(s)
        if color:
            c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
        r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t)
    return r
def ins_wrap(run):
    e = insmark(); e.append(run); return e

# --- F) superiority sentence into para 3 ---
SENT = ("In a contested logistics environment, adapting faster, resolving supply-chain friction earlier, and "
        "preserving decision time are essential to mission assurance, operational resiliency, and U.S. military "
        "superiority. ")
p3 = next(p for p in d.paragraphs
          if 'Combat Support Agency' in "".join(x.text or "" for x in p._p.iter(qn('w:t'))))
done = False
for t in p3._p.iter(qn('w:t')):
    if t.text and 'exposed to disruption. Current planning' in t.text:
        t.text = t.text.replace('exposed to disruption. Current planning',
                                'exposed to disruption. ' + SENT + 'Current planning')
        done = True; break
print("superiority sentence inserted:", done)

# --- G) Figure 1 picture + caption after para 3 ---
# add picture to a temp paragraph so python-docx creates the image part + drawing
tmp_p = d.add_paragraph(); tmp_r = tmp_p.add_run(); tmp_r.add_picture(IMG, width=Inches(6.5))
drawing = copy.deepcopy(tmp_r._r.find(qn('w:drawing')))
# unique docPr id/name
dp = drawing.find('.//' + qn('wp:docPr'))
if dp is not None:
    dp.set('id', '9001'); dp.set('name', 'Figure 1 From Metrics to Mission')
# figure paragraph (centered, tracked)
figp = OxmlElement('w:p'); pPr = OxmlElement('w:pPr')
jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)
rPr = OxmlElement('w:rPr'); rPr.append(insmark()); pPr.append(rPr); figp.append(pPr)
runel = OxmlElement('w:r'); runel.append(drawing); figp.append(ins_wrap(runel))
# caption paragraph (centered, italic gray, tracked)
capp = OxmlElement('w:p'); cpPr = OxmlElement('w:pPr')
cjc = OxmlElement('w:jc'); cjc.set(qn('w:val'), 'center'); cpPr.append(cjc)
crPr = OxmlElement('w:rPr'); crPr.append(insmark()); cpPr.append(crPr); capp.append(cpPr)
capp.append(ins_wrap(_run("Figure 1 - From Metrics to Mission Decision Intelligence", italic=True, sz='18', color='595959')))
# place after para 3
p3._p.addnext(capp); p3._p.addnext(figp)
# remove temp paragraph (image part stays, referenced by the copied drawing)
tmp_p._p.getparent().remove(tmp_p._p)

o = REV; k = 3
while True:
    try:
        d.save(o); break
    except PermissionError:
        root, ext = os.path.splitext(REV); o = f"{root.replace('_v2','')}_v{k}{ext}" if '_v2' in root else f"{root}_v{k}{ext}"; k += 1
print("Saved:", o)
