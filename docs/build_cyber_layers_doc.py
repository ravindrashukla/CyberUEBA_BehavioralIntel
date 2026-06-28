# -*- coding: utf-8 -*-
"""Advanced layered-cybersecurity whitepaper for 22CT / V-Intelligence.
Five DEEP standalone layer sections + embedded defense-in-depth diagram.
Protects proprietary internals (no embedding math, dimensions, CUSUM/drift formulas,
5-phase recipe, 7-innovation blueprint). No em/en dashes. Output: WP DLA/Presentation/.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6); GRAY=RGBColor(0x55,0x5F,0x6B)
d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(10.5)

def H1(t):
    p=d.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
    return p
def H2(t):
    p=d.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=BLUE
    return p
def P(t,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold; return p
def MISSION(t):
    p=d.add_paragraph(); r=p.add_run("Mission: "); r.bold=True; r.font.color.rgb=NAVY; p.add_run(t); return p
def B(label,rest=""):
    p=d.add_paragraph(style='List Bullet')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def Bn(t): return d.add_paragraph(t,style='List Bullet')
def table(headers,rows,widths=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9.5)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""; r=cells[i].paragraphs[0].add_run(v); r.font.size=Pt(9.5)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    d.add_paragraph(); return t
def figure(path,width=6.6,caption=None):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path,width=Inches(width))
    if caption:
        c=d.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=c.add_run(caption); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=GRAY

# ===== TITLE =====
tp=d.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("Defense in Depth for the Age of AI-Enabled Attacks"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
for txt,sz,col,it in [("A Five-Layer Cybersecurity Architecture for Critical Infrastructure",13,BLUE,False),
                      ("Layers, Devices, Applications, Unresolved Challenges, and the Innovation Still Needed",11,GRAY,True)]:
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=p.add_run(txt); rr.font.size=Pt(sz); rr.font.color.rgb=col; rr.italic=it
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
ar=p.add_run("Ravindra Shukla"); ar.font.size=Pt(11); ar.bold=True
p2=d.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
p2.add_run("V-Intelligence  |  Preemptive Defense  +  Behavioral Intelligence  +  API Security  +  Data Governance  +  Code Assurance").font.size=Pt(10)
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Technical Whitepaper  |  For Critical-Infrastructure and Federal Mission Defenders").italic=True
d.add_paragraph()

# ===== EXEC SUMMARY =====
H1("Executive Summary")
P("The cybersecurity contest has changed shape. Apex nation-state adversaries are already resident inside critical "
  "infrastructure, and they no longer need malware to get there. They enter through unproven configuration gaps and "
  "unpatched edge devices, operate with valid credentials and living-off-the-land techniques that cross no threshold, "
  "fragment a single identity across ten or more systems, and stay below every alert for years. At the same time, "
  "offensive artificial intelligence has collapsed the timeline: autonomous agents now reconnoiter, find "
  "vulnerabilities, and exploit them at machine speed, and a single operator can act with the reach of a team.")
P("No single product closes this. Prevention tools sample the attack surface but never prove it is closed. Detection "
  "tools watch for spikes that slow, credential-borne intrusions are designed never to produce. Application and API "
  "defenses rely on signatures that valid-but-malicious workflows walk past. Data controls leak through the new AI "
  "pipelines. And code reaches production faster than humans can review it. Each gap is exactly where the adversary "
  "lives.")
P("This paper presents a layered answer organized around one shared vocabulary (MITRE ATT&CK and ATLAS), covering the "
  "full attack lifecycle from the network edge to the codebase. Each of the five layers is treated as a deep, "
  "standalone section: the threat at that layer, the active devices and components, the existing applications and "
  "tools, the standard process, the unresolved challenges, the V-Intelligence approach, the federal fit, "
  "and the innovation still needed. It closes with a dedicated treatment of countering AI-enabled attacks by fighting "
  "AI with AI across every layer, under human control. The proprietary internals of our methods are deliberately "
  "summarized at the capability level: the proprietary internals (representations, formulas, dimensions, and scoring "
  "logic) are not disclosed here, while architectural structure is described only where it aids understanding.")

# ===== 1 THREAT INFLECTION =====
H1("1. The Threat Inflection: Why the Prior Model Is Failing")
H2("1.1 Apex Adversaries Are Already Inside")
P("Two campaigns reset expectations for what nation-state intrusion looks like.")
B("Volt Typhoon (CISA AA24-038A). ","A People's Republic of China state actor pre-positioned in U.S. energy, water, "
  "transportation, and communications for potential disruption during a crisis. It used only legitimate Windows "
  "tooling (wmic, ntdsutil, netsh, rundll32), produced no malware, no signatures, and no indicators of compromise, "
  "and entered through misconfigured or unpatched small-office and edge devices, including a Fortinet flaw. Dwell time "
  "was assessed at five years or more.")
B("Salt Typhoon (CISA AA25-239A). ","A parallel PRC campaign conducted espionage at carrier scale across U.S. and "
  "allied telecommunications, reaching into lawful-intercept systems, by chaining known, unpatched edge-device "
  "vulnerabilities (including a Cisco IOS XE flaw) with valid-account access. It was slow, sub-threshold, and patient, "
  "and went undetected for years.")
P("The structural lesson of both Typhoons is the same: defenders must prove the door is closed, and then watch the "
  "room for what slips through. Neither preventive proof nor behavioral watch alone is sufficient.")
H2("1.2 Offensive AI Has Collapsed the Timeline")
B("Autonomous, agentic operations. ","In a campaign disclosed in November 2025, a China-linked actor jailbroke an "
  "agentic AI coding assistant and used it to run a largely autonomous espionage operation against roughly thirty "
  "technology, finance, chemical, and government targets, issuing thousands of requests at machine speed with humans "
  "intervening at only a handful of decision points. (The autonomy figure is a vendor self-report and has been "
  "publicly debated, but the direction is not in dispute.)")
B("Self-modifying and AI-written malware ","that calls a model at runtime to rewrite itself for evasion, defeating "
  "signatures.")
B("Autonomous vulnerability discovery and exploitation. ","Frontier models have autonomously exploited the large "
  "majority of one-day vulnerabilities from the public description alone, and autonomous systems at the DARPA AI Cyber "
  "Challenge finals (August 2025) discovered real zero-day vulnerabilities and patched many of them without human "
  "help. The same capability is dual-use: defenders can find their own zero-days first.")
B("Deepfakes and AI social engineering. ","A single deepfake video call drove a finance employee to transfer roughly "
  "25 million U.S. dollars (about HK$200 million) in the Arup case; AI-generated phishing and voice fraud have surged; "
  "the human element remains involved in roughly six of ten breaches (Verizon DBIR).")
P("Independent measurement is catching up: a leading industry breach study attributed roughly one in six breaches in "
  "2025 to attacker use of AI, while adversary breakout time has fallen into the tens of minutes, with extreme cases "
  "under a minute. Gartner projects preemptive cybersecurity will grow from under five percent of security spending in "
  "2024 to about half by 2030, rebalancing away from a detection-and-response-only posture. The rest of this paper operationalizes "
  "that shift across five layers.")
H2("1.3 Two Structural Gaps the Adversary Exploits")
table(["Where defense fails","Why it fails","What the adversary does"],
  [["Prevention","Pentests and scans sample; they never prove. Attack-surface tools see what exists, not what is reachable. Config review is manual and stale within hours. Known-exploited-vulnerability patching covers only a fraction of real paths.",
    "Enters through configuration gaps no one has proven closed, and through unpatched or misconfigured edge devices that cannot host an endpoint agent."],
   ["Detection","Threshold analytics never fire on living-off-the-land. Signature and rule sets lag retooling. Per-source silos make one attacker look like many. Scores say how anomalous, not in which direction. Snapshots replace trajectory.",
    "Uses valid credentials indistinguishable from legitimate use, allowed-by-policy tools that cannot be blanket-blocked, and slow drift that stays below every fixed threshold for months or years."]],
  widths=[1.4,3.0,2.1])

# ===== 2 LAYERED ANSWER + DIAGRAM =====
H1("2. A Layered Answer: One Shared Vocabulary")
P("Each layer catches what the others miss. Together they cover the full attack lifecycle, from the network edge to "
  "the codebase, and they speak one shared language (MITRE ATT&CK for adversary techniques, MITRE ATLAS for "
  "adversarial AI) so that a finding in one layer is intelligible in the next.")
figure("docs/_layer_diagram.png",width=6.9,
       caption="Figure 1. Defense in depth across five layers, one shared MITRE vocabulary, fighting AI with AI under human control.")
table(["Layer","Mission","What it proves or catches"],
  [["1. Network and Preemptive","Close the door before traffic flows","No configuration gap goes unproven; every known attack path is blocked or compensated."],
   ["2. Behavioral and UEBA","Watch the room for what slips through","No malicious behavior goes undetected, including slow, credential-borne, living-off-the-land intrusion."],
   ["3. Application and API","Guard the applications","No API goes undiscovered or unprotected; valid-but-malicious workflows are caught."],
   ["4. Data and Agentic Governance","Protect the data and govern the agents","No sensitive data leaves, or is reached by an ungoverned agent, without detection."],
   ["5. Code, CVE and Supply Chain","Secure the code that builds the mission","No vulnerable or tampered code reaches production."]],
  widths=[1.7,2.0,2.8])
P("The note in this diagram applies throughout: the document expands the application and data aspects of the "
  "architecture into two standalone sections (Layers 3 and 4) for depth, and treats code and supply chain as Layer 5. "
  "Each of the five sections that follow uses the same deep structure.")

# ---- deep layer renderer ----
def deep_layer(secnum, lname, title, mission, intro, threat_intro, threat_b,
               devices, tools, process, challenges, appr_intro, appr_b, fed_intro, fed_b, future):
    H1(f"{secnum}. Layer {lname}: {title}")
    MISSION(mission)
    for x in intro: P(x)
    H2("The Threat at This Layer")
    P(threat_intro)
    for lab,rest in threat_b: B(lab,rest)
    H2("Active Devices and Components")
    for x in devices: Bn(x)
    H2("Existing Applications and Tools")
    for x in tools: Bn(x)
    H2("How It Works Today")
    for x in process: P(x)
    H2("Unresolved Challenges")
    for lab,rest in challenges: B(lab,rest)
    H2("The V-Intelligence Approach")
    P(appr_intro)
    for lab,rest in appr_b: B(lab,rest)
    H2("Federal and Critical-Infrastructure Fit")
    P(fed_intro)
    for x in fed_b: Bn(x)
    H2("Future Innovation Needed")
    for x in future: Bn(x)

# ===== LAYER 1 =====
deep_layer("3","1","Network Perimeter and Preemptive Defense",
  "prove every known attack path is closed before any traffic flows.",
  ["The network perimeter is the set of enforcement points that mediate traffic between trusted and untrusted zones. "
   "Preemptive defense means acting before an attacker is detected, by closing exposures and proving they are closed, "
   "rather than waiting for an alert. This is the shift from sampling to proving and from reactive to preemptive.",
   "Both Typhoon campaigns entered here, through misconfigured or unpatched edge devices using valid accounts. The "
   "perimeter is therefore both where the adversary gets in and the one place where defenders can offer formal, "
   "model-checked proof, over the modeled configuration, that they cannot."],
  "Exploitation of public-facing devices is now the leading initial-access vector, and misconfiguration, not product "
  "flaws, accounts for the large majority of firewall breaches.",
  [("Misconfiguration and drift. ","Industry analysis attributes the overwhelming majority of firewall breaches to misconfiguration; rule sets grow into the tens of thousands of lines and drift from intent within hours."),
   ("Edge-device exploitation. ","Recent campaigns chained vulnerabilities in Ivanti, Palo Alto, Fortinet, and Cisco edge and VPN devices that cannot host an endpoint agent."),
   ("The combinatorial state space. ","Firewall reachability is computationally hard; a multi-vendor path spans an astronomically large configuration state space that sampling cannot cover, so shadowed rules and two-vendor blind spots survive review."),
   ("Reactive lag. ","Signature and rule tools cannot anticipate a path not yet seen, while disclosed vulnerabilities grow toward an order of magnitude more by 2030 and time-to-exploit has gone negative.")],
  ["Next-generation firewalls and the emerging hybrid-mesh firewall (stateful inspection, deep packet inspection, intrusion prevention, TLS decryption).",
   "Intrusion prevention and detection systems; SASE and its security half SSE (secure web gateway, cloud access security broker, zero-trust network access).",
   "Web application firewalls; network segmentation and microsegmentation; network access control.",
   "Identity providers as the zero-trust anchor; routers, switches, VPN concentrators; and internet-facing edge devices."],
  ["Firewall and SASE platforms: Palo Alto Networks, Fortinet, Cisco (including AI-native distributed firewalling), Check Point, Zscaler, Netskope.",
   "Network security policy management: Tufin, AlgoSec, FireMon (rule cleanup, change automation, compliance).",
   "Attack-surface management: CrowdStrike Falcon Surface, Palo Alto Cortex Xpanse, Microsoft Defender EASM, Censys, Tenable, Qualys, RunZero, Axonius.",
   "Breach-and-attack simulation and adversarial exposure validation: SafeBreach, AttackIQ, Cymulate, Picus, Pentera; automated penetration testing: Horizon3.ai, Pentera, Synack, Bishop Fox.",
   "Configuration verification research and the Continuous Threat Exposure Management program model (scope, discover, prioritize, validate, mobilize)."],
  ["Defenders review configurations periodically, run point-in-time penetration tests and vulnerability scans, and "
   "patch against known-exploited-vulnerability lists, now on a risk-based schedule. Attack-surface management "
   "discovers external assets continuously, and policy-management tools attempt to keep rule sets clean.",
   "The common thread is that coverage is sampled and time-bounded, not proven and continuous. A path that was never "
   "sampled, or a rule that drifted after the last review, is exactly where the adversary operates."],
  [("Misconfiguration dominates. ","The overwhelming majority of firewall breaches stem from misconfiguration, not flaws; stolen credentials and edge entry then bypass configuration checks entirely."),
   ("The state space is too large to sample. ","Sampling even a billion paths per second covers effectively none of a multi-vendor configuration state space, so shadowed rules, weak-inspection paths, and cross-vendor blind spots survive."),
   ("Configuration drift. ","Rule sets accumulate shadowed and stale rules and diverge from intent within hours of a change."),
   ("Edge and valid-credential entry. ","Edge devices cannot host agents, and valid accounts block no rule."),
   ("Reactive-only posture. ","Coverage lags the application and threat state; the path not yet seen cannot be pre-blocked.")],
  "V-Intelligence addresses Layer 1 with a preemptive approach that builds a formal model of the defensive "
  "configuration and reasons exhaustively over it, rather than sampling. The objective is completeness: for every "
  "known technique, prove whether any combination of rules allows the traffic. The internal modeling and query "
  "language are proprietary and are summarized here only by outcome.",
  [("Attack intelligence. ","Continuously ingested vulnerability, threat-intelligence, and advisory feeds are compiled into MITRE ATT&CK-enriched attack graphs that update on each new advisory."),
   ("Defense intelligence. ","A symbolic model of the firewall, identity-provider, IPS, SASE, and WAF estate is reasoned over completely, with no sampling, to pinpoint shadowed rules, conflicts, weak inspection, and intent errors across vendors."),
   ("Remediation intelligence. ","A guardrailed, agentic reasoner prescribes precise, risk-prioritized fixes that introduce no new errors, reviewed by humans and then continuously re-verified."),
   ("Operational fit. ","Telemetry-free and agentless, configuration-only, with time-to-value measured in days, and one formal model answering many use cases from posture and drift to vendor migration and compliance.")],
  "Preemptive verification maps cleanly to federal pre-attack expectations.",
  ["NIST Cybersecurity Framework pre-attack functions (Govern, Identify, Protect) and NIST 800-53 boundary-protection controls (SC-7).",
   "CISA continuous-diagnostics and continuous-assurance programs, and policy-as-code.",
   "Deployable as SaaS, on-premises, or air-gapped for sensitive enclaves."],
  ["Scalable formal verification of the full rule state space across heterogeneous vendor estates.",
   "Continuous adversarial exposure validation tied directly to safe, automated remediation.",
   "Automated moving-target defense that proactively reconfigures to deny reachability.",
   "Intent-based verification that catches drift the moment a change diverges from approved baseline."])

# ===== LAYER 2 =====
deep_layer("4","2","Behavioral Intelligence and UEBA (User and Entity Behavior Analytics)",
  "watch the room for what slips through.",
  ["Once the attack surface is provably reduced, the residual threat is the adversary already inside with valid "
   "access. This layer detects malicious behavior, including the slow, credential-borne, living-off-the-land intrusion "
   "that thresholds never catch. As a market, standalone UEBA has folded into SIEM, XDR, and identity threat detection "
   "and response; as a capability, it is indispensable.",
   "The core limitation of the prior generation is that it counts rather than comprehends. Two users may each touch "
   "the same number of files in a week; one reads their own team's documents, the other scans a dozen departments "
   "after hours. A scalar feature sees the same number and the attacker is invisible."],
  "Most intrusions are now malware-free and identity-based. Valid accounts are among the most-used adversary "
  "techniques, stolen credentials are the top initial-access vector, and credential-borne breaches dwell the longest "
  "of any category.",
  [("Living-off-the-land. ","The large majority of high-severity intrusions abuse legitimate tools, so every individual event looks legitimate and no threshold fires."),
   ("Valid-account abuse (MITRE T1078). ","Stolen-credential logins are indistinguishable from normal administration, and static thresholds also break on legitimate spikes such as quarter-end or hiring-season activity."),
   ("Archetype campaigns. ","Volt Typhoon and Salt Typhoon are the reference pattern: slow, valid-account, sub-threshold, multi-year.")],
  ["Security information and event management and extended detection and response platforms.",
   "Endpoint detection and response sensors; identity providers (Active Directory, Microsoft Entra ID, Okta).",
   "VPN, proxy, and firewall logs; PKI and authentication events; cloud access security broker and data-loss-prevention feeds.",
   "Raw authentication, process, network, file, and identity telemetry across on-premises and multi-cloud estates."],
  ["UEBA and SIEM analytics: Splunk (native UEBA in Enterprise Security), Exabeam, Securonix, Microsoft Sentinel UEBA, IBM QRadar (transitioning to Cortex XSIAM), Rapid7, Gurucul.",
   "Identity threat detection and response: CrowdStrike Falcon Identity Protection, Microsoft Defender for Identity, SentinelOne Singularity Identity, Silverfort, Semperis, Vectra AI.",
   "Agentic SOC triage: CrowdStrike Charlotte AI, Microsoft Security Copilot agents."],
  ["UEBA engines baseline normal behavior per entity with unsupervised learning, score deviations against the entity's "
   "own history and a peer group, map findings to MITRE ATT&CK, and feed risk-ranked alerts to a tiered SOC.",
   "The dominant limitation is that scoring is largely scalar and point-in-time: it measures how much a count changed, "
   "not what the entity is becoming, and it analyzes one account at a time rather than the whole subject over time."],
  [("Living-off-the-land and low-and-slow evasion. ","Every event in a slow campaign looks legitimate, so thresholds never fire."),
   ("Valid-account abuse. ","Stolen-credential access mimics normal administration; thresholds also break on legitimate spikes."),
   ("Identity fragmentation. ","One subject is split across ten or more systems, so one attacker appears as many users and the cross-domain chain is hidden; non-human identities now far outnumber humans."),
   ("Black-box, directionless scores. ","An alert that a user is some percent anomalous does not say toward what, and cannot be triaged or audited."),
   ("Alert fatigue and dwell time. ","Roughly half of SOC alerts are false positives, while credential-borne breaches dwell the longest of any vector.")],
  "V-Intelligence treats the same security logs as living behavioral entities and reads behavior as meaning rather "
  "than as scalar counts. The same trajectory engine fielded in the Entity Digital Model for logistics was retargeted "
  "to cyber. The representation, scoring, and detection methods are proprietary; the following describes only what the "
  "capability delivers.",
  [("Behavior as meaning. ","Each entity's activity is represented in a high-dimensional behavioral space, so two identical counts with different meaning separate from one another where a scalar feature would see them as the same."),
   ("Trajectory and direction. ","The system tracks where an entity is heading over time, not just where it stands today, and expresses that direction in MITRE ATT&CK terms (for example, drifting toward lateral movement) rather than as a bare anomaly percentage."),
   ("Identity fusion. ","Fragmented identities across many systems are unified into one behavioral entity, so an attacker stops looking like ten users."),
   ("Multi-front known-bad profiles. ","A multi-front threat-profile detector, the primary detector, matches measurable known-bad behavioral profiles (command-and-control beacon, domain-generation algorithm, living-off-the-land process, cohort-rare access, reconnaissance fan-out, insider collection) scored against cohort peers and fired only on corroboration, with composite embedding scoring as a complementary discovery layer."),
   ("Illustrative result. ","In an internal blind 250-entity evaluation, the threat-profile detector caught all four embedded stealth nation-state-style campaigns at zero false positives, while composite embedding scoring also surfaced all four (at an 8.1 percent false-positive operating point, cleanly separating two of the four) that point-anomaly baselines ranked as normal, with every alert naming the behavior and mapping it to a technique. These are controlled synthetic-data results, indicative rather than a guarantee of field performance, and under composite scoring the stealthiest attackers surface by rank in a risk-ordered list rather than by a clean threshold gap.")],
  "The approach is built for federal sensitivity and explainability.",
  ["Aligns with MITRE ATT&CK and with insider-threat maturity expectations; every verdict names what is happening, where it lives, and who else is affected.",
   "Deployable across IL5, IL6, and JWICS, with a local-model option for air-gapped enclaves.",
   "Privacy-preserving by design, with no personally identifiable information held in the behavioral representation."],
  ["Semantic, embedding-based behavioral detection moving from research into production, where most commercial UEBA is still threshold-based.",
   "Trajectory and sequence modeling replacing point-in-time outlier scoring.",
   "Identity and entity resolution that stitches accounts to one real subject by default.",
   "Explainability mapped to MITRE for every verdict, extended to non-human and AI-agent identities."])

# ===== LAYER 3 =====
deep_layer("5","3","Application and API Security",
  "discover, protect, and govern every application and API before attackers find what defenders do not know exists.",
  ["APIs are the fastest-growing attack surface, and the defining shift of the last two years is that most API abuse "
   "now uses valid, authenticated calls in malicious sequences that signatures do not recognize. Defense is moving "
   "from north-south perimeter filtering to behavioral and runtime detection of legitimate-looking abuse.",
   "The application layer is also where the agentic-AI era lands first: AI agents and the connections that wire them "
   "to enterprise APIs are multiplying faster than governance."],
  "The OWASP API Security Top 10 frames the threat, led by broken object-level authorization. The large majority of "
  "API attacks originate from authenticated sources and target external-facing APIs, and a majority involve "
  "business-logic abuse rather than technical exploits.",
  [("Shadow and zombie APIs. ","Undocumented and deprecated-but-live endpoints are unknown attack surface; several large breaches traced to a single forgotten endpoint exposing tens of millions of records."),
   ("Business-logic abuse. ","Valid calls sequenced maliciously (scraping, transaction automation, resource exhaustion) defeat signatures."),
   ("Agentic AI and Model Context Protocol risk. ","Agent-to-API connections introduce tool poisoning, prompt injection, and insecurely stored tokens, with documented protocol-level vulnerabilities and federal guidance now issued.")],
  ["API gateways and load balancers; web application firewalls.",
   "Application servers and microservices; databases and data-layer services.",
   "Kubernetes and container orchestration; service mesh and east-west service-to-service traffic, typically secured with mutual TLS."],
  ["API security and web-application-and-API protection: Salt Security, Wallarm, Traceable, Akamai, Cloudflare, Imperva, F5, Fortinet.",
   "API gateways: Kong, Google Apigee, AWS API Gateway, Azure API Management.",
   "Bot management and the OWASP API Security Top 10 as the working threat baseline; service meshes such as Istio and Linkerd for east-west control."],
  ["Teams build an API inventory, apply WAF rules and rate limiting, validate requests against published schemas, and "
   "enforce mutual TLS for internal calls.",
   "The control set is largely signature and policy based: effective against known technical exploits, weak against "
   "business-logic abuse, undiscovered endpoints, and ungoverned agent connections."],
  [("Shadow and zombie APIs. ","Unknown, unmonitored endpoints are unprotected attack surface."),
   ("Business-logic abuse. ","Valid-but-malicious sequences bypass signature defenses; most API attacks come from authenticated sources."),
   ("Automated bots and fraud. ","Credential stuffing, account-opening abuse, and scraping, increasingly AI-driven, scale faster than rules adapt."),
   ("Sensitive-data exposure. ","PII and credentials leak in API responses; data exposure is among the top incident classes."),
   ("Ungoverned agentic AI connections. ","Agent workflows and protocol connections wire models to APIs with weak governance.")],
  "V-Intelligence extends behavioral protection to the application and API layer, applying the same "
  "behavioral principle used for users to API consumers and AI agents. The detection logic is proprietary.",
  [("Automated API discovery. ","Finds shadow, zombie, and undocumented APIs across the environment from observed traffic."),
   ("Behavioral bot and abuse defense. ","Distinguishes human, bot, and AI-agent traffic and identifies business-logic abuse that WAFs miss, without CAPTCHAs or friction."),
   ("Runtime API protection. ","Real-time detection at the API layer rather than static rules, with sensitive-data exposure monitoring across internal, external, and third-party traffic."),
   ("Agentic AI gateway. ","Security, governance, and control for AI agent-to-API interactions, so every agent connection is observed and policed.")],
  "The layer aligns to the recognized application-security baselines.",
  ["OWASP API Security Top 10 coverage and unified visibility across internal, external, and third-party traffic.",
   "Alignment with federal guidance on securing the Model Context Protocol and AI-agent connections.",
   "Zero-trust principles extended from users to APIs and agents."],
  ["Behavioral, runtime API protection as the default rather than signature lists.",
   "Agentic-AI gateways that enforce policy at the agent-connection level and generate compliance evidence.",
   "Emerging runtime-security standards for AI agents and LLM applications.",
   "Continuous, traffic-derived inventory that keeps pace with microservice and agent sprawl."])

# ===== LAYER 4 =====
deep_layer("6","4","Data Protection through Agentic AI",
  "protect the data itself, and govern the AI agents that now reach it.",
  ["As enterprises wire large language models and autonomous agents into their data, the protection problem shifts "
   "from guarding files to governing behavior around data, in real time, at machine speed. Agentic AI is hitting this "
   "layer in two directions at once: AI inside the tooling (autonomous classification, triage, and remediation), and "
   "the governance of the AI agents that now touch data.",
   "Retrieval-augmented-generation pipelines and vector stores have become new repositories of sensitive data, and "
   "the prompt has become a new, fast-moving exfiltration channel."],
  "A meaningful share of breaches now involve data nobody knew was there, a large share of attacks use valid accounts, "
  "and shadow AI use has become a leading new exposure, with employees routinely moving corporate data into "
  "general-purpose AI tools through non-corporate accounts.",
  [("Shadow data and sprawl. ","Sensitive data spreads across multi-cloud and SaaS faster than discovery."),
   ("Data into and through LLMs. ","Over-permissioned retrieval and indirect prompt injection have produced zero-click data exfiltration from enterprise AI assistants; the OWASP Top 10 for LLM applications ranks prompt injection and sensitive-information disclosure at the top."),
   ("Ungoverned agents and non-human identity sprawl. ","Most agent incidents trace to over-privileged non-human identities, and a majority of organizations report agents already acting beyond intended scope.")],
  ["Databases, data lakes and warehouses, and object storage.",
   "Data-loss-prevention engines, cloud access security broker and SSE; encryption and key management.",
   "Data catalogs; and increasingly vector databases and retrieval-augmented-generation pipelines as new stores of sensitive data."],
  ["Data security posture management and data detection and response: Cyera, BigID, Securiti, Sentra, Varonis, Wiz, Microsoft Purview.",
   "Data loss prevention: Microsoft Purview, Broadcom/Symantec, Forcepoint, Zscaler, Palo Alto Enterprise DLP, Google Sensitive Data Protection.",
   "Emerging agentic data governance: AI-agent inventory and entitlement control tied to agent identities, and AI-aware data-loss prevention that inspects GenAI prompts."],
  ["Defenders classify data, author and enforce data-loss-prevention policies, encrypt at rest and in transit, and run "
   "periodic access reviews; newer programs extend data-loss prevention to GenAI prompts and plan post-quantum migration.",
   "The model is largely periodic and rule based, which struggles with unstructured data, shadow data, and the new, "
   "fast-moving channel of AI prompts and agent actions."],
  [("Data sprawl and shadow data. ","Discovery cannot keep pace with multi-cloud and SaaS growth."),
   ("Data-loss-prevention weaknesses. ","Pattern-based controls generate false positives, are weak on unstructured data, are bypassable, and miss chat and GenAI channels."),
   ("Insider and valid-access exfiltration. ","A large share of attacks use valid accounts, and a large share of insider incidents are data theft, often near resignation."),
   ("Sensitive data flowing through LLMs. ","Prompt injection, over-permissioned retrieval, and shadow AI move data out through new channels."),
   ("Ungoverned agentic AI and non-human identity sprawl. ","Over-privileged agents act beyond scope; encryption-in-use remains the weakest leg.")],
  "V-Intelligence applies agentic AI to data protection in two directions: AI that governs data-access "
  "behavior, and governance of the AI agents that touch data. Implementation specifics are proprietary.",
  [("Behavioral data-access monitoring. ","Detects exfiltration patterns and abnormal access by valid accounts that volume-based controls miss, by reading the behavior around data rather than only matching content."),
   ("Real-time exfiltration detection. ","Moves from periodic scans to streaming detection and response at the moment data moves."),
   ("AI pipeline governance. ","Monitors sensitive data flowing into prompts, retrieval, and agent actions, and enforces entitlement and guardrails on enterprise GenAI and agent identities."),
   ("Sensitive-data exposure monitoring. ","Flags PII and credential leakage across data stores, APIs, and AI responses, mapped to the same behavioral vocabulary as the other layers.")],
  "The layer aligns to data-centric federal requirements.",
  ["NIST 800-53 data-protection controls and federal guidance on AI data security.",
   "OWASP Top 10 for LLM applications as the working baseline for AI data risk.",
   "Support for post-quantum cryptography migration planning as standards finalize."],
  ["Autonomous data-governance agents that classify, decide access, and remediate continuously.",
   "Data-centric behavioral analytics as a standard control.",
   "End-to-end security for AI and LLM data pipelines, including retrieval entitlement enforcement and vector-store protection.",
   "Practical encryption-in-use through maturing confidential computing."])

# ===== LAYER 5 =====
deep_layer("7","5","Code, CVE, and Software Supply Chain (Mythos)",
  "find vulnerabilities before deployment, not after breach, and secure the code and dependencies that build the mission.",
  ["This layer matters more than ever because exploitation now frequently precedes the patch, and because AI can find "
   "and weaponize vulnerabilities autonomously. The disclosure-and-patch cycle that human cadence can sustain is no "
   "longer fast enough.",
   "Software supply-chain compromise has become a patient, automated discipline, reaching core open-source libraries "
   "and registries that the entire ecosystem depends on."],
  "Disclosed vulnerabilities are growing toward an order of magnitude more by the end of the decade, national "
  "enrichment has fallen behind, and for a growing share of exploited vulnerabilities, time-to-exploit is effectively negative, with exploitation often "
  "preceding patch availability. In parallel, autonomous AI systems can now find and exploit vulnerabilities faster "
  "than human review.",
  [("Supply-chain injection. ","The XZ Utils backdoor (CVE-2024-3094) was a multi-year social-engineering compromise of a core library; self-replicating package worms have poisoned hundreds of packages with billions of weekly downloads; dependency confusion and typosquatting remain structurally unsolved."),
   ("Noise over signal. ","Most container images carry a high or critical vulnerability, but only a fraction are runtime-reachable, so reachable, exploitable findings are buried."),
   ("AI autonomous discovery. ","Models exploit the majority of one-day vulnerabilities from the description alone, and autonomous systems now discover real zero-days.")],
  ["Source code repositories and CI/CD build and deploy pipelines.",
   "Package registries and dependency managers; container image registries.",
   "Infrastructure-as-code templates; runtime application environments."],
  ["Code and dependency security: Snyk, Checkmarx, Veracode, GitHub Advanced Security (CodeQL), Semgrep, SonarQube, Black Duck.",
   "Reachability and supply-chain analysis: Endor Labs, Socket, JFrog; cloud-native and container security: Wiz, Aqua, Sysdig.",
   "Application security posture management to correlate findings; software bill of materials and provenance through SPDX and CycloneDX formats, supply-chain levels for software artifacts, and signed attestation."],
  ["Teams run static, dynamic, and composition analysis in CI/CD, scan containers and infrastructure-as-code, generate "
   "a software bill of materials, patch against known-exploited-vulnerability lists, and prioritize with exploit-"
   "prediction and severity scoring.",
   "The process is increasingly automated, but it is overwhelmed by volume and false positives, and it assumes a patch "
   "exists in time, which the negative time-to-exploit now denies."],
  [("Vulnerability volume and negative time-to-exploit. ","Disclosure outpaces enrichment and patching; exploitation often precedes the fix."),
   ("Supply-chain injection. ","Core libraries, registries, and build systems are targeted by patient, automated campaigns."),
   ("Vulnerable containers and insecure infrastructure-as-code. ","Most images carry critical vulnerabilities and misconfigured templates deploy weaknesses at scale, while only a fraction are reachable, drowning defenders in noise."),
   ("False positives. ","Untuned analysis buries the exploitable findings."),
   ("AI autonomous discovery and exploitation. ","Code that ships unverified is exposed sooner in the new era of machine-speed vulnerability discovery.")],
  "Mythos is the V-Intelligence code and supply-chain assurance capability. It turns AI-driven autonomous "
  "vulnerability discovery to the defender's advantage, finding an organization's own zero-days before adversaries do, "
  "and enforcing that vulnerable code cannot reach production. The engine internals are proprietary.",
  [("Autonomous vulnerability discovery at scale. ","In reported use, the capability surfaced more than two thousand previously-unknown vulnerability findings across production codebases in a multi-week run, a pace human review cannot match. (Findings, not all confirmed exploitable zero-days; reported internally and to be independently validated.)"),
   ("Seven-gate CI/CD pipeline. ","Code, static analysis, dynamic analysis, software composition analysis, container, infrastructure-as-code, and runtime gates each enforce pass or fail, so vulnerable code cannot reach production."),
   ("Supply-chain and provenance assurance. ","Dependency analysis flags compromised or abandoned packages, container scanning catches vulnerabilities before images enter the registry, and a complete software bill of materials is generated and tracked for every release."),
   ("Runtime drift detection. ","Continuous monitoring detects divergence between deployed code and the approved baseline.")],
  "The layer aligns to federal software-supply-chain mandates.",
  ["NIST Secure Software Development Framework and software bill of materials minimum elements.",
   "Supply-chain provenance through artifact levels and signed attestation.",
   "Risk-based vulnerability handling consistent with current federal directives."],
  ["Defensive autonomous vulnerability discovery as standard practice, finding your own zero-days first.",
   "Reachability analysis to cut static-analysis and composition-analysis false positives by an order of magnitude.",
   "End-to-end provenance through signed attestation and supply-chain levels for software artifacts.",
   "Risk-based prioritization that closes the gap left by patch cycles human cadence can no longer meet."])

# ===== 8 AI-ENABLED DEFENSE =====
H1("8. Countering AI-Enabled Cyber Attacks")
P("Innovation at every layer must fight AI with AI across the full attack surface, under human control. The defining "
  "asymmetry is speed: machine-speed attacks outpace human-speed response, so defenders must act autonomously while "
  "keeping strategic control in human hands.")
H2("8.1 How the Adversary Uses AI")
B("Autonomous, agentic operations ","that reconnoiter, find vulnerabilities, exploit, and move laterally at machine speed, so one operator acts with the reach of a team.")
B("Self-modifying and AI-written malware ","that rewrites itself at runtime, defeating signatures.")
B("AI-accelerated vulnerability discovery ","from public descriptions, and autonomous discovery of new zero-days.")
B("AI phishing, deepfake voice and video, and social engineering ","at scale and quality that defeat human judgment, including multi-million-dollar deepfake fraud.")
B("Attacks on AI systems themselves ","through prompt injection, jailbreaks, data and model poisoning, and AI-supply-chain compromise such as malicious models and stolen AI keys.")
H2("8.2 How to Defend, Layer by Layer")
table(["Adversary AI move","Defensive answer","Layer"],
  [["Polymorphic, signatureless malware and valid-credential abuse","Behavioral, meaning-based detection of what code and identities do, not what they are","2, 3, 4"],
   ["Exploiting unproven configuration and edge gaps at speed","Preemptive formal proof that paths are closed, with agentic, guardrailed remediation","1"],
   ["Ungoverned agent-to-API and agent-to-data connections","Agentic AI gateways and agent-identity governance with least privilege","3, 4"],
   ["Autonomous vulnerability discovery and exploitation","Defensive autonomous discovery (find your own zero-days first) and hard CI/CD gates","5"],
   ["Machine-speed campaigns","Agentic SOC triage and response under human-in-the-loop control for high-impact actions","2 (cross-layer)"]],
  widths=[2.4,2.9,1.2])
H2("8.3 Frameworks and Governance")
B("MITRE ATT&CK and MITRE ATLAS ","as the shared vocabulary for adversary techniques and adversarial machine learning.")
B("NIST AI Risk Management Framework and its Generative AI profile ","treating prompt injection and data poisoning as information-security risks.")
B("OWASP Top 10 for LLM applications and OWASP API Security Top 10 ","as working baselines for AI and API risk.")
B("CISA, NSA, and Five Eyes guidance ","on deploying AI securely and on secure deployment of AI agents (verified agent identities, short-lived credentials, human sign-off for high-impact actions), and the security provisions of emerging AI regulation.")
H2("8.4 Securing the Defensive AI Itself")
P("The same jailbreak and indirect prompt injection that subvert an attacker's tool can subvert a defender's SOC "
  "agent. Defensive agentic AI must run with least-privilege agent identities, sandboxed tools, mediated gateways, and "
  "human-in-the-loop approval for safety-critical actions. Explainability is now a security control: analysts trust an "
  "AI verdict that is backed by evidence and mapped to a technique, and regulators increasingly require it. Every "
  "layer in this architecture is therefore built to explain what is happening, where it lives, and who else is "
  "affected, not merely to score.")

# ===== 9 CROSS-CUTTING =====
H1("9. Cross-Cutting Challenges Still Unresolved")
B("Speed asymmetry. ","Attacks break out in minutes; defense at human speed cannot keep pace, yet full autonomy without human control is unacceptable for high-impact actions.")
B("Proof versus sampling. ","Most prevention still samples the attack surface; proving it closed at scale across heterogeneous estates remains hard.")
B("Meaning versus magnitude. ","Most detection still counts; reading behavior as meaning and direction is still moving from research into production.")
B("Identity and agent sprawl. ","Human, machine, and AI-agent identities are fragmented and over-privileged; resolving them to real subjects and governing agent actions is unfinished.")
B("Supply-chain trust. ","Provenance for code, models, and data is immature, while injection attacks grow more patient and automated.")
B("Securing defensive AI. ","The defender's own agents are now attack surface; governing them without losing their speed advantage is an open problem.")
B("Explainability and trust. ","Black-box defense is untenable operationally and increasingly legally; every verdict must be auditable.")

# ===== 10 ROADMAP =====
H1("10. The Innovation Roadmap")
P("The through-line across all five layers is the same: move from sampling to proving, from counting to understanding "
  "meaning and direction, and from human-speed reaction to governed, machine-speed action that a human and an auditor "
  "can both trust. The priority innovations are:")
B("Preemptive proof at scale (Layer 1) ","with formal verification tied to safe, agentic remediation.")
B("Semantic, trajectory-aware behavioral detection (Layer 2) ","with identity fusion and MITRE-mapped explainability, extended to non-human and agent identities.")
B("Behavioral, runtime application and API protection (Layer 3) ","with agentic-AI gateways governing every agent connection.")
B("Agentic data governance (Layer 4) ","with real-time exfiltration detection and end-to-end security for AI data pipelines.")
B("Defensive autonomous vulnerability discovery (Layer 5) ","with reachability-driven prioritization and verifiable supply-chain provenance.")
B("A shared MITRE vocabulary and human-in-the-loop governance ","binding the layers into one explainable, auditable system of defense.")

# ===== 11 CONCLUSION =====
H1("11. Conclusion")
P("The adversary has changed the game on two axes at once: nation-state operators are already inside, using valid "
  "access and living-off-the-land techniques that cross no threshold, and offensive AI has collapsed the timeline to "
  "machine speed. No single tool answers this. A layered architecture does: prove the door is closed, watch the room "
  "for what slips through, guard the applications and APIs, govern the data and the agents that reach it, and secure "
  "the code that builds the mission, all in one shared vocabulary and under human control. Fighting AI with AI across "
  "every layer, with explainable, auditable verdicts, is how critical-infrastructure defenders move from years of "
  "undetected access to immediate identification, and from reactive response to preemptive assurance.")

# ===== REFERENCES =====
H1("References")
H2("Related V-Intelligence Publications")
for x in [
 "V-Intelligence UEBA: A Layered Answer for Critical-Infrastructure Defense (earlier four-layer architecture brief, expanded to five layers in this paper).",
 "Behavioral Entity Intelligence for Cyber Defense (academic whitepaper).",
 "V-Intelligence versus Traditional UEBA (technical and business-friendly editions).",
 "When Normal Hides the Threat (customer whitepaper).",
 "UEBA Behavioral Intelligence Whitepaper (federal editions: Army, DISA, USSOCOM, DLA).",
 "APT Threat Intelligence Business Guide.",
 "Rigor AI: Preemptive Cybersecurity for the Age of AI-Enabled Fast-Moving Attacks (Layer 1 reference).",
 "Decision Intelligence for Defense Logistics: Entity Digital Model and Behavioral Entity Intelligence (cross-domain foundation).",
]: Bn(x)
H2("Selected External References")
for x in [
 "CISA Advisory AA24-038A, PRC State-Sponsored Actors (Volt Typhoon) Compromise of U.S. Critical Infrastructure.",
 "CISA Advisory AA25-239A, PRC State-Sponsored Actors (Salt Typhoon) Targeting of Telecommunications.",
 "Anthropic, Disrupting the First Reported AI-Orchestrated Cyber-Espionage Campaign (GTG-1002), November 2025.",
 "Gartner, Preemptive Cybersecurity press release, September 2025; projection of CVE growth toward one million by 2030.",
 "DARPA AI Cyber Challenge (AIxCC) final results, August 2025; Google Project Zero Big Sleep zero-day discovery.",
 "MITRE ATT&CK; MITRE ATLAS (Adversarial Threat Landscape for AI Systems).",
 "NIST AI Risk Management Framework 1.0 and Generative AI Profile (AI 600-1).",
 "OWASP API Security Top 10 (2023); OWASP Top 10 for LLM Applications (2025).",
 "Verizon Data Breach Investigations Report 2025; IBM Cost of a Data Breach 2025.",
 "XZ Utils backdoor, CVE-2024-3094; CISA Known Exploited Vulnerabilities catalog and risk-based directives.",
]: Bn(x)
d.add_paragraph()
fn=d.add_paragraph(); fr=fn.add_run("Proprietary note: the internal methods, models, formulas, dimensions, and scoring "
  "logic underlying the V-Intelligence, Rigor AI, and Mythos capabilities are proprietary to V-Intelligence "
  "and are summarized in this document at the capability level only.")
fr.italic=True; fr.font.size=Pt(9); fr.font.color.rgb=GRAY

import sys
out=sys.argv[1] if len(sys.argv)>1 else "WP DLA/Presentation/V-Intelligence_Layered_Cybersecurity_Whitepaper.docx"
os.makedirs("WP DLA/Presentation",exist_ok=True)
d.save(out)
print("saved:",out,"| paragraphs:",len(d.paragraphs),"| tables:",len(d.tables))
