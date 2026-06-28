"""Tracked fix: 'pilot Element' -> 'Pilot Element' in Table 7 (casing artifact from an
earlier MVP->pilot replacement)."""
import os, copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REV = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_final_REV_v2.docx"
AUTHOR = "ReviewPanel"; DATE = "2026-06-19T22:00:00Z"
d = Document(REV)
_id = [25000]
def nid():
    _id[0] += 1; return str(_id[0])
def _run(text, rpr=None):
    r = OxmlElement('w:r')
    if rpr is not None: r.append(copy.deepcopy(rpr))
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t); return r
def _wrap(tag, run):
    e = OxmlElement(tag); e.set(qn('w:id'), nid()); e.set(qn('w:author'), AUTHOR); e.set(qn('w:date'), DATE)
    e.append(run); return e

OLD, NEW = "pilot Element", "Pilot Element"
done = False
for tb in d.tables:
    for row in tb.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if not done and run.text and OLD in run.text:
                        full = run.text; i = full.index(OLD)
                        before, after = full[:i], full[i + len(OLD):]
                        rpr = run._r.find(qn('w:rPr'))
                        for t in run._r.findall(qn('w:t')): t.text = before
                        delrun = _run(OLD, rpr)
                        for t in list(delrun.iter(qn('w:t'))):
                            dt = OxmlElement('w:delText'); dt.set(qn('xml:space'), 'preserve'); dt.text = t.text
                            t.getparent().replace(t, dt)
                        run._r.addnext(_run(after, rpr))
                        run._r.addnext(_wrap('w:ins', _run(NEW, rpr)))
                        run._r.addnext(_wrap('w:del', delrun))
                        done = True
o = REV; k = 3
while True:
    try:
        d.save(o); break
    except PermissionError:
        root, ext = os.path.splitext(REV); o = f"{root.replace('_v2','')}_v{k}{ext}"; k += 1
print("fixed 'pilot Element' -> 'Pilot Element':", done, "| saved:", o)
