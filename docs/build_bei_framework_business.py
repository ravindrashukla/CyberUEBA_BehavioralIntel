"""Business-friendly edition of the Behavioral Entity Intelligence (BEI) framework paper.
Keeps the architecture-level technical substance (nine dimensions, three layers,
composite scoring, validated cyber results, novelty/prior-art argument) but redacts
exact implementation logic (formulas, thresholds, scoring math, code). Abbreviations
clarified on first use. ~7-8 pages.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "BEI_Framework_Academic.docx")

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
ORG  = RGBColor(0xC2, 0x5A, 0x12)
GREY = RGBColor(0x55, 0x5F, 0x6B)
GREEN= RGBColor(0x1E, 0x7A, 0x44)
WHITE= RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.1
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(doc.sections[0], m, Inches(0.85))


def H1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    return p

def H2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ORG
    return p

def para(text, size=10.5, color=None, bold=False, italic=False, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def bullet(text, lead=None):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run(text); r2.font.size = Pt(10.5)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def setcell(cell, text, bold=False, size=9.5, color=None, fill=None):
    if fill: shade(cell, fill)
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
    if color: r.font.color.rgb = color

def table(headers, rows, widths=None, hdr_fill='0B1F3A'):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        setcell(t.rows[0].cells[j], h, bold=True, color=WHITE, fill=hdr_fill)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            setcell(t.rows[i].cells[j], v, bold=(j == 0), fill='F4F6F8' if j == 0 else None)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ================= FIGURE GENERATION =================
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

FIGDIR = os.path.join(os.path.dirname(__file__), 'figs_business_bei')
os.makedirs(FIGDIR, exist_ok=True)

C_LEG = '#8A93A0'   # traditional grey
C_LEGD = '#C0392B'  # traditional red accent
C_NEW = '#0891B2'   # BEI teal
C_NAVY = '#0B1F3A'
plt.rcParams.update({'font.family': 'Calibri', 'font.size': 10, 'axes.edgecolor': '#C9CFD6',
                     'axes.labelcolor': C_NAVY, 'xtick.color': '#55607B', 'ytick.color': '#55607B'})


def _box(ax, x, y, w, h, text, fc, ec, tc='white', fs=9.0, bold=True, lw=1.2):
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.012',
                                facecolor=fc, edgecolor=ec, linewidth=lw))
    ax.text(x + w/2, y + h/2, text, ha='center', va='center', fontsize=fs,
            color=tc, fontweight='bold' if bold else 'normal')

def _arrow(ax, x1, y1, x2, y2, color=C_NAVY, lw=1.4):
    ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle='-|>',
                                 mutation_scale=11, color=color, lw=lw))

def fig1_concept():
    """One framework, many domains — concept diagram."""
    fig, ax = plt.subplots(figsize=(7.2, 3.7))
    ax.set_xlim(0, 10); ax.set_ylim(0, 10); ax.axis('off')

    # Left column: domain chips
    domains = ['Cyber security\n(insider / APT)', 'Defense supply chain\n(SCRM / SBOM)',
               'Maintenance\n(equipment health)', 'Healthcare & insurance\n(fraud)',
               'Financial crime\n(AML)', 'Counter-intel /\ncounter-terror / energy']
    y0 = 8.55
    for i, d in enumerate(domains):
        yy = y0 - i * 1.55
        _box(ax, 0.15, yy, 2.35, 1.25, d, '#EEF2F6', '#C9CFD6', tc=C_NAVY, fs=7.8)
        _arrow(ax, 2.55, yy + 0.62, 3.25, 5.0 + (yy + 0.62 - 5.0) * 0.28, color='#8A93A0', lw=1.1)

    # Central framework pipeline (vertical stack)
    steps = [('Digital entity per subject\n(a living behavioral "medical chart")', C_NEW),
             ('Embedded in one unified semantic space\n(a shared "map of meaning")', C_NEW),
             ('Tracked over time → trajectory & drift\n(speed, persistence, direction)', C_NEW),
             ('Drift direction NAMED against\nknown threat / risk concepts', C_NAVY)]
    ys = 7.15
    for i, (s, c) in enumerate(steps):
        yy = ys - i * 1.83
        _box(ax, 3.35, yy, 3.55, 1.45, s, c, c, fs=8.0)
        if i < len(steps) - 1:
            _arrow(ax, 5.12, yy - 0.04, 5.12, yy - 0.36, color=C_NAVY)
    ax.text(5.12, 9.55, 'ONE FRAMEWORK — Behavioral Entity Intelligence',
            ha='center', fontsize=10.5, fontweight='bold', color=C_NAVY)
    ax.text(5.12, 9.02, '9 dimensions  ·  3 layers  ·  5 scoring phases',
            ha='center', fontsize=8.6, color='#55607B')

    # Right: outputs
    outs = ['Ranked anomaly list\n(composite score)', 'Which behavioral zone\nchanged — and toward\nwhat named pattern',
            'Self + peer-cohort\nverdict per entity']
    for i, o in enumerate(outs):
        yy = 7.3 - i * 2.1
        _box(ax, 7.55, yy, 2.3, 1.55, o, '#FFFFFF', C_NEW, tc=C_NAVY, fs=8.0)
        _arrow(ax, 6.95, yy + 0.77 - (i - 1) * 0.0, 7.5, yy + 0.77, color=C_NEW, lw=1.3)
    ax.text(9.92, 0.12, 'Illustrative', ha='right', fontsize=8, style='italic', color='#8A93A0')
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig1_concept.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig2_detection():
    """Traditional methods vs BEI: attacks detected out of 4 (real system results)."""
    methods = ['One-Class SVM\n(traditional)', 'Isolation Forest\n(traditional)',
               'Local Outlier Factor\n(traditional)', 'Z-Score\n(traditional)',
               'BEI composite\nscoring']
    vals = [0, 0, 0, 1, 4]
    colors = [C_LEG, C_LEG, C_LEG, C_LEGD, C_NEW]
    fig, ax = plt.subplots(figsize=(6.8, 3.0))
    bars = ax.barh(methods, vals, color=colors, height=0.62)
    for b, v in zip(bars, vals):
        if v == 0:
            ax.text(0.06, b.get_y() + b.get_height()/2, '0 of 4', va='center',
                    fontsize=9, color='#55607B', fontweight='bold')
        elif v == 4:
            ax.text(v - 0.07, b.get_y() + b.get_height()/2,
                    'all 4 of 4 detected  ·  8.1% false-positive rate', va='center',
                    ha='right', fontsize=9.5, color='white', fontweight='bold')
        else:
            ax.text(v - 0.07, b.get_y() + b.get_height()/2, f'{v} of 4', va='center',
                    ha='right', fontsize=9.5, color='white', fontweight='bold')
    ax.set_xlabel('Attack campaigns detected (of 4 embedded in 250 users, 485 days)')
    ax.set_xlim(0, 4.3); ax.set_xticks([0, 1, 2, 3, 4])
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig2_detection.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig3_direction():
    """The core insight: volume flat, composition shifts (illustrative)."""
    wk = np.arange(1, 17)
    total = 100 + 3 * np.sin(wk / 2.1) + np.array([1,-2,2,0,-1,1,-2,0,2,-1,1,0,-2,1,0,-1])
    restricted_share = np.clip(0.04 + 0.05 * np.maximum(wk - 4, 0), 0.04, 0.62)
    restricted = total * restricted_share
    public = total - restricted
    fig, ax = plt.subplots(figsize=(7.0, 2.9))
    ax.bar(wk, public, color=C_LEG, alpha=0.75, label='Public / routine file access')
    ax.bar(wk, restricted, bottom=public, color=C_LEGD, alpha=0.85,
           label='Restricted / sensitive file access')
    ax.plot(wk, total, color=C_NAVY, lw=2.0, ls='--',
            label='Total volume — all traditional monitoring sees (flat → no alarm)')
    ax.annotate('WHAT changed,\nnot HOW MUCH', xy=(13, 48), xytext=(9.0, 22),
                fontsize=9.5, fontweight='bold', color=C_NAVY,
                arrowprops=dict(arrowstyle='->', color=C_NAVY, lw=1.3))
    ax.set_xlabel('Week'); ax.set_ylabel('File accesses')
    ax.set_ylim(0, 145)
    ax.legend(loc='upper left', fontsize=8.2, frameon=False)
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig3_direction.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig4_ranking():
    """Composite-score ranking: 4 attackers among 250 users (all points = system output)."""
    csv_path = os.path.join(os.path.dirname(__file__), '..', 'data', 'tier3_results', 'composite_scores.csv')
    attackers = {1: (51.3, 'USR-118\nSalt Typhoon'), 2: (46.2, 'USR-156\nInsider'),
                 7: (19.4, 'USR-234\nSlow APT (C2)'), 24: (13.7, 'USR-042\nVolt Typhoon LOTL')}
    curve_label = 'All 250 users ranked by composite score (system output)'
    try:
        import csv as _csv
        with open(csv_path, newline='') as fh:
            scores = sorted((float(r['composite']) for r in _csv.DictReader(fh)), reverse=True)
        scores = np.array(scores)
    except OSError:
        # fallback: monotone curve pinned through the attacker points (illustrative)
        ranks_f = np.arange(1, 251)
        scores = 19.0 * np.exp(-(ranks_f - 1) / 30.0) + 2.0
        for rk, (sc, _) in attackers.items():
            scores[rk - 1] = sc
        scores = np.sort(scores)[::-1]
        curve_label = 'All 250 users ranked by composite score (illustrative curve)'
    ranks = np.arange(1, len(scores) + 1)
    fig, ax = plt.subplots(figsize=(7.0, 3.1))
    ax.plot(ranks, scores, color=C_LEG, lw=1.6, label=curve_label)
    offsets = {1: (18, 51.0), 2: (30, 43.5), 7: (55, 25.0), 24: (88, 16.5)}
    for rk, (sc, lab) in attackers.items():
        ax.plot(rk, sc, 'o', ms=8, color=C_LEGD, mec='white', mew=1.2, zorder=5)
        ox, oy = offsets[rk]
        ax.annotate(f'{lab}  ·  score {sc}  ·  rank #{rk}', xy=(rk, sc), xytext=(ox, oy),
                    fontsize=8.4, color=C_NAVY, fontweight='bold',
                    arrowprops=dict(arrowstyle='-', color='#8A93A0', lw=0.9))
    ax.plot([], [], 'o', color=C_LEGD, label='Attack users (system output)')
    ax.set_xlabel('Rank among 250 users (by composite anomaly score)')
    ax.set_ylabel('Composite score')
    ax.set_xlim(0, 255); ax.set_ylim(-2, 58)
    ax.legend(loc='upper right', fontsize=8.4, frameon=False)
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig4_ranking.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

FIG1, FIG2, FIG3, FIG4 = fig1_concept(), fig2_detection(), fig3_direction(), fig4_ranking()

def add_fig(path, caption, width=6.4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY


# ================= TITLE =================
r = doc.add_paragraph().add_run('Behavioral Entity Intelligence:\nOne Framework for Detecting What Threshold Analytics Miss')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
para('A business and technology overview of the framework, its operation, its areas of application, and the '
     'supporting evidence',
     12.5, GREY, after=2)
para('22nd Century Technologies, Inc.  ·  V-Intelligence Program  ·  June 2026  ·  Business Edition',
     9.5, GREY, after=2)
para('Disclosure note: this edition describes the architecture and reports measured results. The '
     'exact model logic — mathematical formulas, scoring weights, and detection thresholds — is proprietary and '
     'withheld here; it is available in the full technical edition under NDA (non-disclosure agreement).',
     9.5, ORG, italic=True, after=4)
para('Data note: all quantitative results come from a synthetic (simulated) cybersecurity dataset built to prove the '
     'concept — 250 users, 485 days, ~14 million events, 4 embedded attack campaigns — and must be re-validated on '
     'real operational data before any production claim. Of the twelve application domains described, only '
     'cybersecurity is empirically validated to date.', 9.5, ORG, italic=True, after=8)

# ================= 1. EXEC SUMMARY =================
H1('1.  Executive Summary')
para('Most organizations monitor entities of some kind: users, suppliers, patients, equipment, accounts, or '
     'personnel. The monitoring tools they use largely ask one question, namely whether any single metric is '
     'abnormally large or small. That question tends to miss the most damaging threats, because capable actors '
     'operate within normal statistical ranges. An insider may access the same number of files while shifting from '
     'public to restricted content. A supplier may deliver the same volume while moving sourcing to '
     'adversary-nation origins. A fraudulent provider may submit the same number of claims while steadily coding '
     'them higher. In each case the volume appears normal; the entity has changed what it does rather than how much.')
para('Behavioral Entity Intelligence (BEI) is an analytical framework organized around a single premise: represent '
     'any observable subject as a digital entity, a living behavioral profile comparable to a patient\'s medical '
     'chart for a user, a supplier, or a machine. Every entity is placed in one unified semantic space, a '
     'mathematical "map of meaning" in which entities that behave alike sit near one another. The framework tracks '
     'each entity\'s movement through that space over time and measures the direction of its behavioral drift, not '
     'only its magnitude. That direction is then matched against named threat concepts, so the system reports not '
     'merely that an entity is unusual but what the entity appears to be becoming.')
para('Stated concisely: where traditional analytics measure the magnitude of change, Behavioral Entity Intelligence '
     'measures its direction and assigns it a name.', bold=True)
bullet('In a cybersecurity proof of concept (250 simulated users, 485 days, 4 embedded multi-month attack '
       'campaigns), BEI detected 4 of 4 attacks at an 8.1% false-positive rate. The strongest traditional method '
       'caught 1 of 4, and three standard machine-learning detectors caught none.', lead='Validated detection. ')
bullet('BEI also reports which behavioral dimension changed and toward what threat pattern, which allows an analyst '
       'to distinguish an insider from an external campaign at triage time. None of the traditional methods tested '
       'provides this.', lead='Directional intelligence. ')
bullet('Twelve mission domains, ranging from supply-chain risk to anti-money laundering to clinical-trial '
       'integrity, map onto the same architecture using data that organizations already collect.',
       lead='One framework, many domains. ')

# ================= 2. PROBLEM =================
H1('2.  The Problem: Why Traditional Analytics Fail')
para('Across domains, the most damaging threats share a common characteristic: they operate within the bounds of '
     'authorized activity. The insider uses valid credentials. The fraudulent provider submits properly formatted '
     'claims. The money launderer keeps each transaction below reporting thresholds. The compromised supplier '
     'delivers on time and on specification while gradually substituting components. Traditional analytics tend to '
     'fail against this class of threat for three structural reasons.')
bullet('Rule-based systems alarm on individual metrics, such as transactions above $10,000, logins above 50, or '
       'delivery delays above 5 days. A capable actor can stay below every individual threshold while achieving the '
       'objective through small, simultaneous changes across many dimensions.', lead='Threshold blindness. ')
bullet('Statistical anomaly detection (z-scores, isolation forests, clustering) measures how far a metric deviates '
       'from the average. A benign seasonal spike and a malicious diversion can produce the same statistical '
       'magnitude, so the kind of deviation remains invisible.', lead='Magnitude-only measurement. ')
bullet('Traditional systems analyze each entity in isolation. The supplier appears normal and the part appears '
       'normal, but the relationship between them has changed, because that specific part is now sourced from a '
       'different country. No single-entity analysis can detect this.', lead='Single-entity tunnel vision. ')
para('The gap is structural rather than parametric. Better thresholds, additional features, and larger training '
     'sets do not close it. Closing it requires measuring behavioral direction, across entity relationships, over '
     'time.')
add_fig(FIG3, 'Figure 1 — The core insight. Total volume stays flat (dashed line), so every traditional monitor stays '
              'quiet — while the composition shifts from public to restricted content. (Illustrative.)')

# ================= 3. THE BIG IDEA =================
H1('3.  The Core Concept: Digital Entities in a Shared Map of Meaning')
para('Consider the central concepts in plain terms. A digital entity, sometimes called a digital twin, is a living '
     'software profile of a real-world subject. A patient\'s medical chart is a useful analogy: it does not merely '
     'record today\'s temperature but accumulates history, vitals, and context, and a clinician reads its trend in '
     'order to act before a condition becomes acute. BEI maintains such a chart for every entity it watches, '
     'whether a user, supplier, machine, account, or person of interest.')
para('The second concept is the semantic space. Each entity\'s behavior is periodically written out as a structured '
     'description and converted by an embedding model into a numerical vector. This is the same representational '
     'machinery that allows large language models to interpret language, applied here to behavior. The result is a '
     'unified, high-dimensional "map of meaning" in which entities whose behavior is similar sit near one another '
     'even when their raw numbers differ. Where a language model uses that map to interpret words, BEI uses it to '
     'interpret what entities do.')
para('The third concept is drift. Because every entity is re-profiled on a regular cadence, weekly for users and '
     'monthly for suppliers and equipment, it traces a trajectory across the map. Drift is that movement, '
     'characterized by its speed, whether it is accelerating, whether it persists, and, most importantly, its '
     'direction. The direction is compared against a library of named threat and risk concepts placed on the same '
     'map, and the strongest match names what the entity is becoming, for example "insider data staging," "supply '
     'disruption," or "command-and-control beaconing." This concept-direction step converts an anomaly score into '
     'an explanation.')
para('The underlying vector mathematics, including how zone profiles compose into an entity vector, how drift and '
     'its direction are computed, how slow drift is accumulated so that months of small steps still trigger an '
     'alarm, and how pairwise relationship signatures are formed, is proprietary and withheld in this edition. The '
     'relevant point here is that the geometry is identical in every domain; only what is described into the '
     'profile changes.', italic=True)
add_fig(FIG1, 'Figure 2 — One framework, many domains. Each domain supplies its own entities and data; the analytical '
              'core never changes. (Illustrative.)')

# ================= 4. NINE DIMENSIONS =================
H1('4.  Nine Dimensions of Entity Behavior')
para('BEI characterizes every entity through nine complementary dimensions: six aspects of entity state and '
     'interaction, together with three analytical lenses that turn raw observations into intelligence.')
table(['Dimension', 'Question answered', 'Example signal'],
      [('1. Identity', 'WHO or WHAT is it?', 'Role, type, clearance, registration — and changes to them.'),
       ('2. Behavior', 'WHAT does it do?', 'Activity decomposed into independent zones (a user\'s access, data, network, risk zones) so one drifting zone is not diluted by the others.'),
       ('3. Time', 'HOW is it evolving?', 'Drift speed, acceleration, sudden regime shifts, slow cumulative change.'),
       ('4. Space', 'WHERE does it act?', 'A login from home is routine; from a never-visited country it is not.'),
       ('5. Relationship', 'HOW do pairs behave together?', 'User on an unfamiliar device; supplier sourcing one part from a new country — while both individuals look stable. Each pair gets its own signature, deliberately not an average of the two profiles.'),
       ('6. Network', 'WHAT is its ecosystem position?', 'A clean Tier-1 supplier whose Tier-3 sub-source sits in an adversary nation; "independent" entities moving in lockstep.'),
       ('7. Sequence', 'In what ORDER do events occur?', 'Order→Accept skipping inspection; kill-chain stage progression; treatment coded before diagnosis.'),
       ('8. Periodicity', 'WHAT is its rhythm?', 'Malware beaconing every 6 hours with clockwork regularity — the rhythm, not the volume, is the signal.'),
       ('9. Context', 'WHAT does it MEAN here?', 'The same demand spike is expected during deployment prep, alarming during drawdown.')],
      widths=[1.2, 1.7, 3.9])
para('The relationship dimension detects what no single-entity system can. Entity A appears normal and entity B '
     'appears normal, yet the A×B pairing has drifted. This is the primary detection signal for supply-chain '
     'substitution, credential sharing, and referral fraud.')

# ================= 5. THREE LAYERS =================
H1('5.  Three Layers and Two Baselines')
table(['Layer', 'Analyzes', 'Example finding'],
      [('1. Entity', 'Individual behavioral profiles', 'A user\'s data-access pattern shifts from public to restricted files at constant volume.'),
       ('2. Relationship', 'Pairwise interaction dynamics', 'A supplier\'s quality for ONE specific part degrades while every other metric stays clean — targeted substitution.'),
       ('3. Network', 'The full graph of relationships', 'A Tier-3 manufacturer with adversary-nation ties behind two clean intermediaries; "independent" providers billing in lockstep — a kickback ring.')],
      widths=[1.2, 2.2, 3.4])
para('Every entity is judged against two baselines, and an entity that fails both represents the framework\'s '
     'highest-confidence finding.')
bullet('Whether the entity is behaving as it previously did, measured as drift from its own established pattern.',
       lead='Self-baseline. ')
bullet('Whether the entity is behaving like others of its kind. Entities are grouped into cohorts from stable '
       'identity attributes, with users grouped by role (administrator, security, developer, business, executive) '
       'and suppliers by commodity, geography, and tier, and each is scored against the normal members of its own '
       'cohort. Behavior that is unremarkable for a system administrator may be alarming for an accountant, and the '
       'wrong reference group will hide the signal.', lead='Peer-cohort baseline. ')
para('Cohorts also produce a signal of their own. When many members of one cohort drift together in the same '
     'direction, the correlated movement points to a shared cause, such as a common sub-tier dependency, a market '
     'shock, or a coordinated campaign.')

# ================= 6. COMPOSITE SCORING =================
H1('6.  Multi-Phase Composite Scoring')
para('At every layer, detection fuses five independent phases into a single ranked score. A normal entity may score '
     'high on one phase by chance; a genuinely anomalous entity scores high on several at once. The phases (exact '
     'weights and thresholds proprietary):')
table(['Phase', 'What it measures'],
      [('Signal strength', 'How extreme the entity\'s strongest deviations are, relative to its peer group.'),
       ('Breadth', 'How many behavioral dimensions deviate at once — multi-dimension patterns vs. single-dimension noise.'),
       ('Sustained deviation', 'Whether the anomaly persists across periods — campaigns and degradation vs. transient events.'),
       ('Context divergence', 'Whether the entity flags consistently under threat-specific lenses (an "insider investigation" lens vs. an "APT hunt" lens; APT = Advanced Persistent Threat).'),
       ('Novelty persistence', 'Whether previously-unseen behaviors (new contacts, routes, components) keep recurring.')],
      widths=[1.7, 5.1])
para('No single phase is sufficient on its own, which is by design. An actor who suppresses signal strength will '
     'still fail on breadth, and one who limits breadth will still fail on sustained deviation or novelty '
     'persistence. The combination forms a detection surface that no single evasion technique defeats.', bold=False)

# ================= 7. DOMAINS =================
H1('7.  One Framework, Twelve Domains')
para('Twelve mission domains have been mapped onto the architecture. In each, only three things change: what the '
     'entities are, what behavioral zones describe them, and which relationships matter. The data required is data '
     'that organizations already collect, with no new instrumentation in most cases. Four examples follow.')
table(['Domain', 'Entities & cadence', 'The anomaly traditional analytics miss'],
      [('Cybersecurity\n(validated)', 'Users, devices, applications, network segments — weekly snapshots from authentication, file, network, DNS (Domain Name System — internet address lookups), and endpoint logs.',
        'Insider shifting from public to restricted access at constant volume; a command-and-control beacon whose persistence, not volume, is the tell; "living off the land" abuse of admin tools.'),
       ('Defense supply chain', 'Suppliers, parts (NSNs — National Stock Numbers), depots, manufacturers — weekly/monthly snapshots from procurement, delivery, quality, and country-of-origin data.',
        'Supplier gradually shifts sourcing to adversary-nation origins while delivery and price stay clean; a bill of materials drifts without notification; two "independent" suppliers fail in correlation.'),
       ('Prescriptive maintenance', 'Equipment units, components, facilities — daily/weekly sensor aggregates.',
        'The relationship between engine speed (RPM), temperature, and fuel flow shifts 8% over 6 weeks — degradation invisible to every individual sensor threshold; drift also predicts WHAT parts will be needed WHEN.'),
       ('Healthcare fraud', 'Providers, patients, procedures, diagnoses — monthly claim aggregates.',
        'A provider upcodes step by step over 18 months — each claim individually valid, the drift is the fraud; a referral relationship suddenly intensifies (kickback); patient clusters too similar to be real.')],
      widths=[1.3, 2.4, 3.1])
para('The remaining domains, namely insurance fraud, clinical-trial integrity, counter-intelligence, anti-money '
     'laundering, counter-terrorism, counter-narcotics, energy and critical infrastructure, and workforce '
     'analytics, follow the same template. In anti-money laundering, for instance, an account maintains normal '
     'total volume while its counterparty geography drifts toward sanctioned jurisdictions, with every transaction '
     'kept below reporting thresholds.')

# ================= 8. RESULTS =================
H1('8.  Validated Results: the Cybersecurity Proof of Concept')
para('The framework was validated end-to-end in the V-Intelligence UEBA program (UEBA = User and Entity Behavior '
     'Analytics, the cybersecurity discipline of profiling how users and machines behave). The test bed is '
     'synthetic and intentionally demanding: 250 simulated users over 485 days (January 2025 to May 2026), '
     'approximately 14 million events across 5 log sources, with 4 long-duration attack campaigns embedded. These '
     'comprise an insider threat, a slow APT running command-and-control (C2) beaconing, a Volt Typhoon-style '
     '"living-off-the-land" (LOTL) campaign abusing legitimate admin tools, and a Salt Typhoon-style telecom '
     'intrusion. By construction the attacks are statistically quiet across all 23 individual behavioral features. '
     'The top-ranked attacker\'s largest single-feature deviation is a z-score of 1.71 (a z-score counts how many '
     'standard deviations a value sits from its average), well below the conventional three-standard-deviation '
     'alarm line, and only one of the four (the Volt Typhoon actor, at 3.04) crosses that line on any feature.')
table(['Attack user', 'Campaign', 'Duration', 'Rank (of 250)', 'Primary detection signal'],
      [('USR-118', 'Salt Typhoon (telecom)', '412 days', '#1 (score 51.3)', 'Signal strength + breadth across multiple behavioral zones.'),
       ('USR-156', 'Insider threat', '14 months', '#2 (score 46.2)', 'Sustained deviation over a 14-month escalation; data-behavior drift.'),
       ('USR-234', 'Slow APT (C2 beacon)', '417 days', '#7 (score 19.4)', 'Novelty persistence: the same beacon destination (IP address) appears in 60 of 60 post-baseline weeks.'),
       ('USR-042', 'Volt Typhoon (LOTL)', '412 days', '#24 (score 13.7)', 'Context divergence + signal strength from LOTL tool patterns.')],
      widths=[0.95, 1.5, 0.85, 1.25, 2.25])
para('On the same data, BEI\'s composite scoring detected 4 of 4 campaigns at an 8.1% false-positive rate. Z-Score '
     'analysis caught 1 of 4 (only the Volt Typhoon actor) at a higher 9.8% false-positive rate. Local Outlier '
     'Factor (4.5% false positives), Isolation Forest (5.3%), and One-Class SVM (Support Vector Machine, 14.6%) '
     'each caught 0 of 4, and none offers any directional explanation for triage. BEI additionally reports which '
     'behavioral zone changed and toward what named threat pattern, which allows an analyst to distinguish the '
     'insider from the APT before opening a single log.')
add_fig(FIG2, 'Figure 3 — Detection head-to-head on the same test bed: traditional detectors catch at most one of the '
              'four campaigns; BEI\'s composite catches all four. (System output, synthetic data.)')
add_fig(FIG4, 'Figure 4 — Where the four attackers land among all 250 users: the full ranked composite-score curve, '
              'with the four attack users marked. (System output, synthetic data.)')
para('A candid reading of the numbers is warranted. An 8.1% false-positive rate means that 20 of the 246 normal '
     'users also rank high enough to warrant review, which is acceptable for a weekly triage workflow and is the '
     'declared trade-off for catching slow, low-magnitude campaigns. One attacker (living off the land) ranks #24 '
     'rather than #1; it is found because composite scoring aggregates weak signals, but it would not survive a '
     '"top-10 only" review policy. Another (the slow C2 beacon) is surfaced principally by the novelty-persistence '
     'phase, and removing that one phase lowers its score from 19.4 to 6.4 and its rank from #7 to roughly #105. '
     'All figures are synthetic and must be re-validated on production telemetry.')

# ================= 9. NOVELTY =================
H1('9.  Novelty and Prior Art')
para('A structured, source-verified review of commercial products and the published literature (2024–2026) found no '
     'commercially marketed system, and no single published method, that combines the elements of this framework. '
     'Commercial UEBA platforms (Exabeam, Securonix, Microsoft Sentinel UEBA) detect via statistical profiling, '
     'machine-learned baselines, and rarity- and magnitude-based scoring; their generative-AI features are '
     'conversational copilots, not the detection mechanism. Commercial supply-chain platforms (SAP IBP, o9, Kinaxis) '
     'forecast with machine learning over tabular features, again with generative AI as a conversational layer.')
para('We are explicit about prior art. The constituent techniques exist individually in the literature, including '
     'embedding serialized behavior with language models for anomaly detection, converting raw metrics to prose '
     'before embedding, embed-then-detect pipelines, and measuring embedding drift over time (citations in the '
     'technical edition). Each of them, however, detects by magnitude and operates in a single domain. None '
     'projects the direction of drift onto named concepts, and none unifies entities of every kind in one semantic '
     'space.')
para('The contribution lies in the integrated combination rather than in any single technique. It comprises one '
     'semantic space for all entity types, drift measured as a trajectory, the drift direction named against domain '
     'concepts, relationship and network layers above the entity layer, and multi-phase composite scoring. To the '
     'best of our knowledge, no commercial platform or single published system offers this combination.', bold=True)
para('Scoping notes: vendor internals are not fully public (the comparison is against marketed methodology), and a '
     'formal patent / freedom-to-operate review is recommended before any contractual assertion of uniqueness.',
     9.5, GREY, italic=True)

# ================= 10. DEPLOYMENT =================
H1('10.  Implementation and Deployment')
bullet('A containerized, database-backed service comprising a vector-enabled database (for fast similarity search '
       'over behavioral profiles), an API layer for ingestion and query, an automated profiling-and-scoring '
       'pipeline, and an interactive investigation dashboard.', lead='Platform. ')
bullet('The same platform serves every domain through configuration of entity definitions, behavioral zones, and '
       'relationship mappings, rather than through re-engineering.', lead='Domain-agnostic by design. ')
bullet('Deployable in enterprise data centers, secure enclaves, tactical edge, or cloud, and consumes data sources '
       'that organizations already operate.', lead='Footprint. ')
bullet('A staged rollout: configuration (1-2 weeks), data integration (2-4 weeks), behavioral baselining (4-6 '
       'weeks), detection tuning and false-positive validation (2-4 weeks), and then continuous production '
       'monitoring.', lead='Path to production. ')

# ================= 11. LIMITATIONS =================
H1('11.  Limitations')
bullet('All quantitative results are synthetic; detection and false-positive rates must be re-measured on real '
       'operational telemetry.')
bullet('Only one of the twelve domains (cybersecurity) is empirically validated; the other eleven are designed '
       'applications, not yet tested on data.')
bullet('The 8.1% false-positive rate is workable for weekly triage but is not alert-grade precision; production '
       'tuning will be domain- and customer-specific.')
bullet('One attack class (living off the land) ranks #24 of 250 — detected, but dependent on a review process that '
       'looks beyond a short top-N list.')
bullet('Each new domain requires a baselining period (weeks to a year, by the domain\'s natural cadence) before '
       'detection is meaningful.')
bullet('Profiling depends on an external or locally hosted embedding model; air-gapped deployment requires '
       'characterizing a local replacement.')

# ================= 12. CONCLUSION =================
H1('12.  Conclusion')
para('Sophisticated threats change what they do rather than how much. They remain inside every volume threshold '
     'while shifting behavioral direction, which is precisely what threshold- and magnitude-based analytics cannot '
     'see. Behavioral Entity Intelligence builds a behavioral chart for every entity, places all of them in one '
     'shared map of meaning, tracks each entity\'s trajectory, and names the direction of its drift against known '
     'threat patterns at the entity, relationship, and network levels. In its first validated domain it found all '
     'four embedded attack campaigns where the strongest traditional method found one, and it identified for the '
     'analyst what kind of attack each one was.')
para('The framework combines nine behavioral dimensions, three analytical layers, and five scoring phases in one '
     'architecture applicable from cybersecurity to supply chain to healthcare to counter-intelligence. Where '
     'traditional analytics measure the magnitude of change, Behavioral Entity Intelligence measures its '
     'direction.', bold=True)

# ================= GLOSSARY =================
H1('Appendix — Glossary')
gl = [
    ('BEI', 'Behavioral Entity Intelligence — digital behavioral profiles of entities, tracked in a shared semantic space, with the direction of change measured and named.'),
    ('Digital entity / digital twin', 'A living software profile of a real-world subject — a patient\'s medical chart for a user, supplier, or machine: history, vitals, relationships, trend.'),
    ('Semantic space / embedding', 'A mathematical "map of meaning" built from numerical vectors (embeddings) — the core representation inside large language models; alike behavior sits near alike.'),
    ('Drift', 'An entity\'s movement through the semantic space over time — speed, persistence, and direction. BEI\'s key signal is the direction.'),
    ('Behavioral zone', 'One independent dimension of an entity\'s activity (access patterns vs. data behavior), profiled separately so one drifting zone is not diluted by stable ones.'),
    ('Concept-direction naming', 'Matching the direction of an entity\'s drift against named threat/risk concepts to explain what the entity is becoming.'),
    ('UEBA', 'User and Entity Behavior Analytics — profiling how users and machines normally behave to spot deviations.'),
    ('APT', 'Advanced Persistent Threat — a long-duration, stealthy intrusion campaign, typically state-sponsored.'),
    ('C2 (command and control)', 'The covert channel malware uses to phone home; "beaconing" is its rhythmic check-in.'),
    ('LOTL (living off the land)', 'Attacks that abuse legitimate admin tools already on the system, leaving no malware to detect.'),
    ('Volt / Salt Typhoon', 'Publicly reported state-sponsored campaign styles emulated in the test bed: LOTL pre-positioning in critical infrastructure (Volt); telecom intrusion (Salt).'),
    ('Composite scoring', 'Fusing five independent detection phases (signal strength, breadth, persistence, context divergence, novelty) into one ranked score per entity.'),
    ('Peer cohort', 'A group of entities expected to behave alike (users in the same role; suppliers in the same commodity class); each entity is scored against its own cohort.'),
    ('Relationship signature', 'A distinct behavioral profile for a PAIR of entities (user×device, supplier×part), able to drift even when both individuals look stable.'),
    ('Network layer', 'Analysis of the full relationship graph: multi-hop risk chains, hidden shared dependencies, coordinated (lockstep) behavior.'),
    ('False-positive rate', 'The share of normal entities incorrectly flagged for review — 8.1% in the proof of concept.'),
    ('Z-score', 'A standard statistic counting how many standard deviations a value sits from its average; a common traditional alarm convention is three standard deviations.'),
    ('SCRM / SBOM', 'Supply Chain Risk Management / Software Bill of Materials — the component composition of an assembly, tracked here as an entity relationship.'),
    ('NSN', 'National Stock Number — the unique identifier of a part in the defense supply system.'),
    ('AML', 'Anti-Money Laundering — detecting financial transactions structured to disguise illicit funds.'),
    ('Synthetic data', 'Simulated data built to prove the concept before access to real operational data.'),
]
gt = doc.add_table(rows=len(gl) + 1, cols=2); gt.style = 'Table Grid'
setcell(gt.rows[0].cells[0], 'Term', bold=True, color=WHITE, fill='0B1F3A')
setcell(gt.rows[0].cells[1], 'Meaning', bold=True, color=WHITE, fill='0B1F3A')
for i, (t, d) in enumerate(gl, start=1):
    setcell(gt.rows[i].cells[0], t, bold=True, size=9, color=NAVY, fill='F4F6F8')
    setcell(gt.rows[i].cells[1], d, size=9)
for row in gt.rows:
    row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(4.9)

para('')
para('Business Edition — architecture-level disclosure. Exact model logic, formulas, scoring weights, and detection '
     'thresholds are proprietary and available in the full technical edition under NDA. All results are synthetic '
     'and must be re-validated on real operational data. Do not redistribute.', 9, GREY, italic=True)

out = OUT; n = 2
while True:
    try:
        doc.save(out); break
    except PermissionError:
        root, ext = os.path.splitext(OUT); out = f"{root}_v{n}{ext}"; n += 1
print('Saved:', out)
