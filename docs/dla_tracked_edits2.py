"""Second pass of tracked edits on the DLA REV file:
  D) Sharpen the embedding-generates-42-features point (append to the features paragraph).
  E) Reframe the cyber result as a transferable detection principle / hypothesis
     (append to the transfer-logic paragraph + a note under the Table 5 caption).
Also re-extracts the ACCEPTED text (insertions applied, deletions removed) for review.
"""
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REV = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_final_REV.docx"
AUTHOR = "ReviewPanel"; DATE = "2026-06-19T13:00:00Z"
d = Document(REV)
_id = [7000]
def nid():
    _id[0] += 1; return str(_id[0])

def _run(text, bold=False, italic=False):
    r = OxmlElement('w:r')
    if bold or italic:
        rPr = OxmlElement('w:rPr')
        if bold: rPr.append(OxmlElement('w:b'))
        if italic: rPr.append(OxmlElement('w:i'))
        r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t)
    return r

def _ins_wrap(run):
    ins = OxmlElement('w:ins'); ins.set(qn('w:id'), nid()); ins.set(qn('w:author'), AUTHOR); ins.set(qn('w:date'), DATE)
    ins.append(run); return ins

def append_ins(p, text, bold=False, italic=False):
    p._p.append(_ins_wrap(_run(text, bold, italic)))

def new_para(parts):
    p = OxmlElement('w:p'); pPr = OxmlElement('w:pPr')
    rPr = OxmlElement('w:rPr')
    ins = OxmlElement('w:ins'); ins.set(qn('w:id'), nid()); ins.set(qn('w:author'), AUTHOR); ins.set(qn('w:date'), DATE)
    rPr.append(ins); pPr.append(rPr); p.append(pPr)
    for (txt, b, it) in parts:
        p.append(_ins_wrap(_run(txt, b, it)))
    return p

# D) embedding -> 42 features
p70 = next(p for p in d.paragraphs if p.text.strip().startswith('Engineered features:'))
append_ins(p70, " Of the 67 features, 42 (the 23 zone-derived, 13 drift, and 6 cohort families) are generated "
    "from the embedding representation rather than from raw tabular fields; that is the embedding's specific "
    "contribution, and the tabular-only baseline (V1, 25 features) versus the embedding-enriched variant (V2, 29 "
    "features) isolates it.")

# E) cyber reframe: body sentence + note under Table 5 caption
p139 = next(p for p in d.paragraphs if 'meaningful evidence that the same mechanism' in p.text)
append_ins(p139, " This is offered as a transferable detection principle, that slow, persistent, sub-threshold "
    "signals are discoverable, and as a hypothesis to validate on DLA logistics data, not as a logistics performance "
    "result.")

t5 = next(p for p in d.paragraphs if p.text.strip().startswith('Table 5 - Cybersecurity'))
note = new_para([("Note: ", True, False),
    ("these are cybersecurity proof-of-concept figures that establish the detection principle; they are not a "
     "logistics performance claim. The 8.1% false-positive operating point was selected with knowledge of the "
     "ground truth and is a reporting benchmark, not a deployable threshold.", False, True)])
t5._p.addnext(note)

o = REV; k = 2
while True:
    try:
        d.save(o); break
    except PermissionError:
        root, ext = os.path.splitext(REV); o = f"{root}_v{k}{ext}"; k += 1
print("Saved:", o)

# ---- extract ACCEPTED text (w:t kept incl. insertions, delText skipped) ----
d2 = Document(o)
from docx.table import Table
from docx.text.paragraph import Paragraph
out = []
def para_text(p_el):
    return "".join(t.text or "" for t in p_el.iter(qn('w:t')))
for child in d2.element.body.iterchildren():
    if child.tag == qn('w:p'):
        t = para_text(child).strip()
        if t: out.append(t)
    elif child.tag == qn('w:tbl'):
        out.append("\n[TABLE]")
        for row in Table(child, d2).rows:
            out.append(" | ".join(c.text.strip().replace("\n", " ") for c in row.cells))
        out.append("[/TABLE]\n")
open("WP DLA/_DLA_paper_REV_accepted.txt", "w", encoding="utf-8").write("\n".join(out))
print("accepted-text chars:", sum(len(x) for x in out), "| lines:", len(out))
