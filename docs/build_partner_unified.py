# -*- coding: utf-8 -*-
"""Partner-augmented, UNIFIED-22CT-BRAND working-session documents.
All presenters represent 22nd Century; customer sees one integrated platform.
Two docs:
 1) Working_Session_RunOfShow_Unified22CT_PARTNER_INTERNAL.docx
 2) Partner_Briefing_Packet_Represent_22CT_INTERNAL.docx
Honesty guardrail baked in: unified brand front-of-room, honest teaming line on
direct provenance/IP/contract questions (acquisition + ARCYBER audience).
No em-dashes.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6); GRAY=RGBColor(0x55,0x5F,0x6B)
RED=RGBColor(0xA6,0x1B,0x1B); GREEN=RGBColor(0x1F,0x7A,0x33); AMBER=RGBColor(0xB5,0x5A,0x00)

def new_doc():
    d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(10.5); return d
def H1(d,t):
    p=d.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
    return p
def H2(d,t):
    p=d.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=BLUE
    return p
def P(d,t,italic=False,bold=False,color=None,size=None):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold
    if color: r.font.color.rgb=color
    if size: r.font.size=Pt(size)
    return p
def B(d,label,rest="",sub=False):
    p=d.add_paragraph(style='List Bullet 2' if sub else 'List Bullet')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def meta(d,label,val):
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(label+"  "); r.bold=True; r.font.color.rgb=GRAY; r.font.size=Pt(9.5)
    r2=p.add_run(val); r2.font.size=Pt(9.5); return p
def say(d,lines):
    lp=d.add_paragraph(); lp.paragraph_format.space_after=Pt(1)
    lr=lp.add_run("SAY:"); lr.bold=True; lr.font.color.rgb=GREEN; lr.font.size=Pt(9.5)
    for ln in lines:
        p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.25); p.paragraph_format.space_after=Pt(3)
        r=p.add_run("“"+ln+"”"); r.italic=True; r.font.size=Pt(10.5)
def callout(d,text,color=RED,size=9.5):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2)
    r=p.add_run(text); r.bold=True; r.font.color.rgb=color; r.font.size=Pt(size); return p
def table(d,headers,rows,widths=None,fs=8.5,hdr_fs=8.5):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(hdr_fs)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""
            for j,part in enumerate(v.split("||")):
                para=cells[i].paragraphs[0] if j==0 else cells[i].add_paragraph()
                para.add_run(part).font.size=Pt(fs)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    d.add_paragraph(); return t
def title_block(d,title,subtitle,banner):
    tp=d.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=tp.add_run(title); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=NAVY
    sp=d.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=sp.add_run(subtitle); r.font.size=Pt(11.5); r.font.color.rgb=BLUE
    wn=d.add_paragraph(); wn.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rw=wn.add_run(banner); rw.bold=True; rw.font.color.rgb=RED; rw.font.size=Pt(9.5)

# =====================================================================
# DOC 1: UNIFIED RUN-OF-SHOW
# =====================================================================
d=new_doc()
title_block(d,"Working Session Run-of-Show",
  "Unified 22nd Century Four-Layer Platform (Partner-Augmented)  |  Army Cyber: Esquibel - DiGrezio - Herrmann",
  "INTERNAL. All presenters represent 22nd Century. Customer sees ONE platform.")

H1(d,"1. The mandate: one platform, one brand")
P(d,"The customer sees a single integrated offering: the 22nd Century V-Intelligence four-layer platform, delivered "
  "and integrated by 22nd Century. Rigor AI and Cequence present as part of the 22nd Century team. 22CT is the prime "
  "and integrator and owns the open, every transition, and the close.")
P(d,"Internal only (never shown or stated to the customer):",bold=True)
table(d,["Layer","Capability","Driven by (internal)","Brand shown to customer"],
 [["1 - Preemptive Network Assurance","Formal-proof network assurance","Rigor AI","22nd Century / V-Intelligence Layer 1"],
  ["2 - Behavioral Entity Intelligence","The behavioral digital twin (22CT IP)","22nd Century","22nd Century / V-Intelligence Layer 2"],
  ["3 - Application, API & Data Runtime","API + data + agentic-AI protection","Cequence","22nd Century / V-Intelligence Layer 3"],
  ["4 - Code & Supply-Chain","Novel zero-day discovery","22nd Century (in development)","22nd Century / V-Intelligence Layer 4 (roadmap)"]],
 widths=[1.4,1.5,1.2,2.0],fs=8)

H1(d,"2. The unified-but-honest guardrail (read before the room)")
callout(d,"Unified brand front-of-room. Honest teaming line on direct questions. These are not in tension.")
B(d,"Present capability as ours: ","\"our platform's Layer 1 / Layer 3.\" That is fair: 22nd Century integrates and "
  "delivers it as one platform.")
B(d,"Do NOT claim 22CT engineered every layer. ","With ARCYBER (Esquibel) and an acquisition officer (DiGrezio) in "
  "the room, sole-authorship overclaim is both untrue and fails diligence.")
B(d,"If asked who builds it / is this your IP / what is the contract vehicle: ","the 22CT lead gives the teaming line, "
  "calmly and as a strength:")
say(d,["22nd Century delivers this as one integrated platform under a single accountable prime. The behavioral core, "
 "Layer 2, is our IP. Layers 1 and 3 are delivered through our technology partners inside the 22nd Century platform: "
 "one integrator, one prime, one contract vehicle. That is the point. You buy and field one platform, not four tools."])
meta(d,"Why this matters:","DiGrezio needs the teaming truth to field it; Esquibel rewards candor. Unified brand does "
 "not mean hiding the teaming structure when directly asked. It means the customer experiences one platform.")
meta(d,"Partner discipline:","On any provenance, IP, pricing, or competitive question, partners DEFER to the 22CT lead. "
 "They never freelance the teaming story.")

H1(d,"3. One vocabulary, one scenario, one format")
B(d,"One vocabulary: ","MITRE ATT&CK across all four layers.")
B(d,"One scenario: ","the shared Volt Typhoon / Salt Typhoon narrative. Every layer maps to the SAME campaigns.")
B(d,"One results format (all layers identical): ","threat scenario -> what incumbents miss -> live demo -> measured result.")

H1(d,"4. Agenda (half-day, ~4.5h with discipline)")
table(d,["#","Block","Brand-facing owner","Driven by (internal)","Time"],
 [["0","Open: their goals, set the arc, the platform thesis","22nd Century","22CT","20 min"],
  ["1","Layer 2 - Behavioral digital twin (theory, live demo, results)","22nd Century","22CT","60 min"],
  ["2","Layer 1 - Preemptive Network Assurance (theory, live demo, results)","22nd Century","Rigor AI","60 min"],
  ["3","Layer 3 - API, Data & Agentic-AI Runtime (theory, live demo, results)","22nd Century","Cequence","60 min"],
  ["4","Layer 4 - Code & Supply-Chain (roadmap, honest)","22nd Century","22CT","15 min"],
  ["5","Integration: one Typhoon across all four layers + federal fit + the ask","22nd Century","22CT","30 min"],
  ["-","Buffer for hard questions","All","-","20 min"]],
 widths=[0.35,3.0,1.2,1.0,0.6],fs=8)
meta(d,"Sequence note:","Lead with the digital twin (Layer 2). They asked for it, it is our true IP, and it sets the "
 "credibility bar the partner blocks inherit. 22CT bookends so it reads as one architecture.")

H1(d,"5. Block-by-block")
H2(d,"Block 0 - Open (22CT)")
say(d,["What you will see today is one platform, the 22nd Century V-Intelligence four-layer architecture, covering the "
 "full attack lifecycle under one MITRE vocabulary. Every layer runs live. We will be explicit about what is "
 "operational today versus on the roadmap. Let's start with the behavioral core."])
H2(d,"Block 1 - Layer 2, the digital twin (22CT, anchor)")
meta(d,"Content:","The 7-part anatomy + the live four-attacker demo + the 8.1% derivation. Real numbers (Section 7).")
meta(d,"Honesty line:","\"The behavioral core is our IP and runs live. Identity fusion, forensics, and response are our roadmap.\"")
H2(d,"Block 2 - Layer 1, Preemptive (22CT brand, Rigor drives)")
meta(d,"Content:","Formal-proof network assurance. Scenario: prove Volt's misconfigured-edge entry path is closed before traffic flows.")
meta(d,"Format:","theory 15, live demo 30, results 15. Same results frame as every other layer.")
meta(d,"Handoff in:","22CT lead introduces it as \"our Layer 1\" and hands to the presenter as a 22nd Century team member.")
H2(d,"Block 3 - Layer 3, API/Data/Agentic (22CT brand, Cequence drives)")
meta(d,"Content:","API, data, and agentic-AI runtime protection. Scenario: catch Salt's API/data collection and agent abuse at runtime.")
meta(d,"Format:","theory 15, live demo 30, results 15.")
meta(d,"Handoff in:","22CT lead introduces it as \"our Layer 3.\"")
H2(d,"Block 4 - Layer 4, Code/Supply-Chain (22CT, roadmap)")
meta(d,"Honesty line:","\"This layer is in development. I am presenting the approach, not a running product.\" Substantiate or drop the 2,000-zero-days figure.")
H2(d,"Block 5 - Integration + the ask (22CT)")
meta(d,"The moat:","Only the integrated platform shows the layers handing off. Trace one Typhoon intrusion: Layer 4 removes the "
 "exploitable code, Layer 1 proves the edge is closed, Layer 2 catches the valid-account behavior inside, Layer 3 catches the API/data exfil.")
P(d,"The ask: \"You saw one platform, four layers, every layer live, under one MITRE vocabulary, delivered by a single "
  "accountable prime. The honest next step is a bounded pilot on real Army telemetry, advisory and human-in-the-loop, "
  "measured side-by-side against current tooling.\"",italic=True)

H1(d,"6. Per-attendee lenses")
B(d,"Dr. Judy Esquibel (ARCYBER operator, your ally): ","detection fidelity, LOTL/valid-account, FP economics, baseline drift. Go deep on the twin; let the partners go deep on their layers.")
B(d,"LTC Micah DiGrezio (acquisition): ","one prime, one contract vehicle, one fieldable platform. Data/telemetry requirements, footprint (IL5/6, JWICS), integration, sustainment, acquisition path (OTA -> PoR). The teaming line is FOR him.")
B(d,"Kellsie Herrmann (policy): ","why it matters: Volt/Salt, valid-account critical-infra gap, authorities, no PII in the behavioral vector, readiness. Keep the math light.")

H1(d,"7. Numbers to keep straight (digital twin)")
B(d,"Dataset: ","250 entities, 485 days, ~14M events, 5 log sources.")
B(d,"Threat-profile detector (primary): ","4/4 at 0 FP by named technique (C2-beacon, DGA, LOTL-process, cohort-rare access, recon-fanout, insider-collection).")
B(d,"FP (composite): ","8.1% composite (~20/246 normals above Volt's 13.70), 4/4 caught but cleanly separating only 2/4 (USR-118, USR-156). Use 8.1, never 8.5.")
B(d,"Ranks: ","Salt #1 (51.27); Insider #2 (46.24); Slow APT #7 (19.44); Volt #24 (13.70).")
B(d,"Baselines: ","Isolation Forest / OC-SVM / LOF each 0/4; z-score 1/4.")
B(d,"Separation is by RANK: ","top normal (21.5) outscores Slow APT (19.4) and Volt (13.7). Present a ranked list.")

H1(d,"8. Logistics (in-person, all partners present)")
callout(d,"In-person, all three teams in the room. The #1 failure mode is network access at an Army facility. Plan for no internet.")
B(d,"Network is the top risk: ","assume the room has restricted or no internet. Every demo (Rigor L1, 22CT L2, Cequence "
  "L3) must run locally/offline, over the presenter's own cellular hotspot, or from a full backup recording. Do NOT "
  "assume cloud access to any partner demo environment.")
B(d,"Backup recordings are mandatory ","for all three layers, in case a live demo stalls or the network is locked down.")
B(d,"Bring-your-own-hardware: ","each presenter brings their own tested demo laptop plus HDMI and USB-C adapters.")
B(d,"One shared display, ","with a clean driver-switch protocol between the three teams. 22CT lead keeps the 60-minute boxes.")
B(d,"Arrive early: ","60 to 90 minutes for an AV check and a full demo dry-run in the actual room before the customer arrives.")
B(d,"Out of scope here: ","teaming, NDA, and contract admin are handled separately by the business team.")

P(d,"")
P(d,"Internal preparation document. Unified brand front-of-room; honest teaming line on direct questions. Not for distribution.",
  italic=True,size=9,color=GRAY)

out1="WP DLA/Presentation/Working_Session_RunOfShow_Unified22CT_PARTNER_INTERNAL.docx"
os.makedirs("WP DLA/Presentation",exist_ok=True)
d.save(out1)

# =====================================================================
# DOC 2: PARTNER BRIEFING PACKET
# =====================================================================
d=new_doc()
title_block(d,"Partner Briefing Packet",
  "Representing 22nd Century in the Army Working Session  |  For Rigor AI (Layer 1) and Cequence (Layer 3)",
  "Please read before the session. You present as part of the 22nd Century team.")

H1(d,"1. The mission")
P(d,"One session, one platform. You are presenting as part of the 22nd Century team. The customer experiences a single "
  "integrated four-layer platform (V-Intelligence), not three companies. 22nd Century is the prime and integrator and "
  "runs the open, the transitions, and the close. Your job is to make your layer land as a seamless part of that one platform.")

H1(d,"2. The audience (and why candor is non-negotiable)")
B(d,"Dr. Judy Esquibel ","- Senior Cyber Advisor, U.S. Army Cyber Command (ARCYBER); PhD; critical-infrastructure exercise lead. Deeply technical. She will probe detection fidelity and provenance. She rewards candor and detects overclaim instantly.")
B(d,"LTC Micah DiGrezio ","- Army acquisition (PM Cyber & Space). He judges whether this can be bought, integrated, and fielded. He will ask about data, footprint, sustainment, and contract structure.")
B(d,"Kellsie Herrmann ","- Army cyber/technology policy. She judges strategic significance and authorities. Keep the math light for her.")
callout(d,"One overclaim from any presenter burns the whole team's credibility with this room. When in doubt, under-claim and defer to the 22CT lead.")

H1(d,"3. Brand and vocabulary rules")
B(d,"Present your capability as 22nd Century: ","\"our platform's Layer 1\" (Rigor) or \"our platform's Layer 3\" (Cequence). Do not present your company brand front-of-room.")
B(d,"Use MITRE ATT&CK ","as the shared vocabulary, and map your layer to the shared Volt Typhoon / Salt Typhoon scenario.")
B(d,"No competitive cross-talk, ","no pricing, no naming other vendors, no internal company positioning.")
B(d,"On 'who builds this / is this your IP / what is the contract vehicle': ","stop and hand to the 22CT lead. Do not freelance the teaming story. The 22CT lead has the agreed teaming line.")

H1(d,"4. Your block: 60 minutes, fixed format")
P(d,"Every layer uses the identical structure so the platform feels like one product:",bold=True)
table(d,["Segment","Time","What to cover"],
 [["Theory","15 min","The threat at your layer, what incumbents do, and the core shortcoming. Map to MITRE."],
  ["Live demo","30 min","Run it live against the shared scenario. No slideware-only. Have a backup recording."],
  ["Results","15 min","The measured result: what you caught that incumbents miss, with real numbers."]],
 widths=[1.1,0.8,4.3],fs=9)
B(d,"Results frame (use this exact arc): ","threat scenario -> what incumbents miss -> live demo -> measured result.")

H1(d,"5. Your layer and your scenario")
table(d,["Partner","Layer (22CT brand)","Map your demo to this scenario"],
 [["Rigor AI","Layer 1 - Preemptive Network Assurance","Prove Volt Typhoon's misconfigured-edge / cross-VLAN entry path is closed before traffic flows. Formal proof over the full config space, not sampling."],
  ["Cequence","Layer 3 - API, Data & Agentic-AI Runtime","Catch Salt Typhoon's API and data collection and exfiltration, plus agent/bot abuse, at runtime. Authenticated-abuse and business-logic attacks that signatures miss."]],
 widths=[0.9,1.7,3.6],fs=8.5)

H1(d,"6. Federal-readiness you must each cover")
P(d,"DiGrezio judges fieldability across the whole stack, so every layer must speak to:")
B(d,"Data / telemetry: ","what your layer needs, ideally from telemetry already collected.")
B(d,"Footprint and accreditation: ","containerized; deployable across IL5, IL6, and JWICS; air-gap / offline option.")
B(d,"Integration: ","how your layer fits an existing Army cyber stack without rip-and-replace.")
B(d,"Sustainment: ","how updates and new threat coverage are delivered.")

H1(d,"7. Logistics and discipline (IN-PERSON)")
callout(d,"This is an in-person session at an Army facility. Assume NO reliable internet in the room. Your demo must work offline.")
B(d,"Bring-your-own-hardware: ","your own tested demo laptop plus HDMI and USB-C adapters.")
B(d,"Network: ","run your demo locally/offline or over your own cellular hotspot, and ALWAYS bring a full backup recording. Do not rely on the room network reaching your cloud.")
B(d,"On-site dry-run: ","arrive 60 to 90 minutes early and test your demo on the actual room display before the customer arrives.")
B(d,"Time box is hard: ","60 minutes. The 22CT lead keeps time. Practice to fit.")
B(d,"Q&A handoff: ","answer your-layer technical questions; hand integration, provenance, contract, and cross-layer questions to the 22CT lead.")
B(d,"Honesty boundary: ","label anything not in your current build as roadmap, out loud. Do not present roadmap as fielded.")

H1(d,"8. What success looks like")
P(d,"At the end, the room believes they saw one coherent platform, every layer live and measured, under one vocabulary, "
  "delivered by one accountable prime, with nothing left to doubt. That is the standard.")

P(d,"")
P(d,"Prepared by 22nd Century for partner coordination. Internal to the team.",italic=True,size=9,color=GRAY)

out2="WP DLA/Presentation/Partner_Briefing_Packet_Represent_22CT_INTERNAL.docx"
d.save(out2)

print("saved:")
print("  ",out1)
print("  ",out2)
