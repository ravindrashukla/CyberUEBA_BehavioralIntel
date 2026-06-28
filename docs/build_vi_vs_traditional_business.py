"""Business-friendly edition of the V-Intelligence vs Traditional UEBA paper.
Keeps real technical substance (three-tier story, blind-test results, per-attacker
tells, honest caveats, novelty) but redacts exact implementation logic (formulas,
thresholds, scoring weights, attention tables, code). Abbreviations clarified on
first use. ~7-8 pages.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "VIntelligence_vs_Traditional_UEBA_Academic.docx")

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
ORG  = RGBColor(0xC2, 0x5A, 0x12)
GREY = RGBColor(0x55, 0x5F, 0x6B)
GREEN= RGBColor(0x1E, 0x7A, 0x44)
WHITE= RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.1
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(doc.sections[0], m, Inches(0.85))


def H1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    return p

def H2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ORG
    return p

def para(text, size=10.5, color=None, bold=False, italic=False, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def bullet(text, lead=None):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run(text); r2.font.size = Pt(10.5)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def setcell(cell, text, bold=False, size=9.5, color=None, fill=None):
    if fill: shade(cell, fill)
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
    if color: r.font.color.rgb = color

def table(headers, rows, widths=None, hdr_fill='0B1F3A'):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        setcell(t.rows[0].cells[j], h, bold=True, color=WHITE, fill=hdr_fill)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            setcell(t.rows[i].cells[j], v, bold=(j == 0), fill='F4F6F8' if j == 0 else None)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ================= FIGURE GENERATION =================
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

FIGDIR = os.path.join(os.path.dirname(__file__), 'figs_business_vi')
os.makedirs(FIGDIR, exist_ok=True)

C_TRAD = '#8A93A0'   # traditional grey
C_TRADR = '#C0392B'  # traditional red accent
C_VI = '#0891B2'     # V-Intelligence teal
C_NAVY = '#0B1F3A'
plt.rcParams.update({'font.family': 'Calibri', 'font.size': 10, 'axes.edgecolor': '#C9CFD6',
                     'axes.labelcolor': C_NAVY, 'xtick.color': '#55607B', 'ytick.color': '#55607B'})

def fig1_scoreboard():
    """Blind-test scoreboard: campaigns detected out of 4, with FP rates."""
    methods = ['Isolation\nForest', 'LOF', 'One-Class\nSVM', 'Z-Score', 'Feature\nCUSUM',
               'V-Intelligence\nComposite']
    caught = [0, 0, 0, 1, 0, 4]
    fp = ['5.3% FP', '4.5% FP', '14.6% FP', '9.8% FP', '10.2% FP', '8.1% FP']
    colors = [C_TRAD, C_TRAD, C_TRAD, C_TRADR, C_TRAD, C_VI]
    x = np.arange(len(methods))
    fig, ax = plt.subplots(figsize=(7.0, 3.2))
    bars = ax.bar(x, caught, 0.62, color=colors)
    for xi, c, f in zip(x, caught, fp):
        ax.text(xi, c + 0.14, f, ha='center', va='bottom', fontsize=8.4, color='#55607B')
        ax.text(xi, c + 0.56, f'{c} of 4', ha='center', va='bottom', fontsize=10,
                fontweight='bold', color=C_NAVY)
    ax.axhline(4, color=C_VI, lw=1.0, ls=':', alpha=0.6)
    ax.text(0.05, 4.10, 'All four stealth campaigns', fontsize=8.6, color=C_VI)
    ax.set_xticks(x, methods, fontsize=8.8)
    ax.set_ylim(0, 5.3); ax.set_yticks([0, 1, 2, 3, 4])
    ax.set_ylabel('Stealth campaigns detected (of 4)')
    ax.spines[['top', 'right']].set_visible(False)
    ax.tick_params(axis='x', length=0)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig1_scoreboard.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig2_spaces():
    """Feature space vs semantic space — why the insider is invisible to one and
    visible to the other. Illustrative."""
    rng = np.random.default_rng(11)
    fig, axes = plt.subplots(1, 2, figsize=(7.0, 3.1))
    # (a) feature space: attacker buried inside the normal cloud
    ax = axes[0]
    nx, ny = rng.normal(0, 1, 160), rng.normal(0, 1, 160)
    ax.scatter(nx, ny, s=14, color=C_TRAD, alpha=0.45, label='Normal users')
    ax.scatter([0.4], [-0.2], s=80, color=C_TRADR, zorder=5, marker='D', label='Insider (same file count)')
    ax.annotate('indistinguishable\nfrom the crowd', xy=(0.4, -0.2), xytext=(1.6, -2.6),
                fontsize=8.6, color=C_TRADR, ha='center',
                arrowprops=dict(arrowstyle='->', color=C_TRADR, lw=1.1))
    ax.set_title('Numeric feature space\n(what traditional tools measure)', fontsize=9.5, color=C_NAVY)
    ax.set_xlim(-3.6, 3.6); ax.set_ylim(-3.6, 3.6)
    ax.legend(loc='upper left', fontsize=7.8, frameon=False)
    # (b) semantic space: trajectory drifts toward the threat region
    ax = axes[1]
    nx2, ny2 = rng.normal(-1.2, 0.8, 160), rng.normal(0, 0.8, 160)
    ax.scatter(nx2, ny2, s=14, color=C_TRAD, alpha=0.45, label='Normal behavior region')
    from matplotlib.patches import Ellipse
    e = Ellipse((2.4, -2.0), 2.5, 1.8, angle=-15, color=C_TRADR, alpha=0.13)
    ax.add_patch(e)
    ax.text(2.55, -2.35, 'data-exfiltration\nconcept region', ha='center', va='center',
            fontsize=8.2, color=C_TRADR)
    traj_x = np.array([-1.3, -0.9, -0.4, 0.3, 1.0, 1.7])
    traj_y = np.array([0.1, -0.1, -0.5, -0.9, -1.3, -1.6])
    ax.plot(traj_x, traj_y, color=C_VI, lw=2.0, marker='o', ms=4.5, label='Insider, week by week')
    ax.annotate('', xy=(2.0, -1.72), xytext=(1.7, -1.6),
                arrowprops=dict(arrowstyle='-|>', color=C_VI, lw=2.0))
    ax.set_title('Semantic space — a map of meaning\n(what V-Intelligence measures)', fontsize=9.5, color=C_NAVY)
    ax.set_xlim(-3.6, 3.9); ax.set_ylim(-3.6, 2.6)
    ax.legend(loc='upper left', fontsize=7.8, frameon=False)
    for ax in axes:
        ax.set_xticks([]); ax.set_yticks([])
        ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig2_spaces.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig3_ranks():
    """Per-attacker composite score vs the normal population (real run output)."""
    rng = np.random.default_rng(5)
    # Normal population backdrop: 246 normal users, bulk low, upper tail toward ~14
    normals = np.concatenate([rng.gamma(2.2, 1.8, 220), rng.gamma(3.0, 3.0, 26)])
    normals = np.clip(normals, 0.2, 17.5)
    attackers = [('USR-118  Salt Typhoon', 51.3, '#1 of 250'),
                 ('USR-156  Insider Threat', 46.2, '#2 of 250'),
                 ('USR-234  Slow APT', 19.4, '#7 of 250'),
                 ('USR-042  Volt Typhoon LoTL', 13.7, '#24 of 250')]
    fig, ax = plt.subplots(figsize=(7.0, 3.0))
    ax.scatter(normals, rng.uniform(-0.32, 0.32, len(normals)), s=12, color=C_TRAD, alpha=0.4,
               label='246 normal users (synthetic backdrop)')
    for i, (name, score, rank) in enumerate(attackers):
        ax.scatter([score], [0], s=110, color=C_VI, marker='D', zorder=5,
                   edgecolor=C_NAVY, linewidth=0.8)
        ytxt = 0.62 if i % 2 == 0 else -0.72
        ax.annotate(f'{name}\nscore {score} — rank {rank}', xy=(score, 0), xytext=(score, ytxt),
                    ha='center', fontsize=8.4, color=C_NAVY, fontweight='bold',
                    arrowprops=dict(arrowstyle='-', color=C_NAVY, lw=0.8))
    ax.axvline(13.7, color=C_TRADR, lw=1.2, ls='--', alpha=0.8)
    ax.text(13.4, -1.18, 'operating point: threshold that catches all 4\nadmits 20 normal users = 8.1% FP',
            fontsize=8.2, color=C_TRADR, ha='right')
    ax.set_xlim(-1, 56); ax.set_ylim(-1.45, 1.25)
    ax.set_yticks([])
    ax.set_xlabel('Multi-phase composite score (higher = riskier)')
    ax.legend(loc='upper right', fontsize=8.2, frameon=False)
    ax.spines[['top', 'right', 'left']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig3_ranks.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

def fig4_fusion():
    """Ranking quality (ROC AUC) per phase vs fused composite."""
    phases = ['Novelty\npersistence', 'Signal\nstrength', 'Breadth', 'Sustained', 'Context',
              'Fused\ncomposite']
    auc = [0.734, 0.840, 0.858, 0.893, 0.935, 0.976]
    colors = [C_TRAD] * 5 + [C_VI]
    x = np.arange(len(phases))
    fig, ax = plt.subplots(figsize=(6.8, 2.9))
    ax.bar(x, auc, 0.6, color=colors)
    for xi, a in zip(x, auc):
        ax.text(xi, a + 0.012, f'{a:.2f}', ha='center', fontsize=9, fontweight='bold', color=C_NAVY)
    ax.axhline(0.5, color=C_TRADR, lw=1.1, ls='--')
    ax.text(-0.38, 0.515, 'coin flip (0.50)', fontsize=8.2, color=C_TRADR)
    ax.set_xticks(x, phases, fontsize=8.8)
    ax.set_ylim(0.4, 1.04)
    ax.set_ylabel('Ranking quality (ROC AUC)')
    ax.spines[['top', 'right']].set_visible(False)
    ax.tick_params(axis='x', length=0)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig4_fusion.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p

FIG1, FIG2, FIG3, FIG4 = fig1_scoreboard(), fig2_spaces(), fig3_ranks(), fig4_fusion()

def add_fig(path, caption, width=6.4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY


# ================= TITLE =================
r = doc.add_paragraph().add_run('Insider Threats and Nation-State Intrusions:\nTraditional Anomaly Detection vs. V-Intelligence UEBA')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
para('A business and technology overview of the approach, its measured results, and the evidence behind them',
     12.5, GREY, after=2)
para('22nd Century Technologies, Inc.  ·  V-Intelligence UEBA Program  ·  June 2026  ·  Business Edition',
     9.5, GREY, after=2)
para('Disclosure note: this edition describes the architecture and the innovation, and reports measured results. The '
     'exact detection logic — formulas, thresholds, and scoring weights — is proprietary and is withheld here; it is '
     'available in the full technical edition under NDA (non-disclosure agreement).', 9.5, ORG, italic=True, after=4)
para('Data note: all results in this paper come from a synthetic (simulated) enterprise built to prove the concept — '
     '250 users, 485 days, roughly 145 million log events across seven sources, with four hidden long-duration attack '
     'campaigns. They must be re-validated on real production telemetry before any operational claim.', 9.5, ORG,
     italic=True, after=8)

# ================= 1. EXEC SUMMARY =================
H1('1.  Executive Summary')
para('The most damaging attacks rarely resemble attacks. An insider with valid credentials opens the usual '
     'number of files, but gradually shifts from public documents to restricted ones. A nation-state '
     'implant contacts its operator two or three times a day for more than a year, never often enough to trip any '
     'threshold. The Volt Typhoon actors, who according to the U.S. Cybersecurity and Infrastructure Security '
     'Agency (CISA) dwelled inside U.S. critical infrastructure undetected for at least five years, used only the '
     'operating system\'s own signed tools, so nothing ever resembled malware. Such adversaries remain inside '
     'authorized activity and inside normal statistical ranges, which is precisely where traditional security '
     'analytics cannot see.')
para('UEBA (User and Entity Behavior Analytics) is the discipline of finding such threats by modeling how each user '
     'and system normally behaves. Traditional UEBA reduces behavior to numbers, such as counts, ratios, and byte '
     'volumes, and flags statistical outliers. V-Intelligence UEBA takes a different approach. It maintains a living '
     'digital twin of every entity, comparable to a patient\'s medical chart, which records not only today\'s '
     'temperature but the accumulated history and context a doctor needs to read the trend. The twin captures the '
     'meaning of the behavior rather than the size of the numbers alone, and the system tracks the direction in '
     'which that meaning is moving. It expresses that direction in the vocabulary of MITRE ATT&CK, the '
     'industry-standard catalog of adversary techniques.')
para('The supporting evidence is a blind test: one simulated enterprise of 250 users with four stealth campaigns '
     'hidden inside, namely a 14-month insider, a 417-day slow APT (Advanced Persistent Threat), a Volt '
     'Typhoon-style "living off the land" (LoTL) campaign, and a Salt Typhoon-style telecom infrastructure pivot. '
     'Three generations of detection ran on the same data:', after=4)
bullet(' four of the five traditional detectors caught 0 of 4 campaigns; the best (Z-Score) caught 1 of 4 — '
       'and even then said only "anomalous," with no indication of why.', lead='Traditional (Tier 1): ')
bullet(' measuring behavioral drift in semantic space begins to surface the insider and the slow APT that '
       'magnitude-based methods cannot see.', lead='Intermediate (Tier 2): ')
bullet(' the digital-twin composite ranked all 4 of 4 campaigns in the top tier of 250 entities, at a '
       'false-positive (FP) rate of 8.1% — 20 of 246 normal users above the operating threshold — with each alert '
       'explaining which behavioral zone changed, in what direction, and toward which named threat pattern.',
       lead='V-Intelligence (Tier 3): ')
para('The Salt Typhoon-style campaign, representing the class of intrusion that compromised U.S. telecommunications '
     'carriers in the real world, ranked #1 of all 250 entities, even though its largest measured deviation on any '
     'traditional feature was 1.71 standard deviations, deep inside the normal statistical band and invisible to '
     'every traditional detector. The Volt Typhoon-style campaign, among the hardest threat classes to detect, '
     'surfaced in the top ~10%. The only traditional catch on the entire benchmark was Z-Score\'s bare "anomalous" '
     'flag on this same entity, which was correct but unexplained and indistinguishable from a busy week. The '
     'capability this approach adds is the ability to surface a five-year-invisible threat together with an '
     'explanation of why it was flagged.', bold=True)

# ================= 2. PROBLEM =================
H1('2.  The Problem: Attacks That Hide Inside Normal')
para('Security teams already face more alerts than they can review. The question UEBA must answer is not whether a '
     'metric spiked, but which entities are quietly becoming dangerous. The four campaigns in the blind test were '
     'engineered to represent the hardest real-world classes:')
bullet(' valid credentials and the same volume of files, with a gradual shift toward restricted and confidential '
       'content and off-hours activity, escalating over 14 months.', lead='The insider: ')
bullet(' command-and-control (C2) beaconing, in which the implant checks in with its operator just two or three '
       'times a day to new infrastructure, for 417 days. No single day exceeds any threshold.', lead='The slow APT: ')
bullet(' the attacker uses only native, signed administrative tools (the Volt Typhoon pattern), so no malware '
       'signature ever exists; 412 days.', lead='Living off the land (LoTL): ')
bullet(' quiet movement across internal network segments (the Salt Typhoon pattern observed in telecom carriers); '
       '412 days.', lead='The infrastructure pivot: ')
para('Each of these operates within authorized activity and within normal statistical ranges. The defensive '
     'question is therefore what the entity is becoming and whether that direction is dangerous. The central claim '
     'of this paper is that no amount of tuning makes a magnitude detector answer a direction question; doing so '
     'requires a different representation of behavior.')

# ================= 3. TRADITIONAL =================
H1('3.  The Traditional Approach (Tier 1)')
para('The traditional pipeline is mature and well understood, and it serves as the baseline against which '
     'V-Intelligence is measured. Raw logs from seven sources (logins, file access, network traffic, DNS or Domain '
     'Name System lookups, endpoint telemetry, application activity, and privileged operations) are condensed into '
     '31 numeric features per user, consisting of counts, ratios, and volumes. Five families of statistical and '
     'machine-learning outlier detectors then flag '
     'users whose numbers sit far from the crowd, including Isolation Forest, One-Class SVM (Support Vector '
     'Machine), Local Outlier Factor (LOF), Z-Score (how many standard deviations from baseline), and a '
     'feature-level CUSUM (cumulative-sum change detector, which accumulates small week-over-week changes to catch '
     'slow trends).')
H2('Why it is not enough')
bullet(' a shift from public to restricted files at a constant file count produces almost no numeric movement. The '
       'insider changed what they accessed, not how much, which leaves no signal for a magnitude detector.',
       lead='No sense of meaning: ')
bullet(' even when a detector fires, the alert says only "anomalous." It cannot say which aspect of behavior '
       'changed, in which direction, or toward what threat. A genuine insider and a developer who had a busy '
       'week therefore generate identical alerts, and each demands the same manual triage.', lead='No direction: ')
bullet(' each feature is an independent axis; "off-hours logins + restricted files + high clearance" is never read '
       'as one combined pattern.', lead='No combination: ')
bullet(' high-variance features such as bytes transferred numerically swamp the subtle ratios that actually carry '
       'the insider signal.', lead='Scale dominance: ')
para('The blind-test results follow directly: Isolation Forest 0 of 4 (5.3% FP), LOF 0 of 4 (4.5% FP), '
     'One-Class SVM 0 of 4 (14.6% FP), Z-Score 1 of 4 (9.8% FP), and feature-CUSUM 0 of 4 (10.2% FP). The campaigns '
     'were engineered to stay within per-feature norms, and they did.')

# ================= 4. TIER 2 =================
H1('4.  The Intermediate Step (Tier 2): Drift With Direction')
para('The first advance is in how behavior is represented. Instead of feeding raw numbers to a detector, each '
     'user\'s weekly behavior is written out as a short analyst-style narrative. Rather than "restricted_ratio = '
     '0.15," the description reads "elevated data-access anomaly: restricted-file ratio at a high percentile, '
     'roughly 12 times this employee\'s own baseline, trend rising." That narrative is converted into an embedding, '
     'a long numeric vector that captures the meaning of the text. Embeddings are the core representational '
     'machinery inside large language models, repurposed here to interpret behavior rather than language. Behaviors '
     'that mean similar things land near each other on this "map of meaning," even when their raw numbers differ.')
para('On that map, two measurements become possible that no threshold can produce: how far the entity moved this '
     'week (drift magnitude), and in which direction it moved (drift direction), measured against a library of '
     'pre-positioned threat concepts such as data exfiltration, C2 beaconing, and lateral movement, each mapped to '
     'MITRE ATT&CK techniques. A cumulative drift detector accumulates the small weekly steps that a per-week test '
     'misses.')
para('This step was still not sufficient, because Tier 2 collapses the whole entity into one blended profile. A '
     'real intrusion typically changes one aspect of behavior, such as the network footprint, while everything '
     'else stays normal. Averaged into a single vector, that one-zone signal is diluted by the four stable zones. '
     'This dilution problem is what motivates the full digital twin.')

# ================= 5. TIER 3 =================
H1('5.  The V-Intelligence Approach (Tier 3): the Living Digital Twin')
para('V-Intelligence decomposes every entity into five behavioral zones, each described and embedded independently '
     'so a signal in one zone is never averaged away by the stable ones: identity (role, department, and clearance, '
     'which are mostly static, so a change is itself a signal), access pattern (logins), data behavior (files), '
     'network '
     'footprint (traffic and DNS), and risk posture (endpoint telemetry). The zone descriptions deliberately '
     'retain qualitative detail that numeric features discard, including actual directory names, external '
     'addresses, and domain names, which is part of why direction becomes measurable.')
para('The twin is re-profiled weekly, so over months it traces a trajectory through the map of meaning: how fast it '
     'is moving, whether it is accelerating, whether one zone is drifting while the others hold still ("identity '
     'stable but network footprint drifting" is a classic intrusion signature). Relationships between entities are '
     'profiled too, so a change in who talks to what is visible.')
add_fig(FIG2, 'Figure 1 — Why representation matters (illustrative). Left: in numeric feature space the insider is '
              'indistinguishable from the crowd — the file count never changed. Right: on the semantic "map of '
              'meaning," the same insider drifts week by week toward the data-exfiltration concept region. The '
              'insider changed what they accessed, not how much: invisible to magnitude, visible to meaning.')
H2('Five scoring phases combined into one verdict')
para('A single drifting signal is not proof of compromise. The composite scorer interrogates each twin with five '
     'independent questions, then fuses the answers into one ranked score. (What each phase asks is described below; '
     'the exact weights and thresholds are proprietary and withheld in this edition.)')
table(['Scoring phase', 'The question it asks'],
      [('Signal strength', 'Is this entity\'s deviation stronger than its true peer group\'s? (Administrators are compared to administrators, analysts to analysts — not to the whole company.)'),
       ('Breadth', 'How many aspects of behavior are elevated at once? One busy metric is noise; many together are a pattern.'),
       ('Sustained', 'Does the deviation persist week after week, or was it one bad day?'),
       ('Context divergence', 'Does the entity look anomalous from multiple investigative viewpoints (insider lens, APT-hunt lens, privilege-audit lens) at the same time?'),
       ('Novelty persistence', 'Is the entity in recurring contact with infrastructure it has never touched before? (The slow-APT tell: tiny volumes, but the same never-before-seen destinations week after week.)')],
      widths=[1.7, 5.2])
para('A normal user may score high on one question by chance; an attacker scores high on several. Fusing the five '
     'into a single composite returns the entire population as one risk-ranked list, so analysts work from the top '
     'down rather than triaging a flood of disconnected alerts. Every flagged entity also carries its explanation: '
     'which zone changed, in which direction, and aligned to which named threat concept and MITRE ATT&CK '
     'techniques. The result is an auditable record of why the entity was flagged.', bold=False)

# ================= 6. RESULTS =================
H1('6.  The Blind Test: Three Generations, One Dataset (Synthetic)')
para('All three tiers ran on the same synthetic enterprise: 250 users, 485 days, roughly 145 million events across '
     'seven log sources, with the four campaigns of Section 2 embedded. Embeddings were real (OpenAI '
     'text-embedding-3-small), not simulated. The false-positive rate is reported at a defined operating '
     'point: the score threshold that recalls all four campaigns. That threshold was set with knowledge of the '
     'ground truth and is used as a transparent reporting device rather than a deployable setting. At that '
     'threshold, 20 of the 246 normal users rank above the lowest-ranked campaign, an 8.1% false-positive rate.')
add_fig(FIG1, 'Figure 2 — The scoreboard. Five traditional detectors against the four stealth campaigns: four catch '
              'none, Z-Score catches one (at a comparable 9.8% false-positive rate, and with no explanation). The '
              'V-Intelligence multi-phase composite ranks all four in the top tier at 8.1% false positives — with '
              'zone, direction, and MITRE ATT&CK context attached to every alert. (Synthetic data.)')
H2('6.1  Per-attacker results: each caught by a different tell')
para('The four campaigns are not caught by a single mechanism applied four times. Each one relies on a different '
     'mix of scoring phases, which is why fusion rather than any single detector carries the result:')
table(['Entity', 'Campaign', 'Rank / score', 'How it was caught'],
      [('USR-118', 'Salt Typhoon — telecom infrastructure pivot (412 days)', '#1 of 250\nscore 51.3',
        'Composite drift + breadth: quiet movement across segments lights up several zones at once. Its largest measured z-score on any traditional feature was 1.71 — deep inside the normal band — yet it ranked #1.'),
       ('USR-156', 'Insider threat — 14-month escalation', '#2 of 250\nscore 46.2',
        'Signal strength, breadth, and sustained deviation concentrated in the data-behavior zone while identity stayed flat.'),
       ('USR-234', 'Slow APT — 417-day C2 beaconing', '#7 of 250\nscore 19.4',
        'Novelty persistence: among the lowest raw signal of any entity (2–3 beacons/day), but the same never-before-seen infrastructure recurring week after week — the axis no threshold measures.'),
       ('USR-042', 'Volt Typhoon — living off the land (412 days)', '#24 of 250\nscore 13.7',
        'Composite drift + breadth: the hardest case, sitting in the upper normal tail — surfaced in the top ~10%. The one traditional catch on the benchmark was Z-Score\'s flag on this entity (measured z = 3.04, barely over the common z = 3 alarm line) — a bare "anomalous" with no indication of why.')],
      widths=[0.9, 2.2, 1.2, 2.6])
add_fig(FIG3, 'Figure 3 — Where the four campaigns land on the composite ranking (real run output; normal-user '
              'backdrop illustrative). Salt Typhoon and the insider stand far clear of the population; the slow APT '
              'and the Volt Typhoon LoTL campaign sit above the bulk but nearer the tail — recalling all four admits '
              '20 normal users, the 8.1% operating point.')
H2('6.2  The role of fusion')
para('Taken individually, the system\'s underlying detectors are each noisy; a naive union of their individual '
     'alarms would flag roughly 19% of normal users. Treated as a ranking problem, each scoring phase is only a '
     'partial discriminator: ranking quality (ROC AUC, short for the area under the receiver operating characteristic curve, i.e. the probability a random campaign outranks a random normal '
     'user, where 0.5 is a coin flip and 1.0 is perfect) runs from 0.73 for novelty alone to 0.94 for context '
     'alone. The fused composite reaches 0.976. Combining many weak, independent signals yields one reliable '
     'ranking, and an attacker cannot evade it by defeating a single phase, because staying hidden requires '
     'evading all five at once. This is the property that makes the design resistant to evasion.')
add_fig(FIG4, 'Figure 4 — Ranking quality by scoring phase versus the fused composite (250 users, real embeddings). '
              'Every individual phase is only a partial discriminator; fusion lifts ranking quality to 0.98. For '
              'context, the best traditional discriminator on a separate 50-user run reached about 0.81.')
H2('6.3  What 8.1% false positives actually means')
para('It does not mean the system blocks 8.1% of employees. The composite is a ranked triage list. At the '
     'threshold that catches all four campaigns, an analyst reviewing the list top-down would work through about 24 '
     'entities (4 real campaigns plus 20 normal users) out of 250, each with its zone, direction, and '
     'threat-concept explanation attached rather than a bare "anomalous" flag. A stricter threshold trades the '
     'hardest case (the Volt Typhoon entity, closest to the boundary) for fewer reviews; that '
     'detection-versus-false-alarm curve will be published on larger data.')

# ================= 7. DEPLOYMENT =================
H1('7.  Scale, Deployment, and Other Uses')
bullet(' cost scales with entities and weeks, not raw event volume. The 145 million events reduce to weekly '
       'behavioral profiles before embedding, and embeddings are cached so repeated behavior incurs no new cost.',
       lead='Economics. ')
bullet(' containerized (Docker) over PostgreSQL with a vector store and full bi-temporal history; deployable in an '
       'enterprise data center, at a tactical edge, or in an air-gapped enclave (the external embedding model swaps '
       'for a locally hosted one).', lead='Footprint. ')
bullet(' ingests from existing SIEM (Security Information and Event Management) and log sources, requiring no new '
       'instrumentation. Rollout runs in shadow mode alongside the existing SIEM before operation.',
       lead='Integration. ')
bullet(' the method is domain-agnostic: the same machinery has been instantiated for defense supply-chain '
       'supplier risk, and applies naturally to healthcare fraud, financial crime, equipment-degradation '
       'monitoring, and counter-intelligence, applying anywhere an entity\'s behavior can be observed over time.',
       lead='Beyond cyber. ')

# ================= 8. NOVELTY =================
H1('8.  Why This Is New')
para('Behavioral analytics is a mature field, and we are precise about what is and is not new. Commercial UEBA '
     'platforms, according to their publicly described methodology, baseline entities statistically and score '
     'deviations on numeric features and peer groups. They are effective on volume anomalies but cannot detect a '
     'change of kind that holds volume constant. Their recent generative-AI features are conversational copilots '
     'layered on top, not the detection mechanism. '
     'Academic work has separately established each ingredient used here: embedding serialized behavior, '
     'detecting anomalies in embedding space, and measuring embedding drift. Every published system we found, '
     'however, still detects by magnitude, that is, by how unusual or how far from baseline a behavior is.')
para('What no commercial platform and no single published system does, to the best of our knowledge, is detect by '
     'the direction of behavioral drift in a unified semantic space: naming what the entity is becoming against '
     'pre-positioned threat concepts mapped to MITRE ATT&CK, localized to a behavioral zone, scored against true '
     'peers, and fused across independent phases into one explainable ranked verdict. The innovation is the '
     'integrated architecture and the emphasis on direction rather than magnitude, not any single technique.',
     bold=True)
para('Scoping notes: vendor internals are not fully public (the comparison is against marketed methodology), and a '
     'formal freedom-to-operate review is recommended before any contractual assertion of uniqueness.', 9.5, GREY,
     italic=True)

# ================= 9. LIMITATIONS =================
H1('9.  Limitations, Stated Plainly')
bullet('All results are from synthetic telemetry; real-data validation with recalibrated thresholds is the planned '
       'next step.')
bullet('Four embedded campaigns demonstrate the concept end to end but are not a statistically powered evaluation; '
       'a larger, more varied campaign set is next.')
bullet('The 8.1% figure is measured at the threshold that catches all four campaigns; the full '
       'detection-versus-false-alarm curve on larger data will let each operator set their own alert budget.')
bullet('The behavioral narratives include interpretive wording ("restricted," "rising") that helps the model but '
       'may also prime it; a plain-numbers control run is planned to confirm the signal holds without it.')
bullet('A model-aware adversary will adapt, and detection is an ongoing contest. Multi-phase fusion, novelty '
       'persistence, and relationship signals raise the bar today and are designed to evolve.')
bullet('Air-gapped and classified environments are supported via a self-hosted embedding model; that local model '
       'should be re-characterized before relying on these exact numbers.')

# ================= 10. CONCLUSION =================
H1('10.  Conclusion')
para('Three generations of detection on one dataset make the case clearly. Traditional detection largely misses '
     'the stealthy campaigns, catching 0 of 4 for most methods and 1 of 4 for the best, and even when it fires the '
     'alert says only "anomalous," leaving the analyst the same manual triage. An intermediate drift approach '
     'begins to recover direction. The V-Intelligence digital twin, combining zones, trajectories, peer cohorts, '
     'relationships, and multi-phase fusion, finds all four campaigns at an 8.1% false-positive rate, including a '
     'Salt Typhoon-class intrusion ranked #1 of 250 and a Volt Typhoon-class campaign of the kind that stayed '
     'invisible in U.S. critical infrastructure for five years.')
para('Traditional analytics can detect anomalous users, but detection without direction leaves the analyst with '
     'triage and no interpretation. V-Intelligence identifies which behavioral zone changed and toward what threat '
     'pattern, mapped to MITRE ATT&CK, which is what turns an alert into an actionable investigation.', bold=True)

# ================= GLOSSARY =================
H1('Appendix — Glossary')
gl = [
    ('UEBA', 'User and Entity Behavior Analytics — finding threats by modeling how each user and system normally behaves, rather than matching known attack signatures.'),
    ('Digital twin', 'A living software profile of a real-world entity — like a patient\'s medical chart: history, vitals, context, and trend, updated weekly.'),
    ('Embedding / semantic space', 'A numerical vector capturing the meaning of a description; the resulting "map of meaning" places similar behaviors near each other even when raw numbers differ. The core representation inside large language models.'),
    ('Drift', 'How an entity\'s behavioral profile moves through the semantic space over time — how far, how fast, and in which direction.'),
    ('Behavioral zone', 'One of five independently profiled aspects of an entity: identity, access pattern, data behavior, network footprint, risk posture.'),
    ('MITRE ATT&CK', 'The industry-standard public catalog of adversary tactics and techniques; alerts here are named in its vocabulary.'),
    ('APT', 'Advanced Persistent Threat — a patient, well-resourced (typically nation-state) intruder operating quietly over months or years.'),
    ('LoTL', 'Living off the Land — attacking with only the system\'s own legitimate, signed tools, so no malware signature ever exists (the Volt Typhoon pattern).'),
    ('C2 beaconing', 'Command-and-control check-ins: an implant periodically contacting its operator\'s infrastructure, often only a few times a day.'),
    ('Volt / Salt Typhoon', 'Real nation-state campaigns against U.S. critical infrastructure (CISA advisory AA24-038A) and telecommunications carriers; the blind test embeds campaigns of both classes.'),
    ('SIEM / SOC', 'Security Information and Event Management (the log-collection platform) / Security Operations Center (the team that triages its alerts).'),
    ('CUSUM', 'Cumulative-sum change detection — accumulating small week-over-week changes so a slow trend becomes visible even when each step is tiny.'),
    ('Z-Score', 'How many standard deviations a value sits from its own baseline average; z = 3 is a common alarm line. The Salt Typhoon entity never exceeded z = 1.71 — statistically "normal" throughout.'),
    ('False positive (FP)', 'A normal user flagged for review. Here 8.1% means 20 of 246 normal users ranked above the threshold that catches all four campaigns — a triage list, not auto-blocking.'),
    ('Operating point', 'The score threshold at which results are reported; here, the threshold that recalls 100% of the embedded campaigns.'),
    ('ROC AUC', 'Ranking quality: the probability a randomly chosen attacker outranks a randomly chosen normal user. 0.5 is a coin flip; 1.0 is perfect.'),
    ('Peer cohort', 'The entity\'s true comparison group (administrators vs. administrators, analysts vs. analysts) rather than the whole population.'),
    ('Novelty persistence', 'Recurring contact with never-before-seen infrastructure — the slow-APT tell that volume-based tools cannot measure.'),
    ('Synthetic data', 'Simulated enterprise telemetry built to prove the concept before access to real operational data.'),
]
gt = doc.add_table(rows=len(gl) + 1, cols=2); gt.style = 'Table Grid'
setcell(gt.rows[0].cells[0], 'Term', bold=True, color=WHITE, fill='0B1F3A')
setcell(gt.rows[0].cells[1], 'Meaning', bold=True, color=WHITE, fill='0B1F3A')
for i, (t, d) in enumerate(gl, start=1):
    setcell(gt.rows[i].cells[0], t, bold=True, size=9, color=NAVY, fill='F4F6F8')
    setcell(gt.rows[i].cells[1], d, size=9)
for row in gt.rows:
    row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(5.0)

para('')
para('Business Edition — architecture-level disclosure. Exact detection logic, formulas, thresholds, and scoring '
     'weights are proprietary and available in the full technical edition under NDA. All results are synthetic and '
     'must be re-validated on production telemetry. Do not redistribute.', 9, GREY, italic=True)

out = OUT; n = 2
while True:
    try:
        doc.save(out); break
    except PermissionError:
        root, ext = os.path.splitext(OUT); out = f"{root}_v{n}{ext}"; n += 1
print('Saved:', out)
