"""Executive briefing for a U.S. Army Cyber leader — V-Intelligence UEBA innovation.
Run:  python docs/build_army_cyber_brief.py
Out:  docs/VIntelligence_Army_Cyber_Brief.docx

6 pages: (1) the problem, (2) Salt & Volt Typhoon detection difficulty, (3) why current
tools fall short, (4) how we solved it (high level), (5) results, (6) preemptive cyber +
advanced use cases. All results are SYNTHETIC and labeled as such.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
NAVY = RGBColor(0x0D, 0x1B, 0x2A)
BLUE = RGBColor(0x1B, 0x4F, 0x72)
RED = RGBColor(0xB9, 0x1C, 0x1C)
GREEN = RGBColor(0x1E, 0x8A, 0x49)
GRAY = RGBColor(0x55, 0x55, 0x55)


def _base(doc):
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6)


def h1(doc, text, n=None):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
    if n:
        r = p.add_run(f"{n}  "); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = RED
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY


def h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BLUE


def body(doc, text, italic=False, color=None):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic = italic
    if color: r.font.color.rgb = color
    return p


def bullet(doc, text, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if prefix:
        r = p.add_run(prefix); r.bold = True
    p.add_run(text)


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill); tcPr.append(shd)


def callout(doc, text, fill="EAF2FB", color=None):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0); _shade(cell, fill)
    p = cell.paragraphs[0]; r = p.add_run(text); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = color or NAVY
    doc.add_paragraph()


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, hd in enumerate(headers):
        t.rows[0].cells[i].paragraphs[0].add_run(hd).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def build():
    doc = Document(); _base(doc)

    # ---------- Title ----------
    for _ in range(2): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("V-Intelligence UEBA"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = NAVY
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("Catching the threats that stay inside normal —\nby magnitude AND direction of behavioral change")
    rs.italic = True; rs.font.size = Pt(14); rs.font.color.rgb = BLUE
    for _ in range(2): doc.add_paragraph()
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = c.add_run("22nd Century Technologies, Inc. (TSCTI)"); rc.bold = True; rc.font.size = Pt(13)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rsub = sub.add_run("Behavioral Entity Intelligence — a US-developed detection capability")
    rsub.font.size = Pt(11); rsub.font.color.rgb = GRAY
    for _ in range(6): doc.add_paragraph()
    pf = doc.add_paragraph(); pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rpf = pf.add_run("Prepared for: Brandon Pugh  ·  U.S. Army Cyber")
    rpf.font.size = Pt(11); rpf.font.color.rgb = NAVY
    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = note.add_run("Executive briefing. Quantitative results herein are from a controlled "
                      "synthetic evaluation and require re-validation on operational data.")
    rn.italic = True; rn.font.size = Pt(9); rn.font.color.rgb = GRAY
    doc.add_page_break()

    # ---------- Page 1: The Problem ----------
    h1(doc, "The Problem: the most dangerous threats stay inside “normal”", "1.")
    body(doc, "The intrusions and insiders that do the most damage are engineered to keep every "
              "individual signal inside its normal range. No single metric breaks. The damage is in "
              "the DIRECTION and ACCUMULATION of behavioral change over time — what an entity is "
              "becoming — not in any threshold being crossed.")
    bullet(doc, "logs in to different things, not more things; shifts access from public to "
                "restricted content at the same volume.", prefix="An insider ")
    bullet(doc, "beacons a few times a day for months using legitimate tools and credentials — no "
                "malware, no flood, nothing loud.", prefix="A patient nation-state actor ")
    body(doc, "Conventional detection scores each metric on its own axis, so a change of KIND at "
              "constant volume is invisible. The defender is left with two bad outcomes: the threat "
              "is missed entirely, or it produces an undifferentiated “anomalous” alert "
              "with no direction — feeding alert fatigue and manual triage while the adversary dwells.")
    callout(doc, "The gap adversaries exploit: defenders ask “is any single metric abnormal?” "
                 "The right question is “what is this entity becoming?”")
    doc.add_page_break()

    # ---------- Page 2: Salt & Volt Typhoon ----------
    h1(doc, "Why Salt & Volt Typhoon Are So Hard to Detect", "2.")
    body(doc, "These PRC state-sponsored campaigns are the defining example of the problem above — "
              "and exactly the class V-Intelligence is built to surface.")
    h2(doc, "Volt Typhoon — living-off-the-land in U.S. critical infrastructure")
    bullet(doc, "dwelled in U.S. critical-infrastructure networks UNDETECTED for at least five "
                "years (CISA Advisory AA24-038A).", prefix="Dwell time: ")
    bullet(doc, "living-off-the-land — uses built-in tools and valid credentials, leaves no "
                "malware signature; every action looks like routine administration.", prefix="Tradecraft: ")
    h2(doc, "Salt Typhoon — compromise of U.S. telecommunications infrastructure")
    bullet(doc, "deep, persistent access into U.S. telecom networks (U.S. Government reporting, "
                "2024–2025), again via legitimate-looking access patterns.", prefix="Scope: ")
    body(doc, "Why signature- and threshold-based tools fail on both:")
    bullet(doc, "no malware signature to match (living-off-the-land).", prefix="No signature: ")
    bullet(doc, "activity stays within per-metric norms — constant volume, valid accounts.", prefix="No threshold breach: ")
    bullet(doc, "the only signal is a slow change in behavioral DIRECTION over months, which "
                "per-metric tools never compute.", prefix="The signal is directional: ")
    callout(doc, "A five-year dwell time is not bad luck — it is the predictable result of "
                 "watching magnitudes instead of direction.", fill="FDECEA", color=RED)
    doc.add_page_break()

    # ---------- Page 3: Why current apps fall short ----------
    h1(doc, "Why Current Tools Are Not Good Enough", "3.")
    body(doc, "Commercial UEBA and SIEM analytics are mature and useful — but, by their "
              "publicly described methodology, they measure MAGNITUDE and rarity, not behavioral "
              "direction.")
    bullet(doc, "Exabeam, Securonix, Microsoft Sentinel UEBA — statistical / p-value "
                "baselining, rarity- and magnitude-based scores, count- or TF-IDF-based peer "
                "grouping.", prefix="How they detect: ")
    bullet(doc, "their 2024–2026 generative-AI features are SOC copilots and conversational "
                "layers — assistants, not the detection mechanism.", prefix="The GenAI in them: ")
    body(doc, "The consequences for exactly the threats above:")
    bullet(doc, "on a controlled benchmark of four advanced, long-duration campaigns, traditional "
                "magnitude-based detectors caught at most ONE of four — the stealthy ones "
                "ranked among ordinary users.", prefix="Detection gap: ")
    bullet(doc, "even when one fires, the alert says only “anomalous” — no zone, no "
                "direction, no threat mapping — so every alert costs the same manual triage.",
           prefix="Intelligence gap: ")
    callout(doc, "Two gaps, one root cause: tools that score each metric independently cannot see a "
                 "coordinated change of behavior, and cannot explain the rare hit.")
    doc.add_page_break()

    # ---------- Page 4: How we solved it ----------
    h1(doc, "How We Solved It (High Level)", "4.")
    body(doc, "V-Intelligence builds a DIGITAL TWIN of every entity — user, device, "
              "application, network segment, session — and detects by the direction that twin "
              "is drifting.")
    bullet(doc, "each entity is a living, multi-dimensional behavioral profile, tracked as a "
                "trajectory over time, not a single snapshot.", prefix="Digital twin: ")
    bullet(doc, "behavior is serialized into language and embedded into one unified semantic space "
                "— bringing the representational power of large language models into the "
                "detection layer, coupled with rigorous mathematics.", prefix="Behavior-as-language: ")
    bullet(doc, "we measure the DIRECTION of drift and project it onto named threat concepts mapped "
                "to MITRE ATT&CK — so an alert says WHICH behavior changed and TOWARD WHAT "
                "pattern.", prefix="Direction, not magnitude: ")
    bullet(doc, "peer-cohort comparison, change-point detection, and multi-phase fusion combine "
                "into one ranked, explained verdict that resists evasion.", prefix="Fused & explainable: ")
    callout(doc, "Where an LLM uses that semantic space to understand language, V-Intelligence uses "
                 "it to understand behavior — and to say what an entity is becoming.", fill="E6F4EA", color=GREEN)
    doc.add_page_break()

    # ---------- Page 5: Results ----------
    h1(doc, "Results (Controlled Synthetic Evaluation)", "5.")
    body(doc, "Evaluation: 250 entities over ~485 days, with four embedded advanced campaigns — "
              "a patient insider, a slow command-and-control intrusion, and Volt- and Salt-Typhoon-"
              "class nation-state campaigns.", italic=True, color=GRAY)
    table(doc, ["Approach", "Campaigns caught (of 4)", "False-positive rate", "Directional explanation"],
          [["Traditional magnitude-based\n(LOF, IForest, OCSVM, Z-Score)", "0 – 1 of 4", "4.5% – 14.6%", "None"],
           ["V-Intelligence composite", "4 of 4", "10.6%", "Zone + direction + MITRE concept"],
           ["V-Intelligence threat-profile detector", "4 of 4", "0%", "Named technique (C2-beacon, DGA, LOTL, etc.)"]],
          widths=[2.4, 1.5, 1.4, 1.7])
    bullet(doc, "the multi-front threat-profile detector — the primary detector — caught all four by "
                "named known-bad technique (C2-beacon, DGA, LOTL-process, cohort-rare access, "
                "recon-fanout, insider-collection) at ZERO false positives.", prefix="Primary detection: ")
    bullet(doc, "embedding composite scoring also caught all four — including BOTH Volt and Salt Typhoon — "
                "at the operating point that recalls every campaign (10.6% false positives), though it "
                "cleanly separates only two of the four (USR-118, USR-156); ranking AUC ≈ 0.98.", prefix="Composite: ")
    bullet(doc, "traditional detectors caught at most one of the four and ranked the stealthy "
                "campaigns among ordinary users.", prefix="The contrast: ")
    callout(doc, "Surfacing a five-year-invisible, living-off-the-land campaign AT ALL — with "
                 "an explanation an analyst can act on — is the capability gap this closes.")
    body(doc, "Scope: all figures are from synthetic data and must be re-validated on "
              "operational telemetry. A scoped pilot on real data is the recommended next step.",
         italic=True, color=GRAY)
    doc.add_page_break()

    # ---------- Page 6: Preemptive + advanced use cases ----------
    h1(doc, "Preemptive Cybersecurity & Advanced Use Cases", "6.")
    h2(doc, "Preemptive cybersecurity — act before compromise completes")
    body(doc, "Because the signal is a directional drift that builds over weeks and months, "
              "V-Intelligence is an EARLY-WARNING capability, not just post-hoc detection. The "
              "five-year Volt Typhoon dwell time is precisely the window this is designed to close.")
    bullet(doc, "insider-risk and account-takeover early warning before exfiltration.", prefix="Get-left-of-impact: ")
    bullet(doc, "configuration-drift and continuous-monitoring support for RMF / ATO and ongoing "
                "authorization.", prefix="Continuous monitoring: ")
    bullet(doc, "explainable, MITRE-mapped outputs that fit human-on-the-loop decision authority.",
           prefix="Analyst-ready: ")
    h2(doc, "One architecture, many missions")
    body(doc, "The same digital-twin + behavioral-drift engine is domain-agnostic and has been "
              "applied beyond cyber:")
    bullet(doc, "supplier and supply-chain risk (counterfeit, adversary-nation sourcing, "
                "concentration).", prefix="Defense supply chain: ")
    bullet(doc, "drift-aware demand forecasting and capacity planning.", prefix="Logistics / DLA: ")
    bullet(doc, "prescriptive maintenance — predict equipment failure from behavioral drift.",
           prefix="Readiness: ")
    bullet(doc, "fraud / improper-payment and counter-intelligence pattern detection.", prefix="Mission integrity: ")
    callout(doc, "Model- and cloud-agnostic, air-gap capable, US-developed and patent-pending — "
                 "a capability that complements existing tools rather than replacing them.", fill="E6F4EA", color=GREEN)

    out = os.path.join(HERE, "VIntelligence_Army_Cyber_Brief.docx")
    doc.save(out)
    print("Saved:", out)


if __name__ == "__main__":
    build()
