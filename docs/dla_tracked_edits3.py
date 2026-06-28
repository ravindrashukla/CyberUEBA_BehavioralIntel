"""Third pass of tracked edits (quick wins from the second-round panel) on the DLA REV file.
  1) Bind worked-example numbers + calibration to the validated config (V2, 29 features).
  2) Competitive-landscape paragraph (Kinaxis/o9/Blue Yonder; Interos/Everstream/Resilinc).
  3) Glossary: add SHAP, MASE, pinball loss, SBOM; seasonality caveat; name EBS/DPDS/DLMS.
  4) Exec-summary plain-language lead paragraph + split para-4's worst run-on sentence.
All tracked (author=ReviewPanel). Re-extracts accepted text for verification.
"""
import os, copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REV = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_final_REV.docx"
AUTHOR = "ReviewPanel"; DATE = "2026-06-19T14:00:00Z"
d = Document(REV)
_id = [9000]
def nid():
    _id[0] += 1; return str(_id[0])

def _run(text, rpr=None, bold=False, italic=False):
    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    elif bold or italic:
        rPr = OxmlElement('w:rPr')
        if bold: rPr.append(OxmlElement('w:b'))
        if italic: rPr.append(OxmlElement('w:i'))
        r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t)
    return r

def _ins_wrap(run):
    ins = OxmlElement('w:ins'); ins.set(qn('w:id'), nid()); ins.set(qn('w:author'), AUTHOR); ins.set(qn('w:date'), DATE)
    ins.append(run); return ins

def _del_wrap(run):
    de = OxmlElement('w:del'); de.set(qn('w:id'), nid()); de.set(qn('w:author'), AUTHOR); de.set(qn('w:date'), DATE)
    de.append(run); return de

def append_ins(p, text):
    p._p.append(_ins_wrap(_run(text)))

def new_para(parts, style=None):
    p = OxmlElement('w:p'); pPr = OxmlElement('w:pPr')
    if style:
        ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'), style); pPr.append(ps)
    rPr = OxmlElement('w:rPr')
    ins = OxmlElement('w:ins'); ins.set(qn('w:id'), nid()); ins.set(qn('w:author'), AUTHOR); ins.set(qn('w:date'), DATE)
    rPr.append(ins); pPr.append(rPr); p.append(pPr)
    for (txt, b, it) in parts:
        p.append(_ins_wrap(_run(txt, bold=b, italic=it)))
    return p

def find(sub):
    return next(p for p in d.paragraphs if sub in p.text)

# ---- 1) config binding ----
p63 = find('The worked examples below show both failure')
append_ins(p63, " These worked-example forecasts and the calibration figures reported below come from the validated "
    "embedding-enriched configuration (V2, 29 features); the 67-feature count is the full architecture, of which not "
    "all layers were active in these measurements.")

# ---- 2) competitive landscape (after the novelty/composition paragraph) ----
pnov = find('claimed to be novel in isolation')
comp = new_para([("Positioning relative to existing tools. ", True, False),
    ("EDM / BEI is intended to complement, not replace, the established demand-planning and supply-chain-risk markets. "
     "Demand-sensing and planning platforms (for example Kinaxis, o9, and Blue Yonder) and supplier-risk and n-tier "
     "visibility platforms (for example Interos, Everstream, and Resilinc) already provide control-tower visibility, "
     "probabilistic planning, and supplier monitoring. The differentiation claimed here is not any single technique but "
     "the composition: item-location behavioral drift, calibrated planning bands, peer-cohort transfer for sparse "
     "items, and a cross-domain early-warning method applied together at the application layer on DLA data. This "
     "positioning should be validated against incumbent tooling during the pilot.", False, False)])
pnov._p.addnext(comp)

# ---- 3a) seasonality caveat on the snapshots paragraph ----
snap = None
for key in ['monthly snapshots', '16 to 23', 'forming the trajectory']:
    try:
        snap = find(key); break
    except StopIteration:
        continue
if snap is not None:
    append_ins(snap, " Because this proof-of-concept window spans roughly two years, it covers only about two annual "
        "cycles; full seasonal and exercise-cycle validation requires the longer DLA demand history and is a pilot "
        "data-readiness item.")
print("seasonality anchor:", "ok" if snap is not None else "NOT FOUND")

# ---- 3b) EBS/DPDS/DLMS integration surface ----
pimpl = find('implementation details for EDM')
append_ins(pimpl, " It will also name the target integration surface, including the Enterprise Business System (EBS), "
    "the demand-planning system, and the DLMS transaction layer, and the closed-loop path by which a validated "
    "planning band re-enters the system of record rather than remaining a side-by-side view.")

# ---- 4) exec-summary plain-language lead + split worst sentence ----
exec_h = next(p for p in d.paragraphs if p.text.strip() == 'Executive Summary' and p.style.name == 'Heading 1')
lead = new_para([("In plain terms, EDM / BEI models every part, supplier, depot, and route as a living profile and "
    "watches which way its behavior is heading, not just where it stands today, so DLA can act before a threshold "
    "breaks. On worked demand items it cut the 8-week buy by 40 to 54 percent (about $460,000 less over-buy on one "
    "part) while still covering realized demand, and it produced forecasts for items that legacy methods report as "
    "zero. The results below are a proof of concept on synthetic data and would be re-validated on DLA data.", False, False)])
exec_h._p.addnext(lead)

# split para-4 worst sentence (single-run paragraph)
p4 = find('This paper presents Entity Digital Model')
r0 = p4.runs[0]; full = r0.text
rpr = r0._r.find(qn('w:rPr'))
s_start = full.index('Each item-location')
s_end = full.index('explainable planning or risk signals.') + len('explainable planning or risk signals.')
old = full[s_start:s_end]
new = ("Each item-location, supplier, depot, route, relationship, and network is represented as a behavioral entity. "
       "Its movement is measured over time and compared against its own history and peer cohorts. It is aligned to "
       "known supply-chain concepts and translated into explainable planning or risk signals.")
before, after = full[:s_start], full[s_end:]
for t in r0._r.findall(qn('w:t')):
    t.text = before
delw = _del_wrap(_run(old, rpr=rpr))
# convert delete run's w:t -> w:delText (w:t is nested inside the run, use iter)
for t in list(delw.iter(qn('w:t'))):
    dt = OxmlElement('w:delText'); dt.set(qn('xml:space'), 'preserve'); dt.text = t.text
    t.getparent().replace(t, dt)
insw = _ins_wrap(_run(new, rpr=rpr))
afterrun = _run(after, rpr=rpr)
r0._r.addnext(afterrun); r0._r.addnext(insw); r0._r.addnext(delw)

# ---- 3c) glossary additions (before sectPr) ----
body = d.element.body
sectPr = body.find(qn('w:sectPr'))
EXTRA = [
    ("SHAP", "a method that explains a model's output by attributing it to individual input features, so each forecast or risk score is traceable."),
    ("MASE (mean absolute scaled error)", "a scale-free forecast-accuracy measure suited to intermittent demand; below 1 means better than a naive baseline."),
    ("Pinball loss", "the standard scoring rule for quantile (range) forecasts such as the P90; lower is better."),
    ("SBOM (software bill of materials)", "a structured list of the components in a software product, used to track supply-chain and component risk."),
]
for term, defn in EXTRA:
    el = new_para([(term + ". ", True, False), (defn, False, False)])
    if sectPr is not None: sectPr.addprevious(el)
    else: body.append(el)

o = REV; k = 2
while True:
    try:
        d.save(o); break
    except PermissionError:
        root, ext = os.path.splitext(REV); o = f"{root}_v{k}{ext}"; k += 1
print("Saved:", o)

# re-extract accepted text
d2 = Document(o)
from docx.table import Table
out = []
for child in d2.element.body.iterchildren():
    if child.tag == qn('w:p'):
        t = "".join(x.text or "" for x in child.iter(qn('w:t'))).strip()
        if t: out.append(t)
    elif child.tag == qn('w:tbl'):
        out.append("[TABLE]")
        for row in Table(child, d2).rows:
            out.append(" | ".join(c.text.strip().replace("\n", " ") for c in row.cells))
        out.append("[/TABLE]")
open("WP DLA/_DLA_paper_REV_accepted.txt", "w", encoding="utf-8").write("\n".join(out))
print("accepted-text lines:", len(out))
