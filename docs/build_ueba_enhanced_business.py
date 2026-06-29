"""Revised business edition of the V-Intelligence UEBA whitepaper.

References the CyberUEBA_BehavioralIntel_Enhanced app and adopts the DLA 25
EDM Executive narrative framework (operating environment -> the real challenge
-> why legacy falls short -> bottom line in one slide -> two ways to detect the
same attack -> inside the approach (two layers) -> the pipeline -> plain terms
-> results -> roadmap -> deployment -> novelty -> limitations -> appendix).

Two-layer positioning (verified against the enhanced app):
  Layer 1 (precision / production) -- Multi-front threat-profile detector:
     named known-bad fingerprints, cohort-relative + raw-event + self-drift.
     VERIFIED 4/4 injected attacks at 0 false positives (data/threat_profile_alerts.csv).
  Layer 2 (discovery / differentiation) -- behavioral embeddings + drift +
     novelty: the semantic 'map of meaning' that names direction (MITRE ATT&CK)
     and is the unsupervised net for attacks not yet profiled.

Exact detection logic (formulas, thresholds, weights) withheld; NDA edition.
All numbers are synthetic-data results and must be re-validated on real telemetry.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "UEBA_Behavioral_Intelligence_Enhanced_BusinessEdition.docx")

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
ORG  = RGBColor(0xC2, 0x5A, 0x12)
GREY = RGBColor(0x55, 0x5F, 0x6B)
GREEN= RGBColor(0x1E, 0x7A, 0x44)
TEAL = RGBColor(0x08, 0x91, 0xB2)
WHITE= RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.1
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(doc.sections[0], m, Inches(0.85))


def H1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    return p

def H2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ORG
    return p

def para(text, size=10.5, color=None, bold=False, italic=False, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def bullet(text, lead=None):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run(text.lstrip()); r2.font.size = Pt(10.5)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def setcell(cell, text, bold=False, size=9.5, color=None, fill=None):
    if fill: shade(cell, fill)
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
    if color: r.font.color.rgb = color

def table(headers, rows, widths=None, hdr_fill='0B1F3A'):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        setcell(t.rows[0].cells[j], h, bold=True, color=WHITE, fill=hdr_fill)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            setcell(t.rows[i].cells[j], v, bold=(j == 0), fill='F4F6F8' if j == 0 else None)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def callout(text):
    """DLA-deck 'Executive point' shaded box."""
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    setcell(t.rows[0].cells[0], '', fill='EAF3F6')
    c = t.rows[0].cells[0]; c.paragraphs[0].clear() if hasattr(c.paragraphs[0], 'clear') else None
    p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run('Executive point   '); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = TEAL
    r2 = p.add_run(text); r2.font.size = Pt(10); r2.font.color.rgb = NAVY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def fourbox(items):
    """2x2 bottom-line grid: list of (label, body)."""
    t = doc.add_table(rows=2, cols=2); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = [t.rows[0].cells[0], t.rows[0].cells[1], t.rows[1].cells[0], t.rows[1].cells[1]]
    for cell, (lab, body) in zip(cells, items):
        shade(cell, 'F4F6F8'); cell.text = ''
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        r = p.add_run(lab); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = ORG
        p2 = cell.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(body); r2.font.size = Pt(9.5); r2.font.color.rgb = NAVY
    for row in t.rows:
        for cc in row.cells:
            cc.width = Inches(3.35)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ================= FIGURE GENERATION =================
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

FIGDIR = os.path.join(os.path.dirname(__file__), 'figs_ueba_enhanced')
os.makedirs(FIGDIR, exist_ok=True)

C_LEG = '#8A93A0'   # traditional grey
C_LEGD = '#C0392B'  # traditional red accent
C_VI = '#0891B2'    # V-Intelligence teal
C_NAVY = '#0B1F3A'
C_GOLD = '#B5651D'
plt.rcParams.update({'font.family': 'Calibri', 'font.size': 10, 'axes.edgecolor': '#C9CFD6',
                     'axes.labelcolor': C_NAVY, 'xtick.color': '#55607B', 'ytick.color': '#55607B'})

# ---- VERIFIED per-attacker fingerprints (data/threat_profile_alerts.csv) ----
# (id, attack, cohort, [cohort-z fronts], [raw-event fronts], self_drift, elevated)
FP_DATA = [
    ('USR-156', 'Insider',       'business',  [('mass_collection', 5.9)],
        [('cohort_rare_dst', 76)], 19.6, True),
    ('USR-234', 'Slow APT',      'developer', [],
        [('c2_beacon', 386), ('dga_dns', 160)], 1.4, False),
    ('USR-042', 'LOTL',          'security',  [('lotl_process', 4.5)],
        [], 1.8, False),
    ('USR-118', 'Salt Typhoon',  'developer', [('recon_fanout', 8.2)],
        [], 7.3, True),
]


def fig_slowlow():
    """The slow-and-low evasion: each week stays under the alarm; accumulated drift crosses it."""
    rng = np.random.default_rng(11)
    wk = np.arange(1, 61)
    weekly = 0.18 + 0.10 * (wk / 60) ** 1.4 + rng.normal(0, 0.035, len(wk))
    weekly = np.clip(weekly, 0.05, 0.46)
    thresh = 0.55
    fig, axes = plt.subplots(1, 2, figsize=(7.0, 3.0))
    ax = axes[0]
    ax.bar(wk, weekly, color=C_LEG, alpha=0.75, width=0.9, label='Weekly behavioral change (attacker)')
    ax.axhline(thresh, color=C_LEGD, lw=2.0, ls='--', label='Fixed alert threshold')
    ax.text(30, thresh + 0.03, 'Never crossed — no alert, ever', ha='center', fontsize=8.8,
            color=C_LEGD, fontweight='bold')
    ax.set_title('What a threshold-based tool sees\n(each week judged alone)', fontsize=9.5, color=C_NAVY)
    ax.set_xlabel('Week'); ax.set_ylabel('Change vs. baseline')
    ax.set_ylim(0, 0.75); ax.legend(loc='upper left', fontsize=8, frameon=False)
    ax = axes[1]
    drift = np.cumsum(weekly - 0.16)
    normal_hi = np.cumsum(np.full(len(wk), 0.045))
    ax.fill_between(wk, 0, normal_hi, color=C_LEG, alpha=0.30, label='Normal-user range')
    ax.plot(wk, drift, color=C_VI, lw=2.4, label='Accumulated behavioral drift (attacker)')
    above = np.where(drift > normal_hi)[0]; cross = above[above > 10][0]
    ax.plot(wk[cross], drift[cross], 'o', color=C_LEGD, ms=7, zorder=5)
    ax.annotate('Drift alarm —\nmonths, not years', xy=(wk[cross], drift[cross]),
                xytext=(12, drift[cross] + 0.35), fontsize=8.8, fontweight='bold', color=C_NAVY,
                arrowprops=dict(arrowstyle='->', color=C_NAVY, lw=1.2))
    ax.set_title('What behavioral intelligence sees\n(the same weeks, accumulated)', fontsize=9.5, color=C_NAVY)
    ax.set_xlabel('Week'); ax.set_ylabel('Accumulated drift')
    ax.set_ylim(0, drift[-1] * 1.32); ax.legend(loc='upper left', fontsize=8, frameon=False)
    for ax in axes:
        ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig_slowlow.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p


def fig_scoreboard():
    """Detection scoreboard — threat-profile detector is the hero: 4/4 at 0 FP."""
    methods = ['Local Outlier\nFactor', 'Isolation\nForest', 'One-Class\nSVM', 'Z-Score',
               'Semantic layer\n(on its own)', 'Multi-front\nThreat-Profile']
    det = [0, 0, 0, 1, 2, 4]
    fp  = [4.5, 5.3, 14.6, 9.8, None, 0.0]
    colors = [C_LEG, C_LEG, C_LEG, C_LEG, '#5FA8B8', C_VI]
    x = np.arange(len(methods))
    fig, axes = plt.subplots(1, 2, figsize=(7.2, 3.1))
    ax = axes[0]
    ax.bar(x, det, color=colors, width=0.64)
    for xi, d in zip(x, det):
        ax.text(xi, d + 0.08, f'{d}/4', ha='center', fontsize=9.5, fontweight='bold',
                color=(C_NAVY if d else C_LEGD))
    ax.set_xticks(x, methods, fontsize=7.2)
    ax.set_ylim(0, 4.8); ax.set_ylabel('Attack campaigns detected (of 4)')
    ax.set_title('Detection — traditional methods catch at most 1 of 4...', fontsize=9.2, color=C_NAVY)
    ax = axes[1]
    xv = [i for i, f in enumerate(fp) if f is not None]
    fv = [fp[i] for i in xv]; cv = [colors[i] for i in xv]
    ax.bar(xv, fv, color=cv, width=0.64)
    for xi, f in zip(xv, fv):
        ax.text(xi, f + 0.35, f'{f:.0f}%' if f else '0%', ha='center', fontsize=9, fontweight='bold',
                color=('#1E7A44' if f == 0 else '#0B1F3A'))
    ax.text(5, 1.3, '0 false\npositives', ha='center', fontsize=8.4, fontweight='bold', color='#1E7A44')
    ax.set_xticks(x, methods, fontsize=7.2)
    ax.set_ylim(0, 18); ax.set_ylabel('False-positive rate (%)')
    ax.set_title('...the threat-profile detector catches all four —\nat zero false positives', fontsize=9.2, color=C_NAVY)
    for ax in axes:
        ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig_scoreboard.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p


def fig_fingerprint():
    """The four named-profile matches: cohort-relative z (left) and raw-event counts (right)."""
    fig, axes = plt.subplots(1, 2, figsize=(7.2, 3.0))
    # LEFT: cohort-relative z fronts vs the 4.5 flag line
    ax = axes[0]
    labs = ['Insider\nmass_collection\n(USR-156)', 'LOTL\nlotl_process\n(USR-042)', 'Salt Typhoon\nrecon_fanout\n(USR-118)']
    zs = [5.9, 4.5, 8.2]
    cols = ['#C0392B', '#8E44AD', '#2980B9']
    xb = np.arange(len(labs))
    ax.bar(xb, zs, color=cols, width=0.6)
    ax.axhline(4.5, color=C_NAVY, lw=1.6, ls='--')
    ax.text(2.35, 4.5 + 0.25, 'flag line (4.5)', ha='right', fontsize=8.2, color=C_NAVY, fontweight='bold')
    for xi, z in zip(xb, zs):
        ax.text(xi, z + 0.15, f'{z}', ha='center', fontsize=9.5, fontweight='bold', color=C_NAVY)
    ax.set_xticks(xb, labs, fontsize=7.6)
    ax.set_ylabel('Cohort-relative deviation (robust z)')
    ax.set_title('Front A — cohort-relative profile match\n(entity vs role-group peers)', fontsize=9.0, color=C_NAVY)
    ax.set_ylim(0, 9.4); ax.spines[['top', 'right']].set_visible(False)
    # RIGHT: raw-event counts (log scale, mixed units annotated)
    ax = axes[1]
    rl = ['Slow APT\nC2 beacon\n(USR-234)', 'Slow APT\nDGA domains\n(USR-234)', 'Insider\nrare-IP fan\n(USR-156)']
    rv = [386, 160, 76]
    units = ['days persistent\n(robotic rhythm)', 'high-entropy\ndomains', 'cohort-rare\nexternal IPs']
    rc = ['#E67E22', '#E67E22', '#C0392B']
    xr = np.arange(len(rl))
    ax.bar(xr, rv, color=rc, width=0.6)
    for xi, v, u in zip(xr, rv, units):
        ax.text(xi, v + 6, f'{v}', ha='center', fontsize=9.5, fontweight='bold', color=C_NAVY)
        ax.text(xi, v / 2, u, ha='center', va='center', fontsize=6.8, color='white', fontweight='bold')
    ax.set_xticks(xr, rl, fontsize=7.6)
    ax.set_ylabel('Raw-event evidence (count)')
    ax.set_title('Front B — raw-event profile match\n(label-free attack fingerprints)', fontsize=9.0, color=C_NAVY)
    ax.set_ylim(0, 430); ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig_fingerprint.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p


def fig_beacon():
    """Label-free C2-beacon detection: robotic regularity vs a jittery legitimate service."""
    rng = np.random.default_rng(7)
    fig, axes = plt.subplots(1, 2, figsize=(7.2, 2.9))
    # build two call-out timelines over ~10 days
    base = 12.0  # hours between callouts
    beacon_gaps = base + rng.normal(0, base * 0.18, 40)   # CV ~ 0.18 -> robotic
    svc_gaps = np.abs(rng.exponential(base, 40))          # bursty legit service -> CV ~ 1.0
    bt = np.cumsum(beacon_gaps); st = np.cumsum(svc_gaps)
    ax = axes[0]
    ax.eventplot([bt, st], colors=[C_LEGD, C_LEG], lineoffsets=[1, 0], linelengths=0.7, linewidths=1.4)
    ax.set_yticks([0, 1], ['Legitimate\nservice', 'C2 beacon\n(USR-234)'], fontsize=8.4)
    ax.set_xlabel('Hours'); ax.set_title('Call-out timing over ~20 days', fontsize=9.2, color=C_NAVY)
    ax.spines[['top', 'right', 'left']].set_visible(False)
    ax = axes[1]
    cv_b = float(np.std(np.diff(bt)) / np.mean(np.diff(bt)))
    cv_s = float(np.std(np.diff(st)) / np.mean(np.diff(st)))
    bars = ax.bar(['C2 beacon\n(USR-234)', 'Legitimate\nservice'], [cv_b, cv_s], color=[C_LEGD, C_LEG], width=0.55)
    ax.axhline(0.65, color=C_NAVY, lw=1.6, ls='--')
    ax.text(1.42, 0.67, 'flag line (0.65)', ha='right', fontsize=8.0, color=C_NAVY, fontweight='bold')
    for b, v in zip(bars, [cv_b, cv_s]):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.02, f'{v:.2f}', ha='center', fontsize=9.2, fontweight='bold', color=C_NAVY)
    ax.set_ylabel('Rhythm irregularity (CV of gaps)')
    ax.set_title('Below the line = machine-like = beacon', fontsize=9.2, color=C_NAVY)
    ax.set_ylim(0, max(cv_s, 1.1) * 1.25); ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig_beacon.png'); fig.savefig(p, dpi=200); plt.close(fig)
    return p


import csv as _csv
# Figures are driven by the ENHANCED app's real data (70-week features + trajectories,
# composite scores, threat-profile alerts). Falls back to the local repo data if absent.
_ENH = r'C:\Users\shuklar\ClaudeCode\CyberUEBA_BehavioralIntel_Enhanced\data'
_DATA = _ENH if os.path.isdir(_ENH) else os.path.join(os.path.dirname(__file__), '..', 'data')
ATK_STYLE = [('USR-156', '#C0392B', 'Insider threat (USR-156)'),
             ('USR-234', '#E67E22', 'Slow APT (USR-234)'),
             ('USR-042', '#8E44AD', 'Volt Typhoon / LOTL (USR-042)'),
             ('USR-118', '#2980B9', 'Salt Typhoon (USR-118)')]
ATK_IDS = {a for a, _, _ in ATK_STYLE}


def _band(series_by_user, normal_ids):
    """Per-week 5th/50th/95th percentile band across normal users (ragged-safe)."""
    W = max(len(v) for v in series_by_user.values())
    lo, mid, hi = [], [], []
    for w in range(W):
        vals = [series_by_user[u][w] for u in normal_ids if len(series_by_user[u]) > w]
        a = np.array(vals)
        lo.append(np.percentile(a, 5)); mid.append(np.median(a)); hi.append(np.percentile(a, 95))
    return np.array(lo), np.array(mid), np.array(hi)


def _first_sep(s, hi, after=2):
    """First week a series clears the 95th-percentile band (sustained marker)."""
    above = np.where(s > hi[:len(s)])[0]
    above = above[above > after]
    return int(above[0]) if len(above) else None


def fig_separation():
    """Signal Separation (real data): feature-space drift hides most attackers;
    semantic embedding drift separates three of four. Stars = first sustained detection.
    Reproduces the app's dual-lens view with the actual detection weeks."""
    # ---- LEFT: feature-space CUSUM (raw-magnitude), real weekly features ----
    wf = {}
    with open(os.path.join(_DATA, 'comparison_results', 'weekly_features.csv')) as f:
        rdr = _csv.DictReader(f); FC = [c for c in rdr.fieldnames if c not in ('user_id', 'week_idx', 'week_start', 'week_end')]
        for r in rdr:
            wf.setdefault(r['user_id'], []).append((int(r['week_idx']), [float(r[c] or 0) for c in FC]))
    feat = {}
    for u, rows in wf.items():
        rows.sort(); X = np.array([v for _, v in rows]); n = len(X)
        bm, bs = X[:n // 2].mean(0), X[:n // 2].std(0); bs[bs == 0] = 1.0
        wd = np.abs((X - bm) / bs).mean(1)
        cu = np.cumsum(np.maximum(wd[1:] - 0.5, 0)); feat[u] = np.insert(cu, 0, 0.0)
    # ---- RIGHT: semantic embedding CUSUM, real trajectories ----
    sem_raw = {}
    _tr = _ueba_db_rows("SELECT user_id, week_idx, composite_drift FROM weekly_trajectories")
    if _tr is None:
        with open(os.path.join(_DATA, 'tier3_results', 'all250_trajectories.csv')) as f:
            _tr = list(_csv.DictReader(f))
    for r in _tr:
        sem_raw.setdefault(r['user_id'], []).append((int(r['week_idx']), float(r['composite_drift'])))
    sem = {u: np.cumsum([d for _, d in sorted(v)]) for u, v in sem_raw.items()}

    fig, axes = plt.subplots(1, 2, figsize=(7.4, 3.4))
    for ax, series, title, sub, subcol in (
            (axes[0], feat, 'Raw-magnitude lens — feature-space drift',
             'the loud attacks separate — but only late (weeks 36–47)', C_LEGD),
            (axes[1], sem, 'V-Intelligence "meaning" lens — semantic embedding drift',
             'the meaning lens catches the subtle insider 35 weeks sooner', '#1E7A44')):
        normal_ids = [u for u in series if u not in ATK_IDS]
        lo, mid, hi = _band(series, normal_ids)
        wk = np.arange(len(hi))
        ax.fill_between(wk, lo, hi, color='#BDC3C7', alpha=0.35, lw=0, label='Normal range (5–95%)')
        ax.plot(wk, mid, color='#9AA3AD', lw=1.3, ls='--', label='Normal median')
        for u, col, lab in ATK_STYLE:
            s = series[u]; ax.plot(np.arange(len(s)), s, color=col, lw=2.2, label=lab)
            fw = _first_sep(s, hi)
            if fw is not None:
                ax.plot(fw, s[fw], marker='*', color=col, ms=13, mec='white', mew=0.6, zorder=6)
        ax.set_title(title, fontsize=8.8, color=C_NAVY, pad=20)
        ax.text(0.5, 1.045, sub, transform=ax.transAxes, ha='center', fontsize=7.6,
                color=subcol, fontweight='bold')
        ax.set_xlabel('Week'); ax.set_ylabel('Cumulative drift (CUSUM)')
        ax.spines[['top', 'right']].set_visible(False)
        ax.legend(loc='upper left', fontsize=6.6, frameon=False)
    axes[0].annotate('$\\star$  marks first sustained detection (clears the normal band)',
                     xy=(0.5, -0.30), xytext=(0.5, -0.30),
                     xycoords='axes fraction', ha='center', fontsize=7.4, color=C_NAVY, style='italic')
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig_separation.png'); fig.savefig(p, dpi=200, bbox_inches='tight'); plt.close(fig)
    return p


def _ueba_db_rows(query):
    """Operational DB first (single source of truth); returns list of dict rows or None."""
    try:
        import psycopg2
        from dotenv import load_dotenv
        load_dotenv(os.path.join(os.path.dirname(__file__), '..', '.env'))
        url = os.getenv('DATABASE_URL_HOST') or os.getenv('DATABASE_URL')
        if not url:
            return None
        conn = psycopg2.connect(url, connect_timeout=3)
        cur = conn.cursor(); cur.execute(query)
        cols = [d[0] for d in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]
        conn.close()
        return rows or None
    except Exception:
        return None


def fig_radar():
    """5-Phase Percentile (real data, app min-max scaling): each attacker breaks
    out on a different combination; normal users stay a small central shape."""
    rows = _ueba_db_rows("SELECT uid, is_attack, grp, role, signal_strength, breadth_15, sustained_signal, ctx_max_z, novelty_score FROM composite_scores")
    if rows is None:
        rows = list(_csv.DictReader(open(os.path.join(_DATA, 'tier3_results', 'composite_scores.csv'))))
    cols = ['signal_strength', 'breadth_15', 'sustained_signal', 'ctx_max_z', 'novelty_score']
    labels = ['Signal\nStrength', 'Breadth', 'Sustained\nDeviation', 'Context\nDivergence', 'Novelty\nPersistence']
    mn = {c: min(float(r[c]) for r in rows) for c in cols}
    mx = {c: max(float(r[c]) for r in rows) for c in cols}

    def scale(v, c):
        rng = mx[c] - mn[c]
        return 100.0 * (float(v) - mn[c]) / rng if rng > 0 else 50.0

    normal = [r for r in rows if r['uid'] not in ATK_IDS]
    med = [scale(np.median([float(r[c]) for r in normal]), c) for c in cols]
    p75 = [scale(np.percentile([float(r[c]) for r in normal], 75), c) for c in cols]

    def user_vec(uid):
        r = next(r for r in rows if r['uid'] == uid); return [scale(r[c], c) for c in cols]

    ang = np.linspace(0, 2 * np.pi, len(cols), endpoint=False).tolist(); ang += ang[:1]
    fig, ax = plt.subplots(figsize=(5.8, 4.6), subplot_kw=dict(polar=True))
    series = ([(med, '#9AA3AD', 'Normal median', '--', 1.4),
               (p75, '#707A86', 'Normal 75th pct', ':', 1.4)] +
              [(user_vec(u), c, l, '-', 2.0) for u, c, l in ATK_STYLE])
    for vec, col, lab, ls, lw in series:
        v = list(vec) + [vec[0]]
        ax.plot(ang, v, color=col, ls=ls, lw=lw, label=lab)
        if ls == '-':
            ax.fill(ang, v, color=col, alpha=0.06)
    ax.set_xticks(ang[:-1]); ax.set_xticklabels(labels, fontsize=8.8, color=C_NAVY)
    ax.set_ylim(0, 100); ax.set_yticks([25, 50, 75, 100])
    ax.set_yticklabels(['normal', 'halfway', '', 'most extreme'], fontsize=6.8, color='#8A93A0')
    ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.06), ncol=3, fontsize=7.4, frameon=False)
    fig.tight_layout()
    p = os.path.join(FIGDIR, 'fig_radar.png'); fig.savefig(p, dpi=200, bbox_inches='tight'); plt.close(fig)
    return p


FIG_SLOW = fig_slowlow()
FIG_SCORE = fig_scoreboard()
FIG_FP = fig_fingerprint()
FIG_BEACON = fig_beacon()
FIG_SEP = fig_separation()
FIG_RADAR = fig_radar()


def add_fig(path, caption, width=6.4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY


# ================= TITLE =================
r = doc.add_paragraph().add_run('Catching the Cyber Threats That Never Trip an Alarm\n'
                                'Behavioral Intelligence vs. Traditional Security Monitoring')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
para('A business and technology overview of V-Intelligence UEBA: what it is, what it detects that current tools '
     'miss, how the two-layer detection engine works, and what the measured evidence shows.', 12.5, GREY, after=2)
para('22nd Century Technologies, Inc.  ·  V-Intelligence UEBA Program  ·  June 2026  ·  Business Edition (Enhanced)',
     9.5, GREY, after=2)
para('Disclosure note: this edition describes the architecture and the innovation and reports measured results. The '
     'exact detection logic — named-profile definitions, formulas, thresholds, weights, and parameters — is '
     'proprietary and is withheld here; it is available in the full technical edition under NDA (non-disclosure '
     'agreement).', 9.5, ORG, italic=True, after=4)
para('Data note: all results in this paper come from a synthetic (simulated) dataset built to prove the concept — '
     '250 users, 485 days (70 weeks), approximately 15 million log events across five log sources, with four embedded '
     'long-duration attack campaigns modeled on real, publicly documented intrusions. Results must be re-validated on '
     'real operational telemetry before any production claim.', 9.5, ORG, italic=True, after=8)

# ================= 1. EXECUTIVE SUMMARY =================
H1('1.  Executive Summary')
para('The most damaging cyber attacks of the past decade did not trip a single alarm. Volt Typhoon — a Chinese '
     'state-sponsored campaign against U.S. critical infrastructure — sat inside victim networks for at least five '
     'years undetected. Salt Typhoon — described by the chairman of the Senate Intelligence Committee as "the worst '
     'telecom hack in our nation\'s history" — operated inside major U.S. telecommunications providers for months. In '
     'both cases the victims\' own security tools never fired, and the campaigns were discovered by outside '
     'intelligence agencies. Such attackers succeed by operating with valid credentials and legitimate tools, '
     'staying within the statistical bounds of normal activity while shifting their behavior gradually over months.')
para('This paper presents the business case for User and Entity Behavior Analytics (UEBA) — a detection approach '
     'that learns a behavioral baseline for every user and system, then watches both how much behavior changes and '
     'the direction in which it is trending. The V-Intelligence UEBA program validated the approach on a realistic '
     'simulated enterprise: 250 users, 485 days, roughly 15 million events, and four embedded attack campaigns — a '
     '14-month insider threat, a 417-day slow nation-state intrusion with covert call-home traffic, a Volt '
     'Typhoon-style living-off-the-land campaign, and a Salt Typhoon-style telecom campaign.')
para('The detection engine has two layers, and that is the central message of this edition:', bold=True)
bullet('a library of measurable known-bad behavioral fingerprints — cohort-relative profiles (how an entity behaves '
       'versus its role-group peers) plus raw-event profiles (covert call-home beacons, algorithmically generated '
       'domains, rare external destinations). Each flag is a named profile match, so it is precise and explainable. '
       'On the validation set it detected all four attack campaigns at zero false positives.',
       lead='Layer 1 — Multi-front threat-profile detector (precision). ')
bullet('a behavioral "map of meaning" built from the same representation that powers large language models, which '
       'measures drift and novelty and names the direction of change against MITRE ATT&CK. It is the discovery net '
       'for attacks not yet in the fingerprint library, and it is what makes the precision layer\'s beacon and '
       'domain fingerprints possible. On its own it cleanly separates two of the four campaigns and supplies the '
       'explanation for all of them.', lead='Layer 2 — Behavioral embeddings, drift, and novelty (discovery). ')
para('The headline result, stated plainly:', bold=True)
bullet('Of four industry-standard anomaly-detection algorithms run against the same data, three detected none of the '
       'four attack campaigns and the best detected exactly one — a statistical spike with no explanation of which '
       'behavior changed or what kind of threat it resembled.', lead='Traditional methods miss these attacks. ')
bullet('Layer 1 matched all four campaigns to named attack fingerprints — and flagged no normal user. Every '
       'detection carries the specific profile it matched and the MITRE ATT&CK technique it represents.',
       lead='Four of four detected, zero false positives, each one named. ')
para('The underlying distinction is straightforward. An insider changes what they access, not how much; a slow '
     'intrusion changes where it communicates, not how often. Traditional algorithms measure magnitude; '
     'V-Intelligence UEBA measures direction, peer-relative deviation, and named-profile match.', bold=True)

# ================= 2. OPERATING ENVIRONMENT =================
H1('2.  The Operating Environment Requires Faster, Better-Explained Detection')
para('A modern security operations center (SOC — the analyst team that watches alerts and runs investigations) is '
     'not short of data. It is short of context: which of today\'s thousands of alerts is the early stage of a real '
     'intrusion, and what exactly is the evidence. Analysts drown in low-context alerts while the attacks that '
     'matter most generate no alert at all, because every individual action is authorized.')
callout('The bottleneck is not missing logs — it is missing behavioral context. The enterprise already collects the '
        'telemetry needed to catch these attacks; what is missing is an engine that reads each entity\'s behavior in '
        'context, names what changed, and hands the analyst a short, explained list instead of a long, raw one.')

# ================= 3. THE CHALLENGE =================
H1('3.  The Challenge Is Unseen Intent, Not Missing Data')
para('Enterprise security monitoring rests on two pillars. The first is the SIEM (Security Information and Event '
     'Management system), the central log platform that applies expert-written rules to spot known attack patterns. '
     'The second is signature-based detection: antivirus and intrusion-detection tools that match activity against '
     'IOCs (Indicators of Compromise — known-bad file fingerprints, addresses, and domains). Both answer the same '
     'question: does this activity match something already known to be malicious? Modern adversaries are engineered '
     'to make the answer "no." Three threat families account for most of these cases.')
bullet('A trusted employee or contractor who misuses access they legitimately hold. Industry data (the Verizon Data '
       'Breach Investigations Report) attributes roughly 20% of breaches to insiders, with a median time to '
       'detection exceeding 200 days. Every individual action uses valid credentials; only the pattern over months '
       'reveals intent.', lead='Insider threats. ')
bullet('An APT (Advanced Persistent Threat) is a patient, well-resourced intrusion — typically nation-state — that '
       'gains a foothold and stays for months or years. Its covert command-and-control (C2) channel, through which '
       'the implant calls home for instructions, is throttled to a trickle: a check-in every several hours, only '
       'kilobytes of traffic, invisible inside enterprise volume.', lead='Advanced persistent threats. ')
bullet('Living-off-the-land (LOTL) attacks use only the administration tools already installed on every system — '
       'PowerShell, remote desktop, management utilities. There is no malware to fingerprint and no rule to violate, '
       'because every action is, in isolation, an authorized one.', lead='Living-off-the-land attacks. ')
para('The common technique is "slow and low": behavior changes so gradually that no single week looks unusual, while '
     'the cumulative change over months achieves the objective. The activity is visible only in the trajectory, '
     'never in a single snapshot.')
add_fig(FIG_SLOW, 'Figure 1 — The "slow and low" evasion. Left: a threshold-based tool judges each week alone, and '
                  'the attacker\'s weekly change never crosses the fixed alert line — no alert is ever raised. Right: '
                  'the same weeks accumulated as behavioral drift separate the attacker from the normal-user range in '
                  'months, not years. (Illustrative, based on the simulated 60-week slow-APT campaign.)')
para('The real-world consequence is documented. Neither Volt Typhoon nor Salt Typhoon was caught by any victim\'s '
     'own monitoring; SIEM, intrusion detection, and endpoint tools all failed because there was no signature to '
     'match and no threshold crossed. Both were found by external intelligence operations, years or months after '
     'compromise. This is a structural gap, not a tuning problem — and it is the gap behavioral analytics is built '
     'to close.')

# ================= 4. BOTTOM LINE =================
H1('4.  Bottom Line — the Opportunity in One Page')
para('Same enterprise logs, read a different way — a behavioral detection legacy tools cannot produce.', bold=True)
fourbox([
    ('THE PROBLEM',
     'Signature and threshold tools ask "does this match known-bad?" The attacks that matter answer "no" — they use '
     'valid credentials and legitimate tools, and change behavior too slowly to cross any line.'),
    ('WHAT V-INTELLIGENCE DOES',
     'Reads the same logs as a behavioral profile per entity: where it sits versus its peers, where it is heading '
     '(drift), and whether it matches a named known-bad fingerprint — then names the threat and the ATT&CK technique.'),
    ('THE PROOF',
     'Four embedded campaigns — insider, slow APT, LOTL, Salt Typhoon — all four matched to named attack fingerprints '
     'at ZERO false positives, where the best traditional method caught one of four with no explanation.'),
    ('THE ASK',
     'Validate on real telemetry through a bounded pilot — uses five log sources most enterprises already collect; '
     'runs in containers (cloud or air-gapped); integrates with the existing SIEM/SOAR stack.'),
])
para('Figures throughout are synthetic / representative demo data, to be re-validated on real operational telemetry '
     'before any production claim.', 9, GREY, italic=True)

# ================= 5. TWO WAYS =================
H1('5.  Two Ways to Detect the Same Attack')
para('Consider the slow nation-state intrusion in the validation set (user USR-234): a 417-day campaign whose '
     'implant calls home roughly every several hours in kilobyte-sized packets, with total traffic, sign-ins, and '
     'lookups all inside normal ranges. Here is how each approach handles exactly the same data.')
table(['', 'Traditional monitoring (threshold / signature)', 'V-Intelligence UEBA (behavioral)'],
      [('What it looks at',
        'Each day\'s totals against fixed thresholds; each connection against a known-bad list.',
        'Every entity\'s behavior versus its own past and its role-group peers, week over week, in a shared map of meaning.'),
       ('Key steps',
        'Volume thresholds → signature match → rule correlation in the SIEM.',
        'Baseline per entity → measure drift and novelty → match named known-bad profiles → name the ATT&CK technique.'),
       ('On USR-234',
        'Nothing crosses a threshold; no signature matches; no alert is ever raised.',
        'The call-home channel is caught by its robotic rhythm (a beacon fingerprint) and 160 algorithmically '
        'generated domains — a named, explained match.'),
       ('What the analyst gets',
        'Either silence, or a context-free spike with no "why."',
        'A short ranked list, each item naming the matched profile, the behavioral zone that moved, and the threat it resembles.')],
      widths=[1.2, 2.9, 2.9])
callout('The difference is not a better threshold. It is a different question: not "did this cross a line?" but '
        '"does this entity\'s behavior match a known-bad pattern its peers do not show?"')

# ================= 6. INSIDE THE APPROACH =================
H1('6.  Inside the Approach — One Behavioral Profile, Two Detection Layers')
H2('6.1  The digital entity — a medical chart for behavior')
para('V-Intelligence maintains a "digital entity" — a living behavioral profile — for every monitored user and '
     'system. The closest analogy is a patient\'s medical chart: it does not record only today\'s temperature; it '
     'accumulates history, vitals, and context, and a clinician reads the trend to act before a condition becomes '
     'acute. Each week, raw activity from five log sources (sign-ins, file access, endpoint activity, network '
     'traffic, and domain lookups) is distilled into roughly two dozen behavioral measurements per user, organized '
     'into five behavioral zones.')
table(['Behavioral zone', 'What it captures', 'What it catches when it drifts'],
      [('Identity', 'Who the entity is — role, department, clearance, tenure.', 'Stability here is itself a clue: insiders and credential thieves still look like the same person.'),
       ('Access pattern', 'How the entity signs in — frequency, failures, timing, sources.', 'Credential abuse and account takeover.'),
       ('Data behavior', 'What data it touches — volume, sensitivity mix, read/write balance.', 'Insider data theft: a shift toward restricted and confidential material.'),
       ('Network footprint', 'Where it communicates — destinations, external traffic, domain lookups.', 'Covert call-home channels and data-removal routes.'),
       ('Risk posture', 'Endpoint health — suspicious process activity and risk indicators.', 'Compromised machines and privilege-escalation tooling.')],
      widths=[1.3, 2.7, 2.9])

H2('6.2  Layer 1 — the multi-front threat-profile detector (precision)')
para('The production detection layer is a library of measurable known-bad fingerprints. Each fingerprint is a '
     'precise, named description of how a specific attack technique shows up in behavior, and a flag is raised only '
     'when an entity matches one. Because a flag means "this entity matches the recon-fan-out profile" rather than '
     '"this entity is statistically unusual," the result is both high-precision and self-explaining. The detector '
     'fuses three fronts.')
bullet('How an entity behaves relative to its role-group peers, not relative to the whole company. A developer is '
       'compared with developers, security staff with security staff. A named profile fires only when the right '
       'combination of measurements is jointly elevated against peers — for example, "mass collection" requires both '
       'a surge in unique files and in total file volume; "recon fan-out" requires an unusual spread of network '
       'destinations. Requiring the combination, not any single metric, is what keeps false positives near zero.',
       lead='Front A — cohort-relative profiles. ')
bullet('Fingerprints read straight from the raw events, with no labels. The covert call-home beacon is found by its '
       'machine-like regularity (a robotic, evenly spaced rhythm that legitimate services do not have); '
       'algorithmically generated domains are found by their high randomness and shared infrastructure; rare '
       'external destinations are found by being contacted by this entity alone within its cohort.',
       lead='Front B — raw-event profiles. ')
bullet('Whether the entity has drifted from its own past. This front never raises a flag by itself; it corroborates '
       'and strengthens a match found by the other two.', lead='Front C — self-drift (supporting). ')

H2('6.3  Layer 2 — behavioral embeddings, drift, and novelty (discovery)')
para('The discovery layer is what makes the precision layer possible — and what protects against attacks the '
     'fingerprint library does not yet contain. Each week\'s measurements are first written out as a structured '
     'plain-language description (for example, "30 files, 15% restricted, 8% confidential" rather than just "30 '
     'files"), and that description is converted into a high-dimensional numerical vector called an embedding — the '
     'same representation that powers large language models. The result is a semantic space: a map of meaning in '
     'which behaviors that mean similar things sit near each other even when their raw numbers differ.')
para('This matters because raw counts can hide intent. The simulated insider accessed about 30 files per week both '
     'before and during the campaign — so on raw counts nothing changed — but week one was public files and week '
     'forty was restricted files, and on the map of meaning those weeks sit far apart. Tracking each entity\'s '
     'weekly position produces a trajectory; the trajectory yields drift (the direction and rate of behavioral '
     'change) and novelty (a never-before-seen behavior that keeps recurring). Drift is then named against a library '
     'of reference threat concepts aligned with MITRE ATT&CK, the industry-standard public catalog of adversary '
     'techniques — which is how an alert can say "drifting toward covert beaconing" and attach the matching '
     'techniques.')
para('The two lenses are complementary, and that is the point. Over the full 70-week campaign, the raw-magnitude '
     'lens separates the noisy, high-volume attacks earliest — it lifts the Salt Typhoon campaign clear of the '
     'normal band by week 36 — but is slow on the subtle ones. The semantic lens catches the quiet insider 35 weeks '
     'sooner (week 4 versus week 39) by reading meaning rather than volume. Neither lens, on its own, ever separates '
     'the slow APT — which is exactly why the precision layer\'s label-free beacon fingerprint (Section 9.3) matters.')
add_fig(FIG_SEP, 'Figure 2 — Signal separation, both lenses, over the full 70-week campaign (measured on the '
                 'enhanced-app data; ★ = first sustained separation from the normal band). Left: the raw-magnitude '
                 'lens separates the high-volume attacks late (Salt Typhoon wk 36, insider wk 39, LOTL wk 47). '
                 'Right: the semantic "meaning" lens catches the insider at wk 4 and the LOTL at wk 15 — far sooner '
                 '— but trails on the volume attack. Each lens wins on different attacks; neither catches the slow '
                 'APT, which the threat-profile beacon does.')
para('The same five behavioral questions the discovery layer asks — how strong the deviation is versus peers, across '
     'how many measurements, how sustained, under how many analytical lenses, and how persistent any novel behavior '
     'is — also explain why each attacker looks different. No normal user reaches the extremes on several questions '
     'at once.')
add_fig(FIG_RADAR, 'Figure 3 — The discovery layer\'s five behavioral questions, as a population-scaled profile '
                   '(measured on the enhanced-app data). Normal users stay a small shape near the center (grey). '
                   'Each attacker breaks out on a different combination — the slow APT spikes almost only on novelty '
                   'persistence, the insider dominates signal strength and breadth, Salt Typhoon extends on nearly '
                   'every question at once. No normal user reaches these extremes on multiple questions '
                   'simultaneously.', width=5.4)

# ================= 7. PIPELINE =================
H1('7.  How It Works — the Five-Step Pipeline')
table(['Step', 'What happens', 'In plain terms'],
      [('1.  Profile', 'Distill each week of raw logs into ~two dozen measurements per entity, in five zones.',
        'Build the medical chart for every user and system.'),
       ('2.  Embed', 'Write the measurements as text and convert to a 1,536-dimension semantic vector (MVP setting).',
        'Place each week on a shared "map of meaning."'),
       ('3.  Drift & novelty', 'Track each entity\'s trajectory; accumulate slow change; flag recurring new behavior.',
        'Notice where behavior is heading, even when each week looks normal.'),
       ('4.  Match profiles', 'Test each entity against the named known-bad fingerprints, cohort-relative and raw-event.',
        'Ask: does this match a recognized attack pattern its peers do not show?'),
       ('5.  Name & rank', 'Attach the matched profile and ATT&CK technique; return one short, ranked, explained list.',
        'Hand the analyst leads, not raw alerts.')],
      widths=[1.2, 3.1, 2.5])

# ================= 8. PLAIN TERMS =================
H1('8.  What "Cohort", "Drift", "Novelty", and "Profile" Mean — in Plain Terms')
table(['Term', 'Plain meaning', 'Why it matters here'],
      [('Cohort', 'An entity\'s peer group — people in the same role (developers, security staff, business users).',
        'Behavior is judged against peers, so "normal for a developer" is not confused with "normal for an executive."'),
       ('Drift', 'The direction and speed an entity\'s behavior is trending over weeks.',
        'Catches "slow and low" change that no single week reveals.'),
       ('Novelty', 'A never-before-seen behavior that then keeps recurring (e.g., a new external contact every week).',
        'The fingerprint of covert command-and-control infrastructure.'),
       ('Profile (fingerprint)', 'A named, measurable description of how a known attack technique shows up in behavior.',
        'A match is precise and self-explaining — "recon fan-out," not just "anomaly."'),
       ('Cohort-relative z', 'How far an entity sits from its peers, on a robust scale that ignores normal spread.',
        'Lets a profile fire only on genuine peer-relative extremes, keeping false positives low.')],
      widths=[1.5, 2.7, 2.7])

# ================= 9. RESULTS =================
H1('9.  What the Evidence Shows (Synthetic Data)')
H2('9.1  The scoreboard')
para('Four industry-standard anomaly detectors (Z-Score statistics, Isolation Forest, Local Outlier Factor, and '
     'One-Class SVM, or Support Vector Machine) were run against the full 250-user population alongside the two '
     'V-Intelligence layers. Every number below is reproduced directly from the evaluation data.')
table(['Method', 'Attacks detected', 'False positives', 'What the analyst gets'],
      [('Local Outlier Factor', '0 of 4', '4.5%', 'Low false positives — and zero detections. All four campaigns sit inside its normal range.'),
       ('Isolation Forest', '0 of 4', '5.3%', 'All four attackers look statistically normal in feature space.'),
       ('One-Class SVM', '0 of 4', '14.6%', 'Many false positives and zero true detections.'),
       ('Z-Score', '1 of 4', '9.8%', 'Catches only one spike (the LOTL campaign, barely over the line); the other three never cross it.'),
       ('Semantic layer (on its own)', '2 of 4', '—', 'Discovery layer: cleanly separates the insider and Salt Typhoon, and explains the direction of all four.'),
       ('Multi-front Threat-Profile', '4 of 4', '0', 'All four matched to a named attack fingerprint and ATT&CK technique. No normal user flagged.')],
      widths=[1.9, 1.2, 1.0, 2.8], hdr_fill='0B1F3A')
para('Stated precisely: the best traditional method catches 1 of 4, and its single detection is a context-free '
     'spike. The threat-profile detector catches all four — each matched to a named, measurable fingerprint — at '
     'zero false positives, while the semantic layer supplies the explanation and the unsupervised net for novel '
     'attacks. The result is not an alert but an investigation lead with the "why" already attached.', bold=True)
add_fig(FIG_SCORE, 'Figure 4 — The detection scoreboard across 250 users with four embedded campaigns (measured, '
                   'synthetic data). Traditional methods catch at most 1 of 4; the semantic layer separates 2 of 4 '
                   'on its own; the multi-front threat-profile detector catches 4 of 4 at zero false positives, each '
                   'detection named.')

H2('9.2  Four attacks, four named fingerprints')
para('Each campaign matched a different combination of profiles — empirical proof that no single detector would '
     'have sufficed, and that a flag here is a specific, explainable accusation rather than a statistical shrug.')
table(['Attack (user)', 'Cohort', 'Named profile match', 'How many fronts'],
      [('Insider (USR-156)', 'Business', 'mass_collection (peer-z 5.9) + cohort_rare_dst (76 rare external IPs)', '2 fronts'),
       ('Slow APT (USR-234)', 'Developer', 'c2_beacon (386-day persistent, robotic rhythm) + dga_dns (160 generated domains)', '2 fronts'),
       ('LOTL / Volt Typhoon (USR-042)', 'Security', 'lotl_process (peer-z 4.5 — unusual admin-tool breadth)', '1 front'),
       ('Salt Typhoon (USR-118)', 'Developer', 'recon_fanout (peer-z 8.2 — network-destination fan-out)', '1 front')],
      widths=[1.7, 1.0, 3.5, 0.9])
add_fig(FIG_FP, 'Figure 5 — The four named-profile matches (measured, synthetic data). Left: cohort-relative '
                'profiles fire only when an entity\'s peer-relative deviation clears the flag line (4.5). Right: '
                'raw-event profiles are read straight from the logs — the slow APT\'s persistent beacon (386 days) '
                'and generated domains (160), and the insider\'s rare-IP fan-out (76).')

H2('9.3  Case study: the attack that was "undetectable"')
para('The slow APT (USR-234) is the hardest detection problem in the study. For 417 days its implant called home '
     'roughly every several hours, and its total traffic, lookups, and sign-ins all stayed within normal ranges. On '
     'every magnitude-based method its scores are negligible — it is invisible to thresholds and to statistics '
     'alike. Yet the threat-profile detector matched it on two independent raw-event fingerprints, with no labels '
     'and no thresholds crossed.')
bullet('Among all of this user\'s external contacts, one destination was contacted persistently — on 386 distinct '
       'days — and, crucially, with a robotic, evenly spaced rhythm. Legitimate persistent services (software '
       'updates, cloud sync) are bursty and irregular; a beacon is metronomic. The detector flags the regularity, '
       'not the destination, so it needs no prior knowledge of the attacker\'s infrastructure.',
       lead='C2 beacon. ')
bullet('The same implant resolved 160 algorithmically generated domains — high-randomness names sharing back-end '
       'infrastructure — the signature of domain-generation malware locating its controller.', lead='Generated domains. ')
add_fig(FIG_BEACON, 'Figure 6 — Label-free beacon detection (illustrative, modeled on the simulated slow-APT '
                    'campaign). Left: the attacker\'s call-outs are evenly spaced; a legitimate service is bursty. '
                    'Right: the rhythm-irregularity measure (coefficient of variation of the gaps) — below the flag '
                    'line means machine-like, the fingerprint of a covert beacon. The detector reads the rhythm, not '
                    'the address, so it generalizes to infrastructure it has never seen.')

H2('9.4  Zero false positives, and what that costs')
para('At the operating point that matches all four attackers, no normal user was flagged. This is the direct '
     'payoff of the named-profile design: because a flag requires matching a specific known-bad fingerprint — and, '
     'for cohort-relative profiles, a jointly elevated combination against peers — ordinary high-activity users do '
     'not trip it. Two caveats apply. First, a fingerprint library only catches techniques it describes; the '
     'discovery layer (Section 6.3) is the net for novel attacks, and it is noisier by design. Second, these are '
     'synthetic results, and a real environment will produce some false positives that must be measured and tuned. '
     'The relevant point is the structure: precise, named, peer-relative matching is what drives the false-positive '
     'count down without sacrificing detection.')

# ================= 10. ROADMAP =================
H1('10.  Phased Engagement — Validate on Real Telemetry')
table(['Phase', 'Duration', 'What happens', 'Exit criterion'],
      [('1.  Baseline', '4–6 weeks', 'Ingest five standard log sources; build per-entity profiles and cohorts.', 'Stable baselines and cohorts established.'),
       ('2.  Calibrate', '2–4 weeks', 'Tune named-profile thresholds to the environment; measure the real false-positive rate.', 'False positives within agreed budget.'),
       ('3.  Validate', '2–4 weeks', 'Red-team or historical-incident replay to confirm detection on this environment\'s data.', 'Known-bad activity matched and named.'),
       ('4.  Production', 'Ongoing', 'Feed ranked, explained leads into the existing SIEM/SOAR workflow.', 'Analyst-accepted leads in steady state.')],
      widths=[1.3, 1.0, 3.0, 1.7])

# ================= 11. DEPLOYMENT =================
H1('11.  Deployment in Practice')
para('V-Intelligence UEBA is designed to complement, not replace, the existing stack. The recommended operating '
     'model layers the fast magnitude-based tools as a familiar pre-filter, the threat-profile detector as the '
     'precision engine producing the named ranked list, and the semantic layer as the discovery net and the source '
     'of investigation context.')
bullet('Five standard log sources most enterprises already collect (sign-ins, file access, endpoint telemetry, '
       'network flow, and domain lookups), plus data-sensitivity labels where available. No new sensors required.',
       lead='Data requirements. ')
bullet('Containerized, standard infrastructure — no specialized hardware — suitable for enterprise data centers, '
       'cloud, or disconnected secure environments. Profiles update incrementally, so steady-state cost is small.',
       lead='Footprint. ')
bullet('Alerts feed existing SIEM/SOAR (Security Orchestration, Automation and Response) workflows with the matched '
       'profile and ATT&CK technique attached; behavioral risk can feed Zero Trust continuous-verification policies, '
       'the federal architecture mandate that no user or device is trusted by default.', lead='Integration. ')
bullet('Validation used 250 users; larger populations improve the statistics, because peer baselines and cohort '
       'comparisons get richer.', lead='Scale. ')

# ================= 12. NOVELTY =================
H1('12.  Why This Is New')
para('Commercial UEBA products exist, and the better ones perform per-entity baselining with statistical or '
     'machine-learning anomaly scores. What they produce is magnitude — a number saying behavior is different — '
     'without the direction that distinguishes insider exfiltration from benign change, and without a named, '
     'peer-relative profile that an analyst can act on directly. The individual techniques used here (serializing '
     'structured data to text for embedding, anomaly detection in embedding space, cohort-relative scoring, '
     'label-free beacon and domain detection, novelty tracking) each exist separately in the literature.')
para('What is not offered, commercially or in any single published system we have found, is the composition: every '
     'entity maintained as a behavioral profile in one shared semantic space; a precision layer of named, '
     'measurable known-bad fingerprints scored relative to role-group peers and read label-free from raw events; a '
     'discovery layer that accumulates slow drift and recurring novelty and names the direction against MITRE '
     'ATT&CK; and the two layers fused into one short, explained, ranked list. The innovation is the integrated '
     'architecture, not any single technique.', bold=True)
para('Scoping notes: vendor internals are not fully public (the comparison is against marketed capability), and a '
     'formal freedom-to-operate review is recommended before any contractual assertion of uniqueness.', 9.5, GREY,
     italic=True)

# ================= 13. LIMITATIONS =================
H1('13.  Limitations — Stated Plainly')
bullet('All results are from synthetic telemetry calibrated to public CISA (Cybersecurity and Infrastructure '
       'Security Agency) advisories. Real enterprise behavior is messier; every figure must be re-measured on real '
       'data, including the zero-false-positive result.')
bullet('A fingerprint library only catches techniques it describes. Novel attacks fall to the discovery layer, which '
       'is more sensitive and therefore noisier — its operating threshold is a real analyst-cost trade-off.')
bullet('Named-profile thresholds were validated knowing the answers; in production they must be set blind and tuned '
       'per environment.')
bullet('Baselines need 4–6 weeks of history before detection is meaningful, and legitimate behavioral shifts — role '
       'changes, project transitions — are the main source of false flags.')
bullet('Data-sensitivity labels on file activity materially strengthen the insider profile; environments without '
       'classification metadata lose part of that signal.')
bullet('The semantic layer relies on an embedding model; fully disconnected deployments require a locally hosted '
       'replacement, which must be re-validated.')

# ================= 14. CONCLUSION =================
H1('14.  Conclusion')
para('The threats that matter most — insiders, nation-state intrusions, living-off-the-land campaigns, and '
     'infrastructure compromises — share one property: they operate inside authorized access, beneath every '
     'signature and threshold, and reveal themselves only as a gradual change of behavioral direction over months. '
     'The documented record is unambiguous: the victims of Volt Typhoon and Salt Typhoon did not detect these '
     'campaigns; outside agencies did, years and months after the fact.')
para('V-Intelligence UEBA reads the same enterprise logs as a living behavioral profile for every entity, and runs '
     'two complementary detection layers over it: a precision layer of named, measurable known-bad fingerprints '
     'scored against role-group peers, and a discovery layer of semantic drift and novelty that names the direction '
     'against known adversary behavior. On a 250-user, 485-day validation with four embedded campaigns, it matched '
     '4 of 4 attacks to named attack fingerprints at zero false positives, against 1 of 4 with no explanation for '
     'the best traditional method.', bold=True)
para('An insider changes what they access, not how much; an APT changes where it communicates, not how often. '
     'Traditional tools measure magnitude. V-Intelligence UEBA measures direction, peer-relative deviation, and '
     'named-profile match — the dimensions in which these attacks are actually visible.', bold=True)

# ================= GLOSSARY =================
H1('Appendix — Glossary')
gl = [
    ('UEBA', 'User and Entity Behavior Analytics — detecting threats by learning each entity\'s normal behavior and flagging departures, rather than matching known-bad patterns.'),
    ('Digital entity', 'A living software profile of a user or system — like a patient\'s medical chart for behavior: history, vitals, context, and trend, updated weekly.'),
    ('Threat-profile detector', 'The precision detection layer: a library of named, measurable known-bad fingerprints; a flag means an entity matched a specific attack profile, not merely that it looked unusual.'),
    ('Cohort-relative', 'Scored against the entity\'s role-group peers (developer vs developers, etc.), so "normal for the role" is not mistaken for an anomaly.'),
    ('C2 beacon', 'A covert "call-home" channel an implanted attacker tool uses for instructions; detected here label-free by its robotic, evenly spaced rhythm.'),
    ('DGA (domain-generation algorithm)', 'Malware that invents many random-looking domain names to find its controller; detected by high name-randomness and shared back-end infrastructure.'),
    ('Semantic space / embedding', 'A numerical "map of meaning" produced by the representation that powers large language models; behaviors that mean similar things sit near each other.'),
    ('Behavioral zone', 'One of five dimensions of an entity\'s profile: identity, access pattern, data behavior, network footprint, risk posture.'),
    ('Drift', 'The movement of an entity\'s behavioral profile over time — how fast, how consistently, and in which direction it is trending.'),
    ('Novelty persistence', 'A never-before-seen behavior (e.g., a new external contact) that keeps recurring week after week — the fingerprint of covert C2 infrastructure.'),
    ('MITRE ATT&CK', 'The industry-standard public catalog of adversary tactics and techniques, used as a common language for describing attacker behavior.'),
    ('SIEM / SOC / SOAR', 'Security Information and Event Management (the central log/rule platform); Security Operations Center (the analyst team); Security Orchestration, Automation and Response (incident-response automation).'),
    ('IOC', 'Indicator of Compromise — a known-bad artifact (file fingerprint, address, domain) used by signature-based tools.'),
    ('APT / LOTL', 'Advanced Persistent Threat (a patient, well-resourced, usually nation-state intrusion); Living off the land (attacking with only the legitimate admin tools already present — nothing to fingerprint).'),
    ('Z-Score / Isolation Forest / Local Outlier Factor / One-Class SVM', 'The four industry-standard statistical and machine-learning anomaly detectors used as the traditional comparison baseline.'),
    ('Zero Trust', 'The federal security architecture mandate that no user or device is trusted by default; every access is continuously verified.'),
    ('False-positive rate', 'The share of normal users incorrectly flagged; the threat-profile layer flagged none of the 246 normal users on this synthetic set.'),
    ('Synthetic data', 'Simulated telemetry built to prove the concept — calibrated against public CISA advisories — before access to real operational data.'),
]
gt = doc.add_table(rows=len(gl) + 1, cols=2); gt.style = 'Table Grid'
setcell(gt.rows[0].cells[0], 'Term', bold=True, color=WHITE, fill='0B1F3A')
setcell(gt.rows[0].cells[1], 'Meaning', bold=True, color=WHITE, fill='0B1F3A')
for i, (t, d) in enumerate(gl, start=1):
    setcell(gt.rows[i].cells[0], t, bold=True, size=9, color=NAVY, fill='F4F6F8')
    setcell(gt.rows[i].cells[1], d, size=9)
for row in gt.rows:
    row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(5.0)

para('')
para('Business Edition (Enhanced) — architecture-level disclosure. Named-profile definitions, formulas, thresholds, '
     'weights, and parameters are proprietary and available in the full technical edition under NDA. All results are '
     'synthetic and must be re-validated on real operational telemetry. Do not redistribute.', 9, GREY, italic=True)

out = OUT; n = 2
while True:
    try:
        doc.save(out); break
    except PermissionError:
        root, ext = os.path.splitext(OUT); out = f"{root}_v{n}{ext}"; n += 1
print('Saved:', out)
