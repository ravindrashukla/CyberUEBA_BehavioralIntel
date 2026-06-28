# -*- coding: utf-8 -*-
"""Build the four session gap documents for the Army cyber working session."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
NAVY=RGBColor(0x1f,0x38,0x64); BLUE=RGBColor(0x2e,0x75,0xb6); GRAY=RGBColor(0x55,0x60,0x6b)
RED=RGBColor(0xa6,0x1b,0x1b); GREEN=RGBColor(0x1f,0x7a,0x33); AMBER=RGBColor(0xb5,0x57,0x0a)
BASE="WP DLA/Presentation"

def newdoc():
    d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(10.5); return d
def H1(d,t):
    p=d.add_heading(t,1)
    for r in p.runs: r.font.color.rgb=NAVY
    return p
def H2(d,t):
    p=d.add_heading(t,2)
    for r in p.runs: r.font.color.rgb=BLUE
    return p
def P(d,t,**k):
    p=d.add_paragraph(); r=p.add_run(t)
    r.italic=k.get('i',False); r.bold=k.get('b',False)
    if k.get('c'): r.font.color.rgb=k['c']
    if k.get('s'): r.font.size=Pt(k['s'])
    return p
def B(d,label,rest="",sub=False):
    p=d.add_paragraph(style='List Bullet 2' if sub else 'List Bullet')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def NUM(d,label,rest=""):
    p=d.add_paragraph(style='List Number')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def code(d,t):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2)
    r=p.add_run(t); r.font.name='Consolas'; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0x33,0x33,0x33)
    return p
def callout(d,t,c=RED,s=9.5):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2)
    r=p.add_run(t); r.bold=True; r.font.color.rgb=c; r.font.size=Pt(s)
    return p
def title(d,t,sub,banner):
    tp=d.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=tp.add_run(t); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=NAVY
    sp=d.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=sp.add_run(sub); r.font.size=Pt(11.5); r.font.color.rgb=BLUE
    wn=d.add_paragraph(); wn.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=wn.add_run(banner); r.bold=True; r.font.color.rgb=RED; r.font.size=Pt(9.5)
def table(d,headers,rows,widths=None,fs=9):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(fs)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""; cells[i].paragraphs[0].add_run(v).font.size=Pt(fs)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    d.add_paragraph()

# ============ 1) DEMO RUN CARD ============
d=newdoc()
title(d,"Layer 2 Live Demo - Run Card","Behavioral Digital Twin  |  for the Twin lead (Block 1, 60 min)","INTERNAL. Hold this during the live demo.")
H2(d,"Before the room: open and tested")
code(d,"data/tier3_results/composite_scores.csv  |  weekly_zone_trajectories.csv\n"
      "detection/ (cusum, drift_direction, reference_concepts, composite_scorer)\n"
      "comparison/run_comparison.py  |  UI up on host 8001  |  backup recording ready")
H2(d,"Run order")
NUM(d,"Setup: ","four nation-state attackers hidden among 246 normal users; 250 entities, 485 days, ~14M events, full ground truth, blind and gradeable.")
NUM(d,"Show the four composites: ","Salt USR-118 #1 (51.27), Insider USR-156 #2 (46.24), Slow APT USR-234 #7 (19.44), Volt USR-042 #24 (13.70).")
NUM(d,"Re-derive live from the formula: ","composite = signal_strength + 0.5 breadth + 0.3 sustained + 0.5 ctx_spread + 0.3 ctx_max + novelty.")
code(d,"Volt:  6.40 + 5.50 + 1.14 + 0.23 + 0.43 + 0.00 = 13.70\n"
      "Salt: 29.89 + 9.00 + 2.87 + 4.32 + 2.38 + 2.80 = 51.27")
NUM(d,"Derive the 8.1% live: ","sort all 250 by composite. Volt (13.70) is the stealthiest scorer, so the line must reach 13.70 to catch all four. 20 of 246 normals score above it: 20/246 = 8.1% false positive, 4 of 4 caught.")
NUM(d,"Baselines on the same data: ","Isolation Forest, One-Class SVM, LOF each 0 of 4; z-score 1 of 4.")
NUM(d,"The two data points that carry the room: ","Slow APT is invisible to thresholds (signal 4.47, breadth 1) and caught by novelty alone (a novel C2 IP every week for 60 weeks). Volt is caught by breadth of small signals, the living-off-the-land signature.")
callout(d,"Say the hard truth FIRST: the top normal user (21.5) outscores Slow APT (19.4) and Volt (13.7). That is exactly why we deliver a ranked list, not a threshold. The honest claim is recall at the top of the ranking.",GREEN)
H2(d,"If the demo stalls")
B(d,"Recite the numbers ","from this card and switch to the backup recording. Do not improvise numbers.")
B(d,"Honesty line: ","detection core is live; identity fusion, forensics, response are roadmap.")
d.save(f"{BASE}/DemoRunCard_Layer2_Twin_INTERNAL.docx")

# ============ 2) LOGISTICS / AV CHECKLIST ============
d=newdoc()
title(d,"In-Person Session - Logistics and AV Checklist","Army Cyber Working Session  |  half-day, ~4.5 hours","INTERNAL. Run this list before the customer arrives.")
H2(d,"The #1 risk: network")
callout(d,"Assume the room has restricted or NO internet. Every live demo must run offline, over a presenter's own cellular hotspot, or from a full backup recording. Do NOT assume cloud access.")
B(d,"Backup recordings ","are mandatory for every live layer (Twin / Layer 1 / Layer 3).")
H2(d,"Hardware and AV")
B(d,"Bring-your-own: ","each presenter brings their own tested demo laptop plus HDMI and USB-C adapters.")
B(d,"One shared display ","with a clean driver-switch protocol; the session lead keeps the 60-minute boxes.")
B(d,"Arrive 60 to 90 minutes early ","for an AV check and a full demo dry-run on the actual room display.")
H2(d,"Per-presenter readiness")
table(d,["Presenter","Ready check","Done"],
 [["Twin lead (Layer 2)","demo files open, 8.1% derivation rehearsed, backup recording","[ ]"],
  ["Layer 1 lead","preemptive demo offline-ready, proof points loaded, backup recording","[ ]"],
  ["Layer 3 lead","API/Data/agentic walkthrough ready (architecture, no live claim)","[ ]"],
  ["Session lead","open, transitions, integration, the ask; timekeeping","[ ]"],
  ["Scribe","captures questions and pilot-scope asks to a parking lot","[ ]"]],
 widths=[2.2,3.6,0.7],fs=9)
H2(d,"Discipline")
B(d,"One company, one platform. ","Present every layer as V-Intelligence. No partner names.")
B(d,"Label roadmap out loud. ","This room rewards candor and punishes overclaim.")
B(d,"Route integration, provenance, and contract questions ","to the session lead.")
d.save(f"{BASE}/InPerson_Logistics_AV_Checklist_INTERNAL.docx")

# ============ 3) MASTER INDEX ============
d=newdoc()
title(d,"Session Materials - Master Index","Army Cyber Working Session (Esquibel / DiGrezio / Herrmann)","INTERNAL. The single index the team works from.")
H2(d,"Present / show")
table(d,["Artifact","Purpose","Used in"],
 [["Pentagon_Working_Session_Deck_v2.pptx","The deck (26 slides incl. detailed Digital Twin section)","All blocks"],
  ["cyber_whitepaper_figures.html","Figures: architecture, preemptive flow, twin pipeline, Typhoon-across-layers","Open / per layer"]],
 widths=[4.3,3.5,1.4],fs=8.5)
H2(d,"Run the room (internal)")
table(d,["Artifact","Purpose","Used in"],
 [["Working_Session_RunOfShow_FourLayer_22CT_INTERNAL.docx","Master playbook: agenda, talking tracks, why-it-matters, numbers","All blocks"],
  ["Presenter_Briefing_FourLayer_Session_INTERNAL.docx","Per-presenter brief","Prep"],
  ["DemoRunCard_Layer2_Twin_INTERNAL.docx","Live-demo derivation card (8.1%, four attackers)","Block 1"],
  ["InPerson_Logistics_AV_Checklist_INTERNAL.docx","Go-checklist: network, AV, hardware, roles","Pre-session"],
  ["Layer1_Preemptive_TalkingTrack_INTERNAL.docx","L1 lead talking track + proof points","Block 2"]],
 widths=[4.3,3.5,1.4],fs=8.5)
H2(d,"Technical depth (for hard Q&A)")
table(d,["Artifact","Purpose","Used in"],
 [["Behavioral_Digital_Twin_DeepDive_DataPoints_INTERNAL_v2.docx","Twin anatomy with real data points (Esquibel)","Block 1 Q&A"],
  ["Digital_Twin_Design_and_Build_Plan_INTERNAL.docx","Design, gaps, build plan (DiGrezio fielding)","Block 1/5"],
  ["Cyber_DeepDive_Briefing_Book_*.docx","Slide-by-slide internal prep","Prep"],
  ["Tough_Questions_Cheat_Sheet_*.docx","Q&A answers","All blocks"]],
 widths=[4.3,3.5,1.4],fs=8.5)
H2(d,"Leave-behinds (shareable)")
table(d,["Artifact","Purpose","Used in"],
 [["V-Intelligence_Layered_Cybersecurity_Whitepaper_US_Army_Pentagon.docx","Flagship 5-layer whitepaper, Army-framed","Leave-behind"],
  ["V-Intelligence_UEBA_Technical_Leave-Behind.docx","One-page shareable summary","Leave-behind"]],
 widths=[4.3,3.5,1.4],fs=8.5)
callout(d,"Numbers locked across all materials: 4 of 4 at 8.1% FP (baselines 0 of 4, z-score 1 of 4); ranks Salt #1 / Insider #2 / Slow APT #7 / Volt #24; 250 entities, ~14M events. No Rigor/Cequence names (one company).",NAVY)
d.save(f"{BASE}/Session_Materials_Master_Index_INTERNAL.docx")

# ============ 4) LAYER 1 TALKING TRACK ============
d=newdoc()
title(d,"Layer 1 - Preemptive Network Assurance","Talking track for the L1 lead (Block 2, 60 min)","INTERNAL. Theory, demo, results - the same arc as every layer.")
H2(d,"The mission")
P(d,"Prove every known attack path through the network controls is closed before traffic flows. This is the prevention half of the gap: Layer 2 catches the adversary already inside; Layer 1 proves the edge was provably shut.")
H2(d,"What incumbents miss")
B(d,"Sampling, not proof: ","breach-simulation and policy tools sample an astronomically large configuration state space; they never prove segmentation is correct.")
B(d,"Misconfiguration dominates: ","99 percent of firewall breaches are misconfiguration, not product flaws; rule sets drift within hours of a change.")
B(d,"Reactive lag: ","signature and rule tools cannot pre-block a path not yet seen, while disclosed vulnerabilities grow toward an order of magnitude more.")
H2(d,"The approach")
B(d,"Attack intelligence: ","continuously ingested threat feeds compiled into MITRE ATT&CK-enriched attack graphs.")
B(d,"Defense intelligence: ","an exhaustive formal, symbolic model of the firewall, identity, IPS, SASE, and WAF estate, reasoned over completely, with no sampling.")
B(d,"Remediation intelligence: ","a guardrailed reasoner prescribes precise, risk-prioritized fixes that introduce no new errors, reviewed by humans.")
B(d,"Operational fit: ","agentless and configuration-only; uses residual capacity in the estate to block attacks; first actionable findings in under 15 days; drift caught within the hour.")
H2(d,"Proof points (verified findings)")
B(d,"Credential and IP-theft path closed ","from a shadowed, incomplete rule set allowing an unsanctioned app from internal subnets.")
B(d,"Lateral-movement path closed: ","three higher-priority rules formed a group shadow letting legacy SMBv1 pass without inspection.")
B(d,"External attack surface formally assured ","for a national cyber-defense agency, beyond what EASM can simulate.")
B(d,"Regulatory compliance verified ","against a live national malicious-IP and domain feed; infosec compliance gaps diagnosed.")
H2(d,"Federal fit and the line")
B(d,"Deployable ","across IL5, IL6, and JWICS enclaves; relevant to ARCYBER continuous-monitoring operations.")
P(d,"The line: \"We do not sample the door and hope; we prove it is shut, before traffic flows, and we catch the drift within the hour.\"",i=True)
d.save(f"{BASE}/Layer1_Preemptive_TalkingTrack_INTERNAL.docx")

print("built: DemoRunCard, Logistics_AV_Checklist, Master_Index, Layer1_TalkingTrack")
