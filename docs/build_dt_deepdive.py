# -*- coding: utf-8 -*-
"""The Behavioral Digital Twin - Technical Deep-Dive, with worked examples and
REAL data points pulled from the code and result files. INTERNAL (contains the
scoring formula / dimensions). No em-dashes.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6); GRAY=RGBColor(0x55,0x5F,0x6B)
RED=RGBColor(0xA6,0x1B,0x1B); GREEN=RGBColor(0x1F,0x7A,0x33); MONO=RGBColor(0x33,0x33,0x33)
d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(10.5)

def H1(t):
    p=d.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
    return p
def H2(t):
    p=d.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=BLUE
    return p
def P(t,italic=False,bold=False,color=None,size=None):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold
    if color: r.font.color.rgb=color
    if size: r.font.size=Pt(size)
    return p
def B(label,rest="",sub=False):
    p=d.add_paragraph(style='List Bullet 2' if sub else 'List Bullet')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def NUM(label,rest=""):
    p=d.add_paragraph(style='List Number')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def code(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2)
    r=p.add_run(text); r.font.name='Consolas'; r.font.size=Pt(9.5); r.font.color.rgb=MONO
    return p
def callout(text,color=GREEN,size=9.5):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2)
    r=p.add_run(text); r.bold=True; r.font.color.rgb=color; r.font.size=Pt(size); return p
def table(headers,rows,widths=None,fs=8.5,hdr_fs=8.5):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(hdr_fs)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""
            for j,part in enumerate(v.split("||")):
                para=cells[i].paragraphs[0] if j==0 else cells[i].add_paragraph()
                para.add_run(part).font.size=Pt(fs)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    d.add_paragraph(); return t

# ===== TITLE =====
tp=d.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("The Behavioral Digital Twin"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
sp=d.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Technical Deep-Dive with Worked Examples and Real Data Points"); r.font.size=Pt(12.5); r.font.color.rgb=BLUE
sp2=d.add_paragraph(); sp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp2.add_run("22nd Century  |  V-Intelligence Behavioral Entity Intelligence"); r.font.size=Pt(10); r.italic=True; r.font.color.rgb=GRAY
wn=d.add_paragraph(); wn.alignment=WD_ALIGN_PARAGRAPH.CENTER
rw=wn.add_run("INTERNAL. Contains proprietary scoring logic and dimensions. Not for distribution."); rw.bold=True; rw.font.color.rgb=RED; rw.font.size=Pt(9.5)
P("Every number in this document is taken directly from the code and the result files (composite_scores.csv, "
  "novelty_metrics.csv, weekly_zone_trajectories.csv, entity_structures.json). Each composite re-derives from the "
  "formula in Part 6.",italic=True,size=9.5,color=GRAY)

# ===== 0 WHAT IT IS =====
H1("0. What a digital twin is here")
P("It is not a database record. For each of 250 entities the twin is a living behavioral representation, rebuilt every "
  "week and tracked as a trajectory over 70 weeks. Every step below exists for a reason:")
code("raw logs -> 5 zone prose descriptions -> 5 embeddings (1536-d each)\n"
     "        -> 1 attention-weighted composite vector -> temporal trajectory of composites\n"
     "        -> drift (magnitude + direction) -> CUSUM -> 5-phase composite score -> ranked verdict")

H2("Step 1 - Raw logs to five zone prose descriptions")
P("The entity's behavior is aggregated into 23 numeric features and grouped into four action zones (access_pattern, "
  "data_behavior, network_footprint, risk_posture); a fifth zone, identity, carries profile context (role, department, "
  "clearance, type), not measured behavior. Within each zone, its attributes are combined into one plain-language "
  "sentence, so the meaning of every feature in that zone is carried into one embedding. Prose first forces meaning over "
  "magnitude: context such as restricted-vs-public or off-hours survives, where a raw number would lose it. The four "
  "action zones are the behavioral surfaces an intruder must cross; identity anchors what normal looks like. All are "
  "derivable from logs every organization already collects:")
table(["Zone","What it captures","What it is built to catch"],
 [["identity","role, department, clearance, tenure, type","the anchor: what normal should look like for THIS entity"],
  ["access_pattern","auth volume, fail rate, off-hours, unique sources/dests, methods","credential abuse, lateral movement, off-hours access"],
  ["data_behavior","file count, restricted/confidential ratios, write ratio, unique paths, bytes","data collection and exfiltration staging"],
  ["network_footprint","bytes out, unique destinations, external ratio, DNS domains, NXDOMAIN","C2 beaconing, tunneling, exfiltration"],
  ["risk_posture","endpoint events, suspicious ratio, max/mean risk, unique processes","living-off-the-land tooling, malware, LOLBins"]],
 widths=[1.3,2.6,2.5],fs=8)
P("Three of those terms, expanded:",bold=True)
B("Exfiltration (data_behavior + network_footprint): ","the unauthorized movement of data out of the organization: "
  "copying restricted or confidential files to a staging area, compressing them, and sending them to an external "
  "destination (cloud storage, a personal email account, or an attacker server). It shows as rising restricted and "
  "confidential file access and elevated write (staging) activity in data_behavior, and as asymmetric, outbound-heavy "
  "volume to new external destinations in network_footprint.")
B("C2 beaconing (network_footprint): ","a compromised host quietly phoning home to command-and-control infrastructure "
  "at regular intervals to receive instructions. It shows as repeated, often low-volume, periodic connections to a small "
  "set of external or newly-seen destinations and domains. The network_footprint zone plus novelty persistence (Part 7), "
  "a novel IP recurring week after week, are built to catch exactly this.")
B("LOLBins / risk_posture: ","LOLBins (living-off-the-land binaries) are legitimate, signed operating-system tools "
  "(PowerShell, certutil, wmic, netsh, schtasks) repurposed for malicious actions, so no malware file is dropped and "
  "signature tools see nothing. The risk_posture zone tracks endpoint process behavior, the suspicious-process ratio, "
  "the peak and mean risk score, and the count of unique processes, so a spike in new or risky processes flags tooling "
  "or malware even when the binaries themselves are trusted.")
P("These zones are not fixed: each can be expanded, and new behavioral zones added, to capture additional insight as "
  "needed; nothing downstream changes.",italic=True)

H2("Step 2 - Five zone prose to five embeddings (1536-d each)")
P("Each zone's combined sentence is mapped into a shared 1536-dimensional semantic space. Because all of a zone's "
  "attributes were folded into that one sentence, the meaning of every feature in the zone is carried into that single "
  "zone embedding. Significance: two behaviors with identical counts but different intent separate, and we can later "
  "measure direction against named threat concepts.")

H2("Step 3 - Five embeddings to one attention-weighted composite vector")
P("This is where the five zone vectors are aggregated into one, once per week. The weighting is attention-based: each "
  "zone's contribution is its magnitude times a context weight, through a softmax. By default the weights are uniform "
  "(each zone ~0.2 in normal operations); investigation contexts re-weight them (an insider investigation emphasizes "
  "data_behavior at 0.4; an APT hunt emphasizes network_footprint and risk_posture). Significance: one behavioral "
  "representation in the same 1536-d space as the threat concepts, with the ability to re-weight which zones matter. "
  "This is the weighted aggregate of the five zone vectors.")

H2("Step 4 - Composite to temporal trajectory")
P("The weekly composite vectors, in sequence, form the trajectory. Significance: the twin is a path, not a snapshot; "
  "slow campaigns only reveal themselves over time.")

H2("Step 5 - Trajectory to drift (magnitude and direction)")
P("Magnitude says how much the entity changed week to week; direction says toward what, meaning which MITRE-mapped "
  "threat concept. Direction is the meaning layer that replaces '87% anomalous' (Part 5).")

H2("Step 6 - Drift to CUSUM")
P("CUSUM accumulates slow, sub-threshold drift so the years-long dwell of an APT triggers even though no single week "
  "crosses a line (Part 4).")

H2("Step 7 - Five-phase composite score to ranked verdict")
P("Everything collapses into one comparable, explainable, ranked number: a risk-ranked list, not a flood of threshold "
  "alarms (Part 6).")

B("Test design: ","four attackers hidden among 246 normal users; 250 entities, 485 days (70 weekly snapshots), "
  "~14M events, seven raw log streams feeding five zones; full ground truth, blind-gradeable.")

# ===== DATA PIPELINE =====
H1("Data Pipeline - How the Data Is Generated and Transformed")
P("This is what happens before the twin exists: how raw logs become the structured record the twin is built from. The "
  "same transformation runs on real customer logs; only the source of the raw events changes.")

H2("A. Data generation (the synthetic test bed)")
P("The simulator builds a synthetic enterprise and writes raw event logs in the same shape a real SIEM or log pipeline "
  "produces. Scale (simulator/config.py): 250 users, 400 devices, 25 network segments, 60 applications, 30 roles; 485 "
  "days (2025-01-01 to 2026-05-01).")
P("Seven raw event streams, one file per day per source, which consolidate into the five behavioral zones:",bold=True)
table(["Raw stream","Rate","Approx per day"],
 [["auth (sign-in)","8 / user / day","2,000"],
  ["network flows","40 / device / day","16,000"],
  ["dns","10 / device / day","4,000"],
  ["endpoint","8 / device / day","3,200"],
  ["file_access","10 / user / day","2,500"],
  ["app","5 / user / day","1,250"],
  ["privilege","~5 / day","5"]],widths=[1.7,1.9,1.5],fs=9)
B("Total volume: ","about 29,000 events/day x 485 days = ~14 million events (this is where the ~14M figure comes from).")
B("Ground truth: ","normal entities follow role-group baselines; four attackers are injected with MITRE-mapped, phased "
  "behavior (Volt LOTL, Salt telecom pivot, Insider, Slow APT). Full ground truth is retained for blind grading.")
B("For a real deployment: ","these same seven streams come from the customer's own logs. Nothing about the "
  "transformation below changes; only the source of the raw events does.")

H2("B. Transformation: raw logs -> one structured JSON record per entity per window")
P("Raw per-event logs are aggregated for each entity over a time window (weekly for the trajectory). The aggregation "
  "produces 23 numeric behavioral features per entity per week, combined into a single structured JSON record:")
code("{ entity_id, entity_type,\n"
     "  profile { user_id, user_type, department, role, clearance, primary_device_id },\n"
     "  is_attack, week_idx,\n"
     "  raw_features { ...23 aggregated metrics... },\n"
     "  zone_features { identity, access_pattern, data_behavior, network_footprint, risk_posture },\n"
     "  zone_serialized_text, zone_embedding_dims, phase_state, context_weights }")
P("The 23 numeric features map to four action zones (access_pattern, data_behavior, network_footprint, risk_posture); "
  "a fifth zone, identity, carries profile context (department, role, clearance, type) rather than measured behavior:",bold=True)
table(["Zone","Aggregated features"],
 [["identity","department, role, clearance, user_type"],
  ["access_pattern","auth_total, auth_failed, auth_fail_rate, auth_unique_sources, auth_unique_dests, auth_off_hours_ratio, auth_methods_used"],
  ["data_behavior","file_total, file_unique_paths, file_restricted_ratio, file_confidential_ratio, file_write_ratio, file_total_bytes"],
  ["network_footprint","net_bytes_out, net_unique_dsts, net_external_ratio, dns_unique_domains, dns_nxdomain_ratio"],
  ["risk_posture","endpoint_total, endpoint_suspicious_ratio, endpoint_max_risk, endpoint_mean_risk, endpoint_unique_processes"]],
 widths=[1.4,5.0],fs=8)
P("This JSON is the single combined record of who the entity is and how it behaved that week, across every log source. "
  "It is the bridge between raw logs and the embedding.",italic=True)

H2("C. Transformation: JSON record -> behavioral embedding (the twin)")
B("Serialize: ","each zone/signal of the JSON record is turned into a natural-language description "
  "(zone_serialized_text). Meaning over magnitude (Part 1).")
B("Embed: ","each signal description is embedded into a 1536-d vector (Part 1).")
B("Compose: ","the five zone vectors are combined into one composite vector by attention-weighted average and "
  "L2-normalized (Part 2). That composite IS the entity's behavioral digital twin for that week.")
B("Over time: ","one composite per week forms the trajectory, which feeds drift, CUSUM, drift-direction, and the "
  "composite score (Parts 4 to 9).")
code("raw logs -> aggregate per entity/week -> structured JSON (23 features -> 4 action zones + identity)\n"
     "        -> serialized text -> 5 embeddings (1536-d) -> 1 composite twin vector -> trajectory")

H2("D. Future enhancement - additional logs and data we can capture")
P("The current twin uses 23 features across five zones, all derivable from logs most organizations already collect. "
  "The architecture is extensible: additional log sources can enrich existing zones or add new zones, with no change to "
  "the embedding, composition, drift, or scoring downstream. Candidate enhancements:")
table(["Additional log / data source","What it adds","Zone (new or enriched)"],
 [["IdP / IAM (Okta, Entra ID), MFA, conditional access","sign-in risk, group and role changes, MFA-fatigue patterns","identity, access_pattern"],
  ["Privilege / PAM / sudo, admin actions","privilege escalation, just-in-time grants, admin command use","new: privilege"],
  ["Email + collaboration (mail metadata, chat, file sharing)","external communication, data sharing, exfil over email","new: communication"],
  ["Web proxy / Secure Web Gateway / CASB","uploads, risky categories, sanctioned vs shadow cloud apps","network_footprint, data_behavior"],
  ["Cloud / SaaS audit (CloudTrail, Azure Activity) + API gateway","cloud control-plane actions, API behavior, service-account use","new: cloud / API"],
  ["EDR detail (process trees, command line, script-block, registry)","LOLBin command lines, persistence, parent-child process","risk_posture (enriched)"],
  ["NetFlow / IPFIX, TLS fingerprints (JA3), beacon timing","C2 cadence, encrypted-channel fingerprints, flow durations","network_footprint (enriched)"],
  ["DNS detail (subdomain entropy / DGA, TXT and MX volume)","DNS tunneling, domain-generation-algorithm probing","network_footprint (enriched)"],
  ["DLP + removable media + print logs","data-loss events, USB exfiltration, bulk printing","data_behavior (enriched)"],
  ["Geolocation / VPN / device posture","impossible travel, unmanaged or unpatched device","new: location / device"],
  ["Peer-group / organizational graph","deviation from peers, not just from self","new: relationship"],
  ["Temporal cadence (time-of-day, day-of-week)","rhythm and periodicity anomalies","new: temporal"],
  ["HR / asset context (tenure, role change, asset criticality)","insider-risk context, asset-weighted scoring","identity (enriched)"]],
 widths=[2.5,2.4,1.5],fs=7.5)
P("Each addition is just another prose-to-embedding zone, or extra attributes in an existing zone's sentence. The "
  "five-zone model is the proven core, not a ceiling.",italic=True)

# ===== PART 1 =====
H1("Part 1. The five zones, serialized (the twin's senses)")
P("Each zone bundles several raw attributes into ONE natural-language sentence, and that sentence is embedded into 1536 "
  "dimensions. The bundling is the point: all of a zone's features carry their meaning into a single zone embedding.")
B("The key design decision: ","raw metrics become natural language before embedding. This is meaning over magnitude. "
  "Two users who both touched 47 files embed differently if the prose context differs (restricted vs public, off-hours "
  "vs business hours). A scalar feature cannot see that; the embedding can.")
P("Real zone serializations (verbatim from entity_structures.json, USR-156):",bold=True)
code("identity:          User USR-156 identity: role=Analyst, department=Marketing,\n"
     "                   clearance=internal, tenure_days=150, type=employee\n"
     "access_pattern:    User USR-156 access: auth_events=16, fail_rate=0.0625,\n"
     "                   off_hours=0.1250, unique_sources=3, unique_dests=5, methods=3\n"
     "data_behavior:     User USR-156 data: file_accesses=10, restricted_ratio=0.0,\n"
     "                   confidential_ratio=0.2, write_ratio=0.2, unique_paths=10, bytes=124933027\n"
     "network_footprint: User USR-156 network: bytes_out=125407529, unique_dsts=22,\n"
     "                   external_ratio=0.1042, dns_domains=14, nxdomain_ratio=0.0\n"
     "risk_posture:      User USR-156 risk: endpoint_events=8, suspicious_ratio=0.0,\n"
     "                   max_risk=30.0, mean_risk=12.25, unique_processes=7")
P("Each of the five lines above becomes one 1536-d vector. Five zones, five vectors.",italic=True)

H2("Why fuse the whole zone, not score attributes one by one")
P("An attribute carries a number; the zone sentence carries meaning, because meaning lives in the relationships among "
  "the attributes. Two illustrations:")
P("Example 1 - same number, opposite meaning (data_behavior).",bold=True)
P("Take file_total_bytes = 250 MB in isolation. It is identical for two very different users. Fused with the rest of the zone:")
B("User A: ","\"250 MB across 180 unique public files; restricted_ratio=0.0, confidential_ratio=0.0, write_ratio=0.1\" "
  "-> routine bulk reading of reference material. Benign.")
B("User B: ","\"250 MB from 6 unique files; restricted_ratio=0.7, confidential_ratio=0.5, write_ratio=0.6\" "
  "-> collection of sensitive data with local staging. Looks like exfil staging.")
P("Same 250 MB. The single attribute cannot tell them apart, but the fused sentences carry opposite meaning, and the "
  "two embeddings land in different regions of the space (one near normal, one near data_exfiltration / insider_threat).")
P("Example 2 - individually-benign attributes that only mean something together (access_pattern).",bold=True)
P("Each attribute alone is unremarkable: off_hours=0.6 (perhaps a night-shift worker), unique_dests=8 (normal for many "
  "roles), fail_rate=0.0 (clean, no fumbling), methods=3 (fine). Fused into one sentence:")
code("\"8 systems accessed, 60% off-hours, 3 auth methods, zero failures\"")
P("That combination reads as confident, deliberate, off-hours movement across many systems with valid credentials and "
  "no errors: the signature of living-off-the-land lateral movement. No single attribute flags it; the fused sentence "
  "does, and its embedding sits near lateral_movement / LOTL.")
P("The principle: serialize the whole zone, then embed, to preserve the relationships among attributes (few files AND "
  "restricted AND high-write; or off-hours AND broad AND error-free). Per-attribute scoring throws those relationships away.",italic=True)

# ===== PART 2 =====
H1("Part 2. Composition (attention-weighted, not concatenation)")
B("","The five 1536-d zone vectors fuse into ONE 1536-d composite by attention-weighted average, then L2-normalize to a unit vector.")
B("Attention weighting: ","each zone's contribution is its magnitude times a context weight, passed through a softmax. "
  "In normal operations the weights are uniform (~0.2 each); investigation contexts re-weight (insider investigation "
  "emphasizes data_behavior; APT hunt emphasizes network_footprint and risk_posture). Divergence between an entity's "
  "composite under different contexts is itself a signal (the ctx_spread and ctx_max terms in Part 6).")
B("Why average, not concatenate: ","the composite stays in the SAME 1536-d space as the individual zones and the "
  "reference concepts. That is what lets us later measure direction by cosine similarity against threat concepts. "
  "Concatenation would break that geometry.")

# ===== PART 3 =====
H1("Part 3. The behavioral zones (decomposition for drift)")
P("The twin is decomposed into five behavioral zones, tracked independently in the trajectory:")
code("identity | access_pattern | data_behavior | network_footprint | risk_posture")
B("zone_divergence ","is the living-off-the-land tell: some zones drift while others stay frozen. A SOC operator whose "
  "network_footprint and access_pattern drift while identity stays flat is the signature we want.")

# ===== PART 4 =====
H1("Part 4. Trajectory and CUSUM (the time dimension)")
B("","Each week produces a fresh composite; the ordered sequence is the trajectory. Between consecutive weeks, "
  "drift_magnitude = 1 - cosine(week t, week t+1), a cosine distance in [0, 2].")
B("CUSUM ","accumulates excess drift: cusum = max(0, cusum + mag - 0.02), and fires when cusum >= 0.05 for >= 2 "
  "consecutive periods. A slow campaign drifting 0.02/week never trips a per-week threshold, but CUSUM sums it and fires.")
P("Live example. Volt (USR-042) trajectory, final 6 weeks (drift values are real, from weekly_zone_trajectories.csv; "
  "the CUSUM column is computed from composite_drift over this window):",bold=True)
table(["Week","access_pattern_drift","composite_drift","CUSUM (composite)"],
 [["64","0.068","0.013","0.000"],["65","0.021","0.021","0.001"],["66","0.017","0.020","0.001"],
  ["67","0.021","0.017","0.000"],["68","0.232","0.027","0.007"],["69","0.212","0.033","0.020"]],widths=[0.8,1.9,1.7,1.7],fs=9)
P("How these numbers come: each week the entity's composite, and each zone, is a 1536-d vector. The drift is the cosine "
  "distance to the previous week (1 - cosine). access_pattern_drift is that distance for the access_pattern zone alone; "
  "composite_drift is it for the full composite vector; CUSUM is the running sum of composite_drift above the 0.02 noise "
  "floor (the Part 4 formula), floored at zero and shown from a zero start at week 64.")
P("Read it honestly: the access_pattern zone ramps about 10x in the final two weeks, but Volt's COMPOSITE drift only "
  "creeps (0.013 to 0.033) and its CUSUM reaches just 0.020, below the 0.05 fire line. Volt does not separate by "
  "dramatic drift; CUSUM is the broad net, and Volt is ultimately caught by BREADTH in the composite score (Part 6, "
  "rank #30). That low drift is exactly why it is the stealthiest of the four.",italic=True)

# ===== PART 5 =====
H1("Part 5. Drift direction and MITRE projection (meaning, not magnitude)")
P("This is the part that replaces \"87% anomalous\" with an actual diagnosis:")
NUM("drift_vector = normalize(v_new - v_old) ","a unit vector pointing in the direction of behavioral change.")
NUM("Project it (cosine) onto 14 threat + 2 benign concepts, ","each itself a prose description embedded into the SAME 1536-d space.")
NUM("If the top-aligned concept is a threat and alignment >= 0.15, ","flag it and emit the concept name plus its MITRE techniques.")
P("Zone-specific filtering compares each zone only against relevant concepts (network_footprint vs c2_beacon, not vs "
  "data_exfiltration), which kills cross-domain false alignments. Output reads like:")
code("\"identity STABLE (0.003), network_footprint DRIFTING (0.082) toward c2_beacon [T1071, T1573]\"")
P("The 14 threat concepts and the four attackers they match:",bold=True)
table(["Concept","What it is","MITRE","Matches"],
 [["living_off_the_land","built-in OS tools (PowerShell, certutil, wmic) used maliciously, no malware; blends with admin activity","T1218, T1053, T1059.001, T1036","Volt"],
  ["telecom_infrastructure_pivot","targeting telecom infra for surveillance: CDR databases, lawful-intercept, router/switch config, restricted zones","T1557, T1040, T1005, T1556","Salt"],
  ["insider_threat_slow","a trusted user gradually expanding file access, off-hours activity, and out-of-role documents over months","T1078, T1083, T1005, T1052","Insider / Slow APT"],
  ["c2_beacon","covert command-and-control: regular periodic callbacks to external infra, small packets, newly-seen domains","T1071, T1573, T1568, T1102","Slow APT (novel IP)"],
  ["data_exfiltration","systematic collection and transfer of sensitive data out: staging, compression, external upload; volume rising","T1005, T1074, T1048, T1567","(data theft)"],
  ["privilege_escalation","progressively gaining higher access: elevated-permission requests, misconfig exploitation, scope expanding org-wide","T1078, T1068, T1134, T1548","(priv-esc)"],
  ["lateral_movement","moving host to host: SMB/WinRM/RDP to many hosts, credential reuse, admin-share access outside normal scope","T1021, T1570, T1550, T1072","(lateral)"],
  ["compromised_endpoint","a device showing compromise: encoded PowerShell, C2 connections, registry persistence, security-tool disabling","T1059, T1547, T1562, T1071","(endpoint)"],
  ["credential_harvesting_slow","collecting credentials via native tools over time: LSASS access, SAM dumps, Kerberoasting (not brute force)","T1003, T1558, T1110.002, T1555","(cred theft)"],
  ["dns_tunneling_exfil","DNS used as a covert exfil channel: high-entropy subdomains encoding data, elevated TXT/MX to attacker domains","T1071.004, T1048.003, T1041","(DNS exfil)"]],
 widths=[1.3,2.9,1.5,0.8],fs=7)
P("The remaining four threat concepts and the two benign anchors:",bold=True)
table(["Concept","What it is","Type / MITRE"],
 [["reconnaissance","systematic discovery and mapping of the network: port scanning, LDAP / share / account enumeration, methodical progression through segments","threat: T1046, T1018, T1087, T1135"],
  ["credential_stuffing","mass authentication attempts across many accounts from external IPs; high failure rate, automated timing, proxy / VPN / botnet origin","threat: T1110, T1078"],
  ["supply_chain_compromise","a trusted application behaving differently after an update: new connections to unfamiliar domains, new process spawning, access to data stores not touched before","threat: T1195, T1059, T1071"],
  ["insider_threat_fast","a user rapidly collecting and exfiltrating sensitive data (often pre-termination): sudden download spike, bulk restricted access, USB or large email, within days","threat: T1005, T1052, T1048, T1567"],
  ["normal_role_change","legitimate shift from a role transfer or promotion: new app access, new hours / timezone, new peer group, matching HR records","benign anchor"],
  ["seasonal_variation","behavior fluctuating with business cycles: end-of-quarter finance spikes, audit-driven changes, holiday lulls, deadline-driven off-hours work","benign anchor"]],
 widths=[1.5,3.6,1.4],fs=7.5)
P("The two benign anchors are embedded alongside the threat concepts. If an entity's drift aligns most strongly with a "
  "benign concept, the alert is suppressed: the change looks like a legitimate role change or a seasonal pattern rather "
  "than a threat. This is how the system damps false positives.",italic=True)

H2("Worked example: the living_off_the_land concept (embedding and alignment)")
P("The concept description, verbatim from detection/reference_concepts.py:")
code("An entity using only built-in operating system tools for malicious purposes\n"
     "with no malware deployed. Execution of certutil for downloads, wmic for remote\n"
     "process execution, netsh for port forwarding, schtasks for persistence, and\n"
     "PowerShell with obfuscated commands. ... Behavioral metrics show\n"
     "endpoint_unique_processes increasing, endpoint_suspicious_ratio elevated,\n"
     "auth_unique_dests expanding, net_unique_dsts growing, and data_behavior plus\n"
     "network_footprint drifting simultaneously.")
B("The embedding that is created: ","this entire paragraph is passed once through the SAME embedding model used for the "
  "entity zones, producing one fixed 1536-d concept vector (cached in data/concept_embeddings.npz). Same model and same "
  "space is what makes a cosine between concept and entity meaningful; and the description deliberately names the same "
  "features as the zone serializations (unique_processes, suspicious_ratio, unique_dests), so a real living-off-the-land "
  "entity embeds near it by design.")
P("How the alignment is derived for Volt:",bold=True)
NUM("Drift direction: ","for Volt's relevant zones (access_pattern, network_footprint, risk_posture), compute d = normalize(V_new - V_old), the unit vector of where Volt is heading.")
NUM("Cosine: ","alignment = cosine(d, c_LOTL) = (d . c_LOTL) / (||d|| ||c_LOTL||).")
NUM("Decide: ","living_off_the_land gives Volt the highest cosine among the zone's relevant concepts and is >= 0.15, so the system labels Volt's direction 'living_off_the_land' and emits T1218, T1053, T1059.001, T1036.")
P("Why they align: Volt's real behavior (new admin processes, expanding destinations, simultaneous data and network "
  "drift, no malware) moves in a direction whose embedding sits close to the concept's embedding, because both are "
  "expressed in the same behavioral terms. That closeness is the high cosine.",italic=True)

# ===== PART 6 =====
H1("Part 6. The composite scorer (5 phases to one ranked number)")
P("Per user, features are extracted, converted to GROUP-RELATIVE z-scores (each user vs the normals in their own role "
  "group: admin, security, developer, business, executive), then fused by this exact formula:")
code("composite = signal_strength    x 1.0   (sum of top-3 feature z-scores)\n"
     "          + breadth_15         x 0.5   (# features with z > 1.5)\n"
     "          + sustained          x 0.3   (sum of top-2 zone sustained z)\n"
     "          + max(ctx_spread_z,0) x 0.5  (context divergence, group-relative z)\n"
     "          + max(ctx_max_z,0)    x 0.3  (peak context score, group-relative z)\n"
     "          + novelty            x 1.0   (C2-beacon persistence, Part 7)")
P("Verbatim from detection/composite_scorer.py (compute_composite_scores). The context terms are the group-relative "
  "z-scores of ctx_spread and ctx_max (the ctx_spread_z / ctx_max_z columns in composite_scores.csv, and the values "
  "shown in the Part 8 table); breadth_15 and novelty are the only non-z terms.",italic=True,size=9.5)
B("Why group-relative: ","a SOC operator is scored against other security users, not against HR. An attacker has to be "
  "abnormal for their OWN role, which is a far higher bar than a global threshold.")

# ===== PART 7 =====
H1("Part 7. Novelty persistence (the C2 catch)")
P("Alongside the 23 numeric features, the system keeps qualitative features: the actual external IP addresses each "
  "entity contacted each week, not just the count. Novelty persistence works on those IPs.")
B("Baseline: ","the first 10 weeks establish each entity's normal set of external destinations. From week 10 on, any "
  "external IP not in that baseline is flagged NOVEL. The baseline is per-entity, so each entity is judged against its own history.")
P("Three metrics it computes (novelty_metrics.csv):",bold=True)
B("persistent_novel_ips ","- count of novel IPs that recur in 5 or more post-baseline weeks.",sub=True)
B("max_persistence (weeks) ","- the most weeks any single novel IP recurs. This is the key C2 signal.",sub=True)
B("weeks_frac ","- fraction of post-baseline weeks that contain at least one novel IP.",sub=True)
B("The score (composite_scorer.py): ","novelty = min(max_persistence / 5, 10), but only when max_persistence > 10 weeks; "
  "plus a bonus of weeks_frac x 3 when weeks_frac > 0.5.")
B("Why this is the C2 catch: ","a one-off new IP (a new website visited once) appears and disappears. A command-and-control "
  "beacon is a new IP the host calls back to repeatedly, week after week. The persistence of a novel IP is the beacon "
  "signature, which is exactly what max_persistence measures.")
P("Real novelty metrics for the four attackers (novelty_metrics.csv):",bold=True)
table(["Attacker","UID","persistent_novel_ips","max_persistence (weeks)","weeks_frac","novelty_score"],
 [["Volt","USR-042","0","0","0.00","0.0"],
  ["Salt","USR-118","1","14","0.23","2.8"],
  ["Insider","USR-156","0","1","0.05","0.0"],
  ["Slow APT","USR-234","1","60","1.00","13.0"]],widths=[0.9,0.9,1.4,1.4,0.9,1.0],fs=8.5)
P("Slow APT: a novel IP every week for 60 weeks. max_persistence 60 > 10 gives min(60/5,10)=10; weeks_frac 1.00 > 0.5 "
  "adds 1.00 x 3 = 3; total 13.0. Volt is 0 (LOTL uses internal admin tooling, no external beacon); Salt 2.8 (one novel "
  "IP for 14 weeks); Insider 0 (a novel IP once, below the 10-week gate). For Slow APT, signal_strength is only 4.50 and "
  "breadth is 2, so this persistent novel IP is the ONLY signal that lifts it to rank #7; without it, Slow APT is missed entirely.",italic=True)

# ===== PART 8: THE FOUR ATTACKERS =====
H1("Part 8. The four attackers: why each landed where it did")
P("All numbers exact from composite_scores.csv. Every composite re-derives from the Part 6 formula.")
table(["Attacker","UID / role","Rank","signal_str","breadth z>1.5 (z>2.0)","sustained","ctx_spread","ctx_max","novelty","COMPOSITE"],
 [["Salt","USR-118 / QA Eng (Engineering)","#1","30.60","17 (16)","9.58","9.00","8.10","2.8","51.71"],
  ["Insider","USR-156 / Analyst","#2","30.02","19 (17)","8.37","4.34","6.80","0.0","46.24"],
  ["Slow APT","USR-234 / ML Eng (Data Science)","#7","4.50","2 (0)","1.57","1.48","0.85","13.0","20.00"],
  ["Volt","USR-042 / SOC Op (Security)","#30","6.20","10 (3)","3.70","0.40","1.40","0.0","12.90"]],
 widths=[0.7,1.55,0.45,0.75,0.95,0.7,0.75,0.65,0.65,0.85],fs=7.5,hdr_fs=7.5)

H2("Salt Typhoon - USR-118 - Rank #1 (51.71)")
B("Caught by: ","breadth and context. QA Engineer whose department is Engineering. 17 features elevated, the "
  "highest context spread of anyone (9.00).")
code("51.71 = 30.60 + 17x0.5 + 9.58x0.3 + 9.00x0.5 + 8.10x0.3 + 2.8\n"
     "      = 30.60 + 8.50 + 2.87 + 4.50 + 2.43 + 2.80")
H2("Insider - USR-156 - Rank #2 (46.24)")
B("Caught by: ","raw breadth and strength. The highest signal_strength of anyone (30.02), 19 features elevated. A classic broad insider.")
code("46.24 = 30.02 + 9.50 + 2.51 + 2.17 + 2.04 + 0.00")
H2("Slow APT - USR-234 - Rank #7 (20.00) - the 'invisible to thresholds' case")
B("Caught by: ","novelty ALONE. signal_strength 4.50, breadth(z>1.5) = 2, breadth(z>2.0) = 0. On every magnitude or "
  "threshold method this user is invisible. The novelty score of 13.0 (a novel C2 IP every week for 60 weeks) is what "
  "lifts it to rank 7. This is the proof that the system catches what thresholds structurally cannot.")
code("20.00 = 4.50 + 1.00 + 0.47 + 0.74 + 0.25 + 13.00")
H2("Volt Typhoon - USR-042 - Rank #30 (12.90) - the hardest, living-off-the-land")
B("Caught by: ","breadth of small elevations. signal_strength 6.20, breadth(z>1.5)=10 but breadth(z>2.0)=only 3, "
  "novelty 0. Translation: many small elevations, no single dramatic one. That is the signature of legitimate admin "
  "tooling used maliciously.")
code("12.90 = 6.20 + 5.00 + 1.11 + 0.20 + 0.42 + 0.00")
P("Raw behavioral features, week 69 (entity_structures.json), showing how ordinary each attacker looks per metric:",bold=True)
table(["Metric","Volt USR-042","Salt USR-118","Slow APT USR-234"],
 [["auth_off_hours_ratio","0.39","0.60","0.24"],
  ["auth_unique_dests","7","4","7"],
  ["file_total_bytes","264.6 MB","150.6 MB","814.0 MB"],
  ["file_restricted_ratio","0.00","0.00","0.00"],
  ["endpoint_suspicious_ratio","0.00","0.00","0.17"],
  ["endpoint_max_risk","26","23","50"],
  ["net_bytes_out","42.9 MB","146.9 MB","308.8 MB"],
  ["net_unique_dsts","24","23","24"],
  ["dns_unique_domains","17","13","19"]],widths=[2.0,1.6,1.6,1.6],fs=8.5)
P("No single feature screams attack. That is the whole point: the twin separates them by trajectory, direction, "
  "breadth, and novelty, not by any one number.",italic=True)

# ===== PART 9: THE 10.6% =====
H1("Part 9. The live 10.6% derivation")
B("","Sort all 250 entities by composite. Volt (12.90) is the stealthiest scorer, so the operating line must reach down "
  "to 12.95 to catch all four. 26 of 246 normal users score above that line, so 26/246 = 10.6% false positive, 4/4 caught.")
B("Baselines on the same data: ","Isolation Forest, One-Class SVM, and Local Outlier Factor each catch 0 of 4; a "
  "z-score method catches 1 of 4.")
B("Separation is by RANK, not a clean gap: ","the top normal user (USR-081, HR Manager, 21.52) outscores Slow APT "
  "(20.00) and Volt (12.90). About 26 normals outscore Volt. This is why we deliver a ranked list, not a threshold; the "
  "honest claim is recall at the top of the ranking.")
table(["Rank","UID","Type","Composite"],
 [["1","USR-118","Salt Typhoon","51.71"],["2","USR-156","Insider","46.24"],
  ["3","USR-081","normal (HR Manager)","21.52"],["4","USR-075","normal (Data Scientist)","21.03"],
  ["5","USR-220","normal (Senior Engineer)","20.83"],["6","USR-021","normal (ML Engineer)","20.50"],
  ["7","USR-234","Slow APT","20.00"],["...","...","~23 more normals","12.90 to 20.0"],
  ["30","USR-042","Volt Typhoon","12.90"]],widths=[0.7,1.2,2.4,1.2],fs=8.5)

# ===== PART 10: BUILD/ROADMAP =====
H1("Part 10. What is built versus roadmap (honest split)")
B("Built and runs live: ","the 5 signals, 1536-d embedding, weighted composition, zone drift, CUSUM, "
  "drift-direction/MITRE projection, the 5-phase composite, and novelty persistence.")
B("Roadmap (not running today): ","identity fusion across systems, the versioned twin store, the behavioral graph, "
  "tamper-evident forensics, and automated response.")
callout("With an ARCYBER audience, state this split out loud at the start of the demo. Candor is the advantage.")

# ===== GLOSSARY =====
pb=d.add_paragraph(); pb.add_run().add_break(WD_BREAK.PAGE)
H1("Glossary - Terms and Acronyms")
P("Plain-language expansion of the terms and acronyms used above.",italic=True,size=9.5,color=GRAY)
H2("Math and machine learning")
table(["Term","Meaning"],
 [["Embedding","a vector (here 1536 numbers) representing the meaning of text in a high-dimensional space"],
  ["1536-d","the dimensionality of each embedding vector"],
  ["Cosine similarity","cosine of the angle between two vectors: +1 identical direction, 0 unrelated, -1 opposite"],
  ["Cosine distance","1 - cosine similarity: 0 identical, 1 unrelated, 2 opposite"],
  ["z-score","how many standard deviations a value sits from the (group) mean"],
  ["L2-normalize","scale a vector to unit length (magnitude 1)"],
  ["Softmax","turns a set of scores into positive weights that sum to 1"],
  ["Attention","weighting that emphasizes the more relevant inputs (here, the louder zones)"],
  ["CUSUM","Cumulative Sum; a change-point method that accumulates small deviations to detect slow, sustained drift"],
  ["Drift","the change in an entity's behavioral vector between time periods"]],
 widths=[1.5,5.0],fs=8.5)
H2("Detection and cyber")
table(["Term","Meaning"],
 [["MITRE ATT&CK","an industry-standard knowledge base of adversary tactics and techniques (the T-codes)"],
  ["T-code (e.g., T1078)","a numbered MITRE ATT&CK technique; T1078 = Valid Accounts"],
  ["UEBA","User and Entity Behavior Analytics"],
  ["APT","Advanced Persistent Threat; a stealthy, long-dwell adversary"],
  ["Volt / Salt Typhoon","Chinese state-sponsored campaigns against U.S. critical infrastructure (CISA advisories)"],
  ["LOTL / Living off the Land","attacks using only built-in OS tools, with no malware deployed"],
  ["LOLBins","Living-off-the-Land Binaries; signed OS tools (PowerShell, certutil, wmic) repurposed maliciously"],
  ["C2 / C2 beacon","Command and Control; a compromised host's periodic callback to attacker infrastructure"],
  ["DGA","Domain Generation Algorithm; malware auto-generating many domains to reach C2"],
  ["Kerberoasting","extracting service-account credentials by requesting and cracking Kerberos tickets"],
  ["LSASS","Local Security Authority Subsystem Service; Windows process holding credentials in memory"],
  ["SAM","Security Account Manager; the Windows local credential database"]],
 widths=[1.5,5.0],fs=8.5)
H2("Protocols, data, and infrastructure")
table(["Term","Meaning"],
 [["DNS","Domain Name System; resolves names to IP addresses (abused for tunneling)"],
  ["NXDOMAIN","a DNS response meaning the queried domain does not exist"],
  ["TXT / MX","DNS record types (text / mail-exchange), abused as covert channels"],
  ["SMB / WinRM / RDP","Windows file-sharing / remote-management / remote-desktop protocols; common lateral-movement paths"],
  ["LDAP","Lightweight Directory Access Protocol; directory queries (enumeration)"],
  ["CDR","Call Detail Record; telecom metadata of who called whom, when"],
  ["Lawful intercept","legally-authorized communications-interception systems in telecom networks"],
  ["IdP / MFA / PAM / EDR","Identity Provider / Multi-Factor Authentication / Privileged Access Management / Endpoint Detection and Response"],
  ["API","Application Programming Interface"],
  ["IL5 / IL6","DoD cloud security Impact Levels (IL5 CUI/NSS; IL6 up to SECRET)"],
  ["JWICS","Joint Worldwide Intelligence Communications System; the Top-Secret/SCI network"]],
 widths=[1.5,5.0],fs=8.5)

P("")
P("Internal technical reference. Contains proprietary scoring logic. Not for distribution.",italic=True,size=9,color=GRAY)

out="WP DLA/Presentation/Behavioral_Digital_Twin_DeepDive_DataPoints_INTERNAL.docx"
os.makedirs("WP DLA/Presentation",exist_ok=True)
d.save(out)
print("saved:",out,"| paragraphs:",len(d.paragraphs),"| tables:",len(d.tables))
