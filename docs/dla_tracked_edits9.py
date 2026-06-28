"""Ninth pass: restore the 'Digital Twin Demonstration' figure (from previous_final) into
the EDM Forecasting Pipeline subsection of Use Case 1, as a TRACKED insertion, and
renumber the downstream figures + references (direct, mechanical):
  new sequence: ... 5 Legacy/EDM | 6 Digital Twin (restored) | 7 Supplier Zones |
                8 Network | 9 Evidence | 10 Pilot
"""
import os, copy
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PREV = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_previous_final.docx"
REV = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_final_REV_v2.docx"
AUTHOR = "ReviewPanel"; DATE = "2026-06-19T20:00:00Z"
IMG = os.path.join("docs", "_fig_digitaltwin.png")

def acc(p): return "".join(x.text or "" for x in p._p.iter(qn('w:t')))

# --- extract Digital Twin image from previous_final ---
prev = Document(PREV)
idxs = [i for i, p in enumerate(prev.paragraphs) if 'Digital Twin demonstration' in acc(p) or 'Digital twin:' in acc(p)]
blip = None; width_in = 6.5
for i in idxs:
    for j in range(max(0, i - 2), min(len(prev.paragraphs), i + 3)):
        bl = prev.paragraphs[j]._p.findall('.//' + qn('a:blip'))
        if bl:
            blip = bl[0]
            ext = prev.paragraphs[j]._p.findall('.//' + qn('wp:extent'))
            if ext: width_in = min(int(ext[0].get('cx')) / 914400, 6.5)
            break
    if blip is not None: break
if blip is None:
    raise SystemExit("Digital Twin image not found in previous_final")
open(IMG, 'wb').write(prev.part.related_parts[blip.get(qn('r:embed'))].blob)
print("extracted Digital Twin figure:", os.path.getsize(IMG), "bytes, width_in=%.2f" % width_in)

# --- open current ---
d = Document(REV)
_id = [21000]
def nid():
    _id[0] += 1; return str(_id[0])
def insmark():
    e = OxmlElement('w:ins'); e.set(qn('w:id'), nid()); e.set(qn('w:author'), AUTHOR); e.set(qn('w:date'), DATE); return e
def ins_wrap(run):
    e = insmark(); e.append(run); return e

# --- renumber (direct), descending captions then refs ---
OPS = [
    ("Figure 9 - Bounded Pilot", "Figure 9", "Figure 10"),
    ("Figure 8 - Technical Evidence", "Figure 8", "Figure 9"),
    ("Figure 7 - Logistics Network", "Figure 7", "Figure 8"),
    ("Figure 6 - Our Supplier Risk Zones", "Figure 6", "Figure 7"),
    ("as shown in Figure 6, is better treated", "Figure 6", "Figure 7"),
    ("detailed in Figure 6, below", "Figure 6", "Figure 7"),
    ("Figure 6 shows both. The top row", "Figure 6", "Figure 7"),
    ("behavioral twin of each item-and-depot as shown in Figure 4", "Figure 4", "Figure 6"),
]
def renumber(context, old, new):
    for p in d.paragraphs:
        if context in acc(p):
            for t in p._p.iter(qn('w:t')):
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new, 1); return True
    return False
applied = sum(1 for c, o, n in OPS if renumber(c, o, n))
print("renumber ops applied:", applied, "of", len(OPS))

# --- restore Digital Twin figure (tracked) after the EDM Forecasting Pipeline heading ---
cap_src = next(p for p in d.paragraphs if 'Legacy Approach and EDM Vector Intelligence' in acc(p) and acc(p).strip().startswith('Figure'))
src_pPr = copy.deepcopy(cap_src._p.find(qn('w:pPr')))
src_rPr = cap_src.runs[0]._r.find(qn('w:rPr'))
src_rPr = copy.deepcopy(src_rPr) if src_rPr is not None else None

heading = next(p for p in d.paragraphs if acc(p).strip() == 'EDM Forecasting Pipeline')
tmp_p = d.add_paragraph(); tmp_r = tmp_p.add_run(); tmp_r.add_picture(IMG, width=Inches(width_in))
drawing = copy.deepcopy(tmp_r._r.find(qn('w:drawing')))
dp = drawing.find('.//' + qn('wp:docPr'))
if dp is not None:
    dp.set('id', '9006'); dp.set('name', 'Figure 6 Digital Twin Demonstration')
figp = OxmlElement('w:p'); pPr = OxmlElement('w:pPr')
jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)
rPr = OxmlElement('w:rPr'); rPr.append(insmark()); pPr.append(rPr); figp.append(pPr)
runel = OxmlElement('w:r'); runel.append(drawing); figp.append(ins_wrap(runel))
capp = OxmlElement('w:p'); capp.append(copy.deepcopy(src_pPr))
crPr = capp.find(qn('w:pPr')).find(qn('w:rPr'))
if crPr is None:
    crPr = OxmlElement('w:rPr'); capp.find(qn('w:pPr')).append(crPr)
crPr.append(insmark())
crun = OxmlElement('w:r')
if src_rPr is not None: crun.append(copy.deepcopy(src_rPr))
ct = OxmlElement('w:t'); ct.set(qn('xml:space'), 'preserve'); ct.text = "Figure 6 - Digital Twin Demonstration"
crun.append(ct); capp.append(ins_wrap(crun))
heading._p.addnext(capp); heading._p.addnext(figp)
tmp_p._p.getparent().remove(tmp_p._p)

o = REV; k = 3
while True:
    try:
        d.save(o); break
    except PermissionError:
        root, ext = os.path.splitext(REV); o = f"{root.replace('_v2','')}_v{k}{ext}"; k += 1
print("Saved:", o)
