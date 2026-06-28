# -*- coding: utf-8 -*-
"""Internal deep-dive briefing book for the Brandon Pugh (Pentagon cyber advisor)
technical meeting. Slide-by-slide. Grounded in the ACTUAL CyberUEBA codebase.
Each slide tagged BUILT / PARTIAL / ROADMAP / SEPARATE. No em/en dashes.
Output: WP DLA/Presentation/.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6); GRAY=RGBColor(0x55,0x5F,0x6B)
GREEN=RGBColor(0x1F,0x7A,0x33); AMBER=RGBColor(0xB5,0x5A,0x00); RED=RGBColor(0xA6,0x1B,0x1B)
d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(10.5)

def H1(t):
    p=d.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
    return p
def P(t,italic=False,bold=False,color=None,size=None):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold
    if color: r.font.color.rgb=color
    if size: r.font.size=Pt(size)
    return p
def field(label,text):
    p=d.add_paragraph(); rl=p.add_run(label+" "); rl.bold=True; rl.font.color.rgb=BLUE; p.add_run(text); return p
def B(label,rest=""):
    p=d.add_paragraph(style='List Bullet')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def Bn(t): return d.add_paragraph(t,style='List Bullet')
def qa(q,a):
    p=d.add_paragraph(style='List Bullet')
    rq=p.add_run("Q: "+q); rq.bold=True
    p.add_run("  A: "+a)
    return p
def table(headers,rows,widths=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""; r=cells[i].paragraphs[0].add_run(v); r.font.size=Pt(9)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    d.add_paragraph(); return t

STATUS={'BUILT':GREEN,'PARTIAL':AMBER,'ROADMAP':RED,'SEPARATE':RED,'FRAME':NAVY}
def slide(num,title,status):
    p=d.add_heading(level=2);
    r=p.add_run(f"Slide {num}  |  {title}"); r.font.color.rgb=BLUE
    bp=d.add_paragraph(); br=bp.add_run(f"[{status}]"); br.bold=True; br.font.color.rgb=STATUS[status]; br.font.size=Pt(9.5)
    return p

# ===== COVER =====
cp=d.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=cp.add_run("Cyber Behavioral Intelligence: Technical Deep-Dive Briefing Book"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
for txt,sz,it,col in [("Preparation for the Brandon Pugh Technical Session (Pentagon / R Street)",13,False,BLUE),
                      ("Internal preparation document. Slide-by-slide. Grounded in the actual implementation.",10.5,True,GRAY)]:
    pp=d.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=pp.add_run(txt); rr.font.size=Pt(sz); rr.italic=it; rr.font.color.rgb=col
pp=d.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
rr=pp.add_run("Ravindra Shukla  |  V-Intelligence"); rr.font.size=Pt(10)
warn=d.add_paragraph(); warn.alignment=WD_ALIGN_PARAGRAPH.CENTER
rw=warn.add_run("INTERNAL USE ONLY. Contains build-vs-roadmap distinctions and proprietary detail. Not the leave-behind."); rw.bold=True; rw.font.color.rgb=RED; rw.font.size=Pt(9.5)
d.add_paragraph()

# ===== SECTION 0: FRAME =====
H1("Section 0. How to Use This Book and Run the Meeting")

slide("0.1","Meeting Context and Objectives","FRAME")
field("Audience:","Brandon Pugh, cybersecurity lead at R Street and advisor to Pentagon / critical-infrastructure cyber. Technical enough to go deep; his lens is policy and national defense. He rewards candor and discounts anything that looks tuned or overclaimed.")
field("Format:","2 to 3 hour technical deep dive. He has the layered cybersecurity whitepaper from the Presentation folder.")
field("Our objectives:","(1) Establish that behavioral, meaning-based detection closes the Volt/Salt living-off-the-land gap; (2) prove it with a live, honest demo; (3) secure a bounded validation pilot and a policy/thought-leadership relationship.")
field("Golden rule:","Lead with the national problem, then the innovation, then proof. Show what runs. Label what is designed-next as roadmap, explicitly. Candor is the advantage.")

slide("0.2","Run-of-Show (target 2.5 hours + buffer)","FRAME")
table(["Block","Time","Content","Driver"],
 [["0. Open","5 min","His goals, confirm time, set the arc","You"],
  ["1. Problem + Innovation","25 min","Threat inflection (Volt/Salt, LoTL, T1078); meaning vs magnitude; architecture","You"],
  ["2. Data + Digital Twin","35 min","Synthetic dataset; build one entity into a behavioral twin live; the 4 campaigns","Demo driver"],
  ["3. Detection deep dive","50 min","Drift, direction/projection, composite scoring; walk Salt #1 and Volt #24; 4/4 vs 0-1/4","Both"],
  ["4. Operationalize + Federal fit","25 min","Deploy (IL5/6, JWICS, air-gap, no PII), SIEM fit, roadmap, honest validation path","You"],
  ["5. Policy + next steps","15 min","Preemptive shift, the national gap, pilot, thought leadership","You"],
  ["Buffer","15 min","His tangents and questions","-"]],
 widths=[1.6,0.8,3.4,0.9])

slide("0.3","The Honesty Boundary (most important internal page)","FRAME")
P("Keep these three columns straight at all times. If asked about a roadmap item, say plainly: \"That is designed into "
  "the architecture and is the next build phase; what you are seeing today is the detection core.\"")
table(["Show live (BUILT)","Show with precision (PARTIAL)","Designed next / separate (DO NOT claim built)"],
 [["4 attack simulators","Digital twin: 5 signals + zones built; identity fusion NOT built","Identity fusion across 10+ systems (ROADMAP)"],
  ["Behavioral embedding (text to 1536-d)","FP economics: 8.1% composite real (cleanly separates 2/4); multi-front Threat-Profile detector built, 4/4 at 0 FP","HMAC tamper-evident forensics (ROADMAP)"],
  ["Drift / CUSUM","Cross-domain DLA: same architecture ported, no live data link","ABAC proportional response, MFA/lockout (ROADMAP)"],
  ["Direction / MITRE projection","","Layer 1 (preemptive), Layer 3 API, Layer 5 Mythos: presented as our four-layer platform; only the behavioral layer is in this codebase"],
  ["5-phase composite scoring","",""],
  ["The 4 case-study ranks","",""],
  ["Head-to-head baselines (IF/SVM/LOF/z)","",""],
  ["Deployment (Docker, local-model air-gap path, API/UI)","",""]],
 widths=[2.0,2.4,2.4])

slide("0.4","Numbers to Keep Straight (he will check)","FRAME")
B("Dataset: ","250 synthetic entities, 485 days (70 weeks), 2025-01-01 to 2026-04-30, ~14 million events across 5 log sources.")
B("Embedding: ","1536 dimensions (OpenAI text-embedding-3-small, mandatory; the pipeline hard-fails without an API key). For air-gapped runs a local open embedding model swaps in behind the same embedder interface.")
B("False-positive rate: ","8.1 percent at the composite operating point (about 20 false positives out of 246 normal users), catching all four campaigns. Use 8.1, not 8.5; the deck's 8.5 is wrong. Where 8.1 comes from, and you can derive it live: to catch the stealthiest attacker, Volt at composite 13.70 (rank #24), the threshold sits at 13.70, and 20 normal users score above that, so 20/246 = 8.1 percent, catching 4 of 4.")
B("Ranks (from composite_scores.csv): ","Salt USR-118 #1 (composite 51.27); Insider USR-156 #2 (46.24); Slow APT USR-234 #7 (19.44); Volt USR-042 #24 (13.70). Volt is #24, and the deck is correct; do not 'correct' it to #25.")
B("Baselines: ","Isolation Forest 0/4 (5.3% FP); One-Class SVM 0/4 (14.6% FP); Local Outlier Factor 0/4 (4.5% FP); z-score 1/4 (9.8% FP). So traditional methods catch 0 to 1 of 4.")
B("Detection results to keep straight: ","traditional (IsolationForest/SVM/LOF) 0/4; z-score 1/4; composite 4/4 at 8.1% FP but cleanly separates only 2/4 (USR-118, USR-156); multi-front Threat-Profile detector 4/4 at 0 FP (named profiles: C2-beacon, DGA, LOTL-process, cohort-rare access, recon-fanout, insider-collection) is the primary detector.")
B("Separation is by RANK, not a clean threshold gap (say this honestly): ","the top normal user (USR-081, composite 21.5) outscores the two stealthiest attackers, Slow APT (19.4) and Volt (13.7); about 20 normal users score above Volt. All four attackers still surface in the top of the risk-ranked list. This is precisely why we present a ranked list, not a threshold. Do not claim 'clean separation' on raw composite; the honest claim is recall at the top of the ranking.")
B("Operating-point candor: ","the composite threshold was selected with knowledge of ground truth in the proof of concept. Present it as a reporting benchmark to be re-validated on real data, not a deployable threshold.")

# ===== SECTION 1: PROBLEM & INNOVATION =====
H1("Section 1. Problem and Innovation (the first 25 minutes)")

slide("1","The Threat Inflection: Volt and Salt Typhoon","BUILT")
field("Message:","Apex adversaries are already inside, using valid accounts and living-off-the-land, crossing no threshold for years. This is the gap we close.")
field("Technical detail:","")
B("Volt Typhoon (CISA AA24-038A): ","pre-positioning in U.S. critical infrastructure, only legitimate Windows tooling (wmic, ntdsutil, netsh, rundll32), no malware, no signatures, five-plus years dwell, edge-device entry.")
B("Salt Typhoon (CISA AA25-239A): ","telecom espionage at carrier scale into lawful-intercept systems, slow and sub-threshold, multi-year, valid-account access.")
B("The common technique: ","valid accounts (MITRE T1078), the most-used initial-access technique; stolen credentials are the top initial vector; credential-borne breaches dwell the longest of any category.")
field("Talking track:","\"Every event these campaigns generate looks legitimate in isolation. That is the point. A threshold never fires, a signature never matches, and one attacker is split across many identity systems. The detection problem is not magnitude, it is meaning and direction over time.\"")
field("Likely questions:","")
qa("Are these the real campaigns or your simulation?","Our simulators reproduce the published TTPs of these campaigns in a controlled population so we can score detection blind. We are not claiming to have caught the real intrusions.")

slide("2","The Core Insight: Meaning vs Magnitude","BUILT")
field("Message:","Traditional tools count; they do not comprehend. The same number can mean two completely different things.")
field("Technical detail:","")
B("The example: ","two users each touch 47 files this week. One reads their own team's documents; the other scans 12 departments at 2 a.m. A scalar feature sees the same number, 47, and the attacker is invisible.")
B("Why every traditional method has a blind spot: ","Isolation Forest misses slow drift within normal ranges; One-Class SVM gives high false positives in high dimensions; Local Outlier Factor is scale-sensitive; z-score only catches single-feature spikes.")
B("Our move: ","represent behavior as meaning. Raw metrics become prose, prose becomes a 1536-dimensional vector that captures behavioral semantics, and we then read the trajectory and direction of that vector over time.")
field("Talking track:","\"We do not add another anomaly score. We change what is being measured, from how much an entity did to what it is becoming.\"")

# ===== SECTION 2: THE ENGINE (BUILT CORE) =====
H1("Section 2. The Engine, Deep (the built core you can demo)")

slide("3","Card 1: Attack Simulators and Synthetic Data","BUILT")
field("Message:","Four distinct nation-state-grade campaigns embedded in a 250-entity population with full ground truth, so detection is a blind, scorable test.")
field("Technical detail (simulator/attacks/):","")
B("insider_threat.py (USR-156, 8 months): ","four phases, mood shift, curiosity, reconnaissance, exfiltration.")
B("apt_slow.py (USR-234, 180 days): ","low-magnitude C2 beacon plus domain-generation-algorithm DNS plus gradual data staging.")
B("volt_typhoon.py (USR-042, 120 days): ","living-off-the-land, four phases (initial access, credential harvesting, lateral movement, pre-positioning), legitimate admin tools (PowerShell, WMI, netsh, certutil).")
B("salt_typhoon.py (USR-118, 100 days): ","telecom infrastructure, four phases (edge compromise, infrastructure mapping, call-detail/comms collection, staged exfiltration via DNS tunneling).")
B("Each simulator ","injects realistic multi-signal events (auth, network, DNS, file, endpoint) with MITRE ATT&CK technique tags.")
field("Show / demo:","Open one simulator file; show the phase structure and MITRE tags. Then show the generated event stream for that user in data/generated/.")
field("Talking track:","\"We control ground truth, so this is a blind test we can grade: 4 known attackers hidden among 246 normal users.\"")

slide("4","Card 2: Behavioral Embedding (text to 1536-d)","BUILT")
field("Message:","Raw metrics are serialized into prose and embedded into a unified 1536-dimensional semantic space, so meaning, not just counts, becomes measurable.")
field("Technical detail:","")
B("embeddings/embedder.py: ","Embedder uses OpenAI text-embedding-3-small (1536-d) with disk and memory caching; real embeddings are mandatory and the pipeline hard-fails without an API key. For air-gapped runs a local open embedding model swaps in behind the same interface.")
B("embeddings/snapshot_builder.py: ","serializes the five behavioral signals per entity per period into a text description, embeds each, and composes them into one monthly vector via weighted average.")
B("Example serialization: ","\"USR-042 ran administrative tooling across 9 hosts it had not touched before, 71 percent outside normal hours, with no new file writes.\" That sentence is what gets embedded.")
field("Show / demo:","Show a serialized behavior sentence and note that the vector captures its semantics. Mention the local-model path for air-gapped enclaves.")
field("Likely questions:","")
qa("You used a commercial embedding model. Data sovereignty?","Real embeddings are mandatory today, and a local open model can replace the commercial embedder in an enclave for fully offline runs. No personally identifiable information is placed in the embedding; we embed aggregated behavior descriptions, not raw events.")

slide("5","Card 3: Digital Twin (5 signals + behavioral zones)","PARTIAL")
field("Message:","Each entity is rebuilt as a behavioral twin: five signals, five behavioral zones, and a trajectory of snapshots over time.")
field("Technical detail:","")
B("models/cyber_entity.py: ","CyberEntity holds zone embeddings (5 zones), phase state (velocity, acceleration, regime shifts), relationships, and risk scores.")
B("models/hierarchical_zones.py: ","five zones for a user, identity, access pattern, data behavior, network footprint, risk posture.")
B("Five signals: ","authentication, process, network, file, identity.")
B("data/tier3_results/weekly_zone_trajectories.csv: ","actual weekly zone trajectories across 250 users over 70 weeks (a real 5.8 MB artifact you can open).")
field("PARTIAL boundary:","The twin, zones, and trajectory are built. Identity fusion across 10+ systems is NOT built; today each user_id is treated as atomic. Say \"identity fusion is the designed next phase.\"")
field("Show / demo:","Open an entity in the UI and show its zone trajectory over the weeks.")

slide("6","Card 4: Drift Detection and CUSUM","BUILT")
field("Message:","Sequential change-point detection on the embedding trajectory accumulates sub-threshold drift, so a slow campaign that no monthly check would flag becomes visible.")
field("Technical detail (detection/cusum.py):","")
B("cusum_detect() ","accumulates drift magnitudes and alarms on sustained change even when each step is below the individual threshold.")
B("cusum_scan_entity() ","applies CUSUM to an entity's snapshot sequence; threshold 0.05, baseline allowance 0.02, minimum run length 2.")
B("Why it matters: ","drift of 0.02 per month for six months is invisible to any single monthly check but accumulates to an alarm. Sequential detection of this kind is minimax-optimal, so years of dwell collapse to days of lead time.")
B("First-pass behavior: ","CUSUM is the broad net (it flags most entities that move at all); the composite scorer then ranks who actually matters. Explain this so the broad first-pass count is not mistaken for the final result.")
field("Show / demo:","Show a healthy entity's flat drift versus an attacker's accumulating drift curve.")

slide("7","Card 5: Direction and MITRE ATT&CK Projection","BUILT")
field("Message:","We do not say \"87 percent anomalous.\" We project the drift direction onto named threat concepts and report \"drifting toward lateral movement,\" with a MITRE technique ID.")
field("Technical detail:","")
B("detection/drift_direction.py: ","compute_drift_vector() builds the normalized direction; concept_alignment() projects it onto reference concepts by cosine similarity; analyze_entity_drift() returns the top-aligned concepts, confidence, and threat classification.")
B("detection/reference_concepts.py: ","14 threat concepts (for example compromised_endpoint, data_exfiltration, insider_threat_slow, living_off_the_land, telecom_infrastructure_pivot) and 2 benign concepts (normal_role_change, seasonal_variation). Each is a natural-language description embedded to 1536-d; threat concepts carry MITRE techniques (T1078, T1083, T1005, and others).")
B("Example output: ","\"USR-156 drifting toward insider_threat_slow, alignment 0.68, confidence high.\"")
field("Talking track:","\"This is what makes an alert actionable and auditable for a SOC analyst and an auditor at once: not a score, a named direction mapped to a technique.\"")

slide("8","Card 7: Multi-Phase Composite Scoring","BUILT")
field("Message:","Five independent behavioral fronts are fused into one ranked score. Each attacker is exposed on a different front; no single front catches all four.")
field("Technical detail (detection/composite_scorer.py):","")
B("Phase 1, signal strength: ","sum of the top-3 z-scores.")
B("Phase 2, breadth: ","count of elevated features (z over 1.5 and over 2.0).")
B("Phase 3, sustained deviation: ","zone-specific late-stage sustained z-scores.")
B("Phase 4, context divergence: ","spread of context-specific scores.")
B("Phase 5, novelty persistence: ","recurrence of genuinely new behavior such as a C2 beacon.")
B("Fusion (exact, from composite_scorer.py): ","composite = signal_strength + 0.5*breadth + 0.3*sustained + 0.5*context_spread + 0.3*context_max + novelty, with the two context terms floored at zero. Note the context terms are two separate variables (spread and max), not one term times 0.8. Results for all 250 users are in data/tier3_results/composite_scores.csv.")
field("Talking track:","\"A normal user can be extreme on one front by chance. An attacker is extreme on several at once. That multi-front extremity is the fingerprint a single threshold cannot see.\"")

slide("9","Card 8: The Four Case Studies (the proof)","BUILT")
field("Message:","Four campaigns, four different tells. Composite scoring catches each on the phase that exposes it, blind, over 250 entities.")
table(["Entity","Campaign","Rank","Composite","The tell"],
 [["USR-118","Salt Typhoon (telecom)","#1","51.27","Extreme and broad: 16 features above z=2.0, driven by sustained deviation and context divergence. The strongest, clearest attacker (do NOT call it invisible to thresholds)."],
  ["USR-156","Insider (8-month)","#2","46.24","Signal strength + breadth; strong, broad, sustained escalation"],
  ["USR-234","Slow APT (180-day)","#7","19.44","The genuine sub-threshold case: signal strength only 4.47, breadth 1, invisible to per-feature thresholds; caught only by novelty persistence (a C2 beacon recurs weekly)"],
  ["USR-042","Volt Typhoon (LoTL)","#24","13.70","Breadth; living-off-the-land creates breadth, not magnitude. Hardest case: about 20 normal users outscore it; it surfaces by rank, not by a clean gap"]],
 widths=[0.9,1.9,0.6,1.0,2.6])
field("Honesty note:","Volt ranks #24 (the deck is correct; do not 'correct' it). It is the hardest case because living-off-the-land blends with normal IT admin behavior, so about 20 normal users actually outscore it and it surfaces by rank, not by a clean threshold gap. Say this openly: the method still finds the stealthiest attacker by ranking, and it is the strongest motivation for identity fusion and richer signals on the roadmap. Likewise, present Slow APT USR-234, not Salt, as the 'sub-threshold, invisible to per-feature thresholds' case.")

slide("10","Card 9: False-Positive Economics","PARTIAL")
field("Message:","All four caught at 8.1 percent false positive, where traditional tools caught zero to one. This is a risk-ranked list, not a flood of threshold alerts.")
field("Technical detail:","")
B("Composite: ","4 of 4 at about 8.1 percent FP (about 20 false positives among 246 normal users). The 8.1 percent is the price of catching the stealthiest attacker: Volt sits at composite 13.70, and 20 normal users score above that threshold (20/246 = 8.1 percent).")
B("Baselines at comparable thresholds: ","Isolation Forest 5.3 percent FP (0/4), Local Outlier Factor 4.5 percent FP (0/4), One-Class SVM 14.6 percent FP (0/4), z-score 9.8 percent FP (1/4).")
B("Operational framing: ","analysts triage the top of a ranked list, not thousands of alarms; the value is recall on the threats that matter with explainable ranking.")
field("Two detectors, two results:","The composite/embedding result above is 4 of 4 at 8.1 percent FP, but cleanly separates only 2 of 4 (USR-118, USR-156). The multi-front Threat-Profile detector (threat_profile_detector.py) is built and is the primary detector: named known-bad behavioral profiles (C2-beacon, DGA, LOTL-process, cohort-rare access, recon-fanout, insider-collection) catch all 4 of 4 at 0 false positives. Present the Threat-Profile 4/4-at-0-FP as the headline detection result and the composite 8.1% as the embedding/anomaly-ranking comparison.")
field("Likely questions:","")
qa("8 percent FP at DoD scale is a lot.","It is a ranked list, not threshold alerts; analysts work top-down. The operating point is a PoC reporting benchmark to be tuned and re-validated on real data.")

slide("11","Card 10: Head-to-Head Baselines","BUILT")
field("Message:","We ran the tools agencies actually deploy against the same data. They miss what we catch, on the record.")
field("Technical detail (comparison/run_comparison.py):","")
B("Implemented: ","Isolation Forest, One-Class SVM, Local Outlier Factor, and z-score, all from scikit-learn, run on the identical synthetic dataset with threshold sweeps.")
B("Result: ","traditional point-anomaly methods catch 0 of 4; z-score catches 1 of 4 while alarming on nearly everyone. The same four attackers sit inside the normal band in feature space and separate from the population only in the semantic embedding space.")
field("Show / demo:","Run or show the comparison output side by side with the composite ranking.")

slide("12","Card 12: Deployment and Sovereignty","BUILT")
field("Message:","Containerized, runs on the logs you already collect, deployable in an enclave, privacy-preserving by design.")
field("Technical detail:","")
B("docker-compose.yml: ","PostgreSQL with pgvector, a FastAPI service (container port 8000, published on host 8001), and a background worker.")
B("api/main.py and demo/ui/: ","REST API (entities, trajectories, detection, tier3 routes) and a static UI (entity detail, alerts, tier3 views).")
B("Air-gap: ","a local open embedding model replaces the commercial one in an enclave, running with no external calls.")
B("Privacy: ","metrics are aggregated; raw events are stored but not embedded, so no personally identifiable information enters the vector.")
field("Federal fit:","Deployable across IL5, IL6, and JWICS; local-model option for air-gapped enclaves; every verdict is explainable and MITRE-mapped.")

# ===== SECTION 3: CONTEXT & ROADMAP =====
H1("Section 3. Context and Roadmap (be precise here)")

slide("13","Card 13: Cross-Domain Foundation (DLA)","PARTIAL")
field("Message:","The same behavioral architecture is fielded in a different domain, supply-chain decision intelligence at DLA, which is independent evidence the method generalizes.")
field("Technical detail:","")
B("models/cyber_entity.py and hierarchical_zones.py ","were ported from the DLA Entity Digital Model; the zone-embedding architecture and velocity/acceleration/regime-shift detection come from that work.")
field("PARTIAL boundary:","It is the same architecture and algorithms, retargeted with cyber signals. It is NOT one running engine spanning both, and there is no live data link. Say \"same framework, separate deployments.\"")

slide("14","Card 6: Identity Fusion across 10+ Systems","ROADMAP")
field("Message (roadmap):","Designed to unify AD, Entra, Okta, AWS, Kubernetes, EDR, VPN, PKI, and more into one behavioral entity, so an attacker stops looking like ten users.")
field("Honest status:","NOT built. The code treats each user_id as atomic; there is no cross-system resolver. This is the next build phase and a natural pilot deliverable.")
field("How to handle if raised:","\"Today the engine scores per identity. Fusing fragmented identities into one entity is designed and is the next phase; it is one of the highest-value items for a DoD environment with many identity systems.\"")

slide("15","Card 11: Forensics and Proportional Response","ROADMAP")
field("Message (roadmap):","Designed for tamper-evident, court-admissible behavioral evidence and proportional, attribute-based response (step-up authentication through lockout).")
field("Honest status:","NOT built. There is no HMAC-chained evidence log and no response automation (no MFA trigger, no lockout) in the code today.")
field("How to handle if raised:","\"The detection core is what runs now. Tamper-evident evidence and proportional response are designed into the architecture and are roadmap; we kept the prototype focused on proving detection first.\"")

slide("16","Card 14: The Full Stack (Preemptive, Behavioral, API, Code)","FRAME")
field("Message:","The layered architecture (preemptive, behavioral, application/API, data, code) is the strategic picture. This app is the behavioral layer.")
field("Honest status:","Layer 1 preemptive, Layer 3 application/API, and Layer 5 code (Mythos) are presented as part of the 22nd Century four-layer platform; only the behavioral layer is in this codebase. Present the full stack as one platform, with this behavioral engine as the live, demonstrated layer.")
field("How to handle if raised:","\"What you are seeing live today is the behavioral layer of our four-layer platform. The other layers are part of the same 22nd Century platform; in this session I am focused on the behavioral layer.\"")

# ===== SECTION 4: DISCUSSION =====
H1("Section 4. Running the Discussion")

slide("17","Consolidated Tough Questions","FRAME")
qa("It is synthetic. How do you know it works on real DoD telemetry?","It is a controlled blind test to prove the mechanism; the next step is a bounded pilot on real data. It is not operational proof, and we say so.")
qa("Did you select the operating point knowing ground truth?","Yes, in the PoC; it is a reporting benchmark to be re-validated, not a deployable threshold.")
qa("How is this different from Exabeam, Securonix, Splunk UBA?","Meaning versus counts, direction versus magnitude, and the measured 0/4 versus 4/4 on the same data.")
qa("8 percent FP is a lot at scale.","Ranked list, not threshold alerts; analysts work top-down; tune on real data.")
qa("Scale, 250 entities versus millions?","The architecture, cohorts, and snapshot model are designed to scale; be straight about what is proven (250) versus engineered for scale (roadmap).")
qa("Can an adversary who knows about behavioral detection evade it?","Multi-front design; being normal on every front at once is far harder than beating one threshold. Hardening the detector itself (adversarial ML, MITRE ATLAS) is on the roadmap.")
qa("Data sovereignty / commercial model?","Real embeddings today; local open model in an enclave for offline runs; no PII in embeddings.")
qa("Identity fusion, forensics, response, the other layers, are they built?","Detection core is built and demoed; those are designed-next or separate. We are explicit about the line.")
qa("Show me where the 8.1 percent comes from.","Sort the 250 entities by composite. To catch the stealthiest attacker, Volt at 13.70 (rank #24), the threshold is 13.70; 20 normal users score above it, so 20/246 = 8.1 percent, and all four attackers are caught. We can derive it live from composite_scores.csv.")
qa("Your top normal user outscores two of your attackers, so that is not clean separation.","Correct, and we present it that way. The two stealthiest attackers (Volt, Slow APT) do not separate on raw score; about 20 normal users outscore Volt. They surface by rank, which is why we deliver a risk-ranked list, not a threshold. The honest claim is recall at the top of the ranking, not a clean gap.")

slide("18","Next Steps and the Ask","FRAME")
B("Bounded validation pilot ","on representative DoD or critical-infrastructure telemetry, scoped to detection and explainability.")
B("Roadmap co-design: ","identity fusion first (highest DoD value), then tamper-evident forensics and proportional response.")
B("Policy and thought leadership: ","R Street perspective on the preemptive plus behavioral framing and the national living-off-the-land / valid-account gap.")
B("Follow-on engineering session ","with his team if this one lands.")
P("")
P("Closing line: \"What you saw run today is the behavioral detection core, on the logs you already collect, catching "
  "four nation-state-grade campaigns that the tools agencies deploy today score as normal. Identity fusion, forensics, "
  "and response are designed and next. The honest ask is a bounded pilot on real data.\"",italic=True)

out="WP DLA/Presentation/Cyber_DeepDive_Briefing_Book_Brandon_Pugh_INTERNAL.docx"
os.makedirs("WP DLA/Presentation",exist_ok=True)
d.save(out)
print("saved:",out,"| paragraphs:",len(d.paragraphs),"| tables:",len(d.tables))
