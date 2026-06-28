#!/usr/bin/env python3
"""Build Prescriptive Maintenance Whitepaper (Word Document).

Behavioral Entity Intelligence (BEI) applied to military equipment
prescriptive maintenance — predicting failures before they impact
mission readiness. Targets DLA leadership, Army logistics, and USSOCOM.

Output: docs/Prescriptive_Maintenance_Whitepaper.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
BLUE = RGBColor(0x1B, 0x4F, 0x72)
TEAL = RGBColor(0x0E, 0x6B, 0x8A)
GOLD = RGBColor(0xB7, 0x95, 0x0B)
RED_ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREEN_ACCENT = RGBColor(0x1E, 0x8A, 0x49)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)

OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
                           "Prescriptive_Maintenance_Whitepaper.docx")


# ── Helper functions (shared pattern from build_tech_spec.py) ────────


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
    return h


def add_body_text(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def add_callout(doc, text, border_color_hex="0E6B8A"):
    """Callout box with left border styling (from build_ueba_whitepaper.py)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="12" w:color="{border_color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EAF4F7" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    run.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return p


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "0D1B2A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EAF0F6")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── TITLE PAGE ─────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Behavioral Entity Intelligence\nfor Prescriptive Maintenance"
    )
    run.font.size = Pt(30)
    run.font.color.rgb = NAVY
    run.bold = True

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Predicting Equipment Failures Through Behavioral Profiling,\n"
        "Temporal Trajectory Analysis, and Supply Chain Integration"
    )
    run.font.size = Pt(15)
    run.font.color.rgb = BLUE

    for _ in range(4):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "22nd Century Technologies, Inc.\n\n"
        "June 2026 — Version 1.0\n\n"
        "Prepared for: Defense Logistics Agency | US Army | USSOCOM"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY

    add_page_break(doc)

    # ── TABLE OF CONTENTS ──────────────────────────────────────────
    add_section_heading(doc, "Table of Contents", level=1)

    toc_entries = [
        "1. Executive Summary",
        "2. The Maintenance Challenge",
        "   2.1. Why Condition-Based Maintenance Misses Failures",
        "   2.2. Mission Impact of Unplanned Failures",
        "   2.3. The Prescriptive Maintenance Vision",
        "3. Equipment as a Behavioral Entity",
        "   3.1. The Entity Model",
        "   3.2. Equipment Behavioral Zones (5 Zones)",
        "   3.3. Zone Divergence — The Key Detection Signal",
        "4. Temporal Trajectory and Regime Detection",
        "   4.1. Velocity and Acceleration",
        "   4.2. Regime Classification",
        "   4.3. Remaining Useful Life Estimation",
        "5. Relationship-Level Analysis",
        "   5.1. Equipment × Component Relationships",
        "   5.2. Equipment × Environment Relationships",
        "   5.3. Component × Supplier Relationships",
        "   5.4. Fleet-Level Network Analysis",
        "6. Supply Chain Integration",
        "   6.1. Predicted Maintenance → Parts Pre-Positioning",
        "   6.2. Demand Decomposition for Maintenance",
        "   6.3. The Maintenance-Supply Chain Feedback Loop",
        "7. Platform-Specific Applications",
        "8. Operational Deployment",
        "   8.1. Data Requirements",
        "   8.2. Deployment Architecture",
        "   8.3. Deployment Phases",
        "9. Business Value and Mission Impact",
        "10. Conclusion",
    ]
    for entry in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(entry)
        run.font.size = Pt(10.5)
        run.font.color.rgb = NAVY if not entry.startswith("   ") else BLUE
        p.paragraph_format.space_after = Pt(2)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "1. Executive Summary", level=1)

    add_body_text(doc, (
        "Equipment failures in deployed environments directly degrade mission readiness. "
        "When a generator fails at a Forward Operating Base, when a Stryker drivetrain "
        "seizes during a movement-to-contact, when an F-35 engine requires unplanned "
        "maintenance during sustained combat operations — the cost is not measured "
        "solely in dollars. It is measured in mission capability, operational tempo, "
        "and the safety of the men and women who depend on that equipment."
    ))

    add_body_text(doc, (
        "Traditional condition-based maintenance (CBM) sets thresholds on individual "
        "sensor readings — temperature must stay below 200°F, vibration below "
        "0.5g, oil pressure above 40 PSI. Equipment fails when the RELATIONSHIP between "
        "readings changes, not when any single value exceeds a limit. A turbine running "
        "below its temperature threshold whose vibration-to-temperature relationship has "
        "drifted steadily over weeks is approaching failure — yet no single sensor "
        "crosses a threshold."
    ))

    add_callout(doc,
        "The turbine is not too hot. It is not vibrating too much. But the relationship "
        "between temperature and vibration has changed — and that behavioral shift "
        "is the precursor to failure."
    )

    add_body_text(doc, (
        "Behavioral Entity Intelligence (BEI) applies the same behavioral intelligence "
        "framework proven in cybersecurity and supply chain analysis to equipment health "
        "monitoring. Each piece of equipment is treated as a behavioral entity with its "
        "own unique profile, built from operating history across multiple sensor dimensions. "
        "BEI tracks how equipment behaves over time — not just whether individual "
        "readings are within spec — using entity profiles, temporal trajectory "
        "analysis, relationship dynamics, and multi-phase scoring to detect degradation "
        "patterns weeks before failure occurs."
    ))

    add_body_text(doc, (
        "The critical link: equipment behavioral drift predicts WHAT parts will be "
        "needed WHEN — connecting maintenance prediction directly to supply chain "
        "pre-positioning. When BEI detects that a specific engine's thermal-to-vibration "
        "relationship is following the same trajectory as engines that required turbine "
        "blade replacement, DLA can initiate parts procurement and depot pre-positioning "
        "weeks before the maintenance order is generated."
    ), bold=True)

    add_callout(doc,
        "Scope of this paper: This is a concept paper. It applies the Behavioral Entity "
        "Intelligence framework — proven in an operational cybersecurity deployment — to "
        "equipment maintenance by analogy. No maintenance-specific proof of concept has "
        "been conducted. The behavioral signatures, regime classifications, and value "
        "ranges described here are illustrative of the approach and drawn from established "
        "predictive-maintenance industry experience; BEI-specific performance would be "
        "established through a proof of concept on platform sensor data."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  2. THE MAINTENANCE CHALLENGE
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "2. The Maintenance Challenge", level=1)

    add_section_heading(doc, "2.1. Why Condition-Based Maintenance Misses Failures", level=2)

    add_body_text(doc, (
        "Condition-based maintenance represents a significant improvement over "
        "time-based maintenance — replacing parts based on actual condition "
        "rather than arbitrary schedules. However, CBM as currently implemented "
        "in military applications has four fundamental limitations that allow "
        "failures to occur despite continuous monitoring."
    ))

    add_body_text(doc, "Individual Sensor Thresholds", bold=True, space_after=2)
    add_body_text(doc, (
        "CBM systems monitor temperature, vibration, pressure, fluid levels, and "
        "other parameters independently. Each sensor has an upper and lower threshold "
        "defined by the manufacturer or by fleet-wide statistical analysis. An alert "
        "fires only when an individual reading exceeds its threshold. This approach "
        "fundamentally misses failures caused by changes in the relationship between "
        "sensors — not by any individual sensor exceeding its limit."
    ))

    add_body_text(doc, "Correlation Blindness", bold=True, space_after=2)
    add_body_text(doc, (
        "The relationship between RPM, exhaust gas temperature, and fuel flow rate "
        "predicts compressor degradation in turbine engines. The relationship between "
        "transmission temperature, drivetrain vibration frequency, and vehicle speed "
        "predicts gearbox failure in ground vehicles. No single value in these "
        "relationships is abnormal — the abnormality exists only in how the "
        "values relate to each other and how that relationship has changed over time."
    ))

    add_body_text(doc, "Static Baselines", bold=True, space_after=2)
    add_body_text(doc, (
        "Manufacturer thresholds are designed for a generic unit operating under "
        "reference conditions. They do not account for THIS specific unit's unique "
        "operating history, environmental exposure, maintenance record, and duty "
        "cycle. A generator that has operated at high altitude for 18 months develops "
        "different thermal characteristics than the same model at sea level. "
        "Population-level thresholds cannot capture unit-specific degradation patterns."
    ))

    add_body_text(doc, "Regime Blindness", bold=True, space_after=2)
    add_body_text(doc, (
        "Gradual degradation over weeks is invisible to point-in-time readings. A "
        "bearing whose vibration amplitude increases by 0.002g per week will not "
        "cross a threshold for months — but the steady upward trajectory is "
        "a reliable failure predictor. CBM systems that evaluate each reading "
        "independently cannot detect this gradual regime shift."
    ))

    add_section_heading(doc, "2.2. Mission Impact of Unplanned Failures", level=2)

    add_body_text(doc, (
        "Unplanned equipment failures create cascading impacts across the mission "
        "chain. The following table illustrates the operational consequences across "
        "major platform categories."
    ))

    create_table(doc,
        ["Platform", "Example Systems", "Failure Impact", "Readiness Effect"],
        [
            ["Aircraft Engine", "F-35 F135,\nF-16 F110,\nC-130 T56",
             "Aircraft grounded. Sortie generation\nrate drops. Spare engine\nrotation disrupted.",
             "Every grounded aircraft reduces\nsquadron combat capability.\n"
             "Engine swap: 72-96 hours."],
            ["Ground Vehicle", "Stryker, MRAP,\nJLTV drivetrain",
             "Vehicle deadlined. Unit loses\nmobility and firepower.\nRecovery assets diverted.",
             "Deadlined vehicle in a BCT\ndirectly reduces maneuver\nflexibility."],
            ["Generator", "MEP-series,\nFOB power systems",
             "Power loss at FOB. Communications,\nmedical, C2 systems affected.\nBackup gen capacity stressed.",
             "Single generator failure can\ncascade to mission-critical\nsystem outages."],
            ["Naval Propulsion", "Gas turbines,\nshaft assemblies",
             "Reduced speed capability.\nMission abort or delay.\nDrydock required for major repairs.",
             "Propulsion casualty limits\noperational range and\nresponse time."],
            ["Communications", "SATCOM terminals,\nradar systems",
             "Loss of command and control.\nDegraded situational awareness.\nForced to backup systems.",
             "Comms failure in contested\nenvironment is a direct\nthreat to force protection."],
        ],
        col_widths=[1.0, 1.2, 2.2, 2.2],
    )

    add_body_text(doc, (
        "The cost differential between planned and unplanned maintenance is well "
        "documented across military and industrial applications: unplanned maintenance "
        "typically costs 3-10x more than equivalent planned maintenance. This multiplier "
        "accounts for expedited parts shipping, overtime labor, disrupted maintenance "
        "schedules, and the secondary damage caused by operating equipment past its "
        "failure point."
    ))

    add_body_text(doc, (
        "Mission readiness rates are directly linked to maintenance effectiveness. "
        "Every 1% improvement in Operational Readiness (OR) rate translates to "
        "significant combat capability — more aircraft available for tasking, "
        "more vehicles available for movement, more generators available for power. "
        "According to GAO analyses, approximately 40% of maintenance delays are "
        "parts-related: the maintenance team identifies the failure but cannot execute "
        "the repair because the required part is not in stock at the right location."
    ), bold=True)

    add_section_heading(doc, "2.3. The Prescriptive Maintenance Vision", level=2)

    add_body_text(doc, (
        "Prescriptive maintenance extends beyond prediction to action. While predictive "
        "maintenance answers 'WHEN will failure occur?', prescriptive maintenance "
        "answers three questions simultaneously:"
    ))

    add_bullet(doc, "WHEN will failure occur? — Temporal trajectory analysis "
               "extrapolates current degradation velocity and acceleration to estimate "
               "remaining useful life with confidence intervals.",
               bold_prefix="Predictive: ")
    add_bullet(doc, "WHAT action should be taken NOW? — Zone-specific diagnosis "
               "identifies which component or subsystem is degrading, directing "
               "maintenance crews to the specific intervention required.",
               bold_prefix="Prescriptive: ")
    add_bullet(doc, "WHERE should parts be staged? — Behavioral drift predictions "
               "automatically generate demand signals for DLA, triggering parts "
               "pre-positioning at the depot closest to the affected unit's location.",
               bold_prefix="Supply-Integrated: ")

    add_callout(doc,
        "Prescriptive maintenance does not just predict failure — it tells the "
        "maintainer what to do, tells the supply chain what to ship, and tells the "
        "commander how long until the equipment is back in service."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  3. EQUIPMENT AS A BEHAVIORAL ENTITY
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "3. Equipment as a Behavioral Entity", level=1)

    add_section_heading(doc, "3.1. The Entity Model", level=2)

    add_body_text(doc, (
        "In the BEI framework, each piece of equipment is treated as a behavioral "
        "entity — identified by its unique serial number, not just its model or "
        "type designation. This distinction is critical: two F-35 engines of identical "
        "manufacture will develop different behavioral profiles based on their "
        "individual operating histories, environmental exposures, maintenance records, "
        "and duty cycles."
    ))

    add_body_text(doc, (
        "Each entity develops its own behavioral baseline from operating history. "
        "This baseline captures HOW this specific unit operates — not just WHETHER "
        "it is within manufacturer specifications. The baseline includes the "
        "relationships between sensor readings, the typical response curves under "
        "different load conditions, the characteristic thermal signatures during "
        "startup and shutdown, and the vibration spectrum under normal operation."
    ))

    add_body_text(doc, (
        "When a maintainer asks 'Is this engine healthy?', traditional CBM answers "
        "'All readings are within spec.' BEI is designed to answer with behavioral "
        "context — illustratively: 'This engine's behavioral profile has drifted "
        "measurably from its established baseline over recent weeks, concentrated in "
        "the thermal-vibration relationship zone, following a trajectory similar to "
        "other engines that subsequently required compressor blade replacement.' The "
        "specific drift magnitudes and lead times such a system reports would be "
        "established and calibrated through a proof of concept on platform data."
    ))

    add_section_heading(doc, "3.2. Equipment Behavioral Zones (5 Zones)", level=2)

    add_body_text(doc, (
        "Equipment behavior is decomposed into five independent zones, each capturing "
        "a distinct dimension of operational health. Zone decomposition enables the "
        "system to identify WHICH subsystem is degrading — not just that "
        "degradation is occurring."
    ))

    create_table(doc,
        ["Zone", "Sensor Inputs", "What It Captures"],
        [
            ["Vibration\nProfile",
             "Frequency spectrum, amplitude\npatterns, harmonic relationships,\n"
             "bearing signature frequencies",
             "Mechanical health of rotating components.\n"
             "Bearing wear, shaft misalignment, blade\n"
             "damage, gear mesh degradation."],
            ["Thermal\nSignature",
             "Temperature distributions, heating\nand cooling rates, cross-component\n"
             "thermal relationships",
             "Heat transfer efficiency and fluid system\n"
             "health. Coolant flow restrictions, exhaust\n"
             "path blockage, insulation degradation."],
            ["Power /\nPerformance",
             "Fuel consumption, output efficiency,\n"
             "RPM-to-output ratios, power factor,\n"
             "load response curves",
             "Energy conversion efficiency.\n"
             "Compressor degradation, injector fouling,\n"
             "generator winding degradation."],
            ["Fluid\nChemistry",
             "Oil analysis (metals, viscosity),\ncoolant chemistry, hydraulic\n"
             "fluid degradation markers",
             "Internal wear products and fluid\n"
             "degradation. Metal particle counts indicate\n"
             "specific component wear origins."],
            ["Operational\nDuty Cycle",
             "Utilization patterns, load distribution,\nstart/stop frequency,\n"
             "environmental exposure",
             "Operating stress history. High duty-cycle\n"
             "units degrade differently than intermittent-\n"
             "use units of the same model."],
        ],
        col_widths=[1.0, 2.5, 3.0],
    )

    add_section_heading(doc, "3.3. Zone Divergence — The Key Detection Signal", level=2)

    add_body_text(doc, (
        "The most powerful diagnostic signal is not the magnitude of drift in any "
        "single zone, but the pattern of divergence between zones. Zone divergence "
        "reveals the specific failure mode developing and directs both maintenance "
        "action and parts procurement."
    ))

    add_body_text(doc, "When ALL zones drift together:", bold=True, space_after=2)
    add_body_text(doc, (
        "Normal aging. The equipment is gradually wearing across all subsystems "
        "simultaneously. This is expected, predictable, and addressable through "
        "scheduled overhaul at the manufacturer-recommended interval. No specific "
        "part needs to be procured ahead of schedule."
    ))

    add_body_text(doc, "When ONE zone drifts while others remain stable:", bold=True,
                  space_after=2)
    add_body_text(doc, (
        "Specific component degradation. The drifting zone identifies which subsystem "
        "is failing, and the stable zones confirm that the rest of the equipment is "
        "healthy. This zone-specific diagnosis directs maintenance action and predicts "
        "which SPECIFIC part is needed."
    ))

    create_table(doc,
        ["Drift Pattern", "Diagnosis", "Maintenance Action", "Parts Required"],
        [
            ["Thermal drifting +\nVibration stable",
             "Fluid system degradation.\nCoolant flow restriction\nor heat exchanger fouling.",
             "Inspect fluid system.\nFlush coolant lines.\nCheck heat exchanger.",
             "Coolant pump, heat\nexchanger core,\nthermostat assembly"],
            ["Vibration drifting +\nThermal stable",
             "Bearing wear or shaft\nmisalignment. Mechanical\ndegradation without\nthermal impact yet.",
             "Inspect bearings and\nshaft alignment.\nReplace bearings before\ncascade to thermal.",
             "Bearing sets, shaft\nseals, alignment\nshims"],
            ["Power drifting +\nAll others stable",
             "Efficiency loss without\nmechanical cause.\nFuel system or\nelectrical degradation.",
             "Inspect fuel injectors,\nfilters, winding\nresistance.",
             "Fuel injectors, filters,\nstator windings"],
            ["Vibration + Thermal\nboth drifting",
             "Cascading failure.\nMechanical wear causing\nsecondary thermal effects.",
             "Accelerated inspection.\nSchedule within 2 weeks.\nMultiple subsystems\nat risk.",
             "Full bearing kit,\ncoolant components,\nseals — order\nimmediately"],
        ],
        col_widths=[1.3, 1.8, 1.8, 1.6],
    )

    add_callout(doc,
        "The capability this paper describes would transform maintenance from "
        "'something is wrong' to, illustratively: 'the bearing in the #2 accessory "
        "gearbox is degrading, and based on its behavioral trajectory it will require "
        "replacement on a predictable timeline rather than failing without warning — "
        "and the bearing kit can be ordered and pre-positioned before the failure.'"
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  4. TEMPORAL TRAJECTORY AND REGIME DETECTION
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "4. Temporal Trajectory and Regime Detection", level=1)

    add_section_heading(doc, "4.1. Velocity and Acceleration", level=2)

    add_body_text(doc, (
        "BEI tracks behavioral change over time using the same trajectory analysis "
        "concepts applied in cybersecurity behavioral monitoring. For each equipment "
        "entity, the system computes:"
    ))

    add_bullet(doc, "How fast is behavioral change occurring? The rate of drift from "
               "baseline measured across all zones and their inter-relationships. A "
               "bearing whose vibration profile is changing slowly is on a very different "
               "trajectory than one changing several times faster, even if both are "
               "currently within specification.",
               bold_prefix="Drift Velocity: ")
    add_bullet(doc, "Is the rate of degradation itself increasing? Positive acceleration "
               "indicates a failure cascade is approaching — degradation feeds on "
               "itself as component wear creates secondary stress on adjacent systems. "
               "Acceleration is the strongest predictor of imminent failure.",
               bold_prefix="Acceleration: ")
    add_bullet(doc, "Is degradation monotonic (steady decline) or oscillating "
               "(intermittent fault)? Monotonic drift indicates progressive wear; "
               "oscillating patterns suggest intermittent faults such as loose "
               "connections, sticking valves, or thermal cycling damage.",
               bold_prefix="Trend Consistency: ")

    add_section_heading(doc, "4.2. Regime Classification", level=2)

    add_body_text(doc, (
        "Based on trajectory analysis, each equipment entity is classified into "
        "one of four operational regimes. The regime classification drives maintenance "
        "scheduling and supply chain action."
    ))

    create_table(doc,
        ["Regime", "Criteria", "Maintenance Response", "Supply Chain Action"],
        [
            ["STABLE", "All zones within baseline\nenvelope. Low velocity,\nno acceleration.",
             "Normal operations.\nMaintenance at next\nscheduled interval.",
             "No action required.\nStandard inventory\nlevels sufficient."],
            ["DRIFTING", "One or more zones showing\nsteady drift from baseline.\nLow-moderate velocity.",
             "Schedule inspection at\nnext available opportunity.\nMonitor weekly.",
             "Verify parts availability.\nFlag for potential\nprocurement."],
            ["ACCELERATING", "Degradation rate itself\nis increasing. Positive\nacceleration detected.",
             "Schedule maintenance\nwithin 2 weeks.\nRestrict operational\ntempo if possible.",
             "Initiate parts order.\nPre-position at nearest\ndepot. Expedite if lead\ntime exceeds 2 weeks."],
            ["REGIME_SHIFT", "Fundamental behavioral\nchange detected. Before/\nafter discontinuity\nin profile.",
             "Immediate inspection\nrequired. Ground or\ndeadline equipment\nuntil inspected.",
             "Emergency parts request.\nCross-level from other\nlocations if not in\nstock."],
        ],
        col_widths=[1.1, 1.8, 1.8, 1.8],
    )

    add_section_heading(doc, "4.3. Remaining Useful Life Estimation", level=2)

    add_body_text(doc, (
        "BEI estimates remaining useful life (RUL) by extrapolating the current "
        "behavioral trajectory forward. Unlike traditional RUL estimation based on "
        "time-since-overhaul or cycle counts, BEI RUL is based on actual behavioral "
        "state — accounting for the specific operating history, environmental "
        "stress, and degradation pattern of each individual unit."
    ))

    add_bullet(doc, "Given current drift velocity and acceleration, when does the unit "
               "enter the ACCELERATING or REGIME_SHIFT regime? This intersection point "
               "provides the estimated time-to-failure.",
               bold_prefix="Behavioral Trajectory Extrapolation: ")
    add_bullet(doc, "RUL estimates include uncertainty bounds. A unit with highly "
               "consistent (monotonic) drift has tight confidence intervals; a unit "
               "with oscillating patterns has wider bounds, reflecting greater "
               "uncertainty.",
               bold_prefix="Confidence Intervals: ")
    add_bullet(doc, "BEI identifies other units in the fleet at similar behavioral "
               "stages — units whose zone profiles and drift trajectories most "
               "closely resemble the unit under analysis. What happened to those peer "
               "units provides empirical validation of the RUL estimate.",
               bold_prefix="Comparison to Fleet Peers: ")

    add_callout(doc,
        "Traditional RUL answers from manufacturer cycle limits: 'This engine has a "
        "fixed number of flight hours remaining.' A behavioral RUL approach would "
        "instead answer from the unit's own trajectory and its fleet peers — "
        "illustratively: 'This engine's behavioral trajectory resembles peer units that "
        "subsequently required compressor blade replacement, indicating a finite and "
        "estimable window of useful life rather than a generic cycle count.' The "
        "precision of such estimates would be established through a proof of concept."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  5. RELATIONSHIP-LEVEL ANALYSIS
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "5. Relationship-Level Analysis", level=1)

    add_body_text(doc, (
        "BEI extends beyond individual equipment analysis to track relationships "
        "between equipment, components, environments, suppliers, and fleet cohorts. "
        "Relationship analysis reveals patterns invisible when analyzing individual "
        "entities in isolation."
    ))

    add_section_heading(doc, "5.1. Equipment × Component Relationships", level=2)

    add_body_text(doc, (
        "Equipment-Component relationship analysis tracks which specific components "
        "within each equipment entity are contributing to behavioral drift. This "
        "enables root-cause diagnosis and predicts cascading failures."
    ))

    add_bullet(doc, "Which components are degrading on THIS specific unit? "
               "Zone divergence patterns map to specific Line Replaceable Units (LRUs) "
               "and subsystems based on the sensor inputs feeding each zone.",
               bold_prefix="Component Isolation: ")
    add_bullet(doc, "Historical analysis reveals which components tend to fail "
               "together. When Component A degrades, Component B follows within "
               "a predictable window. This enables proactive replacement of the "
               "secondary component during the primary repair.",
               bold_prefix="Failure Correlation: ")
    add_bullet(doc, "If the #1 bearing in the accessory gearbox fails, what else is "
               "at risk? BEI maintains a cascade probability model built from fleet "
               "behavioral history, enabling maintainers to inspect and replace "
               "adjacent components before secondary failures occur.",
               bold_prefix="Cascade Prediction: ")

    add_section_heading(doc, "5.2. Equipment × Environment Relationships", level=2)

    add_body_text(doc, (
        "The operating environment directly influences degradation patterns. BEI "
        "tracks how equipment behavior varies by location, climate, altitude, and "
        "operational tempo."
    ))

    add_bullet(doc, "Does the operating environment accelerate degradation? Equipment "
               "deployed in desert environments may show faster fluid chemistry "
               "degradation due to sand ingestion and thermal stress. Equipment in "
               "maritime environments may show accelerated corrosion signatures in "
               "the vibration zone.",
               bold_prefix="Equipment × Location: ")
    add_bullet(doc, "Does the usage pattern affect the failure mode? Equipment operated "
               "at sustained high load degrades differently than equipment with "
               "frequent start-stop cycles. The same model in different duty cycles "
               "develops different behavioral profiles and different failure modes.",
               bold_prefix="Equipment × Duty Cycle: ")

    add_body_text(doc, (
        "Two identical Stryker engines may show dramatically different degradation "
        "patterns: one deployed in sustained convoy operations in high-altitude desert "
        "terrain, the other in intermittent patrol operations in temperate lowland. "
        "The relationship between equipment and environment explains why the same "
        "model fails differently in different operating contexts."
    ), italic=True)

    add_section_heading(doc, "5.3. Component × Supplier Relationships", level=2)

    add_body_text(doc, (
        "BEI tracks the performance of components by supplier, identifying quality "
        "patterns that may not be visible in individual failure reports but become "
        "apparent through behavioral analysis across the fleet."
    ))

    add_bullet(doc, "Do parts from Supplier A degrade faster than parts from Supplier B? "
               "BEI compares the behavioral trajectories of equipment using components "
               "from different suppliers, revealing quality differences invisible in "
               "individual inspection reports.",
               bold_prefix="Supplier Quality Comparison: ")
    add_bullet(doc, "Is a supplier's component quality degrading over time for a "
               "specific part number? BEI tracks supplier performance longitudinally, "
               "detecting gradual quality erosion before it manifests as increased "
               "failure rates.",
               bold_prefix="Quality Drift Detection: ")

    add_body_text(doc, (
        "This analysis directly connects to supply chain risk assessment. When BEI "
        "detects that components from a particular supplier are associated with "
        "accelerated behavioral drift across multiple equipment entities, it generates "
        "a supplier quality signal that feeds into procurement decisions and vendor "
        "evaluation processes."
    ), bold=True)

    add_section_heading(doc, "5.4. Fleet-Level Network Analysis", level=2)

    add_body_text(doc, (
        "Fleet-level analysis examines patterns across the entire population of "
        "equipment entities, enabling predictions and interventions that would be "
        "impossible from individual unit analysis alone."
    ))

    add_bullet(doc, "Units with similar behavioral profiles are automatically clustered "
               "into cohorts. These cohorts may share common characteristics "
               "(same model, same operating environment, same maintenance history) "
               "or may reveal previously unknown groupings.",
               bold_prefix="Cohort Analysis: ")
    add_bullet(doc, "If 10% of units in a behavioral cohort have already experienced "
               "a specific failure, the remaining units in that cohort are at elevated "
               "risk for the same failure. This enables fleet-wide predictive action "
               "based on leading indicators from peer units.",
               bold_prefix="Fleet-Wide Failure Prediction: ")
    add_bullet(doc, "When multiple units across different locations develop similar "
               "behavioral drift simultaneously, BEI identifies the common cause: "
               "shared defect, shared component batch, shared environmental stressor, "
               "or shared maintenance procedure.",
               bold_prefix="Common Cause Identification: ")

    add_callout(doc,
        "Fleet-level analysis turns individual equipment monitoring into an "
        "enterprise intelligence capability. When BEI detects that 15 UH-60 engines "
        "across 3 installations are developing similar vibration-zone drift, it "
        "identifies the common component batch and triggers proactive inspection "
        "of all affected units — before the first failure occurs."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  6. SUPPLY CHAIN INTEGRATION
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "6. Supply Chain Integration", level=1)

    add_body_text(doc, (
        "The integration of equipment behavioral intelligence with supply chain "
        "operations is the force multiplier that distinguishes prescriptive "
        "maintenance from predictive maintenance. Predicting failure without "
        "ensuring parts availability achieves nothing — the maintainer knows "
        "the equipment will fail but cannot prevent it because the part is not "
        "in stock."
    ))

    add_section_heading(doc, "6.1. Predicted Maintenance → Parts Pre-Positioning", level=2)

    add_body_text(doc, (
        "BEI behavioral drift predictions generate three demand signals that "
        "directly feed DLA supply chain operations:"
    ))

    create_table(doc,
        ["Demand Signal", "Source", "DLA Action"],
        [
            ["WHAT parts\nwill be needed",
             "Zone divergence pattern maps to\nspecific LRUs and components.\n"
             "Thermal drifting → coolant pump,\nheat exchanger.\n"
             "Vibration drifting → bearing kit.",
             "Identify NSN (National Stock\nNumber) for predicted parts.\n"
             "Check current inventory\nposition across depots."],
            ["WHEN parts\nwill be needed",
             "Regime trajectory and RUL\nestimation. DRIFTING regime\n"
             "indicates a longer horizon;\nACCELERATING indicates a\nnear-term window.",
             "Align procurement lead time\nto predicted need date.\n"
             "Initiate order if lead time\nexceeds available window."],
            ["WHERE parts\nshould be staged",
             "Equipment location data from\nGCSS-Army or unit reporting.\n"
             "Current and projected unit\nlocation over RUL window.",
             "Pre-position at depot or SSA\nclosest to affected unit.\n"
             "Cross-level from excess\nlocations if needed."],
        ],
        col_widths=[1.2, 2.5, 2.8],
    )

    add_body_text(doc, (
        "This is the critical link between prescriptive maintenance and supply "
        "chain optimization. Today, the supply chain responds to maintenance "
        "orders — reactive. With BEI, the supply chain responds to behavioral "
        "predictions — proactive. Parts are moving toward the depot before "
        "the maintainer submits the work order."
    ), bold=True)

    add_section_heading(doc, "6.2. Demand Decomposition for Maintenance", level=2)

    add_body_text(doc, (
        "Parts demand for military equipment is composed of multiple independent "
        "demand components. Accurate demand forecasting requires decomposing total "
        "demand into its constituent drivers."
    ))

    create_table(doc,
        ["Demand Component", "Description", "BEI Contribution"],
        [
            ["Baseline", "Routine consumption from\nnormal operations and\nscheduled maintenance",
             "BEI validates baseline assumptions\nby comparing predicted vs. actual\nconsumption rates."],
            ["Exercise", "Surge demand during major\nexercises and training events",
             "BEI adjusts RUL estimates based\non increased operational tempo\nduring exercise periods."],
            ["Seasonal", "Climate-driven demand patterns\n(winter vs. summer, wet vs. dry)",
             "BEI detects environmental\ninfluences on degradation rate\nacross the fleet."],
            ["Maintenance", "Planned maintenance actions\ndriving parts demand",
             "BEI predicts maintenance demand\nweeks before work orders are\ngenerated."],
            ["Combat", "Operational demand during\ndeployment and contingency",
             "BEI monitors accelerated\ndegradation under combat\noperational tempo."],
            ["Failure", "Unplanned failure-driven\ndemand (emergency orders)",
             "BEI REDUCES this component by\npredicting failures before they\noccur — converting failure\ndemand to maintenance demand."],
        ],
        col_widths=[1.1, 2.2, 3.2],
    )

    add_body_text(doc, (
        "BEI-predicted failures feed the 'failure' and 'maintenance' demand components "
        "with greater accuracy and longer lead time than traditional approaches. The "
        "key value proposition: converting unplanned failure demand into planned "
        "maintenance demand reduces emergency procurement costs, eliminates expedited "
        "shipping charges, and improves demand forecasting accuracy across all "
        "components."
    ))

    add_section_heading(doc, "6.3. The Maintenance-Supply Chain Feedback Loop", level=2)

    add_body_text(doc, (
        "BEI creates a closed-loop optimization cycle connecting equipment behavior, "
        "supply chain operations, and maintenance execution."
    ))

    add_bullet(doc, "Equipment behavioral drift generates demand signals "
               "for specific parts at specific locations within specific timeframes.",
               bold_prefix="Equipment Behavior → Parts Demand: ")
    add_bullet(doc, "Parts availability at the right location determines "
               "whether maintenance can be scheduled optimally or must wait for "
               "parts delivery.",
               bold_prefix="Parts Availability → Maintenance Scheduling: ")
    add_bullet(doc, "Completed maintenance resets the equipment's behavioral "
               "baseline. BEI monitors the post-maintenance behavioral profile "
               "to verify that the intervention was effective and that the "
               "equipment has returned to a healthy operating regime.",
               bold_prefix="Maintenance Execution → Equipment Behavior: ")

    add_callout(doc,
        "The closed loop: Equipment behavior predicts parts demand. Parts "
        "availability enables timely maintenance. Maintenance restores "
        "equipment health. BEI monitors the result. The cycle repeats — "
        "each iteration improving prediction accuracy through learned "
        "outcomes."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  7. PLATFORM-SPECIFIC APPLICATIONS
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "7. Platform-Specific Applications", level=1)

    add_body_text(doc, (
        "BEI prescriptive maintenance applies across military platform categories. "
        "The following table details the entity zones, key relationships, detection "
        "signals, and supply chain impacts for each major platform type."
    ))

    create_table(doc,
        ["Platform", "Entity Zones", "Key\nRelationships", "Detection\nSignal",
         "Parts Impact"],
        [
            ["Aircraft Engine\n(F-35, F-16, C-130)",
             "Turbine vibration spectrum,\nexhaust gas temperature,\nfuel flow efficiency,\noil metal particle count",
             "Vibration × Temperature\n× Fuel flow\nrelationship drift",
             "RPM-to-EGT ratio\nshift indicates\ncompressor blade\ndegradation",
             "Turbine blades, bearing\nkits, fuel nozzles,\noil cooler assemblies"],
            ["Ground Vehicle\n(Stryker, MRAP, JLTV)",
             "Drivetrain vibration,\ntransmission thermal,\npower train efficiency,\nhydraulic pressure",
             "Drivetrain vibration\n× Transmission temp\n× Speed ratio",
             "Transmission temp\nrise rate vs. load\nindicates gear wear",
             "Transmission assembly,\nfinal drive seals,\nCV joints, filters"],
            ["Generator\n(MEP-series, FOB)",
             "Load-to-fuel ratio,\nbearing vibration onset,\nexhaust temperature,\noutput voltage stability",
             "Load × Fuel flow\n× Bearing vibration\ncorrelation",
             "Load-to-fuel ratio\ndrift signals\ninjector fouling\nor winding loss",
             "Fuel injectors, voltage\nregulators, bearing\nsets, stator windings"],
            ["Naval Propulsion\n(Gas turbine, shaft)",
             "Shaft vibration spectrum,\nbearing temperature,\nspeed-to-power ratio,\nlube oil analysis",
             "Shaft vibration ×\nSpeed × Sea state\ncorrelation",
             "Bearing temp rise\nnormalized to speed\nand sea state\ndetects wear",
             "Shaft bearings, seals,\nreduction gear\ncomponents, lube oil\ncoolers"],
            ["Communications/Radar\n(SATCOM, ground radar)",
             "RF output drift, power\nsupply thermal signature,\nconnector degradation,\nreceiver sensitivity",
             "RF output × Power\nsupply temp ×\nAmbient conditions",
             "RF output decline\nnormalized to power\ninput signals\namplifier degradation",
             "Power amplifier tubes,\nwaveguides, power\nsupply modules,\nconnector assemblies"],
            ["Rotary Wing\n(UH-60, CH-47)",
             "Rotor vibration spectrum,\ngearbox thermal signature,\nhydraulic pressure,\nswashplate wear",
             "Rotor vibration ×\nGearbox temp ×\nFlight regime",
             "Gearbox thermal\nrise vs. rotor\nvibration detects\ngear mesh wear",
             "Main rotor bearings,\ntail rotor gearbox,\nhydraulic actuators,\nswashplate assembly"],
        ],
        col_widths=[1.2, 1.6, 1.4, 1.3, 1.4],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  8. OPERATIONAL DEPLOYMENT
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "8. Operational Deployment", level=1)

    add_section_heading(doc, "8.1. Data Requirements", level=2)

    add_body_text(doc, (
        "BEI prescriptive maintenance requires integration with existing military "
        "data sources. No new sensor installation is required for initial deployment "
        "— BEI leverages data already collected by existing monitoring systems, "
        "maintenance databases, and logistics applications."
    ))

    create_table(doc,
        ["Data Source", "System of Record", "Data Elements", "Integration Method"],
        [
            ["IoT Sensors /\nSCADA", "Platform-specific\nmonitoring systems",
             "Temperature, vibration,\npressure, flow rate,\nvoltage, current",
             "Automated feed via\nIoT gateway or SCADA\nhistorian export"],
            ["Maintenance Logs", "GCSS-Army, IETMS,\nIMDS (Air Force)",
             "Work orders, fault codes,\nparts replaced, labor hours,\ncompletion dates",
             "Database integration\nvia standard DoD\ndata interfaces"],
            ["Fluid Analysis", "Oil Analysis Program\n(JOAP/AOAP)",
             "Metal particle counts,\nviscosity, water content,\nacid number",
             "Lab result import\nvia standardized\nreporting format"],
            ["CMMS Records", "DPAS, MAXIMO,\nor equivalent",
             "Asset registry, maintenance\nhistory, configuration,\nserial number tracking",
             "API integration or\nscheduled data export"],
            ["Operational Tempo", "Unit reporting,\nflight/drive logs",
             "Hours operated, miles driven,\nsorties flown, load profiles,\ndeployment status",
             "Automated from existing\nusage tracking systems"],
            ["Environmental\nConditions", "Weather services,\ndeployment location data",
             "Temperature, humidity,\naltitude, sand/dust\nexposure, salt air",
             "Geographic lookup\nfrom unit location\nand weather API"],
            ["Parts Usage\nHistory", "DLA EBS,\nFedMall, GCSS-Army",
             "NSN, quantity consumed,\norder lead times,\nsupplier, depot location",
             "DLA data feed via\nestablished logistics\ninterfaces"],
        ],
        col_widths=[1.1, 1.3, 1.8, 1.5],
    )

    add_section_heading(doc, "8.2. Deployment Architecture", level=2)

    add_body_text(doc, (
        "BEI prescriptive maintenance deploys in a tiered architecture that "
        "balances edge responsiveness with enterprise-scale analysis."
    ))

    add_body_text(doc, "Edge Processing (Platform / Base Level)", bold=True, space_after=2)
    add_body_text(doc, (
        "Behavioral snapshot computation occurs at the platform or installation "
        "level. Edge processing ingests sensor data, computes zone-level behavioral "
        "profiles, and performs local regime classification. This enables real-time "
        "REGIME_SHIFT detection and immediate maintenance alerts even when "
        "connectivity to the enterprise cloud is limited or denied."
    ))

    add_body_text(doc, "Cloud Aggregation (Enterprise Level)", bold=True, space_after=2)
    add_body_text(doc, (
        "Fleet-wide analysis, cohort detection, and demand signal generation "
        "occur at the enterprise cloud level. Cloud processing enables cross-fleet "
        "comparison, fleet-wide failure prediction, supplier quality analysis, "
        "and integrated demand forecasting that requires visibility across all "
        "units and installations."
    ))

    add_body_text(doc, "Integration Layer", bold=True, space_after=2)
    add_body_text(doc, (
        "BEI integrates with existing military logistics and maintenance systems "
        "including GCSS-Army (maintenance and supply), DLA's Enterprise Business "
        "System (EBS) for parts procurement, and unit-level maintenance management "
        "systems. Integration is achieved through standard DoD data interfaces "
        "and APIs — no replacement of existing systems is required."
    ))

    add_section_heading(doc, "8.3. Deployment Phases", level=2)

    create_table(doc,
        ["Phase", "Duration", "Activities", "Deliverables"],
        [
            ["Phase 1:\nSensor Integration\n& Entity Registration",
             "2-4 weeks",
             "Connect to existing sensor data\nsources. Register equipment entities\n"
             "by serial number. Validate data\nquality and completeness.",
             "Data pipeline operational.\nEquipment entities registered.\n"
             "Sensor data flowing to BEI."],
            ["Phase 2:\nBehavioral Baseline\nEstablishment",
             "3-6 months\nof operating\nhistory",
             "Collect operating data across\nnormal duty cycles. Build per-entity\n"
             "behavioral profiles. Establish zone\nbaselines and normal operating\nenvelopes.",
             "Per-entity behavioral baselines.\n"
             "Zone-specific normal ranges.\n"
             "Fleet cohort identification."],
            ["Phase 3:\nDetection Deployment\n& Validation",
             "2-4 weeks",
             "Deploy drift detection and regime\nclassification. Validate against\n"
             "known maintenance events.\nCalibrate thresholds per platform type.",
             "Detection algorithms active.\n"
             "Validation report with detection\naccuracy metrics."],
            ["Phase 4:\nSupply Chain\nIntegration",
             "4-8 weeks",
             "Connect BEI demand signals to\nDLA EBS. Implement parts pre-\npositioning "
             "workflow. Establish\nclosed-loop feedback from\nmaintenance outcomes.",
             "End-to-end prescriptive\nmaintenance operational.\n"
             "Demand signals flowing to DLA.\n"
             "Closed-loop feedback active."],
        ],
        col_widths=[1.3, 1.0, 2.5, 2.2],
    )

    add_body_text(doc, (
        "Total deployment timeline from initial sensor integration to fully "
        "operational prescriptive maintenance with supply chain integration: "
        "approximately 6-9 months. The longest phase is baseline establishment "
        "(Phase 2), which requires sufficient operating history to build "
        "statistically meaningful behavioral profiles for each equipment entity."
    ))

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  9. BUSINESS VALUE AND MISSION IMPACT
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "9. Business Value and Mission Impact", level=1)

    add_body_text(doc, (
        "A BEI prescriptive maintenance capability targets value across five "
        "dimensions: reduced unplanned failures, lower maintenance costs, improved "
        "mission readiness, optimized parts inventory, and enhanced supply chain "
        "responsiveness."
    ))

    add_body_text(doc, (
        "The improvement ranges in the table below are not BEI-measured results — no "
        "maintenance proof of concept has been conducted. They are documented outcome "
        "ranges from predictive and condition-based maintenance programs across defense "
        "and industry (as reported in DoD, DOE, and commercial reliability studies), and "
        "they represent the opportunity that a behavioral approach is designed to "
        "capture. BEI-specific outcomes would be established through a proof of concept."
    ), italic=True)

    create_table(doc,
        ["Impact Area", "Current State", "Behavioral Approach",
         "Predictive Maintenance\nBenchmark Range*"],
        [
            ["Unplanned\nFailures",
             "Failures detected at point\nof occurrence. No advance\nwarning. Emergency\nmaintenance required.",
             "Behavioral early warning\nprovides weeks of lead time.\nFailures prevented or\nconverted to planned events.",
             "40-60% reduction in\nunplanned failures\nacross monitored fleet"],
            ["Maintenance\nCost",
             "Unplanned maintenance costs\n3-10x planned equivalent.\nExpedited parts, overtime\nlabor, secondary damage.",
             "Prescriptive action enables\nplanned maintenance at\noptimal time and cost.\nSecondary damage prevented.",
             "25-40% reduction in\ntotal maintenance cost\nper platform"],
            ["Mission\nReadiness",
             "OR rates limited by\nunplanned downtime.\nEquipment unavailable\nwhen needed most.",
             "Fewer equipment-down days.\nMaintenance scheduled during\nplanned downtime windows.\nHigher fleet availability.",
             "3-8% improvement in\nOperational Readiness\nrates"],
            ["Parts\nInventory",
             "High safety stock levels\nto buffer against demand\nuncertainty. Excess\ninventory ties up capital.",
             "Demand forecasting accuracy\nimprovement reduces safety\nstock requirements. Right\nparts at right locations.",
             "15-25% reduction in\nsafety stock value\nwithout service\nlevel degradation"],
            ["Supply Chain\nResponsiveness",
             "Supply chain responds to\nmaintenance orders after\nfailure is identified.\nReactive posture.",
             "Predicted demand signals\nenable pre-positioning.\nParts moving before work\norder is generated.",
             "2-4 week improvement\nin parts availability\nlead time"],
        ],
        col_widths=[1.0, 1.8, 1.8, 1.5],
    )

    add_body_text(doc, (
        "*Benchmark ranges from predictive and condition-based maintenance programs "
        "across defense and industry; not BEI-measured results."
    ), italic=True, space_after=2)

    add_callout(doc,
        "The compound effect a behavioral approach is designed to achieve: not just "
        "predicting failures, but converting unplanned failure demand into planned "
        "maintenance demand, reducing inventory carrying costs, improving depot "
        "utilization, and increasing the number of mission-ready platforms available "
        "to the warfighter at any given time."
    )

    add_body_text(doc, (
        "The platform figures below are illustrative, derived from published DoD "
        "maintenance cost structures and the predictive-maintenance benchmark ranges "
        "above. They indicate the scale of opportunity per platform category — not "
        "projected BEI results."
    ))

    create_table(doc,
        ["Platform Category", "Avg Unplanned\nMaintenance Cost", "Illustrative Annual\nSavings Range*",
         "Readiness\nImprovement Range"],
        [
            ["Aircraft Engine\n(fighter / transport)", "$150K - $500K\nper event",
             "$75K - $250K", "4-8% OR\nimprovement"],
            ["Ground Vehicle\n(Stryker / JLTV)", "$25K - $100K\nper event",
             "$15K - $50K", "3-6% OR\nimprovement"],
            ["Generator\n(MEP-series / FOB)", "$10K - $40K\nper event",
             "$5K - $20K", "5-10% availability\nimprovement"],
            ["Naval Propulsion", "$200K - $1M+\nper event",
             "$100K - $500K", "2-5% OR\nimprovement"],
            ["Communications / Radar", "$30K - $150K\nper event",
             "$15K - $75K", "5-8% availability\nimprovement"],
            ["Rotary Wing\n(UH-60 / CH-47)", "$100K - $400K\nper event",
             "$50K - $200K", "4-7% OR\nimprovement"],
        ],
        col_widths=[1.5, 1.5, 1.5, 1.3],
    )

    add_body_text(doc, (
        "Note: Savings estimates are based on published DoD maintenance cost data "
        "and industry benchmarks for predictive maintenance programs. Actual results "
        "will vary by platform, fleet size, operating environment, and current "
        "maintenance maturity level."
    ), italic=True)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  10. CONCLUSION
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "10. Conclusion", level=1)

    add_body_text(doc, (
        "Equipment changes HOW it operates, not WHETHER it operates. A turbine "
        "approaching failure does not suddenly stop producing thrust — it "
        "gradually shifts the relationships between temperature, vibration, fuel "
        "flow, and power output. A gearbox approaching seizure does not suddenly "
        "overheat — its thermal-to-vibration relationship drifts over weeks "
        "as metal-to-metal contact increases incrementally."
    ))

    add_body_text(doc, (
        "Traditional condition-based maintenance measures individual sensors — "
        "and misses failures caused by relationship changes between sensors. A "
        "behavioral approach measures direction across zones and the relationships "
        "between zones. When the vibration zone is drifting but the thermal zone is "
        "stable, it would diagnose bearing wear and indicate the specific bearing kit "
        "needed; when both zones drift together, it would identify cascading failure "
        "and escalate the maintenance priority. This is the same zone-divergence logic "
        "proven to isolate threat signatures in the cybersecurity domain, applied here "
        "to equipment health."
    ))

    add_body_text(doc, (
        "The maintenance-supply chain integration is the force multiplier. "
        "Behavioral intelligence does not just predict WHEN equipment will fail — "
        "it predicts WHAT parts will be needed, WHERE they should be staged, and "
        "HOW LONG before the maintenance action must occur. This transforms the "
        "supply chain from reactive (responding to work orders) to proactive "
        "(pre-positioning based on behavioral predictions), reducing parts-related "
        "maintenance delays and improving operational readiness across the fleet."
    ), bold=True)

    add_callout(doc,
        "BEI prescriptive maintenance is available for a 4-week proof of concept "
        "on a selected platform and installation. The proof of concept will "
        "demonstrate behavioral profiling, zone divergence detection, regime "
        "classification, and supply chain demand signal generation using existing "
        "sensor and maintenance data."
    )

    add_page_break(doc)

    # ── END MATTER ─────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("22nd Century Technologies, Inc.")
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("V-Intelligence UEBA Program — Prescriptive Maintenance")
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("June 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY

    # ── SAVE ───────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"Built: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
