"""Tenth pass: restore the 'Supplier Risk and Supply Chain Risk Management' overview
figure into Use Case 2, but using the CROPPED version (stale AUC/C-index/ECE metrics
bar removed) so no stale data is re-introduced. Tracked insertion + figure renumber:
  ... 6 Digital Twin | 7 SCM Overview (restored, cropped) | 8 Supplier Zones |
      9 Network | 10 Evidence | 11 Pilot
"""
import os, copy
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REV = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_final_REV_v2.docx"
IMG = os.path.join("docs", "_fig_scm_clean_82.png")   # cropped: metrics bar removed
AUTHOR = "ReviewPanel"; DATE = "2026-06-19T21:00:00Z"
d = Document(REV)
_id = [23000]
def nid():
    _id[0] += 1; return str(_id[0])
def insmark():
    e = OxmlElement('w:ins'); e.set(qn('w:id'), nid()); e.set(qn('w:author'), AUTHOR); e.set(qn('w:date'), DATE); return e
def ins_wrap(run):
    e = insmark(); e.append(run); return e
def acc(p): return "".join(x.text or "" for x in p._p.iter(qn('w:t')))

# --- renumber (direct), descending; leave the intro 'Figure 7' ref (it will point to the restored overview) ---
OPS = [
    ("Figure 10 - Bounded Pilot", "Figure 10", "Figure 11"),
    ("Figure 9 - Technical Evidence", "Figure 9", "Figure 10"),
    ("Figure 8 - Logistics Network", "Figure 8", "Figure 9"),
    ("Figure 7 - Our Supplier Risk Zones", "Figure 7", "Figure 8"),
    ("detailed in Figure 7, below", "Figure 7", "Figure 8"),
    ("Figure 7 shows both. The top row", "Figure 7", "Figure 8"),
]
def renumber(context, old, new):
    for p in d.paragraphs:
        if context in acc(p):
            for t in p._p.iter(qn('w:t')):
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new, 1); return True
    return False
applied = sum(1 for c, o, n in OPS if renumber(c, o, n))
print("renumber ops applied:", applied, "of", len(OPS))

# --- insert cropped SCM overview figure as Figure 7 after the SCRM governance paragraph ---
cap_src = next(p for p in d.paragraphs if 'Our Supplier Risk Zones and Relationship Map' in acc(p) and acc(p).strip().startswith('Figure'))
src_pPr = copy.deepcopy(cap_src._p.find(qn('w:pPr')))
src_rPr = cap_src.runs[0]._r.find(qn('w:rPr'))
src_rPr = copy.deepcopy(src_rPr) if src_rPr is not None else None

anchor = next(p for p in d.paragraphs if acc(p).strip().startswith('We understand Supply Chain Risk Management (SCRM)'))
tmp_p = d.add_paragraph(); tmp_r = tmp_p.add_run(); tmp_r.add_picture(IMG, width=Inches(6.3))
drawing = copy.deepcopy(tmp_r._r.find(qn('w:drawing')))
dp = drawing.find('.//' + qn('wp:docPr'))
if dp is not None:
    dp.set('id', '9007'); dp.set('name', 'Figure 7 Supplier Risk and SCM')
figp = OxmlElement('w:p'); pPr = OxmlElement('w:pPr')
jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)
rPr = OxmlElement('w:rPr'); rPr.append(insmark()); pPr.append(rPr); figp.append(pPr)
runel = OxmlElement('w:r'); runel.append(drawing); figp.append(ins_wrap(runel))
capp = OxmlElement('w:p'); capp.append(copy.deepcopy(src_pPr))
crPr = capp.find(qn('w:pPr')).find(qn('w:rPr'))
if crPr is None:
    crPr = OxmlElement('w:rPr'); capp.find(qn('w:pPr')).append(crPr)
crPr.append(insmark())
crun = OxmlElement('w:r')
if src_rPr is not None: crun.append(copy.deepcopy(src_rPr))
ct = OxmlElement('w:t'); ct.set(qn('xml:space'), 'preserve')
ct.text = "Figure 7 - Supplier Risk and Supply Chain Risk Management"
crun.append(ct); capp.append(ins_wrap(crun))
anchor._p.addnext(capp); anchor._p.addnext(figp)
tmp_p._p.getparent().remove(tmp_p._p)

o = REV; k = 3
while True:
    try:
        d.save(o); break
    except PermissionError:
        root, ext = os.path.splitext(REV); o = f"{root.replace('_v2','')}_v{k}{ext}"; k += 1
print("Saved:", o)
