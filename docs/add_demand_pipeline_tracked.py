"""Add the EDM demand-forecast front-end (digital twin + 67 engineered features
+ cohort cold-start) and the pipeline slide to the WP DLA paper, as TRACKED CHANGES.

The paper already covers steps 2-5 (two-stage LightGBM, snapshots, L2 drift, 8 concepts,
alpha+bands). Missing is slide 8's core: how raw records become one digital twin and the
67-feature decomposition. We add three crisp bullets + the (clean) pipeline figure, all as
Word tracked insertions, and turn track-changes on so the reviewer sees every addition.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

F = os.path.abspath(r'WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC- GY1.docx')
# write the tracked output to its own file; never overwrite the clean base GY1
OUT_FILE = os.path.abspath(r'WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC- GY1_tracked.docx')
IMG = os.path.abspath(r'WP DLA/_demand_slides/fig_demand_pipeline.png')          # pipeline figure (chrome stripped)
IMG2 = os.path.abspath(r'WP DLA/_demand_slides/fig_demand_compare.png')          # legacy-vs-EDM (overlap fixed, chrome stripped)
FIG6 = os.path.abspath(r'WP DLA/_demand_slides/fig6_supplier_zones_shaded.png')  # supplier zones, implemented-vs-planned shading
AUTHOR = 'Claude (V-Intelligence)'
DATE = '2026-06-17T12:00:00Z'
GREY = RGBColor(0x55, 0x5F, 0x6B)

doc = Document(F)

# ---- 0) remove ALL yellow/highlight change-marking (formatting only, not tracked) ----
def strip_highlights(container):
    n = 0
    for p in container.paragraphs:
        for r in p.runs:
            if r.font.highlight_color is not None:
                r.font.highlight_color = None; n += 1
    return n
removed = strip_highlights(doc)
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            removed += strip_highlights(c)
print('highlights removed:', removed)
_id = [1000]
def nid():
    _id[0] += 1; return _id[0]

def wrap_ins(run):
    """Move a run's <w:r> inside a tracked <w:ins>."""
    r = run._r
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(nid())); ins.set(qn('w:author'), AUTHOR); ins.set(qn('w:date'), DATE)
    r.addprevious(ins); ins.append(r)

def mark_para_ins(para):
    """Mark the paragraph mark itself as an inserted change."""
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr'); pPr.insert(0, rPr)
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(nid())); ins.set(qn('w:author'), AUTHOR); ins.set(qn('w:date'), DATE)
    rPr.insert(0, ins)

def find_para(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise SystemExit('anchor not found: ' + prefix)

# ---- 0b) legacy-vs-EDM overview slide before the "EDM Forecasting Pipeline" heading ----
def insert_figure(anchor_para, img_path, caption_text, width=6.3):
    figp = anchor_para.insert_paragraph_before()
    figp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = figp.add_run(); run.add_picture(img_path, width=Inches(width))
    wrap_ins(run); mark_para_ins(figp)
    cap = anchor_para.insert_paragraph_before()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption_text); cr.font.size = Pt(8.5); cr.italic = True; cr.font.color.rgb = GREY
    wrap_ins(cr); mark_para_ins(cap)

pipeline_hd = find_para('EDM Forecasting Pipeline')
insert_figure(pipeline_hd, IMG2,
              'Figure. Same data, two paths: the legacy baseline and EDM Vector Intelligence run on one shared '
              'data foundation. On a held-out 8-week window, EDM (798) lands closer to actual (627) than the legacy '
              'forecast (930). (Source: DLA EDM Vector Intelligence; single-item illustration.)')

# ---- 1) three crisp bullets before the "Baseline" step (front-end of the pipeline) ----
anchor = find_para('Baseline – a two-stage')
bullets = [
    "Digital twin: the same records are re-modeled as a behavioral twin of each item-and-depot. "
    "Five signal families (demand, supply, inventory, volatility, and context) are serialized and "
    "embedded into 1,536-dimension vectors, composited into one monthly vector, with 16 to 23 monthly "
    "snapshots per entity (the history available in the synthetic dataset) forming the trajectory.",
    "Engineered features: the twin is distilled into up to 67 features: 25 base tabular (demand "
    "history, exercise, item and depot, fiscal), 23 zone-derived (each behavioral zone reduced to its "
    "distance, isolation, and first two principal components, plus cross-zone terms), 13 drift signals, "
    "and 6 cohort features drawn from the nearest behavioral twins. The validated configuration used 29.",
    "Cohort transfer: for items with little history, a nearest-behavioral-twins (k-NN) transfer "
    "produces a usable day-one forecast, addressing cold-start.",
]
for text in bullets:
    p = anchor.insert_paragraph_before(style='List Bullet')
    r = p.add_run(text); r.font.size = Pt(11); r.font.name = 'Calibri'
    wrap_ins(r); mark_para_ins(p)

# ---- 2) the pipeline figure + caption before "Demand Technical Proof Points" ----
proof = find_para('Demand Technical Proof Points')
figp = proof.insert_paragraph_before()
figp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = figp.add_run(); run.add_picture(IMG, width=Inches(6.3))
wrap_ins(run); mark_para_ins(figp)

cap = proof.insert_paragraph_before()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cap.add_run('Figure. EDM demand-forecast pipeline: one digital twin distilled into engineered '
                 'features, then drift-nudged and Monte-Carlo simulated into drift-adjusted P10 to P90 bands. '
                 '(Source: DLA EDM Vector Intelligence; illustrative.)')
cr.font.size = Pt(8.5); cr.italic = True; cr.font.color.rgb = GREY
wrap_ins(cr); mark_para_ins(cap)

# ---- 2b) Use Case 2 honesty: mark the supplier figure as target architecture + preliminary framing ----
from docx.text.paragraph import Paragraph as _Para

def new_para_after(para):
    el = OxmlElement('w:p'); para._p.addnext(el)
    return _Para(el, para._parent)

def find_contains(substr):
    for p in doc.paragraphs:
        if substr in p.text:
            return p
    return None

# (i) append a tracked clarifier to the Figure 6 title
cap = find_contains('Our Supplier Risk Zones and Relationship Map')
if cap is not None:
    r = cap.add_run('  (Target Architecture)'); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
    wrap_ins(r)
    # (ii) tracked caption note distinguishing implemented vs planned
    note = new_para_after(cap); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run('Planned behavioral framework for supplier risk. The preliminary model implements the '
                      'performance-trajectory and network-position zones and the Supplier×NSN and NSN×Depot links; '
                      'the remaining zones and relationships (geopolitical sourcing, financial/SPRS, sub-tier and '
                      'bill-of-materials links) are the planned build-out.')
    nr.italic = True; nr.font.size = Pt(8.5); nr.font.color.rgb = GREY
    wrap_ins(nr); mark_para_ins(note)

# (iii) preliminary-phase framing sentence before the supplier proof-point numbers
pp = None
for p in doc.paragraphs:
    if p.text.strip().startswith('Supplier-Risk Technical Proof Points'):
        pp = p; break
if pp is not None:
    fr = new_para_after(pp)
    ft = fr.add_run('In a preliminary modeling phase, a supplier-risk classifier (XGBoost) and a survival model '
                    '(Random Survival Forest) were trained on synthetic data. The early results below (AUC ≈ 0.78, '
                    'survival concordance or C-index ≈ 0.79, calibrated at 30/60/90 days) indicate that the '
                    'behavioral approach transfers from demand to supplier risk. These are proof-of-approach figures, '
                    'not a fielded capability, and full framework coverage remains to be built and validated on DLA '
                    'operational data.')
    ft.font.size = Pt(11); ft.font.name = 'Calibri'
    wrap_ins(ft); mark_para_ins(fr)

# ---- 2c) tracked replacement of the flat Figure 6 with the shaded (implemented-vs-planned) version ----
def prev_para_with_drawing(cap_para):
    el = cap_para._p.getprevious()
    while el is not None:
        if el.tag == qn('w:p') and el.find('.//' + qn('w:drawing')) is not None:
            return _Para(el, cap_para._parent)
        el = el.getprevious()
    return None

# DISABLED: keep the original polished Figure 6; honesty is carried by the title clarifier,
# the caption note, and the preliminary-phase framing text (no shaded-graphic swap).
print('Figure 6 swap: disabled (keeping original image)')

# ---- 2d) §3.2 Three-Layer Architecture: add explanatory wording (plain, no em-dashes) ----
arch_next = None
for p in doc.paragraphs:
    if p.text.strip().startswith('Mathematical and Model Design Principles'):
        arch_next = p; break
if arch_next is not None:
    arch_paras = [
        "What makes the architecture useful is the work each layer does. The first layer profiles every entity on "
        "its own, whether a supplier, an item, a depot, or a route. Most analytics stop at this level. In practice, a "
        "large share of supply chain and security risk does not show up when entities are examined one at a time. The "
        "second layer treats the relationship between two entities as something to model in its own right, since a "
        "dependency can weaken while both entities still look healthy on their own. Two nominally independent "
        "suppliers that begin to move together is a relationship signal, not an entity signal. The third layer steps "
        "back to the whole network, where risk propagation, hidden common dependencies, and coordinated behavior "
        "become apparent only in the graph. Taken together, the layers describe individual behavior, the links "
        "between entities, and the way risk spreads across the system.",
        "Time is treated as a dimension that cuts across all three layers rather than as a fourth layer. Each entity, "
        "each relationship, and the network itself is stored as a series of time-stamped snapshots, and the model "
        "reads drift, velocity, and direction of change from that series. This is the part that separates the method "
        "from point-in-time reporting, because it lets the system act on where something is heading rather than only "
        "on its current state.",
        "The architecture is best read as a direction of travel rather than a finished schema. Each layer is built to "
        "take on more data and more attributes over time, including additional attributes at the entity layer, new "
        "relationship types at the relationship layer, and richer edges and propagation rules at the network layer. "
        "As the data foundation grows, the same three-layer structure can be widened and re-tuned without a redesign.",
    ]
    for text in arch_paras:
        p = arch_next.insert_paragraph_before()
        r = p.add_run(text); r.font.size = Pt(11); r.font.name = 'Calibri'
        wrap_ins(r); mark_para_ins(p)
    print('§3.2 architecture wording added:', len(arch_paras), 'paragraphs')

# ---- 2e) rewrite the demand "value" paragraph to lead with the innovation (tracked) ----
def mark_runs_deleted(para):
    for r in list(para._p.findall(qn('w:r'))):
        for t in r.findall(qn('w:t')):
            t.tag = qn('w:delText')
        de = OxmlElement('w:del'); de.set(qn('w:id'), str(nid()))
        de.set(qn('w:author'), AUTHOR); de.set(qn('w:date'), DATE)
        r.addprevious(de); de.append(r)

VAL_P1 = ("The core innovation is not another point forecast but a new way of representing the problem. Each "
          "item and depot is rebuilt as a digital twin using EDM Vector Intelligence, which brings the "
          "mathematics behind today's foundation models and large language models, transformer-style embeddings "
          "into a high-dimensional vector space, down to the application layer of defense logistics. Raw "
          "requisition, supply, and inventory records become a living behavioral entity instead of a flat table. "
          "That representation is what makes advanced operations possible, including behavioral drift detection, "
          "concept alignment, peer-cohort transfer, and full distributional simulation, techniques that are "
          "difficult or impossible to run on generic tabular data and that let the approach answer questions a "
          "conventional forecaster cannot.")
VAL_P2 = ("Because the accuracy of a single number was never the objective, the synthetic-data comparison should "
          "be read carefully. The EDM Vector Intelligence based Digital Twin matched a strong engineered baseline "
          "on point accuracy, with a Weighted Absolute Percentage Error (WAPE) of 0.7893 against 0.7908 for the "
          "baseline. Reaching the same accuracy is the point, not exceeding it by a fraction. The twin delivers "
          "that accuracy while also producing what the baseline cannot, namely calibrated uncertainty, early "
          "warning from behavioral drift, and a usable forecast for items with little or no history.")
VAL_P3 = ("The distributions it produces are trustworthy, which is what allows a planner to act on a range rather "
          "than on a single estimate. On the occurrence stage, which predicts whether demand will appear at all, "
          "the model is well calibrated, with an Expected Calibration Error (ECE, how far predicted probabilities "
          "sit from observed frequencies, where lower is better) of 0.0322 and a Brier score (the mean squared "
          "error of a probabilistic prediction, also lower being better and zero being perfect) of 0.0885. These "
          "calibrated probabilities are not a side metric. They are the basis of the planning range that drives "
          "procurement and early-warning decisions, where overplanning, underplanning, or a late response carry "
          "real operational and cost impact. The value of the EDM Vector Intelligence based Digital Twin is "
          "therefore in what its representation enables, a decision-support layer with calibrated confidence and "
          "advanced behavioral analytics, not a marginal change in conventional forecast accuracy. All figures "
          "here are from synthetic data and would be re-validated on DLA operational data.")

val_old = None
for p in doc.paragraphs:
    if 'specified rather than generic' in p.text:
        val_old = p; break
if val_old is not None:
    mark_runs_deleted(val_old)
    r1 = val_old.add_run(VAL_P1); r1.font.size = Pt(11); r1.font.name = 'Calibri'; wrap_ins(r1)
    p2 = new_para_after(val_old); r2 = p2.add_run(VAL_P2); r2.font.size = Pt(11); r2.font.name = 'Calibri'
    wrap_ins(r2); mark_para_ins(p2)
    p3 = new_para_after(p2); r3 = p3.add_run(VAL_P3); r3.font.size = Pt(11); r3.font.name = 'Calibri'
    wrap_ins(r3); mark_para_ins(p3)
    print('demand value paragraph rewritten (tracked)')
else:
    print('demand value paragraph not found')

# ---- 3) turn Track Changes ON so the document opens in markup ----
settings = doc.settings.element
if settings.find(qn('w:trackChanges')) is None:
    tc = OxmlElement('w:trackChanges'); settings.append(tc)

out = OUT_FILE; n = 2
while True:
    try:
        doc.save(out); break
    except PermissionError:
        root, ext = os.path.splitext(OUT_FILE); out = f"{root}_v{n}{ext}"; n += 1
print('Saved:', out)
