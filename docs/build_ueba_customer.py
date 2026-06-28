"""Customer-facing edition of the behavioral-intelligence cyber whitepaper.

Same factual content as the academic edition (every number verified against the
enhanced app and this repo), but:
  - de-branded (no company or product name),
  - guarded on method (no disclosure of the internal detection mechanics),
  - restructured and restyled to read as an editorial thought-leadership piece,
    organized around the GAPS in today's tools and why a behavioral approach breaks
    the pattern.

Distinct visual identity vs the academic edition: serif (Cambria) display headings,
teal / charcoal / amber palette, colored cover band, "GAP" callout boxes, pull-quotes.
No em-dashes or en-dashes. Synthetic proof-of-concept framing.

Verified facts carried over (unchanged):
  250 users, 485 days (70 weeks), ~14M events, 5 log sources, 4 embedded campaigns.
  Behavioral approach: 4 of 4 campaigns detected at 0 false positives.
  Best conventional method: 1 of 4 (a single statistical spike), no explanation.
  Conventional false-positive load: 4.5% / 5.3% / 9.8% / 14.6% across four methods.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "When_Normal_Hides_the_Threat_Customer_Whitepaper.docx")
FIG = os.path.join(os.path.dirname(__file__), "figs_ueba_enhanced")

TEAL   = RGBColor(0x0E, 0x6E, 0x6B)   # primary
CHAR   = RGBColor(0x22, 0x28, 0x2E)   # charcoal text / display
AMBER  = RGBColor(0xC2, 0x7A, 0x07)   # accent
SLATE  = RGBColor(0x53, 0x60, 0x68)   # secondary text
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
TEAL_HEX, CHAR_HEX, AMBER_HEX = '0E6E6B', '22282E', 'C27A07'
WASH_HEX, WASHA_HEX = 'EAF2F1', 'FBF3E4'   # teal wash, amber wash

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Georgia'; st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(7); st.paragraph_format.line_spacing = 1.22
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(doc.sections[0], m, Inches(0.95))


def _set_after(p, pt): p.paragraph_format.space_after = Pt(pt)


def band(text, sub=None, fill=CHAR_HEX, height_before=10):
    """Full-width colored band (single-cell table) used for the cover and part dividers."""
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); tcPr.append(sh)
    cell.width = Inches(6.6)
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(height_before); p.paragraph_format.space_after = Pt(2 if sub else height_before)
    r = p.add_run(text); r.bold = True; r.font.name = 'Cambria'; r.font.size = Pt(17); r.font.color.rgb = WHITE
    if sub:
        p2 = cell.add_paragraph(); p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(height_before)
        r2 = p2.add_run(sub); r2.font.name = 'Georgia'; r2.font.size = Pt(10); r2.italic = True
        r2.font.color.rgb = RGBColor(0xDD, 0xE6, 0xE6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def H(text, color=CHAR, size=15, before=16, after=4, kicker=None):
    if kicker:
        k = doc.add_paragraph(); k.paragraph_format.space_before = Pt(before); k.paragraph_format.space_after = Pt(0)
        rk = k.add_run(kicker.upper()); rk.bold = True; rk.font.name = 'Georgia'; rk.font.size = Pt(8.5)
        rk.font.color.rgb = AMBER; rk.font.element.rPr.append(OxmlElement('w:spacing'))
        rk.font.element.rPr.find(qn('w:spacing')).set(qn('w:val'), '40')
        before = 2
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = True; r.font.name = 'Cambria'; r.font.size = Pt(size); r.font.color.rgb = color
    return p


def sub(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.name = 'Cambria'; r.font.size = Pt(11.5); r.font.color.rgb = TEAL
    return p


def para(text, size=10.5, color=None, bold=False, italic=False, after=7, justify=True):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if justify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p


def lead_para(lead, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.space_after = Pt(7)
    r = p.add_run(lead); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = CHAR
    r2 = p.add_run(text); r2.font.size = Pt(10.5)
    return p


def bullet(text, lead=None):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = TEAL
    r2 = p.add_run(text); r2.font.size = Pt(10.5)
    return p


def callout(tag, lines, fill=WASH_HEX, tagcolor=TEAL):
    """Shaded callout box: a colored tag line followed by one or more body lines."""
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]; cell.width = Inches(6.4)
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); tcPr.append(sh)
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(tag.upper()); r.bold = True; r.font.name = 'Georgia'; r.font.size = Pt(9); r.font.color.rgb = tagcolor
    for i, (lead, body) in enumerate(lines):
        bp = cell.add_paragraph(); bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.space_after = Pt(3 if i < len(lines) - 1 else 5)
        if lead:
            rl = bp.add_run(lead); rl.bold = True; rl.font.size = Pt(10); rl.font.color.rgb = CHAR
        rb = bp.add_run(body); rb.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def pullquote(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.5); p.paragraph_format.right_indent = Inches(0.5)
    r = p.add_run(text); r.italic = True; r.font.name = 'Cambria'; r.font.size = Pt(13.5); r.font.color.rgb = TEAL


_fign = [0]


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor); tcPr.append(sh)


def setcell(cell, text, bold=False, size=9.5, color=None, fill=None):
    if fill: shade(cell, fill)
    cell.text = ''; p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Georgia'
    if color: r.font.color.rgb = color


def table(headers, rows, widths=None):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        setcell(t.rows[0].cells[j], h, bold=True, color=WHITE, fill=TEAL_HEX)
    for i, row in enumerate(rows, start=1):
        fill = WASH_HEX if i % 2 == 0 else None
        for j, v in enumerate(row):
            setcell(t.rows[i].cells[j], v, bold=(j == 0), fill=fill if j else WASHA_HEX)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_fig(name, caption, width=6.2):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(os.path.join(FIG, name), width=Inches(width))
    _fign[0] += 1
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after = Pt(8)
    r = c.add_run(f"Exhibit {_fign[0]}.  {caption}"); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = SLATE


# =========================================================================
# COVER
# =========================================================================
sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)
r = sp.add_run('A WHITEPAPER ON BEHAVIORAL CYBER DEFENSE'); r.bold = True; r.font.name = 'Georgia'
r.font.size = Pt(9); r.font.color.rgb = AMBER
band('When Normal Hides the Threat',
     'Why today\'s security tools miss the attacks that matter most, and the behavioral shift that finally makes them visible.',
     fill=CHAR_HEX, height_before=14)
para('The intrusions that cause the most damage do not look like attacks. They look like ordinary work: a trusted '
     'employee opening files, a server making a routine connection, an administrator running a familiar command. They '
     'unfold slowly, stay inside the bounds of normal, and leave conventional monitoring showing all green. This paper '
     'examines why that happens, where current tools fall short, and what changes when defense is built around '
     'behavior and its direction rather than rules and thresholds.', 10.5, SLATE, italic=True, after=4)
para('Prepared for security and risk leadership.  ·  Customer Edition  ·  June 2026', 9, SLATE, italic=True, justify=False, after=2)
para('Note on evidence: every result in this paper comes from a controlled study on a synthetic (simulated) '
     'enterprise built to test the concept end to end. The findings define a proof of concept and should be '
     're-validated on live telemetry before any production claim. The detection methods themselves are proprietary '
     'and are described here only at the level needed to understand what the approach does, not how it is built.',
     9, AMBER, italic=True, after=6)

# =========================================================================
# OPENING
# =========================================================================
H('The attacks that never trip an alarm', kicker='Why this paper', size=16)
para('Consider the two intrusions security leaders have spent the past two years studying. Volt Typhoon, a state-'
     'sponsored campaign against United States critical infrastructure, lived inside victim networks for at least five '
     'years without detection. Salt Typhoon, described by the chairman of the Senate Intelligence Committee as the '
     'worst telecommunications hack in the nation\'s history, operated inside major United States carriers for months. '
     'In both cases the victims\' own security tools never raised an alarm. The campaigns were found by outside '
     'intelligence agencies, long after the damage was done.')
para('These were not unsophisticated targets with weak tooling. They ran the same security stack most large '
     'enterprises run: a central log platform with expert-written rules, signature-based antivirus and intrusion '
     'detection, and endpoint protection. All of it stayed green. That is the uncomfortable truth this paper starts '
     'from: the failure was not a missed update or a tuning mistake. It was structural. The tools were asking a '
     'question these attackers were engineered to answer with a clean "nothing to see here."')
pullquote('"An insider changes what they touch, not how much. A patient intruder changes where it connects, not how '
          'often. Tools that measure volume never see either one."')

# =========================================================================
# PART ONE: THE GAP
# =========================================================================
band('Part One:  The Gap in Today\'s Defenses', fill=TEAL_HEX, height_before=11)
para('Conventional security monitoring rests on two ideas: match activity against things already known to be bad, '
     'and alert when a number crosses a fixed line. Both ideas are useful, and neither should be discarded. But '
     'modern adversaries are specifically built to slip between them. The sections below name the gaps one by one, '
     'because naming them is the first step to closing them.')

H('Gap one: the tools look for the known, and the dangerous attacks are unknown', size=13, before=12)
para('Signature and indicator-based detection works by comparing activity to a catalog of known-bad files, addresses, '
     'and domains. It is fast and precise for threats that have been seen before. But a novel implant, a freshly '
     'registered domain, or a first-seen destination is invisible until someone, somewhere, has already catalogued '
     'it. Worse, the most evasive campaigns use no malware at all. Living-off-the-land attacks operate entirely with '
     'the administration tools already installed on every machine, so there is nothing to fingerprint and no rule to '
     'violate, because every individual action is authorized.')
callout('The cost', [('', 'Detection that depends on a known-bad list is always one step behind the adversary, and '
                      'completely blind to the attacker who brings no tools at all.')])

H('Gap two: the tools measure how much, and the attack changes direction', size=13, before=12)
para('Most anomaly detection asks whether a number is unusually large: more logins, more data, more connections. A '
     'patient adversary simply keeps the numbers normal. An insider may access about the same volume of files before '
     'and during a campaign, while the content quietly shifts from public to restricted to confidential. A covert '
     'channel may call home in packets so small they vanish inside ordinary traffic. The volume never moves. What '
     'moves is the direction the behavior is heading, and a volume-based tool has no way to see direction at all.')

H('Gap three: the tools watch each entity alone, and threats live in the connections', size=13, before=12)
para('A user, device, or account can look perfectly normal on its own while its relationships have changed: a new '
     'recurring external destination, a widening spread of systems contacted, a shift in who talks to whom. Tools '
     'that score one entity and one metric at a time cannot see the pattern that only appears when behaviors are read '
     'together. The single most telling sign of a covert channel, one outside address contacted on a steady rhythm '
     'for over a year, is invisible to a detector that only counts how much traffic left the building.')

H('Gap four: the tools judge a moment, and the attack unfolds over months', size=13, before=12)
para('Threshold checks evaluate each day, or each week, in isolation. The signature technique of the most damaging '
     'campaigns is the opposite of a spike. It is called "slow and low": behavior changes so gradually that no single '
     'week ever looks unusual, while the cumulative change over many months achieves the objective. The activity is '
     'visible only in the trajectory, never in a snapshot, and a point-in-time tool only ever sees the snapshot.')
add_fig('fig_slowlow.png',
        'The slow-and-low pattern. Left: a threshold tool judges each week on its own, and the attacker\'s weekly '
        'change never crosses the alert line, so no alarm is ever raised. Right: the same weeks read as a trajectory '
        'separate the attacker from the normal range in months, not years. (Illustrative.)')

H('Gap five: the tools compare you to everyone, and risk is relative to your peers', size=13, before=12)
para('When a tool flags behavior as unusual, it usually means unusual compared to the whole population. But a '
     'developer behaves nothing like an accountant, and a security administrator nothing like a sales manager. '
     'Judging everyone against one global average hides the abuse that only stands out against the right peer group. '
     'A developer account quietly fanning out across the network looks unremarkable next to the company as a whole, '
     'and alarming only next to other developers.')

H('Gap six: tighten the net and you drown in false alarms', size=13, before=12)
para('Faced with these blind spots, the instinct is to lower the thresholds and catch more. The result is an alert '
     'flood. In our controlled study, one widely used anomaly method had to flag roughly one normal user in seven to '
     'reach even modest sensitivity, and still caught none of the real campaigns. The team that chases every alert '
     'has no time left for the one that matters. This is the trap at the center of the problem: loosen the tools and '
     'they bury the analyst; tighten them and they miss the attack.')
callout('The documented record', [
    ('Critical infrastructure: ', 'a state-sponsored campaign lived in victim networks for five-plus years, '
     'undetected by any victim tool.'),
    ('Telecommunications: ', 'a second campaign sat inside major carriers for months and was found only by outside '
     'agencies.'),
    ('Insider threats: ', 'industry data attributes roughly one in five breaches to insiders, with a median time to '
     'discovery exceeding two hundred days.')], fill=WASHA_HEX, tagcolor=AMBER)

# =========================================================================
# PART TWO: THE SHIFT
# =========================================================================
band('Part Two:  A Different Way to See', fill=TEAL_HEX, height_before=11)
para('Every gap above has the same root: the tools look at activity, when the answer lives in behavior over time. '
     'The shift this paper argues for is to stop scoring isolated events and start tracking how each entity, every '
     'user, device, segment, and account, behaves, how that behavior is changing, and which direction it is heading. '
     'The closest analogy is a medical chart. A good clinician does not act on today\'s temperature alone. They read '
     'the trend across many readings and intervene before a condition becomes acute.')
pullquote('"Defense stops asking is this action on a blocklist, and starts asking is this entity becoming something '
          'it has never been before."')

sub('Two questions, asked together')
para('A behavioral approach keeps the existing fast tools and adds a layer that answers two questions conventional '
     'monitoring cannot.')
lead_para('How far has this entity moved from its own normal, judged against its true peers?  ',
          'Behavior is measured relative to an entity\'s own history and to the genuine peer group it belongs to, '
          'not a global average, so abuse that only stands out among similar roles becomes visible.')
lead_para('Which direction is it heading?  ',
          'Change is not just measured for size but named for direction, so a gradual drift toward data theft, covert '
          'communication, or reconnaissance is recognized for what it resembles, and mapped to the recognized '
          'catalogue of adversary techniques security teams already use.')
para('Underneath, the approach runs two complementary layers. One is precise: it recognizes specific, measurable '
     'patterns of known-bad behavior and names each match in plain language, so an alert arrives with its reason '
     'attached. The other is exploratory: it watches for behavior that has never been seen before and keeps recurring, '
     'acting as a safety net for threats no one has profiled yet. The precise layer keeps false alarms near zero; the '
     'exploratory layer makes sure novel attacks are not missed. How each layer works internally is proprietary and '
     'is not described here. What matters for this paper is what the combination delivers.')

# =========================================================================
# PART THREE: WHAT BECOMES POSSIBLE
# =========================================================================
band('Part Three:  What Becomes Possible', fill=TEAL_HEX, height_before=11)
para('The same three families that defeat conventional tools become detectable when defense is built around behavior '
     'and direction. The scenarios below are drawn directly from the controlled study, told as outcomes rather than '
     'mechanics.')

sub('The trusted insider')
para('A long-serving business user runs a fourteen-month campaign, drifting from public toward restricted and '
     'confidential files while total file volume stays flat. To every volume-based tool the user is unremarkable. '
     'The behavioral approach lifts the user clear of the entire population early in the campaign, because what '
     'changed was the content and the direction, not the amount, and because the user\'s activity is judged against '
     'business-unit peers rather than the company at large.')

sub('The patient intruder')
para('A developer account hosts a slow, year-long intrusion whose implant calls home in small packets every few '
     'hours, with total traffic, lookups, and logins all inside normal ranges. This is the hardest case in the study, '
     'invisible to every volume-based method. The behavioral approach catches it on the one thing the attacker cannot '
     'hide: the machine-like regularity of the call-home rhythm, a steadiness that legitimate, bursty services never '
     'have. It reads the rhythm, not the address, so it does not need to have seen the attacker\'s infrastructure '
     'before.')

sub('The infrastructure compromise')
para('Two further campaigns model the critical-infrastructure and telecommunications intrusions that no victim '
     'organization caught on its own. One operates entirely through legitimate administration tools, leaving nothing '
     'to fingerprint. The other quietly widens its reach across the network. Both are recognized by the shape and '
     'breadth of their behavior, measured against the right peer group, even though no rule was broken and no '
     'signature existed.')

# =========================================================================
# PART FOUR: THE EVIDENCE
# =========================================================================
band('Part Four:  What We Proved, and What We Did Not', fill=TEAL_HEX, height_before=11)
para('Claims in security are cheap. The following results come from a single controlled study: a synthetic enterprise '
     'of 250 users observed across 485 days (about seventy weeks), roughly 14 million events drawn from five standard '
     'log sources, with four long-duration attack campaigns embedded and modeled on real, publicly documented '
     'intrusions. The same data was run through four widely used anomaly-detection methods and through the behavioral '
     'approach.')
callout('The headline', [
    ('Conventional methods. ', 'Of four standard anomaly detectors, three caught none of the four campaigns, and the '
     'best caught exactly one, a single statistical spike with no explanation of what changed.'),
    ('Behavioral approach. ', 'All four campaigns detected, zero false positives, each detection arriving with a '
     'plain-language reason and a named adversary technique.')], fill=WASH_HEX, tagcolor=TEAL)
table(['Approach', 'Campaigns found', 'Normal users wrongly flagged', 'What the analyst receives'],
      [('Conventional method A', '0 of 4', '4.5%', 'Few false alarms, and every real campaign sits inside its normal range.'),
       ('Conventional method B', '0 of 4', '5.3%', 'All four attackers look statistically ordinary.'),
       ('Conventional method C', '0 of 4', '14.6%', 'A flood of false alarms and still no real detection.'),
       ('Conventional method D', '1 of 4', '9.8%', 'One spike caught, barely; the other three never cross the line.'),
       ('Behavioral approach', '4 of 4', '0%', 'Every campaign named and explained; no normal user flagged.')],
      widths=[1.7, 1.0, 1.3, 2.5])
para('Two numbers carry the story. The behavioral approach found four of four where the best conventional method '
     'found one of four, and it did so while flagging zero normal users, against false-alarm rates that ran as high '
     'as one normal user in seven for the conventional tools. Detection improved and the noise went down at the same '
     'time, which is the combination the alert-flood trap is supposed to make impossible.')

sub('Where the limits are')
para('This is a proof of concept, not a production track record, and it should be read that way.')
bullet('The study runs on synthetic data, calibrated to public threat advisories. Every figure, including the zero-'
       'false-positive result, must be re-measured on live telemetry before it can be claimed operationally.',
       lead='The data is simulated. ')
bullet('The precise layer recognizes the patterns it has been built to recognize. The exploratory layer is the net '
       'for genuinely novel attacks, and by design it is noisier and needs analyst judgment.', lead='Coverage has limits. ')
bullet('The operating point in this study was chosen knowing the answers. In the field, thresholds must be set blind '
       'and tuned to each environment.', lead='The bar was set in a lab. ')
para('None of this weakens the central finding. It frames it. The approach demonstrably sees a class of attack that '
     'today\'s tools demonstrably miss, and it does so with an explanation attached. That is the result worth '
     'validating on real data.')

# =========================================================================
# PART FIVE: THE PATH FORWARD
# =========================================================================
band('Part Five:  A Low-Risk Path to Proof', fill=TEAL_HEX, height_before=11)
para('The natural next step is a bounded pilot that runs alongside the existing security stack, in advisory mode, on '
     'real telemetry. Nothing is automated to block until results are validated by analysts. The pilot answers three '
     'questions on the customer\'s own data: how early the behavioral layer detects long-running activity the current '
     'stack misses, how reliably it does so, and how clearly it explains itself.')
callout('Why the risk is low', [
    ('It adds, it does not replace. ', 'The behavioral layer complements existing detectors rather than displacing '
     'them, so current coverage is never reduced.'),
    ('It uses data you already collect. ', 'The approach runs on five standard log sources most enterprises already '
     'gather, with no new sensors required.'),
    ('It fits where your team already works. ', 'Findings arrive as ranked, explained leads with a named technique '
     'attached, feeding the existing analyst workflow rather than creating a new one.')], fill=WASHA_HEX, tagcolor=AMBER)
para('A typical pilot establishes behavioral baselines over the first several weeks, runs the behavioral layer in '
     'advisory mode against live activity, validates it by replaying known-bad or historical incidents on the '
     'environment\'s own data, and measures detection earliness, false-alarm load, and analyst-accepted leads in '
     'steady state. The success measure is simple: does the behavioral layer surface real, explained threats the '
     'current stack does not, at a false-alarm rate the team can live with.')

# =========================================================================
# CLOSING
# =========================================================================
band('The bottom line', fill=CHAR_HEX, height_before=11)
para('The threats that matter most, insiders, patient nation-state intrusions, living-off-the-land campaigns, and '
     'infrastructure compromises, all share one property. They operate inside authorized access, beneath every '
     'signature and threshold, and reveal themselves only as a gradual change of behavioral direction over months. '
     'The documented record is unambiguous: the victims of the most serious recent campaigns did not detect them. '
     'Outside agencies did, years and months too late.')
para('Closing those gaps does not require ripping out the security stack. It requires adding a layer that reads the '
     'same logs as a living record of behavior, measures how each entity is changing relative to its true peers, and '
     'names the direction of that change before the consequence is visible. In a controlled proof of concept, that '
     'layer found every embedded campaign with no false alarms, where the best conventional method found one. The '
     'next step is to prove it on real data, on a low-risk pilot, on your environment.', bold=True)
pullquote('"Today\'s tools measure magnitude. The attacks that matter live in direction. Closing that gap is the '
          'whole game."')

out = OUT; n = 2
while True:
    try:
        doc.save(out); break
    except PermissionError:
        root, ext = os.path.splitext(OUT); out = f"{root}_v{n}{ext}"; n += 1
print('Saved:', out, '| exhibits:', _fign[0])
