# -*- coding: utf-8 -*-
"""Digital Twin: design, gap assessment, and phased build plan (DESIGN ONLY, no code).
Grounded in the actual CyberUEBA implementation. No em/en dashes.
Output: WP DLA/Presentation/.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6); GRAY=RGBColor(0x55,0x5F,0x6B)
GREEN=RGBColor(0x1F,0x7A,0x33); AMBER=RGBColor(0xB5,0x5A,0x00); RED=RGBColor(0xA6,0x1B,0x1B)
d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(10.5)

def H1(t):
    p=d.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
    return p
def H2(t):
    p=d.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=BLUE
    return p
def P(t,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold; return p
def B(label,rest=""):
    p=d.add_paragraph(style='List Bullet')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def Bn(t): return d.add_paragraph(t,style='List Bullet')
def table(headers,rows,widths=None,fs=9):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(fs)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""; r=cells[i].paragraphs[0].add_run(v); r.font.size=Pt(fs)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    d.add_paragraph(); return t

# ===== COVER =====
cp=d.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=cp.add_run("The Behavioral Digital Twin"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
for txt,sz,it,col in [("Design, Gap Assessment, and Phased Build Plan",13,False,BLUE),
                      ("Grounded in the current implementation. Design only, no code.",10.5,True,GRAY)]:
    pp=d.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=pp.add_run(txt); rr.font.size=Pt(sz); rr.italic=it; rr.font.color.rgb=col
pp=d.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run("Ravindra Shukla  |  V-Intelligence  |  Internal Engineering Design").font.size=Pt(10)
d.add_paragraph()

# ===== 1 PURPOSE =====
H1("1. Purpose and Scope")
P("This document defines what the behavioral digital twin must become, assesses the current implementation against "
  "that target, identifies the gaps, and proposes a phased build plan that preserves the proven detection core. It is "
  "a design document only. No code is written here. The detection engine (embedding, drift, CUSUM, direction, "
  "composite scoring) is already built and validated on synthetic data; the work ahead is to turn that engine into an "
  "operational, scalable, real-data digital-twin platform.")

# ===== 2 TARGET DEFINITION =====
H1("2. What a Behavioral Digital Twin Must Be")
P("A behavioral digital twin is a living, time-aware representation of an entity (a user, device, service, segment, "
  "or relationship) that captures not just what the entity is, but what it is doing, how that behavior is changing, "
  "and in which direction. A complete twin has eight properties:")
B("One entity, many identifiers. ","All identifiers for a real subject (across AD, Entra, Okta, AWS, Kubernetes, EDR, VPN, PKI, and more) resolve to one twin.")
B("Multi-signal. ","Behavior is read across authentication, privilege, data, network, and endpoint signals, not a single feature.")
B("Semantic. ","Behavior is represented as meaning in a high-dimensional space, so identical counts with different meaning separate.")
B("Trajectory-aware. ","The twin is a sequence of time-stamped snapshots, so drift, velocity, and direction are measurable.")
B("Relational. ","The twin knows its relationships and its place in the network, so correlated and propagated risk is visible.")
B("Persistent and versioned. ","Every past state of the twin is stored, queryable, and auditable.")
B("Explainable. ","Every signal is named and mapped to MITRE ATT&CK, with evidence retained.")
B("Deployable and private. ","It runs in an enclave with a local model and holds no personally identifiable information in the vector.")

# ===== 3 CURRENT STATE =====
H1("3. Current State (as built)")
P("The following is grounded in the codebase. It describes what exists today.")
H2("3.1 Entity model")
B("models/cyber_entity.py: ","the CyberEntity object holds the entity type and id, a static profile, five zone embeddings, a composite embedding (attention-weighted fusion of the zones), a phase state (velocity, acceleration, stability, regime shifts, trend consistency, total drift), pairwise relationship embeddings, and per-zone and composite risk scores.")
B("Entity types: ","the model nominally supports user, device, network segment, application, and session, but the signal and zone logic is built out for users; non-user twins are scaffolding, not full implementations.")
H2("3.2 Signals and zones")
B("Five behavioral signal families ","(authentication, privilege and access, data access, network, and endpoint behavior) are computed per entity per period in embeddings/snapshot_builder.py, each serialized to a text description and embedded to a 1536-dimensional vector.")
B("Five behavioral zones ","(identity, access pattern, data behavior, network footprint, risk posture) are defined in models/hierarchical_zones.py; the signals compose into these zones, and the zones compose into the entity's composite embedding.")
H2("3.3 Trajectory and phase")
B("Snapshots: ","each entity is re-embedded per period to form a sequence (the trajectory). Weekly zone trajectories for 250 users across 70 weeks exist as a stored artifact (data/tier3_results/weekly_zone_trajectories.csv).")
B("Phase state: ","velocity, acceleration, and regime shifts are derived from the trajectory, the basis for drift detection.")
H2("3.4 Detection core (built and validated)")
B("Drift and CUSUM (detection/cusum.py), ","direction and MITRE projection (detection/drift_direction.py with 14 threat and 2 benign reference concepts in detection/reference_concepts.py), and five-phase composite scoring (detection/composite_scorer.py) operate on the twin and produce the ranked, explainable detections.")
H2("3.5 Persistence and serving")
B("Persisted today: ","trajectory and result artifacts as CSV files (composite scores, zone trajectories, z-scored features); a PostgreSQL with pgvector store and a FastAPI service plus a static UI exist (api/main.py, demo/ui/).")
B("Runtime-only today: ","the in-memory snapshot series, composite snapshots, and relationship snapshots are computed at run time and are not persisted as a versioned twin store.")
table(["Twin property","Status today"],
 [["One entity, many identifiers","Not built (each user_id is atomic; no resolver)"],
  ["Multi-signal","Built (five signal families for users)"],
  ["Semantic embedding","Built (1536-d, OpenAI text-embedding-3-small, mandatory)"],
  ["Trajectory-aware","Built (weekly/monthly snapshots, phase state)"],
  ["Relational / network","Partial (pairwise relationship embeddings; no graph store or propagation)"],
  ["Persistent and versioned","Partial (CSV artifacts; no versioned twin history store)"],
  ["Explainable","Built at run time (MITRE-mapped); evidence not persisted per alert"],
  ["Deployable and private","Partial (Docker; no local semantic model for air-gap yet; PII governance informal)"]],
 widths=[2.6,4.3])

# ===== 4 GAP ASSESSMENT =====
H1("4. Gap Assessment")
P("Each gap below lists what is missing, why it matters for a Pentagon-grade deployment, and a design-level build "
  "approach. Severity reflects pilot and operational impact.")
table(["#","Gap","Severity","Why it matters","Design-level approach"],
 [["1","Identity fusion / entity resolution","High","One real subject is split across 10+ systems, so one attacker looks like many users and the cross-domain chain is hidden. Highest-value gap for DoD.","An identity-resolution layer that maps many identifiers to one entity key using deterministic links (UPN, SID, email, cert) plus probabilistic behavioral matching, feeding one twin per subject."],
  ["2","Real-data ingestion connectors","High","Today only a synthetic generator feeds the twin. No connectors to real SIEM, IdP, EDR, VPN, or cloud logs. This is the pilot blocker.","Normalized ingestion adapters per source (Splunk/Sentinel, AD/Entra/Okta, EDR, VPN, cloud) mapping raw events to the common signal schema."],
  ["3","Streaming / online updates","High","The pipeline is batch (weekly/monthly). Operational SOC use needs near-real-time twin updates and detection.","Incremental snapshot updates on a rolling window, with online baseline maintenance, so the twin updates continuously rather than per batch."],
  ["4","Scale (250 to millions)","High","Validated at 250 entities. DoD enclaves have orders of magnitude more entities and identities.","pgvector index tuning (HNSW), cohort partitioning, and per-entity incremental compute; prove scaling on a representative real population in the pilot."],
  ["5","Persistent, versioned twin store","Medium","Snapshots are runtime-only; there is no queryable history of past twin states for replay or audit.","A versioned twin store (entity, zone embeddings, composite, phase, evidence per period) with slowly-changing-dimension history."],
  ["6","Cross-entity graph / network layer","Medium-High","Relationships exist only as pairwise embeddings. Hidden dependencies, correlated movement, and risk propagation (the network early-warning value) are not computed.","A graph layer over entities and relationships with propagation and cohort-movement analytics, feeding zone-level network intelligence."],
  ["7","Non-user entity twins","Medium","Device, service, segment, and session twins are scaffolding only; lateral movement and infrastructure attacks need them.","Per-type signal and zone schemas for device, service, and segment twins, reusing the same embedding and drift core."],
  ["8","Configurable / richer signals","Medium","The signal set is fixed in code; new data sources or mission contexts cannot be added by configuration.","A declarative signal schema (source, fields, serialization template) so signals are added by configuration, not code."],
  ["9","Cohort / peer-group construction","Medium","Cohort features exist but peer-group construction is basic; at scale, behavior must be judged against the right peers.","Role-based and behavioral clustering to construct cohorts, with cohort baselines maintained online."],
  ["10","Explainability and evidence storage","Medium","Direction and MITRE mapping are produced at run time but not persisted as queryable, per-alert evidence.","An evidence store linking each detection to its drifted signals, aligned concepts, MITRE techniques, and contributing snapshots."],
  ["11","Tamper-evident forensics","Medium","No HMAC-chained, court-admissible evidence log (a deck claim, not built).","An append-only, hash-chained evidence log over the persisted twin states and detections."],
  ["12","Proportional response (ABAC)","Medium","No response automation (a deck claim, not built).","A policy layer that turns a ranked verdict into a proportional, human-approved action (step-up authentication through lockout), kept human-in-the-loop for high-impact actions."],
  ["13","Local embedding model for air-gap","Medium","Real OpenAI embeddings are mandatory but not air-gapped. Enclave deployments need semantic quality offline.","Swap in a local open embedding model behind the existing embedder interface, validated for behavioral semantics."],
  ["14","PII governance","Low-Medium","Embeddings hold aggregated text (no PII), but raw events are stored; DoD needs formal PII handling and data lineage.","Formal PII classification, minimization, and lineage over the ingestion and storage layers."]],
 widths=[0.3,1.7,0.8,2.0,2.1],fs=8)

# ===== 5 TARGET ARCHITECTURE =====
H1("5. Target Architecture (conceptual)")
P("The target keeps the proven detection core and adds the layers around it that make it operational, scalable, and "
  "real-data ready. Reading top to bottom:")
B("Ingestion layer: ","normalized connectors from SIEM, IdP, EDR, VPN, and cloud to a common event schema.")
B("Identity resolution layer: ","fuse identifiers to one entity key, producing one twin per subject.")
B("Twin builder: ","signals to zones to composite, incremental, with online baselines and cohorts. (Existing core, made incremental.)")
B("Versioned twin store: ","every period's twin state, evidence, and detections, queryable and auditable.")
B("Graph / network layer: ","entity relationships, propagation, cohort movement, hidden-dependency analytics.")
B("Detection core: ","drift, CUSUM, direction and MITRE projection, composite scoring. (Built and validated.)")
B("Explainability and evidence: ","per-alert evidence, MITRE-mapped, optionally tamper-evident.")
B("Serving and response: ","API, UI, SOC and SIEM integration, and human-in-the-loop proportional response.")

# ===== 6 BUILD PLAN =====
H1("6. Phased Build Plan")
P("The plan preserves the proven core and sequences the gaps by pilot value and dependency.")
table(["Phase","Theme","Delivers","Closes gaps"],
 [["Phase 0 (now)","Proven detection core","Embedding, drift, CUSUM, direction/MITRE, composite scoring, validated 4 of 4 on synthetic data; Docker, API, UI.","(baseline)"],
  ["Phase 1 (pilot-enabling)","Run on real data","Real-data ingestion connectors; basic identity resolution; a versioned twin store; a local embedding model for the enclave. Outcome: the engine runs on a bounded set of real DoD telemetry.","2, 1 (basic), 5, 13"],
  ["Phase 2","Real-time and scale","Streaming/incremental twin updates and online baselines; scale to a large population (cohorting, index tuning); first non-user twins (device, service).","3, 4, 9, 7 (start)"],
  ["Phase 3","Network and breadth","Graph/network layer with propagation and cohort movement; full identity fusion; configurable signal schema; non-user twins completed. Outcome: network early-warning (Layer 3 of the architecture).","6, 1 (full), 8, 7 (finish)"],
  ["Phase 4","Trust and response","Persisted, MITRE-mapped evidence store; tamper-evident forensics; human-in-the-loop proportional response; formal PII governance. Outcome: operational SOC integration.","10, 11, 12, 14"]],
 widths=[1.2,1.4,3.3,1.0],fs=8.5)
P("Dependencies: Phase 1 (ingestion, identity, persistence, local model) is the prerequisite for everything "
  "operational and is the natural scope of a bounded pilot. Real-time and scale (Phase 2) depend on the versioned "
  "store and incremental builder from Phase 1. The graph and full fusion (Phase 3) depend on resolved identities and "
  "persisted relationships. Trust and response (Phase 4) depend on the evidence store.")

# ===== 7 PILOT ALIGNMENT =====
H1("7. Pilot Alignment")
P("Phase 1 is the pilot. The two highest-value first builds are identity fusion (so an attacker stops looking like "
  "many users) and real-data ingestion (so the validated engine runs on real telemetry). A bounded pilot would: "
  "stand up connectors for a selected enclave's identity and endpoint logs; resolve identities for a defined "
  "population; run the existing detection core on real behavior; and measure detection, explainability, and "
  "false-positive economics against the agency's current tooling, all in advisory, side-by-side mode under "
  "human-in-the-loop control. This keeps the proven core intact while proving it on operational data.")

# ===== 8 APPENDIX =====
H1("Appendix. Implementation Map (current)")
table(["Component","File(s)","Status"],
 [["Entity model","models/cyber_entity.py","Built (user-centric)"],
  ["Zones","models/hierarchical_zones.py","Built (5 zones)"],
  ["Signals and snapshots","embeddings/snapshot_builder.py","Built (5 signal families, serialization, embedding)"],
  ["Embedding","embeddings/embedder.py","Built (OpenAI text-embedding-3-small 1536-d, mandatory)"],
  ["Drift / CUSUM","detection/cusum.py","Built"],
  ["Direction / MITRE","detection/drift_direction.py, reference_concepts.py","Built (14 threat + 2 benign concepts)"],
  ["Composite scoring","detection/composite_scorer.py","Built (5 phases)"],
  ["Serving","api/main.py, demo/ui/, docker-compose.yml","Built (API, UI, container; host 8001 / container 8000)"],
  ["Identity fusion","(none)","Not built"],
  ["Graph / network layer","(pairwise embeddings only)","Partial"],
  ["Ingestion connectors","simulator/ (synthetic only)","Not built (real-data)"],
  ["Forensics / response","(none)","Not built"]],
 widths=[1.7,2.9,2.3],fs=8.5)
P("")
fn=d.add_paragraph(); fr=fn.add_run("Internal engineering design. Build-vs-roadmap distinctions in this document are "
  "the honest basis for scoping a pilot and must not be overstated as currently operational capability.")
fr.italic=True; fr.font.size=Pt(9); fr.font.color.rgb=GRAY

out="WP DLA/Presentation/Digital_Twin_Design_and_Build_Plan_INTERNAL.docx"
os.makedirs("WP DLA/Presentation",exist_ok=True)
d.save(out)
print("saved:",out,"| paragraphs:",len(d.paragraphs),"| tables:",len(d.tables))
