"""Apply tracked-changes edits to the final DLA academic whitepaper.

Edits (all as Word tracked changes, author=ReviewPanel):
  A) Three notes after the Table 2 caption: realized-demand/accuracy evidence,
     criticality-tiered service level, and the illustrative-synthetic-NSN footnote.
  B) Tracked deletion of the trailing-dash artifacts (paras 19, 82, 83, 84, 88).
  C) Appendix A glossary of key terms at the end.
Track changes is turned ON in settings so the markup shows and further edits track.
"""
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_final.docx"
OUT = "WP DLA/Decision_Intelligence_for_Defense_Logistics_White_Paper_ACADEMIC_final_REV.docx"
AUTHOR = "ReviewPanel"; DATE = "2026-06-19T12:00:00Z"

d = Document(SRC)
_id = [5000]
def nid():
    _id[0] += 1; return str(_id[0])

def _run(text, bold=False, italic=False):
    r = OxmlElement('w:r')
    if bold or italic:
        rPr = OxmlElement('w:rPr')
        if bold: rPr.append(OxmlElement('w:b'))
        if italic: rPr.append(OxmlElement('w:i'))
        r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t)
    return r

def _ins_wrap(run):
    ins = OxmlElement('w:ins'); ins.set(qn('w:id'), nid()); ins.set(qn('w:author'), AUTHOR); ins.set(qn('w:date'), DATE)
    ins.append(run); return ins

h1 = next(p for p in d.paragraphs if p.style.name == 'Heading 1')
H1ID = h1.style.style_id

def new_para(parts, style=None):
    p = OxmlElement('w:p'); pPr = OxmlElement('w:pPr')
    if style:
        ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'), style); pPr.append(ps)
    rPr = OxmlElement('w:rPr')
    ins = OxmlElement('w:ins'); ins.set(qn('w:id'), nid()); ins.set(qn('w:author'), AUTHOR); ins.set(qn('w:date'), DATE)
    rPr.append(ins); pPr.append(rPr); p.append(pPr)
    for (txt, b, it) in parts:
        p.append(_ins_wrap(_run(txt, b, it)))
    return p

# ---- A) notes after Table 2 caption ----
anchor = next(p for p in d.paragraphs if p.text.strip().startswith('Table 2 - Demand worked examples'))
notes = [
    new_para([("Measured against realized demand on these items, EDM's planning bands track actual requisitions far "
        "more closely than the legacy bands. For the adapter plate, the EDM 8-week P90 of 225 units sits at the "
        "observed peak 8-week demand of 224 units, while the legacy P90 of 379 carried roughly 155 units of excess. "
        "For the needle bearing, the EDM P90 of 545 is far nearer realized demand than the legacy P90 of 1,184, which "
        "over-forecasts by several hundred units. In each case EDM right-sizes toward real demand rather than padding "
        "the band.", False, False)]),
    new_para([("Because the proof-of-concept data carries an item criticality code (the adapter plate and needle "
        "bearing are coded C and the valve train kit B), the planning quantile can be tiered by criticality: a higher "
        "service level (for example P95 to P99) for sole-source, critical, or war-reserve items, and a lower level for "
        "routine items, rather than a uniform P90.", False, False)]),
    new_para([("Note: ", True, False),
        ("the National Stock Numbers shown in these examples are illustrative synthetic identifiers used for the proof "
         "of concept; they do not correspond to actual catalog records.", False, True)]),
]
cur = anchor._p
for np_ in notes:
    cur.addnext(np_); cur = np_

# ---- B) trailing-dash tracked deletions (match by text, not index) ----
import re
def del_trailing_dashes(p):
    n = 0
    for r in reversed(p.runs):
        tx = r.text
        if tx and set(tx) <= {'-'}:
            re = r._r
            for t in re.findall(qn('w:t')):
                dt = OxmlElement('w:delText'); dt.set(qn('xml:space'), 'preserve'); dt.text = t.text
                re.replace(t, dt)
            delw = OxmlElement('w:del'); delw.set(qn('w:id'), nid()); delw.set(qn('w:author'), AUTHOR); delw.set(qn('w:date'), DATE)
            re.addprevious(delw); delw.append(re); n += 1
        elif tx and tx.endswith('-'):
            stripped = tx.rstrip('-'); dashes = tx[len(stripped):]
            for t in r._r.findall(qn('w:t')):
                t.text = stripped
            delr = _run(dashes)
            for t in delr.findall(qn('w:t')):
                dt = OxmlElement('w:delText'); dt.set(qn('xml:space'), 'preserve'); dt.text = t.text; delr.replace(t, dt)
            delw = OxmlElement('w:del'); delw.set(qn('w:id'), nid()); delw.set(qn('w:author'), AUTHOR); delw.set(qn('w:date'), DATE)
            delw.append(delr); r._r.addnext(delw); n += 1
            break
        else:
            break
    return n
# artifact pattern: text ends with a period (or word char) followed by one or more stray hyphens
dash_re = re.compile(r'[\w.]-+$')
ddel = 0; dpar = 0
for p in d.paragraphs:
    t = p.text.rstrip()
    if t.endswith('-') and dash_re.search(t) and len(t) > 20:
        c = del_trailing_dashes(p)
        if c: dpar += 1; ddel += c

# ---- C) glossary appendix ----
body = d.element.body
sectPr = body.find(qn('w:sectPr'))
GLOSS = [
    ("EDM (Entity Digital Model) Vector Intelligence", "an approach that rebuilds each item, supplier, depot, and route as a time-aware digital twin and represents its state as a high-dimensional vector, so it can be compared over time and against peers."),
    ("BEI (Behavioral Entity Intelligence)", "the cross-domain early-warning framework that analyzes entities, their relationships, and the network over time to surface slow, directional change."),
    ("NSN (National Stock Number)", "the standard identifier for a supply item; the first four digits are the Federal Supply Class."),
    ("DODAAC", "Department of Defense Activity Address Code; identifies the activity or location associated with a requisition."),
    ("P10 / P50 / P90", "points on the forecast range: the P90 is the level demand is not expected to exceed about 90 percent of the time, used here as the planning (buy) quantile; the P50 is the midpoint."),
    ("Calibration (ECE, Brier score)", "measures of whether predicted probabilities match real frequencies; lower values mean the forecast ranges can be trusted."),
    ("Hurdle model", "a two-stage forecast that first predicts whether demand will occur at all, then how much, which suits intermittent demand."),
    ("Drift / CUSUM", "drift is gradual, directional change in behavior; CUSUM is a change-detection method that accumulates small shifts until they become significant."),
    ("Cohort", "a peer group of similar items or suppliers; behavior is judged relative to the cohort, and new items borrow from their nearest behavioral twins."),
    ("Survival model", "a method that estimates the time until an event such as a supplier disruption and how that risk develops over 30, 60, and 90 days."),
    ("Embedding", "a numerical vector that captures the meaning of serialized entity state; here it generates most of the engineered features (zone, drift, and cohort features)."),
    ("Criticality code", "an item designation (for example C or B) reflecting mission importance, used to set how protective the planning level should be."),
]
def insert_before_sectpr(el):
    if sectPr is not None: sectPr.addprevious(el)
    else: body.append(el)
insert_before_sectpr(new_para([("Appendix A: Glossary of Key Terms", False, False)], style=H1ID))
for term, defn in GLOSS:
    insert_before_sectpr(new_para([(term + ". ", True, False), (defn, False, False)]))

# ---- enable track changes ----
settings = d.settings.element
if settings.find(qn('w:trackChanges')) is None:
    settings.append(OxmlElement('w:trackChanges'))

o = OUT; k = 2
while True:
    try:
        d.save(o); break
    except PermissionError:
        root, ext = os.path.splitext(OUT); o = f"{root}_v{k}{ext}"; k += 1
print("Saved:", o)
print("notes inserted:", len(notes), "| dash paragraphs fixed:", dpar, "| dash runs deleted:", ddel, "| glossary terms:", len(GLOSS))
