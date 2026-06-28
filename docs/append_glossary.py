# -*- coding: utf-8 -*-
"""Append a consistent 'Glossary of Terms and Acronyms' section to all relevant
cyber-session documents. Idempotent: skips a doc that already has the glossary.
Robust: catches PermissionError (file open in Word) and reports it.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6); GRAY=RGBColor(0x55,0x5F,0x6B)

GLOSSARY = [
 ("Network and Infrastructure", [
  ("ABAC","Attribute-Based Access Control","Authorization from user/resource attributes and policy, not fixed roles."),
  ("C2","Command and Control","Attacker infrastructure that directs compromised hosts."),
  ("CASB","Cloud Access Security Broker","Policy enforcement point between users and cloud services (part of SASE)."),
  ("DGA","Domain Generation Algorithm","Malware technique that auto-generates many domains to reach C2."),
  ("DNS","Domain Name System","Resolves names to IP addresses; abused for tunneling and exfiltration."),
  ("IdP","Identity Provider","Authenticates users and issues identity tokens (e.g., Okta, Microsoft Entra ID)."),
  ("IPS","Intrusion Prevention System","Inline sensor that detects and blocks known malicious traffic."),
  ("NGFW","Next-Generation Firewall","Firewall adding application awareness, IPS, and deep packet inspection."),
  ("RDP","Remote Desktop Protocol","Windows remote-desktop protocol; common lateral-movement path."),
  ("SASE","Secure Access Service Edge","Cloud-delivered networking plus security (SD-WAN, ZTNA, SWG, CASB)."),
  ("SD-WAN","Software-Defined Wide Area Network","Policy-driven WAN connectivity (part of SASE)."),
  ("SMB","Server Message Block","Windows file/printer-sharing protocol; common lateral-movement path."),
  ("SWG","Secure Web Gateway","Filters web traffic and enforces policy (part of SASE)."),
  ("VLAN","Virtual Local Area Network","Logical segmentation of a physical network."),
  ("WAF","Web Application Firewall","Filters and monitors HTTP traffic to protect web apps and APIs."),
  ("WinRM","Windows Remote Management","Windows remote-management protocol; lateral-movement path."),
  ("ZTNA","Zero Trust Network Access","Identity-aware, least-privilege application access (part of SASE)."),
 ]),
 ("Detection and Analytics", [
  ("ATT&CK","MITRE ATT&CK","MITRE's knowledge base of adversary tactics and techniques."),
  ("ATLAS","MITRE ATLAS","MITRE's adversarial-threat framework for AI/ML systems."),
  ("BEI","Behavioral Entity Intelligence","22nd Century's behavioral-twin detection approach (the Layer 2 product)."),
  ("CUSUM","Cumulative Sum","Change-point method that accumulates small drifts to detect slow, sustained change."),
  ("iForest","Isolation Forest","Classical tree-based anomaly-detection algorithm (a baseline here)."),
  ("LOF","Local Outlier Factor","Classical density-based anomaly-detection algorithm (a baseline here)."),
  ("LOTL","Living off the Land","Attacks using built-in OS tools only, with no malware deployed."),
  ("OC-SVM","One-Class Support Vector Machine","Classical anomaly-detection algorithm (a baseline here)."),
  ("SOC","Security Operations Center","The team and facility that monitor and respond to threats."),
  ("T-codes","MITRE technique IDs","Numbered ATT&CK techniques; e.g., T1078 = Valid Accounts."),
  ("UEBA","User and Entity Behavior Analytics","Detecting threats by modeling normal behavior and flagging deviation."),
  ("z-score","Standard score","Number of standard deviations a value sits from its group mean."),
 ]),
 ("Application, API and Data", [
  ("API","Application Programming Interface","Interface through which systems and AI agents exchange data."),
  ("ASPM","Application Security Posture Management","Aggregates and prioritizes application-security findings."),
  ("BOLA","Broken Object Level Authorization","Top API flaw: accessing another user's objects through the API."),
  ("HMAC","Hash-based Message Authentication Code","Keyed hash used to make records tamper-evident."),
  ("MFA","Multi-Factor Authentication","Authentication requiring more than one independent factor."),
  ("PII","Personally Identifiable Information","Data that identifies an individual; not held in the behavioral vector."),
 ]),
 ("Code and Supply Chain", [
  ("CI/CD","Continuous Integration / Continuous Delivery","Automated build, test, and deploy pipeline."),
  ("CVE","Common Vulnerabilities and Exposures","Public catalog of known, disclosed vulnerabilities."),
  ("DAST","Dynamic Application Security Testing","Tests a running application for vulnerabilities."),
  ("IaC","Infrastructure as Code","Provisioning infrastructure via machine-readable definition files."),
  ("RCE","Remote Code Execution","Vulnerability allowing an attacker to run code remotely."),
  ("SAST","Static Application Security Testing","Analyzes source code for vulnerabilities."),
  ("SCA","Software Composition Analysis","Identifies vulnerable open-source and third-party dependencies."),
 ]),
 ("Federal and Acquisition", [
  ("ARCYBER","U.S. Army Cyber Command","The Army's cyber operations command."),
  ("DFARS","Defense Federal Acquisition Regulation Supplement","DoD supplement to federal acquisition rules."),
  ("FAR","Federal Acquisition Regulation","Primary U.S. federal acquisition regulation."),
  ("IL5 / IL6","Impact Level 5 / 6","DoD cloud security levels (IL5 CUI/NSS; IL6 up to SECRET)."),
  ("JWICS","Joint Worldwide Intelligence Communications System","Top-Secret/SCI classified network."),
  ("NDAA","National Defense Authorization Act","Annual U.S. defense policy and funding law."),
  ("NIST SP 800-53 / 800-161","NIST Special Publications","Security-control and supply-chain-risk-management catalogs."),
  ("OTA","Other Transaction Authority","Flexible non-FAR mechanism for prototypes and pilots."),
  ("PCA","Principal Cyber Advisor","Senior cyber advisor to a Service Secretary (Brandon Pugh, for the Army)."),
  ("PEO IEW&S","Program Executive Office, Intelligence Electronic Warfare and Sensors","Army acquisition organization."),
  ("PM C&S","Project Manager, Cyber and Space","Army acquisition office (LTC DiGrezio's seat)."),
  ("PoR","Program of Record","A formally funded, sustained acquisition program."),
  ("SCRM","Supply Chain Risk Management","Managing risk across the supplier base."),
  ("TRL","Technology Readiness Level","A 1-to-9 scale of technology maturity."),
 ]),
]

HEADING="Glossary of Terms and Acronyms"

def has_glossary(doc):
    for p in doc.paragraphs:
        if HEADING.lower() in (p.text or "").lower():
            return True
    return False

def append_glossary(path):
    doc=Document(path)
    if has_glossary(doc):
        return "SKIP (already has glossary)"
    # page break + heading
    pb=doc.add_paragraph(); pb.add_run().add_break(WD_BREAK.PAGE)
    h=doc.add_heading(HEADING,level=1)
    for r in h.runs: r.font.color.rgb=NAVY
    intro=doc.add_paragraph()
    ri=intro.add_run("Plain-language expansions of the acronyms used in this document and the wider session materials.")
    ri.italic=True; ri.font.size=Pt(9.5); ri.font.color.rgb=GRAY
    for cat,rows in GLOSSARY:
        ch=doc.add_paragraph(); cr=ch.add_run(cat); cr.bold=True; cr.font.color.rgb=BLUE; cr.font.size=Pt(11)
        ch.paragraph_format.space_before=Pt(6); ch.paragraph_format.space_after=Pt(2)
        t=doc.add_table(rows=1,cols=3); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
        hdr=["Acronym","Term","Meaning"]
        for i,htext in enumerate(hdr):
            c=t.rows[0].cells[i]; c.text=""; rr=c.paragraphs[0].add_run(htext); rr.bold=True; rr.font.size=Pt(8.5)
        for acro,term,mean in rows:
            cells=t.add_row().cells
            for i,val in enumerate([acro,term,mean]):
                cells[i].text=""; rr=cells[i].paragraphs[0].add_run(val); rr.font.size=Pt(8.5)
                if i==0: rr.bold=True
        for row in t.rows:
            row.cells[0].width=Inches(0.95); row.cells[1].width=Inches(2.25); row.cells[2].width=Inches(3.3)
    doc.save(path)
    return "OK"

TARGETS=[
 "V-Intelligence_Layered_Cybersecurity_Whitepaper_v2.docx",
 "V-Intelligence_UEBA_Technical_Leave-Behind.docx",
 "Cyber_DeepDive_Briefing_Book_Brandon_Pugh_INTERNAL.docx",
 "Brandon_Pugh_Tough_Questions_Cheat_Sheet_INTERNAL.docx",
 "Brandon_Pugh_Working_Session_Run-Sheet_INTERNAL.docx",
 "Digital_Twin_Working_Session_RunOfShow_INTERNAL.docx",
 "Digital_Twin_Design_and_Build_Plan_INTERNAL.docx",
 "Working_Session_RunOfShow_FourLayer_22CT_INTERNAL.docx",
 "Presenter_Briefing_FourLayer_Session_INTERNAL.docx",
]
base="WP DLA/Presentation"
n_terms=sum(len(r) for _,r in GLOSSARY)
print(f"Glossary: {n_terms} terms across {len(GLOSSARY)} categories\n")
for fn in TARGETS:
    path=os.path.join(base,fn)
    if not os.path.exists(path):
        print(f"  MISSING   {fn}"); continue
    try:
        status=append_glossary(path)
        print(f"  {status:9s} {fn}")
    except PermissionError:
        print(f"  LOCKED    {fn}  <-- close it in Word, then re-run")
    except Exception as e:
        print(f"  ERROR     {fn}: {e}")
