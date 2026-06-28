# -*- coding: utf-8 -*-
"""ONE-COMPANY four-layer working-session docs. No mention of partners anywhere.
All four layers presented as 22nd Century. In-person, all presenters present.
Two docs:
 1) Working_Session_RunOfShow_FourLayer_22CT_INTERNAL.docx
 2) Presenter_Briefing_FourLayer_Session_INTERNAL.docx
No em-dashes.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6); GRAY=RGBColor(0x55,0x5F,0x6B)
RED=RGBColor(0xA6,0x1B,0x1B); GREEN=RGBColor(0x1F,0x7A,0x33)

def new_doc():
    d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(10.5); return d
def H1(d,t):
    p=d.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
    return p
def H2(d,t):
    p=d.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=BLUE
    return p
def P(d,t,italic=False,bold=False,color=None,size=None):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold
    if color: r.font.color.rgb=color
    if size: r.font.size=Pt(size)
    return p
def B(d,label,rest="",sub=False):
    p=d.add_paragraph(style='List Bullet 2' if sub else 'List Bullet')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def meta(d,label,val):
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(label+"  "); r.bold=True; r.font.color.rgb=GRAY; r.font.size=Pt(9.5)
    r2=p.add_run(val); r2.font.size=Pt(9.5); return p
def say(d,lines):
    lp=d.add_paragraph(); lp.paragraph_format.space_after=Pt(1)
    lr=lp.add_run("SAY:"); lr.bold=True; lr.font.color.rgb=GREEN; lr.font.size=Pt(9.5)
    for ln in lines:
        p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.25); p.paragraph_format.space_after=Pt(3)
        r=p.add_run("“"+ln+"”"); r.italic=True; r.font.size=Pt(10.5)
def callout(d,text,color=RED,size=9.5):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2)
    r=p.add_run(text); r.bold=True; r.font.color.rgb=color; r.font.size=Pt(size); return p
def table(d,headers,rows,widths=None,fs=8.5,hdr_fs=8.5):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(hdr_fs)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""
            for j,part in enumerate(v.split("||")):
                para=cells[i].paragraphs[0] if j==0 else cells[i].add_paragraph()
                para.add_run(part).font.size=Pt(fs)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    d.add_paragraph(); return t
def title_block(d,title,subtitle,banner):
    tp=d.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=tp.add_run(title); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=NAVY
    sp=d.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=sp.add_run(subtitle); r.font.size=Pt(11.5); r.font.color.rgb=BLUE
    wn=d.add_paragraph(); wn.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rw=wn.add_run(banner); rw.bold=True; rw.font.color.rgb=RED; rw.font.size=Pt(9.5)

# =====================================================================
# DOC 1: RUN-OF-SHOW (one company)
# =====================================================================
d=new_doc()
title_block(d,"Working Session Run-of-Show",
  "22nd Century V-Intelligence Four-Layer Platform  |  Army Cyber: Esquibel - DiGrezio - Herrmann",
  "INTERNAL prep. One company, one platform, every layer live.")

H1(d,"1. The platform")
P(d,"The customer sees one integrated offering: the 22nd Century V-Intelligence four-layer platform, covering the full "
  "attack lifecycle under one MITRE vocabulary. The behavioral digital twin (Layer 2) is the centerpiece they asked "
  "for. 22nd Century owns the open, every transition, and the close.")
table(d,["Layer","Capability","Maturity"],
 [["1 - Preemptive Network Assurance","Formal-proof network assurance over the full config space","Live"],
  ["2 - Behavioral Entity Intelligence","The behavioral digital twin (our core IP)","Live"],
  ["3 - Application, API & Data Runtime","API, data, and agentic-AI runtime protection","Live"],
  ["4 - Code & Supply-Chain","Novel zero-day discovery pre-deploy","In development"]],
 widths=[2.0,3.0,1.0],fs=9)

H1(d,"2. Honest about maturity (say it out loud)")
B(d,"Layer 4 is in development. ","Present the approach, not a running product. Substantiate or drop the 2,000-zero-days figure.")
B(d,"Within Layer 2, ","the detection core runs live; identity fusion across systems, the versioned twin store, the "
  "behavioral graph, tamper-evident forensics, and automated response are roadmap.")
callout(d,"With ARCYBER (Esquibel) in the room, candor is the advantage. Label roadmap as roadmap, every time.")

H1(d,"3. One vocabulary, one scenario, one format")
B(d,"One vocabulary: ","MITRE ATT&CK across all four layers.")
B(d,"One scenario: ","the shared Volt Typhoon / Salt Typhoon narrative. Every layer maps to the SAME campaigns.")
B(d,"One results format (every layer identical): ","threat scenario -> what incumbents miss -> live demo -> measured result.")

H1(d,"4. Agenda (half-day, ~4.5h with discipline)")
table(d,["#","Block","Lead","Time"],
 [["0","Open: their goals, set the arc, the platform thesis","Session lead","20 min"],
  ["1","Layer 2 - Behavioral digital twin (theory, live demo, results)","Twin lead","60 min"],
  ["2","Layer 1 - Preemptive Network Assurance (theory, live demo, results)","L1 lead","60 min"],
  ["3","Layer 3 - API, Data & Agentic-AI Runtime (theory, live demo, results)","L3 lead","60 min"],
  ["4","Layer 4 - Code & Supply-Chain (roadmap, honest)","Session lead","15 min"],
  ["5","Integration: one Typhoon across all four layers + federal fit + the ask","Session lead","30 min"],
  ["-","Buffer for hard questions","All","20 min"]],
 widths=[0.35,3.6,1.0,0.7],fs=8.5)
meta(d,"Sequence note:","Lead with the digital twin (Layer 2). They asked for it, it is the core IP, and it sets the "
 "credibility bar the other blocks inherit. The session lead bookends so it reads as one architecture.")

H1(d,"5. Block-by-block")
H2(d,"Block 0 - Open")
say(d,["What you will see today is one platform, the 22nd Century V-Intelligence four-layer architecture, covering the "
 "full attack lifecycle under one MITRE vocabulary. Every layer runs live. We will be explicit about what is "
 "operational today versus on the roadmap. Let's start with the behavioral core."])
meta(d,"Why:","Sets the stakes (Volt/Salt, valid-account gap) and states the honesty boundary before any mechanics, so the demos that follow are answering a real campaign, not floating as a pitch.")
H2(d,"Block 1 - Layer 2, the digital twin (anchor)")
meta(d,"Content:","The 7-part anatomy + the live four-attacker demo + the 8.1% derivation. Real numbers in Section 9.")
meta(d,"Why:","Proves the hardest claim live: detecting a nation-state living off the land on valid accounts. Slow APT is invisible to thresholds and caught by novelty alone; Volt is caught by breadth of small signals. Win this block and the other layers inherit the credibility.")
meta(d,"Honesty line:","\"The behavioral core runs live. Identity fusion, forensics, and response are roadmap.\"")
H2(d,"Block 2 - Layer 1, Preemptive")
meta(d,"Content:","Formal-proof network assurance. Scenario: prove Volt's misconfigured-edge entry path is closed before traffic flows.")
meta(d,"Why:","The prevention half of the gap. Layer 2 catches the adversary already inside; Layer 1 proves the edge was provably closed. Detection alone is reactive; prevention plus detection is a posture.")
meta(d,"Format:","theory 15, live demo 30, results 15. Same results frame as every other layer.")
meta(d,"Handoff:","Twin lead hands to the L1 lead; session lead keeps it framed as Layer 1 of the one platform.")
H2(d,"Block 3 - Layer 3, API/Data/Agentic")
meta(d,"Content:","API, data, and agentic-AI runtime protection. Scenario: catch Salt's API/data collection and agent abuse at runtime.")
meta(d,"Why:","The exfiltration altitude, where Salt actually steals data, plus the agentic-AI blind spot. Layer 2 sees the behavior; Layer 3 sees the data movement and the AI agents.")
meta(d,"Format:","theory 15, live demo 30, results 15.")
H2(d,"Block 4 - Layer 4, Code/Supply-Chain (roadmap)")
meta(d,"Why:","Closes the lifecycle at the source, the exploitable code that grants initial access (SolarWinds / XZ style) that no other layer addresses. Even as roadmap, it shows the architecture is complete end to end.")
meta(d,"Honesty line:","\"This layer is in development. I am presenting the approach, not a running product.\"")
H2(d,"Block 5 - Integration + the ask")
meta(d,"Why:","Here the layers become more than their sum. The hand-off across layers is the moat no single tool can claim, and it converts the demo into a decision: a bounded pilot.")
meta(d,"The moat:","Only the integrated platform shows the layers handing off. Trace one Typhoon intrusion: Layer 4 removes the "
 "exploitable code, Layer 1 proves the edge is closed, Layer 2 catches the valid-account behavior inside, Layer 3 catches the API/data exfil.")
P(d,"The ask: \"You saw one platform, four layers, every layer live, under one MITRE vocabulary, from one vendor. The "
  "honest next step is a bounded pilot on real Army telemetry, advisory and human-in-the-loop, measured side-by-side "
  "against current tooling.\"",italic=True)

H1(d,"6. Why covering all four layers makes it relevant")
B(d,"The adversary uses all four layers, ","so a single-layer answer is structurally incomplete. Covering all four maps the defense to how Volt and Salt actually operate.")
B(d,"Each layer catches what the others cannot: ","Layer 1 cannot see valid-account behavior (the credentials are valid); Layer 2 cannot prove the edge configuration; Layer 3 sees runtime data movement but not the behavioral trajectory; Layer 4 prevents the foothold none of the others reach. The coverage is complementarity, not redundancy.")
B(d,"One MITRE vocabulary ","makes it one defense, not four tools with blind spots between them.")
B(d,"It satisfies all three seats at once: ","Esquibel gets measured detection fidelity, DiGrezio gets one fieldable platform from one vendor, Herrmann gets a complete answer to the national valid-account / critical-infrastructure gap.")

H1(d,"7. The overall message (the close)")
P(d,"\"No single box stops a Volt or a Salt; they log in, they do not break in, and they move across the whole stack. "
  "What you saw today is one platform, four layers, under one MITRE vocabulary, covering the full attack lifecycle from "
  "code to edge to behavior to data. The behavioral core is ours and runs live, on the logs you already collect, "
  "catching four nation-state campaigns your current tools score as normal. The other layers close the rest of the "
  "lifecycle; Layer 4 is the honest roadmap. The right next step is a bounded pilot on real Army telemetry, advisory "
  "and human-in-the-loop, measured side-by-side against what you run today.\"",italic=True)
B(d,"Walk-out belief: ","the most complete and the most honest answer to the valid-account / living-off-the-land gap; the core is proven, not promised; it is fieldable as one platform.")
B(d,"Walk-out action: ","scope the bounded pilot.")

H1(d,"8. Per-attendee lenses")
B(d,"Dr. Judy Esquibel (ARCYBER operator, your ally): ","detection fidelity, LOTL/valid-account, FP economics, baseline drift. Go deep on the twin and on each layer.")
B(d,"LTC Micah DiGrezio (acquisition): ","one fieldable platform, one vendor. Data/telemetry requirements, footprint (IL5/6, JWICS), integration, sustainment, acquisition path (OTA -> PoR).")
B(d,"Kellsie Herrmann (policy): ","why it matters: Volt/Salt, valid-account critical-infra gap, authorities, no PII in the behavioral vector, readiness. Keep the math light.")

H1(d,"9. Numbers to keep straight (digital twin)")
B(d,"Dataset: ","250 entities, 485 days, ~14M events, 5 log sources.")
B(d,"FP: ","8.1% composite (~20/246 normals above Volt's 13.70), 4/4 caught. Use 8.1, never 8.5.")
B(d,"Ranks: ","Salt #1 (51.27); Insider #2 (46.24); Slow APT #7 (19.44); Volt #24 (13.70).")
B(d,"Baselines: ","Isolation Forest / OC-SVM / LOF each 0/4; z-score 1/4.")
B(d,"Separation is by RANK: ","top normal (21.5) outscores Slow APT (19.4) and Volt (13.7). Present a ranked list.")

H1(d,"10. Logistics (in-person, all presenters present)")
callout(d,"In-person at an Army facility. The #1 failure mode is network access. Plan for no internet in the room.")
B(d,"Network is the top risk: ","assume the room has restricted or no internet. Every demo (all three live layers) must "
  "run locally/offline, over the presenter's own cellular hotspot, or from a full backup recording. Do NOT assume "
  "cloud access to any demo environment.")
B(d,"Backup recordings are mandatory ","for all three live layers, in case a demo stalls or the network is locked down.")
B(d,"Bring-your-own-hardware: ","each presenter brings their own tested demo laptop plus HDMI and USB-C adapters.")
B(d,"One shared display, ","with a clean driver-switch protocol. The session lead keeps the 60-minute boxes.")
B(d,"Arrive early: ","60 to 90 minutes for an AV check and a full demo dry-run in the actual room before the customer arrives.")

P(d,"")
P(d,"Internal preparation document. One company, one platform. Not for distribution.",italic=True,size=9,color=GRAY)
d.save("WP DLA/Presentation/Working_Session_RunOfShow_FourLayer_22CT_INTERNAL.docx")

# =====================================================================
# DOC 2: PRESENTER BRIEFING (one company)
# =====================================================================
d=new_doc()
title_block(d,"Presenter Briefing",
  "Four-Layer Working Session  |  For every presenter on the 22nd Century team",
  "Please read before the session. The customer sees one company, one platform.")

H1(d,"1. The mission")
P(d,"One session, one platform. Everyone presents as 22nd Century. The customer experiences a single integrated "
  "four-layer platform (V-Intelligence). The session lead runs the open, the transitions, and the close. Your job is "
  "to make your layer land as a seamless part of that one platform.")

H1(d,"2. The audience (and why candor is non-negotiable)")
B(d,"Dr. Judy Esquibel ","- Senior Cyber Advisor, U.S. Army Cyber Command (ARCYBER); PhD; critical-infrastructure exercise lead. Deeply technical. She will probe detection fidelity and rewards candor; she detects overclaim instantly.")
B(d,"LTC Micah DiGrezio ","- Army acquisition (PM Cyber & Space). He judges whether this can be bought, integrated, and fielded. He will ask about data, footprint, and sustainment.")
B(d,"Kellsie Herrmann ","- Army cyber/technology policy. She judges strategic significance and authorities. Keep the math light for her.")
callout(d,"One overclaim from any presenter burns the whole team's credibility with this room. When in doubt, under-claim and defer to the session lead.")

H1(d,"3. Brand and vocabulary rules")
B(d,"Present as 22nd Century: ","your layer is \"our Layer 1\" or \"our Layer 3\" of the one platform.")
B(d,"Use MITRE ATT&CK ","as the shared vocabulary, and map your layer to the shared Volt Typhoon / Salt Typhoon scenario.")
B(d,"No pricing, no competitive cross-talk, ","no internal positioning.")
B(d,"Overall-platform, integration, and acquisition questions ","go to the session lead. Answer your-layer technical questions yourself.")

H1(d,"4. Your block: 60 minutes, fixed format")
P(d,"Every layer uses the identical structure so the platform feels like one product:",bold=True)
table(d,["Segment","Time","What to cover"],
 [["Theory","15 min","The threat at your layer, what incumbents do, and the core shortcoming. Map to MITRE."],
  ["Live demo","30 min","Run it live against the shared scenario. No slideware-only. Have a backup recording."],
  ["Results","15 min","The measured result: what you caught that incumbents miss, with real numbers."]],
 widths=[1.1,0.8,4.3],fs=9)
B(d,"Results frame (use this exact arc): ","threat scenario -> what incumbents miss -> live demo -> measured result.")

H1(d,"5. Your layer and your scenario")
table(d,["Layer","Map your demo to this scenario"],
 [["Layer 1 - Preemptive Network Assurance","Prove Volt Typhoon's misconfigured-edge / cross-VLAN entry path is closed before traffic flows. Formal proof over the full config space, not sampling."],
  ["Layer 2 - Behavioral digital twin","Catch the valid-account, living-off-the-land behavior once inside. The four-attacker demo and the 8.1% derivation."],
  ["Layer 3 - API, Data & Agentic-AI Runtime","Catch Salt Typhoon's API and data collection and exfiltration, plus agent/bot abuse, at runtime. Authenticated-abuse and business-logic attacks that signatures miss."]],
 widths=[1.8,4.4],fs=8.5)

H1(d,"6. Federal-readiness you must each cover")
P(d,"DiGrezio judges fieldability across the whole stack, so every layer must speak to:")
B(d,"Data / telemetry: ","what your layer needs, ideally from telemetry already collected.")
B(d,"Footprint and accreditation: ","containerized; deployable across IL5, IL6, and JWICS; air-gap / offline option.")
B(d,"Integration: ","how your layer fits an existing Army cyber stack without rip-and-replace.")
B(d,"Sustainment: ","how updates and new threat coverage are delivered.")

H1(d,"7. Logistics and discipline (IN-PERSON)")
callout(d,"In-person at an Army facility. Assume NO reliable internet in the room. Your demo must work offline.")
B(d,"Bring-your-own-hardware: ","your own tested demo laptop plus HDMI and USB-C adapters.")
B(d,"Network: ","run your demo locally/offline or over your own cellular hotspot, and ALWAYS bring a full backup recording.")
B(d,"On-site dry-run: ","arrive 60 to 90 minutes early and test your demo on the actual room display before the customer arrives.")
B(d,"Time box is hard: ","60 minutes. The session lead keeps time. Practice to fit.")
B(d,"Honesty boundary: ","label anything not in your current build as roadmap, out loud. Do not present roadmap as fielded.")

H1(d,"8. What success looks like")
P(d,"At the end, the room believes they saw one coherent platform, every layer live and measured, under one vocabulary, "
  "from one vendor, with nothing left to doubt. That is the standard.")

P(d,"")
P(d,"Internal team preparation document. One company, one platform. Not for distribution.",italic=True,size=9,color=GRAY)
d.save("WP DLA/Presentation/Presenter_Briefing_FourLayer_Session_INTERNAL.docx")

print("saved one-company docs:")
print("  Working_Session_RunOfShow_FourLayer_22CT_INTERNAL.docx")
print("  Presenter_Briefing_FourLayer_Session_INTERNAL.docx")
