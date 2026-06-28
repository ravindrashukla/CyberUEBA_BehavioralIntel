# -*- coding: utf-8 -*-
"""Brandon Pugh full-stack working-session run-sheet (INTERNAL).
Faithful build of the user's run-sheet: frame, architecture table, roles, agenda,
block-by-block, numbers, tough-Q card, pre-meeting fix list. Output: Presentation/.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6); GRAY=RGBColor(0x55,0x5F,0x6B)
RED=RGBColor(0xA6,0x1B,0x1B); GREEN=RGBColor(0x1F,0x7A,0x33); AMBER=RGBColor(0xB5,0x5A,0x00)
d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(10.5)

def H1(t):
    p=d.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
    return p
def H2(t):
    p=d.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=BLUE
    return p
def P(t,italic=False,bold=False,color=None,size=None):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold
    if color: r.font.color.rgb=color
    if size: r.font.size=Pt(size)
    return p
def B(label,rest="",sub=False):
    p=d.add_paragraph(style='List Bullet 2' if sub else 'List Bullet')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def NUM(label,rest=""):
    p=d.add_paragraph(style='List Number')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p
def callout(text,color=RED):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2)
    r=p.add_run(text); r.bold=True; r.font.color.rgb=color; r.font.size=Pt(9.5)
    return p
def table(headers,rows,widths=None,fs=8.5,hdr_fs=8.5):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(hdr_fs)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""
            parts=v.split("||")
            for j,part in enumerate(parts):
                para=cells[i].paragraphs[0] if j==0 else cells[i].add_paragraph()
                rr=para.add_run(part); rr.font.size=Pt(fs)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    d.add_paragraph(); return t

# ===== TITLE =====
tp=d.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("Full-Stack Working-Session Run-Sheet"); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=NAVY
sp=d.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Brandon Pugh Technical Deep-Dive"); r.font.size=Pt(13); r.font.color.rgb=BLUE
sp2=d.add_paragraph(); sp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp2.add_run("22nd Century  |  V-Intelligence Four-Layer Critical-Infrastructure Defense"); r.font.size=Pt(10); r.italic=True; r.font.color.rgb=GRAY
wn=d.add_paragraph(); wn.alignment=WD_ALIGN_PARAGRAPH.CENTER
rw=wn.add_run("INTERNAL — not the leave-behind"); rw.bold=True; rw.font.color.rgb=RED; rw.font.size=Pt(10)

# ===== 1 FRAME =====
H1("1. Frame (read before you walk in)")
B("Audience: ","Brandon Pugh, cybersecurity lead at R Street, advisor on Pentagon / critical-infrastructure cyber. Technical enough to go deep; his lens is policy and national defense. Rewards candor; discounts anything tuned or overclaimed.")
B("Format: ","2.5 to 3 hour technical working session. He has the layered whitepaper.")
B("The story: ","22nd Century's integrated four-layer defense-in-depth architecture under one MITRE vocabulary, covering the full attack lifecycle from code to network edge to behavior to API/data runtime. The behavioral layer (V-Intelligence) is demoed live; the others are presented as architecture and capability.")
B("Three objectives: ","(1) establish the four-layer architecture closes the Volt/Salt gap end to end; (2) prove the behavioral layer live and honestly; (3) secure a bounded Phase-1 pilot plus a policy / thought-leadership relationship.")
B("Golden rule: ","national problem, then architecture, then proof. Show what runs (Layer 2). Label what is integrated or in development, out loud. Candor is the advantage.")

# ===== 2 ARCHITECTURE =====
H1("2. The architecture at a glance: incumbents, shortcomings, our wedge")
table(["Layer","What it secures","Incumbent approach today","Core shortcoming","22nd Century wedge","Maturity"],
 [["1 - Preemptive Network Assurance","Firewall/NGFW, IdP, IPS, SASE, WAF configuration","Legacy SIEM and monitoring; policy-analysis and breach-simulation tooling; vulnerability management / pen testing","Reactive, or samples the ~2^8000 config space; vuln mgmt covers only ~20% of attacks; none can prove segmentation","Formal mathematical proof over the full state space, deterministic and verifiable; residual-capacity threat blocking","Integrated"],
  ["2 - Behavioral Entity Intelligence (UEBA)","Valid-account, living-off-the-land behavior once inside","Traditional UEBA; point-anomaly ML (Isolation Forest, One-Class SVM, LOF, z-score)","Threshold-based (LoTL never crosses a line); magnitude not meaning; per-source silos; '87% anomalous,' no direction","Meaning over magnitude: 1536-d behavioral embeddings, CUSUM on trajectory, MITRE-direction projection, multi-front composite","LIVE demo (22CT-built)"],
  ["3 - Application, API & Data Runtime Protection","APIs, microservices, data, AI-agent pathways","Traditional WAFs; standalone API-security/ASPM tools; API gateways; bot-management point tools","WAFs miss API nuance and business-logic abuse; signatures miss authenticated abuse (95% of API attacks); agentic AI ungoverned","Unified discover+comply+protect on a behavioral/identity core; runtime-native mitigation; agentic-AI governance","Integrated"],
  ["4 - Code & Supply-Chain Assurance","Code, dependencies, containers, IaC, CI/CD","SAST/DAST; software composition analysis; container/cloud scanners; secrets scanners","Signature / known-CVE matching; can't find novel zero-days; reactive to disclosure; alert fatigue","AI-driven novel zero-day discovery pre-deploy; multi-gate CI/CD assurance","In development"]],
 widths=[1.05,1.0,1.45,1.45,1.45,0.7],fs=7.5,hdr_fs=8)
callout("Internal guardrail (do not blur this in the room): only Layer 2 is 22CT-built and demoable live today. Layers 1, "
 "3, and 4 are presented as architecture and integrated/in-development capability, not as live 22CT demos. The four "
 "layers share a MITRE vocabulary and a vision; they are NOT one turnkey running platform today, and full cross-layer "
 "fusion is the roadmap. Keeping this straight is what protects credibility with Brandon.")

# ===== 3 ROLES =====
H1("3. Roles and pre-flight")
table(["Role","Owner","Job"],
 [["Lead / narrator","You (Ravindra)","Architecture, Layers 1/3/4 framing, the ask, the honesty line"],
  ["Demo driver","(assign)","Layer 2 live demo; files open and tested in advance"],
  ["Scribe","(assign)","His questions and pilot-scope asks to a parking lot"]],
 widths=[1.3,1.7,3.9],fs=9.5)
P("Demo machine (Layer 2), open and tested: simulator/attacks/volt_typhoon.py; data/generated/ (USR-042); "
  "data/tier3_results/weekly_zone_trajectories.csv; composite_scores.csv; detection/ (cusum.py, drift_direction.py, "
  "reference_concepts.py, composite_scorer.py); comparison/run_comparison.py; UI up (host 8001). Dry-run the live 8.1 "
  "percent derivation once.",bold=False)

# ===== 4 AGENDA =====
H1("4. Agenda (target 3h incl. buffer)")
table(["#","Block","Time","Driver"],
 [["0","Open: his goals, time, set the arc","5 min","Lead"],
  ["1","Problem + 4-layer architecture (full attack lifecycle, shared vocabulary)","25 min","Lead"],
  ["2","Layer 1 - Preemptive Network Assurance: incumbents, gap, formal-proof wedge, evidence","20 min","Lead"],
  ["3","Layer 2 - Behavioral Entity Intelligence: LIVE demo; the centerpiece","45 min","Demo driver"],
  ["4","Layer 3 - Application, API & Data Runtime Protection: incumbents, gap, agentic-AI wedge","20 min","Lead"],
  ["5","Layer 4 - Code & Supply-Chain Assurance: incumbents, gap, zero-day wedge","15 min","Lead"],
  ["6","How the layers work together + federal fit","20 min","Both"],
  ["7","Roadmap + the ask","15 min","Lead"],
  ["-","Buffer: his tangents","15 min","-"]],
 widths=[0.4,4.2,0.8,0.9],fs=9)

# ===== 5 BLOCK BY BLOCK =====
H1("5. Block-by-block")
H2("Block 1 - Problem + architecture (25 min)")
B("The inflection: ","Volt Typhoon (CISA AA24-038A): pre-positioning in U.S. critical infrastructure, legitimate Windows tooling only, no malware, 5+ year dwell. Salt Typhoon (CISA AA25-239A): telecom espionage at carrier scale, slow, sub-threshold, valid-account. Common technique: valid accounts (MITRE T1078).")
B("The two structural gaps: ","prevention fails (sampling, not proof) and detection fails (thresholds, not meaning).")
B("The answer: ","four layers, one shared MITRE vocabulary, covering the full lifecycle: code, network edge, behavior, API/data runtime. Each layer catches what the others miss.")
B("Set the honesty frame now: ","\"The behavioral layer is ours and you'll see it run. The other layers are architecture and integrated capability. Full cross-layer integration is what we'd build together, and I'll flag every seam.\"")

H2("Block 2 - Layer 1: Preemptive Network Assurance (20 min)")
B("Mission: ","prove every known attack path through the network controls is closed before traffic flows.")
B("Incumbents and shortcomings: ","legacy SIEM/monitoring is reactive (fires after); breach-simulation and policy tooling only sample the ~2^8000 config space (resource-bound, never complete); vulnerability management covers ~20% of attacks and perfect patching is impractical; none can prove segmentation is correct, so cross-VLAN / shadowed-rule misconfigs are exploited before any alarm. Industry data: 99% of firewall breaches are misconfiguration; >60% of incidents through 2029 traced to misconfigured controls.")
B("Our wedge: ","a formal mathematical model of the entire config, exhaustive, deterministic, verifiable; MITRE-enriched attack models queried against it; uses residual control capacity to add threat-blocking rules from live intel; drift caught within the hour.")
B("Evidence to cite: ","credential/IP-theft path from a shadowed rule plus incomplete object; SMB lateral movement from a 3-rule group shadow; multi-vendor update-traffic RCE path with no inspection; ~25% of internet-facing rules weak/no inspection, proven reachable; SMTP trusted to untrusted = NIST 800-53 SC-7 gap; 45% rule compression; verified attack-surface compliance for a national agency.")
B("Maturity: ","integrated capability. Show the evidence; not the live demo.")
B("Likely Q - 'Isn't this firewall policy analysis?' ","Policy tools lint rules heuristically; this proves reachability and segmentation across the full state space and against MITRE attacks. Proof, not linting.")

H2("Block 3 - Layer 2: Behavioral Entity Intelligence (45 min, LIVE)")
B("Mission: ","catch the apex adversary already inside on valid accounts, the gap Layer 1 cannot see because the credentials are valid.")
B("Incumbents and shortcomings: ","traditional UEBA and the point-anomaly ML inside it (Isolation Forest, One-Class SVM, LOF, z-score) are threshold-based (LoTL never crosses a line); they measure magnitude not meaning (47 files = 47 files); per-source silos make one attacker look like ten users; output is '87% anomalous' with no direction. On identical data, the classical methods catch 0/4; z-score 1/4.")
P("Our wedge, demo it:",bold=True)
NUM("Synthetic data: ","volt_typhoon.py (4 phases, legitimate admin tooling, MITRE tags) to USR-042 event stream. Four attackers hidden among 246 normal users, full ground truth, blind gradeable.")
NUM("Digital twin (lead with the 8-property scorecard): ","5 of 8 built (multi-signal, semantic embedding, trajectory, run-time explainability, partial deploy); identity fusion / versioned store / graph are partial-or-next. Say it out loud; candor is the advantage.")
NUM("Drift / CUSUM: ","healthy flat curve vs the attacker's accumulating drift; minimax-optimal; years of dwell to days. CUSUM is the broad net; composite ranks who matters.")
NUM("Direction / MITRE projection: ","not '87% anomalous' but 'drifting toward insider_threat_slow, alignment 0.68' plus technique ID (14 threat + 2 benign concepts).")
NUM("Composite (5 phases): ","signal strength, breadth, sustained, context divergence, novelty persistence, to one ranked verdict.")
NUM("The proof - Salt #1 (51.27), Volt #24 (13.70): ","walk both. Get ahead of it: 'Our top normal user outscores two of our four attackers, which is exactly why we deliver a ranked list, not a threshold.'")
NUM("Baselines + live 8.1 percent: ","sort by composite; to catch Volt at 13.70, 20/246 normal users score above, so 8.1% FP, 4/4 caught.")
B("Maturity: ","22CT-built. The one layer that runs live. Identity fusion, forensics, response = roadmap.")
B("Likely Q - 'vs traditional UEBA?' ","Meaning vs counts, direction vs magnitude, measured 0/4 vs 4/4 on identical data, ranked explainable verdict mapped to MITRE.")

H2("Block 4 - Layer 3: Application, API & Data Runtime Protection (20 min)")
B("Mission: ","discover, govern, and protect every API and AI-agent pathway to apps and data, at runtime. This is the altitude where Salt Typhoon-style data collection and exfiltration actually happen.")
B("Incumbents and shortcomings: ","traditional WAFs don't understand API nuance and miss business-logic abuse (BOLA); signature-based detection misses authenticated abuse (95% of API attacks come from authenticated sources); standalone point tools (discovery + WAF + bot + testing) stitched together create blind spots and overhead; long baselining and alert noise; and the new blind spot, agentic AI is ungoverned (AI agents talk via APIs; ~99% of AI-related vulnerabilities are API-related; rogue agent connectors and over-permissioned agents).")
B("Our wedge: ","unified discover+comply+protect in one platform (no stitched point tools, no app instrumentation); a behavioral/identity foundation that profiles user/entity/agent behavior, the same philosophy as Layer 2 applied at runtime; real-time native mitigation (block, rate-limit, header injection, deception, inline or passive); agentic-AI governance (secure agent access, govern internal agents, block AI-driven bots, real-time PII/credential redaction). Built for hyperscale and for the public sector.")
B("Maturity: ","integrated capability. Architecture story, not a 22CT demo.")
B("Likely Q - 'How is this different from API-security point tools?' ","Unified behavioral/identity core extended natively to agentic-AI governance with no-instrumentation runtime mitigation, and it shares the behavioral philosophy with our Layer 2, which is what makes the cross-layer story coherent.")

H2("Block 5 - Layer 4: Code & Supply-Chain Assurance (15 min)")
B("Mission: ","find and remove the exploitable vulnerability before deploy, the code/dependency/container/IaC flaw that grants initial access in the first place.")
B("Incumbents and shortcomings: ","SAST/DAST, software composition analysis, container/cloud scanners, secrets scanners. Largely signature / known-CVE matching, can't find novel zero-days; reactive to disclosure (the exploit-to-patch window is exactly what Volt/Salt use); high false positives / developer fatigue; supply-chain injection and adversary backdoors slip past dependency scanners checking only known-bad versions.")
B("Our wedge: ","AI-driven novel zero-day discovery across code, deps, containers, IaC, pre-deploy not post-disclosure; multi-gate CI/CD assurance embedded in the build.")
B("Maturity: ","in development.")
callout("FLAG: the '2,000+ zero-days in 7 weeks' figure must be substantiated/sourced before stating, or label it roadmap. Brandon will check.",color=AMBER)
B("Likely Q - 'Isn't this SAST/SCA?' ","Those match known CVEs; this targets novel zero-day discovery before disclosure.")

H2("Block 6 - How the layers work together + federal fit (20 min)")
B("One vocabulary: ","MITRE ATT&CK across all four layers.")
P("A Typhoon intrusion across the stack:",bold=True)
B("Layer 4 ","removes the exploitable code/dependency that enables initial access.",sub=True)
B("Layer 1 ","proves the misconfigured edge / cross-VLAN path is closed (Volt enters via misconfigured edge devices).",sub=True)
B("Layer 2 ","catches the valid-account, living-off-the-land behavior once inside, the gap Layer 1 can't see.",sub=True)
B("Layer 3 ","catches API/data exfiltration plus agent/bot abuse at runtime, Salt's collection/exfil altitude.",sub=True)
B("Federal fit (Layer 2 specifics): ","containerized, runs on logs already collected; IL5/6, JWICS; no PII in the vector; offline embedding today (non-semantic) with a local semantic model as a Phase-1 build.")
B("The honest integration line: ","\"No single box does this. We've built an integrated four-layer architecture under one MITRE vocabulary. The behavioral layer is ours and runs live; the rest is architecture and integrated capability; full cross-layer fusion is the roadmap.\"")

H2("Block 7 - Roadmap + the ask (15 min)")
B("Phased build (Layer 2 to operational): ","Phase 0 detection core (now); Phase 1 pilot (real-data ingestion, basic identity resolution, versioned twin store, local semantic model); Phase 2 streaming/scale; Phase 3 graph + full identity fusion; Phase 4 evidence/forensics/response.")
B("The ask: ","a bounded Phase-1 pilot: connectors for one enclave's identity + endpoint logs; resolve identities for a defined population; run the validated core on real telemetry; measure detection, explainability, FP economics side-by-side vs current tooling; advisory, human-in-the-loop. Identity fusion is the first co-designed build.")
B("Also: ","R Street thought-leadership on preemptive-plus-behavioral and the national valid-account / LoTL gap; a follow-on engineering session.")
P("Close: \"What you saw run today is the behavioral detection core, on the logs you already collect, catching four "
  "nation-state campaigns the tools agencies deploy score as normal. It sits inside a four-layer architecture: "
  "preemptive, behavioral, API/data, and code, under one MITRE vocabulary. The behavioral layer is ours and runs "
  "live; the rest is architecture and integrated capability; full fusion is what we'd build. The honest ask is a "
  "bounded pilot on real data.\"",italic=True)

# ===== 6 NUMBERS =====
H1("6. Numbers to keep straight (he will check)")
B("Dataset: ","250 entities, 485 days (70 weeks), ~14M events, 5 log sources.")
B("FP rate: ","8.1% composite (about 20/246 normal users above Volt's 13.70), 4/4 caught. Use 8.1, not 8.5.")
B("Ranks: ","Salt USR-118 #1 (51.27); Insider USR-156 #2 (46.24); Slow APT USR-234 #7 (19.44); Volt USR-042 #24 (13.70).")
B("Baselines: ","classical anomaly methods 0/4 (FP 4.5 to 14.6%); z-score 1/4 (9.8%).")
B("Separation is by RANK, not a clean gap: ","top normal user (USR-081, 21.5) outscores Slow APT (19.4) and Volt (13.7). Present a ranked list.")
B("Operating point ","chosen knowing ground truth in the PoC, a reporting benchmark to re-validate, not a deployable threshold.")
B("Context facts (current): ","~99% of AI-related vulnerabilities are API-related; 95% of API attacks come from authenticated sources.")

# ===== 7 TOUGH Q CARD =====
H1("7. Tough-question quick card")
H2("Cross-cutting")
B("'Is this one product or a bundle?' ","Integrated four-layer architecture under one MITRE vocabulary. Behavioral layer runs live; the rest is architecture and integrated capability; full fusion is roadmap. Explicit about the seams.")
B("'What can you actually show me today?' ","The behavioral layer, live, end to end. The other layers as architecture and evidence.")
B("'How integrated is it today?' ","Shared vocabulary and vision now; turnkey cross-layer fusion is roadmap. Don't overclaim.")
H2("Layer-specific")
B("L1 'vs policy-analysis tools?' ","Proof over the full state space + MITRE, not rule-linting.")
B("L2 'vs traditional UEBA?' ","Meaning vs counts; 0/4 vs 4/4 on identical data.")
B("L2 '8% FP at scale?' ","Ranked list, not threshold alerts; analysts work top-down; tune on real data.")
B("L2 'synthetic, real telemetry?' ","Blind test to prove the mechanism; next step is the bounded pilot.")
B("L3 'vs API-security point tools?' ","Unified behavioral/identity core + native agentic-AI governance, no instrumentation; coherent with Layer 2.")
B("L4 'vs SAST/SCA?' ","Novel zero-day discovery pre-disclosure, not known-CVE matching. (Substantiate the 2,000 figure.)")
B("'Data sovereignty?' ","Offline embedding today (non-semantic); local semantic model is Phase 1; no PII in embeddings.")

# ===== 8 FIX LIST =====
H1("8. Pre-meeting fix list (do before he sees the deck)")
P("Internal docs contradict the v6 leave-behind; reconcile or brief the room:")
B("8.5% FP to 8.1% ","(slide 14). [DONE in the deck]")
B("'Known-bad profiles: 4 of 4 at 0 false positives' (slide 20): ","not implemented. Remove / relabel roadmap.")
B("Tamper-evident / HMAC forensics: ","a deck claim, not built. Relabel roadmap.")
B("Proportional response (MFA/lockout): ","a deck claim, not built. Relabel roadmap.")
B("Layer-4 '2,000+ zero-days in 7 weeks': ","substantiate/source or label roadmap.")
B("'Integrated platform' language: ","present as architecture + integrated/in-development capability; cross-layer fusion is roadmap, not a running product.")
B("Baseline results: ","reconcile to one canonical set; retire the others.")
callout("Internal-doc guidance: build-vs-integrated-vs-roadmap distinctions must not be overstated as currently operational "
 "capability. Letting the handout overclaim is the one thing that loses this room.",color=RED)

fp=d.add_paragraph(); fr=fp.add_run("Internal preparation document. Build-vs-integrated-vs-roadmap distinctions and the "
 "reconciled number set. Not for distribution.")
fr.italic=True; fr.font.size=Pt(9); fr.font.color.rgb=GRAY

out="WP DLA/Presentation/Brandon_Pugh_Working_Session_Run-Sheet_INTERNAL.docx"
os.makedirs("WP DLA/Presentation",exist_ok=True)
d.save(out)
print("saved:",out,"| paragraphs:",len(d.paragraphs),"| tables:",len(d.tables))
