# -*- coding: utf-8 -*-
"""Two Brandon-Pugh deliverables:
 1) V-Intelligence_UEBA_Technical_Leave-Behind.docx  (shareable, ~1 page, no secret sauce)
 2) Brandon_Pugh_Tough_Questions_Cheat_Sheet_INTERNAL.docx (internal Q&A)
Grounded numbers: 250 entities/485 days, 4/4 at 10.6% FP, baselines 0-1/4, Volt #30.
No em/en dashes.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6); GRAY=RGBColor(0x55,0x5F,0x6B); RED=RGBColor(0xA6,0x1B,0x1B)

def new_doc():
    d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(10.5); return d
def H1(d,t):
    p=d.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
    return p
def H2(d,t):
    p=d.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=BLUE
    return p
def P(d,t,**k):
    p=d.add_paragraph(); r=p.add_run(t)
    r.italic=k.get('italic',False); r.bold=k.get('bold',False)
    if k.get('color'): r.font.color.rgb=k['color']
    if k.get('size'): r.font.size=Pt(k['size'])
    return p
def B(d,label,rest=""):
    p=d.add_paragraph(style='List Bullet')
    if label: p.add_run(label).bold=True
    if rest: p.add_run(rest)
    return p

# ============ 1) LEAVE-BEHIND (shareable) ============
d=new_doc()
tp=d.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("V-Intelligence UEBA"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
sp=d.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Behavioral Detection for Living-off-the-Land and Valid-Account Threats"); r.font.size=Pt(12); r.font.color.rgb=BLUE
mp=d.add_paragraph(); mp.alignment=WD_ALIGN_PARAGRAPH.CENTER
mp.add_run("Ravindra Shukla  |  V-Intelligence  |  Technical Summary").font.size=Pt(9.5)

H2(d,"The gap")
P(d,"Apex adversaries (Volt Typhoon, Salt Typhoon) are already inside critical infrastructure using valid credentials "
  "and living-off-the-land techniques. Every individual event looks legitimate, so threshold-based UEBA never fires, "
  "signatures never match, and one attacker is fragmented across many identity systems. These campaigns dwell for "
  "years below every alert.")
H2(d,"The approach")
B(d,"Behavior as meaning. ","Each entity's activity is read as meaning in a high-dimensional behavioral space, so two identical counts with different intent separate, where a scalar feature would see them as the same.")
B(d,"Trajectory and direction. ","The system tracks where an entity is heading over time, not just where it stands, and names that direction in MITRE ATT&CK terms (for example, drifting toward lateral movement) rather than a bare anomaly percentage.")
B(d,"Identity fusion and multi-front scoring. ","Fragmented identities are unified into one entity, and independent behavioral fronts fuse into one ranked, explainable verdict, so an attacker extreme on several fronts at once separates from any normal user extreme on one by chance.")
B(d,"Same logs you already collect. ","Sign-in, network, endpoint, file, and identity telemetry, read as living behavioral entities. No new data sources required.")
H2(d,"The proof (blind test)")
P(d,"On a synthetic population of 250 entities over 485 days (about 14 million events across five log sources), four "
  "nation-state-style campaigns were embedded with full ground truth: an 8-month insider, a 180-day slow APT, a Volt "
  "Typhoon-style living-off-the-land campaign, and a Salt Typhoon-style telecom campaign.")
B(d,"Result: ","the multi-front threat-profile detector (the primary detector) caught all four by named known-bad technique (C2-beacon, DGA, LOTL-process, cohort-rare access, recon-fanout, insider-collection) at zero false positives; embedding composite scoring also surfaced all four (ranks 1, 2, 7, and 30 of 250) at a 10.6 percent false-positive operating point, where Isolation Forest, One-Class SVM, and Local Outlier Factor each caught zero of four and a z-score method caught one of four.")
B(d,"Explainable: ","every alert names the behavior, maps it to a MITRE technique, and is delivered as a risk-ranked list rather than a flood of threshold alarms.")
H2(d,"Deployability")
B(d,"Federal-ready. ","Containerized; deployable across IL5, IL6, and JWICS; a local-model option for air-gapped enclaves; no personally identifiable information held in the behavioral representation.")
B(d,"Proven foundation. ","The same behavioral trajectory engine is fielded for supply-chain decision intelligence at the Defense Logistics Agency (Entity Digital Model), independent evidence that the method generalizes.")
H2(d,"Framing and next step")
P(d,"These are controlled, synthetic-data proof-of-concept results, indicative rather than a guarantee of field "
  "performance; the operating point is a reporting benchmark to be re-validated. The recommended next step is a "
  "bounded pilot on representative operational telemetry, run in advisory, side-by-side mode under human control.")
fp=d.add_paragraph(); fr=fp.add_run("Proprietary note: the internal representations, scoring logic, formulas, and "
  "dimensions of V-Intelligence are proprietary and are summarized here at the capability level only.")
fr.italic=True; fr.font.size=Pt(8.5); fr.font.color.rgb=GRAY
out1="WP DLA/Presentation/V-Intelligence_UEBA_Technical_Leave-Behind.docx"
d.save(out1); print("saved:",out1,"| paragraphs:",len(d.paragraphs))

# ============ 2) TOUGH-QUESTIONS CHEAT SHEET (internal) ============
d=new_doc()
tp=d.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("Tough-Questions Cheat Sheet"); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=NAVY
sp=d.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Brandon Pugh Technical Deep-Dive  |  Internal Prep"); r.font.size=Pt(11); r.font.color.rgb=BLUE
wn=d.add_paragraph(); wn.alignment=WD_ALIGN_PARAGRAPH.CENTER
rw=wn.add_run("INTERNAL USE ONLY. Lead with candor; he rewards it."); rw.bold=True; rw.font.color.rgb=RED; rw.font.size=Pt(9.5)

def QA(q,a):
    p=d.add_paragraph(); rq=p.add_run("Q. "+q); rq.bold=True; rq.font.color.rgb=NAVY
    p2=d.add_paragraph(); p2.add_run("A. "+a)
    d.add_paragraph()

H1(d,"Numbers to keep straight")
B(d,"Dataset: ","250 entities, 485 days, about 14 million events, 5 log sources.")
B(d,"Detection: ","threat-profile detector (primary) 4 of 4 at 0 FP by named technique; composite 4 of 4 at 10.6 percent false positive, cleanly separating only 2 of 4 (USR-118, USR-156). Baselines: Isolation Forest, SVM, LOF each 0 of 4; z-score 1 of 4.")
B(d,"Ranks: ","Salt #1, Insider #2, Slow APT #7, Volt #30. Volt is #30, the deck is correct.")
B(d,"Where 10.6 comes from: ","to catch the stealthiest attacker, Volt at composite 12.95, the threshold sits at 12.95; about 26 normal users score above it, so 26 of 246 = 10.6 percent.")
B(d,"Separation is by rank, not a clean gap: ","the top normal user (composite 21.5) outscores Slow APT (20.0) and Volt (12.9); all four still surface at the top of the ranked list. Present it that way.")

H1(d,"Questions and answers")
QA("It is synthetic. How do you know it works on real DoD telemetry?",
   "It is a controlled blind test to prove the mechanism, not operational proof, and we say so. The next step is a bounded pilot on real data.")
QA("10.6 percent false positive is a lot at DoD scale.",
   "It is a risk-ranked list, not threshold alerts; analysts work top-down. The operating point is a proof-of-concept reporting benchmark to be tuned and re-validated on real data.")
QA("Show me where 10.6 percent comes from.",
   "Sort the 250 entities by composite. To catch Volt at 12.95 (rank #30), the threshold is 12.95; 26 normal users score above it, so 26 of 246 = 10.6 percent, catching all four. We can derive it live from the results file.")
QA("Your top normal user outscores two of your attackers, so that is not clean separation.",
   "Correct, and we present it that way. The two stealthiest attackers (Volt, Slow APT) do not separate on raw score; they surface by rank, which is why we deliver a risk-ranked list, not a threshold. The claim is recall at the top of the ranking.")
QA("How is this different from Exabeam, Securonix, or Splunk UBA?",
   "Meaning versus counts, direction versus magnitude, and the measured 0 of 4 versus 4 of 4 on the same data.")
QA("Did you select the operating point with knowledge of ground truth?",
   "Yes, in the proof of concept; it is a reporting benchmark to be re-validated, not a deployable threshold. We are explicit about it.")
QA("You used a commercial embedding model. Data sovereignty and air-gap?",
   "A deterministic offline mode runs today, and a local open embedding model can replace the commercial one in an enclave. No personally identifiable information is placed in the embedding.")
QA("Scale: 250 entities versus millions?",
   "The architecture, cohorts, and snapshot model are designed to scale; we are straight about what is proven (250) versus engineered for scale (roadmap). Scale is a pilot deliverable.")
QA("Can an adversary who knows about behavioral detection evade it?",
   "Multi-front design: being normal on every front at once is far harder than beating one threshold. Hardening the detector itself (adversarial machine learning, MITRE ATLAS) is on the roadmap.")
QA("Are identity fusion, tamper-evident forensics, and automated response built?",
   "The detection core is built and demoed. Identity fusion across many systems, tamper-evident evidence, and proportional response are designed and are the next build phase. We do not present them as running today.")
QA("Are the preemptive, API, and code layers part of this demo?",
   "No. What runs live today is the behavioral detection layer. The other layers are part of the architecture; some are partner or separate capabilities, and we will not present them as running in this demo.")

out2="WP DLA/Presentation/Brandon_Pugh_Tough_Questions_Cheat_Sheet_INTERNAL.docx"
d.save(out2); print("saved:",out2,"| paragraphs:",len(d.paragraphs))
