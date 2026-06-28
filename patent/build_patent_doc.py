"""Generate the provisional-patent invention-disclosure Word document.

Run:  python patent/build_patent_doc.py
Out:  patent/Provisional_Patent_Invention_Disclosure.docx

Innovation is made prominent: a Statement of Innovation box up front, a
novel-vs-prior-art table, and embedded diagrams. Content mirrors the .md disclosure.
NOT legal advice — an inventor's technical disclosure for attorney drafting.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIGDIR = os.path.join(HERE, "..", "docs", "figures")

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
BLUE = RGBColor(0x1B, 0x4F, 0x72)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GRAY = RGBColor(0x55, 0x55, 0x55)


def _base(doc):
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BLUE


def body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, text, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if prefix:
        r = p.add_run(prefix)
        r.bold = True
    p.add_run(text)


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def callout(doc, text, fill="EAF2FB", bold=True):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    _shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10.5)
    r.font.color.rgb = NAVY
    doc.add_paragraph()


def code_block(doc, text):
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    _shade(cell, "F4F6F8")
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    doc.add_paragraph()


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].add_run(hdr).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def figure(doc, fname, caption, width=6.2):
    p = os.path.join(FIGDIR, fname)
    if not os.path.exists(p):
        return False
    doc.add_picture(p, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY
    return True


def build():
    doc = Document()
    _base(doc)

    # ---------- Title ----------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Invention Disclosure for U.S. Provisional Patent Application")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Behavioral Entity Intelligence: Anomaly Detection by the DIRECTION of "
                     "Behavioral Drift in a Unified Cross-Domain Semantic Embedding Space")
    rs.italic = True
    rs.font.size = Pt(13)
    rs.font.color.rgb = BLUE
    warn = doc.add_paragraph()
    warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rw = warn.add_run("FOR ATTORNEY REVIEW — NOT LEGAL ADVICE")
    rw.bold = True
    rw.font.color.rgb = ACCENT
    body(doc, "An inventor's technical disclosure prepared to support drafting and filing of a "
              "U.S. provisional patent application. Counsel should refine claims, run a formal "
              "patentability / freedom-to-operate (FTO) search, and confirm inventorship and "
              "ownership before filing.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------- Statement of Innovation (prominent) ----------
    h1(doc, "Statement of Innovation (What Is Novel)")
    body(doc, "The damaging anomalies — insider threats, slow advanced persistent threat (APT) campaigns, living-off-the-land "
              "intrusions, gradual sourcing shifts, demand regime changes — stay within normal "
              "ranges on every individual metric. They are detectable only by the DIRECTION and "
              "ACCUMULATION of behavioral change. The innovation detects exactly that, and does it "
              "across domains in one space:")
    callout(doc,
            "THE INNOVATION, IN ONE SENTENCE:  represent every entity (user, device, supplier, "
            "part, account) as time-evolving points in ONE semantic embedding space built by "
            "serializing behavior into language, and detect risk by the DIRECTION of behavioral "
            "drift projected onto NAMED concepts (e.g., MITRE ATT&CK) — not by the magnitude of "
            "any single metric.",
            fill="FDECEA")
    bullet(doc, "a single, unified LLM text-embedding space spanning multiple entity types AND "
                "multiple domains (cybersecurity behavior analytics + supply-chain forecasting).",
           prefix="Novel element 1 — Unified cross-domain space: ")
    bullet(doc, "detection by the DIRECTION of drift, via cosine projection onto zone-filtered, "
                "pre-embedded concept vectors mapped to MITRE ATT&CK / SCM taxonomies — answering "
                "'what is this entity becoming,' not 'how big is the number.'",
           prefix="Novel element 2 — Concept-direction detection: ")
    bullet(doc, "three-tier detection architecture integrating traditional ML, embedding-drift "
                "change-point detection, and digital-entity models with 9 detection methods — "
                "velocity magnitude, acceleration, regime detection, zone divergence, contextual "
                "anomaly, relationship drift, behavioral progression, concept-direction alignment, "
                "and novelty persistence — fused via multi-phase composite scoring.",
           prefix="Novel element 3 — Three-tier integrated architecture: ")
    bullet(doc, "semantic gap bypass: direct extraction of high-discrimination features (novel IPs, "
                "domain generation algorithm (DGA) entropy, novel processes) before embedding, combined with z-score-tiered "
                "serialization that makes the embedding quantitatively grounded.",
           prefix="Novel element 4 — Semantic gap bypass + quantitative grounding: ")
    bullet(doc, "the full end-to-end stack: prose serialization + unified space + concept-direction "
                "drift + cohort-relative scoring + zone decomposition + multi-phase fusion + "
                "bi-temporal SCD2 embedding history + drift-augmented dynamic forecast bands.",
           prefix="Novel element 5 — The complete integrated system: ")

    h2(doc, "What Is Novel vs. What Is Prior Art (stated honestly)")
    table(doc, ["Established prior art (NOT claimed as new)", "Claimed as novel (the combination)"],
          [["LLM-embedding of serialized behavior for\nsingle entity type, single domain "
            "(APT-LLM, LogBERT)",
            "ONE UNIFIED space across multiple entity types\n(users, devices, suppliers, items) "
            "AND multiple domains"],
           ["Serialize raw metrics to prose then embed\n(TabLLM — single-table, no zones, no "
            "drift tracking)",
            "Detection by DIRECTION of drift onto NAMED\nconcept vectors with zone-specific "
            "filtering"],
           ["Embed-then-detect via MAGNITUDE in embedding space\n(TAD-Bench benchmarks "
            "IForest/LOF on LLM vectors)",
            "Detection by DIRECTION of drift onto NAMED concept vectors mapped to MITRE ATT&CK "
            "— the concept alignment IS the detector, not a post-hoc enrichment"],
           ["Embedding drift via cosine distance (ZEDD)",
            "Full velocity vectors + regime state machine + behavioral progression (Kendall tau)"],
           ["Peer-cohort baselining; MITRE mapping as enrichment",
            "Co-drifting cohort detection via drift-direction similarity clustering"],
           ["CUSUM change-point detection; kill chain models",
            "Inter-zone divergence as attack-type signature (which zones drift vs. stay stable)"],
           ["Bi-temporal data models; feature stores",
            "Bi-temporal SCD2 embedding history with leak-free point-in-time retrieval"],
           ["Anomaly scoring pipelines",
            "Semantic gap bypass: direct extraction of novel IPs/DGA/processes before embedding"],
           ["Single-tier or two-tier detection systems",
            "Three-tier integrated architecture (traditional ML + embedding drift + digital entity)"],
           ["Statistical time-series forecasting",
            "Drift-augmented DYNAMIC, asymmetric forecast bands shaped by named behavioral concepts"]],
          widths=[3.1, 3.1])
    body(doc, "The concept-direction detection paradigm is fundamentally distinct from all "
              "surveyed prior art. While individual building blocks (embeddings, CUSUM, peer cohorts) "
              "exist in isolation, no system combines them into a unified architecture that detects "
              "by DIRECTION of drift onto NAMED concepts. A competitive review (2024–2026) confirmed "
              "no commercial product or published system replicates this approach.", italic=True,
         color=GRAY)
    callout(doc,
            "SYNERGISTIC RESULTS — The integrated system produces results no individual component "
            "achieves alone:\n"
            "• Novelty persistence (semantic gap bypass) elevated a slow-APT entity from rank #99 "
            "(undetectable) to rank #7 — neither embeddings nor statistics alone accomplished this\n"
            "• Zone divergence correctly distinguished exfiltration from credential compromise "
            "sharing identical aggregate anomaly scores\n"
            "• Co-drifting cohort detection identified lateral movement campaigns invisible to "
            "per-entity analysis\n"
            "• Three-tier architecture achieved 4/4 true-positive detection at 8.5% FP WITH "
            "directional explainability — individual tiers detect anomalies but cannot name "
            "the attack direction or distinguish categories\n"
            "These results demonstrate non-obvious synergy beyond the sum of known parts.",
            fill="E8F5E9")

    # ---------- Brief description of drawings ----------
    h1(doc, "Brief Description of the Drawings")
    bullet(doc, "FIG. 1 is a block diagram of the end-to-end pipeline showing raw behavioral "
                "data serialized to natural language, embedded into a unified semantic space, "
                "decomposed into behavioral zones, and drift direction projected onto named "
                "concept vectors to produce a ranked detection verdict.")
    bullet(doc, "FIG. 2 is a signal chart illustrating that individual behavioral signals "
                "remain below alert thresholds while the fused composite entity score crosses the "
                "detection threshold, demonstrating that the risk is visible only in combination.")
    bullet(doc, "FIG. 3 is a data flow diagram of the forecasting embodiment showing how raw "
                "events become point-in-time features and embeddings, feed the detection and "
                "forecasting models, and yield a calibrated, drift-aware procurement plan.")
    bullet(doc, "FIG. 4 is a comparative chart showing frozen-shape forecast bands (conventional "
                "approach) versus drift-augmented dynamic forecast bands (present invention), "
                "where band shape adapts to the named concept independently of the median.")
    bullet(doc, "FIG. 5 is a cumulative drift trajectory chart comparing an insider-threat "
                "entity (USR-156) whose cumulative semantic drift diverges sharply from the normal "
                "population band (±2σ), and a slow-APT entity (USR-234) whose drift remains within "
                "the normal band — demonstrating that magnitude-based detection alone cannot "
                "reliably distinguish slow threats from normal behavioral variation.")
    bullet(doc, "FIG. 6 is a zone divergence bar chart comparing per-zone drift magnitudes for "
                "two attack types: the insider threat shows dominant drift in the data and access "
                "zones while the slow APT shows dominant drift in the network zone — demonstrating "
                "that inter-zone divergence patterns serve as attack-type signatures.")
    bullet(doc, "FIG. 7 is a layered architecture diagram showing three intelligence layers: "
                "Layer 1 (entity profiles with per-entity zones, embedded and tracked over time), "
                "Layer 2 (relationship dynamics via pairwise Hadamard interaction drift), and "
                "Layer 3 (network intelligence with multi-hop chains, co-drift, and cascade "
                "detection).")
    bullet(doc, "FIG. 8 is a radar chart showing five-phase composite scoring for four attack "
                "entities and the normal-entity median, demonstrating that each attacker is "
                "detected by a DIFFERENT phase mix — proving that no single phase suffices and "
                "multi-phase fusion is necessary for robust detection.")

    # ---------- Visuals ----------
    h1(doc, "Illustrative Diagrams")
    figure(doc, "fig_innov_pipeline.png",
           "FIG. 1 — The pipeline: raw behavior is serialized to language, embedded into one "
           "semantic space, decomposed into zones, and the drift direction is named against MITRE "
           "concepts before a composite ranks the entity.")
    figure(doc, "fig_innov_signal.png",
           "FIG. 2 — The core insight: each individual signal stays below the alert threshold, yet "
           "the fused entity score crosses it — the risk is visible only in combination.")
    figure(doc, "fig_dla_dataflow.png",
           "FIG. 3 — End-to-end data flow (forecasting embodiment): raw events become point-in-time "
           "features and embeddings, feed the models, and yield a calibrated, drift-aware plan.")
    figure(doc, "fig_dla_bands.png",
           "FIG. 4 — Frozen-shape bands (conventional) vs. drift-augmented dynamic bands (this "
           "invention): the band shape adapts to the named concept, independently of the median.")
    figure(doc, "fig_drift.png",
           "FIG. 5 — Cumulative drift trajectory: the insider threat (USR-156, red) diverges "
           "sharply from the normal population band (±2σ, shaded), while the slow APT (USR-234, "
           "green) remains within normal — demonstrating the challenge of magnitude-based detection.")
    figure(doc, "fig_zone.png",
           "FIG. 6 — Zone divergence: per-zone drift magnitudes for two attack types. The insider "
           "threat (red) dominates in data and access zones; the slow APT (green) dominates in the "
           "network zone. The divergence pattern IS the attack-type signature.")
    figure(doc, "fig_bei_layers.png",
           "FIG. 7 — Three-layer intelligence architecture: entity profiles (zones, embeddings, "
           "drift) → relationship dynamics (Hadamard interaction drift) → network intelligence "
           "(multi-hop chains, co-drift, cascade).")
    figure(doc, "fig_radar.png",
           "FIG. 8 — Five-phase composite radar: each attacker is detected by a DIFFERENT phase "
           "mix. USR-234 (slow APT) is dominated by novelty persistence; USR-118 (Salt Typhoon) "
           "by signal strength and breadth; USR-042 (Volt Typhoon) by context divergence. No "
           "single phase detects all four — multi-phase fusion is essential.")

    # ---------- 0 Filing metadata ----------
    h1(doc, "0.  Filing Metadata (to be completed)")
    table(doc, ["Field", "Value"],
          [["Proposed title", "Behavioral Entity Intelligence: Anomaly Detection by the Direction "
            "of Behavioral Drift in a Unified Cross-Domain Semantic Embedding Space"],
           ["Inventor(s)", "[full legal name(s), address(es)] — TO COMPLETE"],
           ["Assignee / applicant", "22nd Century Technologies, Inc. (confirm)"],
           ["Conception / reduction-to-practice date", "[date(s)] — TO COMPLETE (attach commit "
            "history / notebook as evidence)"],
           ["Prior public disclosure?", "[CONFIRM any public demo/whitepaper/RFI — affects "
            "12-month statutory bar and foreign rights] — TO CONFIRM"]],
          widths=[2.0, 4.2])
    body(doc, "Note for counsel: provisional applications are not examined and need not include "
              "formal claims, but the draft claims in Section 9 anchor the eventual non-provisional. "
              "Confirm the public-disclosure timeline (on-sale / public-use bar).", italic=True,
         color=GRAY)

    # ---------- 1 Field ----------
    h1(doc, "1.  Field of the Invention")
    body(doc, "Machine-learning systems for anomaly and risk detection over the behavior of "
              "monitored entities. More particularly, representing heterogeneous entities (users, "
              "devices, network segments, applications, suppliers, inventory items, accounts) as "
              "time-evolving points in a unified semantic vector space derived from natural-language "
              "serialization of behavior, and detecting anomalies by the direction of behavioral "
              "change projected onto named reference concepts — across cybersecurity behavior "
              "analytics and supply-chain demand forecasting.")

    # ---------- 2 Background ----------
    h1(doc, "2.  Background")
    h2(doc, "2.1  The problem")
    body(doc, "The most consequential anomalies remain within normal statistical ranges on every "
              "individual metric. The damaging signal is in the DIRECTION and ACCUMULATION of "
              "change, not the MAGNITUDE of any single metric crossing a threshold.")
    h2(doc, "2.2  Limitations of existing approaches")
    bullet(doc, "baseline entities statistically (p-value, histogram/profile deviation, isolation "
                "forests, one-class methods) on numeric features; recent generative-AI is an "
                "analyst assistant, not the detector.", prefix="Behavioral security analytics: ")
    bullet(doc, "statistical + gradient-boosted/deep tabular models over numeric features, with "
                "fixed-shape uncertainty bands and generative-AI as a conversational layer.",
           prefix="Demand / supply-chain forecasting: ")
    bullet(doc, "tokenize numeric series directly rather than representing behavior as language.",
           prefix="Foundation time-series models: ")
    body(doc, "None (a) represent heterogeneous entities and multiple domains in one semantic "
              "space, (b) detect by the DIRECTION of drift projected onto named concepts, or "
              "(c) adjust forecasts/uncertainty by that direction.")
    h2(doc, "2.3  Why direction carries more information than magnitude")
    body(doc, "A scalar anomaly score (magnitude) collapses a D-dimensional embedding change into "
              "one number, discarding the directional information. The direction of drift lies on "
              "a (D−1)-dimensional unit hypersphere and preserves all directional information. "
              "Projecting this direction onto K named concept vectors yields K "
              "alignment scores that simultaneously answer (a) whether the entity is anomalous, "
              "(b) what kind of anomaly it is, and (c) which behavioral zone drives the change — "
              "collapsing detection, classification, and triage into a single computation. "
              "Magnitude-based methods require a separate classification step after detection, "
              "with no guarantee the anomaly explanation is causally related to the detection "
              "signal.")

    # ---------- 3 Summary ----------
    h1(doc, "3.  Summary of the Invention")
    for n, txt in enumerate([
        "Serializes behavior to language with z-score tiering: aggregates raw metrics into "
        "behavioral zones, annotates each metric with a statistical tier (normal/elevated/critical "
        "relative to population and entity baseline), and renders as natural-language text.",
        "Embeds into a unified semantic space common to all entity types and domains.",
        "Composes an entity embedding from zone embeddings (weighted average and/or context-adaptive "
        "attention), with automatic exclusion of static (near-zero variance) zones.",
        "Tracks a trajectory: stores successive entity embeddings with bi-temporal SCD2 versioning "
        "(valid-time + knowledge-time) for leak-free point-in-time retrieval.",
        "Detects by DIRECTION of drift: projects the full-dimensional drift vector onto "
        "zone-filtered, pre-embedded concept vectors mapped to MITRE ATT&CK / SCM concepts.",
        "Detects inter-zone divergence: identifies attack-type signatures from which zones drift "
        "while others remain stable (e.g., data zone drifting + identity stable = exfiltration).",
        "Accumulates slow drift via CUSUM change-point detection; detects behavioral progression "
        "via Kendall tau rank correlation; classifies regime states (stable/drifting/accelerating/"
        "regime_shift) via velocity state machine.",
        "Scores relative to a peer cohort; detects co-drifting cohorts via drift-direction "
        "similarity clustering (coordinated insider rings, lateral movement campaigns).",
        "Extracts high-discrimination features (novel IPs, DGA entropy, novel processes) directly "
        "from text before embedding, bypassing the semantic compression layer.",
        "Fuses multiple phases (signal strength, breadth, sustained deviation, context divergence, "
        "novelty persistence) into one ranked verdict; models pairwise relationships via Hadamard "
        "products and session-based relationship entities.",
        "Reconstructs multi-step kill chains via temporal correlation across entities, ordered by "
        "MITRE ATT&CK tactic sequence, with automated natural-language narrative generation.",
        "Drift-augments forecasts: bounded plan adjustment + DYNAMIC, asymmetric uncertainty bands.",
        "Is domain-portable: same architecture, only serialization/zones/concepts change.",
    ], 1):
        bullet(doc, txt, prefix=f"({n}) ")

    # ---------- 4 Three-tier detection architecture ----------
    h1(doc, "4.  Three-Tier Detection Architecture")
    body(doc, "The system operates as a layered detection pipeline where each tier addresses a "
              "different class of anomaly. Tiers run in parallel; their outputs are fused into a "
              "unified verdict.")
    bullet(doc, "Six established statistical/ML algorithms (Isolation Forest, One-Class SVM, "
                "Local Outlier Factor, Z-Score, Temporal-Z, Feature Cumulative Sum (CUSUM)) operating on 23 scalar "
                "behavioral features. Detects magnitude anomalies but provides no directional "
                "intelligence for triage.",
           prefix="Tier 1 — Traditional ML: ")
    bullet(doc, "Embedding-drift analysis via CUSUM on cosine distances between successive "
                "behavioral embeddings. Detects entities whose behavioral meaning is changing, "
                "even when all scalar features remain within normal ranges.",
           prefix="Tier 2 — Embedding Drift: ")
    bullet(doc, "Full digital-entity model with 9 detection methods — velocity magnitude, "
                "acceleration, regime detection, zone divergence, contextual anomaly, relationship "
                "drift, behavioral progression, concept-direction alignment, and novelty "
                "persistence — fused via multi-phase composite scoring into one ranked verdict.",
           prefix="Tier 3 — Digital Entity Intelligence: ")
    body(doc, "The three-tier architecture is novel in its integration: no prior system combines "
              "traditional statistical methods, embedding-drift change-point detection, and "
              "entity-level digital models with concept-direction detection in a single pipeline.",
         italic=True, color=GRAY)

    # ---------- 5 Detailed description (math) ----------
    h1(doc, "5.  Detailed Description")
    h2(doc, "5.1  Entity digital model and behavioral zones")
    body(doc, "Each entity is modeled as a structured, multi-dimensional record. Top-level "
              "groupings are called 'behavioral zones' (also referred to as 'dimensions' or "
              "'slots' in implementation variants; the term 'behavioral zone' is used throughout "
              "this disclosure). Each behavioral zone holds sub-dimensions (individual features); "
              "the record is captured as a TIME SERIES of snapshots. Cyber-user zones: identity, "
              "access pattern, data behavior, network footprint, risk posture. Supplier/item "
              "zones: identity, performance, geographic/sourcing, network position, risk/financial.")
    body(doc, "In an alternative embodiment, the entity's behavioral state may be serialized as a "
              "single unified natural-language description without decomposition into separate "
              "behavioral zones, and the resulting entity vector may be used directly for drift "
              "computation and concept projection without a zone composition step. This zoneless "
              "embodiment retains the core concept-direction detection mechanism (serialization → "
              "embedding → drift → concept projection) while omitting zone-level granularity, and "
              "is applicable when zone-specific attack-type signatures are not required.")
    h2(doc, "5.2  Natural-language serialization with z-score tiering")
    body(doc, "Each zone's metrics are serialized to text; in the preferred interpretive form each "
              "metric is expressed relative to population, the entity's own baseline, and trend, "
              "preserving qualitative tokens (directories, IPs, domains, country-of-origin).")
    body(doc, "Critically, each metric is annotated with a statistical tier before embedding:")
    code_block(doc, "NORMAL    |z| < 1.5    metric is within expected variation\n"
                    "ELEVATED  1.5 <= |z| < 2.5    metric is noteworthy\n"
                    "CRITICAL  |z| >= 2.5    metric is a strong statistical outlier\n\n"
                    "Example:  'auth_fail_rate is CRITICAL at 4.2x population average (z=3.1),\n"
                    "           up from NORMAL last period'")
    body(doc, "This tiering ensures the embedding captures statistical significance, not just "
              "raw values — a subtle but essential innovation that makes the semantic space "
              "quantitatively grounded. In an alternative embodiment, the natural-language "
              "description may be generated from raw metric values without statistical-tier "
              "annotations, relying on the text-embedding model to capture behavioral state "
              "from the raw values alone.")
    h2(doc, "5.3  Unified cross-domain embedding")
    body(doc, "Each description is embedded into a fixed-dimensional vector (D=1536 in the preferred embodiment) in ONE "
              "semantic space common to all entity types and domains; vectors are cached by content "
              "hash and stored in a vector database with approximate-nearest-neighbor search.")
    h2(doc, "5.3.1  Bi-temporal embedding history")
    body(doc, "Embeddings are stored with two independent time axes following slowly-changing "
              "dimension type 2 (SCD2) semantics:")
    code_block(doc, "valid_from / valid_to      — real-world behavioral truth\n"
                    "knowledge_from / knowledge_to  — when the system learned the information\n\n"
                    "Three-step SCD2 upsert protocol:\n"
                    "  1. Supersede current open row (set knowledge_to = now on the\n"
                    "     existing open-ended record)\n"
                    "  2. Insert closed-valid historical row with the OLD embedding\n"
                    "     vectors (valid_to = now, preserving the prior behavioral state)\n"
                    "  3. Insert new open row with the NEW embedding vectors\n"
                    "     (valid_from = now, valid_to = NULL, knowledge_from = now)\n\n"
                    "Point-in-time query:\n"
                    "  WHERE valid_from <= t AND (valid_to IS NULL OR valid_to > t)")
    body(doc, "This enables leak-free point-in-time retrieval and retroactive correction of "
              "embeddings — essential for reproducible training and regulatory audit trails.")
    h2(doc, "5.4  Composition with static zone exclusion")
    code_block(doc, "weighted:   V = normalize( SUM_z  w_z * e_z )\n"
                    "attention:  a_z = ( ||e_z|| * b_z(c) ) / SUM_j( ||e_j|| * b_j(c) )\n"
                    "            V = normalize( SUM_z  a_z * e_z )\n\n"
                    "'Linear normalization' as used herein: dividing each zone's raw\n"
                    "attention score by the sum of all zones' raw attention scores, i.e.,\n"
                    "  a_z = raw_z / SUM_j(raw_j)   where raw_z = ||e_z|| * b_z(c)\n"
                    "This preserves intended weight differentiation across investigation\n"
                    "contexts (unlike softmax, which compresses outlier weights).\n\n"
                    "Static zone exclusion:  zones designated as static (e.g., identity for\n"
                    "cyber-user entities) are excluded from the weighted composition, preventing\n"
                    "unchanging zones from diluting anomaly signals in the composed entity embedding.\n\n"
                    "Example context weight profiles (cyber-user entity):\n"
                    "  normal_ops:            identity=0.20  access=0.20  data=0.20  network=0.20  risk=0.20\n"
                    "  insider_investigation: identity=0.10  access=0.15  data=0.40  network=0.15  risk=0.20\n"
                    "  apt_hunt:              identity=0.05  access=0.15  data=0.10  network=0.40  risk=0.30\n"
                    "  privilege_audit:       identity=0.10  access=0.25  data=0.10  network=0.15  risk=0.40")
    body(doc, "Multi-context detection: in the preferred embodiment, zone vectors are composed "
              "under each of the plurality of context weight profiles to produce a respective "
              "entity vector per context. Concept-direction alignment (Section 5.6) is then "
              "computed for each entity vector. The investigation context whose entity vector "
              "yields the strongest concept-direction alignment is selected as the primary "
              "detection context for that entity. This enables a single entity to be evaluated "
              "under, for example, an insider_investigation context and an apt_hunt context "
              "simultaneously, with the context producing the clearest threat signal being "
              "surfaced for analyst triage.")
    h2(doc, "5.5  Drift — full velocity vectors and regime detection")
    code_block(doc, "magnitude:  d_t = 1 - cos( V_{t-1}, V_t )\n"
                    "direction:  u_t = ( V_t - V_{t-1} ) / || V_t - V_{t-1} ||\n"
                    "            (when || V_t - V_{t-1} || < epsilon, the entity is\n"
                    "             classified as stable; concept projection is skipped)\n\n"
                    "Velocity vector (full 1536-d, not scalar):\n"
                    "  vel_t = V_t - V_{t-1}   (preserves direction + magnitude)\n\n"
                    "Six trajectory features per entity per window:\n"
                    "  velocity_magnitude, acceleration, stability, regime_shifts,\n"
                    "  trend_consistency (mean cosine similarity of consecutive\n"
                    "  direction vectors), total_drift")
    body(doc, "A regime detection state machine classifies each entity into one of four states "
              "(evaluated in priority order):")
    code_block(doc, "REGIME_SHIFT     regime_shifts > 0  (consecutive pairs with similarity < 0.7)\n"
                    "ACCELERATING     acceleration > 0.01  AND  trend_consistency > 0.3\n"
                    "DRIFTING         total_drift > 0.05   AND  stability > 0.7\n"
                    "STABLE           default (none of the above triggered)\n\n"
                    "Thresholds are configurable per deployment; defaults above are from\n"
                    "the reduction-to-practice embodiment.")
    h2(doc, "5.6  Concept projection with zone-specific filtering")
    body(doc, "A library of reference concepts is pre-embedded from natural-language descriptions. "
              "The drift direction is projected onto each concept; the strongest alignment names "
              "what the entity is becoming and maps to a taxonomy (e.g., MITRE ATT&CK).")
    code_block(doc, "align_k = cos( u_t , Embed(concept_k) )   -> rank -> MITRE / SCM concept")
    body(doc, "In the preferred embodiment, the similarity metric is cosine similarity. In "
              "alternative embodiments, the similarity metric may comprise dot product similarity, "
              "Euclidean distance-based similarity, or other distance-based similarity measures "
              "computable in the embedding space; cosine similarity is preferred because it is "
              "invariant to vector magnitude and directly measures angular alignment.")
    body(doc, "In a further alternative embodiment, the drift vector may be classified using a "
              "trained model (e.g., a supervised classifier or neural network) that has been trained "
              "to associate drift vectors with named behavioral concepts, producing a detection "
              "output identifying a named concept. This classifier-based approach captures the same "
              "drift-to-concept mapping as concept projection but learns the mapping from labeled "
              "examples rather than computing geometric similarity.")
    body(doc, "Zone-specific concept filtering maps each behavioral zone to a subset of "
              "semantically relevant reference concepts, preventing false concept alignments:")
    code_block(doc, "ZONE_RELEVANT_CONCEPTS = {\n"
                    "  'identity':          [],                         # excluded (static zone)\n"
                    "  'access_pattern':    ['credential_stuffing', 'privilege_escalation', ...],\n"
                    "  'data_behavior':     ['data_exfiltration', 'insider_threat_slow', ...],\n"
                    "  'network_footprint': ['c2_beacon', 'lateral_movement', 'reconnaissance', ...],\n"
                    "  'risk_posture':      ['compromised_endpoint', 'privilege_escalation', ...],\n"
                    "}\nConcept names map to MITRE ATT&CK technique IDs via a reference table.\n"
                    "align_k = cos( u_t , Embed(concept_k) )  only where k in ZONE_RELEVANT[zone]")
    h2(doc, "5.6.1  Zone divergence detection")
    body(doc, "Inter-zone divergence — the difference in drift magnitudes across behavioral zones "
              "— serves as an attack-type signature. When one zone drifts strongly while others "
              "remain stable, the divergence pattern names the attack category:")
    code_block(doc, "data_behavior drifting + identity stable  ->  exfiltration (not role change)\n"
                    "network_footprint drifting + access stable  ->  C2 channel (not workflow change)\n"
                    "identity drifting + data stable  ->  credential compromise\n\n"
                    "divergence_score = max_zone_drift - median_zone_drift")
    h2(doc, "5.7  Slow-drift accumulation (CUSUM)")
    code_block(doc, "S_0 = 0 ;  S_t = max( 0, S_{t-1} + (d_t - k) ) ;  alarm when S_t >= h for r periods\n\n"
                    "Parameters and exemplary defaults:\n"
                    "  k = baseline drift allowance (exemplary default: 0.02)\n"
                    "  h = detection threshold on cumulative sum (exemplary default: 0.05)\n"
                    "  r = minimum consecutive periods above threshold (exemplary default: 2)\n\n"
                    "The key insight: individual drift magnitudes may be small (below any\n"
                    "single-period alert threshold), but the CUSUM accumulates them until\n"
                    "the cumulative sum triggers detection.")
    h2(doc, "5.7.1  Behavioral progression detection (Kendall tau)")
    body(doc, "Monotonic behavioral progressions (e.g., insider escalation over months) are "
              "detected via Kendall tau rank correlation across sequential observations:")
    code_block(doc, "tau = kendall_tau( [1, 2, ..., N], [drift_mag_w1, drift_mag_w2, ..., drift_mag_wN] )\n"
                    "progression_alert when tau > 0.6 over >= 8 periods\n\n"
                    "Note: this is a bivariate rank correlation between drift magnitudes and\n"
                    "their temporal ordering (Mann-Kendall trend test), detecting monotonic\n"
                    "escalation that individual-point anomaly methods cannot capture.")
    body(doc, "This catches the gradual escalation pattern that individual-point anomaly "
              "detection methods fundamentally cannot see.")
    h2(doc, "5.8  Peer-cohort relative scoring and co-drifting cohort detection")
    code_block(doc, "z[i,f] = ( x[i,f] - mu_cohort[f] ) / sigma_cohort[f]")
    body(doc, "Beyond individual cohort scoring, the system detects co-drifting cohorts — groups "
              "of entities whose drift direction vectors exhibit high cosine similarity:")
    code_block(doc, "co_drift_sim(a,b) = cos( drift_vec_a, drift_vec_b )\n"
                    "Greedy single-linkage clustering where sim > T_codrift\n"
                    "Cohort coherence = mean pairwise cosine similarity within cluster\n\n"
                    "Application: coordinated insider rings, supply-chain collusion,\n"
                    "             lateral movement campaigns affecting multiple entities")
    body(doc, "Identified co-drifting cohorts and their coherence scores are transmitted to a "
              "security orchestration, automation, and response (SOAR) system via a structured "
              "alert payload containing the cluster member identifiers, coherence score, dominant "
              "concept alignment, and recommended investigation actions, enabling coordinated-"
              "threat investigation and automated response workflows.")
    h2(doc, "5.9  Relationship embeddings and session entities")
    code_block(doc, "R[a x b] = normalize( V_a (*) V_b )      # element-wise (Hadamard) product")
    body(doc, "Additionally, sessions are modeled as relationship entity types that bridge user "
              "and device embeddings via temporal co-occurrence within a bounded time window. "
              "The session relationship vector is computed as the normalized Hadamard product of "
              "the user embedding and the device embedding observed during the session window, "
              "following the same formula R[a x b] = normalize(V_a (*) V_b). Drift detection on "
              "the session relationship vector captures changes in user-device interaction "
              "patterns.")
    h2(doc, "5.10  Multi-phase composite scoring")
    body(doc, "A final verdict fuses independent phases into one ranked score, so no single "
              "evasion technique defeats detection:")
    code_block(doc, "composite = w1 * signal_strength      # peak concept-alignment score\n"
                    "          + w2 * breadth              # count of zones with drift > threshold\n"
                    "          + w3 * sustained_deviation   # CUSUM accumulated value\n"
                    "          + w4 * context_divergence    # entity z-score vs peer cohort\n"
                    "          + w5 * novelty_persistence   # duration of novel features\n\n"
                    "Threshold: rank by composite; alert above Nth percentile.\n"
                    "Operating point: 90th percentile yields ~8.5% FP rate.\n\n"
                    "Cross-tier fusion (geometric-mean in log-space):\n"
                    "  pct_floored = max(0.01, pct)  for each rank percentile (avoids log(0))\n"
                    "  depth = exp( 0.05*log(velocity_pct) + 0.30*log(zone_pct)\n"
                    "             + 0.05*log(context_pct)  + 0.30*log(cusum_pct)\n"
                    "             + 0.30*log(progression_pct) )\n"
                    "  breadth = count of distinct methods flagging the entity\n"
                    "  final = 0.7 * rank_normalize(depth) + 0.3 * rank_normalize(breadth)\n\n"
                    "Per-alert confidence fusion:\n"
                    "  base = min(0.6, drift_magnitude / 0.5)\n"
                    "  alignment_boost = min(0.3, mean_concept_similarity * 0.4)\n"
                    "  method_boost = {cusum: 0.10, drift_direction: 0.08, cohort: 0.05}\n"
                    "  confidence = min(1.0, base + alignment_boost + method_boost)")
    h2(doc, "5.10.1  Semantic gap bypass for high-discrimination features")
    body(doc, "Certain high-discrimination features — novel IP addresses, novel process "
              "executions, DGA domain entropy — are extracted directly from behavioral text prior "
              "to embedding, bypassing the embedding layer. This addresses the semantic gap problem "
              "where embedding compresses novelty signals into a dense vector that dilutes their "
              "detection sensitivity:")
    code_block(doc, "Direct extraction (bypasses embedding):\n"
                    "  novel_ips = IPs not seen in baseline period\n"
                    "  dga_score = shannon_entropy(domain_name)  ->  flag if > 3.5\n"
                    "  novel_processes = processes not in baseline inventory\n\n"
                    "Novelty persistence metrics:\n"
                    "  persistent_novel_ips:  count of novel IPs recurring > 4 weeks\n"
                    "  novel_ip_max_persistence:  maximum total weeks any single novel IP appeared\n"
                    "  novel_ip_weeks_frac:  fraction of post-baseline weeks with novel IP contact")
    h2(doc, "5.10.2  Role-appropriateness annotation")
    body(doc, "File access paths are compared against role-expected and sensitive directory "
              "mappings before embedding. Access outside expected patterns is tagged as NOVEL or "
              "ATYPICAL, ensuring the embedding captures role-inappropriate access as semantically "
              "distinct from routine access at the same volume.")
    h2(doc, "5.11  Kill chain reconstruction")
    body(doc, "Correlated alerts across entities and time windows are assembled into multi-step "
              "kill chains ordered by MITRE ATT&CK tactic sequence:")
    code_block(doc, "1. Temporal correlation: group alerts within 72-hour sliding windows\n"
                    "2. Entity-hop linkage: connect alerts where one entity's target\n"
                    "   becomes another entity's source\n"
                    "3. Tactic ordering: sort by MITRE ATT&CK kill chain phase\n"
                    "   (Recon -> Initial Access -> Execution -> ... -> Exfiltration)\n"
                    "4. Narrative generation: automated natural-language summary of the\n"
                    "   reconstructed attack sequence\n"
                    "5. Staleness: mark chains with no new events in > 2x correlation\n"
                    "   window (default 144h) as stale\n\n"
                    "Chain confidence score:\n"
                    "  tactic_coverage = unique_tactics_in_chain / total_kill_chain_phases\n"
                    "  temporal_coherence = 1 / (1 + mean_inter_alert_gap_hours)\n"
                    "  linkage_strength = entity_hop_count / total_alerts_in_chain\n"
                    "  confidence = w_t * tactic_coverage + w_c * temporal_coherence + w_l * linkage_strength\n"
                    "  (weights w_t, w_c, w_l are predetermined and sum to 1.0)\n\n"
                    "In the preferred embodiment, the chain confidence score is computed\n"
                    "using the weighted formula above. In an alternative simplified\n"
                    "embodiment, tactic count serves as a severity proxy.")
    body(doc, "The reconstructed kill chain, its confidence score, and the natural-language "
              "narrative are transmitted via a structured alert payload to a security orchestration, "
              "automation, and response (SOAR) system to trigger automated containment actions "
              "based on chain confidence and tactic coverage. The payload includes entity "
              "identifiers, tactic sequence, confidence sub-scores, and the generated narrative. "
              "Chains marked as stale are deprioritized in the SOAR queue.")
    h2(doc, "5.12  Forecasting embodiment — drift-augmented dynamic bands")
    code_block(doc, "direction(concept) = +1 if concept implies increased demand, -1 if decreased\n"
                    "alpha   = clamp( direction(concept) * alignment * drift_norm , -0.25, +0.25 )\n"
                    "adj_P50 = max(0, P50 * (1 + alpha))     # median shifts up to ±25%\n"
                    "adj_P90 = adj_P50 + (P90 - P50) * beta  # bands widen/narrow INDEPENDENTLY of median\n"
                    "  (beta > 0 required; beta > 1 widens bands, beta < 1 narrows, beta = 1 unchanged)\n\n"
                    "Clamp bounds (±0.25) are predetermined safety gates preventing\n"
                    "drift from dominating the baseline forecast.")
    body(doc, "Worked example with per-concept beta values:")
    code_block(doc, "Concept: 'supplier_disruption'  -> beta = 1.4  (widen bands: more uncertainty)\n"
                    "  adj_P90 = adj_P50 + (P90 - P50) * 1.4   # bands expand 40%\n\n"
                    "Concept: 'demand_surge'          -> beta = 0.7  (narrow bands: confident uptick)\n"
                    "  adj_P90 = adj_P50 + (P90 - P50) * 0.7   # bands contract 30%\n\n"
                    "Concept: 'seasonal_normalization' -> beta = 1.0 (unchanged: expected pattern)\n"
                    "  adj_P90 = adj_P50 + (P90 - P50) * 1.0   # bands unchanged")
    body(doc, "The adjusted forecast, including the drift-aware quantile bands and the identified "
              "concept direction, is transmitted via a structured data interface to an automated "
              "procurement or inventory management system, which adjusts resource allocation "
              "quantities based on the drift-aware forecast bands.")
    h2(doc, "5.13  Cross-domain portability")
    body(doc, "The same pipeline instantiates in a new domain by changing only serialization "
              "templates, zone definitions, and the concept library. Disclosed embodiments: "
              "(a) user and entity behavior analytics (UEBA) for cybersecurity and (b) demand forecasting / supplier-risk for defense "
              "logistics — both in the same unified embedding space.")
    body(doc, "Additional applicable domains demonstrating portability (zone/concept/entity "
              "mappings differ, architecture is identical):")
    bullet(doc, "entities = healthcare providers/patients; zones = billing, prescribing, referral, "
                "demographics; concepts = upcoding, phantom billing, kickback patterns.",
           prefix="Healthcare fraud detection: ")
    bullet(doc, "entities = accounts/transactions; zones = transaction pattern, counterparty "
                "network, geographic flow; concepts = layering, structuring, shell company pattern.",
           prefix="Anti-money laundering: ")
    bullet(doc, "entities = equipment/components; zones = vibration, thermal, operational load; "
                "concepts = bearing wear, seal degradation, impeller imbalance.",
           prefix="Prescriptive maintenance: ")
    bullet(doc, "entities = personnel/access events; zones = access pattern, travel, "
                "communication, financial; concepts = recruitment indicators, exfiltration tradecraft.",
           prefix="Counter-intelligence: ")

    # ---------- 6 Advantages ----------
    h1(doc, "6.  Advantages")
    for txt in [
        "Detects constant-volume, direction-only changes that magnitude/threshold methods miss.",
        "Explainable output (named concept + taxonomy + zone attribution) instead of an opaque score.",
        "Zone divergence patterns name the attack category without additional analysis.",
        "One unified space enables cross-entity similarity, cohort cold-start transfer, and co-drift.",
        "Co-drifting cohort detection identifies coordinated campaigns across multiple entities.",
        "Kill chain reconstruction automates multi-step attack narrative generation.",
        "Bi-temporal embedding history enables leak-free training and regulatory audit trails.",
        "Semantic gap bypass preserves detection sensitivity for high-discrimination features.",
        "Three-tier layered architecture enables defense-in-depth detection.",
        "Domain-portable architecture reduces re-engineering to serialization/zones/concepts only.",
        "Dynamic, asymmetric forecast bands improve probabilistic decision support.",
        "Regime detection state machine classifies behavioral trajectories enabling anticipatory response.",
        "Multi-context detection evaluates each entity under multiple investigation lenses simultaneously.",
        "Session relationship entities capture interaction-level behavioral changes invisible to individual entity analysis.",
    ]:
        bullet(doc, txt)

    # ---------- 7 Reduction to practice ----------
    h1(doc, "7.  Reduction to Practice (synthetic validation — summary)")
    body(doc, "All figures are from SYNTHETIC datasets and illustrate operability; they are not "
              "performance claims for real-world deployment.", italic=True, color=ACCENT)
    bullet(doc, "250 entities × ~485 days; 4 embedded long-duration campaigns (insider threat "
                "14 months, slow APT 417 days, Volt Typhoon LOTL 412 days, Salt Typhoon telecom "
                "412 days).",
           prefix="Security dataset: ")
    bullet(doc, "Composite scores: USR-118 (Salt Typhoon) = 51.27 rank #1; USR-156 (insider) "
                "= 46.24 rank #2; USR-234 (slow APT) = 19.44 rank #7; USR-042 (Volt Typhoon) "
                "= 13.70 rank #24. All 4 attackers detected at 90th-percentile threshold "
                "(8.5% FP, 21 of 246 normal entities).",
           prefix="Detection results: ")
    bullet(doc, "Five-phase scoring: signal strength (peak concept alignment), breadth (zones "
                "above threshold), sustained deviation (CUSUM accumulated value), context "
                "divergence (cohort z-score), novelty persistence (novel IP/process duration). "
                "Key finding: novelty persistence elevated USR-234 from rank #99 to #7 — the "
                "semantic gap bypass was the decisive feature.",
           prefix="Scoring detail: ")
    bullet(doc, "Traditional Tier 1 detectors: Z-Score detected 4/4 at 4.9% FP, LOF 3/4 at "
                "3.3% FP, IForest 3/4 at 7.3% FP, OCSVM 4/4 at 29.7% FP — demonstrating that "
                "magnitude-based detection CAN flag anomalous entities but provides no directional "
                "intelligence to distinguish attack category or triage priority.",
           prefix="Tier comparison: ")
    bullet(doc, "500 items × 200 suppliers × 4 depots × 24 months. Calibrated two-stage "
                "forecaster (occurrence expected calibration error (ECE) ~0.03; near-nominal P10-P90 coverage); drift-augmented "
                "dynamic bands right-size procurement. Supplier-risk area under the curve (AUC) ~0.77, survival "
                "C-index ~0.79.",
           prefix="Forecasting embodiment: ")

    # ---------- 8 Abstract ----------
    h1(doc, "8.  Abstract")
    body(doc, "A computer-implemented system and method for detecting behavioral anomalies by "
              "representing entities as time-evolving points in a unified semantic embedding space "
              "derived from natural-language serialization of behavioral metrics with statistical-"
              "tier annotations. Anomalies are detected by the DIRECTION of behavioral drift "
              "projected onto named concept vectors mapped to a structured threat taxonomy, rather "
              "than by the magnitude of any individual metric. A three-tier architecture integrates "
              "traditional statistical methods, embedding-drift change-point detection, and "
              "digital-entity models with concept-direction alignment, zone divergence, regime "
              "detection, co-drifting cohort detection, kill chain reconstruction, and multi-phase "
              "composite scoring into a unified verdict. Bi-temporal embedding history enables "
              "leak-free training. A semantic gap bypass extracts high-discrimination features "
              "before embedding. Relationship and session entities capture interaction-level "
              "behavioral changes. The architecture is domain-portable across cybersecurity, "
              "supply-chain forecasting, healthcare fraud, and other domains by varying only "
              "serialization templates, zone definitions, and concept libraries.")

    # ---------- 9 Claims ----------
    h1(doc, "9.  Draft Claims (illustrative — for attorney refinement)")
    claims = [
        ("Claim 1 (method).", "A computer-implemented method for detecting behavioral anomalies "
         "in a monitored computing environment, the method comprising, for each entity of a "
         "plurality of entities spanning two or more entity types: "
         "(a) receiving, by a processor, observational data from one or more data sources and "
         "aggregating observational metrics over an observation period into a plurality of "
         "behavioral zones stored in a data store; "
         "(b) generating, for each zone, a natural-language description of the zone's state, "
         "including statistical-tier annotations (normal, elevated, critical) relative to "
         "population statistics and per-entity baselines; "
         "(c) encoding each description, using a text-embedding model executed on the processor, "
         "into a vector stored in a vector database with approximate-nearest-neighbor indexing, "
         "said vector occupying a semantic vector space common to all said entity types; "
         "(d) composing the zone vectors into an entity vector via normalized weighting, excluding "
         "zones designated as static; "
         "(e) storing each entity vector with bi-temporal metadata comprising a valid-time axis "
         "and a knowledge-time axis, and repeating (a)-(d) over successive periods to form a time "
         "series of entity vectors; "
         "(f) computing, by the processor, a drift vector from at least two entity vectors, said "
         "drift vector being a full-dimensional direction vector preserving directional information "
         "in the embedding space; "
         "(g) computing a cosine similarity between the drift vector and each of a plurality of "
         "pre-embedded natural-language concept vectors in said space, filtered to zone-relevant "
         "concepts only, wherein each concept vector corresponds to a named threat pattern or "
         "behavioral concept; and "
         "(h) producing a detection output stored in the data store, identifying a named concept "
         "based on said similarity, indicating a direction of behavioral change, and mapping the "
         "identified concept to a structured threat taxonomy, thereby providing directional "
         "detection that simultaneously identifies the type of behavioral anomaly and the "
         "behavioral zone driving the change without requiring a separate classification step."),
        ("Claim 2.", "The method of claim 1, wherein the entities span two or more operational "
         "domains and the semantic vector space is common across said domains."),
        ("Claim 3.", "The method of claim 1, wherein the natural-language description applies three "
         "statistical tiers — normal (|z| < 1.5), elevated (1.5 ≤ |z| < 2.5), and critical "
         "(|z| ≥ 2.5) — to annotate behavioral features relative to population statistics, "
         "per-entity historical baselines, and recent trend direction before embedding, thereby "
         "ensuring the embedding captures statistical significance."),
        ("Claim 4.", "The method of claim 1, further comprising accumulating drift magnitude over "
         "successive periods with a cumulative-sum procedure and signaling a change point when the "
         "accumulated value exceeds a threshold for a minimum run length; and further comprising "
         "computing velocity as a full-dimensional vector (V_t − V_{t−1}) preserving direction, "
         "acceleration as the change in velocity, stability as the inverse of velocity variance, "
         "and trend consistency as the mean cosine similarity of consecutive direction vectors."),
        ("Claim 5.", "The method of claim 1, wherein composing comprises a context-adaptive "
         "attention weighting over zones determined by linear normalization of zone-vector "
         "magnitudes biased by a context weight per zone, preserving intended weight "
         "differentiation across zones."),
        ("Claim 6.", "The method of claim 1, further comprising standardizing entity features "
         "against statistics of a cohort of comparable entities and scoring relative to said "
         "cohort; and further detecting co-drifting cohorts by clustering entities whose drift "
         "direction vectors exhibit cosine similarity exceeding a predetermined similarity "
         "threshold, and computing cohort coherence as the mean pairwise cosine similarity "
         "within each cluster."),
        ("Claim 7.", "The method of claim 1, further comprising emitting a human-readable "
         "explanation of the identified concept and its taxonomy mapping; reconstructing "
         "multi-step kill chains by temporally correlating alerts across entities within a "
         "predetermined time window; ordering correlated events by tactic sequence; and generating "
         "a natural-language attack narrative from the reconstructed chain."),
        ("Claim 8.", "The method of claim 1, further comprising forming a relationship vector as a "
         "normalized element-wise product of two entity vectors and detecting drift of the "
         "relationship vector; and further modeling sessions as relationship entity types that "
         "bridge user and device embeddings via temporal co-occurrence."),
        ("Claim 9.", "The method of claim 1, further comprising fusing a plurality of detection "
         "phases — two or more of signal strength, breadth, sustained deviation, context "
         "divergence, and novelty persistence — into a single ranked score."),
        ("Claim 10 (forecasting).", "A computer-implemented method for generating a drift-augmented "
         "probabilistic forecast and controlling resource allocation in a supply-chain management "
         "system, the method performed by a computing system comprising a processor, a memory, "
         "and a data store, the method comprising: "
         "(a) producing, by the processor, a baseline forecast having quantile bands from historical "
         "entity data stored in the data store; "
         "(b) aggregating, by the processor, observational metrics over successive periods into "
         "behavioral zones and encoding each zone as a vector stored in a vector database with "
         "approximate-nearest-neighbor indexing in a semantic vector space using a text-embedding "
         "model executed on the processor; "
         "(c) composing, by the processor, zone vectors into entity vectors and storing the entity "
         "vectors with bi-temporal metadata in the data store to form a time series; "
         "(d) computing, by the processor, a drift vector from at least two entity vectors and "
         "projecting the drift vector onto pre-embedded concept vectors to determine a direction "
         "of behavioral drift; "
         "(e) classifying, by the processor, each entity's behavioral state via a regime detection "
         "state machine into one of: stable, drifting, accelerating, or regime shift; "
         "(f) adjusting, by the processor, the baseline forecast by (i) a bounded scalar adjustment "
         "to a central quantile based on the direction and magnitude of drift, and (ii) a per-concept "
         "band multiplier that widens or narrows quantile band spread independently of the central "
         "quantile, thereby producing asymmetric, drift-aware forecast bands; and "
         "(g) transmitting the adjusted forecast to an automated procurement or inventory "
         "management system to control resource allocation quantities based on the drift-aware "
         "forecast bands."),
        ("Claim 11 (system).", "A system for detecting behavioral anomalies in a monitored "
         "computing environment, comprising: a processor; a memory storing instructions that, "
         "when executed by the processor, cause the processor to: receive observational data from "
         "one or more data sources and aggregate observational metrics over an observation period "
         "into a plurality of behavioral zones for each entity of a plurality of entities spanning "
         "two or more entity types; generate a natural-language description of each zone's state "
         "including statistical-tier annotations relative to population statistics and per-entity "
         "baselines; encode each description into a vector in a semantic vector space common to all "
         "entity types using a text-embedding model, and store said vector in a vector database with "
         "approximate-nearest-neighbor indexing; compose zone vectors into entity vectors via "
         "normalized weighting, excluding zones designated as static; store entity vectors with "
         "bi-temporal metadata comprising a valid-time axis and a knowledge-time axis; compute a "
         "drift vector from successive entity vectors; compute cosine similarity between the drift "
         "vector and pre-embedded concept vectors filtered to zone-relevant concepts, wherein each "
         "concept vector corresponds to a named threat pattern; and produce a detection output "
         "identifying a named concept indicating a direction of behavioral change and mapping the "
         "identified concept to a structured threat taxonomy, thereby providing directional "
         "detection that simultaneously identifies the type of behavioral anomaly and the "
         "behavioral zone driving the change."),
        ("Claim 12 (medium).", "A non-transitory computer-readable medium storing instructions "
         "that, when executed by a processor, cause the processor to perform a method comprising: "
         "aggregating observational metrics from a monitored computing environment over an "
         "observation period into behavioral zones for each entity of a plurality of entities "
         "spanning two or more entity types; generating natural-language descriptions of zone "
         "states with statistical-tier annotations relative to population statistics and per-entity "
         "baselines; encoding descriptions into vectors in a unified semantic vector space "
         "common to all said entity types using a text-embedding model and storing said vectors "
         "in a vector database with "
         "approximate-nearest-neighbor indexing; composing zone vectors into entity vectors via "
         "normalized weighting, excluding static zones; storing entity vectors with bi-temporal "
         "metadata comprising a valid-time axis and a knowledge-time axis; computing drift vectors "
         "from successive entity vectors; computing cosine similarity between drift vectors and "
         "pre-embedded concept vectors filtered to zone-relevant concepts, wherein each concept "
         "vector corresponds to a named threat pattern; and producing detection outputs identifying "
         "named concepts indicating directions of behavioral change and mapping identified concepts "
         "to a structured threat taxonomy, thereby providing directional detection that "
         "simultaneously identifies the type of behavioral anomaly and the behavioral zone "
         "driving the change."),
    ]
    h2(doc, "Independent and core dependent claims (1-12)")
    for label, txt in claims:
        bullet(doc, txt, prefix=label + " ")

    h2(doc, "Extended dependent claims (13-15)")
    ext_claims = [
        ("Claim 13.", "The method of claim 1, further comprising executing detection in three "
         "parallel tiers: a first tier comprising statistical and machine-learning algorithms "
         "operating on scalar behavioral features to produce magnitude-based anomaly scores; a "
         "second tier comprising change-point detection on cosine distances between successive "
         "behavioral embeddings to identify entities whose behavioral meaning is changing; and a "
         "third tier comprising a digital-entity model with concept-direction alignment, zone "
         "divergence, regime detection, behavioral progression, and novelty persistence; and "
         "fusing outputs of the three tiers via a geometric-mean composite in log-space blended "
         "with a corroboration breadth count into a unified detection verdict."),
        ("Claim 14.", "The method of claim 7, further comprising maintaining a directed graph "
         "data structure storing entity-hop linkages, computing a chain confidence score as a "
         "weighted combination of tactic coverage fraction, temporal coherence inversely "
         "proportional to inter-alert time gaps, and entity linkage count; marking chains stored "
         "in the directed graph data structure as stale when no new events occur within a staleness window equal to a "
         "predetermined multiple of the correlation window; and transmitting the reconstructed "
         "chain to a security orchestration system to trigger automated containment actions."),
        ("Claim 15.", "The method of claim 6, further comprising computing pairwise cosine "
         "similarity between drift direction vectors of entities in the semantic vector space; "
         "clustering entities whose drift direction similarity exceeds a predetermined similarity "
         "threshold using single-linkage clustering; computing cohort coherence as the mean "
         "pairwise cosine similarity within each cluster; and producing a coordinated-threat "
         "detection output identifying clusters of entities exhibiting synchronized directional "
         "behavioral changes indicative of coordinated campaigns, and transmitting the output to "
         "a security orchestration system."),
    ]
    for label, txt in ext_claims:
        bullet(doc, txt, prefix=label + " ")

    h2(doc, "Dependent claims (16-25)")
    new_claims = [
        ("Claim 16.", "The method of claim 1, further comprising detecting inter-zone divergence "
         "by comparing drift magnitudes across behavioral zones, wherein a pattern of one zone "
         "exhibiting drift above a predetermined zone-drift threshold while another zone remains "
         "below said threshold serves as an attack-type signature distinguishing exfiltration "
         "from credential compromise."),
        ("Claim 17.", "The method of claim 4, further comprising a regime detection state machine "
         "that classifies each entity's behavioral state into one of: stable, drifting, "
         "accelerating, or regime shift, evaluated in priority order based on regime shift count, "
         "acceleration with trend consistency, and total drift with stability thresholds."),
        ("Claim 18.", "The method of claim 1, wherein the bi-temporal metadata of step (e) is "
         "maintained using a three-step slowly-changing-dimension type 2 (SCD2) upsert protocol "
         "comprising: (i) superseding the current open row by setting a knowledge-to timestamp on "
         "the existing open-ended record; (ii) inserting a closed-valid historical row with the "
         "prior embedding vectors and a valid-to timestamp preserving the prior behavioral state; "
         "and (iii) inserting a new open row with the new embedding vectors, a valid-from timestamp, "
         "and a null valid-to timestamp; and supporting point-in-time queries by filtering on the "
         "valid-time axis, thereby enabling retroactive correction of embeddings for leak-free "
         "training and regulatory audit trails."),
        ("Claim 19.", "The method of claim 1, further comprising computing domain name entropy "
         "using Shannon's formula to detect algorithmically generated domains (DGA), and annotating "
         "domain-related features as NOVEL prior to embedding when entropy exceeds a predetermined "
         "entropy threshold."),
        ("Claim 20.", "The method of claim 1, further comprising role-appropriateness checking "
         "wherein file access paths are compared against role-expected and sensitive directory "
         "mappings, with access outside expected patterns tagged as NOVEL or ATYPICAL features "
         "before embedding, and computing novelty persistence metrics tracking how long novel "
         "behaviors persist across observation windows."),
        ("Claim 21.", "The method of claim 1, wherein certain high-discrimination features "
         "including novel IP addresses, novel process executions, and DGA domain entropy scores "
         "are extracted directly from behavioral text prior to embedding, bypassing the embedding "
         "layer to preserve detection sensitivity for features whose novelty signal would be "
         "diluted by semantic compression."),
        ("Claim 22.", "The method of claim 1, further comprising detecting behavioral progression "
         "by computing Kendall tau rank correlation across sequential behavioral measurements over "
         "a minimum of eight observation periods, identifying monotonic escalation patterns that "
         "individual-point anomaly detection methods cannot capture."),
        ("Claim 23.", "The method of claim 9, wherein the fusing of detection phases comprises "
         "computing signal strength as the peak concept-alignment score, breadth as the number of "
         "zones exhibiting above-threshold drift, sustained deviation as the CUSUM accumulated "
         "value, context divergence as the entity's z-score relative to its peer cohort, and "
         "novelty persistence as the duration of novel behavioral features, and combining said "
         "phases via weighted summation into a single ranked score."),
        ("Claim 24.", "The method of claim 8, wherein sessions are modeled as relationship entity "
         "types bridging user and device embeddings via temporal co-occurrence within a bounded "
         "time window, creating a relationship embedding that captures the interaction pattern "
         "between two entities and enables drift detection on the interaction itself."),
        ("Claim 25.", "The method of claim 1, wherein the entities span two or more operational "
         "domains selected from: cybersecurity behavioral analytics, supply-chain demand "
         "forecasting, healthcare fraud detection, anti-money laundering, prescriptive "
         "maintenance, and counter-intelligence, and the same pipeline architecture serves all "
         "domains with only serialization templates, zone definitions, and concept libraries "
         "varying per domain."),
    ]
    for label, txt in new_claims:
        bullet(doc, txt, prefix=label + " ")
    h2(doc, "Claims from implemented innovations (26-45)")
    impl_claims = [
        ("Claim 26.", "The method of claim 5, further comprising maintaining a plurality of "
         "distinct investigation context weight profiles, each profile assigning different "
         "zone attention weights to the behavioral zones; composing the zone vectors under each "
         "of the plurality of context weight profiles to produce a respective entity vector per "
         "context; computing concept-direction alignment for each entity vector; and selecting "
         "the investigation context whose entity vector yields the strongest concept-direction "
         "alignment as the primary detection context for that entity."),
        ("Claim 27.", "The method of claim 9, wherein the fusing comprises computing a geometric "
         "mean of rank-normalized detection method scores in log-space with calibrated per-method "
         "weights, and blending the geometric-mean depth score with a corroboration breadth score "
         "counting the number of independent detection methods flagging the entity."),
        ("Claim 28.", "The method of claim 3, wherein the natural-language description generated "
         "at each statistical tier employs progressively escalating threat-pattern language, with "
         "critical-tier annotations explicitly naming consistent attack patterns to ensure the "
         "text-embedding model produces semantically distinct vectors for anomalous versus normal "
         "behavior."),
        ("Claim 29.", "The method of claim 1, wherein step (g) further comprises filtering the "
         "plurality of concept vectors to a zone-relevant subset for each behavioral zone using "
         "a zone-to-concept mapping that restricts each zone's drift-concept alignment computation "
         "to semantically appropriate concepts and excludes cross-domain concepts, and selecting "
         "the behavioral zone exhibiting the strongest concept-direction alignment as the primary "
         "attack indicator for that entity."),
        ("Claim 30 (broader method).", "A computer-implemented method for detecting behavioral "
         "anomalies in a monitored computing environment, the method comprising, for each entity "
         "of a plurality of entities spanning two or more entity types: "
         "(a) receiving, by a processor, observational data from one or more data sources and "
         "aggregating observational metrics over an observation period into a plurality of "
         "behavioral zones stored in a data store; "
         "(b) generating, for each zone, a natural-language description of the zone's behavioral "
         "state from raw metric values; "
         "(c) encoding each description, using a text-embedding model executed on the processor, "
         "into a vector stored in a vector database, said vector occupying a semantic vector space "
         "common to all said entity types; "
         "(d) composing the zone vectors into an entity vector via normalized weighting; "
         "(e) storing each entity vector with temporal metadata and repeating (a)-(d) over "
         "successive periods to form a time series of entity vectors; "
         "(f) computing, by the processor, a drift vector from at least two entity vectors; "
         "(g) computing a similarity metric between the drift vector and each of a plurality of "
         "pre-embedded natural-language concept vectors in said space, wherein each concept vector "
         "corresponds to a named behavioral concept; and "
         "(h) producing a detection output stored in the data store, identifying a named concept "
         "based on said similarity metric, indicating a direction of behavioral change."),
        ("Claim 31.", "The method of claim 1, further comprising caching embedding vectors in a "
         "content-hash-indexed sharded storage structure, checking the cache before invoking the "
         "text-embedding model, and preloading a consolidated cache into memory at system startup "
         "to minimize redundant embedding computation."),
        ("Claim 32 (zoneless method).", "A computer-implemented method for detecting behavioral "
         "anomalies in a monitored computing environment, the method comprising, for each entity "
         "of a plurality of entities: "
         "(a) receiving, by a processor, observational data from one or more data sources "
         "describing the entity's behavior over an observation period and generating a "
         "natural-language description of the entity's behavioral state; "
         "(b) encoding the description, using a text-embedding model executed on the processor, "
         "into an entity vector stored in a vector database with approximate-nearest-neighbor "
         "indexing, said entity vector occupying a semantic vector space; "
         "(c) repeating (a)-(b) over successive periods to form a time series of entity vectors "
         "stored with temporal metadata in a data store; "
         "(d) computing, by the processor, a drift vector from at least two entity vectors in "
         "said time series; "
         "(e) computing a similarity metric between the drift vector and each of a plurality of "
         "pre-embedded natural-language concept vectors in said space, wherein each concept vector "
         "corresponds to a named behavioral concept; and "
         "(f) producing a detection output stored in the data store, identifying a named concept "
         "based on said similarity metric, indicating a direction of behavioral change from "
         "the entity vector directly."),
        ("Claim 33.", "The method of claim 30, further comprising adding statistical-tier "
         "annotations to the natural-language descriptions of step (b) prior to encoding."),
        ("Claim 34.", "The method of claim 30, further comprising filtering the plurality of "
         "concept vectors to a zone-relevant subset for each behavioral zone using a zone-to-concept "
         "mapping that restricts alignment computation to semantically appropriate concepts."),
        ("Claim 35.", "The method of claim 32, further comprising aggregating the observational "
         "data into a plurality of behavioral zones and generating a zone-level natural-language "
         "description for each zone prior to encoding."),
        ("Claim 36.", "The method of claim 32, further comprising accumulating drift magnitude "
         "over successive periods with a cumulative-sum procedure and signaling a change point "
         "when the accumulated value exceeds a threshold for a minimum run length."),
        ("Claim 37.", "The method of claim 10, further comprising computing a bounded scalar "
         "adjustment to a central quantile using a clamped product of concept direction, concept "
         "alignment strength, and drift magnitude, wherein the clamp bounds are predetermined "
         "safety gates preventing drift from dominating the baseline forecast."),
        ("Claim 38 (classifier method).", "A computer-implemented method for detecting behavioral "
         "anomalies in a monitored computing environment, the method comprising, for each entity "
         "of a plurality of entities: "
         "(a) receiving, by a processor, observational data from one or more data sources and "
         "generating a natural-language description of the entity's behavioral state; "
         "(b) encoding the description, using a text-embedding model executed on the processor, "
         "into an entity vector stored in a vector database with approximate-nearest-neighbor "
         "indexing, said entity vector occupying a semantic vector space; "
         "(c) repeating (a)-(b) over successive periods to form a time series of entity vectors "
         "stored with temporal metadata in a data store; "
         "(d) computing, by the processor, a drift vector from at least two entity vectors in "
         "said time series; and "
         "(e) classifying, by the processor, the drift vector using a model trained to associate "
         "drift vectors with named behavioral concepts, said model producing a detection output "
         "stored in the data store identifying a named concept indicating a direction of behavioral "
         "change."),
        ("Claim 39.", "The system of claim 11, further comprising executing detection in three "
         "parallel tiers: a first tier comprising statistical and machine-learning algorithms "
         "operating on scalar behavioral features, a second tier comprising change-point detection "
         "on cosine distances between successive behavioral embeddings, and a third tier comprising "
         "a digital-entity model with concept-direction alignment and zone divergence."),
        ("Claim 40.", "The system of claim 11, further comprising detecting co-drifting cohorts by "
         "clustering entities whose drift direction vectors exhibit cosine similarity exceeding a "
         "predetermined similarity threshold and transmitting identified cohorts to a security "
         "orchestration system."),
        ("Claim 41.", "The non-transitory computer-readable medium of claim 12, wherein the method "
         "further comprises accumulating drift magnitude over successive periods with a "
         "cumulative-sum procedure and signaling a change point when the accumulated value exceeds "
         "a threshold for a minimum run length."),
        ("Claim 42.", "The non-transitory computer-readable medium of claim 12, wherein the method "
         "further comprises detecting inter-zone divergence by comparing drift magnitudes across "
         "behavioral zones, wherein a pattern of zone-specific drift serves as an attack-type "
         "signature."),
        ("Claim 43.", "The method of claim 38, wherein the model is a neural network trained on "
         "labeled pairs of drift vectors and named behavioral concepts from a structured threat "
         "taxonomy."),
        ("Claim 44.", "The method of claim 38, further comprising computing a confidence score for "
         "the classification output and suppressing detection outputs below a predetermined "
         "confidence threshold."),
        ("Claim 45.", "The method of claim 10, further comprising adjusting the per-concept band "
         "multiplier of step (f)(ii) independently for upper and lower quantile bounds, thereby "
         "producing asymmetric forecast bands that widen in one direction while narrowing in "
         "another based on the direction of behavioral drift."),
    ]
    for label, txt in impl_claims:
        bullet(doc, txt, prefix=label + " ")
    body(doc, "Further dependent variations for counsel: embedding dimensionality/model class; "
              "domain-specific zone sets; forecast safety gates; two-stage 'hurdle' demand "
              "baseline; alert deduplication within sliding time windows.",
         italic=True, color=GRAY)

    # ---------- 10 Prior art ----------
    h1(doc, "10.  Distinction Over Prior Art (source-verified)")
    body(doc, "A structured, adversarially-verified review (2024–2026) surveyed both commercial "
              "products and academic literature. No system was found to combine concept-direction "
              "detection, a unified cross-domain embedding space, and the full integrated stack.")
    h2(doc, "10.1  Commercial UEBA / SIEM (Security Information and Event Management) products (magnitude-based)")
    bullet(doc, "Statistical ML (session risk scoring) + TF-IDF peer grouping. "
                "Generative AI added as analyst copilot (2024), not as detector. "
                "No embedding-based drift detection; no concept-direction alignment.",
           prefix="Exabeam Advanced Analytics: ")
    bullet(doc, "Peer-group baselining + rule engine + statistical anomaly scoring. "
                "MITRE ATT&CK used for post-hoc alert enrichment, not as detection signal. "
                "No behavioral embeddings; no zone decomposition.",
           prefix="Securonix UEBA: ")
    bullet(doc, "Log-based anomaly scoring, fusion rules, identity timeline. "
                "Generative AI layer for investigation summaries. "
                "No semantic embedding space; magnitude-threshold detection only.",
           prefix="Microsoft Sentinel UEBA: ")
    bullet(doc, "Rule + statistical engines with risk-adaptive scoring. "
                "No embedding representations; fixed-schema feature engineering.",
           prefix="Splunk UBA / IBM QRadar UEBA: ")
    bullet(doc, "Event-driven correlation + behavioral profiling. "
                "Network-centric anomaly scoring. No text-embedding pipeline.",
           prefix="Palo Alto Cortex XDR / CrowdStrike: ")
    bullet(doc, "Self-learning Bayesian probabilistic models operating across network, email, "
                "cloud, and endpoint telemetry. Detects anomalies by deviation probability "
                "(magnitude), not by direction of drift. No concept vector projection; no "
                "zone-specific filtering; no named-concept explanations.",
           prefix="Darktrace Antigena: ")
    bullet(doc, "ML/AI-driven UEBA with behavioral risk scoring and identity analytics. "
                "Statistical anomaly detection on tabular features. No behavioral embeddings; "
                "no directional drift detection; no concept-direction alignment.",
           prefix="Gurucul UEBA: ")
    body(doc, "All commercial UEBA/SIEM/XDR products — including Darktrace's 'self-learning AI' "
              "— operate on magnitude-based statistical scoring, Bayesian probability deviation, "
              "or rule-based correlation. None represent behavior as language, embed into a semantic "
              "vector space, or detect by the direction of drift onto named concept vectors.",
         italic=True, color=GRAY)
    h2(doc, "10.2  Commercial demand/SCM forecasting")
    bullet(doc, "Statistical + gradient-boosted models over tabular features with fixed-shape "
                "uncertainty bands. Generative AI as conversational analysis layer.",
           prefix="SAP IBP / o9 Solutions / Kinaxis: ")
    body(doc, "No system uses behavioral embeddings to shape forecast uncertainty bands by "
              "named supply-chain concepts.", italic=True, color=GRAY)
    h2(doc, "10.3  Academic prior art (single building blocks)")
    bullet(doc, "APT-LLM, LogBERT: LLM-embedding of log sequences for single entity type, "
                "single domain. Magnitude-based anomaly scoring (reconstruction error, probability).",
           prefix="Log-embedding anomaly detection: ")
    bullet(doc, "TabLLM: serialize tabular data to prose then embed. Single-table, no zones, "
                "no drift tracking, no concept projection.",
           prefix="Tabular serialization: ")
    bullet(doc, "TAD-Bench: benchmarks IForest/LOF on LLM embedding vectors. Detects by "
                "MAGNITUDE (distance from centroid), not by DIRECTION onto named concepts.",
           prefix="Embedding anomaly benchmarks: ")
    bullet(doc, "ZEDD: embedding drift via cosine distance over time. Scalar distance only — "
                "no directional decomposition, no concept projection, no zone structure.",
           prefix="Embedding drift: ")
    bullet(doc, "Kim et al. (ICML 2018): Concept Activation Vectors (CAVs/TCAV) project neural "
                "network activations onto concept directions for interpretability of classifiers. "
                "DISTINGUISHED: CAVs explain an existing classifier's internal representations; "
                "in this invention, concept-direction projection IS the detector itself, not a "
                "post-hoc explanation layer, and operates on behavioral embeddings (not convolutional neural network (CNN) "
                "activations) in a temporal drift context.",
           prefix="Concept Activation Vectors: ")
    bullet(doc, "Lu et al. survey (2019): 'concept drift' in ML refers to distributional shift "
                "in data or model performance over time. DISTINGUISHED: this invention detects "
                "'concept-direction drift' — the direction of change in a behavioral embedding "
                "space projected onto named semantic concepts — a fundamentally different signal "
                "from statistical distributional shift.",
           prefix="Concept drift (ML term): ")
    bullet(doc, "Chronos (Amazon, 2024), TimeGPT (Nixtla, 2024), Lag-Llama (2024), MOMENT "
                "(CMU, 2024): foundation models for time-series forecasting that process "
                "numeric series directly (via quantization, patching, or lagged feature "
                "extraction). DISTINGUISHED: these models treat time series as numeric "
                "sequences, not as natural-language behavioral descriptions. "
                "None embed behavior into a semantic space, compute drift direction, or "
                "project onto named concept vectors.",
           prefix="Foundation time-series models: ")
    bullet(doc, "WhyLabs, Arize (2023–2024): commercial ML observability platforms that monitor "
                "embedding drift via cosine distance as a model-quality signal. DISTINGUISHED: "
                "these track scalar drift magnitude for model monitoring, not directional drift "
                "projected onto named behavioral concepts for anomaly detection.",
           prefix="Embedding drift monitoring: ")
    body(doc, "TEACHING AWAY: All surveyed commercial systems use MITRE ATT&CK exclusively as a "
              "post-hoc enrichment layer applied AFTER detection to label alerts for analyst "
              "consumption. This established practice teaches away from the present invention's "
              "approach of using pre-embedded MITRE ATT&CK concept vectors as the PRIMARY "
              "detection signal — projecting drift direction onto concept vectors to determine "
              "WHAT an entity is becoming, rather than applying taxonomy labels after an anomaly "
              "has already been detected by other means.", italic=True, color=BLUE)
    callout(doc,
            "DEFENSIBLE NOVELTY: A comprehensive competitive review (2024–2026) confirmed NO "
            "commercial product — neither UEBA/SIEM (Exabeam, Securonix, Microsoft Sentinel, "
            "Splunk, IBM QRadar, CrowdStrike, Palo Alto, Darktrace, Gurucul) nor SCM forecasting "
            "(SAP IBP, o9 Solutions, Kinaxis) — and NO published academic system "
            "(APT-LLM, LogBERT, TabLLM, TAD-Bench, ZEDD, TCAV) combines:\n"
            "(i) a unified semantic embedding space across entity types AND operational domains,\n"
            "(ii) detection by the DIRECTION of behavioral drift onto named concept vectors "
            "mapped to MITRE ATT&CK as the primary detection signal,\n"
            "(iii) a three-tier integrated architecture with zone divergence, co-drifting cohort "
            "detection, behavioral progression, and kill chain reconstruction, and\n"
            "(iv) semantic gap bypass with novelty persistence scoring.\n\n"
            "ALL surveyed competitors use magnitude-based statistical scoring, rule engines, or "
            "fixed-schema ML — fundamentally different from the concept-direction paradigm.")
    h2(doc, "10.4  Cited patents")
    bullet(doc, "US11258814B2 — behavioral analytics patent (likely Exabeam or similar vendor); "
                "covers statistical baseline comparison for user behavior. Distinguished: no "
                "embedding space, no concept-direction detection, no zone decomposition.",
           prefix="US11258814B2: ")
    bullet(doc, "US12250239B2 — recently granted; relates to ML-based security analytics. "
                "Distinguished: operates on tabular features with fixed-schema ML, not "
                "behavioral embeddings with directional drift detection.",
           prefix="US12250239B2: ")
    bullet(doc, "US20230336571A1 (CrowdStrike, 2023) — generating embeddings for security event "
                "data. Distinguished: embeds raw security events for classification, not behavioral "
                "zones with statistical-tier annotations; no drift-direction detection, no concept "
                "projection, no zone decomposition.",
           prefix="US20230336571A1: ")
    bullet(doc, "US11790085B2 (Abnormal Security, 2023) — behavioral analysis using entity "
                "embeddings for email security. Distinguished: as disclosed in the patent, "
                "single-domain (email), single entity type; uses embedding similarity for "
                "classification, not drift-direction projection onto named concept vectors.",
           prefix="US11790085B2: ")
    body(doc, "Counsel action: confirm exact scope of cited patents via formal patentability/FTO "
              "search; verify no claims overlap with the concept-direction detection paradigm.",
         italic=True, color=GRAY)

    # ---------- 11 Pre-filing ----------
    h1(doc, "11.  Notes / To-Do Before Filing")
    for txt in [
        "Complete Section 0 metadata (inventors, assignee, conception/RTP dates, public-disclosure "
        "timeline).",
        "Re-render FIG. 1-8 as formal patent drawings.",
        "Counsel: patentability + FTO search; confirm Section 101 framing (tie claims to a specific "
        "technical improvement, not an abstract idea).",
        "Decide claim scope and which embodiments to emphasize (security / forecasting / both).",
        "Confirm whether any enabling public disclosure has occurred (statutory-bar clock).",
    ]:
        bullet(doc, txt)
    body(doc, "Prepared as an inventor's technical disclosure to accompany attorney drafting of a "
              "U.S. provisional patent application. Not legal advice.", italic=True, color=GRAY)

    out = os.path.join(HERE, "Provisional_Patent_Invention_Disclosure.docx")
    try:
        doc.save(out)
    except PermissionError:
        import pathlib
        p = pathlib.Path(out)
        for i in range(2, 20):
            alt = p.with_stem(f"{p.stem}_v{i}")
            try:
                doc.save(str(alt))
                out = str(alt)
                break
            except PermissionError:
                continue
        else:
            raise
    print("Saved:", out)


if __name__ == "__main__":
    build()
