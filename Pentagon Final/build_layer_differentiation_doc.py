# -*- coding: utf-8 -*-
"""Technical write-up: the gap in current tools at each layer, and how we cover it.
Gap-first, hardcore-technical (mechanism, formal property, boundary). Not sales. No em/en dashes."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0E,0x28,0x41); STEEL=RGBColor(0x1F,0x5E,0x7A); AMBER=RGBColor(0xB9,0x6A,0x12)
RISK=RGBColor(0xA5,0x34,0x2F); MUTED=RGBColor(0x5A,0x66,0x75); INK=RGBColor(0x14,0x1D,0x2B)
WHITE=RGBColor(0xFF,0xFF,0xFF)
BODYF="Aptos"; HEADF="Aptos Display"; MONOF="Consolas"

doc=Document()
st=doc.styles['Normal']; st.font.name=BODYF; st.font.size=Pt(10.5); st.font.color.rgb=INK
st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.13
for section in doc.sections:
    section.top_margin=Inches(0.8); section.bottom_margin=Inches(0.8)
    section.left_margin=Inches(0.9); section.right_margin=Inches(0.9)

def shade(cell,hex_):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hex_); tcPr.append(sh)
def run(p,text,size=10.5,color=INK,bold=False,italic=False,font=BODYF):
    r=p.add_run(text); r.font.size=Pt(size); r.font.color.rgb=color; r.font.bold=bold
    r.font.italic=italic; r.font.name=font; return r
def para(text=None,size=10.5,color=INK,bold=False,italic=False,after=4,before=0,align=None,font=BODYF,indent=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    if align is not None: p.alignment=align
    if indent is not None: p.paragraph_format.left_indent=Inches(indent)
    if text is not None: run(p,text,size,color,bold,italic,font)
    return p
def rule(color=NAVY,size=18):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); p.paragraph_format.space_before=Pt(0)
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),str(size)); bottom.set(qn('w:space'),'1')
    bottom.set(qn('w:color'),'%02X%02X%02X'%(color[0],color[1],color[2])); pbdr.append(bottom); pPr.append(pbdr)
    return p
def h1(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(15); p.paragraph_format.space_after=Pt(2)
    run(p,text,size=16.5,color=NAVY,bold=True,font=HEADF); rule(NAVY,14); return p
def h2(text,color=NAVY):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(9); p.paragraph_format.space_after=Pt(3)
    run(p,text,size=12,color=color,bold=True,font=HEADF); return p
def eyebrow(text,color=AMBER):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(7); p.paragraph_format.space_after=Pt(2)
    r=run(p,text.upper(),size=8.5,color=color,bold=True); return p
def gapb(lead,rest,lead_color=AMBER):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.28); p.paragraph_format.space_after=Pt(5)
    p.paragraph_format.line_spacing=1.1
    run(p,"- ",size=10,color=lead_color,bold=True)
    run(p,lead,size=10.5,color=lead_color,bold=True); run(p," "+rest,size=10.5,color=INK)
    return p
def step(n,lead,rest):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.28); p.paragraph_format.space_after=Pt(5)
    p.paragraph_format.line_spacing=1.1
    run(p,f"{n}   ",size=10,color=NAVY,bold=True,font=MONOF)
    run(p,lead+"  ",size=10.5,color=NAVY,bold=True); run(p,rest,size=10.5,color=INK)
    return p
def prop(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.0); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr')
    for side in ('left',):
        e=OxmlElement('w:'+side); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'18'); e.set(qn('w:space'),'6')
        e.set(qn('w:color'),'%02X%02X%02X'%(STEEL[0],STEEL[1],STEEL[2])); pbdr.append(e)
    pPr.append(pbdr); p.paragraph_format.left_indent=Inches(0.1)
    run(p,"NET PROPERTY.  ",size=10,color=STEEL,bold=True); run(p,text,size=10.5,color=INK,bold=True)
    return p
def boundary(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.28); p.paragraph_format.space_after=Pt(6)
    run(p,"Boundary / honesty.  ",size=9.5,color=RISK,bold=True,italic=True)
    run(p,text,size=9.5,color=MUTED,italic=True)
    return p

# ============================== TITLE ==============================
eyebrow("360 Closed-Loop Cyber Defense  |  Pentagon technical deep-dive  |  July 2026",MUTED)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
run(p,"The Gap at Each Layer, and How We Close It",size=24,color=NAVY,bold=True,font=HEADF)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
run(p,"A technical differentiation review of the five-layer platform",size=13,color=STEEL,italic=True,font=HEADF)
rule(NAVY,18)
para("CONFIDENTIAL / Internal / Competitive-Technical          5 layers  |  mechanism, formal property, and boundary per layer",
     size=8.5,color=MUTED,after=8)
para("This document is a mechanism-level review of the five-layer platform. For each layer it does three "
     "things: it states the technical gap that current tools leave open (named vendors, and why the "
     "limitation is structural, not incidental); it describes how we cover that gap (the pipeline and the "
     "formal or algorithmic property it discharges); and it states the boundary of the guarantee so a "
     "knowledgeable evaluator cannot turn a superlative into a counterexample. It is written for a technical "
     "audience and deliberately avoids market and adoption claims.",after=6)

# ============================== SUMMARY ==============================
h1("1.  Summary")
para("Two families of mechanism carry the platform, and one discipline governs every claim.",after=4)
para("Provable, not probable (Layers 1 and 5).",bold=True,color=STEEL,after=1)
para("Both layers discharge a proof over a bounded model. Layer 1 proves that no modeled attack technique "
     "evades the composed control configuration; Layer 5 proves that a vulnerable code path is unreachable. "
     "The value is the true negative (soundness), which sampling (breach-and-attack simulation) and heuristic "
     "call-graph reachability cannot provide.",after=4)
para("Meaning, not surface signal (Layers 2, 3, and 4).",bold=True,color=STEEL,after=1)
para("Each reads semantics rather than magnitude. Layer 2 embeds behavior and detects the direction and "
     "persistence of drift; Layer 3 classifies the request's intent; Layer 4 makes a declared persona the "
     "enforcement spec. In each case the detection is on meaning or intent, not on a threshold crossing.",after=4)
para("The discipline: every guarantee is stated with its boundary (the modeled config and technique corpus "
     "for L1; the verification boundary or TCB for L5; the cohort and concept basis for L2), probabilistic AI "
     "is kept outside any formal guarantee, and every number is taken from a primary source (the live "
     "database for L2, the current datasheet for L3), never from an unverified slide.",before=2,after=4)

# ============================== THREADS + DISCIPLINE ==============================
h1("2.  The two mechanism families and the discipline")
h2("2.1  Provable, not probable  (L1, L5)")
para("Formal methods are the through-line. A proof answers a question sampling cannot: it certifies a true "
     "negative. Layer 1 returns UNSAT (no evasion path exists in the model); Layer 5 returns UNSAT (the "
     "vulnerable sink is unreachable). Both results are soundness statements bounded to an explicit model, "
     "which is the honest and defensible form of the claim.",after=4)
h2("2.2  Meaning, not surface signal  (L2, L3, L4)")
para("The second family reads intent and behavior in a semantic representation rather than counting events. "
     "The twin embeds behavior as prose and reads drift direction in representation space; the API layer "
     "classifies the request's intent instead of its deviation from a learned normal; the agent gateway "
     "compiles a declared persona into the enforcement spec so the baseline is the specification itself.",after=4)
h2("2.3  The discipline that governs every claim",color=RISK)
para("1.  State the boundary of every proof. Soundness is meaningful only over a stated model; \"no false "
     "negatives\" must read \"over the modeled configuration and corpus\" (L1) or \"over the analysis "
     "boundary / TCB\" (L5). Unbounded, each invites a one-name counterexample.",indent=0.28,after=3,color=INK)
para("2.  Keep probabilistic AI outside any formal guarantee. Where an LLM or ML step produces a finding, "
     "the \"provable\" claim cannot cover it. In L1 the LLM is confined to configuration ingestion and "
     "human-readable remediation; the proof core is symbolic.",indent=0.28,after=3,color=INK)
para("3.  Take every number from a primary source. L2 metrics are computed live from the database; L3 scale "
     "figures must come from the current datasheet, not from the deck's uncorroborated numbers.",indent=0.28,after=4,color=INK)

# ============================== L1 ==============================
h1("3.  Layer 1: control-efficacy verification")
para("The gap: current tools verify reachability or sample attacks; none prove that the control "
     "configuration defeats the technique, and none can certify a true negative.",italic=True,color=MUTED,after=4)
eyebrow("The gap in current tools")
gapb("Reachability verifiers (Forward Networks, Batfish):","are SMT-based, but the model is the "
     "packet-forwarding graph (ACLs, routes, NAT). They answer whether host A can reach service B on port P, "
     "not whether the IPS signature plus EPP rule plus WAF policy actually stops a technique. They carry no "
     "model of prevention or detection semantics.")
gapb("Breach-and-attack simulation (SafeBreach, AttackIQ):","executes a curated, finite scenario library "
     "against live infrastructure. Coverage equals the library; a technique not in it is untested, and a "
     "blocked result is a single observation, not a proof over all paths.")
gapb("Scanners and CTEM (Tenable, XM Cyber, RedSeal):","enumerate CVEs and build best-effort attack graphs, "
     "then simulate. The graphs are heuristic; there is no soundness guarantee and no true-negative.")
gapb("The structural gap:","all of these are positive-evidence methods - they find a hit. None can certify "
     "a true negative (prove that no evasion path exists), which is exactly what \"our controls defeat this "
     "technique\" requires.")
eyebrow("How we cover it",NAVY)
step(1,"Config ingest and normalization.","NGFW, IPS, EPP, WAF, and IdP rule-sets are parsed and lifted into "
     "a single symbolic policy model - a normalized representation of allow / deny / detect semantics across the stack.")
step(2,"Technique encoding.","Each MITRE ATT&CK technique is modeled as a state transition with pre-conditions "
     "(attacker state required) and post-conditions (state achieved) over that environment.")
step(3,"Symbolic path search.","An SMT / symbolic solver searches for the existence of a technique chain - a "
     "sequence of transitions - that the composed control set fails to block. A satisfying assignment is an evasion path.")
step(4,"Verdict.","UNSAT means no evasion path exists in the model: a soundness certificate of control "
     "efficacy. SAT returns a concrete counter-example path and the specific control that fails.")
prop("a soundness certificate that no modeled technique chain evades the composed control configuration, "
     "derived from configuration alone (no agents, no live traffic).")
boundary("Soundness and completeness hold over the modeled configuration and the encoded technique corpus; "
         "you cannot prove absence of a technique you never modeled. Reachability verifiers are the nearest "
         "formal analog but model forwarding only. The LLM is not in the proof path (ingestion and "
         "remediation only), so \"provable\" remains valid.")

# ============================== L2 ==============================
h1("4.  Layer 2: the math behind LLMs, applied to behavior (the Digital Twin)")
para("In plain terms. Most security tools watch individual numbers - how many files someone downloaded, how "
     "many systems they logged into - and raise an alarm when one number gets too big. A careful insider or a "
     "patient nation-state intruder keeps every individual number inside the normal range, so those tools stay "
     "silent. This layer takes the same mathematics that gives large language models their power to read "
     "meaning and applies it to behavior: it turns each person's entire pattern of activity into one evolving "
     "signature and watches which way that signature is drifting. No single action looks wrong; the danger is "
     "visible only when the whole is read together, the way a sentence means something its individual words do "
     "not. That is what \"the whole exceeds the sum of the parts\" means here.",after=5)
para("The gap: current analytics threshold on magnitude; the direction and persistence of behavioral drift "
     "are invisible to them, and a slow, low-magnitude drift stays under every single-window threshold.",
     italic=True,color=MUTED,after=4)
eyebrow("The gap in current tools")
gapb("SIEM and UEBA (Splunk, Sentinel, Exabeam):","represent behavior as numeric aggregates and detect with "
     "a threshold or supervised score on magnitude. A metric can be normal-sized while the kind of activity "
     "has shifted - they measure how much, never in what direction.")
gapb("Network anomaly (Darktrace):","scores unsupervised deviation on network telemetry - still "
     "magnitude-of-deviation, and computed on flows and packets rather than a semantic model of the entity's role behavior.")
gapb("Global-baseline ML (Vectra, Gurucul):","compares each entity to a population threshold. A shift that is "
     "abnormal for one role but ordinary population-wide is missed.")
gapb("The structural gap:","none of these embed behavior itself as the representation - embeddings elsewhere "
     "are used for search and RAG, not as the detection substrate - so a slow patient drift and the "
     "direction of change both go unseen.")
eyebrow("How we cover it",NAVY)
step(1,"Zone serialization.","Each entity's weekly activity is reduced to 5 behavioral zones - identity, "
     "access_pattern, data_behavior, network_footprint, risk_posture - each rendered as prose that encodes the "
     "salient metrics together with their cohort z-scores and meaning. (This serialization is intentional: "
     "encoding z-scores and semantics into text before embedding is the design, not leakage.)")
step(2,"Per-zone embedding into one semantic space.","5 independent 1536-d vectors (OpenAI "
     "text-embedding-3-small) are composed by context-adaptive attention into one composite behavioral vector "
     "per weekly snapshot, in a single unified semantic space spanning every entity type. The embedding "
     "captures the meaning and context of the behavior, not its counts.")
step(3,"Trajectory and drift.","The composite vector traces a weekly trajectory in embedding space. Drift is "
     "the cosine distance week over week; feature-CUSUM and embedding-CUSUM accumulate small persistent drifts "
     "that single-week thresholds miss.")
step(4,"Direction - the concept alignment IS the detector.","The drift vector is projected by cosine onto 16 "
     "reference-concept embeddings (14 threat, 2 benign) to name which threat direction the behavior is moving "
     "toward - answering 'what is this entity becoming', not 'how big is the number'. This concept-direction "
     "alignment is the detector itself, not a post-hoc enrichment; the top matches map to MITRE ATT&CK (top-3).")
step(5,"Two detectors on one twin.","A cohort-relative cumulative composite score (discovery, breadth) and a "
     "known-bad threat-profile classifier (precision) that fires on measurable profiles - C2-beacon, DGA, "
     "LOTL process, cohort-rare access, recon fan-out, insider collection - each corroborated by raw-event evidence.")

h2("4.1  Why the whole exceeds the sum of the parts")
para("The mechanism is a mathematical model operating on top of the LLM embedding space. The LLM contributes "
     "the semantic representation - the meaning and context of behavior serialized to language; the "
     "mathematics (attention composition, drift velocity, cosine projection onto named concept directions, "
     "CUSUM change-point, and multi-phase fusion) is what converts that representation into a ranked verdict. "
     "Neither half detects alone: the embedding without the math is a descriptive vector, and the math "
     "without the embedding is another magnitude threshold.",after=4)
para("The composite is more than the sum of its signals. For a patient adversary each individual signal - "
     "velocity, zone divergence, cohort deviation, novelty - stays below its own alert threshold; only the "
     "fused multi-phase composite crosses the detection line. Risk is visible in combination and invisible in "
     "isolation. The clearest instance is novelty persistence lifting a slow-APT entity from an otherwise "
     "undetectable rank to a top-tier verdict that no embedding or statistic reached alone.",after=4)
para("Two design choices keep the semantic space quantitatively grounded rather than merely descriptive: "
     "z-score-tiered serialization writes each metric together with how abnormal it is for the entity's "
     "cohort, so the embedding encodes deviation and not just activity; and a semantic-gap bypass extracts "
     "high-discrimination features (novel IPs, DGA entropy, novel processes) directly, before embedding, so a "
     "rare but decisive event is not averaged away inside the vector.",after=4)
prop("risk is the direction and persistence of behavioral drift in a semantic space, where the whole "
     "(behavioral meaning) exceeds the sum of the parts - the slow APT below every single-metric threshold is "
     "caught by the fused profile and novelty persistence.")
para("Evidence (computed live from PostgreSQL, never hardcoded): classical baselines (Isolation Forest, "
     "One-Class SVM, LOF) catch 0 of 4 injected attacks; a z-score method catches 1 of 4 and is noisy; the "
     "cumulative composite catches 4 of 4 but cleanly separates only 2 of 4 at a 10.6% false-positive rate; "
     "the threat-profile classifier catches 4 of 4 at 0 false positives.",before=2,after=3,indent=0.28)
boundary("The acid test is the slow APT (USR-234): it stays inside both drift bands (feature- and "
         "embedding-CUSUM) and is caught only by the known-bad profile and novelty persistence. Any claim that "
         "\"drift catches all four\" is false. Guarantees are relative to the cohort and the 16-concept basis.")

# ============================== L3 ==============================
h1("5.  Layer 3: intent classification, enforced inline")
para("The gap: current WAAP learns a normal and enforces at a vendor edge; request intent and air-gapped "
     "enforcement both go unmet.",italic=True,color=MUTED,after=4)
eyebrow("The gap in current tools")
gapb("Anomaly-baseline WAAP (Salt, Imperva, Akamai):","must first converge a learned normal per API, then "
     "alert on deviation. Two consequences follow: a warm-up requirement, and an output that is a deviation "
     "score, not the request's intent.")
gapb("Edge / CDN-coupled enforcement (Akamai, Cloudflare, F5):","places the enforcement point in the "
     "vendor's cloud, so traffic must transit it - incompatible with air-gapped and ATO-bounded enclaves.")
gapb("Detect-only designs (Salt):","split detection from blocking and need a downstream WAF or gateway to enforce.")
gapb("Stateless per-request rules:","miss business-logic abuse that only appears across an API call sequence.")
eyebrow("How we cover it",NAVY)
step(1,"Dual-model scoring.","A cross-tenant global model and a per-endpoint local model score every transaction.")
step(2,"Intent determination.","The classifier determines the request's intent rather than its distance from "
     "a learned baseline, so there is no normal to converge first.")
step(3,"Inline enforcement.","Block, rate-limit, or deceive is executed in the native request path with no "
     "redirection to a vendor cloud - deployable inside air-gapped enclaves.")
step(4,"Sequence scoring and attestation.","Requests are judged within their session sequence, exposing "
     "business-logic abuse; agent and bot traffic is checked with a hardware Secure-Enclave attestation signal "
     "a cloud-hosted agent cannot virtualize.")
prop("intent is classified and enforced in the data path - no learned baseline to converge, no traffic leaving the enclave.")
boundary("Keep claims at the mechanism level (dual-model, inline, sequence scoring, enclave attestation). "
         "Number hygiene: verify every scale figure against the current datasheet, and do not cite the deck's "
         "uncorroborated figures (117M bot IPs, \"12+ telcos\").")

# ============================== L4 ==============================
h1("6.  Layer 4: the persona as the enforcement spec")
para("The gap: current tools hand-author agent scopes and learn a baseline; the specification itself is "
     "never the control.",italic=True,color=MUTED,after=4)
eyebrow("The gap in current tools")
gapb("Hand-authored scopes (Entra Agent ID, Zscaler registry, Noma):","write the least-privilege policy "
     "manually, per agent. The scope does not follow from the agent's declared job, so it drifts from intent "
     "and scales poorly.")
gapb("Learned-baseline monitors:","need a warm-up window to converge a normal before they can flag an "
     "off-spec action, leaving the agent unmonitored during onboarding.")
gapb("SASE and transport controls:","inspect the network, not the agent's tool-call semantics - they cannot "
     "see which MCP tool was invoked with which parameters.")
gapb("MCP supply-chain risk (rug-pull, tool-poisoning):","is caught in a periodic scan, not at the moment of the call.")
eyebrow("How we cover it",NAVY)
step(1,"Persona synthesis.","A plain-English job description is compiled into a least-privilege MCP tool-scope "
     "(the allowed tools and parameters). The allow-list is generated from declared intent, not hand-written.")
step(2,"Scoped virtual endpoint.","A per-persona virtual MCP endpoint proxies the agent; every tool call is "
     "checked against that scope.")
step(3,"Spec-based detection.","The first off-spec tool call is a deterministic violation - no warm-up, and "
     "no probabilistic score crossing a threshold.")
step(4,"Registry integrity.","A crypto-signed, version-pinned MCP registry is enforced inline, blocking "
     "rug-pull and tool-poisoning at call time rather than in a scan report.")
prop("the declared persona is the enforcement spec; a violation is the first tool call outside scope, decided "
     "deterministically with no learning window.")
boundary("The durable technical slivers are natural-language-to-scope compilation and spec-as-baseline "
         "(deterministic, zero warm-up). Do not claim the AI-gateway category is unique (Palo Alto Prisma AIRS "
         "and Zscaler AI Broker exist); the depth inside the gateway is the point. Have a per-call enforcement "
         "latency figure ready, since inline checking cost is the obvious objection.")

# ============================== L5 ==============================
h1("7.  Layer 5: provable un-reachability and unified containment")
para("The gap: current tools estimate reachability heuristically and split host containment from agent "
     "containment; neither yields a proof nor one policy.",italic=True,color=MUTED,after=4)
eyebrow("The gap in current tools")
gapb("SCA reachability (Endor, Snyk, Contrast):","uses heuristic or observed call graphs and markets fewer "
     "false positives. It cannot claim zero false negatives, so it cannot justify leaving a flagged CVE unpatched.")
gapb("Virtual patching (Trend Micro, Fortinet):","applies a signature block at the perimeter - not a proof "
     "that the vulnerable code path is actually closed.")
gapb("Microsegmentation (Illumio, Guardicore):","contains hosts at the network but cannot revoke a token or "
     "tool permission. An agent operating inside a trusted host never crosses a segment, so segmentation cannot reach it.")
gapb("Agent-identity tools:","revoke tokens but enforce nothing at the network - two separate control planes "
     "for a single incident.")
eyebrow("How we cover it",NAVY)
step(1,"Sound reachability.","Call-graph analysis determines whether the vulnerable sink is reachable from an "
     "entry point under the application's actual configuration.")
step(2,"Un-reachability proof.","UNSAT reachability is provable un-reachability: soundness implies no false "
     "negatives over the boundary, so a CVE can be deferred defensibly on a mission system.")
step(3,"Proof-carrying virtual patch.","A reachable sink is closed by a patch that carries a proof it "
     "preserves behavior - not a WAF signature at the perimeter.")
step(4,"Unified containment.","Network microsegmentation and identity / token / tool-permission revocation are "
     "expressed under one verified policy, sequenced by attack path: unreachable deprioritize, reachable "
     "virtual-patch or IOC-block or full-patch.")
prop("provable un-reachability (soundness over the analysis boundary) and one policy that revokes a host "
     "segment or an agent's token.")
boundary("\"Formally verified\" needs a named, bounded verification boundary (an seL4-style trusted computing "
         "base) - an audience that funded HACMS and seL4 will probe it, so state the boundary explicitly. "
         "Soundness (zero false negatives) is the wedge over heuristic reachability; never use \"reachability\" "
         "or \"virtual patching\" unqualified.")

# ============================== SUMMARY TABLE ==============================
h1("8.  At a glance: the gap and the coverage")
rows=[
 ("L1","Verifiers prove packet reachability; BAS samples a scenario library; none certify control efficacy or a true negative","Symbolic proof the config defeats every modeled technique (soundness over config + corpus)"),
 ("L2","SIEM / UEBA threshold on magnitude; slow low-magnitude drift and drift direction go unseen","Behavior embedded as meaning; detect drift direction + persistence, cohort-relative"),
 ("L3","WAAP alerts on deviation from a learned normal and enforces at a vendor edge","Classify request intent; enforce inline with no traffic egress"),
 ("L4","Scopes are hand-authored RBAC; learned baselines need a warm-up; supply-chain caught in scans","Persona compiled to a least-privilege MCP scope; spec is the baseline; registry enforced inline"),
 ("L5","Reachability is heuristic (no soundness); virtual patch is a signature; host and agent on separate planes","Sound un-reachability proof; proof-carrying patch; one host + agent containment policy"),
]
tbl=doc.add_table(rows=1,cols=3); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.autofit=False; tbl.style='Table Grid'
widths=[Inches(0.5),Inches(3.7),Inches(3.5)]
hdr=tbl.rows[0].cells
for i,t in enumerate(["","THE GAP IN CURRENT TOOLS","HOW WE COVER IT"]):
    hdr[i].width=widths[i]; shade(hdr[i],'0E2841')
    run(hdr[i].paragraphs[0],t,size=9,color=WHITE,bold=True,font=HEADF)
for code,gap,cover in rows:
    c=tbl.add_row().cells
    for i in range(3): c[i].width=widths[i]
    shade(c[0],'0E2841')
    p0=c[0].paragraphs[0]; p0.alignment=WD_ALIGN_PARAGRAPH.CENTER; run(p0,code,size=10,color=WHITE,bold=True,font=HEADF)
    run(c[1].paragraphs[0],gap,size=9,color=INK)
    run(c[2].paragraphs[0],cover,size=9,color=NAVY,bold=True)
para("Formal-methods layers (L1, L5) discharge a proof over a bounded model; semantic layers (L2, L3, L4) "
     "classify intent and behavior, not surface counts.",size=9.5,color=STEEL,italic=True,before=6,after=4)

# ============================== APPENDIX ==============================
h1("9.  Appendix: named tools and their technical limitation")
refs=[
 ("L1","Forward Networks, Batfish (formal reachability, forwarding only); SafeBreach, AttackIQ (BAS, finite "
       "scenario library); XM Cyber, RedSeal, Cymulate (heuristic attack graphs); Tenable, Qualys (CVE enumeration)."),
 ("L2","Splunk, Microsoft Sentinel, Exabeam (UEBA, magnitude thresholds); Darktrace (network deviation); "
       "Vectra, Gurucul (global-baseline ML)."),
 ("L3","Salt Security (detect-only); Akamai, Imperva, Cloudflare, F5 (edge / CDN enforcement); Traceable "
       "(tracing agents)."),
 ("L4","Microsoft Entra Agent ID, Zscaler registry, Noma (hand-authored scopes); Palo Alto Prisma AIRS, "
       "Zscaler AI Broker (gateway category peers); SASE vendors (transport-only)."),
 ("L5","Endor Labs, Snyk, Contrast, Oligo (heuristic reachability / ADR); Trend Micro, Fortinet (signature "
       "virtual patching); Illumio, Guardicore (host microsegmentation); agent-identity tools (token revocation only)."),
]
for code,txt in refs:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.28); p.paragraph_format.space_after=Pt(5)
    run(p,code+"   ",size=10,color=NAVY,bold=True,font=HEADF); run(p,txt,size=10,color=INK)
para(" ",after=6)
rule(NAVY,10)
para("CONFIDENTIAL  |  22nd Century Technologies  |  Technical differentiation review, gap and coverage per "
     "layer  |  Verify all vendor and number claims against current sources before use.",size=8,color=MUTED,after=0)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"Layer_Differentiation_Findings.docx")
try:
    doc.save(out); print("saved",os.path.abspath(out))
except PermissionError:
    alt=r"C:/Users/shuklar/AppData/Local/Temp/claude/c--Users-shuklar-ClaudeCode-CyberUEBA-BehavioralIntel-Enhanced/a6818674-bde5-4462-be21-ffe7ac102bf2/scratchpad/Layer_Differentiation_Findings.docx"
    doc.save(alt); print("LOCKED - saved fallback",alt)
